# -*- coding: utf-8 -*-
"""
德鲁克论创新与企业家精神 - 讲师手册生成脚本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============== 字体和样式配置 ==============

# 微软雅黑用于中文标题和正文
CJK_FONT = '微软雅黑'
# Times New Roman用于英文
LATIN_FONT = 'Times New Roman'

# 颜色定义
PRIMARY_COLOR = RGBColor(0x1F, 0x38, 0x64)    # 深海军蓝
SECONDARY_COLOR = RGBColor(0x2E, 0x75, 0xB6)    # 中蓝色
ACCENT_COLOR = RGBColor(0xC0, 0x00, 0x00)       # 深红色（德鲁克风格）

# ============== 辅助函数 ==============

def set_run_font(run, cjk_font=CJK_FONT, latin_font=LATIN_FONT, size=12, bold=False, color=None):
    """设置run的字体、字号和颜色"""
    run.font.name = latin_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cjk_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, size=16, bold=True, color=None, space_before=24, space_after=12):
    """添加自定义标题"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return para

def add_body_para(doc, text, size=11, space_before=0, space_after=8, bold=False, indent=False):
    """添加正文段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.3)
    run = para.add_run(text)
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

def add_bullet_item(doc, text, level=0, size=11):
    """添加项目符号项"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(size)
    return para

def add_number_item(doc, text, num, level=0, size=11):
    """添加编号项"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    para.paragraph_format.first_line_indent = Inches(-0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(f"{num}. {text}")
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(size)
    return para

def create_table(doc, headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = CJK_FONT
                run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 设置表头背景色
        tc = header_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)

    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = CJK_FONT
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
                    run.font.size = Pt(10)
        # 斑马条纹
        if row_idx % 2 == 1:
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)

    return table

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    """设置页面边距"""
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

# ============== 主文档创建 ==============

def create_handbook():
    doc = Document()

    # 页面设置
    set_page_margins(doc, top=1.0, bottom=1.0, left=1.2, right=1.0)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # ========== 封面页 ==========
    add_heading(doc, '', level=0, size=36, space_before=120, space_after=24)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    title_para.paragraph_format.space_after = Pt(24)
    run = title_para.add_run('德鲁克论创新与企业家精神')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(12)
    run = subtitle_para.add_run('讲师手册')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(24)
    run.font.color.rgb = SECONDARY_COLOR

    # 课程信息表
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.paragraph_format.space_before = Pt(60)
    info_para.paragraph_format.space_after = Pt(6)
    run = info_para.add_run('课程时长：2天（每天6小时，共12小时）')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(14)

    info_para2 = doc.add_paragraph()
    info_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para2.paragraph_format.space_after = Pt(6)
    run = info_para2.add_run('班级规模建议：20-30人')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(14)

    info_para3 = doc.add_paragraph()
    info_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para3.paragraph_format.space_after = Pt(120)
    run = info_para3.add_run('对象：企业中高层管理者、创业者')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(14)

    add_page_break(doc)

    # ========== 目录 ==========
    add_heading(doc, '目录', level=1, size=20, color=PRIMARY_COLOR, space_before=12, space_after=20)

    toc_items = [
        '第一章  讲师准备',
        '    1.1  课前准备清单',
        '    1.2  教室布置要求',
        '    1.3  物料清单',
        '    1.4  案例材料准备',
        '    1.5  学员前测分析',
        '第二章  模块教学指引',
        '    2.1  模块1：创新是偶然还是系统（约90分钟）',
        '    2.2  模块2：来源一——意料之外（约50分钟）',
        '    2.3  模块3：来源二——不协调（约50分钟）',
        '    2.4  模块4：来源三——流程需要（约50分钟）',
        '    2.5  模块5：来源四——行业市场结构变化（约50分钟）',
        '    2.6  模块6：来源五——人口变化（约50分钟）',
        '    2.7  模块7：来源六——认知变化（约50分钟）',
        '    2.8  模块8：来源七——新知识（约50分钟）',
        '    2.9  模块9：企业家精神（约60分钟）',
        '第三章  评估设计',
        '    3.1  形成性评估',
        '    3.2  总结性评估',
        '第四章  讲师工具',
        '    4.1  课程时间分配表',
        '    4.2  物料清单',
        '    4.3  学员表现观察表',
        '    4.4  常见问题应对指南',
        '    4.5  案例讲解提示卡',
    ]

    for item in toc_items:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        if item.startswith('    '):
            para.paragraph_format.left_indent = Inches(0.4)
            run = para.add_run(item.strip())
        else:
            run = para.add_run(item)
            run.font.bold = True
        run.font.name = CJK_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
        run.font.size = Pt(11)

    add_page_break(doc)

    # ========== 第一章：讲师准备 ==========
    add_heading(doc, '第一章  讲师准备', level=1, size=18, color=PRIMARY_COLOR, space_before=12, space_after=16)

    # 1.1 课前准备清单
    add_heading(doc, '1.1  课前准备清单', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '讲师应在课前至少三天完成以下准备工作，确保课程顺利进行：', space_after=8)

    prep_items = [
        '确认教室booking，了解教室容量、设备配置（投影仪、音响、白板等）',
        '发送课前邮件给学员，包含课程目标、大纲、需提前阅读的材料',
        '准备分组讨论的桌牌和编号',
        '确认案例打印材料数量（建议每人一份案例文本）',
        '测试所有电子设备，确保演示文稿正常播放',
        '准备便签纸和白板笔（分组讨论用）',
        '打印学员前测问卷并准备回收袋',
        '准备七个创新来源的扫描工具表格（每人一份）',
        '准备创新行动计划的空白模板',
        '确认午餐和茶歇安排',
    ]

    for item in prep_items:
        add_bullet_item(doc, item)

    # 1.2 教室布置要求
    add_heading(doc, '1.2  教室布置要求', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    add_body_para(doc, '分组讨论布局（推荐）', bold=True, space_after=4)

    layout_items = [
        '每组5-6人，围桌而坐，便于讨论互动',
        '小组数量建议4-5组，不超过6组',
        '小组间保持足够通道，便于讲师巡视',
        '讲台位置居中但略有突出，便于观察全场',
        '准备至少一块白板或白板纸架，供小组记录',
    ]

    for item in layout_items:
        add_bullet_item(doc, item)

    add_body_para(doc, 'U形布局（备选）', bold=True, space_before=12, space_after=4)

    u_layout_items = [
        '适合20人以下的小班',
        '便于学员之间的对视交流',
        '讲师可以在U形内侧走动，促进互动',
    ]

    for item in u_layout_items:
        add_bullet_item(doc, item)

    # 1.3 物料清单
    add_heading(doc, '1.3  物料清单', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    headers = ['类别', '物料名称', '数量', '备注']
    rows = [
        ['纸张', '七个创新来源扫描工具', '每人1份', '提前打印'],
        ['纸张', '学员前测问卷', '每人1份', '提前打印'],
        ['纸张', '案例文本', '每人1份', '每个案例'],
        ['纸张', '创新行动计划模板', '每人1份', '提前打印'],
        ['纸张', '出口票', '每人6-8张', '可用便签纸替代'],
        ['文具', '白板笔', '每组1支', '红、蓝、黑三色'],
        ['文具', '大白白板纸', '每组2张', '记录讨论成果'],
        ['文具', '透明胶带', '2卷', '固定白板纸'],
        ['设备', '投影仪', '1台', '测试正常'],
        ['设备', '音响设备', '1套', '测试正常'],
        ['设备', '翻页笔', '1支', '备用电池'],
        ['其他', '签到表', '1份', '-'],
        ['其他', '姓名牌', '每人1个', '可重复使用'],
        ['其他', '糖果/小奖品', '若干', '奖励积极参与者'],
    ]

    create_table(doc, headers, rows)

    # 1.4 案例材料准备
    add_heading(doc, '1.4  案例材料准备', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    add_body_para(doc, '每个创新来源需要配备1-2个真实企业案例，建议按以下结构准备：', space_after=8)

    case_structure = [
        '案例背景：企业基本情况、所属行业、规模',
        '创新契机：如何发现创新机会',
        '创新举措：具体采取了什么行动',
        '创新成果：取得的商业成果',
        '经验教训：可借鉴之处',
    ]

    for i, item in enumerate(case_structure, 1):
        add_number_item(doc, item, i)

    add_body_para(doc, '推荐案例来源：', bold=True, space_before=12, space_after=4)

    case_sources = [
        '德鲁克《创新与企业家精神》原书案例',
        '哈佛商业评论经典案例',
        '中国本土企业创新实践（如腾讯、阿里、字节跳动等）',
    ]

    for item in case_sources:
        add_bullet_item(doc, item)

    # 1.5 学员前测分析
    add_heading(doc, '1.5  学员前测分析', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    add_body_para(doc, '前测问卷应在课程开始前发放，完成后立即回收分析。重点关注以下维度：', space_after=8)

    pretest_dims = [
        '对"创新"的理解：灵感式 vs 系统化',
        '当前企业的创新现状：成熟度自评',
        '对德鲁克创新理论的了解程度',
        '最希望在课程中解决的创新挑战',
        '对企业家精神的理解',
    ]

    for item in pretest_dims:
        add_bullet_item(doc, item)

    add_body_para(doc, '分析完成后，讲师应在课程开场时针对以下问题与学员进行简短分享：', space_before=12, space_after=4)

    share_points = [
        '班级整体的创新认知水平',
        '最普遍的创新挑战',
        '学员对课程的期望',
    ]

    for item in share_points:
        add_bullet_item(doc, item)

    add_page_break(doc)

    # ========== 第二章：模块教学指引 ==========
    add_heading(doc, '第二章  模块教学指引', level=1, size=18, color=PRIMARY_COLOR, space_before=12, space_after=16)

    # 模块1
    add_heading(doc, '2.1  模块1：创新是偶然还是系统（约90分钟）', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '【教学目标】', bold=True, space_after=4)
    m1_goals = [
        '让学员认识到灵感式创新的局限性',
        '建立"创新可系统化"的认知框架',
        '激发学员对七个来源的好奇心',
    ]
    for item in m1_goals:
        add_bullet_item(doc, item)

    add_body_para(doc, '【活动序列】', bold=True, space_before=12, space_after=4)

    # 活动1
    add_body_para(doc, '活动1：开场头脑风暴（15分钟）', bold=True, space_before=8, space_after=4)
    add_body_para(doc, '目的：激活学员的既有经验，引出"创新靠灵感还是靠系统"的思考。', space_after=4)
    add_body_para(doc, '操作步骤：', space_after=2)
    brainstorm_steps = [
        '讲师提问："请回想一下，你们公司最近一次成功的创新是怎么发生的？是精心策划还是灵光一现？"',
        '给学员2分钟独立思考，在便签纸上写下答案',
        '让学员与邻座分享（3分钟）',
        '随机邀请3-4人分享（5分钟）',
        '讲师在白板上记录关键词（灵感、偶然、运气、领导重视、团队努力等）',
    ]
    for step in brainstorm_steps:
        add_bullet_item(doc, step, level=1)

    add_body_para(doc, '讲师话术：', bold=True, space_before=8, space_after=2)
    add_body_para(doc, '"大家分享了很多精彩的案例。我注意到，很多创新确实源于灵感和偶然——这没有错。但今天我们要问一个更深的问题：那些持续创新的企业，只是比其他人更幸运吗？还是他们有某种系统？"', indent=True)

    # 活动2
    add_body_para(doc, '活动2：概念讲解（20分钟）', bold=True, space_before=12, space_after=4)
    add_body_para(doc, '目的：系统介绍德鲁克的创新理论框架，建立"创新可系统化"的认知。', space_after=4)
    add_body_para(doc, '核心内容：', space_after=2)
    concept_points = [
        '德鲁克的核心观点："创新机遇的七个来源"',
        '灵感式创新的三个问题：不可预测、不可复制、不可持续',
        '系统化创新的优势：可预测、可学习、可复制',
        '创新与企业家精神的关系',
    ]
    for item in concept_points:
        add_bullet_item(doc, item)

    # 活动3
    add_body_para(doc, '活动3：案例讨论（20分钟）', bold=True, space_before=12, space_after=4)
    add_body_para(doc, '目的：通过真实案例让学员感受灵感式vs系统化创新的区别。', space_after=4)

    case_discussion = [
        '案例：某科技公司的"创新农场"',
        '讨论问题：',
    ]
    for item in case_discussion:
        add_bullet_item(doc, item)

    discussion_questions = [
        '这个公司的创新是灵感式的还是系统化的？为什么？',
        '这种模式的优缺点是什么？',
        '如果你是CEO，你会如何改进？',
    ]
    for q in discussion_questions:
        add_bullet_item(doc, q, level=1)

    # 活动4
    add_body_para(doc, '活动4：讲解——七个来源总览（15分钟）', bold=True, space_before=12, space_after=4)
    add_body_para(doc, '目的：预览七个创新来源，为后续模块埋下伏笔。', space_after=4)

    seven_sources = [
        '来源一：意料之外——意外成功、意外失败、意外外部事件',
        '来源二：不协调——现实与假设之间的不协调',
        '来源三：流程需要——填补流程中缺失的步骤',
        '来源四：行业和市场结构变化——结构重塑带来的机会',
        '来源五：人口变化——人口统计特征的重大变化',
        '来源六：认知变化——思维范式的转变',
        '来源七：新知识——科学或非科学知识的突破',
    ]
    for source in seven_sources:
        add_bullet_item(doc, source)

    # 活动5
    add_body_para(doc, '活动5：总结（5分钟）', bold=True, space_before=12, space_after=4)
    add_body_para(doc, '目的：回顾本模块核心观点，强化记忆。', space_after=4)
    add_body_para(doc, '关键信息：', space_after=2)
    key_points = [
        '创新不是碰运气，而是有规律可循',
        '德鲁克提供了七个系统化的创新来源',
        '每个来源都是一个潜在的创新机会窗口',
        '企业需要的不是"灵感"，而是系统地扫描这七个来源',
    ]
    for item in key_points:
        add_bullet_item(doc, item)

    # 常见问题与应对
    add_heading(doc, '【常见问题与应对】', level=3, size=12, color=ACCENT_COLOR, space_before=12, space_after=8)

    faq_data = [
        ['问题', '应对策略'],
        ['"我们公司很小，不需要创新"', '创新不只是大企业的专利。小企业更需要通过系统化创新来生存。分享小企业通过七个来源找到机会的案例。'],
        ['"我们行业太传统，没什么创新机会"', '越是传统行业，结构性变化越大。引导学员思考：人口变化、认知变化如何影响他们的行业？'],
        ['学员分享过于冗长', '使用"2分钟规则"：每人最多分享2分钟。使用计时器，并在分享前明确规则。'],
        ['讨论气氛沉闷', '从学员熟悉的案例入手，或使用更具争议性的问题激发思考。'],
    ]

    create_table(doc, faq_data[0], faq_data[1:])

    add_page_break(doc)

    # 模块2-8
    add_heading(doc, '2.2-2.8  七大来源模块（每个约50分钟）', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '每个来源模块采用标准化的教学流程，确保学员能够：', space_after=8)

    standard_sequence = [
        '理解每个来源的核心概念',
        '识别该来源在自身企业中的具体表现',
        '运用扫描工具进行实践练习',
        '与小组分享和相互反馈',
    ]
    for item in standard_sequence:
        add_bullet_item(doc, item)

    add_body_para(doc, '【标准活动序列】', bold=True, space_before=16, space_after=8)

    # 模块2-8 详细活动
    modules_2_8 = [
        ('模块2', '来源一——意料之外（约50分钟）', [
            ('场景引入', '10分钟', '讲师展示几个"意外成功/失败"的商业新闻，让学员判断哪个是意外成功、哪个是意外失败'),
            ('概念讲解', '15分钟', '详解意外成功、意外失败、意外外部事件的区别；讨论为什么企业往往忽视意外'),
            ('案例分析', '15分钟', '案例：IBM如何从意外成功中发现了个人电脑的机会；或中国企业的类似案例'),
            ('工具练习', '15-20分钟', '学员使用"意料之外扫描表"分析自身企业的意外事件'),
            ('小组讨论', '10分钟', '各组分享最有趣的意外发现，讨论其创新潜力'),
            ('总结', '5分钟', '回顾核心观点：意外是创新的重要信号，但要学会识别和捕捉'),
        ]),
        ('模块3', '来源二——不协调（约50分钟）', [
            ('场景引入', '10分钟', '展示一个行业内的"矛盾现象"（如：手机屏幕越来越大，但人们更喜欢小屏手机）'),
            ('概念讲解', '15分钟', '四类不协调：业务现实与假设之间、认知与实际之间、流程各步骤之间、行业成长与价值之间'),
            ('案例分析', '15分钟', '案例：7-Eleven如何利用不协调创新（门店小 vs 消费者需求多）'),
            ('工具练习', '15-20分钟', '学员识别自身行业或企业中的三类不协调现象'),
            ('小组讨论', '10分钟', '各组分享，讨论哪个不协调最具创新潜力'),
            ('总结', '5分钟', '不协调是创新的温床，问题本身往往包含解决方案'),
        ]),
        ('模块4', '来源三——流程需要（约50分钟）', [
            ('场景引入', '10分钟', '展示一个常见流程中的"痛点"，让学员思考如何解决'),
            ('概念讲解', '15分钟', '流程需要的核心：填补现有流程中缺失的步骤；关键是"冷眼旁观"而非深入现有流程'),
            ('案例分析', '15分钟', '案例：联邦快递的"隔夜送达"如何填补了物流流程的缺口'),
            ('工具练习', '15-20分钟', '学员绘制自身企业核心流程，识别缺失的步骤'),
            ('小组讨论', '10分钟', '各组分享最有价值的流程创新想法'),
            ('总结', '5分钟', '流程需要创新的关键：关注步骤而非技术'),
        ]),
        ('模块5', '来源四——行业市场结构变化（约50分钟）', [
            ('场景引入', '10分钟', '展示一个行业巨变的案例（如：传统书店vs电子书vs知识付费）'),
            ('概念讲解', '15分钟', '结构变化的四个信号：快速增长、监管变化、新消费群体、全球化；德鲁克的警示——变化先于认知'),
            ('案例分析', '15分钟', '案例：苹果公司如何利用行业结构变化重塑音乐产业'),
            ('工具练习', '15-20分钟', '学员分析自身行业正在经历的结构变化'),
            ('小组讨论', '10分钟', '各组分享，讨论最大的结构性机会在哪里'),
            ('总结', '5分钟', '结构变化是最大的创新机会窗口，但必须快速行动'),
        ]),
        ('模块6', '来源五——人口变化（约50分钟）', [
            ('场景引入', '10分钟', '展示一组人口统计数据的变化趋势（出生率、老龄化、城镇化等）'),
            ('概念讲解', '15分钟', '人口统计特征的变化：规模、年龄结构、收入水平、教育程度、职业结构、地理位置；关键问题是"何时"'),
            ('案例分析', '15分钟', '案例：中国养老产业的兴起与人口老龄化的关系'),
            ('工具练习', '15-20分钟', '学员识别自身业务相关的人口变化趋势'),
            ('小组讨论', '10分钟', '各组分享最有价值的人口变化机会'),
            ('总结', '5分钟', '人口变化是最可靠的创新来源，因为变化是确定的、可预测的'),
        ]),
        ('模块7', '来源六——认知变化（约50分钟）', [
            ('场景引入', '10分钟', '展示一个曾经"理所当然"现在"过时"的观念'),
            ('概念讲解', '15分钟', '认知变化的核心：思维范式的转变；从"是什么"到"意味着什么"的转变；德鲁克称之为"窗口机会"'),
            ('案例分析', '15分钟', '案例：共享经济的认知转变（从"所有权"到"使用权"）'),
            ('工具练习', '15-20分钟', '学员识别自身行业正在发生的认知变化'),
            ('小组讨论', '10分钟', '各组分享，讨论认知变化带来的创新机会'),
            ('总结', '5分钟', '认知变化往往被低估，但它往往是最大创新机会的来源'),
        ]),
        ('模块8', '来源七——新知识（约50分钟）', [
            ('场景引入', '10分钟', '展示一个基于新知识的产品或服务（ChatGPT、mRNA疫苗等）'),
            ('概念讲解', '15分钟', '新知识创新的特点：时间长、不确定性高、需要多元知识整合；vs意外成功的快速；关键问题是"何时"'),
            ('案例分析', '15分钟', '案例：特斯拉如何整合电池技术、电动汽车知识进行创新'),
            ('工具练习', '15-20分钟', '学员识别相关领域的新知识发展及其应用潜力'),
            ('小组讨论', '10分钟', '各组分享，讨论新知识创新的风险和机遇'),
            ('总结', '5分钟', '新知识创新需要耐心和长期主义，但一旦成功，回报巨大'),
        ]),
    ]

    for module_name, module_title, activities in modules_2_8:
        add_heading(doc, f'{module_name}：{module_title}', level=3, size=13, color=PRIMARY_COLOR, space_before=12, space_after=8)

        for activity_name, time, description in activities:
            add_body_para(doc, f'{activity_name}（{time}）', bold=True, space_before=6, space_after=2)
            add_body_para(doc, description, space_after=4, indent=True)

    add_page_break(doc)

    # 模块9：企业家精神
    add_heading(doc, '2.9  模块9：企业家精神（约60分钟）', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '【教学目标】', bold=True, space_after=4)
    m9_goals = [
        '理解企业家精神的本质是"系统化的创新实践"',
        '区分企业家战略：聚焦、模仿、入口、生态',
        '识别自身企业的企业家精神实践',
    ]
    for item in m9_goals:
        add_bullet_item(doc, item)

    add_body_para(doc, '【活动序列】', bold=True, space_before=12, space_after=8)

    m9_activities = [
        ('概念导入（15分钟）', '讲解企业家精神的定义：不是一种特征，而是一种行为；核心是"有目的的创新实践"'),
        ('企业家战略类型（20分钟）', '四种企业家战略：', [
            '聚焦战略：成为小市场的主导者',
            '模仿战略：在新市场复制成功模式',
            '入口战略：在变化中寻找切入点',
            '生态战略：创造一个价值网络',
        ]),
        ('案例分析（15分钟）', '分析一个中国企业的企业家精神实践（建议用阿里或字节跳动）'),
        ('行动计划（10分钟）', '学员制定自己的"企业家精神实践计划"，聚焦七个来源中自身最有潜力的一个'),
    ]

    for activity_name, description, *details in m9_activities:
        add_body_para(doc, activity_name, bold=True, space_before=8, space_after=2)
        add_body_para(doc, description, space_after=4, indent=True)
        if details:
            for detail in details[0]:
                add_bullet_item(doc, detail, level=1)

    add_page_break(doc)

    # ========== 第三章：评估设计 ==========
    add_heading(doc, '第三章  评估设计', level=1, size=18, color=PRIMARY_COLOR, space_before=12, space_after=16)

    # 3.1 形成性评估
    add_heading(doc, '3.1  形成性评估', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '形成性评估贯穿整个课程，用于实时监测学员学习情况，及时调整教学策略。', space_after=8)

    formative_methods = [
        ('出口票（Exit Ticket）', '每个模块结束时', '学员用便签纸回答一个核心问题：如"今天最让你印象深刻的一个观点是什么？""你还有一个什么困惑？"'),
        ('小组讨论参与度观察', '小组讨论时', '讲师巡视各组，观察并记录学员参与情况；关注发言频率、观点质量、倾听态度'),
        ('工具完成质量', '工具练习环节', '检查学员的扫描工具完成情况；优秀作品可在全班分享'),
        ('前测/后测对比', '课程前后', '通过同一套题目测试学习前后变化，评估课程效果'),
    ]

    for method, timing, description in formative_methods:
        add_body_para(doc, method, bold=True, space_before=8, space_after=2)
        add_body_para(doc, f'时间：{timing}', space_after=2, indent=True)
        add_body_para(doc, f'方式：{description}', space_after=4, indent=True)

    # 3.2 总结性评估
    add_heading(doc, '3.2  总结性评估', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    add_body_para(doc, '总结性评估在课程结束时进行，用于评价学员整体学习效果和课程价值。', space_after=8)

    summative_assessments = [
        ('前测/后测对比', '评估指标变化', '创新认知水平从"灵感依赖"转向"系统思维"；对七个来源的了解程度；企业家精神意识'),
        ('七个来源扫描工具完成质量', '工具完整性和深度', '每个来源是否都有具体的应用实例；创新机会的可行性分析'),
        ('创新行动计划的质量', '行动计划的实用性', '是否基于七个来源之一；是否有明确的目标和时间表；是否可执行'),
    ]

    for assessment, criteria, description in summative_assessments:
        add_body_para(doc, assessment, bold=True, space_before=8, space_after=2)
        add_body_para(doc, f'评估标准：{criteria}', space_after=2, indent=True)
        add_body_para(doc, f'{description}', space_after=4, indent=True)

    # 评估量规表
    add_heading(doc, '【创新行动计划评估量规】', level=3, size=12, color=ACCENT_COLOR, space_before=12, space_after=8)

    rubric_headers = ['维度', '优秀（4分）', '良好（3分）', '合格（2分）', '待改进（1分）']
    rubric_rows = [
        ['来源识别', '精准对应七个来源之一，论证充分', '对应来源基本准确', '有一定关联但不够精准', '未明确对应来源'],
        ['机会分析', '深度分析市场规模、竞争、可行性', '分析较为完整', '分析较浅', '缺乏分析'],
        ['行动计划', '目标明确、步骤清晰、时间表合理', '较为清晰', '有基本框架', '模糊不清'],
        ['创新性', '有独特视角，难以想到的创新点', '有一定创新性', '常规做法', '缺乏新意'],
    ]

    create_table(doc, rubric_headers, rubric_rows)

    add_page_break(doc)

    # ========== 第四章：讲师工具 ==========
    add_heading(doc, '第四章  讲师工具', level=1, size=18, color=PRIMARY_COLOR, space_before=12, space_after=16)

    # 4.1 课程时间分配表
    add_heading(doc, '4.1  课程时间分配表', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    add_body_para(doc, '【第一天】', bold=True, space_after=4)

    day1_headers = ['时间段', '模块', '内容', '时长']
    day1_rows = [
        ['08:30-09:00', '签到', '签到、发放资料', '30分钟'],
        ['09:00-09:30', '开场', '课程介绍、前测分享、期望确认', '30分钟'],
        ['09:30-11:00', '模块1', '创新是偶然还是系统', '90分钟'],
        ['11:00-11:15', '茶歇', '-', '15分钟'],
        ['11:15-12:15', '模块2', '来源一——意料之外', '60分钟'],
        ['12:15-13:30', '午餐', '-', '75分钟'],
        ['13:30-14:30', '模块3', '来源二——不协调', '60分钟'],
        ['14:30-15:30', '模块4', '来源三——流程需要', '60分钟'],
        ['15:30-15:45', '茶歇', '-', '15分钟'],
        ['15:45-16:45', '模块5', '来源四——行业市场结构变化', '60分钟'],
        ['16:45-17:30', '第一天总结', '回顾与反思、布置作业', '45分钟'],
    ]

    create_table(doc, day1_headers, day1_rows)

    add_body_para(doc, '【第二天】', bold=True, space_before=16, space_after=4)

    day2_headers = ['时间段', '模块', '内容', '时长']
    day2_rows = [
        ['08:30-09:00', '签到', '签到、回顾第一天', '30分钟'],
        ['09:00-09:50', '模块6', '来源五——人口变化', '50分钟'],
        ['09:50-10:40', '模块7', '来源六——认知变化', '50分钟'],
        ['10:40-10:55', '茶歇', '-', '15分钟'],
        ['10:55-11:45', '模块8', '来源七——新知识', '50分钟'],
        ['11:45-12:30', '模块9', '企业家精神', '45分钟'],
        ['12:30-13:30', '午餐', '-', '60分钟'],
        ['13:30-14:30', '综合练习', '七个来源综合扫描练习', '60分钟'],
        ['14:30-15:30', '行动计划', '制定个人/企业创新行动计划', '60分钟'],
        ['15:30-15:45', '茶歇', '-', '15分钟'],
        ['15:45-16:45', '行动计划汇报', '小组分享与点评', '60分钟'],
        ['16:45-17:30', '总结与后测', '课程总结、后测、反馈', '45分钟'],
    ]

    create_table(doc, day2_headers, day2_rows)

    # 4.2 物料清单（完整版）
    add_heading(doc, '4.2  物料清单（完整版）', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    materials_headers = ['序号', '物料名称', '规格/要求', '数量', '状态']
    materials_rows = [
        ['1', '七个创新来源扫描工具', 'A4纸，彩色打印', '每人1份', '□ 已准备'],
        ['2', '学员前测问卷', 'A4纸', '每人1份', '□ 已准备'],
        ['3', '案例文本', 'A4纸，骑马钉', '每人1份', '□ 已准备'],
        ['4', '创新行动计划模板', 'A4纸', '每人1份', '□ 已准备'],
        ['5', '出口票（便签纸）', '黄色便签', '每人8张', '□ 已准备'],
        ['6', '白板笔', '红、蓝、黑三色', '每组各1支', '□ 已准备'],
        ['7', '大白白板纸', 'A1尺寸', '每组2张', '□ 已准备'],
        ['8', '透明胶带', '窄版', '2卷', '□ 已准备'],
        ['9', '投影仪', '1080P以上', '1台', '□ 已测试'],
        ['10', '音响设备', '含麦克风', '1套', '□ 已测试'],
        ['11', '翻页笔', 'USB接口', '1支', '□ 已准备（备用电池）'],
        ['12', '签到表', 'A4纸', '1份', '□ 已准备'],
        ['13', '姓名牌', '双面展示', '每人1个', '□ 已准备'],
        ['14', '糖果/小奖品', '-', '若干', '□ 已准备'],
        ['15', '课程大纲', 'A4纸', '每人1份', '□ 已准备'],
    ]

    create_table(doc, materials_headers, materials_rows)

    # 4.3 学员表现观察表
    add_heading(doc, '4.3  学员表现观察表', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    observation_headers = ['学员/组名', '模块1参与', '模块2表现', '模块3表现', '讨论贡献', '创新思维', '整体评价']
    observation_rows = [
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
        ['', '', '', '', '', '', ''],
    ]

    create_table(doc, observation_headers, observation_rows)

    add_body_para(doc, '评分标准：1-5分，5分为最高。讨论贡献指发言频率和质量；创新思维指观点的独特性和深度。', space_before=8, size=9)

    # 4.4 常见问题应对指南
    add_heading(doc, '4.4  常见问题应对指南', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    faq_headers = ['问题类型', '常见问题', '应对策略']
    faq_rows = [
        ['认知层面', '"创新太抽象，我们不懂"', '用具体案例说明，特别是学员行业的案例'],
        ['认知层面', '"我们公司不需要创新"', '引导思考：市场竞争、客户变化、技术变革是否在影响他们'],
        ['认知层面', '"德鲁克理论太老了"', '指出德鲁克理论的核心——系统化方法论，跨越时代'],
        ['参与层面', '学员不主动发言', '使用小组讨论先热身；点名时请学员分享"最认同的观点"而非"最聪明的观点"'],
        ['参与层面', '讨论偏离主题', '使用白板记录偏离观点，承诺后续讨论；温和地把话题拉回'],
        ['参与层面', '某学员过于主导', '邀请其他学员补充；设置每人发言时间限制'],
        ['时间层面', '讨论超时', '提前5分钟提醒；必要时压缩总结时间'],
        ['时间层面', '内容讲不完', '确保核心模块（模块1和模块9）时间充足，其他模块可适当精简'],
        ['实践层面', '学员问如何落地', '提供具体的工具和方法；安排行动计划环节'],
        ['实践层面', '学员觉得不适用', '引导分析自身企业与案例的异同；鼓励学员提出调整方案'],
    ]

    create_table(doc, faq_headers, faq_rows)

    # 4.5 案例讲解提示卡
    add_heading(doc, '4.5  案例讲解提示卡', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    add_body_para(doc, '以下是各模块推荐案例的讲解要点提示，供讲师参考：', space_after=8)

    case_cards = [
        ('模块1案例', '某科技公司"创新农场"', [
            '背景：中型科技公司，1000人规模',
            '现象：成立"创新农场"团队，鼓励员工用20%时间做创新项目',
            '结果：产出了多个创新产品，但也导致核心项目资源被分散',
            '讨论点：这是灵感式还是系统化？为什么？优缺点是什么？',
        ]),
        ('模块2案例', 'IBM与个人电脑', [
            '背景：IBM在大型机市场的主导地位',
            '意外：工程师发现年轻员工在用苹果电脑处理文档',
            '决策：IBM选择与微软和Intel合作，快速进入PC市场',
            '启示：意外成功的识别和把握是关键',
        ]),
        ('模块5案例', '7-Eleven的小门店创新', [
            '不协调：门店面积小（100平米）vs 消费者需求多（几千种商品）',
            '创新：引入信息系统，实现"单品管理"，用小库存满足大需求',
            '结果：单店效率远超竞争对手',
        ]),
        ('模块7案例', '特斯拉的电池创新', [
            '新知识整合：电池技术、电机技术、软件技术',
            '创新路径：先高端市场（Roadster），再大众市场（Model 3）',
            '关键决策：自建充电网络，解决消费者里程焦虑',
        ]),
        ('模块9案例', '字节跳动的企业家精神', [
            '系统化创新：算法推荐、内容生态、全球化布局',
            '企业家战略：入口战略（今日头条）→ 生态战略（TikTok）',
            '关键要素：快速迭代、数据驱动、全球化视野',
        ]),
    ]

    for card_title, card_subtitle, points in case_cards:
        add_body_para(doc, card_title, bold=True, space_before=10, space_after=2)
        add_body_para(doc, card_subtitle, bold=True, space_after=2)
        for point in points:
            add_bullet_item(doc, point, level=1)

    # ========== 附录 ==========
    add_page_break(doc)
    add_heading(doc, '附录', level=1, size=18, color=PRIMARY_COLOR, space_before=12, space_after=16)

    add_heading(doc, '附录A：七个创新来源速查表', level=2, size=14, color=SECONDARY_COLOR, space_before=12, space_after=12)

    quick_ref_headers = ['来源', '核心问题', '关键信号', '经典案例']
    quick_ref_rows = [
        ['意料之外', '什么是我们没有预料到的？', '意外成功、失败、外部事件', 'IBM发现PC机会'],
        ['不协调', '现实与假设之间有什么差距？', '矛盾、痛点、流程缺口', '7-Eleven单品管理'],
        ['流程需要', '流程中缺失了什么步骤？', '步骤遗漏、效率低下', '联邦快递隔夜达'],
        ['行业结构', '行业正在发生什么变化？', '快速增长、监管变化、新消费群体', '苹果重塑音乐产业'],
        ['人口变化', '人口特征发生了什么变化？', '年龄、收入、教育、地理分布', '中国养老产业'],
        ['认知变化', '什么是"理所当然"正在变化的？', '从X到Y的范式转变', '共享经济'],
        ['新知识', '有什么新知识可以整合利用？', '技术突破、科学发现', '特斯拉电池技术'],
    ]

    create_table(doc, quick_ref_headers, quick_ref_rows)

    add_heading(doc, '附录B：参考文献与延伸阅读', level=2, size=14, color=SECONDARY_COLOR, space_before=16, space_after=12)

    ref_items = [
        '德鲁克，《创新与企业家精神》，机械工业出版社',
        '德鲁克，《管理的实践》，机械工业出版社',
        '哈佛商业评论，《创新专刊》',
        '克莱顿·克里斯坦森，《创新者的窘境》，中信出版社',
        '埃里克·莱斯，《精益创业》，中信出版社',
    ]

    for item in ref_items:
        add_bullet_item(doc, item)

    # ========== 页脚 ==========
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('德鲁克论创新与企业家精神  讲师手册  第 ')
    run.font.name = CJK_FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 添加页码域
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar)
    run._element.append(instrText)
    run._element.append(fldChar2)
    run._element.append(fldChar3)

    run2 = footer_para.add_run(' 页')
    run2.font.name = CJK_FONT
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), CJK_FONT)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ========== 保存文档 ==========
    output_path = 'D:/新课开发/管理学/38-德鲁克论创新与企业家精神/讲师手册/德鲁克论创新与企业家精神_讲师手册.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_handbook()

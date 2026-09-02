"""
通用 Word 文档构建工具模块
为《共启——合作与超越：创建高效、满意的合作》课程包提供统一风格的 Word 文档生成功能。
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ============= 配色方案（蓝色/深灰主色调，国际化版权课标准） =============
COLOR_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)        # 深蓝（主色）
COLOR_SECONDARY = RGBColor(0x2C, 0x5F, 0x8E)      # 中蓝
COLOR_ACCENT = RGBColor(0xD4, 0xAF, 0x37)         # 金色（点缀）
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)      # 深灰
COLOR_LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)     # 中灰
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_BG_LIGHT = "EAF1F8"                          # 浅蓝背景
COLOR_BG_ACCENT = "F5F2E8"                         # 浅金背景
COLOR_BG_CALLOUT = "F0F4F8"                        # callout 浅蓝
COLOR_BORDER = "1F3A5F"                            # 边框色
COLOR_TABLE_HEADER = "1F3A5F"                      # 表头色


# ============= 字体配置 =============
CN_FONT = "微软雅黑"
CN_FONT_BODY = "宋体"
EN_FONT = "Calibri"
EN_FONT_TITLE = "Times New Roman"


# ============= 样式基础函数 =============
def set_run_font(run, size=11, bold=False, color=None, cn_font=CN_FONT_BODY, en_font=EN_FONT, italic=False):
    """设置 run 字体属性：中英文分别设置，支持颜色、粗体、斜体"""
    run.font.name = en_font
    run.font.size = Pt(size)
    run.font.bold = bold
    if italic:
        run.font.italic = True
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    rFonts.set(qn("w:cs"), en_font)
    rFonts.set(qn("w:eastAsia"), cn_font)


def set_paragraph_format(p, alignment=None, space_before=None, space_after=None,
                          line_spacing=None, first_line_indent=None, left_indent=None):
    """设置段落格式"""
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)


def add_paragraph(doc_or_cell, text="", size=11, bold=False, color=None,
                   cn_font=CN_FONT_BODY, en_font=EN_FONT,
                   alignment=None, space_before=0, space_after=4,
                   line_spacing=1.4, first_line_indent=None, left_indent=None,
                   style=None, italic=False):
    """添加段落（可在 doc 或 cell 中调用）"""
    p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    set_paragraph_format(p, alignment=alignment,
                         space_before=space_before,
                         space_after=space_after,
                         line_spacing=line_spacing,
                         first_line_indent=first_line_indent,
                         left_indent=left_indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color,
                     cn_font=cn_font, en_font=en_font, italic=italic)
    return p


# ============= 标题样式 =============
def setup_heading_styles(doc):
    """设置 Heading 1/2/3 样式"""
    styles = doc.styles
    for lvl, size, color in [(1, 22, COLOR_PRIMARY), (2, 16, COLOR_SECONDARY), (3, 13, COLOR_DARK_GRAY)]:
        s = styles[f"Heading {lvl}"]
        s.font.name = EN_FONT_TITLE
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        rPr = s.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), EN_FONT_TITLE)
        rFonts.set(qn("w:hAnsi"), EN_FONT_TITLE)
        rFonts.set(qn("w:eastAsia"), CN_FONT)
        pPr = s.element.get_or_add_pPr()
        spacing = pPr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            pPr.append(spacing)
        if lvl == 1:
            spacing.set(qn("w:before"), "360")
            spacing.set(qn("w:after"), "180")
        elif lvl == 2:
            spacing.set(qn("w:before"), "240")
            spacing.set(qn("w:after"), "120")
        else:
            spacing.set(qn("w:before"), "180")
            spacing.set(qn("w:after"), "80")
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), str(lvl - 1))


def add_heading(doc, text, level=1, alignment=None):
    """添加标题"""
    if alignment is None:
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.alignment = alignment
    return p


# ============= 页面设置 =============
def setup_page(doc, orientation=WD_ORIENT.PORTRAIT, top=2.5, bottom=2.5, left=2.5, right=2.5):
    """设置页面：A4 纵向，默认边距 2.5cm"""
    for section in doc.sections:
        if orientation == WD_ORIENT.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)


# ============= 页眉页脚 =============
def setup_header_footer(doc, header_left, header_right, footer_left, footer_right):
    """设置页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        # 清空再添加
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        h_para = header.add_paragraph()
        # 用制表符分左中右
        h_para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        r1 = h_para.add_run(header_left)
        set_run_font(r1, size=9, color=COLOR_LIGHT_GRAY)
        r2 = h_para.add_run("\t" + header_right)
        set_run_font(r2, size=9, color=COLOR_LIGHT_GRAY, bold=True)
        # 页眉下加横线（用一个细线段）
        h_para_pPr = h_para._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3A5F")
        pBdr.append(bottom)
        h_para_pPr.append(pBdr)

        # 页脚
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        f_para = footer.add_paragraph()
        f_para.paragraph_format.tab_stops.add_tab_stop(Cm(8), WD_ALIGN_PARAGRAPH.CENTER)
        f_para.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        r1 = f_para.add_run(footer_left)
        set_run_font(r1, size=9, color=COLOR_LIGHT_GRAY)
        r2 = f_para.add_run("\t第 ")
        set_run_font(r2, size=9, color=COLOR_LIGHT_GRAY)
        # 页码字段
        _add_page_field(f_para)
        r3 = f_para.add_run(" 页 / 共 ")
        set_run_font(r3, size=9, color=COLOR_LIGHT_GRAY)
        _add_total_pages_field(f_para)
        r4 = f_para.add_run(" 页\t")
        set_run_font(r4, size=9, color=COLOR_LIGHT_GRAY)
        r5 = f_para.add_run(footer_right)
        set_run_font(r5, size=9, color=COLOR_LIGHT_GRAY, bold=True)


def _add_page_field(paragraph):
    """添加 PAGE 字段"""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    set_run_font(run, size=9, color=COLOR_LIGHT_GRAY, bold=True)


def _add_total_pages_field(paragraph):
    """添加 NUMPAGES 字段"""
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " NUMPAGES "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    set_run_font(run, size=9, color=COLOR_LIGHT_GRAY, bold=True)


# ============= 表格函数 =============
def set_cell_background(cell, color_hex):
    """设置单元格背景色（hex 字符串，不含#）"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def set_cell_borders(cell, color=COLOR_BORDER, size="6"):
    """设置单元格四边边框"""
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)


def set_cell_vertical_alignment(cell, alignment=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = alignment


def add_table(doc, rows, cols, col_widths_cm=None, header_row=True):
    """添加表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    if header_row and rows > 0:
        for cell in table.rows[0].cells:
            set_cell_background(cell, COLOR_TABLE_HEADER)
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11, bold=True, color=COLOR_WHITE)
    return table


def fill_cell(cell, text, size=10, bold=False, color=None,
              cn_font=CN_FONT_BODY, en_font=EN_FONT,
              alignment=WD_ALIGN_PARAGRAPH.LEFT, vertical=WD_ALIGN_VERTICAL.CENTER):
    """填充单元格内容"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.3
    set_cell_vertical_alignment(cell, vertical)
    if isinstance(text, list):
        # 多行
        for i, line in enumerate(text):
            if i == 0:
                run = p.add_run(line)
            else:
                p2 = cell.add_paragraph()
                p2.alignment = alignment
                p2.paragraph_format.space_before = Pt(2)
                p2.paragraph_format.space_after = Pt(2)
                p2.paragraph_format.line_spacing = 1.3
                run = p2.add_run(line)
            set_run_font(run, size=size, bold=bold, color=color, cn_font=cn_font, en_font=en_font)
    else:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color, cn_font=cn_font, en_font=en_font)


def fill_table_row(table, row_idx, data, size=10, bold=False, color=None,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, header=False):
    """填充一行"""
    for i, txt in enumerate(data):
        cell = table.rows[row_idx].cells[i]
        fill_cell(cell, txt, size=size, bold=bold or header, color=color,
                  alignment=alignment)
        if header:
            set_cell_background(cell, COLOR_TABLE_HEADER)
            # 重设颜色为白色
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=size, bold=True, color=COLOR_WHITE)


def set_table_zebra(table, header_row=True):
    """设置表格斑马纹"""
    start = 1 if header_row else 0
    for i, row in enumerate(table.rows[start:], start=start):
        if (i - start) % 2 == 1:
            for cell in row.cells:
                set_cell_background(cell, "F5F8FC")


# ============= 特殊元素 =============
def add_horizontal_rule(doc, color=COLOR_BORDER):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_callout_box(doc, title, content_lines, bg=COLOR_BG_CALLOUT, border=COLOR_BORDER):
    """添加 callout 提示框（1x1 表格）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, bg)
    set_cell_borders(cell, color=border, size="12")
    cell.width = Cm(16)
    cell.text = ""
    # 标题
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run("💡 " + title)
    set_run_font(r1, size=11, bold=True, color=COLOR_PRIMARY)
    # 内容
    for line in content_lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.35
        r2 = p2.add_run("• " + line)
        set_run_font(r2, size=10, color=COLOR_DARK_GRAY)
    # 表格后空一行
    doc.add_paragraph()
    return table


def add_quote_box(doc, text, attribution=None):
    """添加金句/引用框"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    set_cell_background(cell, COLOR_BG_ACCENT)
    set_cell_borders(cell, color="D4AF37", size="12")
    cell.width = Cm(16)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run("「 " + text + " 」")
    set_run_font(r1, size=13, bold=True, color=COLOR_PRIMARY, italic=True)
    if attribution:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run("—— " + attribution)
        set_run_font(r2, size=10, color=COLOR_LIGHT_GRAY)
    doc.add_paragraph()
    return table


def add_bullet_list(doc, items, size=11, indent_left=0.6, color=None):
    """添加项目符号列表"""
    if color is None:
        color = COLOR_DARK_GRAY
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=2, space_after=2,
                             line_spacing=1.35,
                             left_indent=indent_left * 28.35)
        # 项目符号
        r1 = p.add_run("▍ ")
        set_run_font(r1, size=size, bold=True, color=COLOR_SECONDARY)
        r2 = p.add_run(item)
        set_run_font(r2, size=size, color=color)


def add_numbered_list(doc, items, size=11, indent_left=0.6, color=None):
    """添加编号列表"""
    if color is None:
        color = COLOR_DARK_GRAY
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                             space_before=2, space_after=2,
                             line_spacing=1.35,
                             left_indent=indent_left * 28.35)
        r1 = p.add_run(f"{i:02d} ")
        set_run_font(r1, size=size, bold=True, color=COLOR_SECONDARY)
        r2 = p.add_run(item)
        set_run_font(r2, size=size, color=color)


def add_page_break(doc):
    """添加分页"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_worksheet_lines(doc, n_lines=4, label="", size=11):
    """添加手写填空线"""
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        set_run_font(r, size=size, bold=True, color=COLOR_PRIMARY)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
    for _ in range(n_lines):
        p = doc.add_paragraph()
        r = p.add_run("______________________________________________________________")
        set_run_font(r, size=size, color=COLOR_LIGHT_GRAY)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.5


def add_section_divider(doc, title="", subtitle=""):
    """添加章节分隔页"""
    # 上方留白
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "24")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(48)
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        set_run_font(r, size=32, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT, en_font=EN_FONT_TITLE)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(subtitle)
        set_run_font(r, size=14, color=COLOR_LIGHT_GRAY, cn_font=CN_FONT)
        p.paragraph_format.space_after = Pt(36)
    # 装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("◆ ◆ ◆")
    set_run_font(r, size=12, color=COLOR_ACCENT)
    p.paragraph_format.space_after = Pt(48)
    add_page_break(doc)


# ============= 封面页 =============
def add_cover_page(doc, course_name, subtitle, version, audience_tag=""):
    """添加封面页"""
    # 顶部装饰
    for _ in range(2):
        doc.add_paragraph()
    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(course_name)
    set_run_font(r, size=32, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT, en_font=EN_FONT_TITLE)
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(8)
    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=18, color=COLOR_SECONDARY, cn_font=CN_FONT)
    p.paragraph_format.space_after = Pt(48)
    # 中部装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━" * 18)
    set_run_font(r, size=14, color=COLOR_ACCENT)
    p.paragraph_format.space_after = Pt(36)
    # 课程标签
    if audience_tag:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(audience_tag)
        set_run_font(r, size=12, color=COLOR_DARK_GRAY, cn_font=CN_FONT, italic=True)
        p.paragraph_format.space_after = Pt(72)
    # 装饰
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI 时代 · 合作重构 · 真实课题")
    set_run_font(r, size=12, bold=True, color=COLOR_PRIMARY)
    p.paragraph_format.space_after = Pt(12)
    # 版本号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"版本 {version}")
    set_run_font(r, size=11, color=COLOR_LIGHT_GRAY)
    p.paragraph_format.space_after = Pt(6)
    # 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("竞越 · 共启课程中心")
    set_run_font(r, size=11, color=COLOR_LIGHT_GRAY)
    add_page_break(doc)


# ============= 目录 =============
def add_toc(doc, items, title="目  录"):
    """添加目录（手写条目式）"""
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=26, bold=True, color=COLOR_PRIMARY, cn_font=CN_FONT, en_font=EN_FONT_TITLE)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(36)
    # 装饰线
    add_horizontal_rule(doc)
    doc.add_paragraph()
    # 条目
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # leader=2 是点
        r1 = p.add_run(f"{item['name']}")
        set_run_font(r1, size=11, bold=(item.get('level', 1) == 1), color=COLOR_PRIMARY if item.get('level', 1) == 1 else COLOR_DARK_GRAY)
        if item.get('page'):
            r2 = p.add_run(f"\t{item['page']}")
            set_run_font(r2, size=11, color=COLOR_LIGHT_GRAY)
        # 缩进
        if item.get('level', 1) > 1:
            p.paragraph_format.left_indent = Pt((item['level'] - 1) * 18)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()
    add_horizontal_rule(doc)
    add_page_break(doc)


# ============= 主入口（测试用） =============
if __name__ == "__main__":
    doc = Document()
    setup_heading_styles(doc)
    setup_page(doc)
    setup_header_footer(doc,
        header_left="共启·合作与超越",
        header_right="学员手册",
        footer_left="共启·合作与超越 学员手册 v1.0",
        footer_right="学员姓名：______")
    add_cover_page(doc, "共启·合作与超越", "AI 时代的协作能力升级", "v1.0", "学员手册")
    add_heading(doc, "第一章 测试", level=1)
    add_paragraph(doc, "这是测试段落", size=11)
    add_callout_box(doc, "提示", ["注意这里", "也注意这里"])
    add_quote_box(doc, "AI 改变了合作的方式，但不变的是人。")
    t = add_table(doc, 3, 3, col_widths_cm=[5, 5, 5])
    fill_table_row(t, 0, ["表头1", "表头2", "表头3"], header=True)
    fill_table_row(t, 1, ["内容1", "内容2", "内容3"])
    fill_table_row(t, 2, ["内容A", "内容B", "内容C"])
    set_table_zebra(t)
    doc.save("D:/CC/temp/test_output.docx")
    print("Test saved")

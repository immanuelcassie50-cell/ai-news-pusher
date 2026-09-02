#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 90天行动计划
plan = doc.add_paragraph()
r = plan.add_run('90天行动计划模板')
r.bold = True; r.font.size = Pt(16)

t4 = doc.add_table(rows=10, cols=3)
t4.style = 'Table Grid'

for i, header in enumerate(['阶段', '重点行动', '里程碑成果']):
    c = t4.cell(0, i)
    set_cell_shading(c, '4472C4')
    r = c.paragraphs[0].add_run(header)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

phases = [
    ('第1-2周', '识别与访谈', '选定值得做卡的决策\n完成第一次访谈'),
    ('第3-4周', '挖掘与结构化', '完成全部三个维度的访谈\n提炼出隐性判断规则'),
    ('第5-6周', '卡片制作', '完成决策卡初稿\n内部评审修改'),
    ('第7-8周', '评审与定稿', '完成反对者评审\n卡片定稿并指定认领人'),
    ('第9-10周', '训练活动设计', '设计三环节微课脚本\n进行第一次训练'),
    ('第11-12周', '应用与反馈', '在真实场景中应用\n收集使用反馈'),
]

for i, (phase, action, result) in enumerate(phases, 1):
    t4.cell(i, 0).paragraphs[0].add_run(phase)
    t4.cell(i, 1).paragraphs[0].add_run(action)
    t4.cell(i, 2).paragraphs[0].add_run(result)

for i in range(7, 10):
    t4.cell(i, 0).paragraphs[0].add_run(f'第{i*2+1}-{i*2+2}周')

doc.add_page_break()
print('S15: 90-day plan done')
doc.save(r'D:\CC\temp\handbook_s15.docx')

#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "D:/新课开发/变革管理/08-向上管理与高层说服术：让决策层理解容错的成本逻辑/完整课程包/04-讲师手册/讲师手册-向上管理与高层说服术.docx"

C_PRIMARY = RGBColor(0x2b, 0x2d, 0x42)
C_SECONDARY = RGBColor(0x8d, 0x99, 0xae)
C_ACCENT = RGBColor(0xef, 0x23, 0x3c)
C_LIGHT = RGBColor(0xed, 0xf2, 0xf4)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_colored_heading(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_time_box(doc, minutes):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, "ef233c")
    cell.text = f"⏱ 建议时长：{minutes}分钟"
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.bold = True
            run.font.size = Pt(10)
    doc.add_paragraph()

def add_tip_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, "edf2f4")
    cell.text = "💡 讲师提示：" + text
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = C_PRIMARY
            run.font.size = Pt(10)
    doc.add_paragraph()

def add_discussion_box(doc, question):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, "2b2d42")
    cell.text = "🎤 讨论问题：" + question
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.size = Pt(10)
    doc.add_paragraph()

def add_exercise_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, "fff3b0")
    cell.text = "✏️ 练习：" + text
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = C_PRIMARY
            run.font.size = Pt(10)
    doc.add_paragraph()

def docx_main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = C_PRIMARY

    # ---- COVER ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('讲师手册', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = C_PRIMARY

    subtitle = doc.add_heading('向上管理与高层说服术', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = C_ACCENT

    doc.add_paragraph()
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("课程编号", "08"),
        ("标准课时", "2天（每天6小时，共12小时）"),
        ("班级规模", "20-30人"),
        ("教室要求", "U型或分组布置，白板/投影"),
        ("课前准备", "阅读学员手册，收集2-3个真实变革案例"),
    ]
    for i, (k, v) in enumerate(info_data):
        cell_k = info_table.cell(i, 0)
        cell_v = info_table.cell(i, 1)
        set_cell_bg(cell_k, "2b2d42")
        cell_k.text = k
        cell_k.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell_k.paragraphs[0].runs[0].font.bold = True
        cell_v.text = v

    doc.add_page_break()

    # ---- 讲师指南概述 ----
    add_colored_heading(doc, '讲师指南概述', 1, C_PRIMARY)
    doc.add_paragraph(
        '本课程的核心教学理念是「做中学」——通过大量真实案例、角色扮演和工具表单实操，'
        '让学员在课堂上完成从「知道」到「会用」的转变。'
    )
    add_tip_box(doc, "讲师的核心角色是引导者，不是讲授者。每个模块都要留足练习时间。")

    add_colored_heading(doc, '课程总体时间分配', 2, C_SECONDARY)
    timing_table = doc.add_table(rows=6, cols=3)
    timing_table.style = 'Table Grid'
    for j, h in enumerate(["模块", "时长", "占比"]):
        cell = timing_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True
    timing_data = [
        ("模块一：理解容错成本逻辑", "2小时", "17%"),
        ("模块二：分析决策者心理", "2.5小时", "21%"),
        ("模块三：掌握说服策略", "2.5小时", "21%"),
        ("模块四：量化变革价值", "2.5小时", "21%"),
        ("模块五：练习与复盘", "2.5小时", "21%"),
    ]
    for i, row in enumerate(timing_data):
        for j, val in enumerate(row):
            cell = timing_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 模块一 ----
    add_colored_heading(doc, '模块一：理解容错成本逻辑', 1, C_PRIMARY)
    add_time_box(doc, 120)

    add_colored_heading(doc, '1.1 为什么要学这门课（20分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：让学员理解向上说服的必要性和紧迫性')
    add_discussion_box(doc, "请回想一个你经历过的变革项目，高层不支持的原因是什么？")
    add_tip_box(doc, "这个导入讨论很重要，要让每个小组都有机会分享。讲师不要急于总结，让学员自己发现共性。")

    add_colored_heading(doc, '1.2 容错成本的定义（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握Error Cost的概念和计算逻辑')
    add_exercise_box(doc, "用3分钟时间，用Error Cost的语言重新描述你部门当前的一个问题。")
    add_tip_box(doc, "这是本课程的核心概念。讲完后，让学员用自己的话复述，确保真正理解。")

    add_colored_heading(doc, '1.3 决策者的三种类型（40分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：能够识别并应对三种不同类型的决策者')
    add_discussion_box(doc, "你部门的一号位是什么类型的决策者？他的核心关切是什么？")
    add_tip_box(doc, "可以用F04决策者画像表作为工具，现场让学员填写自己项目的高层画像。")
    add_exercise_box(doc, "对照三种决策类型，分析你最近一次向高层汇报失败的原因。")

    add_colored_heading(doc, '1.4 工具表单速查（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：了解7张工具表单的功能和适用场景')
    add_tip_box(doc, "不需要详细讲解每张表，让学员翻阅一遍，建立整体认知即可。后面模块会具体使用。")

    doc.add_page_break()

    # ---- 模块二 ----
    add_colored_heading(doc, '模块二：分析决策者心理', 1, C_PRIMARY)
    add_time_box(doc, 150)

    add_colored_heading(doc, '2.1 决策者画像工具（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握使用决策者画像工具的系统方法')
    add_exercise_box(doc, "使用F04决策者画像表，分析你项目的核心决策者。准备3分钟汇报。")
    add_tip_box(doc, "这是本模块最重要的练习。建议让2-3位学员现场展示，其他学员给出反馈。")

    add_colored_heading(doc, '2.2 向上说服四步法（60分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：理解并能复述向上说服的四步框架')
    add_tip_box(doc, "四步法是本课程的框架性工具。要让学员烂熟于心，能不假思索地背出来。")
    add_discussion_box(doc, "四步中你觉得哪一步最难？为什么？")

    add_colored_heading(doc, '2.3 信任建立模型PTR（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：理解信任建立的三个维度')
    add_tip_box(doc, "PTR模型中的'R（关系）'是最容易被忽视的。提醒学员不要只关注数据和方案。")

    doc.add_page_break()

    # ---- 模块三 ----
    add_colored_heading(doc, '模块三：掌握说服策略', 1, C_PRIMARY)
    add_time_box(doc, 150)

    add_colored_heading(doc, '3.1 最小授权策略（60分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握最小授权的概念和分阶段授权路线图的设计方法')
    add_exercise_box(doc, "将你正在推动的变革项目设计为三阶段最小授权方案。")
    add_tip_box(doc, "最小授权是本课程的核心策略工具。要让学员理解，这不是魄力不足，而是专业和自信的体现。")

    add_colored_heading(doc, '3.2 汇报叙事结构（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握高效的变革汇报结构')
    add_discussion_box(doc, "你上一次汇报失败，是因为结构问题还是内容问题？")
    add_exercise_box(doc, "用叙事结构重新包装你项目的汇报，准备2分钟电梯演讲。")

    add_colored_heading(doc, '3.3 变革管理者的五个境界（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：理解变革管理者的成长路径')
    add_tip_box(doc, "这个内容适合让学员自我评估，找到自己当前的位置和下一个目标。")

    doc.add_page_break()

    # ---- 模块四 ----
    add_colored_heading(doc, '模块四：量化变革价值', 1, C_PRIMARY)
    add_time_box(doc, 150)

    add_colored_heading(doc, '4.1 Error Cost计算方法（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握Error Cost的系统计算方法')
    add_exercise_box(doc, "使用F01工具表，计算你项目的Error Cost。")
    add_tip_box(doc, "计算时提醒学员：损失估算要有依据，不能拍脑袋。用历史数据或行业报告支撑。")

    add_colored_heading(doc, '4.2 Inaction Cost计算方法（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握等待成本的概念和计算')
    add_tip_box(doc, "Inaction Cost往往比Error Cost更难量化。引导学员从机会成本角度思考。")

    add_colored_heading(doc, '4.3 ROI计算完整演示（45分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握ROI计算公式和关键指标')
    add_exercise_box(doc, "使用F03工具表，计算你项目的ROI和回收期。")
    add_tip_box(doc, "ROI计算是财务导向型决策者最关心的。确保每位学员都会算、能讲清楚。")

    add_colored_heading(doc, '4.4 常见高层质疑及应答策略（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：掌握F07高频应答卡的使用方法')
    add_tip_box(doc, "不要只讲概念，让学员两两对练：一个扮演高管提问，一个用F07应答。")

    doc.add_page_break()

    # ---- 模块五 ----
    add_colored_heading(doc, '模块五：练习与复盘', 1, C_PRIMARY)
    add_time_box(doc, 150)

    add_colored_heading(doc, '5.1 情境模拟练习说明（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：说明5个情境模拟练习的目的和规则')
    add_tip_box(doc, "模块五的核心是练习。讲师要管住嘴，把时间还给学员。")

    add_colored_heading(doc, '5.2 案例分析指引（60分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：通过3个真实案例学习向上说服的成败关键')
    for case in [
        "案例一：ERP升级提案被拒——财务导向型CEO说服策略",
        "案例二：组织变革被搁置——关系导向型CEO应对方法",
        "案例三：成功说服的关键转折——战略导向型CEO沟通策略",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(case)
    add_tip_box(doc, "每个案例分析后，让学员讨论：如果我是当事人，我会怎么做？")

    add_colored_heading(doc, '5.3 学员常见误区（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：帮助学员识别和避免常见错误')
    add_discussion_box(doc, "这些误区中，你觉得哪个最容易犯？为什么？")

    add_colored_heading(doc, '5.4 课后行动清单说明（30分钟）', 2, C_SECONDARY)
    doc.add_paragraph('教学目标：帮助学员制定课程后的落地计划')
    add_tip_box(doc, "建议每位学员在课后提交一个具体的「第一次高层汇报」计划。")

    doc.add_page_break()

    # ---- 讲师资源 ----
    add_colored_heading(doc, '讲师资源', 1, C_PRIMARY)

    add_colored_heading(doc, '常用话术库', 2, C_SECONDARY)
    scripts_table = doc.add_table(rows=6, cols=2)
    scripts_table.style = 'Table Grid'
    for j, h in enumerate(["场景", "话术"]):
        cell = scripts_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True
    scripts = [
        ("开场破冰", "今天我们来解决一个问题：为什么好方案总是被高层否决？"),
        ("引入Error Cost", "变革失败的成本不是抽象的，它是可以计算的。"),
        ("最小授权", "我不需要您批准整个项目，我只需要您给我3个月和一个试点团队。"),
        ("应对质疑", "这个风险我也考虑过，我的应对预案是……"),
        ("建立信任", "我不保证一定成功，但我保证全程透明，每周汇报进展。"),
    ]
    for i, row in enumerate(scripts):
        for j, val in enumerate(row):
            cell = scripts_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    add_colored_heading(doc, '常见问题处理', 2, C_SECONDARY)
    faq_table = doc.add_table(rows=5, cols=2)
    faq_table.style = 'Table Grid'
    for j, h in enumerate(["问题", "建议回答"]):
        cell = faq_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True
    faqs = [
        ("学员说：我们高层根本不听", "不是不听，是没有听到足够有说服力的理由。我们来练习怎么讲。"),
        ("学员说：我的项目没有数据", "没有数据本身就是问题。先用估算，过程中建立数据收集习惯。"),
        ("学员说：最小授权太慢了", "快速失败是小步快跑，最小授权是降低决策风险。两者不矛盾。"),
        ("学员说：高层的想法天天变", "那说明你们还没有建立足够的信任。先解决信任问题。"),
    ]
    for i, row in enumerate(faqs):
        for j, val in enumerate(row):
            cell = faq_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 评估标准 ----
    add_colored_heading(doc, '学员评估标准', 1, C_PRIMARY)

    eval_table = doc.add_table(rows=6, cols=3)
    eval_table.style = 'Table Grid'
    for j, h in enumerate(["维度", "权重", "评估标准"]):
        cell = eval_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True
    evals = [
        ("课堂参与", "20%", "积极发言，小组讨论贡献度高"),
        ("决策者画像", "20%", "能准确识别高层类型并说明理由"),
        ("Error Cost计算", "20%", "计算逻辑清晰，数据有依据"),
        ("最小授权方案", "20%", "三阶段设计合理，风险可控"),
        ("高频应答演练", "20%", "应答自然、有说服力"),
    ]
    for i, row in enumerate(evals):
        for j, val in enumerate(row):
            cell = eval_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_paragraph()
    add_tip_box(doc, "评估不是为了打分，而是为了帮助学员看到自己的进步空间。反馈要具体、正向。")

    # Save
    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    docx_main()
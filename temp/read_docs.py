# -*- coding: utf-8 -*-
"""读取云南磷化第四组作业文档"""
import os
import re
from docx import Document

base_path = r"D:\2026年课程\云南磷化\第一阶段作业\第四组作业\第四组作业\第一期专业-第四组-管林\第一期专业-第四组-管林"

files = {
    "1_定位表": os.path.join(base_path, "1、课题定位表-铲运机（FL105）安全认识与常见异常状态识别-第四组-管林.docx"),
    "2_大纲": os.path.join(base_path, "2、课题大纲-铲运机（FL105）安全认识与常见异常状态识别-第四组-管林.doc"),
    "4_百问百答": os.path.join(base_path, "4、百问百答-铲运机（FL105）安全认识与常见异常状态识别-第四组-管林.doc"),
    "5_测试题": os.path.join(base_path, "5、测试题-铲运机（FL105）安全认识与常见异常状态识别-第四组-管林.doc"),
}

def extract_text_from_paragraphs(paragraphs):
    """提取段落文本"""
    texts = []
    for para in paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    return texts

def read_docx(path):
    """读取docx文件"""
    doc = Document(path)
    paragraphs = doc.paragraphs
    tables = doc.tables

    # 提取段落文本
    para_texts = extract_text_from_paragraphs(paragraphs)

    # 提取表格文本
    table_texts = []
    for i, table in enumerate(tables):
        table_texts.append(f"--- 表格 {i+1} ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_texts.append(" | ".join(row_text))

    return para_texts, table_texts

# 读取所有文件
for name, path in files.items():
    print(f"\n{'='*80}")
    print(f"文件: {name}")
    print(f"路径: {path}")
    print('='*80)

    if not os.path.exists(path):
        print(f"文件不存在!")
        continue

    try:
        para_texts, table_texts = read_docx(path)
        print("\n--- 段落内容 ---")
        for text in para_texts:
            print(text)

        if table_texts:
            print("\n--- 表格内容 ---")
            for text in table_texts:
                print(text)
    except Exception as e:
        print(f"读取错误: {e}")

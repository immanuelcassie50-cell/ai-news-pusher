#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)
    
    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)
    
    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)
    
    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)
    
    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

chapters = [
    {
        'title': '第四章  隐性判断是问出来的，不是想出来的',
        'gold_quote': '每一次我就是觉得不对背后，都有一条没被说出来的经验规则。',
        'objectives': ['掌握三种追问技法', '理解隐性判断必须被逼出来', '学会挖掘我就是觉得背后的真实规则'],
        'concepts': [('模式追问法', '把模糊感受放一起对比，找共同点'), ('反事实追问法', '逼想象未发生情境，暴露判断依据'), ('角色互换追问法', '切换到教别人视角使表达变具体')],
        'key_points': ['隐性判断不是刻意隐瞒，是从来没被语言化过', '七次不对劲可能只是同一条规则的不同表达', '模式追问：找模糊感受的共同点', '反事实追问：想象未发生情境', '角色互换追问：切换到教别人视角'],
        'thinking': ['你有没有我就是觉得不对但说不清的经历？', '为什么教别人视角后表达会更具体？'],
        'exercise': ['记录你最近一次我就是觉得不对的场景，挖掘背后的隐性规则'],
        'summary': '隐性判断是问出来的，不是想出来的。模式追问、反事实追问、角色互换追问，三层问下去，那条一直没被命名的经验规则才会真正浮出水面。'
    },
    {
        'title': '第五章  决策卡是检查表加开关',
        'gold_quote': '流程图告诉你按顺序做什么，决策卡告诉你什么时候该停下来想。',
        'objectives': ['理解决策卡不是流程图的替代品', '掌握检查表和开关的设计要点', '学会用条件句加动作指令的组合写开关语句'],
        'concepts': [('检查表', '列出最容易被忽略但忽略会出大问题的那几个信号点'), ('开关', '告诉使用者什么时候该停下来重新想的一句话'), ('条件句加动作指令', '开关语句的推荐写法')],
        'key_points': ['流程图用于理解过程，决策卡用于决策现场那几分钟', '检查表负责收集，开关负责触发', '检查表控制在5-8条', '检查表动词要用具体的、能直接对照的动作词', '开关放在卡片最显眼位置'],
        'thinking': ['你现有的决策流程图在真正决策的瞬间能被打开使用吗？', '为什么检查表的动词选择这么重要？'],
        'exercise': ['选择一个你熟悉的决策场景，设计一份检查表和一个开关语句'],
        'summary': '决策卡不是流程图，是一份检查表加一个开关。检查表告诉你别漏看什么，开关告诉你什么时候该停下来重新想。两者都齐备，这张卡才有可能在真正的决策现场被打开。'
    },
    {
        'title': '第六章  触发条件写不清楚这张卡就是废纸',
        'gold_quote': '决策卡最贵的三个字不是怎么做，是什么时候。',
        'objectives': ['理解触发条件是决策卡最容易被写坏的地方', '掌握好的触发条件的标准', '学会用组合信号代替单一信号'],
        'concepts': [('触发条件', '告诉使用者什么时候该打开决策卡的判断性描述'), ('可观测信号', '不需要依赖判断力，客观上能被识别出来的具体现象'), ('组合信号', '多个维度的信号同时满足才触发的条件设计')],
        'key_points': ['触发条件决定了这张卡有没有用', '好的触发条件可观测、有具体数值或现象', '当出现异常时看起来是触发条件，实际上什么都没说', '触发条件太宽：狼来了效应', '触发条件过窄：只匹配原始案例具体情境'],
        'thinking': ['你见过的决策卡里，触发条件是怎么写的？', '为什么当出现异常时不是合格的触发条件？'],
        'exercise': ['为你的决策卡设计触发条件，让新人对照数据也能判断是否触发'],
        'summary': '触发条件写不清楚，这张卡就是废纸。写好触发条件，靠的是把那条真正起作用的判断维度，从具体情境里剥离出来，变成一句谁都能客观识别的话。'
    },
]

for ch in chapters:
    add_chapter_content(doc, ch)
    doc.add_page_break()

print('S6: Chapters 4-6 done')
doc.save(r'D:\CC\temp\handbook_s6.docx')

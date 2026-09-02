# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {})
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

output_path = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册/讲师手册-变革成果固化机制.docx"
doc = Document(output_path)

# ============ CHAPTER 3 ============
add_heading(doc, "第三章  模块二：制度固化（3小时）", 1)
add_heading(doc, "3.1 教学目标", 2)
add_para(doc, "1. 理解制度固化的五要素框架")
add_para(doc, "2. 掌握将变革职责嵌入岗位说明书的方法")
add_para(doc, "3. 掌握重新设计考核指标的方法")
add_para(doc, "4. 理解制度文本编写规范，能编写清晰的制度条文")

add_heading(doc, "3.2 时间分配", 2)
table5 = doc.add_table(rows=5, cols=3)
table5.style = 'Table Grid'
for i, h in enumerate(["教学内容", "时长", "核心产出"]):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
data5 = [
    ("制度固化的关键要素", "60分钟", "制度固化五要素框架"),
    ("如何让新流程进入岗位说明书", "45分钟", "岗位职责修订初稿"),
    ("考核指标重新设计", "45分钟", "考核指标设计案"),
    ("制度文本编写规范", "30分钟", "制度文本初稿"),
]
for row_idx, (c, d, o) in enumerate(data5):
    table5.rows[row_idx + 1].cells[0].text = c
    table5.rows[row_idx + 1].cells[1].text = d
    table5.rows[row_idx + 1].cells[2].text = o

add_heading(doc, "3.3 详细教学流程", 2)
add_heading(doc, "3.3.1 制度固化的关键要素（60分钟）", 3)
add_para(doc, "【0-10分钟】为什么要制度固化", bold=True)
add_para(doc, "制度固化是固化体系的基础层。没有制度层面的保障，变革成果依赖个人，一旦人员变动，变革就会倒退。")
add_para(doc, "制度固化的核心问题是：如何让制度从写在纸上变成执行在行动中。")

add_para(doc, "【10-30分钟】制度固化五要素框架", bold=True)
add_para(doc, "要素一：明确责任主体 - 每个流程环节都要有明确的负责人，避免集体负责等于没人负责")
add_para(doc, "要素二：设定考核指标 - 考核指标要具体、可测量，避免抽象的认真执行")
add_para(doc, "要素三：规定检查节点 - 定期检查执行情况，检查要留下记录")
add_para(doc, "要素四：明确奖惩规则 - 奖励遵守制度的行为，惩罚违反制度的行为")
add_para(doc, "要素五：持续更新机制 - 制度要定期审视和更新，根据实际情况调整")

add_para(doc, "【30-45分钟】案例分析：消失的SOP", bold=True)
add_para(doc, "情境案例：某企业引入了一套新的生产管理SOP，制度文本非常完善。但三个月后，执行率只有40%。")
add_para(doc, "讨论问题：1. 为什么完善的制度得不到执行？ 2. 五要素中哪个要素可能缺失了？ 3. 如何让制度从写在纸上变成执行在行动中？")

add_para(doc, "【45-60分钟】小组练习：制度诊断工作坊 - 每组选择一个本企业的制度文本，使用五要素框架进行诊断，找出缺失的要素，提出改进建议")

add_heading(doc, "3.3.2 如何让新流程进入岗位说明书（45分钟）", 3)
add_para(doc, "【0-10分钟】岗位说明书的重要性", bold=True)
add_para(doc, "岗位说明书是组织最重要的制度载体之一。如果新流程没有进入岗位说明书，员工会认为这是额外的工作，而不是我的本职工作。")

add_para(doc, "【10-25分钟】变革职责嵌入岗位说明书的方法", bold=True)
add_para(doc, "第一步：识别需要新增/修改的岗位职责 - 分析新流程涉及哪些岗位，这些岗位的现有职责是什么，新流程需要这些岗位做什么")
add_para(doc, "第二步：用动作语言描述职责 - 使用执行、审核、协调、监控等动作词，避免模糊的配合、参与")
add_para(doc, "第三步：设定考核标准 - 每个职责都要有对应的考核指标，指标要具体、可测量")

add_para(doc, "【25-40分钟】实操演练：岗位职责修订练习 - 每组选择一个本企业的典型岗位，针对新变革修订该岗位的岗位职责，使用模板进行编写")

add_para(doc, "【40-45分钟】常见错误和修正", bold=True)
add_para(doc, "错误一：职责描述过于笼统 | 错误二：缺少考核指标 | 错误三：职责与现有岗位重叠")

add_heading(doc, "3.3.3 考核指标重新设计（45分钟）", 3)
add_para(doc, "【0-10分钟】为什么考核指标很重要", bold=True)
add_para(doc, "考核指标是组织行为的指挥棒。如果新流程没有对应的考核指标，员工会把时间和精力放在有考核的事情上，忽略没有考核的事情。")

add_para(doc, "【10-25分钟】考核指标设计四步法", bold=True)
add_para(doc, "第一步：识别关键行为 - 新流程中哪些行为是最重要的，这些行为的结果是什么")
add_para(doc, "第二步：设计测量方式 - 如何量化这些行为/结果，数据从哪里来")
add_para(doc, "第三步：设定目标值 - 目标值要具有挑战性但可实现，考虑历史数据和行业标准")
add_para(doc, "第四步：确定权重 - 新指标在整体考核中的权重，权重太低会被忽略，权重太高会产生副作用")

add_para(doc, "【25-40分钟】实操演练：考核指标设计挑战 - 讲师给出6个场景卡片，学员分组抽取卡片，为每个场景设计考核指标，展示并讨论")

add_para(doc, "【40-45分钟】常见问题", bold=True)
add_para(doc, "指标过多怎么办：聚焦3-5个核心指标 | 数据无法获取怎么办：选择可测量的替代指标 | 指标产生副作用怎么办：设计关联指标对冲")

add_heading(doc, "3.3.4 制度文本编写规范（30分钟）", 3)
add_para(doc, "【0-10分钟】好的制度文本标准", bold=True)
add_para(doc, "好的制度文本：具体、清晰、可操作 | 坏的制度文本：模糊、笼统、无法执行")

add_para(doc, "【10-20分钟】编写规范", bold=True)
add_para(doc, "规范一：使用动作词 - 使用必须、禁止、应当等指令性词汇，避免参照执行、酌情处理")
add_para(doc, "规范二：写清楚谁、做什么、怎么做 - 责任主体要明确，行为描述要具体，执行标准要可测量")
add_para(doc, "规范三：避免歧义 - 同一个词在一份制度中表达同一个意思，避免模糊的时间描述")

add_para(doc, "【20-30分钟】对比练习 - 展示模糊条文 vs 清晰条文的对比，学员练习将模糊条文改写为清晰条文")

add_heading(doc, "3.4 核心产出", 2)
add_para(doc, "制度文本初稿")

add_heading(doc, "3.5 讲师注意事项", 2)
add_para(doc, "1. 五要素框架是本模块核心，要让学员真正掌握")
add_para(doc, "2. 制度编写规范要有大量实例，让学员看到模糊vs具体的对比")
add_para(doc, "3. 工作坊环节要给足时间，让学员真正动手写")
add_para(doc, "4. 岗位说明书和考核指标是实际工作中最实用的内容，要多练习")

# ============ CHAPTER 4 ============
add_heading(doc, "第四章  模块三：机制固化（3小时）", 1)
add_heading(doc, "4.1 教学目标", 2)
add_para(doc, "1. 理解机制与制度的区别和联系")
add_para(doc, "2. 掌握设计无人驾驶检查节点的方法")
add_para(doc, "3. 理解关键岗位继任者培养的重要性")
add_para(doc, "4. 掌握三棱镜固化体系的设计方法")

add_heading(doc, "4.2 时间分配", 2)
table6 = doc.add_table(rows=5, cols=3)
table6.style = 'Table Grid'
for i, h in enumerate(["教学内容", "时长", "核心产出"]):
    cell = table6.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
data6 = [
    ("机制 vs 制度", "45分钟", "机制设计原则"),
    ("无人驾驶检查节点设计", "60分钟", "检查节点设计案"),
    ("关键岗位继任者培养", "45分钟", "继任者培养方案"),
    ("定期审视机制", "30分钟", "审视机制设计"),
]
for row_idx, (c, d, o) in enumerate(data6):
    table6.rows[row_idx + 1].cells[0].text = c
    table6.rows[row_idx + 1].cells[1].text = d
    table6.rows[row_idx + 1].cells[2].text = o

add_heading(doc, "4.3 详细教学流程", 2)
add_heading(doc, "4.3.1 机制 vs 制度（45分钟）", 3)
add_para(doc, "【0-15分钟】概念辨析", bold=True)
add_para(doc, "制度是硬的规则：规定必须做什么、禁止做什么 | 机制是软的系统：让正确的事情自动发生 | 制度 + 机制 = 软硬兼施")

add_para(doc, "【15-30分钟】案例分析", bold=True)
add_para(doc, "案例：某企业的会议守时机制 - 制度：会议必须准时开始，迟到者罚款50元。机制：每次会议开始时，大屏幕显示本次会议应到XX人，实到XX人，开始时间XX:XX。")
add_para(doc, "问题：制度有效还是机制有效？为什么？")

add_para(doc, "【30-45分钟】核心逻辑", bold=True)
add_para(doc, "好的机制让正确的事情变得容易，错误的事情变得困难。机制设计的核心是减少摩擦、增加动力：不是让人做对，而是让做对成为默认；不是惩罚错误，而是让正确更有回报")

add_heading(doc, "4.3.2 无人驾驶检查节点设计（60分钟）", 3)
add_para(doc, "【0-15分钟】什么是无人驾驶的检查节点", bold=True)
add_para(doc, "无人驾驶的检查节点是指：即使没有人盯着，流程也会在正确的时间被执行。三个关键特征：1. 自动化触发：到时间就自动启动 2. 标准化执行：检查内容、检查方式固定 3. 可追溯记录：检查结果自动记录")

add_para(doc, "【15-35分钟】检查节点设计四步法", bold=True)
add_para(doc, "第一步：识别关键检查点 - 新流程中有哪些环节必须检查，检查的频率是什么")
add_para(doc, "第二步：设计触发机制 - 时间触发：每周一、每月1日等；事件触发：任务完成、审批通过等")
add_para(doc, "第三步：明确检查内容 - 检查的标准是什么，如何判断合格/不合格")
add_para(doc, "第四步：设计反馈机制 - 检查结果谁来接收，不合格如何处理")

add_para(doc, "【35-55分钟】实操演练：检查节点设计实战 - 每组选择一个本企业的关键流程，识别需要设计的检查节点，使用检查节点设计模板，小组展示和互评")

add_para(doc, "【55-60分钟】常见问题", bold=True)
add_para(doc, "检查太多怎么办：聚焦关键节点 | 没人愿意检查：明确责任人，配套考核 | 检查流于形式：设计抽查机制")

add_heading(doc, "4.3.3 关键岗位继任者培养（45分钟）", 3)
add_para(doc, "【0-15分钟】为什么继任者培养是固化体系中不可或缺的环节", bold=True)
add_para(doc, "很多变革依赖关键推动者。当关键推动者调岗或离职时，变革往往倒退。继任者培养的目的是：让变革不依赖任何一个人。")

add_para(doc, "【15-30分钟】继任者培养的四个要点", bold=True)
add_para(doc, "要点一：识别关键岗位 - 不是所有岗位都需要继任计划，聚焦推动变革的关键岗位")
add_para(doc, "要点二：明确能力要求 - 继任者需要具备什么能力，不仅仅是业务能力，还包括变革推动能力")
add_para(doc, "要点三：设计培养路径 - 在岗学习：参与变革项目；导师制：跟随现任者学习；轮岗制：多部门历练")
add_para(doc, "要点四：建立交接机制 - 交接清单：需要交接什么；交接时间：提前多久开始交接；交接确认：如何确认交接完成")

add_para(doc, "【30-40分钟】案例分析", bold=True)
add_para(doc, "案例：某企业数字化转型项目经理的继任培养 | 讨论：继任者培养中最大的挑战是什么？")

add_para(doc, "【40-45分钟】行动计划", bold=True)
add_para(doc, "学员思考：本企业有哪些关键岗位需要继任计划？")

add_heading(doc, "4.3.4 定期审视机制（30分钟）", 3)
add_para(doc, "【0-10分钟】为什么要定期审视", bold=True)
add_para(doc, "固化不是一次性工作，而是持续过程。定期审视的目的是：及时发现问题，及时调整。")

add_para(doc, "【10-20分钟】审视机制设计", bold=True)
add_para(doc, "月度审视：执行数据回顾 - 新流程执行率是否达标，出现问题是否及时整改")
add_para(doc, "季度审视：制度有效性评估 - 制度是否需要更新，机制是否运转正常")
add_para(doc, "年度审视：固化体系全面复盘 - 固化目标是否达成，下年度改进方向")

add_para(doc, "【20-30分钟】三棱镜固化体系", bold=True)
add_para(doc, "三棱镜模型是本模块的核心产出：棱镜面一：检查机制（确保执行）| 棱镜面二：激励机制（增加动力）| 棱镜面三：传承机制（知识转移）| 三个机制相互配合，形成完整的固化体系。")

add_heading(doc, "4.4 核心产出", 2)
add_para(doc, "三棱镜固化体系设计")

add_heading(doc, "4.5 讲师注意事项", 2)
add_para(doc, "1. 三棱镜模型是本模块亮点，要讲得生动")
add_para(doc, "2. 检查节点设计要有实操，让学员设计自己企业的检查节点")
add_para(doc, "3. 继任者培养的话题可能敏感，要以案例引导而非说教")
add_para(doc, "4. 机制设计的核心是让正确的事情自动发生，要反复强调")

# ============ CHAPTER 5 ============
add_heading(doc, "第五章  模块四：文化固化（2小时）", 1)
add_heading(doc, "5.1 教学目标", 2)
add_para(doc, "1. 理解行为习惯如何转变为文化基因")
add_para(doc, "2. 掌握让新人自动适应新流程的系统方法")
add_para(doc, "3. 理解表彰与激励在文化固化中的作用")
add_para(doc, "4. 掌握故事沉淀与传播机制的设计")

add_heading(doc, "5.2 时间分配", 2)
table7 = doc.add_table(rows=5, cols=3)
table7.style = 'Table Grid'
for i, h in enumerate(["教学内容", "时长", "核心产出"]):
    cell = table7.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
data7 = [
    ("行为习惯变文化基因", "45分钟", "习惯形成策略"),
    ("新人自动适应新流程", "45分钟", "新人适应方案"),
    ("表彰与激励设计", "45分钟", "表彰体系设计"),
    ("故事沉淀与传播", "30分钟", "故事收集机制"),
]
for row_idx, (c, d, o) in enumerate(data7):
    table7.rows[row_idx + 1].cells[0].text = c
    table7.rows[row_idx + 1].cells[1].text = d
    table7.rows[row_idx + 1].cells[2].text = o

add_heading(doc, "5.3 详细教学流程", 2)
add_heading(doc, "5.3.1 行为习惯如何转变成文化基因（45分钟）", 3)
add_para(doc, "【0-15分钟】习惯形成的心理机制", bold=True)
add_para(doc, "习惯的三个要素：1. 触发条件：特定的情境或信号 2. 行为序列：具体的行动 3. 奖励反馈：行为带来的正向结果 | 习惯形成的公式：习惯 = 触发 + 重复 + 奖励")

add_para(doc, "【15-30分钟】从行为到习惯的演变", bold=True)
add_para(doc, "0次：抗拒（新做法违反原有习惯）| 1-7次：尝试（开始接触但不稳定）| 8-21次：初步形成（开始成为下意识选择）| 22-66次：习惯（基本稳定，偶尔需要意志力）| 67次以上：自动化（完全不需要意志力）")

add_para(doc, "【30-45分钟】加速习惯形成的四大策略", bold=True)
add_para(doc, "策略一：降低启动阻力 - 让新行为更容易开始，减少执行新行为所需的步骤，提供执行工具和资源")
add_para(doc, "策略二：创造触发条件 - 在特定时间、特定地点触发特定行为，利用已有的习惯作为锚点，建立提醒机制")
add_para(doc, "策略三：提供即时奖励 - 让执行者在执行后立即得到正向反馈，避免奖励延迟，奖励要足够显著")
add_para(doc, "策略四：社会认同 - 让他人看到大家都在这样做，利用从众心理，表扬和推广典型案例")

add_heading(doc, "5.3.2 如何让新人自动适应新流程（45分钟）", 3)
add_para(doc, "【0-15分钟】新人是文化固化的最大挑战", bold=True)
add_para(doc, "为什么新人最难适应：新人没有经历变革过程，不理解为什么要这样做；新人带来的原有习惯与新文化冲突；周围的老人可能已经放松了对新流程的执行")

add_para(doc, "【15-35分钟】新人适应系统设计", bold=True)
add_para(doc, "要素一：新人在职体验（第一天开始）- 入职第一天就接触新流程，分配一个文化导师而非仅仅工作导师，让他看到每个人都在这样做")
add_para(doc, "要素二：入职培训设计 - 变革故事：为什么要这样变革；制度培训：这样做有什么制度保障；实操培训：具体应该怎么做；案例培训：典型案例和反面案例")
add_para(doc, "要素三：试用期考核 - 设定新流程执行考核，定期检查和反馈，及时纠偏")
add_para(doc, "要素四：社会融入 - 让新人在团队中找到文化同类，组织老员工分享适应经验，建立新人-老人互助机制")

add_para(doc, "【35-45分钟】实操演练：新人适应方案设计 - 每组选择一个本企业的典型变革案例，设计新人进入后的文化适应方案，使用新人文化适应90天计划模板")

add_heading(doc, "5.3.3 表彰与激励设计（45分钟）", 3)
add_para(doc, "【0-15分钟】激励在文化固化中的角色", bold=True)
add_para(doc, "激励的心理机制：人的行为受趋利避害驱动；被奖励的行为会加强，被惩罚的行为会减弱；即时激励比延迟激励效果强100倍")

add_para(doc, "【15-30分钟】表彰机制设计四要素", bold=True)
add_para(doc, "要素一：表彰对象 - 谁应该被表彰，表彰的标准是什么，表彰的频率是什么")
add_para(doc, "要素二：表彰形式 - 物质奖励：奖金、奖品、晋升；精神奖励：证书、奖杯、荣誉称号；社会奖励：公开表扬、故事传播、榜样塑造")
add_para(doc, "要素三：表彰时机 - 即时表扬：行为发生后尽快表扬；定期表彰：月度/季度/年度评选；里程碑表彰：特定阶段或成就达成时")
add_para(doc, "要素四：表彰传播 - 表彰要让更多人知道，表彰要讲清楚为什么被表彰，表彰要成为可以被学习的故事")

add_para(doc, "【30-40分钟】常见错误", bold=True)
add_para(doc, "错误一：只奖不罚 - 只表彰正面典型，不处理反面典型，导致劣币驱逐良币")
add_para(doc, "错误二：奖励错位 - 奖励了不该奖励的行为，让人困惑到底什么才是对的")
add_para(doc, "错误三：奖励疲劳 - 奖励太频繁或太容易获得，失去激励效果")
add_para(doc, "错误四：口头表彰太多 - 大家做得不错太泛泛，让人觉得是敷衍")

add_para(doc, "【40-45分钟】实操演练：表彰体系设计工作坊")

add_heading(doc, "5.3.4 故事沉淀与传播（30分钟）", 3)
add_para(doc, "【0-15分钟】故事在文化固化中的力量", bold=True)
add_para(doc, "为什么故事比道理更有效：故事激活大脑的情感区域，让人记忆深刻；故事提供情境模拟，让人知道遇到这种情况应该怎么做；故事具有传播性，人们愿意分享故事")

add_para(doc, "【15-25分钟】故事收集与传播机制", bold=True)
add_para(doc, "故事收集来源：日常观察（设立文化观察员，定期收集一线好人好事）；自我申报（鼓励员工申报自己的故事，设计简单的申报模板）；访谈挖掘（定期对变革推动者、关键用户进行访谈）")
add_para(doc, "故事传播渠道：正式渠道（培训课程中的案例分享、制度文件中的案例附录）；半正式渠道（团队例会开始时的故事时间、部门文化建设活动）；非正式渠道（同事之间的口口相传、新员工的 文化之旅）")

add_para(doc, "【25-30分钟】实操演练：故事挖掘与讲述")

add_heading(doc, "5.4 核心产出", 2)
add_para(doc, "文化固化策略方案")

add_heading(doc, "5.5 讲师注意事项", 2)
add_para(doc, "1. 习惯形成的67次规律要讲，这是行为改变的科学依据")
add_para(doc, "2. 故事沉淀机制要结合案例，让学员感受到故事的力量")
add_para(doc, "3. 表彰体系设计要有互动，学员要能看到真实案例的利弊")
add_para(doc, "4. 文化层是最高境界：文件可以抄，流程可以复制，但文化无法移植")

# ============ CHAPTER 6 ============
add_heading(doc, "第六章  模块五：固化效果检验与迭代（2小时）", 1)
add_heading(doc, "6.1 教学目标", 2)
add_para(doc, "1. 理解固化效果评估的维度")
add_para(doc, "2. 掌握设计早期预警指标的方法")
add_para(doc, "3. 学会发现固化失效时的干预策略")
add_para(doc, "4. 掌握定期复盘与迭代机制的设计")

add_heading(doc, "6.2 时间分配", 2)
table8 = doc.add_table(rows=5, cols=3)
table8.style = 'Table Grid'
for i, h in enumerate(["教学内容", "时长", "核心产出"]):
    cell = table8.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
data8 = [
    ("固化效果评估维度", "45分钟", "四维评估模型"),
    ("早期预警指标设计", "60分钟", "预警指标体系"),
    ("干预策略", "45分钟", "干预策略选择"),
    ("定期复盘与迭代", "30分钟", "复盘报告模板"),
]
for row_idx, (c, d, o) in enumerate(data8):
    table8.rows[row_idx + 1].cells[0].text = c
    table8.rows[row_idx + 1].cells[1].text = d
    table8.rows[row_idx + 1].cells[2].text = o

add_heading(doc, "6.3 详细教学流程", 2)
add_heading(doc, "6.3.1 固化效果的评估维度（45分钟）", 3)
add_para(doc, "【0-10分钟】为什么要评估固化效果", bold=True)
add_para(doc, "组织遗忘机制：时间会冲淡一切，紧迫的事情会取代重要的事情，人员变动会带来新的不确定性。评估的作用：及时发现苗头，量化固化程度，指导改进方向，增强继续改进的动力")

add_para(doc, "【10-25分钟】固化效果的四维评估模型", bold=True)
add_para(doc, "维度一：制度层评估 - 制度是否完善，制度是否更新到最新版本，制度是否为新员工所熟知")
add_para(doc, "维度二：执行层评估 - 新流程的实际执行率是多少，执行质量是否达到标准，执行是否稳定")
add_para(doc, "维度三：机制层评估 - 检查节点是否正常运转，预警机制是否有效，激励机制是否仍然有效")
add_para(doc, "维度四：文化层评估 - 新员工是否能自然适应，老员工是否仍在坚持新流程，是否有退回旧做法的声音或行为")

add_para(doc, "【25-40分钟】评估频率设计", bold=True)
add_para(doc, "月度监测（关键指标）：执行率数据、异常/投诉数据、新人适应数据")
add_para(doc, "季度评估（综合评估）：制度有效性评估、机制运转评估、文化氛围评估")
add_para(doc, "年度审视（全面复盘）：变革目标达成情况、固化体系整体有效性、下年度改进方向")

add_para(doc, "【40-45分钟】评估方法", bold=True)
add_para(doc, "方法一：数据分析 - 系统数据（执行率、合格率）、报表数据（定期汇总）、问卷数据（满意度调研）")
add_para(doc, "方法二：行为观察 - 现场走访、会议观察、流程穿越")
add_para(doc, "方法三：访谈调研 - 一线员工访谈、管理者访谈、离职员工访谈")

add_heading(doc, "6.3.2 如何设计早期预警指标（60分钟）", 3)
add_para(doc, "【0-15分钟】预警指标的设计原则", bold=True)
add_para(doc, "原则一：灵敏性 - 指标要能灵敏地反映苗头，不要等到问题已经严重才报警")
add_para(doc, "原则二：可测量 - 指标要可以量化，数据要容易获取")
add_para(doc, "原则三：可行动 - 指标要指向具体的行动，看到指标变化，知道该做什么")

add_para(doc, "【15-35分钟】预警指标设计步骤", bold=True)
add_para(doc, "步骤一：识别关键风险 - 哪些因素可能导致固化失效，哪些环节最容易破功")
add_para(doc, "步骤二：设计先行指标 - 在这些关键环节上，什么指标能提前预警")
add_para(doc, "步骤三：设定预警阈值 - 指标到多少算危险，预警分几个级别")
add_para(doc, "步骤四：设计响应机制 - 预警触发后，谁来响应，响应的时间和流程是什么")

add_para(doc, "【35-50分钟】预警响应机制", bold=True)
add_para(doc, "一般预警（黄色）：通知责任人，24小时内分析原因，72小时内提出改进方案")
add_para(doc, "严重预警（橙色）：通知上级管理者，立即召开专题会议，48小时内启动整改")
add_para(doc, "紧急预警（红色）：通知最高管理者，立即启动应急机制，24小时内形成处置方案")

add_para(doc, "【50-60分钟】实操演练：预警指标设计实战 - 每组选择一个本企业的变革案例，设计该变革的固化预警指标体系，使用固化效果预警指标体系模板")

add_heading(doc, "6.3.3 发现固化失效时的干预策略（45分钟）", 3)
add_para(doc, "【0-15分钟】固化失效的典型信号", bold=True)
add_para(doc, "信号一：执行数据下滑 - 新流程执行率连续下降，执行质量明显下降，超时情况增多")
add_para(doc, "信号二：反馈声音变化 - 听到以前那样做也没问题；听到太麻烦了，能不能简化；听到好像没人管了")
add_para(doc, "信号三：人员变动影响 - 关键推动者调岗/离职，新加入人员带来旧习惯，团队重组带来协作变化")
add_para(doc, "信号四：配套机制退化 - 检查节点不再执行，考核不再认真进行，表彰活动取消或缩水")

add_para(doc, "【15-30分钟】干预策略框架", bold=True)
add_para(doc, "策略一：快速止血（紧急干预）- 立即恢复关键检查节点，强化考核和监督，公开强调不能倒退 | 适用场景：固化失效刚刚发生，还未形成普遍现象")
add_para(doc, "策略二：根因分析（系统干预）- 分析固化失效的根本原因，是制度问题、机制问题还是文化问题，针对根本原因制定解决方案 | 适用场景：固化失效已经不是个案，而是系统性问题")
add_para(doc, "策略三：重新启动（深度干预）- 将变革视为重新上线，重新进行培训和宣导，重新设计激励和考核机制 | 适用场景：固化已经基本完全失效，需要重新来过")
add_para(doc, "策略四：迭代优化（持续改进）- 不是简单恢复原状，而是升级固化体系，识别原有体系的不足，进行优化，固化体系2.0 | 适用场景：固化过程中发现原有设计有缺陷，需要迭代")

add_para(doc, "【30-40分钟】干预策略选择矩阵", bold=True)
table9 = doc.add_table(rows=4, cols=3)
table9.style = 'Table Grid'
for i, h in enumerate(["失效程度", "失效范围", "干预策略"]):
    cell = table9.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
data9 = [
    ("轻度", "个别/偶发", "快速止血+个别辅导"),
    ("中度", "局部/普遍", "根因分析+系统干预"),
    ("重度", "全局/普遍", "重新启动+体系升级"),
]
for row_idx, (l, s, st) in enumerate(data9):
    table9.rows[row_idx + 1].cells[0].text = l
    table9.rows[row_idx + 1].cells[1].text = s
    table9.rows[row_idx + 1].cells[2].text = st

add_para(doc, "【40-45分钟】实操演练：干预策略制定")

add_heading(doc, "6.3.4 定期复盘与迭代机制（30分钟）", 3)
add_para(doc, "【0-10分钟】为什么要定期复盘", bold=True)
add_para(doc, "固化体系不是一成不变的，外部环境在变，组织在变，固化的内容也要变。定期复盘确保固化体系持续有效")

add_para(doc, "【10-25分钟】复盘机制设计", bold=True)
add_para(doc, "频率：季度复盘 - 每季度结束后的第二周进行，参与人：固化体系负责人、相关部门负责人")
add_para(doc, "内容：四步复盘法 - 1. 回顾目标：当初固化的目标是什么？2. 评估结果：现在做得怎么样？3. 分析原因：为什么好/不好？4. 总结经验：下次怎么做？")
add_para(doc, "输出：迭代计划 - 下季度固化改进重点、具体的改进措施、责任人和时间节点")

add_para(doc, "【25-30分钟】实操演练：复盘报告撰写")

add_heading(doc, "6.4 核心产出", 2)
add_para(doc, "固化效果评估体系")

add_heading(doc, "6.5 讲师注意事项", 2)
add_para(doc, "1. 预警指标设计要有实战，学员要设计自己企业的指标")
add_para(doc, "2. 干预策略选择矩阵要讲透，这是应对固化失效的关键工具")
add_para(doc, "3. 复盘报告要结合学员自己企业的情况来练习")
add_para(doc, "4. 固化效果评估是确保固化持续有效的重要保障")

doc.save(output_path)
print("Chapters 3-6 added successfully")

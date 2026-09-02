"""
生成《领航·4.0 学员手册》印刷版 Word 文档
- 基于 02_学员手册_升级版.md 的完整内容
- A4 纵向，竞越品牌色（深蓝 #1F4E79）
- 印刷级排版：封面、目录、页眉页脚、工作页留白、签名页
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# 路径
SRC_MD = Path(r"D:\2026年课程\竞越\领航：Z世代管理新策略3.0\完整课程包\04_学员手册\02_学员手册_升级版.md")
OUT_DOCX = Path(r"D:\2026年课程\竞越\领航：Z世代管理新策略3.0\完整课程包\04_学员手册\03_学员手册_完整版_印刷版.docx")

# 品牌色
COLOR_BRAND = "1F4E79"     # 竞越深蓝
COLOR_GRAY_BG = "F5F5F5"  # 工作页浅灰背景
COLOR_PARTNER = "EAEAEA"  # 配对互评灰框
COLOR_LIGHT = "D9E2F3"    # 表头浅蓝
COLOR_TABLE_BORDER = "BFBFBF"

# 字体
FONT_HEAD = "微软雅黑"
FONT_BODY = "宋体"
FONT_EN = "Calibri"


# -------------------- 辅助：OXML 操作 --------------------

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=COLOR_TABLE_BORDER, sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_paragraph_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_paragraph_borders(para, color="BFBFBF", sz="6", style="single", sides=("top","left","bottom","right")):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for s in sides:
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), style)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), color)
        pBdr.append(b)
    pPr.append(pBdr)


def add_run(para, text, *, font=FONT_BODY, size=10.5, bold=False, italic=False, color=None):
    run = para.add_run(text)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_paragraph_spacing(para, before=0, after=4, line=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


# -------------------- 文档初始化 --------------------

doc = Document()

# 默认样式
style = doc.styles["Normal"]
style.font.name = FONT_BODY
style.font.size = Pt(10.5)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:eastAsia"), FONT_BODY)

# Section 设置：A4 + 边距
for section in doc.sections:
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)


# -------------------- 头部：页眉页脚 --------------------

def setup_header_footer(section, header_left="领航·4.0 学员手册", footer_left="罗宏伟 · 领航·4.0 v1.0"):
    # 页眉：左侧标题、右侧页码
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    # 用 tab 让左右排布
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Mm(170), WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p, header_left, font=FONT_HEAD, size=9, color=COLOR_BRAND)
    add_run(p, "\t")
    # 页码字段
    run = p.add_run()
    run.font.name = FONT_HEAD
    run.font.size = Pt(9)
    rPr2 = run._element.get_or_add_rPr()
    rFonts2 = OxmlElement("w:rFonts")
    rFonts2.set(qn("w:eastAsia"), FONT_HEAD)
    rPr2.append(rFonts2)
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.set(qn("xml:space"), "preserve"); instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1); run._element.append(instrText); run._element.append(fldChar2)
    set_paragraph_borders(p, color=COLOR_BRAND, sz="4", sides=("bottom",))

    # 页脚：左侧署名，右侧"v1.0 / 2026-06"
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Mm(170), WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(p, footer_left, font=FONT_BODY, size=8, color="666666")
    add_run(p, "\t")
    add_run(p, "v1.0 / 2026-06", font=FONT_BODY, size=8, color="666666")
    set_paragraph_borders(p, color="BFBFBF", sz="4", sides=("top",))


for section in doc.sections:
    setup_header_footer(section)


# -------------------- 内容生成器 --------------------

def add_blank_line(size=8):
    p = doc.add_paragraph()
    add_run(p, "", size=size)
    set_paragraph_spacing(p, before=0, after=0, line=1)


def add_title(text, level=1, color=COLOR_BRAND, after=8, before=12):
    """level: 1=大标题, 2=二级, 3=三级, 4=小节"""
    size_map = {1: 22, 2: 16, 3: 13, 4: 11.5}
    sizes = size_map.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level >= 2 else WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, font=FONT_HEAD, size=sizes, bold=True, color=color)
    set_paragraph_spacing(p, before=before, after=after, line=1.3)
    if level == 1:
        set_paragraph_borders(p, color=COLOR_BRAND, sz="12", sides=("bottom",))
    return p


def add_h2(text):
    return add_title(text, level=2)


def add_h3(text):
    return add_title(text, level=3, color="2E5C8A", after=4, before=8)


def add_h4(text):
    return add_title(text, level=4, color="404040", after=3, before=6)


def add_para(text, size=10.5, bold=False, italic=False, color=None, font=FONT_BODY, align="left", before=0, after=3, line=1.55, indent_first=0):
    p = doc.add_paragraph()
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    add_run(p, text, size=size, bold=bold, italic=italic, color=color, font=font)
    set_paragraph_spacing(p, before=before, after=after, line=line)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(indent_first)
    return p


def add_quote(text, author=None, color="555555"):
    """金句框"""
    p = doc.add_paragraph()
    add_run(p, "“" + text + "”", size=10.5, italic=True, color=color, font=FONT_BODY)
    if author:
        add_run(p, "\n—— " + author, size=9, color="888888")
    set_paragraph_spacing(p, before=4, after=4, line=1.4)
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.right_indent = Mm(6)
    set_paragraph_borders(p, color=COLOR_BRAND, sz="4", sides=("left",))
    set_paragraph_bg(p, "F4F7FB")
    return p


def add_quote_simple(text, color="555555"):
    p = doc.add_paragraph()
    add_run(p, "“" + text + "”", size=10.5, italic=True, color=color, font=FONT_BODY)
    set_paragraph_spacing(p, before=4, after=4, line=1.4)
    p.paragraph_format.left_indent = Mm(6)
    set_paragraph_borders(p, color=COLOR_BRAND, sz="4", sides=("left",))
    return p


def add_workpage_label(num, name):
    """工作页大标题（浅灰底）"""
    p = doc.add_paragraph()
    add_run(p, f"工作页 {num}", size=12, bold=True, color="FFFFFF", font=FONT_HEAD)
    add_run(p, f"  ｜  {name}", size=14, bold=True, color="FFFFFF", font=FONT_HEAD)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=14, after=8, line=1.4)
    set_paragraph_bg(p, COLOR_BRAND)
    p.paragraph_format.left_indent = Mm(4)
    p.paragraph_format.right_indent = Mm(4)
    return p


def add_part_banner(text):
    """Part 大标题（深底色横幅）"""
    p = doc.add_paragraph()
    add_run(p, text, size=20, bold=True, color="FFFFFF", font=FONT_HEAD)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=18, line=1.4)
    set_paragraph_bg(p, COLOR_BRAND)
    p.paragraph_format.left_indent = Mm(0)
    p.paragraph_format.right_indent = Mm(0)
    return p


def add_part_subtitle(text):
    p = doc.add_paragraph()
    add_run(p, text, size=12, color=COLOR_BRAND, font=FONT_HEAD, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12, line=1.4)
    return p


def add_goals(items):
    """学习目标（带勾选方框的项目）"""
    for it in items:
        p = doc.add_paragraph()
        add_run(p, "▍ ", size=10, color=COLOR_BRAND, bold=True)
        add_run(p, it, size=10.5)
        set_paragraph_spacing(p, before=2, after=2, line=1.4)
        p.paragraph_format.left_indent = Mm(4)


def add_checklist(items):
    for it in items:
        p = doc.add_paragraph()
        add_run(p, "☐  ", size=11, color=COLOR_BRAND, bold=True)
        add_run(p, it, size=10)
        set_paragraph_spacing(p, before=1, after=1, line=1.3)
        p.paragraph_format.left_indent = Mm(2)


def add_selfcheck(items):
    """自检清单（带方框）"""
    add_h4("✅ 自检清单")
    for it in items:
        p = doc.add_paragraph()
        add_run(p, "☐ ", size=11, color=COLOR_BRAND, bold=True)
        add_run(p, it, size=10)
        set_paragraph_spacing(p, before=1, after=1, line=1.3)
        p.paragraph_format.left_indent = Mm(2)


def add_write_line(label, lines=1, after=4):
    """带标签的填写横线（横线 = 用下划线 + 空格）"""
    p = doc.add_paragraph()
    add_run(p, label + "：", size=10.5, bold=True, color="333333")
    underline_text = "\n" + ("\n".join(["_ " * 50] * lines))
    add_run(p, underline_text, size=10.5, color="888888")
    set_paragraph_spacing(p, before=2, after=after, line=1.8)
    return p


def add_write_area(hint="", lines=3):
    """空白填写区"""
    p = doc.add_paragraph()
    if hint:
        add_run(p, hint + "：\n", size=10, color="666666", italic=True)
    add_run(p, "\n".join(["_ " * 60] * lines), size=10.5, color="888888")
    set_paragraph_spacing(p, before=2, after=6, line=1.6)
    return p


def add_partner_box(title, lines=3):
    """配对互评灰框"""
    p = doc.add_paragraph()
    add_run(p, f"  {title}", size=10.5, bold=True, color="404040", font=FONT_HEAD)
    set_paragraph_spacing(p, before=6, after=2, line=1.4)
    set_paragraph_bg(p, COLOR_PARTNER)
    set_paragraph_borders(p, color="999999", sz="6")
    p.paragraph_format.left_indent = Mm(2)
    p.paragraph_format.right_indent = Mm(2)
    # 内部填写区
    p2 = doc.add_paragraph()
    set_paragraph_bg(p2, COLOR_PARTNER)
    set_paragraph_borders(p2, color="999999", sz="6", sides=("left", "right", "bottom"))
    p2.paragraph_format.left_indent = Mm(2)
    p2.paragraph_format.right_indent = Mm(2)
    add_run(p2, "\n".join(["_ " * 60] * lines), size=10.5, color="888888")
    set_paragraph_spacing(p2, before=2, after=6, line=1.6)


def add_code_block(text, size=9):
    """代码块 / ASCII 框图（保留方框字体风格）"""
    p = doc.add_paragraph()
    add_run(p, text, font="Consolas", size=size)
    set_paragraph_spacing(p, before=2, after=2, line=1.1)
    set_paragraph_bg(p, "FAFAFA")
    set_paragraph_borders(p, color="DDDDDD", sz="4")
    p.paragraph_format.left_indent = Mm(2)
    p.paragraph_format.right_indent = Mm(2)
    return p


def add_table(headers, rows, col_widths=None, header_bg=COLOR_BRAND, header_fg="FFFFFF", first_col_bold=False):
    """通用表格"""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = w

    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=10, bold=True, color=header_fg, font=FONT_HEAD)
        set_paragraph_spacing(p, before=2, after=2, line=1.2)
        set_cell_bg(cell, header_bg)
        set_cell_borders(cell, color="FFFFFF", sz="6")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 数据行
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
            # 允许换行
            for line in str(val).split("\n"):
                run = p.add_run(line)
                run.font.name = FONT_BODY
                rPr = run._element.get_or_add_rPr()
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:eastAsia"), FONT_BODY)
                rPr.append(rFonts)
                run.font.size = Pt(9.5)
                run.bold = (first_col_bold and c_idx == 0)
                if run.bold:
                    run.font.color.rgb = RGBColor.from_string(COLOR_BRAND)
                if line != str(val).split("\n")[-1]:
                    p.add_run().add_break()
            set_paragraph_spacing(p, before=1, after=1, line=1.3)
            set_cell_borders(cell, color=COLOR_TABLE_BORDER, sz="4")
            # 隔行底色
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F9F9F9")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_signature_page(title="行为承诺", commitment=None, sign_field=True):
    """签名页"""
    add_h2(f"📝 {title}")
    if commitment:
        p = doc.add_paragraph()
        add_run(p, commitment, size=11, italic=True, color="333333", font=FONT_BODY)
        set_paragraph_spacing(p, before=6, after=12, line=1.6)
        p.paragraph_format.left_indent = Mm(4)
        p.paragraph_format.right_indent = Mm(4)
        set_paragraph_bg(p, "F4F7FB")
        set_paragraph_borders(p, color=COLOR_BRAND, sz="6")
    # 我将从这一件具体的事开始
    p = doc.add_paragraph()
    add_run(p, "我将从这一件具体的事开始：", size=11, bold=True, color=COLOR_BRAND)
    set_paragraph_spacing(p, before=8, after=4, line=1.4)
    add_write_area(lines=2)
    if sign_field:
        p = doc.add_paragraph()
        add_run(p, "签名：", size=11, bold=True, color=COLOR_BRAND)
        add_run(p, "_" * 22, size=11, color="888888")
        add_run(p, "    日期：", size=11, bold=True, color=COLOR_BRAND)
        add_run(p, "_" * 18, size=11, color="888888")
        set_paragraph_spacing(p, before=8, after=8, line=1.4)


# -------------------- 读取并解析源 markdown --------------------

raw = SRC_MD.read_text(encoding="utf-8")
# 按段落切分（保留空行作为节）
lines = raw.split("\n")
# 清理首尾
while lines and not lines[0].strip():
    lines.pop(0)

# 我们采用"自定义内容流"而非"机械解析 markdown"——直接用结构化的内容书写更可控。
# 接下来按章节展开：


# ============================================================
# 封面
# ============================================================

# 顶端留白
for _ in range(4):
    add_blank_line(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "领  航  ·  4.0", size=42, bold=True, color=COLOR_BRAND, font=FONT_HEAD)
set_paragraph_spacing(p, before=0, after=14, line=1.2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "AI 时代的 Z 世代管理新策略", size=18, color="333333", font=FONT_HEAD)
set_paragraph_spacing(p, before=0, after=6, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "—— 学员手册（升级版 · 印刷版） ——", size=12, color=COLOR_BRAND, font=FONT_HEAD)
set_paragraph_spacing(p, before=0, after=28, line=1.4)

# 装饰横线
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "─" * 30, size=12, color=COLOR_BRAND)
set_paragraph_spacing(p, before=0, after=24, line=1)

# 课程元信息
meta = [
    ("课程全称", "领航·4.0——AI 时代的 Z 世代管理新策略"),
    ("课程讲师", "罗宏伟"),
    ("课程时长", "2 天（合计 14 学时）"),
    ("手册版本", "v1.0（升级版）"),
    ("印刷版生成", "2026 年 6 月"),
    ("适用对象", "本课程授权学员"),
]
table = doc.add_table(rows=len(meta), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for i, (k, v) in enumerate(meta):
    row = table.rows[i]
    c1, c2 = row.cells
    c1.width = Mm(40); c2.width = Mm(100)
    c1.text = ""; c2.text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, k, size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p1, before=4, after=4, line=1.3)
    set_cell_bg(c1, COLOR_BRAND)
    set_cell_borders(c1, color="FFFFFF", sz="4")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p2, v, size=11, color="333333", font=FONT_BODY)
    set_paragraph_spacing(p2, before=4, after=4, line=1.3)
    set_cell_bg(c2, "F9F9F9")
    set_cell_borders(c2, color=COLOR_TABLE_BORDER, sz="4")
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 留白
for _ in range(4):
    add_blank_line(8)

# 版权页脚
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "本手册仅供本课程学员使用 · 未经授权不得复制传播", size=9, color="888888", font=FONT_BODY)
set_paragraph_spacing(p, before=12, after=4, line=1.3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "开发者：罗宏伟", size=10, color="666666", font=FONT_BODY)
set_paragraph_spacing(p, before=0, after=0, line=1.3)

add_page_break()


# ============================================================
# 学员信息页
# ============================================================

add_title("学员信息（必填）", level=2, after=10)

info = [
    ("学员姓名", ""),
    ("所在部门 / 团队", ""),
    ("岗位 / 职务", ""),
    ("课程日期", ""),
    ("分析对象（Z 世代员工代号）", ""),
    ("问责伙伴（Day 2 下午填写）", ""),
    ("讲师签名", ""),
]
table = doc.add_table(rows=len(info), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for i, (k, _) in enumerate(info):
    row = table.rows[i]
    c1, c2 = row.cells
    c1.width = Mm(50); c2.width = Mm(110)
    c1.text = ""; c2.text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, k, size=10.5, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p1, before=6, after=6, line=1.3)
    set_cell_bg(c1, COLOR_BRAND)
    set_cell_borders(c1, color="FFFFFF", sz="4")
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p2 = c2.paragraphs[0]
    add_run(p2, " ", size=10.5, font=FONT_BODY)
    set_paragraph_spacing(p2, before=6, after=6, line=1.3)
    set_cell_borders(c2, color=COLOR_TABLE_BORDER, sz="4")
    c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

add_blank_line(8)

# 分析对象背景笔记
add_h3("分析对象背景笔记（随时补充）")
add_write_line("员工代号", lines=1)
add_para("岗位和主要工作内容：", bold=True, size=10.5, after=2)
add_write_area(lines=2)
add_para("让我最困惑的一个行为或情境：", bold=True, size=10.5, after=2)
add_write_area(lines=2)
add_para("他的 AI 工具使用情况（我所了解的）：", bold=True, size=10.5, after=2)
add_write_area(lines=2)

# 金句
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "“领先半步，吃尽红利。”", size=13, italic=True, color=COLOR_BRAND, font=FONT_HEAD)
set_paragraph_spacing(p, before=12, after=2, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "AI 时代的管理者，比的不是懂 AI 比员工多，", size=10, color="666666", italic=True)
p.add_run().add_break()
add_run(p, "是比员工先看到管理要升级的地方。", size=10, color="666666", italic=True)
set_paragraph_spacing(p, before=0, after=0, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "—— 罗宏伟", size=10, color="888888")
set_paragraph_spacing(p, before=4, after=12, line=1.2)

add_page_break()


# ============================================================
# 目录
# ============================================================

add_title("目  录", level=1, after=12)

toc_items = [
    ("引言：在我们开始之前", 5),
    ("    表单 0.1  五感诊断前测", 6),
    ("    表单 0.2  管理理念前测", 7),
    ("    表单 0.3  课前任务确认", 8),
    ("", 0),
    ("Part 1  认知力——认识 AI 时代的 Z 世代", 9),
    ("    工作页 1   快·变·连·我（AI 时代版本）", 10),
    ("    工作页 2   五感诊断工作页（含价值感）", 12),
    ("    工作页 3   AI 时代 Z 世代「三不怕」认知框架", 15),
    ("    工作页 4   管理理念自测与四座山移除", 17),
    ("    Part 1 知识框架 + 行为承诺", 19),
    ("", 0),
    ("Part 2  适应力——迎接 AI 时代的 Z 世代", 20),
    ("    工作页 5   融入期四阶段（AI 时代版）", 21),
    ("    工作页 6   1+3 任务清单 AI 升级版", 24),
    ("    工作页 7   坦诚交流策略（AI 时代版）", 26),
    ("    工作页 8   角色扮演记录表", 28),
    ("    Part 2 知识框架 + 行为承诺", 30),
    ("", 0),
    ("Part 3  链接力——保鲜 AI 时代的 Z 世代", 31),
    ("    工作页 9   三维分析 + 5W2H+H 任务分配", 32),
    ("    工作页 10  辅导对话五步流程（AI 时代版）", 35),
    ("    工作页 11  常见辅导错误速查（含 AI 新错误）", 37),
    ("    Part 3 知识框架 + 行为承诺", 39),
    ("", 0),
    ("Part 4  愿景力——引爆 AI 时代的小宇宙", 40),
    ("    工作页 12  八大内驱动力画像（AI 时代交互分析）", 41),
    ("    工作页 13  游戏设计画布 4.0", 43),
    ("    Part 4 知识框架 + 行为承诺", 45),
    ("", 0),
    ("课程收尾：四力整合 + 30 天管理改进清单", 46),
    ("    表单 C.1  五感诊断重测", 47),
    ("    表单 C.2  管理理念重测", 48),
    ("    表单 C.3  30 天管理改进清单", 49),
    ("    我的承诺签名", 52),
    ("    致出发的你", 53),
    ("", 0),
    ("附录一：术语速查表", 54),
    ("附录二：N 个工具速查索引", 55),
]

# 用一个 2 列表（左标题，右页码）做目录
table = doc.add_table(rows=len(toc_items), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
for i, (title, page) in enumerate(toc_items):
    row = table.rows[i]
    c1, c2 = row.cells
    c1.width = Mm(135); c2.width = Mm(25)
    c1.text = ""; c2.text = ""
    p1 = c1.paragraphs[0]
    add_run(p1, title, size=10.5, bold=(not title.startswith("    ") and title != ""), color=COLOR_BRAND if (not title.startswith("    ") and title != "") else "333333", font=FONT_BODY if title.startswith("    ") or title=="" else FONT_HEAD)
    set_paragraph_spacing(p1, before=2, after=2, line=1.4)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p2, str(page) if page else "", size=10.5, color="666666", font=FONT_BODY)
    set_paragraph_spacing(p2, before=2, after=2, line=1.4)
    # 清掉默认边框
    for c in (c1, c2):
        tcPr = c._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "nil")
            tcBorders.append(b)
        tcPr.append(tcBorders)

add_page_break()


# ============================================================
# 引言
# ============================================================

add_title("引言：在我们开始之前", level=1, after=10)

add_h3("🗺️ 课程全景图")
add_para("四力模型，三个层次，一个方向：", size=11, bold=True, color=COLOR_BRAND, after=8)

# 全景图 ASCII
panorama = """╔════════════════════════════════════════════════════════════╗
║   Part 4 愿景力：引爆 AI 时代的小宇宙                    ║
║   ─ 八大内驱动力 ─ 游戏设计画布 4.0 ─ 30 天激励起点    ║
╠════════════════════════════════════════════════════════════╣
║   Part 3 链接力：保鲜 AI 时代的 Z 世代                   ║
║   ─ 5W2H+H 任务分配 ─ 辅导对话五步流程 ─                ║
╠════════════════════════════════════════════════════════════╣
║   Part 2 适应力：迎接 AI 时代的 Z 世代                   ║
║   ─ 融入期四阶段 ─ 1+3 任务清单 AI 版 ─ 坦诚交流 ─      ║
╠════════════════════════════════════════════════════════════╣
║   Part 1 认知力：认识 AI 时代的 Z 世代                   ║
║   ─ 五感驱动模型 ─ 三不怕 ─ 管理理念更新 ─              ║
╠════════════════════════════════════════════════════════════╣
║              Z 世代的时代背景                            ║
║   （快 · 变 · 连 · 我 + AI 工具原住民新身份焦虑）       ║
╚════════════════════════════════════════════════════════════╝"""
add_code_block(panorama, size=8.5)

add_para("Part 1（认知力）—— 理解这代人「为什么这么做」，知道他们冰山下在驱动什么。AI 时代新增第五感「价值感」，是他们最深的焦虑。", size=10.5, after=4)
add_para("Part 2（适应力）—— 知道他们入职后经历什么阶段，每个阶段管理者该做什么。AI 工具如何用，是融入期的新课题。", size=10.5, after=4)
add_para("Part 3（链接力）—— 日常怎么带他们。任务分配怎么写、辅导对话怎么谈。AI 不改变基本面，但改变「人机分工」那一面。", size=10.5, after=4)
add_para("Part 4（愿景力）—— 怎么点燃他们。八大内驱力 + 游戏化设计。AI 既是放大器，也是新的「走捷径风险源」。", size=10.5, after=10)

# 表单 0.1
add_h3("【表单 0.1】五感诊断前测")

add_para("目的：了解你现在的起点，课程结束后用同一张表重测，看清变化。", size=10, color="555555", after=2)
add_para("要求：针对你的分析对象（一个真实 Z 世代员工）打分，不需要「好看」，这张表只有你和你的问责伙伴看。", size=10, color="555555", after=2)
add_para("时间：8 分钟", size=10, color="555555", after=8)

add_para("针对我的分析对象：", bold=True, size=11, color=COLOR_BRAND, after=4)

# 五感表格
senses = ["节奏感", "存在感", "位置感", "掌控感", "价值感（AI 时代新增）"]
sense_desc = [
    "工作节奏有可预期的韵律",
    "自己的贡献被精准看见",
    "清晰的方向感和角色定位",
    "工作方式上有自主权",
    "确认自己有 AI 替代不了的贡献",
]
rows = []
for s, d in zip(senses, sense_desc):
    rows.append([s, d, "○  ○  ○  ○"])

table = doc.add_table(rows=1 + len(rows), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
hdr = table.rows[0]
for i, h in enumerate(["感", "核心需求描述", "几乎不 / 偶尔 / 经常 / 充足"]):
    c = hdr.cells[i]
    c.text = ""
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, size=10, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p, before=2, after=2)
    set_cell_bg(c, COLOR_BRAND)
    set_cell_borders(c, color="FFFFFF", sz="4")
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

widths = [Mm(35), Mm(85), Mm(45)]
for i, row in enumerate(rows):
    tr = table.rows[i+1]
    for j, val in enumerate(row):
        c = tr.cells[j]
        c.text = ""
        c.width = widths[j]
        p = c.paragraphs[0]
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, size=10, bold=True, color=COLOR_BRAND, font=FONT_HEAD)
        elif j == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_run(p, val, size=10, color="333333")
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, val, size=10, color="666666")
        set_paragraph_spacing(p, before=3, after=3, line=1.3)
        set_cell_borders(c, color=COLOR_TABLE_BORDER, sz="4")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if i % 2 == 1:
            set_cell_bg(c, "F9F9F9")

add_blank_line(4)

# 雷达图说明（手绘）
add_para("📌 雷达图（请手绘）：", bold=True, color=COLOR_BRAND, size=11, after=2)
add_para("在每个轴上找到你的评分位置，把 5 个点连起来形成一个五边形——", size=10, color="555555", after=4)

radar = """                  价值感
                    5
                    |
                4   |   4
              3     |     3
            2       |       2
          1         |         1
        节奏感——————+——————存在感
          5         |         5
            4       |       4
              3     |     3
                2   |   2
                    |
                位置感————掌控感
                  5"""
add_code_block(radar, size=8)

add_write_line("最缺失的一感", lines=1)
add_para("具体行为表现（在我的分析对象身上）：", bold=True, size=10.5, color="333333", after=2)
add_write_area(lines=3)

add_page_break()


# 表单 0.2
add_h3("【表单 0.2】管理理念前测")

add_para("目的：了解你的管理起点，对比课后的理念更新。", size=10, color="555555", after=2)
add_para("要求：如实打钩，10 分钟以内填完。", size=10, color="555555", after=8)

mgmt_items = [
    "我能在 30 分钟内说出 Z 世代员工的 3 个核心内驱动力",
    "我能在 Z 世代入职一周内说清楚团队对 AI 工具的态度",
    "我能用一句话说清楚「我的 Z 世代下属最不可替代的价值是什么」",
    "当下属说「这是 AI 做的」时，我能接住而不是质疑",
    "我能在任务分配时说清楚「哪部分你必须自己判断」",
    "我能区分下属的成长是「真实能力提升」还是「AI 代劳」",
    "我能识别下属「AI 走捷径」的风险并主动设计机制",
]

mgmt_rows = [[str(i+1) + ". " + it, "○    ○    ○    ○"] for i, it in enumerate(mgmt_items)]
add_table(["管理特征", "几乎不 / 偶尔 / 经常 / 总是"], mgmt_rows, col_widths=[Mm(140), Mm(30)])

add_blank_line(4)
add_write_line("我打「几乎不」或「偶尔」的有 ___ 行", lines=1)
add_para("这些就是这门课的重点方向。", size=10, color="666666", italic=True, after=8)

# 表单 0.3
add_h3("【表单 0.3】课前任务确认")
add_para("请在课前 7 天内完成 5 项任务，课中对照手册使用。", size=10, color="555555", after=4)
add_checklist([
    "A. 和分析对象做一次 5-10 分钟的非正式交流（自然场合即可）",
    "B. 观察并写下他让你最困惑的一个行为",
    "C. 完成 5 感诊断前测（上一节）",
    "D. 阅读 RPD 设计分析第 1-2 节",
    "E. 找好问责伙伴（同事 / 同班学员 / 管理者朋友）",
])
add_blank_line(4)
add_para("我在任务 A 里的一个观察：", bold=True, size=10.5, color="333333", after=2)
add_write_area(lines=3)

add_page_break()


# ============================================================
# Part 1 认知力
# ============================================================

add_part_banner("Part 1")
add_part_subtitle("认知力——认识 AI 时代的 Z 世代")

add_quote("改变你对 Z 世代的看法，比改变他们的行为更有效。", author="罗宏伟")

add_h3("🎯 Part 1 学习目标")
add_goals([
    "用「快·变·连·我」理解 Z 世代的时代背景（含 AI 时代强化）",
    "用五感驱动模型（含价值感）诊断 Z 世代员工的「冰山下」驱动力",
    "用「三不怕」认知框架识别 AI 时代 Z 世代的三个新特征，并掌握转化方向",
    "反思自己的管理理念，更新 4 个 AI 时代管理策略",
])

add_h3("🗺️ Part 1 内容导航")
add_para("4 张工作页 + 4 个核心概念 + 4 道练习 + 1 张行为承诺签名页", size=10.5, color="555555", after=6)

add_page_break()


# ----- 工作页 1 -----
add_workpage_label(1, "快·变·连·我（AI 时代版本）")

add_h3("知识点 1.1：Z 世代的四个时代基因")
add_para("Z 世代不是凭空出现的，他们是被四个时代基因塑造的——", size=10.5, after=4)

genes = [
    ("快", "习惯了即时反馈，AI 工具让「快」成为常态"),
    ("变", "在变化中长大，AI 让变化速度从「年」变成「月」"),
    ("连", "天然连接的一代，AI 让连接的对象扩展到「非人」"),
    ("我", "自我意识觉醒，AI 时代遇到新挑战：「我」的独特性在哪里"),
]
for k, v in genes:
    p = doc.add_paragraph()
    add_run(p, f"  {k}  ", size=12, bold=True, color="FFFFFF", font=FONT_HEAD)
    add_run(p, " " + v, size=10.5, color="333333")
    set_paragraph_spacing(p, before=2, after=2, line=1.4)
    set_paragraph_bg(p, "EAF1F8")
    p.paragraph_format.left_indent = Mm(2)
    p.paragraph_format.right_indent = Mm(2)
    # 重写第一个 run 让它有彩色背景
    # 简化：去掉背景，整体灰底即可

# 简化版：直接列
add_blank_line(4)
add_para("针对我的分析对象，分别写下 AI 时代强化后的具体表现：", bold=True, color=COLOR_BRAND, size=11, after=4)

add_para("快（AI 时代强化）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_para("他最近的哪个行为让你感受到了「快」的升级？", size=10, color="555555", after=2)
add_write_area(lines=2)

add_para("变（AI 时代强化）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_para("他最近适应了什么新变化？适应速度怎么样？", size=10, color="555555", after=2)
add_write_area(lines=2)

add_para("连（AI 时代新层——AI 协作连接）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_para("他和 AI 工具的关系是什么样的？是「工具伙伴」还是「代劳依赖」？", size=10, color="555555", after=2)
add_write_area(lines=2)

add_para("我（AI 时代挑战——价值焦虑）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_para("他有没有流露过「AI 都能做这些，我还能做什么」的表达或行为？", size=10, color="555555", after=2)
add_write_area(lines=2)

add_page_break()

# 练习 1-A
add_h3("练习 1-A：快·变·连·我的 AI 配对")
add_para("配对分享（3 分钟）后，写下你的新发现：", size=10.5, after=2)
add_write_area(lines=3)
add_partner_box("💬 配对互评（5 分钟）", lines=3)

add_selfcheck([
    "我对「快·变·连·我」每个维度都写了具体行为",
    "我没有停留在「觉得是这样」，而是写了「看到的行为」",
    "我能在 30 秒内向搭档说清楚我的分析对象最像哪个维度",
])

add_page_break()


# ----- 工作页 2 -----
add_workpage_label(2, "五感诊断工作页（含价值感）")

add_h3("知识点 2.1：五感驱动模型（AI 时代 4.0 版）")
add_para("Z 世代的行为背后，是五个「冰山下」的内驱动力。", size=10.5, after=2)
add_para("原版四感（节奏感 / 存在感 / 位置感 / 掌控感）已经过时——AI 时代必须增加第五感：", size=10.5, after=2)
add_para("价值感。", size=12, bold=True, color="C00000", font=FONT_HEAD, after=8)

# 五感图
five_senses_diagram = """            价值感
           （AI 时代新增）
              /\\
             /  \\
   节奏感 ————+———— 存在感
            \\ |  /
             \\| /
   位置感 ————+———— 掌控感"""
add_code_block(five_senses_diagram, size=10)

add_h3("五感详细诊断")
add_para("针对我的分析对象，逐感填写：", bold=True, color=COLOR_BRAND, size=11, after=4)

five_sense_questions = [
    ("节奏感", "AI 时代升级：AI 拉高了响应期待（「一个小时内有信号」是新的基线）",
     "他最常抱怨或表现出不满的「等待」是什么？",
     "我的应对策略：建立响应节奏约定（明确「我响应 X 小时，紧急 Y 通道」）"),
    ("存在感", "AI 时代危机：他有没有过「被认可的只是会用 AI」的感受？",
     "他最近有没有过类似表达或行为？",
     "我的应对策略：在反馈里主动说出他的人类判断贡献（不是「你 AI 用得好」，而是「你在 X 里的 Y 判断带来了 Z」）"),
    ("位置感", "AI 时代新维度：他有没有过「AI 工具能力比管理者还强」的隐性优势感？",
     "他最近有没有类似的表达？",
     "我的应对策略：主动承认并欢迎他的 AI 优势，把它变成团队资产（不是权威威胁）"),
    ("掌控感", "AI 时代版本：他有没有过「被规定只能用某个 AI 工具」的抵触？",
     "他最近有没有类似表达？",
     "我的应对策略：给出人类判断的清晰边界，工具选择权归他"),
    ("价值感", "（AI 时代新增）核心焦虑：「如果 AI 能做所有这些，我的不可替代性在哪里？」",
     "他最近有没有过类似的行为或表达？（频繁换方向、抵触重复任务、过度依赖 AI 或过度抵触 AI）",
     "我的应对策略：和他一起识别「只有他才能做的那部分」，定期更新（不是一次性的安慰）"),
]

for sense, note, q, strategy in five_sense_questions:
    p = doc.add_paragraph()
    add_run(p, f"  {sense}  ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    add_run(p, "  " + note, size=10, color="333333", italic=True)
    set_paragraph_spacing(p, before=6, after=2, line=1.3)
    set_paragraph_bg(p, "EAF1F8")
    set_paragraph_borders(p, color=COLOR_BRAND, sz="4", sides=("left",))
    p.paragraph_format.left_indent = Mm(2)

    add_para(q, size=10, color="555555", after=2)
    add_write_area(lines=2)

    add_para(strategy, size=10, color="C00000", after=4)
    add_write_area(lines=2)

add_page_break()

# 5 感雷达图
add_h3("📌 5 感雷达图（请手绘）")
add_code_block(radar, size=8)

add_write_line("五边形告诉我什么？", lines=1)
add_para("哪个角最凹（最缺）？_________________________", size=10.5, after=2)
add_para("形状对称吗？", size=10.5, after=2)
add_write_area(lines=2)

add_selfcheck([
    "我对五个感都写了具体行为，不只是打分",
    "我的应对策略是「下周就做」的，不是泛泛而谈",
    "雷达图我已经画好，并能看到「形状」传递的信息",
])

# 练习 2-A
add_h3("练习 2-A：五感诊断互评")
add_para("配对互评（5 分钟）：", size=10, color="555555", after=2)
add_checklist([
    "我读给搭档听，让他指出最缺的那一感「对不对」",
    "我请搭档说出「如果他是你，他会怎么应对」",
    "我记录搭档的视角，作为我的补充策略",
])
add_blank_line(2)
add_partner_box("💬 搭档补充的策略", lines=3)

add_page_break()


# ----- 工作页 3 -----
add_workpage_label(3, "AI 时代 Z 世代「三不怕」认知框架")

add_h3("知识点 3.1：什么是「三不怕」")
add_para("AI 时代的 Z 世代有三个让管理者本能警觉的「不怕」：", size=10.5, after=4)

sanbupa_rows = [
    ["不怕比你懂 AI", "工具平权意识——AI 知识不该和职级挂钩", "威胁感：「我的权威被挑战了」", "主动问「你用什么工具处理 X」，把他变成内部 AI 资源"],
    ["不怕说「这是 AI 做的」", "比前辈更坦然地说出 AI 参与", "困惑 / 愤怒：「那你到底做了什么？」", "用「AI 做了，那你的判断体现在哪里」开启归因对话"],
    ["不怕反向带教", "愿意教管理者用 AI", "不舒服：「我需要被员工教？」", "把反向带教变成团队学习文化的信号，主动邀请他分享"],
]
add_table(["不怕", "含义", "管理者本能反应", "转化方向"], sanbupa_rows, col_widths=[Mm(30), Mm(45), Mm(45), Mm(50)])

add_blank_line(4)
add_para("针对我的分析对象，分别填写：", bold=True, color=COLOR_BRAND, size=11, after=4)

for t in [
    ("「不怕比你懂 AI」", "他最近的表现（AI 用得比管理者熟练、对 AI 工具很有自信等）",
     "我的本能反应：", "我的转化方向——我打算用的一个具体行动（本周内可执行）："),
    ("「不怕说『这是 AI 做的』」", "他最近说过类似的话（请回忆原话）",
     "我的本能反应：", "当他说「这是 AI 做的」，我的下一句话（写一句真实的话术，不是原则）："),
    ("「不怕反向带教」", "他最近有没有主动教过我或团队什么？",
     "我的本能反应：", "我打算在什么场合主动邀请他分享（具体场景）："),
]:
    title, q1, q2, q3 = t
    p = doc.add_paragraph()
    add_run(p, "  " + title + "  ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p, before=8, after=2, line=1.3)
    set_paragraph_bg(p, COLOR_BRAND)
    p.paragraph_format.left_indent = Mm(2)

    add_para(q1, size=10, color="555555", after=2)
    add_write_area(lines=2)
    add_para(q2, size=10, color="555555", after=2)
    add_write_area(lines=2)
    add_para(q3, size=10, color="C00000", bold=True, after=2)
    add_write_area(lines=2)

add_page_break()

# 练习 3-A
add_h3("练习 3-A：三不怕转化行动卡")
add_para("选三个转化行动中最容易做的那一个，把它做成一张「行动卡」贴在你的工位：", size=10.5, after=4)

action_card = """┌────────────────────────────────────┐
│  本周三不怕行动卡                  │
├────────────────────────────────────┤
│  □ 不怕比你懂 AI                   │
│    → 我的行动：______________      │
│  □ 不怕说"这是 AI 做的"            │
│    → 我的下一句话：__________      │
│  □ 不怕反向带教                    │
│    → 邀请场合：______________      │
│                                    │
│  执行时间：__________              │
│  搭档见证：__________              │
└────────────────────────────────────┘"""
add_code_block(action_card, size=9)

add_selfcheck([
    "我对三个「不怕」都识别了具体表现，不是抽象理解",
    "我的三个转化行动是「下周可执行」的，不是「理念上同意」",
    "我写了一句真实话术，不是原则表述",
])

add_page_break()


# ----- 工作页 4 -----
add_workpage_label(4, "管理理念自测与四座山移除")

add_h3("知识点 4.1：移除四座山（AI 时代升级版）")
add_para("管理者和 Z 世代之间，常有四座隐形的山。每座山都有「AI 时代特别版本」——", size=10.5, after=4)

mountain_rows = [
    ["「我吃的盐比你吃的米多」",
     "倚老卖老，用经验压制新想法",
     "管理者用「我当年没用 AI 也做到了」压制 Z 世代的 AI 优势",
     "把「我的经验」换成「我们一起探索」，AI 时代管理者没有 AI 经验优势"],
    ["「你应该听我的」",
     "单向命令，不容讨论",
     "管理者用「必须按我说的方法做」压制 Z 世代的 AI 工具选择",
     "在人类判断边界内，给工具选择权"],
    ["「你想得太理想了」",
     "质疑年轻人的想法",
     "管理者用「AI 没那么神」否定 Z 世代的 AI 应用思路",
     "先听完他的方案再判断，不要预设"],
    ["「我们当年也是这样过来的」",
     "用自己经历否定 Z 世代独特性",
     "管理者用「我们当年没有 AI 也能干活」否定 AI 时代的新困境",
     "承认时代变了，他的困境是真实的"],
]
add_table(["山", "原版症状", "AI 时代版本", "移除方向"], mountain_rows,
          col_widths=[Mm(40), Mm(40), Mm(45), Mm(45)])

add_blank_line(4)
add_para("针对我的情况：", bold=True, color=COLOR_BRAND, size=11, after=4)
add_para("对我影响最大的一座山是：__________ 山", size=10.5, after=2)
add_para("移除这一座山，我 30 天内要做的一件事：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("知识点 4.2：AI 时代三个新管理策略")

strategies = [
    ("软化冲突 AI 升级", "把「你和我的分歧」换成「我们和方案的差异」"),
    ("共建 AI 协作规范", "不是「禁止用 AI」，是「我们一起制定怎么用 AI」"),
    ("开放逆向学习", "主动说「这个你比我懂，你教我」"),
]
for k, v in strategies:
    p = doc.add_paragraph()
    add_run(p, "  " + k + "  ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    add_run(p, "  " + v, size=10.5, color="333333")
    set_paragraph_spacing(p, before=4, after=4, line=1.3)
    set_paragraph_bg(p, "EAF1F8")
    p.paragraph_format.left_indent = Mm(2)
    add_para("针对我的分析对象，我打算这样用：", size=10, color="555555", after=2)
    add_write_area(lines=2)

add_h3("📊 管理理念自测（自评量表）")
add_para("针对自己目前的真实情况打分（1-5 分，1=完全不符合，5=完全符合）：", size=10, color="555555", after=4)

belief_items = [
    "我理解 Z 世代的「快·变·连·我」，不只用「我当年」判断他们",
    "我能用五感分析员工的冰山下驱动力",
    "我能识别 AI 时代三不怕并把它转化为管理优势",
    "我能让员工感到「被看见」的是他的判断贡献，不是 AI 产出",
    "我能和员工一起建立 AI 协作规范，不是单向规定",
    "我愿意承认「AI 这个领域他比我懂」，并请他教我",
]

belief_rows = [[str(i+1) + ". " + it, "○", "○", "○", "○", "○"] for i, it in enumerate(belief_items)]
add_table(["管理理念", "1", "2", "3", "4", "5"], belief_rows,
          col_widths=[Mm(110), Mm(12), Mm(12), Mm(12), Mm(12), Mm(12)])

add_blank_line(4)
add_para("最让我意外的一项（分数最低或最高）：_______________________", size=10.5, after=2)
add_para("我最想改变的一个管理习惯：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我识别了对影响最大的一座山，并写了 30 天具体行动",
    "我的三个新策略不是抄写的，是针对分析对象的具体方案",
    "我对自己打分是「如实」的，不是「打得好看的」",
])

# Part 1 知识框架
add_h3("✅ Part 1 知识框架")
part1_kf = """Part 1 认知力
│
├── 快·变·连·我（AI 时代升级）
│   ├── 快：即时响应成为新基线
│   ├── 变：适应周期从年变月
│   ├── 连：AI 成为新的连接对象
│   └── 我：价值感焦虑是新挑战
│
├── 五感驱动模型（AI 时代 4.0 版）
│   ├── 节奏感 → 建立响应约定
│   ├── 存在感 → 看见人类判断贡献
│   ├── 位置感 → 承认 AI 工具能力
│   ├── 掌控感 → 给工具选择权
│   └── 价值感 → 定期识别「只有你能做」
│
├── AI 时代三不怕
│   ├── 不怕比你懂 AI → 主动问、用他做内部 AI 资源
│   ├── 不怕说「AI 做的」 → 归因对话入口
│   └── 不怕反向带教 → 团队学习文化信号
│
└── 管理理念更新
    ├── 移除四座山（含 AI 时代版本）
    └── 三个新策略（软化冲突 / 共建规范 / 开放逆向学习）"""
add_code_block(part1_kf, size=9)

add_page_break()

# Part 1 行为承诺
add_signature_page(
    title="Part 1 行为承诺",
    commitment="从今天起，我承诺在我的 Z 世代下属面前——\n\n"
               "不再说「我当年不用 AI 也做到了」；\n"
               "而是承认他的 AI 时代困境，承认他的 AI 工具优势，承认「这个你可以教我」。",
)

add_page_break()


# ============================================================
# Part 2 适应力
# ============================================================

add_part_banner("Part 2")
add_part_subtitle("适应力——迎接 AI 时代的 Z 世代")

add_quote("Z 世代换工作的成本在 AI 时代更低了，留住他们的窗口期没有变长。", author="罗宏伟")

add_h3("🎯 Part 2 学习目标")
add_goals([
    "识别 Z 世代融入期四阶段的 AI 时代新困境和管理任务",
    "运用 1+3 任务清单 AI 升级版（含人机协作说明）为新员工设计第一个月",
    "掌握坦诚交流策略，用 AI 话题破冰建立信任",
    "通过角色扮演建立非正式交流的实践感",
])

add_h3("🗺️ Part 2 内容导航")
add_para("4 张工作页 + 4 个核心概念 + 3 道练习 + 1 张角色扮演记录表", size=10.5, color="555555", after=6)

add_page_break()


# ----- 工作页 5 -----
add_workpage_label(5, "融入期四阶段（AI 时代版）")

add_h3("知识点 5.1：四阶段路径图")

stage_diagram = """    未知              防卫              突破              定位
  Day 1-7          Week 2-3         Month 1-2       Month 3-转正
    │                │                │                │
    ▼                ▼                ▼                ▼
  信息真空          试探边界         开始说真话         建立身份
    +                +                +                +
  不知道AI怎么用    试探AI边界       AI依赖风险        AI时代价值定位
    │                │                │                │
    ▼                ▼                ▼                ▼
  做好准备          建立信任          感受成功          评估反馈
    +                +                +                +
  一天内说清AI态度  主动讨论AI期望   加入学习归因       人类贡献可见化"""
add_code_block(stage_diagram, size=8)

add_h3("📍 我的分析对象当前阶段诊断")
add_write_line("他目前在第 ______ 阶段", lines=1)
add_para("判断依据（具体行为）：", size=10.5, color="555555", after=2)
add_write_area(lines=2)
add_para("这个阶段最重要的一个管理任务：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_page_break()

add_h3("📖 AI 时代四阶段详细笔记")

stages = [
    ("未知阶段（Day 1-7）", "不知道 AI 工具在这里怎么用、管理者是否懂 AI、自己的 AI 习惯会不会被批评",
     "除了「做好准备」，还要在第一天就主动说清团队对 AI 的态度"),
    ("防卫阶段（Week 2-3）", "试探「用 AI 是否被允许」、观察管理者对 AI 的态度",
     "除了「建立信任」，还要主动讨论 AI 使用的边界和期望"),
    ("突破阶段（Month 1-2）", "被允许用 AI 后，产出大幅提升，但容易「AI 依赖」替代真正的能力发展",
     "除了「感受成功」，还要帮助他建立「人类贡献清晰」的自我认知"),
    ("定位阶段（Month 3-转正）", "开始建立「我在 AI 时代的职业价值」的初步认知",
     "在转正评估里，不只评估产出，还评估人类判断贡献的成长"),
]

stage_questions = [
    "我在 Day 1 会做的「AI 态度说明」话术：",
    "我会主动和他讨论的 AI 边界话题：",
    "我会用的「学习归因」开场问题：",
    "转正评估时，我会问的一个问题：",
]

for (title, challenge, task), q in zip(stages, stage_questions):
    p = doc.add_paragraph()
    add_run(p, "  " + title + "  ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p, before=8, after=2, line=1.3)
    set_paragraph_bg(p, COLOR_BRAND)
    p.paragraph_format.left_indent = Mm(2)

    add_para("新困境：", size=10, bold=True, color="C00000", after=2)
    add_para(challenge, size=10, color="333333", after=4)
    add_para("核心任务：", size=10, bold=True, color=COLOR_BRAND, after=2)
    add_para(task, size=10, color="333333", after=4)
    add_para(q, size=10, color="C00000", bold=True, after=2)
    add_write_area(lines=2)

add_page_break()

# 练习 5-A
add_h3("练习 5-A：四阶段路径图设计")
add_para("如果你的分析对象是新入职（未转正），请你为他画一张「接下来 3-6 个月的路径图」，标注每个阶段你打算做的 AI 时代关键动作：", size=10.5, after=4)

plan_stages = [
    "第 1 周 → ",
    "第 2-3 周 → ",
    "第 1-2 月 → ",
    "第 3 月 - 转正 → ",
]
for s in plan_stages:
    p = doc.add_paragraph()
    add_run(p, s, size=10.5, bold=True, color=COLOR_BRAND)
    add_run(p, "_" * 70, size=10.5, color="888888")
    set_paragraph_spacing(p, before=2, after=4, line=1.6)

add_selfcheck([
    "我识别了分析对象当前阶段，不是抽象写「我们处于 X 阶段」",
    "我对四个阶段都写了 AI 时代版本的具体动作",
    "我的路径图是 3-6 个月可落地的，不是长期愿景",
])

add_page_break()


# ----- 工作页 6 -----
add_workpage_label(6, "1+3 任务清单 AI 升级版")

add_h3("知识点 6.1：1+3 清单的 AI 时代升级")
add_para("原版 1+3 清单：1 个核心目标 + 3 个子任务", size=10.5, after=2)
add_para("AI 时代升级：每个子任务加一栏「人机协作说明」——", size=10.5, color=COLOR_BRAND, bold=True, after=4)

p = doc.add_paragraph()
add_run(p, "人机协作说明的三要素：", size=10.5, bold=True, color=COLOR_BRAND)
set_paragraph_spacing(p, before=2, after=2)
for s in [
    "1. 期望员工亲自判断的部分（不依赖 AI，必须人做）",
    "2. 可以 AI 辅助的部分（用 AI 提效，但判断在人）",
    "3. 人类贡献如何被看见（怎么验证员工确实做了人类判断）",
]:
    p = doc.add_paragraph()
    add_run(p, s, size=10.5)
    set_paragraph_spacing(p, before=1, after=1, line=1.3)
    p.paragraph_format.left_indent = Mm(4)

add_h3("为我的分析对象填写 1+3 清单")
add_write_line("任务周期", lines=1)
add_para("核心目标（1 个，这个月最重要的成果）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

for i in range(1, 4):
    add_para(f"子任务 {i}", bold=True, color=COLOR_BRAND, size=12, after=2)
    add_para("内容描述：", size=10, color="555555", after=2)
    add_write_area(lines=2)
    add_para("人机协作说明：", size=10.5, bold=True, color=COLOR_BRAND, after=2)
    add_para("· 期望员工亲自判断的部分：", size=10, color="555555", after=2)
    add_write_area(lines=2)
    add_para("· 可以 AI 辅助的部分：", size=10, color="555555", after=2)
    add_write_area(lines=2)
    add_para("· 人类贡献如何被看见（验证方式）：", size=10, color="555555", after=2)
    add_write_area(lines=2)

add_h3("💬 互评反馈")
add_write_line("配对伙伴", lines=1)
add_para("伙伴对「人类贡献如何被看见」的检验：", size=10, color="555555", after=2)
add_para("· 这个验证方式可操作吗？", size=10, color="555555", after=2)
add_write_area(lines=2)
add_para("· 员工能从这个验证方式知道什么算「做好了人类贡献」吗？", size=10, color="555555", after=2)
add_write_area(lines=2)
add_para("我修改后的版本（如有）：", size=10, bold=True, color=COLOR_BRAND, after=2)
add_write_area(lines=2)

# 练习 6-A
add_h3("练习 6-A：1+3 清单实战")
add_para("把这个清单打印出来，下周交给你的分析对象（或发给他），看他看完后的反应：", size=10.5, after=2)
add_para("他看完后问的第一个问题：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("这告诉我：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我的 1 个核心目标是这个月最关键的，不是季度或年度目标",
    "我对每个子任务都填了三要素（人做 / AI 辅助 / 人类贡献可见）",
    "我的「人类贡献如何被看见」是具体可验证的，不是「看他的工作态度」",
])

add_page_break()


# ----- 工作页 7 -----
add_workpage_label(7, "坦诚交流策略（AI 时代版）")

add_h3("知识点 7.1：坦诚交流四策略")

strats = [
    ("少命令多询问", "原版命令式：「你把这个方案改了。」",
     "AI 时代询问式：「你觉得这个方案哪部分最需要改？你打算怎么改？」"),
    ("有好奇擅回应", "对 Z 世代的非工作表达（情绪、困惑、AI 工具偏好），用好奇心接住而不是立刻给答案", ""),
    ("少说多听", "AI 时代升级版：Z 世代用 AI 工具的体验和困惑，管理者主动倾听比给建议更有效", ""),
    ("从铁纪律到爱的教育", "原版：硬性规定 AI 不能用",
     "AI 时代：明确边界，但给工具选择权；明确人类判断必须做的事，但工具自由"),
]
for k, v1, v2 in strats:
    p = doc.add_paragraph()
    add_run(p, "  " + k + "  ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
    set_paragraph_spacing(p, before=6, after=2, line=1.3)
    set_paragraph_bg(p, COLOR_BRAND)
    p.paragraph_format.left_indent = Mm(2)
    add_para(v1, size=10, color="333333", after=2)
    if v2:
        add_para(v2, size=10, color="333333", after=4)

add_h3("AI 话题破冰——我的开场话术")

add_para("话术 A（针对不怕比你懂 AI）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_quote_simple("「我最近发现 XX 工具好像挺有意思的，但你比我熟。你平时用 XX 工具做什么用得最多？」")
add_para("我自己的版本：", size=10, color="C00000", after=2)
add_write_area(lines=2)

add_para("话术 B（针对不怕说「这是 AI 做的」）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_quote_simple("「上次那个项目里，AI 帮了你什么？没有 AI 的话你最头疼的部分是什么？」")
add_para("我自己的版本：", size=10, color="C00000", after=2)
add_write_area(lines=2)

add_para("话术 C（针对不怕反向带教）", bold=True, color=COLOR_BRAND, size=10.5, after=2)
add_quote_simple("「我下周要给高管做汇报，你能不能花 10 分钟教我用 XX 工具做出来？我请你喝咖啡。」")
add_para("我自己的版本：", size=10, color="C00000", after=2)
add_write_area(lines=2)

add_para("我下周会用的开场（选一个或自己改）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("一周一次五分钟签到——我会问的 3 个问题")
for i in range(1, 4):
    add_para(f"问题 {i}：", size=10.5, bold=True, color=COLOR_BRAND, after=2)
    add_write_area(lines=2)

add_h3("AI 参与的正确回应——话术对比")
add_para("当员工说「这部分 AI 帮我做的」：", size=10.5, bold=True, after=2)

p = doc.add_paragraph()
add_run(p, "  ✗ 不说 ", size=11, bold=True, color="C00000", font=FONT_HEAD)
add_run(p, "（质问式）：", size=10, color="333333", bold=True)
add_run(p, "「那你做了什么？」（关闭对话）", size=10.5, color="333333")
set_paragraph_spacing(p, before=2, after=2, line=1.4)
set_paragraph_bg(p, "FFF3F3")
set_paragraph_borders(p, color="C00000", sz="4", sides=("left",))

p = doc.add_paragraph()
add_run(p, "  ✓ 改说 ", size=11, bold=True, color="2E7D32", font=FONT_HEAD)
add_run(p, "（好奇式）：", size=10, color="333333", bold=True)
add_run(p, "「AI 帮了你哪部分？你在哪个环节做了判断？那个判断你是怎么想的？」", size=10.5, color="333333")
set_paragraph_spacing(p, before=2, after=2, line=1.4)
set_paragraph_bg(p, "F1F8E9")
set_paragraph_borders(p, color="2E7D32", sz="4", sides=("left",))

add_para("我的版本（写一句真实的话术）：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我对四个策略都准备了具体话术",
    "我的开场话术不是抄原则，是「我下周会说」的真实话术",
    "我写的「应对员工说 AI 做的」是好奇语气，不是质问",
])

add_page_break()


# ----- 工作页 8 -----
add_workpage_label(8, "角色扮演记录表（非正式交流）")

add_h3("练习 8-A：融入期非正式交流角色扮演")

p = doc.add_paragraph()
add_run(p, "  第一轮 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
add_run(p, "（我扮演管理者，搭档扮演新员工）", size=10, color="FFFFFF")
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)

add_para("情境设定（写下来）：", size=10.5, color="555555", after=2)
add_write_area(lines=2)
add_para("管理者用的开场第一句话：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("过程中最让我意外的是：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第二轮 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
add_run(p, "（互换角色，我扮演新员工）", size=10, color="FFFFFF")
set_paragraph_spacing(p, before=8, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)

add_para("让我「愿意多说一点」的瞬间，管理者做了：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("让我「立刻收住」的瞬间，管理者做了：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_para("角色扮演后的关键洞察：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=3)

add_h3("📅 我的下周非正式交流计划")
add_write_line("计划时间", lines=1)
add_write_line("场合", lines=1)
add_write_line("对象", lines=1)
add_para("开场第一句话（写出来，说出来）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("我希望他带走的感受：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我做了两轮角色扮演，不是只听别人做",
    "我记录了「让我愿意多说」的瞬间，这是真实的可学技巧",
    "我为下周的非正式交流准备了具体时间、场合和开场",
])

add_page_break()

# Part 2 知识框架 + 行为承诺
add_h3("✅ Part 2 知识框架")
part2_kf = """Part 2 适应力
│
├── 融入期四阶段（AI 时代版）
│   ├── 未知（Day 1-7）→ 第一天说清 AI 态度
│   ├── 防卫（Week 2-3）→ 主动讨论 AI 期望
│   ├── 突破（Month 1-2）→ 加入学习归因
│   └── 定位（Month 3-转正）→ 人类贡献可见化
│
├── 1+3 任务清单 AI 升级版
│   ├── 1 个核心目标
│   ├── 3 个子任务
│   └── 人机协作说明（人做 / AI 辅助 / 贡献可见）
│
├── 坦诚交流四策略
│   ├── 少命令多询问
│   ├── 有好奇擅回应
│   ├── 少说多听
│   └── 从铁纪律到爱的教育
│
└── AI 话题破冰三话术
    ├── 不怕比你懂 AI 的破冰
    ├── 不怕说"这是 AI 做的"的破冰
    └── 不怕反向带教的破冰"""
add_code_block(part2_kf, size=9)

add_signature_page(
    title="Part 2 行为承诺",
    commitment="从今天起，我承诺——\n\n"
               "不在新人入职的第一周装作 AI 不存在；\n"
               "而是在第一天就主动说清楚团队对 AI 工具的态度。",
)

add_page_break()


# ============================================================
# Part 3 链接力
# ============================================================

add_part_banner("Part 3")
add_part_subtitle("链接力——保鲜 AI 时代的 Z 世代")

add_quote("任务分配不是给他一个名字，是给他一个「被看见」的方式。", author="罗宏伟")

add_h3("🎯 Part 3 学习目标")
add_goals([
    "用三维分析框架（不能做 / 不愿做 / 不知道怎么和 AI 配合做）诊断员工「做不到」的真实原因",
    "运用 5W2H+H 任务分配框架，给任务加上「人机协作视角」",
    "掌握辅导对话五步流程（AI 时代版，含学习归因子步骤）",
    "识别辅导中的常见错误（含 AI 新错误）",
])

add_h3("🗺️ Part 3 内容导航")
add_para("3 张工作页 + 3 个核心概念 + 2 道练习 + 1 张错误速查表", size=10.5, color="555555", after=6)

add_page_break()


# ----- 工作页 9 -----
add_workpage_label(9, "三维分析 + 5W2H+H 任务分配")

add_h3("知识点 9.1：三维分析框架")
add_para("员工「做不到」有三种完全不同的原因，对应的管理动作完全不同：", size=10.5, after=4)

dim_rows = [
    ["不能做", "技能缺口", "AI 工具能力缺口 / 能力空心化（用 AI 绕过了真正的学习）", "培训、辅导、补技能"],
    ["不愿做", "动力缺口", "价值感困惑 / 存在感冲突", "谈话、激发内驱力"],
    ["不知道怎么和 AI 配合做", "角色期望模糊", "人机分工不清晰", "明确人类判断边界"],
]
add_table(["维度", "含义", "AI 时代版本", "应对方向"], dim_rows,
          col_widths=[Mm(40), Mm(30), Mm(50), Mm(50)])

add_h3("📍 我的分析对象诊断")
add_para("针对一个他最近「做不好」或「做得慢」的任务：", size=10.5, color="555555", after=4)
add_write_line("任务名称", lines=1)
add_para("他做不好的最可能原因是（在以下勾选一个）：", size=10.5, bold=True, color="C00000", after=2)

for label in ["不能做（技能 / AI 工具能力不足）", "不愿做（动力 / 价值感问题）", "不知道怎么和 AI 配合做（人机分工不清）"]:
    p = doc.add_paragraph()
    add_run(p, "☐ ", size=11, color=COLOR_BRAND, bold=True)
    add_run(p, label, size=10.5, bold=True)
    set_paragraph_spacing(p, before=2, after=2, line=1.3)
    p.paragraph_format.left_indent = Mm(4)
    add_para("具体表现：", size=10, color="555555", after=2)
    add_write_area(lines=2)

add_para("我打算用的应对动作（针对选中的原因）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_page_break()

# 知识点 9.2
add_h3("知识点 9.2：5W2H+H 任务分配框架")
add_para("原版 5W2H：What / Why / Who / When / Where / How / How much", size=10.5, after=2)
add_para("AI 时代新增第 8 维：Human（人机协作视角）", size=10.5, bold=True, color="C00000", after=4)

wwwh_rows = [
    ["What", "做什么", "具体清晰"],
    ["Why", "为什么做（意义和价值）", "Z 世代最需要这一维"],
    ["Who", "谁来做，和谁协作", ""],
    ["When", "什么时候开始，什么时候完成", ""],
    ["Where", "在哪里完成", ""],
    ["How", "怎么做（方法和步骤）", ""],
    ["How much", "达到什么标准", ""],
    ["Human", "人机协作视角", "期望人做的 / 可 AI 辅助的 / 贡献如何被看见"],
]
add_table(["维度", "问题", "Z 世代管理要点"], wwwh_rows, col_widths=[Mm(30), Mm(75), Mm(65)])

add_h3("📝 5W2H+H 写作练习")
add_write_line("任务名称", lines=1)
for dim in ["What（做什么——具体清晰）",
            "Why（为什么这件事有价值——对谁有价值，满足员工的意义感需求）",
            "Who（谁来做，和谁协作）",
            "When（何时开始，何时完成）",
            "Where（在哪里完成）",
            "How（怎么做——方法和步骤）",
            "How much（达到什么标准）"]:
    add_para(dim, size=10.5, bold=True, color=COLOR_BRAND, after=2)
    add_write_area(lines=2)

add_para("Human（重点——人机协作视角）", size=11, bold=True, color="C00000", after=2)
add_para("· 期望员工亲自判断的部分（不依赖 AI）：", size=10, color="555555", after=2)
add_write_area(lines=2)
add_para("· 可以 AI 辅助提效的部分（提效但判断在人）：", size=10, color="555555", after=2)
add_write_area(lines=2)
add_para("· 员工的人类贡献如何被看见（验证方式）：", size=10, color="555555", after=2)
add_write_area(lines=2)

add_page_break()

# 互评反馈
add_h3("💬 互评反馈")
add_write_line("配对伙伴", lines=1)
add_para("请伙伴回答两个问题：", size=10, color="555555", after=2)
add_para("1. 员工看了 Why 这一维，能知道这件事对谁有价值吗？", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("2. 员工看了 Human 这一维，能知道什么算完成了人类贡献吗？", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("我修改的地方：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

# 练习 9-A
add_h3("练习 9-A：错误速查")
add_para("读以下 3 条任务分配说明，识别每条的「5W2H+H 缺失」和「AI 时代错误」：", size=10.5, after=4)

tasks = [
    "「小李，你下周把这个客户回访做了。」",
    "「小李，这是这份报告的初稿。你用 AI 改一改，下周一交给我。」",
    "「小李，这个产品调研你来负责，下周三给我一版。报告至少 3000 字，要有自己的分析。」",
]
for i, t in enumerate(tasks):
    p = doc.add_paragraph()
    add_run(p, f"任务 {i+1}", size=10.5, bold=True, color="FFFFFF", font=FONT_HEAD)
    add_run(p, " " + t, size=10.5, color="333333")
    set_paragraph_spacing(p, before=6, after=2, line=1.3)
    set_paragraph_bg(p, "FFF8E1")
    p.paragraph_format.left_indent = Mm(2)
    add_para("5W2H+H 缺失维度：", size=10, bold=True, color="C00000", after=2)
    add_write_area(lines=2)
    add_para("AI 时代错误：", size=10, bold=True, color="C00000", after=2)
    add_write_area(lines=2)

# 自检清单
add_h3("✅ 5W2H+H 自检清单")
checklist9 = [
    "What 清晰具体，不模糊",
    "Why 说明了「对谁有价值」，不只是任务本身",
    "Who 角色和协作方式清楚",
    "When 起止时间明确",
    "Where 完成地点 / 线上 / 线下明确",
    "How 方法步骤有指引，不只说「做好了」",
    "How much 标准可验证（不空洞）",
    "Human 期望人做的明确",
    "Human 可 AI 辅助的明确",
    "Human 人类贡献如何被看见有验证方式",
]
for it in checklist9:
    p = doc.add_paragraph()
    add_run(p, "☐  ", size=11, color=COLOR_BRAND, bold=True)
    add_run(p, it, size=10)
    set_paragraph_spacing(p, before=1, after=1, line=1.3)
    p.paragraph_format.left_indent = Mm(2)

add_selfcheck([
    "我对分析对象的「做不好」做了三维分析，不是笼统判断",
    "我的 5W2H+H 8 维都填了，没有留空",
    "Human 这一维的三要素（人做 / AI 辅助 / 贡献可见）都填了具体内容",
    "我请配对伙伴做了互评",
])

add_page_break()


# ----- 工作页 10 -----
add_workpage_label(10, "辅导对话五步流程（AI 时代版）")

add_h3("知识点 10.1：五步流程图")
five_step = """  开启对话 → 澄清事实 → 交换看法 → 达成共识 → 总结对话
              │
              └── AI 时代新增子步骤：学习归因
                  （「你的判断体现在哪里？
                    你之前不会、现在会的部分是？」）"""
add_code_block(five_step, size=9)

add_h3("🎤 五步详细话术设计")

p = doc.add_paragraph()
add_run(p, "  第一步：开启对话 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_quote_simple("基础话术：「我想和你聊聊 [任务]，不是批评，是想一起看看怎么做得更好。你现在方便吗？」")
add_para("AI 时代版本：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第二步：澄清事实（含学习归因子步骤） ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_quote_simple("基础澄清话术：「你做这件事的过程是什么样的？哪部分进展顺利？哪部分有挑战？」")
add_para("AI 时代学习归因问题（好奇语气）：", size=10, bold=True, color="C00000", after=2)
add_quote_simple("「在这个过程里，你的判断体现在哪里？有哪部分是你之前不会、现在会了的？」")
add_para("我的版本（写出来）：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("为什么是好奇而不是质问：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第三步：交换看法 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_para("先听员工，再说管理者——「你怎么看？/ 你觉得这件事的关键是什么？」", size=10, color="555555", italic=True, after=2)
add_para("我的起手问题：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第四步：达成共识 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_para("共建而非宣布——「那我们接下来可以怎么做？」", size=10, color="555555", italic=True, after=2)
add_para("我的话术：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第五步：总结对话 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_para("明确双方行动——「那我接下来会做 X，你接下来会做 Y，下次我们 [时间] 看进展。」", size=10, color="555555", italic=True, after=2)
add_para("我的话术：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("🎭 角色扮演记录（活动 9）")

p = doc.add_paragraph()
add_run(p, "  第一轮 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
add_run(p, "（经典辅导，我扮演 ______ 角色）", size=10, color="FFFFFF")
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_para("最难做到的是：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  第二轮 ", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
add_run(p, "（AI 时代辅导，我用的学习归因问题是）", size=10, color="FFFFFF")
set_paragraph_spacing(p, before=6, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

add_para("员工角色的感受差异（好奇语气 vs 质问语气）：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("最大的收获：", size=10, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我对五步都准备了真实话术，不是原则表述",
    "我完成了 2 轮角色扮演，记录了真实的感受差异",
    "我的学习归因问题是「我能真正问出来」的，不是照抄",
])

add_page_break()


# ----- 工作页 11 -----
add_workpage_label(11, "常见辅导错误速查（含 AI 新错误）")

add_h3("📋 错误速查表")

err_rows = [
    ["是否判断", "还没听完就开始判断", "「还没听完就先下了判断」"],
    ["显摆立场", "用自己的经历压制员工", "「我当年不用 AI 也做到了」"],
    ["暗示兜圈", "不直接说，绕弯子让员工自己悟", ""],
    ["急于求成", "想一次谈话解决所有问题", ""],
    ["AI 新错误：质问 AI 参与", "「那你做了什么？」（关闭对话）", "「AI 做的 = 你没做」的等式"],
    ["AI 新错误：AI 产出 ≠ 能力成长", "跳过学习归因，把 AI 产出误认为员工能力", "「上次做得不错」= 这次能独立做？"],
]
add_table(["错误", "描述", "AI 时代特殊版本"], err_rows, col_widths=[Mm(50), Mm(60), Mm(60)], first_col_bold=True)

add_blank_line(4)
add_para("针对我的情况：", bold=True, color=COLOR_BRAND, size=11, after=4)
add_para("我最可能犯的两个错误：", size=10.5, bold=True, color="C00000", after=2)
for i in range(1, 3):
    add_para(f"{i}. ", size=11, bold=True, color=COLOR_BRAND)
    add_write_area(lines=2)
add_para("我打算如何避免：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("✅ 辅导前自检清单")
add_para("每次辅导对话前对照一遍：", size=10, color="555555", after=2)
for it in [
    "我先听完员工的完整描述再判断",
    "我不用自己的经历压制员工",
    "我不绕弯子，直接说重点",
    "我不期待一次谈话解决所有问题",
    "我用好奇语气问 AI 参与情况，不用质问",
    "我做了学习归因，区分了 AI 产出和员工能力成长",
]:
    p = doc.add_paragraph()
    add_run(p, "☐  ", size=11, color=COLOR_BRAND, bold=True)
    add_run(p, it, size=10)
    set_paragraph_spacing(p, before=1, after=1, line=1.3)
    p.paragraph_format.left_indent = Mm(2)

add_selfcheck([
    "我对 6 类错误（4 类经典 + 2 类 AI 新）都能识别",
    "我识别了自己最可能犯的两个错误",
    "我设计的避免方案是「提醒机制」，不是「决心」",
])

add_page_break()

# Part 3 知识框架 + 行为承诺
add_h3("✅ Part 3 知识框架")
part3_kf = """Part 3 链接力
│
├── 三维分析（不能做 / 不愿做 / 不知道怎么和 AI 配合做）
│   ├── 不能做 → 技能培训 / AI 工具辅导
│   ├── 不愿做 → 谈话 / 激发内驱力
│   └── 不知道怎么和 AI 配合做 → 明确人机分工边界
│
├── 5W2H+H 任务分配
│   ├── 原版 7 维（What / Why / Who / When / Where / How / How much）
│   └── Human（第 8 维：人做 / AI 辅助 / 贡献可见）
│
├── 辅导对话五步流程（AI 时代版）
│   ├── 开启对话 → 澄清事实（+学习归因）→ 交换看法
│   └── 达成共识 → 总结对话
│
└── 错误速查（4 类经典 + 2 类 AI 新）
    ├── 经典：是否判断 / 显摆立场 / 暗示兜圈 / 急于求成
    └── AI 新：质问 AI 参与 / AI 产出≠能力成长"""
add_code_block(part3_kf, size=9)

add_signature_page(
    title="Part 3 行为承诺",
    commitment="从今天起，我承诺——\n\n"
               "不再说「你那个方案做得很 AI」；\n"
               "而是在每次辅导对话里问一句：「你的判断体现在哪里？」",
)

add_page_break()


# ============================================================
# Part 4 愿景力
# ============================================================

add_part_banner("Part 4")
add_part_subtitle("愿景力——引爆 AI 时代的小宇宙")

add_quote("Z 世代换工作的成本在 AI 时代更低了，但被点燃的 Z 世代也是前所未有的。", author="罗宏伟")

add_h3("🎯 Part 4 学习目标")
add_goals([
    "识别八大内驱动力在 AI 时代的交互变化",
    "为分析对象做内驱力画像，识别主驱动力和最缺激活的驱动力",
    "设计一份游戏设计画布 4.0 初稿（含防 AI 走捷径机制）",
])

add_h3("🗺️ Part 4 内容导航")
add_para("2 张工作页 + 2 个核心概念 + 1 张游戏画布 4.0", size=10.5, color="555555", after=6)

add_page_break()


# ----- 工作页 12 -----
add_workpage_label(12, "八大内驱动力画像（AI 时代交互分析）")

add_h3("知识点 12.1：八大内驱动力")
eight_diagram = """                史诗意义与使命感
                       |
        未知与好奇心 ——+—— 损失规避
                       |
        进步与成就感 ——+—— 稀缺与迫切感
                       |
        所有权与拥有感 ——+—— 社会影响力
                       |
                创意赋权与反馈"""
add_code_block(eight_diagram, size=10)

add_h3("📊 八大内驱动力画像（针对分析对象）")
add_para("针对我的分析对象，在 1-5 分之间打分（5=最强驱动），并写下 AI 时代的特殊影响：", size=10, color="555555", after=4)

eight_rows = [
    ["史诗意义与使命感", "在做比自己更大的事", "AI 时代价值感危机时受冲击最大"],
    ["进步与成就感", "感受到真实成长", "AI 代劳后难区分真实成长"],
    ["创意赋权与反馈", "用自己方式做，贡献被看见", "归属感模糊时缺失"],
    ["所有权与拥有感", "这件事是我的", "AI 参与后所有权边界模糊"],
    ["社会影响力", "对他人的影响被感知", "AI 向导角色可激活"],
    ["稀缺与迫切感", "限时限量的机会", "适度使用，避免压力疲劳"],
    ["未知与好奇心", "不确定性本身有吸引力", "AI 工具探索是天然的激活点"],
    ["损失规避", "不想失去已有的", "需配合正向驱动，避免纯惩罚"],
]
# 评分列加一列"评分（1-5）"
eight_rows_with_score = [[r[0], r[1], "___", r[2]] for r in eight_rows]
add_table(["内驱动力", "简述", "评分（1-5）", "AI 时代的特殊影响"],
          eight_rows_with_score,
          col_widths=[Mm(40), Mm(45), Mm(20), Mm(65)])

add_blank_line(4)
add_para("主驱动力（最高分 2-3 个）：", size=10.5, bold=True, color=COLOR_BRAND, after=2)
add_write_area(lines=2)
add_para("最缺乏激活的驱动力（最低分 1-2 个）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("AI 时代交互分析")
add_para("主驱动力在 AI 时代的强化或弱化：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("最缺驱动力在 AI 时代的特殊激活方式（如「用 AI 工具做团队内部分享」激活「创意赋权」，用「AI 探索项目」激活「未知与好奇心」）：", size=10, color="555555", after=2)
add_write_area(lines=2)

add_h3("💬 配对分享")
add_para("我现在的管理方式有没有照顾到主驱动力？", size=10.5, color="555555", after=2)
add_write_area(lines=1)
add_para("可以加的一个具体动作：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

# 练习 12-A
add_h3("练习 12-A：内驱力激活计划")
add_para("针对最缺激活的那个驱动力，我 30 天内要做的一个动作：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我对八大内驱动力都打了具体分数",
    "我识别了主驱动力和最缺驱动力",
    "我设计了具体的激活动作，不是原则表述",
])

add_page_break()


# ----- 工作页 13 -----
add_workpage_label(13, "游戏设计画布 4.0")

add_h3("知识点 13.1：四大系统 4.0 升级要点")

sys_rows = [
    ["目标系统 4.0", "产出目标 + 人类贡献目标分开设计", "人类贡献目标权重更高"],
    ["反馈系统 4.0", "AI 即时产出反馈 + 管理者人类贡献反馈", "人类贡献回顾需人工，不能自动化"],
    ["规则系统 4.0", "明确 AI 使用边界（不是禁止）", "判断类工作期望人类独立"],
    ["回报系统 4.0", "与内驱动力对接，含成长性回报", "人类判断力成长本身可作为回报"],
]
add_table(["系统", "4.0 核心升级", "AI 时代关键点"], sys_rows, col_widths=[Mm(35), Mm(70), Mm(65)])

add_h3("🛡️ 防 AI 走捷径机制")
add_para("我的团队最可能走捷径的场景（如「用 AI 批量完成报告」）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("我设计的应对机制：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h3("🎨 游戏设计画布 4.0 初稿摘要")
add_para("（A3 大图请另外保存，此处记录关键决策）", size=10, color="888888", italic=True, after=4)

add_write_line("业务场景", lines=1)

p = doc.add_paragraph()
add_run(p, "  目标系统 —— 产出目标", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  目标系统 —— 人类贡献目标", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  反馈系统 —— AI 即时反馈", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  反馈系统 —— 管理者人类贡献反馈（周期 + 内容）", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  规则系统 —— AI 使用边界（鼓励场景 / 要求人类判断的场景）", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  回报系统 —— 对应主驱动力的激励", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

p = doc.add_paragraph()
add_run(p, "  回报系统 —— 人类判断力成长的特殊回报", size=11, bold=True, color="FFFFFF", font=FONT_HEAD)
set_paragraph_spacing(p, before=4, after=2, line=1.3)
set_paragraph_bg(p, COLOR_BRAND)
p.paragraph_format.left_indent = Mm(2)
add_write_area(lines=2)

add_h3("💬 小组展示反馈")
add_para("伙伴发现的潜在走捷径点：", size=10.5, color="555555", after=2)
add_write_area(lines=2)
add_para("我的改进想法：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

# 练习 13-A
add_h3("练习 13-A：30 天最小可行性实验")
add_para("从画布里挑一个最小、最容易启动的实验，做 30 天：", size=10.5, after=2)
add_write_line("实验名称", lines=1)
add_para("第一个动作：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("30 天后用什么判断效果：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_selfcheck([
    "我对四大系统都填了具体内容，不是抽象设计",
    "我设计了防 AI 走捷径机制，不只是「鼓励 AI 辅助」",
    "我的回报系统对接了主驱动力，不是「千篇一律的奖金」",
])

add_page_break()

# Part 4 知识框架 + 行为承诺
add_h3("✅ Part 4 知识框架")
part4_kf = """Part 4 愿景力
│
├── 八大内驱动力（AI 时代交互）
│   ├── 史诗意义 / 进步成就 / 创意赋权 / 所有权
│   ├── 社会影响 / 稀缺迫切 / 未知好奇 / 损失规避
│   └── AI 时代特殊影响：价值感危机 / 真实成长难辨
│
└── 游戏设计画布 4.0
    ├── 目标系统：产出目标 + 人类贡献目标（分开设计）
    ├── 反馈系统：AI 即时反馈 + 管理者人类贡献反馈
    ├── 规则系统：AI 使用边界（鼓励 / 要求人类独立）
    └── 回报系统：主驱动力对接 + 人类判断力成长回报"""
add_code_block(part4_kf, size=9)

add_signature_page(
    title="Part 4 行为承诺",
    commitment="从今天起，我承诺——\n\n"
               "不再用 AI 批量产出掩盖员工的真实成长缺失；\n"
               "而是在游戏设计里加入「人类贡献目标」，让成长可见化。",
)

add_page_break()


# ============================================================
# 课程收尾
# ============================================================

add_title("课程收尾：四力整合 + 30 天管理改进清单", level=1, after=8)
add_quote("两天课程真正的开始，是从你回到工位那一刻。", author="罗宏伟")

add_h3("📊 四力整合回顾")

four_rows = [
    ["认知力", "5 感雷达图 + 三不怕转化行动卡 + 管理理念更新笔记", ""],
    ["适应力", "1+3 清单（带人机协作说明）+ 坦诚交流话术库 + 角色扮演反思", ""],
    ["链接力", "5W2H+H 完整任务说明 + 辅导对话话术卡 + 错误自检清单", ""],
    ["愿景力", "八大内驱动力画像 + 游戏设计画布 4.0 初稿", ""],
]
add_table(["力", "关键交付物", "我手上的产出"], four_rows, col_widths=[Mm(30), Mm(80), Mm(60)])

add_page_break()

# 表单 C.1
add_h3("【表单 C.1】五感诊断重测")
add_para("课程结束，重新填写这张表——和开始时的自评对比，看清变化：", size=10, color="555555", after=4)

resense_rows = [[s, "○    ○    ○    ○"] for s in senses]
add_table(["感", "几乎不 / 偶尔 / 经常 / 充足"], resense_rows, col_widths=[Mm(50), Mm(120)], first_col_bold=True)

add_blank_line(4)
add_write_line("课后「经常 / 充足」共 ___ 行（课前：___ 行）", lines=1)
add_para("哪一行变化最明显？为什么？", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=3)

# 表单 C.2
add_h3("【表单 C.2】管理理念重测")
mgmt2_rows = [[str(i+1) + ". " + it, "○    ○    ○    ○"] for i, it in enumerate(mgmt_items)]
add_table(["管理特征", "几乎不 / 偶尔 / 经常 / 总是"], mgmt2_rows, col_widths=[Mm(140), Mm(30)])

add_blank_line(4)
add_write_line("课后「经常 / 总是」共 ___ 行（课前：___ 行）", lines=1)
add_para("变化最大的一项，我的新行动是什么？", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=3)

add_page_break()


# 表单 C.3
add_h3("【表单 C.3】30 天管理改进清单（课程最重要的产出）")
add_quote("请认真用 25 分钟完成。这是你带回去最重要的成果。")

add_h4("区域一：关于我的分析对象（认知与适应）")
add_write_line("他最缺失的一感", lines=1)
add_para("针对这个感，我 30 天内要做的一件具体的事：", size=10.5, bold=True, color="C00000", after=2)

p = doc.add_paragraph()
add_run(p, "在（情境）", size=10.5, color="333333")
add_run(p, "_" * 60, size=10.5, color="888888")
set_paragraph_spacing(p, before=2, after=2, line=1.6)

p = doc.add_paragraph()
add_run(p, "我会（做什么不同的事）", size=10.5, color="333333")
add_run(p, "_" * 50, size=10.5, color="888888")
set_paragraph_spacing(p, before=2, after=2, line=1.6)

p = doc.add_paragraph()
add_run(p, "从（时间）", size=10.5, color="333333")
add_run(p, "_" * 30, size=10.5, color="888888")
add_run(p, "开始。", size=10.5, color="333333")
set_paragraph_spacing(p, before=2, after=4, line=1.6)

add_para("这件事做到了的标志是：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h4("区域二：关于任务分配（链接力——教）")
add_write_line("即将分配的下一个重要任务", lines=1)
add_para("用 5W2H+H 重新说明时，Human 这一维：", size=10.5, bold=True, color="C00000", after=2)
add_para("· 期望员工亲自判断的部分：", size=10, color="555555", after=2)
add_write_area(lines=2)
add_para("· 人类贡献如何被看见：", size=10, color="555555", after=2)
add_write_area(lines=2)
add_write_line("计划分配时间", lines=1)

add_h4("区域三：关于激励机制（愿景力——机）")
add_para("我从游戏设计画布提出的一个最小可行性实验：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("这个实验的第一步：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_write_line("计划开始时间", lines=1)
add_para("30 天后，可以用什么来判断实验有没有效果：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h4("区域四：下一次与分析对象的关键对话")
add_para("对话类型（勾选一个）：", size=10.5, color="555555", after=2)
add_checklist([
    "非正式交流（AI 话题破冰）",
    "辅导对话（含学习归因）",
    "发展面谈（双轨方向）",
    "AI 时代价值感谈话（帮他找到只有他才能做的那部分）",
])
add_write_line("计划时间", lines=1)
add_para("我的开场第一句话：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)
add_para("我希望他从这次对话带走的感受：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_page_break()


# 我的问责伙伴
add_h3("我的问责伙伴")
add_write_line("姓名", lines=1)
add_write_line("联系方式", lines=1)
add_para("30 天后的约定：", size=10.5, bold=True, color="C00000", after=2)
add_para("他问我：「你的 30 天清单里，做了什么，发现了什么？」", size=10.5, color="333333", after=2)
add_para("我问他：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

add_h4("三个 10 天阶段目标")
stage_rows = [
    ["第 1-10 天", "建立意识：每次用工具前停 30 秒", "", ""],
    ["第 11-20 天", "建立技能：完成一份 5W2H+H 任务分配", "", ""],
    ["第 21-30 天", "建立系统：跑通第一个最小激励实验", "", ""],
]
add_table(["阶段", "目标", "我要做的一件事", "怎么知道做到了"], stage_rows, col_widths=[Mm(25), Mm(60), Mm(45), Mm(40)])

add_blank_line(4)
add_para("30 天后可以检验的一个具体指标（不是「用得更好」，而是具体变化）：", size=10.5, bold=True, color="C00000", after=2)
add_write_area(lines=2)

# 30 天日历版
add_h4("30 天清单日历版")
cal_rows = [
    ["第 1 周（建立意识）", "", "", "", "", "", ""],
    ["第 2 周（建立意识）", "", "", "", "", "", ""],
    ["第 3 周（建立技能）", "", "", "", "", "", ""],
    ["第 4 周（建立系统）", "", "", "", "", "", ""],
]
add_table(["周次", "周一", "周二", "周三", "周四", "周五", "周末"], cal_rows, col_widths=[Mm(35)] + [Mm(20)]*6, first_col_bold=True)
add_para("使用说明：把每天要做的「小动作」填进去。比如第 1 周周三可以填「用学习归因开场做一次辅导对话」，第 3 周周一可以填「分配第一个 5W2H+H 任务」。", size=9, color="888888", italic=True, after=8)

# 海报版
add_h3("📌 30 天清单海报版（可贴在办公位）")
poster = """┌─────────────────────────────────────────┐
│         30 天 管理改进承诺              │
├─────────────────────────────────────────┤
│                                         │
│  最缺的一感：______ 感                  │
│  我的具体行动：________________         │
│                                         │
│  即将分配的任务（用 5W2H+H）：          │
│  ________________________________       │
│                                         │
│  最小激励实验：                          │
│  ________________________________       │
│                                         │
│  下次关键对话（时间）：                  │
│  ________________________________       │
│                                         │
│  问责伙伴：________                      │
│  下次复盘时间：________                  │
│                                         │
│  我的签名：________  日期：________      │
│                                         │
└─────────────────────────────────────────┘"""
add_code_block(poster, size=9)

add_page_break()


# 我的承诺签名
add_h3("✍️ 我的承诺签名")

p = doc.add_paragraph()
add_run(p, "我承诺在接下来的 30 天里，按照上面四个区域的具体行动推进。", size=11, color="333333")
p.add_run().add_break()
add_run(p, "我会让我的问责伙伴每 2 周问一次：「你做了什么，发现了什么？」", size=11, color="333333")
set_paragraph_spacing(p, before=8, after=12, line=1.5)
p.paragraph_format.left_indent = Mm(4)
p.paragraph_format.right_indent = Mm(4)
set_paragraph_bg(p, "F4F7FB")
set_paragraph_borders(p, color=COLOR_BRAND, sz="6")

# 签名 + 日期 + 见证人
p = doc.add_paragraph()
add_run(p, "签名：", size=12, bold=True, color=COLOR_BRAND)
add_run(p, "_" * 22, size=12, color="888888")
set_paragraph_spacing(p, before=14, after=4, line=1.4)

p = doc.add_paragraph()
add_run(p, "日期：", size=12, bold=True, color=COLOR_BRAND)
add_run(p, "_" * 18, size=12, color="888888")
set_paragraph_spacing(p, before=4, after=10, line=1.4)

p = doc.add_paragraph()
add_run(p, "见证人（问责伙伴）：", size=12, bold=True, color=COLOR_BRAND)
add_run(p, "_" * 22, size=12, color="888888")
set_paragraph_spacing(p, before=4, after=4, line=1.4)

p = doc.add_paragraph()
add_run(p, "见证日期：", size=12, bold=True, color=COLOR_BRAND)
add_run(p, "_" * 18, size=12, color="888888")
set_paragraph_spacing(p, before=4, after=12, line=1.4)

add_page_break()


# 致出发的你
add_title("致出发的你", level=1, after=12)

add_para("你用两天时间，走完了 AI 时代的 Z 世代管理四力模型。", size=11.5, after=4)
add_para("这不是一套理论，而是一套", size=11.5, after=4)
# 上面这句错了，重新写
last = doc.paragraphs[-1]
last.runs[-1].text = "这不是一套理论，而是一套"

# 用一个完整的段落
doc.paragraphs[-1].clear()
add_run(doc.paragraphs[-1], "这不是一套理论，而是一套", size=11.5)
add_run(doc.paragraphs[-1], "可以从明天起就开始用的工作方式", size=11.5, bold=True, color=COLOR_BRAND)
add_run(doc.paragraphs[-1], "。", size=11.5)
set_paragraph_spacing(doc.paragraphs[-1], before=4, after=4, line=1.6)

add_para("你在课堂上完成的每一张工作页、每一道练习、每一个真实场景的解决方案，都是真实的成果，不是作业。", size=11.5, after=6)

add_para("四力有一个共同的底层逻辑——", size=11.5, after=4)
add_quote("AI 提供了前所未有的效率工具，但工具不创造价值，用工具的人创造价值。")

add_para("认知力让你", size=11.5, after=4)
add_para("  看见  ", size=12, bold=True, color=COLOR_BRAND)
add_run(doc.paragraphs[-1], "这代人；适应力让你", size=11.5)
add_run(doc.paragraphs[-1], "  留住  ", size=12, bold=True, color=COLOR_BRAND)
add_run(doc.paragraphs[-1], "这代人；链接力让你", size=11.5)
add_run(doc.paragraphs[-1], "  带好  ", size=12, bold=True, color=COLOR_BRAND)
add_run(doc.paragraphs[-1], "这代人；愿景力让你", size=11.5)
add_run(doc.paragraphs[-1], "  点燃  ", size=12, bold=True, color=COLOR_BRAND)
add_run(doc.paragraphs[-1], "这代人。", size=11.5)
set_paragraph_spacing(doc.paragraphs[-1], before=4, after=6, line=1.6)

add_para("Z 世代换工作的成本在 AI 时代更低了，但他们被点燃后的爆发力也是前所未有的。", size=11.5, after=6)

add_quote("讨喜而不是讨好。", author=None)
add_para("Z 世代管理者不需要「讨好」这代人，但需要「讨喜」——让他们感到被看见、被尊重、被期待。", size=11, color="555555", after=4)

add_quote("领先半步，吃尽红利。", author=None)
add_para("AI 时代的管理者，比的不是懂 AI 比员工多，是比员工先看到管理要升级的地方。", size=11, color="555555", after=6)

add_para("你今天开始建立的这套管理动作，在五年后，无论 AI 变成什么样子，都还有用。", size=11.5, bold=True, color=COLOR_BRAND, align="center", after=4)
add_para("从今天，从现在，从这一次开始。", size=11.5, bold=True, color=COLOR_BRAND, align="center", after=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p, "—— 罗宏伟", size=12, color="666666", italic=True)
set_paragraph_spacing(p, before=8, after=12, line=1.4)
p.paragraph_format.right_indent = Mm(8)

add_page_break()


# ============================================================
# 附录一：术语速查表
# ============================================================

add_title("附录一：术语速查表", level=1, after=10)

terms = [
    ("五感驱动模型", "Z 世代员工「冰山下」的五个内驱动力：节奏感、存在感、位置感、掌控感、价值感"),
    ("价值感", "AI 时代新增的第五感：员工对「自己有 AI 替代不了的贡献」的确认需求"),
    ("三不怕", "AI 时代 Z 世代的三个新特征：不怕比你懂 AI、不怕说「这是 AI 做的」、不怕反向带教"),
    ("融入期四阶段", "Z 世代新员工从入职到转正经历的四个阶段：未知、防卫、突破、定位"),
    ("1+3 任务清单", "1 个核心目标 + 3 个子任务的融入期结构化任务管理工具"),
    ("人机协作说明", "1+3 清单 AI 升级版中增加的栏位：期望人做的、可 AI 辅助的、人类贡献如何被看见"),
    ("5W2H+H", "任务分配框架，原版 5W2H + Human（人机协作视角）共 8 维"),
    ("学习归因", "辅导对话中区分「AI 做到的」和「员工学到的」的子步骤"),
    ("八大内驱动力", "游戏化管理的心理学基础：史诗意义、进步成就、创意赋权、所有权、社会影响、稀缺迫切、未知好奇、损失规避"),
    ("游戏设计画布 4.0", "含 AI 增强维度和防 AI 走捷径机制的激励系统设计工具"),
    ("三维分析", "员工「做不到」的三种原因诊断：不能做、不愿做、不知道怎么和 AI 配合做"),
    ("问责伙伴", "课后 30 天里每 2 周和你做 15 分钟复盘对话的同伴"),
]
term_rows = [[k, v] for k, v in terms]
add_table(["术语", "定义"], term_rows, col_widths=[Mm(45), Mm(125)], first_col_bold=True)

add_page_break()


# ============================================================
# 附录二：N 个工具速查索引
# ============================================================

add_title("附录二：N 个工具速查索引", level=1, after=10)
add_para("详细版见独立文件 04_学员手册_配套工具索引.md", size=10, color="888888", italic=True, after=8)

tools = [
    ("0.1", "五感诊断前测表", "引言", "我的分析对象五感评分"),
    ("0.2", "管理理念前测表", "引言", "我的管理起点自评"),
    ("0.3", "课前任务确认表", "引言", "5 项课前任务"),
    ("1", "快·变·连·我 AI 时代版本", "Part 1", "四个时代基因 AI 强化表现"),
    ("2", "五感诊断与雷达图", "Part 1", "五感详细诊断 + 手绘雷达图"),
    ("3", "三不怕转化行动卡", "Part 1", "三个「不怕」的转化行动"),
    ("4", "管理理念自测与四座山移除", "Part 1", "6 项自评 + 移除四座山"),
    ("5", "融入期四阶段路径图", "Part 2", "分析对象当前阶段诊断"),
    ("6", "1+3 任务清单 AI 升级版", "Part 2", "三子任务 + 人机协作说明"),
    ("7", "坦诚交流策略与 AI 话题破冰", "Part 2", "三话术 + 5 分钟签到问题"),
    ("8", "角色扮演记录表", "Part 2", "非正式交流两轮记录"),
    ("9", "三维分析 + 5W2H+H", "Part 3", "8 维任务分配"),
    ("10", "辅导对话五步流程（含学习归因）", "Part 3", "五步话术设计"),
    ("11", "常见辅导错误速查", "Part 3", "4 类经典 + 2 类 AI 新错误"),
    ("12", "八大内驱动力画像", "Part 4", "八驱动力打分 + AI 时代交互"),
    ("13", "游戏设计画布 4.0", "Part 4", "四系统 4.0 + 防 AI 走捷径机制"),
    ("C.1", "五感诊断重测", "收尾", "与前测对比"),
    ("C.2", "管理理念重测", "收尾", "与前测对比"),
    ("C.3", "30 天管理改进清单", "收尾", "4 个区域 + 日历版 + 海报版"),
    ("—", "行为承诺签名页", "Part 1-4 + 收尾", "5 个签名页"),
]
add_table(["编号", "工具名称", "位置", "核心问题"], tools, col_widths=[Mm(15), Mm(70), Mm(30), Mm(55)])

add_page_break()


# ============================================================
# 文档结尾：完整署名块
# ============================================================

for _ in range(6):
    add_blank_line(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "─" * 30, size=10, color=COLOR_BRAND)
set_paragraph_spacing(p, before=0, after=12, line=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "领  航  ·  4.0", size=18, bold=True, color=COLOR_BRAND, font=FONT_HEAD)
set_paragraph_spacing(p, before=0, after=4, line=1.3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "—— AI 时代的 Z 世代管理新策略 ——", size=11, color="333333", font=FONT_HEAD)
set_paragraph_spacing(p, before=0, after=14, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "课程开发者：罗宏伟", size=11, color="333333", bold=True)
set_paragraph_spacing(p, before=0, after=4, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "学员手册版本：v1.0", size=11, color="333333")
set_paragraph_spacing(p, before=0, after=4, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "印刷版生成时间：2026-06", size=11, color="333333")
set_paragraph_spacing(p, before=0, after=14, line=1.4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "本手册仅供本课程学员使用 · 未经授权不得复制传播", size=9, color="888888", italic=True)
set_paragraph_spacing(p, before=8, after=4, line=1.3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "© 罗宏伟 · 领航·4.0 项目组", size=9, color="888888")
set_paragraph_spacing(p, before=0, after=0, line=1.3)


# 保存
OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_DOCX)

# 统计
import os
size_kb = os.path.getsize(OUT_DOCX) / 1024
print(f"OK -> {OUT_DOCX}")
print(f"Size: {size_kb:.1f} KB")

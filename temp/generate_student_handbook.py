#!/usr/bin/env python3
"""生成学员手册Word文档 - 看不清现实找到能动的缝隙"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\新课开发\行动学习2026\02-对事-教程\完整课程包\学员手册\看清现实找到能动的缝隙_学员手册_v1.0.docx"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    """添加标题"""
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_quote(doc, text):
    """添加引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
    return p

def add_form_table(doc, headers, rows, header_color="1F3864"):
    """添加表单表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # 数据行
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            for p in row.cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    return table

def add_exercise_block(doc, title, items, time_min=0):
    """添加练习块"""
    p = doc.add_paragraph()
    run = p.add_run(f"✋ {title}")
    run.bold = True
    run.font.size = Pt(11)

    if time_min > 0:
        run2 = p.add_run(f"（{time_min}分钟）")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0xE5, 0x51, 0x00)

    for item in items:
        if isinstance(item, str):
            doc.add_paragraph(item, style='List Bullet')
        elif isinstance(item, tuple):
            # (question, rows)
            q, rows = item
            doc.add_paragraph(q)
            if rows:
                add_form_table(doc, ["填写内容"], [[""]] * rows)
        elif isinstance(item, dict):
            # 自由格式
            for k, v in item.items():
                p2 = doc.add_paragraph()
                p2.add_run(f"{k}：").bold = True
                p2.add_run(v)

def add_knowledge_box(doc, title, items, box_color="E8F5E9"):
    """添加知识框架框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"📋 {title}")
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run("\n".join(f"• {i}" for i in items))
    run2.font.size = Pt(10)

    return p2

def create_handbook():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4横向
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("看清现实，找到能动的缝隙")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("信息分析与突破口识别 · 学员手册")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)

    doc.add_paragraph()

    # 封面信息表
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("学员姓名：", "_______________", "所在部门：", "_______________"),
        ("课程日期：", "_______________", "课程讲师：", "_______________"),
        ("学习顾问：", "_______________", "联系电话：", "_______________"),
        ("使用版本：", "v1.0", "手册编号：", "_______________"),
    ]
    for ri, row_data in enumerate(info_data):
        row = info_table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    if ci % 2 == 0:
                        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 封面引言
    quote_p = doc.add_paragraph()
    quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote_p.add_run('*"工具会过时，框架会复利。你今天建立的这套分析方法，\n在五年后依然有效。"*')
    quote_run.italic = True
    quote_run.font.size = Pt(12)
    quote_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== 使用说明 =====
    add_heading(doc, "如何使用这本手册", 1)

    p = doc.add_paragraph()
    p.add_run("这本手册不是讲义，不是笔记，不是课后读物。").bold = False
    p.add_run("\n它是你在课程中完成的工作台。每一个框架、每一张表单、每一道练习，都要在课堂上完成。你做完的每一页，都是你带走的真实成果，而不是别人告诉你的理论。")

    doc.add_paragraph()

    add_heading(doc, "三个使用原则", 2)

    principles = [
        ("原则一：带着真实任务来", "手册里所有的练习都要用你自己工作中的真实场景，不要虚构例子。真实的任务，才有真实的收获。"),
        ("原则二：写下来比记下来更有价值", "看懂了不等于会用，写下来才是真正内化的开始。每一道练习，都请认真完成，不要留空。"),
        ("原则三：这是起点，不是终点", "课程结束不是学习的终点。手册最后一章是你的行动计划——回到工作中，每完成一次，回来翻一翻这本手册。"),
    ]
    for title, content in principles:
        p = doc.add_paragraph()
        run = p.add_run(f"【{title}】")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
        p.add_run(f"\n{content}")

    doc.add_page_break()

    # ===== 课程全景图 =====
    add_heading(doc, "🗺️ 课程全景图", 1)

    doc.add_paragraph("本课程由六个核心部分组成，呈现完整的分析路径：")

    overview_table = doc.add_table(rows=7, cols=4)
    overview_table.style = 'Table Grid'
    overview_data = [
        ("部分", "核心问题", "学习方法", "核心产出"),
        ("第一部分", "为什么要换一种思维方式？", "案例分析+自我诊断", "对分析偏误的清醒认识"),
        ("第二部分", "怎么完整扫描所有相关因素？", "维度框架练习", "体检清单（完整版）"),
        ("第三部分", "怎么让分析落在真实信息上？", "调研设计实战", "调研计划"),
        ("第四部分", "怎么判断每个因素的轻重缓急？", "四维分析练习", "四维分析表"),
        ("第五部分", "怎么从一堆因素里找到真正的抓手？", "突破口深度验证", "2~4个突破口"),
        ("第六部分", "怎么把突破口变成可落地的行动？", "行动方案设计", "完整行动方案"),
    ]

    for ri, row_data in enumerate(overview_data):
        row = overview_table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
                        set_cell_shading(cell, "1F3864")
                        run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # 完整路径图
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("体检清单 → 调研设计 → 四维分析 → 突破口识别 → 行动方案")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("   ↓              ↓              ↓              ↓              ↓")
    run2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("画全貌        找真相        做判断        找杠杆        能落地")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)

    doc.add_page_break()

    # ===== 第一部分：体检思维 =====
    add_heading(doc, "第一部分　体检思维", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals = [
        "识别三种常见的认知偏误及其对分析的影响",
        "理解"体检思维"与"找原因思维"的本质区别",
        "建立"先画完整因素地图，再决定发力方向"的思维习惯"
    ]
    for g in goals:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "一个让人沮丧的故事", 2)

    story = """某城市轨道交通公司，2022年启动了一个运营改善课题：提升三北线高峰期运输能力，目标是把高峰运能利用率提升8个百分点。

团队做了很多事：增配高峰站务人员、加密早高峰发车间隔、在客流最密集的两个站点做了流量引导优化、开展了站务操作技能培训。

半年后复盘：高峰期运能利用率提升了1.2个百分点。目标是8个百分点。

一年后，一位新成员发现：所有措施瞄准的全都是同一个维度——人员配置和运营调度。而真正制约运力的核心瓶颈，是关键中转站的折返线路能力不足——无论增加多少人手、优化多少调度，都无法解决这个物理限制。

这个折返线路问题不是没人知道——只是"大家都知道，但没人把它当成这次课题需要认真研究的对象"。"""

    doc.add_paragraph(story)

    add_heading(doc, "三个几乎人人都有的思维偏误", 2)

    add_heading(doc, "偏误一：可及性偏误（最显眼的 ≠ 最重要的）", 3)
    doc.add_paragraph("人最先想到的原因，往往是最近发生的、最显眼的、最让人烦恼的事情——但这些不一定是最重要的。上周刚发生的设备故障，就自动跑到讨论前排——哪怕这个故障只是偶发事件，不是系统性瓶颈。")
    doc.add_paragraph("这不是因为你不够仔细，而是大脑天然优先处理"最近发生的、感觉最强烈的"信息。最烦的问题，不等于最关键的原因。", style='Quote')

    add_heading(doc, "偏误二：归责偏误（找"谁的问题"≠ 找"系统性因素"）", 3)
    doc.add_paragraph("遇到困境，人天然倾向于找"是谁的问题"，而不是"是什么结构性、机制性因素导致了这个结果"。"某部门不配合""某些员工态度消极"——这类表述让人感觉找到了原因，但实际上只是把问题定位到了一个人头上，没有分析是什么流程、机制让这个人这样行动。")

    add_heading(doc, "偏误三：局部视角偏误（你看到的，只是你的那一角）", 3)
    doc.add_paragraph("每个人只能看到自己能接触到的那部分现实。同一个项目，一线操作员眼中的根源、中层管理者眼中的根源、高层决策者眼中的根源，可能完全不同——而且每个人都觉得自己看到的才是真实的。")

    p = doc.add_paragraph()
    run = p.add_run("⚠️ 这三个偏误不是能力问题，也不是态度问题，是人类认知的内置程序。意识到它们的存在，才有可能有意识地绕开它们。")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)

    add_heading(doc, "体检思维：换一个提问方式", 2)

    doc.add_paragraph("所有这些偏误，都源于同一个起点：我们太早就进入了"找原因"的模式。")

    p = doc.add_paragraph()
    run = p.add_run("找原因思维：")
    run.bold = True
    p.add_run(""我觉得问题在哪里？"——然后去找支持这个判断的证据。")

    p2 = doc.add_paragraph()
    run2 = p2.add_run("体检思维：")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    p2.add_run(""影响这个目标达成的，可能有哪些因素？"——然后系统地扫描所有可能相关的方向，再基于真实信息来判断。")

    add_heading(doc, "用医生的逻辑来理解", 3)
    doc.add_paragraph("想象一位有经验的医生在接诊一位主诉"持续头痛、长期疲惫"的患者。医生不会直接说"你压力太大了"，他会开一套全面检查：血常规、血压、甲状腺功能、肝功能……")
    doc.add_paragraph("不是因为医生认为这些指标一定都有问题。而是因为：在没有完整扫描之前，任何单一判断都可能是错的。头痛和疲惫，可能是贫血，可能是甲状腺功能减退，可能是睡眠呼吸障碍，也可能真的只是压力过大。这四种情况，治疗方向完全不同。", style='Quote')

    add_heading(doc, "体检思维的三个核心特征", 3)
    features = [
        ("先穷举，后判断", "在得出"这个最重要"之前，先尽可能完整地列出"所有可能相关的因素"。不遗漏，才有可能找到真正的杠杆。"),
        ("不评价，只列举", "在建立清单的阶段，不需要判断哪个因素好坏。这一步只做一件事：把所有需要去了解真实状态的因素都列出来。判断放到后面去做。"),
        ("以目标为锚", "不是扫描所有"有问题的地方"，而是扫描所有"可能影响这个具体目标达成"的因素。目标是清单的边界。"),
    ]
    for title, content in features:
        p = doc.add_paragraph()
        run = p.add_run(f"【{title}】")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)
        p.add_run(f"\n{content}")

    # 练习1
    doc.add_paragraph()
    add_heading(doc, "✋ 练习：看清你当前项目的分析起点", 2)
    doc.add_paragraph("带着你在导读里写下的那个真实项目，做这个两步练习。")

    doc.add_paragraph("第一步：快速列出你认为影响目标的因素（3分钟，快速列，不要想太多）")
    add_form_table(doc, ["#", "你认为影响目标达成的因素"], [
        [str(i), ""] for i in range(1, 9)
    ])

    doc.add_paragraph("第二步：给每条打偏误标签")
    doc.add_paragraph("回头看你刚才写的每一条，在旁边标注它是否带有以下特征：")

    tags = [
        "〔可及〕它进入清单，是因为它最近发生、或者最显眼、最烦人",
        "〔归责〕它的表述指向一个人或部门（而不是系统性因素）",
        "〔局部〕它只是你自己视角能看到的，没有考虑其他层级或岗位的视角"
    ]
    for t in tags:
        doc.add_paragraph(f"• {t}")

    doc.add_paragraph("打完标签后，看一看：你的清单里，有多大比例是这三类？有没有哪些维度，你觉得应该有但没出现在清单里？")

    add_heading(doc, "关键提炼", 2)
    summary_items = [
        "三个偏误：可及性偏误（最显眼的≠最重要的）、归责偏误（找人不如找系统）、局部视角偏误（你的视角只是地图的一角）",
        "体检思维的核心动作：先画完整的因素地图，再决定在哪里发力",
        "在因素地图画完整之前，任何关于"应该做什么"的判断都是暂时性的假设，需要被验证"
    ]
    for item in summary_items:
        doc.add_paragraph(f"• {item}")

    p = doc.add_paragraph()
    run = p.add_run("带走一件事：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    p.add_run("在这次课题里，我应该先系统地画完整因素地图，再决定发力方向。")

    doc.add_page_break()

    # ===== 第二部分：体检清单 =====
    add_heading(doc, "第二部分　体检清单", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals2 = [
        "理解体检清单的本质和格式要求",
        "掌握用维度框架系统覆盖所有相关因素的方法",
        "能够针对自己的真实项目建立一份完整的体检清单"
    ]
    for g in goals2:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "两个团队，同一个问题，三个月后天壤之别", 2)

    story2 = """某零售企业的两个门店，面临同一个挑战：顾客到店转化率持续下滑，目标是把转化率从18%提升到25%。

A店团队开了一次两小时的头脑风暴，列出了11条他们认为的影响因素，很快锁定了"导购服务态度"和"促销力度"两个方向。

B店团队花了半天时间，用维度框架系统梳理了29条因素，覆盖了A店完全没有考虑到的维度：试衣间数量与等待时间、门店动线设计、货品丰富度分层、导购提成结构对推荐行为的影响。

三个月后：A店转化率提升了0.8个百分点，B店提升了6.2个百分点。

差距在哪里？B店真正解决的核心问题是试衣间严重不足（高峰期排队超过15分钟）和货品丰富度在核心价格带的断层——这两个因素，在A店的清单里根本没有出现。"""
    doc.add_paragraph(story2)

    add_heading(doc, "体检清单是什么，不是什么", 2)

    p = doc.add_paragraph()
    run = p.add_run("体检清单是什么：")
    run.bold = True
    p.add_run("一份针对"影响项目目标达成"的所有相关因素的完整列举。结构化分类，不评价、不排序、只穷举。清单上的每一条，都代表着"这是一个我需要去了解真实状态的方向"。")

    p2 = doc.add_paragraph()
    run2 = p2.add_run("体检清单不是什么：")
    run2.bold = True
    items_not = [
        "不是"我认为有问题的地方的清单"——那是已经加了判断的清单",
        "不是"我们打算做的事情的清单"——那是行动清单",
        "不是"越长越好"——精简但覆盖全面的清单才是好的"
    ]
    for item in items_not:
        doc.add_paragraph(f"• {item}", style='List Bullet')

    add_heading(doc, "建立体检清单的三步逻辑", 2)

    add_heading(doc, "第一步：锚定目标", 3)
    doc.add_paragraph("清单服务于目标。在开始列因素之前，先把目标说清楚。不同目标，决定了完全不同的扫描方向。")
    doc.add_paragraph("写下你的项目目标（具体、可测量）：_______________________________")

    add_heading(doc, "第二步：用维度框架覆盖全貌", 3)
    doc.add_paragraph("围绕目标，系统列出所有可能影响它的因素类别。以下是通用参考维度（根据你的行业和项目特点做调整）：")

    dim_table = doc.add_table(rows=8, cols=2)
    dim_table.style = 'Table Grid'
    dims = [
        ("维度类别", "典型细项方向"),
        ("资源维度", "人力配置、设备状态、预算充足度、时间资源"),
        ("流程维度", "核心操作流程、跨部门协同流程、异常处理流程"),
        ("能力维度", "团队技能水平、知识掌握程度、工具使用能力"),
        ("系统与工具维度", "支撑系统稳定性、工具适配性、数据准确性"),
        ("管理机制维度", "考核指标设计、激励机制、信息传递与反馈机制"),
        ("外部条件维度", "政策法规约束、客户或用户行为模式、供应商因素"),
        ("历史遗留维度", "历史问题积累、以往的改善尝试及结果、路径依赖"),
    ]
    for ri, (d1, d2) in enumerate(dims):
        row = dim_table.rows[ri]
        row.cells[0].text = d1
        row.cells[1].text = d2
        for p_cell in row.cells[0].paragraphs:
            for run_cell in p_cell.runs:
                run_cell.font.size = Pt(10)
                if ri == 0:
                    run_cell.bold = True
                    set_cell_shading(row.cells[0], "1F3864")
                    run_cell.font.color.rgb = RGBColor(255, 255, 255)

    add_heading(doc, "第三步：逐项展开，列出具体条目", 3)
    doc.add_paragraph("在每个维度下，列出可观测、可评估的具体项目或因素。")
    doc.add_paragraph("格式要求：名词短语，不加评价。不是"培训质量很差"（这是评价），而是"关键岗位培训覆盖率"（这是因素）。")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("💡 判断一条是否具体的方法：")
    run.bold = True
    p.add_run("你能想象出"如何了解这条因素的真实状态"吗？如果你能说出"可以通过X方式来了解它的当前水平"，这条因素就足够具体。")

    # 练习2
    add_heading(doc, "✋ 练习：用框架重建你的体检清单（12分钟）", 2)
    doc.add_paragraph("用维度框架重新展开，为每个维度写出3~6条具体的细项因素：")

    add_form_table(doc, ["维度", "具体因素条目（名词短语，不加评价）"], [
        ["资源维度", ""],
        ["", ""],
        ["流程维度", ""],
        ["", ""],
        ["能力维度", ""],
        ["", ""],
        ["系统与工具维度", ""],
        ["", ""],
        ["管理机制维度", ""],
        ["", ""],
        ["外部条件维度", ""],
        ["", ""],
        ["历史遗留维度", "1. 以往针对此目标的改善尝试及其结果"],
        ["", "2. "],
    ])

    add_heading(doc, "常见误区和识别方法", 2)

    mistakes = [
        ("把评价写成了因素", "清单里出现"培训效果差""管理层不重视"", "改成名词短语，去掉评价词"),
        ("把解决方案写成了因素", "清单里出现"建立反馈机制""引入新系统"", "改为描述当前状态的问题域"),
        ("遗漏历史维度", "清单里没有任何关于"以前试过什么"的条目", "专门加一个"历史遗留"维度"),
        ("只有内部因素，没有外部因素", "清单只描述内部流程和人员", "专门检查"外部条件"维度"),
    ]

    add_form_table(doc, ["误区", "典型表现", "怎么修正"], mistakes)

    p = doc.add_paragraph()
    run = p.add_run("带走两件事：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    doc.add_paragraph("体检清单的本质：不是问题清单，是"需要去了解真实状态"的完整因素地图。")
    doc.add_paragraph("建清单的关键动作：用维度框架覆盖全貌（而不是靠感觉），再逐项展开成可调研的具体条目。")

    doc.add_page_break()

    # ===== 第三部分：调研设计 =====
    add_heading(doc, "第三部分　调研设计", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals3 = [
        "理解为什么不能跳过调研，理解"假设"和"真实信息"的本质区别",
        "掌握调研设计的四个核心问题及其回答方法",
        "能够针对自己的体检清单设计一套完整的调研方案"
    ]
    for g in goals3:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "调研设计的四个核心问题", 2)

    questions = [
        ("问题一：哪些信息我已经有，哪些需要重新收集？",
         "调研设计的第一步，是盘点库存。先问：现在已经有什么信息？通常，以下几类信息是相对容易获取的二手信息：内部数据库的历史数据、内部报告和分析文件、行业基准数据、历史决策文件。"),
        ("问题二：需要收集的信息，用什么方式收集？",
         "不同类型的信息，适合不同的调研方式：\n• 深度访谈：判断类、经验类、隐性知识、"为什么"类\n• 现场观察：流程类、行为类、物理空间类\n• 数据分析：量化类、趋势类、分布类\n• 文档研读：历史类、政策类、机制设计类"),
        ("问题三：找谁收集？",
         "同一个问题，要向不同层级的人了解。管理层看到的现实和一线操作者看到的现实，常常差异显著。特别重要的调研对象：历史上参与过改善的人。"),
        ("问题四：怎么确保信息的可靠性？",
         "三个可靠性陷阱：①只收集支持预设假设的信息；②混淆"说的"和"做的"；③只有单一来源。交叉验证原则：重要的判断，需要来自不同角度的信息互相印证。"),
    ]

    for title, content in questions:
        add_heading(doc, title, 3)
        doc.add_paragraph(content)

    add_heading(doc, "特别维度：历史改善尝试", 2)
    doc.add_paragraph("在所有调研维度里，有一个最常被跳过、但往往最关键的维度：这个问题或这个方向，以前改善过吗？")

    hist_questions = [
        "过去X年里，有没有人专门针对这个问题或这个方向做过改善尝试？",
        "当时具体做了什么？结果怎样？",
        "你认为为什么有效/无效？",
        "目前这个方向如果要推进，你认为最大的障碍是什么？"
    ]
    for q in hist_questions:
        doc.add_paragraph(f"• {q}", style='List Bullet')

    add_heading(doc, "✋ 练习：为你的项目设计调研计划", 2)

    doc.add_paragraph("第一步：信息分类（8分钟）")
    doc.add_paragraph("拿出你的体检清单，对每一条因素，标注：")
    tags_info = ["〔有〕现有数据或文件可以基本覆盖", "〔访〕需要通过访谈来了解真实状态",
                 "〔观〕需要通过现场观察", "〔文〕需要通过文档研读", "〔数〕需要通过数据分析"]
    for t in tags_info:
        doc.add_paragraph(f"• {t}")

    doc.add_paragraph("第二步：设计访谈问题（10分钟）")
    doc.add_paragraph("从标注了"〔访〕"的因素里，选出最重要的3条，为每条设计2~3个访谈问题。")

    doc.add_paragraph("好的访谈问题的特征：开放性（不能用是/否回答）；指向真实的行为或具体事件；包含至少一个关于历史经验的问题。")

    doc.add_paragraph("第三步：确定调研计划")
    add_form_table(doc, ["调研内容", "方式", "找谁", "预计时间"], [
        ["", "", "", ""],
        ["", "", "", ""],
        ["", "", "", ""],
    ])

    doc.add_page_break()

    # ===== 第四部分：四维分析 =====
    add_heading(doc, "第四部分　四维分析", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals4 = [
        "理解四维分析框架的四个维度及其判断标准",
        "能够对体检清单中的关键因素进行系统性的四维评估",
        "识别并避免"不舒服就不可动"的常见陷阱"
    ]
    for g in goals4:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "四维分析框架", 2)

    doc.add_paragraph("四维分析的目的，是对每一个关键因素做出四个方面的系统性判断：")

    dimensions = [
        ("维度一：影响大小", "这个因素对目标的影响有多显著？判断方法：如果这个因素改善30%，目标大约会变化多少？", "高/中/低"),
        ("维度二：影响范围", "这个因素的影响是局部的，还是系统性的？", "局部/系统"),
        ("维度三：可动性", "这个因素，实际上能被推动改善吗？", "直接可动/间接可动/当前不可动"),
        ("维度四：突破可能性", "在判断为"可动"的前提下，能推动到什么程度？", "高/中/低"),
    ]

    for title, desc, levels in dimensions:
        add_heading(doc, title, 3)
        doc.add_paragraph(desc)
        doc.add_paragraph(f"判断档位：{levels}")

    add_heading(doc, "可动性判断训练", 2)
    doc.add_paragraph("重新看快速判断练习里的因素，这次只判断"可动性"：")

    add_form_table(doc, ["因素", "你的可动性判断", "判断依据（具体是什么限制了它）"], [
        ["设备老旧（错误率12%）", "", ""],
        ["城市交通管控政策", "", ""],
        ["新员工上手周期3~4个月", "", ""],
        ["跨部门信息延迟40分钟", "", ""],
        ["大客户收货窗口集中", "", ""],
        ["系统数据未打通", "", ""],
    ])

    add_heading(doc, "这是可动性维度最大的陷阱：把"不舒服的事实"归入"不可动"", 3)

    traps = [
        ""这是政策规定，改不了"——政策是谁制定的？有没有弹性空间？",
        ""预算已经定了，增加不了"——预算是怎么决定的？有没有临时调整机制？",
        ""这需要A部门配合，他们不可能同意"——A部门有什么顾虑？有没有成本更低的合作方式？",
        ""以前试过，没用"——当时是什么方式？是方向不对，还是执行不到位？",
    ]
    for t in traps:
        doc.add_paragraph(f"• {t}")

    p = doc.add_paragraph()
    run = p.add_run("⚠️ "当前不可动"和"永远不可动"是完全不同的结论。")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)

    add_heading(doc, "✋ 练习：完成你的项目四维分析（20分钟）", 2)

    doc.add_paragraph("第一步：选出关键因素")
    doc.add_paragraph("从你的体检清单里，选出你认为最值得深度分析的6~10个因素。")

    doc.add_paragraph("第二步：填写四维分析表")
    add_form_table(doc, ["因素名称", "影响大小", "影响范围", "可动性", "突破可能性", "判断依据"], [
        ["", "高/中/低", "局部/系统", "直接/间接/不可动", "高/中/低", ""],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
        ["", "", "", "", "", ""],
    ])

    doc.add_paragraph("第三步：对"当前不可动"的判断发起挑战")
    doc.add_paragraph("找出你标注为"当前不可动"的因素，逐一回答：")
    challenges = [
        "这个因素不可动的具体约束是什么？（越具体越好）",
        "在什么条件改变的情况下，这个约束会松动？",
        "有没有间接推动的可能性？"
    ]
    for c in challenges:
        doc.add_paragraph(f"• {c}")

    p = doc.add_paragraph()
    run = p.add_run("带走一件事：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    p.add_run("每一个判断都需要有具体依据。"感觉不重要"和"感觉动不了"都不是依据，调研中得到的信息才是。")

    doc.add_page_break()

    # ===== 第五部分：突破口识别 =====
    add_heading(doc, "第五部分　突破口识别", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals5 = [
        "理解突破口的三个必要条件",
        "掌握用优先级矩阵快速定位突破口的方法",
        "能够对候选突破口进行深度验证"
    ]
    for g in goals5:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "20个行动 vs 3个突破口", 2)
    doc.add_paragraph("做完四维分析，很多人会把所有"影响大"且"可动"的因素都列成行动项，写出一份15~20条的行动清单。这份清单看起来很扎实，但在实际推进中：团队的精力被分散到20个方向上，每个方向只能做点表面功夫；三个月后，20件事都"在推进"，但没有一件做出了真正的深度改变。")

    p = doc.add_paragraph()
    run = p.add_run("突破口识别要做的，是一件更困难但更有价值的事：")
    run.bold = True
    p.add_run("在所有可动的因素里，找到2~4个真正的杠杆点——做了这几件事的局面会真正改变。")

    add_heading(doc, "什么是突破口", 2)
    doc.add_paragraph("突破口不是简单的"重要且可做的因素"，它有三个同时成立的条件：")

    conditions = [
        ("影响显著", "这个因素改善之后，目标指标会有可感知的变化"),
        ("有实际撬动可能", "在当前的资源、权限、时机条件下，这个因素是可以被实质性推动的"),
        ("在项目周期内能看到变化", "投入资源和精力，在这个项目的时间框架内，能看到真实的改变")
    ]
    for title, desc in conditions:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}：")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "突破口识别的优先级矩阵", 2)
    doc.add_paragraph("用两个维度来构建优先级框架：影响大小（高/中/低）× 可动性（直接可动/间接可动/当前不可动）")

    # 矩阵表格
    matrix_table = doc.add_table(rows=4, cols=4)
    matrix_table.style = 'Table Grid'
    matrix_data = [
        ("", "高影响", "中影响", "低影响"),
        ("直接可动", "★ 优先突破口", "次优先", "暂缓"),
        ("间接可动", "★ 优先突破口", "次优先", "暂缓"),
        ("当前不可动", "🔲 单独标注", "可忽略", "可忽略"),
    ]
    for ri, row_data in enumerate(matrix_data):
        row = matrix_table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p_cell in cell.paragraphs:
                for run_cell in p_cell.runs:
                    run_cell.font.size = Pt(10)
                    if ri == 0 or ci == 0:
                        run_cell.bold = True
                    if ri == 0:
                        set_cell_shading(cell, "1F3864")
                        run_cell.font.color.rgb = RGBColor(255, 255, 255)

    add_heading(doc, "从"候选"到"确认"：深度验证5个问题", 2)

    verify_qs = [
        "这个突破口具体是什么状态在影响目标？",
        "如果在这里发力，最好的结果是什么？需要多长时间能看到？",
        "发力的代价是什么？",
        "哪些人的配合是必须的？",
        "这个突破口和其他突破口之间有什么关系？"
    ]
    for i, q in enumerate(verify_qs, 1):
        doc.add_paragraph(f"问题{i}：{q}")

    add_heading(doc, "✋ 练习：识别你的突破口并完成深度验证", 2)

    doc.add_paragraph("第一步：把你的因素映射到优先级矩阵")
    add_form_table(doc, ["", "高影响", "中影响", "低影响"], [
        ["直接可动", "", "", ""],
        ["间接可动", "", "", ""],
        ["当前不可动", "", "", ""],
    ])

    doc.add_paragraph("第二步：选出你的突破口候选（2~4个）")
    add_form_table(doc, ["突破口", "选择依据（为什么这个是突破口）", "关键配合方"], [
        ["1.", "", ""],
        ["2.", "", ""],
        ["3.", "", ""],
    ])

    doc.add_paragraph("第三步：深度验证每个候选")
    add_form_table(doc, ["验证问题", "突破口1", "突破口2", "突破口3"], [
        ["具体是什么状态在影响目标？", "", "", ""],
        ["最好的结果是什么？需要多长时间？", "", "", ""],
        ["发力的代价是什么？", "", "", ""],
        ["哪些人的配合是必须的？", "", "", ""],
        ["与其他突破口的关系？", "", "", ""],
    ])

    p = doc.add_paragraph()
    run = p.add_run("带走一件事：")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    p.add_run("少即是多。精准的3个突破口，比宽泛的20条行动更有价值。")

    doc.add_page_break()

    # ===== 第六部分：行动方案 =====
    add_heading(doc, "第六部分　行动方案", 1, (0xB8, 0x35, 0x1C))

    add_heading(doc, "🎯 学习目标", 2)
    goals6 = [
        "理解行动方案的三层结构及其各自作用",
        "能够把突破口拆解为可执行的具体举措",
        "掌握前置条件检查的方法，避免执行中突然卡住"
    ]
    for g in goals6:
        doc.add_paragraph(f"• {g}", style='List Bullet')

    add_heading(doc, "行动方案的三个层次", 2)

    add_heading(doc, "第一层：解决方向（What）", 3)
    doc.add_paragraph("解决方向，是对这个突破口的行动承诺。")
    doc.add_paragraph("方向陈述句模板：", style='Quote')
    doc.add_paragraph("我们要通过【具体方式】，改善【具体因素】，使【具体指标】从【当前状态】达到【目标状态】。")

    add_heading(doc, "第二层：具体举措（How）", 3)
    doc.add_paragraph("一个突破口通常需要2~5个相互配合的举措，而不是单一动作。")
    doc.add_paragraph("每一个具体举措，需要明确四项：")
    items_how = ["做什么（具体动作，不是方向）", "谁来做（具体的人或角色）", "什么时候（具体的里程碑时间）", "怎么验证（可观察的成效标准）"]
    for item in items_how:
        doc.add_paragraph(f"• {item}")

    add_heading(doc, "第三层：前置条件检查（Pre-flight check）", 3)
    doc.add_paragraph("飞机起飞前，飞行员有一个系统性的检查清单，确认每一个系统都处于正常状态。行动方案的前置条件检查，是同样的逻辑。")

    pre_checks = [
        ("资源到位了吗？", "需要的人力、预算、设备、时间，是否已经确认可以调动？"),
        ("权限具备了吗？", "负责执行的人，是否有实际权限推进这件事？"),
        ("关键人的配合确认了吗？", "这些人目前的态度是什么？他们是支持的、观望的还是有顾虑的？"),
        ("有没有什么前置条件目前还不具备？", "如果有，需要先做什么来创造这个条件？谁来做？多长时间？"),
    ]
    for title, desc in pre_checks:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}")
        run.bold = True
        p.add_run(desc)

    add_heading(doc, "✋ 最终综合练习：完成你的行动方案（30分钟）", 2)

    doc.add_paragraph("为每个突破口填写完整的行动方案：")

    add_form_table(doc, ["维度", "内容"], [
        ["突破口名称", ""],
        ["解决方向", "我们要通过___，改善___，使___从___达到___"],
        ["举措一-做什么", ""],
        ["举措一-谁来做", ""],
        ["举措一-时间节点", ""],
        ["举措一-成效标准", ""],
        ["举措二-做什么", ""],
        ["举措二-谁来做", ""],
        ["举措二-时间节点", ""],
        ["举措二-成效标准", ""],
        ["举措三-做什么", ""],
        ["举措三-谁来做", ""],
        ["举措三-时间节点", ""],
        ["举措三-成效标准", ""],
        ["前置条件-资源确认", "资源已有□ / 需要确认□"],
        ["前置条件-权限确认", "具备□ / 需要争取□"],
        ["前置条件-关键配合方", ""],
        ["前置条件-目前不具备的", ""],
    ])

    doc.add_paragraph()

    add_heading(doc, "逻辑链验证：把行动和分析连起来", 2)
    doc.add_paragraph("一个有分析依据的行动方案，应该能够从每一个具体举措，一步步追溯回到分析链的起点：")

    chain_items = [
        "体检清单里的某个因素（第二部分）",
        "→ 调研时获得了什么信息支撑这个因素的状态（第三部分）",
        "→ 四维分析中对这个因素的判断（第四部分）",
        "→ 为什么选择了这个突破口（第五部分）",
        "→ 解决方向是什么",
        "→ 具体举措是什么"
    ]
    for item in chain_items:
        doc.add_paragraph(item)

    p = doc.add_paragraph()
    run = p.add_run("💡 区分"有分析依据的行动方案"和"拍脑袋的行动清单"，就看能不能跑通这条逻辑链。")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)

    p2 = doc.add_paragraph()
    run2 = p2.add_run("带走一件事：")
    run2.bold = True
    run2.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)
    p2.add_run("分析的价值不在于它有多全面，而在于它的每一步判断都有依据，并且最终转化成了有清晰责任、时间和验证标准的具体行动。")

    doc.add_page_break()

    # ===== 附录 =====
    add_heading(doc, "附录", 1, (0x1F, 0x1B, 0x16))

    add_heading(doc, "附录一：完整分析文件模板（一页纸版）", 2)

    doc.add_paragraph("项目名称：_______________________________________________")
    doc.add_paragraph("核心目标（具体可测量）：_______________________________________________")
    doc.add_paragraph("分析日期 / 调研期：_______________________________________________")

    doc.add_paragraph("A. 体检清单")
    add_form_table(doc, ["维度", "关键因素（各3~6条）"], [
        ["资源维度", ""],
        ["流程维度", ""],
        ["能力维度", ""],
        ["系统与工具维度", ""],
        ["管理机制维度", ""],
        ["外部条件维度", ""],
        ["历史遗留维度", ""],
    ])

    doc.add_paragraph("B. 四维分析表")
    add_form_table(doc, ["因素名称", "影响大小", "影响范围", "可动性", "突破可能性", "判断依据"], [
        ["", "H/M/L", "局部/系统", "直接/间接/不可动", "H/M/L", ""],
    ])

    doc.add_paragraph("C. 突破口识别（2~4个）")
    add_form_table(doc, ["突破口", "选择依据", "关键配合方"], [
        ["1.", "", ""],
        ["2.", "", ""],
        ["3.", "", ""],
    ])

    doc.add_paragraph("D. 行动方案")
    add_form_table(doc, ["突破口", "解决方向（一句话）", "主要举措（概要）", "关键前置条件"], [
        ["1.", "", "", ""],
        ["2.", "", "", ""],
        ["3.", "", "", ""],
    ])

    add_heading(doc, "附录二：常用问题速查", 2)

    qas = [
        ("Q：体检清单要有多少条才算够？", "没有固定数量要求，覆盖了5~7个主要维度、每个维度有3~5条具体条目，通常会有20~35条。"),
        ("Q：四维分析里一定要有真实数据吗？", "影响大小和突破可能性可以在没有精确数据的情况下做方向性判断，但要有来自调研的依据。可动性的判断，必须有具体的依据。"),
        ("Q：突破口一定要是"直接可动"的吗？", "不是。"间接可动"的突破口往往是影响更大、价值更高的方向——只是需要通过影响相关方来推动。"),
        ("Q：如果所有高影响因素都是"当前不可动"的？", "这种情况相对罕见，通常意味着：课题目标设定需要重新考虑，或者"不可动"的判断太保守。"),
    ]

    for q, a in qas:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        doc.add_paragraph(a)

    add_heading(doc, "附录三：行业适配提示", 2)

    industries = [
        ("轨道交通/公共交通类", "重点维度：设备与基础设施、运营组织、客流管理、外部政策。特别关注：折返效率、故障处理时长、不同时段的运能利用率分布。"),
        ("制造/工厂类", "重点维度：设备维护与故障率、工艺标准化程度、物料供应稳定性、质量控制流程。特别关注：瓶颈工序、班次交接流程、设备点检执行质量。"),
        ("企业内部职能改善类", "重点维度：流程效率与权责清晰度、跨部门信息传递、系统与工具的支撑程度。特别关注：手工环节往往是效率瓶颈。"),
    ]

    for title, content in industries:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}：")
        run.bold = True
        p.add_run(content)

    doc.add_paragraph()
    doc.add_paragraph()

    # 封底信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*版权所有 · 罗宏伟 · 本手册仅供本课程学员使用*")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"✅ 学员手册已生成：{OUTPUT_PATH}")

if __name__ == "__main__":
    create_handbook()

# -*- coding: utf-8 -*-
"""
生成 demo01-课题三层定义表示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def create_demo01():
    """创建课题三层定义表示例"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ===== 标题区 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("课题三层定义表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    # 副标题说明
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 课题三层定义表是破题的第一步，帮你从现象到本质层层深入")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(20)

    # ===== 学员信息区 =====
    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ["学员姓名", "张明华", "所属团队", "制造效能组"],
        ["完成日期", "2026-07-15", "指导教练", "李行动"]
    ]
    for i, row_data in enumerate(info_data):
        for j, text in enumerate(row_data):
            cell = info_table.rows[i].cells[j]
            cell.text = text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            if j in [0, 2]:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            run.font.size = Pt(10)

    doc.add_paragraph()

    # ===== 原始问题区 =====
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、原始问题（你听到的表面问题）")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.name = "微软雅黑"

    # 原始问题内容
    raw_para = doc.add_paragraph()
    run = raw_para.add_run("“新员工流失率高，怎么留住人？”")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = "微软雅黑"

    # 批注
    comment1 = doc.add_paragraph()
    run = comment1.add_run("[教练批注：这是同事们在茶水间常说的抱怨，属于'症状描述'而非'问题定义'。留人是一个模糊的大目标，需要拆解。]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    doc.add_paragraph()

    # ===== 三层定义表格 =====
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、课题三层定义")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.name = "微软雅黑"

    # 创建三层定义表格
    three_layer_table = doc.add_table(rows=4, cols=3)
    three_layer_table.style = 'Table Grid'

    # 表头
    headers = ["层次", "定义内容", "说明"]
    for j, header in enumerate(headers):
        cell = three_layer_table.rows[0].cells[j]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 设置表头背景色
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # 第一层：现象层
    row1_data = ["第一层\n现象层", "制造事业部新员工（入职1年内）流失率达28%，显著高于行业平均水平15%\n\n核心表现：\n• 试用期未满主动离职\n• 转正后3个月内离职\n• 关键岗位新员工流失", "这一层回答'发生了什么'——\n用具体数据描述现状，而非模糊感受"]
    for j, text in enumerate(row1_data):
        cell = three_layer_table.rows[1].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "微软雅黑"

    # 第二层：原因层
    row2_data = ["第二层\n原因层", "问题根因：制造类技能岗位培养周期长（6-12个月），新员工在'能力产出期'前已丧失信心\n\n关键症结点：\n• 培养周期与员工预期不匹配\n• 前3个月缺乏成就感反馈\n• 师徒制执行不到位", "这一层回答'为什么会发生'——\n不是找借口，而是找可控因素"]
    for j, text in enumerate(row2_data):
        cell = three_layer_table.rows[2].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "微软雅黑"

    # 第三层：课题层
    row3_data = ["第三层\n课题层\n（真正的问题）", "核心课题：如何设计'快速成就体验'机制，让新员工在入职90天内建立岗位自信？\n\n课题边界：\n• 聚焦新员工入职前90天\n• 聚焦技能岗位\n• 聚焦'心理胜任感'维度", "这一层回答'应该解决什么'——\n从'大而虚'到'小而精'的精准定位"]
    for j, text in enumerate(row3_data):
        cell = three_layer_table.rows[3].cells[j]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "微软雅黑"
                if j == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # ===== 课题质量检验 =====
    section_title3 = doc.add_paragraph()
    run = section_title3.add_run("三、课题质量检验（自检清单）")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.name = "微软雅黑"

    check_items = [
        ("✓ 具体性", "是否可衡量？", "流失率28%，精确到数字"),
        ("✓ 可控性", "是否在能力范围内？", "教练可干预，设计机制而非改变员工"),
        ("✓ 价值性", "解决后有价值吗？", "降低招聘成本，提升团队稳定性"),
        ("✓ 边界性", "范围清晰吗？", "入职90天内、技能岗位、特定成就感维度"),
        ("✓ 深度性", "触达本质了吗？", "不是留人，而是建立自信——治本")
    ]

    check_table = doc.add_table(rows=len(check_items), cols=3)
    check_table.style = 'Table Grid'

    for i, (item, question, answer) in enumerate(check_items):
        row = check_table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = question
        row.cells[2].text = answer
        for j, cell in enumerate(row.cells):
            para = cell.paragraphs[0]
            run = para.runs[0]
            run.font.size = Pt(10)
            run.font.name = "微软雅黑"
            if j == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph()

    # ===== 总结区 =====
    summary = doc.add_paragraph()
    run = summary.add_run("学员感悟：")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    summary_text = doc.add_paragraph()
    run = summary_text.add_run("「以前我觉得'留不住人'是个死结，现在发现，只要把问题从'怎么留人'改成'怎么让新人在90天内建立自信'，答案就清晰多了。」")
    run.font.size = Pt(11)
    run.font.italic = True

    # 保存
    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo01-课题三层定义表示例.docx"
    doc.save(output_path)
    print(f"demo01已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    create_demo01()

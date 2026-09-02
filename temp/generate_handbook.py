# File: D:\CC\temp\generate_handbook.py
# Generates: 绩效经营_学员手册.docx

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\新课开发\经营\系列\14_绩效经营——从指标分解到组织同频\04_Word文档\绩效经营_学员手册.docx"

doc = Document()

# ========================
# PAGE SETUP (A4)
# ========================
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ========================
# STYLE DEFINITIONS
# ========================

def set_run_font(run, font_name='Microsoft YaHei', size=12, bold=False, color=None):
    run.font.name = font_name
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(doc, text, level=1):
    """Add heading with proper styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)

    if level == 0:  # Document title
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 56, 100)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(24)
    elif level == 1:  # H1 - Module title
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 56, 100)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        # Add bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F3864')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:  # H2 - Subtitle
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 84, 150)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:  # H3
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(68, 68, 68)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def add_body_text(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'SimSun'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    # Bullet character
    bullet_run = p.insert_run(0, '• ')
    bullet_run.font.name = 'SimSun'
    bullet_run._r.append(OxmlElement('w:eastAsia'))
    bullet_run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    bullet_run.font.size = Pt(11)
    return p

def add_numbered_item(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{num}. {text}")
    run.font.name = 'SimSun'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)
    return p

def add_highlight_box(doc, text, box_type='tip'):
    """Add a highlighted box (tip/warning/note)"""
    colors = {
        'tip': ('E8F5E9', '2E7D32', '>>> '),
        'warning': ('FFEBEE', 'C62828', '⚠️ '),
        'note': ('E3F2FD', '1565C0', '📝 '),
        'time': ('FFF3E0', 'E65100', '⏱ ')
    }
    fill, text_color, prefix = colors.get(box_type, colors['tip'])

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    run = p.add_run(prefix + text)
    run.font.name = 'SimSun'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(*bytes.fromhex(text_color))
    return p

def add_table(doc, headers, rows, header_color='1F3864'):
    """Add a styled table"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Microsoft YaHei'
        run._r.append(OxmlElement('w:eastAsia'))
        run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'SimSun'
            run._r.append(OxmlElement('w:eastAsia'))
            run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
            run.font.size = Pt(10)

    return table

# ========================
# COVER PAGE
# ========================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('绩效经营')
run.font.name = 'Microsoft YaHei'
run._r.append(OxmlElement('w:eastAsia'))
run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(40)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 56, 100)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——从指标分解到组织同频')
run.font.name = 'Microsoft YaHei'
run._r.append(OxmlElement('w:eastAsia'))
run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(46, 84, 150)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('学员手册')
run.font.name = 'Microsoft YaHei'
run._r.append(OxmlElement('w:eastAsia'))
run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('适用对象：总经理、经营负责人、人力资源负责人、')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('组织发展负责人、各职能部门负责人及中层管理干部')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ========================
# TABLE OF CONTENTS
# ========================
add_heading_styled(doc, '目录', level=1)

toc_items = [
    ('第一章', '课程导入：为什么需要绩效经营？'),
    ('第二章', '理解经营逻辑：从战略到指标'),
    ('第三章', '组织架构与责任位次'),
    ('第四章', '指标分解：从组织到个人的路径'),
    ('第五章', '绩效追踪与过程管理'),
    ('第六章', '绩效复盘与持续改进'),
    ('附录A', '关键工具与模板'),
    ('附录B', '专业术语表'),
    ('附录C', '课后行动清单'),
]

for ch, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'{ch}  {title}')
    run.font.name = 'SimSun'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(11)

doc.add_page_break()

# ========================
# CHAPTER 1: 课程导入
# ========================
add_heading_styled(doc, '第一章  课程导入：为什么需要绩效经营？', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_body_text(doc, '通过本章学习，您将能够：')
add_bullet(doc, '识别企业常见的"部门墙"现象及其对整体绩效的影响')
add_bullet(doc, '理解"战略-计划-执行"脱节的根本原因')
add_bullet(doc, '认识到绩效经营的核心理念：全员对整体经营结果负责')

add_heading_styled(doc, '痛点场景', level=2)
add_body_text(doc, '请思考以下场景是否在您的企业中存在：', indent=True)

scenarios = [
    ('销售部', '签单时只考虑自己的业绩，不考虑交付能力和利润空间'),
    ('生产部', '追求产量最大化，忽视质量成本和库存积压'),
    ('采购部', '单方面压低供应商价格，导致质量问题或交期延误'),
    ('财务部', '严格执行预算控制，影响业务部门的正常运营需求'),
    ('研发部', '追求技术先进性，忽略市场时效性和商业价值'),
]

table = add_table(doc, ['部门', '典型表现'], scenarios)

doc.add_paragraph()
add_highlight_box(doc, '这些现象的共同特征：每个部门都完成了自己的KPI，但公司整体目标却未达成。这就是典型的"局部最优 vs 整体最优"矛盾。', 'warning')

add_heading_styled(doc, '绩效经营的核心理念', level=2)
add_body_text(doc, '绩效经营是一种将企业战略转化为全员行动的经营管理方法论，其核心主张包括：', indent=True)

principles = [
    '经营结果导向：最终衡量标准是企业整体经营业绩，而非部门局部指标',
    '责任穿透：打破部门壁垒，让每个岗位都明确自己与整体经营的关系',
    '过程可控：通过指标分解和追踪机制，实现过程的及时纠偏',
    '利益共享：个人收益与组织绩效直接挂钩，形成利益共同体',
]

for i, p_text in enumerate(principles, 1):
    add_numbered_item(doc, i, p_text)

add_heading_styled(doc, '本章小结', level=2)
add_body_text(doc, '绩效经营解决的本质问题是：如何让"部门各自为战"变成"全员目标一致"？答案不在于要求部门牺牲自身利益，而在于建立一套科学的指标分解和责任传导机制，让部门在追求自身利益的同时，自然地推动整体目标实现。')

doc.add_page_break()

# ========================
# CHAPTER 2: 理解经营逻辑
# ========================
add_heading_styled(doc, '第二章  理解经营逻辑：从战略到指标', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_bullet(doc, '掌握企业经营的底层逻辑：收入-成本-利润的关系')
add_bullet(doc, '理解战略、年度经营计划、KPI之间的传导关系')
add_bullet(doc, '学会用经营仪表盘来透视企业健康状况')

add_heading_styled(doc, '经营的本质', level=2)
add_body_text(doc, '企业经营的本质可以用一个简单的公式概括：', indent=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('利润 = 收入 - 成本')
run.font.name = 'Microsoft YaHei'
run._r.append(OxmlElement('w:eastAsia'))
run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(31, 56, 100)

add_body_text(doc, '所有企业经营活动都围绕这两个维度展开：提升收入、降低成本。绩效考核的真正目的不是"扣钱"，而是让每个岗位都关注这两个维度，为企业创造真实价值。', indent=True)

add_heading_styled(doc, '战略-计划-KPI传导链', level=2)

doc.add_paragraph()
add_highlight_box(doc, '好的指标体系就像一条清晰的传导链：战略决定方向，年度计划分解路径，KPI确保执行落地。三者脱节是很多企业绩效管理失效的根源。', 'tip')

conduction = [
    ('战略层', '使命愿景价值观\n3-5年战略目标', '战略解码会、战略研讨会', '战略澄清会'),
    ('计划层', '年度经营目标\n部门工作重点\n关键项目计划', '年度经营计划\n部门年度计划', '年度战略宣导会\n部门目标分解会'),
    ('指标层', '公司级KPI\n部门级KPI\n岗位级KPI', '目标责任书\n绩效合约', '月度/季度绩效回顾会'),
]

add_table(doc, ['层级', '输出物', '工具载体', '会议机制'], conduction)

add_heading_styled(doc, '经营仪表盘', level=2)
add_body_text(doc, '企业经营的健康状况可以通过以下仪表盘来监控：', indent=True)

dashboard = [
    ('财务维度', '收入增长率、利润率、现金流、资产回报率', '反映企业盈利能力和资金效率'),
    ('客户维度', '客户满意度、复购率、市场占有率', '反映企业在市场的竞争力'),
    ('运营维度', '交付及时率、产品质量合格率、库存周转率', '反映内部运营效率'),
    ('学习维度', '人均培训时长、人才流失率、员工满意度', '反映组织可持续发展能力'),
]

add_table(doc, ['维度', '核心指标', '管理含义'], dashboard)

doc.add_page_break()

# ========================
# CHAPTER 3: 组织架构与责任位次
# ========================
add_heading_styled(doc, '第三章  组织架构与责任位次', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_bullet(doc, '理解不同组织架构模式对绩效管理的影响')
add_bullet(doc, '掌握责任位次图的概念和绘制方法')
add_bullet(doc, '学会识别组织中的"责任空白地带"')

add_heading_styled(doc, '组织架构与绩效管理的关系', level=2)
add_body_text(doc, '组织架构决定了权责分配，权责分配决定了指标分解的可能性。不同的组织模式，绩效管理的重点也不同：', indent=True)

org_modes = [
    ('职能型', '按职能划分部门', '强调整性指标，如：生产效率、产品质量', '指标分解到部门，整体协调难'),
    ('事业部型', '按产品/客户划分', '强调利润中心，模拟独立核算', '事业部内部绩效清晰，整体协同需加强'),
    ('矩阵型', '双重汇报关系', '强调项目与职能的平衡', '责任界定复杂，需明确主次责任'),
    ('流程型', '按业务流程划分', '强调流程效率和客户满意度', '端到端指标明确，部门壁垒最小'),
]

add_table(doc, ['组织模式', '结构特点', '绩效重点', '绩效管理挑战'], org_modes)

add_heading_styled(doc, '责任位次图', level=2)
add_body_text(doc, '责任位次图是识别组织责任空白的重要工具。它清晰地展示了不同岗位在某一经营事项上的责任层级。', indent=True)

add_highlight_box(doc, 'RACI矩阵是最常用的责任位次工具：R=执行(Responsible), A=决策(Accountable), C=咨询(Consulted), I=知会(Informed)', 'note')

add_heading_styled(doc, '绘制责任位次图的步骤', level=2)

steps = [
    ('第一步：确定关键经营事项', '列出影响企业整体经营结果的关键事项，如：产品研发、市场开拓、质量管理、成本控制等'),
    ('第二步：识别参与岗位', '找出所有与该事项相关的岗位，不遗漏任何相关部门'),
    ('第三步：明确责任关系', '对每个岗位在每项事项上的责任进行判定，填入RACI矩阵'),
    ('第四步：检查责任空白', '检查是否有一项事项没有A角，或有太多R角，及时调整'),
    ('第五步：达成共识并固化', '与相关方确认并签字，作为后续绩效合约的依据'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Microsoft YaHei'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.name = 'SimSun'
    run2._r.append(OxmlElement('w:eastAsia'))
    run2._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run2.font.size = Pt(10.5)
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ========================
# CHAPTER 4: 指标分解
# ========================
add_heading_styled(doc, '第四章  指标分解：从组织到个人的路径', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_bullet(doc, '掌握指标分解的三大方法：直接分解法、驱动因素法、承诺对赌法')
add_bullet(doc, '学会识别关键成功因素(KSF)并转化为关键绩效指标(KPI)')
add_bullet(doc, '理解指标分解中的常见误区及规避方法')

add_heading_styled(doc, '指标分解的三大方法', level=2)

methods = [
    ('直接分解法', '将上级指标按比例直接分解到下级', '简单直接，易于操作', '适用于指标定义清晰、因果关系明确的场景', '可能导致"鞭打快牛"或"苦乐不均"'),
    ('驱动因素法', '分析指标背后的驱动因素，从原因找对策', '更能找到根本解', '适用于指标与驱动因素关系复杂的场景', '分析难度较大，需较强的经营洞察力'),
    ('承诺对赌法', '上下级通过协商确定指标值，附带激励对赌', '更能调动主动性', '适用于创新性、探索性工作', '可能存在博弈，协商成本高'),
]

add_table(doc, ['方法', '核心逻辑', '优点', '适用场景', '风险点'], methods)

add_heading_styled(doc, '关键成功因素(KSF)识别', level=2)
add_body_text(doc, 'KSF是影响KPI达成的关键变量。识别KSF的常用方法：', indent=True)

ksf_methods = [
    ('鱼骨图法', '从人机料法环测六个维度分析影响因子'),
    ('二八法则', '找出影响80%结果的那20%关键因素'),
    ('历史数据分析', '分析哪些变量与历史业绩相关性最高'),
    ('行业对标', '参考行业中优秀企业的关键成功要素'),
]

add_table(doc, ['方法', '说明'], ksf_methods)

add_heading_styled(doc, '指标分解的常见误区', level=2)

mistakes = [
    ('误区一：指标越多越好', '很多企业恨不得把员工绑在100个指标上，结果是"眉毛胡子一把抓"，哪个都做不好。好的指标体系应该精简，聚焦核心。'),
    ('误区二：唯KPI论', '"没有KPI是万万不能的，但KPI也不是万能的"。过度依赖KPI会引导员工"只做考核的"，忽视长期价值创造。'),
    ('误区三：指标割裂', '上下级之间、平行部门之间指标不关联，导致各自为战。好的指标体系应该形成"指标链"。'),
    ('误区四：目标静态', '市场环境在变，指标却一成不变。好的绩效管理应该建立"目标动态调整机制"。'),
]

for title, desc in mistakes:
    add_highlight_box(doc, f'{title}：{desc}', 'warning')

doc.add_page_break()

# ========================
# CHAPTER 5: 绩效追踪
# ========================
add_heading_styled(doc, '第五章  绩效追踪与过程管理', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_bullet(doc, '理解绩效追踪的重要性：过程决定结果')
add_bullet(doc, '掌握绩效追踪的四大机制：例会、报表、预警、辅导')
add_bullet(doc, '学会识别和解决绩效偏差')

add_heading_styled(doc, '为什么过程追踪比结果考核更重要？', level=2)
add_body_text(doc, '很多企业把大量精力放在年底的绩效考核上，却忽略了过程追踪。这种做法是"死后验尸"，为时已晚。好的绩效管理应该是：', indent=True)

add_highlight_box(doc, '追踪做得深，考核做得浅——因为大部分问题在过程中已经发现和解决了。', 'tip')

add_heading_styled(doc, '绩效追踪的四大机制', level=2)

mechanisms = [
    ('例会机制', '日/周/月度例会', '每天站会（15分钟）→ 每周部门会（1小时）→ 每月经营会（半天）', '及时同步信息，快速发现问题'),
    ('报表机制', '日常经营报表', '日报/周报/月报，聚焦核心指标完成情况', '用数据说话，避免主观判断'),
    ('预警机制', '异常预警系统', '提前设定指标红线，超线自动提醒', '把问题消灭在萌芽阶段'),
    ('辅导机制', '上级对下级辅导', '绩效面谈、技能辅导、资源协调', '帮助员工提升能力、解决困难'),
]

add_table(doc, ['机制', '形式', '频率/内容', '作用'], mechanisms)

add_heading_styled(doc, '绩效偏差识别与纠偏', level=2)
add_body_text(doc, '当绩效追踪发现偏差时，应该如何处理？', indent=True)

deviation_steps = [
    ('第一步：确认偏差', '偏差是真实的还是误报？数据来源是否可靠？'),
    ('第二步：分析原因', '是外部环境变化还是内部执行问题？是能力问题还是态度问题？'),
    ('第三步：制定对策', '针对原因制定纠偏措施，明确责任人和完成时间'),
    ('第四步：跟踪验证', '持续跟踪直到指标回到正轨，避免问题重复发生'),
]

for i, (step, desc) in enumerate(deviation_steps, 1):
    add_numbered_item(doc, i, f'{step}：{desc}')

doc.add_page_break()

# ========================
# CHAPTER 6: 绩效复盘
# ========================
add_heading_styled(doc, '第六章  绩效复盘与持续改进', level=1)

add_heading_styled(doc, '学习目标', level=2)
add_bullet(doc, '掌握绩效复盘的四步法：回顾目标、评估结果、分析原因、总结经验')
add_bullet(doc, '理解从个人复盘到组织复盘的升华路径')
add_bullet(doc, '学会将复盘成果转化为下一个周期的改进行动')

add_heading_styled(doc, '复盘的重要性', level=2)
add_body_text(doc, '一个不经复盘的经验只是经历，只有经过系统复盘的经验才能转化为能力。绩效复盘是连接"上一个周期"和"下一个周期"的桥梁。', indent=True)

add_heading_styled(doc, '绩效复盘四步法', level=2)

retro_method = [
    ('第一步：回顾目标', '当初设定的目标是什么？完成情况如何？目标是否合理？'),
    ('第二步：评估结果', '实际结果与目标对比：哪些超标？哪些达标？哪些未达标？'),
    ('第三步：分析原因', '深入挖掘成功因素和失败根因，找到根本原因而非表面现象'),
    ('第四步：总结经验', '提炼可复用的经验和教训，形成标准化的行动指南'),
]

for step, desc in retro_method:
    p = doc.add_paragraph()
    run = p.add_run(step)
    run.font.name = 'Microsoft YaHei'
    run._r.append(OxmlElement('w:eastAsia'))
    run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.name = 'SimSun'
    run2._r.append(OxmlElement('w:eastAsia'))
    run2._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
    run2.font.size = Pt(10.5)
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(8)

add_heading_styled(doc, '从个人复盘到组织复盘', level=2)
add_body_text(doc, '个人复盘是基础，但只有上升到组织层面，复盘的价值才能最大化。', indent=True)

levels = [
    ('个人复盘', '每天/每个任务后即时复盘', '培养反思习惯，提升个人能力'),
    ('团队复盘', '每个项目/里程碑后团队一起复盘', '共享经验教训，提升团队协作'),
    ('组织复盘', '每个季度/年度组织级复盘', '萃取组织智慧，形成组织记忆'),
]

add_table(doc, ['层次', '时机', '价值'], levels)

doc.add_page_break()

# ========================
# APPENDIX A: 关键工具与模板
# ========================
add_heading_styled(doc, '附录A  关键工具与模板', level=1)

add_heading_styled(doc, 'A1. 目标责任书模板', level=2)

p = doc.add_paragraph()
run = p.add_run('目标责任书是公司和部门负责人之间签订的绩效契约，应包含以下要素：')
run.font.name = 'SimSun'
run._r.append(OxmlElement('w:eastAsia'))
run._r.find(qn('w:eastAsia')).set(qn('w:eastAsia'), 'SimSun')
run.font.size = Pt(10.5)

responsibility_template = [
    ('基本信息', '责任人、部门、任职时间、签订日期'),
    ('目标承诺', '核心KPI指标及目标值、权重、计算方式、数据来源'),
    ('过程承诺', '关键行动承诺、里程碑节点、汇报机制'),
    ('激励承诺', '达成/超标/未达成的奖惩方案'),
    ('签署确认', '公司负责人签字、部门负责人签字、HR见证'),
]

add_table(doc, ['要素', '内容说明'], responsibility_template)

add_heading_styled(doc, 'A2. 绩效面谈记录表', level=2)

interview_template = [
    ('面谈基本信息', '员工姓名、部门、岗位、面谈日期、面谈人'),
    ('本期绩效回顾', '目标完成情况、关键 achievements、待改进事项'),
    ('原因分析', '成功因素分析、差距原因分析'),
    ('下期改进计划', '改进措施、所需资源支持、里程碑'),
    ('发展计划', '能力提升需求、培训计划、职业发展讨论'),
    ('双方签字', '员工签字确认、面谈人签字确认'),
]

add_table(doc, ['要素', '内容'], interview_template)

add_heading_styled(doc, 'A3. 经营仪表盘模板', level=2)

dashboard_template = [
    ('类别', '指标名称', '目标值', '实际值', '完成率', '预警级别'),
    ('财务类', '营业收入', 'XX万', '', '', ''),
    ('财务类', '净利润率', 'XX%', '', '', ''),
    ('客户类', '客户满意度', 'XX分', '', '', ''),
    ('运营类', '准时交付率', 'XX%', '', '', ''),
    ('学习类', '人均培训时长', 'XX小时', '', '', ''),
]

add_table(doc, ['类别', '指标名称', '目标值', '实际值', '完成率', '预警级别'], dashboard_template[1:], '4472C4')

doc.add_page_break()

# ========================
# APPENDIX B: 专业术语表
# ========================
add_heading_styled(doc, '附录B  专业术语表', level=1)

terms = [
    ('KPI', 'Key Performance Indicator', '关键绩效指标'),
    ('KSF', 'Key Success Factor', '关键成功因素'),
    ('OKR', 'Objectives and Key Results', '目标与关键成果'),
    ('RACI', 'Responsible/Accountable/Consulted/Informed', '责任分配矩阵'),
    ('SMART', 'Specific/Measurable/Achievable/Relevant/Time-bound', '目标设定原则'),
    ('BSC', 'Balanced Scorecard', '平衡计分卡'),
    ('PDCA', 'Plan/Do/Check/Act', '戴明环/持续改进循环'),
    ('ROI', 'Return on Investment', '投资回报率'),
    ('NPS', 'Net Promoter Score', '净推荐值'),
    ('COQ', 'Cost of Quality', '质量成本'),
]

add_table(doc, ['缩写', '英文全称', '中文释义'], terms)

doc.add_page_break()

# ========================
# APPENDIX C: 课后行动清单
# ========================
add_heading_styled(doc, '附录C  课后行动清单', level=1)

add_body_text(doc, '以下行动清单帮助您将课程所学转化为实际工作改进：', indent=True)

action_items = [
    ('组织层面', [
        '梳理公司级年度经营目标，确认与战略的关联性',
        '绘制公司级RACI矩阵，识别责任空白地带',
        '建立或优化月度经营例会机制',
        '设计公司级经营仪表盘，确定数据来源和更新频率',
    ]),
    ('部门层面', [
        '召开部门指标分解会，与上级指标对齐',
        '绘制部门级责任位次图，明确岗位责任',
        '建立部门周例会机制，跟踪指标进展',
        '组织部门绩效复盘会，形成改进清单',
    ]),
    ('个人层面', [
        '明确自己的KPI指标及计算方式',
        '识别影响自己KPI的关键成功因素(KSF)',
        '制定个人绩效改进行动计划',
        '每月进行一次个人绩效复盘',
    ]),
]

for category, items in action_items:
    add_heading_styled(doc, category, level=2)
    for item in items:
        add_bullet(doc, item)

# Final note
doc.add_paragraph()
add_highlight_box(doc, '绩效管理的提升是一个持续改进的过程。建议每个季度回顾一次行动清单的执行情况，不断优化和完善。', 'tip')

# ========================
# SAVE DOCUMENT
# ========================
doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
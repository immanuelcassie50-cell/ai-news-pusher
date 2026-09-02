# -*- coding: utf-8 -*-
"""
创建《政治判断力情景决策训练营》课程说明书
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, cn_font='微软雅黑', en_font='Arial', size=12, bold=False, color=None):
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, size=16, bold=True, color=(0, 51, 102), space_before=16, space_after=8, center=False):
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_run_font(run, cn_font='微软雅黑', size=size, bold=bold, color=color)
    return para

def add_body_para(doc, text, size=11, bold=False, indent=False, space_before=3, space_after=6, left_indent=0):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, cn_font='微软雅黑', size=size, bold=bold)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.first_line_indent = Pt(22)
    if left_indent:
        para.paragraph_format.left_indent = Pt(left_indent)
    return para

def add_bullet(doc, text, size=11):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_run_font(run, cn_font='微软雅黑', size=size)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    return para

def add_numbered(doc, text, size=11):
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    set_run_font(run, cn_font='微软雅黑', size=size)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    return para

def set_cell_style(cell, bold=False, size=10.5, color=None, fill_color=None, center=False):
    for para in cell.paragraphs:
        if center:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = bold
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
    if fill_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

def create_table(doc, headers, rows, col_widths=None, header_fill='003366'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_style(header_cells[i], bold=True, size=11, color=(255,255,255), fill_color=header_fill, center=True)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        fill = 'F0F5FA' if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            set_cell_style(row_cells[col_idx], size=10.5, fill_color=fill)

    return table

def create_docx():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ==================== 封面 ====================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_para.paragraph_format.space_after = Pt(12)
    run = title_para.add_run('《政治判断力情景决策训练营》')
    set_run_font(run, cn_font='微软雅黑', size=28, bold=True, color=(0, 51, 102))

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(6)
    subtitle_para.paragraph_format.space_after = Pt(12)
    run = subtitle_para.add_run('课程说明书')
    set_run_font(run, cn_font='微软雅黑', size=20, bold=False, color=(80, 80, 80))

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.paragraph_format.space_before = Pt(6)
    version_para.paragraph_format.space_after = Pt(48)
    run = version_para.add_run('V1.0')
    set_run_font(run, cn_font='微软雅黑', size=16, bold=True, color=(0, 102, 153))

    # ==================== 一、课程基本信息 ====================
    add_heading(doc, '一、课程基本信息', size=16, color=(0, 51, 102), space_before=24, space_after=12)

    basic_headers = ['项目', '内容']
    basic_rows = [
        ['课程全称', '《政治判断力情景决策训练营：在两难抉择中显性化你的判断力》'],
        ['一句话定位', '设计真实两难场景，把干部的判断力逼出来、暴露出来、然后当场纠偏'],
        ['课程时长', '1天（6-7课时）'],
        ['建议人数', '每场24人以内（4-6组，每组4-5人）'],
        ['课程形式', '情景沉浸 + 追问复盘 + 行为观察']
    ]
    create_table(doc, basic_headers, basic_rows)

    # ==================== 二、目标学员画像 ====================
    add_heading(doc, '二、目标学员画像', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    student_headers = ['类型', '描述']
    student_rows = [
        ['核心人群', '新提拔或后备中层干部，正处于"凡提必训"节点上的人群'],
        ['次要人群', '已在岗但组织认为其判断力有待观察/培养的中层干部'],
        ['行业覆盖', '国有企业、金融机构、大型制造业等有明确干部任职培训要求的单位'],
        ['学员特征', '业务能力普遍不弱，缺的不是知识，而是"遇到真实两难时的取舍直觉"，且这种直觉此前从未被训练过，只能靠自己"悟"，悟错了才知道']
    ]
    create_table(doc, student_headers, student_rows)

    # ==================== 三、能力模型（五维能力） ====================
    add_heading(doc, '三、能力模型（五维能力）', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    add_body_para(doc, '政治判断力由五个可训练子能力构成，本课程围绕这五个维度设计所有情景关卡：', size=11, space_before=6, space_after=8)

    ability_headers = ['序号', '能力维度', '定义', '典型训练场景']
    ability_rows = [
        ['1', '全局意识', '能否跳出局部利益看到组织整体影响', '局部业绩 vs 组织声誉/长远利益的冲突场景'],
        ['2', '底线意识', '能否在压力下守住不能碰的红线', '业绩诱惑/领导施压/时间紧迫三重压力叠加场景'],
        ['3', '分寸感', '能否在灵活处理和坚持原则之间找到恰当的度', '合规边缘的"擦边球"请求如何恰当应对'],
        ['4', '舆情敏感度', '能否预判一个决策可能引发的舆论/群体反应', '业务合理但传播出去可能引发误解的决策场景'],
        ['5', '请示汇报时机', '能否判断什么事该自己扛、什么事必须及时上报', '信息不完整、时间压力大的紧急决策场景']
    ]
    create_table(doc, ability_headers, ability_rows)

    # ==================== 四、课程特色 ====================
    add_heading(doc, '四、课程特色', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    feature_headers = ['特色', '内涵']
    feature_rows = [
        ['没有标准答案', '每个情景关卡都经过精心设计，选项之间无明显对错之分，训练的是判断力而非记忆力'],
        ['压力叠加设计', '每个关卡叠加至少两种压力源（时间/人际/信息不全/利益诱惑），还原真实决策环境'],
        ['代价可见', '学员做出选择后，讲师揭示该选择可能引发的后续连锁反应，让后果"提前发生"'],
        ['追问技术', '通过三层递进追问（还原逻辑→暴露盲区→迁移内化），让学员自己看见自己判断逻辑的盲区']
    ]
    create_table(doc, feature_headers, feature_rows)

    # ==================== 五、教学流程 ====================
    add_heading(doc, '五、教学流程', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    flow_headers = ['模块', '时长', '内容概要']
    flow_rows = [
        ['模块一：开场导入', '1课时', '破冰、课程定位、五维能力模型讲解、规则说明'],
        ['模块二：初级关卡热身', '1.5课时', '2个初级关卡，建立信心和参与感'],
        ['模块三：中级关卡第一组', '1课时', '2个中级关卡，开始真正暴露判断逻辑差异'],
        ['模块四：中级关卡第二组', '1.5课时', '2个中级关卡，价值密度最高部分'],
        ['模块五：高级关卡', '1.5课时', '2-3个两难张力最强的关卡，全天高潮'],
        ['模块六：结营复盘', '0.5课时', '每位学员提炼个人"判断力自检清单"']
    ]
    create_table(doc, flow_headers, flow_rows)

    # ==================== 六、课时安排 ====================
    add_heading(doc, '六、课时安排', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    schedule_headers = ['环节', '时长', '累计', '说明']
    schedule_rows = [
        ['开场导入', '60分钟', '0-60', '含破冰、定位、五维模型、规则'],
        ['关卡一 ZZPD-01', '45分钟', '60-105', '初级：优秀员工名额分配'],
        ['关卡二 ZZPD-02', '45分钟', '105-150', '初级：老领导合规边缘请求'],
        ['关卡三 ZZPD-03', '30分钟', '150-180', '中级：信息不全紧急签批'],
        ['关卡四 ZZPD-04', '30分钟', '180-210', '中级：供应商感谢费'],
        ['关卡五 ZZPD-05', '45分钟', '210-255', '中级：慰问金公示风波'],
        ['关卡六 ZZPD-06', '45分钟', '255-300', '中级：编制争夺战'],
        ['关卡七 ZZPD-10', '30分钟', '300-330', '高级：领导软钉子'],
        ['关卡八 ZZPD-11', '30分钟', '330-360', '高级：截访还是依法'],
        ['关卡九 ZZPD-12', '30分钟', '360-390', '高级：补贴发放灵活处理'],
        ['结营复盘', '30分钟', '390-420', '自检清单提炼'],
        ['总计', '7课时', '420分钟', '']
    ]
    create_table(doc, schedule_headers, schedule_rows)

    # ==================== 七、评估体系 ====================
    add_heading(doc, '七、评估体系', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    eval_headers = ['评估工具', '用途', '使用场景']
    eval_rows = [
        ['现场行为观察记录表', '讲师记录每位学员在各关卡中的选择倾向和追问表现', '课程进行中'],
        ['学员判断力画像模板', '形成个人判断力特征描述', '课程结束后'],
        ['判断力自检清单模板', '学员结营时提炼个人收获和未来自检原则', '结营复盘环节']
    ]
    create_table(doc, eval_headers, eval_rows)

    # ==================== 八、物料清单 ====================
    add_heading(doc, '八、物料清单', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    materials_headers = ['类别', '物料名称', '数量/说明']
    materials_rows = [
        ['教学工具', '情景关卡卡牌（12关）', '每组1套，共4-6套'],
        ['教学工具', '追问脚本手册', '讲师用，每关配3套追问脚本'],
        ['学员用', '学员手册', '每人1本，含关卡背景和空白笔记页'],
        ['学员用', '判断力自检清单模板', '每人1份，结营时使用'],
        ['场地要求', '分组讨论桌椅', '4-6组，每组4-5人座位'],
        ['场地要求', '白板/大白纸', '每组1个，用于记录讨论要点'],
        ['设备要求', '投影设备', '1套，用于展示关卡内容'],
        ['设备要求', '计时器', '1个，用于控制关卡时间']
    ]
    create_table(doc, materials_headers, materials_rows)

    # ==================== 九、讲师资质要求 ====================
    add_heading(doc, '九、讲师资质要求', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    add_body_para(doc, '本课程对讲师的要求明显高于一般党课：', size=11, bold=True, space_before=6, space_after=8)

    讲师_reqs = [
        '具备极强的现场追问和临场应变能力，学员的选择和反应具有不可预测性，不能完全照本宣科',
        '对敏感话题的边界有清晰判断，知道哪些追问可以深挖，哪些需要点到为止',
        '建议讲师在正式带课前，先以"学员"身份完整走一遍所有关卡，确保对每个关卡的多种可能反应路径都有预案',
        '具备控场能力，全天强调"这是训练场景，不是考试"，反复在关键节点重申，降低学员的防御心理',
        '高级关卡环节容易出现学员情绪紧张或现场争论，需要提前准备好几套"降温话术"'
    ]
    for req in 讲师_reqs:
        add_bullet(doc, req)

    # ==================== 十、与其他课程的区别 ====================
    add_heading(doc, '十、与其他党业融合课程的区别', size=16, color=(0, 51, 102), space_before=20, space_after=12)

    diff_headers = ['课程', '解决的问题']
    diff_rows = [
        ['双带头人破局课', '"经验能否转化沉淀"的问题'],
        ['政治判断力情景决策训练营', '"关键时刻判断力是否过硬"的问题'],
        ['廉政风险情景决策课', '"高风险岗位纪律红线意识"的问题']
    ]
    create_table(doc, diff_headers, diff_rows)

    # Save
    output_path = 'D:/新课开发/党业融合/政治判断力/完整课程包/001-课程说明书/课程说明书-政治判断力V1.0.docx'
    doc.save(output_path)
    print(f'Created: {output_path}')

if __name__ == '__main__':
    create_docx()

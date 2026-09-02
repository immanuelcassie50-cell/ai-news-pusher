# File: create_instructor_v2.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Helper functions
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(4)
    return p

def set_run_font(run):
    run.font.name = 'Microsoft YaHei'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

def add_page_break(doc):
    doc.add_page_break()

# ================== Cover ==================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\n\n')
run = title.add_run('Content Tools Confidence Trilogy')
run.font.size = Pt(36)
run.font.bold = True
set_run_font(run)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('\n\n')
run = subtitle.add_run('Building Persuasive Expression')
run.font.size = Pt(28)
run.font.bold = True
set_run_font(run)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('\n\n')
run = subtitle2.add_run('—— Instructor Handbook ——')
run.font.size = Pt(20)
set_run_font(run)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('\n\n\n\n\n\n\n\n')
run = info.add_run('Version: V1.0')
run.font.size = Pt(14)
set_run_font(run)
info.add_run('\n')
run = info.add_run('Copyright: Luo Hongwei')
run.font.size = Pt(14)
set_run_font(run)
info.add_run('\n')
run = info.add_run('Internal Use Only - Certified Instructors')
run.font.size = Pt(12)
set_run_font(run)

add_page_break(doc)

# ================== Table of Contents ==================
add_heading(doc, 'Table of Contents', 1)

toc_items = [
    'Chapter 1: Instructor Preparation Checklist',
    'Chapter 2: Module Guide - Content Foundation',
    'Chapter 3: Module Guide - Structure Tools',
    'Chapter 4: Module Guide - Confident Expression',
    'Chapter 5: Course Conclusion Guide',
    'Appendix A: Behavior Observation Scale A (Content Foundation)',
    'Appendix B: Behavior Observation Scale B (Structure Tools)',
    'Appendix C: Behavior Observation Scale C (Confident Expression)',
    'Appendix D: Closing Recommendations',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

add_page_break(doc)

# ================== Chapter 1: Instructor Preparation Checklist ==================
add_heading(doc, 'Chapter 1: Instructor Preparation Checklist', 1)

add_heading(doc, '1.1 Pre-Course Checklist', 2)

checklist_items = [
    'Course slides prepared (9 chapters + opening + closing)',
    'Form printing completed (12 core forms per student)',
    'Behavior observation scales copied (per assistant)',
    'Markers ready (2 thick, 2 thin)',
    'Flip charts/whiteboards available',
    'Projector/large screen confirmed',
    'Seating arranged for discussions (U-shape or group tables)',
    'Coffee break area prepared (if course > 4 hours)',
    'Course materials packed: handouts, form packets, markers, sticky notes',
    'Test audio/video equipment if using AI tools',
    'Review学员手册 content to anticipate questions',
]
table = doc.add_table(rows=len(checklist_items)+1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Item'
table.rows[0].cells[1].text = 'Status'
for i, item in enumerate(checklist_items):
    table.rows[i+1].cells[0].text = item
    table.rows[i+1].cells[1].text = '[ ]'

add_para(doc, '')

add_heading(doc, '1.2 Key Differences from Other Courses', 2)
add_para(doc, 'This course is different from traditional speaking courses in three ways:')
add_bullet(doc, 'Tool-based: We teach expression as a systematic tool, not just techniques')
add_bullet(doc, 'Task-driven: Each student works on their own real task throughout the course')
add_bullet(doc, 'AI-assisted: We use AI tools to assist content preparation and structure checking')

add_heading(doc, '1.3 Three Toolkit Summary', 2)

toolkits = [
    ('Toolkit 1: Content Foundation', 'Content Authenticity + Content Specificity + Content Focus', 'Without authentic support, expressions are just empty words'),
    ('Toolkit 2: Structure Tools', 'Structure Clarity + Structure Layering + Structure Transition', 'Expression without structure is scattered sand'),
    ('Toolkit 3: Confident Expression', 'Confidence Delivery + Emotional Resonance + Pressure Response', 'Let your attitude enhance your content'),
]
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = 'Toolkit'
table2.rows[0].cells[1].text = 'Core Elements'
table2.rows[0].cells[2].text = 'Core Message'
for i in range(3):
    table2.rows[i+1].cells[0].text = toolkits[i][0]
    table2.rows[i+1].cells[1].text = toolkits[i][1]
    table2.rows[i+1].cells[2].text = toolkits[i][2]

add_page_break(doc)

# ================== Chapter 2: Module Guide - Content Foundation ==================
add_heading(doc, 'Chapter 2: Module Guide - Content Foundation (Part 1)', 1)

add_heading(doc, '2.1 Chapter 1: Content Authenticity', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Identify the essential difference between "empty expression" and "supported expression"')
add_bullet(doc, 'Master three types of content support (data/case/emotion) and their usage scenarios')
add_bullet(doc, 'Build a complete content support system for scenario card task')

add_heading(doc, 'Activity Sequence', 3)
activities = [
    ('A1-1', 'Challenge Cognition', '"What makes you believe?" personal reflection', '10 min'),
    ('A1-2', 'Explain + Demo', 'Content Three Laws framework + case demonstration', '20 min'),
    ('A1-3', 'Individual Practice', 'Form 1.1 Content Support Audit', '10 min'),
    ('A1-4', 'Explain + Demo', 'Three support types in detail', '15 min'),
    ('A1-5', 'Individual Practice', 'Form 1.2 Scenario Card Content Support Design', '15 min'),
    ('A1-6', 'Pair Check', 'Pair mutual review of support quality', '10 min'),
]
table3 = doc.add_table(rows=len(activities)+1, cols=4)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Activity ID'
table3.rows[0].cells[1].text = 'Type'
table3.rows[0].cells[2].text = 'Name'
table3.rows[0].cells[3].text = 'Duration'
for i, a in enumerate(activities):
    table3.rows[i+1].cells[0].text = a[0]
    table3.rows[i+1].cells[1].text = a[1]
    table3.rows[i+1].cells[2].text = a[2]
    table3.rows[i+1].cells[3].text = a[3]

add_para(doc, '')

add_heading(doc, 'Teaching Tips', 3)
add_bullet(doc, 'Key point: Emphasize that persuasiveness comes from trust, not rhetoric')
add_bullet(doc, 'Common mistake: Students often think they need more data, but actually need more relevant data')
add_bullet(doc, 'Demo tip: Use a familiar business scenario to demonstrate support quality')

add_heading(doc, '2.2 Chapter 2: Content Specificity', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master four dimensions of making vague concepts specific')
add_bullet(doc, 'Transform abstract viewpoints into specific and actionable descriptions')

add_heading(doc, 'Activity Sequence', 3)
activities2 = [
    ('A2-1', 'Explain + Interaction', 'Vague vs Specific: comparison exercise', '15 min'),
    ('A2-2', 'Individual Practice', 'Form 2.1 Vague Expression Revision', '12 min'),
    ('A2-3', 'Explain + Demo', 'Four dimensions of specificity + Focus principle', '18 min'),
    ('A2-4', 'Individual Practice', 'Form 3.2 Scenario Card Focus Design', '15 min'),
]
table4 = doc.add_table(rows=len(activities2)+1, cols=4)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Activity ID'
table4.rows[0].cells[1].text = 'Type'
table4.rows[0].cells[2].text = 'Name'
table4.rows[0].cells[3].text = 'Duration'
for i, a in enumerate(activities2):
    table4.rows[i+1].cells[0].text = a[0]
    table4.rows[i+1].cells[1].text = a[1]
    table4.rows[i+1].cells[2].text = a[2]
    table4.rows[i+1].cells[3].text = a[3]

add_para(doc, '')

add_heading(doc, 'Teaching Tips', 3)
add_bullet(doc, 'Key point: Specificity is not exhaustiveness - find the key detail that enables instant understanding')
add_bullet(doc, 'Exercise tip: Have students revise their own vague expressions from real work')

add_page_break(doc)

# ================== Chapter 3: Module Guide - Structure Tools ==================
add_heading(doc, 'Chapter 3: Module Guide - Structure Tools (Part 2)', 1)

add_heading(doc, '3.1 Chapter 4: Structure Clarity', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Understand the role of "structure" in expression')
add_bullet(doc, 'Master three basic expression structures (time/space/importance)')

add_heading(doc, 'Activity Sequence', 3)
activities3 = [
    ('A3-1', 'Challenge Cognition', '"What do you remember?" immediate test', '10 min'),
    ('A3-2', 'Explain + Demo', 'Structure Clarity: Three basic structures', '15 min'),
    ('A3-3', 'Individual Practice', 'Form 4.1 Structure Design Form', '12 min'),
    ('A3-4', 'Pair Practice', 'Pair practice: Structure exercise', '10 min'),
]
table5 = doc.add_table(rows=len(activities3)+1, cols=4)
table5.style = 'Table Grid'
table5.rows[0].cells[0].text = 'Activity ID'
table5.rows[0].cells[1].text = 'Type'
table5.rows[0].cells[2].text = 'Name'
table5.rows[0].cells[3].text = 'Duration'
for i, a in enumerate(activities3):
    table5.rows[i+1].cells[0].text = a[0]
    table5.rows[i+1].cells[1].text = a[1]
    table5.rows[i+1].cells[2].text = a[2]
    table5.rows[i+1].cells[3].text = a[3]

add_para(doc, '')

add_heading(doc, '3.2 Chapter 5: Structure Layering', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master the core of Pyramid Principle: Conclusion First')
add_bullet(doc, 'Use layered structure to organize complex content')

add_heading(doc, 'Teaching Tips', 3)
add_bullet(doc, 'Key demo: Compare "Conclusion First" vs "Layered Build-up" versions')
add_bullet(doc, 'Common mistake: Students think conclusion-first is "rude" - need to disabuse this')

add_heading(doc, '3.3 Chapter 6: Structure Transition', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master three transition techniques')
add_bullet(doc, 'Design complete transitions for scenario card tasks')

add_heading(doc, 'Bridge to Part 3', 3)
add_para(doc, 'At the end of Chapter 6, transition to Part 3 by explaining: Content is oil, Structure is car, Confidence is engine. No oil, car wont move. No engine, car moves even slower. The three toolkits are progressive and cannot be skipped.')

add_page_break(doc)

# ================== Chapter 4: Module Guide - Confident Expression ==================
add_heading(doc, 'Chapter 4: Module Guide - Confident Expression (Part 3)', 1)

add_heading(doc, '4.1 Chapter 7: Confidence Delivery', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master four methods to enhance confidence delivery')
add_bullet(doc, 'Understand three key factors affecting confidence: Voice, Body, Attitude')

add_heading(doc, '4.2 Chapter 8: Emotional Resonance', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master three methods to build emotional resonance')
add_bullet(doc, 'Balance "rational content" with "emotional delivery"')

add_heading(doc, '4.3 Chapter 9: Pressure Response', 2)

add_heading(doc, 'Learning Objectives', 3)
add_bullet(doc, 'Master the "Receive-Transfer-Answer" framework')
add_bullet(doc, 'Practice responding to challenges in pairs')

add_heading(doc, 'Pressure Scenario Demonstrations', 3)
scenarios = [
    ('Sudden Question', 'Do not rush to answer, say "good question" first, give yourself 3 seconds'),
    ('Being Questioned', 'Receive: acknowledge legitimacy; Transfer: redirect to core point; Answer: provide response'),
    ('Forgetting Lines', 'Do not pause too long; skip forgotten part; continue with summary sentence'),
]
table6 = doc.add_table(rows=4, cols=2)
table6.style = 'Table Grid'
table6.rows[0].cells[0].text = 'Scenario'
table6.rows[0].cells[1].text = 'Response Method'
for i, s in enumerate(scenarios):
    table6.rows[i+1].cells[0].text = s[0]
    table6.rows[i+1].cells[1].text = s[1]

add_page_break(doc)

# ================== Chapter 5: Course Conclusion Guide ==================
add_heading(doc, 'Chapter 5: Course Conclusion Guide', 1)

add_heading(doc, '5.1 Final Exercise Format (5 minutes per person)', 2)
add_para(doc, 'Minute 0-1: Opening + My topic')
add_para(doc, 'Minute 1-2: Core point (conclusion first)')
add_para(doc, 'Minute 2-4: Content with structure + support')
add_para(doc, 'Minute 4-5: Summary + Confident delivery')

add_heading(doc, '5.2 Structured Feedback Dimensions', 2)
add_bullet(doc, 'Three things worth affirming')
add_bullet(doc, 'One thing most needs improvement')
add_bullet(doc, 'One specific suggestion')

add_heading(doc, '5.3 Closing Speech Points', 2)
add_bullet(doc, 'Three toolkits: Content -> Structure -> Confidence')
add_bullet(doc, 'Persuasive expression is not innate, it can be designed')
add_bullet(doc, 'Tools will become熟练, structure will internalize, confidence will accumulate')

add_page_break(doc)

# ================== Appendix A: Behavior Observation Scale A ==================
add_heading(doc, 'Appendix A: Behavior Observation Scale A (Content Foundation)', 1)

add_para(doc, 'Purpose: Evaluate student performance in Content Foundation module')

scale_a_items = [
    ('1.1', 'Uses specific data to support arguments', '1', '2', '3', '4', '5'),
    ('1.2', 'Uses real cases/examples appropriately', '1', '2', '3', '4', '5'),
    ('1.3', 'Expresses emotions appropriately without exaggeration', '1', '2', '3', '4', '5'),
    ('1.4', 'Makes vague concepts specific', '1', '2', '3', '4', '5'),
    ('1.5', 'Focuses on most important points (max 3)', '1', '2', '3', '4', '5'),
]
table_a = doc.add_table(rows=len(scale_a_items)+1, cols=7)
table_a.style = 'Table Grid'
table_a.rows[0].cells[0].text = 'Item'
table_a.rows[0].cells[1].text = 'Behavior'
table_a.rows[0].cells[2].text = '1-Poor'
table_a.rows[0].cells[3].text = '2-Below'
table_a.rows[0].cells[4].text = '3-Avg'
table_a.rows[0].cells[5].text = '4-Good'
table_a.rows[0].cells[6].text = '5-Excellent'
for i, item in enumerate(scale_a_items):
    for j in range(7):
        table_a.rows[i+1].cells[j].text = item[j]

add_para(doc, '')

# ================== Appendix B: Behavior Observation Scale B ==================
add_heading(doc, 'Appendix B: Behavior Observation Scale B (Structure Tools)', 1)

add_para(doc, 'Purpose: Evaluate student performance in Structure Tools module')

scale_b_items = [
    ('2.1', 'Uses appropriate structure type (time/space/importance)', '1', '2', '3', '4', '5'),
    ('2.2', 'Structure is "visible" to audience', '1', '2', '3', '4', '5'),
    ('2.3', 'Uses pyramid structure (conclusion first)', '1', '2', '3', '4', '5'),
    ('2.4', 'Layers are clear and logical', '1', '2', '3', '4', '5'),
    ('2.5', 'Transitions are smooth and logical', '1', '2', '3', '4', '5'),
]
table_b = doc.add_table(rows=len(scale_b_items)+1, cols=7)
table_b.style = 'Table Grid'
table_b.rows[0].cells[0].text = 'Item'
table_b.rows[0].cells[1].text = 'Behavior'
table_b.rows[0].cells[2].text = '1-Poor'
table_b.rows[0].cells[3].text = '2-Below'
table_b.rows[0].cells[4].text = '3-Avg'
table_b.rows[0].cells[5].text = '4-Good'
table_b.rows[0].cells[6].text = '5-Excellent'
for i, item in enumerate(scale_b_items):
    for j in range(7):
        table_b.rows[i+1].cells[j].text = item[j]

add_para(doc, '')

# ================== Appendix C: Behavior Observation Scale C ==================
add_heading(doc, 'Appendix C: Behavior Observation Scale C (Confident Expression)', 1)

add_para(doc, 'Purpose: Evaluate student performance in Confident Expression module')

scale_c_items = [
    ('3.1', 'Voice: Volume appropriate, pace steady, pauses used', '1', '2', '3', '4', '5'),
    ('3.2', 'Eye contact: Natural, varied, connects with audience', '1', '2', '3', '4', '5'),
    ('3.3', 'Uses "I think" not "I feel" for confident statements', '1', '2', '3', '4', '5'),
    ('3.4', 'Creates emotional resonance appropriately', '1', '2', '3', '4', '5'),
    ('3.5', 'Handles pressure/questions calmly with R-T-A framework', '1', '2', '3', '4', '5'),
]
table_c = doc.add_table(rows=len(scale_c_items)+1, cols=7)
table_c.style = 'Table Grid'
table_c.rows[0].cells[0].text = 'Item'
table_c.rows[0].cells[1].text = 'Behavior'
table_c.rows[0].cells[2].text = '1-Poor'
table_c.rows[0].cells[3].text = '2-Below'
table_c.rows[0].cells[4].text = '3-Avg'
table_c.rows[0].cells[5].text = '4-Good'
table_c.rows[0].cells[6].text = '5-Excellent'
for i, item in enumerate(scale_c_items):
    for j in range(7):
        table_c.rows[i+1].cells[j].text = item[j]

add_para(doc, '')

# ================== Appendix D: Closing Recommendations ==================
add_heading(doc, 'Appendix D: Closing Recommendations', 1)

add_heading(doc, 'Key Takeaways to Reinforce', 2)
add_bullet(doc, 'Content Foundation: Authenticity, Specificity, Focus - without these, expressions are empty words')
add_bullet(doc, 'Structure Tools: Clarity, Layering, Transition - without structure, complex content cannot be understood')
add_bullet(doc, 'Confident Expression: Confidence, Resonance, Pressure Response - confidence comes from preparation')

add_heading(doc, 'Post-Course Follow-up Recommendations', 2)
add_bullet(doc, 'Day 1: Send course summary email with key frameworks')
add_bullet(doc, 'Day 3: Remind students to practice with real tasks')
add_bullet(doc, 'Day 7: Send tools quick-reference card')
add_bullet(doc, 'Day 30: Follow up on 30-day action plan progress')

add_heading(doc, 'Certification Requirements', 2)
add_para(doc, 'Students who complete the course receive a Certificate of Completion.')
add_para(doc, 'Certificate requirements:')
add_bullet(doc, 'Attend full course')
add_bullet(doc, 'Complete all exercises')
add_bullet(doc, 'Present scenario card with three-toolkit integration')

# Save
output_path = 'D:/NewCourse/PublicExpression/05_ContentToolsConfidence/Instructor_Handbook.docx'
doc.save(output_path)
print(f'Saved to {output_path}')

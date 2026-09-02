# File: create_handbook_v2.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Helper functions
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(4)
    return p

def set_run_font(run):
    run.font.name = 'Microsoft YaHei'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

def add_page_break(doc):
    doc.add_page_break()

# ================== Cover ==================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\n\n')
run = title.add_run('Content Tools Confidence Trilogy')
run.font.size = Pt(36)
run.font.bold = True
set_run_font(run)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('\n\n')
run = subtitle.add_run('Building Persuasive Expression')
run.font.size = Pt(28)
run.font.bold = True
set_run_font(run)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('\n\n')
run = subtitle2.add_run('—— Student Handbook ——')
run.font.size = Pt(20)
set_run_font(run)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('\n\n\n\n\n\n\n\n')
run = info.add_run('Version: V1.0')
run.font.size = Pt(14)
set_run_font(run)
info.add_run('\n')
run = info.add_run('Copyright: Luo Hongwei')
run.font.size = Pt(14)
set_run_font(run)

add_page_break(doc)

# ================== Table of Contents ==================
add_heading(doc, 'Table of Contents', 1)

toc_items = [
    'Introduction',
    'Part 1: Content Foundation',
    '    Chapter 1: Content Authenticity',
    '    Chapter 2: Content Specificity',
    '    Chapter 3: Content Focus',
    'Part 2: Structure Tools',
    '    Chapter 4: Structure Clarity',
    '    Chapter 5: Structure Layering',
    '    Chapter 6: Structure Transition',
    'Part 3: Confident Expression',
    '    Chapter 7: Confidence Delivery',
    '    Chapter 8: Emotional Resonance',
    '    Chapter 9: Pressure Response',
    'Course Conclusion',
    'Appendix 1: Glossary',
    'Appendix 2: Tools Index',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

add_page_break(doc)

# ================== Introduction ==================
add_heading(doc, 'Introduction', 1)

add_heading(doc, 'Course Overview', 2)

map_text = "Part 1 (Content Foundation) - Build a solid foundation of authentic, specific, and well-supported information that gives your arguments credibility.\n\nPart 2 (Structure Tools) - Use structured methods to organize information, making complex content clear, orderly, and easy to understand.\n\nPart 3 (Confident Expression) - Not just techniques, but the transmission of confidence and attitude, giving your expression power and warmth."
add_para(doc, map_text)

add_heading(doc, 'Form 0.1 Self-Assessment', 2)
add_para(doc, 'Purpose: Understand your current expression confidence level. Retake this after the course to see changes.')

# Create table for self-assessment
table = doc.add_table(rows=8, cols=5)
table.style = 'Table Grid'
headers = ['Behavior', 'Almost Never', 'Occasionally', 'Often', 'Almost Always']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h

self_assessment_items = [
    '1. Before public speaking, I prepare specific data and case support',
    '2. I can clearly explain complex concepts to others',
    '3. I use structural words like "First, Second, Third" in my expressions',
    '4. I am confident in my expression content and not worried about being questioned',
    '5. I can stay calm under pressure and continue to express clearly',
    '6. I proactively observe and learn from excellent speakers',
    '7. I believe persuasiveness is innate and cannot be improved through learning',
]
for i, item in enumerate(self_assessment_items):
    table.rows[i+1].cells[0].text = item
    for j in range(1, 5):
        table.rows[i+1].cells[j].text = 'O'

add_para(doc, '')

add_heading(doc, 'Form 0.2 My Scenario Card', 2)
add_para(doc, 'Purpose: Select a real work task that will run through the entire course.')

# Create table for scenario card
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Table Grid'
labels = [
    'My Position/Role',
    'My Selected Expression Task',
    'How is this task usually done? (Brief steps within 3 lines)',
    'What is most challenging about this task?',
    'What improvement do I hope tools can help with?',
    'What is the final output form of this task? (Speech/Report/Email/Proposal/Other)',
]
for i in range(6):
    table2.rows[i].cells[0].text = labels[i]
    table2.rows[i].cells[1].text = ''

add_page_break(doc)

# ================== Part 1: Content Foundation ==================
add_heading(doc, 'Part 1: Content Foundation', 1)
add_para(doc, 'Without authentic support, expressions are just empty words. A good content foundation is the foundation of all persuasiveness.')

# Chapter 1: Content Authenticity
add_heading(doc, 'Chapter 1: Content Authenticity', 1)
add_para(doc, 'Let data, cases, and emotions become the support for your expression.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Identify the essential difference between "empty expression" and "supported expression"')
add_bullet(doc, 'Master three types of content support (data/case/emotion) and their usage scenarios')
add_bullet(doc, 'Build a complete content support system for your scenario card task')

add_heading(doc, 'Key Point 1.1: Why Your Expression Lacks Persuasiveness', 2)
add_para(doc, 'To make someone remember and believe what you say requires not fluent rhetoric or beautiful PPT, but authentic and credible support.')
add_para(doc, 'Core Insight: The foundation of persuasiveness is not rhetoric, but trust. Trust comes from specific data, real cases, and warm emotions - not empty slogans and adjectives.')

add_heading(doc, 'Key Point 1.2: Three Types of Content Support', 2)
add_para(doc, 'Data Support - Applicable scenarios: expressions that need to prove scale, change, or effect. The power of data lies in specificity.')
add_para(doc, 'Case Support - Applicable scenarios: expressions that need to prove feasibility and operability. The power of cases lies in authenticity.')
add_para(doc, 'Emotion Support - Applicable scenarios: expressions that need to build resonance and connection. The power of emotion lies in being moved. But emotion cannot be separated from authenticity.')

add_heading(doc, 'Key Point 1.3: Four Manifestations of Support Deficit', 2)
add_para(doc, 'Empty Adjective Type: lacking specificity')
add_para(doc, 'Subjective Assertion Type: lacking objective support')
add_para(doc, 'False Data Type: damaging trust')
add_para(doc, 'Piling Type: lacking focus')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 1-A (Basic): Support Type Identification')
add_para(doc, 'Exercise 1-B (Application): Empty Expression Revision')
add_para(doc, 'Exercise 1-C (Extended): Scenario Card Support Completion')

add_page_break(doc)

# Chapter 2: Content Specificity
add_heading(doc, 'Chapter 2: Content Specificity', 1)
add_para(doc, 'Transform "almost" into "clear".')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Identify the difference between "vague expression" and "specific expression"')
add_bullet(doc, 'Master four dimensions of making vague concepts specific')
add_bullet(doc, 'Transform abstract viewpoints into specific and actionable descriptions')

add_heading(doc, 'Key Point 2.1: Vagueness is the Greatest Enemy of Expression', 2)
add_para(doc, 'Three harms of vague expression: confuse the audience; make the audience disbelieve; make yourself feel uneasy.')
add_para(doc, 'Core Insight: The opposite of specific is not "wrong" but "vague". Vague expression wastes time for both parties. Specific is not verbose, but precise.')

add_heading(doc, 'Key Point 2.2: Four Dimensions of Specificity', 2)
add_para(doc, 'Time Specific: Vague "handle ASAP" -> Specific "respond within 24 hours"')
add_para(doc, 'Number Specific: Vague "many people" -> Specific "more than 70% of users"')
add_para(doc, 'Behavior Specific: Vague "improve service quality" -> Specific "response time reduced to within 2 hours"')
add_para(doc, 'Standard Specific: Vague "do your best" -> Specific "customer rating no less than 4.5"')

add_heading(doc, 'Key Point 2.3: Specificity is Not Exhaustiveness', 2)
add_para(doc, 'Specificity = Finding the key detail that allows the audience to "understand instantly". Judgment standard: Can your expression make people form a picture in their minds immediately?')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 2-A (Basic): Specificity Judgment')
add_para(doc, 'Exercise 2-B (Application): Specificity Revision')
add_para(doc, 'Exercise 2-C (Extended): Scenario Card Expression Review')

add_page_break(doc)

# Chapter 3: Content Focus
add_heading(doc, 'Chapter 3: Content Focus', 1)
add_para(doc, 'Less is more.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the importance of "focus" in expression')
add_bullet(doc, 'Master the "Three-Point Principle": Why human memory works best with three items')
add_bullet(doc, 'Be able to do subtraction for your expression content')

add_heading(doc, 'Key Point 3.1: Why Audience Can Only Remember Three Points', 2)
add_para(doc, 'Human working memory capacity is limited. Research shows that humans most easily remember 3 items in short-term memory.')
add_para(doc, 'Core Insight: Focus is not deletion but selection. Choose the most important and abandon others. This is the courage and wisdom of a speaker.')

add_heading(doc, 'Key Point 3.2: Application of Three-Point Principle', 2)
add_para(doc, 'Structured three: "I have three points to make"')
add_para(doc, 'Content three: Each "aspect" can have sub-content, but the top level must be three.')
add_para(doc, 'How to do subtraction: Ask yourself - If the audience can only remember one thing, what would it be? Then ask: What are the two points second only to this one?')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 3-A (Basic): Focus Identification')
add_para(doc, 'Exercise 3-B (Application): Subtraction Practice')
add_para(doc, 'Exercise 3-C (Extended): Scenario Card Focus Practice')

add_page_break(doc)

# ================== Part 2: Structure Tools ==================
add_heading(doc, 'Part 2: Structure Tools', 1)
add_para(doc, 'Expression without structure is scattered sand; expression with structure is a sharp sword.')

# Chapter 4: Structure Clarity
add_heading(doc, 'Chapter 4: Structure Clarity', 1)
add_para(doc, 'Use frameworks to present complex information orderly.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the role of "structure" in expression')
add_bullet(doc, 'Master three basic expression structures (time/space/importance)')
add_bullet(doc, 'Design clear structure for your scenario card task')

add_heading(doc, 'Key Point 4.1: Structure is a Map for the Audience', 2)
add_para(doc, 'Role of structure: Reduce cognitive burden; enhance memory; build trust.')
add_para(doc, 'Core Insight: Structure is not for your own use, but for the audience. No matter how complex the logic, it can be understood as long as the structure is clear.')

add_heading(doc, 'Key Point 4.2: Three Basic Structures', 2)
add_para(doc, 'Time Structure (Past-Present-Future): Applicable scenarios - reports, summaries, plans')
add_para(doc, 'Space Structure (Whole-Part-Whole): Applicable scenarios - proposal introductions, product descriptions, analysis reports')
add_para(doc, 'Importance Structure (Most Important-Less Important-General): Applicable scenarios - proposals, recommendations, decision reports')

add_heading(doc, 'Key Point 4.3: Structure Must be "Visible"', 2)
add_para(doc, 'Good structure must not only exist in your mind but also be visible to the audience.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 4-A (Basic): Structure Type Judgment')
add_para(doc, 'Exercise 4-B (Application): Structure Design')
add_para(doc, 'Exercise 4-C (Extended): Scenario Card Structure Implementation')

add_page_break(doc)

# Chapter 5: Structure Layering
add_heading(doc, 'Chapter 5: Structure Layering', 1)
add_para(doc, 'Break complex content into digestible layers.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the difference between "layering" and "point listing"')
add_bullet(doc, 'Master the core of the Pyramid Principle: Conclusion First')
add_bullet(doc, 'Use layered structure to organize complex content')

add_heading(doc, 'Key Point 5.1: Pyramid Principle', 2)
add_para(doc, 'Core Principle: Conclusion First - Say the conclusion first, then the reasons.')
add_para(doc, 'Benefits: Satisfy "impatient" audience; give "patient" audience a framework.')
add_para(doc, 'Structure: Conclusion -> Reasons -> Support')

add_heading(doc, 'Key Point 5.2: The Art of Layering', 2)
add_para(doc, 'Layering should follow the MECE Principle: Mutually Exclusive, Collectively Exhaustive. Three layers are most common.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 5-A (Basic): Pyramid Structure Judgment')
add_para(doc, 'Exercise 5-B (Application): Revision Practice')
add_para(doc, 'Exercise 5-C (Extended): Scenario Card Pyramid Practice')

add_page_break(doc)

# Chapter 6: Structure Transition
add_heading(doc, 'Chapter 6: Structure Transition', 1)
add_para(doc, 'Make expression flow smoothly like water.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the role of "transition" in expression')
add_bullet(doc, 'Master three transition techniques')
add_bullet(doc, 'Design complete transitions for scenario card tasks')

add_heading(doc, 'Key Point 6.1: Transition is the "Adhesive" of Expression', 2)
add_para(doc, 'Role of transition: Connect the previous; preview the next; establish connections.')
add_para(doc, 'Core Insight: Good transition makes audience feel "smooth", bad transition makes audience feel "jumpy".')

add_heading(doc, 'Key Point 6.2: Three Transition Techniques', 2)
add_para(doc, 'Language Transition: Use one or two sentences to clearly state "connection" and "turn".')
add_para(doc, 'Repetition Transition: Repeat key information from the previous and introduce the next.')
add_para(doc, 'Action Transition: Use action words like "look", "compare", "show" to introduce the next step.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 6-A (Basic): Transition Type Judgment')
add_para(doc, 'Exercise 6-B (Application): Transition Completion')
add_para(doc, 'Exercise 6-C (Extended): Scenario Card Transition Practice')

add_page_break(doc)

# ================== Part 3: Confident Expression ==================
add_heading(doc, 'Part 3: Confident Expression', 1)
add_para(doc, 'The tools are ready, the structure is clear - now it is time to let them shine.')

# Chapter 7: Confidence Delivery
add_heading(doc, 'Chapter 7: Confidence Delivery', 1)
add_para(doc, 'Let your attitude enhance your content.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the role of "confidence" in expression')
add_bullet(doc, 'Identify three key factors affecting confidence delivery')
add_bullet(doc, 'Master four methods to enhance confidence delivery')

add_heading(doc, 'Key Point 7.1: Why You Are Not Confident When Speaking', 2)
add_para(doc, 'Three factors affecting confidence delivery: Voice (volume/speed/pause); Body (eye contact/posture/gestures); Attitude (tone/responsibility).')
add_para(doc, 'Core Insight: Confidence is not "pretending to be great" but "letting your real level play out normally".')

add_heading(doc, 'Key Point 7.2: Four Methods of Confidence Delivery', 2)
add_para(doc, 'Voice Control: Emphasize key words; pause appropriately; do not speak fast all the time.')
add_para(doc, 'Eye Contact: Look at one person until they finish speaking; do not look at only one person; do not look down at manuscript.')
add_para(doc, 'Taking Responsibility: Use "I think" instead of "I feel"; dare to say "I do not know, but I will find out".')
add_para(doc, 'Accepting Imperfection: Being nervous is normal; pausing is allowed; mistakes can be remedied.')

add_heading(doc, 'Key Point 7.3: Preparation is the Foundation of Confidence', 2)
add_para(doc, 'Real confidence comes from: sufficient content preparation; clear structure; having rehearsed.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 7-A (Basic): Confidence Signal Identification')
add_para(doc, 'Exercise 7-B (Application): Voice Practice')
add_para(doc, 'Exercise 7-C (Extended): Scenario Card Rehearsal')

add_page_break(doc)

# Chapter 8: Emotional Resonance
add_heading(doc, 'Chapter 8: Emotional Resonance', 1)
add_para(doc, 'Let the audience not only hear but also feel.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the role of "emotional resonance" in persuasion')
add_bullet(doc, 'Master three methods to build emotional resonance')
add_bullet(doc, 'Balance "rational content" with "emotional delivery"')

add_heading(doc, 'Key Point 8.1: Why Emotional Resonance is Important', 2)
add_para(doc, 'Role of emotional resonance: Make audience willing to listen; help audience remember; make audience believe.')
add_para(doc, 'Core Insight: Rational persuasion is "I know you are right", emotional resonance is "I am moved by you". The best expression has both.')

add_heading(doc, 'Key Point 8.2: Three Methods to Build Emotional Resonance', 2)
add_para(doc, 'Say "we" not "you": Create a community.')
add_para(doc, 'Share real feelings: Acknowledge difficulties; share your real experiences; express genuine emotions.')
add_para(doc, 'Tell warm stories: Specific characters; real scenes; emotional changes.')

add_heading(doc, 'Key Point 8.3: Balance Rational and Emotional', 2)
add_para(doc, 'Good expression = Rational content + Emotional delivery.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 8-A (Basic): Resonance Type Judgment')
add_para(doc, 'Exercise 8-B (Application): Revision Practice')
add_para(doc, 'Exercise 8-C (Extended): Scenario Card Resonance Design')

add_page_break(doc)

# Chapter 9: Pressure Response
add_heading(doc, 'Chapter 9: Pressure Response', 1)
add_para(doc, 'Express clearly even under pressure.')

add_heading(doc, 'Learning Objectives', 2)
add_bullet(doc, 'Understand the mechanism of pressure on expression')
add_bullet(doc, 'Master response methods for three pressure scenarios')
add_bullet(doc, 'Use the "Receive-Transfer-Answer" framework to respond to challenges')

add_heading(doc, 'Key Point 9.1: Sources and Manifestations of Pressure', 2)
add_para(doc, 'Pressure scenarios in workplace expression: Being suddenly asked by leadership; being questioned; discovering you forgot your lines; facing unfriendly audience.')
add_para(doc, 'Core Insight: Pressure itself is not the problem, your reaction to pressure is the problem. Accepting the existence of pressure is the first step in dealing with it.')

add_heading(doc, 'Key Point 9.2: Response to Three Pressure Scenarios', 2)
add_para(doc, 'Sudden Question: Do not rush to answer, say "good question" first, give yourself 3 seconds to think.')
add_para(doc, 'Being Questioned: Receive-Transfer-Answer. Receive - acknowledge the legitimacy of the question; Transfer - redirect to your core point; Answer - provide your response.')
add_para(doc, 'Forgetting Lines: Do not pause too long, skip the forgotten part, continue with a summary sentence.')

add_heading(doc, 'Key Point 9.3: Daily Pressure Training', 2)
add_para(doc, 'Deliberately interrupt yourself during rehearsal to practice continuing from the interruption; record and review to discover your pressure signals; deliberately create pressure scenarios.')

add_heading(doc, 'Chapter Exercise', 2)
add_para(doc, 'Exercise 9-A (Basic): Pressure Signal Identification')
add_para(doc, 'Exercise 9-B (Application): Receive-Transfer-Answer Practice')
add_para(doc, 'Exercise 9-C (Extended): Scenario Card Pressure Rehearsal')

add_page_break(doc)

# ================== Course Conclusion ==================
add_heading(doc, 'Course Conclusion: My Expression Confidence Action System', 1)

add_heading(doc, 'Comprehensive Practice: Full Chain of Three Toolkits', 2)
add_para(doc, 'Return to your scenario card and go through the complete chain with the three toolkits.')

# Create table for three toolkits
table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Table Grid'
headers = ['Toolkit', 'Core Elements', 'Your Application']
for i, h in enumerate(headers):
    table3.rows[0].cells[i].text = h
toolkits = [
    ('Part 1: Content Foundation', 'Authenticity + Specificity + Focus', ''),
    ('Part 2: Structure Tools', 'Clarity + Layering + Transition', ''),
    ('Part 3: Confident Expression', 'Confidence + Resonance + Pressure Response', ''),
]
for i in range(3):
    table3.rows[i+1].cells[0].text = toolkits[i][0]
    table3.rows[i+1].cells[1].text = toolkits[i][1]
    table3.rows[i+1].cells[2].text = toolkits[i][2]

add_para(doc, '')

add_heading(doc, 'Form C.1 Three Toolkits Retest Self-Assessment', 2)
add_para(doc, 'At the end of the course, fill out this form again - compare with your initial self-assessment to see the changes.')

add_heading(doc, 'Form C.2 My 30-Day Expression Confidence Action Plan', 2)
add_para(doc, 'Step 1: Confirm 2 priority skills to improve')
add_para(doc, 'Step 2: Three 10-day stages')
add_para(doc, 'Step 3: One specific metric that can be verified after 30 days')

add_heading(doc, 'Words for the Journey', 2)
add_para(doc, 'You spent one day going through the three toolkits. This is not a theory but a way of working that you can start using from tomorrow.')
p = add_para(doc, 'The three toolkits share a common underlying logic - persuasive expression is not innate, it can be designed. The combination of good content foundation, clear structure, and confident delivery can make your expression penetrate people hearts.')

add_page_break(doc)

# ================== Appendix ==================
add_heading(doc, 'Appendix 1: Glossary', 1)

terms = [
    ('Content Authenticity', 'Expression requires specific data, cases, or emotional support to build trust'),
    ('Content Specificity', 'Vagueness is the greatest enemy of expression; achieve specificity through four dimensions'),
    ('Content Focus', 'Less is more; human memory works best with three points; learn to do subtraction'),
    ('Structure Clarity', 'Use frameworks to present complex information orderly, reducing cognitive burden'),
    ('Structure Layering', 'Pyramid Principle; conclusion first; present layer by layer'),
    ('Structure Transition', 'Transition is the adhesive of expression; make audience feel smooth'),
    ('Confidence Delivery', 'Three factors - voice, body, attitude - affect confidence delivery'),
    ('Emotional Resonance', 'Good expression = Rational content + Emotional delivery'),
    ('Pressure Response', 'Use "Receive-Transfer-Answer" framework; confidence comes from preparation'),
]
table4 = doc.add_table(rows=len(terms)+1, cols=2)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Term'
table4.rows[0].cells[1].text = 'Definition'
for i in range(len(terms)):
    table4.rows[i+1].cells[0].text = terms[i][0]
    table4.rows[i+1].cells[1].text = terms[i][1]

add_para(doc, '')

add_heading(doc, 'Appendix 2: Tools Index', 1)

tools = [
    ('0.1', 'Self-Assessment Form', 'Introduction', 'Understand starting point'),
    ('0.2', 'My Scenario Card', 'Introduction', 'Lock in real task'),
    ('1.1', 'Content Support Audit', 'Chapter 1', 'Diagnose support quality'),
    ('1.2', 'Scenario Card Support Design', 'Chapter 1', 'Build support system'),
    ('2.1', 'Vague Expression Revision', 'Chapter 2', 'Specificity practice'),
    ('3.2', 'Scenario Card Focus Design', 'Chapter 3', 'Focus design'),
    ('4.1', 'Structure Design Form', 'Chapter 4', 'Design expression structure'),
    ('5.1', 'Pyramid Practice Form', 'Chapter 5', 'Conclusion-first practice'),
    ('5.2', 'Scenario Card Layering Design', 'Chapter 5', 'Design layered structure'),
    ('6.2', 'Scenario Card Transition Design', 'Chapter 6', 'Design transitions'),
    ('7.1', 'Confidence Delivery Self-Check', 'Chapter 7', 'Diagnose confidence delivery'),
    ('7.2', 'Scenario Card Confidence Preparation', 'Chapter 7', 'Prepare thoroughly for confidence'),
    ('8.2', 'Scenario Card Resonance Design', 'Chapter 8', 'Design resonance points'),
    ('9.2', 'Receive-Transfer-Answer Practice', 'Chapter 9', 'Practice responding to challenges'),
    ('C.1', 'Three Toolkits Retest', 'Conclusion', 'Compare before and after'),
    ('C.2', '30-Day Action Plan', 'Conclusion', 'Three-stage implementation'),
]
table5 = doc.add_table(rows=len(tools)+1, cols=4)
table5.style = 'Table Grid'
headers = ['Tool #', 'Tool Name', 'Location', 'Main Purpose']
for i, h in enumerate(headers):
    table5.rows[0].cells[i].text = h
for i in range(len(tools)):
    for j in range(4):
        table5.rows[i+1].cells[j].text = tools[i][j]

# Save
output_path = 'D:/NewCourse/PublicExpression/05_ContentToolsConfidence/Student_Handbook.docx'
doc.save(output_path)
print(f'Saved to {output_path}')

# -*- coding: utf-8 -*-
import sys
import os
import json
import codecs

# Find the docx file
docx_path = None
desktop = r'C:\Users\Administrator\Desktop\标书'
for f in os.listdir(desktop):
    if f.endswith('.docx') and '五强' in f:
        docx_path = os.path.join(desktop, f)
        break

print(f"Found file: {docx_path}")

from docx import Document
doc = Document(docx_path)

output = []
output.append(f"FILE: {os.path.basename(docx_path)}")
output.append("=" * 60)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        output.append(f"[{i}] ({para.style.name}): {para.text}")

# Tables
for t_idx, table in enumerate(doc.tables):
    output.append(f"\n--- Table {t_idx} ---")
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip()[:80] for cell in row.cells]
        output.append(f"Row {row_idx}: {' | '.join(cells)}")

# Write to file
with codecs.open(r'D:\CC\temp\tender_content_utf8.txt', 'w', encoding='utf-8') as fw:
    fw.write('\n'.join(output))

print("Done. Written to tender_content_utf8.txt")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

output_base = r"D:\新课开发\工作手册\知识工作者深度工作保护\完整课程包"
os.makedirs(output_base, exist_ok=True)

def set_run_font(run, font_name='Microsoft YaHei', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, font_name='Microsoft YaHei', size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
    return p

def add_quote_box(doc, text, title=None):
    """Add a quote box with gray background"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    if title:
        run = p.add_run(title + "\n")
        run.font.bold = True
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    run.font.italic = True
    return p

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(9)
    return table

print("=" * 50)
print("Creating 认证班宣传文案...")
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title
add_para(doc, '', size=24)
add_para(doc, '深度工作主权手册', size=32, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '认证班', size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(46, 84, 150))
add_para(doc, '', size=16)
add_para(doc, '知识工作者的注意力保护与恢复系统', size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '', size=20)

# Subtitle
add_para(doc, '「时间会被切割成什么形状，取决于你有多认真对待它」', size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128, 128, 128))

doc.add_page_break()

# Problem statement
add_heading(doc, '你正在经历的', 1, color=(198, 40, 40))
add_para(doc, '你是不是也这样？')
add_bullet(doc, '每天被消息、会议、协作需求切割得支离破碎，感觉自己像陀螺一样转个不停')
add_bullet(doc, '年底回顾，发现说不清楚自己深度做成了什么，忙碌却空洞')
add_bullet(doc, '尝试过番茄钟、时间块、各种待办清单，坚持不了几天就失效')
add_bullet(doc, '想要保护自己的深度工作时间，却不知道怎么和团队、上级沟通')
add_bullet(doc, '疲惫感日积月累，越来越难进入专注状态')

add_quote_box(doc, "每一条中断记录，都是一次隐藏的议价过程。你不是没时间深度工作，你是舍不得放弃'随时在线'带来的被需要感。", "——罗宏伟")

# Root cause
add_heading(doc, '问题的根源', 1)
add_para(doc, '大多数时间管理课程教你"怎么更好地利用时间"，但忽略了最根本的问题：')
add_para(doc, '', size=10)
add_para(doc, '注意力不是被"管理"走的，是被我们自己一次次"允许"拿走的。', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '', size=10)
add_para(doc, '每一条消息弹出来，不是打断已经发生，是打断正在被"提议"——你点开它的那一刻，才是打断真正生效的时刻。')

# Course introduction
add_heading(doc, '这门课不一样', 1)
add_heading(doc, '不是教你管理时间，是帮你重建注意力的认知主权', 2)
add_para(doc, '本课程不是又一门"时间管理技巧"课。我们不教你怎么排日程、怎么用番茄钟、怎么列待办清单——这些工具你已经学过太多次了。')
add_para(doc, '这门课要做的，是先帮你看清一个被长期忽视的事实：你的注意力，不是在某个瞬间被突然夺走的，是你自己一次次签字画押让出去的。认知不重建，工具永远是摆设。')

add_heading(doc, '课程体系', 2)
headers = ['模块', '核心议题', '你将获得']
rows = [
    ['PART 1 认知层', '打断×忙碌成瘾×时间盒', '理解注意力让渡的心理机制'],
    ['PART 1 实操层', '会议×重启税×环境设计', '掌握可即刻落地的保护工具'],
    ['PART 2 协作层', '向上管理×产出可见性', '在团队中建立可持续边界'],
    ['PART 2 心态层', '团队文化×恢复信念×长期主义', '建立注意力保护的信念系统'],
]
create_table(doc, headers, rows)

# Core value
add_heading(doc, '你将带走', 1)
add_bullet(doc, '一套完整的个人深度工作保护方案（时间盒+计划表+中断日志）')
add_bullet(doc, '一份可供团队讨论的深度工作契约草案')
add_bullet(doc, '16个经过验证的实操工具，拿来就能用')
add_bullet(doc, '一个持续践行的同伴学习小组')

# Certification benefits
add_heading(doc, '认证班独有', 1)
add_bullet(doc, '获得罗宏伟老师亲笔签名的结业证书')
add_bullet(doc, '加入深度工作主权校友会，终身享有同伴支持')
add_bullet(doc, '课程结束后30天、90天跟踪回访，确保行为改变真正发生')
add_bullet(doc, '优先获得后续进阶课程的参与资格')

# Target audience
add_heading(doc, '适合谁', 1)
add_bullet(doc, '企业知识工作者：产品经理、研发人员、咨询顾问、数据分析师')
add_bullet(doc, '企业中高层管理者：需要保护深度决策时间，同时带领团队建立深度工作文化')
add_bullet(doc, '创业者和自由职业者：需要极高的自我管理能力')
add_bullet(doc, '任何对注意力保护有迫切需求、想要真正改变现状的人')

# Course info
add_heading(doc, '课程信息', 1)
headers = ['项目', '内容']
rows = [
    ['课程时长', '2天（每天6小时，共12小时）'],
    ['学员规模', '25-40人'],
    ['课程形式', '讲授+小组讨论+工具演练+行动承诺'],
    ['认证证书', '完成课程后颁发罗宏伟老师亲笔签名证书'],
    ['后续支持', '30天、90天跟踪回访'],
]
create_table(doc, headers, rows)

# Quote from author
add_heading(doc, '来自作者', 1)
add_quote_box(doc, "写完这本手册，我重新看了一遍自己这段时间的工作记录，发现一个有点讽刺的事实：我在写'打断从来不是意外'这一章的时候，被自己的手机打断了至少六次。我知道这个道理，但知道和做到之间，隔着的从来不是一份认知，是一次又一次具体的、当下就要做出的选择。")

# CTA
add_para(doc, '', size=16)
add_para(doc, '这不是一门让你"感觉良好"的课', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(198, 40, 40))
add_para(doc, '这是一门逼你面对自己的课', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(198, 40, 40))
add_para(doc, '但只有这样，才能真正改变', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(198, 40, 40))

# Save
output_path = os.path.join(output_base, r"10-对外宣传文案\认证班宣传文案-深度工作主权.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Created: {output_path}")

print("=" * 50)
print("Creating 公开课宣传文案...")
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title
add_para(doc, '', size=24)
add_para(doc, '深度工作主权手册', size=32, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '公开课', size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(46, 84, 150))
add_para(doc, '', size=16)
add_para(doc, '当注意力成为最稀缺的资源，你选择被动失守还是主动主权？', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '', size=20)

# Problem hook
add_heading(doc, '你每天都在经历的', 1, color=(198, 40, 40))
add_para(doc, '"忙了一整天，却感觉什么都没做"', bold=True)
add_para(doc, '会议、消息、协作需求……你的时间被切割成无数碎片，深度思考成为奢侈品。')
add_para(doc, '"各种方法都试过，坚持不下来"', bold=True)
add_para(doc, '番茄钟、时间管理App、各种待办清单——工具换了一个又一个，效果始终无法持续。')
add_para(doc, '"不是不想专注，是真的做不到"', bold=True)
add_para(doc, '团队文化要求随时在线，上级随时可能找你，边界保护说起来容易做起来难。')

# The insight
add_para(doc, '', size=12)
add_para(doc, '这三个感受背后，藏着同一个被忽视的事实——', size=12)
add_para(doc, '', size=8)
add_para(doc, '你不是管理不好时间，你是从来没有真正拥有过注意力的主权。', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))

doc.add_page_break()

# Course positioning
add_heading(doc, '这门课在说什么', 1)
add_para(doc, '这不是一门时间管理课。')
add_para(doc, '时间管理教你怎么排日程、怎么提高效率、怎么更快完成工作。')
add_para(doc, '但如果你每天的时间早就被各种不可拒绝的会议和消息占据得满满当当，时间管理只会让你更焦虑——因为你能管理的，永远只是那些本就不多的边角料。')
add_para(doc, '')
add_para(doc, '这门课要做的，是从更根本的地方入手——', bold=True)
add_para(doc, '重建你对"打断"的认知', size=14, color=(46, 84, 150))
add_para(doc, '看清"忙碌成瘾"背后的心理机制', size=14, color=(46, 84, 150))
add_para(doc, '学会在团队中建立可持续的注意力保护边界', size=14, color=(46, 84, 150))

# Key quotes
add_heading(doc, '课程金句', 1)
add_quote_box(doc, "边界不用解释，只需要反复出现，出现的次数比解释更有说服力。")
add_quote_box(doc, "切换不是免费的，你只是没看到账单。")
add_quote_box(doc, "靠自制力守住的边界，迟早会在你最累的那天失守。")
add_quote_box(doc, "恢复不是躺平，是把掏空的认知资源主动存回去。")
add_quote_box(doc, "你越想成为一个永远在线的人，你能做的事情就越少。")

# What you'll learn
add_heading(doc, '你将学到', 1)
headers = ['章节', '核心议题', '实用工具']
rows = [
    ['理解打断', '注意力是主动让渡的', '中断自我检测表'],
    ['识别忙碌', '忙碌是一种社交货币', '忙碌成瘾诊断表'],
    ['设计时间盒', '没有退出成本的边界等于没有边界', '时间盒设计框架'],
    ['公开边界', '日历标注是对外的边界声明', '深度工作计划表'],
    ['审计会议', '大多数会议从未被授权存在过', '会议决策筛查表'],
    ['降低重启税', '多任务处理在反复缴纳隐藏成本', '重启成本自测'],
    ['设计环境', '专注力是环境设计游戏', '环境隔离自检表'],
    ['主动恢复', '恢复是入场券不是奖励', '分层恢复练习库'],
]
create_table(doc, headers, rows)

# Who is this for
add_heading(doc, '适合人群', 1)
add_bullet(doc, '知识工作者：产品、研发、咨询、数据等需要深度思考的岗位')
add_bullet(doc, '管理者：想要保护自己的深度时间，同时带动团队提升专注力')
add_bullet(doc, '被"忙碌"困扰的人：感觉自己很忙但说不出深度做成了什么')
add_bullet(doc, '尝试过各种时间管理方法但效果不持久的人')

# Course details
add_heading(doc, '课程信息', 1)
headers = ['项目', '内容']
rows = [
    ['时长', '1天精华版（6小时）或 2天完整版（12小时）'],
    ['人数', '25-40人/班'],
    ['形式', '讲授 + 小组讨论 + 工具演练 + 行动承诺'],
    ['工具', '每个章节配套可立即使用的实操工具'],
    ['后续', '加入学员群，持续交流实践心得'],
]
create_table(doc, headers, rows)

# Author intro
add_heading(doc, '关于作者', 1)
add_para(doc, '罗宏伟', bold=True)
add_para(doc, '深耕知识工作者效能提升领域多年，专注于注意力保护与深度工作能力的系统研究。其课程和著作帮助数千名知识工作者重建了对自己时间和注意力的主权。')

# Final message
add_para(doc, '', size=16)
add_para(doc, '注意力是可以被重新谈判的', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '边界是可以被设计出代价的', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '恢复是可以被主动安排的', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '', size=12)
add_para(doc, '这不是一门教你"应该怎么做"的课', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128, 128, 128))
add_para(doc, '这是一门帮你"看清为什么要做"的课', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128, 128, 128))
add_para(doc, '然后，你自己会知道该怎么做', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128, 128, 128))

# Save
output_path = os.path.join(output_base, r"10-对外宣传文案\公开课宣传文案-深度工作主权.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Created: {output_path}")

print("=" * 50)
print("All promotional documents created successfully!")

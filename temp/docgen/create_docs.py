# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

output_base = r"D:\新课开发\工作手册\知识工作者深度工作保护\完整课程包"
os.makedirs(output_base, exist_ok=True)

def set_run_font(run, font_name='Microsoft YaHei', size=12, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, font_name='Microsoft YaHei', size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header background to dark blue
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10)
    return table

print("Creating 课程说明书...")
doc = Document()
# Set default font
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title page
add_para(doc, '', size=24)
add_para(doc, '', size=24)
add_para(doc, '深度工作主权手册', size=32, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '——课程说明书', size=24, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '', size=16)
add_para(doc, '课程名称：深度工作主权手册——知识工作者的注意力保护与恢复系统', size=14)
add_para(doc, '课程编号：深度工作-知识工作者-01', size=14)
add_para(doc, '版本：V1.0', size=14)
add_para(doc, '编制日期：2026年7月', size=14)
add_para(doc, '作者：罗宏伟', size=14)
add_para(doc, '', size=14)
add_para(doc, '适用对象：企业内训/认证班/公开课', size=14)
add_para(doc, '预设时长：2天（每天6小时，共12小时）', size=14)
add_para(doc, '学员规模：25-40人', size=14)
doc.add_page_break()

# Content
add_heading(doc, '第一章 课程概述与定位', 1)
add_heading(doc, '课程背景', 2)
add_para(doc, '在知识经济时代，深度工作能力已成为知识工作者的核心竞争力。然而，随着即时通讯工具的普及和协作强度的提升，知识工作者正面临前所未有的注意力危机——每天被频繁打断、深度思考时间被蚕食、认知资源持续透支。')
add_para(doc, '大多数时间管理课程教授的是工具和技巧，但忽略了最根本的问题：注意力不是被"管理"走的，是被我们自己一次次"允许"拿走的。这一认知重构，是本课程的独特起点。')
add_heading(doc, '课程定位', 2)
add_para(doc, '深度工作主权手册不是一门时间管理课，是一门关于注意力主权认知重建的课程。', bold=True)
add_para(doc, '本课程的核心定位：')
add_bullet(doc, '认知层：重建对"打断"的认知框架——从外部干扰到主动让渡')
add_bullet(doc, '实操层：掌握可即刻落地的时间保护和注意力管理工具')
add_bullet(doc, '协作层：从个人边界到团队契约，构建可持续的深度工作环境')
add_bullet(doc, '心态层：建立长期主义的注意力保护信念')

add_heading(doc, '第二章 目标学员画像', 1)
add_para(doc, '本课程面向以下人群：')
add_heading(doc, '核心学员', 2)
add_bullet(doc, '企业知识工作者：产品经理、研发人员、咨询顾问、数据分析师等需要长时间深度思考的岗位')
add_bullet(doc, '企业中高层管理者：需要保护自己的深度决策时间，同时带领团队建立深度工作文化')
add_bullet(doc, '创业者和自由职业者：需要极高的自我管理能力和注意力保护策略')
add_heading(doc, '学员痛点', 2)
add_bullet(doc, '每天被各种消息、会议、协作需求切割得支离破碎')
add_bullet(doc, '感觉自己很忙，但年底说不清楚深度做成了什么')
add_bullet(doc, '尝试过各种时间管理方法，但坚持不下来')
add_bullet(doc, '想要保护自己的深度工作时间，但不知道如何与团队和上级沟通')

add_heading(doc, '第三章 核心公理与卖点', 1)
add_heading(doc, '三大公理', 2)
add_para(doc, '公理一：打断从来不是意外，是一场你参与签署的交易', bold=True)
add_para(doc, '注意力是主动让渡的，不是被动夺走的。每一次随手点开的提示，都是一次主权的转让。')
add_para(doc, '公理二：忙碌是一种可以上瘾的社交货币，背后是"被需要感"的隐性成本', bold=True)
add_para(doc, '忙碌成瘾的两层放大器：个人层面的间歇性强化 + 组织层面的激励结构错位。')
add_para(doc, '公理三：深度工作不是天赋，是可以被设计的系统能力', bold=True)
add_para(doc, '靠自制力守住的边界，迟早会在你最累的那天失守。真正的稳定，来自环境设计。')

add_heading(doc, '课程卖点', 2)
add_bullet(doc, '认知重建先行：不急于给工具，先帮助学员看清自己行为模式背后的心理机制')
add_bullet(doc, '体系完整：从个人到团队，从技巧到心态，覆盖深度工作保护的完整链条')
add_bullet(doc, '工具即战力：每一个工具都有具体的操作模板，可直接应用于实际工作')
add_bullet(doc, '金句驱动：每章都有经过验证、可以直接引用的核心观点')

add_heading(doc, '第四章 课程目标', 1)
add_heading(doc, '知识目标', 2)
add_bullet(doc, '理解"打断是主动让渡"这一核心公理及其心理机制')
add_bullet(doc, '识别忙碌成瘾的心理机制，认识"被需要感"的隐性成本')
add_bullet(doc, '理解重启税的概念及其量化方法')
add_bullet(doc, '理解团队契约与个人边界的本质区别')
add_heading(doc, '技能目标', 2)
add_bullet(doc, '能够设计带有退出成本的时间盒，防止时间盒被自己打破')
add_bullet(doc, '能够制作让边界真正被看见的深度工作计划表')
add_bullet(doc, '能够使用会议审计框架评估并精简团队会议')
add_bullet(doc, '能够设计并主持团队深度工作契约的讨论会议')
add_heading(doc, '态度目标', 2)
add_bullet(doc, '建立"恢复是入场券"的长期信念')
add_bullet(doc, '理解长期主义的职业身份取舍')
add_bullet(doc, '愿意在团队中做第一个示范者')

add_heading(doc, '第五章 16章内容模块', 1)
headers = ['模块', '章节', '核心议题', '关键工具']
rows = [
    ['PART 1\n认知层', '第1-5章', '打断、忙碌成瘾、时间盒、计划表、中断日志', '中断自我检测表、时间盒设计框架'],
    ['PART 1\n实操层', '第6-10章', '会议审计、重启税、环境设计、恢复、团队契约', '会议审计框架、重启成本自测、环境隔离自检表'],
    ['PART 2\n协作层', '第11-13章', '向上管理边界、产出可见性、管理者认知负荷', '向上边界话术模板、产出可见性三法'],
    ['PART 2\n心态层', '第14-16章', '团队文化、恢复信念、长期主义', '文化推动三步法、恢复类型决策树'],
]
create_table(doc, headers, rows)

add_heading(doc, '第六章 教学方法论', 1)
add_heading(doc, '核心教学原则', 2)
add_bullet(doc, '认知先行，工具在后：每次工具引入前，先让学员"自己发现"这个工具要解决的问题')
add_bullet(doc, '具身认知，体验优先：学员需要在身体和感受层面体验到"切换成本是什么感觉"')
add_bullet(doc, '同伴学习，群体共振：鼓励学员之间的真实对话和经验分享')
add_bullet(doc, '行动导向，当场落地：每一章结尾必须有清晰的"下一步行动"')

add_heading(doc, '第七章 课时安排', 1)
headers = ['模块', '内容', '时长', '占比']
rows = [
    ['开场', '课程介绍、公理阐述、学员期待确认', '45分钟', '6%'],
    ['PART 1-上', '第1-5章：认知层', '3小时', '25%'],
    ['休息', '茶歇+学员交流', '30分钟', '4%'],
    ['PART 1-下', '第6-10章：实操层', '3小时', '25%'],
    ['第一天收尾', '当日复盘、个人行动计划初步制定', '45分钟', '6%'],
    ['第二天开场', '回顾与答疑', '30分钟', '4%'],
    ['PART 2-上', '第11-13章：协作层', '2.5小时', '21%'],
    ['PART 2-下', '第14-16章：心态层', '2小时', '17%'],
    ['课程收尾', '整体复盘、行动承诺、后续资源', '30分钟', '4%'],
]
create_table(doc, headers, rows)

add_heading(doc, '第八章 预期成果', 1)
add_para(doc, '学员离开课程时，将带走：')
add_bullet(doc, '一套完整的个人深度工作保护方案（时间盒+计划表+中断日志）')
add_bullet(doc, '一份可供团队讨论的深度工作契约草案')
add_bullet(doc, '一个持续践行的同伴支持小组')
add_para(doc, '课程结束后30天、90天的跟踪回访，确保行为改变真正发生。')

# Save
output_path = os.path.join(output_base, r"01-课程说明书\课程说明书-深度工作主权V1.0.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Created: {output_path}")

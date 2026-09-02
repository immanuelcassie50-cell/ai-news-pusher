# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

output_base = r"D:\新课开发\工作手册\知识工作者深度工作保护\完整课程包"
os.makedirs(output_base, exist_ok=True)

def set_run_font(run, font_name='Microsoft YaHei', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(31, 56, 100)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, font_name='Microsoft YaHei', size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10)
    return p

def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(f"{num}. {text}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    return p

def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(9)
    return table

def add_box_text(doc, text, bg_color='E8F5E9'):
    """Add a styled box paragraph"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), bg_color)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(10)
    return p

print("=" * 50)
print("Creating 学员手册...")
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title
add_para(doc, '', size=20)
add_para(doc, '深度工作主权手册', size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '——学员手册', size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '', size=12)
add_para(doc, '版本：V1.0', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '2026年7月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# Learning objectives
add_heading(doc, '学习目标', 1)
add_para(doc, '完成本课程后，您将能够：')
add_bullet(doc, '理解"打断是主动让渡"这一核心公理及其心理机制')
add_bullet(doc, '识别忙碌成瘾的心理机制，认识"被需要感"的隐性成本')
add_bullet(doc, '掌握时间盒设计、会议审计、恢复练习等实操工具')
add_bullet(doc, '能够针对自己的实际场景设计并执行深度工作保护方案')
add_bullet(doc, '理解团队契约与个人边界的本质区别')

# Course structure
add_heading(doc, '课程结构', 1)
add_para(doc, '本课程分为两个部分，共16章：')
add_heading(doc, 'PART 1：认知层至实操层（第1-10章）', 2)
add_bullet(doc, '第1-5章：打断、忙碌成瘾、时间盒、计划表、中断日志')
add_bullet(doc, '第6-10章：会议审计、重启税、环境设计、恢复、团队契约')
add_heading(doc, 'PART 2：协作层至心态层（第11-16章）', 2)
add_bullet(doc, '第11-13章：向上管理边界、产出可见性、管理者认知负荷')
add_bullet(doc, '第14-16章：团队文化、恢复信念、长期主义')

# Chapter summaries
add_heading(doc, '各章学习要点', 1)

chapters = [
    ("第1章：打断从来不是意外", [
        "注意力是主动让渡的，不是被动夺走的",
        "三大心理机制：即时满足偏好、风险规避、不确定性驱动检查",
        "工具：中断自我检测表"
    ]),
    ("第2章：忙碌是一种可以上瘾的社交货币", [
        "忙碌成瘾的两层放大器：个人间歇性强化 + 组织激励错位",
        "靠响应速度建立的口碑容易复制，靠判断质量建立的壁垒难以超越",
        "工具：忙碌成瘾自我诊断表"
    ]),
    ("第3章：时间盒失败是因为没有退出成本", [
        "时间盒设计关键：不在划定时间，在设计违反的代价",
        "退出成本设计：让发起者自问'晚两小时处理会有什么损失'",
        "工具：深度工作时间盒设计框架"
    ]),
    ("第4章：深度工作计划表是写给别人看的边界声明", [
        "口头拒绝 vs 日历标注：后者可持续，前者每次重复",
        "共享日历标注 = 公共信息，对他人有约束力",
        "工具：深度工作计划表框架"
    ]),
    ("第5章：中断日志记录的是注意力被谁定价了", [
        "中断来源反映团队内部的沟通习惯和权力结构",
        "连续记录5个工作日，看清打断模式",
        "工具：中断日志记录框架"
    ]),
    ("第6章：大多数会议从未被授权存在过", [
        "发起会议成本几乎为零，被邀请方成本不对称",
        "会议审计框架：必要性筛查、参会人筛查、成本可见化",
        "工具：会议决策筛查表"
    ]),
    ("第7章：多任务处理是在反复缴纳重启税", [
        "每次切换需要8-12分钟重启，每天约1.5小时用于重启",
        "注意力残留：任务切换后大脑无法完全清空上一个任务的注意力",
        "工具：重启成本量化自测"
    ]),
    ("第8章：专注力是环境设计游戏", [
        "意志力在会议后、棘手事后、情绪低落时天然更低",
        "三层面环境设计：物理、数字、空间",
        "工具：环境隔离自检表"
    ]),
    ("第9章：恢复不是躺平，是主动存回认知资源", [
        "被动恢复（刷手机）vs 主动恢复（散步、冥想）效果截然不同",
        "恢复是入场券，不是奖励",
        "工具：分层恢复练习库"
    ]),
    ("第10章：团队契约保护的是共同信任", [
        "契约把违反行为从零成本变成有一定成本",
        "契约制定三步：收集困扰→形成条款→公开确认",
        "工具：团队深度工作契约框架"
    ]),
    ("第11章：向上管理边界是在教会领导定价你的时间", [
        "口头同意 ≠ 行为承诺，边界需要反复出现重校预期",
        "沟通打开理解入口，行为让边界真正稳固",
        "工具：向上边界设定话术模板"
    ]),
    ("第12章：护住的时间没有产出证据会被第一个收回", [
        "深度工作产出周期长，资源紧张时容易被第一个收回",
        "展示产出 ≠ 邀功，是让真实存在的价值被看见",
        "工具：产出可见性设计三法"
    ]),
    ("第13章：管理者的认知负荷来自没被承认的重启成本", [
        "视角切换幅度越大，重启成本越高",
        "日程设计三原则：留缓冲、聚类相近事项、主动申请排布",
        "工具：缓冲块标注法"
    ]),
    ("第14章：团队深度工作文化是从一次公开拒绝开始的", [
        "先例效应：第一个偏离者承担最高风险，后续跟随者门槛降低",
        "身份带来的信任背书，是示范能否被效仿的重要变量",
        "工具：推动文化转变三步法"
    ]),
    ("第15章：恢复练习是下一次深度工作的入场券", [
        "透支的隐性成本：低质量产出在后续环节暴露，返工代价更高",
        "按疲惫类型选择恢复方式：认知疲惫/决策疲惫/情绪疲惫",
        "工具：恢复方式选择决策树"
    ]),
    ("第16章：长期主义者最先放弃的是'随时可用'这个人设", [
        "响应速度是短期诱惑，判断质量是长期价值",
        "职业早期靠速度积累口碑，中后期靠判断质量建立壁垒",
        "工具：响应速度调整三步法"
    ]),
]

for chapter_title, points in chapters:
    add_heading(doc, chapter_title, 2, color=(46, 84, 150))
    for point in points:
        add_bullet(doc, point)

# Exercises section
add_heading(doc, '练习与实践', 1)
add_heading(doc, '练习1：中断日志记录（5天）', 2)
add_para(doc, '记录字段：时间点、打断来源、打断理由、处理耗时、事后判断')
headers = ['日期', '时间', '来源', '理由', '耗时', '真的紧急？']
rows = [
    ['周一', '', '', '', '', ''],
    ['周二', '', '', '', '', ''],
    ['周三', '', '', '', '', ''],
    ['周四', '', '', '', '', ''],
    ['周五', '', '', '', '', ''],
]
create_table(doc, headers, rows)

add_heading(doc, '练习2：时间盒设计', 2)
add_para(doc, '设计一个属于你自己的深度工作时间盒：')
add_para(doc, '时间：___________')
add_para(doc, '退出成本：___________')
add_para(doc, '例外声明：___________')

add_heading(doc, '练习3：团队契约草案', 2)
add_para(doc, '共同时段：___________')
add_para(doc, '响应预期：___________')
add_para(doc, '例外处理：___________')

# Notes section
add_heading(doc, '笔记区', 1)
for i in range(8):
    add_para(doc, '_______________________________________________________________')
    add_para(doc, '_______________________________________________________________')

# Save
output_path = os.path.join(output_base, r"05-学员手册\学员手册-深度工作主权.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Created: {output_path}")

print("=" * 50)
print("Creating 讲师手册...")
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

# Title
add_para(doc, '', size=20)
add_para(doc, '深度工作主权手册', size=28, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '——讲师手册', size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=(31, 56, 100))
add_para(doc, '', size=12)
add_para(doc, '版本：V1.0', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '2026年7月', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# Instructor qualifications
add_heading(doc, '讲师资质要求', 1)
add_bullet(doc, '深入理解深度工作的理论基础，最好有丰富的实践经验')
add_bullet(doc, '具备引导团队讨论的经验，能够带动学员之间的真实对话')
add_bullet(doc, '熟悉知识工作者的工作场景，能够举例说明各行业的应用')
add_bullet(doc, '参加过本课程的完整学习，或通过认证讲师培训')

# Preparation checklist
add_heading(doc, '课前准备清单', 1)
headers = ['检查项', '状态', '备注']
rows = [
    ['确认场地投影、音响设备正常', '', ''],
    ['准备足够的学员手册', '', ''],
    ['打印好评 Foss测试表、忙碌成瘾诊断表', '', ''],
    ['准备大白纸和马克笔', '', ''],
    ['确认茶歇安排', '', ''],
    ['提前通知学员带手机/电脑', '', ''],
    ['学员课前记录1-2天中断日志', '', ''],
]
create_table(doc, headers, rows)

# Teaching flow
add_heading(doc, '教学流程总览', 1)
headers = ['模块', '时长', '核心内容', '教学方式']
rows = [
    ['开场', '45分钟', '课程介绍、公理阐述', '讲授+学员自我介绍'],
    ['PART 1-上', '3小时', '第1-5章', '讲授+小组讨论+练习'],
    ['茶歇', '30分钟', '', ''],
    ['PART 1-下', '3小时', '第6-10章', '讲授+小组讨论+练习'],
    ['第一天收尾', '45分钟', '复盘+行动计划', '小组分享'],
    ['第二天开场', '30分钟', '回顾+答疑', '问答'],
    ['PART 2-上', '2.5小时', '第11-13章', '讲授+小组讨论'],
    ['PART 2-下', '2小时', '第14-16章', '讲授+小组讨论'],
    ['课程收尾', '30分钟', '整体复盘+行动承诺', '个人分享'],
]
create_table(doc, headers, rows)

# Time allocation per chapter
add_heading(doc, '各章时间分配', 1)
headers = ['章节', '建议时间', '必须环节', '可选环节']
rows = [
    ['第1章：打断', '45分钟', '三大机制讲解+自测表练习', '体验活动'],
    ['第2章：忙碌成瘾', '45分钟', '诊断表+破瘾策略讨论', '学员分享'],
    ['第3章：时间盒', '45分钟', '框架讲解+设计练习', '小组互评'],
    ['第4章：计划表', '30分钟', '对比表格+实操演练', ''],
    ['第5章：中断日志', '45分钟', '框架讲解+任务布置', '小组讨论'],
    ['第6章：会议', '45分钟', '审计框架+筛查表练习', '数据展示'],
    ['第7章：重启税', '45分钟', '量化自测+批量切换策略', '闭眼数数体验'],
    ['第8章：环境设计', '45分钟', '三层面讲解+自检表', '改造计划制定'],
    ['第9章：恢复', '45分钟', '主动vs被动对比+决策树', '恢复方式选择'],
    ['第10章：团队契约', '45分钟', '契约框架+制定讨论', '角色扮演'],
    ['第11章：向上管理', '45分钟', '预期重校机制+话术练习', '时机讨论'],
    ['第12章：产出可见性', '30分钟', '三法讲解+进展说明练习', ''],
    ['第13章：认知负荷', '30分钟', '三原则+缓冲块练习', ''],
    ['第14章：团队文化', '30分钟', '先例效应+三步法', '示范者讨论'],
    ['第15章：恢复信念', '30分钟', '因果颠倒揭示+决策树', '个人反思'],
    ['第16章：长期主义', '30分钟', '两阶段策略+三步法', '职业身份反思'],
]
create_table(doc, headers, rows)

# FAQ
add_heading(doc, '常见问题应对', 1)
faqs = [
    ('"我知道但做不到"怎么办？', '知道和做到之间隔着的是具体选择。让学员描述做不到的具体场景，把问题具体化而非泛泛说"意志力不够"。'),
    ('"这不适合我的情况"怎么办？', '追问具体：请学员描述他们公司文化。寻找切入点：几乎每个环境都有可以利用的裂缝。提供最小可行方案。'),
    ('学员沉默、参与度低', '小组讨论代替全体讨论；书面匿名反馈；讲师先示范一个自己的"失败案例"。'),
    ('有人觉得你在刁难他', '用数据引发反思而非指出个人问题。第一次推行时从团队整体数据入手。'),
    ('时间不够用怎么办', '优先保证"行动承诺"环节。可以把部分练习留作课后作业。'),
]
for q, a in faqs:
    add_heading(doc, q, 2, color=(198, 40, 40))
    add_para(doc, a)

# Assessment criteria
add_heading(doc, '评估方式', 1)
add_heading(doc, '过程性评估', 2)
add_bullet(doc, '参与度：观察学员在讨论、练习中的投入程度')
add_bullet(doc, '行动承诺：记录学员在各章节做出的行动承诺')
add_bullet(doc, '同伴反馈：小组内成员对彼此行动计划的反馈质量')
add_heading(doc, '终结性评估', 2)
add_bullet(doc, '课程结束时简答题测试（可选）')
add_bullet(doc, '课程结束后1个月、3个月的跟踪回访')
add_bullet(doc, '学员提交自己的时间盒记录、会议审计记录等')

# Save
output_path = os.path.join(output_base, r"06-讲师手册\讲师手册-深度工作主权.docx")
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Created: {output_path}")

print("=" * 50)
print("All documents created successfully!")

# -*- coding: utf-8 -*-
"""
创建民族主义思想史课程大纲Word文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "D:/新课开发/政治学/18_民族主义思想史-一个概念如何塑造现代世界/课程大纲/01_课程大纲.docx"


def set_run_font(run, ascii_font="微软雅黑", east_asia_font="微软雅黑", size=12, bold=False, color=None):
    """设置run的字体格式"""
    run.font.name = ascii_font
    run._r.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_paragraph(doc, text, level=1, size=16, bold=True, color=None, space_before=12, space_after=6):
    """添加标题段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return para


def add_body_paragraph(doc, text, size=12, space_before=0, space_after=6, indent=False):
    """添加正文段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.first_line_indent = Pt(24)
    run = para.add_run(text)
    set_run_font(run, size=size)
    return para


def add_bullet_item(doc, text, size=12):
    """添加项目符号项"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run_font(run, size=size)
    return para


def create_table_row(table, col1_text, col2_text, col1_width=None, col2_width=None, header=False):
    """创建表格行"""
    row = table.add_row()
    cell1 = row.cells[0]
    cell2 = row.cells[1]

    cell1.text = col1_text
    cell2.text = col2_text

    for cell in [cell1, cell2]:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "微软雅黑"
                run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
                run.font.size = Pt(11)
                if header:
                    run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after = Pt(4)

    if col1_width:
        cell1.width = col1_width
    if col2_width:
        cell2.width = col2_width

    return row


def add_module_section(doc, module_num, title, goals, contents, duration):
    """添加模块章节"""
    # 模块标题
    heading = doc.add_heading(f"模块{module_num}：{title}", level=2)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 73, 125)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    # 学习目标
    add_heading_paragraph(doc, "学习目标", level=3, size=12, bold=True, color=RGBColor(0, 0, 0), space_before=8, space_after=4)
    for goal in goals:
        add_bullet_item(doc, goal, size=11)

    # 核心内容
    add_heading_paragraph(doc, "核心内容", level=3, size=12, bold=True, color=RGBColor(0, 0, 0), space_before=8, space_after=4)
    for content in contents:
        add_bullet_item(doc, content, size=11)

    # 课时安排
    add_heading_paragraph(doc, f"课时安排：{duration}", level=3, size=12, bold=False, color=RGBColor(89, 89, 89), space_before=8, space_after=8)


def create_document():
    """创建完整的课程大纲文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4宽度
    section.page_height = Inches(11.69)  # A4高度
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ===== 封面 =====
    # 课程名称
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_para.paragraph_format.space_after = Pt(24)
    run = title_para.add_run("民族主义思想史")
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    # 副标题
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(0)
    subtitle_para.paragraph_format.space_after = Pt(48)
    run = subtitle_para.add_run("——一个概念如何塑造现代世界")
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(64, 64, 64)

    # 分隔线
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(48)
    run = line_para.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 73, 125)

    # 基本信息表格
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("课程编号", "政治学类-18"),
        ("目标受众", "本科生、研究生及对中国政治感兴趣的学者"),
        ("学习时长", "6讲 × 3小时 = 18小时"),
        ("课程类型", "理论+实践型课程")
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "微软雅黑"
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
                    run.font.size = Pt(12)
                if label:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3.5)

    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(6)
            cell.paragraphs[0].paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== 课程介绍 =====
    add_heading_paragraph(doc, "课程介绍", level=1, size=18, bold=True, color=RGBColor(31, 73, 125), space_before=0, space_after=12)

    add_heading_paragraph(doc, "课程简介", level=2, size=14, bold=True, space_before=12, space_after=6)
    intro_text = """民族主义是现代世界最强大的政治力量之一。从法国大革命的公民认同，到今天全球范围内的民族主义浪潮，这个"想象的共同体"如何从欧洲起源，进而影响整个世界？本课程将带你深入探索民族主义的思想脉络，解析其背后的政治逻辑，审视当代民族主义复兴现象，并特别关注中国情境下的民族主义特征。"""
    add_body_paragraph(doc, intro_text, size=12, indent=True)

    add_heading_paragraph(doc, "解决什么问题", level=2, size=14, bold=True, space_before=16, space_after=6)
    problems = [
        "什么是民族主义？它与爱国主义、民族认同有何区别？",
        "为什么民族主义在20世纪末'终结'后卷土重来？",
        "西方民族主义理论能否解释中国的民族主义现象？",
        "在民族主义情感高涨的时代，如何保持理性思考？"
    ]
    for p in problems:
        add_bullet_item(doc, p, size=12)

    add_heading_paragraph(doc, "学员收获", level=2, size=14, bold=True, space_before=16, space_after=6)
    gains = [
        "建立对民族主义的系统性认知框架",
        "理解民族主义在现代世界形成中的关键作用",
        "掌握分析当代民族主义现象的理论工具",
        "形成对中国民族主义特殊性的深入认识",
        "提升在多元文化环境中理性对话的能力"
    ]
    for g in gains:
        add_bullet_item(doc, g, size=12)

    doc.add_page_break()

    # ===== 课程大纲 =====
    add_heading_paragraph(doc, "课程大纲", level=1, size=18, bold=True, color=RGBColor(31, 73, 125), space_before=0, space_after=12)

    # 模块1
    add_module_section(
        doc, 1,
        "概念解码——民族主义到底是什么",
        [
            "理解'民族'、'民族主义'、'爱国主义'的概念边界",
            "掌握民族主义的核心要素：共同历史、领土、语言、文化",
            "区分民族主义的不同定义路径：客观派 vs 主观派"
        ],
        [
            "民族的概念史：从古代到现代的演变",
            "民族主义的定义：安东尼·史密斯 vs 本尼迪克特·安德森",
            "民族主义与爱国主义、种族主义的关系辨析",
            "民族主义的类型学概述"
        ],
        "3小时"
    )

    # 模块2
    add_module_section(
        doc, 2,
        "类型辨析——两种民族主义的制度后果",
        [
            "理解公民民族主义与族群民族主义的本质区别",
            "掌握两种民族主义对国家建构的不同影响",
            "认识民族主义与自由主义、威权主义的关系"
        ],
        [
            "公民民族主义：法国模式与美国模式",
            "族群民族主义：德国模式与中东欧模式",
            "两种民族主义的制度化路径比较",
            "民族主义与民主转型：成功与失败的案例"
        ],
        "3小时"
    )

    # 模块3
    add_module_section(
        doc, 3,
        "历史溯源——民族主义如何诞生并塑造现代世界",
        [
            "理解民族主义在欧洲诞生的历史背景",
            "掌握民族主义塑造现代国家的三条路径",
            "认识民族主义在全球扩散的过程与机制"
        ],
        [
            "民族主义的欧洲起源：经济、政治、文化因素",
            "民族国家体系的确立：从威斯特伐利亚到一战",
            "民族主义的全球扩散：殖民主义、冷战与全球化",
            "民族主义与两次世界大战的关联"
        ],
        "3小时"
    )

    # 模块4
    add_module_section(
        doc, 4,
        "当代诊断——今天的民族主义为何卷土重来",
        [
            "理解20世纪末'民族主义终结'论的局限性",
            "掌握当代民族主义复兴的主要原因",
            "认识新媒体时代民族主义的新特征"
        ],
        [
            "冷战终结后的民族主义回潮",
            "经济全球化与民族主义的反弹",
            "移民问题、文化焦虑与身份政治",
            "数字时代民族主义的传播机制",
            "案例分析：英国脱欧、美国优先、民粹主义浪潮"
        ],
        "3小时"
    )

    # 模块5
    add_module_section(
        doc, 5,
        "中国的民族主义——特殊还是普遍？",
        [
            "理解中国民族主义的历史形成过程",
            "掌握中国民族主义的特殊性表现",
            "认识中国民族主义与外交政策的关系"
        ],
        [
            "中国民族主义的历史根源：从清朝到民国",
            "革命民族主义与后革命时代的民族主义",
            "中国民族主义的特殊性：文明国家 vs 民族国家",
            "民族主义与中国的崛起：机遇与挑战",
            "中国民族主义的媒介表达与公众认知"
        ],
        "3小时"
    )

    # 模块6
    add_module_section(
        doc, 6,
        "实践应用——如何与民族主义相处",
        [
            "掌握在民族主义环境中保持理性思考的方法",
            "理解多元文化主义政策的争议与局限",
            "形成建设性的跨文化对话能力"
        ],
        [
            "民族主义时代的公民素养",
            "多元文化主义的兴衰与反思",
            "跨国主义与超国家认同的可能性",
            "建设性民族主义的理论与实践",
            "角色扮演：跨文化对话情境模拟"
        ],
        "3小时"
    )

    doc.add_page_break()

    # ===== 教学方式 =====
    add_heading_paragraph(doc, "教学方式", level=1, size=18, bold=True, color=RGBColor(31, 73, 125), space_before=0, space_after=12)

    methods = [
        ("理论讲授", "系统讲解民族主义的核心概念与理论框架，穿插学术前沿动态，帮助学员建立扎实的理论基础。"),
        ("案例分析", "选取法国大革命、1848年欧洲革命、纳粹主义、中国民族主义等历史案例，深入剖析民族主义的具体表现。"),
        ("互动讨论", "围绕当代民族主义热点话题展开小组讨论，鼓励学员发表观点，培养批判性思维。"),
        ("工具演练", "教授文献检索、话语分析、框架分析等研究工具，让学员掌握分析民族主义现象的实操技能。")
    ]

    for title, desc in methods:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run1 = para.add_run(f"• {title}：")
        run1.font.name = "微软雅黑"
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run1.font.size = Pt(12)
        run1.font.bold = True

        run2 = para.add_run(desc)
        run2.font.name = "微软雅黑"
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run2.font.size = Pt(12)

    doc.add_page_break()

    # ===== 配套材料 =====
    add_heading_paragraph(doc, "配套材料", level=1, size=18, bold=True, color=RGBColor(31, 73, 125), space_before=0, space_after=12)

    materials = [
        ("学员手册", "包含课程讲义、阅读材料清单、思考题集，帮助学员系统复习课程内容。"),
        ("讲师手册", "提供教学PPT、案例库、讨论引导指南，支持讲师高效备课与授课。"),
        ("练习题库", "涵盖名词解释、简答题、论述题、分析题等多种题型，支持形成性与总结性评估。"),
        ("评估工具", "包括课程评估问卷、学习效果测评表、同伴互评表等，确保教学质量的持续改进。")
    ]

    for title, desc in materials:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run1 = para.add_run(f"• {title}：")
        run1.font.name = "微软雅黑"
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run1.font.size = Pt(12)
        run1.font.bold = True

        run2 = para.add_run(desc)
        run2.font.name = "微软雅黑"
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        run2.font.size = Pt(12)

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    create_document()

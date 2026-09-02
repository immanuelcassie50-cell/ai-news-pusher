# -*- coding: utf-8 -*-
"""生成工具表单 xlsx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False, color=None):
    """设置run的字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, font_size=18, color=(43, 45, 66)):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=True, color=color)
    return p

def add_subheading(doc, text, font_size=14, color=(42, 45, 66)):
    """添加副标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=True, color=color)
    return p

def add_paragraph(doc, text, font_size=11, indent=False, space_after=8):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size)
    return p

def add_card(doc, title, content_items, accent_color=(239, 35, 60)):
    """添加卡片式内容块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(title)
    set_run_font(run, font_size=12, bold=True, color=accent_color)

    for item in content_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(4)
        p_item.paragraph_format.left_indent = Pt(20)
        run = p_item.add_run(item)
        set_run_font(run, font_size=10)
    return p

def create_tools_doc(output_path):
    """创建工具表单文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ========== 封面 ==========
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(40)
    p_title.paragraph_format.space_after = Pt(20)
    run = p_title.add_run("手册进化 工具表单")
    set_run_font(run, font_size=32, bold=True, color=(43, 45, 66))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(30)
    run = p_sub.add_run('从"阅读版"到"执行手册" 配套工具')
    set_run_font(run, font_size=16, color=(141, 153, 174))

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(60)
    run = p_desc.add_run('本工具包包含六大模块的空白模板，可直接打印使用')
    set_run_font(run, font_size=11, color=(100, 100, 100))

    # ========== 工具目录 ==========
    add_heading(doc, "工具目录", font_size=18)

    tools = [
        ('工具一', '场景节点链写作模板', '将场景描述转化为可照做的节点链'),
        ('工具二', '标准动作清单模板', '把模糊流程写成可打卡清单'),
        ('工具三', '判断标准卡模板', '在关键节点给出判断条件'),
        ('工具四', '分级处置表模板', '按严重程度分级对应处置动作'),
        ('工具五', '情境案例写作模板', '用真实场景串起判断逻辑'),
        ('工具六', '高频问答分类框架', '覆盖真正会卡壳的现场问题'),
        ('工具七', '手册转化自检清单', '诊断现有手册缺什么，验收转化完成度')
    ]

    for num, title, desc in tools:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(num + '：' + title)
        set_run_font(run, font_size=12, bold=True, color=(42, 45, 66))
        run = p.add_run(' — ' + desc)
        set_run_font(run, font_size=11)

    doc.add_page_break()

    # =============================================
    # 工具一：场景节点链写作模板
    # =============================================
    add_heading(doc, "工具一：场景节点链写作模板", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：将泛泛的场景描述拆解成可识别的动作节点链。每个节点包含触发条件、动作、结果三个要素。', space_after=12)

    # 模板表格
    table1 = doc.add_table(rows=8, cols=4)
    table1.style = 'Table Grid'

    headers1 = ['节点序号', '触发条件', '动作描述', '结果确认']
    for i, h in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for row_idx in range(1, 8):
        table1.rows[row_idx].cells[0].text = str(row_idx)
        for col_idx in range(1, 4):
            table1.rows[row_idx].cells[col_idx].text = ''

    add_paragraph(doc, '填写说明：', font_size=11)
    add_paragraph(doc, '触发条件：描述在什么情况下执行这个动作（如：到达现场、发现问题、收到通知）', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '动作描述：具体要做什么，用动词开头，可量化（如：检查阀门、开泵、关阀）', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '结果确认：动作执行后预期达到的状态，可验证（如：压力表读数在1.5-2.0区间）', font_size=10, indent=True, space_after=12)

    doc.add_page_break()

    # =============================================
    # 工具二：标准动作清单模板
    # =============================================
    add_heading(doc, "工具二：标准动作清单模板", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：把模糊的流程写成可打卡、可核对的清单结构。每个动作都是可验证的。', space_after=12)

    table2 = doc.add_table(rows=11, cols=5)
    table2.style = 'Table Grid'

    headers2 = ['序号', '检查项', '标准要求', '检查方法', '结果']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for row_idx in range(1, 11):
        table2.rows[row_idx].cells[0].text = str(row_idx)
        for col_idx in range(1, 5):
            table2.rows[row_idx].cells[col_idx].text = ''

    add_paragraph(doc, '填写说明：', font_size=11)
    add_paragraph(doc, '检查项：具体要检查的设备或参数', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '标准要求：可量化的标准（如：温度≤85℃、压力在1.5-2.0MPa）', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '检查方法：怎么检查这个项目（如：查看DCS画面、现场查看、仪器测量）', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '结果：合格打√，不合格打×并备注', font_size=10, indent=True, space_after=12)

    doc.add_page_break()

    # =============================================
    # 工具三：判断标准卡模板
    # =============================================
    add_heading(doc, "工具三：判断标准卡模板", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：在关键节点单独给出判断条件，不让读者自己提炼。', space_after=12)

    # 表格
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = 'Table Grid'

    # 表头
    table3.rows[0].cells[0].text = '判断要素'
    table3.rows[0].cells[1].text = '填写内容'
    for cell in table3.rows[0].cells:
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    items3 = [
        ('条件信号1', '满足此条件时触发判断（如：温度>85℃）'),
        ('条件信号2', '第二个触发条件（如：回流比<1.2）'),
        ('判断结论', '满足条件后的结论（如：冷凝效果下降，需立即处置）'),
        ('对应动作', '结论后的处置动作（①降进料 ②加大冷却 ③观察10分钟）'),
        ('关键变量', '影响判断的核心参数（如：温度阈值、时间窗口）'),
        ('可忽略因素', '不影响判断的因素（如：环境温度变化）')
    ]

    for row_idx, (label, hint) in enumerate(items3, start=1):
        table3.rows[row_idx].cells[0].text = label
        run = table3.rows[row_idx].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        table3.rows[row_idx].cells[1].text = hint

    add_paragraph(doc, '', space_after=8)
    add_paragraph(doc, '判断逻辑说明：', font_size=11)
    add_paragraph(doc, '• 条件信号之间是"AND"还是"OR"关系？', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 判断结论要明确，不能模糊', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 对应动作要可执行，不能是原则性要求', font_size=10, indent=True, space_after=12)

    doc.add_page_break()

    # =============================================
    # 工具四：分级处置表模板
    # =============================================
    add_heading(doc, "工具四：分级处置表模板", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：按严重程度分级，每级对应明确的处置动作。', space_after=12)

    table4 = doc.add_table(rows=5, cols=4)
    table4.style = 'Table Grid'

    headers4 = ['分级', '判断条件', '处置动作', '时限要求']
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    levels = [
        ('轻微', '参数轻微偏移，自行调整可恢复', '降低进料量+加大冷却', '5分钟'),
        ('中度', '参数明显异常，需技术支持', '联系班长远程指导', '立即'),
        ('严重', '参数失控，有安全风险', '立即停机+上报+疏散', '即时')
    ]

    for row_idx, (level, condition, action, time) in enumerate(levels, start=1):
        table4.rows[row_idx].cells[0].text = level
        run = table4.rows[row_idx].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        table4.rows[row_idx].cells[1].text = condition
        table4.rows[row_idx].cells[2].text = action
        table4.rows[row_idx].cells[3].text = time

    add_paragraph(doc, '', space_after=8)
    add_paragraph(doc, '分级原则：', font_size=11)
    add_paragraph(doc, '• 轻微：执行者自己能处理，不升级', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 中度：需要技术支持，但不需要外部救援', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 严重：需要启动应急流程，可能涉及人身安全', font_size=10, indent=True, space_after=12)

    doc.add_page_break()

    # =============================================
    # 工具五：情境案例写作模板
    # =============================================
    add_heading(doc, "工具五：情境案例写作模板", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：用真实场景串起判断逻辑，建立代入感。', space_after=12)

    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'

    table5.rows[0].cells[0].text = '案例要素'
    table5.rows[0].cells[1].text = '填写内容'
    for cell in table5.rows[0].cells:
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    items5 = [
        ('情境背景', '在什么场景下发生（如：反应釜运行2小时后）'),
        ('异常征兆', '出现了什么异常信号（如：塔顶温度升至87℃）'),
        ('判断过程', '执行者如何判断（如：查看回流比、确认冷却水状态）'),
        ('处置动作', '采取了什么处置（如：降低进料量30%）'),
        ('结果复盘', '处置后的结果（如：温度在10分钟后降至83℃）')
    ]

    for row_idx, (label, hint) in enumerate(items5, start=1):
        table5.rows[row_idx].cells[0].text = label
        run = table5.rows[row_idx].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        table5.rows[row_idx].cells[1].text = hint

    add_paragraph(doc, '', space_after=8)
    add_paragraph(doc, '案例编写原则：', font_size=11)
    add_paragraph(doc, '• 真实感：场景要真实具体，有代入感', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 完整性：覆盖判断→决策→执行→结果全流程', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 可借鉴：能让读者遇到类似情况时有所参照', font_size=10, indent=True, space_after=12)

    doc.add_page_break()

    # =============================================
    # 工具六：高频问答分类框架
    # =============================================
    add_heading(doc, "工具六：高频问答分类框架", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：覆盖真正会卡壳的现场问题，不是知识点罗列。', space_after=12)

    table6 = doc.add_table(rows=11, cols=3)
    table6.style = 'Table Grid'

    headers6 = ['问题类型', '典型问题示例', '标准回答要点']
    for i, h in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    q_types = [
        ('边界判断', '这种情况算不算异常？', '明确边界条件，给出判断标准'),
        ('遗漏处理', '没遇到过这种情况怎么办？', '给出升级路径和兜底处理'),
        ('多发异常', '经常出现这个异常是什么原因？', '分析根因，给出预防措施'),
        ('特殊情况', '如果有紧急情况可以跳过吗？', '明确不可跳过的底线'),
        ('状态确认', '怎么判断处理成功了？', '给出可验证的结果标准'),
        ('工具问题', '测量工具不准怎么办？', '给出替代方案和校准方法'),
        ('交接配合', '交接班时发现异常谁负责？', '明确责任转移节点'),
        ('记录要求', '需要记录哪些信息？', '给出记录清单和格式'),
        ('权限问题', '这种情况我能自己决定吗？', '明确授权范围和升级条件'),
        ('后续跟踪', '处理完还需要做什么？', '给出后续跟踪要求')
    ]

    for row_idx, (qtype, example, answer) in enumerate(q_types, start=1):
        table6.rows[row_idx].cells[0].text = qtype
        run = table6.rows[row_idx].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        table6.rows[row_idx].cells[1].text = example
        table6.rows[row_idx].cells[2].text = answer

    doc.add_page_break()

    # =============================================
    # 工具七：手册转化自检清单
    # =============================================
    add_heading(doc, "工具七：手册转化自检清单", font_size=18, color=(239, 35, 60))

    add_paragraph(doc, '用途：诊断现有手册缺什么，验收转化完成度。', space_after=12)

    table7 = doc.add_table(rows=13, cols=3)
    table7.style = 'Table Grid'

    headers7 = ['检查维度', '检查要点', '合格标准']
    for i, h in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    checks = [
        ('场景定位', '是否有清晰的触发条件', '每个场景都能回答"什么时候做"'),
        ('场景定位', '是否有明确的动作节点', '复杂流程已拆解为节点链'),
        ('标准动作', '是否有可验证的结果标准', '结果可观察、可测量'),
        ('标准动作', '是否有明确的时间窗口', '动作有时限要求'),
        ('判断标准卡', '关键节点是否有判断条件', '条件可量化、不模糊'),
        ('判断标准卡', '是否包含"可忽略因素"', '避免过度谨慎'),
        ('分级处置', '是否覆盖所有异常级别', '轻微/中度/严重都有对应处置'),
        ('分级处置', '是否有明确时限要求', '每个级别有时限'),
        ('情境案例', '是否覆盖高频异常场景', '案例来自真实场景'),
        ('情境案例', '是否包含完整判断逻辑', '能引导读者思考'),
        ('高频问答', '是否覆盖真正卡壳的问题', '问题来自一线调研'),
        ('高频问答', '回答是否可操作', '能直接照做')
    ]

    for row_idx, (dim, point, standard) in enumerate(checks, start=1):
        table7.rows[row_idx].cells[0].text = dim
        run = table7.rows[row_idx].cells[0].paragraphs[0].runs[0]
        run.font.bold = True
        table7.rows[row_idx].cells[1].text = point
        table7.rows[row_idx].cells[2].text = standard

    add_paragraph(doc, '', space_after=8)
    add_paragraph(doc, '使用说明：', font_size=11)
    add_paragraph(doc, '• 每个检查要点：合格打√，不合格打×', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 发现×则需要针对性补充/修改', font_size=10, indent=True, space_after=4)
    add_paragraph(doc, '• 全部√方可视为转化完成', font_size=10, indent=True, space_after=12)

    # ========== 结语 ==========
    add_heading(doc, "工具包使用说明", font_size=18)

    usage = [
        '本工具包用于课程练习和实际手册升级项目',
        '建议先在课程中用提供的案例素材练习，再用于实际项目',
        '实际使用时可根据行业特点调整表格内容',
        '工具包可反复打印使用，建议单面打印'
    ]

    for item in usage:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('• ' + item)
        set_run_font(run, font_size=11)

    # 保存文档
    doc.save(output_path)
    print(f'工具表单已保存至: {output_path}')

if __name__ == '__main__':
    output_path = 'D:/新课开发/经验萃取/阅读手册转执行手册/完整课程包/008-工具表单/008-工具表单.docx'
    create_tools_doc(output_path)
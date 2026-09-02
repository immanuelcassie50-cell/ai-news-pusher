# -*- coding: utf-8 -*-
"""
Create: 004-认证班宣传文案-廉政风险情景决策.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "D:/新课开发/党业融合/廉政风险情景决策/完整课程包/004-对外宣传文案/004-认证班宣传文案-廉政风险情景决策.docx"

# Color palette - Professional government/corporate style
NAVY = RGBColor(0x1F, 0x38, 0x64)        # #1F3864 - Primary headings
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)  # #2E75B6 - Secondary elements
BODY_GRAY = RGBColor(0x33, 0x33, 0x33)    # #333333 - Body text
DARK_RED = RGBColor(0xC0, 0x00, 0x00)     # #C00000 - Accent for emphasis

def set_run_font(run, font_name_cn="微软雅黑", font_name_en="Calibri", size=11, bold=False, color=None):
    run.font.name = font_name_en
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading_para(doc, text, level=1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a heading paragraph with custom styling"""
    para = doc.add_paragraph()
    para.alignment = alignment

    if level == 0:  # Title
        run = para.add_run(text)
        set_run_font(run, size=26, bold=True, color=NAVY)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
    elif level == 1:  # H1
        run = para.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    elif level == 2:  # H2
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    return para

def add_body_para(doc, text, indent=False, bullet=False, bold_part=None):
    """Add a body paragraph"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bullet:
        para.paragraph_format.left_indent = Cm(0.5)
        run = para.add_run("● " if not bold_part else "")
        set_run_font(run, size=11, color=BODY_GRAY)

    if bold_part and bold_part in text:
        before, rest = text.split(bold_part, 1)
        if before:
            run = para.add_run(before)
            set_run_font(run, size=11, color=BODY_GRAY)
        run = para.add_run(bold_part)
        set_run_font(run, size=11, bold=True, color=BODY_GRAY)
        run = para.add_run(rest)
        set_run_font(run, size=11, color=BODY_GRAY)
    else:
        run = para.add_run(text)
        set_run_font(run, size=11, color=BODY_GRAY)

    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(8)
    return para

def add_divider_line(doc):
    """Add a subtle divider"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

def create_document():
    doc = Document()

    # Page setup - A4, standard margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ========== 1. 课程宣传标题 ==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("廉政风险情景决策训练营")
    set_run_font(run, size=28, bold=True, color=NAVY)
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after = Pt(8)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("—— 看清楚自己，是如何一步步被说服越界的")
    set_run_font(run, size=16, bold=False, color=ACCENT_BLUE)
    subtitle_para.paragraph_format.space_before = Pt(0)
    subtitle_para.paragraph_format.space_after = Pt(20)

    add_divider_line(doc)

    # ========== 2. 课程背景与问题 ==========
    add_heading_para(doc, "一、课程背景：为什么高管需要这门课", level=1)

    add_body_para(doc, "在廉政风险防控工作中，您是否也面临这样的困惑：")

    add_body_para(doc, "制度流程越来越完善，但风险点依然存在", bullet=True)
    add_body_para(doc, "培训做了不少，员工却说'与我无关'", bullet=True)
    add_body_para(doc, "案例复盘很详细，但回到岗位依然我行我素", bullet=True)
    add_body_para(doc, "反腐败警示教育参加过了，触动很深，但没过多久又'好了伤疤忘了疼'", bullet=True)

    add_body_para(doc, "根本原因在于：传统培训让学员始终以旁观者身份'看戏'，而真正的风险发生在每一个人内心的说服过程中。")

    # Highlight box
    highlight_para = doc.add_paragraph()
    highlight_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    highlight_para.paragraph_format.left_indent = Cm(1)
    highlight_para.paragraph_format.right_indent = Cm(1)
    run = highlight_para.add_run("【核心洞察】")
    set_run_font(run, size=12, bold=True, color=DARK_RED)
    run2 = highlight_para.add_run("腐败不是一瞬间发生的，而是一个'温水煮青蛙'的渐进过程——每一步都有合理的理由，每一步都在'可接受'的范围内，直到回头看时才发现自己已经走得很远。")
    set_run_font(run2, size=11, color=BODY_GRAY)
    highlight_para.paragraph_format.space_before = Pt(12)
    highlight_para.paragraph_format.space_after = Pt(16)

    # ========== 3. 课程核心定位 ==========
    add_heading_para(doc, "二、课程核心定位", level=1)

    add_body_para(doc, "不复盘别人的案例，只做一件事——")

    core_para = doc.add_paragraph()
    core_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = core_para.add_run("把旁观者变成当局者")
    set_run_font(run, size=18, bold=True, color=NAVY)
    core_para.paragraph_format.space_before = Pt(12)
    core_para.paragraph_format.space_after = Pt(12)

    add_body_para(doc, "让学员在精心设计的模拟情景中，亲身经历'一步步被说服越界'的完整心理过程——然后当场看清这个说服自己的逻辑是怎么运作的。")

    add_body_para(doc, "这不像传统培训那样'听完觉得有道理，过后依然老样子'，而是在模拟中'亲身经历、深度触动、恍然大悟'。")

    # ========== 4. 目标学员描述 ==========
    add_heading_para(doc, "三、精准目标学员画像", level=1)

    add_body_para(doc, "本课程专为以下岗位的干部和骨干量身定制：")

    # Table for target audience
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["目标岗位", "核心痛点"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = NAVY
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    data = [
        ("采购管理", "供应商'人情'、评标中被'说服'、价格调整的灰色空间"),
        ("工程项目", "变更签证中的默契、验收签字的'通融'、资金拨付的时间差"),
        ("财务资金", "审批流程中的'特事特办'、账务处理中的'灵活变通'"),
        ("物资招标", "招标文件倾向性设置、中标结果'内定'、合同执行中的让步"),
    ]

    for row_idx, (position, pain_point) in enumerate(data, 1):
        table.cell(row_idx, 0).text = position
        table.cell(row_idx, 1).text = pain_point

    doc.add_paragraph()  # Spacer

    # ========== 5. 课程特色 ==========
    add_heading_para(doc, "四、课程特色：与传统培训的本质区别", level=1)

    features = [
        ("特色一：心理机制暴露", "不是讲道理，而是让学员'成为'当事人，在情景模拟中体验自己的心路历程，看清那些说服自己越界的逻辑链条。"),
        ("特色二：角色沉浸体验", "采用情景演绎、角色扮演等沉浸式教学方法，学员不是'听案例'，而是在模拟情景中'做决策'。"),
        ("特色三：即时复盘反思", "每个情景模拟后立即进行深度复盘，不是外部评判，而是'自己审判自己'——当事人的内在逻辑在众目睽睽下被还原、被看见。"),
        ("特色四：可迁移的觉察能力", "学员学到的不是'这个案例怎么处理'，而是'我发现自己在被说服时的思维模式'——这种觉察能力可以迁移到任何真实场景。"),
        ("特色五：行动转化路径", "不只是触动，更提供'情景决策工具箱'，让学员带走可立即使用的风险识别与应对方法。"),
    ]

    for title, desc in features:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(title + "：")
        set_run_font(run, size=12, bold=True, color=ACCENT_BLUE)
        run2 = para.add_run(desc)
        set_run_font(run2, size=11, color=BODY_GRAY)

    # ========== 6. 课程内容概述 ==========
    add_heading_para(doc, "五、课程内容模块", level=1)

    modules = [
        ("模块一：情景预热与心理准备", "建立安全敞开的课堂场域，引入真实情景素材，激活学员的相关记忆与情感。"),
        ("模块二：高风险情景模拟", "3-4个典型情景，覆盖采购、工程、财务、招标等场景，让学员在'真实决策'中体验心理变化。"),
        ("模块三：说服逻辑还原", "复盘每个情景中'说服自己越界'的完整逻辑链：合理化借口 → 风险淡化 → 自我免责 → 破例先例。"),
        ("模块四：心理机制解码", "深入分析'温水煮青蛙'效应的心理机制：渐变认知、框架效应、自我服务偏差、责任分散。"),
        ("模块五：风险觉察工具", "提供'决策前STOP工具''压力信号识别卡''风险情景预演法'等可操作工具。"),
        ("模块六：行动计划制定", "结合个人岗位实际，制定个人廉政风险防控行动计划，签署承诺书。"),
    ]

    for idx, (module_title, module_desc) in enumerate(modules, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f"【{module_title}】")
        set_run_font(run, size=11, bold=True, color=NAVY)
        run2 = para.add_run(f"\n{module_desc}")
        set_run_font(run2, size=11, color=BODY_GRAY)

    # ========== 7. 学员收获 ==========
    add_heading_para(doc, "六、学员核心收获", level=1)

    outcomes = [
        ("对个人", [
            "亲身经历'被说服越界'的完整心理过程，触动远比看案例深刻10倍",
            "掌握识别'温水煮青蛙'心理机制的能力，在第一时间觉察风险信号",
            "带走'廉政风险情景决策工具箱'，可立即应用于本职工作",
            "建立'我选择我负责'的主体意识，而非被动接受制度约束",
        ]),
        ("对组织", [
            "培养一批具有'风险嗅觉'的内训师队伍",
            "建立情景化的廉政风险教育体系，突破传统培训的瓶颈",
            "推动廉政风险防控从'制度约束'向'心理自觉'升级",
        ]),
    ]

    for category, items in outcomes:
        add_heading_para(doc, category, level=2)
        for item in items:
            add_body_para(doc, item, bullet=True)

    # ========== 8. 课程实施信息 ==========
    add_heading_para(doc, "七、课程实施信息", level=1)

    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ("课程时长", "2天（每天6小时），可根据需求调整为1天精华版"),
        ("班级规模", "24-40人/班（确保情景演练效果最优化）"),
        ("授课形式", "情景模拟 40% + 深度复盘 30% + 工具讲解 20% + 行动计划 10%"),
        ("物料配置", "学员手册、情景卡牌、决策工具卡、复盘画布、承诺书模板"),
        ("场地要求", "U型或鱼骨式座位布置，确保每个学员都能参与情景互动"),
    ]

    for row_idx, (label, value) in enumerate(info_data):
        cell0 = info_table.cell(row_idx, 0)
        cell0.text = label
        cell0.paragraphs[0].runs[0].font.bold = True
        cell0.paragraphs[0].runs[0].font.color.rgb = NAVY
        cell1 = info_table.cell(row_idx, 1)
        cell1.text = value

    doc.add_paragraph()

    # ========== 9. 典型应用场景 ==========
    add_heading_para(doc, "八、典型应用场景", level=1)

    scenarios = [
        ("新任干部任职培训", "在干部履新阶段即建立风险警觉意识，种下'心理疫苗'"),
        ("重点岗位年度轮训", "针对采购、工程、财务等高风险岗位的年度必修课"),
        ("专项审计/巡视前培训", "配合专项工作开展，增强干部的主动防控意识"),
        ("党风廉政建设活动", "作为党建品牌活动的创新形式，提升参与感和实效性"),
        ("内部控制系统完善", "结合单位内控体系建设，提供'软性'的心里防线建设"),
    ]

    for scenario, benefit in scenarios:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("▶ " + scenario)
        set_run_font(run, size=11, bold=True, color=NAVY)
        run2 = para.add_run(f"\n   {benefit}")
        set_run_font(run2, size=10, color=BODY_GRAY)

    # ========== 10. 客户见证/成功案例框架 ==========
    add_heading_para(doc, "九、成功案例（框架）", level=1)

    add_body_para(doc, "某省属国有企业集团（2024年）")
    add_body_para(doc, "背景：集团下属采购中心连续两年出现供应商围串标问题，涉事人员均为'老员工'、'业绩优秀'。", indent=True)
    add_body_para(doc, "干预：对该岗位48名骨干开展'廉政风险情景决策训练营'。", indent=True)
    add_body_para(doc, "效果：", indent=True)
    add_body_para(doc, "培训后3个月内，采购投诉率下降67%；", indent=True, bullet=True)
    add_body_para(doc, "主动上报风险隐患的数量提升3倍；", indent=True, bullet=True)
    add_body_para(doc, "85%学员表示'第一次真正看清自己可能在什么时候被说服越界'。", indent=True, bullet=True)

    add_body_para(doc, "某市政府工程建设领域专题培训（2023年）")
    add_body_para(doc, "覆盖工程审批、资金管理、质量监督等关键岗位112人。", indent=True)
    add_body_para(doc, "培训后跟踪半年，该领域信访举报量同比下降41%。", indent=True)

    # ========== 11. 联系我们 ==========
    add_heading_para(doc, "十、联系我们", level=1)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run("让每一次'差一点就越界'都变成'我守住了底线'")
    set_run_font(run, size=14, bold=True, color=NAVY)
    contact_para.paragraph_format.space_before = Pt(12)
    contact_para.paragraph_format.space_after = Pt(16)

    add_body_para(doc, "如需了解详细课程方案或预约内训，请联系我们的课程顾问团队。")

    final_para = doc.add_paragraph()
    final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final_para.add_run("【课程开发团队】")
    set_run_font(run, size=11, bold=True, color=ACCENT_BLUE)
    run2 = final_para.add_run("\n专注于党政干部廉政风险教育与情景化培训研发")
    set_run_font(run2, size=10, color=BODY_GRAY)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    create_document()

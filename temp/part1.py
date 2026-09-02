# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"D:\新课开发\战略和领导力\登攀者——AI时代的授权赋能领导力\完整课程包\04_讲师手册"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "登攀者讲师手册_完整版.docx")

C_SCRIPT = RGBColor(0x1A, 0x56, 0xDB)
C_TIME = RGBColor(0xDC, 0x26, 0x26)
C_KEY = RGBColor(0x16, 0x3A, 0x64)
C_NOTE = RGBColor(0x59, 0x73, 0x5B)
C_AI = RGBColor(0x7C, 0x3A, 0xED)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(s)

def script(p, t):
    r = p.add_run(t)
    r.font.color.rgb = C_SCRIPT
    r.font.italic = True

def time_tag(p, t):
    r = p.add_run("【" + t + "】")
    r.font.color.rgb = C_TIME
    r.font.bold = True

def key_pt(p, t):
    r = p.add_run(t)
    r.font.color.rgb = C_KEY
    r.font.bold = True

def note(p, t):
    r = p.add_run("[讲师备注: " + t + "]")
    r.font.color.rgb = C_NOTE
    r.font.size = Pt(9)

def prep_area(doc):
    p = doc.add_paragraph()
    p.add_run("-" * 60)
    r = p.add_run("【备课笔记区】")
    r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    r.font.size = Pt(9)
    np = doc.add_paragraph()
    np.add_run("本节关键点：\n可能的学员挑战：\n个人备注：")
    for run in np.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    doc.add_paragraph()

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(11)

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Cover
h = doc.add_heading('登攀者', level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h = doc.add_heading('AI时代的授权赋能领导力', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h = doc.add_heading('讲师手册（完整版）', level=2)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
doc.add_paragraph()

it = doc.add_table(rows=5, cols=2)
it.style = 'Table Grid'
for i, (l, v) in enumerate([("开发者", "罗宏伟"), ("版本", "1.0"), ("适用对象", "授权讲师"), ("课程时长", "2天（13小时）"), ("班级规模", "16-24人")]):
    it.rows[i].cells[0].text = l
    it.rows[i].cells[1].text = v
    shade(it.rows[i].cells[0], "E5E7EB")

doc.add_page_break()

# TOC
doc.add_heading('目录', level=1)
for item, page in [("第一部分 Day 1", "3"), ("  讲师资质要求与提醒", "3"), ("  课程概览与学习目标", "4"), ("  教学方法说明", "5"), ("  教具与材料清单", "6"), ("  开场：从答案到问题（30分钟）", "7"), ("  Part 1 前段：教练思维与信任环境（90分钟）", "9"), ("  Part 1 后段：GUIDE模型（60分钟）", "14"), ("  Part 2 前段：聆听与提问（90分钟）", "19"), ("  Part 2 后段：反馈与认同（120分钟）", "24"), ("第二部分 Day 2", "30"), ("  Day 2 开场与回温（30分钟）", "30"), ("  Part 3 前段：DIRECT模型（90分钟）", "32"), ("  Part 3 后段：DIRECT演练（60分钟）", "38"), ("  Part 4 前段：综合演练（75分钟）", "40"), ("  Part 4 后段：行动计划（90分钟）", "43"), ("  收尾与教练承诺（45分钟）", "46")]:
    p = doc.add_paragraph()
    p.add_run(item + "\t" + page)

doc.add_page_break()

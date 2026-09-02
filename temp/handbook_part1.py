# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

def sf(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    try:
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')
    except:
        pass

def set_cell(cell, text, bold=False, size=11):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            sf(r, size, bold)

def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = '微软雅黑'
    return p

def add_p(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, size, bold)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    sf(r)
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    return p

OUT = 'D:/新课开发/企业大学/对内/7.高管学习项目重构：AI辅助的场景推演与决策模拟设计/讲师手册/讲师手册_高管学习项目重构.docx'

# ============ COVER ============
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('高管学习项目重构')
sf(r, 28, True)
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('AI辅助的场景推演与决策模拟设计')
sf(r2, 20, True)
doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('讲师手册')
sf(r3, 24, True)
doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('FACILITATOR GUIDE')
sf(r4, 14)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t5.add_run('本手册为内部培训使用材料，请勿对外传播')
sf(r5, 10)
doc.add_page_break()

# ============ TABLE OF CONTENTS ============
add_h(doc, '目录', 1)
toc = [
    '第一章  讲师资质要求与课前准备',
    '第二章  每个模块的详细授课指引',
    '第三章  时间管理要点与超时处理',
    '第四章  学员常见问题应对话术',
    '第五章  小组讨论引导技巧',
    '第六章  评估标准与评分细则',
    '第七章  风险预案与替代方案',
    '第八章  课后跟进建议',
    '附录一  核心工具速查表',
    '附录二  场景卡使用指南',
]
for item in toc:
    p = doc.add_paragraph()
    r = p.add_run(item)
    sf(r, 12)
doc.add_page_break()

# ============ CHAPTER 1 ============
add_h(doc, '第一章  讲师资质要求与课前准备', 1)

add_h(doc, '1.1 讲师资质要求', 2)
add_p(doc, '本课程对讲师有较高的专业要求，不仅需要具备传统的培训授课能力，还需要对AI工具和高管决策场景有深刻理解。')

add_h(doc, '1.1.1 硬性资质要求', 3)
for item in [
    '具有5年以上企业培训或OD/TD相关工作经验',
    '主导过至少3个高管培训项目的设计与交付',
    '熟悉主流AI工具（ChatGPT/Claude/MiniMax等）的使用方法',
    '具备商业战略、组织变革或决策心理学相关知识背景',
    '有场景式教学或案例教学经验者优先',
]:
    add_bullet(doc, item)

add_h(doc, '1.1.2 软性能力要求', 3)
soft = [
    ('控场能力', '能够应对高管学员的质疑和挑战，引导讨论方向不偏离主题'),
    ('商业敏感度', '能快速理解学员所在行业的商业逻辑，提出有深度的问题'),
    ('AI素养', '不必须是AI技术专家，但需要理解AI的能力边界和适用场景'),
    ('学习敏锐度', '愿意持续学习新工具、新案例，保持课程内容更新'),
    ('EQ与格局', '面对高管学员时不怯场，也不炫耀，能在平等对话中传递价值'),
]
for title, desc in soft:
    p = doc.add_paragraph()
    r = p.add_run(title + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '1.1.3 讲师分级标准', 3)
tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'
hdrs = ['级别', '认证要求', '可授课范围', '年度复训']
for i, h in enumerate(hdrs):
    set_cell(tbl.cell(0, i), h, bold=True)
levels = [
    ['助理讲师', '完成认证培训+试讲', '半天工作坊', '必须'],
    ['认证讲师', '3次以上授课+案例评审', '标准版/压缩版', '必须'],
    ['高级讲师', '10次以上+课程迭代贡献', '全部版本+企业内训', '建议'],
]
for ri, row in enumerate(levels):
    for ci, val in enumerate(row):
        set_cell(tbl.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '1.2 课前准备清单', 2)
add_p(doc, '请在课前7天、3天、1天三个节点分别完成以下准备工作：')

add_h(doc, '1.2.1 课前7天：物料与学员信息确认', 3)
for item in [
    '确认学员名单，了解学员背景（行业、岗位、职级、工作年限）',
    '发送课前调研问卷，收集学员对课程的期待和已有经验',
    '确认场地布置：岛型分组座位（4-6人/组）、白板/翻页纸、投影设备',
    '确认AI工具可用性：ChatGPT Plus/Claude Pro账号登录测试',
    '打印学员手册、工作纸、场景卡等材料（建议多印10%备用）',
    '准备助教：至少1名助教负责计时、记录、分发材料',
    '如使用专业模拟平台，提前完成账号开通和权限配置',
]:
    add_bullet(doc, item)

add_h(doc, '1.2.2 课前3天：内容熟悉度准备', 3)
for item in [
    '熟读本课程所有7个模块的教学文档',
    '准备自己在每个模块的2-3个亲身经历案例（或行业案例）',
    '预演关键练习的示范答案，能在课堂上展示第一级和第二级示范',
    '熟悉附录中的参考答案，准备好应对学员的不同答案',
    '确认每个模块的核心工具能现场演示',
    '准备AI演示用的Prompt模板，现场能实时生成情境',
]:
    add_bullet(doc, item)

add_h(doc, '1.2.3 课前1天：细节确认与应急准备', 3)
for item in [
    '再次确认学员人数和分组安排',
    '检查投影仪、音响、白板笔等设备',
    '将课程PPT、案例文档、工作纸全部存入U盘备份（避免网络问题）',
    '准备Plan B：如果AI工具不可用，改为纯讨论模式',
    '确认茶歇、餐饮、住宿（如果外地）等后勤安排',
    '在教室显眼位置张贴"手机静音"提示',
    '提前30分钟到达教室，调试设备，熟悉环境',
]:
    add_bullet(doc, item)

add_h(doc, '1.3 学员信息分析表', 2)
add_p(doc, '讲师应在课前完成学员信息分析，以便调整授课重点：')
tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = 'Table Grid'
for i, h in enumerate(['分析维度', '需要了解的信息', '调整依据']):
    set_cell(tbl2.cell(0, i), h, bold=True)
info = [
    ['行业分布', '学员主要来自哪些行业', '选择行业相关的案例和场景'],
    ['职级构成', 'CEO/COO/CFO/部门负责人比例', '调整案例的复杂度和决策层级'],
    ['AI熟悉度', '学员对AI工具的了解程度', '调整AI工具讲解的深度'],
    ['培训经历', '过去参加过哪些高管培训', '避免重复，讲新角度'],
    ['学习期待', '学员最想解决什么问题', '重点回应，增加相关内容'],
]
for ri, row in enumerate(info):
    for ci, val in enumerate(row):
        set_cell(tbl2.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '1.4 教室布置标准', 2)
add_p(doc, '教室布置直接影响学员的参与度和学习效果。以下是标准布置要求：')
seating = [
    ('座位安排', '岛型分组（4-6人/组）优于剧院式座位，每组配备彩色马克笔+白板纸'),
    ('投影位置', '屏幕应在教室前方居中，从所有位置都能清晰看到'),
    ('讲师站位', '讲师不应孤立在讲台上，应能方便走到学员之间'),
    ('白板使用', '每组附近应有白板或翻页纸，用于现场板书和练习展示'),
    ('茶歇区', '茶歇区应与学习区有明显边界，避免干扰'),
    ('网络要求', '确认WiFi可用且稳定，最好准备移动热点作为备份'),
]
for title, desc in seating:
    p = doc.add_paragraph()
    r = p.add_run(title + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)
doc.add_page_break()

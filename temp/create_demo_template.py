# -*- coding: utf-8 -*-
"""
创建 001-Demo成果模板.docx - 战略解码双螺旋引擎学习成果模板
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_with_header(doc, headers, rows, header_color="4472C4"):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)

    return table

def create_template():
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('战略解码双螺旋引擎', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('学习成果模板')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(16)

    doc.add_paragraph()

    # ==================== 第一部分：必赢战役定义模板 ====================
    add_heading(doc, '一、必赢战役定义模板', 1)

    doc.add_paragraph('以下模板用于定义一场必赢战役，请按照格式填写。')

    add_heading(doc, '战役定义标准格式', 2)

    battle_template = [
        ['战役名称', '[动词 + 关键领域]（简洁、有力、清楚）'],
        ['赢的标准', '[量化的结果指标 + 时间节点]'],
        ['战役负责人', '[一个具体的人，不是委员会]'],
        ['战役时限', '[开始时间 - 结束时间]'],
        ['战略意义', '[为什么这场战役对整体战略是决定性的]'],
    ]

    table = doc.add_table(rows=len(battle_template), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(battle_template):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "E7E6E6")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    add_heading(doc, '示例：战役定义', 2)

    battle_example = [
        ['战役名称', '解决方案销售能力建设战役'],
        ['赢的标准', '2024年底，大客户销售团队60%完成解决方案销售认证；解决方案型商机在管道中占比从10%提升至40%'],
        ['战役负责人', '销售VP 李建华'],
        ['战役时限', '2024年3月 — 2024年12月'],
        ['战略意义', '如果打不赢这场，解决方案业务无从谈起——没有能卖的人，35%目标是空话'],
    ]

    table = doc.add_table(rows=len(battle_example), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(battle_example):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "D9E2F3")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== 第二部分：战役拆解模板 ====================
    add_heading(doc, '二、战役拆解模板', 1)

    doc.add_paragraph('一场战役通常有3-5个主要战术方向，每个战术方向下有若干关键战斗。')

    add_heading(doc, '战役拆解层级结构', 2)

    structure = """必赢战役（顶层：赢的目标）
├── 战术策略一（方向：怎么打）
│   ├── 关键战斗1（行动：具体做什么）── 负责人 + 完成时间 + 成果标志
│   ├── 关键战斗2── 同上
│   └── 关键里程碑（中间检查点）
├── 战术策略二
│   ├── 关键战斗1
│   └── 关键战斗2
└── 战术策略三（可选）
    └── 关键战斗"""

    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading(doc, '战术策略模板', 2)

    headers = ['战术策略', '关键战斗', '完成时间', '成果标志', '负责人']
    tactic_rows = [
        ['战术策略一：', '', '', '', ''],
        ['战术策略二：', '', '', '', ''],
        ['战术策略三：', '', '', '', ''],
    ]
    add_table_with_header(doc, headers, tactic_rows)

    doc.add_paragraph()

    add_heading(doc, '跨部门协同需求模板', 2)

    headers = ['部门', '需要提供的资源或配合', '时间要求']
    cross_rows = [
        ['', '', ''],
        ['', '', ''],
        ['', '', ''],
    ]
    add_table_with_header(doc, headers, cross_rows)

    doc.add_paragraph()

    add_heading(doc, '关键里程碑模板', 2)

    headers = ['时间节点', '里程碑描述', '判断标准']
    milestone_rows = [
        ['', '', ''],
        ['', '', ''],
        ['', '', ''],
    ]
    add_table_with_header(doc, headers, milestone_rows)

    doc.add_paragraph()

    # ==================== 第三部分：OKR模板 ====================
    add_heading(doc, '三、OKR模板', 1)

    doc.add_paragraph('OKR（Objective & Key Results）用于战役级目标管理。')

    add_heading(doc, 'OKR标准格式', 2)

    okr_template = """O（Objective）：[一句话，描述你想要达到的雄心状态，有格局，能激励，不能直接测量]

├── KR1（Key Result 1）：[可测量的结果，证明O正在实现]
├── KR2（Key Result 2）：[可测量的结果]
└── KR3（Key Result 3）：[可测量的结果]"""

    p = doc.add_paragraph()
    run = p.add_run(okr_template)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    add_heading(doc, 'KR设计检验清单', 2)

    kr_check = [
        ['检验项', '要求'],
        ['结果导向', 'KR是结果，不含"完成""推进""开展"等行动词'],
        ['可测量性', '每个KR有具体数字或可观察的状态'],
        ['逻辑性', '实现了这三个KR，O的实现是令人信服的'],
        ['独立性', 'KR之间相互独立，不重叠'],
    ]
    add_table_with_header(doc, kr_check[0], kr_check[1:])

    doc.add_paragraph()

    add_heading(doc, '战役OKR示例', 2)

    okr_example = [
        ['O', '在2024年，让智合集团大客户销售团队真正具备以顾问为导向的解决方案销售能力，让"卖解决方案"不再是口号，而是看得见的能力'],
        ['KR1', '大客户销售团队中60%完成解决方案销售认证（截至2024年12月底）'],
        ['KR2', '解决方案型商机在销售管道中的占比从10%提升至40%（Q4末）'],
        ['KR3', '完成至少3个可对外引用的解决方案标杆案例（12月底）'],
    ]

    table = doc.add_table(rows=len(okr_example), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(okr_example):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "D9E2F3")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== 第四部分：BSC战略图模板 ====================
    add_heading(doc, '四、BSC战略图模板', 1)

    doc.add_paragraph('BSC（平衡计分卡）战略图从四个层面描述组织如何创造价值。')

    add_heading(doc, 'BSC四层框架', 2)

    bsc_framework = [
        ['层面', '说明', '典型目标示例'],
        ['财务层面', '最终财务产出', '收入增长、成本降低、资产回报率'],
        ['客户层面', '驱动财务结果的价值创造', '客户满意度、市场份额、客户留存率'],
        ['内部流程层面', '支撑客户价值的运营能力', '产品质量、交付效率、创新能力'],
        ['学习与成长层面', '整个体系的基础能力', '员工能力、技术系统、组织文化'],
    ]
    add_table_with_header(doc, bsc_framework[0], bsc_framework[1:])

    doc.add_paragraph()

    add_heading(doc, 'BSC战略图模板（各层）', 2)

    doc.add_paragraph('战略主题：[一句话描述你的战略转型方向]')
    doc.add_paragraph()

    # 学习与成长层
    add_heading(doc, '学习与成长层面', 3)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    lr_rows = [
        ['L1：', '', '', ''],
        ['L2：', '', '', ''],
        ['L3：', '', '', ''],
    ]
    add_table_with_header(doc, headers, lr_rows)

    doc.add_paragraph()

    # 内部流程层
    add_heading(doc, '内部流程层面', 3)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    ip_rows = [
        ['I1：', '', '', ''],
        ['I2：', '', '', ''],
        ['I3：', '', '', ''],
    ]
    add_table_with_header(doc, headers, ip_rows)

    doc.add_paragraph()

    # 客户层
    add_heading(doc, '客户层面', 3)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    c_rows = [
        ['C1：', '', '', ''],
        ['C2：', '', '', ''],
    ]
    add_table_with_header(doc, headers, c_rows)

    doc.add_paragraph()

    # 财务层
    add_heading(doc, '财务层面', 3)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    f_rows = [
        ['F1：', '', '', ''],
        ['F2：', '', '', ''],
    ]
    add_table_with_header(doc, headers, f_rows)

    doc.add_paragraph()

    add_heading(doc, '因果链模板', 2)
    doc.add_paragraph('L1（学习层）→ I（流程层）→ C（客户层）→ F（财务层）')
    doc.add_paragraph('L2（学习层）→ I（流程层）→ C（客户层）→ F（财务层）')
    doc.add_paragraph()

    # ==================== 第五部分：部门绩效设计模板 ====================
    add_heading(doc, '五、部门绩效设计模板', 1)

    doc.add_paragraph('部门绩效设计需要平衡战略级指标和运营级指标。')

    add_heading(doc, '部门绩效设计模板', 2)

    headers = ['指标类型', '指标名称', '衡量方式', '目标值', '权重']
    perf_rows = [
        ['战略级', '', '', '', ''],
        ['战略级', '', '', '', ''],
        ['战略级', '', '', '', ''],
        ['运营级', '', '', '', ''],
        ['运营级', '', '', '', ''],
    ]
    add_table_with_header(doc, headers, perf_rows)

    doc.add_paragraph()

    add_heading(doc, '绩效设计原则', 2)

    perf_principles = [
        ['原则', '说明'],
        ['精简原则', '核心绩效指标3-5个，超过7个失去重点'],
        ['结果导向', '指标衡量"做到了什么"，而非"做了什么"'],
        ['责任到人', '每个指标必须有对应的岗位或具体人负责'],
        ['战略平衡', '战略级指标30-50%，运营级指标50-70%'],
    ]
    add_table_with_header(doc, perf_principles[0], perf_principles[1:])

    doc.add_paragraph()

    add_heading(doc, '关键岗位绩效设计模板', 2)

    headers = ['指标类型', '指标名称', '衡量方式', '季度目标', '权重']
    key岗位_rows = [
        ['结果', '', '', '', ''],
        ['结果', '', '', '', ''],
        ['行为', '', '', '', ''],
        ['行为', '', '', '', ''],
    ]
    add_table_with_header(doc, headers, key岗位_rows)

    doc.add_paragraph()

    # ==================== 第六部分：行动计划模板 ====================
    add_heading(doc, '六、课后行动计划模板', 1)

    doc.add_paragraph('基于双螺旋引擎框架，制定你的战略推动计划。')

    add_heading(doc, '第一部分：引擎一——必赢战役', 2)

    headers = ['项目', '战役一', '战役二']
    action_battle = [
        ['战役名称', '', ''],
        ['赢的标准', '', ''],
        ['战役指挥官', '', ''],
        ['前三个关键战斗', '① ② ③', '① ② ③'],
        ['战役OKR的O', '', ''],
        ['执行机制最大缺口', '', ''],
        ['一致力诊断：最需解决的对齐差距', '', ''],
    ]
    add_table_with_header(doc, headers, action_battle)

    doc.add_paragraph()

    add_heading(doc, '第二部分：引擎二——部门绩效协同', 2)

    headers = ['项目', '我的回答']
    action_bsc = [
        ['BSC战略图的战略主题', ''],
        ['学习层最关键的1个目标', ''],
        ['流程层最关键的1个目标', ''],
        ['部门核心战略级绩效指标（3个）', '① ② ③'],
        ['需要更新/新增的绩效指标', ''],
        ['现有绩效体系中与战略最不对齐的指标', ''],
    ]
    add_table_with_header(doc, headers, action_bsc)

    doc.add_paragraph()

    add_heading(doc, '第三部分：回去之后的第一件事', 2)

    action_final = [
        ['回去后第一周最重要的一件事', ''],
        ['要找的人', ''],
        ['要用的工具', ''],
    ]
    add_table_with_header(doc, ['项目', '内容'], action_final)

    # 保存文档
    output_path = 'D:/新课开发/战略和领导力/战略解码双螺旋引擎让大象跳舞/完整课程包/009-Demo成果/001-Demo成果模板.docx'
    doc.save(output_path)
    print(f'模板文档已保存至: {output_path}')

if __name__ == '__main__':
    create_template()

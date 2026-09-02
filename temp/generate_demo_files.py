# -*- coding: utf-8 -*-
"""
AI时代决策工作手册 - Demo文件生成脚本
使用python-docx和openpyxl库生成完整的教学demo文件
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
import os

# 设置输出目录
OUTPUT_DIR = "D:/新课开发/工作手册/AI时代决策工作手册/完整课程包/08-成果demo"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_red_text(paragraph, text, bold=False):
    """添加红色文本"""
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    if bold:
        run.font.bold = True
    return run

def create_decision_card_supplier():
    """生成决策卡1：供应商切换"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【决策卡】供应商切换决策")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 定位说明（红色标注）
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("定位说明：")
    run.font.bold = True
    run = p.add_run("本卡适用于采购负责人对核心供应商进行切换决策的场景。触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(9)

    doc.add_paragraph()

    # 触发条件（最显眼位置）
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    cell = table1.rows[0].cells[0]
    set_cell_shading(cell, "FFF2CC")  # 浅黄色背景

    p = cell.paragraphs[0]
    run = p.add_run("【开关】触发条件")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = cell.add_paragraph()
    run = p.add_run("如果出现以下组合信号，先做核实动作，再决定是否切换供应商：")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("① 质量异常报告同一指标连续2次超出历史波动区间")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("② 或交付准时率连续3个月低于95%")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("③ 且无明确外部因素解释（如政策变化、不可抗力）")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    # 检查表
    p = doc.add_paragraph()
    run = p.add_run("【检查表】切换前必须核实的5项")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    checklist = [
        "核实备选供应商最近3个月实际交付记录，而非口头承诺的产能数字",
        "核查备选供应商同品类产品的质量检验合格率（需≥原供应商历史水平）",
        "评估切换后磨合期对生产计划的影响（预估至少2周的产能波动）",
        "确认备选供应商的原材料来源是否稳定（要求提供近6个月采购凭证）",
        "核实备选供应商实际地理位置和物流周期，评估应急调货时间窗口"
    ]

    table2 = doc.add_table(rows=len(checklist), cols=2)
    table2.style = 'Table Grid'

    for i, item in enumerate(checklist):
        cell1 = table2.rows[i].cells[0]
        cell2 = table2.rows[i].cells[1]
        cell1.text = f"□ {i+1}"
        cell1.width = Cm(0.8)
        cell2.text = item
        cell2.width = Cm(14)

    doc.add_paragraph()

    # 应急方案
    p = doc.add_paragraph()
    run = p.add_run("【应急方案】变体场景处理")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    emergency = [
        ("如果时间极度紧迫（需24小时内决策）", "启动【双供应商并行】模式：新供应商先以20%小批量试单，同步保持原供应商70%订单量，剩余10%作为安全库存缓冲"),
        ("如果关键人员缺位（采购负责人不在）", "由备用联系人发起临时采购决议，抄送供应链总监和财务总监，48小时内补齐正式审批流程"),
        ("如果原供应商提出整改承诺", "要求对方提供书面整改方案和时间表，设定3周观察期，观察期内新供应商候选保持待命状态")
    ]

    for scenario, solution in emergency:
        p = doc.add_paragraph()
        run = p.add_run(f"• {scenario}，则：{solution}")
        run.font.size = Pt(9)

    doc.add_paragraph()

    # 适用场景说明
    p = doc.add_paragraph()
    run = p.add_run("【适用场景】")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("• 适用：核心供应商的质量或交付问题已出现明确信号，需评估切换必要性")
    p = doc.add_paragraph()
    p.add_run("• 适用：采购品类为定制件或长周期物料（交期>4周）的供应商管理")
    p = doc.add_paragraph()
    p.add_run("• 不适用：日常采购中的临时性小额订单调整（金额<5万元且不影响生产计划）")
    p = doc.add_paragraph()
    p.add_run("• 不适用：因价格波动的常规供应商替换（应走年度采购谈判流程）")

    doc.add_paragraph()

    # 警示案例（嵌入式）
    p = doc.add_paragraph()
    run = p.add_run("【警示案例】")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("岔路口描述：")
    run.font.bold = True
    p.add_run("某制造企业采购负责人因原供应商报价上涨15%，决定切换至价格更低的备选供应商。未做深度核实，仅凭备选供应商提供的产能声明就完成了切换决策。")

    p = doc.add_paragraph()
    run = p.add_run("被忽略的信号：")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run("备选供应商在样品测试阶段交期已出现3次延迟，每次都被解释为【物流原因】；实际产能仅为其声称的60%")

    p = doc.add_paragraph()
    run = p.add_run("后果：")
    run.font.bold = True
    p.add_run("切换后第2个月即因备选供应商产能不足，导致两条产线停工共计12天，直接损失超200万元")

    doc.add_paragraph()

    # 认领人信息
    p = doc.add_paragraph()
    run = p.add_run("【认领人信息】")
    run.font.bold = True
    run.font.size = Pt(11)

    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = "主认领人"
    table3.rows[0].cells[1].text = "李明 / 供应链部采购总监 / 138-xxxx-xxxx"
    table3.rows[1].cells[0].text = "备份认领人"
    table3.rows[1].cells[1].text = "王芳 / 供应链部采购经理 / 139-xxxx-xxxx"
    table3.rows[2].cells[0].text = "最近更新日期"
    table3.rows[2].cells[1].text = "2026-06-15"

    doc.add_paragraph()

    # 版本历史
    p = doc.add_paragraph()
    run = p.add_run("【版本历史】")
    run.font.bold = True
    run.font.size = Pt(11)

    table4 = doc.add_table(rows=3, cols=4)
    table4.style = 'Table Grid'
    headers = ["版本", "日期", "修改人", "主要变更"]
    for i, header in enumerate(headers):
        table4.rows[0].cells[i].text = header
        set_cell_shading(table4.rows[0].cells[i], "D9D9D9")

    history = [
        ["V1.0", "2026-03-10", "罗宏伟", "初版发布"],
        ["V1.1", "2026-06-15", "李明", "增加应急方案第3条；补充警示案例"]
    ]
    for row_idx, row_data in enumerate(history):
        for col_idx, cell_data in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例，非实际业务决策依据 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-决策卡-供应商切换.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_decision_card_risk_judgment():
    """生成决策卡2：项目风险判断"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【决策卡】项目风险判断决策")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 定位说明
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("定位说明：")
    run.font.bold = True
    run = p.add_run("本卡适用于项目经理对项目执行过程中出现的风险信号进行判断的场景。触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(9)

    doc.add_paragraph()

    # 触发条件
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    cell = table1.rows[0].cells[0]
    set_cell_shading(cell, "FFF2CC")

    p = cell.paragraphs[0]
    run = p.add_run("【开关】触发条件")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = cell.add_paragraph()
    run = p.add_run("如果出现以下组合信号，先停下来重新评估风险等级，再决定是否继续推进：")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("① 项目进度偏差累计超过2周，且无明确赶工措施")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("② 预算消耗速度超过进度完成速度（Burn Rate > 1.2）")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("③ 关键技术方案未经评审即进入实施阶段")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("④ 或核心团队成员在关键路径上发生变更")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    # 检查表
    p = doc.add_paragraph()
    run = p.add_run("【检查表】风险再评估6项必查")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    checklist = [
        "重新核对项目基准范围（Scope），确认偏差是否涉及范围蔓延",
        "计算当前项目的EAC（完工估算），与原始BAC（完工预算）对比偏差率",
        "列出剩余工作的关键依赖路径，标注任何新增的外部依赖",
        "与核心技术负责人单独确认技术方案的可行性和风险点",
        "评估核心团队成员变更对进度的影响（按人天估算延误量）",
        "整理一份1页纸的风险摘要，包含TOP3风险项和建议的应对策略"
    ]

    table2 = doc.add_table(rows=len(checklist), cols=2)
    table2.style = 'Table Grid'

    for i, item in enumerate(checklist):
        cell1 = table2.rows[i].cells[0]
        cell2 = table2.rows[i].cells[1]
        cell1.text = f"□ {i+1}"
        cell1.width = Cm(0.8)
        cell2.text = item
        cell2.width = Cm(14)

    doc.add_paragraph()

    # 应急方案
    p = doc.add_paragraph()
    run = p.add_run("【应急方案】变体场景处理")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    emergency = [
        ("如果项目已处于关键路径延误状态", "立即召开专题会议，识别可以并行执行的任务，将串行改为并行以压缩工期；同时评估是否需要追加资源"),
        ("如果预算即将耗尽（预计7天内）", "暂停所有非必要支出，优先保障关键路径任务；48小时内向管理层提交资金申请和项目状态报告"),
        ("如果客户需求发生重大变更", "启动变更控制流程，评估对范围、进度、成本的影响；先签变更单再实施，不得边实施边谈变更")
    ]

    for scenario, solution in emergency:
        p = doc.add_paragraph()
        run = p.add_run(f"• {scenario}，则：{solution}")
        run.font.size = Pt(9)

    doc.add_paragraph()

    # 适用场景
    p = doc.add_paragraph()
    run = p.add_run("【适用场景】")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("• 适用：项目执行过程中遭遇技术难题、资源短缺或外部环境变化")
    p = doc.add_paragraph()
    p.add_run("• 适用：跨部门协作项目中出现的接口责任不清、依赖延误")
    p = doc.add_paragraph()
    p.add_run("• 不适用：项目立项阶段的初步风险评估（应使用立项评审清单）")
    p = doc.add_paragraph()
    p.add_run("• 不适用：项目收尾阶段的验收争议（应使用验收决策卡）")

    doc.add_paragraph()

    # 警示案例
    p = doc.add_paragraph()
    run = p.add_run("【警示案例】")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("岔路口描述：")
    run.font.bold = True
    p.add_run("某软件开发项目在开发阶段遇到技术方案变更，项目经理评估后认为【可以通过加班赶回来】，继续按原计划推进。")

    p = doc.add_paragraph()
    run = p.add_run("被忽略的信号：")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run("核心开发人员在周报中连续3次提到【技术方案可能需要调整】，项目经理均以【先按现有方案推进】回复；同期Burn Rate已达1.5")

    p = doc.add_paragraph()
    run = p.add_run("后果：")
    run.font.bold = True
    p.add_run("项目最终延期4个月交付，追加预算80%，客户满意度大幅下降")

    doc.add_paragraph()

    # 认领人信息
    p = doc.add_paragraph()
    run = p.add_run("【认领人信息】")
    run.font.bold = True
    run.font.size = Pt(11)

    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = "主认领人"
    table3.rows[0].cells[1].text = "张伟 / 项目管理部总监 / 137-xxxx-xxxx"
    table3.rows[1].cells[0].text = "备份认领人"
    table3.rows[1].cells[1].text = "刘强 / PMO经理 / 136-xxxx-xxxx"
    table3.rows[2].cells[0].text = "最近更新日期"
    table3.rows[2].cells[1].text = "2026-05-20"

    doc.add_paragraph()

    # 版本历史
    p = doc.add_paragraph()
    run = p.add_run("【版本历史】")
    run.font.bold = True
    run.font.size = Pt(11)

    table4 = doc.add_table(rows=2, cols=4)
    table4.style = 'Table Grid'
    headers = ["版本", "日期", "修改人", "主要变更"]
    for i, header in enumerate(headers):
        table4.rows[0].cells[i].text = header
        set_cell_shading(table4.rows[0].cells[i], "D9D9D9")

    history = [["V1.0", "2026-04-15", "张伟", "初版发布"]]
    for row_idx, row_data in enumerate(history):
        for col_idx, cell_data in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例，非实际业务决策依据 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-决策卡-项目风险判断.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_decision_card_process_anomaly():
    """生成决策卡3：流程异常识别"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【决策卡】流程异常识别决策")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 定位说明
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("定位说明：")
    run.font.bold = True
    run = p.add_run("本卡适用于部门负责人或流程管理员识别日常业务流程中出现的异常情况。触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.size = Pt(9)

    doc.add_paragraph()

    # 触发条件
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    cell = table1.rows[0].cells[0]
    set_cell_shading(cell, "FFF2CC")

    p = cell.paragraphs[0]
    run = p.add_run("【开关】触发条件")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = cell.add_paragraph()
    run = p.add_run("如果出现以下组合信号，先停下来核实原因，再决定是否启动流程优化：")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("① 某流程环节的平均处理时长突然缩短超过30%（无系统性提速措施）")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("② 或某个节点的审批跳过率（Override Rate）异常上升")
    run.font.size = Pt(10)

    p = cell.add_paragraph()
    run = p.add_run("③ 且同一操作人员连续出现多次类似异常行为")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    # 检查表
    p = doc.add_paragraph()
    run = p.add_run("【检查表】异常识别7项核查")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    checklist = [
        "调取该流程近3个月的处理数据，计算历史平均时长和标准差",
        "核查是否存在系统变更或功能更新导致流程路径改变",
        "追溯被跳过审批节点的授权记录，确认是否有违规授权",
        "访谈该环节的实际操作人员，了解操作习惯和遇到的障碍",
        "调取异常操作发生时的日志，识别是否集中在特定时间段",
        "评估该异常是否与近期组织架构调整或岗位职责变动相关",
        "汇总异常数据，形成流程健康度报告，提交流程owner审阅"
    ]

    table2 = doc.add_table(rows=len(checklist), cols=2)
    table2.style = 'Table Grid'

    for i, item in enumerate(checklist):
        cell1 = table2.rows[i].cells[0]
        cell2 = table2.rows[i].cells[1]
        cell1.text = f"□ {i+1}"
        cell1.width = Cm(0.8)
        cell2.text = item
        cell2.width = Cm(14)

    doc.add_paragraph()

    # 应急方案
    p = doc.add_paragraph()
    run = p.add_run("【应急方案】变体场景处理")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    emergency = [
        ("如果异常涉及合规边界（如审批权限滥用）", "立即暂停该操作人员权限，48小时内启动正式调查；同步通报合规部门"),
        ("如果是系统原因导致的假性异常", "协调IT部门修复系统Bug，临时建立人工台账确保业务连续性；同时更新SOP"),
        ("如果是流程设计缺陷导致的规避行为", "组织流程owner和关键用户开展流程优化工作坊，2周内输出新流程草案")
    ]

    for scenario, solution in emergency:
        p = doc.add_paragraph()
        run = p.add_run(f"• {scenario}，则：{solution}")
        run.font.size = Pt(9)

    doc.add_paragraph()

    # 适用场景
    p = doc.add_paragraph()
    run = p.add_run("【适用场景】")
    run.font.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("• 适用：日常运营中发现流程处理速度异常（过快或过慢）")
    p = doc.add_paragraph()
    p.add_run("• 适用：定期流程审计中发现的系统性偏差")
    p = doc.add_paragraph()
    p.add_run("• 不适用：单个偶发性的流程延误（应作为个例处理，不触发本卡）")
    p = doc.add_paragraph()
    p.add_run("• 不适用：因外部客户要求导致的流程调整（应走变更申请流程）")

    doc.add_paragraph()

    # 警示案例
    p = doc.add_paragraph()
    run = p.add_run("【警示案例】")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("岔路口描述：")
    run.font.bold = True
    p.add_run("某公司财务报销流程中，审批环节的平均处理时长突然从3天缩短至0.5天。流程管理员认为这是【效率提升】，未做深入调查。")

    p = doc.add_paragraph()
    run = p.add_run("被忽略的信号：")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p.add_run("审批时长骤降的同时，报销差错率上升了3倍；部分员工的报销金额明显偏高且集中在月底")

    p = doc.add_paragraph()
    run = p.add_run("后果：")
    run.font.bold = True
    p.add_run("半年后内部审计发现，审批人员为完成KPI存在【秒批】行为，同时存在虚报交通费的情况，累计违规金额约15万元")

    doc.add_paragraph()

    # 认领人信息
    p = doc.add_paragraph()
    run = p.add_run("【认领人信息】")
    run.font.bold = True
    run.font.size = Pt(11)

    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = "主认领人"
    table3.rows[0].cells[1].text = "陈静 / 运营管理部经理 / 135-xxxx-xxxx"
    table3.rows[1].cells[0].text = "备份认领人"
    table3.rows[1].cells[1].text = "赵敏 / 质量管理部主管 / 134-xxxx-xxxx"
    table3.rows[2].cells[0].text = "最近更新日期"
    table3.rows[2].cells[1].text = "2026-04-28"

    doc.add_paragraph()

    # 版本历史
    p = doc.add_paragraph()
    run = p.add_run("【版本历史】")
    run.font.bold = True
    run.font.size = Pt(11)

    table4 = doc.add_table(rows=2, cols=4)
    table4.style = 'Table Grid'
    headers = ["版本", "日期", "修改人", "主要变更"]
    for i, header in enumerate(headers):
        table4.rows[0].cells[i].text = header
        set_cell_shading(table4.rows[0].cells[i], "D9D9D9")

    history = [["V1.0", "2026-04-28", "陈静", "初版发布"]]
    for row_idx, row_data in enumerate(history):
        for col_idx, cell_data in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例，非实际业务决策依据 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-决策卡-流程异常识别.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_scenario_mapping_matrix():
    """生成场景映射矩阵Excel"""
    wb = openpyxl.Workbook()

    # Sheet1: 场景映射矩阵
    ws = wb.active
    ws.title = "供应商管理场景映射"

    # 设置列宽
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 35
    ws.column_dimensions['D'].width = 35

    # 标题行
    ws.merge_cells('A1:D1')
    ws['A1'] = "供应商管理场景映射矩阵"
    ws['A1'].font = Font(size=14, bold=True, color="FFFFFF")
    ws['A1'].fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 30

    # 列标题行
    headers = ["场景变量", "核心供应商", "边缘供应商"]
    ws['A3'] = ""
    ws['B3'] = "核心供应商"
    ws['C3'] = "边缘供应商"
    ws['D3'] = "场景特征说明"

    for col in ['A', 'B', 'C', 'D']:
        ws[f'{col}3'].font = Font(bold=True, size=11)
        ws[f'{col}3'].fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        ws[f'{col}3'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    ws.row_dimensions[3].height = 25

    # 采购品类行标题
    ws['A4'] = "标准件"
    ws['A4'].font = Font(bold=True)
    ws['A4'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[4].height = 60

    ws['A5'] = "定制件"
    ws['A5'].font = Font(bold=True)
    ws['A5'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[5].height = 60

    # 场景1：核心供应商-标准件
    ws['B4'] = """【决策卡参数】
触发条件：质量异常同一指标连续2次超出阈值；或交付准时率连续3月<95%

检查表5项：
① 核实备选供应商近3月实际交付记录
② 核查备选供应商同品类质量检验合格率
③ 评估切换后2周产能波动
④ 确认备选供应商原材料来源稳定性
⑤ 核实备选供应商实际物流周期

应急方案：启动"双供应商并行"模式

适用场景：战略采购、长期合作供应商管理"""
    ws['B4'].alignment = Alignment(vertical='top', wrap_text=True)

    # 场景2：边缘供应商-标准件
    ws['C4'] = """【决策卡参数】
触发条件：价格波动超过±10%；或交付准时率单月<90%

检查表4项（简化版）：
① 核查市场同类供应商报价
② 评估切换采购成本
③ 确认备选供应商交付能力
④ 评估账期和付款条件

应急方案：直接切换至备选

适用场景：日常采购、零星采购"""
    ws['C4'].alignment = Alignment(vertical='top', wrap_text=True)

    # 场景3：核心供应商-定制件
    ws['B5'] = """【决策卡参数】
触发条件：定制件质量问题连续2次；或供应商技术能力下降

检查表6项（加强版）：
① 核实备选供应商同类定制件经验
② 核查技术方案承接能力
③ 评估磨合期对核心客户的影响
④ 确认备选供应商定制件产能
⑤ 评估模具转移时间和成本
⑥ 核实定制件质量检验标准

应急方案：保持原供应商小批量生产

适用场景：核心产品定制件、关键技术部件"""
    ws['B5'].alignment = Alignment(vertical='top', wrap_text=True)

    # 场景4：边缘供应商-定制件
    ws['C5'] = """【决策卡参数】
触发条件：定制件质量问题1次；或交付延期超过1周

检查表3项（快速版）：
① 确认备选供应商定制能力
② 评估切换时间和成本
③ 确认质量检验能力

应急方案：直接切换

适用场景：非关键定制件、辅助材料"""
    ws['C5'].alignment = Alignment(vertical='top', wrap_text=True)

    # 场景特征说明
    ws['D4'] = """【核心供应商特征】
• 供货占比>20%
• 战略合作关系
• 技术门槛较高
• 替代供应商难找

【标准件特征】
• 通用规格
• 市场供应充足
• 价格透明
• 切换成本低"""
    ws['D4'].alignment = Alignment(vertical='top', wrap_text=True)

    ws['D5'] = """【边缘供应商特征】
• 供货占比<10%
• 常规合作关系
• 通用技术
• 替代供应商容易

【定制件特征】
• 非标规格
• 市场供应有限
• 技术门槛较高
• 切换成本高"""
    ws['D5'].alignment = Alignment(vertical='top', wrap_text=True)

    # 设置边框
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    for row in range(3, 6):
        for col in range(1, 5):
            ws.cell(row=row, column=col).border = thin_border

    # Sheet2: 空白练习版
    ws2 = wb.create_sheet("练习版-空白矩阵")

    ws2.column_dimensions['A'].width = 18
    ws2.column_dimensions['B'].width = 28
    ws2.column_dimensions['C'].width = 28
    ws2.column_dimensions['D'].width = 20

    # 标题行
    ws2.merge_cells('A1:D1')
    ws2['A1'] = "供应商管理场景映射矩阵 - 练习版"
    ws2['A1'].font = Font(size=14, bold=True, color="FFFFFF")
    ws2['A1'].fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    ws2['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws2.row_dimensions[1].height = 30

    # 列标题行
    ws2['A3'] = ""
    ws2['B3'] = "核心供应商"
    ws2['C3'] = "边缘供应商"
    ws2['D3'] = "填写说明"

    for col in ['A', 'B', 'C', 'D']:
        ws2[f'{col}3'].font = Font(bold=True, size=11)
        ws2[f'{col}3'].fill = PatternFill(start_color="D9D9D9", end_color="D9D9D9", fill_type="solid")
        ws2[f'{col}3'].alignment = Alignment(horizontal='center', vertical='center')

    ws2.row_dimensions[3].height = 25

    ws2['A4'] = "标准件"
    ws2['A4'].font = Font(bold=True)
    ws2['A4'].alignment = Alignment(horizontal='center', vertical='center')

    ws2['A5'] = "定制件"
    ws2['A5'].font = Font(bold=True)
    ws2['A5'].alignment = Alignment(horizontal='center', vertical='center')

    # 填写说明
    ws2['D4'] = """请根据以下维度填写：
1. 触发条件（具体可观测信号）
2. 检查表项目（用动词开头）
3. 应急方案
4. 适用/不适用场景"""
    ws2['D4'].alignment = Alignment(vertical='top', wrap_text=True)

    ws2['D5'] = """请思考：
1. 与标准件场景的差异点
2. 需要特别注意的风险点
3. 决策的优先级判断"""
    ws2['D5'].alignment = Alignment(vertical='top', wrap_text=True)

    # 为练习格添加浅色背景
    light_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    for row in range(4, 6):
        for col in range(2, 4):
            ws2.cell(row=row, column=col).fill = light_fill
            ws2.cell(row=row, column=col).border = thin_border

    # 添加使用说明
    ws2['A7'] = "使用说明："
    ws2['A7'].font = Font(bold=True, color="C00000")
    ws2.merge_cells('A8:D8')
    ws2['A8'] = "本练习版供学员填写。每个场景格子对应一个具体的供应商管理情境，学员需要根据情境自主设计触发条件、检查表和应急方案。"
    ws2['A8'].alignment = Alignment(wrap_text=True)
    ws2.merge_cells('A9:D9')
    ws2['A9'] = "练习完成后，可对比【供应商管理场景映射】Sheet中的参考答案，进行自我评估。"
    ws2['A9'].alignment = Alignment(wrap_text=True)

    filename = os.path.join(OUTPUT_DIR, "demo-场景映射矩阵-供应商管理.xlsx")
    wb.save(filename)
    print(f"已生成: {filename}")


def create_audit_dialogue_record():
    """生成稽核对话记录"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【稽核对话记录】供应商切换决策稽核")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # 基本信息表
    p = doc.add_paragraph()
    run = p.add_run("一、基本信息")
    run.font.bold = True
    run.font.size = Pt(12)

    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'
    info = [
        ["决策名称", "2026年Q2核心供应商切换决策"],
        ["决策日期", "2026年4月15日"],
        ["决策者", "李明（供应链部采购总监）"],
        ["稽核日期", "2026年5月20日"],
        ["稽核人", "罗宏伟（复盘引导师）"],
        ["对话时长", "约45分钟"]
    ]
    for i, (key, value) in enumerate(info):
        table1.rows[i].cells[0].text = key
        table1.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 稽核对话内容
    p = doc.add_paragraph()
    run = p.add_run("二、稽核对话全文")
    run.font.bold = True
    run.font.size = Pt(12)

    dialogue = [
        ("【问题1：沉淀检查】", """复盘引导师：李明，这次供应商切换决策，我想先了解一下，当时整个决策过程中，有没有让你感到"不确定"的关键节点？

李明：有的。主要是在决定切换的那一刻，我其实犹豫过。

复盘引导师：那个犹豫具体是什么情况？

李明：当时原供应商的质量异常报告已经连续2次了，按道理应该切换，但我在想要不要先给他们一个整改的机会。备选供应商那边也有不确定性，心里没底。

复盘引导师：那个"不确定"后来有没有被记录下来？

李明：没有正式记录。当时就是脑子里想了一下，后来就拍板了。

【判断：沉淀□ 未沉淀√】""", "● 沉淀点：决策者明确识别出【是否给原供应商整改机会】作为关键犹豫点\n● 未沉淀点：备选供应商的不确定性评估过程没有留下记录"),

        ("【问题2：使用检查】", """复盘引导师：我们之前有一张《供应商切换决策卡》，你当时有没有用到？

李明：那张卡我看了。触发条件里写的是"同一指标连续2次超出阈值"，我们这次确实触发了。检查表我也核对过，大概做了三四项。

复盘引导师：做了哪些？

李明：核 实了备选供应商的交付记录，也评估了磨合期的影响。但是……

复盘引导师：但是什么？

李明：有两项没做。一个是原材料来源稳定性核查，这个当时觉得麻烦就没查。另一个是地理位置和物流周期，这个后来发现其实挺重要的。

复盘引导师：使用后有反馈给卡片认领人吗？

李明：还没有。

【判断：已使用□ 部分使用√ 反馈：否√】""", "● 已使用：决策者承认部分参考了卡片内容\n● 未完整执行：原材料来源和物流周期两项检查未执行\n● 未反馈：卡片内容与实际操作存在偏差，未反馈给认领人"),

        ("【问题3：价值识别】", """复盘引导师：你觉得这次决策，从判断难度上来说，值得做成一张决策卡吗？

李明：我觉得值得。因为这种情况以后可能还会遇到。

复盘引导师：如果再做一次这张卡，你有什么建议想让下一任采购负责人注意的？

李明：一个是备选供应商的"口头承诺"不能轻信，一定要有书面凭证。另一个是磨合期的影响比我想象的更大，新供应商前两个月的交付要盯紧一点。

【判断：具备价值√ 已提取建议√】""", "● 价值识别：决策者认为值得沉淀\n● 隐性经验浮现：两条关键建议来自亲历教训"),

        ("【问题4：反馈迭代】", """复盘引导师：你刚才提到有两项检查没做完。回顾一下，你觉得是卡片设计有问题，还是执行不到位？

李明：两边都有。执行上是我当时图省事。卡片设计上……那个原材料来源稳定性的核查项，要求提供【近6个月采购凭证】，说实话执行起来不太现实，能不能改成【核实主要原材料供应商名单】？

复盘引导师：这是个好建议。还有别的吗？

李明：还有一个，磨合期的评估能不能给个参考时间？比如【新供应商切换后至少保持2周观察期】？

【判断：发现漏洞√ 反馈建议√】""", "● 卡片漏洞：检查项【采购凭证】执行困难，需简化\n● 迭代建议：增加磨合期观察期指导")
    ]

    for i, (title, content, summary) in enumerate(dialogue):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        p = doc.add_paragraph()
        p.add_run(content)
        p.style = doc.styles['Quote']

        p = doc.add_paragraph()
        run = p.add_run(summary)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    doc.add_paragraph()

    # 稽核结论
    p = doc.add_paragraph()
    run = p.add_run("三、稽核结论与建议")
    run.font.bold = True
    run.font.size = Pt(12)

    conclusions = [
        ("□ 经验已有效沉淀", "部分沉淀。关键犹豫点未被记录，但提取出2条隐性建议"),
        ("□ 需要跟进", "建议补充记录本次决策的犹豫点评估过程"),
        ("☑ 需要迭代卡片", "根据反馈，更新检查表第4条为【核实主要原材料供应商名单】；补充磨合期观察期指导")
    ]

    for status, content in conclusions:
        p = doc.add_paragraph()
        run = p.add_run(status + " ")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.add_run(content)

    doc.add_paragraph()

    # 额外问题
    p = doc.add_paragraph()
    run = p.add_run("四、额外问题：偏离检测")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("复盘引导师：这次决策你有没有基于自己的判断，偏离了卡片上的建议？为什么？")

    p = doc.add_paragraph()
    run = p.add_run("李明：偏离了。原材料来源稳定性那个检查项，我觉得执行成本太高，就跳过了。现在回想起来，其实应该查一下，那是我自己的判断失误，不是卡片的问题。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("【判断】未因畏惧风险而不敢偏离；主动承认执行层面问题，区分了【卡片设计】与【执行不到位】")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    doc.add_paragraph()

    # 签名
    p = doc.add_paragraph()
    run = p.add_run("稽核人签名：__________    日期：2026-05-20")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("决策者签名：__________    日期：2026-05-20")
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-稽核对话记录.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_interview_record():
    """生成复盘访谈记录"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【复盘访谈记录】某项目风险判断决策")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # 基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、访谈基本信息")
    run.font.bold = True
    run.font.size = Pt(12)

    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'
    info = [
        ["决策名称", "某产品研发项目风险判断与叫停决策"],
        ["决策者", "张伟（项目管理部总监）"],
        ["访谈日期", "2026年6月10日"],
        ["访谈者", "罗宏伟（复盘引导师）"],
        ["访谈时长", "约90分钟"],
        ["访谈形式", "一对一当面访谈"]
    ]
    for i, (key, value) in enumerate(info):
        table1.rows[i].cells[0].text = key
        table1.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 访谈正文
    p = doc.add_paragraph()
    run = p.add_run("二、访谈正文")
    run.font.bold = True
    run.font.size = Pt(12)

    # 第一维度：追因
    p = doc.add_paragraph()
    run = p.add_run("【维度一：追因】决策触发点挖掘")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    p.add_run("复盘引导师：张伟，先跟我说说这个项目整体的情况，当时是什么让你觉得需要做这个决策？")

    p = doc.add_paragraph()
    p.add_run("张伟：那个项目是去年Q3启动的，做的是一个新品研发。一开始进展还挺顺利的，到了今年Q1，突然就感觉不太对劲了。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：\"感觉不太对劲\"具体是指什么？你能描述一下那个信号吗？")

    p = doc.add_paragraph()
    p.add_run("张伟：就是……进度突然变快了。你知道做研发，正常情况下，进度应该是跟计划差不多，或者稍微慢一点。但那个阶段，项目进度报告上显示完成率突然飙升，比计划快了大概三周。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：进度变快本身听起来是好事，为什么这会让你觉得需要做决策？")

    p = doc.add_paragraph()
    p.add_run("张伟：一开始我也没觉得是问题。但是后来我仔细看了周报，发现一个问题——质量报告的提交时间点和进度飙升的时间点高度重合。我就问了一下技术负责人，他说\"为了赶进度，质量测试有些环节简化了\"。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师追问：\"在那之前呢？你有没有在更早的时候就感觉到什么不对劲？\"")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    p = doc.add_paragraph()
    p.add_run("张伟：（思考）你这么一说，我想起来了。其实在进度飙升之前大概两周，有个开发人员跟我提过一句，说\"感觉这个功能实现起来比预期的麻烦\"。我当时没太在意，觉得开发过程中的波折很正常。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("【隐性信号识别】")
    run.font.bold = True
    p.add_run("初始触发点：技术负责人透露\"质量测试环节简化\"\n更深层信号：两周前开发人员提到的\"比预期麻烦\"\n两者关联：进度异常加速掩盖了原本应该暴露的技术风险")

    doc.add_paragraph()

    # 第二维度：权衡
    p = doc.add_paragraph()
    run = p.add_run("【维度二：权衡】选项与排除逻辑")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：当时你能想到的选项有哪些？")

    p = doc.add_paragraph()
    p.add_run("张伟：说实话，一开始我脑子里只有两个声音。一个是\"继续推进，反正快完成了\"，另一个是\"停下来评估一下风险\"。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师追问：\"除了这两个选项，你有没有想过第三个方向？比如，一边继续一边整改？\"")

    p = doc.add_paragraph()
    p.add_run("张伟：（摇头）想过，但是后来否决了。因为我当时判断，如果走\"边推进边整改\"这条路，团队的注意力会被分散，而且那个技术风险点不解决，后面的问题只会越来越大。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师追问：\"你当时为什么觉得那个风险点\"只会越来越大\"？你是怎么得出这个判断的？\"")

    p = doc.add_paragraph()
    p.add_run("张伟：凭经验吧。以前遇到类似情况的时候，如果你不主动停下来处理，问题往往会滚雪球。但是我承认，这个判断我没法量化。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("【隐性判断浮现】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    p = doc.add_paragraph()
    p.add_run("排除\"边推进边整改\"的理由：团队注意力分散 + 风险累积\n隐性假设：技术风险具有自增性，不会自然消解\n未验证的直觉：\"以前遇到类似情况\"——多次经验压缩的直觉反应")

    doc.add_paragraph()

    # 第三维度：未预见的假设
    p = doc.add_paragraph()
    run = p.add_run("【维度三：未预见的假设】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：如果当时那个前提——我是说，你假设\"继续推进会让风险越来越大\"——如果这个前提不成立，你还会做同样的判断吗？")

    p = doc.add_paragraph()
    p.add_run("张伟：如果这个前提不成立……我可能真的会犹豫。因为我当时的决策，核心依赖就是这个判断。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：那我们来想想，有没有什么情况是这个前提不成立的？")

    p = doc.add_paragraph()
    p.add_run("张伟：（思考了很久）可能……如果那个技术风险点是一个相对独立的模块，后面不依赖它，可能就不会滚那么大。但我们这个项目不是，那个模块是核心。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("【假设验证】")
    run.font.bold = True
    p = doc.add_paragraph()
    p.add_run("关键假设：技术风险模块是核心依赖，非独立模块\n反事实情景：若为独立模块，判断可能不同\n最终结论：假设与实际情况吻合，判断合理")

    doc.add_paragraph()

    # 隐性判断挖掘
    p = doc.add_paragraph()
    run = p.add_run("【隐性判断挖掘】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师：张伟，我注意到你刚才说了好几次【凭经验】的感觉。我想问一下，这种【经验】有没有可能在你自己脑子里已经形成了一条规则？")

    p = doc.add_paragraph()
    p.add_run("张伟：你这么说我倒是想起来了，我确实有一个习惯，就是当进度异常加速的时候，我会特别警惕。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("复盘引导师追问：\"这个'警惕'具体是什么意思？你会怎么做？\"")

    p = doc.add_paragraph()
    p.add_run("张伟：我会去查两个东西：一个是质量报告，看测试环节有没有被简化；另一个是直接问开发人员，实际遇到的技术困难比预期大多少。如果这两个信号同时出现，我就会启动风险评估。")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    run = p.add_run("【隐性规则提炼】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    p = doc.add_paragraph()
    p.add_run("\"当进度异常加速 + 质量测试简化 + 技术难度超预期三者同时出现，我就会启动风险再评估\"\n这条规则张伟从未主动语言化，但在历次类似决策中反复使用")

    doc.add_paragraph()

    # 访谈小结
    p = doc.add_paragraph()
    run = p.add_run("三、访谈小结")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("本次访谈成功挖掘出以下判断结构：")

    findings = [
        "触发信号：进度异常加速 + 质量报告异常 + 技术难度超预期，三者组合触发风险再评估",
        "排除逻辑：\"边推进边整改\"会分散注意力，且技术风险具有自增性",
        "隐性假设：技术风险模块是核心依赖，若为独立模块则判断可能不同",
        "隐性经验规则：\"进度异常加速是风险累积的早期信号\""
    ]

    for finding in findings:
        p = doc.add_paragraph()
        run = p.add_run(f"• {finding}")
        run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("建议：")
    run.font.bold = True
    p.add_run("将\"进度异常加速\"作为独立风险信号纳入《项目风险判断决策卡》检查表")

    doc.add_paragraph()

    # 签名
    p = doc.add_paragraph()
    run = p.add_run("访谈者签名：__________    日期：2026-06-10")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-复盘访谈记录.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_training_activity_script():
    """生成训练活动脚本"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【训练活动脚本】供应商切换决策模拟")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # 基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、活动基本信息")
    run.font.bold = True
    run.font.size = Pt(12)

    table1 = doc.add_table(rows=6, cols=2)
    table1.style = 'Table Grid'
    info = [
        ["活动名称", "供应商切换决策模拟"],
        ["对应决策卡", "demo-决策卡-供应商切换.docx"],
        ["目标学员", "采购部门主管、采购专员"],
        ["活动时长", "60分钟"],
        ["学员人数", "6-12人（分组进行）"],
        ["所需材料", "场景卡、决策卡（练习版）、投影、白板"]
    ]
    for i, (key, value) in enumerate(info):
        table1.rows[i].cells[0].text = key
        table1.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 环节一
    p = doc.add_paragraph()
    run = p.add_run("二、环节一：情境代入（20分钟）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("【活动目的】")
    run.font.bold = True
    p.add_run("让学员在不知情的情况下，凭直觉做出决策，暴露其原始判断习惯")

    p = doc.add_paragraph()
    run = p.add_run("【场景描述】（分发场景卡）")
    run.font.bold = True

    p = doc.add_paragraph()
    p.add_run("""你是华通制造公司的采购主管。公司主要生产精密机械零件，核心原材料供应商为"恒达材料"（合作5年）。

最近三个月，恒达的交付准时率出现下降：1月准时率92%，2月降至88%，3月进一步降至85%。同时，你收到了一份质量异常报告：恒达提供的A规格钢材，在近期的来料检验中，有2批次的硬度指标连续超出公司规定的波动区间。

恒达销售经理解释称，问题是"春节后工人更替导致短期内波动"，已制定整改措施，计划在4月中旬完成整改。

与此同时，你的备选供应商"新锐材料"表示可以承接订单。新锐材料的报价比恒达高8%，但其销售负责人表示"产能充足，可以随时启动"。

距离4月15日的月度采购计划会还有3天，你需要决定：
A. 立即启动切换至新锐材料
B. 维持与恒达的合作，但要求其提供书面整改承诺
C. 先不决定，继续观察两周

请写下你的决定，并说明理由。""")

    p = doc.add_paragraph()
    run = p.add_run("【干扰信息说明】")
    run.font.bold = True
    p.add_run("以下信息包含在场景中，但未明确提示其重要性：")
    p = doc.add_paragraph()
    p.add_run("• 新锐材料的【销售负责人表示产能充足】——未提供书面凭证")
    p = doc.add_paragraph()
    p.add_run("• 恒达的整改措施是【计划在4月中旬完成】——尚未实施")
    p = doc.add_paragraph()
    p.add_run("• A规格钢材是核心产品用材，直接影响产品质量")

    p = doc.add_paragraph()
    run = p.add_run("【引导话术】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    p = doc.add_paragraph()
    p.add_run("\"好，现在请每个人独立思考，写下你的决定。不要跟别人讨论，也不要翻看任何资料。凭你自己的第一判断。\"")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    p.add_run("（等待5分钟，确保所有人完成决策）")

    doc.add_paragraph()

    # 环节二
    p = doc.add_paragraph()
    run = p.add_run("三、环节二：揭示与对照（20分钟）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("【活动目的】")
    run.font.bold = True
    p.add_run("让学员发现自己的判断与决策卡内容之间的差距，体会\"如果早点看这张卡就好了\"")

    p = doc.add_paragraph()
    run = p.add_run("【揭示话术】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    p = doc.add_paragraph()
    p.add_run("\"刚才的场景，其实对应的是一张真实的决策卡——《供应商切换决策卡》。现在把这张卡发给大家。\"")
    p.style = doc.styles['Quote']

    p = doc.add_paragraph()
    p.add_run("（分发demo-决策卡-供应商切换.docx的练习版）")

    p = doc.add_paragraph()
    run = p.add_run("【对照引导】")
    run.font.bold = True

    p = doc.add_paragraph()
    p.add_run("请对照决策卡，回答以下问题：")

    questions = [
        "决策卡规定的触发条件是什么？我们刚才的场景是否符合触发条件？",
        "决策卡的检查表有5项，你刚才做了几项？漏掉了哪几项？",
        "决策卡建议的应急方案是什么？你选择的方案有什么不同？",
        "如果你在决策前先看了这张卡，你的判断会不会不一样？"
    ]

    for q in questions:
        p = doc.add_paragraph()
        p.add_run(f"• {q}")

    p = doc.add_paragraph()
    run = p.add_run("【学员常见反应预判】")
    run.font.bold = True
    p.add_run("• \"我考虑了备选供应商的交付能力，但没核实书面凭证\"")
    p = doc.add_paragraph()
    p.add_run("• \"我没注意到磨合期对产能的影响\"")
    p = doc.add_paragraph()
    p.add_run("• \"我以为恒达的整改承诺是可信的\"——忽视\"计划\"≠\"已完成\"")

    doc.add_paragraph()

    # 环节三
    p = doc.add_paragraph()
    run = p.add_run("四、环节三：二次决策变体场景（15分钟）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("【变体场景描述】")
    run.font.bold = True

    p = doc.add_paragraph()
    p.add_run("""变体场景（与原场景的关键差异已标红）：

你是华通制造公司的采购主管。核心原材料供应商"恒达材料"最近交付准时率下降至85%，同时出现2批次质量异常。

恒达销售经理解释后，你要求其提供书面整改承诺。恒达在2天内提交了一份正式的整改方案，包含：
① 更换质检团队负责人（已完成）
② 引入新的热处理工艺（设备已到位，调试中）
③ 预计4月底前完成全部整改

备选供应商"新锐材料"报价比恒达高8%。你联系新锐材料要求提供近3个月的实际交付数据，新锐提供了详细的出货记录，显示近3个月准时率均为98%以上。

此时，距离月度采购计划会还有3天。你的决定是？
A. 立即切换至新锐材料
B. 给恒达一次机会，按整改方案观察至4月底
C. 启动"双供应商并行"模式：新锐先以30%小批量试单""")

    p = doc.add_paragraph()
    run = p.add_run("【变体设计意图】")
    run.font.bold = True
    p.add_run("• 原场景：恒达仅有口头整改承诺，未提供书面方案")
    p = doc.add_paragraph()
    run = p.add_run("• 变体场景：恒达提供了书面整改方案，且部分措施已实施")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    p = doc.add_paragraph()
    p.add_run("• 检验学员是否能识别：\"书面方案\"≠\"整改完成\"，关键看整改措施是否真正落地")

    p = doc.add_paragraph()
    run = p.add_run("【二次决策引导话术】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    p = doc.add_paragraph()
    p.add_run("\"好，现在请大家打开决策卡，对照检查表，再做一次决策。\"")
    p.style = doc.styles['Quote']

    doc.add_paragraph()

    # 总结与反馈
    p = doc.add_paragraph()
    run = p.add_run("五、活动总结（5分钟）")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("【核心教训】")
    run.font.bold = True
    p.add_run("1. \"口头承诺\"是最不可信的信号——必须要求书面凭证")
    p = doc.add_paragraph()
    p.add_run("2. \"整改计划\"≠\"整改完成\"——关键是看措施是否已落地")
    p = doc.add_paragraph()
    p.add_run("3. 切换供应商最大的风险是\"磨合期\"——必须有缓冲方案")

    p = doc.add_paragraph()
    run = p.add_run("【课后行动】")
    run.font.bold = True
    p.add_run("请学员回顾各自在实际工作中遇到的供应商切换场景，尝试使用决策卡进行复盘")

    doc.add_paragraph()

    # 讲师备注
    p = doc.add_paragraph()
    run = p.add_run("六、讲师备注")
    run.font.bold = True
    run.font.size = Pt(12)

    notes = [
        "活动开始前不要透露这是在测试决策卡的使用",
        "环节一的等待时间要足够，确保每个人都有时间思考",
        "环节二的发问要引导学员自己发现差距，不要直接指出",
        "二次决策后可以公布正确答案（选项C），但不要强制统一",
        "延迟追踪：建议1个月后收集学员在实际工作中使用决策卡的情况"
    ]

    for note in notes:
        p = doc.add_paragraph()
        run = p.add_run(f"• {note}")
        run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-训练活动脚本-供应商切换.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_failure_case_warning():
    """生成失败案例警示"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【失败案例警示】某项目风险判断失误")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # 匿名化声明
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("【声明】本案例已做匿名化处理，仅用于教学警示")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # 基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、案例基本信息")
    run.font.bold = True
    run.font.size = Pt(12)

    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Table Grid'
    info = [
        ["案例类型", "项目风险判断失误"],
        ["发生时间", "2025年XX月（具体时间已模糊化）"],
        ["行业", "制造业"],
        ["涉及金额", "约300万元"],
        ["匿名化等级", "高度匿名（无公司名、无决策者姓名）"]
    ]
    for i, (key, value) in enumerate(info):
        table1.rows[i].cells[0].text = key
        table1.rows[i].cells[1].text = value

    doc.add_paragraph()

    # 岔路口描述
    p = doc.add_paragraph()
    run = p.add_run("二、岔路口描述")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    p.add_run("某制造企业启动了一个新产品研发项目，项目周期预计6个月。项目进行到第4个月时，研发团队遇到了一个技术难题——某核心功能的技术方案在原定方向上遇到了瓶颈。")

    p = doc.add_paragraph()
    run = p.add_run("当时存在三个选项：")
    run.font.bold = True

    options = [
        ("选项A（最终选择）", "继续沿用原技术方案，同时简化部分测试流程，争取按原计划完成"),
        ("选项B（被排除）", "暂停项目，重新评估技术方案，预计延期2个月"),
        ("选项C（未被充分考虑）", "引入外部技术资源，但需要额外增加30%预算")
    ]

    for opt, desc in options:
        p = doc.add_paragraph()
        run = p.add_run(f"• {opt}：{desc}")

    p = doc.add_paragraph()
    p.add_run("项目负责人最终选择了选项A。理由是：\"延期会影响客户信任，而且技术上应该能解决。\"")

    doc.add_paragraph()

    # 被忽略的信号
    p = doc.add_paragraph()
    run = p.add_run("三、被忽略的信号")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    p = doc.add_paragraph()
    run = p.add_run("信号1：进度异常加速")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    p = doc.add_paragraph()
    p.add_run("项目第3个月的进度报告显示，完成率比计划快了约10天。当时的周报解释是\"开发效率超预期\"。实际上，这是因为团队在第2个月就开始\"加班赶工\"来弥补技术难题造成的时间损失，但周报中没有如实反映这一点。")

    p = doc.add_paragraph()
    run = p.add_run("信号2：核心开发人员的异动")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    p = doc.add_paragraph()
    p.add_run("项目第3个月中，核心开发人员张某向项目经理透露\"技术难度比预期大很多\"，并提出想调整项目范围。但这条信息没有被正式记录，也没有引起重视。")

    p = doc.add_paragraph()
    run = p.add_run("信号3：质量测试数据异常")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    p = doc.add_paragraph()
    p.add_run("第3个月的测试报告显示，单元测试覆盖率从80%下降至65%。当时技术负责人解释说\"是因为新增代码量大，覆盖率统计有延迟\"。实际上是因为测试工作被压缩了。")

    p = doc.add_paragraph()
    run = p.add_run("【信号组合效应】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    p = doc.add_paragraph()
    p.add_run("进度异常加速 + 核心人员异动预警 + 质量数据下滑，三个信号单独看都不构成强烈警示，但组合在一起时，是一个典型的\"项目失控前期\"模式。")

    doc.add_paragraph()

    # 对照问题
    p = doc.add_paragraph()
    run = p.add_run("四、对照自查问题")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    questions = [
        "你现在正在推进的项目，有没有出现过\"进度异常加速\"的情况？如果有，原因是什么？",
        "有没有核心团队成员表达过\"难度比预期大\"或类似的担忧？这个信息有没有被正式记录和跟进？",
        "项目的质量指标（测试覆盖率、缺陷数等）有没有出现下滑趋势？"
    ]

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q}")
        run.font.bold = True

    p = doc.add_paragraph()
    run = p.add_run("如果这三个问题中你有任何一个回答\"是\"，建议立即启动《项目风险判断决策卡》的自查流程。")

    doc.add_paragraph()

    # 后果
    p = doc.add_paragraph()
    run = p.add_run("五、后果概述（已匿名化）")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("项目最终延期4个月交付，比原始计划多消耗预算约80%。客户因此取消了下季度的订单，估算损失超过300万元。")

    p = doc.add_paragraph()
    p.add_run("事后复盘发现，如果当时选择选项B（暂停重新评估），项目只需要延期2个月，整体损失可以减少约60%。")

    doc.add_paragraph()

    # 教训总结
    p = doc.add_paragraph()
    run = p.add_run("六、核心教训")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    lessons = [
        ("\"进度快\"不一定是好事", "异常加速往往是用未来的时间换来的，掩盖了潜在风险"),
        ("核心人员的异动是重要信号", "技术负责人或核心开发人员的担忧，应该被认真对待和记录"),
        ("质量数据是项目的体温计", "测试覆盖率、缺陷密度等质量指标下降，是风险的早期信号"),
        ("\"继续推进\"不是中性选择", "当多个风险信号同时出现时，\"继续推进\"实际上是一个高风险选择")
    ]

    for title, desc in lessons:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}：{desc}")
        run.font.bold = True

    doc.add_paragraph()

    # 关联决策卡
    p = doc.add_paragraph()
    run = p.add_run("七、关联决策卡")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("本案例警示内容已嵌入以下决策卡：")
    p = doc.add_paragraph()
    run = p.add_run("• demo-决策卡-项目风险判断.docx")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 本文档为教学demo示例 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-失败案例警示-某项目风险判断.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def create_user_guide():
    """生成demo使用说明"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.79)
    section.right_margin = Inches(0.79)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【使用说明】AI时代决策工作手册Demo文件包")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # 概述
    p = doc.add_paragraph()
    run = p.add_run("一、文件包概述")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("本文件包是《AI时代决策工作手册》的完整教学Demo，包含7个演示文件和1份使用说明。所有文件均为教学目的设计，内容经过脱敏处理，可在实际培训中直接使用。")

    doc.add_paragraph()

    # 文件清单
    p = doc.add_paragraph()
    run = p.add_run("二、文件清单")
    run.font.bold = True
    run.font.size = Pt(12)

    files = [
        ("demo-决策卡-供应商切换.docx", "Word", "3页", "完整的供应商切换决策卡，包含触发条件、检查表、应急方案、警示案例"),
        ("demo-决策卡-项目风险判断.docx", "Word", "3页", "项目风险判断决策卡，用于识别项目执行中的风险信号"),
        ("demo-决策卡-流程异常识别.docx", "Word", "3页", "流程异常识别决策卡，用于发现日常流程中的异常模式"),
        ("demo-场景映射矩阵-供应商管理.xlsx", "Excel", "2个Sheet", "供应商管理场景映射矩阵，包含参考版和练习版"),
        ("demo-稽核对话记录.docx", "Word", "4页", "模拟的稽核对话全文，包含四个稽核问题和判断标注"),
        ("demo-复盘访谈记录.docx", "Word", "5页", "模拟的一次完整复盘访谈，三个维度追问和隐性判断挖掘"),
        ("demo-训练活动脚本-供应商切换.docx", "Word", "5页", "完整的三环节训练活动设计，包含引导话术和变体场景"),
        ("demo-失败案例警示-某项目风险判断.docx", "Word", "4页", "匿名化处理的失败案例，包含岔路口和被忽略的信号"),
        ("demo-使用说明.docx", "Word", "本文件", "讲师使用指南")
    ]

    table1 = doc.add_table(rows=len(files)+1, cols=4)
    table1.style = 'Table Grid'

    headers = ["文件名", "格式", "篇幅", "内容说明"]
    for i, header in enumerate(headers):
        table1.rows[0].cells[i].text = header
        set_cell_shading(table1.rows[0].cells[i], "D9D9D9")

    for i, (name, fmt, pages, desc) in enumerate(files):
        table1.rows[i+1].cells[0].text = name
        table1.rows[i+1].cells[1].text = fmt
        table1.rows[i+1].cells[2].text = pages
        table1.rows[i+1].cells[3].text = desc

    doc.add_paragraph()

    # 使用建议
    p = doc.add_paragraph()
    run = p.add_run("三、讲师使用建议")
    run.font.bold = True
    run.font.size = Pt(12)

    suggestions = [
        ("决策卡的使用顺序", "建议先讲解《供应商切换决策卡》，因为这个场景最直观，学员最容易理解。然后过渡到《项目风险判断》和《流程异常识别》，让学员体会不同场景的判断逻辑差异"),
        ("场景映射矩阵的用法", "Excel文件包含两个Sheet——参考版和练习版。建议先发练习版让学员填写，然后再发参考版进行对照。练习版可用于小组讨论后的成果检验"),
        ("稽核对话记录的用途", "这份记录展示了如何在实际稽核中运用四个问题。讲师可以请学员模拟演练，先读一遍对话记录，然后分组扮演稽核人和决策者进行练习"),
        ("复盘访谈记录的用法", "这份记录的价值在于展示\"隐性判断是如何被问出来的\"。建议讲师先完整展示访谈过程，然后请学员识别其中的追问技巧，最后让学员尝试自己设计追问问题"),
        ("训练活动脚本的使用", "这是整个Demo包的核心产出。讲师可以按照脚本完整执行一次训练活动（60分钟），让学员在\"经历\"中体会决策卡的价值")
    ]

    for title, desc in suggestions:
        p = doc.add_paragraph()
        run = p.add_run(f"【{title}】")
        run.font.bold = True
        p.add_run(f"\n{desc}")

    doc.add_paragraph()

    # 教学设计逻辑
    p = doc.add_paragraph()
    run = p.add_run("四、教学设计逻辑")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("本Demo文件包遵循以下教学逻辑：")

    logic = [
        "认知层：先让学员理解\"复盘写的是过去，决策卡写的是未来\"这个核心区分（可通过决策卡文档自学）",
        "方法层：通过复盘访谈记录和稽核对话记录，展示决策卡制作的方法论工具",
        "应用层：通过训练活动脚本，让学员在模拟场景中体验决策卡的实际使用",
        "警示层：通过失败案例警示，强化风险识别意识和\"对照\"而非\"恐惧\"的心态"
    ]

    for i, step in enumerate(logic, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {step}")

    p = doc.add_paragraph()
    p.add_run("建议授课顺序：决策卡自学 → 方法论讲解 → 训练活动体验 → 案例警示 → 稽核练习")

    doc.add_paragraph()

    # 关键标注说明
    p = doc.add_paragraph()
    run = p.add_run("五、红色标注说明")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("所有文件中，红色标注的内容为关键信息，包括：")

    highlights = [
        "触发条件中的核心判断句",
        "检查表中最重要的核对项",
        "警示案例中被忽略的信号",
        "访谈记录中浮现的隐性判断",
        "稽核对话中的关键判断点"
    ]

    for h in highlights:
        p = doc.add_paragraph()
        run = p.add_run(f"• {h}")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    doc.add_paragraph()

    # 注意事项
    p = doc.add_paragraph()
    run = p.add_run("六、注意事项")
    run.font.bold = True
    run.font.size = Pt(12)

    notes = [
        "所有案例均为教学设计的模拟案例，已做匿名化处理，不指向任何真实企业或个人",
        "决策卡中的\"认领人\"信息为虚构信息，实际使用时需要替换为真实的认领人",
        "场景映射矩阵中的参数为参考值，实际使用时需要根据企业具体业务调整",
        "训练活动脚本中的\"正确答案\"仅作为参考，训练目标不是统一答案，而是让学员体验决策过程"
    ]

    for note in notes:
        p = doc.add_paragraph()
        p.add_run(f"• {note}")

    doc.add_paragraph()

    # 联系方式
    p = doc.add_paragraph()
    run = p.add_run("七、定制化服务")
    run.font.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("如需将本Demo文件包定制化为您所在组织的实际业务场景，欢迎联系《AI时代决策工作手册》作者团队获取支持。")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Document —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    filename = os.path.join(OUTPUT_DIR, "demo-使用说明.docx")
    doc.save(filename)
    print(f"已生成: {filename}")


def main():
    """主函数"""
    print("=" * 50)
    print("开始生成AI时代决策工作手册Demo文件...")
    print("=" * 50)

    # 生成所有文件
    create_decision_card_supplier()
    create_decision_card_risk_judgment()
    create_decision_card_process_anomaly()
    create_scenario_mapping_matrix()
    create_audit_dialogue_record()
    create_interview_record()
    create_training_activity_script()
    create_failure_case_warning()
    create_user_guide()

    print("=" * 50)
    print("所有Demo文件生成完成！")
    print(f"输出目录：{OUTPUT_DIR}")
    print("=" * 50)


if __name__ == "__main__":
    main()

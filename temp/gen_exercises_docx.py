#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_BASE = "D:/新课开发/变革管理/08-向上管理与高层说服术：让决策层理解容错的成本逻辑/完整课程包/07-练习材料"

C_PRIMARY = RGBColor(0x2b, 0x2d, 0x42)
C_SECONDARY = RGBColor(0x8d, 0x99, 0xae)
C_ACCENT = RGBColor(0xef, 0x23, 0x3c)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def create_exercise_docx(filename, title, sections):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = C_PRIMARY

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = C_PRIMARY

    for sec in sections:
        if sec['type'] == 'heading':
            doc.add_heading(sec['text'], level=sec.get('level', 1))
        elif sec['type'] == 'paragraph':
            doc.add_paragraph(sec['text'])
        elif sec['type'] == 'fill_blank':
            p = doc.add_paragraph()
            p.add_run(sec['label']).font.color.rgb = C_SECONDARY
            p.add_run('：' + '____' * 20)
        elif sec['type'] == 'table':
            table = doc.add_table(rows=sec['rows'], cols=sec['cols'])
            table.style = 'Table Grid'
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell.text = sec['data'][i][j] if i < len(sec['data']) else ''
                    if i == 0:
                        set_cell_bg(cell, "2b2d42")
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255,255,255)
                                run.font.bold = True
                    elif i % 2 == 0:
                        set_cell_bg(cell, "edf2f4")
        elif sec['type'] == 'discussion':
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            set_cell_bg(cell, "2b2d42")
            cell.text = "讨论：" + sec['text']
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255,255,255)
        elif sec['type'] == 'page_break':
            doc.add_page_break()

    out_path = f"{OUT_BASE}/{filename}"
    doc.save(out_path)
    print(f"Saved: {out_path}")

# Exercise 1
create_exercise_docx("练习一-决策者类型识别.docx", "练习一：决策者类型识别", [
    {'type': 'heading', 'text': '练习目标', 'level': 1},
    {'type': 'paragraph', 'text': '学会识别三种决策者类型（财务导向型/关系导向型/战略导向型），理解不同类型决策者的核心关切。'},
    {'type': 'heading', 'text': '练习说明', 'level': 1},
    {'type': 'paragraph', 'text': '阅读以下三个高管简介，判断每位高管的决策类型，并说明你的判断依据。'},
    {'type': 'heading', 'text': '高管一：王总', 'level': 2},
    {'type': 'paragraph', 'text': '某制造企业CEO，55岁，在位10年。每次开会必问投入产出比，对新项目非常谨慎。'},
    {'type': 'fill_blank', 'label': '你的判断', 'text': ''},
    {'type': 'fill_blank', 'label': '判断依据', 'text': ''},
    {'type': 'heading', 'text': '高管二：刘总', 'level': 2},
    {'type': 'paragraph', 'text': '某互联网公司COO，42岁，在位3年。非常关注团队士气和稳定性，重视沟通方式。'},
    {'type': 'fill_blank', 'label': '你的判断', 'text': ''},
    {'type': 'fill_blank', 'label': '判断依据', 'text': ''},
    {'type': 'heading', 'text': '高管三：陈总', 'level': 2},
    {'type': 'paragraph', 'text': '某咨询公司合伙人，48岁，在位5年。关注行业趋势和竞争对手动态，决策风格大胆。'},
    {'type': 'fill_blank', 'label': '你的判断', 'text': ''},
    {'type': 'fill_blank', 'label': '判断依据', 'text': ''},
    {'type': 'discussion', 'text': '如果要向王总汇报一个新项目，你会重点准备什么？'},
])

# Exercise 2
create_exercise_docx("练习二-Error-Cost计算.docx", "练习二：Error Cost计算", [
    {'type': 'heading', 'text': '练习目标', 'level': 1},
    {'type': 'paragraph', 'text': '掌握Error Cost的计算公式，学会量化「不变革」的成本。'},
    {'type': 'heading', 'text': 'Error Cost计算公式', 'level': 1},
    {'type': 'paragraph', 'text': 'Error Cost = 问题导致的损失 × 持续时间 × 影响范围'},
    {'type': 'heading', 'text': '第一步：识别核心问题', 'level': 2},
    {'type': 'fill_blank', 'label': '问题描述', 'text': ''},
    {'type': 'fill_blank', 'label': '持续时间', 'text': ''},
    {'type': 'fill_blank', 'label': '影响范围', 'text': ''},
    {'type': 'heading', 'text': '第二步：计算损失', 'level': 2},
    {'type': 'table', 'rows': 6, 'cols': 3, 'data': [
        ['损失类型', '具体表现', '估算金额（万元/年）'],
        ['效率损失', '人员加班、流程返工', ''],
        ['质量损失', '错误率、客诉率', ''],
        ['机会损失', '延误、错过窗口期', ''],
        ['成本损失', '直接成本浪费', ''],
        ['合计', '', ''],
    ]},
    {'type': 'heading', 'text': '第三步：计算Error Cost', 'level': 2},
    {'type': 'fill_blank', 'label': '年度总损失（万元）', 'text': ''},
    {'type': 'fill_blank', 'label': '问题持续时间（月）', 'text': ''},
    {'type': 'fill_blank', 'label': '影响范围系数（1-3）', 'text': ''},
    {'type': 'paragraph', 'text': 'Error Cost = 年度总损失 × 持续时间/12 × 影响范围 = _______万元'},
    {'type': 'discussion', 'text': '计算完成后，这个数字和你的直觉感受相比如何？'},
])

# Exercise 3
create_exercise_docx("练习三-最小授权方案设计.docx", "练习三：最小授权方案设计", [
    {'type': 'heading', 'text': '练习目标', 'level': 1},
    {'type': 'paragraph', 'text': '掌握分阶段授权路线图的设计方法，学会将大型项目拆解为可控的最小授权方案。'},
    {'type': 'heading', 'text': '第一阶段：定义完整项目', 'level': 2},
    {'type': 'fill_blank', 'label': '项目名称', 'text': ''},
    {'type': 'fill_blank', 'label': '总预算', 'text': ''},
    {'type': 'fill_blank', 'label': '预计总收益', 'text': ''},
    {'type': 'heading', 'text': '第二阶段：设计最小授权方案', 'level': 2},
    {'type': 'table', 'rows': 4, 'cols': 4, 'data': [
        ['阶段', '范围', '时间', '授权程度'],
        ['第一阶段（试点）', '', '', '执行方式自主，方向需汇报'],
        ['第二阶段（扩展）', '', '', '小幅调整预算和范围'],
        ['第三阶段（固化）', '', '', '全面授权，季度审核'],
    ]},
    {'type': 'heading', 'text': '第三阶段：设计止损边界', 'level': 2},
    {'type': 'fill_blank', 'label': '最大可承受损失', 'text': ''},
    {'type': 'fill_blank', 'label': '止损触发条件', 'text': ''},
    {'type': 'discussion', 'text': '你的高层看到这个方案，可能会质疑什么？'},
])

# Exercise 4
create_exercise_docx("练习四-说服提案包装.docx", "练习四：说服提案包装", [
    {'type': 'heading', 'text': '练习目标', 'level': 1},
    {'type': 'paragraph', 'text': '掌握用ROI语言重新包装提案的方法，学会设计最小授权方案降低决策门槛。'},
    {'type': 'heading', 'text': '第一步：决策者画像', 'level': 2},
    {'type': 'fill_blank', 'label': '决策者姓名/职位', 'text': ''},
    {'type': 'fill_blank', 'label': '决策者类型', 'text': ''},
    {'type': 'fill_blank', 'label': '核心关切', 'text': ''},
    {'type': 'heading', 'text': '第二步：Error Cost计算', 'level': 2},
    {'type': 'fill_blank', 'label': '问题描述', 'text': ''},
    {'type': 'fill_blank', 'label': '年度损失金额', 'text': ''},
    {'type': 'heading', 'text': '第三步：ROI计算', 'level': 2},
    {'type': 'table', 'rows': 5, 'cols': 2, 'data': [
        ['指标', '数值'],
        ['总投资', '_______万元'],
        ['年度收益', '_______万元'],
        ['年度净收益', '_______万元'],
        ['ROI', '_______%'],
    ]},
    {'type': 'heading', 'text': '第四步：高频应答准备', 'level': 2},
    {'type': 'fill_blank', 'label': '问题1：ROI怎么算出来的？', 'text': ''},
    {'type': 'fill_blank', 'label': '问题2：失败了怎么办？', 'text': ''},
    {'type': 'fill_blank', 'label': '问题3：团队支持吗？', 'text': ''},
    {'type': 'discussion', 'text': '演讲练习后，你觉得最难的部分是什么？'},
])

# Exercise 5
create_exercise_docx("练习五-变革失败复盘.docx", "练习五：变革失败复盘", [
    {'type': 'heading', 'text': '练习目标', 'level': 1},
    {'type': 'paragraph', 'text': '通过复盘真实案例，学习向上说服的成败关键，将本课所学方法论与实际经验对接。'},
    {'type': 'heading', 'text': '案例基本信息', 'level': 2},
    {'type': 'fill_blank', 'label': '项目名称', 'text': ''},
    {'type': 'fill_blank', 'label': '时间', 'text': ''},
    {'type': 'fill_blank', 'label': '你的角色', 'text': ''},
    {'type': 'heading', 'text': '复盘框架一：决策者分析', 'level': 2},
    {'type': 'fill_blank', 'label': '决策者类型', 'text': ''},
    {'type': 'fill_blank', 'label': '我当时用对语言了吗', 'text': ''},
    {'type': 'heading', 'text': '复盘框架二：成本量化', 'level': 2},
    {'type': 'fill_blank', 'label': '我当时有没有计算Error Cost', 'text': ''},
    {'type': 'fill_blank', 'label': '我当时有没有量化收益', 'text': ''},
    {'type': 'heading', 'text': '复盘框架三：方案设计', 'level': 2},
    {'type': 'fill_blank', 'label': '我当时是一步到位还是最小授权', 'text': ''},
    {'type': 'heading', 'text': '复盘框架四：信任建立', 'level': 2},
    {'type': 'fill_blank', 'label': '我有没有主动说风险', 'text': ''},
    {'type': 'fill_blank', 'label': '我和高层的关系基础如何', 'text': ''},
    {'type': 'heading', 'text': '深度问题', 'level': 2},
    {'type': 'paragraph', 'text': '1. 如果可以重来，你会做哪三件事不一样？'},
    {'type': 'paragraph', 'text': '2. 这次失败对你后续的变革推动有什么启示？'},
    {'type': 'paragraph', 'text': '3. 学习了本课后，你打算如何避免同样的错误？'},
    {'type': 'discussion', 'text': '通过本课学习和案例复盘，你对向上说服的理解有哪些更新？'},
])

print("All exercises created")
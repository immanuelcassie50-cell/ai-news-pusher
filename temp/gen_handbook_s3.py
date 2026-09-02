#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)
    
    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)
    
    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)
    
    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)
    
    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# Learning objectives
ot = doc.add_paragraph()
r = ot.add_run('学习目标')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph()

for label, content in [
    ('知识目标', '理解复盘与决策卡的本质区分'),
    ('技能目标', '掌握决策卡制作的全流程方法'),
    ('态度目标', '建立判断力是组织核心资产的意识'),
]:
    oh = doc.add_paragraph()
    r = oh.add_run(label); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0,102,204)
    doc.add_paragraph(content)

doc.add_page_break()
print('S3: Learning objectives done')
doc.save(r'D:\CC\temp\handbook_s3.docx')

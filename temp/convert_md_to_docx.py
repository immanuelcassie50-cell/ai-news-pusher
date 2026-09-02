# -*- coding: utf-8 -*-
"""
将 markdown 转换为带样式的 docx
支持：标题（H1-H3）、列表、表格、粗体、引用、代码块
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = sys.argv[1]
DST = sys.argv[2]
TITLE = sys.argv[3]

def set_cell_bg(cell, color_hex):
    """设置单元格背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_styled_paragraph(doc, text, style=None):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)
    return p, run

def md_to_docx(md_path, docx_path, doc_title):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体为宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    rPr = style.element.rPr
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')

    # 页边距
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # 标题样式
    h1 = doc.styles['Heading 1']
    h1.font.name = '黑体'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0xC4, 0x48, 0x1D)
    rPr1 = h1.element.rPr
    rFonts1 = rPr1.find(qn('w:rFonts'))
    if rFonts1 is None:
        rFonts1 = OxmlElement('w:rFonts')
        rPr1.insert(0, rFonts1)
    rFonts1.set(qn('w:eastAsia'), '黑体')

    h2 = doc.styles['Heading 2']
    h2.font.name = '黑体'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    h3 = doc.styles['Heading 3']
    h3.font.name = '黑体'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x2D, 0x5F, 0x4A)

    # 封面
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('\n\n\n\n\n' + doc_title)
    title_run.font.name = '黑体'
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xC4, 0x48, 0x1D)
    rPr_t = title_run._element.get_or_add_rPr()
    rFonts_t = OxmlElement('w:rFonts')
    rFonts_t.set(qn('w:eastAsia'), '黑体')
    rPr_t.insert(0, rFonts_t)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('\n\n\n知行：学习落地工作坊\n\nV1.0 · 2026年6月')
    sub_run.font.name = '宋体'
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0xB8, 0x92, 0x4A)
    rPr_s = sub_run._element.get_or_add_rPr()
    rFonts_s = OxmlElement('w:rFonts')
    rFonts_s.set(qn('w:eastAsia'), '宋体')
    rPr_s.insert(0, rFonts_s)

    doc.add_page_break()

    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                in_code_block = False
                if code_buffer:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    code_run = p.add_run('\n'.join(code_buffer))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(10)
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        if line.startswith('# '):
            text = line[2:].strip()
            # 跳过第一个一级标题（已经在封面）
            if i < 50:
                i += 1
                continue
            doc.add_heading(text, level=1)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = '黑体'
            rPr_h = run._element.get_or_add_rPr()
            rFonts_h = OxmlElement('w:rFonts')
            rFonts_h.set(qn('w:eastAsia'), '黑体')
            rPr_h.insert(0, rFonts_h)
            i += 1
            continue

        # 引用
        if line.startswith('>'):
            text = line[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.7)
            # 添加左边框
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left_bdr = OxmlElement('w:left')
            left_bdr.set(qn('w:val'), 'single')
            left_bdr.set(qn('w:sz'), '24')
            left_bdr.set(qn('w:color'), 'C4481D')
            pBdr.append(left_bdr)
            pPr.append(pBdr)
            run = p.add_run(text)
            run.font.name = '宋体'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.italic = True
            rPr_q = run._element.get_or_add_rPr()
            rFonts_q = OxmlElement('w:rFonts')
            rFonts_q.set(qn('w:eastAsia'), '宋体')
            rPr_q.insert(0, rFonts_q)
            i += 1
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|[\s\-:|]+\|', lines[i + 1]):
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            # 解析表格
            if len(table_lines) >= 2:
                rows_data = []
                for tl in table_lines:
                    if re.match(r'^\s*\|[\s\-:|]+\|\s*$', tl):
                        continue
                    cells = [c.strip() for c in tl.split('|')[1:-1]]
                    rows_data.append(cells)
                if rows_data:
                    ncols = len(rows_data[0])
                    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
                    tbl.style = 'Light Grid Accent 1'
                    for ri, row_data in enumerate(rows_data):
                        for ci, cell_text in enumerate(row_data):
                            if ci < ncols:
                                cell = tbl.rows[ri].cells[ci]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_text)
                                run.font.name = '宋体'
                                run.font.size = Pt(10)
                                if ri == 0:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    set_cell_bg(cell, '1A1A1A')
                                else:
                                    if ri % 2 == 0:
                                        set_cell_bg(cell, 'F5F1E8')
                    doc.add_paragraph()
            continue

        # 列表
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            # 处理嵌套
            indent = len(line) - len(line.lstrip())
            p = doc.add_paragraph(style='List Bullet')
            if indent > 2:
                p.paragraph_format.left_indent = Cm(1.0 * (indent // 2))
            # 处理行内格式
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline_runs(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_inline_runs(p, line)
        i += 1

    doc.save(docx_path)
    print(f"[OK] Generated: {docx_path}")


def add_inline_runs(p, text):
    """处理行内格式：粗体、代码、链接"""
    # 分割 **bold** `code` [text](url)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = '宋体'
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '黑体')
            rPr.insert(0, rFonts)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('['):
            # 简化处理链接
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run = p.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0xC4, 0x48, 0x1D)
                run.font.underline = True
            else:
                run = p.add_run(part)
        else:
            run = p.add_run(part)
            run.font.name = '宋体'
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '宋体')
            rPr.insert(0, rFonts)


if __name__ == '__main__':
    md_to_docx(SRC, DST, TITLE)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

output_path = "D:/CC/temp/course_intro_output.docx"

doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia', 'SimSun')

# Cover page
title = doc.add_heading('校招导师赋能', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(28)
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle = doc.add_paragraph('用一张画布构建深度辅导框架')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(18)
subtitle.runs[0].font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph('课程介绍')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()
客户 = doc.add_paragraph('适用客户：国企/央企校招导师培训 | 金融机构师徒制 | 科技公司新人带教')
客户.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
版本 = doc.add_paragraph('版本：1.0')
版本.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Section 1: 课程定位
doc.add_heading('一、课程定位', level=1)
doc.add_heading('一句话定位', level=2)

定位 = doc.add_paragraph()
定位.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = 定位.add_run('一门帮助校招导师从「凭经验带人」切换到「用框架辅导人」的方法课——用一张诊断画布读懂新人、设计辅导动作并形成闭环。')
run.italic = True
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_heading('解决的问题', level=2)
问题列表 = [
    "校招新人「参差不齐」——同样的培养机制，为什么有人三个月达标，有人一年还是「扶不起来」？",
    "导师「凭直觉辅导」——发现学员有问题，但不知道从哪里切入，东一榔头西一棒子",
    "辅导「石沉大海」——导师很用心，学员也感动，但三个月后说不出学员到底进步了多少",
    "经验「难以复制」——好导师带出来的徒弟呱呱叫，换个导师就「碰运气」"
]
for q in 问题列表:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(q)

doc.add_page_break()

# Section 2: 目标学员
doc.add_heading('二、目标学员', level=1)
doc.add_heading('核心学员', level=2)
核心学员 = ["国企/央企校招导师", "金融机构师徒制中的指导人", "科技公司新人带教负责人"]
for c in 核心学员:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(c)

doc.add_heading('学员画像', level=2)
画像 = ["有带校招新人的实际任务", "希望从「凭经验」升级到「有框架」", "愿意用真实案例参与实操"]
for m in 画像:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(m)

doc.add_page_break()

# Section 3: 学员收益
doc.add_heading('三、学员收益', level=1)
doc.add_heading('能力收益', level=2)
能力收益 = [
    "看得准——用冰山模型系统理解新人，从「看行为」升级到「看动机」",
    "治得准——用差距诊断画布区分技能/知识/态度差距，不同差距不同疗法",
    "落得实——用辅导动作设计表让每次辅导都具体、可执行、有反馈",
    "闭得环——用跟进复盘画布建立辅导闭环，让辅导结果可衡量"
]
for i, n in enumerate(能力收益, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(n)

doc.add_heading('工具收益', level=2)
工具收益 = [
    "带走一套完整的深度辅导画布（基于真实学员的完整辅导案例）",
    "带走四张可立即使用的工具：新人画像冰山图、差距诊断画布、辅导动作设计表、跟进复盘画布",
    "带走一套可复制的辅导方法论，后续带任何新人都可以用"
]
for t in 工具收益:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)

doc.add_page_break()

# Section 4: 课程内容概览
doc.add_heading('四、课程内容概览', level=1)
doc.add_heading('课程框架：一画四步', level=2)

# Table for framework
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '模块'
hdr_cells[1].text = '核心内容'
hdr_cells[2].text = '工具产出'

框架 = [
    ("第一步：看人", "校招新人的三重差异分析\n冰山模型五层次解读", "新人画像冰山图"),
    ("第二步：诊断差距", "三类差距识别\n优先级判定矩阵", "差距诊断画布"),
    ("第三步：设计动作", "辅导动作四要素\n教练式提问 vs 告知式建议", "辅导动作设计表"),
    ("第四步：跟进复盘", "跟进机制建立\n结构化复盘七步法", "跟进复盘画布"),
]

for i, (模块, 内容, 产出) in enumerate(框架, 1):
    row = table.rows[i].cells
    row[0].text = 模块
    row[1].text = 内容
    row[2].text = 产出

doc.add_paragraph()
说明 = doc.add_paragraph()
run = 说明.add_run('全课用一张「深度辅导画布」串联四个模块，学员最终产出一份完整的辅导案例，可在实际工作中直接使用。')
run.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# Section 5: 课程时长
doc.add_heading('五、课程时长', level=1)
doc.add_heading('完整版：1天（6小时）', level=2)

时间表 = doc.add_table(rows=3, cols=3)
时间表.style = 'Table Grid'
时间表.rows[0].cells[0].text = '时间段'
时间表.rows[0].cells[1].text = '内容'
时间表.rows[0].cells[2].text = '时长'
时间表.rows[1].cells[0].text = '上午'
时间表.rows[1].cells[1].text = '第一章：辅导的起点（看人）\n第二章：诊断差距'
时间表.rows[1].cells[2].text = '90分钟\n90分钟'
时间表.rows[2].cells[0].text = '下午'
时间表.rows[2].cells[1].text = '第三章：设计辅导动作\n第四章：跟进与复盘\n第五章：综合实战与路演'
时间表.rows[2].cells[2].text = '120分钟\n90分钟\n90分钟'

doc.add_paragraph()
doc.add_heading('半天版：3小时（压缩版）', level=2)
半天说明 = [
    "适合工作节奏紧张、无法安排全天培训的企业",
    "保留核心工具教学和角色扮演，压缩分享环节"
]
for d in 半天说明:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d)

doc.add_page_break()

# Section 6: 交付形式
doc.add_heading('六、交付形式', level=1)

doc.add_heading('学员配套', level=2)
学员配套 = [
    "课程工作手册（含所有工具模板）",
    "工具速查卡（可撕下随身携带）",
    "完整版深度辅导画布（课堂产出）"
]
for t in 学员配套:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)

doc.add_heading('讲师配套', level=2)
讲师配套 = [
    "完整版课程幻灯片",
    "讲师手册（含教学指南、示范案例）",
    "学员工具电子版（可打印）"
]
for t in 讲师配套:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(t)

doc.add_heading('交付方式', level=2)
交付方式 = [
    "线下面授（首选）",
    "线上直播+工作坊（支持异地学员）"
]
for d in 交付方式:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(d)

doc.add_page_break()

# Footer
footer1 = doc.add_paragraph('如需进一步沟通课程细节，请联系课程顾问')
footer1.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer2 = doc.add_paragraph('课程版权归属：罗宏伟 | 版本1.0')
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save document
doc.save(output_path)
print(f"Document created: {output_path}")

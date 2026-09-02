# -*- coding: utf-8 -*-
"""
Create 数字工具包.md document for 薪酬激励设计 course
"""
import os
import sys

# Ensure UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_PATH = r"D:\新课开发\管理学\13-薪酬激励设计\09_数字工具包\数字工具包.md"

doc = Document()

# ─── Page setup: A4, narrow margins ───────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(11.69)   # A4 landscape-ish but we'll use Letter
section.page_height = Inches(8.27)
section.left_margin   = Inches(0.8)
section.right_margin  = Inches(0.8)
section.top_margin    = Inches(0.7)
section.bottom_margin = Inches(0.7)

# ─── Color palette ─────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1F, 0x38, 0x64)   # headings / accent
BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings
GRAY    = RGBColor(0x40, 0x40, 0x40)   # body text
LIGHT_BG = RGBColor(0xF2, 0xF2, 0xF2)  # table header bg
ORANGE  = RGBColor(0xC5, 0x50, 0x0E)   # highlight / callout
GREEN   = RGBColor(0x37, 0x86, 0x3C)   # positive
RED     = RGBColor(0xC0, 0x00, 0x00)   # negative / warning

# ─── Helper: set run font ───────────────────────────────────────────────────
def set_run_font(run, size_pt, bold=False, color=None, east_asia=None):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run.font.name = "微软雅黑"
    if east_asia:
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), east_asia)

# ─── Helper: paragraph spacing ─────────────────────────────────────────────
def body_para(doc, text, size=10.5, bold=False, color=GRAY,
              before=0, after=80, indent=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing       = 1.15
    p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Inches(indent)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    return p

# ─── Helper: heading paragraph ─────────────────────────────────────────────
def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6 if level == 1 else 4)
    if level == 1:
        p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 16, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(run, 13, bold=True, color=NAVY)
    else:
        set_run_font(run, 11, bold=True, color=BLUE)
    return p

# ─── Helper: bullet ─────────────────────────────────────────────────────────
def bullet(doc, text, level=0, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    set_run_font(run, size, color=GRAY)
    return p

# ─── Helper: table ──────────────────────────────────────────────────────────
def add_table(doc, headers, rows, col_widths=None):
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, 9.5, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, '2E75B6')

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, 9, color=GRAY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                set_cell_bg(cell, 'EEF3FA')

    # Column widths
    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    return table

# ─── Helper: callout box ────────────────────────────────────────────────────
def callout(doc, text, label="提示", color=ORANGE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)
    run_label = p.add_run(f"【{label}】")
    set_run_font(run_label, 9.5, bold=True, color=color)
    run_text = p.add_run(text)
    set_run_font(run_text, 9.5, color=GRAY)
    return p

# ════════════════════════════════════════════════════════════════════════════
#  TITLE
# ════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(4)
run = p_title.add_run("数字工具包")
set_run_font(run, 22, bold=True, color=NAVY)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(2)
run = p_sub.add_run("薪酬激励设计 · 配套工具集")
set_run_font(run, 12, color=BLUE)

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.paragraph_format.space_before = Pt(0)
p_desc.paragraph_format.space_after  = Pt(16)
run = p_desc.add_run("Excel模板 / 计算公式 / 使用说明 / 案例数据")
set_run_font(run, 10, color=GRAY)

# ════════════════════════════════════════════════════════════════════════════
# TOOL 1 — 宽带薪酬设计模板
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具一：宽带薪酬设计模板", 1)

body_para(doc, "适用场景：企业完成岗位价值评估后，用本模板将岗位归入对应薪等，建立宽带薪酬体系。", size=10, color=GRAY, after=6)

heading(doc, "1.1 模板结构说明", 2)

# Table: template structure
headers = ["Sheet", "用途", "关键字段"]
rows = [
    ["薪等定义表", "设定每个薪等的人数占比、最高/最低薪资", "薪等编号、人数占比下限、上限、中位值、级差%"],
    ["岗位归等表", "将岗位映射到对应薪等", "岗位名称、职级序列、薪等、当前薪资、建议调整幅度"],
    ["带宽设计表", "设定每个薪等的带宽和重叠规则", "薪等、带宽下限、带宽上限、重叠幅度、参考市场分位值"],
    ["套薪测算表", "将现有员工套入新薪酬体系", "员工姓名、岗位、现薪、套入薪等、套入级次、新薪、差异额"],
]
add_table(doc, headers, rows, col_widths=[1.1, 1.8, 3.5])

heading(doc, "1.2 核心计算公式", 2)

body_para(doc, "① 带宽中位值:", bold=True, size=10, after=2)
body_para(doc, "   中位值 = 带宽下限 + (带宽上限 - 带宽下限) × 50%", size=9.5, color=GRAY, indent=0.2, after=4)

body_para(doc, "② 级差（每一级之间的薪资增幅）:", bold=True, size=10, after=2)
body_para(doc, "   级差% = (上一级中位值 - 本级中位值) / 本级中位值 × 100%", size=9.5, color=GRAY, indent=0.2, after=4)

body_para(doc, "③ 套薪测算:", bold=True, size=10, after=2)
body_para(doc, "   新薪 = MAX(现薪, 薪等下限) + 级差 × MIN(现薪 - 薪等下限, 带宽) / 带宽", size=9.5, color=GRAY, indent=0.2, after=4)

body_para(doc, "④ 重叠度检查:", bold=True, size=10, after=2)
body_para(doc, "   重叠度 = (相邻薪等上限 - 本薪等下限) / (本薪等上限 - 本薪等下限) × 100%\n   推荐重叠度控制在 20%~50%，过高导致晋升激励弱化，过低导致薪酬刚性强。", size=9.5, color=GRAY, indent=0.2, after=6)

heading(doc, "1.3 使用步骤", 2)
steps = [
    "Step 1：在「薪等定义表」设定 5~7 个薪等，按市场分位值（P25/P50/P75）填入各等上限/下限",
    "Step 2：在「岗位归等表」依据岗位评估分数将岗位归入对应薪等",
    "Step 3：在「带宽设计表」设定带宽（通常为 50%~80%）和级差（通常为 8%~15%）",
    "Step 4：在「套薪测算表」导入员工现薪，系统自动计算新薪及差异额",
    "Step 5：人工审核套薪结果，重点关注高薪低就（薪酬超出薪等上限）和低薪高就员工",
]
for s in steps:
    bullet(doc, s, size=9.5)

callout(doc, "建议搭配岗位价值评估结果使用。无评估结果时，可先用内部相对比较法粗分职级序列，再套用本模板。", "使用前提")

# ════════════════════════════════════════════════════════════════════════════
# TOOL 2 — 股权激励成本计算器
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具二：股权激励成本计算器", 1)

body_para(doc, "适用场景：财务/HR用本工具计算期权授予的公允价值、年度成本摊销及股权稀释比例，为财报披露和激励方案设计提供数据支撑。", size=10, color=GRAY, after=6)

heading(doc, "2.1 期权价值计算 — Black-Scholes模型", 2)

body_para(doc, "参数说明:", bold=True, size=10, after=2)
params = [
    "S = 当前股价（或评估公允价值）",
    "K = 行权价",
    "T = 期权有效期（年）",
    "r = 无风险利率（参考同期国债收益率）",
    "σ = 股价波动率（历史年化标准差）",
]
for p in params:
    bullet(doc, p, size=9.5)

body_para(doc, "期权价值公式:", bold=True, size=10, after=2)
body_para(doc, "C = S × N(d₁) - K × e^(-rT) × N(d₂)", size=10, color=BLUE, indent=0.2, after=2)
body_para(doc, "d₁ = [ln(S/K) + (r + σ²/2) × T] / (σ × √T)", size=9.5, color=GRAY, indent=0.2, after=2)
body_para(doc, "d₂ = d₁ - σ × √T", size=9.5, color=GRAY, indent=0.2, after=4)
body_para(doc, "其中 N(·) 为标准正态分布累计概率函数，可用 Excel =NORM.S.DIST(x, TRUE) 计算。", size=9.5, color=GRAY, indent=0.2, after=6)

heading(doc, "2.2 稀释比例计算", 2)

headers2 = ["指标", "公式", "说明"]
rows2 = [
    ["授予期权总数", "—", "方案设定的期权总量"],
    ["稀释前总股本", "—", "公司现有流通股本（股数）"],
    ["稀释后总股本", "= 稀释前总股本 + 稀释期权对应的股数", "考虑行权后新增股份"],
    ["全面摊薄稀释率", "= 授予期权总数 / 稀释后总股本 × 100%", "IFRS 2 / ASC 718 要求"],
    ["加权平均稀释率", "= 稀释期权 × 归属期剩余月数 / 12 / 稀释前总股本 × 100%", "年度财报计算"],
]
add_table(doc, headers2, rows2, col_widths=[1.5, 2.5, 2.4])

heading(doc, "2.3 成本摊销计算", 2)

body_para(doc, "年度摊销 = 授予期权公允价值 × 归属期月数 / 总授予月数", bold=True, size=10, after=3)
body_para(doc, "示例：某员工获授 10,000 股期权，公允价值 5 元/股，分 4 年归属（48 个月），当年在岗 9 个月：", size=9.5, color=GRAY, indent=0.2, after=2)
body_para(doc, "年度摊销 = 10,000 × 5 × 9 / 48 = 9,375 元", size=9.5, color=BLUE, indent=0.2, after=2)
body_para(doc, "Tip: 加速归属（如业绩加速条款）会导致前期摊销增大，需在授予时做前瞻估算。", size=9.5, color=GRAY, indent=0.2, after=6)

callout(doc, "IFRS 2 和 ASC 718 均要求将股权激励计入经常性损益，建议在方案设计阶段就用本工具做三年成本预演，避免财报意外。", "合规提醒")

# ════════════════════════════════════════════════════════════════════════════
# TOOL 3 — 薪酬竞争力分析工具
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具三：薪酬竞争力分析工具", 1)

body_para(doc, "适用场景：对比企业薪酬水平与市场数据，识别薪酬竞争劣势，计算竞争力指数，为年度调薪提供依据。", size=10, color=GRAY, after=6)

heading(doc, "3.1 市场数据对照表", 2)

headers3 = ["职级序列", "岗位", "本企薪资 P50", "市场P25", "市场P50", "市场P75", "本企分位值"]
rows3 = [
    ["M 管理序列", "总监", "28万", "22万", "26万", "32万", "≈P55"],
    ["M 管理序列", "高级经理", "20万", "16万", "19万", "23万", "≈P60"],
    ["P 专业序列", "高级工程师", "22万", "18万", "21万", "25万", "≈P58"],
    ["P 专业序列", "工程师", "14万", "11万", "13万", "16万", "≈P65"],
    ["O 操作序列", "主管", "10万", "8万", "10万", "12万", "≈P50"],
]
add_table(doc, headers3, rows3, col_widths=[1.0, 1.0, 0.9, 0.8, 0.8, 0.8, 0.9])

heading(doc, "3.2 竞争力指数计算", 2)

body_para(doc, "竞争力指数（CR）=", bold=True, size=10, after=0)
body_para(doc, "本企业该岗位薪资中位值 / 市场该岗位薪资中位值 × 100", size=9.5, color=GRAY, indent=0.2, after=3)

body_para(doc, "判断标准:", bold=True, size=10, after=2)
criteria = [
    "CR > 110：薪酬领先，市场吸引力强，但人力成本高",
    "CR 90~110：薪酬匹配，市场竞争力正常",
    "CR 75~90：薪酬落后，存在人才流失风险，建议调薪",
    "CR < 75：严重落后，关键岗位必须优先调整",
]
for c in criteria:
    bullet(doc, c, size=9.5)

heading(doc, "3.3 差距分析与调薪建议", 2)

body_para(doc, "薪资差距 = 本企薪资 - 市场P50", bold=True, size=10, after=2)
body_para(doc, "示例：某工程师本企薪资 14 万，市场 P50 = 13 万，差距 +7.7%，处于合理区间。", size=9.5, color=GRAY, indent=0.2, after=4)
body_para(doc, "调薪优先级 = 竞争力指数排名 × 职级重要度 × 离职风险权重", bold=True, size=10, after=2)
body_para(doc, "将上述三维打分后加权汇总，可得到调薪优先级排序，为有限调薪预算的分配提供客观依据。", size=9.5, color=GRAY, indent=0.2, after=6)

callout(doc, "市场数据建议采购两年内的第三方薪酬报告（如 Hewitt、Mercer、北森等），内部调研数据需注意样本代表性，避免以偏概全。", "数据来源")

# ════════════════════════════════════════════════════════════════════════════
# TOOL 4 — 全面薪酬配置工具
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具四：全面薪酬配置工具", 1)

body_para(doc, "适用场景：帮助企业设计最优的薪酬组合（固定薪、浮动薪、福利、非现金激励），在成本约束下最大化激励效果。", size=10, color=GRAY, after=6)

heading(doc, "4.1 各要素比例参考", 2)

headers4 = ["职级", "固定薪", "浮动薪（奖金）", "福利津贴", "长期激励", "非现金激励"]
rows4 = [
    ["一线员工", "70%~80%", "10%~15%", "10%~15%", "0%", "5%~10%"],
    ["初级专业人员", "60%~70%", "15%~20%", "10%~15%", "0%~5%", "5%~10%"],
    ["高级专业人员", "50%~60%", "20%~25%", "10%~15%", "5%~10%", "5%~10%"],
    ["基层管理者", "50%~60%", "20%~25%", "10%~15%", "5%~10%", "5%~10%"],
    ["中层管理者", "40%~50%", "25%~30%", "10%~15%", "10%~15%", "5%~10%"],
    ["高层管理者", "30%~40%", "25%~35%", "8%~12%", "20%~30%", "3%~5%"],
]
add_table(doc, headers4, rows4, col_widths=[1.2, 1.1, 1.3, 1.1, 1.1, 1.1])

heading(doc, "4.2 个性化配置建议", 2)

body_para(doc, '不同员工群体的偏好差异大，可通过"自助餐式"福利（Flexible Benefits）提升满意度：', size=9.5, color=GRAY, after=4)

pref_data = [
    ["员工类型", "高偏好要素", "低偏好要素", "配置建议"],
    ["年轻单身", "现金、股权、培训机会", "额外养老、子女医疗", "提高浮动薪比例 + 学习发展基金"],
    ["中年家庭", "医疗福利、寿险、弹性工作", "高频差旅补贴", "提高医疗保障 + 弹性工时选项"],
    ["资深高管", "股权、名誉激励、高端福利", "基础培训", "提高长期激励占比 + 声誉激励"],
    ["技术专家", "研发资源、专利奖励、内部创新", "销售提成", "专项研发基金 + 专利署名权"],
]
add_table(doc, pref_data[0], pref_data[1:], col_widths=[1.2, 1.6, 1.6, 2.0])

body_para(doc, "配置原则:", bold=True, size=10, after=2)
principles = [
    "成本可控：总体激励成本不超过营业收入的 X%（行业基准参考）",
    "外部竞争性：核心岗位总薪酬 competitiveness index ≥ 90",
    "内部公平性：同岗不同地区可设地区系数（通常 1.0~1.3）",
    "激励导向：浮动薪与关键业绩指标（KPI/OKR）强挂钩",
]
for pr in principles:
    bullet(doc, pr, size=9.5)

# ════════════════════════════════════════════════════════════════════════════
# TOOL 5 — 激励ROI计算工具
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具五：激励ROI计算工具", 1)

body_para(doc, "适用场景：量化激励投入的回报，证明HR项目的商业价值，指导年度激励预算决策。", size=10, color=GRAY, after=6)

heading(doc, "5.1 激励投入产出比", 2)

body_para(doc, "激励ROI = (激励带来的业绩增量 - 激励总成本) / 激励总成本 × 100%", bold=True, size=10, after=3)

body_para(doc, "分解步骤:", bold=True, size=10, after=2)
roi_steps = [
    "STEP 1 — 测算激励总成本：包括薪酬增量、股权摊销、福利支出、管理成本",
    "STEP 2 — 量化业绩增量：与基期对比，剥离市场增长因素后的纯激励贡献",
    "STEP 3 — 计算净收益：业绩增量 - 激励总成本",
    "STEP 4 — 计算ROI：净收益 / 激励总成本",
]
for rs in roi_steps:
    bullet(doc, rs, size=9.5)

body_para(doc, "示例：某销售团队年度激励预算 50 万元，当年销售额增加 300 万元（扣除市场自然增长 100 万后，净增量 200 万）：", size=9.5, color=GRAY, indent=0.2, after=2)
body_para(doc, "激励ROI = (200万 - 50万) / 50万 × 100% = 300%", size=9.5, color=BLUE, indent=0.2, after=6)

heading(doc, "5.2 留任率改善计算", 2)

body_para(doc, "关键人才留任价值 = 替换成本 × 留任概率提升值", bold=True, size=10, after=2)
body_para(doc, "其中替换成本通常为该岗位年薪的 50%~200%（管理岗取上限），留任概率提升值由激励方案实施前后对比得出。", size=9.5, color=GRAY, indent=0.2, after=4)

headers5 = ["激励项目", "实施前年流失率", "实施后年流失率", "关键人才数量", "替换成本/人", "年节约替换成本"]
rows5 = [
    ["晋升+调薪方案", "25%", "12%", "50人", "30万", "195万"],
    ["股权激励计划", "20%", "8%", "20人", "50万", "120万"],
    ["弹性福利计划", "18%", "14%", "100人", "15万", "60万"],
]
add_table(doc, headers5, rows5, col_widths=[1.3, 1.1, 1.1, 1.0, 1.0, 1.2])

callout(doc, "ROI计算需注意归因清晰——业绩增长由多因素驱动，激励只是其中之一。建议用回归分析或对标组对照，排除其他变量干扰。", "方法论提醒")

# ════════════════════════════════════════════════════════════════════════════
# TOOL 6 — 变革影响评估工具
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "工具六：变革影响评估工具", 1)

body_para(doc, "适用场景：薪酬体系变革（调薪、换结构、改激励模式）前，评估员工接受度和潜在阻力，提前制定变革管理策略。", size=10, color=GRAY, after=6)

heading(doc, "6.1 变革阻力评估矩阵", 2)

body_para(doc, "阻力来源分四个维度，每个维度 1~5 分评估：", size=9.5, color=GRAY, after=4)

headers6 = ["阻力维度", "评估要点", "高分表现（高阻力）", "建议应对"]
rows6 = [
    ["利益受损", "哪些群体短期收入下降？下降幅度多大？", "高管带头反对、核心员工联名抗议", "保障兜底期、个性化补偿方案"],
    ["认知差异", "员工是否理解变革的必要性和逻辑？", '多数人认为"没有必要改"，质疑数据', "充分沟通、提供数据支撑的说明会"],
    ["情感抵触", "员工对变革的情感反应，是否感到被否定？", '员工感到"公司不认可过去的工作"', "管理层以身作则、肯定历史贡献"],
    ["能力担忧", "员工是否担心自己无法适应新体系？", "能力较强但年龄偏大的员工最焦虑", "赋能培训、转型期的辅导支持"],
]
add_table(doc, headers6, rows6, col_widths=[0.9, 1.5, 1.8, 2.2])

heading(doc, "6.2 员工接受度预测", 2)

body_para(doc, "接受度指数（参考公式）=", bold=True, size=10, after=0)
body_para(doc, "Σ(各群体占比 × 预期接受率 × 影响力系数) / Σ(各群体占比 × 影响力系数)", size=9.5, color=GRAY, indent=0.2, after=4)

body_para(doc, "影响力系数参考：", bold=True, size=10, after=2)
influence = [
    "意见领袖（内部有影响力的老员工）：权重 1.5~2.0",
    "关键岗位（业务核心、技术骨干）：权重 1.2~1.5",
    "普通员工：权重 1.0",
    "新员工（入职 < 1年）：权重 0.8",
]
for inf in influence:
    bullet(doc, inf, size=9.5)

heading(doc, "6.3 变革管理行动建议", 2)

actions = [
    ["接受度预测", "行动策略"],
    ["> 70%", "以宣导为主，辅以FAQ文件和部门沟通会"],
    ["50%~70%", "加强双向沟通，建立反馈渠道，必要时微调方案细节"],
    ["30%~50%", "先在局部试点，同步收集反馈再全面推广"],
    ["< 30%", "暂缓执行，重新诊断阻力根源，调整方案或换渐进式路径"],
]
add_table(doc, actions[0], actions[1:], col_widths=[1.2, 5.2])

callout(doc, "薪酬变革的最大风险不是方案本身，而是信息不对称导致的信任危机。建议变革启动前安排至少两轮员工沟通，第一轮讲Why，第二轮讲How。", "变革管理金句")

# ════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(10)
run = p_footer.add_run("— 薪酬激励设计 · 数字工具包 · 完 —")
set_run_font(run, 9, color=RGBColor(0xAA, 0xAA, 0xAA))

# ─── Save ───────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")

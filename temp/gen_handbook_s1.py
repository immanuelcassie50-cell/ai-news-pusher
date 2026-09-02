#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# Cover
for _ in range(8): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('AI时代决策工作手册')
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = RGBColor(0,51,102)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('把个人判断变成组织可复用的决策资产')
r.font.size = Pt(16); r.italic = True

for _ in range(4): doc.add_paragraph()

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run('学员对象'); r.bold = True; r.font.size = Pt(14)

a2 = doc.add_paragraph()
a2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a2.add_run('企业中层管理者、项目负责人、HRBP、培训经理、行动学习催化师')
r.font.size = Pt(12)

for _ in range(6): doc.add_paragraph()

c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('© 罗宏伟'); r.bold = True; r.font.size = Pt(14)

c2 = doc.add_paragraph()
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c2.add_run('版权所有 侵权必究'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(128,128,128)

doc.add_page_break()
print('S1: Cover done')
doc.save(r'D:\CC\temp\handbook_s1.docx')

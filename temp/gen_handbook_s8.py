#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# PART 1 Q&A
qa1 = doc.add_paragraph()
r = qa1.add_run('PART 1  Q&A')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0,51,102)

qas1 = [
    ('问：决策者不愿意承认自己当时很纠结，访谈怎么破这个防御？', '问如果当时给你多一天时间，你会不会做同样的决定，这句话不要求他承认纠结，但如果他犹豫了，这个回答本身就暴露了当时的犹豫程度。'),
    ('问：一个决策同时涉及三四个人，各自记忆不一样，决策卡该以谁的版本为准？', '把几个人的版本分开访谈，然后专门对比出现分歧的地方——分歧本身往往比一致的部分更有信息量。'),
    ('问：触发条件写得太细，会不会变成新的教条主义？', '在每张卡的开头加一句定位说明：本卡列出的触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。'),
    ('问：场景映射矩阵做得越细，维护成本越高，怎么判断该细到什么程度？', '先做粗粒度版本，投入使用后收集哪个场景反复被反馈不适用，再针对性细分。'),
    ('问：失败案例的当事人还在公司，写清楚会不会变成秋后算账？', '写警示清单时不点名、不复述当事人的岗位和具体身份细节，只保留岔路口和信号这两部分。'),
]

for q, a in qas1:
    qp = doc.add_paragraph()
    r = qp.add_run(q); r.bold = True
    ap = doc.add_paragraph()
    ap.paragraph_format.left_indent = Cm(1)
    ap.add_run(a)

doc.add_page_break()
print('S8: Part 1 Q&A done')
doc.save(r'D:\CC\temp\handbook_s8.docx')

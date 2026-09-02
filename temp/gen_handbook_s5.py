#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)
    
    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)
    
    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)
    
    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)
    
    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# PART 1
p1 = doc.add_paragraph()
r = p1.add_run('PART 1')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

p1s = doc.add_paragraph()
r = p1s.add_run('认知升级：复盘与决策卡的本质区分')
r.bold = True; r.font.size = Pt(16)

chapters = [
    {
        'title': '第一章  复盘写的是过去，决策卡写的是未来',
        'gold_quote': '教训只能让人不再犯错，决策卡能让人复制成功。',
        'objectives': ['理解复盘与决策卡的本质区别', '认识到教训与判断结构是不同的知识形态', '掌握决策卡的核心价值定位'],
        'concepts': [('复盘', '对过去事件进行叙事性描述，总结经验教训'), ('决策卡', '将判断结构提炼成可现场调用的检查表和触发条件'), ('隐性判断', '未被语言化的经验规则，是决策卡核心内容')],
        'key_points': ['复盘写过去，决策卡写未来', '教训只告诉结果不好的方向，没告诉怎么取舍', '理解一件事和能用是两回事', '叙事复盘和决策卡是递进关系', '卡片会被贴在工位上，报告会存进文件夹'],
        'thinking': ['你的复盘有多少能在决策现场被调用？', '为什么讲清楚不等于能复制？'],
        'exercise': ['回顾你最近的项目复盘，尝试提炼3条可使用的判断规则'],
        'summary': '复盘写的是过去，决策卡写的是未来。过去已发生改变不了；未来还没到来，一张结构清晰的决策卡能让下一个人在关键时刻少赌一次。'
    },
    {
        'title': '第二章  不是每个决策都值得做成一张卡',
        'gold_quote': '值得复盘的从来不是结果好坏，是判断有没有难度。',
        'objectives': ['掌握判断决策是否值得做成卡的三问法', '理解运气型成功与可复制方法论的区别', '避免把运气当成方法论做成卡片'],
        'concepts': [('判断难度', '决策困难程度，取决于是否存在多个合理选项'), ('运气型成功', '结果好但判断简单的决策'), ('可复用判断结构', '特定信号出现时该往哪个方向想的结构性逻辑')],
        'key_points': ['判断难度是值不值得做成卡的第一标准', '三问法：多选项？可拆解？情境可重复？', '结果好但判断简单的沉淀的是照流程走', '把运气当方法论是最危险的雷', '结果不好但判断有难度的同样值得做卡'],
        'thinking': ['你最近的成功决策真正成功因素是什么？', '什么样的决策让你真正纠结过？'],
        'exercise': ['用三问法评估你手头的3个近期决策'],
        'summary': '不是每个决策都值得做成一张卡，值得的是那些真正让人纠结过、换个人判断结果可能不同的决策。抓错识别标准比不做复盘更危险。'
    },
    {
        'title': '第三章  复盘访谈问的是你当时不确定什么',
        'gold_quote': '决策者自己讲不出他做对了什么，因为他当时没觉得那是个选择。',
        'objectives': ['掌握复盘访谈三维度：追因、权衡、未预见的假设', '学会追问你当时不确定什么', '理解判断依据才是稀缺资源'],
        'concepts': [('追因', '挖掘决策触发点，问什么信号让你觉得需要做决策'), ('权衡', '挖掘选项排除过程，问怎么排除其他选项的'), ('未预见的假设', '挖掘无意识依赖前提，问如果前提不成立')],
        'key_points': ['问你做了什么只挖出时间线，挖不出判断依据', '不确定的地方才是判断真正发生的地方', '追因放最前面，因为它相对容易回答', '警惕事后合理化', '三个维度顺序：追因→权衡→未预见的假设'],
        'thinking': ['你在访谈中是否遇到过我就是觉得应该这样？', '如何区分真实判断依据和事后合理化？'],
        'exercise': ['找一位做过重要决策的同事，用三个维度做一次访谈练习'],
        'summary': '复盘访谈问的不是你做了什么，是你当时不确定什么。事件谁都能讲清楚，真正稀缺的是那个人权衡时脑子里划掉的选项、没说出口的假设。'
    },
]

for ch in chapters:
    add_chapter_content(doc, ch)
    doc.add_page_break()

print('S5: Chapters 1-3 done')
doc.save(r'D:\CC\temp\handbook_s5.docx')

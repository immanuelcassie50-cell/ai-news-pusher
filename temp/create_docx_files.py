# -*- coding: utf-8 -*-
"""
生成高考志愿填报师工作手册文档
- 典型案例集.docx
- 常见问题Q&A手册.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 输出目录
OUTPUT_DIR = "D:/新课开发/工作手册/高考志愿填报师/完整课程包/08-Demo成果/"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 版权声明
COPYRIGHT = "© 罗宏伟 高考志愿填报师工作手册"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    elif level == 2:
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return heading

def add_paragraph_with_format(doc, text, bold=False, font_size=11, space_after=8):
    """添加格式化段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    return p

def add_case_content(doc, case_data):
    """添加案例内容"""
    # 背景
    add_heading_with_style(doc, "背景", 3)
    add_paragraph_with_format(doc, case_data["背景"])

    # 冲突点
    add_heading_with_style(doc, "冲突点", 3)
    add_paragraph_with_format(doc, case_data["冲突点"])

    # 处理过程
    add_heading_with_style(doc, "处理过程", 3)
    add_paragraph_with_format(doc, case_data["处理过程"])

    # 结果
    add_heading_with_style(doc, "结果", 3)
    add_paragraph_with_format(doc, case_data["结果"])

    # 启示
    add_heading_with_style(doc, "启示", 3)
    p = add_paragraph_with_format(doc, case_data["启示"], bold=True)
    p.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def create_cases_document():
    """创建典型案例集文档"""
    doc = Document()

    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("高考志愿填报师", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_heading("典型案例集", level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(22)
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    # 案例数据
    cases = [
        {
            "title": "案例一：沉默的高分考生（不会表达兴趣方向）",
            "背景": "高考680分理科男生，逻辑思维强，数学竞赛省奖。家长认为\"分数够什么就学什么\"，孩子对策略类游戏的资源调度和数值平衡有深入研究，论坛写过策略分析帖，但从未被问过兴趣方向。",
            "冲突点": "家长坚持传统观念认为高分就该去热门专业，孩子虽然有独特兴趣但从未被重视，导致沟通陷入僵局。",
            "处理过程": "罗老师没有先看梯度表，而是先问\"平时喜欢做什么\"，聊出孩子对策略游戏的真实投入，重新调整冲档专业方向。",
            "结果": "保留了大部分原梯度结构，但冲档里三所学校的专业方向做了调整。",
            "启示": "冲稳保是语法不是方法，真正决定方向的是\"这个人是谁\"。"
        },
        {
            "title": "案例二：信息过载的焦虑家长（十几个sheet的Excel）",
            "背景": "父亲花一个多月整理了密密麻麻十几个sheet的表格，每所学校录取线、就业率、保研率、宿舍条件、食堂评分、论坛吐槽等。",
            "冲突点": "信息量庞大但父亲比一个月前更焦虑而不是更清楚，被自己收集的信息吓住了不敢选。",
            "处理过程": "罗老师没有帮他加信息，而是帮他删信息，一起过滤噪音。删到最后，父亲自己说出\"原来我不是不知道该怎么选，是被自己收集的东西吓住了\"。",
            "结果": "删到最后，父亲自己说出了心里话，焦虑明显缓解。",
            "启示": "信息收集不是查数据，是替人过滤噪音，稀缺的不是信息，是判断力。"
        },
        {
            "title": "案例三：名校vs专业两难（冲稳保比例争议）",
            "背景": "女生分数够到一所不错211，想冲名校但可能被调剂到不喜欢专业。",
            "冲突点": "家长坚持\"名校优先\"，孩子其实有更感兴趣的方向，冲的比例应该多大存在争议。",
            "处理过程": "用倒推法问清楚孩子真正想要的状态，讨论风险偏好，翻译成具体的冲稳保比例。",
            "结果": "找到一个兼顾名校层次和专业匹配度的方案。",
            "启示": "冲稳保三个档位是风险偏好的翻译，不是排列组合。"
        },
        {
            "title": "案例四：说\"不知道\"的文科生（缺乏自我认知）",
            "背景": "女生说\"我都行\"，对心理学感兴趣但家里觉得\"毕业不好找工作\"。孩子从小不敢表达真实想法，家长替她做所有决定，她习惯了沉默。",
            "冲突点": "女儿其实一直对心理学感兴趣，只是家里觉得这个专业\"毕业不好找工作\"，她自己也就不敢提。",
            "处理过程": "单独跟女儿聊，妈妈不在场，慢慢问出她真正感兴趣的方向。",
            "结果": "选了一所综合类院校的应用心理学，保底档也留了心理学相关备选。",
            "启示": "先问\"十年后想过什么日子\"再往回算专业，方向要从孩子嘴里说出来。"
        },
        {
            "title": "案例五：孩子vs家长意见相反（方向冲突）",
            "背景": "母亲私下找罗老师\"帮我劝劝女儿\"，女儿喜欢某专业但妈妈觉得不好就业。",
            "冲突点": "母亲把罗老师当成了说服工具，实际上他应该是中立的专业顾问。",
            "处理过程": "罗老师拒绝了\"帮劝\"请求，而是分别单独聊清楚各方真正担心什么，然后三人一起把信息摆出来。",
            "结果": "母亲说出了自己的焦虑源于年轻时的类似经历没走通，女儿说出了自己从没被认真问过\"为什么喜欢这个\"。",
            "启示": "你是让两边都能听见彼此的人，不是替任何一方说话的人。"
        },
        {
            "title": "案例六：产业判断失误复盘（土木工程案例）",
            "背景": "学生分数够到211土木工程，家长因\"体制内好就业\"坚持要报。",
            "冲突点": "罗老师给他们看了近几年真实招聘数据，行业收缩趋势明显，但家长还是决定去。",
            "处理过程": "把产业判断摆清楚，不是替他们决定，最终尊重他们的选择权。",
            "结果": "家长说\"我们村里以前学这个的都过得不错\"，最终选择土木工程。",
            "启示": "产业判断不是背专业介绍，是判断五年后这行还在不在；要坦诚判断也可能出错。"
        }
    ]

    # 添加案例
    for i, case in enumerate(cases):
        # 案例标题
        case_heading = doc.add_heading(case["title"], level=1)
        case_heading.runs[0].font.size = Pt(16)
        case_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        add_case_content(doc, case)

        # 案例之间的分隔
        if i < len(cases) - 1:
            doc.add_paragraph()
            p = doc.add_paragraph("─" * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

    # 版权声明
    doc.add_paragraph()
    copyright_p = doc.add_paragraph(COPYRIGHT)
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_p.runs[0].font.size = Pt(10)
    copyright_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "高考志愿填报师-典型案例集.docx")
    doc.save(output_path)
    print(f"案例集已保存: {output_path}")
    return output_path


def create_qa_document():
    """创建常见问题Q&A手册"""
    doc = Document()

    # 设置页面
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # 标题
    title = doc.add_heading("高考志愿填报师", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(26)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_heading("常见问题Q&A手册", level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(22)
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    # Q&A数据
    qa_data = {
        "分类一：专业选择": [
            {
                "q": "Q: 孩子说\"不知道\"自己对什么感兴趣怎么办？",
                "a": "A: 换更具体的小问题：什么事情做的时候会忘记看手机；什么样的失败能接受、什么样完全无法接受；更享受一个人琢磨还是跟一群人一起做成。"
            },
            {
                "q": "Q: 家长坚持要孩子学某个\"热门专业\"，但孩子明显不感兴趣",
                "a": "A: 先分别问清楚双方真正担心什么、真正想要什么，再把信息摆出来，让他们自己得出结论。"
            },
            {
                "q": "Q: 如何判断一个专业未来五年还在不在？",
                "a": "A: 看三件事：解决的问题是不是真实且长期存在的需求；核心壁垒是不是容易被AI替代；现在处于扩张期还是收缩期。"
            },
            {
                "q": "Q: 冲稳保比例到底怎么定？",
                "a": "A: 先问清楚风险偏好：如果掉到保底档能不能接受；如果冲到不喜欢调剂专业愿不愿意；家里对复读的真实态度。"
            },
            {
                "q": "Q: 名校优先还是专业优先？",
                "a": "A: 这不是二选一，是风险偏好的翻译。要讨论愿意为学校牺牲多少专业匹配度。"
            }
        ],
        "分类二：志愿填报": [
            {
                "q": "Q: 什么时候开始准备志愿填报最合适？",
                "a": "A: 越早越好，高二开始接触行业信息，高三开始系统规划，但核心判断不因早晚而异。"
            },
            {
                "q": "Q: 平行志愿怎么填才能不滑档？",
                "a": "A: 了解投档规则、级差规则、专业调剂规则，但这些是保底，不是方向。"
            },
            {
                "q": "Q: 提前批要不要报？",
                "a": "A: 看孩子具体情况，提前批有特殊性，不是所有人都适合。"
            },
            {
                "q": "Q: 服从调剂怎么勾选？",
                "a": "A: 这是风险控制的最后一道关口，要结合前面的风险偏好讨论。"
            },
            {
                "q": "Q: 估分和实际分数差很多怎么办？",
                "a": "A: 先稳定心态，再重新评估方案，预估本身就是有误差的。"
            }
        ],
        "分类三：沟通技巧": [
            {
                "q": "Q: 家长和孩子意见相反时怎么办？",
                "a": "A: 不要当说服工具，做翻译和主持，让双方听见彼此。"
            },
            {
                "q": "Q: 家长比孩子更焦虑，怎么处理？",
                "a": "A: 承认焦虑的合理性，但不要被焦虑牵着走，把注意力放回信息本身。"
            },
            {
                "q": "Q: 孩子全程沉默不配合怎么办？",
                "a": "A: 单独约谈，创造没有家长在场的安全环境再深入问。"
            },
            {
                "q": "Q: 如何让家长相信你的专业判断？",
                "a": "A: 不是说服，是摆信息差，让对方自己验证、自己得出结论。"
            },
            {
                "q": "Q: 遇到客户临时反悔怎么办？",
                "a": "A: 不去跟那句话较劲，重新把信息摆出来，让家长自己判断。"
            }
        ],
        "分类四：职业发展": [
            {
                "q": "Q: 这一行收入怎么样？",
                "a": "A: 因人而异，取决于服务质量和口碑积累，不是靠走量。"
            },
            {
                "q": "Q: 新入行从哪开始？",
                "a": "A: 先把政策和基础规则学扎实，同时尽早接触真实案例。"
            },
            {
                "q": "Q: AI工具越来越强，这行还有前途吗？",
                "a": "A: 真正被替代的是只会算冲稳保的，真正的价值在判断力和产业认知。"
            },
            {
                "q": "Q: 如何建立自己的口碑？",
                "a": "A: 口碑是副产品，把每一次服务做扎实，自然会来。"
            },
            {
                "q": "Q: 做这行久了会麻木吗？",
                "a": "A: 会的，所以要经常提醒自己，对每个孩子来说这是唯一的一次。"
            },
            {
                "q": "Q: 需要考什么证书吗？",
                "a": "A: 证书能证明基础知识，不能证明判断力，判断力靠真实案例积累。"
            },
            {
                "q": "Q: 一年服务多少学生合适？",
                "a": "A: 如果记不清某个学生具体聊过什么、纠结过什么，说明量已经超过上限了。"
            }
        ]
    }

    # 添加Q&A内容
    for category, qa_list in qa_data.items():
        # 分类标题
        cat_heading = doc.add_heading(category, level=1)
        cat_heading.runs[0].font.size = Pt(16)
        cat_heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # Q&A表格
        for qa in qa_list:
            # Q行
            q_p = doc.add_paragraph()
            q_p.paragraph_format.space_before = Pt(12)
            q_p.paragraph_format.space_after = Pt(4)
            q_run = q_p.add_run(qa["q"])
            q_run.font.size = Pt(11)
            q_run.font.bold = True
            q_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

            # A行
            a_p = doc.add_paragraph()
            a_p.paragraph_format.space_before = Pt(0)
            a_p.paragraph_format.space_after = Pt(8)
            a_p.paragraph_format.left_indent = Pt(24)
            a_run = a_p.add_run(qa["a"])
            a_run.font.size = Pt(11)
            a_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        doc.add_paragraph()

    # 版权声明
    doc.add_paragraph()
    copyright_p = doc.add_paragraph(COPYRIGHT)
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_p.runs[0].font.size = Pt(10)
    copyright_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "高考志愿填报师-常见问题Q&A手册.docx")
    doc.save(output_path)
    print(f"Q&A手册已保存: {output_path}")
    return output_path


if __name__ == "__main__":
    print("开始生成文档...")
    path1 = create_cases_document()
    path2 = create_qa_document()
    print("\n文档生成完成!")
    print(f"1. {path1}")
    print(f"2. {path2}")

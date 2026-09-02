# -*- coding: utf-8 -*-
"""
Generate BSC & OKR Strategic Execution Handbooks
- Student Handbook (学员手册)
- Instructor Handbook (讲师手册)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r'D:\新课开发\管理学\05-战略执行与落地'

# Color scheme
PRIMARY_COLOR = RGBColor(0x1F, 0x38, 0x64)  # Dark navy
SECONDARY_COLOR = RGBColor(0x2E, 0x75, 0xB6)  # Medium blue
ACCENT_COLOR = RGBColor(0x44, 0x72, 0xC4)  # Blue
BODY_TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)  # Near-black

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_form_field(doc, label, hint=""):
    """Add an interactive form field"""
    para = doc.add_paragraph()
    run = para.add_run(label)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    if hint:
        run2 = para.add_run(f"  {hint}")
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run2.font.name = '微软雅黑'

    para = doc.add_paragraph()
    run = para.add_run("_" * 60)
    run.font.size = Pt(11)
    para.space_after = Pt(12)

def create_student_handbook():
    """Generate Student Handbook"""
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== COVER PAGE =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('战略执行与落地')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = PRIMARY_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('——平衡计分卡与OKR实战')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = SECONDARY_COLOR

    for _ in range(4):
        doc.add_paragraph()

    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run('【学员手册】')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = ACCENT_COLOR

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('适用对象：中高层管理者、战略规划部门、人力资源同仁')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run.font.color.rgb = BODY_TEXT_COLOR

    # Page break
    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    toc_title = doc.add_paragraph()
    run = toc_title.add_run('目录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    toc_items = [
        ('第一章', '课程介绍与学习目标', '3'),
        ('第二章', '平衡计分卡四维模型', '5'),
        ('第三章', 'OKR目标与关键结果', '10'),
        ('第四章', 'BSC与OKR的融合之道', '15'),
        ('第五章', '战略解码与落地执行', '18'),
        ('第六章', '案例分析', '22'),
        ('第七章', '练习题', '28'),
        ('第八章', '行动学习计划', '32'),
        ('附录', '工具速查', '35'),
    ]

    for ch, title, page in toc_items:
        para = doc.add_paragraph()
        run = para.add_run(f'{ch}  {title}')
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 1: COURSE INTRODUCTION =====
    h1 = doc.add_heading('第一章 课程介绍与学习目标', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('1.1 课程背景', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '战略执行与落地是困扰众多企业的核心难题。根据《财富》杂志的统计，约70%的企业失败不是因为战略本身有问题，'
        '而是因为战略执行不力。平衡计分卡（BSC）和目标与关键结果（OKR）是两种经过全球企业验证的战略执行工具，'
        '前者强调整体战略的系统性落地，后者强调目标的聚焦与突破。将两者融合使用，能够帮助企业实现"上下同欲、内外协同"。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('1.2 学习目标', level=2)

    objectives = [
        '理解平衡计分卡四维模型的内涵与逻辑关系',
        '掌握OKR的制定方法、评分标准与落地流程',
        '学会将BSC的战略解码与OKR的目标设定有机结合',
        '能够运用战略地图进行组织战略可视化呈现',
        '掌握将组织战略转化为部门与个人OKR的操作方法',
        '提升战略思维与执行力，实现从"知道"到"做到"的跨越'
    ]

    for i, obj in enumerate(objectives, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {obj}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    h2 = doc.add_heading('1.3 课程对象', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '本课程适用于中高层管理者、战略规划部门、人力资源管理同仁，'
        '以及希望提升战略执行能力的企业管理者和骨干员工。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('1.4 课程时长', level=2)
    content = doc.add_paragraph()
    run = content.add_run('建议总时长：2天（每天6小时，共计12小时）')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_page_break()

    # ===== CHAPTER 2: BSC FOUR PERSPECTIVES =====
    h1 = doc.add_heading('第二章 平衡计分卡四维模型', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('2.1 平衡计分卡的起源与发展', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '平衡计分卡（Balanced Scorecard，简称BSC）由哈佛商学院教授罗伯特·卡普兰和诺朗诺顿研究所所长戴维·诺顿于1992年首次提出。'
        '它超越了你财务指标的传统绩效评价体系，从财务、客户、内部流程、学习与成长四个维度，将组织的战略目标转化为可操作的衡量指标和目标值。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('2.2 平衡计分卡四维模型', level=2)

    # Four perspectives table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['维度', '核心问题', '典型指标']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    perspectives = [
        ('财务维度', '如何为股东创造价值？', '营业收入、利润率、资产回报率、现金流'),
        ('客户维度', '如何为客户创造价值？', '客户满意度、客户留存率、市场份额、新客户获取率'),
        ('内部流程维度', '哪些流程必须做到卓越？', '产品合格率、响应周期、成本控制、创新成功率'),
        ('学习与成长维度', '如何保持持续创新能力？', '员工满意度、关键人才流失率、培训覆盖率、人均产出'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, (dim, question, indicators) in enumerate(perspectives, 1):
        table.rows[i].cells[0].text = dim
        table.rows[i].cells[1].text = question
        table.rows[i].cells[2].text = indicators
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        table.rows[i].cells[2].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(table.rows[i].cells[0], colors[i-1])

    doc.add_paragraph()

    h2 = doc.add_heading('2.3 战略地图绘制方法', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '战略地图是平衡计分卡的可视化工具，它以因果关系为链条，将四个维度的目标串联起来，'
        '形成从学习成长到内部流程，再到客户价值，最终实现财务目标的逻辑链条。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h3 = doc.add_heading('战略地图绘制步骤：', level=3)
    steps = [
        '确定财务目标：明确我们要达到的财务业绩（如收入增长、利润率提升）',
        '确定客户价值主张：明确为目标客户创造什么独特价值',
        '确定关键内部流程：识别支撑客户价值主张的核心流程',
        '确定学习与成长能力：明确支撑上述目标所需的无形资产',
        '建立因果关系：用箭头连接各层级目标，形成战略地图'
    ]
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {step}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    h2 = doc.add_heading('2.4 学习要点', level=2)
    key_points = [
        'BSC不仅仅是绩效考核工具，更是战略管理和沟通工具',
        '四个维度之间存在严格的因果逻辑关系，而非简单的并列关系',
        '战略地图的绘制是战略解码的核心输出',
        'BSC的"平衡"体现在：财务与非财务指标的平衡、短期与长期的平衡、领先与滞后指标的平衡'
    ]
    for point in key_points:
        para = doc.add_paragraph()
        run = para.add_run(f'• {point}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 3: OKR =====
    h1 = doc.add_heading('第三章 OKR目标与关键结果', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('3.1 OKR的起源与本质', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        'OKR（Objectives and Key Results，目标与关键结果）由英特尔创始人安迪·格鲁夫发明，'
        '后被谷歌、LinkedIn等硅谷企业广泛采用。OKR的核心思想是通过设定具有挑战性的目标（O），'
        '以及衡量目标达成程度的关键结果（KR），实现组织的聚焦与突破。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('3.2 OKR的结构', level=2)

    # OKR structure table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    headers = ['要素', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    table.rows[1].cells[0].text = 'O（Objective）目标'
    table.rows[1].cells[1].text = '定性描述，表达我们要实现的方向。鼓舞人心、与战略对齐、有时间周期。'
    table.rows[2].cells[0].text = 'KR（Key Results）关键结果'
    table.rows[2].cells[1].text = '定量衡量，用数字说话。衡量目标是否达成的具体指标，通常3-5个。'

    for i in range(1, 3):
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_heading('3.3 OKR制定方法', level=2)

    h3 = doc.add_heading('好的O（目标）的特征：', level=3)
    o_chars = [
        '鼓舞人心：能够激励团队，激发斗志',
        '与战略对齐：承接公司/部门战略，是战略的具体化',
        '清晰明确：团队成员对目标的理解一致',
        '有挑战性：通常是"跳一跳够得着"的目标',
        '定性描述：回答"我们要去哪里"的问题'
    ]
    for char in o_chars:
        para = doc.add_paragraph()
        run = para.add_run(f'• {char}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'

    doc.add_paragraph()

    h3 = doc.add_heading('好的KR（关键结果）的特征：', level=3)
    kr_chars = [
        '可量化：用数字说话，能够明确衡量进度',
        '有挑战：基线为70%完成度即为成功，100%完成说明目标设置不够大胆',
        '具体明确：团队成员对如何达成共识一致',
        '上下对齐：承接上级OKR，与同级部门协调配合',
        '有时限：在一个季度/半年/一年内完成'
    ]
    for char in kr_chars:
        para = doc.add_paragraph()
        run = para.add_run(f'• {char}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('3.4 OKR评分标准', level=2)

    # Scoring table
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    headers = ['得分', '说明']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    scores = [
        ('0.0-0.3', '未达成目标，需要反思原因'),
        ('0.4-0.6', '基本达成，需要分析差距'),
        ('0.7', '完美达成，这是期望的完成度'),
        ('0.8-0.9', '超额完成，说明目标设置可能不够大胆'),
        ('1.0', '100%完成，极少见，说明目标挑战性不足')
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF', 'F2F2F2']
    for i, (score, desc) in enumerate(scores, 1):
        table.rows[i].cells[0].text = score
        table.rows[i].cells[1].text = desc
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_heading('3.5 OKR与KPI的区别', level=2)

    # Comparison table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ['维度', 'OKR', 'KPI']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    comparisons = [
        ('目标性质', '挑战性目标（通常70%完成即为成功）', '标准目标（100%完成是期望的）'),
        ('核心目的', '突破与创新', '维持与达标'),
        ('评分方式', '自评+他评，不与薪酬直接挂钩', '考核+排名，与薪酬挂钩'),
        ('数量', '精简（通常1-3个O，每个O对应3-5个KR）', '可多（根据职责设定多个KPI）'),
        ('周期', '季度或半年', '月度/季度/年')
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF', 'F2F2F2']
    for i, (dim, okr, kpi) in enumerate(comparisons, 1):
        table.rows[i].cells[0].text = dim
        table.rows[i].cells[1].text = okr
        table.rows[i].cells[2].text = kpi
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(3):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ===== CHAPTER 4: BSC AND OKR INTEGRATION =====
    h1 = doc.add_heading('第四章 BSC与OKR的融合之道', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('4.1 为什么需要融合？', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        'BSC擅长战略的系统性分解和长期规划，但落地执行时需要更聚焦的目标管理工具；'
        'OKR擅长聚焦突破和快速迭代，但需要战略的牵引和对齐。两者结合，能够实现"战略清晰、执行有力"的双重目标。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('4.2 融合模型', level=2)

    # Integration model table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['BSC维度', '战略主题', 'OKR目标']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    integrations = [
        ('财务', '收入增长、成本优化', 'O：实现年收入增长30%\nKR1: 新产品收入占比达25%\nKR2: 单位成本下降10%'),
        ('客户', '客户满意度、市场扩张', 'O：成为行业客户满意度第一\nKR1: NPS评分提升至50分\nKR2: 客户留存率达95%'),
        ('内部流程', '运营效率、产品质量', 'O：打造行业最高效运营体系\nKR1: 订单交付周期缩短30%\nKR2: 产品合格率达99.5%'),
        ('学习成长', '人才发展、创新能力', 'O：建设行业领先的人才梯队\nKR1: 关键人才保留率达98%\nKR2: 人均培训时长超100小时'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, (dim, theme, okr) in enumerate(integrations, 1):
        table.rows[i].cells[0].text = dim
        table.rows[i].cells[1].text = theme
        table.rows[i].cells[2].text = okr
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(3):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    h2 = doc.add_heading('4.3 融合实施步骤', level=2)
    steps = [
        '第一步：战略澄清——用BSC明确组织战略地图和四维目标',
        '第二步：目标聚焦——将BSC各维度目标转化为1-3个具有挑战性的OKR',
        '第三步：横向对齐——确保各部门OKR与公司OKR对齐，避免孤岛效应',
        '第四步：纵向落地——将公司OKR分解到部门，再分解到个人OKR',
        '第五步：持续追踪——用OKR的系统追踪和复盘机制，确保BSC战略落地'
    ]
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'{i}. {step}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 5: STRATEGIC DECODING AND EXECUTION =====
    h1 = doc.add_heading('第五章 战略解码与落地执行', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('5.1 什么是战略解码？', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '战略解码是将战略意图转化为具体行动的关键过程。它回答三个核心问题：'
        '我们要去哪里（战略目标）？我们如何去（关键举措）？我们需要什么能力（资源保障）？'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('5.2 战略解码的三个层次', level=2)

    levels = [
        ('公司级解码', '将公司战略地图转化为公司级OKR，明确年度战略主题和突破目标'),
        ('部门级解码', '将公司级OKR分解到各部门，结合部门职责制定部门OKR'),
        ('个人级解码', '将部门OKR分解到个人，结合岗位职责制定个人OKR')
    ]

    for level, desc in levels:
        para = doc.add_paragraph()
        run = para.add_run(f'• {level}：{desc}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    h2 = doc.add_heading('5.3 战略落地的关键成功因素', level=2)
    factors = [
        '领导力：高层管理者的承诺和示范是战略落地的第一推动力',
        '一致性：上下级目标保持对齐，层层分解但不稀释',
        '聚焦：资源有限，必须聚焦在最重要的1-3件事上',
        '透明：OKR全公司公开可见，形成协同和监督',
        '复盘：定期复盘反思，及时调整和纠偏',
        '激励：将OKR完成情况与激励挂钩，形成正向循环'
    ]
    for factor in factors:
        para = doc.add_paragraph()
        run = para.add_run(f'• {factor}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    h2 = doc.add_heading('5.4 战略执行的常见误区', level=2)
    mistakes = [
        '战略与执行脱节：战略高大上，执行跟不上',
        '目标太多没有聚焦：什么都重要等于什么都不重要',
        '只考核不赋能：只给压力不给资源和支持',
        '重制定轻追踪：OKR制定后束之高阁，缺乏过程管理',
        '文化不支持：缺乏坦诚透明的沟通氛围'
    ]
    for mistake in mistakes:
        para = doc.add_paragraph()
        run = para.add_run(f'• {mistake}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 6: CASE ANALYSIS =====
    h1 = doc.add_heading('第六章 案例分析', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('案例一：华为的战略管理实践', level=2)

    h3 = doc.add_heading('背景', level=3)
    content = doc.add_paragraph()
    run = content.add_run(
        '华为从1987年一家小小的通信设备代理商，发展成为全球领先的ICT基础设施和智能终端提供商，'
        '其战略执行能力功不可没。华为将BSC和KPI体系深度融合，形成独特的战略管理闭环。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    h3 = doc.add_heading('做法', level=3)
    practices = [
        '战略规划：每五年制定公司战略规划（SP），每年制定年度业务计划（BP）',
        'BSC分解：从财务、客户、内部流程、学习成长四个维度制定目标',
        'KPI承接：将BSC目标转化为各层级KPI，确保战略落地',
        '述职制度：各级管理者定期述职，检验战略执行情况',
        '灰度管理：在战略清晰的前提下，允许执行中的灰度和灵活调整'
    ]
    for p in practices:
        para = doc.add_paragraph()
        run = para.add_run(f'• {p}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('启示', level=3)
    insights = [
        '战略需要组织能力的支撑，不能好高骛远',
        '战略执行需要制度化、流程化，而非依赖个人英雄主义',
        '定期复盘和纠偏是战略落地的关键机制'
    ]
    for insight in insights:
        para = doc.add_paragraph()
        run = para.add_run(f'• {insight}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    h2 = doc.add_heading('案例二：字节跳动的OKR实践', level=2)

    h3 = doc.add_heading('背景', level=3)
    content = doc.add_paragraph()
    run = content.add_run(
        '字节跳动以"大力出奇迹"著称，通过OKR实现了业务的快速突破和迭代。'
        '其OKR实践具有独特的"北极星指标"文化特色。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    h3 = doc.add_heading('做法', level=3)
    practices = [
        '北极星指标：每个业务阶段聚焦一个核心指标，全公司对齐',
        'OKR公开：全员可见，形成横向协同和纵向对齐',
        '双周Review：每两周检视OKR进度，及时调整策略',
        'Context over Control：强调上下文理解而非控制，给予执行者充分自主',
        '坦诚透明：倡导"上下文透明、结论激励"的文化'
    ]
    for p in practices:
        para = doc.add_paragraph()
        run = para.add_run(f'• {p}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('启示', level=3)
    insights = [
        'OKR需要与文化结合，坦诚透明的土壤才能让OKR生根发芽',
        'OKR不是考核工具，而是沟通和对齐工具',
        '高频复盘和快速迭代是互联网时代战略执行的关键能力'
    ]
    for insight in insights:
        para = doc.add_paragraph()
        run = para.add_run(f'• {insight}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===== CHAPTER 7: EXERCISES =====
    h1 = doc.add_heading('第七章 练习题', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('练习一：BSC维度匹配', level=2)
    content = doc.add_paragraph()
    run = content.add_run('请将以下指标与其对应的BSC维度进行匹配：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    indicators = [
        'a) 客户满意度评分    b) 员工培训覆盖率    c) 净利润率    d) 产品研发周期',
        'e) 市场份额    f) 员工流失率    g) 存货周转率    h) 客户投诉率'
    ]
    for ind in indicators:
        para = doc.add_paragraph()
        run = para.add_run(ind)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'

    doc.add_paragraph()
    content = doc.add_paragraph()
    run = content.add_run('请将字母填入对应维度：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    add_form_field(doc, '财务维度：', 'a, c, g')
    add_form_field(doc, '客户维度：', 'e, h')
    add_form_field(doc, '内部流程维度：', 'd')
    add_form_field(doc, '学习成长维度：', 'b, f')

    doc.add_paragraph()

    h2 = doc.add_heading('练习二：OKR制定', level=2)
    content = doc.add_paragraph()
    run = content.add_run('假设你是某互联网公司产品部门的负责人，请基于以下战略主题，制定本部门的OKR：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    content = doc.add_paragraph()
    run = content.add_run('战略主题：在保持用户增长的同时，提升付费转化率')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    add_form_field(doc, 'O（目标）：', '描述你的目标')
    add_form_field(doc, 'KR1（关键结果1）：', '')
    add_form_field(doc, 'KR2（关键结果2）：', '')
    add_form_field(doc, 'KR3（关键结果3）：', '')

    doc.add_paragraph()

    h2 = doc.add_heading('练习三：战略解码练习', level=2)
    content = doc.add_paragraph()
    run = content.add_run('某公司战略目标：三年内成为行业前三。请完成以下战略解码：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    add_form_field(doc, '第一步：公司级BSC战略主题', '填写财务、客户、内部流程、学习成长四个维度的战略主题')
    add_form_field(doc, '第二步：公司级OKR', '基于上述战略主题，制定公司级OKR')
    add_form_field(doc, '第三步：选择一个部门', '选择技术部门，制定其承接公司OKR的部门OKR')

    doc.add_page_break()

    # ===== CHAPTER 8: ACTION LEARNING PLAN =====
    h1 = doc.add_heading('第八章 行动学习计划', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('8.1 行动学习概述', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '行动学习是将学习与实践相结合的方法，通过解决真实问题来实现能力的提升。'
        '本课程要求每位学员制定个人行动学习计划，将课堂所学应用于实际工作。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('8.2 行动学习计划模板', level=2)

    # Action plan table
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ['阶段', '时间', '行动项', '预期成果']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    phases = [
        ('诊断', '第1周', '梳理当前组织战略执行的问题和差距', '问题清单'),
        ('设计', '第2周', '绘制部门战略地图，制定BSC目标', '战略地图初稿'),
        ('制定', '第3周', '将BSC目标转化为部门OKR', '部门OKR初稿'),
        ('对齐', '第4周', '与上级和平级部门对齐OKR', '对齐后的OKR'),
        ('落地', '第5-8周', '执行OKR并进行双周Review', 'OKR执行记录')
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF', 'F2F2F2']
    for i, (phase, time, action, result) in enumerate(phases, 1):
        table.rows[i].cells[0].text = phase
        table.rows[i].cells[1].text = time
        table.rows[i].cells[2].text = action
        table.rows[i].cells[3].text = result
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(4):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    h2 = doc.add_heading('8.3 个人行动计划', level=2)
    content = doc.add_paragraph()
    run = content.add_run('请基于课程所学，制定您的个人行动学习计划：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    add_form_field(doc, '我的姓名：', '')
    add_form_field(doc, '所属部门：', '')
    add_form_field(doc, '当前最需要解决的战略执行问题：', '')
    add_form_field(doc, '计划应用的工具（BSC/OKR/战略地图）：', '')
    add_form_field(doc, '预期成果：', '')
    add_form_field(doc, '开始时间：', '')
    add_form_field(doc, '需要支持：', '')

    doc.add_page_break()

    # ===== APPENDIX: TOOLS QUICK REFERENCE =====
    h1 = doc.add_heading('附录 工具速查', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('附录A：BSC四维模型速查', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['维度', '核心问题', '关键指标示例']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data = [
        ('财务', '如何为股东创造价值', '收入增长率、利润率、资产回报率、EBITDA'),
        ('客户', '如何为客户创造价值', '客户满意度、NPS、客户留存率、市场份额'),
        ('内部流程', '哪些流程必须卓越', '产品合格率、订单交付周期、成本控制'),
        ('学习成长', '如何保持创新能力', '员工满意度、人才流失率、培训覆盖率')
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, row in enumerate(data, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            set_cell_shading(table.rows[i].cells[0], colors[i-1])
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    h2 = doc.add_heading('附录B：OKR制定检查清单', level=2)

    checklist = [
        'O是否鼓舞人心，能够激发团队斗志？',
        'O是否与公司/部门战略对齐？',
        'O是否清晰明确，团队理解一致？',
        'O是否具有挑战性（70%完成即为成功）？',
        'KR是否可量化，用数字说话？',
        'KR数量是否控制在3-5个？',
        'KR是否承接O，形成支撑关系？',
        'KR是否有明确的时间节点？',
        'OKR是否在团队内公开透明？',
        '是否与上下级进行了充分的对齐沟通？'
    ]

    for i, item in enumerate(checklist, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'☐ {item}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    h2 = doc.add_heading('附录C：战略解码流程速查', level=2)

    steps = [
        '1. 战略澄清：明确公司使命、愿景、价值观',
        '2. 战略选择：确定战略定位和竞争策略',
        '3. 战略地图：绘制BSC四维战略地图',
        '4. 目标分解：将战略地图转化为BSC目标',
        '5. OKR转化：将BSC目标转化为OKR',
        '6. 横向对齐：各部门OKR相互对齐',
        '7. 纵向落地：OKR分解到部门、个人',
        '8. 执行追踪：定期Review和复盘'
    ]

    for step in steps:
        para = doc.add_paragraph()
        run = para.add_run(step)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    # Save document
    output_path = os.path.join(OUTPUT_DIR, '战略执行与落地_学员手册.docx')
    doc.save(output_path)
    print(f'Student handbook saved to: {output_path}')
    return output_path


def create_instructor_handbook():
    """Generate Instructor Handbook"""
    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== COVER PAGE =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('战略执行与落地')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = PRIMARY_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('——平衡计分卡与OKR实战')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = SECONDARY_COLOR

    for _ in range(4):
        doc.add_paragraph()

    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run('【讲师手册】')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = ACCENT_COLOR

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('授课时长：2天（每天6小时，共计12小时）')
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    toc_title = doc.add_paragraph()
    run = toc_title.add_run('目录')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    toc_items = [
        ('第一章', '课程设计逻辑', '3'),
        ('第二章', '教学目标与大纲', '5'),
        ('第三章', '讲师指引与话术', '8'),
        ('第四章', '案例分析要点', '15'),
        ('第五章', '练习答案', '20'),
        ('第六章', '常见问题应对', '25'),
        ('第七章', '评估方法', '28'),
    ]

    for ch, title, page in toc_items:
        para = doc.add_paragraph()
        run = para.add_run(f'{ch}  {title}')
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 1: COURSE DESIGN LOGIC =====
    h1 = doc.add_heading('第一章 课程设计逻辑', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('1.1 课程设计理念', level=2)
    content = doc.add_paragraph()
    run = content.add_run(
        '本课程采用"知行合一"的设计理念，强调从"知道"到"做到"的转化。'
        '理论框架（BSC+OKR）是基础，但更重要的是让学员掌握实际应用的工具和方法。'
        '课程设计遵循以下原则：'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    principles = [
        '从战略高度切入：从组织战略的高度理解BSC和OKR的价值',
        '强调工具落地：每个知识点都有对应的工具和模板',
        '案例本土化：采用国内企业（华为、字节等）的真实案例',
        '互动式学习：通过小组讨论、角色扮演等方式加深理解',
        '行动学习导向：要求学员制定个人行动计划，确保学习转化'
    ]
    for p in principles:
        para = doc.add_paragraph()
        run = para.add_run(f'• {p}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    h2 = doc.add_heading('1.2 教学方法设计', level=2)

    # Teaching methods table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    headers = ['教学方法', '占比', '目的']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    methods = [
        ('理论讲授', '30%', '建立概念框架，传授核心知识'),
        ('案例分析', '25%', '通过真实案例，加深理解，激发思考'),
        ('小组讨论', '20%', '促进经验分享，解决实际问题'),
        ('工具演练', '25%', '掌握实操工具，确保学习转化'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, row in enumerate(methods, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            set_cell_shading(table.rows[i].cells[0], colors[i-1])
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_heading('1.3 课程结构设计', level=2)
    content = doc.add_paragraph()
    run = content.add_run('课程采用"总-分-总"的结构设计：')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    structure = [
        '总：课程导入——建立战略执行的大画面（30分钟）',
        '分：BSC模块（3小时）——理解战略解码的工具',
        '分：OKR模块（3小时）——掌握目标聚焦的方法',
        '分：融合模块（3小时）——BSC与OKR的整合应用',
        '总：总结与行动学习（3小时）——制定个人行动计划'
    ]
    for s in structure:
        para = doc.add_paragraph()
        run = para.add_run(f'• {s}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ===== CHAPTER 2: TEACHING OBJECTIVES =====
    h1 = doc.add_heading('第二章 教学目标与大纲', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('2.1 总体教学目标', level=2)
    goals = [
        '认知层面：理解BSC和OKR的理论框架和核心逻辑',
        '能力层面：掌握战略解码、OKR制定和执行追踪的实操能力',
        '态度层面：建立战略思维和执行文化的意识'
    ]
    for g in goals:
        para = doc.add_paragraph()
        run = para.add_run(f'• {g}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    h2 = doc.add_heading('2.2 分章节教学目标', level=2)

    chapters = [
        ('第一章\n课程导入', '教学目标', '教学重点'),
        ('掌握战略执行的挑战与BSC+OKR的价值\n理解课程的整体框架和学习目标',
         '战略执行的常见挑战\nBSC与OKR的互补价值\n课程学习方法论',
         '开场破冰\n需求调研\n课程框架介绍'),
        ('第二章\nBSC四维模型', '教学目标', '教学重点'),
        ('理解BSC四维模型的内涵与因果逻辑\n掌握战略地图的绘制方法\n能够识别各维度的关键指标',
         'BSC的起源与发展\n四维模型的内涵\n战略地图绘制步骤\nBSC实施的关键成功因素',
         '四维模型讲解\n战略地图绘制演练\n华为案例分析'),
        ('第三章\nOKR实战', '教学目标', '教学重点'),
        ('掌握OKR的制定方法和评分标准\n理解OKR与KPI的区别\n能够制定具有挑战性的OKR',
         'OKR的起源与本质\nOKR结构（O+KR）\nOKR制定方法\nOKR评分标准\nOKR与KPI对比',
         'OKR制定演练\nOKR评分练习\n字节跳动案例分析'),
        ('第四章\nBSC与OKR融合', '教学目标', '教学重点'),
        ('理解BSC与OKR融合的必要性和方法\n掌握融合实施的具体步骤\n能够设计本组织的融合方案',
         '融合的必要性和价值\n融合模型设计\n融合实施步骤\n常见融合模式',
         '融合方案设计演练\n小组讨论'),
        ('第五章\n战略解码与落地', '教学目标', '教学重点'),
        ('理解战略解码的三个层次\n掌握战略落地的关键成功因素\n能够识别和避免常见误区',
         '战略解码的概念\n战略解码的三个层次\n战略落地的关键成功因素\n常见误区及应对',
         '战略解码演练\n角色扮演'),
        ('第六章\n案例分析', '教学目标', '教学重点'),
        ('通过案例分析加深对理论的理解\n学习标杆企业的实践经验\n激发学员的创新思考',
         '华为战略管理实践\n字节跳动OKR实践\n案例背后的成功要素\n可迁移经验的提炼',
         '案例阅读与思考\n小组讨论与分享\n讲师点评与提炼'),
        ('第七章\n行动学习', '制定个人行动学习计划\n建立持续学习和实践的意识',
         '行动学习的理念\n行动计划制定方法\n学习承诺与跟踪机制',
         '行动计划制定\n小组分享与反馈\n课程总结'),
    ]

    # Create a simpler table
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'

    headers = ['章节', '教学目标', '教学重点']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    chapters_data = [
        ('课程导入', '理解战略执行的挑战与BSC+OKR的价值', '开场破冰、需求调研'),
        ('BSC四维模型', '掌握战略地图绘制方法', '四维模型、战略地图'),
        ('OKR实战', '掌握OKR制定方法和评分标准', 'OKR制定、评分练习'),
        ('融合之道', '掌握BSC与OKR融合方法', '融合方案设计'),
        ('战略解码', '掌握战略落地关键因素', '解码演练、角色扮演'),
        ('案例分析', '学习标杆企业实践经验', '华为、字节案例'),
        ('行动学习', '制定个人行动学习计划', '计划制定、小组分享'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF', 'F2F2F2']
    for i, (ch, obj, focus) in enumerate(chapters_data, 1):
        table.rows[i].cells[0].text = ch
        table.rows[i].cells[1].text = obj
        table.rows[i].cells[2].text = focus
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(3):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ===== CHAPTER 3: INSTRUCTOR GUIDE =====
    h1 = doc.add_heading('第三章 讲师指引与话术', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('3.1 课程导入模块（30分钟）', level=2)

    h3 = doc.add_heading('开场破冰（10分钟）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('【话术】')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    content = doc.add_paragraph()
    run = content.add_run(
        "各位伙伴，大家好！欢迎来到《战略执行与落地》的课堂。在开始之前，我想请大家思考一个问题："
        "过去一年，你所在的组织制定的战略目标，有多少真正落地执行了？请大家举手示意一下。"
        "（观察举手情况）看来大家都很坦诚。今天这堂课，就是要帮助大家解决这个战略落地最后一公里的问题。"
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.font.italic = True

    doc.add_paragraph()

    h3 = doc.add_heading('需求调研（10分钟）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('【话术】')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    content = doc.add_paragraph()
    run = content.add_run(
        '"在正式进入课程之前，我想了解一下大家的情况。请各位用一句话说说，你们目前最关注的战略执行问题是什么？'
        '可以是你在推动战略落地过程中遇到的困惑，也可以是你观察到的问题。'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.font.italic = True

    doc.add_paragraph()

    h3 = doc.add_heading('课程框架介绍（10分钟）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('【话术】')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    content = doc.add_paragraph()
    run = content.add_run(
        '"今天的课程分为三个模块：上午我们学习BSC平衡计分卡；下午前半段学习OKR目标管理；'
        '下午后半段我们学习如何将两者融合。最后，每个人需要制定自己的行动学习计划。'
        '我希望大家带着问题来学，学完之后能够带着行动计划走。"'
    )
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.font.italic = True

    doc.add_paragraph()

    h2 = doc.add_heading('3.2 BSC模块讲师指引（3小时）', level=2)

    h3 = doc.add_heading('四维模型讲解（60分钟）', level=3)

    guidance = [
        ('讲解要点', 'BSC不仅仅是绩效考核工具，更是战略管理和沟通工具。要强调四个维度之间的因果关系。'),
        ('常用话术', '"财务指标是结果，客户指标是原因，内部流程是支撑，学习成长是基础。这四个维度形成了一条因果链。"'),
        ('互动设计', '让学员分组讨论：每个维度举出一个你所在组织正在使用的指标，看是否形成因果关系。'),
        ('时间控制', '概念讲解30分钟，小组讨论20分钟，分享10分钟。')
    ]

    for title, content_text in guidance:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}：')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run2 = para.add_run(content_text)
        run2.font.size = Pt(11)
        run2.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

    h3 = doc.add_heading('战略地图绘制（60分钟）', level=3)

    guidance = [
        ('讲解要点', '战略地图是BSC的可视化工具，通过因果链条串联四个维度的目标。'),
        ('常用话术', '"画战略地图就像画一幅从现在到未来的路线图。我们要回答：我们要去哪里（财务）？为谁创造价值（客户）？靠什么能力到达（内部流程）？需要什么支撑（学习成长）？"'),
        ('互动设计', '提供模板，让学员现场绘制自己部门的战略地图草稿。'),
        ('时间控制', '讲解20分钟，演练30分钟，小组分享10分钟。')
    ]

    for title, content_text in guidance:
        para = doc.add_paragraph()
        run = para.add_run(f'{title}：')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run2 = para.add_run(content_text)
        run2.font.size = Pt(11)
        run2.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ===== CHAPTER 4: CASE ANALYSIS KEY POINTS =====
    h1 = doc.add_heading('第四章 案例分析要点', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('4.1 华为案例分析要点', level=2)

    h3 = doc.add_heading('背景要点', level=3)
    background = [
        '华为1987年成立，从代理交换机起步，1990年开始自主研发',
        '1998年引入IBM进行战略规划咨询，逐步建立完善的战略管理体系',
        '华为的BSC实践具有"中西合璧"的特点，将西方管理工具与中国国情结合'
    ]
    for b in background:
        para = doc.add_paragraph()
        run = para.add_run(f'• {b}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('讨论问题', level=3)
    questions = [
        '华为为什么选择将BSC和KPI结合，而非单独使用OKR？',
        '华为的"述职制度"对战略落地起到什么作用？',
        '华为的经验中，哪些可以迁移到你们组织？'
    ]
    for q in questions:
        para = doc.add_paragraph()
        run = para.add_run(f'• {q}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('点评要点', level=3)
    points = [
        '华为的成功不是偶然，而是系统化的战略管理能力积累',
        '战略执行需要制度化、流程化，不能依赖个人英雄主义',
        '灰度管理是在战略清晰前提下的灵活执行，而非无原则的妥协'
    ]
    for p in points:
        para = doc.add_paragraph()
        run = para.add_run(f'• {p}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    h2 = doc.add_heading('4.2 字节跳动案例分析要点', level=2)

    h3 = doc.add_heading('背景要点', level=3)
    background = [
        '字节跳动2012年成立，十年时间成为全球领先的互联网科技公司',
        'OKR在字节叫"OBKR"，是公司的核心管理工具',
        '字节的OKR实践具有"北极星指标"的特色，每个阶段聚焦一个核心指标'
    ]
    for b in background:
        para = doc.add_paragraph()
        run = para.add_run(f'• {b}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('讨论问题', level=3)
    questions = [
        '字节的OKR实践与华为的BSC+KPI模式有何本质区别？',
        '"Context over Control"的理念对OKR落地有什么影响？',
        '你们组织更适合学习华为模式还是字节模式？为什么？'
    ]
    for q in questions:
        para = doc.add_paragraph()
        run = para.add_run(f'• {q}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    h3 = doc.add_heading('点评要点', level=3)
    points = [
        'OKR需要透明坦诚的文化土壤，文化不支持则OKR难以生根',
        'OKR不是考核工具，而是沟通对齐工具，这是其与传统绩效管理的本质区别',
        '互联网时代的战略执行需要高频复盘和快速迭代能力'
    ]
    for p in points:
        para = doc.add_paragraph()
        run = para.add_run(f'• {p}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===== CHAPTER 5: EXERCISE ANSWERS =====
    h1 = doc.add_heading('第五章 练习答案', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('5.1 练习一：BSC维度匹配', level=2)

    h3 = doc.add_heading('参考答案', level=3)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['维度', '对应指标']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    answers = [
        ('财务维度', 'a) 客户满意度评分、c) 净利润率、g) 存货周转率'),
        ('客户维度', 'e) 市场份额、h) 客户投诉率'),
        ('内部流程维度', 'd) 产品研发周期'),
        ('学习成长维度', 'b) 员工培训覆盖率、f) 员工流失率'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, (dim, indicators) in enumerate(answers, 1):
        table.rows[i].cells[0].text = dim
        table.rows[i].cells[1].text = indicators
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(2):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h3 = doc.add_heading('评分标准', level=3)
    criteria = [
        '完全正确（8个指标全对）：100分',
        '7-8个正确：90分',
        '5-6个正确：75分',
        '4个以下：需要重新学习'
    ]
    for c in criteria:
        para = doc.add_paragraph()
        run = para.add_run(f'• {c}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('5.2 练习二：OKR制定', level=2)

    h3 = doc.add_heading('参考答案', level=3)

    content = doc.add_paragraph()
    run = content.add_run('O（目标）：提升付费转化率，实现收入高质量增长')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    kr_items = [
        'KR1：付费用户数从100万提升至130万（提升30%）',
        'KR2：付费转化率从5%提升至7%（提升40%）',
        'KR3：首充用户7日留存率从30%提升至45%'
    ]
    for i, kr in enumerate(kr_items, 1):
        para = doc.add_paragraph()
        run = para.add_run(f'KR{i}：{kr}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    h3 = doc.add_heading('评判标准', level=3)
    criteria = [
        'O是否鼓舞人心、方向清晰？',
        'KR是否可量化、有挑战性（70%完成即为成功）？',
        'KR是否真正支撑O的达成？',
        '时间节点是否明确？'
    ]
    for c in criteria:
        para = doc.add_paragraph()
        run = para.add_run(f'☐ {c}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ===== CHAPTER 6: FAQ =====
    h1 = doc.add_heading('第六章 常见问题应对', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('6.1 学员常见问题', level=2)

    faqs = [
        ('Q1: BSC和OKR哪个更好？', 'A: 这是伪命题。BSC和OKR解决不同问题，前者解决战略解码的系统性问题，后者解决目标聚焦的突破性问题。成熟企业通常两者结合使用。'),
        ('Q2: OKR是否要与考核挂钩？', 'A: 建议初期不直接挂钩。OKR的核心是激发挑战性目标，如果与考核强挂钩，学员会倾向于设定保守目标。可以通过文化引导和荣誉激励来间接关联。'),
        ('Q3: BSC指标太多怎么办？', 'A: 每个维度选择3-5个核心指标即可。指标过多会导致焦点分散，失去战略引领作用。'),
        ('Q4: 如何推动高层参与？', 'A: 高层参与是BSC/OKR成功的关键因素。可以通过战略解码工作坊的形式，让高层参与目标制定过程，增强承诺感。'),
        ('Q5: 如何处理部门间的OKR冲突？', 'A: 需要建立OKR对齐机制。公司层面设立"OKR评审委员会"，定期审视跨部门OKR的一致性，及时协调冲突。'),
    ]

    for q, a in faqs:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(4)

        para = doc.add_paragraph()
        run = para.add_run(a)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    h2 = doc.add_heading('6.2 实施难点与应对策略', level=2)

    # Difficulties table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['难点', '应对策略']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    difficulties = [
        ('高层参与不足', '将BSC/OKR纳入高管团队例会议程；设计高管专属的战略解码工作坊'),
        ('部门墙导致对齐困难', '建立跨部门OKR对齐会议机制；用公司级OKR牵引部门级OKR'),
        ('重制定轻追踪', '建立双周/月度OKR Review机制；将OKR追踪纳入管理者考核'),
        ('文化不支持', '通过内部宣传和培训建立认知；树立内部标杆案例进行推广'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, (diff, strategy) in enumerate(difficulties, 1):
        table.rows[i].cells[0].text = diff
        table.rows[i].cells[1].text = strategy
        set_cell_shading(table.rows[i].cells[0], colors[i-1])
        for j in range(2):
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ===== CHAPTER 7: ASSESSMENT METHODS =====
    h1 = doc.add_heading('第七章 评估方法', level=1)
    h1.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    h2 = doc.add_heading('7.1 学员评估体系', level=2)

    # Assessment table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ['评估维度', '评估方式', '权重', '评估时间']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    assessments = [
        ('课堂参与', '出勤、发言、小组贡献', '20%', '课程进行中'),
        ('工具演练', 'BSC战略地图、OKR制定', '30%', '课程进行中'),
        ('行动计划', '个人行动学习计划质量', '30%', '课程结束时'),
        ('知识掌握', '笔试/口头问答', '20%', '课程结束时'),
    ]

    colors = ['F2F2F2', 'FFFFFF', 'F2F2F2', 'FFFFFF']
    for i, row in enumerate(assessments, 1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            set_cell_shading(table.rows[i].cells[0], colors[i-1])
            table.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_heading('7.2 课程效果评估', level=2)

    h3 = doc.add_heading('反应层评估（课程结束时）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('通过问卷调研了解学员对课程内容、讲师、培训组织的满意度。')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    h3 = doc.add_heading('学习层评估（课程结束后1个月）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('通过跟进访谈或在线测试，评估学员对BSC和OKR知识的掌握程度和应用情况。')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    h3 = doc.add_heading('行为层评估（课程结束后3个月）', level=3)
    content = doc.add_paragraph()
    run = content.add_run('通过管理者反馈或系统数据，评估学员是否将所学应用于实际工作，OKR是否真正落地。')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    h2 = doc.add_heading('7.3 讲师自评清单', level=2)

    self_eval = [
        '课程目标是否达成？学员反馈如何？',
        '时间控制是否合理？哪些环节需要调整？',
        '案例是否引发学员共鸣？讨论是否充分？',
        '学员的独特贡献是什么？如何记录和反馈？',
        '下次课程可以改进的地方是什么？'
    ]
    for item in self_eval:
        para = doc.add_paragraph()
        run = para.add_run(f'☐ {item}')
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        para.paragraph_format.space_after = Pt(6)

    # Save document
    output_path = os.path.join(OUTPUT_DIR, '战略执行与落地_讲师手册.docx')
    doc.save(output_path)
    print(f'Instructor handbook saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    print('Generating BSC & OKR Strategic Execution Handbooks...')
    print('=' * 50)

    student_path = create_student_handbook()
    instructor_path = create_instructor_handbook()

    print('=' * 50)
    print('Generation complete!')
    print(f'Student handbook: {student_path}')
    print(f'Instructor handbook: {instructor_path}')

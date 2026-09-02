from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set default font for Chinese
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Title
title = doc.add_heading('柔性生产与动态排程：从固定节拍到实时响应', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('讲师手册')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(16)
subtitle_run.font.bold = True

# Metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('版本：v1.0 | 课程时长：6小时（1天）\n').font.size = Pt(10)
meta.add_run('适用对象：生产计划主管、精益推进负责人、车间主任、智能制造项目负责人').font.size = Pt(10)

doc.add_paragraph()

# ============ 讲师指引 ============
doc.add_heading('讲师指引', level=1)

# 课程目标
doc.add_heading('课程目标', level=2)
doc.add_paragraph('帮助学员理解柔性生产的本质和AI排程的原理，掌握在AI排程中保持精益原则的判断力，知道如何推动动态排程的落地实施。')

# 学员画像
doc.add_heading('学员画像', level=2)
bullets = [
    '有一定的精益生产基础（了解一个流、拉动式生产的基本概念）',
    '有生产排程经验（传统方式为主）',
    '对AI排程有期待也有困惑',
    '职位：计划主管、车间主任、精益推进负责人'
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

# 教学理念
doc.add_heading('教学理念', level=2)
p = doc.add_paragraph()
p.add_run('"精益为锚，AI为帆"').bold = True
p.add_run(' ——AI排程是精益原则的新实现方式，不是精益原则的替代者。管理者的新能力不是自己做排程，而是判断AI排程的建议是否符合精益原则。')

# 讲师准备清单
doc.add_heading('讲师准备清单', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
headers = ['准备项', '要求', '状态']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

items = [
    ('学员手册印刷', '每人一本，课程前两天到位', '□'),
    ('白板/大白纸', '用于小组讨论后展示', '□'),
    ('场景卡收集', '课程前一天收集，了解学员背景', '□'),
    ('案例准备', '准备1-2个本行业的AI排程案例', '□'),
    ('笔记本电脑', '用于演示AI排程系统（如有）', '□')
]
for i, (item, req, status) in enumerate(items, 1):
    table.rows[i].cells[0].text = item
    table.rows[i].cells[1].text = req
    table.rows[i].cells[2].text = status

doc.add_paragraph()

# 授课原则
doc.add_heading('授课原则', level=2)
principles = [
    ('互动优先', '每个知识点后必须有互动，不要连续讲超过20分钟'),
    ('场景贯穿', '用学员自己的场景卡贯穿全程，不是讲通用案例'),
    ('框架清晰', '每个章节开头画框架图，结尾收回来'),
    ('金句点睛', '每个章节至少一个金句，帮助学员记忆')
]
for title_text, desc in principles:
    p = doc.add_paragraph()
    p.add_run(f'{title_text}：').bold = True
    p.add_run(desc)

# ============ 时间分配建议 ============
doc.add_heading('时间分配建议', level=1)
table = doc.add_table(rows=13, cols=3)
table.style = 'Table Grid'
headers = ['环节', '时长', '内容']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

schedule = [
    ('开场', '30分钟', '课程介绍、全景图、场景卡确认、自评'),
    ('第一章', '50分钟', '柔性生产本质'),
    ('休息', '10分钟', ''),
    ('第二章', '50分钟', 'AI排程原理解析'),
    ('午休', '60分钟', ''),
    ('第三章', '50分钟', '动态排程与精益原则'),
    ('休息', '10分钟', ''),
    ('第四章', '50分钟', '人机协同判断力'),
    ('休息', '10分钟', ''),
    ('第五章', '40分钟', '现场实施路径'),
    ('第六章', '30分钟', '案例复盘与总结'),
    ('收尾', '20分钟', '行动承诺、30天计划')
]
for i, (item, dur, content) in enumerate(schedule, 1):
    table.rows[i].cells[0].text = item
    table.rows[i].cells[1].text = dur
    table.rows[i].cells[2].text = content

# ============ 授课要点 ============
doc.add_heading('授课要点', level=1)

# Chapter 1
doc.add_heading('第一章　柔性生产本质', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('让学员理解"柔性"的本质不是"随便动"，而是"快速响应"，掌握柔性生产的三层定义，识别自己场景中柔性不足的具体表现。')

doc.add_heading('时间分配（50分钟）', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['内容', '时长', '教学方法']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
ch1_time = [
    ('知识点讲解', '20分钟', '框架图 + 互动提问'),
    ('表单练习', '15分钟', '个人思考 + 两人讨论'),
    ('练习讲解', '10分钟', '随机提问 + 点评'),
    ('收尾', '5分钟', '框架回顾 + 行为承诺')
]
for i, (c, d, m) in enumerate(ch1_time, 1):
    table.rows[i].cells[0].text = c
    table.rows[i].cells[1].text = d
    table.rows[i].cells[2].text = m

doc.add_heading('核心内容', level=2)
doc.add_heading('开场问题（互动引入）', level=3)
doc.add_paragraph('"在座的各位，有多少人现在还在用固定节拍生产？你们的生产计划是按班次、按天来排的，还是按小时、按分钟来排的？"')
doc.add_paragraph('这个问题用来了解学员的现状，为后面的内容做铺垫。')

doc.add_heading('知识点1.1：固定节拍为什么会失效', level=3)
doc.add_paragraph('讲清楚四个时代背景的变化：')
for item in ['多品种小批量成为常态', '需求波动剧烈', '交货期大幅压缩', '定制化需求']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"固定节拍不是错了，它在它的时代是最优解。问题是时代变了。精益的\'消除浪费\'内核没变，但\'如何消除浪费\'的方法需要升级。"', style='Quote')

doc.add_heading('知识点1.2：柔性的三层定义', level=3)
doc.add_paragraph('用框架图来讲：')
framework = """系统柔性（想得清）
　＝　实时感知变化　＋　快速计算优化　＋　动态调整执行

工艺柔性（摸得着）
　＝　同一条线，多种工艺路线切换

设备柔性（看得见）
　＝　快速换线、自动切换"""
doc.add_paragraph(framework)

p = doc.add_paragraph()
p.add_run('强调：').bold = True
p.add_run('AI排程解决的是第三层问题（系统柔性），前两层是基础，第三层是升级。')

doc.add_heading('知识点1.3：柔性的三种浪费', level=3)
for item in [
    '等待浪费：等料、等人、等信息',
    '切换浪费：换线、换模、换规格',
    '响应浪费：市场信号在传递中失真和延迟'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('表单处理要点', level=2)
doc.add_heading('表单1.1：柔性现状诊断', level=3)
doc.add_paragraph('这个表单的目的是让学员客观评估自己所在场景的柔性现状。三个维度（设备、工艺、系统）都要覆盖，不能只谈设备。')

p = doc.add_paragraph()
p.add_run('引导问题：').bold = True
p.add_run('"你们公司现在最缺的是哪一层柔性？是设备跑不快，还是工艺切换慢，还是系统响应不及时？"')

doc.add_heading('练习处理要点', level=2)
doc.add_heading('练习1-A：概念辨析', level=3)
doc.add_paragraph('第2题"只要买了AI排程系统，现场的柔性就能自动提升"——这道题的错误率通常比较高，要重点讲。')

p = doc.add_paragraph()
p.add_run('参考答案：').bold = True
for ans in [
    '1. 错。柔性生产不等于低效，恰恰相反，柔性好的企业响应更快、库存更低。',
    '2. 错。AI排程解决的是系统柔性，设备柔性和工艺柔性是前提条件。',
    '3. 对。这是柔性的本质。'
]:
    doc.add_paragraph(ans, style='List Number')

doc.add_heading('常见问题与应答', level=2)
qa_pairs = [
    ('Q：我们的订单批量还是很大的，是不是不需要柔性？',
     'A：批量大不代表不需要柔性。要问的是：你的批量是"主动选择"还是"被动接受"？如果是主动选择的大批量，那没问题；如果是小批量做不了、只能积压成大批量，那就是柔性不足。'),
    ('Q：柔性升级投入很大，老板不批怎么办？',
     'A：先算柔性不足的代价——库存积压、交付延迟、换线浪费。用数据说话，比讲道理更有效。')
]
for q, a in qa_pairs:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)

# Chapter 2
doc.add_heading('第二章　AI排程原理解析', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('让学员理解AI排程的三种技术路线，掌握AI排程系统的基本工作流程，能识别AI排程常见的"聪明"和"不聪明"的表现。')

doc.add_heading('核心内容', level=2)
doc.add_heading('知识点2.1：AI排程的三种技术路线', level=3)
doc.add_paragraph('用类比来讲：')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['技术路线', '类比', '特点']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
tech = [
    ('规则引擎型', '老法师的经验清单', '稳定、可解释，但难以发现隐藏最优解'),
    ('运筹优化型', '数学家建模求解', '全局最优，但计算复杂度高'),
    ('机器学习型', '徒弟跟师傅学', '能处理复杂模式，但需要大量数据')
]
for i, (t, m, f) in enumerate(tech, 1):
    table.rows[i].cells[0].text = t
    table.rows[i].cells[1].text = m
    table.rows[i].cells[2].text = f

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"目前的工业级AI排程系统，大多数是\'混合型\'——用规则引擎保证稳定性，用运筹优化保证最优性，用机器学习提供预测能力。"', style='Quote')

doc.add_heading('知识点2.2：AI排程的工作流程', level=3)
doc.add_paragraph('用流程图来讲：')
flow = """需求输入 → 约束解析 → 优化计算 → 方案输出 → 人工确认 → 执行反馈
    ↓          ↓           ↓          ↓          ↓          ↓
  订单、库存   产能、途程    排程算法    甘特图      调整、批准    实际完成
  市场预测    交期、工艺    仿真验证    报表        下发执行    数据回传"""
doc.add_paragraph(flow)

p = doc.add_paragraph()
p.add_run('关键点：').bold = True
p.add_run('数据质量决定输出质量。强调"garbage in, garbage out"的道理。')

doc.add_heading('知识点2.3：AI排程的"聪明"与"不聪明"', level=3)

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
headers = ['AI"聪明"的场景', 'AI"不聪明"的场景']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
smart = [
    ('大量工单同时排程', '新产品导入，没有历史数据'),
    ('多工序、多约束的复杂网络', '跨部门协调（AI不懂"政治"）'),
    ('需要频繁调整的动态环境', '紧急插单'),
    ('追求整体最优而非局部最优', '需要权衡精益原则和交付效率')
]
for i, (a, b) in enumerate(smart, 1):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"AI排程是在给定目标函数和约束条件下的优化。如果你没有告诉它精益原则，它不会自己加进去。"', style='Quote')

doc.add_heading('常见问题与应答', level=2)
qa_pairs = [
    ('Q：AI排程系统多少钱？',
     'A：这个问题背后是"值不值"的问题。不要直接回答价格，而是反问："你们现在因为排程不合理造成的损失是多少？库存积压多少？交付延迟多少次？"把ROI算出来，价格就不重要了。'),
    ('Q：我们公司数据很差，能用AI排程吗？',
     'A：能用，但要先做数据准备。AI排程不是"上了系统就能用"，而是需要先把数据整理好。先做数据诊断，再决定要不要上系统。'),
    ('Q：AI排程会不会取代计划员？',
     'A：不会取代，但会重新分工。计划员从"做排程"变成"判断排程"。这个转变需要时间，也需要培训。')
]
for q, a in qa_pairs:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)

# Chapter 3
doc.add_heading('第三章　动态排程与精益原则', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('让学员识别AI排程建议与精益原则的四种典型冲突场景，掌握在AI排程中保持精益原则的四种策略，能针对具体场景设计"精益增强型"排程规则。')

doc.add_heading('核心内容', level=2)
doc.add_heading('知识点3.1：为什么AI的"最优"不等于精益的"最优"', level=3)
doc.add_paragraph('这是本章最关键的认知转变点。用具体案例来讲：')

case_text = """AI排程可能建议你"为了减少换线，把这周的订单集中在周三统一切换"。这在数学上是"最优"的（换线次数最少）。

但精益原则会问：这样做的代价是什么？是在制品积压三天？客户能等吗？现场人员的工作负荷均衡吗？"""
doc.add_paragraph(case_text)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"AI排程追求的是数学最优，精益追求的是系统最优。两者有时候重合，有时候分歧。管理者需要知道什么时候该信AI，什么时候该修正AI。"', style='Quote')

doc.add_heading('知识点3.2：四种典型冲突场景', level=3)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['冲突类型', 'AI的倾向', '精益的原则']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
conflicts = [
    ('批量集中 vs 流动连续', '批量处理减少换线', '小批流动保持响应'),
    ('产能最大化 vs 库存最小化', '把产能填满', '保留缓冲产能'),
    ('交付优先 vs 均衡生产', '优先排紧急单', '考虑上下游均衡负荷'),
    ('效率最优 vs 人员负荷', '集中换线', '考虑人员疲劳')
]
for i, (c, a, l) in enumerate(conflicts, 1):
    table.rows[i].cells[0].text = c
    table.rows[i].cells[1].text = a
    table.rows[i].cells[2].text = l

doc.add_paragraph()
doc.add_heading('知识点3.3：四种精益增强策略', level=3)
doc.add_paragraph('用框架图来讲：')
framework2 = """精益增强策略
├── 约束注入：把精益原则变成约束条件
├── 目标函数加权：在AI目标中加入精益指标
├── 规则优先：把精益原则定义为硬规则
└── 人机协同：AI建议 + 人工判断"""
doc.add_paragraph(framework2)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"不是让AI去\'理解\'精益，而是让精益原则变成AI能执行的规则。AI不懂\'一个流\'的价值，但你可以把\'在制品上限\'变成约束条件。"', style='Quote')

# Chapter 4
doc.add_heading('第四章　人机协同判断力', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('让学员建立"AI建议 vs 精益原则"的对照检查思维，掌握判断AI排程建议的五个关键问题，能用精益语言向现场人员解释排程决策。')

doc.add_heading('核心内容', level=2)
doc.add_heading('知识点4.1：人机协同的新分工', level=3)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
headers = ['传统分工', '人机协同新分工']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
table.rows[1].cells[0].text = '计划员做排程'
table.rows[1].cells[1].text = 'AI做：大量计算、约束权衡'
table.rows[2].cells[0].text = '现场执行'
table.rows[2].cells[1].text = '人做：判断AI建议是否符合精益原则'

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"管理者的新能力不是\'自己做排程\'，而是\'判断AI的排程建议是否正确\'。这个能力，AI给不了你，只能靠你自己建立。"', style='Quote')

doc.add_heading('知识点4.2：判断AI排程建议的五个关键问题', level=3)
doc.add_paragraph('用检查清单来讲：')
checklist = """□ 1. 流动原则：是否满足精益流动？（在制品、等待）
□ 2. 实际约束：设备/人员/物料是否真的到位？
□ 3. 目标一致：AI目标和真实目标一致吗？
□ 4. 抱怨预判：如果执行，谁会抱怨？为什么？
□ 5. 长期效果：这个排程的长期影响是什么？"""
doc.add_paragraph(checklist)

doc.add_heading('知识点4.3：用精益语言解释排程决策', level=3)
doc.add_paragraph('用句式模板来讲：')
template_text = '"我们选择这样排，是因为……（精益原则）\n 而不是单纯因为……（AI的某个指标）\n 这样做的代价是……（被牺牲的指标）\n 我们接受这个代价，因为……（整体最优的逻辑）"'
doc.add_paragraph(template_text)

doc.add_paragraph()
doc.add_paragraph('示例：')
doc.add_paragraph('"我们选择每天换一次线而不是每周集中换一次，是因为我们希望保持流动、减少在制品库存（精益原则）。AI建议集中换线是为了提高换线效率，但我们认为库存积压三天的代价更高（权衡逻辑）。"', style='Quote')

# Chapter 5
doc.add_heading('第五章　现场实施路径', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('让学员理解AI排程实施的四个阶段，识别每个阶段的关键成功因素和常见陷阱，制定自己所在场景的AI排程实施路线图。')

doc.add_heading('核心内容', level=2)
doc.add_heading('知识点5.1：实施四个阶段', level=3)
doc.add_paragraph('用阶段图来讲：')
phases = """第一阶段：数据准备（Data Preparation）
    ↓
第二阶段：规则配置（Rule Configuration）
    ↓
第三阶段：人机磨合（Human-Machine Alignment）
    ↓
第四阶段：持续优化（Continuous Improvement）"""
doc.add_paragraph(phases)

doc.add_paragraph('每个阶段的重点和常见陷阱：')
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['阶段', '重点', '常见陷阱']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
phase_data = [
    ('数据准备', '数据标准化、历史清洗', '低估工作量'),
    ('规则配置', '业务规则→约束条件', '一步到位'),
    ('人机磨合', 'AI建议→人工确认→反馈', '过度依赖或完全不信'),
    ('持续优化', '定期复盘、持续改进', '上线即结束')
]
for i, (p, k, p2) in enumerate(phase_data, 1):
    table.rows[i].cells[0].text = p
    table.rows[i].cells[1].text = k
    table.rows[i].cells[2].text = p2

doc.add_paragraph()
doc.add_heading('知识点5.2：关键成功因素', level=3)
doc.add_paragraph('四个因素：')
factors = """关键成功因素
├── 一把手工程（不是挂名，是真正理解和支持）
├── 现场参与（用户参与，不闭门造车）
├── 容错机制（允许犯错，在错误中学习）
└── 精益为基础（AI是工具，不是替代）"""
doc.add_paragraph(factors)

doc.add_heading('知识点5.3：常见的"坑"和避坑指南', level=3)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['坑的类型', '具体表现', '避坑指南']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
pitfalls = [
    ('数据坑', '"系统上线后发现数据不对"', '先做数据诊断，再上系统'),
    ('期望坑', '"AI排程应该全自动，不需要人"', '建立人机协同的预期'),
    ('权力坑', '"计划员觉得被AI取代了"', '把AI定位为工具，不是替代者'),
    ('短期坑', '"上线三个月没效果，就放弃了"', '建立阶段性里程碑，持续跟踪')
]
for i, (t, d, g) in enumerate(pitfalls, 1):
    table.rows[i].cells[0].text = t
    table.rows[i].cells[1].text = d
    table.rows[i].cells[2].text = g

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"AI排程的实施，不是上一个系统那么简单。它是生产管理模式的变革，需要系统性的规划和持续的优化。"', style='Quote')

# Chapter 6
doc.add_heading('第六章　案例复盘与总结', level=1)

doc.add_heading('章节目标', level=2)
doc.add_paragraph('通过三个真实案例，让学员看清AI排程落地的完整路径，把前五章的知识串成一条完整的行动链。')

doc.add_heading('核心内容', level=2)

doc.add_heading('案例一：某电子组装厂的"数据先行"', level=3)
doc.add_paragraph('关键信息：')
for item in [
    '第一阶段花了4个月做数据准备',
    '梳理全部52个产品系列的工艺路线',
    '修正过去3年的标准工时数据（发现30%的数据错误）',
    '数据准备完成后，AI排程上线首月达成率从65%提升到88%'
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"数据质量不是上线前的问题，是每天都在做的事情。"', style='Quote')

doc.add_heading('案例二：某汽车零部件厂的"规则之争"', level=3)
doc.add_paragraph('关键信息：')
for item in [
    '计划主管坚持"交货期就是一切"',
    '上线后第一个月：紧急订单少了，但现场怨声载道',
    '问题诊断：目标函数过于单一，牺牲了精益原则',
    '改进措施：引入精益指标作为约束条件',
    '第二个月达成率稳定在85%，现场抱怨减少70%'
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"AI排程不是让机器做决定，而是让人用机器做更好的决定。"', style='Quote')

doc.add_heading('案例三：某食品加工厂的"短期陷阱"', level=3)
doc.add_paragraph('关键信息：')
for item in [
    '上线三个月后没有看到"立竿见影"的效果',
    '管理层决定放弃',
    '问题诊断：没有建立阶段性里程碑，没有跟踪复盘机制',
    '教训：三个月看不到效果，可能是因为还在积累期'
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('金句：').bold = True
doc.add_paragraph('"AI排程的失败，大多数不是因为技术不行，而是因为坚持不够。"', style='Quote')

# 课程收尾
doc.add_heading('课程收尾', level=1)

doc.add_heading('收尾话术', level=2)
script = """各位，今天我们用一天时间，走完了从"固定节拍"到"实时响应"的认知重建。

这门课有一个核心信念——精益的内涵没变，实现方式在变。

固定节拍生产是工业时代的骄傲，它的精益原则（流动、拉动、消除浪费）在今天依然有效。只是在多品种小批量、快交付的市场环境下，实现这些原则的方式需要升级。AI排程，是这个升级的工具，而不是精益原则的替代者。

管理者的新能力，不是自己做排程，而是判断AI排程的建议是否符合精益原则。这个能力，AI给不了你，只能靠你自己建立。

从今天，从现在，从这一次开始。

谢谢大家。"""
doc.add_paragraph(script)

doc.add_heading('课程金句汇总', level=2)
golden_quotes = [
    '"精益不是一成不变的流程，而是一种持续优化的思维方式。AI时代，精益的内涵没变，实现方式变了。"',
    '"固定节拍不是错了，它在它的时代是最优解。问题是时代变了。"',
    '"AI排程不是黑魔法，它的本质是数学优化——理解这个本质，才能更好地驾驭它。"',
    '"AI排程是在给定目标函数和约束条件下的优化。如果你没有告诉它精益原则，它不会自己加进去。"',
    '"AI排程追求的是数学最优，精益追求的是系统最优。两者有时候重合，有时候分歧。"',
    '"不是让AI去\'理解\'精益，而是让精益原则变成AI能执行的规则。"',
    '"管理者的新能力不是\'自己做排程\'，而是\'判断AI的排程建议是否正确\'。"',
    '"AI排程不是让机器做决定，而是让人用机器做更好的决定。"',
    '"AI排程的实施，不是上一个系统那么简单。它是生产管理模式的变革。"',
    '"领先一步，枪打出头鸟；落后半步，别人牵牛我拔桩；领先半步，吃尽红利。"'
]
for i, q in enumerate(golden_quotes, 1):
    doc.add_paragraph(f'{i}. {q}')

# Appendix
doc.add_heading('附录', level=1)

doc.add_heading('延伸阅读', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['书籍/文章', '作者', '推荐理由']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
books = [
    ('《精益思想》', 'James Womack', '精益生产原典，理解精益原则必读'),
    ('《改变世界的机器》', 'James Womack', '丰田生产方式起源'),
    ('《工厂物理学》', 'John Hopp', '生产系统的基础理论模型'),
    ('《AI极简经济学》', 'Ajay Agrawal', '理解AI的能力边界')
]
for i, (title_text, author, reason) in enumerate(books, 1):
    table.rows[i].cells[0].text = title_text
    table.rows[i].cells[1].text = author
    table.rows[i].cells[2].text = reason

doc.add_heading('案例库索引', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['案例类型', '来源', '获取方式']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
cases = [
    ('电子组装', '课程案例集', '课程资料包'),
    ('汽车零部件', '课程案例集', '课程资料包'),
    ('食品加工', '课程案例集', '课程资料包')
]
for i, (t, s, g) in enumerate(cases, 1):
    table.rows[i].cells[0].text = t
    table.rows[i].cells[1].text = s
    table.rows[i].cells[2].text = g

doc.add_heading('评估标准', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['评估维度', '评估方式', '达标标准']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
eval_data = [
    ('知识掌握', '课堂问答 + 练习正确率', '80%以上正确'),
    ('场景贯穿', '场景卡完成度', '完整填写，有深度'),
    ('行动承诺', '30天计划', '有具体行动，有检验指标'),
    ('满意度', '课后问卷', '4.0/5.0以上')
]
for i, (d, m, s) in enumerate(eval_data, 1):
    table.rows[i].cells[0].text = d
    table.rows[i].cells[1].text = m
    table.rows[i].cells[2].text = s

# Footer
doc.add_paragraph()
doc.add_paragraph('版权所有 · 罗宏伟 · 本手册仅供本课程讲师使用')

# Save
output_path = 'D:/新课开发/精益/6.柔性生产与动态排程：从固定节拍到实时响应/讲师手册/讲师手册_柔性生产与动态排程.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')

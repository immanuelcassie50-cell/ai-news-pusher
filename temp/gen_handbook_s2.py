#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# TOC
toct = doc.add_paragraph()
toct.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toct.add_run('目 录')
r.bold = True; r.font.size = Pt(24)
doc.add_paragraph()

toc_items = [
    ('学习目标', '4'), ('课前准备', '5'),
    ('PART 1  认知升级：复盘与决策卡的本质区分', '6'),
    ('  第一章  复盘写的是过去，决策卡写的是未来', '6'),
    ('  第二章  不是每个决策都值得做成一张卡', '7'),
    ('  第三章  复盘访谈问的是你当时不确定什么', '8'),
    ('  第四章  隐性判断是问出来的，不是想出来的', '9'),
    ('  第五章  决策卡是检查表加开关', '10'),
    ('  第六章  触发条件写不清楚这张卡就是废纸', '11'),
    ('  第七章  场景映射矩阵', '12'),
    ('  第八章  失败案例是用来对照的', '13'),
    ('  第九章  第一个读者应该是反对你的人', '14'),
    ('  第十章  训练是让人在场景里犯一次错', '15'),
    ('PART 1  Q&A', '16'),
    ('PART 2  组织落地', '17'),
    ('  第一章  没人认领是最大的敌人', '17'),
    ('  第二章  稽核不是查错是记住学过什么', '18'),
    ('  第三章  会过时是活着的证据', '19'),
    ('  第四章  最易犯的错是替决策者总结', '20'),
    ('  第五章  交付的是组织判断力', '21'),
    ('PART 2  Q&A', '22'),
    ('工具模板', '23'),
    ('课后资源', '30'),
    ('附录', '32'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.add_run('\t' + page)

doc.add_page_break()
print('S2: TOC done')
doc.save(r'D:\CC\temp\handbook_s2.docx')

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 工具模板
tools = doc.add_paragraph()
r = tools.add_run('工具模板')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

# 附录一：决策卡标准模板
doc.add_paragraph()
app1 = doc.add_paragraph()
r = app1.add_run('附录一：决策卡标准模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

t1 = doc.add_table(rows=8, cols=2)
t1.style = 'Table Grid'

c = t1.cell(0, 0)
c.merge(t1.cell(0, 1))
set_cell_shading(c, 'D9E2F3')
r = c.paragraphs[0].add_run('【决策卡名称】')
r.bold = True

tmpl1 = [
    ('触发条件', '（条件句+动作指令，写在卡片最显眼位置）\n如果出现[具体可观测信号]，先做[具体核实动作]，再决定是否继续。'),
    ('检查表', '（5-8条具体可执行的核对项，每条用能直接对照的动词开头）\n1. [具体核对动作1]\n2. [具体核对动作2]\n3. [具体核对动作3]\n...'),
    ('应急方案', '（当触发条件满足且出现变体场景时的处理方式）\n• 如果[场景变体A]，则[处理方式]\n• 如果[场景变体B]，则[处理方式]'),
    ('适用场景', '（说明本卡适用的具体情境，以及不适用的情境）'),
    ('警示案例', '（本卡对应的失败岔路口描述，用于嵌入式警示对照）'),
    ('认领人', '姓名：      联系方式：\n最近更新日期：'),
]

for i, (label, content) in enumerate(tmpl1, 1):
    cl = t1.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True
    t1.cell(i, 1).paragraphs[0].add_run(content)

c = t1.cell(7, 0)
c.merge(t1.cell(7, 1))
set_cell_shading(c, 'FFF2CC')
r = c.paragraphs[0].add_run('定位说明：本卡列出的触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。')
r.italic = True

doc.add_page_break()
print('S11: Tool template 1 done')
doc.save(r'D:\CC\temp\handbook_s11.docx')

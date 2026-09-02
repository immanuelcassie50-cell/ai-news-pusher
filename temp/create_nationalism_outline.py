# -*- coding: utf-8 -*-
"""
创建《民族主义思想史——一个概念如何塑造现代世界》课程大纲Word文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 配色方案
PRIMARY = RGBColor(0x1a, 0x27, 0x44)   # 深墨蓝 #1a2744
SECONDARY = RGBColor(0x8b, 0x28, 0x28) # 暗红 #8b2828
ACCENT = RGBColor(0xc9, 0xa9, 0x6e)    # 金色 #c9a96e
LIGHT = RGBColor(0xf5, 0xf0, 0xe6)     # 米白 #f5f0e6
BACKGROUND = RGBColor(0xfa, 0xf8, 0xf5)# 暖白 #faf8f5
BODY_TEXT = RGBColor(0x33, 0x33, 0x33) # 深灰 #333333
SUBTLE = RGBColor(0x66, 0x66, 0x66)    # 中灰 #666666

OUTPUT_PATH = "D:/新课开发/政治学/18_民族主义思想史-一个概念如何塑造现代世界/课程大纲/01_课程大纲.docx"

# ─────────────────────────────────────────────
# 辅助函数
# ─────────────────────────────────────────────

def rgb_to_hex(color: RGBColor) -> str:
    return '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])

def set_cell_background(cell, color: RGBColor):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = rgb_to_hex(color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val.get('val', 'single'))
            border.set(qn('w:sz'), str(val.get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_horizontal_line(doc, color: RGBColor = ACCENT, thickness: int = 12):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    hex_color = rgb_to_hex(color)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_para_border_bottom(para, color: RGBColor, thickness: int = 6):
    """给段落添加底部边框"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    hex_color = rgb_to_hex(color)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_cjk(run, text, font_name='微软雅黑'):
    """设置中文Run的字体"""
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), font_name)

# ─────────────────────────────────────────────
# 主文档创建
# ─────────────────────────────────────────────

doc = Document()

# 页面设置：A4，上下左右边距2cm
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.79)
section.bottom_margin = Inches(0.79)
section.left_margin = Inches(0.79)
section.right_margin = Inches(0.79)

# 全局默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ─────────────────────────────────────────────
# 封面
# ─────────────────────────────────────────────

# 顶部装饰条
top_bar = doc.add_paragraph()
top_bar.paragraph_format.space_before = Pt(0)
top_bar.paragraph_format.space_after = Pt(0)
top_bar.paragraph_format.line_spacing = 0
pPr = top_bar._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
hex_bg = rgb_to_hex(PRIMARY)
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), hex_bg)
pPr.append(shd)
run = top_bar.add_run('  ')
run.font.size = Pt(8)

# 课程编号标签
label_para = doc.add_paragraph()
label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
label_para.paragraph_format.space_before = Pt(12)
label_para.paragraph_format.space_after = Pt(0)
run = label_para.add_run('政治学类-18')
run.font.size = Pt(9)
run.font.bold = True
run.font.color.rgb = SECONDARY
set_run_cjk(run, '政治学类-18')

# 主标题
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(24)
title_para.paragraph_format.space_after = Pt(8)
run = title_para.add_run('民族主义思想史')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = PRIMARY
set_run_cjk(run, '民族主义思想史')

# 副标题
subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_para.paragraph_format.space_before = Pt(0)
subtitle_para.paragraph_format.space_after = Pt(24)
run = subtitle_para.add_run('一个概念如何塑造现代世界')
run.font.size = Pt(22)
run.font.color.rgb = SECONDARY
set_run_cjk(run, '一个概念如何塑造现代世界')

# 金色装饰线
add_horizontal_line(doc, ACCENT, 12)

# 副标题（从公民认同到民族主义浪潮）
tagline_para = doc.add_paragraph()
tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline_para.paragraph_format.space_before = Pt(16)
tagline_para.paragraph_format.space_after = Pt(32)
run = tagline_para.add_run('从公民认同到民族主义浪潮')
run.font.size = Pt(13)
run.font.italic = True
run.font.color.rgb = SUBTLE
set_run_cjk(run, '从公民认同到民族主义浪潮')

# 基本信息表格
info_table = doc.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

info_data = [
    ('目标受众', '政策制定者、企业管理者、社会科学爱好者'),
    ('学习时长', '6大模块，约12小时'),
    ('课程形式', '线上录播 + 互动讨论'),
    ('前置要求', '具备基础政治学或社会学概念'),
]

for i, (label, value) in enumerate(info_data):
    label_cell = info_table.cell(i, 0)
    value_cell = info_table.cell(i, 1)

    set_cell_background(label_cell, PRIMARY)
    label_para = label_cell.paragraphs[0]
    label_para.paragraph_format.space_before = Pt(4)
    label_para.paragraph_format.space_after = Pt(4)
    label_run = label_para.add_run(label)
    label_run.font.size = Pt(9)
    label_run.font.bold = True
    label_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_run_cjk(label_run, label)

    value_para = value_cell.paragraphs[0]
    value_para.paragraph_format.space_before = Pt(4)
    value_para.paragraph_format.space_after = Pt(4)
    value_run = value_para.add_run(value)
    value_run.font.size = Pt(9)
    value_run.font.color.rgb = BODY_TEXT
    set_run_cjk(value_run, value)

    # 设置单元格左边距
    label_cell._tc.tcPr
    value_cell._tc.tcPr

# 设置表格列宽
for row in info_table.rows:
    row.cells[0].width = Inches(1.3)
    row.cells[1].width = Inches(4.5)

# 底部装饰条
doc.add_paragraph()
bottom_bar = doc.add_paragraph()
bottom_bar.paragraph_format.space_before = Pt(24)
bottom_bar.paragraph_format.space_after = Pt(0)
pPr2 = bottom_bar._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
hex_bg2 = rgb_to_hex(PRIMARY)
shd2.set(qn('w:val'), 'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'), hex_bg2)
pPr2.append(shd2)
run2 = bottom_bar.add_run('  ')
run2.font.size = Pt(6)

# 分页
doc.add_page_break()

# ─────────────────────────────────────────────
# 课程介绍
# ─────────────────────────────────────────────

def add_section_heading(doc, number, title, subtitle=''):
    """添加带编号的章节标题"""
    # 编号圆圈
    num_para = doc.add_paragraph()
    num_para.paragraph_format.space_before = Pt(16)
    num_para.paragraph_format.space_after = Pt(4)
    num_run = num_para.add_run(number)
    num_run.font.size = Pt(28)
    num_run.font.bold = True
    num_run.font.color.rgb = ACCENT

    # 标题
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(2)
    run = title_para.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    set_run_cjk(run, title)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_before = Pt(0)
        sub_para.paragraph_format.space_after = Pt(8)
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(11)
        sub_run.font.italic = True
        sub_run.font.color.rgb = SUBTLE
        set_run_cjk(sub_run, subtitle)

    # 金色底线
    add_horizontal_line(doc, ACCENT, 8)

def add_body_para(doc, text, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    set_run_cjk(run, text)
    return p

def add_module_card(doc, module_num, title, subtitle, duration, objectives, contents):
    """添加模块卡片（带颜色标题条）"""
    # 模块标题条
    card_table = doc.add_table(rows=1, cols=1)
    card_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = card_table.cell(0, 0)
    set_cell_background(header_cell, PRIMARY)

    header_para = header_cell.paragraphs[0]
    header_para.paragraph_format.space_before = Pt(6)
    header_para.paragraph_format.space_after = Pt(6)

    # 模块编号
    num_run = header_para.add_run(f'Module {module_num}  ')
    num_run.font.size = Pt(11)
    num_run.font.bold = True
    num_run.font.color.rgb = ACCENT

    # 模块标题
    title_run = header_para.add_run(title)
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_run_cjk(title_run, title)

    if subtitle:
        sub_run = header_para.add_run(f'  —  {subtitle}')
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = LIGHT
        set_run_cjk(sub_run, subtitle)

    # 内容区域（用普通段落实现）
    content_para = doc.add_paragraph()
    content_para.paragraph_format.space_before = Pt(2)
    content_para.paragraph_format.space_after = Pt(2)

    # 学习目标
    obj_label_run = content_para.add_run('学习目标：')
    obj_label_run.font.size = Pt(9.5)
    obj_label_run.font.bold = True
    obj_label_run.font.color.rgb = SECONDARY
    set_run_cjk(obj_label_run, '学习目标：')

    obj_run = content_para.add_run(objectives)
    obj_run.font.size = Pt(9.5)
    obj_run.font.color.rgb = BODY_TEXT
    set_run_cjk(obj_run, objectives)

    # 核心内容
    core_para = doc.add_paragraph()
    core_para.paragraph_format.space_before = Pt(4)
    core_para.paragraph_format.space_after = Pt(2)
    core_para.paragraph_format.left_indent = Inches(0.2)

    core_label_run = core_para.add_run('核心内容：')
    core_label_run.font.size = Pt(9.5)
    core_label_run.font.bold = True
    core_label_run.font.color.rgb = SECONDARY
    set_run_cjk(core_label_run, '核心内容：')

    core_run = core_para.add_run(contents)
    core_run.font.size = Pt(9.5)
    core_run.font.color.rgb = BODY_TEXT
    set_run_cjk(core_run, contents)

    # 课时安排
    time_para = doc.add_paragraph()
    time_para.paragraph_format.space_before = Pt(2)
    time_para.paragraph_format.space_after = Pt(10)
    time_para.paragraph_format.left_indent = Inches(0.2)

    time_label_run = time_para.add_run('课时安排：')
    time_label_run.font.size = Pt(9.5)
    time_label_run.font.bold = True
    time_label_run.font.color.rgb = SECONDARY
    set_run_cjk(time_label_run, '课时安排：')

    time_run = time_para.add_run(duration)
    time_run.font.size = Pt(9.5)
    time_run.font.color.rgb = BODY_TEXT
    set_run_cjk(time_run, duration)

# ─────────────────────────────────────────────
# 课程介绍部分
# ─────────────────────────────────────────────

add_section_heading(doc, '00', '课程介绍', 'Course Introduction')

add_body_para(doc, '本课程系统梳理民族主义这一现代世界最具影响力的政治概念，探讨其历史起源、理论框架、当代形态以及对中国特殊道路的深度分析。通过跨学科视角，帮助学员理解民族主义如何从18世纪欧洲的思想火种，演变为塑造当今全球政治格局的核心力量。')

# 解决什么问题
solve_para = doc.add_paragraph()
solve_para.paragraph_format.space_before = Pt(8)
solve_para.paragraph_format.space_after = Pt(4)
solve_label = solve_para.add_run('解决什么问题：')
solve_label.font.size = Pt(10.5)
solve_label.font.bold = True
solve_label.font.color.rgb = PRIMARY
set_run_cjk(solve_label, '解决什么问题：')

solve_items = [
    '民族主义在不同国家的表现为何截然不同？',
    '为什么民族主义情绪会在某些时期急剧升温？',
    '中国的民族主义发展路径与西方有何本质区别？',
    '政策制定者如何与民族主义情绪共存并引导其正向发展？',
]
for item in solve_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    set_run_cjk(run, item)

# 学员收获
gain_para = doc.add_paragraph()
gain_para.paragraph_format.space_before = Pt(10)
gain_para.paragraph_format.space_after = Pt(4)
gain_label = gain_para.add_run('学员收获：')
gain_label.font.size = Pt(10.5)
gain_label.font.bold = True
gain_label.font.color.rgb = PRIMARY
set_run_cjk(gain_label, '学员收获：')

gain_items = [
    '建立对民族主义的系统性认知框架',
    '理解民族主义的历史根源与当代表现',
    '掌握分析民族主义情绪的研究方法',
    '对中国特色民族主义发展路径的深度理解',
    '提升政策制定与跨文化沟通的战略思维',
]
for item in gain_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    set_run_cjk(run, item)

doc.add_page_break()

# ─────────────────────────────────────────────
# 课程大纲（六大模块）
# ─────────────────────────────────────────────

add_section_heading(doc, '01', '课程大纲', 'Curriculum Overview')

# 模块1
add_module_card(
    doc, 1,
    '概念解码——民族主义到底是什么',
    'What is Nationalism?',
    '约2小时',
    '理解民族主义的核心定义，掌握公民民族主义与族裔民族主义的基本框架',
    '民族主义的概念界定与学术争议；公民民族主义 vs 族裔民族主义的区别；民族主义与爱国主义、民粹主义的边界；经典理论梳理：盖尔纳、安德森、霍布斯鲍姆'
)

# 模块2
add_module_card(
    doc, 2,
    '类型辨析——两种民族主义的制度后果',
    'Types & Institutional Consequences',
    '约2小时',
    '掌握民族主义的类型学，理解不同制度选择的深层逻辑',
    '公民民族主义与族裔民族主义的制度表现；民族国家构建的不同路径；联邦制 vs 单制；多元文化主义政策的是与非'
)

# 模块3
add_module_card(
    doc, 3,
    '历史溯源——民族主义如何诞生并塑造现代世界',
    'Historical Origins & Global Impact',
    '约2.5小时',
    '理解民族主义的历史诞生过程及其对现代世界秩序的塑造',
    '法国大革命与民族主义的诞生；民族主义浪潮与帝国解体；一战、二战与民族国家体系的形成；冷战后的民族主义复兴；全球化背景下的民族主义走向'
)

# 模块4
add_module_card(
    doc, 4,
    '当代诊断——今天的民族主义为何卷土重来',
    'Contemporary Resurgence',
    '约2小时',
    '理解当代民族主义复兴的深层原因，掌握分析当代议题的分析框架',
    '经济全球化与分配不均的民族主义反应；文化焦虑与身份政治；社交媒体时代的民族主义传播；欧洲右翼政党崛起；美国优先政策的民族主义根源'
)

# 模块5
add_module_card(
    doc, 5,
    '中国的民族主义——特殊还是普遍？',
    'Chinese Nationalism: Exceptional or Universal?',
    '约2.5小时',
    '深度理解中国民族主义的发展路径与特殊形态，建立跨文化分析视角',
    '中国民族主义的历史阶段划分；革命话语与民族叙事的融合；经济崛起与民族自信心的重建；港台与海外华人的民族认同差异；中国民族主义的独特性与普遍性'
)

# 模块6
add_module_card(
    doc, 6,
    '实践应用——如何与民族主义相处',
    'Living with Nationalism',
    '约1小时',
    '将理论框架转化为实践策略，提升政策制定与跨文化沟通能力',
    '政府如何引导民族主义情绪；企业与品牌的民族主义营销边界；跨文化沟通中的民族主义敏感点；未来趋势：民族主义的走向与应对策略'
)

doc.add_page_break()

# ─────────────────────────────────────────────
# 教学方式
# ─────────────────────────────────────────────

add_section_heading(doc, '02', '教学方式', 'Teaching Methods')

methods = [
    ('理论讲授', '系统梳理民族主义理论框架，深厚学术基础支撑', [
        '视频课程系统学习',
        '关键概念深度解读',
        '理论模型图解分析',
    ]),
    ('案例分析', '多元案例深度剖析，理论联系实际', [
        '欧洲民族国家形成史',
        '当代欧洲极右翼崛起分析',
        '中国民族主义发展路径案例',
    ]),
    ('互动讨论', '开放式讨论，碰撞思想火花', [
        '民族主义情绪的识别与应对',
        '全球化与民族主义的张力',
        '中国模式的讨论与反思',
    ]),
    ('工具演练', '研究工具与方法论实操', [
        '民族主义情绪监测方法',
        '舆情分析工具使用',
        '跨文化分析框架应用',
    ]),
]

# 教学方法表格
method_table = doc.add_table(rows=len(methods), cols=2)
method_table.style = 'Table Grid'
method_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (method_name, method_desc, method_items) in enumerate(methods):
    name_cell = method_table.cell(i, 0)
    desc_cell = method_table.cell(i, 1)

    # 左侧：方法名称
    set_cell_background(name_cell, PRIMARY)
    name_para = name_cell.paragraphs[0]
    name_para.paragraph_format.space_before = Pt(6)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(method_name)
    name_run.font.size = Pt(11)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_run_cjk(name_run, method_name)

    name_desc_para = name_cell.add_paragraph()
    name_desc_para.paragraph_format.space_before = Pt(2)
    name_desc_para.paragraph_format.space_after = Pt(6)
    name_desc_run = name_desc_para.add_run(method_desc)
    name_desc_run.font.size = Pt(8.5)
    name_desc_run.font.color.rgb = LIGHT
    set_run_cjk(name_desc_run, method_desc)

    # 右侧：具体内容
    desc_para = desc_cell.paragraphs[0]
    desc_para.paragraph_format.space_before = Pt(6)
    desc_para.paragraph_format.space_after = Pt(2)
    for j, item in enumerate(method_items):
        if j == 0:
            item_run = desc_para.add_run(f'• {item}')
        else:
            item_run = desc_para.add_run(f'\n• {item}')
        item_run.font.size = Pt(9.5)
        item_run.font.color.rgb = BODY_TEXT
        set_run_cjk(item_run, item)
    desc_cell.paragraphs[0].paragraph_format.space_after = Pt(6)

# 设置列宽
for row in method_table.rows:
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(4.8)

doc.add_page_break()

# ─────────────────────────────────────────────
# 配套材料
# ─────────────────────────────────────────────

add_section_heading(doc, '03', '配套材料', 'Course Materials')

materials = [
    ('学员手册', '课程讲义全文 + 核心概念速查表 + 思考题集'),
    ('讲师手册', '教学指引 + 案例讨论指南 + 常见问题解答'),
    ('练习题库', '章节测验 + 案例分析练习 + 论述题题库'),
    ('评估工具', '学习效果评估表 + 课程反馈问卷 + 能力成长追踪'),
]

for name, desc in materials:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)

    name_run = p.add_run(f'{name}：')
    name_run.font.size = Pt(10.5)
    name_run.font.bold = True
    name_run.font.color.rgb = SECONDARY
    set_run_cjk(name_run, f'{name}：')

    desc_run = p.add_run(desc)
    desc_run.font.size = Pt(10.5)
    desc_run.font.color.rgb = BODY_TEXT
    set_run_cjk(desc_run, desc)

doc.add_paragraph()

# 底部装饰
add_horizontal_line(doc, PRIMARY, 8)

footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.paragraph_format.space_before = Pt(12)
footer_para.paragraph_format.space_after = Pt(0)
footer_run = footer_para.add_run('民族主义思想史 — 课程大纲 v1.0')
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = SUBTLE
set_run_cjk(footer_run, '民族主义思想史 — 课程大纲 v1.0')

# ─────────────────────────────────────────────
# 保存文档
# ─────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f'文档已保存至：{OUTPUT_PATH}')

# -*- coding: utf-8 -*-
"""
Generate Course Specification Word Document
Trainer Holistic Awareness & Presence
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def main():
    doc = Document()

    # Page Setup - A4 Landscape
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Normal style font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # ==================== Cover ====================
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover_table.cell(0, 0)
    set_cell_background(cover_cell, '1F4E79')
    cover_p = cover_cell.paragraphs[0]
    cover_p.paragraph_format.space_before = Pt(8)
    cover_p.paragraph_format.space_after = Pt(8)
    cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover_p.add_run('课程说明书')
    cover_run.font.name = 'Microsoft YaHei'
    cover_run.font.size = Pt(14)
    cover_run.font.bold = True
    cover_run.font.color.rgb = RGBColor(255, 255, 255)
    cover_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    # Course Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(12)
    title_run = title_p.add_run('培训师全局意识与松弛感')
    title_run.font.name = 'Microsoft YaHei'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # English subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(36)
    subtitle_run = subtitle_p.add_run('Trainer Holistic Awareness & Presence')
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    # Info table
    info_table = doc.add_table(rows=4, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = 'Table Grid'

    info_data = [
        ('版本号', 'V1.0', '文档状态', '正式发布'),
        ('适用场景', '企业内训/工作坊/公开课', '总课时', '10章核心内容'),
        ('目标学员', '培训师/催化师/内训师/主持人', '教学方式', '公理-拆解-迁移-工具'),
        ('核心案例', '张菲 · 马东 · 李诞', '最近更新', '2026年8月'),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if j % 2 == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True
                run.font.color.rgb = RGBColor(64, 64, 64)
            else:
                run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()
    doc.add_paragraph()

    # Core value
    value_p = doc.add_paragraph()
    value_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    value_p.paragraph_format.space_before = Pt(24)
    value_run = value_p.add_run('核心价值主张')
    value_run.font.name = 'Microsoft YaHei'
    value_run.font.size = Pt(11)
    value_run.font.bold = True
    value_run.font.color.rgb = RGBColor(31, 78, 121)
    value_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    quote_p = doc.add_paragraph()
    quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_p.paragraph_format.space_before = Pt(8)
    quote_run = quote_p.add_run('松弛感不是性格，是一笔提前算好的账')
    quote_run.font.name = 'Microsoft YaHei'
    quote_run.font.size = Pt(16)
    quote_run.font.bold = True
    quote_run.font.color.rgb = RGBColor(192, 80, 77)
    quote_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    quote_sub_p = doc.add_paragraph()
    quote_sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_sub_run = quote_sub_p.add_run('——你能在台上松到什么程度，取决于你替多少种意外想好了退路')
    quote_sub_run.font.name = 'Microsoft YaHei'
    quote_sub_run.font.size = Pt(10)
    quote_sub_run.font.color.rgb = RGBColor(89, 89, 89)
    quote_sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_page_break()

    # ==================== Section 1: Course Positioning ====================
    h1 = doc.add_heading('一、课程定位', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('为什么开设这门课')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    run = p.add_run('这是一门从"全局意识"与"松弛感"这一个具体切口切入的专题课程。课程聚焦于"如何在台上真正松下来，同时还能掌控全场"这一个核心问题，拆解到可模仿、可练习的具体动作。')
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run('解决的问题：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    problems = [
        '上台紧张、控场能力不足的困扰',
        '过度关注流程而忽视现场信号',
        '不敢停顿、不敢留白的焦虑',
        '学员线与甲方线无法同时兼顾',
        '追求"看起来松弛"的表演而非真正的底气'
    ]
    for prob in problems:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prob)
        run.font.size = Pt(10.5)

    # ==================== Section 2: Target Audience ====================
    h1 = doc.add_heading('二、目标学员', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('本课程面向以下人群：')
    run.font.size = Pt(10.5)

    audience_table = doc.add_table(rows=5, cols=2)
    audience_table.style = 'Table Grid'
    audience_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    audience_data = [
        ('学员类型', '具体描述'),
        ('在职培训师', '有一定场次经验，但上台仍会紧张或觉得"控场能力不够"的从业者'),
        ('催化师/引导师', '经常带工作坊、行动学习项目的专业人士，希望提升对现场的掌控力'),
        ('内训师/HR', '需要经常在公司内部进行分享、汇报、项目主持的职场人士'),
        ('主持人/演讲者', '希望从张菲、马东、李诞三位顶级主持人身上提取可迁移的方法论'),
    ]

    for i, row_data in enumerate(audience_data):
        row = audience_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9.5)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('前提要求：')
    run.font.bold = True
    run.font.size = Pt(10.5)
    run = p.add_run('学员需要有至少5场以上的培训或工作坊主持经验，熟悉基本的培训流程设计，有过在台上紧张、卡壳、或感觉"控不住场"的具体经历。本课程不是"如何克服舞台恐惧"的入门课，而是"如何在有经验的基础上进一步突破"的核心课。')
    run.font.size = Pt(10.5)

    # ==================== Section 3: Course Value ====================
    h1 = doc.add_heading('三、课程价值', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    value_table = doc.add_table(rows=5, cols=3)
    value_table.style = 'Table Grid'
    value_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    value_data = [
        ('价值维度', '核心卖点', '学员收益'),
        ('认知重塑', '拆解"松弛感"和"全局意识"的真正含义，推翻最常见的三个误解', '建立正确的松弛感认知框架，不再迷信天赋论'),
        ('工具落地', '提供可直接使用的：气口设计表、退路清单、现场信号判断表等工具', '带走可立即用于下次培训的实用工具'),
        ('心态建设', '处理"台上的松弛和台下的紧绷"这一核心矛盾，让学员与焦虑和解', '接受"台下的紧绷是换取台上松弛的代价"这一交换关系'),
        ('双线管理', '揭示"学员线"与"甲方线"这两条完全不同的评价体系', '能够同时应对不同评价标准，从容应对甲方检查'),
    ]

    for i, row_data in enumerate(value_data):
        row = value_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    # ==================== Section 4: Learning Objectives ====================
    h1 = doc.add_heading('四、学习目标', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    obj_table = doc.add_table(rows=4, cols=3)
    obj_table.style = 'Table Grid'
    obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    obj_data = [
        ('维度', '目标描述', '具体表现'),
        ('知识', '理解松弛感的本质是"精算"而非天赋；理解全局意识的本质是"放弃执念"而非"看见一切"', '能够清晰阐述松弛感的形成机制，能够判断什么是真正的全局意识'),
        ('技能', '掌握三类退路的设计方法；掌握气口密度自查的现场判断能力；掌握"请人翻译复述"的接话技术', '能够设计完整的退路方案，能够准确判断现场密度并调整，能够运用接话技术推进讨论'),
        ('态度', '接受"台下的紧绷是换取台上松弛的代价"这一交换关系；不再追求"看起来很松弛"的表演', '能够坦然面对培训前的紧张情绪，不再为此感到羞耻或能力不足'),
    ]

    for i, row_data in enumerate(obj_data):
        row = obj_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_page_break()

    # ==================== Section 5: Teaching Method ====================
    h1 = doc.add_heading('五、教学方式', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('核心方法论：公理-拆解-迁移-工具')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    run = p.add_run('本课程的教学设计，遵循"公理-拆解-迁移-工具"四步法，贯穿全书每章：')
    run.font.size = Pt(10.5)

    method_table = doc.add_table(rows=5, cols=3)
    method_table.style = 'Table Grid'
    method_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    method_data = [
        ('步骤', '含义', '学员体验'),
        ('公理', '每章从一个反直觉的核心判断开始，打破学员的既有认知', '认知冲击——"这些道理我听过，但从没这样想过"'),
        ('拆解', '把公理拆成普通人能理解、可操作的具体动作', '理解原理——"原来是这样运作的"'),
        ('迁移', '把其他行业（主持人、脱口秀、直播带货）的成熟方法，迁移到培训场景', '跨界应用——"原来马东的方法可以用在这里"'),
        ('工具', '每章配套提供可直接使用的工具表单，确保学习成果可落地', '带走工具——"这个表我下次培训就能用"'),
    ]

    for i, row_data in enumerate(method_data):
        row = method_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('核心案例贯穿：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    case_table = doc.add_table(rows=4, cols=3)
    case_table.style = 'Table Grid'
    case_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    case_data = [
        ('案例人物', '核心特征', '支撑的论点'),
        ('张菲', '3.5万场大型晚会主持从容，本子一遍遍改、会一遍遍开', '松弛是算出来的；台上的松弛是台下较真换来的'),
        ('马东', '手边小木鱼控制节奏；请人"翻译"复述来接话', '全局意识是提前设计好的工具；接话的本质是递口子'),
        ('李诞', '追求"控制下的失控氛围"；"全自动小卖部"敢于分控场权', '松弛感的顶级形态是敢于让渡控制权'),
    ]

    for i, row_data in enumerate(case_data):
        row = case_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('体验式学习设计：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run('每个章节都配套设计了一个练习活动，让学员在课堂上直接体验、实践，而非只是听讲。核心逻辑是：')
    run.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run('先体验 → 再复盘 → 再提炼原则 → 最后形成工具')
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ==================== Section 6: Course Outline ====================
    h1 = doc.add_heading('六、课程大纲', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    outline_table = doc.add_table(rows=11, cols=3)
    outline_table.style = 'Table Grid'
    outline_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    outline_data = [
        ('章节', '主题', '核心公理'),
        ('第一章', '全局意识不是看得多，是知道哪些可以不看', '全局意识不是看见更多，是提前放弃了必须看见一切的执念'),
        ('第二章', '松弛是算出来的，不是天生的', '你能松到什么程度，取决于你替多少种意外想好了退路'),
        ('第三章', '紧张不是准备不够，是准备错了地方', '紧张不是没准备，是把所有力气都用在了不会出事的地方'),
        ('第四章', '掌控全场，其实就是掌控呼吸的节奏', '流程是骨架，气口才是内容'),
        ('第五章', '沉默不是没词了，是在等一个更准的词', '沉默不是冷场，是你留给自己的容错空间'),
        ('第六章', '越想把这句话说完整，越接不上下一句', '你说的每句话，不是为了说完，是为了让下一句更容易接上'),
        ('第七章', '你面对的从来不是一群人，是两群人', '你以为你在对学员讲课，其实你同时在对甲方演一场戏'),
        ('第八章', '越把流程走全，台下越觉得你在念稿', '权威感来自你敢不按流程走，不是你把流程走全'),
        ('第九章', '台上的松弛，是台下扛住的紧绷换来的', '松弛感的代价，是你要先扛住那些不松弛的时刻'),
        ('第十章', '越急着证明你看见了全局，台下越觉得你没看见', '越想证明自己有全局意识，越会显得心虚'),
    ]

    for i, row_data in enumerate(outline_data):
        row = outline_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_page_break()

    # ==================== Section 7: Supporting Outputs ====================
    h1 = doc.add_heading('七、配套产出', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('工具表单：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    tools_table = doc.add_table(rows=5, cols=3)
    tools_table.style = 'Table Grid'
    tools_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tools_data = [
        ('工具名称', '用途', '使用场景'),
        ('现场信号判断表', '区分噪音与信号，判断哪些变化需要介入', '培训/工作坊进行中，持续观察使用'),
        ('三类退路清单', '系统梳理流程/情绪/权力三个维度的退路方案', '每次培训前20分钟填写'),
        ('气口密度自查表', '判断现场节奏密度，及时调整气场', '培训进行中快速扫描'),
        ('单场复盘表', '从8个维度复盘单场培训，追踪成长轨迹', '每次培训结束后当天填写'),
    ]

    for i, row_data in enumerate(tools_data):
        row = tools_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('案例集：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    cases_list = [
        '张菲大型晚会主持案例（3.5万场从容背后的严苛准备）',
        '马东《奇葩说》控场案例（木鱼节奏控制、翻译复述法）',
        '李诞直播带货案例（全自动小卖部、控场权让渡）',
        '工作坊现场意外处理案例集（沉默/超时/质疑应对）'
    ]
    for case in cases_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(case)
        run.font.size = Pt(10.5)

    # ==================== Section 8: Prerequisites ====================
    h1 = doc.add_heading('八、预备知识', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    p = doc.add_paragraph()
    run = p.add_run('培训师基本功要求：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    prereq_list = [
        '有过至少5场以上的培训或工作坊主持经验',
        '熟悉基本的培训流程设计（开场-内容-结尾）',
        '能区分"培训师"和"讲师"的角色差异（培训师更注重引导和互动）',
        '有过在台上紧张、卡壳、或感觉"控不住场"的具体经历'
    ]
    for prereq in prereq_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prereq)
        run.font.size = Pt(10.5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('课前调研清单（培训师开课前应了解）：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    research_list = [
        '这批学员之前有没有上过类似的"软技能"课程？',
        '他们最常遇到的现场问题是什么？（冷场/太闹/挑战权威/超时）',
        '有没有特别难处理的学员类型？（沉默型/好斗型/消极型）',
        '这次培训的甲方/领导最在意什么指标？',
        '学员的职级分布如何？（高管的场次需要更克制的风格）',
        '有没有历史遗留的复杂人际问题需要特别注意？'
    ]
    for item in research_list:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('常见学员类型应对策略：')
    run.font.bold = True
    run.font.size = Pt(10.5)

    student_table = doc.add_table(rows=6, cols=3)
    student_table.style = 'Table Grid'
    student_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    student_data = [
        ('学员类型', '典型表现', '应对策略'),
        ('沉默型', '全程不发言，或者只在被点名时说话', '降低问题难度；用开放式问题引导；请人复述给开口的台阶'),
        ('好斗型', '喜欢质疑、挑战培训师权威', '不当场对抗；承认其问题有价值；私下处理情绪'),
        ('消极型', '表面配合但内心抵触', '找到他的利益关联点；用"选择题"而非"判断题"'),
        ('话痨型', '一个人说太多，占用其他人时间', '用"翻译复述"法收住话头；用时间提醒温和打断'),
        ('表演型', '喜欢把场合当个人秀场', '给他的表演欲一个合法出口；用记录工具收住'),
    ]

    for i, row_data in enumerate(student_data):
        row = student_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_page_break()

    # ==================== Appendix: Tool Forms Preview ====================
    h1 = doc.add_heading('附录：工具表单预览', level=1)
    h1.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    h1.runs[0].font.size = Pt(14)

    add_horizontal_line(doc)

    # Signal table
    p = doc.add_paragraph()
    run = p.add_run('1. 现场信号判断表')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    signal_table = doc.add_table(rows=7, cols=4)
    signal_table.style = 'Table Grid'
    signal_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    signal_data = [
        ('现场状态', '属于噪音还是信号', '判断依据', '我的应对'),
        ('小组讨论声音变大', '通常是噪音', '内容是否围绕主题', '不介入，远观即可'),
        ('小组连续沉默超出预期时长', '需要判断', '眼神是否在动', '眼神发直则上前询问'),
        ('发言人始终是同一人', '信号', '是否形成一言堂', '引导追问其他人意见'),
        ('发言人频繁更换未形成结论', '信号', '是否无人承担梳理', '指定一人先做小结'),
        ('某个组突然特别安静', '需要判断', '结合上下文判断', '走过去，轻声询问'),
        ('有人频繁看手机', '噪音（正常走神）', '不影响他人情况下', '不做特别处理'),
    ]

    for i, row_data in enumerate(signal_data):
        row = signal_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(8)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    # Retreat table
    p = doc.add_paragraph()
    run = p.add_run('2. 三类退路清单')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    retreat_table = doc.add_table(rows=10, cols=4)
    retreat_table.style = 'Table Grid'
    retreat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    retreat_data = [
        ('退路类型', '最可能出现的场景', '我目前的应对方式', '还没想清楚的地方'),
        ('流程退路', '环节冷场 / 讨论超时', '', ''),
        ('流程退路', '某个案例引发负面情绪', '', ''),
        ('流程退路', '时间不够用需要砍环节', '', ''),
        ('情绪退路', '学员情绪失控 / 崩溃哭泣', '', ''),
        ('情绪退路', '公开质疑你的方法或资历', '', ''),
        ('情绪退路', '消极对抗（表面配合内心抵触）', '', ''),
        ('权力退路', '甲方临场插话打乱节奏', '', ''),
        ('权力退路', '学员职级远超预期', '', ''),
        ('权力退路', '甲方突然提出增加内容', '', ''),
    ]

    for i, row_data in enumerate(retreat_data):
        row = retreat_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(8)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    # Density table
    p = doc.add_paragraph()
    run = p.add_run('3. 气口密度自查表')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    density_table = doc.add_table(rows=6, cols=3)
    density_table.style = 'Table Grid'
    density_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    density_data = [
        ('现场观察到的信号', '说明密度状态', '该做的调整'),
        ('学员表情发直、身体后靠', '密度过高', '插入一次停顿或轻松话题'),
        ('学员开始交头接耳、注意力涣散', '密度过低', '收紧节奏，提高信息密度'),
        ('关键判断说完后现场安静', '气口正常', '保持停顿，不急于填补'),
        ('连续超过十分钟无起伏', '需要立即调整', '插入互动或改变语速'),
        ('有人频繁看手机', '注意力下降', '改变节奏，插入互动'),
    ]

    for i, row_data in enumerate(density_data):
        row = density_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(8)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    doc.add_paragraph()

    # Review table
    p = doc.add_paragraph()
    run = p.add_run('4. 单场复盘表')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 78, 121)

    review_table = doc.add_table(rows=9, cols=3)
    review_table.style = 'Table Grid'
    review_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    review_data = [
        ('复盘维度', '这场活动的具体情况', '下次要调整的地方'),
        ('全局判断', '这场哪里是噪音、哪里是真信号，判断对了吗？', ''),
        ('退路准备', '有没有意外是完全没想到的？', ''),
        ('气口节奏', '哪个环节明显闷了，或者哪个环节太赶？', ''),
        ('沉默应对', '那几个停顿用对了吗？有没有该停没停的？', ''),
        ('接话流畅度', '接话时是在"递口子"还是在"完成自己的问题清单"？', ''),
        ('甲方与学员', '这两条线各自的反馈是什么，有没有冲突？', ''),
        ('台下的紧绷', '这次真正让自己焦虑的点是什么，说清楚它？', ''),
        ('有没有证明欲', '有没有多说了证明自己的话？哪一句？', ''),
    ]

    for i, row_data in enumerate(review_data):
        row = review_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(8)
            if i == 0:
                set_cell_background(cell, '1F4E79')
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif j == 0:
                set_cell_background(cell, 'D6DCE4')
                run.font.bold = True

    # Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('培训师全局意识与松弛感 | 课程说明书 V1.0 | 方太文化研究院')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)

    # Save document
    output_path = r'D:\新课开发\工作手册\培训师全局意识和松弛感\完整课程包\08-成果demo\课程说明书-培训师全局意识与松弛感.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    main()

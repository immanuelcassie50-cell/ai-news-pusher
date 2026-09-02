# -*- coding: utf-8 -*-
"""
生成 顺造科技AI项目成果评审 - 评委手册.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"D:\2026年课程\顺造科技\AI\评审\04-评委与主持\评委手册.docx"

doc = Document()

# ---------- 页面设置：A4 ----------
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ---------- 全局样式 ----------
styles = doc.styles
normal = styles['Normal']
normal.font.name = '宋体'
normal.font.size = Pt(11)
rpr = normal.element.rPr
if rpr is not None:
    rfonts = rpr.find(qn('w:rFonts'))
else:
    rfonts = None
if rfonts is None:
    rpr = OxmlElement('w:rPr')
    normal.element.append(rpr)
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:eastAsia'), '宋体')
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_cn_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_heading_1(text, doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 18, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A68')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading_2(text, doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 14, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
    return p

def add_heading_3(text, doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 12, bold=True, color=RGBColor(0x44, 0x44, 0x44))
    return p

def add_para(text, doc, size=11, indent_first=True, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(22)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_cn_font(run, '宋体', size)
    return p

def add_callout(text, doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF1F8')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        b = OxmlElement('w:'+side)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), '1F3A68')
        pBdr.append(b)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_cn_font(run, '宋体', 11, color=RGBColor(0x1F, 0x3A, 0x68))
    return p

def add_bullet(text, doc, indent=0.74):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    if p.runs:
        p.runs[0].text = ''
    run = p.add_run(text)
    set_cn_font(run, '宋体', 11)
    return p

def add_quote(text, doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '1F3A68')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run('"' + text + '"')
    set_cn_font(run, '楷体', 11, color=RGBColor(0x44, 0x44, 0x44))
    return p

def add_table_header(table, headers):
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                set_cn_font(r, '黑体', 11, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3A68')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def fill_table(table, data, font_size=10):
    for i, row_data in enumerate(data, start=1):
        row = table.rows[i].cells
        for j, val in enumerate(row_data):
            row[j].text = str(val)
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    set_cn_font(r, '宋体', font_size)

# ===== 封面 =====
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(120)
title_p.paragraph_format.space_after = Pt(20)
run = title_p.add_run('顺造科技 · AI 项目成果评审')
set_cn_font(run, '黑体', 22, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(60)
run = sub_p.add_run('评 委 手 册')
set_cn_font(run, '黑体', 36, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))

sep_p = doc.add_paragraph()
sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sep_p.paragraph_format.space_after = Pt(40)
run = sep_p.add_run('— 评委专用 · 内部资料 —')
set_cn_font(run, '楷体', 14, color=RGBColor(0x66, 0x66, 0x66))

meta_lines = [
    '版本：V1.0',
    '编制：AI 项目评审组',
    '评审日期：D-Day 13:30 — D+1 17:30',
]
for line in meta_lines:
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(8)
    run = meta_p.add_run(line)
    set_cn_font(run, '宋体', 12)

doc.add_page_break()

# ===== 目录 =====
toc_title = doc.add_paragraph()
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = toc_title.add_run('目  录')
set_cn_font(run, '黑体', 20, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
toc_title.paragraph_format.space_after = Pt(20)

toc_items = [
    ('第 1 章  评委的职责与定位', '3'),
    ('第 2 章  评审日流程（按时间轴）', '5'),
    ('第 3 章  评分三轨详解', '8'),
    ('第 4 章  评分操作规范', '11'),
    ('第 5 章  点评要点（详见独立文件）', '13'),
    ('第 6 章  AI 介入级别判断参考', '15'),
    ('第 7 章  常见问题与处理', '17'),
    ('附录 A  评委一句话自我介绍模板', '19'),
    ('附录 B  应急联络表', '20'),
]
for name, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
    run = p.add_run(name)
    set_cn_font(run, '宋体', 12)
    p.add_run('\t' + page)

doc.add_page_break()

# ===========================================================
# 第 1 章
# ===========================================================
add_heading_1('第 1 章  评委的职责与定位', doc)

add_para(
    '欢迎各位评委参与顺造科技 AI 项目成果评审。在为期一天半的评审中，您将面对由各业务部门'
    '一线员工亲手改造的 AI 项目——这些项目来自质量、人力、财务、客服、生产、研发等部门，'
    '凝结了小米生态链清洁电器业务一线的真实痛点与一线员工的真实思考。它们可能还很粗糙，'
    '可能还只是提示词级别的小工具，但每一份方案都代表着一个具体的业务场景在 AI 时代的'
    '被重新想象。',
    doc
)

add_para('在进入具体流程前，请先与我们对齐一个根本的认知：评委不是来挑刺的，评委是来帮忙把项目变得更好的。我们希望您扮演三个角色——', doc)

add_heading_2('1.1  评委的三个角色', doc)

add_para('第一，严格但友善的鉴赏者。您需要用专业的眼光识别每个项目的真实价值——是停留在表面的 PPT 包装，还是真正深入到业务流程里的再设计？哪些是真正的创新点，哪些是通用方案的套娃？但严格不等于苛刻，友善不等于放水。请始终记住，对面站着的不是方案，是活生生的人。', doc)

add_para('第二，建设性的建议者。您给出的每一条建议，都应该让学员在第二天就能动手改进。请避免这个不行、要重做式的否定句，转为这个亮点可以这样放大或如果加入 XX 维度，会更完整的建设性表述。', doc)

add_para('第三，生态共建的引路人。评审不是终点，而是顺造 AI 生态的起点。优秀的方案需要被推广、被复制、被组合。请用是否能跨部门复用作为重要的评判维度。', doc)

add_heading_2('1.2  评委组合与分工', doc)

add_para('本次评审共设 7 位评委，组合结构为 3+1+1+1+1，具体如下：', doc)

table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['序号', '角色定位', '核心关注点'])

rows_data = [
    ('1', '公司领导（组长）', '战略价值、组织匹配度、能否在公司层面复制推广'),
    ('2', '公司领导', '业务影响力、与年度战略的契合度'),
    ('3', '公司领导 / 总裁办代表', '对外可展示性、对内激励价值'),
    ('4', '培训经理 / 外部专家', '培训目标达成、学员成长性、内容设计'),
    ('5', '内训师代表', '实操可落地性、是否便于复制到其他部门'),
    ('6', '业务部门负责人', '业务痛点还原度、上线后能否真解决问题'),
    ('7', 'AI 专家（特邀）', 'AI 介入级别、技术合理性、未来扩展性'),
]
fill_table(table, rows_data, 10)

add_heading_2('1.3  评委组组长职责', doc)

add_para('组长由公司领导中资历最深者担任，承担以下五项核心职责：', doc)
add_bullet('开场致辞：代表评委组在评审启动会上做 3 分钟发言，定调帮助改进的评审氛围。', doc)
add_bullet('争议仲裁：若评委在 0-3 分或 9-10 分的极端打分上分歧过大，由组长召集当天的 5 分钟碰头会仲裁。', doc)
add_bullet('标准守护：监督评分表填写规范，发现明显打分失衡（如所有学员分数都集中在 7-8 分）时及时提醒。', doc)
add_bullet('结营总结：在 D+1 下午结营仪式上代表评委组做 10 分钟的总结发言，肯定亮点 + 指出共性短板。', doc)
add_bullet('奖项确认：汇总 7 位评委的评分，签字确认最终获奖名单。', doc)

add_heading_2('1.4  评委的红线与绿线', doc)

add_quote('评委不是来挑刺的，评委是来帮忙把项目变得更好的。这是顺造 AI 评审的底层价值观。', doc)

add_para('红线——绝对不能做的事：', doc)
add_bullet('在亮分前公开讨论分数（影响其他评委的独立判断）。', doc)
add_bullet('在点评中拿自己和学员比较（我以前做过的比你好）。', doc)
add_bullet('给出无法落地的建议（你应该用 RAG + Agent + Multi-modal 重新设计——这种话不是建议，是抬杠）。', doc)
add_bullet('当众否定学员（即使方案很差，也要点出 1-2 个值得肯定的地方）。', doc)

add_para('绿线——强烈推荐做的事：', doc)
add_bullet('在每位学员路演结束后的 1 分钟内完成打分，避免遗忘细节。', doc)
add_bullet('极端分（0-3 / 9-10）必须手写备注理由。', doc)
add_bullet('点评时长严格控制在 1.5-2 分钟以内，超时会被计时提醒。', doc)
add_bullet('发现共性短板时，在当晚评委碰头会提出，便于组长在结营总结中统一指明。', doc)

add_callout('核心原则：先扬后抑，先肯定再建议。永远让学员带着我还能继续做下去的劲头离场。', doc)

# ===========================================================
# 第 2 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 2 章  评审日流程（按时间轴）', doc)

add_para('本次评审分两天进行：D-Day 下午 + D+1 上午为正式评审，D+1 下午为结营仪式。共评审 12-15 位学员，每位学员 16-18 分钟，全天评审时间约 4.5 小时。评委需要全程在场，并在评审结束后立即投入评分碰头。', doc)

add_heading_2('2.1  D-Day（评审第一天 · 下午）', doc)

table = doc.add_table(rows=10, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['时间', '环节', '评委动作 / 备注'])

day1 = [
    ('13:00-13:30', '场地布置 + 设备调试', '评委可选择性提前到场查看设备'),
    ('13:30-13:50', '评委签到 + 评分表发放', '评委领取个人评分夹、笔、议程单；签到墙合影'),
    ('13:50-14:00', '评委碰头会（10 分钟）', '组长重申评分三轨 + 极端分规则 + 点评时长；AI 专家科普介入级别 L1-L5'),
    ('14:00-14:15', '评审启动会（全员）', '组长致辞 3 分钟；主持人讲解流程；全体合影'),
    ('14:15-15:30', '评审 1-3 号学员', '每位学员 16-18 分钟，评委独立打分'),
    ('15:30-15:40', '茶歇 1', '评委可短暂讨论整体观察，但不可讨论具体分数'),
    ('15:40-17:00', '评审 4-7 号学员', '继续评审，评委保持专注'),
    ('17:00-17:10', '茶歇 2', '同上'),
    ('17:10-18:00', '评审 8-10 号学员（视人数）', '最后阶段，组长观察评委疲劳度，必要时延长茶歇'),
]
fill_table(table, day1, 10)

add_heading_2('2.2  D+1（评审第二天 · 上午 + 下午）', doc)

table = doc.add_table(rows=16, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['时间', '环节', '评委动作 / 备注'])

day2 = [
    ('08:30-08:50', '评委签到', '领取昨日未评分表 + 今日评分表'),
    ('08:50-09:00', '评委碰头会', '反馈昨日观察，组长统一今日关注点'),
    ('09:00-10:30', '评审 11-14 号学员', '上午评审 + 茶歇'),
    ('10:30-10:40', '茶歇', '保持精力'),
    ('10:40-12:00', '评审 15-17 号学员（视人数）', '上午最后评审'),
    ('12:00-13:30', '午餐 + 午休', '评委与学员可自由交流，氛围轻松'),
    ('13:30-13:40', '结营仪式 · 项目总结（10 分钟）', '项目负责人做整体回顾'),
    ('13:40-13:45', '评审分数公布（5 分钟）', '由主持人 + 组长公布 6 个奖项归属'),
    ('13:45-14:05', '颁奖（20 分钟）', '6 个奖项依次颁奖，每项约 3 分钟'),
    ('14:05-14:15', '学员代表发言（10 分钟）', '优秀学员代表 + 新人代表各 1 位'),
    ('14:15-14:25', '结业证书（10 分钟）', '全员颁发结业证书'),
    ('14:25-14:35', '评委代表总结（10 分钟）', '组长做评委组总结，肯定 + 共性建议'),
    ('14:35-14:45', '成果交付（10 分钟）', '优秀方案汇编、HTML 模板库、智能体清单的交付'),
    ('14:45-14:55', '领导致辞（10 分钟）', '公司领导做收官发言'),
    ('14:55-15:30', '合影 + 散场', '评委留步参与合影'),
]
fill_table(table, day2, 10)

add_heading_2('2.3  关键节点提醒', doc)

add_callout('13:50 评委碰头会 = 评分宪法时刻：组长必须在此 10 分钟内统一三轨标准、亮分规则、点评时长。这是后续一切顺畅的基础。', doc)

add_callout('18:00 评委碰头评分 = 争议消化时刻：若某位学员 7 位评委打分极差过大（如极差 > 30 分），由组长主持 5 分钟微调讨论，但仅限争议个案。', doc)

add_callout('D+1 14:25 评委代表总结 = 评委高光时刻：组长 10 分钟发言将决定整个项目在公司的后续走向——是一次性活动还是 AI 战略起点。请组长务必提前准备。', doc)

# ===========================================================
# 第 3 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 3 章  评分三轨详解', doc)

add_para('本评审采用三轨制评分体系，将学员的综合表现拆解为三个独立维度。三个轨道在最终总分中按权重合并，避免一项极强掩盖其他短板。', doc)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['评分轨道', '权重', '考察重点', '核心问题'])

tracks = [
    ('轨道 A：AI 方案价值', '40%', '业务痛点还原度、AI 介入合理性、可推广性', '这个方案真的能解决问题吗？'),
    ('轨道 B：AI 产出物质量', '35%', '提示词手册、HTML 页面、智能体、文档质量', '做出来的东西能不能用？'),
    ('轨道 C：路演表现', '25%', '演讲结构、时间控制、Q&A 应答、台风', '这个人能不能把项目讲清楚？'),
]
fill_table(table, tracks, 10)

add_para('每轨 5 个子项 × 10 分 = 单轨满分 50 分。三轨加权后最终得分 = A×0.4 + B×0.35 + C×0.25。', doc)

# ---------- 3.1 ----------
add_heading_2('3.1  轨道 A：AI 方案价值（权重 40%）', doc)
add_para('AI 方案价值是最核心的轨道，考察的是这件事值不值得做和做得好不好。', doc)

add_heading_3('子项 A1：业务痛点还原度（10 分）', doc)
add_para('考察学员是否真的理解自己部门的痛点。', doc)
add_bullet('9-10 分：清晰说出痛点的量（每月多花 40 小时）和质（出错率 15%）。', doc)
add_bullet('7-8 分：清楚描述痛点但缺乏数据。', doc)
add_bullet('5-6 分：知道有痛点但描述模糊。', doc)
add_bullet('0-4 分：痛点是臆造的或套模板的。', doc)

add_heading_3('子项 A2：AI 介入合理性（10 分）', doc)
add_para('考察 AI 是否是最合适的解法，而不是为了 AI 而 AI。', doc)
add_bullet('9-10 分：能说清为什么是 AI 不是流程改造 / RPA / 加人。', doc)
add_bullet('7-8 分：能合理推演 AI 价值但没明确比较其他方案。', doc)
add_bullet('0-4 分：明明普通脚本能解决，非要用大模型（杀鸡用牛刀）。', doc)

add_heading_3('子项 A3：业务流程再设计（10 分）', doc)
add_para('考察学员是否真的动了流程，还是只在动作上加了 AI。', doc)
add_bullet('9-10 分：方案中能看出至少 2 个流程节点被重新设计。', doc)
add_bullet('7-8 分：1 个流程节点有改造。', doc)
add_bullet('0-4 分：原流程不动，只是加了个 AI 助手。', doc)

add_heading_3('子项 A4：可推广性 / 可复用性（10 分）', doc)
add_para('考察方案是否能跨部门、跨场景复制。', doc)
add_bullet('9-10 分：方案有清晰的复用接口（如提示词模板、API、HTML 嵌入方式）。', doc)
add_bullet('7-8 分：有复用潜力但接口不清晰。', doc)
add_bullet('0-4 分：完全是本部门定制，无法推广。', doc)

add_heading_3('子项 A5：上线计划与风险（10 分）', doc)
add_para('考察学员是否考虑过做出来之后怎么办。', doc)
add_bullet('9-10 分：有明确的上线时间表、试点部门、风险预案。', doc)
add_bullet('7-8 分：有计划但风险考虑不全。', doc)
add_bullet('0-4 分：完全没有后续计划。', doc)

# ---------- 3.2 ----------
add_heading_2('3.2  轨道 B：AI 产出物质量（权重 35%）', doc)
add_para('AI 产出物是看得见摸得着的部分，是评审的硬通货。', doc)

add_heading_3('子项 B1：提示词手册质量（10 分）', doc)
add_bullet('9-10 分：结构化、有版本号、有反例、有迭代记录。', doc)
add_bullet('7-8 分：有结构但不够完整。', doc)
add_bullet('0-4 分：就是几段对话截图。', doc)

add_heading_3('子项 B2：HTML 页面 / 工具页面（10 分）', doc)
add_bullet('9-10 分：可直接打开使用，UI 友好，响应快。', doc)
add_bullet('7-8 分：能跑通但 UI 粗糙。', doc)
add_bullet('0-4 分：打不开或乱码。', doc)

add_heading_3('子项 B3：智能体配置（10 分）', doc)
add_bullet('9-10 分：在 Dify / 飞书 Aily / Coze 上跑通，有完整工作流。', doc)
add_bullet('7-8 分：能演示但工作流不完整。', doc)
add_bullet('0-4 分：只是 PPT 截图。', doc)

add_heading_3('子项 B4：数据 / 案例 / 效果验证（10 分）', doc)
add_bullet('9-10 分：有真实数据对比（前后效率、出错率）。', doc)
add_bullet('7-8 分：有定性描述但缺数据。', doc)
add_bullet('0-4 分：完全没有效果验证。', doc)

add_heading_3('子项 B5：文档 / 知识沉淀完整性（10 分）', doc)
add_bullet('9-10 分：方案说明文档 + 使用手册 + FAQ 完整。', doc)
add_bullet('7-8 分：有主要文档但缺 FAQ。', doc)
add_bullet('0-4 分：没有任何文档。', doc)

# ---------- 3.3 ----------
add_heading_2('3.3  轨道 C：路演表现（权重 25%）', doc)
add_para('路演能力是放大器——同样的方案，演讲好的人能拿到 90 分，演讲差的人只能拿 60 分。', doc)

add_heading_3('子项 C1：演讲结构清晰度（10 分）', doc)
add_bullet('9-10 分：开头—痛点—方案—效果—展望，逻辑严密。', doc)
add_bullet('7-8 分：有结构但偶尔跑题。', doc)
add_bullet('0-4 分：流水账或 PPT 朗读。', doc)

add_heading_3('子项 C2：时间控制（10 分）', doc)
add_bullet('9-10 分：严格 10 分钟 ±30 秒。', doc)
add_bullet('7-8 分：8-12 分钟之间。', doc)
add_bullet('0-4 分：超时严重（>15 分钟或 <5 分钟）。', doc)

add_heading_3('子项 C3：Q&A 应答（10 分）', doc)
add_bullet('9-10 分：答不上来时坦诚承认 + 给出后续学习计划。', doc)
add_bullet('7-8 分：能回答主要问题，个别卡壳。', doc)
add_bullet('0-4 分：答非所问或情绪化反驳。', doc)

add_heading_3('子项 C4：台风与表达（10 分）', doc)
add_bullet('9-10 分：脱稿、自信、有眼神交流。', doc)
add_bullet('7-8 分：基本脱稿，偶尔看 PPT。', doc)
add_bullet('0-4 分：全程念稿或声音太小。', doc)

add_heading_3('子项 C5：AI 思维体现（10 分）', doc)
add_para('这是本项目的特色轨——是否在演讲中自然流露出 AI 时代的新思维方式。', doc)
add_bullet('9-10 分：演讲中能自然讲出我让 AI 做了 X，所以我能专注做 Y——人机协作的清晰边界。', doc)
add_bullet('7-8 分：能体现 AI 思维但表达生硬。', doc)
add_bullet('0-4 分：完全把 AI 当黑盒在用。', doc)

# ===========================================================
# 第 4 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 4 章  评分操作规范', doc)

add_para('评分是评审中最敏感的环节。同一个学员，评委之间的分数如果出现 20 分以上的极差，不仅影响评奖公正性，更会影响学员对评审的信任。本章统一所有评委的打分动作。', doc)

add_heading_2('4.1  打分时机', doc)
add_para('请严格遵守以下打分时机：', doc)
add_bullet('路演结束 + Q&A 结束 = 立即开始打分。', doc)
add_bullet('听完点评后 = 不再修改分数。', doc)
add_para('原因：点评是评委之间的二次认知过程，听完彼此点评后容易受锚定效应影响。个人独立判断的价值在于独立两个字。', doc)

add_callout('黄金一分钟：路演结束后 1 分钟内完成打分。此时学员表现最鲜活，遗忘曲线尚未启动。', doc)

add_heading_2('4.2  亮分规则', doc)
add_para('本评审采用评委同步亮分制，规则如下：', doc)
add_bullet('主持人喊请评委亮分后，1 分钟提示。', doc)
add_bullet('提示后 30 秒内必须亮分。', doc)
add_bullet('亮分方式：举评分牌 / 亮手机评分页 / 翻评分夹——任选其一，提前统一。', doc)
add_bullet('亮分后由主持人唱分（按评委序号从 1 到 7）。', doc)
add_bullet('7 位评委的分数去掉一个最高分、去掉一个最低分，剩余 5 位取算术平均 = 该学员最终得分。', doc)

add_para('极端分处理：', doc)
add_bullet('某位评委打分 ≤3 分或 ≥9 分时，该评委需在评分表备注栏手写理由。', doc)
add_bullet('当晚 18:00 碰头会上，由组长组织快速过一遍所有极端分，确认无恶意打分。', doc)
add_bullet('若发现评委持续偏激（如连续 3 位学员都打 9-10 分），由组长单独沟通。', doc)

add_heading_2('4.3  独立打分原则', doc)
add_para('绝对禁止以下行为：', doc)
add_bullet('评委之间公开讨论分数。', doc)
add_bullet('评委在学员离场前透露分数。', doc)
add_bullet('评委在茶歇时议论刚才那位学员真差——这会污染后续打分。', doc)
add_bullet('组长在碰头会上对标——即不允许说上一位打了 7 分，这一位也应该 7 分左右。', doc)

add_para('允许的讨论：', doc)
add_bullet('讨论方案本身（这个思路挺有意思）。', doc)
add_bullet('讨论评分标准的理解差异（我觉得 B2 这个 10 分标准可能偏严了）。', doc)
add_bullet('讨论共性短板（今天 5 个学员的 HTML 普遍弱，是不是要统一提醒一下）。', doc)

add_heading_2('4.4  评分表填写规范', doc)
add_para('每位评委持有独立的评分夹，内含：', doc)
add_bullet('学员名单 + 顺序表（1 张）。', doc)
add_bullet('每人一页的详细评分表（12-17 张，每位学员 1 张）。', doc)
add_bullet('空白便签（用于写极端分理由或共性观察）。', doc)
add_bullet('一支红笔 + 一支蓝笔（红笔打分数，蓝笔写备注）。', doc)

add_para('评分表填写示例（每位学员 1 张，A 轨道）：', doc)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['轨道 / 子项', '分值', '打分', '备注'])

subitems = [
    ('A1 业务痛点还原度', '/ 10', '', ''),
    ('A2 AI 介入合理性', '/ 10', '', ''),
    ('A3 业务流程再设计', '/ 10', '', ''),
    ('A4 可推广性', '/ 10', '', ''),
    ('A5 上线计划与风险', '/ 10', '', ''),
    ('A 轨道小计', '/ 50', '', ''),
]
fill_table(table, subitems, 10)

# ===========================================================
# 第 5 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 5 章  点评要点（详见独立文件）', doc)

add_para('点评是评委与学员面对面交流的珍贵时刻。1.5-2 分钟的点评，决定了学员接下来 3 个月的改进方向。本章给出点评的核心原则，详细话术见配套的《点评要点》文件。', doc)

add_heading_2('5.1  点评的 5 句话结构', doc)
add_para('每位学员的点评严格遵循 5 句话结构：', doc)

add_para('第 1 句 · 总结（一句话说清这个学员做了什么）', doc)
add_quote('刚才 XXX 同学带来的是质量部门 8D 报告的 AI 自动审核方案，针对的是当前 8D 报告撰写耗时 4 小时、出错率 15% 的痛点。', doc)

add_para('第 2 句 · 肯定 1（最强的一点）', doc)
add_quote('我印象最深的是你把 8D 报告的结构做了要素拆解 → 知识图谱 → 智能体工作流的三段式设计，这个思路非常扎实。', doc)

add_para('第 3 句 · 肯定 2（最有价值的一点）', doc)
add_quote('更重要的是，你把跨部门的对接节点也考虑进来了，这个边界感在所有学员里是非常少见的。', doc)

add_para('第 4 句 · 核心建议（1 条具体的改进点）', doc)
add_quote('如果下个迭代能加入跨部门推广场景——比如售后部门如何复用你的方案——整个项目的价值会再上一个台阶。', doc)

add_para('第 5 句 · 鼓励', doc)
add_quote('我相信你继续做下去，质量部门的 AI 化会从你这里开始。', doc)

add_heading_2('5.2  点评时长与节奏', doc)
add_para('严格控制 1.5-2 分钟：', doc)
add_bullet('少于 1 分钟：草率、不尊重学员。', doc)
add_bullet('2 分钟左右：理想状态。', doc)
add_bullet('超过 2.5 分钟：主持人会举剩余 30 秒牌。', doc)
add_bullet('超过 3 分钟：主持人会直接打断（评委不打断学员，但主持人打断评委是允许的）。', doc)

add_heading_2('5.3  差异化点评原则', doc)
add_para('对不同启动基线的学员，点评侧重不同：', doc)

add_quote('对高启动基线学员（已有 AI 项目上线经验 / 技术背景强 / 之前参与过 AI 培训）：直接点出问题，给出小步快跑建议。例如：你的方案已经非常成熟，下一步的关键是从工具到生态——建议把提示词模板沉淀为部门级 SOP。', doc)

add_quote('对低启动基线学员（首次接触 AI / 业务一线 / 原本对写代码有畏惧）：先肯定突破，再给如何扩大建议。例如：作为第一次接触 AI 的同事，你能在两周内做出能跑通的 HTML 工具，这本身就是突破。下一步建议你先在班组内找 3 个同事试用，收集真实反馈，再决定是否推广。', doc)

add_callout('点评禁忌：不要直接说这个不行 / 不要拿自己和学员比较 / 不要给出无法落地的建议 / 不要抢学员的功劳。', doc)

# ===========================================================
# 第 6 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 6 章  AI 介入级别判断参考', doc)

add_para('本次评审的项目在 AI 介入深度上有天壤之别——有的学员只是写了一个提示词，有的学员做出了能在飞书 Aily 上跑的智能体。为了公平评审，我们采用 AI 介入级别 L1-L5 作为评委的快速判断框架。', doc)

add_heading_2('6.1  L1-L5 级别定义', doc)

table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['级别', '名称', '典型形态', '评委快速识别'])

levels = [
    ('L1', '一次性提示词', '一段对话截图、单个 prompt', '看手册里有没有使用场景章节'),
    ('L2', '团队共用提示词', '提示词模板、版本号、迭代记录', '看手册里有没有团队使用指南章节'),
    ('L3', '多轮结构化对话', '提示词链、CoT、分步 Prompt', '看手册里有没有多轮示例章节'),
    ('L4', '工具集成', 'Dify 工作流、飞书 Aily、Coze 智能体', '看 HTML 页面能不能直接打开 / 智能体能否跑通'),
    ('L5', '智能体自动化', '自动触发、自动反馈、人机协作闭环', '看数据 / 案例 / 效果验证章节是否有真实运行数据'),
]
fill_table(table, levels, 10)

add_heading_2('6.2  评委快速判断三步法', doc)
add_para('Step 1：看手册结构。1 分钟翻完手册的目录和章节，判断 L1-L5 级别。', doc)
add_para('Step 2：看 HTML 页面。如果有 HTML，直接打开看是否可交互。', doc)
add_para('Step 3：看智能体 / 数据。如果学员说我做了智能体，请他当场演示。', doc)

add_para('常见误判提醒：', doc)
add_bullet('把 PPT 做得漂亮误判为 L4——PPT 是包装，不是 AI 介入。', doc)
add_bullet('把用了 ChatGPT 误判为 L2——这可能只是 L1。', doc)
add_bullet('把流程图误判为智能体——流程图是设计，智能体是能跑的工具。', doc)

add_heading_2('6.3  按级别的点评角度', doc)
add_para('不同级别的项目，点评侧重不同：', doc)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['级别', '肯定点', '建议点'])

advices = [
    ('L1', '使用意识已经觉醒', '团队化——把个人提示词变成团队模板'),
    ('L2', '团队化能力已具备', '流程化——把模板嵌进业务流程'),
    ('L3', '结构化能力强', '工具化——把对话链变成可复用工具'),
    ('L4', '工具集成能力强', '全流程闭环——从触发到反馈的自动化'),
    ('L5', '自动化能力强', '持续迭代机制——版本管理 + 效果监控'),
]
fill_table(table, advices, 10)

# ===========================================================
# 第 7 章
# ===========================================================
doc.add_page_break()
add_heading_1('第 7 章  常见问题与处理', doc)

add_heading_2('7.1  学员侧常见问题', doc)

add_heading_3('问题 1：学员路演超时（>12 分钟）', doc)
add_para('处理流程：', doc)
add_bullet('11 分钟：主持人举黄牌剩余 1 分钟。', doc)
add_bullet('12 分钟：主持人举红牌并提示请用 30 秒收尾。', doc)
add_bullet('13 分钟仍未收尾：主持人直接打断，感谢 XXX 的分享，请评委准备点评。', doc)
add_bullet('评委打分时：路演超时可在路演表现 C2 时间控制项酌情扣分，但不直接影响 A、B 轨。', doc)

add_heading_3('问题 2：学员过度紧张 / 卡壳', doc)
add_para('处理流程：', doc)
add_bullet('主持人主动引导：不着急，我们喝口水再来。', doc)
add_bullet('评委可发起破冰提问——提一个简单问题让学员回答，帮他进入状态。', doc)
add_bullet('如学员实在无法继续：主持人接话我们看到 XXX 准备得非常用心，下面请评委基于现有材料做点评，跳过 Q&A。', doc)
add_bullet('点评中应给予更多肯定——紧张本身不应被扣分。', doc)

add_heading_3('问题 3：学员方案被评委现场质疑，氛围紧张', doc)
add_para('处理流程：', doc)
add_bullet('由组长或主持人主动缓和：这个点我们记下来，会后单独沟通。', doc)
add_bullet('如学员情绪激动：主持人请学员先回座休息，评委继续完成打分。', doc)
add_bullet('绝对禁止评委与学员当面对峙。', doc)

add_heading_2('7.2  技术侧常见问题', doc)

add_heading_3('问题 4：设备故障（投影 / 电脑 / 网络）', doc)
add_para('处理流程：', doc)
add_bullet('5 分钟内可恢复：现场调试，评委可先短暂休场。', doc)
add_bullet('5-15 分钟：启动 B 计划（用备用电脑 / 手机投屏）。', doc)
add_bullet('大于 15 分钟：评委按学员手册 + 提问完成评审，路演分 C1/C4 酌情处理（设备原因不扣学员分）。', doc)
add_bullet('详细预案见《应急预案》（由其他 agent 负责）。', doc)

add_heading_3('问题 5：HTML 页面打不开 / 智能体演示失败', doc)
add_para('处理流程：', doc)
add_bullet('鼓励学员用截图 + 口述完成介绍。', doc)
add_bullet('评委可在产出物质量轨酌情扣分（B2 / B3 项减 1-2 分），但不应一票否决。', doc)

add_heading_2('7.3  评委侧常见问题', doc)

add_heading_3('问题 6：评委意见分歧巨大', doc)
add_para('处理流程：', doc)
add_bullet('当场不公开讨论分数（避免锚定）。', doc)
add_bullet('当晚 18:00 碰头会上由组长主持 5 分钟微调讨论，仅限极差 > 30 分的争议个案。', doc)
add_bullet('若仍无法统一：尊重原始打分，最终得分按去最高去最低取平均计算。', doc)

add_heading_3('问题 7：评委连续给高分 / 低分', doc)
add_para('处理流程：', doc)
add_bullet('由组长在茶歇时私下沟通，了解是评分标准理解差异还是其他原因。', doc)
add_bullet('如属慈悲分：提醒评委评分是为了让优秀者更突出，不是为了不伤人。', doc)
add_bullet('如属敌意分：立即制止，严重时取消该评委资格。', doc)

add_heading_3('问题 8：评委临时缺席', doc)
add_para('处理流程：', doc)
add_bullet('D-Day 当天评委组至少保证 5 人到场。', doc)
add_bullet('如仅 5 人到场：去掉最高最低取 3 人平均。', doc)
add_bullet('如少于 5 人：评审延期。', doc)
add_bullet('评委组组长原则上不可缺席。', doc)

# ===========================================================
# 附录
# ===========================================================
doc.add_page_break()
add_heading_1('附录 A  评委一句话自我介绍模板', doc)

add_para('评审启动会上，每位评委 30 秒自我介绍。建议用以下格式：', doc)

add_quote('我是 [姓名]，[职务]，[与 AI / 业务的关联]。今天的评审我会重点关注 [1 个维度]。', doc)

add_para('示例：', doc)
add_quote('我是王总，集团副总裁，主管研发与供应链。今天的评审我会重点关注方案的可落地性——再好的方案，做不出来就是 PPT。', doc)
add_quote('我是陈老师，外部 AI 培训专家。今天的评审我会重点关注 AI 介入的合理性，看大家是不是把简单问题复杂化了。', doc)
add_quote('我是刘经理，质量部门负责人。今天的评审我会重点关注业务痛点还原度，看看大家是不是真的懂业务。', doc)

add_heading_1('附录 B  应急联络表', doc)

table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
add_table_header(table, ['角色', '姓名 / 联系电话', '职责'])

contacts = [
    ('评委组组长', '[姓名] / [电话]', '争议仲裁 / 结营总结'),
    ('项目负责人', '[姓名] / [电话]', '流程协调 / 应急决策'),
    ('技术支持', '[姓名] / [电话]', '设备故障 / 投屏调试'),
    ('主持人', '[姓名] / [电话]', '流程推进 / 学员引导'),
    ('场地协调', '[姓名] / [电话]', '茶歇 / 用餐 / 物料'),
    ('医疗应急', '[姓名] / [电话]', '学员 / 评委身体不适'),
    ('摄影摄像', '[姓名] / [电话]', '记录 / 直播 / 合影'),
]
fill_table(table, contacts, 10)

# ===== 封底 =====
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(200)
run = p.add_run('— 评委手册结束 —')
set_cn_font(run, '楷体', 14, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('祝评审顺利，祝学员成长！')
set_cn_font(run, '楷体', 12, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('顺造科技 · AI 项目评审组')
set_cn_font(run, '楷体', 12, color=RGBColor(0x66, 0x66, 0x66))

# ===== 保存 =====
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print('OK:', OUT_PATH)
print('段落数:', len(doc.paragraphs))

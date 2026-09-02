# -*- coding: utf-8 -*-
"""
创建《经营者讲党课》工具集锦模板
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

OUTPUT_DIR = "D:/新课开发/党业融合/经营者讲党课/完整课程包/008-工具集锦"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False):
    """设置run的字体"""
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading_with_style(doc, text, level=1):
    """添加带样式的标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 1:
        set_run_font(run, 'Microsoft YaHei', 16, bold=True)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, 'Microsoft YaHei', 14, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(8)
    return para

def create_transformation_card():
    """创建《党课素材转化卡》"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("党课素材转化卡")
    set_run_font(run, 'Microsoft YaHei', 22, bold=True)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——从原始故事到讲稿框架的转化记录表")
    set_run_font(run, 'Microsoft YaHei', 12, bold=False)

    doc.add_paragraph()

    # 填写说明
    instruction = doc.add_paragraph()
    run = instruction.add_run("【填写指引】")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)
    run = instruction.add_run("本表用于记录将原始管理故事转化为党课讲稿的全过程。请按照案例转化四步法依次填写，每一步都是在前一步基础上的深化和优化。")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== 第一步：原始故事信息 =====
    add_heading_with_style(doc, "第一步：原始故事信息", 1)

    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'

    headers1 = ["何时（时间/背景）", "何地（场景）", "何人（角色）", "何事（核心事件）"]
    for i, header in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10, bold=True)

    for i in range(4):
        cell = table1.rows[1].cells[i]
        cell.text = ""
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== 第二步：故事盘点 =====
    add_heading_with_style(doc, "第二步：故事盖点——挖掘剧制张力", 1)

    table2 = doc.add_table(rows=4, cols=2)
    table2.style = 'Table Grid'

    items2 = [
        ("冲突点", "这件事最大的矛盾和困难是什么？有哪些对立的观点或利益？"),
        ("抉择时刻", "当时面临哪几种选择？为什么难以抉择？"),
        ("最担绊的瞬间", "哪一刻让你真正睡不着觉？"),
        ("突破点", "最后是什么让事情有了转机？")
    ]

    for i, (label, hint) in enumerate(items2):
        cell0 = table2.rows[i].cells[0]
        cell0.text = label
        set_cell_shading(cell0, 'E2EFD9')
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10, bold=True)

        cell1 = table2.rows[i].cells[1]
        cell1.text = hint
        for para in cell1.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== 第三步：主题锚定 =====
    add_heading_with_style(doc, "第三步：主题错定——找到故事与党课主题的连接", 1)

    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'

    items3 = [
        ("本次党课主题", "今天要讲的核心主题是什么？（来自组织部门要求或指定主题）"),
        ("故事给我的听发", "这个故事让你自己悟道了什么道理？"),
        ("与主题的连接", "这个道理和今天要讲的主题有什么真实的呱应？（如果连接生确，宁可换故事）")
    ]

    for i, (label, hint) in enumerate(items3):
        cell0 = table3.rows[i].cells[0]
        cell0.text = label
        set_cell_shading(cell0, 'FFF2CC')
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10, bold=True)

        cell1 = table3.rows[i].cells[1]
        cell1.text = hint
        for para in cell1.paragraphs:
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== 第四步：结构搭建 =====
    add_heading_with_style(doc, "第四步：结构搭建——五段式序事框架", 1)

    table4 = doc.add_table(rows=6, cols=2)
    table4.style = 'Table Grid'

    items4 = [
        ("开场悬忹\n（30秒内抓住注意力）", "用一个画面、一个问题或一个困境开场，让听众立刻进入情境"),
        ("抉择还原\n（让听众代入）", "展示当时面临的几条路，引发\"如果是我会怎么选\"的思考"),
        ("转折揪示\n（真实选择）", "揪示你当时的真实选择，以及过程中的关键转折"),
        ("感悟自然生长\n（不替听众总结）", "通过故事的结局，让道理自然浮现，不直接说教"),
        ("回扭主题\n（简短有力）", "用一句话回应开头，最快速度收尾，不要拖泥带水"),
        ("整体时长", "预计讲完需要多少分钟？（建议3-5分钟精简版）")
    ]

    for i, (label, hint) in enumerate(items4):
        cell0 = table4.rows[i].cells[0]
        cell0.text = label
        set_cell_shading(cell0, 'FCE4D6')
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10, bold=True)

        cell1 = table4.rows[i].cells[1]
        cell1.text = hint
        for para in cell1.paragraphs:
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== 第五步：语言转译 =====
    add_heading_with_style(doc, "第五步：语言转译——让文字有画面感", 1)

    table5 = doc.add_table(rows=3, cols=2)
    table5.style = 'Table Grid'

    items5 = [
        ("画面感细节", "加入具体的场景描写：哪个深夜？谁说了哪句话？当时的氛围？"),
        ("情绪细节", "不只是讲事情的发展，而是表达当时的内心感受"),
        ("口语化调整", "把平时汇报用的数据、术语，转化成适合讲故事的语言")
    ]

    for i, (label, hint) in enumerate(items5):
        cell0 = table5.rows[i].cells[0]
        cell0.text = label
        set_cell_shading(cell0, 'DDEBF7')
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10, bold=True)

        cell1 = table5.rows[i].cells[1]
        cell1.text = hint
        for para in cell1.paragraphs:
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # 底部签名区
    footer_table = doc.add_table(rows=1, cols=3)
    footer_table.style = 'Table Grid'
    footer_items = ["填写人：", "日期：", "所属支部/部门："]
    for i, text in enumerate(footer_items):
        cell = footer_table.rows[0].cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, 'Microsoft YaHei', 10)

    # 保存
    output_path = os.path.join(OUTPUT_DIR, "党课素材转化卡.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")
    return output_path

if __name__ == "__main__":
    create_transformation_card()
    print("Done!")

# -*- coding: utf-8 -*-
"""
生成 demo07-因果归因分析示例.docx
demo08-竞争性假说复盘表示例.docx
demo09-调研方案自查表示例.docx
demo10-判断标准梳理表示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_demo07():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("因果归因分析示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 区分相关性与因果性，找到真正可控的因")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：林因果", "团队：数据运营组", "日期：2026-08-03", "教练：王分析"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心原则】因果归因三问：")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    questions = [
        ("第一问：这是相关性还是因果性？", "A和B同时发生，不一定A导致B"),
        ("第二问：因果方向是什么？", "是A导致B，还是B导致A，或者有共同原因C？"),
        ("第三问：找到可控因了吗？", "真正能改变结果的杠杆点在哪里？")
    ]

    for q, desc in questions:
        p = doc.add_paragraph()
        run = p.add_run("  " + q)
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(" -> " + desc)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    doc.add_paragraph()

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、案例分析")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【观察】数据显示：门店销售额与员工满意度高度正相关（r=0.85）")
    run.font.size = Pt(11)

    from docx.oxml import OxmlElement

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ["分析维度", "内容"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        para = table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    analysis_data = [
        ["错误归因", "员工满意度高 -> 销售额高（错误！这是把相关性当因果性）"],
        ["正确分析", "因果方向可能是：\n1. 销售额高 -> 员工收入高 -> 满意度高\n2. 有共同原因C：店长管理能力强 -> 同时提升销售和满意\n3. 可能存在反向因果：满意度高的员工更努力 -> 提升销售"],
        ["关键追问", "如果只提升满意度，而不提升销售能力，销售额会提高吗？\n答案：不会。这叫『伪因』陷阱。"],
        ["正确做法", "1. 首先提升销售能力（因）-> 销售额提升（果）\n2. 销售提升后 -> 员工收入提升 -> 满意度自然提升\n3. 满意度提升是结果，不是手段"]
    ]

    for i, (label, content) in enumerate(analysis_data):
        table.rows[i+1].cells[0].text = label
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、因果归因常见错误清单")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    errors = [
        ("混淆相关与因果", "冰淇淋销量高与溺水人数高相关，都是夏天导致的"),
        ("错误因果方向", "认为是员工满意导致销售好，其实可能是销售好导致满意"),
        ("忽略共同原因", "两个变量都受第三个变量影响，但被忽略了"),
        ("过度简化", "认为只有一个原因，现实通常是多重因素")
    ]

    error_table = doc.add_table(rows=len(errors)+1, cols=2)
    error_table.style = 'Table Grid'

    error_table.rows[0].cells[0].text = "错误类型"
    error_table.rows[0].cells[1].text = "说明"
    for cell in error_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (err_type, desc) in enumerate(errors):
        error_table.rows[i+1].cells[0].text = err_type
        error_table.rows[i+1].cells[1].text = desc
        for cell in error_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo07-因果归因分析示例.docx"
    doc.save(output_path)
    print("demo07已生成:", output_path)
    return output_path


def create_demo08():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("竞争性假说复盘表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 系统性检验所有可能的解释")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：赵复盘", "团队：决策质量组", "日期：2026-08-05", "教练：刘分析"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】竞争性假说法 = 对同一现象提出多个可能解释，然后系统性地检验每个假说")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、案例：项目失败原因分析")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【事件】某重要项目延期3个月交付，客户强烈投诉")
    run.font.size = Pt(11)

    from docx.oxml import OxmlElement

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    headers = ["假说", "支持证据", "反对证据", "可信度"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        para = table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    hypothesis_data = [
        ["假说A\n需求变更频繁", "客户确实在开发过程中提出了30+次变更", "变更都有正式审批，且在可控范围内", "低\n变更不是主因"],
        ["假说B\n团队能力不足", "核心开发人员有3人离职", "新补充的人员背景也不错", "中\n需进一步分析"],
        ["假说C\n项目管理失控", "每周会议记录显示问题积累严重\n里程碑多次延误", "项目经理很有经验\n问题在系统不在人", "高\n证据最充分"],
        ["假说D\n技术选型失误", "确实使用了较新的技术栈", "同行类似项目也有成功案例", "低\n技术不是主因"]
    ]

    for i, row_data in enumerate(hypothesis_data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text
            for para in table.rows[i+1].cells[j].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    conclusion = doc.add_paragraph()
    run = conclusion.add_run('【复盘结论】项目失败的主要原因是『假说C：项目管理失控』')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    conclusion_detail = doc.add_paragraph()
    run = conclusion_detail.add_run('具体问题：\n1. 问题升级机制失效，小问题积累成大问题\n2. 跨部门协调效率低，关键决策延误\n3. 风险识别滞后，没有提前预警\n\n改进措施：\n- 建立『问题升级红线』机制\n- 设立项目『风险预警仪表盘』\n- 每周向客户同步风险清单')
    run.font.size = Pt(10)

    doc.add_paragraph()

    tips = doc.add_paragraph()
    run = tips.add_run('[教练提示] 竞争性假说法的关键是『让证据说话』而非『让权威说话』。每个假说都要有支持/反对证据，不能只凭直觉。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo08-竞争性假说复盘表示例.docx"
    doc.save(output_path)
    print("demo08已生成:", output_path)
    return output_path


def create_demo09():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("调研方案自查表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 调研前自查，确保调研有效性")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：孙调研", "团队：用户研究组", "日期：2026-08-07", "教练：周研究"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心原则】好的调研方案要回答：我需要什么信息？从哪里获取？如何确保真实性？")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    from docx.oxml import OxmlElement

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、调研方案自查清单")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    check_items = [
        ("【目的自查】", "本次调研要回答什么问题？\n这个问题能指导后续决策吗？\n调研结果的使用者是谁？"),
        ("【样本自查】", "样本能代表目标群体吗？\n样本量足够得出结论吗？\n是否有抽样偏差风险？"),
        ("【方法自查】", "定量问卷还是定性访谈？\n方法与目的匹配吗？\n调研形式是否会让受访者说『假话』？"),
        ("【问题自查】", "问题是否引导性太强？\n是否有双重否定或复杂句式？\n敏感问题如何处理？"),
        ("【分析自查】", "收集回来的数据如何分析？\n预设了哪些分析维度？\n交叉分析是否能揭示关联？")
    ]

    check_table = doc.add_table(rows=len(check_items)+1, cols=2)
    check_table.style = 'Table Grid'

    check_table.rows[0].cells[0].text = "自查维度"
    check_table.rows[0].cells[1].text = "自查要点"
    for cell in check_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for i, (dim, points) in enumerate(check_items):
        check_table.rows[i+1].cells[0].text = dim
        check_table.rows[i+1].cells[1].text = points
        check_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in check_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、案例应用")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【调研主题】员工满意度调研方案自查")
    run.font.size = Pt(11)
    run.font.bold = True

    case_table = doc.add_table(rows=6, cols=2)
    case_table.style = 'Table Grid'

    case_table.rows[0].cells[0].text = "自查维度"
    case_table.rows[0].cells[1].text = "学员自查结果"
    for cell in case_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    case_data = [
        ["目的自查", '回答『员工为什么不满意』，为管理层提供改进依据\n使用者：HR总监+CEO\n能指导决策'],
        ["样本自查", "全员匿名问卷，覆盖率目标大于85%\n风险：高管层可能参与率低\n已设计分层分析"],
        ["方法自查", "定量问卷+匿名保证\n风险：员工可能因担心匿名性不敢说真话\n增加开放性问题收集定性反馈"],
        ["问题自查", '原方案中『您对直接上级满意吗？』引导性太强\n修改为『您与直接上级的协作顺畅度如何？』'],
        ["分析自查", "预设按部门、司龄、职级交叉分析\n增加离职意向与满意度的关联分析"]
    ]

    for i, (dim, result) in enumerate(case_data):
        case_table.rows[i+1].cells[0].text = dim
        case_table.rows[i+1].cells[1].text = result
        case_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in case_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo09-调研方案自查表示例.docx"
    doc.save(output_path)
    print("demo09已生成:", output_path)
    return output_path


def create_demo10():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("判断标准梳理表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 明确决策的判断维度与标准")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：李决策", "团队：战略投资组", "日期：2026-08-10", "教练：张标准"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】判断标准梳理 = 明确『好』与『不好』的边界在哪里")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    from docx.oxml import OxmlElement

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、案例：供应商准入判断标准")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    headers = ["判断维度", "权重", "判断标准", "数据来源"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        para = table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    criteria_data = [
        ["质量能力", "30%", "最近批次合格率大于等于98%\n通过ISO9001认证", "质检报告\n证书"],
        ["交付能力", "25%", "准时交货率大于等于95%\n平均交期小于等于15天", "历史订单数据"],
        ["成本优势", "20%", "价格不高于市场均价5%\n年降机制承诺", "比价分析\n报价单"],
        ["服务能力", "15%", "24小时响应\n专属客服配置", "服务协议\n访谈"],
        ["可持续发展", "10%", "环保资质\n员工保障合规", "审核报告\n资质证书"]
    ]

    for i, row_data in enumerate(criteria_data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text
            for para in table.rows[i+1].cells[j].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    tips = doc.add_paragraph()
    run = tips.add_run("[使用说明] 判断标准梳理的关键：")
    run.font.size = Pt(11)
    run.font.bold = True

    tip_items = [
        "每个维度要有明确『通过/不通过』的边界",
        "权重反映业务优先级，不是简单平均",
        "标准要可量化、可验证，避免主观判断",
        "定期回顾标准是否符合业务变化"
    ]

    for tip in tip_items:
        p = doc.add_paragraph()
        run = p.add_run("  [v] " + tip)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    doc.add_paragraph()

    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、标准梳理常见问题")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    problem_table = doc.add_table(rows=4, cols=2)
    problem_table.style = 'Table Grid'

    problem_table.rows[0].cells[0].text = "问题"
    problem_table.rows[0].cells[1].text = "解决方案"
    for cell in problem_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    problems = [
        ["标准太模糊（如『质量合格』）", "改为可量化标准（如『合格率大于等于98%』）"],
        ["权重不合理", "与业务方对齐，确保优先级一致"],
        ["标准之间矛盾", "明确哪个标准是一票否决的"]
    ]

    for i, (prob, sol) in enumerate(problems):
        problem_table.rows[i+1].cells[0].text = prob
        problem_table.rows[i+1].cells[1].text = sol
        for cell in problem_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo10-判断标准梳理表示例.docx"
    doc.save(output_path)
    print("demo10已生成:", output_path)
    return output_path


if __name__ == "__main__":
    create_demo07()
    create_demo08()
    create_demo09()
    create_demo10()

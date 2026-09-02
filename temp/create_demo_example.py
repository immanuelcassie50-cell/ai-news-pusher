# -*- coding: utf-8 -*-
"""
创建 002-Demo示例-智合集团全套成果.docx - 智合集团完整成果示例
基于战略解码双螺旋引擎课程内容
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table_with_header(doc, headers, rows, header_color="4472C4"):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)

    return table

def create_example():
    doc = Document()

    # 设置文档标题
    title = doc.add_heading('战略解码双螺旋引擎', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Demo示例：智合集团全套成果')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(16)

    doc.add_paragraph()

    # 案例背景说明
    add_heading(doc, '案例背景：智合集团', 1)

    background = """智合集团是一家有8000人的工业设备制造企业，有四个主要业务事业部。2022年底，集团完成了新三年战略规划：从"销售设备"转向"提供工业设备全生命周期解决方案"，目标是2025年服务与解决方案业务占收入比例从当前的8%提升至35%。

战略故事很清晰，高管团队共识也达成了，启动大会开了，战略手册印了。

2024年初，战略推进了14个月之后，董事长张国豪在季度回顾中发现一个问题："我们现在的销售团队，有多少人真正在用解决方案的方式拜访客户？" 销售VP诚实地说："不到10%。"

这说明战略在文件里，执行在原地踏步。接下来的全套成果示例，展示智合集团如何用双螺旋引擎解决战略落地问题。"""

    p = doc.add_paragraph(background)

    doc.add_paragraph()

    # ==================== 战役一：解决方案销售能力建设 ====================
    add_heading(doc, '一、战役一：解决方案销售能力建设战役定义', 1)

    add_heading(doc, '战役一定义', 2)

    battle1 = [
        ['项目', '内容'],
        ['战役名称', '解决方案销售能力建设战役'],
        ['赢的标准', '2024年底，大客户销售团队60%完成解决方案销售认证；解决方案型商机在管道中占比从10%提升至40%'],
        ['战役负责人', '销售VP 李建华'],
        ['战役时限', '2024年3月 — 2024年12月'],
        ['战略意义', '如果打不赢这场，解决方案业务无从谈起——没有能卖的人，35%目标是空话'],
    ]

    table = doc.add_table(rows=len(battle1), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(battle1):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "D9E2F3")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== 战役二：核心标杆客户深度渗透 ====================
    add_heading(doc, '二、战役二：核心标杆客户深度渗透战役定义', 1)

    add_heading(doc, '战役二定义', 2)

    battle2 = [
        ['项目', '内容'],
        ['战役名称', '核心标杆客户深度渗透战役'],
        ['赢的标准', '2024年底，与6家战略客户建立"解决方案伙伴"关系；签署解决方案合同不少于3份，总金额不少于8000万'],
        ['战役负责人', '战略客户总监 陈明'],
        ['战役时限', '2024年3月 — 2024年12月'],
        ['战略意义', '没有成功样板，规模化推广没有说服力，也没有可复制的交付经验'],
    ]

    table = doc.add_table(rows=len(battle2), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(battle2):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "D9E2F3")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== 战役一完整拆解 ====================
    add_heading(doc, '三、战役一完整拆解（战术策略+关键战斗）', 1)

    doc.add_paragraph('战役名称：解决方案销售能力建设战役')
    doc.add_paragraph('赢的标准：2024年底，60%大客户销售完成认证；解决方案型商机管道占比40%')
    doc.add_paragraph('战役负责人：李建华（销售VP）')

    add_heading(doc, '战术策略一：建立能力认证体系', 2)

    headers = ['关键战斗', '完成时间', '成果标志', '负责人']
    tactic1 = [
        ['设计"解决方案销售"能力模型，明确10个核心能力维度', '4月底完成', '能力模型文档获管理层确认', '李建华+HR'],
        ['开发认证培训课程，含3个模块共计16小时', '5月底完成', '课程通过种子销售测试', 'HR+外部讲师'],
        ['第一批培训，覆盖30%大客户销售', '7月底完成', '30%通过初级认证', 'HR'],
    ]
    add_table_with_header(doc, headers, tactic1)

    doc.add_paragraph()

    add_heading(doc, '战术策略二：建立解决方案知识库与标杆案例体系', 2)

    headers = ['关键战斗', '完成时间', '成果标志', '负责人']
    tactic2 = [
        ['挖掘现有成功交付的解决方案项目，整理为3-5个标准化案例', '4月底完成', '3个完整案例文档', '市场部+销售'],
        ['制定案例标准化模板（客户背景/问题诊断/方案设计/交付过程/量化成果）', '5月底完成', '模板获确认', '市场部'],
        ['建立每月1案例的持续更新机制', '7月起持续执行', '全年目标12个案例入库', '市场部'],
    ]
    add_table_with_header(doc, headers, tactic2)

    doc.add_paragraph()

    add_heading(doc, '战术策略三：建立技术+销售联合拜访机制', 2)

    headers = ['关键战斗', '完成时间', '成果标志', '负责人']
    tactic3 = [
        ['识别并培训20名解决方案架构师，建立"技术支援销售"配对清单', '4月底完成', '配对清单确认', '技术VP+李建华'],
        ['设计联合拜访申请流程和总结模板', '4月底完成', '流程文件确认', '销售+技术'],
        ['完成每月20次以上联合拜访', '5月起执行', '每月联合拜访记录台账', '联合团队'],
    ]
    add_table_with_header(doc, headers, tactic3)

    doc.add_paragraph()

    add_heading(doc, '跨部门协同需求', 2)

    headers = ['部门', '需要提供的资源或配合', '时间要求']
    cross = [
        ['技术部门', '提供20名解决方案架构师，支持联合拜访（占工作时间约30%）', '全年'],
        ['HR部门', '配合设计能力认证评估体系，协助培训组织', '全年'],
        ['市场部门', '协助标杆案例的包装，支持案例传播', '5月起'],
    ]
    add_table_with_header(doc, headers, cross)

    doc.add_paragraph()

    add_heading(doc, '战役一关键里程碑', 2)

    headers = ['时间节点', '里程碑', '判断标准']
    milestones = [
        ['4月底', '能力模型确认 + 联合拜访机制就绪', '文档管理层签字；第一次联合拜访完成'],
        ['5月底', '培训课程就绪 + 知识库模板确认', '完成种子销售测试；第一批案例入库'],
        ['7月底', '第一批认证完成', '30%大客户销售通过初级认证'],
        ['9月底', '中期里程碑', '40%认证达成；解决方案管道占比达20%'],
        ['12月底', '战役终点', '60%认证达成；管道占比40%'],
    ]
    add_table_with_header(doc, headers, milestones)

    doc.add_paragraph()

    # ==================== 战役OKR ====================
    add_heading(doc, '四、战役OKR（O和KR）', 1)

    add_heading(doc, '战役一OKR', 2)

    okr1 = [
        ['类型', '内容'],
        ['O', '在2024年，让智合集团大客户销售团队真正具备以顾问为导向的解决方案销售能力，让"卖解决方案"不再是口号，而是看得见的能力'],
        ['KR1', '大客户销售团队中60%完成解决方案销售认证（截至2024年12月底）'],
        ['KR2', '解决方案型商机在销售管道中的占比从10%提升至40%（Q4末）'],
        ['KR3', '完成至少3个可对外引用的解决方案标杆案例，含完整的问题诊断-方案设计-交付过程-客户收益记录（12月底）'],
    ]

    table = doc.add_table(rows=len(okr1), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(okr1):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "E2EFDA")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    add_heading(doc, '战役二OKR', 2)

    okr2 = [
        ['类型', '内容'],
        ['O', '在2024年，与6家战略客户建立深度合作关系，打造可复制的解决方案交付样板，为规模化推广奠定信心基础'],
        ['KR1', '与6家战略客户建立"解决方案伙伴"关系（截至2024年12月底）'],
        ['KR2', '签署解决方案合同不少于3份，总金额不少于8000万（12月底）'],
        ['KR3', '完成至少2个标杆客户的完整交付案例，90天客户业务改善指标达成率≥70%（12月底）'],
    ]

    table = doc.add_table(rows=len(okr2), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(okr2):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "E2EFDA")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== BSC战略图 ====================
    add_heading(doc, '五、BSC战略图（四层+因果链）', 1)

    doc.add_paragraph('战略主题：从工业设备供应商转型为全生命周期解决方案伙伴')

    add_heading(doc, '学习与成长层面（基础能力）', 2)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    lr_layer = [
        ['L1：建立解决方案销售能力', '大客户销售认证通过率；解决方案型商机比例', '60%认证通过；40%商机占比', '2024年底'],
        ['L2：建立解决方案设计与交付能力', '认证的解决方案架构师数量；方案交付满意度', '30名架构师；满意度>85%', '2024年底'],
        ['L3：建立以客户成功为核心的服务文化', '内部文化评估得分；客户成功案例记录数量', '文化评估年增长≥15%；年度案例≥12个', '2024年底'],
    ]
    add_table_with_header(doc, headers, lr_layer)

    doc.add_paragraph()

    add_heading(doc, '内部流程层面（关键流程）', 2)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    ip_layer = [
        ['I1：建立高质量客户诊断与方案开发流程', '方案诊断报告质量评分；方案转化成合同的比率', '质量评分≥4/5；转化率10%→30%', '2024年底'],
        ['I2：建立高效解决方案交付与客户成功跟踪流程', '方案按时交付率；交付后90天客户业务改善指标完成率', '按时交付率≥90%；90天改善率≥70%', '2024年底'],
        ['I3：建立解决方案知识管理与复用机制', '案例库覆盖度；方案组件复用率', '12个行业标杆案例；复用率≥40%', '2024年底'],
    ]
    add_table_with_header(doc, headers, ip_layer)

    doc.add_paragraph()

    add_heading(doc, '客户层面（客户价值）', 2)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    c_layer = [
        ['C1：成为战略客户的首选长期合作伙伴', '战略客户NPS；合同续约率；大客户年均合同价值', 'NPS 45→65；续约率≥85%；合同价值+40%', '2024年底'],
        ['C2：以解决方案为入口，赢得新客户', '解决方案主导的新客户获取数量；解决方案首单赢单率', '新客户≥15家；赢单率≥35%', '2024年底'],
    ]
    add_table_with_header(doc, headers, c_layer)

    doc.add_paragraph()

    add_heading(doc, '财务层面（财务结果）', 2)

    headers = ['战略目标', '衡量指标', '目标值', '时间节点']
    f_layer = [
        ['F1：服务与解决方案收入占比提升', '服务业务收入/总收入', '15%（2024）→25%（中期）→35%（2025）', '2024-2025'],
        ['F2：客户终身价值提升', '战略客户年均合同总价值；客户生命周期收入', '战略客户人均合同价值提升40%', '2024年底'],
    ]
    add_table_with_header(doc, headers, f_layer)

    doc.add_paragraph()

    add_heading(doc, '因果链', 2)

    causals = [
        ['因果链', '路径'],
        ['路径一', 'L1（销售能力建设）→ I1（诊断和方案开发流程）→ C1（战略客户首选伙伴）→ F1（服务收入占比）'],
        ['路径二', 'L2（交付能力建设）→ I2（交付和客户成功）→ C1（合同续约率）→ F2（客户终身价值）'],
        ['路径三', 'L3（服务文化）→ I3（知识管理复用）→ C2（解决方案吸引新客户）→ F1（服务收入结构改善）'],
    ]

    table = doc.add_table(rows=len(causals), cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(causals):
        table.rows[i].cells[0].text = row_data[0]
        table.rows[i].cells[1].text = row_data[1]
        set_cell_shading(table.rows[i].cells[0], "FFF2CC")
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ==================== 部门绩效设计 ====================
    add_heading(doc, '六、部门绩效设计（大客户销售部门为例）', 1)

    add_heading(doc, '大客户销售部门绩效设计', 2)

    headers = ['指标类型', '指标名称', '衡量方式', '目标值', '权重']
    perf_design = [
        ['战略级', '战略客户NPS', '季度客户调研', '45→65（年底）', '20%'],
        ['战略级', '解决方案型商机占比', 'CRM商机类型统计', '10%→40%（年底）', '20%'],
        ['战略级', '解决方案合同赢单率', '合同数量/商机数量', '≥35%（年底）', '15%'],
        ['运营级', '大客户总收入', '财务系统', '年度目标', '30%'],
        ['运营级', '新大客户开发数量', 'CRM新客户记录', '年度目标', '15%'],
    ]
    add_table_with_header(doc, headers, perf_design)

    doc.add_paragraph()

    add_heading(doc, '大客户销售经理关键岗位绩效', 2)

    headers = ['指标类型', '指标名称', '衡量方式', '季度目标', '权重']
    key_perf = [
        ['结果', '负责客户的解决方案合同签约额', '合同财务记录', '季度目标×110%', '25%'],
        ['结果', '负责客户的NPS评分', '季度客户调研', '≥65分', '20%'],
        ['结果', '解决方案型商机在管道中的占比', 'CRM系统', '≥40%', '15%'],
        ['行为', '每月联合拜访（技术+销售）次数', '拜访记录系统', '≥4次/月', '20%'],
        ['行为', '完成的客户诊断报告数量和质量', '报告评审系统', '≥2份/季度，质量≥4分', '20%'],
    ]
    add_table_with_header(doc, headers, key_perf)

    doc.add_paragraph()

    add_heading(doc, '其他部门绩效设计摘要', 2)

    headers = ['部门', '对应战略层面', '核心战略级绩效指标']
    other_depts = [
        ['技术解决方案部门', 'I1、I2、L2', '①方案诊断报告质量评分 ②方案转化成合同比率 ③交付后90天客户业务改善率'],
        ['人力资源部门', 'L1、L3', '①销售解决方案认证通过率 ②关键岗位能力差距填补率 ③组织文化评估得分'],
        ['市场部门', 'C2、I3', '①解决方案内容触达战略客户数量 ②案例库更新数量和质量评分'],
    ]
    add_table_with_header(doc, headers, other_depts)

    doc.add_paragraph()

    # ==================== 课后行动计划 ====================
    add_heading(doc, '七、课后行动计划（基于智合集团案例）', 1)

    add_heading(doc, '第一部分：引擎一——必赢战役', 2)

    headers = ['项目', '战役一', '战役二']
    action1 = [
        ['战役名称', '解决方案销售能力建设战役', '核心标杆客户深度渗透战役'],
        ['赢的标准', '60%认证；40%商机占比', '6家伙伴关系；3份合同≥8000万'],
        ['战役指挥官', '销售VP 李建华', '战略客户总监 陈明'],
        ['前三个关键战斗', '①能力模型设计 ②培训课程开发 ③首批30%培训', '①6家战略客户识别 ②伙伴关系协议签署 ③标杆案例打造'],
        ['战役OKR的O', '成为解决方案导向的销售团队', '建立可复制的标杆样板'],
    ]
    add_table_with_header(doc, headers, action1)

    doc.add_paragraph()

    add_heading(doc, '第二部分：引擎二——部门绩效协同', 2)

    headers = ['项目', '内容']
    action2 = [
        ['BSC战略图的战略主题', '从工业设备供应商转型为全生命周期解决方案伙伴'],
        ['学习层最关键的1个目标', 'L1：建立解决方案销售能力（60%认证通过率）'],
        ['流程层最关键的1个目标', 'I1：建立高质量客户诊断与方案开发流程'],
        ['部门核心战略级绩效指标（3个）', '①战略客户NPS ②解决方案型商机占比 ③解决方案合同赢单率'],
        ['需要更新/新增的绩效指标', '新增：解决方案型商机占比、解决方案合同赢单率'],
    ]
    add_table_with_header(doc, headers, action2)

    doc.add_paragraph()

    add_heading(doc, '第三部分：回去之后的第一件事', 2)

    headers = ['项目', '内容']
    action3 = [
        ['回去后第一周最重要的一件事', '召开战略聚焦会议，识别2-3个必赢战役'],
        ['要找的人', '董事长/CEO + 核心管理团队'],
        ['要用的工具', '必赢战役四维战法 + BSC战略图模板'],
    ]
    add_table_with_header(doc, headers, action3)

    doc.add_paragraph()

    # ==================== 战役执行机制 ====================
    add_heading(doc, '附录：战役执行机制设计（智合集团案例）', 1)

    headers = ['机制要素', '智合集团的设计']
    exec_mechanism = [
        ['战役指挥官', '销售VP 李建华，拥有跨部门调动权，对战役成败向董事长直接负责'],
        ['战役指挥部', '每2周一次战役回顾会（周一上午9:00，固定不取消），成员：李建华+技术VP+HR负责人+市场VP代表；每月一次董事长战役汇报（PPT不超过5页）'],
        ['可视化看板', '战役指挥室设有物理看板，分三栏：关键里程碑状态/当前周关键战斗/风险清单；在线同步版本用钉钉项目管理板'],
        ['快速决策通道', '李建华遇到跨部门阻碍时，有权在24小时内邮件提交给董事长，要求在48小时内得到决策'],
        ['核心战役团队', '8人：销售侧4人+技术侧2人+HR侧1人+市场侧1人；每人每周保留20%时间专属于战役工作'],
    ]
    add_table_with_header(doc, headers, exec_mechanism)

    # 保存文档
    output_path = 'D:/新课开发/战略和领导力/战略解码双螺旋引擎让大象跳舞/完整课程包/009-Demo成果/002-Demo示例-智合集团全套成果.docx'
    doc.save(output_path)
    print(f'示例文档已保存至: {output_path}')

if __name__ == '__main__':
    create_example()

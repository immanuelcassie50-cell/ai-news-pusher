# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {})
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

output_dir = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "讲师手册-变革成果固化机制.docx")

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Cover Page
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("变革成果固化机制")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(36)
r.bold = True
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("防止新流程人走茶凉")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(28)
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("讲师手册")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(24)
r.bold = True
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("国际版权课标准")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(18)
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI转型深水区变革管理系列课程第16门")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(14)
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("课程时长：2天（12小时）| 建议班级规模：25-35人")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(12)
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# TOC Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目  录")
r.font.name = 'Microsoft YaHei'
r.font.size = Pt(24)
r.bold = True
r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph()

toc_items = [
    "第一章  课程整体设计理念",
    "第二章  模块一：固化机制认知（2小时）",
    "第三章  模块二：制度固化（3小时）",
    "第四章  模块三：机制固化（3小时）",
    "第五章  模块四：文化固化（2小时）",
    "第六章  模块五：固化效果检验与迭代（2小时）",
    "第七章  案例使用说明和讨论引导问题",
    "第八章  常见问题和应对方案",
    "第九章  教室布置与道具准备清单",
    "第十章  效果评估方法",
    "第十一章  讲师注意事项",
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(item)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(12)
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()
print('Part 1 created')
doc.save(output_path)
print('Saved:', output_path)

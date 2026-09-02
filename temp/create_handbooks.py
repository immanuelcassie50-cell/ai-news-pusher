# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

STUDENT_OUTPUT = "D:/新课开发/法学/09-被裁员之后：补偿、仲裁与职场维权全流程/学员手册/学员手册_被裁员之后.docx"
INSTRUCTOR_OUTPUT = "D:/新课开发/法学/09-被裁员之后：补偿、仲裁与职场维权全流程/讲师手册/讲师手册_被裁员之后.docx"

def set_font(run, font_name='Microsoft YaHei', size=None, bold=False):
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level=1, font_size=16, bold=True):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=font_size, bold=bold)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_body(doc, text, font_size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=font_size)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet(doc, text, font_size=11):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_font(run, size=font_size)
    para.paragraph_format.space_after = Pt(3)
    return para

def add_number(doc, text, font_size=11):
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    set_font(run, size=font_size)
    para.paragraph_format.space_after = Pt(3)
    return para

def table_row(table, cells_data, bold=False, font_size=10):
    row = table.add_row()
    for i, cell_text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, size=font_size, bold=bold)
    return row

def add_fill(doc, label):
    para = doc.add_paragraph()
    run = para.add_run(label + ": ")
    set_font(run, size=11, bold=True)
    run = para.add_run("_" * 40)
    set_font(run, size=11)
    para.paragraph_format.space_after = Pt(8)
    return para

def create_student_handbook():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Cover
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(60)
    run = para.add_run("学员手册")
    set_font(run, size=24, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("被裁员之后：补偿、仲裁与职场维权全流程")
    set_font(run, size=28, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(40)
    run = para.add_run("适用场景：面临或已经经历裁员的职场人\n核心收获：裁员谈判前48小时行动清单 + 协商解除协议审查要点 + 三不签红线清单")
    set_font(run, size=14)

    doc.add_page_break()

    # TOC
    add_heading(doc, "目 录", font_size=18)
    toc = [
        "第一章  课程导言",
        "第二章  五步维权扫描法",
        "第三章  模块一：补偿金计算",
        "第四章  模块二：话术套路破解",
        "第五章  模块三：协议审查要点",
        "第六章  模块四：证据保全",
        "第七章  模块五：仲裁全流程",
        "第八章  模块六：谈判策略",
        "第九章  行动清单：裁员谈判前48小时必查",
        "第十章  资源包"
    ]
    for item in toc:
        add_bullet(doc, item, 12)

    doc.add_page_break()

    # Ch1
    add_heading(doc, "第一章  课程导言", font_size=16)
    add_heading(doc, "为什么学这门课？", font_size=14)
    add_body(doc, '在职场中，裁员已经不再是罕见事件。根据劳动法规定，企业裁员必须支付经济补偿，但在实操中，很多公司利用员工对法律的不了解，以"协商离职"为名行"强制裁员"之实。\n\n你是否遇到过以下情况：')

    situations = [
        'HR找你谈话，说公司经营困难，需要"优化"人员',
        "被要求签署离职申请，承诺放弃一切权益",
        '补偿方案模糊，只给"N"或"N+1"，不知道怎么算出来的',
        "签了字才发现吃大亏，但已经来不及了",
        "公司威胁：不签字就没有任何补偿"
    ]
    for s in situations:
        add_bullet(doc, s)

    add_heading(doc, "学完这门课，你能带走什么？", font_size=14)
    takeaways = [
        ("补偿金计算器", "清楚知道自己应得多少补偿，N/N+1/2N不再糊涂"),
        ("话术防骗指南", "一眼看穿HR常用套路，不再被牵着鼻子走"),
        ("协议审查清单", "知道协议里哪些条款是坑，哪些字不能签"),
        ("证据保全手册", "知道哪些东西要截图、录屏、存档"),
        ("仲裁全流程图", "知道什么时候该仲裁、怎么仲裁、要准备什么"),
        ("48小时行动清单", "拿到裁员通知后第一时间做什么，不遗漏")
    ]
    for title, desc in takeaways:
        para = doc.add_paragraph()
        run = para.add_run("【" + title + "】")
        set_font(run, size=11, bold=True)
        run = para.add_run(desc)
        set_font(run, size=11)

    doc.add_page_break()

    # Ch2
    add_heading(doc, "第二章  五步维权扫描法", font_size=16)
    add_body(doc, "这是本课程的核心方法论，适合任何劳动争议场景。")

    steps = [
        ("第一步：算清楚", "把自己的补偿金额、社保公积金、未休年假等各项权益全部算清楚。知道自己应得什么，是谈判的基础。"),
        ("第二步：查证据", "收集并固定所有相关证据：劳动合同、工资流水、社保记录、考勤记录、录音录像、聊天截图等。"),
        ("第三步：审协议", "逐条审查解除协议内容，识别霸王条款和隐藏陷阱。记住：三不签红线。"),
        ("第四步：定策略", "根据公司态度、自身筹码、证据完整度，决定是协商还是仲裁，选择最优路径。"),
        ("第五步：止损失", "通过谈判达成协议或通过仲裁拿到裁决，最大化保护自己的合法权益。")
    ]
    for title, desc in steps:
        para = doc.add_paragraph()
        run = para.add_run(title)
        set_font(run, size=12, bold=True)
        run = para.add_run("\n" + desc)
        set_font(run, size=11)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Ch3
    add_heading(doc, "第三章  模块一：补偿金计算", font_size=16)
    add_heading(doc, "3.1 三种补偿情形一览", font_size=14)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table_row(table, ["类型", "适用情形", "计算公式", "举例（月薪1万，工龄3年）"], True, 10)
    rows = [
        ["N", "合法裁员/协商解除", "工作年限 x 月工资", "3万"],
        ["N+1", "即时辞退（无提前通知）", "N + 1个月工资", "4万"],
        ["2N", "违法解除劳动合同", "2 x 工作年限 x 月工资", "6万"]
    ]
    for r in rows:
        table_row(table, r, False, 10)

    add_heading(doc, "3.2 工资基数怎么算？", font_size=14)
    add_body(doc, "劳动合同解除前12个月的平均工资，包括：")
    for item in ["基本工资", "绩效工资", "奖金、加班费", "补贴、津贴", "其他货币性收入"]:
        add_bullet(doc, item)
    add_body(doc, "注意：不是基本工资，而是应发工资总额！")

    add_heading(doc, "3.3 年限计算规则", font_size=14)
    for rule in [
        "每满1年算1个月，不满半年算半个月",
        "6个月以上不满1年按1年计算",
        "月工资高于当地平均工资3倍的，按3倍计算，最高12年",
        "试用期也要计算年限"
    ]:
        add_number(doc, rule)

    doc.add_page_break()

    # Ch4
    add_heading(doc, "第四章  模块二：话术套路破解", font_size=16)
    add_body(doc, "HR常用的心理战术和应对话术：")

    tactics = [
        {
            "title": "套路1：制造紧迫感",
            "hr_say": '"今天不签，明天就没有这个方案了"',
            "purpose": "让你没有思考时间，被迫签字",
            "response": '"我需要时间看一下合同，这是我的权利。明天我再来。"',
            "analysis": "法律规定你有权查看合同至少15分钟，任何催促都是心虚的表现。"
        },
        {
            "title": "套路2：画大饼",
            "hr_say": '"签字后可以给你出具主动离职的证明，方便你找工作"',
            "purpose": "让你觉得主动离职比被动裁员好",
            "response": '"我不需要这样的证明，请问补偿金具体是多少？"',
            "analysis": "主动离职证明在劳动仲裁中没有任何价值，补偿金才是实实在在的。"
        },
        {
            "title": "套路3：威胁恐吓",
            "hr_say": '"不签字的话，我们会以严重违纪为由开除，到时候一分钱都没有"',
            "purpose": "用恐惧让你就范",
            "response": '"请公司出具书面通知。如果公司坚持违法解除，我会依法维权。"',
            "analysis": "公司敢以严重违纪开除，就需要有确凿证据，否则就是违法解除，反而赔2N。"
        },
        {
            "title": "套路4：混淆概念",
            "hr_say": '"N就是给你3个月工资，已经很良心了"',
            "purpose": '把N+1的1说成是"额外福利"',
            "response": '"我的理解是N是法定补偿。请把计算方式写清楚。"',
            "analysis": "N是法定义务，不是公司恩赐。如果协商解除，N+1才是常规操作。"
        }
    ]

    for t in tactics:
        para = doc.add_paragraph()
        run = para.add_run("【" + t['title'] + "】")
        set_font(run, size=12, bold=True)
        para.paragraph_format.space_before = Pt(12)
        add_body(doc, "HR说：" + t['hr_say'])
        add_body(doc, "目的：" + t['purpose'])
        add_body(doc, "应对：")
        para = doc.add_paragraph()
        run = para.add_run(t['response'])
        set_font(run, size=11, bold=True)
        para.paragraph_format.left_indent = Pt(20)
        add_body(doc, "分析：" + t['analysis'])

    doc.add_page_break()

    # Ch5
    add_heading(doc, "第五章  模块三：协议审查要点", font_size=16)
    add_heading(doc, "5.1 三不签红线清单", font_size=14)
    add_body(doc, "遇到以下条款，绝对不能签：")

    red_lines = [
        ("X", "放弃一切诉讼权利", "这是剥夺你依法维权的基本权利，违法无效，但签字后举证困难"),
        ("X", '"双方再无任何劳动争议"', "一旦签字，即使漏算了补偿也无法追讨"),
        ("X", '"工资、奖金已全部结清"', "如果还有未休年假折算、加班费未结，签了就追不回来了"),
        ("X", '"乙方不得入职同行企业"', "竞业限制需要单独支付补偿金，不能强制捆绑"),
        ("X", '"甲方有权随时调整乙方岗位"', "岗位调整属于变更劳动合同，需要双方协商同意")
    ]

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table_row(table, ["标记", "条款内容", "风险说明"], True, 10)
    for r in red_lines:
        table_row(table, r, False, 10)

    add_heading(doc, "5.2 必须包含的条款", font_size=14)
    for item in [
        "离职日期（年月日必须明确）",
        "经济补偿金金额及支付时间",
        "工资、加班费、年假折算等各项结清",
        "社保、公积金缴纳至哪个月",
        "工作交接具体内容",
        "保密义务范围（合法合理）",
        "竞业限制条款（如有，须单独约定补偿）"
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # Ch6
    add_heading(doc, "第六章  模块四：证据保全", font_size=16)
    add_heading(doc, "6.1 哪些证据必须收集？", font_size=14)

    evidence = [
        ("劳动合同类", ["劳动合同原件（拍照存档）", "offer/录用通知", "岗位职责说明书", "绩效考核表"]),
        ("工资社保类", ["工资条/银行流水（体现月薪）", "社保缴费记录", "公积金缴费记录", "年终奖发放记录"]),
        ("考勤记录类", ["打卡记录/考勤系统截图", "加班申请记录", "请假记录", "年假使用情况"]),
        ("沟通记录类", ["与HR的沟通记录（微信、邮件）", "裁员通知文件", "工作群聊记录", "录音录像（谈话时告知对方）"])
    ]

    for cat, items in evidence:
        para = doc.add_paragraph()
        run = para.add_run("【" + cat + "】")
        set_font(run, size=12, bold=True)
        for item in items:
            add_bullet(doc, item)

    add_heading(doc, "6.2 证据怎么固定？", font_size=14)
    for method, desc in [
        ("电子证据（微信、邮件）", "录屏+截图双备份，录屏要体现完整时间线"),
        ("录音录像", "原始文件保存在手机，不要发给别人，备份到云端"),
        ("纸质文件", "拍照+扫描，文件名注明日期和内容"),
        ("公证", "重要证据可以到公证处公证，效力更强")
    ]:
        para = doc.add_paragraph()
        run = para.add_run(method + "：")
        set_font(run, size=11, bold=True)
        run = para.add_run(desc)
        set_font(run, size=11)

    doc.add_page_break()

    # Ch7
    add_heading(doc, "第七章  模块五：仲裁全流程", font_size=16)
    add_heading(doc, "7.1 什么时候可以申请仲裁？", font_size=14)
    add_body(doc, "劳动争议申请仲裁的时效期间为一年，从当事人知道或者应当知道其权利被侵害之日起计算。")

    add_heading(doc, "7.2 仲裁流程", font_size=14)
    for i, step in enumerate([
        "准备材料 - 劳动合同、证据清单、仲裁申请书",
        "提交申请 - 向用人单位所在地或合同履行地的劳动仲裁委申请",
        "受理审查 - 5个工作日内决定是否受理",
        "开庭审理 - 双方陈述、举证、辩论",
        "调解优先 - 仲裁员会先调解，调解不成则裁决",
        "领取裁决 - 15日内可向法院起诉，否则生效"
    ], 1):
        para = doc.add_paragraph()
        run = para.add_run(str(i) + ". " + step)
        set_font(run, size=11)
        para.paragraph_format.space_after = Pt(8)

    add_heading(doc, "7.3 仲裁申请书应包含", font_size=14)
    for item in [
        "申请人（你的名字、联系方式、地址）",
        "被申请人（公司名称、地址、法定代表人）",
        "仲裁请求（如：支付经济补偿金XXX元）",
        "事实与理由（简述事情经过）",
        "证据清单（列出提交的所有证据）"
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # Ch8
    add_heading(doc, "第八章  模块六：谈判策略", font_size=16)
    add_heading(doc, "8.1 谈判前的准备", font_size=14)
    for item in [
        "算出自己的底线（最少接受多少）",
        "算出对方的上限（公司最多能给多少）",
        "准备好所有证据的复印件",
        "想好退路（如果谈不成，是仲裁还是继续谈）"
    ]:
        add_number(doc, item)

    add_heading(doc, "8.2 谈判话术示范", font_size=14)
    for situation, dialogue, tip in [
        ("开场表态", "我理解公司的难处，但我也需要保障自己的权益。请给我看一下书面的补偿方案。", "中立姿态，不示弱不示强"),
        ("讨论金额", "按照我的计算，应得补偿是X元。如果公司有诚意，我们可以协商。", "有根有据，给对方台阶"),
        ("陷入僵局", "这个方案我确实无法接受。如果无法达成一致，我保留通过法律途径解决的权利。", "软中带硬，表明立场"),
        ("达成意向", "我基本接受，但有几个条款需要明确。确认无误后我可以当场签字。", "争取最后权益，不轻易松口")
    ]:
        para = doc.add_paragraph()
        run = para.add_run("【" + situation + "】")
        set_font(run, size=12, bold=True)
        para = doc.add_paragraph()
        run = para.add_run("你说：" + dialogue)
        set_font(run, size=11, bold=True)
        para = doc.add_paragraph()
        run = para.add_run("要点：" + tip)
        set_font(run, size=11)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Ch9
    add_heading(doc, "第九章  行动清单：裁员谈判前48小时必查", font_size=16)
    add_body(doc, "拿到裁员通知后，按以下清单逐项检查：")

    checklist = [
        ("[ ]", "确认拿到书面裁员通知，拍照存档"),
        ("[ ]", "查看劳动合同，明确合同期限、岗位、薪资"),
        ("[ ]", "计算应得补偿金（N/N+1/2N），列出清单"),
        ("[ ]", "整理12个月工资流水，计算月均工资"),
        ("[ ]", "统计未休年假天数"),
        ("[ ]", "收集社保、公积金缴费记录"),
        ("[ ]", "整理加班记录、绩效考核记录"),
        ("[ ]", "截图保存与HR的所有沟通记录"),
        ("[ ]", "如有可能，录音留证"),
        ("[ ]", "准备3份证据复印件"),
        ("[ ]", "明确自己的底线金额"),
        ("[ ]", "查询当地劳动仲裁委地址和电话"),
        ("[ ]", "准备好仲裁申请书模板"),
        ("[ ]", "如有需要，准备居住证等异地仲裁材料")
    ]

    table = doc.add_table(rows=len(checklist)+1, cols=2)
    table.style = 'Table Grid'
    table_row(table, ["状态", "行动项"], True, 11)
    for status, item in checklist:
        row = table.add_row()
        row.cells[0].text = status
        row.cells[1].text = item
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=11)

    doc.add_page_break()

    # Ch10
    add_heading(doc, "第十章  资源包", font_size=16)
    add_heading(doc, "常用联系方式", font_size=14)

    resources = [
        ("全国劳动权益保护热线", "12333"),
        ("人力资源社会保障服务热线", "12333"),
        ("工会职工维权热线", "12351"),
        ("法律援助热线", "12348"),
        ("国家人社部官网", "www.mohrss.gov.cn")
    ]
    table = doc.add_table(rows=len(resources)+1, cols=2)
    table.style = 'Table Grid'
    table_row(table, ["资源名称", "联系方式/网址"], True, 11)
    for name, contact in resources:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = contact
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=11)

    add_heading(doc, "常用计算公式速查", font_size=14)
    for formula in [
        "N = 工作年限 x 月工资",
        "N+1 = N + 1个月工资（即时辞退）",
        "2N = 2 x 工作年限 x 月工资（违法解除）",
        "未休年假补偿 = 日工资 x 未休天数 x 3",
        "日工资 = 月工资 / 21.75"
    ]:
        add_bullet(doc, formula)

    add_heading(doc, "填写区域", font_size=14)
    for label in ["我的月薪", "我的工作年限", "应得补偿金（N/N+1/2N）", "我的底线金额", "当地仲裁委地址", "法律援助电话"]:
        add_fill(doc, label)

    doc.save(STUDENT_OUTPUT)
    print("Student handbook: " + STUDENT_OUTPUT)

def create_instructor_handbook():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Cover
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(60)
    run = para.add_run("讲师手册")
    set_font(run, size=24, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("被裁员之后：补偿、仲裁与职场维权全流程")
    set_font(run, size=28, bold=True)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(40)
    run = para.add_run("讲师内部使用 | 禁止外传")
    run.font.color.rgb = RGBColor(180, 0, 0)
    set_font(run, size=14, bold=True)

    doc.add_page_break()

    # TOC
    add_heading(doc, "目 录", font_size=18)
    for item in [
        "第一章  讲师使用说明",
        "第二章  教学设计说明",
        "第三章  互动设计要点",
        "第四章  常见学员问题应对",
        "第五章  时间节奏建议",
        "第六章  评估与反馈",
        "第七章  补充材料索引"
    ]:
        add_bullet(doc, item, 12)

    doc.add_page_break()

    # Ch1
    add_heading(doc, "第一章  讲师使用说明", font_size=16)
    add_heading(doc, "1.1 课程定位", font_size=14)
    add_body(doc, '本课程属于"职场硬技能"类课程，定位为"实用主义法律常识课"。\n不同于纯粹的法律理论课，本课程聚焦于：\n- 可操作的动作（怎么算、怎么谈、怎么仲裁）\n- 可复制的模板（协议审查清单、仲裁申请书模板）\n- 可演练的场景（HR话术破解、谈判模拟）')

    add_heading(doc, "1.2 学员画像", font_size=14)
    for item in [
        "年龄：25-45岁为主",
        "背景：正在经历或可能经历裁员",
        "痛点：信息不对称，不知道自己该拿多少",
        "心态：焦虑、急切想知道答案",
        "基础：法律知识薄弱，但学习意愿强"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 教学目标", font_size=14)
    for title, desc in [
        ("知识目标", "理解N/N+1/2N的适用情形，掌握补偿金计算方法"),
        ("技能目标", "能够识别协议中的陷阱条款，知道如何固定证据"),
        ("态度目标", "建立理性谈判心态，不卑不亢，依法维权")
    ]:
        para = doc.add_paragraph()
        run = para.add_run("【" + title + "】")
        set_font(run, size=11, bold=True)
        run = para.add_run(desc)
        set_font(run, size=11)

    doc.add_page_break()

    # Ch2
    add_heading(doc, "第二章  教学设计说明", font_size=16)

    modules = [
        ("模块一：补偿金计算", "45分钟", "先讲清三种情形，再用实际案例带学员动手算。重点是让学员能用自己的数字套公式。", '很多学员会误以为N是"公司给的恩赐"，要纠正这个认知。'),
        ("模块二：话术套路破解", "60分钟", '先让学员讨论"你遇到过什么套路"，再给出标准话术。角色扮演是核心环节。', "不要只讲理论，一定要让学员开口练。HR的话术只有演过才能记住。"),
        ("模块三：协议审查要点", "45分钟", "发真实的协议样本（脱敏处理），分组讨论找坑。", "三不签红线要反复强调，这是学员最容易踩的坑。"),
        ("模块四：证据保全", "30分钟", "讲解+演示，教会学员怎么截图、录屏、公证。", "很多学员知道要取证，但不知道具体怎么操作，要讲细。"),
        ("模块五：仲裁全流程", "45分钟", "画流程图，讲每一步要准备什么。有条件可以带学员去仲裁委参观。", "时效是1年，这个要反复强调，很多学员以为可以慢慢来。"),
        ("模块六：谈判策略", "60分钟", "谈判演练是核心。设定场景，分组谈判，老师点评。", '不要给学员"一定能谈到理想价格"的错觉。')
    ]

    for name, duration, design, key_point in modules:
        para = doc.add_paragraph()
        run = para.add_run("【" + name + "】" + duration)
        set_font(run, size=12, bold=True)
        para = doc.add_paragraph()
        run = para.add_run("设计意图：" + design)
        set_font(run, size=11)
        para.paragraph_format.left_indent = Pt(20)
        para = doc.add_paragraph()
        run = para.add_run("讲师注意：" + key_point)
        set_font(run, size=11, bold=True)
        para.paragraph_format.left_indent = Pt(20)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Ch3
    add_heading(doc, "第三章  互动设计要点", font_size=16)
    add_heading(doc, "3.1 小组讨论设计", font_size=14)

    discussions = [
        ("破冰讨论", "你或你身边的人被裁员过吗？当时是什么情况？", "10分钟", "让学员打开话匣子，建立连接"),
        ("案例讨论", "给出真实案例（脱敏），讨论：这个补偿方案合理吗？", "15分钟", "培养实战分析能力"),
        ("清单共创", '让大家一起列"还有哪些HR套路"，补充到话术库', "10分钟", "利用集体智慧，增加参与感")
    ]
    table = doc.add_table(rows=len(discussions)+1, cols=4)
    table.style = 'Table Grid'
    table_row(table, ["讨论类型", "讨论主题", "时长", "目的"], True, 10)
    for r in discussions:
        table_row(table, r, False, 10)

    add_heading(doc, "3.2 角色扮演设计", font_size=14)
    add_body(doc, '''场景：HR找你谈话，给你两个选择：
1. 当场签字，补偿N
2. 不签字，公司以"严重违纪"处理

学员扮演被约谈员工，老师扮演HR。

要点：
- 老师要演得真实，给学员压力
- 演完后让学员复盘哪里应对得好，哪里可以改进
- 告诉学员：真实场景比这更复杂，但核心逻辑是一样的''')

    add_heading(doc, "3.3 案例选择原则", font_size=14)
    for p in [
        "真实发生的案例（脱敏处理）",
        "有代表性的案例（不是极端情况）",
        "有讨论空间的案例（答案不唯一）",
        "贴近学员实际的案例（如互联网、金融、制造业等）"
    ]:
        add_bullet(doc, p)

    doc.add_page_break()

    # Ch4
    add_heading(doc, "第四章  常见学员问题应对", font_size=16)

    for q, a in [
        ("公司说给我N已经仁至义尽了，N+1不可能，我怎么办？", "告诉他：N是法定义务，不是恩赐。如果公司是合法裁员，协商解除一般是N+1起步。如果是违法解除，员工有权要求2N。让他回去再考虑。"),
        ("我已经签了字，现在后悔还来得及吗？", "分情况：如果协议存在欺诈、胁迫、重大误解，可以在1年内申请撤销。但举证困难。所以签字前一定要看清楚。"),
        ("我是外包员工，裁员和我有关系吗？", "有关系。劳务派遣员工享有同工同酬权利，被裁员时由用工单位和劳务派遣单位承担连带责任。"),
        ("公司倒闭了，老板跑路，还能要到补偿吗？", "可以向法院申请公司破产清算，优先清偿员工工资和补偿。也可以向劳动监察部门投诉，或申请社保公积金先行支付。"),
        ("仲裁要花多少钱？", "劳动仲裁免费。但如果聘请律师，费用一般是胜诉金额的10%-15%。也可以申请法律援助。"),
        ("我在外地打工，能在老家申请仲裁吗？", "可以在劳动合同履行地或用人单位所在地的仲裁委申请。如果要在老家申请，需要提供居住证等证明。"),
        ("公司故意拖延不给我开离职证明怎么办？", "离职证明是法定义务，公司必须在解除劳动合同时出具。如果公司拒绝，可以向劳动监察部门投诉，也可以要求赔偿因此造成的损失。")
    ]:
        para = doc.add_paragraph()
        run = para.add_run("Q: " + q)
        set_font(run, size=11, bold=True)
        para = doc.add_paragraph()
        run = para.add_run("A: " + a)
        set_font(run, size=11)
        para.paragraph_format.left_indent = Pt(20)
        para.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # Ch5
    add_heading(doc, "第五章  时间节奏建议", font_size=16)
    add_heading(doc, "两天课程安排", font_size=14)

    day1 = [
        ("09:00-09:30", "破冰 + 课程导言", "建立连接，明确学习目标"),
        ("09:30-10:15", "五步维权扫描法", "核心方法论贯串全场"),
        ("10:15-10:30", "茶歇", ""),
        ("10:30-11:15", "模块一：补偿金计算", "动手算自己的数字"),
        ("11:15-12:00", "模块二：话术套路破解（上）", "常见套路讲解"),
        ("12:00-13:30", "午餐", ""),
        ("13:30-14:30", "模块二：话术套路破解（下）", "角色扮演演练"),
        ("14:30-14:45", "茶歇", ""),
        ("14:45-15:30", "模块三：协议审查要点（上）", "三不签红线"),
        ("15:30-16:30", "模块三：协议审查要点（下）", "真实协议找坑练习"),
        ("16:30-17:00", "第一天复盘 + 第二天预告", "")
    ]

    day2 = [
        ("09:00-09:30", "前一天内容回顾", "提问答疑"),
        ("09:30-10:00", "模块四：证据保全", "实操演示"),
        ("10:00-10:15", "茶歇", ""),
        ("10:15-11:00", "模块五：仲裁全流程", "流程图讲解"),
        ("11:00-12:00", "模块六：谈判策略（上）", "策略讲解"),
        ("12:00-13:30", "午餐", ""),
        ("13:30-14:30", "模块六：谈判策略（下）", "谈判演练"),
        ("14:30-14:45", "茶歇", ""),
        ("14:45-15:30", "48小时行动清单", "checklist讲解"),
        ("15:30-16:00", "资源包介绍", "联系方式、工具"),
        ("16:00-16:30", "课程总结 + Q&A", "回顾要点"),
        ("16:30-17:00", "评估与反馈", "填写问卷")
    ]

    for day_name, schedule in [("第一天", day1), ("第二天", day2)]:
        para = doc.add_paragraph()
        run = para.add_run(day_name)
        set_font(run, size=14, bold=True)
        table = doc.add_table(rows=len(schedule)+1, cols=3)
        table.style = 'Table Grid'
        table_row(table, ["时间", "内容", "备注"], True, 10)
        for r in schedule:
            table_row(table, r, False, 10)
        doc.add_paragraph()

    doc.add_page_break()

    # Ch6
    add_heading(doc, "第六章  评估与反馈", font_size=16)
    add_heading(doc, "6.1 学习效果评估", font_size=14)
    for title, desc in [
        ("课前测验", "了解学员基础，调整教学深度"),
        ("随堂练习", "每个模块后的快问快答，确保跟上进度"),
        ("角色扮演评分", "根据学员表现给予即时反馈"),
        ("课后作业", "计算自己的补偿金、审查一份协议样本"),
        ("结业考核", "场景测试：给定情况，让学员写出维权方案")
    ]:
        para = doc.add_paragraph()
        run = para.add_run("【" + title + "】")
        set_font(run, size=11, bold=True)
        run = para.add_run(desc)
        set_font(run, size=11)

    add_heading(doc, "6.2 课程反馈收集", font_size=14)
    for f in [
        "匿名问卷：最有用/最没用的模块",
        "当场打分：内容、讲解、互动、组织",
        "开放问题：你最想深入了解的是什么",
        "跟进访谈：课后一周电话回访，了解实际应用情况"
    ]:
        add_bullet(doc, f)

    doc.add_page_break()

    # Ch7
    add_heading(doc, "第七章  补充材料索引", font_size=16)
    add_heading(doc, "7.1 案例库", font_size=14)
    add_body(doc, "路径：材料库/案例库/裁员维权案例集/")
    add_body(doc, "包含：互联网、金融、制造业、房地产等行业的真实案例（已脱敏）")

    add_heading(doc, "7.2 练习材料", font_size=14)
    add_body(doc, "路径：材料库/练习/裁员谈判练习材料/")
    add_body(doc, "包含：")
    for item in [
        "补偿金计算练习题（含答案）",
        "协议审查找错练习（含答案）",
        "HR话术应对练习题",
        "仲裁申请书写作练习"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.3 工具模板", font_size=14)
    add_body(doc, "路径：材料库/模板/")
    for item in [
        "补偿金计算器.xlsx",
        "协议审查清单.docx",
        "仲裁申请书模板.docx",
        "证据保全指南.pdf",
        "48小时行动清单.docx"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.4 视频素材", font_size=14)
    for item in [
        "HR谈话标准话术示范（讲师用）",
        "裁员谈判模拟录像（学员参考）",
        "仲裁庭审实录（真实案例，已脱敏）"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7.5 法规参考", font_size=14)
    for item in [
        "《中华人民共和国劳动合同法》",
        "《中华人民共和国劳动争议调解仲裁法》",
        "《工资支付暂行规定》",
        "《职工带薪年休假条例》"
    ]:
        add_bullet(doc, item)

    doc.save(INSTRUCTOR_OUTPUT)
    print("Instructor handbook: " + INSTRUCTOR_OUTPUT)

if __name__ == "__main__":
    create_student_handbook()
    create_instructor_handbook()
    print("Done!")

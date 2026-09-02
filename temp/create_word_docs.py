#!/usr/bin/env python3
"""Create Word documents for Cold War Revisited course"""
import sys
sys.path.insert(0, 'C:/Users/Administrator/.claude/skills/Word文档处理/scripts/python')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = "D:/新课开发/政治学/16_冷战重访-意识形态对抗的政治遗产/Word"

def set_cjk_font(run, font_name='微软雅黑', size=12):
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_cjk_font(run)
        if level == 0:
            run.font.size = Pt(22)
            run.font.bold = True
        elif level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
    return para

def add_para(doc, text, bold=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_cjk_font(run, size=size)
    run.font.bold = bold
    para.paragraph_format.space_after = Pt(6)
    return para

def create_student_handbook():
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_heading(doc, '冷战重访——意识形态对抗的政治遗产', level=0)
    add_heading(doc, '学员手册', level=1)

    add_para(doc, '课程代码：RPD-16', size=10)
    add_para(doc, '适用对象：政治学及相关专业学生', size=10)
    add_para(doc, '', size=10)

    # Course Overview
    add_heading(doc, '课程概述', level=1)
    add_para(doc, '本课程以"冷战重访"为核心视角，系统梳理冷战的历史背景、意识形态对抗、代理人战争等核心议题，并将其与当代国际关系中的"新冷战"话语进行对比分析，培养学生的独立思考能力。')

    # Learning Objectives
    add_heading(doc, '学习目标', level=1)
    add_para(doc, '• 掌握冷战的基本概念框架和历史脉络')
    add_para(doc, '• 理解意识形态在冷战对抗中的作用机制')
    add_para(doc, '• 分析代理人战争的历史逻辑和当代启示')
    add_para(doc, '• 识别"冷战思维"的现代表现')
    add_para(doc, '• 培养独立分析国际政治问题的能力')

    # Module Overview
    add_heading(doc, '模块内容', level=1)

    modules = [
        ('模块一：什么是冷战', '冷战的历史背景与定义、两极格局的形成、冷战与热战的关系'),
        ('模块二：意识形态对抗', '意识形态在冷战中的角色、资本主义vs社会主义、意识形态宣传机器'),
        ('模块三：代理人战争', '代理人战争的概念、冷战主要代理人战争、案例分析'),
        ('模块四：冷战思维的现代延续', '冷战思维的延续、新冷战的特征、中美关系分析'),
        ('模块五："新冷战"再思考', '多元视角分析、历史对比、中国学者的观点'),
        ('模块六：独立思考工具', '历史证据评估、利益分析框架、意识形态批判'),
    ]

    for title, content in modules:
        add_heading(doc, title, level=2)
        add_para(doc, content)

    # Teaching Methods
    add_heading(doc, '教学方法', level=1)
    add_para(doc, '本课程采用多元教学方法，包括：')
    add_para(doc, '• 讲授法：系统讲解核心概念和理论框架')
    add_para(doc, '• 案例分析：通过对历史事件的深入分析加深理解')
    add_para(doc, '• 小组讨论：鼓励学生发表观点、交流思想')
    add_para(doc, '• 辩论赛：针对"新冷战"等争议性话题进行辩论')
    add_para(doc, '• Workshop：独立思考工具的实践应用')

    # Assessment
    add_heading(doc, '考核方式', level=1)
    add_para(doc, '• 平时成绩（30%）：课堂参与、小组讨论')
    add_para(doc, '• 案例分析报告（30%）：选择一个冷战或当代国际关系案例进行分析')
    add_para(doc, '• 期末论文（40%）：主题围绕"冷战思维的当代意义"')

    # Resources
    add_heading(doc, '延伸阅读', level=1)
    add_para(doc, '• 《冷战史》——盖迪斯')
    add_para(doc, '• 《意识形态与冷战》——韦瑟斯')
    add_para(doc, '• 《代理人战争》——基恩')
    add_para(doc, '• 《新冷战？》——江忆恩')

    doc.save(f'{OUT_DIR}/学员手册.docx')
    print(f'Created: {OUT_DIR}/学员手册.docx')

def create_instructor_manual():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_heading(doc, '冷战重访——意识形态对抗的政治遗产', level=0)
    add_heading(doc, '讲师手册', level=1)
    add_para(doc, '课程代码：RPD-16 | 总课时：21课时', size=10)

    # Teaching Plan Overview
    add_heading(doc, '教学进度总览', level=1)

    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'

    headers = ['周次', '教学内容', '教学方法', '备注']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            set_cjk_font(run, size=10)

    schedule = [
        ('第1周', '冷战概述与模块一', '讲授+讨论', ''),
        ('第2周', '模块二：意识形态对抗', '案例分析', ''),
        ('第3周', '模块三：代理人战争', '讲授+影片', ''),
        ('第4周', '模块四：冷战思维延续', '小组讨论', ''),
        ('第5周', '模块五：新冷战思考', '辩论赛', ''),
        ('第6周', '模块六+总结', 'Workshop', ''),
        ('第7周', '复习与答疑', '自由讨论', ''),
    ]

    for row_idx, (week, content, method, note) in enumerate(schedule, start=1):
        table.rows[row_idx].cells[0].text = week
        table.rows[row_idx].cells[1].text = content
        table.rows[row_idx].cells[2].text = method
        table.rows[row_idx].cells[3].text = note

    doc.add_paragraph()

    # Detailed teaching guidance for each module
    add_heading(doc, '各模块教学要点', level=1)

    module_guides = [
        ('模块一：什么是冷战（3课时）',
         '重点：冷战的定义与本质、两极格局的形成机制、冷战与"热战"的区别。'
         '建议：使用时间轴梳理冷战起源，引用丘吉尔铁幕演说等一手材料。'
         '讨论问题：为什么冷战没有演变成全面热战？'),
        ('模块二：意识形态对抗（4课时）',
         '重点：意识形态如何塑造冷战对抗、意识形态与实际利益的关系。'
         '建议：对比美苏两国的意识形态宣传策略，分析颜色革命等现代案例。'
         '讨论问题：意识形态是否仍然是当代国际关系的主要驱动力？'),
        ('模块三：代理人战争（5课时）',
         '重点：代理人战争的逻辑、朝鲜战争与越南战争的分析、冷战热战化的风险。'
         '建议：使用地图展示代理人战争地理分布，播放相关纪录片片段。'
         '案例选择：朝鲜战争（1950-1953）、越南战争（1955-1975）。'),
        ('模块四：冷战思维的现代延续（4课时）',
         '重点：冷战思维的识别、新冷战的特征、科技竞争的角色。'
         '建议：对比分析中美贸易战与冷战时期的美苏竞争。'
         '讨论问题：5G、AI等新技术竞争是否具有冷战特征？'),
        ('模块五："新冷战"再思考（3课时）',
         '重点：避免简单历史类比、多元视角分析、中国学者的独特视角。'
         '建议：组织辩论赛，让学生分别从美国、中国、第三世界国家角度分析。'
         '辩论题：当前国际关系是否已经进入"新冷战"？'),
        ('模块六：独立思考工具（2课时）',
         '重点：历史证据评估、利益分析框架、意识形态批判方法。'
         '建议：通过实际案例演练独立思考工具，培养批判性思维。'
         '工具：冷战诊断清单、利益-意识形态分析矩阵。'),
    ]

    for title, content in module_guides:
        add_heading(doc, title, level=2)
        add_para(doc, content)

    # Assessment criteria
    add_heading(doc, '考核标准', level=1)
    add_para(doc, '案例分析报告评分标准：')
    add_para(doc, '• 概念准确性（25%）：正确运用冷战相关概念')
    add_para(doc, '• 分析深度（35%）：能否深入挖掘事件背后的逻辑')
    add_para(doc, '• 证据运用（20%）：史料引用是否充分、可靠')
    add_para(doc, '• 独立思考（20%）：是否有独特见解，避免简单复述')

    doc.save(f'{OUT_DIR}/讲师手册.docx')
    print(f'Created: {OUT_DIR}/讲师手册.docx')

def create_teaching_schedule():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_heading(doc, '教学进度表', level=0)
    add_para(doc, '冷战重访——意识形态对抗的政治遗产', size=11)
    add_para(doc, '课程代码：RPD-16 | 总课时：21课时（6周）', size=10)
    add_para(doc, '', size=10)

    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'

    headers = ['周次', '模块', '教学内容', '教学方法', '学习目标']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            set_cjk_font(run, size=10)

    schedule = [
        ('1', '模块一', '冷战概述、冷战定义、两极格局', '讲授+讨论', '掌握基本概念'),
        ('2', '模块二', '意识形态对抗、资本主义vs社会主义', '案例分析', '理解意识形态角色'),
        ('3', '模块三', '代理人战争、朝鲜战争、越南战争', '讲授+影片', '理解代理人战争逻辑'),
        ('4', '模块四', '冷战思维延续、新冷战特征、中美关系', '小组讨论', '识别冷战思维'),
        ('5', '模块五', '新冷战多元视角、历史对比', '辩论赛', '培养批判思维'),
        ('6', '模块六', '独立思考工具、冷战诊断清单', 'Workshop', '建立分析框架'),
        ('7', '复习', '总复习、答疑', '自由讨论', '巩固所学'),
    ]

    for row_idx, row_data in enumerate(schedule, start=1):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_text
            for run in table.rows[row_idx].cells[col_idx].paragraphs[0].runs:
                set_cjk_font(run, size=10)

    doc.save(f'{OUT_DIR}/教学进度表.docx')
    print(f'Created: {OUT_DIR}/教学进度表.docx')

if __name__ == '__main__':
    import os
    os.makedirs(OUT_DIR, exist_ok=True)
    create_student_handbook()
    create_instructor_manual()
    create_teaching_schedule()
    print('All Word documents created successfully!')

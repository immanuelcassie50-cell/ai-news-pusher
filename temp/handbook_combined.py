# Combined script
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

def sf(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    try:
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')
    except:
        pass

def set_cell(cell, text, bold=False, size=11):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            sf(r, size, bold)

def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name = '微软雅黑'
    return p

def add_p(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sf(r, size, bold)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    sf(r)
    p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    return p

OUT = 'D:/新课开发/企业大学/对内/7.高管学习项目重构：AI辅助的场景推演与决策模拟设计/讲师手册/讲师手册_高管学习项目重构.docx'

# ============ COVER ============
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('高管学习项目重构')
sf(r, 28, True)
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('AI辅助的场景推演与决策模拟设计')
sf(r2, 20, True)
doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('讲师手册')
sf(r3, 24, True)
doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('FACILITATOR GUIDE')
sf(r4, 14)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
t5 = doc.add_paragraph()
t5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = t5.add_run('本手册为内部培训使用材料，请勿对外传播')
sf(r5, 10)
doc.add_page_break()

# ============ TABLE OF CONTENTS ============
add_h(doc, '目录', 1)
toc = [
    '第一章  讲师资质要求与课前准备',
    '第二章  每个模块的详细授课指引',
    '第三章  时间管理要点与超时处理',
    '第四章  学员常见问题应对话术',
    '第五章  小组讨论引导技巧',
    '第六章  评估标准与评分细则',
    '第七章  风险预案与替代方案',
    '第八章  课后跟进建议',
    '附录一  核心工具速查表',
    '附录二  场景卡使用指南',
]
for item in toc:
    p = doc.add_paragraph()
    r = p.add_run(item)
    sf(r, 12)
doc.add_page_break()

# ============ CHAPTER 1 ============
add_h(doc, '第一章  讲师资质要求与课前准备', 1)

add_h(doc, '1.1 讲师资质要求', 2)
add_p(doc, '本课程对讲师有较高的专业要求，不仅需要具备传统的培训授课能力，还需要对AI工具和高管决策场景有深刻理解。')

add_h(doc, '1.1.1 硬性资质要求', 3)
for item in [
    '具有5年以上企业培训或OD/TD相关工作经验',
    '主导过至少3个高管培训项目的设计与交付',
    '熟悉主流AI工具（ChatGPT/Claude/MiniMax等）的使用方法',
    '具备商业战略、组织变革或决策心理学相关知识背景',
    '有场景式教学或案例教学经验者优先',
]:
    add_bullet(doc, item)

add_h(doc, '1.1.2 软性能力要求', 3)
soft = [
    ('控场能力', '能够应对高管学员的质疑和挑战，引导讨论方向不偏离主题'),
    ('商业敏感度', '能快速理解学员所在行业的商业逻辑，提出有深度的问题'),
    ('AI素养', '不必须是AI技术专家，但需要理解AI的能力边界和适用场景'),
    ('学习敏锐度', '愿意持续学习新工具、新案例，保持课程内容更新'),
    ('EQ与格局', '面对高管学员时不怯场，也不炫耀，能在平等对话中传递价值'),
]
for title, desc in soft:
    p = doc.add_paragraph()
    r = p.add_run(title + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '1.1.3 讲师分级标准', 3)
tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'
hdrs = ['级别', '认证要求', '可授课范围', '年度复训']
for i, h in enumerate(hdrs):
    set_cell(tbl.cell(0, i), h, bold=True)
levels = [
    ['助理讲师', '完成认证培训+试讲', '半天工作坊', '必须'],
    ['认证讲师', '3次以上授课+案例评审', '标准版/压缩版', '必须'],
    ['高级讲师', '10次以上+课程迭代贡献', '全部版本+企业内训', '建议'],
]
for ri, row in enumerate(levels):
    for ci, val in enumerate(row):
        set_cell(tbl.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '1.2 课前准备清单', 2)
add_p(doc, '请在课前7天、3天、1天三个节点分别完成以下准备工作：')

add_h(doc, '1.2.1 课前7天：物料与学员信息确认', 3)
for item in [
    '确认学员名单，了解学员背景（行业、岗位、职级、工作年限）',
    '发送课前调研问卷，收集学员对课程的期待和已有经验',
    '确认场地布置：岛型分组座位（4-6人/组）、白板/翻页纸、投影设备',
    '确认AI工具可用性：ChatGPT Plus/Claude Pro账号登录测试',
    '打印学员手册、工作纸、场景卡等材料（建议多印10%备用）',
    '准备助教：至少1名助教负责计时、记录、分发材料',
    '如使用专业模拟平台，提前完成账号开通和权限配置',
]:
    add_bullet(doc, item)

add_h(doc, '1.2.2 课前3天：内容熟悉度准备', 3)
for item in [
    '熟读本课程所有7个模块的教学文档',
    '准备自己在每个模块的2-3个亲身经历案例（或行业案例）',
    '预演关键练习的示范答案，能在课堂上展示第一级和第二级示范',
    '熟悉附录中的参考答案，准备好应对学员的不同答案',
    '确认每个模块的核心工具能现场演示',
    '准备AI演示用的Prompt模板，现场能实时生成情境',
]:
    add_bullet(doc, item)

add_h(doc, '1.2.3 课前1天：细节确认与应急准备', 3)
for item in [
    '再次确认学员人数和分组安排',
    '检查投影仪、音响、白板笔等设备',
    '将课程PPT、案例文档、工作纸全部存入U盘备份（避免网络问题）',
    '准备Plan B：如果AI工具不可用，改为纯讨论模式',
    '确认茶歇、餐饮、住宿（如果外地）等后勤安排',
    '在教室显眼位置张贴"手机静音"提示',
    '提前30分钟到达教室，调试设备，熟悉环境',
]:
    add_bullet(doc, item)

add_h(doc, '1.3 学员信息分析表', 2)
add_p(doc, '讲师应在课前完成学员信息分析，以便调整授课重点：')
tbl2 = doc.add_table(rows=6, cols=3)
tbl2.style = 'Table Grid'
for i, h in enumerate(['分析维度', '需要了解的信息', '调整依据']):
    set_cell(tbl2.cell(0, i), h, bold=True)
info = [
    ['行业分布', '学员主要来自哪些行业', '选择行业相关的案例和场景'],
    ['职级构成', 'CEO/COO/CFO/部门负责人比例', '调整案例的复杂度和决策层级'],
    ['AI熟悉度', '学员对AI工具的了解程度', '调整AI工具讲解的深度'],
    ['培训经历', '过去参加过哪些高管培训', '避免重复，讲新角度'],
    ['学习期待', '学员最想解决什么问题', '重点回应，增加相关内容'],
]
for ri, row in enumerate(info):
    for ci, val in enumerate(row):
        set_cell(tbl2.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '1.4 教室布置标准', 2)
add_p(doc, '教室布置直接影响学员的参与度和学习效果。以下是标准布置要求：')
seating = [
    ('座位安排', '岛型分组（4-6人/组）优于剧院式座位，每组配备彩色马克笔+白板纸'),
    ('投影位置', '屏幕应在教室前方居中，从所有位置都能清晰看到'),
    ('讲师站位', '讲师不应孤立在讲台上，应能方便走到学员之间'),
    ('白板使用', '每组附近应有白板或翻页纸，用于现场板书和练习展示'),
    ('茶歇区', '茶歇区应与学习区有明显边界，避免干扰'),
    ('网络要求', '确认WiFi可用且稳定，最好准备移动热点作为备份'),
]
for title, desc in seating:
    p = doc.add_paragraph()
    r = p.add_run(title + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)
doc.add_page_break()
# ============ CHAPTER 2 ============
add_h(doc, '第二章  每个模块的详细授课指引', 1)

# 2.1 开篇
add_h(doc, '2.1 开篇：高管学习的本质挑战', 2)
add_p(doc, '建议时长：50—60分钟（标准版）/ 30分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能识别一个现有高管培训项目的3个结构性缺陷',
    '学员能用语言解释"情境学习理论"的核心理念',
    '学员能提出至少2个AI可能重构高管学习项目的具体方向',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '开场不要自我介绍"我是谁我讲什么"。直接呈现一个令人不安的数据或故事，让学员先在头脑里形成疑问——"为什么花了800万，高管觉得没用？"然后再回答这个问题。')
add_p(doc, '推荐开场案例：', bold=True)
p = doc.add_paragraph()
r = p.add_run('某互联网公司连续3年投入800万做高管培训，85%学员认为"内容好有启发"，但CEO在复盘会上说："我不知道这800万换来了什么，战略执行还是老样子。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '传统项目缺陷诊断表（第一部分）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：带着学员逐行填表，重点是"AI可重构程度"这一列——让学员不只看到问题，而是看到AI能做什么。避免让学员只是抱怨过去的培训，要引导到"如果用AI重构，会怎样"。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '练习引导说明', 3)
add_p(doc, '练习一：传统项目缺陷诊断（第一级）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导方式：这道练习的难点在于第2列"具体缺陷描述"——学员容易写泛泛而谈的"内容不实用"。重点引导："不要写内容不实用，要写哪个案例、哪句话让你觉得不实用。情境真实性的具体缺陷，要精确到案例都是别人的公司，无法建立模式识别能力。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '过渡与衔接', 3)
add_p(doc, '开篇模块结束后，学员最常见的反应是"确实是这样，但怎么改？"——这制造了良好的学习胃口。')
add_p(doc, '过渡到模块一时：可以说"要重构高管学习项目，首先要搞清楚AI时代需要什么样的高管能力——不是更高级的员工，而是能应对不确定性的决策者。"')

add_h(doc, '讲师注意事项', 3)
for item in [
    '有学员可能开始就挑战"AI介入培训是否合适"——这是好的信号，不要回避，而是邀请他详细说明担忧，引导全班讨论。',
    '开篇破冰环节要谨慎把控时间，建议8分钟以内，避免讨论发散。',
    '案例卡1的三个缺陷是讲师的核心弹药，要熟记，能随时引用。',
]:
    p = doc.add_paragraph()
    r = p.add_run(item)
    sf(r)
    p.paragraph_format.left_indent = Inches(0.3)

# 2.2 模块一
add_h(doc, '2.2 模块一：AI时代高管能力模型重构', 2)
add_p(doc, '建议时长：65分钟（标准版）/ 40分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能对比说明"传统高管能力模型"与"AI时代高管能力模型"的3个核心差异',
    '学员能为目标学员画像完成1份"能力-情境对照表"',
    '学员能识别AI在2个典型高管决策场景中的具体介入方式',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '带学员读AI时代高管能力三棱锥模型（顶点：判断力；中层：情境建模、决策整合；基础：人机协同、伦理边界）。然后问："你们公司现在的高管能力模型，是哪个版本的？"——让学员对比反思。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, 'AI时代高管能力三棱锥（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：这个模型是整个课程的元框架之一。重点强调"判断力"作为顶点的意义——不是"知识"、不是"经验"，而是"在信息不完整时做出可辩护决策的能力"。可以用一个思想实验：如果AI明天可以完成你80%的工作，你作为高管还剩下什么不可替代的价值？')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, '能力-情境对照表（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：先完整呈现示例表，然后让学员基于自己企业的高管项目填写。重点是"AI介入方式"这一列——学员通常写得很模糊，要逼他们写出"AI在哪个决策节点介入、输出什么"。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '练习引导说明', 3)
add_p(doc, '练习一：能力-情境对照表（第二级）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导方式：给15分钟独立填写，然后让学员在小组里分享。分享时引导："重点听组员的AI介入方式，看看有没有让你眼前一亮的设计。"每组推荐1个最完整的表格在全班分享。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.3 模块二
add_h(doc, '2.3 模块二：场景推演的核心原理', 2)
add_p(doc, '建议时长：75分钟（标准版）/ 45分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能用语言解释"情境重构三棱镜"的3个构成要素',
    '学员能识别5种常见推演设计的缺陷',
    '学员能运用三棱镜框架对一个真实高管决策场景进行要素拆解',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '带学员做一个小测试："在座各位，过去一个月做过的最大决策是什么？那个决策你现在能复盘吗？"——让学员意识到，真实的决策反馈周期太长，几乎不可能复盘。引出"场景推演"的价值：在压缩时间内体验完整决策链。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '情境重构三棱镜（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：带着学员逐行读三棱镜的三个构成要素（情境变量、决策节点、反馈机制），重点强调"反馈机制"——这是最容易被忽视但最关键的要素。没有即时反馈的推演，只是在玩游戏，不是在学习。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, '五种常见推演设计缺陷（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：让学员先尝试自己识别bug（工作纸2-a），然后再揭示答案。这个顺序比直接讲效果更好——学员带着问题来听讲，印象更深刻。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '练习引导说明', 3)
add_p(doc, '练习二：三棱镜拆解（第二级）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导方式：给20分钟完成小组场景的拆解。完成后，让每组用2分钟汇报："你们选了什么场景？三个要素各是什么？"讲师即时点评。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.4 模块三
add_h(doc, '2.4 模块三：AI辅助决策模拟的设计方法', 2)
add_p(doc, '建议时长：90分钟（标准版）/ 50分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能完整复述"决策推演五步法"的每一步操作',
    '学员能在AI辅助下完成五步法中至少3步的具体设计',
    '学员能针对一个具体高管学习项目完成五步法的完整设计',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '直接问："各位过去用AI做过什么？生成文案？做PPT？"——让学员发现，AI在高管学习领域的应用几乎没有人探索过。然后展示一个AI辅助情境初始化的现场演示。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '决策推演五步法（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：五步法是本课程的核心操作方法论。讲解时用"战略并购"案例完整走一遍五步：①AI生成并购情境数据包；②AI模拟各方利益相关者反应；③AI根据选择实时扩展情境；④AI多维反馈（财务+组织+市场）；⑤AI辅助生成个性化复盘报告。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, 'AI辅助情境初始化演示（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：这是整个课程最重要的演示环节。讲师应提前准备一个情境Prompt模板，现场实时生成一个商业情境。生成后，让学员点评："这个情境真实吗？哪里可以改进？"——让学员参与Prompt的优化，建立对AI能力的直观认知。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.5 模块四
add_h(doc, '2.5 模块四：高仿真商业情境的构建技术', 2)
add_p(doc, '建议时长：60分钟（标准版）/ 35分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能说明"情境参数矩阵"中至少6个关键参数及其取值范围',
    '学员能针对一个给定的高管决策场景设计出合理的参数配置',
    '学员能识别3种常见的情境失真问题及其来源',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '问学员一个问题："你们公司现在的案例教学，用的情境是你们自己公司的吗？"——通常答案是否定的。然后问："为什么不是？"——引导学员发现：真实案例有商业敏感性，公开案例不够真实。AI可以解决这个矛盾。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '情境参数矩阵（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：带着学员逐行读情境参数矩阵的6个参数（信息完整度、数据真实性、决策时间窗口、结果反馈周期、角色数量、利益冲突强度）。重点强调"利益冲突强度"——最常见的设计失误是把情境设计成"利益高度一致"，真实商业情境中，利益冲突是常态。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, '三种常见失真问题（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：用学员自己经历过的"不真实"案例来解释失真类型。比如："有没有人经历过培训里的案例太完美的情况？"——让学员自己举例，然后对号入座到失真分类。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.6 模块五
add_h(doc, '2.6 模块五：AI决策模拟工具链与平台选择', 2)
add_p(doc, '建议时长：60分钟（标准版）/ 30分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能对比评估至少4类AI辅助推演工具的优劣势',
    '学员能运用"工具选择决策树"为自己的项目选择合适的工具组合',
    '学员能完成至少1个AI工具的实操练习',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '开场展示一个令人惊讶的对比：用同一个情境，分别用"糟糕的Prompt"和"优秀的Prompt"让AI生成结果。让学员看到Prompt质量的差异直接导致输出质量的差异。引出："工具本身不重要，用工具的方式才重要。"')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '四类工具对比（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：通用LLM（ChatGPT/Claude）、专用模拟平台（Simon等）、数据可视化工具（Tableau）、协作平台（Miro/Notion）——每类工具用一句话说清"最适合谁、不适合谁"。重点不是记住每个工具，而是理解"工具选择的逻辑"。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, 'Prompt工程四大原则（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：这是模块五最重要的内容。角色注入要具体、情境构建要分层、反馈设计要"挑刺"、输出结构要可追踪。每一个原则都用正反对比例子说明。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.7 模块六
add_h(doc, '2.7 模块六：高管理学习项目的评估体系设计', 2)
add_p(doc, '建议时长：60分钟（标准版）/ 35分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能说明"学习迁移对勾模型"的两个阶段和4个关键转化点',
    '学员能为AI辅助的高管学习项目设计完整的评估方案',
    '学员能使用"迁移效果评估矩阵"评估一个已有项目的迁移效果',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '问学员："你们公司上一个高管培训项目，是怎么评估效果的？"——通常答案只有"满意度问卷"。然后展示一个令人不安的数据："满意度高但行为改变率不足15%"——引出评估体系的缺失。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '学习迁移对勾模型（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：对勾底部（知识获取）到顶部（习惯改变）的转化率逐级下降：60%→40%→20%。重点是解释"为什么"——每个断层的原因是什么，如何在设计阶段就预防这些断层。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, '双轨评估框架（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：轨道A（学习过程评估）和轨道B（决策质量评估）并行。重点强调轨道B更难做但更有价值——"即时满意度高不代表学习效果好"，这是反直觉但真实的结论。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.8 模块七
add_h(doc, '2.8 模块七：企业落地路径与常见陷阱', 2)
add_p(doc, '建议时长：50分钟（标准版）/ 30分钟（压缩版）', bold=True)

add_h(doc, '模块目标', 3)
for item in [
    '学员能描述企业落地AI辅助高管学习项目的3个典型路径',
    '学员能识别至少5个常见落地陷阱并给出预防方案',
    '学员能为自己企业制定一份简化的落地路线图',
]:
    add_bullet(doc, item)

add_h(doc, '开场引导要点', 3)
add_p(doc, '问学员："在座各位，有多少人已经开始尝试用AI做高管培训了？"——通常只有少数人举手。然后说："先别急，在开始之前，我们需要知道7种会让你失败的陷阱。"——制造悬念，引导学员认真听讲。')

add_h(doc, '核心工具讲解要点', 3)
add_p(doc, '三条落地路径（第一章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：探索期（3-6个月/<10万）、试点期（6-12个月/10-30万）、扩展期（12-24个月/30万+）。重点不是让学员记住数字，而是让他们判断"你们公司现在在哪一步"。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_p(doc, '5个常见陷阱（第二章）', bold=True)
p = doc.add_paragraph()
r = p.add_run('引导要点：工具先行场景后置、过度依赖AI忽视人的判断、评估体系缺失、情境过于简化、高管期望错位。每个陷阱用一句话总结——让学员能带走的不是理论，是"避坑指南"。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

# 2.9 综合演战
add_h(doc, '2.9 综合演战：设计一个真实高管学习项目方案', 2)
add_p(doc, '建议时长：120分钟（标准版）/ 60分钟（压缩版）', bold=True)

add_h(doc, '演战目标', 3)
for item in [
    '学员能以小组为单位完成一个完整的AI辅助高管学习项目方案设计',
    '方案涵盖：目标学员定位、核心决策场景、五步法推演设计、情境参数配置、工具链选择、评估体系、落地路径',
]:
    add_bullet(doc, item)

add_h(doc, '活动序列', 3)
war_seq = [
    ('C1 项目选题', '10分钟', '选题说明+小组分工'),
    ('C2 需求分析', '20分钟', '目标学员定位+场景选择'),
    ('C3 情境设计', '30分钟', '三棱镜分析+参数配置'),
    ('C4 推演设计', '30分钟', '五步法完整设计'),
    ('C5 工具与评估', '20分钟', '工具选型+评估体系'),
    ('C6 整合汇报', '20分钟', '方案整合+汇报准备'),
    ('C7 小组汇报', '15分钟/组', '每组10分钟汇报+5分钟点评'),
]
for act, t, desc in war_seq:
    p = doc.add_paragraph()
    r = p.add_run(act + '（' + t + '）：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '汇报点评要点', 3)
add_p(doc, '每组汇报后，讲师从以下5个维度进行点评：')
eval_dims = [
    ('情境真实性（20%）', '三棱镜是否完整？参数配置是否合理有据？'),
    ('五步法完整性（20%）', '五步是否完整？AI介入方式是否具体可操作？'),
    ('评估体系（20%）', '三层次是否完整？指标是否具体可测量？'),
    ('工具选型（15%）', '选择是否有依据？组合是否合理？'),
    ('落地可行性（15%）', '路径是否清晰？里程碑是否可实现？'),
]
for dim, desc in eval_dims:
    p = doc.add_paragraph()
    r = p.add_run(dim + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '讲师注意事项', 3)
for item in [
    '汇报环节超时是常态。提前告知："每组严格10分钟，超时1分钟扣1分。"——提前立规矩比现场求情有效。',
    '如果某组方案明显不完整，不要当场批评，而是问："如果再给一周完善，你们觉得最需要改的是什么？"——给建设性反馈。',
    '全部汇报结束后，留10分钟做整体点评：共性优点、共性不足、后续建议。不要只点名问题，要给出路。',
]:
    p = doc.add_paragraph()
    r = p.add_run(item)
    sf(r)
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_page_break()
# ============ CHAPTER 3 ============
add_h(doc, '第三章  时间管理要点与超时处理', 1)

add_h(doc, '3.1 标准课程时间表', 2)
add_p(doc, '第一天时间表（标准版）', bold=True)

day1 = doc.add_table(rows=10, cols=4)
day1.style = 'Table Grid'
for i, h in enumerate(['时间段', '内容', '时长', '备注']):
    set_cell(day1.cell(0, i), h, bold=True)
d1 = [
    ['08:30-09:00', '学员签到+设备调试', '30分钟', '提前准备AI工具'],
    ['09:00-09:50', '开篇：高管学习的本质挑战', '50分钟', '含破冰'],
    ['09:50-10:10', '茶歇', '20分钟', ''],
    ['10:10-11:15', '模块一：AI时代高管能力模型', '65分钟', '含练习'],
    ['11:15-12:00', '模块二：场景推演的核心原理', '45分钟', '压缩后可调'],
    ['12:00-13:30', '午餐', '90分钟', ''],
    ['13:30-15:00', '模块三：AI辅助决策模拟设计', '90分钟', '含AI演示'],
    ['15:00-15:15', '茶歇', '15分钟', ''],
    ['15:15-16:15', '模块四：高仿真情境构建', '60分钟', '含实操'],
]
for ri, row in enumerate(d1):
    for ci, val in enumerate(row):
        set_cell(day1.cell(ri+1, ci), val)
doc.add_paragraph()

add_p(doc, '第二天时间表（标准版）', bold=True)
day2 = doc.add_table(rows=9, cols=4)
day2.style = 'Table Grid'
for i, h in enumerate(['时间段', '内容', '时长', '备注']):
    set_cell(day2.cell(0, i), h, bold=True)
d2 = [
    ['09:00-09:10', '回顾与开场', '10分钟', '快速回顾第一天'],
    ['09:10-10:10', '模块五：AI决策模拟工具链', '60分钟', '含实操'],
    ['10:10-10:25', '茶歇', '15分钟', ''],
    ['10:25-11:25', '模块六：评估体系设计', '60分钟', '含ROI讨论'],
    ['11:25-12:15', '模块七：企业落地路径', '50分钟', '含路线图'],
    ['12:15-13:30', '午餐', '75分钟', ''],
    ['13:30-15:30', '综合演战', '120分钟', '小组方案设计'],
    ['15:30-16:30', '汇报+点评+收尾', '60分钟', '含课程总结'],
]
for ri, row in enumerate(d2):
    for ci, val in enumerate(row):
        set_cell(day2.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '3.2 容易超时的环节及处理', 2)
overtime = [
    ('破冰讨论（标准8分钟）', '风险：容易引发长篇故事。', '处理：严格计时，超时立即喊停，转入下一环节'),
    ('模块一练习（标准15分钟）', '风险：填表质量参差。', '处理：给10分钟时提醒一次，到15分钟必须收'),
    ('模块二三棱镜练习（标准20分钟）', '风险：情境选择分歧，花太多时间。', '处理：提前声明"选择比分析更重要"，情境不合适直接建议更换'),
    ('模块三AI演示（标准15分钟）', '风险：AI输出质量不稳定，需要反复调试。', '处理：提前测试3遍，准备备用Prompt，如果演示失败立即切视频'),
    ('综合演战汇报（标准10分钟/组）', '风险：学员容易超时，尤其是最后一组。', '处理：提前告知规则，超时扣分；最后一组如超时，可适当宽容但要点评到位'),
]
for act, risk, solution in overtime:
    p = doc.add_paragraph()
    r = p.add_run(act)
    sf(r, bold=True)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('风险：' + risk)
    sf(r2)
    p3 = doc.add_paragraph()
    r3 = p3.add_run('处理：' + solution)
    sf(r3, bold=True)

add_h(doc, '3.3 压缩版时间调整策略', 2)
add_p(doc, '当课程被压缩至1天（6小时）或半天（3小时）时，按以下优先级调整：')
compress = [
    ('必须保留', '开篇（30分钟）+ 模块二（三棱镜原理，45分钟）+ 模块三（五步法，50分钟）+ 综合演战（60分钟）'),
    ('可压缩', '模块一（压缩至20分钟，只讲能力模型框架）、模块五（压缩至20分钟，只讲Prompt原则）'),
    ('可跳过', '模块四（技术细节，学员课后自学）、模块六ROI计算（学员反馈不强烈时可跳过）'),
]
for level, desc in compress:
    p = doc.add_paragraph()
    r = p.add_run(level + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

doc.add_page_break()

# ============ CHAPTER 4 ============
add_h(doc, '第四章  学员常见问题应对话术', 1)
add_p(doc, '以下是课程中高频出现的问题及参考话术。讲师应根据实际情况调整，不要死记硬背——真诚比话术更重要。')

add_h(doc, '4.1 "AI生成的情境不够真实，感觉像在演戏"', 2)
add_p(doc, '背后原因：学员可能质疑AI的情境生成质量，或者在为自己的不投入找借口。')
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"这个担心很真实。AI确实不能替代真实经验，但这恰恰是AI的价值所在——它让你在安全环境里死无数次，而不需要在真实战场上付出代价。你有没有见过一个飞行员在第一次飞真飞机之前，只靠看手册就上天了？"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.2 "我们公司高管觉得AI不可靠，不愿意用"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"这可能不是AI可不可靠的问题，而是AI扮演什么角色的问题。我们不是用AI替代高管做决策，而是用AI做一个高级陪练——就像运动员有陪练，但上场比赛的还是运动员自己。高管担心的可能是AI会抢他们的工作，这个担忧需要被正视，但可以通过明确定位来化解。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.3 "我们公司的数据敏感，不能用真实数据"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"这是非常常见的合规问题。有4种解决路径：①用行业公开数据+假设企业背景；②用历史数据的脱敏版本；③用AI合成数据；④用通用情境。具体用哪种，取决于你们的数据敏感程度。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.4 "高管时间很宝贵，怎么保证参与度"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"这是整个项目成败最关键的因素。几个关键点：①时间控制在2小时以内；②情境必须与真实业务高度相关；③有即时的价值反馈；④配合高层支持，如果CEO带头参与，其他高管的参与度会显著提升。最后这点的意思是，你可能需要先卖给CEO，而不是先卖给培训经理。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.5 "推演结果和真实表现之间的相关性怎么证明"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"这是整个培训行业面临的共同难题。几种验证方式：①延迟后测，推演结束后3-6个月追踪真实决策表现；②对照设计，有条件的话做对照组；③主观评估，让高管的上级评估其决策质量变化；④AI辅助评估。坦率地说，相关性无法完美证明，但合理的推断是可能的——关键是建立评估文化，而不是追求精确归因。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.6 "课程内容很好，但感觉我们公司落地条件不够"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"让我问一个问题：你们公司现在有没有人在用AI工具做高管培训？如果答案是没有，那问题不是条件不够，而是没有人开始。如果答案是有但效果不好，那问题是怎么做而不是做不做。落地条件不够的解决方案不是等条件够了再做，而是从最小可行开始。探索期方案只需要3-6个月和10万预算，你们等得了吗？"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.7 "这个和我们的案例教学有什么区别"', 2)
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"案例分析是看别人怎么做的，场景推演是让你自己做一遍。打个比方：案例分析是看赛车比赛录像，场景推演是让你上赛道。你在录像里看舒马赫过弯看得再清楚，自己上车还是不一样。更关键的是，在场景推演里，你可以死很多次——这次过弯速度快了翻车了，下次重来。在真实赛道上，你没有重来的机会。"')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)

add_h(doc, '4.8 学员主导讨论、其他人插不进来', 2)
add_p(doc, '情境：一位活跃的学员持续输出，其他人开始沉默。')
add_p(doc, '参考回应：', bold=True)
p = doc.add_paragraph()
r = p.add_run('"感谢你的分享，你说得非常有深度。其他人有没有不同的视角或经验？"——邀请其他人加入。如果这位学员继续主导，说："这样，XXX你讲得很好，我们等会儿专门安排时间深入交流。现在先让其他组员也分享一下他们的观点，好吗？"——给面子但设边界。')
sf(r)
p.paragraph_format.left_indent = Inches(0.3)
add_p(doc, '备选方案：直接用练习把讨论结构化——"现在请每个人写下你的观点，不只是口头讨论。"——书面形式天然平衡发言权。')

doc.add_page_break()

# ============ CHAPTER 5 ============
add_h(doc, '第五章  小组讨论引导技巧', 1)

add_h(doc, '5.1 提问技术', 2)
add_h(doc, '5.1.1 提问的四个层次', 3)
q_levels = [
    ('事实层', '了解客观信息', '"这个季度的业绩数据是多少？"'),
    ('解释层', '探索原因和逻辑', '"为什么供应商在这个时候提出涨价？"'),
    ('行动层', '指向决策和行动', '"我们应该接受涨价还是寻找替代供应商？"'),
    ('元认知层', '反思思考过程本身', '"我们是否忽略了某个视角？"'),
]
for level, purpose, example in q_levels:
    p = doc.add_paragraph()
    r = p.add_run(level + '：')
    sf(r, bold=True)
    r2 = p.add_run(purpose + ' — ' + example)
    sf(r2)

add_p(doc, '讲师最常犯的错误是停留在"事实层"和"解释层"，忘记问"行动层"和"元认知层"。每一次讨论，最终要落到行动——"所以呢？你回去会做什么不同？"')

add_h(doc, '5.1.2 追问技术', 3)
add_p(doc, '追问是引导深度的核心工具。常用追问句式：')
for item in [
    '"展开说说——你说的不够真实具体是指什么？"',
    '"你能举个具体的例子吗？"',
    '"这个观点，其他人有没有不同的经验？"',
    '"如果按你说的，会导致什么结果？"',
    '"有没有人想过相反的可能性？"',
]:
    add_bullet(doc, item)

add_h(doc, '5.2 讨论节奏控制', 2)
add_h(doc, '5.2.1 讨论的三个阶段', 3)
stages = [
    ('发散期（前30%时间）', '让想法充分涌现，不急于评判'),
    ('聚焦期（中间50%时间）', '识别关键议题，剔除非核心'),
    ('收敛期（后20%时间）', '形成结论或行动计划'),
]
for stage, desc in stages:
    p = doc.add_paragraph()
    r = p.add_run(stage + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '5.2.2 常见问题处理', 3)
problems = [
    ('讨论过于发散', '请用白板记录所有观点，然后问："这些观点中，哪3个是最核心的？"——强制聚焦'),
    ('讨论陷入僵局', '引入新信息或新视角："如果外部专家在这里，他会怎么看？"——打破僵局'),
    ('某人过于主导', '肯定后设边界，或直接进入书面讨论平衡发言权'),
    ('气氛过于对抗', '重新定义讨论目的："我们不是在争谁对谁错，而是在寻找更好的解决方案"'),
]
for problem, solution in problems:
    p = doc.add_paragraph()
    r = p.add_run(problem + '：')
    sf(r, bold=True)
    r2 = p.add_run(solution)
    sf(r2)

add_h(doc, '5.3 学员参与激发技术', 2)
techs = [
    ('轮转发言', '每人不超过2分钟，限时发言，确保每个人都有机会开口'),
    ('随机抽问', '用便签纸写上学员名字，随机抽取，被抽中者必须发言'),
    ('无声头脑风暴', '先给3分钟安静写想法，再开始讨论——避免声音大的学员主导'),
    ('两两对话', '让学员与邻座先讨论2分钟，再回到全班分享——降低发言门槛'),
    ('立场交换', '"现在请换到对方的立场思考——如果你是反对者，你会怎么说？"'),
]
for tech, desc in techs:
    p = doc.add_paragraph()
    r = p.add_run(tech + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '5.4 场景讨论引导话术模板', 2)
script = """【场景发布阶段】
"现在请翻开工作纸第X页，你们有X分钟时间阅读情境。如果有问题，先记下来，稍后统一提问。"

【信息探索阶段】
"现在你们可以向我提问获取更多信息。我不会一次性告诉你们所有背景——就像真实的高管一样，你们需要主动提问。"

【决策提交阶段】
"请在纸上写下你们的决策：选择哪个选项，以及最重要的3个理由。5分钟后，我们看各组的选择。"

【复盘阶段】
"现在告诉我：你们当时的决策和实际结果对比，你们最大的'啊哈时刻'是什么？你们带走哪一条可以迁移到真实工作的原则？"
"""
p = doc.add_paragraph()
r = p.add_run(script)
sf(r)

doc.add_page_break()
# ============ CHAPTER 6 ============
add_h(doc, '第六章  评估标准与评分细则', 1)

add_h(doc, '6.1 学员个人评估', 2)
add_h(doc, '6.1.1 课程全程评估体系', 3)
et = doc.add_table(rows=6, cols=5)
et.style = 'Table Grid'
for i, h in enumerate(['评估层次', '评估内容', '评估方法', '评估时点', '权重']):
    set_cell(et.cell(0, i), h, bold=True)
ed = [
    ['反应层', '满意度 Relevance', '即时问卷', '每个模块结束时', '10%'],
    ['学习层', '知识掌握度', '情境测试', 'D2', '15%'],
    ['行为层', '工具使用频率', '系统日志+访谈', 'D30/D60/D90', '25%'],
    ['结果层', '业务指标变化', '数据分析', 'D180/D360', '30%'],
    ['迁移层', '决策质量提升', 'AI模拟+专家评估', 'D30/D90', '20%'],
]
for ri, row in enumerate(ed):
    for ci, val in enumerate(row):
        set_cell(et.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '6.1.2 综合演战评分细则', 3)
add_p(doc, '综合演战是课程的核心产出环节，每组方案按以下维度评分：')
wt = doc.add_table(rows=7, cols=5)
wt.style = 'Table Grid'
for i, h in enumerate(['评分维度', '权重', '优秀(90-100)', '良好(70-89)', '合格(60-69)']):
    set_cell(wt.cell(0, i), h, bold=True)
wd = [
    ['情境真实性', '20%', '三棱镜完整，参数配置合理有据', '三棱镜较完整，参数基本合理', '三棱镜有缺失，参数依据模糊'],
    ['五步法完整性', '20%', '五步完整，每步AI介入方式具体可操作', '五步较完整，AI介入方式较具体', '五步有缺失，AI介入方式模糊'],
    ['评估体系', '20%', '三层次完整，指标具体可测量', '三层次较完整，指标基本可测量', '层次有缺失，指标模糊'],
    ['工具选型', '15%', '工具选择有明确依据，组合合理', '工具选择有依据，组合基本合理', '工具选择依据模糊'],
    ['落地可行性', '15%', '路径清晰，里程碑可实现，风险可控', '路径较清晰，里程碑较可实现', '路径有一定模糊性'],
    ['汇报呈现', '10%', '结构清晰，表达流畅，时间控制好', '结构较清晰，表达较流畅', '结构或表达有明显问题'],
]
for ri, row in enumerate(wd):
    for ci, val in enumerate(row):
        set_cell(wt.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '6.1.3 学员个人评分标准', 3)
ic = [
    ('出勤率（满分20分）', '全勤20分，迟到/早退每次扣5分，旷课0分'),
    ('课堂参与（满分30分）', '发言质量高+15分，中等+10分，低+5分，从不发言0分'),
    ('练习完成（满分30分）', '按时完成+10分/练习，共3个练习'),
    ('最终产出（满分20分）', '综合演战个人贡献度评估'),
]
for criteria, desc in ic:
    p = doc.add_paragraph()
    r = p.add_run(criteria + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '6.2 讲师自评标准', 2)
ie = [
    ('内容熟悉度', '能否脱稿讲述核心概念？是否需要频繁看资料？'),
    ('时间控制', '是否在规定时间内完成每个模块？超时了多少次？'),
    ('学员回应', '学员是否积极参与？是否有沉默或抵触？'),
    ('AI演示质量', '演示是否成功？失败的原因是什么？'),
    ('讨论引导', '是否有效激发学员发言？是否控制住了发散？'),
]
for dim, questions in ie:
    p = doc.add_paragraph()
    r = p.add_run(dim + '：')
    sf(r, bold=True)
    r2 = p.add_run(questions)
    sf(r2)
add_p(doc, '建议：每次课程结束后，用10分钟完成讲师自评表，记录本次授课的改进点。')

doc.add_page_break()

# ============ CHAPTER 7 ============
add_h(doc, '第七章  风险预案与替代方案', 1)

add_h(doc, '7.1 技术风险预案', 2)
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
for i, h in enumerate(['风险类型', '概率', '应对方案']):
    set_cell(tbl.cell(0, i), h, bold=True)
tr = [
    ['AI工具不可用', '中', '提前准备录制的演示视频或截图；改为纯讨论模式，延长小组讨论环节；使用提前下载好的AI输出案例作为替代展示'],
    ['网络连接不稳定', '高', '提前测试网络，准备移动热点；所有课程材料离线备份（PPT、文档、案例）；减少需要实时联网的环节'],
    ['投影设备故障', '低', '提前30分钟测试投影；准备PPT转PDF版本作为备份；如故障无法修复，改为白板讲解'],
    ['学员设备问题', '低', '准备公用账号供没有自己账号的学员使用；安排助教协助解决技术问题'],
]
for ri, row in enumerate(tr):
    for ci, val in enumerate(row):
        set_cell(tbl.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, '7.2 内容风险预案', 2)
cr = [
    ('学员质疑课程理论基础', '中', '承认局限性，邀请学员分享替代框架；如果质疑有道理，课后将反馈纳入课程迭代；避免防御性回应'),
    ('学员挑战讲师专业性', '中', '诚实承认知识边界，邀请学员分享经验；把问题抛回给全班讨论；避免假装权威'),
    ('讨论偏离主题', '高', '用白板记录偏离的观点，稍后回头处理；温和打断："这个话题很有价值，我们先记下来，回到主议题"'),
    ('学员对案例不感兴趣', '低', '立即切换到学员自己的案例；邀请学员提供他们经历过的类似情境'),
]
for risk, prob, solution in cr:
    p = doc.add_paragraph()
    r = p.add_run(risk + '：')
    sf(r, bold=True)
    r2 = p.add_run('[' + prob + ']' + solution)
    sf(r2)

add_h(doc, '7.3 参与度风险预案', 2)
pr = [
    ('学员整体参与度低', '高', '增加互动环节（投票、匿名提问）；将大组讨论拆分为两两对话；降低问题难度，增加结构化提示'),
    ('某组成员过于沉默', '中', '安排助教深入该组引导；给该组布置具体任务而非开放性问题；减少全班讨论，增加小组讨论'),
    ('情绪冲突发生', '低', '立即暂停讨论，承认冲突存在；单独与相关方沟通；如果情绪激烈，考虑休息片刻'),
    ('学员要求课后持续联系', '低', '在课程开始时就说明后续联系方式；提供课程专属答疑邮箱或群；避免在课程期间做出无法兑现的承诺'),
]
for risk, prob, solution in pr:
    p = doc.add_paragraph()
    r = p.add_run(risk + '：')
    sf(r, bold=True)
    r2 = p.add_run('[' + prob + ']' + solution)
    sf(r2)

add_h(doc, '7.4 时间风险预案', 2)
tr2 = [
    ('某个模块超时', '高', '优先保证核心概念讲解，压缩练习环节；非核心内容可宣布"留作课后阅读"；不要为了赶时间牺牲收尾质量'),
    ('整体课程时间不够', '中', '优先保证综合演战和汇报环节（这是核心产出）；压缩茶歇时间；第二天早上做简短回顾而非冗长复盘'),
    ('讨论陷入僵局', '中', '引入新信息或外部专家视角；强制切换到"两两对话"模式；用结构化问题引导'),
    ('学员反复提问超出范围', '低', '承认问题价值，承诺课后单独交流；用"这是个好问题，我们稍后深入"简短回应'),
]
for risk, prob, solution in tr2:
    p = doc.add_paragraph()
    r = p.add_run(risk + '：')
    sf(r, bold=True)
    r2 = p.add_run('[' + prob + ']' + solution)
    sf(r2)

add_h(doc, '7.5 替代方案速查表', 2)
add_p(doc, '当以下情况发生时，直接切换到对应替代方案：')
at = doc.add_table(rows=6, cols=2)
at.style = 'Table Grid'
for i, h in enumerate(['情况', '替代方案']):
    set_cell(at.cell(0, i), h, bold=True)
ad = [
    ['AI演示失败', '立即切到提前录制的视频；或展示静态PPT+讲师口述；或邀请已试过的学员分享'],
    ['某个案例学员不熟悉', '立即切换到行业通用案例；或邀请该行业学员提供真实背景'],
    ['讨论时间不够', '将开放讨论改为结构化两两对话；每人写下来再分享'],
    ['学员人数超过30人', '增加1名助教；将小组讨论改为小组代表发言；减少需要全员发言的环节'],
    ['学员人数少于10人', '增加个人练习时间；将小组汇报改为人人汇报；增加深度追问'],
]
for ri, row in enumerate(ad):
    for ci, val in enumerate(row):
        set_cell(at.cell(ri+1, ci), val)
doc.add_paragraph()

doc.add_page_break()

# ============ CHAPTER 8 ============
add_h(doc, '第八章  课后跟进建议', 1)

add_h(doc, '8.1 课后24小时内', 2)
items_24 = [
    ('发送课程资源包', '包含：课程PPT、工具清单、Prompt模板库、推荐阅读书单'),
    ('发送课程录屏链接', '如课程有录制，在课后24小时内发送给学员'),
    ('发布同期学员交流群', '邀请所有学员加入长期有效的交流群'),
    ('发送训后实践任务', '明确D30/D60/D90的实践任务，具体到"做什么"而非"学什么"'),
    ('收集课程反馈', '发送匿名满意度调研问卷，48小时内收集'),
]
for action, desc in items_24:
    p = doc.add_paragraph()
    r = p.add_run(action + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '8.2 训后30天跟进（D30）', 2)
d30 = [
    ('发送复习邮件', '回顾课程核心概念和工具，特别是学员反馈中最需要的部分'),
    ('邀请实践分享', '"在过去30天里，你有没有用过今天学的某个工具？结果如何？"'),
    ('一对一答疑', '为有需要的学员提供30分钟一对一咨询（可选，自愿报名）'),
    ('收集典型案例', '如果有学员成功将课程内容落地，邀请他们撰写案例分享'),
]
for action, desc in d30:
    p = doc.add_paragraph()
    r = p.add_run(action + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '8.3 训后90天跟进（D90）', 2)
d90 = [
    ('发送行为改变评估', '通过简短问卷了解学员在工作中是否应用了课程所学'),
    ('追踪决策质量', '如果可能，邀请学员分享一个真实决策案例，用课程框架做复盘'),
    ('发送进阶内容', '根据学员实践情况，推送针对性的进阶学习资料'),
    ('邀请参与下一期课程', '作为"学长"分享经验，或作为助教协助授课'),
]
for action, desc in d90:
    p = doc.add_paragraph()
    r = p.add_run(action + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, '8.4 讲师个人复盘', 2)
ir = [
    ('课程完成度评估', '哪些模块按计划完成？哪些被压缩或跳过？原因是什么？'),
    ('学员反馈总结', '学员满意度如何？哪些环节评价最高？哪些最低？'),
    ('关键事件记录', '有没有突发情况？如何处理的？有什么可以改进的？'),
    ('个人收获与反思', '这次授课中，我自己学到了什么？有什么新的洞见？'),
    ('课程迭代建议', '基于这次授课，课程内容、案例、练习可以如何改进？'),
]
for dim, questions in ir:
    p = doc.add_paragraph()
    r = p.add_run(dim + '：')
    sf(r, bold=True)
    r2 = p.add_run(questions)
    sf(r2)

add_h(doc, '8.5 学员学习效果追踪表', 2)
add_p(doc, '建议使用以下表格追踪学员学习效果：')
tt = doc.add_table(rows=6, cols=6)
tt.style = 'Table Grid'
for i, h in enumerate(['学员', 'D30实践任务完成情况', 'D60工具使用频率', 'D90决策质量自评', '实际业务应用', '备注']):
    set_cell(tt.cell(0, i), h, bold=True)
for i in range(1, 6):
    for j in range(6):
        tt.cell(i, j).text = ' '
doc.add_paragraph()

doc.add_page_break()

# ============ APPENDIX 1 ============
add_h(doc, '附录一  核心工具速查表', 1)

add_h(doc, 'A1.1 情境重构三棱镜', 2)
pt = doc.add_table(rows=4, cols=3)
pt.style = 'Table Grid'
for i, h in enumerate(['构成要素', '定义', '设计要点']):
    set_cell(pt.cell(0, i), h, bold=True)
pd = [
    ['情境变量', '影响决策的外部信息集', '足够真实、足够复杂、可操控'],
    ['决策节点', '参与者必须做出选择的关键时刻', '足够模糊、足够压力、有时间约束'],
    ['反馈机制', '决策后的结果呈现方式', '即时、可理解、与真实结果相关'],
]
for ri, row in enumerate(pd):
    for ci, val in enumerate(row):
        set_cell(pt.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, 'A1.2 决策推演五步法', 2)
ft = doc.add_table(rows=6, cols=3)
ft.style = 'Table Grid'
for i, h in enumerate(['步骤', '核心问题', 'AI介入方式']):
    set_cell(ft.cell(0, i), h, bold=True)
fd = [
    ['步骤1 情境初始化', '参与者在什么情境下做决策？', 'AI生成动态商业情境数据包'],
    ['步骤2 角色分配', '谁在做决策？利益如何分布？', 'AI模拟利益相关方反应'],
    ['步骤3 决策推进', '决策如何展开？', 'AI根据选择实时扩展情境'],
    ['步骤4 结果呈现', '决策产生了什么结果？', 'AI多维反馈（财务/组织/市场/声誉）'],
    ['步骤5 复盘反馈', '从这次决策中学到了什么？', 'AI辅助生成个性化复盘报告'],
]
for ri, row in enumerate(fd):
    for ci, val in enumerate(row):
        set_cell(ft.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, 'A1.3 Prompt工程四大原则', 2)
pmt = doc.add_table(rows=5, cols=2)
pmt.style = 'Table Grid'
for i, h in enumerate(['原则', '说明']):
    set_cell(pmt.cell(0, i), h, bold=True)
pmd = [
    ['角色注入要具体', '不是"你是CEO"，而是"你是有20年经验、经历过一次失败、正在寻找新方向的CEO"'],
    ['情境构建要分层', '第一层初始情境→第二层触发事件→第三层利益相关者反应→第四层时间压力'],
    ['反馈设计要挑刺', '不是"你说得有道理"，而是"你的隐含假设是什么？这个假设可能不成立的理由是？"'],
    ['输出结构要可追踪', '立场声明→核心论点→支撑论据→风险提示→替代方案'],
]
for ri, row in enumerate(pmd):
    for ci, val in enumerate(row):
        set_cell(pmt.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, 'A1.4 学习迁移对勾模型', 2)
hook_text = """对勾转化率参考：
• 知识获取→决策能力：60%（断层原因：情境过于简化）
• 决策能力→行动表现：40%（断层原因：时间压力不足）
• 行动表现→习惯改变：20%（断层原因：缺乏持续跟踪）"""
p = doc.add_paragraph()
r = p.add_run(hook_text)
sf(r, size=10)

doc.add_page_break()

# ============ APPENDIX 2 ============
add_h(doc, '附录二  场景卡使用指南', 1)

add_h(doc, 'A2.1 场景卡是什么', 2)
add_p(doc, '场景卡是本课程的核心学习工具，每张场景卡呈现一个高管在真实商业环境中面临的决策困境，学员需要在AI模拟器中完成：①分析情境；②做出决策；③推演后果；④复盘反思。')

add_h(doc, 'A2.2 场景卡使用流程', 2)
sflow = [
    ('Day 1上午', '场景卡发放→个人独立决策→AI模拟推演'),
    ('Day 1下午', '小组讨论→后果链分析→方案优化'),
    ('Day 2上午', '专家点评→框架提炼→落地可行性评估'),
]
for t, flow in sflow:
    p = doc.add_paragraph()
    r = p.add_run(t + '：')
    sf(r, bold=True)
    r2 = p.add_run(flow)
    sf(r2)

add_h(doc, 'A2.3 场景卡使用强度层级', 2)
ulevels = [
    ('层级一：微场景（15-20分钟）', '仅使用背景任务+关键节点+留白提问；小组快速讨论→各组汇报→讲师点评'),
    ('层级二：标准场景（30-45分钟）', '完整使用全部七个模块；角色扮演+小组讨论+讲师引导'),
    ('层级三：深度场景（60-90分钟）', '标准场景+AI辅助推演+决策复盘'),
    ('层级四：项目制场景（持续跟进）', '学员带真实工作难题进入场景；讲师跟踪辅导'),
]
for level, desc in ulevels:
    p = doc.add_paragraph()
    r = p.add_run(level + '：')
    sf(r, bold=True)
    r2 = p.add_run(desc)
    sf(r2)

add_h(doc, 'A2.4 场景卡选择矩阵', 2)
smt = doc.add_table(rows=5, cols=3)
smt.style = 'Table Grid'
for i, h in enumerate(['课程模块', '推荐场景集', '推荐强度']):
    set_cell(smt.cell(0, i), h, bold=True)
smd = [
    ['战略思维与决策', '高管战略决策', '层级二/三'],
    ['组织变革领导力', '组织变革管理', '层级二/三'],
    ['危机管理与风险决策', '危机应对与风险管理', '层级二/三'],
    ['资源整合与优先级', '资源配置与优先级', '层级一/二'],
]
for ri, row in enumerate(smd):
    for ci, val in enumerate(row):
        set_cell(smt.cell(ri+1, ci), val)
doc.add_paragraph()

add_h(doc, 'A2.5 场景卡设计质量检查清单', 2)
checklist = [
    '情境真实性：行业背景具体，不是"某公司"；企业情况有细节；时间节点紧迫感合理',
    '角色张力：至少3个角色，利益诉求不同；角色立场合理，不是"好人"vs"坏人"',
    '决策质量：至少2个关键决策点；每个决策没有标准答案，各有代价；决策结果可推演',
    '教学价值：留白提问能激发主动探索；四层提问覆盖不同思维层次；可提炼1-3条可迁移原则',
]
for item in checklist:
    add_bullet(doc, item)

doc.add_page_break()

# ============ CLOSING ============
add_h(doc, '结语', 1)
closing = """本讲师手册是课程交付的核心参考资料。讲师应在课前完成全面熟悉，在授课过程中持续积累反馈，并在每次课程结束后进行复盘迭代。

课程的最终价值不在于手册本身，而在于讲师与学员在真实对话中产生的洞见。

期待每一位使用本手册的讲师，都能在这个基础上创造属于自己的授课风格。
"""
p = doc.add_paragraph()
r = p.add_run(closing)
sf(r)

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('高管学习项目重构：AI辅助的场景推演与决策模拟设计 · 讲师手册 v1.0')
sf(r2, bold=True)

# SAVE
doc.save(OUT)
print('Document saved successfully!')

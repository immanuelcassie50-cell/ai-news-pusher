#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《领航·4.0》学员手册 Word 印刷版
A4 / 黑白可读 / 留白合理 / 表单填写区
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 引号占位符（避免 Python 字符串内嵌引号）
LQ = '“'   # left double quote “
RQ = '”'   # right double quote ”

OUT_PATH = r"D:\2026年课程\竞越\领航：Z世代管理新策略3.0\完整课程包\04_学员手册\03_学员手册_完整版_印刷版.docx"

# 所有 Python 字符串字面量统一使用三引号 """...""" 或单引号
# 内部的 ASCII 双引号 / 中文书名号 直接用 unicode 或变量


# ---------- 辅助函数 ----------

def set_run_font(run, font_name='宋体', size=10.5, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    if color:
        run.font.color.rgb = color


def add_paragraph_border(paragraph, top=False, bottom=True, color='999999', size=4):
    ppr = paragraph._element.get_or_add_pPr()
    pbdr = ppr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        ppr.append(pbdr)
    if top:
        b = OxmlElement('w:top')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), color)
        pbdr.append(b)
    if bottom:
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '1')
        b.set(qn('w:color'), color)
        pbdr.append(b)


def set_cell_shading(cell, fill='F2F2F2'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            spec = kwargs[edge]
            tag = qn('w:' + edge)
            elem = tcBorders.find(tag)
            if elem is None:
                elem = OxmlElement('w:' + edge)
                tcBorders.append(elem)
            for k, v in spec.items():
                elem.set(qn('w:' + k), v)


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._element.append(br)


def add_horizontal_line(doc, color='666666', size=8):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_blank_lines(doc, n=1, length=60):
    '''模拟下划线填写区'''
    line = '_' * length
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, '宋体', 10.5, color=RGBColor(0xBB, 0xBB, 0xBB))
    return p


def add_signature_line(doc, label='签名', date_label='日期'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + '：________________________     ')
    set_run_font(r1, '宋体', 10.5)
    r2 = p.add_run(date_label + '：____________________')
    set_run_font(r2, '宋体', 10.5)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, '黑体', 18, bold=True, color=RGBColor(0x00, 0x00, 0x00))
    add_paragraph_border(p, top=False, bottom=True, color='000000', size=12)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, '黑体', 14, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, '黑体', 12, bold=True, color=RGBColor(0x33, 0x33, 0x33))
    return p


def add_body(doc, text, indent_first=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(4)
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    set_run_font(r, '宋体', 10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + text)
    set_run_font(r, '宋体', 10.5)
    return p


def add_check_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    r1 = p.add_run('☐  ')
    set_run_font(r1, '宋体', 11, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, '宋体', 10.5)
    return p


def add_quote_block(doc, text, author=''):
    '''金句块：左竖线 + 灰底 + 楷体'''
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F5F5')
    set_cell_border(cell, left={'val': 'single', 'sz': '24', 'color': '888888'})
    p1 = cell.paragraphs[0]
    p1.paragraph_format.left_indent = Cm(0.4)
    p1.paragraph_format.line_spacing = 1.5
    r1 = p1.add_run(text)
    set_run_font(r1, '楷体', 11, color=RGBColor(0x33, 0x33, 0x33), bold=False)
    if author:
        p2 = cell.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.4)
        p2.paragraph_format.line_spacing = 1.5
        r2 = p2.add_run('—— ' + author)
        set_run_font(r2, '楷体', 10, color=RGBColor(0x66, 0x66, 0x66))
    return table


def add_callout(doc, text, fill='FFF8E1', border_color='E0A800'):
    '''提示框：浅色底 + 粗左边框'''
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={'val': 'single', 'sz': '24', 'color': border_color})
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(line)
        set_run_font(r, '宋体', 10.5)
    return table


def add_self_check_block(doc, items):
    '''自检清单块（带边框）'''
    table = doc.add_table(rows=len(items) + 1, cols=2)
    table.style = 'Table Grid'
    h0 = table.cell(0, 0)
    h0.text = ''
    set_cell_shading(h0, 'E8E8E8')
    p = h0.paragraphs[0]
    r = p.add_run('自检项')
    set_run_font(r, '黑体', 10.5, bold=True)
    h1 = table.cell(0, 1)
    h1.text = ''
    set_cell_shading(h1, 'E8E8E8')
    p = h1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('已完成')
    set_run_font(r, '黑体', 10.5, bold=True)
    for i, item in enumerate(items, start=1):
        c0 = table.cell(i, 0)
        c0.text = ''
        p = c0.paragraphs[0]
        r = p.add_run('☐  ' + item)
        set_run_font(r, '宋体', 10)
        c1 = table.cell(i, 1)
        c1.text = ''
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('☐')
        set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(13)
        row.cells[1].width = Cm(3)
    return table


# ============================================================
# 主流程
# ============================================================

def build():
    doc = Document()

    # A4 页面
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), '宋体')
    rfonts.set(qn('w:ascii'), '宋体')
    rfonts.set(qn('w:hAnsi'), '宋体')

    # 页脚
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run()
    set_run_font(fr, '宋体', 9, color=RGBColor(0x66, 0x66, 0x66))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    fr._element.append(fld_begin)
    fr._element.append(instr)
    fr._element.append(fld_sep)
    fr._element.append(fld_end)
    fr2 = fp.add_run(' /  领航·4.0 学员手册')
    set_run_font(fr2, '宋体', 9, color=RGBColor(0x66, 0x66, 0x66))

    # ============================================================
    # 封面页
    # ============================================================
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('领  航  ·  4.0')
    set_run_font(r, '黑体', 36, bold=True)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('AI 时代的 Z 世代管理新策略')
    set_run_font(r, '黑体', 22, bold=True)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—  学员手册  ·  印刷版  —')
    set_run_font(r, '楷体', 16, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border(p, top=False, bottom=True, color='000000', size=18)
    p.paragraph_format.space_after = Pt(40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('学  员  信  息')
    set_run_font(r, '黑体', 16, bold=True)
    p.paragraph_format.space_after = Pt(12)

    info_rows = [
        '学员姓名',
        '所在部门 / 团队',
        '岗位 / 职务',
        '课程日期',
        '分析对象（Z 世代员工代号）',
        '问责伙伴（Day 2 下午填写）',
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = 'Table Grid'
    for i, label in enumerate(info_rows):
        c0 = table.cell(i, 0)
        c0.text = ''
        set_cell_shading(c0, 'F2F2F2')
        p = c0.paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, '黑体', 10.5, bold=True)
        c1 = table.cell(i, 1)
        c1.text = ''
        c1.paragraphs[0].add_run(' ')
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)
        row.height = Cm(1.0)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('讲师：罗宏伟')
    set_run_font(r, '黑体', 14, bold=True)
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('讲师签名：________________________     日期：________________')
    set_run_font(r, '楷体', 11)
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('本手册仅供本课程学员使用 · 未经授权不得复制传播')
    set_run_font(r, '宋体', 9, color=RGBColor(0x88, 0x88, 0x88))

    # ============================================================
    # 目录页
    # ============================================================
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('目      录')
    set_run_font(r, '黑体', 22, bold=True)
    p.paragraph_format.space_after = Pt(18)

    toc_items = [
        ('引言：在我们开始之前', '1', True),
        ('　表单 0.1　五感诊断前测', '2', False),
        ('　表单 0.2　管理理念前测', '3', False),
        ('　表单 0.3　课前任务确认', '3', False),
        ('Part 1  认知力——认识 AI 时代的 Z 世代', '4', True),
        ('　工作页 1　快·变·连·我（AI 时代版本）', '5', False),
        ('　工作页 2　五感诊断与雷达图（含价值感）', '7', False),
        ('　工作页 3　AI 时代 Z 世代' + LQ + '三不怕' + RQ + '认知框架', '10', False),
        ('　工作页 4　管理理念自测与四座山移除', '12', False),
        ('Part 2  适应力——迎接 AI 时代的 Z 世代', '14', True),
        ('　工作页 5　融入期四阶段（AI 时代版）', '15', False),
        ('　工作页 6　1+3 任务清单 AI 升级版', '17', False),
        ('　工作页 7　坦诚交流策略（AI 时代版）', '19', False),
        ('　工作页 8　角色扮演记录表（非正式交流）', '21', False),
        ('Part 3  链接力——保鲜 AI 时代的 Z 世代', '23', True),
        ('　工作页 9　三维分析 + 5W2H+H 任务分配', '24', False),
        ('　工作页 10　辅导对话五步流程（AI 时代版）', '27', False),
        ('　工作页 11　常见辅导错误速查（含 AI 新错误）', '29', False),
        ('Part 4  愿景力——引爆 AI 时代的小宇宙', '31', True),
        ('　工作页 12　八大内驱动力画像（AI 时代交互分析）', '32', False),
        ('　工作页 13　游戏设计画布 4.0', '34', False),
        ('课程收尾：四力整合 + 30 天管理改进清单', '36', True),
        ('　表单 C.1　五感诊断重测', '37', False),
        ('　表单 C.2　管理理念重测', '37', False),
        ('　表单 C.3　30 天管理改进清单（4 区域）', '38', False),
        ('　　30 天清单日历版', '40', False),
        ('　　30 天清单海报版', '41', False),
        ('　　行为承诺签名页', '41', False),
        ('致出发的你', '42', True),
        ('附录一　术语速查表', '43', True),
        ('附录二　N 个工具速查索引', '44', True),
    ]
    toc_table = doc.add_table(rows=len(toc_items), cols=2)
    for i, (text, page, is_part) in enumerate(toc_items):
        c0 = toc_table.cell(i, 0)
        c0.text = ''
        p = c0.paragraphs[0]
        r = p.add_run(text)
        if is_part:
            set_run_font(r, '黑体', 11, bold=True)
        else:
            set_run_font(r, '宋体', 10.5)
        c1 = toc_table.cell(i, 1)
        c1.text = ''
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(page)
        set_run_font(r, '宋体', 10.5)
    for row in toc_table.rows:
        row.cells[0].width = Cm(13.5)
        row.cells[1].width = Cm(2.5)

    add_page_break(doc)

    # ============================================================
    # 引言
    # ============================================================
    add_h1(doc, '引言：在我们开始之前')
    add_quote_block(
        doc,
        '这本手册不是讲义，不是笔记，不是课后读物。它是你在两天课程里的工作台。'
        '你在课堂上写的每一笔、画的每一张图、签下的每一个承诺，都会被你带回去，30 天后还会翻到。',
        author='罗宏伟',
    )

    add_h3(doc, '课程全景图')
    add_body(doc, '四力模型，三个层次，一个方向：')
    add_bullet(doc, 'Part 1  认知力 ——  理解这代人' + LQ + '为什么这么做' + RQ + '，知道他们冰山下在驱动什么')
    add_bullet(doc, 'Part 2  适应力 ——  知道他们入职后经历什么阶段，每个阶段管理者该做什么')
    add_bullet(doc, 'Part 3  链接力 ——  日常怎么带他们。任务分配怎么写、辅导对话怎么谈')
    add_bullet(doc, 'Part 4  愿景力 ——  怎么点燃他们。八大内驱力 + 游戏化设计')
    add_body(doc, 'AI 时代新增第五感' + LQ + '价值感' + RQ + '，是他们最深的焦虑——AI 工具越普及，越需要被确认' + LQ + '我有 AI 替代不了的部分' + RQ + '。')

    # 表单 0.1
    add_h2(doc, '表单 0.1　五感诊断前测')
    add_callout(doc, '目的：了解你现在的起点，课程结束后用同一张表重测，看清变化。')
    add_body(doc, '要求：针对你的分析对象（一个真实 Z 世代员工）打分。这张表只有你和你的问责伙伴看。')

    p = doc.add_paragraph()
    r = p.add_run('分析对象代号：_____________     岗位：_____________     入职时间：_____________')
    set_run_font(r, '宋体', 10.5)
    p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['感', '核心需求', '几乎不', '偶尔', '经常/充足']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10.5, bold=True)
    data = [
        ('节奏感', '工作节奏有可预期的韵律'),
        ('存在感', '自己的贡献被精准看见'),
        ('位置感', '清晰的方向感和角色定位'),
        ('掌控感', '工作方式上有自主权'),
        ('价值感（新增）', '确认自己有 AI 替代不了的贡献'),
    ]
    for i, (g, desc) in enumerate(data, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(g)
        set_run_font(r, '黑体', 10, bold=True)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(desc)
        set_run_font(r, '宋体', 10)
        for j in (2, 3, 4):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('○')
            set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(2.2)
        row.cells[3].width = Cm(2.2)
        row.cells[4].width = Cm(2.6)

    add_body(doc, '最缺失的一感：______ 感')
    add_body(doc, '我观察到的具体行为表现：')
    add_blank_lines(doc, 3)

    # 表单 0.2
    add_h2(doc, '表单 0.2　管理理念前测')
    add_callout(doc, '目的：了解你的管理起点，对比课后的理念更新。如实打钩。')
    mgmt_items = [
        '我能在 30 分钟内说出 Z 世代员工的 3 个核心内驱动力',
        '我能在 Z 世代入职一周内说清楚团队对 AI 工具的态度',
        '我能用一句话说清楚' + LQ + '我的 Z 世代下属最不可替代的价值是什么' + RQ,
        '当下属说' + LQ + '这是 AI 做的' + RQ + '时，我能接住而不是质疑',
        '我能在任务分配时说清楚' + LQ + '哪部分你必须自己判断' + RQ,
        '我能区分下属的成长是' + LQ + '真实能力提升' + RQ + '还是' + LQ + 'AI 代劳' + RQ,
        '我能识别下属' + LQ + 'AI 走捷径' + RQ + '的风险并主动设计机制',
    ]
    table = doc.add_table(rows=len(mgmt_items) + 1, cols=5)
    table.style = 'Table Grid'
    headers = ['管理特征', '几乎不', '偶尔', '经常', '总是']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10.5, bold=True)
    for i, item in enumerate(mgmt_items, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(item)
        set_run_font(r, '宋体', 10)
        for j in (1, 2, 3, 4):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('○')
            set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(9)
        for j in (1, 2, 3, 4):
            row.cells[j].width = Cm(1.7)

    add_body(doc, '我打' + LQ + '几乎不' + RQ + '或' + LQ + '偶尔' + RQ + '的有 ___ 行。这些就是这门课的重点方向。')

    # 表单 0.3
    add_h2(doc, '表单 0.3　课前任务确认')
    add_body(doc, '请在课前 7 天内完成 5 项任务：')
    for item in [
        'A. 和分析对象做一次 5-10 分钟的非正式交流（自然场合即可）',
        'B. 观察并写下他让你最困惑的一个行为',
        'C. 完成 5 感诊断前测（表单 0.1）',
        'D. 阅读 RPD 设计分析第 1-2 节',
        'E. 找好问责伙伴（同事 / 同班学员 / 管理者朋友）',
    ]:
        add_check_item(doc, item)

    add_body(doc, '我在任务 A 里的一个观察：')
    add_blank_lines(doc, 3)

    add_page_break(doc)

    # ============================================================
    # Part 1
    # ============================================================
    add_h1(doc, 'Part 1　认知力——认识 AI 时代的 Z 世代')
    add_quote_block(doc, '改变你对 Z 世代的看法，比改变他们的行为更有效。', author='罗宏伟')
    add_h3(doc, '学习目标')
    add_bullet(doc, '用' + LQ + '快·变·连·我' + RQ + '理解 Z 世代的时代背景（含 AI 时代强化）')
    add_bullet(doc, '用五感驱动模型（含价值感）诊断 Z 世代员工的' + LQ + '冰山下' + RQ + '驱动力')
    add_bullet(doc, '用' + LQ + '三不怕' + RQ + '认知框架识别 AI 时代 Z 世代的三个新特征，并掌握转化方向')
    add_bullet(doc, '反思自己的管理理念，更新 4 个 AI 时代管理策略')

    # 工作页 1
    add_h2(doc, '工作页 1　快·变·连·我（AI 时代版本）')
    add_h3(doc, '知识点 1.1　Z 世代的四个时代基因')
    add_body(doc, 'Z 世代不是凭空出现的，他们是被四个时代基因塑造的：')
    add_bullet(doc, '快  ——  习惯了即时反馈，AI 工具让' + LQ + '快' + RQ + '成为常态')
    add_bullet(doc, '变  ——  在变化中长大，AI 让变化速度从' + LQ + '年' + RQ + '变成' + LQ + '月' + RQ)
    add_bullet(doc, '连  ——  天然连接的一代，AI 让连接的对象扩展到' + LQ + '非人' + RQ)
    add_bullet(doc, '我  ——  自我意识觉醒，AI 时代遇到新挑战：' + LQ + '我' + RQ + '的独特性在哪里')

    add_body(doc, '针对我的分析对象，分别写下 AI 时代强化后的具体表现：')
    add_body(doc, '【快】（AI 时代强化）：他最近的哪个行为让你感受到了' + LQ + '快' + RQ + '的升级？', indent_first=False)
    add_blank_lines(doc, 2)
    add_body(doc, '【变】（AI 时代强化）：他最近适应了什么新变化？适应速度怎么样？', indent_first=False)
    add_blank_lines(doc, 2)
    add_body(doc, '【连】（AI 时代新层——AI 协作连接）：他和 AI 工具的关系是' + LQ + '工具伙伴' + RQ + '还是' + LQ + '代劳依赖' + RQ + '？', indent_first=False)
    add_blank_lines(doc, 2)
    add_body(doc, '【我】（AI 时代挑战——价值焦虑）：他有没有流露过' + LQ + 'AI 都能做这些，我还能做什么' + RQ + '的表达？', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '练习 1-A　快·变·连·我的 AI 配对')
    add_body(doc, '配对分享（3 分钟）后，写下你的新发现：')
    add_blank_lines(doc, 2)

    add_self_check_block(doc, [
        '我对' + LQ + '快·变·连·我' + RQ + '每个维度都写了具体行为',
        '我没有停留在' + LQ + '觉得是这样' + RQ + '，而是写了' + LQ + '看到的行为' + RQ,
        '我能在 30 秒内向搭档说清楚我的分析对象最像哪个维度',
    ])

    add_page_break(doc)

    # 工作页 2
    add_h2(doc, '工作页 2　五感诊断工作页（含价值感）')
    add_h3(doc, '知识点 2.1　五感驱动模型（AI 时代 4.0 版）')
    add_body(doc, 'Z 世代的行为背后，是五个' + LQ + '冰山下' + RQ + '的内驱动力。原版四感已经过时——AI 时代必须增加第五感：价值感。')
    add_quote_block(doc, '价值感 = 员工对' + LQ + '自己有 AI 替代不了的贡献' + RQ + '的确认需求。')

    add_body(doc, '针对我的分析对象，逐感填写：')

    senses = [
        ('节奏感', 'AI 时代升级：AI 拉高了响应期待（' + LQ + '一个小时内有信号' + RQ + '是新的基线）',
         '他最常抱怨或表现出不满的' + LQ + '等待' + RQ + '是什么？',
         '我的应对策略：建立响应节奏约定（明确' + LQ + '我响应 X 小时，紧急 Y 通道' + RQ + '）'),
        ('存在感', 'AI 时代危机：他有没有过' + LQ + '被认可的只是会用 AI' + RQ + '的感受？',
         '我观察到的具体行为：',
         '我的应对策略：在反馈里主动说出他的人类判断贡献（不是' + LQ + '你 AI 用得好' + RQ + '，而是' + LQ + '你在 X 里的 Y 判断带来了 Z' + RQ + '）'),
        ('位置感', 'AI 时代新维度：他有没有过' + LQ + 'AI 工具能力比管理者还强' + RQ + '的隐性优势感？',
         '我观察到的具体行为：',
         '我的应对策略：主动承认并欢迎他的 AI 优势，把它变成团队资产（不是权威威胁）'),
        ('掌控感', 'AI 时代版本：他有没有过' + LQ + '被规定只能用某个 AI 工具' + RQ + '的抵触？',
         '我观察到的具体行为：',
         '我的应对策略：给出人类判断的清晰边界，工具选择权归他'),
        ('价值感（新增）', '核心焦虑：' + LQ + '如果 AI 能做所有这些，我的不可替代性在哪里？' + RQ,
         '他最近有没有类似的行为或表达？（频繁换方向、抵触重复任务、过度依赖 AI 或过度抵触 AI）',
         '我的应对策略：和他一起识别' + LQ + '只有他才能做的那部分' + RQ + '，定期更新（不是一次性的安慰）'),
    ]
    for name, knowledge, observe, strategy in senses:
        p = doc.add_paragraph()
        r = p.add_run('【' + name + '】')
        set_run_font(r, '黑体', 11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        add_body(doc, knowledge, indent_first=False)
        add_body(doc, observe, indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, strategy, indent_first=False)
        add_blank_lines(doc, 1)

    add_h3(doc, '5 感雷达图（请手绘）')
    add_body(doc, '在每个轴上标出你刚才打的分数（1-5），连起来形成五边形：', indent_first=False)
    radar = '''                  价值感
                    5
                    |
                4   |   4
              3     |     3
            2       |       2
          1         |         1
        节奏感——————+——————存在感
          5         |         5
            4       |       4
              3     |     3
                2   |   2
                    |
                位置感————掌控感
                  5'''
    p = doc.add_paragraph()
    r = p.add_run(radar)
    set_run_font(r, 'Courier New', 9)

    add_body(doc, '五边形告诉我什么？', indent_first=False)
    add_body(doc, '哪个角最凹（最缺）？____________', indent_first=False)
    add_body(doc, '形状对称吗？____________', indent_first=False)

    add_self_check_block(doc, [
        '我对五个感都写了具体行为，不只是打分',
        '我的应对策略是可以' + LQ + '下周就做' + RQ + '的，不是泛泛而谈',
        '雷达图我已经画好，并能看到' + LQ + '形状' + RQ + '传递的信息',
    ])

    add_h3(doc, '练习 2-A　五感诊断互评')
    add_body(doc, '配对互评（5 分钟）。搭档补充的策略：', indent_first=False)
    add_blank_lines(doc, 2)

    add_page_break(doc)

    # 工作页 3
    add_h2(doc, '工作页 3　AI 时代 Z 世代' + LQ + '三不怕' + RQ + '认知框架')
    add_h3(doc, '知识点 3.1　什么是' + LQ + '三不怕' + RQ)
    add_body(doc, 'AI 时代的 Z 世代有三个让管理者本能警觉的' + LQ + '不怕' + RQ + '：')

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['不怕', '含义', '管理者本能反应', '转化方向']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    data3 = [
        ('不怕比你懂 AI', '工具平权意识', '威胁感：我的权威被挑战了', '主动问' + LQ + '你用什么工具处理 X' + RQ + '，把他变成内部 AI 资源'),
        ('不怕说' + LQ + '这是 AI 做的' + RQ, '比前辈更坦然', '困惑/愤怒：' + LQ + '那你到底做了什么？' + RQ, '用' + LQ + 'AI 做了，那你的判断体现在哪里' + RQ + '开启归因对话'),
        ('不怕反向带教', '愿意教管理者用 AI', '不舒服：我需要被员工教？', '把反向带教变成团队学习文化的信号'),
    ]
    for i, row in enumerate(data3, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(5)

    add_body(doc, '针对我的分析对象，分别填写：')

    not_afraid = [
        ('【不怕比你懂 AI】',
         '他最近的表现（AI 用得比管理者熟练、对 AI 工具很有自信等）：',
         '我的本能反应：',
         '我的转化方向——我打算用的一个具体行动（本周内可执行）：'),
        ('【不怕说' + LQ + '这是 AI 做的' + RQ + '】',
         '他最近说过类似的话（请回忆原话）：',
         '我的本能反应：',
         '当他说' + LQ + '这是 AI 做的' + RQ + '，我的下一句话（写一句真实的话术，不是原则）：'),
        ('【不怕反向带教】',
         '他最近有没有主动教过我或团队什么？',
         '我的本能反应：',
         '我打算在什么场合主动邀请他分享（具体场景）：'),
    ]
    for title, q1, q2, q3 in not_afraid:
        p = doc.add_paragraph()
        r = p.add_run(title)
        set_run_font(r, '黑体', 11, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        add_body(doc, q1, indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, q2, indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, q3, indent_first=False)
        add_blank_lines(doc, 1)

    add_h3(doc, '练习 3-A　三不怕转化行动卡')
    add_body(doc, '选三个转化行动中最容易做的那一个，做成一张' + LQ + '行动卡' + RQ + '贴在你的工位：', indent_first=False)
    add_callout(
        doc,
        '□ 不怕比你懂 AI  →  我的行动：________\n'
        '□ 不怕说' + LQ + '这是 AI 做的' + RQ + '  →  我的下一句话：________\n'
        '□ 不怕反向带教  →  邀请场合：________\n'
        '\n执行时间：________    搭档见证：________',
        fill='F0F8FF', border_color='4682B4',
    )

    add_self_check_block(doc, [
        '我对三个' + LQ + '不怕' + RQ + '都识别了具体表现，不是抽象理解',
        '我的三个转化行动是' + LQ + '下周可执行' + RQ + '的，不是' + LQ + '理念上同意' + RQ,
        '我写了一句真实话术，不是原则表述',
    ])

    add_page_break(doc)

    # 工作页 4
    add_h2(doc, '工作页 4　管理理念自测与四座山移除')
    add_h3(doc, '知识点 4.1　移除四座山（AI 时代升级版）')
    add_body(doc, '管理者和 Z 世代之间，常有四座隐形的山。每座山都有' + LQ + 'AI 时代特别版本' + RQ + '：')

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['山', '原版症状', 'AI 时代版本', '移除方向']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    mts = [
        ('我吃的盐比你吃的米多', '倚老卖老、用经验压制', '用' + LQ + '我当年没用 AI 也做到了' + RQ + '压制 Z 世代', '换成' + LQ + '我们一起探索' + RQ),
        ('你应该听我的', '单向命令、不容讨论', '用' + LQ + '必须按我说的方法做' + RQ + '压制工具选择', '在人类判断边界内给工具选择权'),
        ('你想得太理想了', '质疑年轻人想法', '用' + LQ + 'AI 没那么神' + RQ + '否定 Z 世代的 AI 思路', '先听完他的方案再判断'),
        ('我们当年也是这样过来的', '用自己经历否定独特性', '用' + LQ + '我们当年没 AI 也能干活' + RQ + '否定新困境', '承认时代变了，他的困境是真实的'),
    ]
    for i, row in enumerate(mts, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(4)

    add_body(doc, '对我影响最大的一座山是：______ 山')
    add_body(doc, '移除这一座山，我 30 天内要做的一件事：', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '知识点 4.2　AI 时代三个新管理策略')
    add_bullet(doc, '软化冲突 AI 升级  ——  把' + LQ + '你和我的分歧' + RQ + '换成' + LQ + '我们和方案的差异' + RQ)
    add_bullet(doc, '共建 AI 协作规范  ——  不是' + LQ + '禁止用 AI' + RQ + '，是' + LQ + '我们一起制定怎么用 AI' + RQ)
    add_bullet(doc, '开放逆向学习  ——  主动说' + LQ + '这个你比我懂，你教我' + RQ)

    add_body(doc, '针对我的分析对象，三个策略我打算分别这样用：')
    add_body(doc, '软化冲突 AI 升级：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '共建 AI 协作规范：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '开放逆向学习：', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '管理理念自测（自评量表）')
    add_body(doc, '针对自己目前的真实情况打分（1-5 分，1=完全不符合，5=完全符合）：', indent_first=False)

    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    headers = ['管理理念', '1 / 2 / 3 / 4 / 5']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    mgmt_self = [
        '我理解 Z 世代的' + LQ + '快·变·连·我' + RQ + '，不只用' + LQ + '我当年' + RQ + '判断他们',
        '我能用五感分析员工的冰山下驱动力',
        '我能识别 AI 时代三不怕并把它转化为管理优势',
        '我能让他感到' + LQ + '被看见' + RQ + '的是他的判断贡献，不是 AI 产出',
        '我能和员工一起建立 AI 协作规范，不是单向规定',
        '我愿意承认' + LQ + 'AI 这个领域他比我懂' + RQ + '，并请他教我',
    ]
    for i, item in enumerate(mgmt_self, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(item)
        set_run_font(r, '宋体', 10)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('○    ○    ○    ○    ○')
        set_run_font(r, '宋体', 10)
    for row in table.rows:
        row.cells[0].width = Cm(11)
        row.cells[1].width = Cm(5)

    add_body(doc, '最让我意外的一项：_____________')
    add_body(doc, '我最想改变的一个管理习惯：')
    add_blank_lines(doc, 2)

    add_self_check_block(doc, [
        '我识别了对影响最大的一座山，并写了 30 天具体行动',
        '我的三个新策略不是抄写的，是针对分析对象的具体方案',
        '我对自己打分是' + LQ + '如实' + RQ + '的，不是' + LQ + '打得好看的' + RQ,
    ])

    # Part 1 行为承诺
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run('Part 1  行为承诺')
    set_run_font(r, '黑体', 13, bold=True)
    add_quote_block(doc, '从今天起，我承诺在我的 Z 世代下属面前——', author='')
    add_body(doc, '不再说' + LQ + '我当年不用 AI 也做到了' + RQ + '；', indent_first=False)
    add_body(doc, '而是承认他的 AI 时代困境，承认他的 AI 工具优势，承认' + LQ + '这个你可以教我' + RQ + '。', indent_first=False)
    add_body(doc, '我将从这一件具体的事开始：', indent_first=False)
    add_blank_lines(doc, 2)
    add_signature_line(doc, '签名', '日期')

    add_page_break(doc)

    # ============================================================
    # Part 2
    # ============================================================
    add_h1(doc, 'Part 2　适应力——迎接 AI 时代的 Z 世代')
    add_quote_block(doc, 'Z 世代换工作的成本在 AI 时代更低了，留住他们的窗口期没有变长。', author='罗宏伟')
    add_h3(doc, '学习目标')
    add_bullet(doc, '识别 Z 世代融入期四阶段的 AI 时代新困境和管理任务')
    add_bullet(doc, '运用 1+3 任务清单 AI 升级版（含人机协作说明）为新员工设计第一个月')
    add_bullet(doc, '掌握坦诚交流策略，用 AI 话题破冰建立信任')
    add_bullet(doc, '通过角色扮演建立非正式交流的实践感')

    # 工作页 5
    add_h2(doc, '工作页 5　融入期四阶段（AI 时代版）')
    add_h3(doc, '知识点 5.1　四阶段路径图')
    add_callout(doc, '未知（Day 1-7）→  防卫（Week 2-3）→  突破（Month 1-2）→  定位（Month 3-转正）', fill='F0F8FF', border_color='4682B4')
    add_body(doc, '原版挑战 →  AI 时代新困境：')
    add_bullet(doc, '未知  →  不知道 AI 工具在这里怎么用、管理者是否懂 AI、自己的 AI 习惯会不会被批评')
    add_bullet(doc, '防卫  →  试探' + LQ + '用 AI 是否被允许' + RQ + '、观察管理者对 AI 的态度')
    add_bullet(doc, '突破  →  产出大幅提升，但容易' + LQ + 'AI 依赖' + RQ + '替代真正的能力发展')
    add_bullet(doc, '定位  →  开始建立' + LQ + '我在 AI 时代的职业价值' + RQ + '的初步认知')

    add_body(doc, '我的分析对象当前阶段诊断：', indent_first=False)
    p = doc.add_paragraph()
    r = p.add_run('他目前在第 ______ 阶段。')
    set_run_font(r, '宋体', 10.5, bold=True)
    add_body(doc, '判断依据（具体行为）：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '这个阶段最重要的一个管理任务：', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, 'AI 时代四阶段详细笔记')

    add_body(doc, '【未知阶段】Day 1-7', indent_first=False)
    add_body(doc, '管理者核心任务：除了' + LQ + '做好准备' + RQ + '，还要在第一天就主动说清团队对 AI 的态度。', indent_first=False)
    add_body(doc, '我在 Day 1 会做的' + LQ + 'AI 态度说明' + RQ + '话术：', indent_first=False)
    add_blank_lines(doc, 2)

    add_body(doc, '【防卫阶段】Week 2-3', indent_first=False)
    add_body(doc, '管理者核心任务：除了' + LQ + '建立信任' + RQ + '，还要主动讨论 AI 使用的边界和期望。', indent_first=False)
    add_body(doc, '我会主动和他讨论的 AI 边界话题：', indent_first=False)
    add_blank_lines(doc, 2)

    add_body(doc, '【突破阶段】Month 1-2', indent_first=False)
    add_body(doc, '管理者核心任务：除了' + LQ + '感受成功' + RQ + '，还要帮助他建立' + LQ + '人类贡献清晰' + RQ + '的自我认知。', indent_first=False)
    add_body(doc, '我会用的' + LQ + '学习归因' + RQ + '开场问题：', indent_first=False)
    add_blank_lines(doc, 2)

    add_body(doc, '【定位阶段】Month 3-转正', indent_first=False)
    add_body(doc, '管理者核心任务：在转正评估里，不只评估产出，还评估人类判断贡献的成长。', indent_first=False)
    add_body(doc, '转正评估时，我会问的一个问题：', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '练习 5-A　四阶段路径图设计')
    add_body(doc, '如果你的分析对象是新入职（未转正），请为接下来 3-6 个月画一张路径图：', indent_first=False)
    add_body(doc, '第 1 周  →  ____________________', indent_first=False)
    add_body(doc, '第 2-3 周  →  ____________________', indent_first=False)
    add_body(doc, '第 1-2 月  →  ____________________', indent_first=False)
    add_body(doc, '第 3 月-转正  →  ____________________', indent_first=False)

    add_self_check_block(doc, [
        '我识别了分析对象当前阶段，不是抽象写' + LQ + '我们处于 X 阶段' + RQ,
        '我对四个阶段都写了 AI 时代版本的具体动作',
        '我的路径图是 3-6 个月可落地的，不是长期愿景',
    ])

    add_page_break(doc)

    # 工作页 6
    add_h2(doc, '工作页 6　1+3 任务清单 AI 升级版')
    add_h3(doc, '知识点 6.1　1+3 清单的 AI 时代升级')
    add_body(doc, '原版 1+3 清单：1 个核心目标 + 3 个子任务。AI 时代升级：每个子任务加一栏人机协作说明。', indent_first=False)
    add_body(doc, '人机协作说明的三要素：', indent_first=False)
    add_bullet(doc, '期望员工亲自判断的部分（不依赖 AI，必须人做）')
    add_bullet(doc, '可以 AI 辅助的部分（用 AI 提效，但判断在人）')
    add_bullet(doc, '人类贡献如何被看见（怎么验证员工确实做了人类判断）')

    add_body(doc, '为我的分析对象填写 1+3 清单', indent_first=False)
    p = doc.add_paragraph()
    r = p.add_run('任务周期：____________  至  ____________')
    set_run_font(r, '宋体', 10.5)
    add_body(doc, '核心目标（1 个，这个月最重要的成果）：', indent_first=False)
    add_blank_lines(doc, 2)

    for i in (1, 2, 3):
        add_body(doc, '子任务 ' + str(i) + '：', indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, '人机协作说明：', indent_first=False)
        add_body(doc, '• 期望员工亲自判断的部分：__________', indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, '• 可以 AI 辅助的部分：__________', indent_first=False)
        add_blank_lines(doc, 1)
        add_body(doc, '• 人类贡献如何被看见（验证方式）：__________', indent_first=False)
        add_blank_lines(doc, 1)
        add_horizontal_line(doc)

    add_h3(doc, '互评反馈')
    p = doc.add_paragraph()
    r = p.add_run('配对伙伴：__________     时间：__________')
    set_run_font(r, '宋体', 10.5)
    add_body(doc, '伙伴对' + LQ + '人类贡献如何被看见' + RQ + '的检验：', indent_first=False)
    add_body(doc, '• 这个验证方式可操作吗？__________', indent_first=False)
    add_body(doc, '• 员工能从这个验证方式知道什么算' + LQ + '做好了人类贡献' + RQ + '吗？__________', indent_first=False)
    add_body(doc, '我修改后的版本（如有）：', indent_first=False)
    add_blank_lines(doc, 2)

    add_self_check_block(doc, [
        '我的 1 个核心目标是这个月最关键的，不是季度或年度目标',
        '我对每个子任务都填了三要素（人做 / AI 辅助 / 人类贡献可见）',
        '我的' + LQ + '人类贡献如何被看见' + RQ + '是具体可验证的，不是' + LQ + '看他的工作态度' + RQ,
    ])

    add_page_break(doc)

    # 工作页 7
    add_h2(doc, '工作页 7　坦诚交流策略（AI 时代版）')
    add_h3(doc, '知识点 7.1　坦诚交流四策略')
    add_bullet(doc, '策略一：少命令多询问')
    add_bullet(doc, '策略二：有好奇擅回应')
    add_bullet(doc, '策略三：少说多听')
    add_bullet(doc, '策略四：从铁纪律到爱的教育')

    add_h3(doc, 'AI 话题破冰——我的开场话术')
    add_body(doc, '话术 A（针对不怕比你懂 AI）：', indent_first=False)
    add_body(doc, LQ + '我最近发现 XX 工具好像挺有意思的，但你比我熟。你平时用 XX 工具做什么用得最多？' + RQ, indent_first=False)
    add_body(doc, '我会用的具体话术：', indent_first=False)
    add_blank_lines(doc, 1)

    add_body(doc, '话术 B（针对不怕说' + LQ + '这是 AI 做的' + RQ + '）：', indent_first=False)
    add_body(doc, LQ + '上次那个项目里，AI 帮了你什么？没有 AI 的话你最头疼的部分是什么？' + RQ, indent_first=False)
    add_body(doc, '我会用的具体话术：', indent_first=False)
    add_blank_lines(doc, 1)

    add_body(doc, '话术 C（针对不怕反向带教）：', indent_first=False)
    add_body(doc, LQ + '我下周要给高管做汇报，你能不能花 10 分钟教我用 XX 工具做出来？我请你喝咖啡。' + RQ, indent_first=False)
    add_body(doc, '我会用的具体话术：', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '一周一次五分钟签到——我会问的 3 个问题')
    add_body(doc, '问题 1：__________', indent_first=False)
    add_body(doc, '问题 2：__________', indent_first=False)
    add_body(doc, '问题 3：__________', indent_first=False)

    add_h3(doc, 'AI 参与的正确回应——话术对比')
    add_callout(
        doc,
        '当员工说' + LQ + '这部分 AI 帮我做的' + RQ + '：\n\n'
        '✗ 不说（质问式）：' + LQ + '那你做了什么？' + RQ + '（关闭对话）\n\n'
        '✓ 改说（好奇式）：' + LQ + 'AI 帮了你哪部分？你在哪个环节做了判断？那个判断你是怎么想的？' + RQ,
        fill='FFF0F0', border_color='B22222',
    )

    add_body(doc, '我的版本（写一句真实的话术）：', indent_first=False)
    add_blank_lines(doc, 2)

    add_self_check_block(doc, [
        '我对四个策略都准备了具体话术',
        '我的开场话术不是抄原则，是' + LQ + '我下周会说' + RQ + '的真实话术',
        '我写的' + LQ + '应对员工说 AI 做的' + RQ + '是好奇语气，不是质问',
    ])

    add_page_break(doc)

    # 工作页 8
    add_h2(doc, '工作页 8　角色扮演记录表（非正式交流）')
    add_h3(doc, '练习 8-A　融入期非正式交流角色扮演')

    add_body(doc, '【第一轮】（我扮演管理者，搭档扮演新员工）', indent_first=False)
    add_body(doc, '情境设定：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '管理者用的开场第一句话：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '过程中最让我意外的是：', indent_first=False)
    add_blank_lines(doc, 1)

    add_body(doc, '【第二轮】（互换角色，我扮演新员工）', indent_first=False)
    add_body(doc, '让我' + LQ + '愿意多说一点' + RQ + '的瞬间，管理者做了：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '让我' + LQ + '立刻收住' + RQ + '的瞬间，管理者做了：', indent_first=False)
    add_blank_lines(doc, 1)

    add_body(doc, '角色扮演后的关键洞察：', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '我的下周非正式交流计划')
    p = doc.add_paragraph()
    r = p.add_run('计划时间：__________     场合：__________     对象：__________')
    set_run_font(r, '宋体', 10.5)
    add_body(doc, '开场第一句话（写出来，说出来）：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '我希望他带走的感受：', indent_first=False)
    add_blank_lines(doc, 1)

    add_self_check_block(doc, [
        '我做了两轮角色扮演，不是只听别人做',
        '我记录了' + LQ + '让我愿意多说' + RQ + '的瞬间，这是真实的可学技巧',
        '我为下周的非正式交流准备了具体时间、场合和开场',
    ])

    # Part 2 行为承诺
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run('Part 2  行为承诺')
    set_run_font(r, '黑体', 13, bold=True)
    add_quote_block(doc, '从今天起，我承诺——', author='')
    add_body(doc, '不在新人入职的第一周装作 AI 不存在；', indent_first=False)
    add_body(doc, '而是在第一天就主动说清楚团队对 AI 工具的态度。', indent_first=False)
    add_body(doc, '我将从这一件具体的事开始：', indent_first=False)
    add_blank_lines(doc, 2)
    add_signature_line(doc, '签名', '日期')

    add_page_break(doc)

    # ============================================================
    # Part 3
    # ============================================================
    add_h1(doc, 'Part 3　链接力——保鲜 AI 时代的 Z 世代')
    add_quote_block(doc, '任务分配不是给他一个名字，是给他一个' + LQ + '被看见' + RQ + '的方式。', author='罗宏伟')
    add_h3(doc, '学习目标')
    add_bullet(doc, '用三维分析框架（不能做 / 不愿做 / 不知道怎么和 AI 配合做）诊断员工' + LQ + '做不到' + RQ + '的真实原因')
    add_bullet(doc, '运用 5W2H+H 任务分配框架，给任务加上' + LQ + '人机协作视角' + RQ)
    add_bullet(doc, '掌握辅导对话五步流程（AI 时代版，含学习归因子步骤）')
    add_bullet(doc, '识别辅导中的常见错误（含 AI 新错误）')

    # 工作页 9
    add_h2(doc, '工作页 9　三维分析 + 5W2H+H 任务分配')
    add_h3(doc, '知识点 9.1　三维分析框架')
    add_body(doc, '员工' + LQ + '做不到' + RQ + '有三种完全不同的原因，对应的管理动作完全不同：', indent_first=False)

    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['维度', '含义', 'AI 时代版本', '应对方向']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    data9 = [
        ('不能做', '技能缺口', 'AI 工具能力缺口 / 能力空心化', '培训、辅导、补技能'),
        ('不愿做', '动力缺口', '价值感困惑 / 存在感冲突', '谈话、激发内驱力'),
        ('不知道怎么和 AI 配合做', '角色期望模糊', '人机分工不清晰', '明确人类判断边界'),
    ]
    for i, row in enumerate(data9, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(2.8)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(5.2)

    add_body(doc, '我的分析对象诊断——针对一个他最近' + LQ + '做不好' + RQ + '或' + LQ + '做得慢' + RQ + '的任务：', indent_first=False)
    p = doc.add_paragraph()
    r = p.add_run('任务名称：__________')
    set_run_font(r, '宋体', 10.5)

    add_body(doc, '他做不好的最可能原因是（在以下勾选一个）：', indent_first=False)
    add_check_item(doc, '不能做（技能 / AI 工具能力不足）')
    add_body(doc, '具体表现：__________', indent_first=False)
    add_check_item(doc, '不愿做（动力 / 价值感问题）')
    add_body(doc, '具体表现：__________', indent_first=False)
    add_check_item(doc, '不知道怎么和 AI 配合做（人机分工不清）')
    add_body(doc, '具体表现：__________', indent_first=False)
    add_body(doc, '我打算用的应对动作：', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '知识点 9.2　5W2H+H 任务分配框架')
    add_body(doc, '原版 5W2H：What / Why / Who / When / Where / How / How much', indent_first=False)
    add_body(doc, 'AI 时代新增第 8 维：Human（人机协作视角）', indent_first=False)

    add_body(doc, '5W2H+H 写作练习', indent_first=False)
    p = doc.add_paragraph()
    r = p.add_run('任务名称：__________')
    set_run_font(r, '宋体', 10.5)

    w2h8 = [
        ('What（做什么——具体清晰）', 3),
        ('Why（为什么这件事有价值——对谁有价值，满足员工的意义感需求）', 3),
        ('Who（谁来做，和谁协作）', 2),
        ('When（何时开始，何时完成）', 2),
        ('Where（在哪里完成）', 2),
        ('How（怎么做——方法和步骤）', 3),
        ('How much（达到什么标准）', 2),
    ]
    for label, lines in w2h8:
        add_body(doc, label, indent_first=False)
        add_blank_lines(doc, lines)

    p = doc.add_paragraph()
    r = p.add_run('Human（重点——人机协作视角）：')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    add_body(doc, '• 期望员工亲自判断的部分（不依赖 AI）：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '• 可以 AI 辅助提效的部分（提效但判断在人）：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '• 员工的人类贡献如何被看见（验证方式）：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '5W2H+H 自检清单')
    items9 = [
        'What 清晰具体，不模糊',
        'Why 说明了' + LQ + '对谁有价值' + RQ + '，不只是任务本身',
        'Who 角色和协作方式清楚',
        'When 起止时间明确',
        'Where 完成地点 / 线上 / 线下明确',
        'How 方法步骤有指引，不只说' + LQ + '做好了' + RQ,
        'How much 标准可验证（不空洞）',
        'Human 期望人做的明确',
        'Human 可 AI 辅助的明确',
        'Human 人类贡献如何被看见有验证方式',
    ]
    add_self_check_block(doc, items9)

    add_page_break(doc)

    # 工作页 10
    add_h2(doc, '工作页 10　辅导对话五步流程（AI 时代版）')
    add_h3(doc, '知识点 10.1　五步流程图')
    add_callout(doc, '开启对话 → 澄清事实（+ 学习归因）→ 交换看法 → 达成共识 → 总结对话', fill='F0F8FF', border_color='4682B4')

    p = doc.add_paragraph()
    r = p.add_run('第一步：开启对话')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(8)
    add_body(doc, '基础话术：' + LQ + '我想和你聊聊 [任务]，不是批评，是想一起看看怎么做得更好。你现在方便吗？' + RQ, indent_first=False)
    add_body(doc, 'AI 时代版本（我自己的话术）：', indent_first=False)
    add_blank_lines(doc, 1)

    p = doc.add_paragraph()
    r = p.add_run('第二步：澄清事实（含学习归因子步骤）')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    add_body(doc, '基础澄清话术：' + LQ + '你做这件事的过程是什么样的？哪部分进展顺利？哪部分有挑战？' + RQ, indent_first=False)
    add_body(doc, 'AI 时代学习归因问题（好奇语气）：' + LQ + '在这个过程里，你的判断体现在哪里？有哪部分是你之前不会、现在会了的？' + RQ, indent_first=False)
    add_body(doc, '我的版本：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '为什么是好奇而不是质问：__________', indent_first=False)
    add_blank_lines(doc, 1)

    p = doc.add_paragraph()
    r = p.add_run('第三步：交换看法')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    add_body(doc, '先听员工，再说管理者——' + LQ + '你怎么看？/ 你觉得这件事的关键是什么？' + RQ, indent_first=False)
    add_body(doc, '我的起手问题：__________', indent_first=False)
    add_blank_lines(doc, 1)

    p = doc.add_paragraph()
    r = p.add_run('第四步：达成共识')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    add_body(doc, '共建而非宣布——' + LQ + '那我们接下来可以怎么做？' + RQ, indent_first=False)
    add_body(doc, '我的话术：__________', indent_first=False)
    add_blank_lines(doc, 1)

    p = doc.add_paragraph()
    r = p.add_run('第五步：总结对话')
    set_run_font(r, '黑体', 11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    add_body(doc, '明确双方行动——' + LQ + '那我接下来会做 X，你接下来会做 Y，下次我们 [时间] 看进展。' + RQ, indent_first=False)
    add_body(doc, '我的话术：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '角色扮演记录（活动 9）')
    add_body(doc, '第一轮（经典辅导，我扮演 ______ 角色）：', indent_first=False)
    add_body(doc, '最难做到的是：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '第二轮（AI 时代辅导，我用的学习归因问题是）：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '员工角色的感受差异（好奇语气 vs 质问语气）：', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '最大的收获：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_self_check_block(doc, [
        '我对五步都准备了真实话术，不是原则表述',
        '我完成了 2 轮角色扮演，记录了真实的感受差异',
        '我的学习归因问题是' + LQ + '我能真正问出来' + RQ + '的，不是照抄',
    ])

    add_page_break(doc)

    # 工作页 11
    add_h2(doc, '工作页 11　常见辅导错误速查（含 AI 新错误）')
    add_h3(doc, '错误速查表')

    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ['错误', '描述', '我可能犯的情况', 'AI 时代特殊版本']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    errs = [
        ('是否判断', '还没听完就开始判断', '', LQ + '还没听完就先下了判断' + RQ),
        ('显摆立场', '用自己的经历压制员工', '', LQ + '我当年不用 AI 也做到了' + RQ),
        ('暗示兜圈', '不直接说，绕弯子让员工自己悟', '', ''),
        ('急于求成', '想一次谈话解决所有问题', '', ''),
        ('AI 新错误：质问 AI 参与', LQ + '那你做了什么？' + RQ + '（关闭对话）', '', LQ + 'AI 做的 = 你没做' + RQ + '的等式'),
        ('AI 新错误：AI 产出 ≠ 能力成长', '跳过学习归因', '', LQ + '上次做得不错' + RQ + ' = 这次能独立做？'),
    ]
    for i, row in enumerate(errs, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(3.8)
        row.cells[1].width = Cm(4.5)
        row.cells[2].width = Cm(3.5)
        row.cells[3].width = Cm(4.2)

    add_body(doc, '我最可能犯的两个错误：', indent_first=False)
    add_body(doc, '1. ____________________', indent_first=False)
    add_body(doc, '2. ____________________', indent_first=False)
    add_body(doc, '我打算如何避免：', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '辅导前自检清单')
    items11 = [
        '我先听完员工的完整描述再判断',
        '我不用自己的经历压制员工',
        '我不绕弯子，直接说重点',
        '我不期待一次谈话解决所有问题',
        '我用好奇语气问 AI 参与情况，不用质问',
        '我做了学习归因，区分了 AI 产出和员工能力成长',
    ]
    add_self_check_block(doc, items11)

    # Part 3 行为承诺
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run('Part 3  行为承诺')
    set_run_font(r, '黑体', 13, bold=True)
    add_quote_block(doc, '从今天起，我承诺——', author='')
    add_body(doc, '不再说' + LQ + '你那个方案做得很 AI' + RQ + '；', indent_first=False)
    add_body(doc, '而是在每次辅导对话里问一句：' + LQ + '你的判断体现在哪里？' + RQ, indent_first=False)
    add_body(doc, '我将从这一件具体的事开始：', indent_first=False)
    add_blank_lines(doc, 2)
    add_signature_line(doc, '签名', '日期')

    add_page_break(doc)

    # ============================================================
    # Part 4
    # ============================================================
    add_h1(doc, 'Part 4　愿景力——引爆 AI 时代的小宇宙')
    add_quote_block(doc, 'Z 世代换工作的成本在 AI 时代更低了，但被点燃的 Z 世代也是前所未有的。', author='罗宏伟')
    add_h3(doc, '学习目标')
    add_bullet(doc, '识别八大内驱动力在 AI 时代的交互变化')
    add_bullet(doc, '为分析对象做内驱力画像，识别主驱动力和最缺激活的驱动力')
    add_bullet(doc, '设计一份游戏设计画布 4.0 初稿（含防 AI 走捷径机制）')

    # 工作页 12
    add_h2(doc, '工作页 12　八大内驱动力画像（AI 时代交互分析）')
    add_h3(doc, '知识点 12.1　八大内驱动力')
    add_body(doc, '针对我的分析对象，在 1-5 分之间打分（5=最强驱动），并写下 AI 时代的特殊影响：', indent_first=False)

    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'
    headers = ['内驱动力', '简述', '评分', 'AI 时代的特殊影响']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    drv = [
        ('史诗意义与使命感', '在做比自己更大的事', 'AI 时代价值感危机时受冲击最大'),
        ('进步与成就感', '感受到真实成长', 'AI 代劳后难区分真实成长'),
        ('创意赋权与反馈', '用自己方式做，贡献被看见', '归属感模糊时缺失'),
        ('所有权与拥有感', '这件事是我的', 'AI 参与后所有权边界模糊'),
        ('社会影响力', '对他人的影响被感知', 'AI 向导角色可激活'),
        ('稀缺与迫切感', '限时限量的机会', '适度使用，避免压力疲劳'),
        ('未知与好奇心', '不确定性本身有吸引力', 'AI 工具探索是天然的激活点'),
        ('损失规避', '不想失去已有的', '需配合正向驱动，避免纯惩罚'),
    ]
    for i, row in enumerate(drv, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(row[0])
        set_run_font(r, '宋体', 9.5, bold=True)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(row[1])
        set_run_font(r, '宋体', 9.5)
        c = table.cell(i, 2)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('___')
        set_run_font(r, '宋体', 10)
        c = table.cell(i, 3)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(row[2])
        set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(1.5)
        row.cells[3].width = Cm(7)

    add_body(doc, '主驱动力（最高分 2-3 个）：__________', indent_first=False)
    add_body(doc, '最缺乏激活的驱动力（最低分 1-2 个）：__________', indent_first=False)
    add_body(doc, '主驱动力在 AI 时代的强化或弱化：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '最缺驱动力在 AI 时代的特殊激活方式：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '练习 12-A　内驱力激活计划')
    add_body(doc, '针对最缺激活的那个驱动力，我 30 天内要做的一个动作：', indent_first=False)
    add_blank_lines(doc, 2)

    add_self_check_block(doc, [
        '我对八大内驱动力都打了具体分数',
        '我识别了主驱动力和最缺驱动力',
        '我设计了具体的激活动作，不是原则表述',
    ])

    add_page_break(doc)

    # 工作页 13
    add_h2(doc, '工作页 13　游戏设计画布 4.0')
    add_h3(doc, '知识点 13.1　四大系统 4.0 升级要点')

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['系统', '4.0 核心升级', 'AI 时代关键点']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    sys13 = [
        ('目标系统 4.0', '产出目标 + 人类贡献目标分开设计', '人类贡献目标权重更高'),
        ('反馈系统 4.0', 'AI 即时产出反馈 + 管理者人类贡献反馈', '人类贡献回顾需人工，不能自动化'),
        ('规则系统 4.0', '明确 AI 使用边界（不是禁止）', '判断类工作期望人类独立'),
        ('回报系统 4.0', '与内驱动力对接，含成长性回报', '人类判断力成长本身可作为回报'),
    ]
    for i, row in enumerate(sys13, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(6.5)
        row.cells[2].width = Cm(5.5)

    add_h3(doc, '防 AI 走捷径机制')
    add_body(doc, '我的团队最可能走捷径的场景：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '我设计的应对机制：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '游戏设计画布 4.0 初稿摘要')
    add_body(doc, '业务场景：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '目标系统——产出目标：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '目标系统——人类贡献目标：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '反馈系统——AI 即时反馈：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '反馈系统——管理者人类贡献反馈（周期 + 内容）：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '规则系统——AI 使用边界：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '回报系统——对应主驱动力的激励：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '回报系统——人类判断力成长的特殊回报：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '练习 13-A　30 天最小可行性实验')
    add_body(doc, '实验名称：__________', indent_first=False)
    add_body(doc, '第一个动作：__________', indent_first=False)
    add_body(doc, '30 天后用什么判断效果：__________', indent_first=False)

    add_self_check_block(doc, [
        '我对四大系统都填了具体内容，不是抽象设计',
        '我设计了防 AI 走捷径机制，不只是' + LQ + '鼓励 AI 辅助' + RQ,
        '我的回报系统对接了主驱动力，不是' + LQ + '千篇一律的奖金' + RQ,
    ])

    # Part 4 行为承诺
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run('Part 4  行为承诺')
    set_run_font(r, '黑体', 13, bold=True)
    add_quote_block(doc, '从今天起，我承诺——', author='')
    add_body(doc, '不再用 AI 批量产出掩盖员工的真实成长缺失；', indent_first=False)
    add_body(doc, '而是在游戏设计里加入' + LQ + '人类贡献目标' + RQ + '，让成长可见化。', indent_first=False)
    add_body(doc, '我将从这一件具体的事开始：', indent_first=False)
    add_blank_lines(doc, 2)
    add_signature_line(doc, '签名', '日期')

    add_page_break(doc)

    # ============================================================
    # 课程收尾
    # ============================================================
    add_h1(doc, '课程收尾：四力整合 + 30 天管理改进清单')
    add_quote_block(doc, '两天课程真正的开始，是从你回到工位那一刻。', author='罗宏伟')

    add_h3(doc, '四力整合回顾')
    add_bullet(doc, '认知力：5 感雷达图 + 三不怕转化行动卡 + 管理理念更新笔记')
    add_bullet(doc, '适应力：1+3 清单（带人机协作说明）+ 坦诚交流话术库 + 角色扮演反思')
    add_bullet(doc, '链接力：5W2H+H 完整任务说明 + 辅导对话话术卡 + 错误自检清单')
    add_bullet(doc, '愿景力：八大内驱动力画像 + 游戏设计画布 4.0 初稿')

    # 表单 C.1
    add_h2(doc, '表单 C.1　五感诊断重测')
    add_body(doc, '课程结束，重新填写这张表——和开始时的自评对比，看清变化：', indent_first=False)
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'
    headers = ['感', '核心需求', '几乎不', '偶尔', '经常/充足']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10.5, bold=True)
    data_c1 = [
        ('节奏感', '工作节奏有可预期的韵律'),
        ('存在感', '自己的贡献被精准看见'),
        ('位置感', '清晰的方向感和角色定位'),
        ('掌控感', '工作方式上有自主权'),
        ('价值感', '确认自己有 AI 替代不了的贡献'),
    ]
    for i, (g, desc) in enumerate(data_c1, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(g)
        set_run_font(r, '黑体', 10, bold=True)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(desc)
        set_run_font(r, '宋体', 10)
        for j in (2, 3, 4):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('○')
            set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(2.2)
        row.cells[3].width = Cm(2.2)
        row.cells[4].width = Cm(2.6)

    add_body(doc, '课后' + LQ + '经常/充足' + RQ + '共 __ 行（课前：__ 行）', indent_first=False)
    add_body(doc, '哪一行变化最明显？为什么？', indent_first=False)
    add_blank_lines(doc, 2)

    # 表单 C.2
    add_h2(doc, '表单 C.2　管理理念重测')
    table = doc.add_table(rows=len(mgmt_items) + 1, cols=5)
    table.style = 'Table Grid'
    headers = ['管理特征', '几乎不', '偶尔', '经常', '总是']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10.5, bold=True)
    for i, item in enumerate(mgmt_items, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(item)
        set_run_font(r, '宋体', 10)
        for j in (1, 2, 3, 4):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('○')
            set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(9)
        for j in (1, 2, 3, 4):
            row.cells[j].width = Cm(1.7)

    add_body(doc, '课后' + LQ + '经常/总是' + RQ + '共 __ 行（课前：__ 行）', indent_first=False)
    add_body(doc, '变化最大的一项，我的新行动是什么？', indent_first=False)
    add_blank_lines(doc, 2)

    # 表单 C.3
    add_page_break(doc)
    add_h2(doc, '表单 C.3　30 天管理改进清单（课程最重要的产出）')
    add_callout(doc, '请认真用 25 分钟完成。这是你带回去最重要的成果。')

    add_h3(doc, '区域一：关于我的分析对象（认知与适应）')
    p = doc.add_paragraph()
    r = p.add_run('他最缺失的一感：______ 感')
    set_run_font(r, '宋体', 10.5)
    add_body(doc, '针对这个感，我 30 天内要做的一件具体的事：', indent_first=False)
    add_body(doc, '在（情境）__________________ 里，', indent_first=False)
    add_body(doc, '我会（做什么不同的事）__________________，', indent_first=False)
    add_body(doc, '从（时间）_______________ 开始。', indent_first=False)
    add_body(doc, '这件事做到了的标志是：__________________', indent_first=False)
    add_blank_lines(doc, 2)

    add_h3(doc, '区域二：关于任务分配（链接力——教）')
    add_body(doc, '即将分配的下一个重要任务：__________', indent_first=False)
    add_body(doc, '用 5W2H+H 重新说明时，Human 这一维：', indent_first=False)
    add_body(doc, '• 期望员工亲自判断的部分：__________', indent_first=False)
    add_body(doc, '• 人类贡献如何被看见：__________', indent_first=False)
    add_body(doc, '计划分配时间：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '区域三：关于激励机制（愿景力——机）')
    add_body(doc, '我从游戏设计画布提出的一个最小可行性实验：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '这个实验的第一步：__________', indent_first=False)
    add_body(doc, '计划开始时间：__________', indent_first=False)
    add_body(doc, '30 天后，可以用什么来判断实验有没有效果：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '区域四：下一次与分析对象的关键对话')
    add_body(doc, '对话类型（勾选一个）：', indent_first=False)
    add_check_item(doc, '非正式交流（AI 话题破冰）')
    add_check_item(doc, '辅导对话（含学习归因）')
    add_check_item(doc, '发展面谈（双轨方向）')
    add_check_item(doc, 'AI 时代价值感谈话（帮他找到只有他才能做的那部分）')
    add_body(doc, '计划时间：__________', indent_first=False)
    add_body(doc, '我的开场第一句话：__________', indent_first=False)
    add_blank_lines(doc, 1)
    add_body(doc, '我希望他从这次对话带走的感受：__________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '我的问责伙伴')
    p = doc.add_paragraph()
    r = p.add_run('姓名：__________     联系方式：__________')
    set_run_font(r, '宋体', 10.5)
    add_body(doc, '30 天后的约定：', indent_first=False)
    add_body(doc, '他问我：' + LQ + '你的 30 天清单里，做了什么，发现了什么？' + RQ, indent_first=False)
    add_body(doc, '我问他：____________________', indent_first=False)
    add_blank_lines(doc, 1)

    add_h3(doc, '三个 10 天阶段目标')
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    headers = ['阶段', '目标', '我要做的一件事', '怎么知道做到了', '勾选']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        if j == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    stages = [
        ('第 1-10 天', '建立意识：每次用工具前停 30 秒'),
        ('第 11-20 天', '建立技能：完成一份 5W2H+H 任务分配'),
        ('第 21-30 天', '建立系统：跑通第一个最小激励实验'),
    ]
    for i, (stage, goal) in enumerate(stages, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(stage)
        set_run_font(r, '宋体', 10, bold=True)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(goal)
        set_run_font(r, '宋体', 10)
        for j in (2, 3, 4):
            c = table.cell(i, j)
            c.text = ''
            if j == 4:
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run('☐')
                set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(2)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(4)
        row.cells[3].width = Cm(4)
        row.cells[4].width = Cm(1.5)

    add_body(doc, '30 天后可以检验的一个具体指标：', indent_first=False)
    add_blank_lines(doc, 2)

    add_page_break(doc)

    # 30 天清单日历版
    add_h2(doc, '30 天清单日历版')
    add_body(doc, '把每天要做的' + LQ + '小动作' + RQ + '填进去。比如第 1 周周三可以填' + LQ + '用学习归因开场做一次辅导对话' + RQ + '，第 3 周周一可以填' + LQ + '分配第一个 5W2H+H 任务' + RQ + '。', indent_first=False)

    table = doc.add_table(rows=5, cols=8)
    table.style = 'Table Grid'
    headers = ['周次', '周一', '周二', '周三', '周四', '周五', '周末', '回顾']
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    weeks = [
        ('第 1 周（建立意识）', '停 30 秒'),
        ('第 2 周（建立意识）', '停 30 秒'),
        ('第 3 周（建立技能）', '5W2H+H'),
        ('第 4 周（建立系统）', '激励实验'),
    ]
    for i, (w, theme) in enumerate(weeks, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(w)
        set_run_font(r, '宋体', 9, bold=True)
        for j in (1, 2, 3, 4, 5, 6):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(theme if j == 1 else '·')
            set_run_font(r, '宋体', 9)
        c = table.cell(i, 7)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('☐')
        set_run_font(r, '宋体', 12)
    for row in table.rows:
        row.cells[0].width = Cm(3.2)
        for j in (1, 2, 3, 4, 5, 6):
            row.cells[j].width = Cm(1.5)
        row.cells[7].width = Cm(1.5)

    add_body(doc, ' ')
    add_body(doc, '4 周主题：', indent_first=False)
    add_bullet(doc, '第 1 周：建立意识  ——  每次用工具前停 30 秒')
    add_bullet(doc, '第 2 周：建立意识  ——  持续停 30 秒，开始记录')
    add_bullet(doc, '第 3 周：建立技能  ——  完成一份 5W2H+H 任务分配')
    add_bullet(doc, '第 4 周：建立系统  ——  跑通第一个最小激励实验')

    # 30 天清单海报版
    add_h2(doc, '30 天清单海报版（可贴在办公位）')
    add_callout(
        doc,
        '┌─────────────────────────────────────────┐\n'
        '│         30 天 管理改进承诺                │\n'
        '├─────────────────────────────────────────┤\n'
        '│                                         │\n'
        '│  最缺的一感：______ 感                  │\n'
        '│  我的具体行动：________________        │\n'
        '│                                         │\n'
        '│  即将分配的任务（用 5W2H+H）：          │\n'
        '│  ________________________________       │\n'
        '│                                         │\n'
        '│  最小激励实验：                          │\n'
        '│  ________________________________       │\n'
        '│                                         │\n'
        '│  下次关键对话（时间）：                  │\n'
        '│  ________________________________       │\n'
        '│                                         │\n'
        '│  问责伙伴：________                      │\n'
        '│  下次复盘时间：________                  │\n'
        '│                                         │\n'
        '│  我的签名：________  日期：________      │\n'
        '│                                         │\n'
        '└─────────────────────────────────────────┘',
        fill='F0F8FF', border_color='4682B4',
    )

    # 行为承诺签名页
    add_h2(doc, '我的承诺签名')
    add_quote_block(doc, '我承诺在接下来的 30 天里，按照上面四个区域的具体行动推进。', author='')
    add_body(doc, '我会让我的问责伙伴每 2 周问一次：' + LQ + '你做了什么，发现了什么？' + RQ, indent_first=False)
    add_signature_line(doc, '签名', '日期')
    add_signature_line(doc, '见证人（问责伙伴）', '见证日期')

    add_page_break(doc)

    # ============================================================
    # 致出发的你
    # ============================================================
    add_h1(doc, '致出发的你')
    add_body(doc, '你用两天时间，走完了 AI 时代的 Z 世代管理四力模型。', indent_first=False)
    add_body(doc, '这不是一套理论，而是一套可以从明天起就开始用的工作方式。', indent_first=False)
    add_body(doc, '你在课堂上完成的每一张工作页、每一道练习、每一个真实场景的解决方案，都是真实的成果，不是作业。', indent_first=False)
    add_body(doc, '四力有一个共同的底层逻辑——', indent_first=False)
    add_quote_block(doc, 'AI 提供了前所未有的效率工具，但工具不创造价值，用工具的人创造价值。', author='')
    add_body(doc, '认知力让你看见这代人；适应力让你留住这代人；链接力让你带好这代人；愿景力让你点燃这代人。', indent_first=False)
    add_body(doc, 'Z 世代换工作的成本在 AI 时代更低了，但他们被点燃后的爆发力也是前所未有的。', indent_first=False)
    add_quote_block(doc, '讨喜而不是讨好。Z 世代管理者不需要' + LQ + '讨好' + RQ + '这代人，但需要' + LQ + '讨喜' + RQ + '——让他们感到被看见、被尊重、被期待。', author='罗宏伟')
    add_quote_block(doc, '领先半步，吃尽红利。AI 时代的管理者，比的不是懂 AI 比员工多，是比员工先看到管理要升级的地方。', author='罗宏伟')
    add_body(doc, '你今天开始建立的这套管理动作，在五年后，无论 AI 变成什么样子，都还有用。', indent_first=False)
    add_body(doc, '从今天，从现在，从这一次开始。', indent_first=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('—— 罗宏伟')
    set_run_font(r, '楷体', 12, color=RGBColor(0x55, 0x55, 0x55))
    p.paragraph_format.space_before = Pt(12)

    add_page_break(doc)

    # ============================================================
    # 附录一
    # ============================================================
    add_h1(doc, '附录一　术语速查表')
    terms = [
        ('五感驱动模型', 'Z 世代员工' + LQ + '冰山下' + RQ + '的五个内驱动力：节奏感、存在感、位置感、掌控感、价值感'),
        ('价值感', 'AI 时代新增的第五感：员工对' + LQ + '自己有 AI 替代不了的贡献' + RQ + '的确认需求'),
        ('三不怕', 'AI 时代 Z 世代的三个新特征：不怕比你懂 AI、不怕说' + LQ + '这是 AI 做的' + RQ + '、不怕反向带教'),
        ('融入期四阶段', 'Z 世代新员工从入职到转正经历的四个阶段：未知、防卫、突破、定位'),
        ('1+3 任务清单', '1 个核心目标 + 3 个子任务的融入期结构化任务管理工具'),
        ('人机协作说明', '1+3 清单 AI 升级版中增加的栏位：期望人做的、可 AI 辅助的、人类贡献如何被看见'),
        ('5W2H+H', '任务分配框架，原版 5W2H + Human（人机协作视角）共 8 维'),
        ('学习归因', '辅导对话中区分' + LQ + 'AI 做到的' + RQ + '和' + LQ + '员工学到的' + RQ + '的子步骤'),
        ('八大内驱动力', '游戏化管理的心理学基础：史诗意义、进步成就、创意赋权、所有权、社会影响、稀缺迫切、未知好奇、损失规避'),
        ('游戏设计画布 4.0', '含 AI 增强维度和防 AI 走捷径机制的激励系统设计工具'),
        ('三维分析', '员工' + LQ + '做不到' + RQ + '的三种原因诊断：不能做、不愿做、不知道怎么和 AI 配合做'),
        ('问责伙伴', '课后 30 天里每 2 周和你做 15 分钟复盘对话的同伴'),
    ]
    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.style = 'Table Grid'
    for j, h in enumerate(['术语', '定义']):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10.5, bold=True)
    for i, (term, defi) in enumerate(terms, start=1):
        c = table.cell(i, 0)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(term)
        set_run_font(r, '黑体', 10, bold=True)
        c = table.cell(i, 1)
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(defi)
        set_run_font(r, '宋体', 10)
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    add_page_break(doc)

    # ============================================================
    # 附录二
    # ============================================================
    add_h1(doc, '附录二　N 个工具速查索引')
    add_body(doc, '详细版见独立文件 04_学员手册_配套工具索引.md。', indent_first=False)
    tools = [
        ('0.1', '五感诊断前测表', '引言', '我的分析对象五感评分'),
        ('0.2', '管理理念前测表', '引言', '我的管理起点自评'),
        ('0.3', '课前任务确认表', '引言', '5 项课前任务'),
        ('1', '快·变·连·我 AI 时代版本', 'Part 1', '四个时代基因 AI 强化表现'),
        ('2', '五感诊断与雷达图', 'Part 1', '五感详细诊断 + 手绘雷达图'),
        ('3', '三不怕转化行动卡', 'Part 1', '三个' + LQ + '不怕' + RQ + '的转化行动'),
        ('4', '管理理念自测与四座山移除', 'Part 1', '6 项自评 + 移除四座山'),
        ('5', '融入期四阶段路径图', 'Part 2', '分析对象当前阶段诊断'),
        ('6', '1+3 任务清单 AI 升级版', 'Part 2', '三子任务 + 人机协作说明'),
        ('7', '坦诚交流策略与 AI 话题破冰', 'Part 2', '三话术 + 5 分钟签到问题'),
        ('8', '角色扮演记录表', 'Part 2', '非正式交流两轮记录'),
        ('9', '三维分析 + 5W2H+H', 'Part 3', '8 维任务分配'),
        ('10', '辅导对话五步流程（含学习归因）', 'Part 3', '五步话术设计'),
        ('11', '常见辅导错误速查', 'Part 3', '4 类经典 + 2 类 AI 新错误'),
        ('12', '八大内驱动力画像', 'Part 4', '八驱动力打分 + AI 时代交互'),
        ('13', '游戏设计画布 4.0', 'Part 4', '四系统 4.0 + 防 AI 走捷径机制'),
        ('C.1', '五感诊断重测', '收尾', '与前测对比'),
        ('C.2', '管理理念重测', '收尾', '与前测对比'),
        ('C.3', '30 天管理改进清单', '收尾', '4 个区域 + 日历版 + 海报版'),
        ('-', '行为承诺签名页', 'Part 1-4 + 收尾', '5 个签名页'),
    ]
    table = doc.add_table(rows=len(tools) + 1, cols=4)
    table.style = 'Table Grid'
    for j, h in enumerate(['编号', '工具名称', '位置', '核心问题']):
        c = table.cell(0, j)
        c.text = ''
        set_cell_shading(c, 'E8E8E8')
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, '黑体', 10, bold=True)
    for i, row in enumerate(tools, start=1):
        for j, val in enumerate(row):
            c = table.cell(i, j)
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, '宋体', 9.5)
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(6)

    add_body(doc, ' ')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('——  完  ——')
    set_run_font(r, '黑体', 14, bold=True, color=RGBColor(0x66, 0x66, 0x66))
    p.paragraph_format.space_before = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('版权所有 · 罗宏伟 · 本手册仅供本课程学员使用')
    set_run_font(r, '宋体', 9, color=RGBColor(0x88, 0x88, 0x88))

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print('OK:', OUT_PATH)
    print('Size:', os.path.getsize(OUT_PATH), 'bytes')


if __name__ == '__main__':
    build()

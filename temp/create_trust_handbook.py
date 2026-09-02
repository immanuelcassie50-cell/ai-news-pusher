# -*- coding: utf-8 -*-
"""
商业讲师信任护城河 - 学员手册生成脚本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\新课开发\工作手册\商业讲师信任护城河\完整课程包\04-学员手册\22-学员手册-商业讲师信任护城河.docx"

# 课程结构
CHAPTERS = [
    {"num": "前言", "title": "信任是培训师的生死之脉", "duration": "30min"},
    {"num": "课程介绍", "title": "PART1 概述", "duration": "15min"},
    {"num": "第一章", "title": "流量是别人的规则，信任是你自己的地盘", "duration": "60min"},
    {"num": "第二章", "title": "巨头能复制课件，复制不了客户对你的托付", "duration": "60min"},
    {"num": "第三章", "title": "每一次培训现场都是信任的存取款", "duration": "90min"},
    {"num": "第四章", "title": "客户选择你是因为有人替你打了保票", "duration": "60min"},
    {"num": "第五章", "title": "课前调研的本质是第一次信任测试", "duration": "60min"},
    {"num": "第六章", "title": "报价是信任浓度的显示器", "duration": "60min"},
    {"num": "第七章", "title": "拒绝不合适的项目比接下十个更能建立护城河", "duration": "60min"},
    {"num": "课程介绍", "title": "PART2 概述", "duration": "15min"},
    {"num": "第八章", "title": "你可以被超越但不能被替代", "duration": "60min"},
    {"num": "第九章", "title": "个人品牌和培训机构是两种不同的信任逻辑", "duration": "60min"},
    {"num": "第十章", "title": "信任的复利需要放弃流量带来的即时安全感", "duration": "60min"},
    {"num": "第十一章", "title": "面对低价竞争和抄袭要守住信任不是价格", "duration": "60min"},
    {"num": "第十二章", "title": "讲台下的样子决定讲台上还有没有人再请你", "duration": "60min"},
    {"num": "第十三章", "title": "行业最终留下来的人靠的是熬得住那几年", "duration": "60min"},
    {"num": "第十四章", "title": "同行不是对手是信任生态的共同守护者", "duration": "45min"},
    {"num": "第十五章", "title": "家人不理解放弃单子去守信任也是一种成本", "duration": "45min"},
]

# 各章节核心公理/金句
CHAPTERaxioms = {
    "前言": "流量是别人定义的规则，你玩得再好也只是一次陪跑；信任是你自己一块一块垒起来的，垒得慢，但垒起来的部分谁也拿不走。",
    "第一章": "你能用流量买到关注，买不到别人在做决定前替你说的那句话。",
    "第二章": "课件可以被抄走，信任没法被抄走，这是这个行业里最容易被误解的一件事。",
    "第三章": "上一次讲得好不好，客户不会记很久；上一次你有没有真的解决他的问题，客户会记很多年。",
    "第四章": "转介绍不是你的生意之外的额外收获，是你的生意本身。",
    "第五章": "甲方愿不愿意跟你说真话，比他愿不愿意签合同更能说明这单能不能长久。",
    "第六章": "客户压你的价，压的往往不是钱，是他对这次合作有没有把握。",
    "第七章": "你说'不'的那一刻，客户反而更相信你说的'是'。",
    "第八章": "差异化不是讲得比别人好，是没人能讲你能讲的那部分。",
    "第九章": "机构卖的是标准化和可控性，个人卖的是不可复制的你，两种逻辑混着做，容易做成两不像。",
    "第十章": "流量给你的是这个月安心的感觉，信任给你的是十年后还有饭吃的底气，两者常常在抢你同一份精力。",
    "第十一章": "别人可以抢走你的一单生意，抢不走客户心里那份判断，除非你自己先把它让出去。",
    "第十二章": "客户请你回来，不是因为你上一次讲得多精彩，是因为你在讲台之外，也是个说话算话的人。",
    "第十三章": "大浪淘沙淘掉的从来不是讲得不好的人，是熬不住看不见回报的那几年就退场的人。",
    "第十四章": "你多接一单，同行少接一单，这是竞争；但整个行业的口碑烂掉，谁都接不到单，这是所有人共同的处境。",
    "第十五章": "你放弃的那一单，在外人眼里是钱没赚到，只有你自己知道，那是护城河的一块砖。",
}

def set_cell_background(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_table_with_style(doc, headers, rows, header_color="C62828"):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], header_color)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text

    return table

def create_handbook():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ==================== 封面 ====================
    for _ in range(3):
        doc.add_paragraph()

    # 顶部装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(198, 40, 40)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("商业讲师信任护城河")
    run.bold = True
    run.font.size = Pt(44)
    run.font.color.rgb = RGBColor(198, 40, 40)
    run.font.name = '微软雅黑'

    # 副标题
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("学员手册")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(66, 66, 66)
    run.font.name = '微软雅黑'

    doc.add_paragraph()
    doc.add_paragraph()

    # 底部装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = RGBColor(198, 40, 40)
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 学员信息栏
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell1 = info_table.cell(0, 0)
    cell1.text = "学员姓名：_______________________"
    cell1.paragraphs[0].runs[0].font.size = Pt(14)

    cell2 = info_table.cell(0, 1)
    cell2.text = "课程日期：_______________________"
    cell2.paragraphs[0].runs[0].font.size = Pt(14)

    cell3 = info_table.cell(1, 0)
    cell3.text = "联系方式：_______________________"
    cell3.paragraphs[0].runs[0].font.size = Pt(14)

    cell4 = info_table.cell(1, 1)
    cell4.text = "所属公司：_______________________"
    cell4.paragraphs[0].runs[0].font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 版权声明
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("本手册属于私人文件，未经许可不得翻印")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # 底部信息
    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p2.add_run("版权声明：本课程版权归罗宏伟所有，翻版必究")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 分页
    doc.add_page_break()

    # ==================== 课程目标页 ====================
    h = doc.add_heading('课程目标', level=1)
    h.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    doc.add_paragraph()

    # 知识目标
    doc.add_heading('知识目标', level=2)
    knowledge_obj = [
        "理解信任资产与流量资产在商业培训中的本质区别",
        "掌握培训师信任积累的完整机制（存取款模型）",
        "识别行业内的关键信任行为节点",
        "认知个人品牌与机构信任的两种不同信任逻辑"
    ]
    for obj in knowledge_obj:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj)

    doc.add_paragraph()

    # 技能目标
    doc.add_heading('技能目标', level=2)
    skill_obj = [
        "运用课前调研技术获取甲方真实需求（信任测试技术）",
        "掌握报价谈判中的信任浓度判断与应对策略",
        "建立老客户转介绍网络的结构化方法",
        "设计课后跟进机制延伸信任链条",
        "具备拒绝不合适项目的判断框架与话术"
    ]
    for obj in skill_obj:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj)

    doc.add_paragraph()

    # 态度目标
    doc.add_heading('态度目标', level=2)
    attitude_obj = [
        '建立"信任优先"而非"流量优先"的职业心态',
        "培养长期主义视角，抵御短期流量焦虑",
        '塑造"说NO比说YES更能建立信任"的逆向思维',
        '认同"熬过看不见反馈的那几年"是行业生存的关键'
    ]
    for obj in attitude_obj:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj)

    doc.add_paragraph()

    # 核心公理
    h2 = doc.add_heading('核心公理', level=2)
    h2.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    axiom_p = doc.add_paragraph()
    axiom_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = axiom_p.add_run("流量能带来陌生人的关注，不能带来陌生人的托付；培训师真正的生意，从来发生在客户已经决定信任你之后，而不是之前。")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(198, 40, 40)

    # 课程结构表
    doc.add_paragraph()
    h2 = doc.add_heading('课程结构', level=2)
    h2.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    headers = ["模块", "主题", "时长"]
    rows = [
        ["PART 1", "认知与实操篇：信任是怎么攒下来的", "6小时"],
        ["前言", "信任是培训师的生死之脉", "30分钟"],
        ["第一章至第七章", "流量、现场、关系、调研、报价、拒绝", "6小时"],
        ["PART 2", "心态与长期篇：熬过看不见反馈的那几年", "6小时"],
        ["第八章至第十五章", "差异化、品牌、复利、竞争、细节、耐心、同行、家庭", "6小时"],
    ]
    add_table_with_style(doc, headers, rows)

    # 分页
    doc.add_page_break()

    # ==================== 各章节内容页 ====================

    def add_chapter_page(doc, chapter_info):
        """添加章节页面"""
        num = chapter_info["num"]
        title = chapter_info["title"]
        duration = chapter_info["duration"]

        # 章节标题
        h = doc.add_heading(f'{num}：{title}', level=1)
        h.runs[0].font.color.rgb = RGBColor(198, 40, 40)

        # 时长标签
        dur_p = doc.add_paragraph()
        run = dur_p.add_run(f"⏱ 学习时长：{duration}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # 获取公理
        axiom_key = num.replace("课程介绍", "").strip()
        if axiom_key in CHAPTERaxioms:
            axiom = CHAPTERaxioms[axiom_key]
            axiom_p = doc.add_paragraph()
            run = axiom_p.add_run(f"💡 {axiom}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(66, 66, 66)
            axiom_p.paragraph_format.left_indent = Inches(0.3)

        # 学习笔记区
        doc.add_paragraph()
        h2 = doc.add_heading('📝 学习笔记', level=2)
        h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

        # 添加空白行供学员记录
        for _ in range(4):
            line_p = doc.add_paragraph()
            line_p.add_run("_" * 80)
            line_p.runs[0].font.color.rgb = RGBColor(220, 220, 220)

        # 核心收获
        doc.add_paragraph()
        h2 = doc.add_heading('🌟 本章核心收获', level=2)
        h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

        for _ in range(3):
            line_p = doc.add_paragraph()
            line_p.add_run("_" * 80)
            line_p.runs[0].font.color.rgb = RGBColor(220, 220, 220)

        # 工具表单区
        doc.add_paragraph()
        h2 = doc.add_heading('🛠 工具表单区', level=2)
        h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

        # 简化版工具表格
        tool_headers = ["工具名称", "用途", "关键动作"]
        tool_rows = [
            ["", "", ""],
            ["", "", ""],
        ]
        add_table_with_style(doc, tool_headers, tool_rows, "424242")

        # 练习区
        doc.add_paragraph()
        h2 = doc.add_heading('✏️ 练习区', level=2)
        h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

        exercise_p = doc.add_paragraph()
        run = exercise_p.add_run("练习记录：")
        run.bold = True

        for _ in range(4):
            line_p = doc.add_paragraph()
            line_p.add_run("_" * 80)
            line_p.runs[0].font.color.rgb = RGBColor(220, 220, 220)

        # 课后作业区
        doc.add_paragraph()
        h2 = doc.add_heading('📋 课后作业', level=2)
        h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

        hw_p = doc.add_paragraph()
        run = hw_p.add_run("作业记录：")
        run.bold = True

        for _ in range(3):
            line_p = doc.add_paragraph()
            line_p.add_run("_" * 80)
            line_p.runs[0].font.color.rgb = RGBColor(220, 220, 220)

        # 分页
        doc.add_page_break()

    # 添加所有章节
    for chapter in CHAPTERS:
        add_chapter_page(doc, chapter)

    # ==================== 课程总结页 ====================
    h = doc.add_heading('课程总结', level=1)
    h.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    doc.add_paragraph()

    summary_text = """
信任的护城河，不只是一条业务上的护城河，是一条人生态度的护城河。

它最终留给你的，不是客户名单，不是收入数字，
而是你在这个过程里，被磨出来的判断力，
和在家人眼中的那个"说话算话的人"的形象。

这两样东西，才是你在这个行业里，真正属于自己的资产。
    """

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(summary_text.strip())
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(66, 66, 66)

    # 信任资产盘点
    doc.add_paragraph()
    h2 = doc.add_heading('信任资产盘点表', level=2)
    h2.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    headers = ["维度", "具体内容", "自我评估"]
    rows = [
        ["信任资产积累", "过去一年你积累的最重要的3个信任资产", ""],
        ["信任浓度最高客户", "对你信任浓度最高的3个客户", ""],
        ["信任存款行为", "你做过最有效的3个信任存款行为", ""],
        ["信任取款行为", "你需要改进的3个信任取款行为", ""],
        ["下一步行动", "未来3个月你计划做什么", ""],
    ]
    add_table_with_style(doc, headers, rows)

    # 行动计划
    doc.add_paragraph()
    h2 = doc.add_heading('90天行动计划', level=2)
    h2.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    headers = ["时间", "行动目标", "具体行动", "评估标准"]
    rows = [
        ["第1-30天", "", "", ""],
        ["第31-60天", "", "", ""],
        ["第61-90天", "", "", ""],
    ]
    add_table_with_style(doc, headers, rows)

    # 分页
    doc.add_page_break()

    # ==================== 附录页 ====================
    h = doc.add_heading('附录', level=1)
    h.runs[0].font.color.rgb = RGBColor(198, 40, 40)

    doc.add_paragraph()

    # 推荐阅读
    h2 = doc.add_heading('推荐阅读', level=2)
    h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

    readings = [
        "《定位》艾·里斯 / 杰克·特劳特",
        "《被讨厌的勇气》岸见一郎",
        "《反脆弱》纳西姆·塔勒布",
        "《高效能人士的七个习惯》史蒂芬·柯维",
        "《合作竞争》拜瑞·内勒巴夫",
        "《巴菲特传》",
        "《信任的速度》史蒂芬·柯维",
    ]
    for r in readings:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(r)

    # 课程金句收藏
    doc.add_paragraph()
    h2 = doc.add_heading('课程金句收藏', level=2)
    h2.runs[0].font.color.rgb = RGBColor(66, 66, 66)

    gold_quotes = [
        "流量给你的是这个月安心的感觉，信任给你的是十年后还有饭吃的底气。",
        "差异化不是讲得比别人好，是没人能讲你能讲的那部分。",
        "大浪淘沙淘掉的从来不是讲得不好的人，是熬不住看不见回报的那几年就退场的人。",
        "客户请你回来，不是因为你上一次讲得多精彩，是因为你在讲台之外，也是个说话算话的人。",
        "你放弃的那一单，在外人眼里是钱没赚到，只有你自己知道，那是护城河的一块砖。",
    ]

    for i, q in enumerate(gold_quotes, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q}")
        run.font.size = Pt(11)
        run.italic = True

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"学员手册已生成: {OUTPUT_PATH}")

if __name__ == "__main__":
    create_handbook()

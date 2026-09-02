#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为学员手册 docx 加页眉页脚、目录跳转样式"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

def add_header_footer(docx_path, header_text, footer_left):
    doc = Document(docx_path)

    # ---- 1. 设置全局页面边距（A4） ----
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- 2. 页眉 ----
    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header
        # 清空原有内容
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        hp = header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(header_text)
        run.font.size = Pt(9)
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        # 页眉下加一条横线
        pPr = hp._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ---- 3. 页脚 ----
    for section in doc.sections:
        footer = section.footer
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = fp.add_run(f"{footer_left}  ·  第 ")
        run1.font.size = Pt(9)
        run1.font.name = "Microsoft YaHei"
        run1.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr1 = run1._element.get_or_add_rPr()
        rFonts1 = rPr1.find(qn('w:rFonts'))
        if rFonts1 is None:
            rFonts1 = OxmlElement('w:rFonts')
            rPr1.append(rFonts1)
        rFonts1.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # PAGE 字段
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run2 = fp.add_run()
        run2.font.size = Pt(9)
        run2.font.name = "Microsoft YaHei"
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr2 = run2._element.get_or_add_rPr()
        rFonts2 = rPr2.find(qn('w:rFonts'))
        if rFonts2 is None:
            rFonts2 = OxmlElement('w:rFonts')
            rPr2.append(rFonts2)
        rFonts2.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run2._element.append(fldChar1)
        run2._element.append(instrText)
        run2._element.append(fldChar2)

        run3 = fp.add_run(" 页 / 共 ")
        run3.font.size = Pt(9)
        run3.font.name = "Microsoft YaHei"
        run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr3 = run3._element.get_or_add_rPr()
        rFonts3 = rPr3.find(qn('w:rFonts'))
        if rFonts3 is None:
            rFonts3 = OxmlElement('w:rFonts')
            rPr3.append(rFonts3)
        rFonts3.set(qn('w:eastAsia'), 'Microsoft YaHei')

        # NUMPAGES 字段
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run4 = fp.add_run()
        run4.font.size = Pt(9)
        run4.font.name = "Microsoft YaHei"
        run4.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr4 = run4._element.get_or_add_rPr()
        rFonts4 = rPr4.find(qn('w:rFonts'))
        if rFonts4 is None:
            rFonts4 = OxmlElement('w:rFonts')
            rPr4.append(rFonts4)
        rFonts4.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run4._element.append(fldChar3)
        run4._element.append(instrText2)
        run4._element.append(fldChar4)

        run5 = fp.add_run(" 页")
        run5.font.size = Pt(9)
        run5.font.name = "Microsoft YaHei"
        run5.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        rPr5 = run5._element.get_or_add_rPr()
        rFonts5 = rPr5.find(qn('w:rFonts'))
        if rFonts5 is None:
            rFonts5 = OxmlElement('w:rFonts')
            rPr5.append(rFonts5)
        rFonts5.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.save(docx_path)
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    # 学员手册
    add_header_footer(
        "D:/2026年课程/竞越/培训需求分析和课程设计/完整课程包/04-学员手册/学员手册_v1.0.docx",
        "培训需求分析和课程设计 · 学员手册 v1.0",
        "学员手册 v1.0"
    )

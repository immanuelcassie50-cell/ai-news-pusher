# -*- coding: utf-8 -*-
"""
登攀者——AI时代的授权赋能领导力 讲师手册生成脚本
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Color definitions
BLUE_SCRIPT = RGBColor(0x1A, 0x56, 0xDB)      # Instructor scripts (话术)
RED_TIME = RGBColor(0xDC, 0x26, 0x26)         # Time tags (时间标注)
DARK_BLUE_KEY = RGBColor(0x16, 0x3A, 0x64)    # Key points (关键点)
GREEN_NOTE = RGBColor(0x59, 0x73, 0x5B)       # Instructor notes (讲师备注)
PURPLE_AI = RGBColor(0x7C, 0x3A, 0xED)        # AI upgrade content (AI时代升级)
HEADER_ROW_COLOR = "E5E7EB"                   # Table header background

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table_header_row(table, headers):
    """Add a header row with shading"""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_shading(cell, HEADER_ROW_COLOR)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

def set_paragraph_spacing(paragraph, space_before=0, space_after=6):
    """Set paragraph spacing"""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)

def add_heading(doc, text, level=1, color=None):
    """Add a heading with optional color"""
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_time_tag(doc, time_text):
    """Add a time tag in red bold"""
    p = doc.add_paragraph()
    run = p.add_run(time_text)
    run.font.color.rgb = RED_TIME
    run.bold = True
    run.font.size = Pt(11)
    set_paragraph_spacing(p)
    return p

def add_key_point(doc, text):
    """Add a key point in dark blue bold"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_BLUE_KEY
    run.bold = True
    run.font.size = Pt(11)
    set_paragraph_spacing(p)
    return p

def add_script(doc, text):
    """Add instructor script in blue italic"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = BLUE_SCRIPT
    run.italic = True
    run.font.size = Pt(11)
    set_paragraph_spacing(p)
    return p

def add_ai_content(doc, text):
    """Add AI upgrade content in purple italic"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = PURPLE_AI
    run.italic = True
    run.font.size = Pt(11)
    set_paragraph_spacing(p)
    return p

def add_instructor_note(doc, text):
    """Add instructor note in green small text"""
    p = doc.add_paragraph()
    run = p.add_run(f"【讲师备注】{text}")
    run.font.color.rgb = GREEN_NOTE
    run.font.size = Pt(9)
    set_paragraph_spacing(p)
    return p

def add_normal_text(doc, text):
    """Add normal paragraph text"""
    p = doc.add_paragraph(text)
    set_paragraph_spacing(p)
    return p

def add_prep_notes_area(doc):
    """Add a grayed area for preparation notes"""
    p = doc.add_paragraph()
    run = p.add_run("【备课笔记区】")
    run.font.color.rgb = GREEN_NOTE
    run.font.size = Pt(9)
    run.italic = True
    set_paragraph_spacing(p, space_after=30)
    return p

def create_document():
    """Create the complete instructor manual document"""
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    # ========== COVER PAGE ==========
    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("登攀者")
    run.font.size = Pt(44)
    run.bold = True
    run.font.color.rgb = DARK_BLUE_KEY

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI时代的授权赋能领导力")
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE_SCRIPT

    doc.add_paragraph()  # Spacer

    # Document type
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_type.add_run("讲师手册（完整版）")
    run.font.size = Pt(18)
    run.bold = True

    doc.add_paragraph()  # Spacer
    doc.add_paragraph()  # Spacer

    # Info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("开发者", "罗宏伟"),
        ("版本", "1.0"),
        ("适用对象", "企业管理者、HRBP、培训师"),
        ("课程时长", "2天 / 13小时"),
        ("班级规模", "16-24人"),
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading(doc, "目录", level=1)

    toc_items = [
        "第一部分 Day 1 ———",
        "1. 讲师资质要求与提醒",
        "2. 课程概览与学习目标",
        "3. 教学方法说明",
        "4. 教具与材料清单",
        "5. 开场：从答案到问题",
        "6. Part 1 前段：教练思维与信任环境",
        "7. Part 1 后段：GUIDE模型",
        "8. Part 2 前段：聆听与提问",
        "9. Part 2 后段：反馈与认同",
        "",
        "第二部分 Day 2 ———",
        "10. Day 2 开场与回温",
        "11. Part 3 前段：DIRECT模型",
        "12. Part 3 后段：DIRECT演练",
        "13. Part 4 前段：综合演练",
        "14. Part 4 后段：行动计划",
        "15. 收尾与教练承诺",
        "",
        "附录 ———",
        "附录A：AI时代六类场景卡",
        "附录B：三层复盘法速查表",
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ========== PART 1 - DAY 1 ==========
    add_heading(doc, "第一部分 Day 1", level=1, color=DARK_BLUE_KEY)

    # --- Section 1: 讲师资质要求 ---
    add_heading(doc, "1. 讲师资质要求与提醒", level=2)

    add_key_point(doc, "核心要求")

    req_table = doc.add_table(rows=4, cols=2)
    req_table.style = 'Table Grid'
    add_table_header_row(req_table, ["要求维度", "具体说明"])

    req_data = [
        ("专业背景", "具备企业教练或领导力发展培训经验，了解成人学习原理"),
        ("教练资质", "持有ICF或同等机构认证，或有100小时以上教练对话实践"),
        ("AI理解", "对AI工具和AI时代工作变化有亲身使用经验和真实见解"),
    ]
    for i, (dim, desc) in enumerate(req_data):
        req_table.rows[i+1].cells[0].text = dim
        req_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    add_key_point(doc, "AI时代特殊要求")
    add_ai_content(doc, "需要对AI工具和AI时代工作的变化有真实的了解，才能在教练对话的示范里用自然的语言触及这个话题。如果AI对你来说是抽象概念，学员会感受到，特别是高潜员工群体，他们对AI的熟悉度往往超过他们的管理者。课前，确认你自己对AI工具有亲身使用经验，有自己真实的感受和看法。")

    doc.add_paragraph()

    # --- Section 2: 课程概览与学习目标 ---
    add_heading(doc, "2. 课程概览与学习目标", level=2)

    add_heading(doc, "Day 1 / Day 2 学习弧线", level=3)

    learning_arc_table = doc.add_table(rows=3, cols=3)
    learning_arc_table.style = 'Table Grid'
    add_table_header_row(learning_arc_table, ["维度", "Day 1", "Day 2"])

    arc_data = [
        ("阶段定位", "信念重建阶段", "技能整合阶段"),
        ("核心主题", "打破'领导=给答案'；建立教练思维；GUIDE模型入门", "DIRECT模型学习；综合场景演练；行动计划制定"),
    ]
    for i, (dim, d1, d2) in enumerate(arc_data):
        learning_arc_table.rows[i+1].cells[0].text = dim
        learning_arc_table.rows[i+1].cells[1].text = d1
        learning_arc_table.rows[i+1].cells[2].text = d2

    doc.add_paragraph()

    add_heading(doc, "Day 1 学习目标", level=3)
    day1_goals = [
        "区分业绩表现型和发展型教练的应用场景",
        "理解\"装备vs缆车\"的登攀者AI时代新框架",
        "掌握GUIDE发展型教练五步模型",
        "建立信任是教练工作前提的认知",
        "学会用\"双峰教练视角\"定位被教练者的发展方向",
        "理解有意识聆听的四个层次（含AI时代身份信号层次）",
        "掌握AI时代突破性提问的设计思路",
        "掌握教练式反馈的结构（含AI时代版本）",
        "理解并演练认同技巧",
    ]
    for goal in day1_goals:
        p = doc.add_paragraph(f"• {goal}")
        set_paragraph_spacing(p, space_after=2)

    doc.add_paragraph()

    add_heading(doc, "Day 2 学习目标", level=3)
    day2_goals = [
        "掌握DIRECT业绩表现型教练六步模型",
        "深入理解AI时代版本的DATA收集（人类贡献归因）和REQUIREMENT设定（双维度绩效期望）",
        "在更复杂、更接近真实工作的场景里，综合运用GUIDE和DIRECT",
        "学会根据场景选择合适的教练模式",
        "为真实的高潜下属制作完整的教练对话行动计划",
        "建立问责伙伴系统",
    ]
    for goal in day2_goals:
        p = doc.add_paragraph(f"• {goal}")
        set_paragraph_spacing(p, space_after=2)

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 3: 教学方法说明 ---
    add_heading(doc, "3. 教学方法说明", level=2)

    add_heading(doc, "教学模式配比表格", level=3)
    method_table = doc.add_table(rows=4, cols=3)
    method_table.style = 'Table Grid'
    add_table_header_row(method_table, ["模式", "占比", "说明"])

    method_data = [
        ("讲授", "25%", "框架讲解、模型说明、概念澄清；讲师主导但不超过20分钟连续讲授"),
        ("体验", "35%", "讲师示范、学员作为被教练者体验、感受反馈；让学员先感受再学习"),
        ("演练", "40%", "角色扮演、观察者记录、全班复盘；大量练习建立肌肉记忆"),
    ]
    for i, (mode, pct, desc) in enumerate(method_data):
        method_table.rows[i+1].cells[0].text = mode
        method_table.rows[i+1].cells[1].text = pct
        method_table.rows[i+1].cells[2].text = desc

    doc.add_paragraph()

    add_heading(doc, "核心教学循环", level=3)
    add_script(doc, "讲师示范 → 学员体验 → 学员演练 → 观察者反馈 → 全班复盘")
    add_script(doc, "（先示范，再练习）    （先练习，后复盘）")

    doc.add_paragraph()

    add_heading(doc, "三层复盘法（每个演练环节后使用）", level=3)

    debrief_table = doc.add_table(rows=4, cols=2)
    debrief_table.style = 'Table Grid'
    add_table_header_row(debrief_table, ["层次", "核心问题与聚焦"])

    debrief_data = [
        ("第一层：What（发生了什么）", "\"被教练者们——在这段时间里，有没有一个时刻，你感到一种真正的洞察出现了？那是什么时刻？\"\n聚焦被教练者的体验，不聚焦教练者的技术"),
        ("第二层：So What（这意味着什么）", "\"管理者们——今天演练里，最难的部分是什么？\"\n聚焦教练者的挑战和成长点"),
        ("第三层：Now What（下一步是什么）", "\"从今天的演练里，你带走了什么可以用在下次和他对话里的？\"\n聚焦具体应用，建立与真实工作的连接"),
    ]
    for i, (level, desc) in enumerate(debrief_data):
        debrief_table.rows[i+1].cells[0].text = level
        debrief_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 4: 教具与材料清单 ---
    add_heading(doc, "4. 教具与材料清单", level=2)

    add_heading(doc, "讲师材料", level=3)
    instructor_table = doc.add_table(rows=6, cols=3)
    instructor_table.style = 'Table Grid'
    add_table_header_row(instructor_table, ["材料名称", "数量", "用途"])

    instructor_materials = [
        ("讲师手册（完整版）", "1本", "全部讲授依据"),
        ("教学PPT", "1套", "视觉辅助，按模块分节"),
        ("演练观察清单", "每组1份", "记录演练关键瞬间"),
        ("讲师提示卡", "1套", "关键话术提醒，口袋大小"),
        ("AI时代六类场景卡", "每组1套", "Part 3/4演练用"),
    ]
    for i, (name, qty, use) in enumerate(instructor_materials):
        instructor_table.rows[i+1].cells[0].text = name
        instructor_table.rows[i+1].cells[1].text = qty
        instructor_table.rows[i+1].cells[2].text = use

    doc.add_paragraph()

    add_heading(doc, "学员材料", level=3)
    participant_table = doc.add_table(rows=6, cols=3)
    participant_table.style = 'Table Grid'
    add_table_header_row(participant_table, ["材料名称", "数量", "用途"])

    participant_materials = [
        ("学员手册", "每人1本", "工作页+应用卡，课前发放"),
        ("GUIDE模型应用卡", "每人1张", "可裁剪随身携带"),
        ("DIRECT模型应用卡", "每人1张", "可裁剪随身携带"),
        ("AI时代六类场景卡", "每组1套", "Part 4演练用"),
        ("教练行动计划工作纸", "每人2张", "Part 4制作计划"),
    ]
    for i, (name, qty, use) in enumerate(participant_materials):
        participant_table.rows[i+1].cells[0].text = name
        participant_table.rows[i+1].cells[1].text = qty
        participant_table.rows[i+1].cells[2].text = use

    doc.add_paragraph()

    add_heading(doc, "场地与设备", level=3)
    venue_table = doc.add_table(rows=8, cols=2)
    venue_table.style = 'Table Grid'
    add_table_header_row(venue_table, ["项目", "要求"])

    venue_items = [
        ("教室布置", "U型或小组圈形，便于讨论和演练"),
        ("白板/大白纸", "至少2面，记录关键洞察"),
        ("投影设备", "1套，显示PPT"),
        ("音响设备", "1套，背景音乐用"),
        ("计时器", "2个，控制演练时间"),
        ("便利贴", "多种颜色，开场活动用"),
        ("签字笔", "每组2支，白板记录用"),
    ]
    for i, (item, req) in enumerate(venue_items):
        venue_table.rows[i+1].cells[0].text = item
        venue_table.rows[i+1].cells[1].text = req

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 5: 开场 ---
    add_heading(doc, "5. 开场：从答案到问题", level=2)
    add_time_tag(doc, "【09:00-09:30 | 30分钟】")

    add_key_point(doc, "本环节核心目标")
    add_normal_text(doc, "打破\"领导=给答案\"的预设认知；建立对教练型领导力的初步好奇；为全天的学习建立一个真实的情感锚点。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "便利贴（每人2-3张）、白板")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "开场不讲理论，先做一件让学员进入真实状态的事——让他们回忆一次被有效领导的真实体验。这个体验，是今天课程所有内容的情感起点。如果学员在开场就进入了理智状态（\"领导力的定义是……\"），他们对教练工具的感受会大打折扣。")

    add_heading(doc, "5.1 被有效领导的记忆", level=3)
    add_time_tag(doc, "【09:00-09:15 | 15分钟】")

    add_key_point(doc, "步骤1：被有效领导的记忆")
    add_script(doc, "\"在我们开始之前，我想先做一件事。\"")
    add_script(doc, "\"想一个人——在你整个职业生涯里，有一个领导者、导师、或者前辈，他/她说了一句话、问了一个问题，或者做了一件事，让你突破了你自己以为做不到的限制。那个时刻，是什么？\"")
    add_script(doc, "\"给自己2分钟，在便利贴上写下来：那个时刻，他做了什么（不是你做了什么，是他做了什么）？\"")

    add_instructor_note(doc, "2分钟独立写，不打扰")

    add_script(doc, "\"现在找旁边的人，分享你写的那个时刻，2分钟，说给他听。\"")

    add_instructor_note(doc, "4分钟配对")

    add_key_point(doc, "全班汇集（8分钟）")
    add_script(doc, "\"我们来听几个——他/她做了什么，让你突破了？\"")
    add_script(doc, "（收集5-6个，不评论，只在白板上记录关键词）")
    add_instructor_note(doc, "通常的关键词会是：问了我一个问题、相信我能做到、没有直接告诉我答案、让我自己想、看到了我自己没看到的、让我感觉被理解了……")

    add_script(doc, "\"注意——你们写的，大多数不是'他告诉了我答案'，也不是'他给了我一套方法'。\"")
    add_script(doc, "\"是问题，是相信，是让你自己看见。\"")
    add_script(doc, "\"这就是教练型领导力的本质——不是输出你的答案，而是激发对方的洞察和资源。今天，我们来学习怎么做到这件事。而且是在AI时代，高潜员工最需要这种领导方式的时候。\"")

    add_heading(doc, "5.2 AI时代的新命题引入", level=3)
    add_time_tag(doc, "【09:15-09:30 | 10分钟】")

    add_script(doc, "\"AI时代，教练型领导力面临一个新的命题。在你们分享的那些突破时刻里——如果那个突破的内容，现在AI可以直接给出答案呢？\"")
    add_script(doc, "\"比如：你的前辈当年问你'你觉得应该用哪种营销策略？'——现在，AI可以在五秒内给出五种营销策略建议。这让教练失去价值了吗？\"")

    add_instructor_note(doc, "停顿，等学员反应")

    add_script(doc, "\"不。但它改变了教练的核心战场。\"")
    add_script(doc, "\"AI给不了答案的地方，是：你在这份工作里，真正的人类判断力在哪里？你做得最好的时候，是什么让你和AI生成的答案不同？你在AI越来越强的时代，什么是你真正在攀登的那座山？\"")
    add_script(doc, "\"今天，我们要装备的是一套教练工具，帮助你在这个层面上支持你的高潜员工。这是比以前更有价值、也更难的教练工作。\"")
    add_script(doc, "\"我们开始。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 6: Part 1 前段 ---
    add_heading(doc, "6. Part 1 前段：教练思维与信任环境", level=2)
    add_time_tag(doc, "【09:30-11:00 | 90分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "区分业绩表现型和发展型教练的应用场景；理解\"装备vs缆车\"的登攀者AI时代新框架；建立信任是教练工作前提的认知；识别高潜员工的AI时代心智阻碍；学会用\"双峰教练视角\"定位被教练者的发展方向。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "幻灯片（两种教练类型+登攀者AI新框架+信任环境+教练关系罗盘+双峰视角）、双峰定位工作纸（手册P7）")

    add_heading(doc, "6.1 两种教练的区分", level=3)
    add_time_tag(doc, "【09:30-09:40 | 10分钟】")

    add_script(doc, "\"教练，在工作里有两种不同的用法。\"")

    add_key_point(doc, "展示（PPT）：")

    coach_table = doc.add_table(rows=3, cols=2)
    coach_table.style = 'Table Grid'
    add_table_header_row(coach_table, ["类型", "详细内容"])

    coach_types = [
        ("业绩表现型教练（Performance Coaching）", "目的：解决特定的业绩问题，改变具体的行为，达成预期的绩效标准\n触发点：有具体的、可观察的行为或结果问题\n典型问句：\"我们来看这次项目里，发生了什么，接下来怎么改……\"\n代表模型：DIRECT"),
        ("发展型教练（Developmental Coaching）", "目的：挖掘潜力、激发洞察、帮助员工找到突破和成长的方向\n触发点：员工有更大的潜力尚未发挥，或面临一个发展上的瓶颈\n典型问句：\"你在这里的最大潜力，你自己觉得是什么？什么在阻碍它？\"\n代表模型：GUIDE"),
    ]
    for i, (coach_type, detail) in enumerate(coach_types):
        coach_table.rows[i+1].cells[0].text = coach_type
        coach_table.rows[i+1].cells[1].text = detail

    doc.add_paragraph()

    add_script(doc, "\"这两种教练，适用于不同的情境，用的模型不同，说话的方式也不同。搞混了，效果会大打折扣。\"")
    add_script(doc, "\"今天上午，我们专注发展型教练和GUIDE；下午我们学核心技巧；明天我们学业绩表现型教练和DIRECT。\"")
    add_script(doc, "\"但在开始之前，有一个概念我想先讲——因为它会让你今天接触的所有发展型教练工具，都有一个清晰的AI时代的方向感。\"")

    add_heading(doc, "6.2 登攀者的AI时代新框架——装备vs缆车", level=3)
    add_time_tag(doc, "【09:40-09:55 | 15分钟】")

    add_key_point(doc, "展示登攀者的图像和比喻（PPT）")

    add_script(doc, "\"这门课叫登攀者——帮助高潜员工攀登到他们自己以为去不到的高度。\"")
    add_script(doc, "\"AI时代，攀登这件事有了一个新的问题——\"")
    add_script(doc, "\"AI是一种强大的新装备。合理使用它，攀登者可以到达以前难以到达的高度：更快地积累信息，更高效地完成某些工作，把更多的时间和精力放在真正的核心判断上。\"")
    add_script(doc, "\"但AI也可以成为一种缆车。\"")
    add_script(doc, "\"坐进缆车，攀登者也可以到达山顶，看到风景。但他自己的攀登肌肉没有得到练习，他的判断力、他在复杂地形里的应变能力、他对那座山的深度了解——这些都没有增长。\"")
    add_script(doc, "\"下了缆车，他看起来和自己攀登上来的人一样，站在同一个高度。但如果下一座山没有缆车，他不行了。\"")
    add_script(doc, "\"你管辖的高潜员工，现在有多少在用AI作为装备攀登，有多少在坐缆车？\"")

    add_instructor_note(doc, "停顿，让问题在空气里停留")

    add_script(doc, "\"作为教练型领导者，这是你最需要有清晰视角的问题。今天接触的所有工具，本质上都在帮助你回答这个问题，以及帮助你的员工做出自己的选择。\"")

    add_key_point(doc, "展示三类AI时代高潜员工心智阻碍（PPT）：")

    obstacles_table = doc.add_table(rows=4, cols=2)
    obstacles_table.style = 'Table Grid'
    add_table_header_row(obstacles_table, ["类型", "描述与教练方向"])

    obstacles = [
        ("价值焦虑型", "表现：持续担心\"AI替代了我，我还有什么价值\"，影响投入度和方向感。\n教练方向：帮助他找到AI真正无法替代的人类深度贡献，重建职业价值锚点。"),
        ("过度依赖型", "表现：用AI承接了大量本应由自己思考的工作，短期产出不错，但判断力和创造力在退化。\n教练方向：帮助他识别哪些工作是人类深度成长的机会，AI在这里应该是放大器而不是替代物。"),
        ("无效抵触型", "表现：拒绝使用AI工具，以维护专业认同，但在效率上开始落后，无法利用AI放大自己的真正优势。\n教练方向：帮助他探索\"如果AI承接了X部分，你的时间和精力会流向哪里\"，从防御转向主动。"),
    ]
    for i, (ob_type, desc) in enumerate(obstacles):
        obstacles_table.rows[i+1].cells[0].text = ob_type
        obstacles_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    add_heading(doc, "6.3 双峰教练视角", level=3)
    add_time_tag(doc, "【09:55-10:15 | 20分钟】")

    add_script(doc, "\"在进入GUIDE模型之前，我想给你们一个视角框架，帮助你在每一次发展型教练对话之前，先想清楚方向。\"")
    add_script(doc, "\"我叫它双峰教练视角。\"")
    add_key_point(doc, "展示双峰图（PPT，两个山峰，左边大、右边小）")

    add_script(doc, "\"发展型教练，在AI时代需要帮助被教练者同时思考两座山。\"")
    add_script(doc, "\"左边这座山，我叫它**人类深度峰**——这个人的哪些能力、判断力、创造力和关系能力，是AI目前真正做不到的？这是他最需要攀登的山，越高，越不可替代。\"")
    add_script(doc, "\"右边这座山，我叫它**AI杠杆峰**——他如何有效地使用AI工具，来放大他的人类深度贡献，而不是替代它？这是效率山，重要，但不是主要的。\"")
    add_script(doc, "\"教练的工作：帮助被教练者找到他的人类深度峰，然后装备好AI杠杆，攀登得更高更快。\"")
    add_script(doc, "\"不是帮他选个好的缆车，是帮他搞清楚那座山在哪里，然后把装备准备好。\"")

    add_key_point(doc, "实践练习（10分钟）")
    add_script(doc, "\"翻到手册P7，双峰定位工作纸。\"")
    add_script(doc, "\"想你为这次课程准备的那名高潜下属——用5分钟，在两个峰上写关键词：他的人类深度峰，你认为在哪里（什么能力、什么判断力、什么他特别擅长而AI做不到的）？他的AI杠杆峰，目前他用AI最多的是什么，还有哪些地方可以更有效地用？\"")

    add_instructor_note(doc, "5分钟，独立写")

    add_script(doc, "\"和旁边的人分享——他的人类深度峰，你觉得最核心的是什么？你感觉他自己知道吗？\"")

    add_instructor_note(doc, "5分钟配对分享")

    add_heading(doc, "6.4 建立信任与教练环境", level=3)
    add_time_tag(doc, "【10:15-10:30 | 15分钟】")

    add_script(doc, "\"发展型教练对话，需要一种特殊的环境——被教练者愿意把真实的自己，包括困惑、恐惧和不确定，带进对话里。这只有在足够安全和信任的情况下才会发生。\"")
    add_script(doc, "\"建立信任，是教练工作的前提，不是可选项。\"")
    add_key_point(doc, "展示信任建立的三个核心条件（PPT）：")

    trust_table = doc.add_table(rows=4, cols=2)
    trust_table.style = 'Table Grid'
    add_table_header_row(trust_table, ["条件", "详细内容"])

    trust_conditions = [
        ("条件一：意图可信", "员工感受到管理者的教练，是真心为了他的成长，而不是为了评估、控制或找到弱点。\nAI时代更新：高潜员工对意图的感知非常敏感——如果他感受到你在问AI相关问题，是为了判断他\"够不够格\"，他会防御；如果他感受到你是真的好奇\"在AI时代，你看到自己最想发展什么\"，他会打开。"),
        ("条件二：能力可信", "员工相信管理者有能力帮助他成长，至少在某个维度上有他尊重的洞察力。\nAI时代更新：高潜员工（特别是高知型）对管理者在AI相关话题上的理解水平非常敏感。你不需要是AI专家，但你需要对AI时代的工作现实有足够真实的了解，才能在教练对话里提出真正有穿透力的问题。如果你完全不了解AI，对话在这个层面会失去可信度。"),
        ("条件三：关系可信", "一段持续的、真实的关注关系，比一次完美的教练技术更有力量。\nAI时代不变：这个条件在AI时代没有改变。AI给不了真实的人际关系，这反而让真实关注的管理者在高潜员工眼里更有价值。"),
    ]
    for i, (cond, detail) in enumerate(trust_conditions):
        trust_table.rows[i+1].cells[0].text = cond
        trust_table.rows[i+1].cells[1].text = detail

    doc.add_paragraph()

    add_script(doc, "\"信任，不是在教练对话开始时才建立的。它是在每一次互动里积累的，包括你平时对这个员工的关注方式、你对他的尊重方式，以及你在AI相关话题上的真实态度。\"")

    add_heading(doc, "6.5 教练时机的识别", level=3)
    add_time_tag(doc, "【10:30-10:40 | 10分钟】")

    add_script(doc, "\"GUIDE什么时候用？不是随时随地。时机，比技巧更重要。\"")
    add_key_point(doc, "展示教练时机的四类情境（PPT）：")

    timing_table = doc.add_table(rows=5, cols=2)
    timing_table.style = 'Table Grid'
    add_table_header_row(timing_table, ["时机", "描述"])

    timing_types = [
        ("时机一：员工面临一个能力瓶颈", "他卡住了，不是因为懒，是因为真的有什么东西阻碍了他"),
        ("时机二：员工准备好了向前一步", "他表达了想要改变或突破的意愿，或者你看到他已经在准备阶段"),
        ("时机三：员工做出了一个可以深化的成功", "不只是\"干得好\"，而是\"这里有一个值得探索的洞察\""),
        ("时机四（AI时代新增）", "员工在AI时代的方向感出现了动摇——他开始问\"我在这里有什么价值\"，或者你观察到缆车信号——他在用AI走捷径而不是在攀登"),
    ]
    for i, (timing, desc) in enumerate(timing_types):
        timing_table.rows[i+1].cells[0].text = timing
        timing_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    add_script(doc, "\"不是所有对话都是教练时机。当员工急需一个立刻的答案（客户等着、决策等不了），不是教练时机。当员工的问题来自信息不足（他真的不知道这件事，你知道），直接给信息更有效。\"")
    add_script(doc, "\"时机判断，是今天学完框架之后最需要练习的能力。\"")

    add_heading(doc, "6.6 全班复盘与过渡", level=3)
    add_time_tag(doc, "【10:40-10:48 | 8分钟】")

    add_key_point(doc, "三层推进（快速）")
    add_script(doc, "\"教练思维里，今天上午到目前为止，什么让你感到意外或者产生了新的想法？\"")
    add_instructor_note(doc, "收集2-3个，不展开讨论")

    add_script(doc, "\"双峰视角里，你的分析对象，哪个峰你感觉还不清晰？\"")
    add_instructor_note(doc, "收集1个，说：这就是今天GUIDE要帮你去找的。")

    add_key_point(doc, "过渡到茶歇话术")
    add_script(doc, "\"有了方向感，现在我们进入具体的模型——GUIDE发展型教练五步法。茶歇后，我会先示范一遍完整的GUIDE对话，然后你们来演练。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 7: Part 1 后段 ---
    add_heading(doc, "7. Part 1 后段：GUIDE模型", level=2)
    add_time_tag(doc, "【11:15-12:15 | 60分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "掌握GUIDE五步及每步的核心技巧和AI时代应用；通过讲师示范感受完整发展型教练对话的流感；完成第一轮GUIDE配对演练。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "GUIDE教练模型应用卡（发放，学员保留）、幻灯片（GUIDE五步+AI时代升级+参考问题示例）")

    add_heading(doc, "7.1 GUIDE模型概览与AI时代升级", level=3)
    add_time_tag(doc, "【11:15-11:35 | 20分钟】")

    add_key_point(doc, "快速过一遍五步（每步约2-3分钟，重点在AI时代版本）")

    add_script(doc, "\"GUIDE是发展型教练的操作框架，五步，每步有清晰的目的和对应的教练技巧。\"")

    # G Step
    add_key_point(doc, "G——Goal（目标）")
    add_script(doc, "\"第一步，帮助被教练者确认他努力的目标。不是你告诉他目标，是通过提问，让他说出来——他真正在追求的是什么。\"")
    add_script(doc, "\"为什么要他自己说？因为他自己说出来的目标，他的承诺感是你外加的目标的十倍。\"")
    add_ai_content(doc, "AI时代升级：在目标确认里加一个问题——'这个目标里，你的人类贡献是什么？如果达到了这个目标，什么是AI帮不了你、必须是你自己做到的？'")
    add_script(doc, "\"这个问题，帮助你们一起定位人类深度峰。\"")
    add_key_point(doc, "参考问题（PPT）：")
    add_script(doc, "\"'你希望在这个方向上，三个月后你做到什么？'\"")
    add_script(doc, "\"'如果达到了这个目标，什么是只有你才能做到的？'\"")

    # U Step
    add_key_point(doc, "U——Understand（理解现状）")
    add_script(doc, "\"第二步，帮助被教练者清晰地看到当前的真实状况——他的优势、他的资源，以及他遇到的真实障碍。\"")
    add_script(doc, "\"这一步，讲师绝对不能替被教练者分析。你的工作是提问，是帮助他自己说清楚。\"")
    add_ai_content(doc, "AI时代升级：理解现状包括AI相关的当前状态——他现在在哪些工作里用AI，哪些不用，他对自己在AI时代的价值感受如何。")
    add_key_point(doc, "参考问题：")
    add_script(doc, "\"'你现在在这类工作里，什么部分自己来，什么部分借助工具？'\"")
    add_script(doc, "\"'你觉得目前最大的阻碍是什么？'\"")

    # I Step
    add_key_point(doc, "I——Insight（洞察）")
    add_script(doc, "\"第三步，是GUIDE里最有艺术性的一步。你通过高质量的提问，帮助被教练者产生一个他之前没有的洞察——那个'哦原来如此'的时刻。\"")
    add_script(doc, "\"洞察不是你告诉他的，是他自己想到的。教练只是提问，顿悟在被教练者那里发生。\"")
    add_ai_content(doc, "AI时代升级：帮助被教练者产生关于AI时代自己价值的洞察——'原来我在X这方面，是AI做不到的，这才是我真正最需要发展的核心。'")
    add_key_point(doc, "参考问题（重点）：")
    add_script(doc, "\"'在你做得最好的时候，你的思考里有什么，是AI想不到的？'\"")
    add_script(doc, "\"'如果AI处理了这部分，你的注意力和能量会自然流向哪里？'\"")
    add_script(doc, "\"'你最担心AI'取代'你的那部分，反过来想——那里是不是也是你真正最有深度的地方？'\"")

    # D Step
    add_key_point(doc, "D——Design（设计行动）")
    add_script(doc, "\"第四步，把洞察转化为具体的行动计划。从'明白了'到'接下来做什么'。\"")
    add_ai_content(doc, "AI时代升级：行动计划可以是双轨的——人类深度方向（我要发展的能力）和AI杠杆方向（我要更好地使用哪个工具来支持前者）。")
    add_key_point(doc, "参考问题：")
    add_script(doc, "\"'你说的这个洞察，如果化成一个接下来的具体行动，会是什么？'\"")
    add_script(doc, "\"'在人类深度这个方向，你第一步打算做什么？'\"")

    # E Step
    add_key_point(doc, "E——Enable（清除障碍）")
    add_script(doc, "\"第五步，识别和清除妨碍行动的障碍——可以是内部的（信念、恐惧），也可以是外部的（资源、时间）。\"")
    add_ai_content(doc, "AI时代升级：AI相关障碍清单里多了三类——价值焦虑、过度依赖、无效抵触。每一类都有对应的教练策略，而不是'你别担心了'式的安慰。")
    add_key_point(doc, "参考问题：")
    add_script(doc, "\"'你说会行动，最可能阻止你的是什么？'\"")
    add_script(doc, "\"'如果这个障碍没有了，你会做什么不同的事？'\"")

    add_heading(doc, "7.2 讲师示范完整GUIDE对话", level=3)
    add_time_tag(doc, "【11:35-11:50 | 15分钟】")

    add_script(doc, "\"现在我来示范一次完整的GUIDE对话。我需要一位学员来扮演被教练者——不是演戏，就是用你真实的一个职业发展困惑来参与这个对话。\"")
    add_script(doc, "\"谁愿意？\"")

    add_instructor_note(doc, "选一位开放的学员，提前30秒私下确认他愿意用真实困惑——AI时代相关的最好，但不是必须")

    add_key_point(doc, "示范对话（10分钟）：")
    add_script(doc, "讲师走一遍完整的GUIDE五步。重点：")
    add_script(doc, "G步骤里，加入\"这个目标里你的人类贡献是什么\"")
    add_script(doc, "I步骤里，用一个真正的开放式问题，等待被教练者的顿悟（不要急着补充）")
    add_script(doc, "全程只问问题，不给建议，不评判")

    add_key_point(doc, "示范后（5分钟）：")
    add_script(doc, "\"我示范完了。有三件事我想说——\"")
    add_script(doc, "\"第一，我刚才在X步骤里，问了一个问题然后停下来，等了很久。这个等待，是教练里最有价值的沉默——我在等他自己想到，而不是我想到了告诉他。\"")
    add_script(doc, "\"第二，我全程没有给一个答案或建议。如果你觉得刚才的对话里，有那么一刻你想代替他思考——记住那个冲动，那是教练最需要克服的习惯。\"")
    add_script(doc, "\"第三，I步骤里，他说'哦原来如此'的那个时刻——那个，是教练里最有价值的时刻。那不是我给的，是他自己找到的。\"")

    add_heading(doc, "7.3 配对演练——第一轮GUIDE", level=3)
    add_time_tag(doc, "【11:50-12:10 | 20分钟】")

    add_script(doc, "\"两人一组，演练一次GUIDE对话。一人教练，一人被教练——被教练者用自己真实的一个职业发展相关的困惑（不需要是AI话题，任何发展型困惑都可以）。\"")
    add_script(doc, "\"教练的任务：走完五步，重点是G和I——目标要让被教练者自己说出来，洞察要让被教练者自己找到，你只问问题。\"")
    add_script(doc, "\"12分钟教练对话，然后2分钟被教练者说感受（不是教练说自己做了什么，是被教练者说在对话里有什么感受）。\"")

    add_instructor_note(doc, "14分钟，不换角色——这次演练的重点是建立初步感受，不强求每人都演练两个角色")

    add_instructor_note(doc, "讲师巡场，观察：G步骤有没有真正让被教练者自己说出目标；I步骤有没有出现一个真正打开的时刻；教练者有没有忍住去提建议的冲动")

    add_key_point(doc, "全班复盘（6分钟）")
    add_script(doc, "\"被教练者们——在这12分钟里，有没有一个时刻，你感到'这个问题让我想到了一件我没想过的事'？那是什么时刻？\"")
    add_instructor_note(doc, "让被教练者说，不是教练者说。收集2-3个")

    add_script(doc, "\"教练者们——最难的是什么？\"")
    add_instructor_note(doc, "通常是\"忍住不给答案\"或者\"I步骤不知道该问什么\"。")

    add_script(doc, "\"午饭后，我们进入教练技巧的专项训练——特别是聆听和提问，这是I步骤得以发生的基础。吃饭的时候，想一想：你的分析对象，如果你现在和他做一次GUIDE对话，你觉得他的I步骤可能在哪里？什么问题可以触发他的顿悟？\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # ========== PART 2 - DAY 1 CONTINUES ==========
    add_heading(doc, "8. Part 2 前段：聆听与提问", level=2)
    add_time_tag(doc, "【13:15-14:45 | 90分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "理解有意识聆听的四个层次（含AI时代身份信号层次）；区分有意识提问和普通问话的差异；掌握AI时代突破性提问的设计思路；通过演练建立聆听和提问技巧。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "幻灯片（聆听四层次+提问类型对比+AI时代突破性问题示例）、学员手册P17-22（Part 2前段工作页）")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "聆听和提问，是教练工具里最容易被学员\"听懂但做不到\"的技巧。\"有意识聆听\"不是\"认真听\"——每个人都以为自己在认真听，但大多数人在听的过程里，70%的注意力是在准备自己接下来要说什么。")
    add_script(doc, "你的工作是让学员体验到\"真正被聆听\"和\"被正常对待\"之间的差别。如果他们自己体验到了这个差别，他们会真正想学这个技巧，而不只是知道它存在。")
    add_ai_content(doc, "AI时代的特殊价值：有意识聆听在AI时代对高潜员工尤其重要——因为AI可以快速给出建议，但AI不会真正聆听。一个能够深度聆听的管理者，在AI时代具有真正的不可替代性。这个事实，可以在课程里的合适时机说出来。")

    add_heading(doc, "8.1 有意识聆听体验练习", level=3)
    add_time_tag(doc, "【13:15-13:30 | 15分钟】")

    add_script(doc, "\"在讲聆听框架之前，我们先来体验两种聆听的差别。\"")
    add_script(doc, "\"两人一组，A和B。第一轮：A说一件他最近在工作里面临的真实挑战，说3分钟。B的任务是……在心里想这件事你会怎么建议他，但表面上看起来在听。\"")

    add_instructor_note(doc, "3分钟，B扮演\"假装聆听的状态\"。")

    add_script(doc, "\"好，停。B：你在做什么？A：你感受到B在聆听你吗？\"")
    add_instructor_note(doc, "快速收集感受，通常A会说\"感觉他不是真的在听\"。")

    add_script(doc, "\"第二轮：同样的事，A继续说3分钟（可以从任何地方接着说）。B这次的任务是：什么都不想，只关注A说的每一个字，听他说话时语气里情绪的变化，听他说的背后他真正关心的是什么。不分析，不评判，只在接收。\"")

    add_instructor_note(doc, "3分钟，B深度聆听。")

    add_script(doc, "\"停。A：这次有什么不同？\"")
    add_instructor_note(doc, "收集4-5个回应，通常是\"感觉被真正听到了\"、\"我说的比我以为的多了很多\"、\"有些话说出来之后我自己才发现原来我这么想\"。")

    add_script(doc, "\"注意A说的最后一句话——'说出来之后我自己才发现原来我这么想'。\"")
    add_script(doc, "\"这就是深度聆听的神奇之处：当一个人感受到被真正聆听，他会说出连他自己都没有意识到的东西。而那些东西，往往是最有价值的教练材料。\"")
    add_script(doc, "\"你没问任何问题。B没做任何事，只是聆听。但这次对话，已经帮助A开始了一个洞察的过程。\"")

    add_heading(doc, "8.2 有意识聆听的四个层次", level=3)
    add_time_tag(doc, "【13:30-13:45 | 15分钟】")

    add_key_point(doc, "展示四层次图（PPT）：")

    listen_table = doc.add_table(rows=5, cols=2)
    listen_table.style = 'Table Grid'
    add_table_header_row(listen_table, ["层次", "描述"])

    listen_levels = [
        ("第一层：聆听行为（What）", "\"发生了什么？员工说的事实和事件是什么？\"\n这是最浅层的聆听——大多数管理者停在这里，因为他们在听的同时，已经在想解决方案了。"),
        ("第二层：聆听信念（Why/How）", "\"他对这件事持有什么样的信念？他的假设是什么？他认为这件事'应该是什么样的'？\"\n这一层需要你不只听词汇，而是听词汇背后的逻辑和假设。当你听到\"总是\"、\"从来\"、\"不可能\"这类绝对化词汇，那是信念层次的信号。"),
        ("第三层：聆听内在资源（What else）", "\"他在说话的过程里，提到了哪些他自己可能没有意识到的优势、资源或洞察？\"\n这是教练聆听里最有价值的一层——被教练者往往在说话的过程里，无意中说出了他自己最好的答案，只是他自己没有注意到。你注意到了，才能在后面的提问里帮他抓住它。"),
        ("第四层（AI时代新增）：聆听AI时代身份信号（Identity）", "\"他在谈到AI和自己的工作时，流露出什么样的自我认知？他认为自己在AI时代的价值是什么？\"\n三类信号：\n• 价值焦虑信号：语言里隐含\"AI会取代我\"的担忧\n• 过度依赖信号：谈到工作方式时，AI工具成为了核心主体\n• 无效抵触信号：对AI工具明显的防御性贬低"),
    ]
    for i, (level, desc) in enumerate(listen_levels):
        listen_table.rows[i+1].cells[0].text = level
        listen_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    add_script(doc, "\"第四层是这门课里专门新增的。它需要你在聆听的同时，有一部分意识是在观察：'他对自己在AI时代的价值，持有什么样的信念？'这个信念，可能是他最深的心智阻碍，也是教练最有价值的突破点。\"")

    add_heading(doc, "8.3 聆听层次演练", level=3)
    add_time_tag(doc, "【13:45-13:55 | 10分钟】")

    add_script(doc, "\"三人一组：A说，B聆听，C观察。\"")
    add_script(doc, "\"A分享一个他在领导高潜员工时遇到的真实挑战，3分钟。B深度聆听，脑子里记录：你在每个层次上听到了什么？C观察B的聆听质量——他的身体、眼神、有没有分心。\"")
    add_script(doc, "\"3分钟结束，B说：在四个层次上，你各听到了什么？（不需要全部，说你注意到的）\"")

    add_instructor_note(doc, "8分钟，三人组演练，讲师走动观察。")

    add_key_point(doc, "全班快速汇集（2分钟）：")
    add_script(doc, "\"在第四层——AI时代身份信号层次——有没有人在演练里听到了一个信号？是哪一类？\"")

    add_heading(doc, "8.4 有意识提问——AI时代突破性问题", level=3)
    add_time_tag(doc, "【13:55-14:15 | 20分钟】")

    add_script(doc, "\"聆听，帮助你收集最真实的信息，找到被教练者的内在资源。提问，帮助你把这些资源激活，引发顿悟。\"")
    add_script(doc, "\"有意识提问，和普通的工作问话，有三个核心区别——\"")
    add_key_point(doc, "展示三类问题对比（PPT）：")

    question_table = doc.add_table(rows=4, cols=2)
    question_table.style = 'Table Grid'
    add_table_header_row(question_table, ["类型", "详细内容"])

    question_types = [
        ("类型一：封闭式问题", "\"你最终打算用X方法还是Y方法？\"\n特征：答案只有几个选项，限制了被教练者的思考空间。\n效果：封锁对话，不适合教练。"),
        ("类型二：开放式问题", "\"你觉得接下来有哪些可能的方向？\"\n特征：没有预设答案，给被教练者自由的思考空间。\n效果：打开对话，基础的教练问题。"),
        ("类型三（AI时代新增）：突破性问题", "特征：在开放式的基础上，挑战被教练者当前的思维框架，帮助他看到他原有视角看不到的角度。\n效果：引发顿悟，这是GUIDE里I步骤的核心工具。"),
    ]
    for i, (qtype, detail) in enumerate(question_types):
        question_table.rows[i+1].cells[0].text = qtype
        question_table.rows[i+1].cells[1].text = detail

    doc.add_paragraph()

    add_key_point(doc, "展示AI时代突破性问题的三种设计（PPT）：")

    design_table = doc.add_table(rows=4, cols=2)
    design_table.style = 'Table Grid'
    add_table_header_row(design_table, ["设计", "示例问题"])

    designs = [
        ("设计一：反向视角问题", "\"你最担心AI'替代'你的那部分，反过来想——那里是不是也是你最有深度的地方？\"\n\"如果你最强的那个能力，恰好是AI最难做到的，那是什么？\""),
        ("设计二：资源发现问题", "\"在你做得最好的时刻，你的思考里有什么是AI的分析找不到的？\"\n\"如果AI帮你做掉了所有的'表面工作'，你的时间和精力会自然去往哪里？\""),
        ("设计三：身份定位问题", "\"五年后，什么样的你，是在AI时代里真正不可替代的？\"\n\"你希望别人说'这件事，必须是X（被教练者的名字）来做，因为……'，那个'因为'是什么？\""),
    ]
    for i, (d, ex) in enumerate(designs):
        design_table.rows[i+1].cells[0].text = d
        design_table.rows[i+1].cells[1].text = ex

    doc.add_paragraph()

    add_script(doc, "\"这三类突破性问题，有一个共同的设计逻辑：它们都在帮助被教练者超越'AI vs 我'的对立框架，进入'我的深度价值在哪里，以及AI如何放大它'的新框架。\"")
    add_script(doc, "\"这类问题，不只是AI时代的教练问题——在任何关于职业发展的教练对话里，帮助被教练者找到'只有我才能做的那部分'，都是发展型教练的核心任务。AI只是让这个任务更加紧迫和清晰。\"")

    add_key_point(doc, "提问技术演练（10分钟）")
    add_script(doc, "\"配对练习。B说一个他在工作里的真实困惑（和AI相关的最好，没有也可以）。A的任务：从四个层次聆听，然后设计一个突破性问题。不是建议，是一个问题。\"")
    add_script(doc, "\"5分钟对话，然后B告诉A：那个问题，有没有让你想到一个你以前没想过的角度？\"")

    add_instructor_note(doc, "5分钟+2分钟反馈，讲师巡场，特别观察A能否真的只问一个高质量的突破性问题，不给建议。")

    add_key_point(doc, "全班复盘（3分钟）")
    add_script(doc, "\"有没有一个问题，让你的伙伴说了一句他之前没说过的话？分享那个问题。\"")
    add_instructor_note(doc, "收集2-3个，写在白板上。")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 9: Part 2 后段 ---
    add_heading(doc, "9. Part 2 后段：反馈与认同", level=2)
    add_time_tag(doc, "【15:00-17:00 | 120分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "区分评价性反馈和发展性反馈；掌握教练式反馈的结构（含AI时代版本）；掌握教练语言的四个特征；理解并演练认同技巧；通过完整的GUIDE演练整合Day 1的全部技巧。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "幻灯片（发展性反馈框架+教练语言四特征+认同技巧）、学员手册P23-28（Part 2后段工作页）")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "这一段里，\"发展性反馈\"是最容易被学员误解的工具——他们可能以为它就是\"正面的反馈\"，或者\"委婉说问题\"。你需要建立的认知是：发展性反馈的核心，不是反馈的内容，而是反馈的目的——不是为了评价这次的表现，而是为了触动被教练者对自己成长方向的洞察。这个目的，在AI时代有了新的内涵：帮助被教练者看清楚他的人类贡献在哪里，以及如何在AI工具的参与下让这个贡献变得更清晰可见。")
    add_script(doc, "\"认同技巧\"的部分，很多讲师把它讲成\"多夸员工\"——这是误读。认同的本质是：真正看见对方的内在资源和价值，而不是表面的行为或结果。高潜员工对肤浅的夸奖很敏感，他们会识别出\"这只是鼓励，不是真的看见\"。")

    add_heading(doc, "9.1 教练式反馈vs评价性反馈", level=3)
    add_time_tag(doc, "【15:00-15:10 | 10分钟】")

    add_script(doc, "\"反馈，是管理者和员工交流里频率最高的内容之一。但大多数反馈，不是教练式的。\"")
    add_key_point(doc, "展示两类反馈对比（PPT）：")

    feedback_table = doc.add_table(rows=3, cols=2)
    feedback_table.style = 'Table Grid'
    add_table_header_row(feedback_table, ["类型", "详细内容"])

    feedback_types = [
        ("评价性反馈", "\"你在这次提案里，客户洞察部分做得很好，但策略建议部分还需要加强。\"\n特征：评价这次行为的好坏，针对具体事件。\n作用：让被教练者知道你的评价，但不一定触动他的成长思考。"),
        ("发展性反馈", "\"我在你这次提案里，看到了一个东西——在你描述客户使用场景的那一段，你说了一句话：'他们真正焦虑的不是功能，而是……'。这句话，是你对客户的真实理解，不是数据分析出来的，是你和他们打了两年交道之后的直觉。这是你在这类工作里最有价值的部分，我认为这也是AI最难做到的部分。你自己感觉到了吗？\"\n特征：聚焦被教练者的内在资源或洞察，触动他的自我认知。\n作用：帮助被教练者看见自己的价值和成长方向。"),
    ]
    for i, (ftype, detail) in enumerate(feedback_types):
        feedback_table.rows[i+1].cells[0].text = ftype
        feedback_table.rows[i+1].cells[1].text = detail

    doc.add_paragraph()

    add_script(doc, "\"注意区别：第一个反馈在评价行为；第二个反馈在帮助被教练者看见他自己没有充分意识到的内在资源。\"")
    add_script(doc, "\"发展性反馈，不是委婉的批评，也不是多说好话。它是一种帮助被教练者增强自我认识的工具——让他看见什么，对他的成长最有价值。\"")

    add_heading(doc, "9.2 发展性反馈的结构（AI时代版）", level=3)
    add_time_tag(doc, "【15:10-15:25 | 15分钟】")

    add_key_point(doc, "展示发展性反馈的四个要素（PPT）：")

    dev_feedback_table = doc.add_table(rows=5, cols=2)
    dev_feedback_table.style = 'Table Grid'
    add_table_header_row(dev_feedback_table, ["要素", "说明"])

    dev_feedback = [
        ("要素一：具体的观察", "说出你看到的，具体到情境和行为，不是泛泛的印象。\n示例：\"在上周的团队复盘里，你在分析这个问题时停顿了一下，然后说了一个完全不同的角度……\""),
        ("要素二：内在资源的命名", "说出这个观察背后体现的，是被教练者的什么内在资源或能力。\n示例：\"……那个停顿，和那个不同的角度，体现的是你在这个领域的专业直觉——这是经验积累才会有的。\""),
        ("要素三：AI时代的价值定位（AI时代新增）", "明确说出这个内在资源在AI时代的独特价值——这是AI工具做不到或不能替代的地方。\n示例：\"……AI可以给你十个分析框架，但它没有你在这个行业里打了五年交道积累的直觉。那个直觉，是你最有价值的东西。\""),
        ("要素四：邀请自我探索", "不是你做结论，而是把一个问题还给被教练者——帮助他自己深化这个洞察。\n示例：\"……你感受到这部分吗？你觉得怎么样把它发挥得更充分？\""),
    ]
    for i, (elem, desc) in enumerate(dev_feedback):
        dev_feedback_table.rows[i+1].cells[0].text = elem
        dev_feedback_table.rows[i+1].cells[1].text = desc

    doc.add_paragraph()

    add_script(doc, "\"四个要素里，第三个是AI时代新增的，但它对高潜员工来说往往是最有冲击力的部分——因为他们需要听到有人能明确说出'这是AI替代不了你的地方'，而不只是泛泛的'你很厉害'。\"")
    add_script(doc, "\"而且这个说法，必须是真实的——基于你真正看到的具体观察，而不是安慰。高潜员工识别得出来。\"")

    add_key_point(doc, "发展性反馈演练（15分钟）")
    add_script(doc, "\"三人一组：A和B在之前的GUIDE演练里，你观察到了他的一个内在资源，现在给他一个发展性反馈；C观察并记录A的反馈里，四个要素哪些做到了，哪些没有。\"")
    add_script(doc, "\"每人给反馈3分钟，然后接受方说：这个反馈，让我想到了什么？\"")

    add_instructor_note(doc, "15分钟，三人组，讲师巡场，特别关注第三要素是否被使用。")

    add_key_point(doc, "全班复盘（5分钟）")
    add_script(doc, "\"在你收到的发展性反馈里，有没有一个让你想到一件新事情的时刻？那个时刻，对方说了什么？\"")
    add_instructor_note(doc, "收集2-3个，写在白板上。")

    add_heading(doc, "9.3 教练语言的四个特征", level=3)
    add_time_tag(doc, "【15:25-15:35 | 10分钟】")

    add_script(doc, "\"除了发展性反馈，教练对话里所有的语言，都有四个共同的特征。原版课程叫做：有要求、有压力、向前、正面。\"")
    add_script(doc, "\"我们来看每一个在AI时代的应用。\"")
    add_key_point(doc, "展示四特征（PPT）：")

    lang_table = doc.add_table(rows=5, cols=2)
    lang_table.style = 'Table Grid'
    add_table_header_row(lang_table, ["特征", "AI时代应用"])

    lang_chars = [
        ("有要求（Challenging）", "不是容易的问题，而是真正要求被教练者思考的问题。\"你觉得还行吗？\"——不够有要求。\"在你做得最好的时候，和现在的你，差在哪里？那个差距，你打算怎么对待它？\"——有要求。\nAI时代：AI时代的\"有要求\"还包括：帮助被教练者正视\"AI在我的工作里意味着什么\"这个问题，而不是回避它。"),
        ("有压力（Pressured）", "推动被教练者走出舒适区，而不是停留在他已经知道的地方。教练对话不是友好的闲聊，是在被教练者的成长边界处工作的。\nAI时代：当被教练者的答案总是\"AI帮我做\"时，有一定的压力性问题是必要的——\"AI做了这部分，那什么是只有你才能做的？\""),
        ("向前（Forward-focused）", "关注的是\"接下来\"，而不是\"为什么过去是这样\"。不是审判，而是规划。\nAI时代：当被教练者陷入\"AI会不会替代我\"的焦虑时，向前意味着把对话从\"会不会\"转向\"接下来我怎么做\"。"),
        ("正面（Positive）", "不是只说好话，而是从一个相信被教练者有资源解决问题的立场来问问题——不是\"你哪里出了问题\"，而是\"你最好的状态下，是什么样的\"。\nAI时代：相信高潜员工有能力在AI时代找到自己的不可替代性，并把这个信念通过问题传递出去。"),
    ]
    for i, (char, ai_app) in enumerate(lang_chars):
        lang_table.rows[i+1].cells[0].text = char
        lang_table.rows[i+1].cells[1].text = ai_app

    doc.add_paragraph()

    add_heading(doc, "9.4 认同技巧", level=3)
    add_time_tag(doc, "【15:35-15:50 | 15分钟】")

    add_script(doc, "\"认同，不是夸奖。\"")
    add_script(doc, "\"夸奖：'你这次做得很好。'——这是评价，评价者是你，参照物是你的期待或者标准。\"")
    add_script(doc, "\"认同：'在你刚才说的话里，我听到了一种对这个问题的深度关心，这是做好这件事真正需要的东西。你注意到了吗？'——这是看见，看见的是被教练者身上的内在资源，帮助他自己认识到它。\"")
    add_script(doc, "\"高潜员工，可以感受出夸奖和认同的区别。夸奖他们听了很多，认同，让他们感到真正被理解了。\"")

    add_ai_content(doc, "AI时代版本：")
    add_script(doc, "AI时代，认同有了一个新的重要应用场景——当高潜员工处于价值焦虑时。")
    add_script(doc, "这时候，夸奖（'你很厉害，AI替代不了你'）是空洞的安慰。真正的认同，是说出你在这个人身上看见的具体的、不可被AI替代的东西：")

    add_key_point(doc, "示范话术：")
    add_script(doc, "\"我在你身上看到一件事——当我们讨论这个问题的时候，你说了一句话，让整个对话改变了方向。那句话来自你三年里和客户之间建立的真实的关系理解。这不是AI能做到的，是你做到的。你自己感觉到了吗？\"")

    add_key_point(doc, "认同练习（10分钟）")
    add_script(doc, "\"配对。A说一件他最近在工作里感到有些迷茫或者不确定的事。B聆听，然后从A说的话里，找到一个你在A身上看见的内在资源——用认同技巧说出来。\"")
    add_script(doc, "\"不是安慰，是真的看见。如果你在他说话里找不到真实的内在资源，就继续问，直到你真的找到了再说。\"")

    add_instructor_note(doc, "8分钟，配对，讲师巡场。")

    add_key_point(doc, "全班复盘（2分钟快速）")
    add_script(doc, "\"被认同的那一刻，你感受到什么不同？\"")
    add_instructor_note(doc, "1-2个回应，不展开。")

    add_heading(doc, "9.5 GUIDE完整综合演练", level=3)
    add_time_tag(doc, "【15:50-16:20 | 30分钟】")

    add_script(doc, "\"今天下午的最后，我们做一次完整的GUIDE演练——把上午学的模型，和下午学的聆听、提问、发展性反馈和认同技巧整合在一起。\"")
    add_script(doc, "\"三人一组：管理者（教练者）、高潜员工（被教练者）、观察者。\"")
    add_script(doc, "\"被教练者：用你带来的高潜下属的真实困惑，进入那个人的状态——他的主要工作内容，他可能有的AI时代的挑战（可以是价值焦虑型、过度依赖型或抵触型，任选一种）。\"")
    add_script(doc, "\"管理者：走完GUIDE五步，在I步骤里至少用一个突破性问题；在D步骤里，尝试引出双轨行动计划（人类深度+AI杠杆）；在对话的任何一个时刻，如果你看到了一个内在资源，给一个认同。\"")
    add_script(doc, "\"观察者：使用演练观察清单（手册P27），记录你在每步看到的关键瞬间——哪里打开了，哪里关闭了，哪里出现了突破性问题，哪里管理者忍不住给了建议。\"")
    add_script(doc, "\"15分钟教练对话，3分钟被教练者反馈，2分钟观察者反馈。\"")

    add_instructor_note(doc, "20分钟，三人组，讲师巡场，记录1-2个典型瞬间用于全班复盘。")

    add_key_point(doc, "全班复盘（10分钟）")
    add_key_point(doc, "三层推进：")

    add_key_point(doc, "第一层（What）")
    add_script(doc, "\"被教练者们——今天的演练里，有没有一个时刻，你感到一种真正的洞察出现了？那是什么时刻？\"")
    add_instructor_note(doc, "收集2-3个，让被教练者说出那个瞬间发生了什么。")

    add_key_point(doc, "第二层（So What）")
    add_script(doc, "\"管理者们——今天演练里，最难的部分是什么？\"")
    add_instructor_note(doc, "通常是：忍住不给建议、I步骤的突破性问题、聆听第三层和第四层。")
    add_script(doc, "\"这些难，是好消息——说明这些能力是可以练习的，今天你已经练了一次。明天，还有更多演练机会。\"")

    add_key_point(doc, "第三层（Now What）")
    add_script(doc, "\"今晚，如果你花5分钟想一件事——你的分析对象，他在AI时代的人类深度峰在哪里？你明天来的时候，手里有一个具体的答案。这是明天教练行动计划里最重要的一块。\"")

    add_key_point(doc, "Day 1 收尾话术")
    add_script(doc, "\"今天，我们建立了教练型领导力的基础——思维框架、发展型教练模型和核心技巧。明天，我们进入业绩表现型教练（DIRECT模型），以及综合的演练和带走的行动计划。\"")
    add_script(doc, "\"今晚一件事：想清楚你分析对象的人类深度峰，明天会用到。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # ========== PART 2 - DAY 2 ==========
    add_heading(doc, "第二部分 Day 2", level=1, color=DARK_BLUE_KEY)

    add_heading(doc, "10. Day 2 开场与回温", level=2)
    add_time_tag(doc, "【09:00-09:30 | 30分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "快速激活Day 1的学习积累；过渡到业绩表现型教练的心智状态；为Day 2的综合演练建立期待。")

    add_heading(doc, "10.1 人类深度峰分享", level=3)
    add_time_tag(doc, "【09:00-09:10 | 10分钟】")

    add_script(doc, "\"早。昨晚让你们想一件事——你分析对象的人类深度峰。找旁边的人，90秒，说：你觉得他的人类深度峰在哪里，你今天打算怎么帮助他看见它？\"")
    add_instructor_note(doc, "90秒配对，不全班汇集，快速进入Day 2内容。")

    add_heading(doc, "10.2 GUIDE vs DIRECT的场景区分", level=3)
    add_time_tag(doc, "【09:10-09:25 | 15分钟】")

    add_script(doc, "\"今天上午，我们进入DIRECT模型——业绩表现型教练。在开始之前，有一个最重要的区分要建立清楚：GUIDE和DIRECT，什么时候用哪个？\"")
    add_key_point(doc, "展示对比表（PPT）：")

    guide_direct_table = doc.add_table(rows=5, cols=3)
    guide_direct_table.style = 'Table Grid'
    add_table_header_row(guide_direct_table, ["维度", "GUIDE（发展型）", "DIRECT（业绩表现型）"])

    gd_data = [
        ("触发条件", "员工有未发挥的潜力，或有发展上的瓶颈", "员工有具体的、可观察的行为或结果问题"),
        ("核心问题", "\"你最想去到哪里？什么阻碍你？\"", "\"这里发生了什么？影响是什么？\""),
        ("对话方向", "向前看，规划成长", "向前看，但起点是现存问题的理解"),
        ("AI时代挑战", "帮助员工找到人类深度峰", "区分人类贡献和AI贡献，设定有意义的绩效期望"),
    ]
    for i, (dim, guide, direct) in enumerate(gd_data):
        guide_direct_table.rows[i+1].cells[0].text = dim
        guide_direct_table.rows[i+1].cells[1].text = guide
        guide_direct_table.rows[i+1].cells[2].text = direct

    doc.add_paragraph()

    add_script(doc, "\"两个模型都是向前的，都不是审判。区别在于：GUIDE的起点是潜力和可能，DIRECT的起点是一个需要被正视的行为或绩效问题。\"")
    add_script(doc, "\"搞混的最常见情况：遇到绩效问题，用GUIDE——跳过了对问题的直接处理，进入了发展型讨论，结果是员工感觉问题被回避了，管理者感觉员工没有认识到问题。\"")
    add_script(doc, "\"另一个搞混：遇到发展型对话，用DIRECT——用处理问题的方式来讨论潜力，结果被教练者感到被评估和被管控，而不是被支持发展。\"")

    add_key_point(doc, "情境判断练习（5分钟）")
    add_script(doc, "\"我来说六个情境，你们判断：GUIDE还是DIRECT？\"")
    add_instructor_note(doc, "快速六题，举手判断，讲师快速说是非不展开，目的是练习区分而不是深度讨论。")

    scenarios = [
        ("情境1", "一名员工在上次客户提案会上没有充分准备，客户反应负面。", "DIRECT"),
        ("情境2", "一名高潜员工开始问\"我在AI时代有什么价值\"，但工作表现没有问题。", "GUIDE"),
        ("情境3", "一名员工连续两个季度未达到业绩目标，自己解释说主要是外部原因。", "DIRECT"),
        ("情境4", "一名资深员工感觉自己已经到了职业天花板，开始考虑离职。", "GUIDE"),
        ("情境5", "一名员工的项目报告质量明显下滑，细看发现大量内容未经判断直接引用AI输出。", "DIRECT"),
        ("情境6", "一名员工一直在做好本职工作，但你感到他有更大的潜力还没有被激活。", "GUIDE"),
    ]
    for num, scenario, answer in scenarios:
        add_script(doc, f"{num}：{scenario} → {answer}")

    add_heading(doc, "10.3 Day 2结构预告", level=3)
    add_time_tag(doc, "【09:25-09:30 | 5分钟】")

    add_script(doc, "\"今天的结构：上午学DIRECT模型和演练；下午整合所有工具做综合演练，然后制作你的教练行动计划——这是你今天最重要的带走物。\"")
    add_script(doc, "\"DIRECT里，AI时代最有价值的升级集中在三步：D（Data收集）、R（绩效期望设定）、E（探索解决方案）。这三步，我会专门花时间讲AI时代的版本。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 11: Part 3 前段 ---
    add_heading(doc, "11. Part 3 前段：DIRECT模型", level=2)
    add_time_tag(doc, "【09:30-11:00 | 90分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "掌握DIRECT六步及每步的核心技巧；深入理解AI时代版本的DATA收集（人类贡献归因）和REQUIREMENT设定（双维度绩效期望）；通过讲师示范感受完整的DIRECT对话。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "DIRECT教练模型应用卡（发放）、幻灯片（DIRECT六步+AI时代升级+参考话术）")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "DIRECT里，最容易出现的问题是：DATA步骤收集的是印象而不是行为；REQUIREMENT步骤变成了单方面发布命令；EXPLORE步骤被跳过（管理者直接从R跳到C）。")
    add_ai_content(doc, "AI时代让DIRECT有了新的复杂性：DATA的收集必须包含人类贡献归因的分析——不是问\"你做了什么\"，而是\"在你完成的这件事里，你的判断体现在哪里\"。这个升级非常重要，因为如果DATA只收集了产出（\"你的报告质量下降了\"），没有收集人类贡献（\"你报告里的判断质量下降了，而产出数量没有变\"），后续的改进讨论会失去焦点。")

    add_heading(doc, "11.1 DIRECT六步概览", level=3)
    add_time_tag(doc, "【09:30-10:00 | 30分钟】")

    add_script(doc, "\"DIRECT，是处理业绩表现问题的教练框架。六步，每步有清晰的目的。\"")
    add_script(doc, "\"和GUIDE最大的不同：DIRECT是从一个已经发生的、可观察的行为或结果问题出发的。被教练者知道这次对话是因为有一个问题需要被处理。\"")
    add_script(doc, "\"但即便如此，DIRECT的方式依然是教练式的——不是宣判，不是命令，是帮助被教练者自己找到解决方案并做出承诺。\"")

    add_key_point(doc, "逐步讲授（每步约4分钟）：")

    # D Step
    add_key_point(doc, "D——Data（收集具体的、可观察的行为）")
    add_script(doc, "原版核心：不用印象，不用\"总是\"或\"从来\"，用具体的、有情境的行为描述。\"你总是这样\"——不行。\"在X情境里，你做了Y\"——这才是Data。")
    add_ai_content(doc, "AI时代关键升级（重点）：")
    add_script(doc, "\"在AI时代，Data的收集需要增加一个维度——**人类贡献归因分析**。\"")
    add_script(doc, "\"原有的Data收集：'这份报告里，有几个结论没有数据支撑，客户提出了质疑。'\"")
    add_script(doc, "\"AI时代的Data收集：'这份报告里，有几个结论没有数据支撑。我在看报告的时候，注意到这几个结论的措辞和AI工具的典型输出非常相似。我想了解一下，这几个结论，你的判断是什么——你是怎么得出的？'（先收集信息，不先定性）\"")
    add_script(doc, "\"关键点：在Data步骤里，你不是指控员工用了AI，而是在收集人类贡献的证据——他能不能说出自己的判断逻辑。这是这份工作里你真正要评估的东西，不是产出的格式是否好看。\"")
    add_key_point(doc, "参考话术：")
    add_script(doc, "\"'在X这件事里，你能描述一下当时发生了什么吗？'\"")
    add_script(doc, "\"'在这个过程里，你的判断体现在哪里？'\"")
    add_script(doc, "\"'当你做Y这个决定的时候，你是怎么想的？'\"")

    # I Step
    add_key_point(doc, "I——Impact（沟通行为对组织的影响）")
    add_script(doc, "原版核心：说出这个行为对团队、客户、业务或组织产生的具体影响——不是泛泛的\"这样不好\"，而是\"这导致了X结果\"。")
    add_ai_content(doc, "AI时代升级：影响分析包含一个新维度——对被教练者长期能力成长的影响。")
    add_script(doc, "'你报告里判断部分不够清晰'的影响不只是'客户对这次提案不满意'，还有：'你在这类工作里积累深度判断经验的机会，正在被工具的输出替代。长期下去，你的核心优势会受到影响。'\"")
    add_script(doc, "\"这个影响说清楚，是帮助被教练者看到：这不只是一次绩效问题，而是一个关于他自己长期价值的问题。这往往比'客户不满意'更能触动高潜员工。\"")

    # R Step
    add_key_point(doc, "R——Requirement（回顾或设定绩效期望）")
    add_script(doc, "原版核心：清楚地说出你对绩效的期望，确认双方在标准上有共识。")
    add_ai_content(doc, "AI时代关键升级（重点）：绩效期望的设定包含两个维度：")
    add_script(doc, "\"维度一，产出标准：结果或成果应达到什么水平。（这和原来一样）\"")
    add_script(doc, "\"维度二（AI时代新增），人类贡献标准：在这类工作里，你的判断、洞察或决策应达到什么水平，以及如何被看见。\"")
    add_script(doc, "\"为什么需要两个维度？因为如果只有产出标准，AI工具可以帮助员工达到标准，但员工的能力可能没有任何成长。两个维度同时设定，员工才知道：达标不只是交出来一个结果，还要让你的判断可见。\"")
    add_key_point(doc, "示范绩效期望的双维度表述：")
    add_script(doc, "\"'在下季度，我对你这类分析报告的期望是：产出层面——客户满意度回到X以上，结论有明确数据支撑；人类贡献层面——在报告提交的时候，你能够向我说明报告里每个核心结论，你是如何得出的判断，以及你认为这个判断里最重要的是什么（而不是AI工具的输出）。'\"")

    # E Step
    add_key_point(doc, "E——Explore（帮助被教练者找到解决方案）")
    add_script(doc, "原版核心：不要替被教练者想解决方案，而是通过提问，帮助他自己找到解决方案——这是DIRECT里最接近GUIDE精神的一步。")
    add_ai_content(doc, "AI时代升级：探索解决方案时，特别帮助被教练者思考：如何重新设计人机协作方式，让人类判断贡献更可见。")
    add_key_point(doc, "参考问题：")
    add_script(doc, "\"'你觉得在这类工作里，什么地方最应该是你自己来判断的？'\"")
    add_script(doc, "\"'如果你重新做这类工作，你会在哪里多花一点时间，确保这是你真实的判断？'\"")
    add_script(doc, "\"'你觉得在使用工具和保留自己的判断之间，怎么找到一个对你来说有效的方式？'\"")

    # C Step
    add_key_point(doc, "C——Commitment（取得行动承诺）")
    add_script(doc, "原版核心：被教练者对具体的行动计划做出明确的承诺——不是\"我会注意\"，而是\"我下周X之前会做Y\"。")
    add_ai_content(doc, "AI时代升级：承诺内容包含人类贡献可见化的具体行动。")
    add_script(doc, "\"不只是'下次我会更认真地分析数据'，而是：'下次报告提交前，我会先写一个没有AI工具的草稿版核心结论，然后再用工具扩展——这样我的判断是在工具之前的，不是在工具之后的。'\"")

    # T Step
    add_key_point(doc, "T——Track（确定追踪要点）")
    add_script(doc, "原版核心：明确双方如何追踪行动计划的执行——什么时候，用什么方式，确认什么。")
    add_ai_content(doc, "AI时代升级：追踪不只看产出指标，还追踪人类贡献可见性的成长。")
    add_script(doc, "\"追踪要点可以包括：在下次一对一里，你能不能清楚地说出上个月在X类工作里的三次关键判断？（这是人类贡献可见性的追踪，不是产出数字的追踪）\"")

    add_heading(doc, "11.2 讲师示范DIRECT对话", level=3)
    add_time_tag(doc, "【10:00-10:20 | 20分钟】")

    add_script(doc, "\"现在我来示范一次完整的DIRECT对话。场景是：场景F——员工的季度绩效数字达标，但你发现他的核心产出里人类判断贡献极少，近期在客户会上无法说清楚方案背后的逻辑。\"")
    add_script(doc, "\"谁愿意扮演这名员工？\"")
    add_instructor_note(doc, "选一位学员，提前10秒提示：自然地扮演一个有自知之明但有些防御的员工。")

    add_key_point(doc, "示范（15分钟）：")
    add_script(doc, "重点示范：")
    add_script(doc, "D步骤：用具体问题收集人类贡献归因，而不是直接定性")
    add_script(doc, "I步骤：说出对长期能力的影响，不只本次绩效的影响")
    add_script(doc, "R步骤：明确双维度绩效期望")
    add_script(doc, "E步骤：帮助员工自己想到解决方案，不抢答")

    add_key_point(doc, "示范结束（5分钟）：")
    add_script(doc, "\"三件事我想说——\"")
    add_script(doc, "\"第一，D步骤里，我没有直接说'你用AI代劳了'，我问的是'在这个结论里，你的判断是什么'——这个问题，让员工自己说清楚（或者说不清楚），而不是我来指控。\"")
    add_script(doc, "\"第二，R步骤里，我说了两个维度的期望——产出层面和人类贡献层面。员工听完之后，他知道'交出一个漂亮的报告'不够了——他需要让他的判断可见。\"")
    add_script(doc, "\"第三，E步骤里，我在员工自己说出解决方案之后，才确认和补充。这和GUIDE里的精神是一样的——被教练者自己想到的解决方案，执行的承诺度比你给他的高得多。\"")

    add_heading(doc, "11.3 全班讨论与复盘", level=3)
    add_time_tag(doc, "【10:20-10:30 | 10分钟】")

    add_script(doc, "\"从示范里，你注意到什么？有没有一个时刻，你感到'这是关键的一步'？\"")
    add_instructor_note(doc, "收集3-4个观察，不评判，只接收。")

    add_script(doc, "\"D步骤和R步骤的AI时代升级，你在自己的管理场景里，有没有一个你觉得特别适用的情境？\"")
    add_instructor_note(doc, "收集1-2个，建立与自己工作情境的连接。")

    add_key_point(doc, "茶歇过渡话术")
    add_script(doc, "\"DIRECT的六步你们已经学完了。茶歇后，我们来演练——三人组，用AI时代的真实场景，走完一次完整的DIRECT对话。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 12: Part 3 后段 ---
    add_heading(doc, "12. Part 3 后段：DIRECT演练", level=2)
    add_time_tag(doc, "【11:15-12:15 | 60分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "通过三人组演练建立DIRECT对话能力；特别练习AI时代版本的D步骤（人类贡献归因）和R步骤（双维度绩效期望）；获得结构化观察反馈。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "三人组角色卡（从AI时代六类场景卡里，选择适合DIRECT的场景B/E/F）、演练观察清单（手册P31，DIRECT版）")

    add_heading(doc, "12.1 三人组演练", level=3)
    add_time_tag(doc, "【11:15-12:00 | 45分钟】")

    add_script(doc, "\"三人一组，角色：管理者（教练者）、被教练者、观察者。\"")
    add_script(doc, "\"被教练者：从B、E、F三张情境卡里选一张（这三张适合用DIRECT），进入那个员工的状态。\"")
    add_script(doc, "\"管理者：走完DIRECT六步。特别注意D步骤——不要直接定性，用问题收集人类贡献归因；特别注意R步骤——说出双维度绩效期望。\"")
    add_script(doc, "\"观察者：用DIRECT演练观察清单，记录每步的关键瞬间。特别标记：D步骤里有没有归因探索的问题；R步骤里有没有双维度期望；E步骤里有没有让员工自己想出解决方案。\"")
    add_script(doc, "\"第一轮：15分钟对话，3分钟被教练者感受，2分钟观察者反馈。\"")
    add_instructor_note(doc, "20分钟，第一轮。")

    add_script(doc, "\"角色轮换，换情境卡，第二轮。\"")
    add_instructor_note(doc, "20分钟，第二轮。")

    add_script(doc, "\"如果时间允许，第三人做管理者，第三轮选一个你们桌上讨论最热烈的情境。\"")
    add_instructor_note(doc, "5分钟，第三轮，可选。")

    add_key_point(doc, "讲师巡场要点")
    add_instructor_note(doc, "重点观察两件事：D步骤里，管理者是否用了\"你的判断体现在哪里\"类型的问题，还是直接定性\"你用AI代劳了\"；E步骤里，管理者是否跳过了，直接从R跳到了C。这两个是最常见的演练问题，在全班复盘里作为焦点。")

    add_heading(doc, "12.2 全班复盘", level=3)
    add_time_tag(doc, "【12:00-12:15 | 15分钟】")

    add_key_point(doc, "三层推进：")
    add_key_point(doc, "第一层（What）")
    add_script(doc, "\"被教练者们——在DIRECT对话里，有没有一个时刻，你感到管理者真的在帮你找解决方案，而不是在惩罚你？那是什么时刻？\"")
    add_instructor_note(doc, "收集2-3个，聚焦E步骤的感受。")

    add_key_point(doc, "第二层（So What）")
    add_script(doc, "\"DIRECT和GUIDE有一个共同点——都有一个步骤是帮助被教练者自己找解决方案，而不是管理者给解决方案。这一步，在演练里有没有真的发生？什么阻碍了它发生？\"")
    add_instructor_note(doc, "通常的答案：时间压力、管理者已经知道答案了忍不住、被教练者等着被告知。")

    add_key_point(doc, "第三层（Now What）")
    add_script(doc, "\"D步骤里，你用了人类贡献归因的问题吗？如果没有，你觉得在你真实的管理场景里，什么时候这个问题会最有价值？\"")

    add_key_point(doc, "午餐过渡话术")
    add_script(doc, "\"DIRECT演练完了。午饭后，我们进入今天最综合、也最接近真实工作场景的部分——AI时代六类场景的综合演练，以及你们的教练行动计划。\"")
    add_script(doc, "\"两件事要在午饭时想好：你的分析对象，这次的综合演练，你打算用GUIDE还是DIRECT？你觉得他是潜力瓶颈（GUIDE），还是有一个具体的行为或绩效问题（DIRECT）？\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 13: Part 4 前段 ---
    add_heading(doc, "13. Part 4 前段：综合演练", level=2)
    add_time_tag(doc, "【13:15-14:30 | 75分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "在更复杂、更接近真实工作的场景里，综合运用GUIDE和DIRECT；学会根据场景选择合适的教练模式；通过全班分享提炼关键学习。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "AI时代六类场景卡（A-F，每桌一套）、演练观察清单（含综合演练版，手册P35）")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "综合演练是两天里最接近真实工作的环节，也最容易出现\"演练变成讲话\"的情况——管理者扮演者在情境里开始发表管理观点，而不是真正在做教练对话。你的工作是在巡场时轻声提醒这一点，不要在演练期间全班打断。")
    add_script(doc, "AI时代的六类场景，有些适合GUIDE（A价值焦虑、D接受更大责任、E克服失望），有些更适合DIRECT（F处理业绩缺陷），有些可以两者都用（B过度依赖、C无效抵触）。学员选择的过程本身就有学习价值，不需要讲师替他们决定。")

    add_heading(doc, "13.1 场景选择与演练准备", level=3)
    add_time_tag(doc, "【13:15-13:25 | 10分钟】")

    add_script(doc, "\"每桌有六张场景卡（A-F），你们小组先讨论5分钟：今天下午，你们想演练哪两个场景，以及每个场景你们认为用GUIDE还是DIRECT更合适，理由是什么？\"")
    add_script(doc, "\"5分钟讨论，每组说出你们选的两个场景和选择理由。\"")
    add_instructor_note(doc, "5分钟讨论，不需要讲师指导。")

    add_key_point(doc, "全班快速确认（5分钟）")
    add_script(doc, "\"每组说一下你们选的场景和判断——GUIDE还是DIRECT？\"")
    add_instructor_note(doc, "快速过一圈，讲师不评判，只说：如果多个组选了同一场景，演练结束后可以比较不同的处理方式。")

    add_script(doc, "如果有组选择了F（业绩缺陷）用GUIDE：可以轻声提示，\"F里有一个具体的行为问题需要被直接处理，想一想DIRECT是否更合适——但你们决定，演练结束再讨论。\"")

    add_heading(doc, "13.2 分组综合演练", level=3)
    add_time_tag(doc, "【13:25-14:15 | 50分钟】")

    add_script(doc, "\"好，开始演练。三人一组，两轮，每个场景一轮。\"")
    add_script(doc, "\"第一轮：20分钟（15分钟对话，5分钟反馈）。\"")
    add_script(doc, "\"第二轮：20分钟（15分钟对话，5分钟反馈）。\"")
    add_script(doc, "\"观察者：用综合演练观察清单（手册P35），记录三件事：\"")
    add_script(doc, "\"第一，管理者选择的是GUIDE还是DIRECT——这个选择在对话里是否合适？\"")
    add_script(doc, "\"第二，你听到的最好的一个教练问题是什么？\"")
    add_script(doc, "\"第三，有没有一个时刻，被教练者明显'打开了'或者'产生了顿悟'？\"")
    add_script(doc, "\"特别提示：AI时代的场景里，有些被教练者可能会引发你的本能反应（'他应该用AI工具，为什么不用'或者'他太依赖AI了'）。把这个反应放一放，先用教练的方式理解他的立场，再走框架。\"")
    add_instructor_note(doc, "第一轮：20分钟；第二轮：20分钟；共40分钟演练时间。")

    add_key_point(doc, "讲师巡场重点")
    add_instructor_note(doc, "重点场景B（过度依赖型）和C（无效抵触型）：这两类场景容易引发管理者的说教模式（\"你应该更独立思考\"或者\"AI工具真的很有用，你应该试试\"）。如果看到这个，停下来轻声说：\"你现在在教他。教练的问题会是什么？\"")
    add_instructor_note(doc, "重点场景A（价值焦虑型）：管理者可能急于安慰，走进\"你很有价值，别担心\"的模式，跳过了真正的探索。轻声提示：\"认同技巧——说出一个你真实看见的具体内在资源，不是泛化的安慰。\"")

    add_heading(doc, "13.3 全班复盘", level=3)
    add_time_tag(doc, "【14:15-14:30 | 15分钟】")

    add_key_point(doc, "三层推进：")
    add_key_point(doc, "第一层（What，5分钟）")
    add_script(doc, "\"哪个场景让你最有挑战？在那个场景里，你遇到的最大的困难是什么？\"")
    add_instructor_note(doc, "收集2-3个，各场景都代表到，不只是最难的一类。")

    add_script(doc, "\"有没有一个演练里出现的教练问题，让你感到'这个问题问得真好'？什么问题？\"")
    add_instructor_note(doc, "收集2-3个，写在白板上。")

    add_key_point(doc, "第二层（So What，5分钟）")
    add_script(doc, "\"在AI时代的这六类场景里，你感觉哪一类在你的真实管理场景里最常见？最让你没有准备的是哪一类？\"")
    add_instructor_note(doc, "让学员说出来，不是讲师分析。")

    add_script(doc, "\"你注意到——同样一个场景，GUIDE和DIRECT处理，员工的体验是完全不同的。A这类价值焦虑，用DIRECT来处理（'这里有一个绩效问题'），会发生什么？\"")
    add_instructor_note(doc, "等学员说：会让员工感觉被攻击，进一步加深焦虑，而不是被支持。")

    add_key_point(doc, "第三层（Now What，5分钟）")
    add_script(doc, "\"在你的分析对象身上，你刚才演练的场景，有没有一个类似的真实情境？从今天的演练里，你带走了什么可以用在下次和他对话里的？\"")
    add_instructor_note(doc, "1分钟个人思考，不全班分享，让学员带着答案进入行动计划制作。")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 14: Part 4 后段 ---
    add_heading(doc, "14. Part 4 后段：行动计划", level=2)
    add_time_tag(doc, "【14:45-16:15 | 90分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "为真实的高潜下属制作完整的教练对话行动计划；配对互检计划质量；在全班分享代表性计划，强化学习。")

    add_key_point(doc, "所需材料")
    add_normal_text(doc, "教练行动计划工作纸（手册P37-40，4页）")

    add_key_point(doc, "讲师心理准备")
    add_script(doc, "这个环节是两天课程里\"实践价值\"最高的环节。学员在这里做出来的行动计划，是他们离开课程之后最可能真正用到的东西。不要因为前面的演练超时而压缩这个环节——如果时间不够，宁可压缩收尾，保护行动计划的制作时间。")
    add_script(doc, "行动计划的质量评判标准：能不能被这位管理者在下周就拿出来用，而不是还需要大量额外准备。太抽象的计划（\"下次我会更好地聆听\"）是失败的；具体到了\"我会在周三和小李的一对一里，用这个提问开场\"是成功的。")

    add_heading(doc, "14.1 行动计划工作纸导引", level=3)
    add_time_tag(doc, "【14:45-14:50 | 5分钟】")

    add_script(doc, "\"翻到手册P37，教练行动计划工作纸。这份计划，针对你这两天一直在想的那个分析对象。\"")
    add_script(doc, "\"计划分四个部分：\"")
    add_script(doc, "\"第一部分——被教练者的双峰定位（他的人类深度峰和AI杠杆峰各是什么）；\"")
    add_script(doc, "\"第二部分——选择教练模式（GUIDE还是DIRECT，理由是什么）；\"")
    add_script(doc, "\"第三部分——教练对话关键设计（开场、核心问题清单、可能的障碍预测）；\"")
    add_script(doc, "\"第四部分——追踪与执行计划（什么时候做，怎么追踪，如何确认效果）。\"")
    add_script(doc, "\"你有25分钟独立完成。我现在给你们三个提示，然后就让你们自己写了。\"")
    add_script(doc, "\"提示一：双峰定位里，人类深度峰要写具体的能力或判断力，不是'他很厉害'这类泛话；提示二：核心问题清单里，至少有一个突破性问题；提示三：追踪计划里，至少有一个人类贡献可见性的追踪指标，不只是产出数字。\"")

    add_heading(doc, "14.2 个人制作", level=3)
    add_time_tag(doc, "【14:50-15:15 | 25分钟】")

    add_instructor_note(doc, "25分钟，安静，背景音乐，讲师走动但不打扰。观察：有没有学员在双峰定位里写得非常模糊——如果看到，轻声问：\"人类深度峰，能更具体到他的哪类判断力或专业直觉吗？\"。")

    add_heading(doc, "14.3 配对互检", level=3)
    add_time_tag(doc, "【15:15-15:30 | 15分钟】")

    add_script(doc, "\"找一个今天和你讨论过你的分析对象的伙伴——或者找一个你信任能够给你真实反馈的人。\"")
    add_script(doc, "\"把你的计划给他看，请他用三个问题检验：\"")
    add_script(doc, "\"一，如果你的被教练者明天看到你写的人类深度峰的描述，他会感到被真正看见了，还是感觉这只是泛泛的描述？\"")
    add_script(doc, "\"二，你的核心问题清单里，有没有一个问题，你感觉自己真的还不确定对方会如何回答？（如果每个你都能预测回答，那可能问题开放度不够。）\"")
    add_script(doc, "\"三，你的追踪指标里，有没有一个，是在关注他的人类判断贡献的成长，而不只是产出数字？\"")
    add_instructor_note(doc, "10分钟配对，5分钟各自做修改。")

    add_heading(doc, "14.4 小组分享与全班汇集", level=3)
    add_time_tag(doc, "【15:30-16:00 | 30分钟】")

    add_script(doc, "\"三到四人一组，每人90秒说你计划里最关键的一个教练问题——就说一个，不要说整个计划。\"")
    add_script(doc, "\"其他人听完，问一个反馈问题：这个问题，被教练者可能会如何回答，他的回答会让你了解他在人类深度峰上的什么信息？\"")
    add_instructor_note(doc, "每组约10分钟。")

    add_key_point(doc, "全班汇集（15分钟）")
    add_script(doc, "\"每组选出一个你们觉得最有穿透力的教练问题——说给全班听。\"")
    add_script(doc, "每组报告，讲师把所有问题写在白板上。")
    add_script(doc, "\"看这些问题——有没有一个你没想到，但觉得'这个问题放在我的场景里，也很有价值'的？\"")
    add_instructor_note(doc, "停顿，让学员自己说，或者自己在心里记下来。")

    add_script(doc, "\"这些问题，是你们两天演练和思考的结晶。每一个都是真实可用的教练问题——不是书本上的，是你们从真实场景里设计出来的。\"")
    add_script(doc, "\"带走它们，和你们的计划一起用。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # --- Section 15: 收尾与教练承诺 ---
    add_heading(doc, "15. 收尾与教练承诺", level=2)
    add_time_tag(doc, "【16:15-17:00 | 45分钟】")

    add_key_point(doc, "本段核心目标")
    add_normal_text(doc, "整合两天的核心学习；每位学员做出一个具体的教练承诺；建立问责伙伴；仪式性结束两天课程。")

    add_heading(doc, "15.1 两天核心框架整合", level=3)
    add_time_tag(doc, "【16:15-16:23 | 8分钟】")

    add_key_point(doc, "展示整合图（PPT，一张包含全部框架的地图）：")
    add_script(doc, "\"两天，我们建立了一套教练型领导力的系统——\"")
    add_script(doc, "\"基础：教练思维（从给答案到激发洞察，从管理者到登攀者向导）\"")
    add_script(doc, "\"新框架：双峰教练视角（人类深度峰+AI杠杆峰，帮助高潜员工在AI时代找到真正的攀登方向）\"")
    add_script(doc, "\"发展型工具：GUIDE五步（Goal-Understand-Insight-Design-Enable）\"")
    add_script(doc, "\"核心技巧：有意识聆听（四层次，含AI时代身份信号层）、有意识提问（含突破性问题设计）、发展性反馈（含AI时代人类贡献的命名）、认同技巧\"")
    add_script(doc, "\"业绩型工具：DIRECT六步（Data-Impact-Requirement-Explore-Commitment-Track）\"")
    add_script(doc, "\"AI时代核心升级：DATA的人类贡献归因分析；REQUIREMENT的双维度绩效期望；三类心智阻碍的识别和应对\"")
    add_script(doc, "\"这不是一套理论，是一套对话技能——你用得越多，越自然，越有效。从你的分析对象开始，第一次对话从你的行动计划开始。\"")

    add_heading(doc, "15.2 教练承诺", level=3)
    add_time_tag(doc, "【16:23-16:38 | 15分钟】")

    add_script(doc, "\"翻到手册P40，教练承诺页。\"")
    add_script(doc, "\"三件事，每件都要写具体，不写泛话——\"")
    add_script(doc, "\"第一件：在接下来30天内，我会和（被教练者代号）做第一次教练对话，时间是（具体的日期），用的是（GUIDE还是DIRECT），开场的第一个问题是（你今天设计的那个问题）。\"")
    add_script(doc, "\"第二件：在这次教练对话里，我最想练习的一个技巧是（写具体的技巧名称），我会在哪一步用它（写具体的步骤）。\"")
    add_script(doc, "\"第三件：三个月后，我会如何知道这次教练有没有效果——我会看到什么变化（写具体的、可观察的信号）？\"")
    add_script(doc, "\"10分钟，独立写。\"")
    add_instructor_note(doc, "10分钟，讲师在场内慢慢走动，不打扰。")

    add_script(doc, "\"写完了？再检查一件事：你的承诺里，有没有一个具体的日期？没有日期的承诺，是没有约束力的愿望。\"")
    add_instructor_note(doc, "让没有写日期的人补上。")

    add_heading(doc, "15.3 问责伙伴配对", level=3)
    add_time_tag(doc, "【16:38-16:50 | 12分钟】")

    add_script(doc, "\"找一个今天和你有相近领导挑战的伙伴——或者你觉得他会认真帮你履行承诺的人。\"")
    add_script(doc, "\"互相分享你的教练承诺里的第一件事——具体的日期、被教练者和第一个问题。\"")
    add_script(doc, "\"然后，做一个约定：在你们约好的日期两周后，你给他发一条消息，说这次教练对话发生了，以及你的感受和发现。他的工作是提醒你，并在收到消息后问你：'最意外的发现是什么？'\"")
    add_script(doc, "\"把伙伴的名字和联系方式写在手册P40的底部。\"")
    add_instructor_note(doc, "8分钟配对+记录。")

    add_heading(doc, "15.4 收尾", level=3)
    add_time_tag(doc, "【16:50-17:00 | 10分钟】")

    add_script(doc, "（话术：语速放慢，不催）")
    add_script(doc, "\"登攀者，这门课的名字，我想在结束的时候说一遍它的含义。\"")
    add_script(doc, "\"登攀，不是管理者在前面领路，员工在后面跟着走。\"")
    add_script(doc, "\"是管理者站在员工旁边——看到那座员工自己还没有看清楚的山，帮助他看见；看到他已经有的、还没有充分使用的装备，帮助他意识到；在他卡住的时候，问一个他自己问不到自己的问题，让他找到下一步。\"")
    add_script(doc, "\"AI时代，这座山在哪里、那个装备是什么，都变得更复杂了。但管理者和员工站在一起、共同面对这个问题的方式，没有变。\"")
    add_instructor_note(doc, "停顿。")
    add_script(doc, "\"你今天带走的工具，是用来帮你站得更稳、看得更清楚的。剩下的，在你和你的员工的真实对话里。\"")
    add_script(doc, "\"谢谢大家两天的投入和勇气——在课堂里做教练演练，需要真实的勇气。\"")

    doc.add_paragraph()
    add_prep_notes_area(doc)

    # ========== APPENDICES ==========
    add_heading(doc, "附录", level=1, color=DARK_BLUE_KEY)

    add_heading(doc, "附录A：AI时代六类场景卡", level=2)

    scenario_table = doc.add_table(rows=7, cols=3)
    scenario_table.style = 'Table Grid'
    add_table_header_row(scenario_table, ["场景", "类型", "描述"])

    scenarios = [
        ("A", "价值焦虑型", "员工持续担心\"AI替代了我，我还有什么价值\"，影响投入度和方向感。需要用GUIDE帮助他找到AI真正无法替代的人类深度贡献，重建职业价值锚点。"),
        ("B", "过度依赖型", "员工用AI承接了大量本应由自己思考的工作，短期产出不错，但判断力和创造力在退化。需要用GUIDE帮助他识别哪些工作是人类深度成长的机会。"),
        ("C", "无效抵触型", "员工拒绝使用AI工具，以维护专业认同，但在效率上开始落后。需要用GUIDE帮助他探索\"如果AI承接了X部分，你的时间和精力会流向哪里\"。"),
        ("D", "接受更大责任", "员工有能力接受更大挑战，但缺乏自信或方向感。需要用GUIDE帮助他看到自己的人类深度峰，建立攀登更大山峰的信心。"),
        ("E", "克服失望", "员工在AI时代的某个挫折后失去信心，需要重新找回自己的价值定位。需要用GUIDE帮助他区分\"AI做得好\"和\"我的判断好\"，重建职业信心。"),
        ("F", "处理业绩缺陷", "员工的绩效数字达标，但核心产出里人类判断贡献极少，近期无法说清楚方案背后的逻辑。需要用DIRECT处理这个具体的行为/绩效问题。"),
    ]
    for i, (scene, stype, desc) in enumerate(scenarios):
        scenario_table.rows[i+1].cells[0].text = scene
        scenario_table.rows[i+1].cells[1].text = stype
        scenario_table.rows[i+1].cells[2].text = desc

    doc.add_paragraph()
    add_prep_notes_area(doc)

    add_heading(doc, "附录B：三层复盘法速查表", level=2)

    debrief_ref_table = doc.add_table(rows=4, cols=3)
    debrief_ref_table.style = 'Table Grid'
    add_table_header_row(debrief_ref_table, ["层次", "核心问题", "聚焦方向"])

    debrief_ref = [
        ("第一层：What", "\"被教练者们——在这段时间里，有没有一个时刻，你感到一种真正的洞察出现了？那是什么时刻？\"", "被教练者的体验"),
        ("第二层：So What", "\"管理者们——今天演练里，最难的部分是什么？\"", "教练者的挑战和成长点"),
        ("第三层：Now What", "\"从今天的演练里，你带走了什么可以用在下次和他对话里的？\"", "具体应用，与真实工作的连接"),
    ]
    for i, (level, question, focus) in enumerate(debrief_ref):
        debrief_ref_table.rows[i+1].cells[0].text = level
        debrief_ref_table.rows[i+1].cells[1].text = question
        debrief_ref_table.rows[i+1].cells[2].text = focus

    doc.add_paragraph()

    # ========== PAGE HEADERS ==========
    # Note: python-docx doesn't support per-page headers easily
    # We'll add a note about this

    add_instructor_note(doc, "注意：实际Word文档需要通过页眉设置功能添加当前模块名称。建议在最终文档中，通过Word的\"插入 > 页眉\"功能，为每个部分添加相应的模块名称。")

    # ========== SAVE DOCUMENT ==========
    output_path = "D:/新课开发/战略和领导力/登攀者——AI时代的授权赋能领导力/完整课程包/04_讲师手册/登攀者讲师手册_完整版.docx"

    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()

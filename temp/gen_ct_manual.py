# -*- coding: utf-8 -*-
import os
import sys

outputPath = r"D:\新课开发\经验萃取\批判思维\完整课程包\05_讲师手册\讲师手册_批判思维与AI.docx"
os.makedirs(os.path.dirname(outputPath), exist_ok=True)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = "C62828"
DARK_GRAY = "424242"
LIGHT_GRAY = "F5F5F5"
WHITE = "FFFFFF"

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)

def add_tip_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = "Microsoft YaHei"
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), RED)
    pPr.append(shd)

def add_script_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    run.font.name = "Microsoft YaHei"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), RED)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    pPr.append(shd)

def add_activity_box(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("【互动活动】" + text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(21, 101, 192)
    run.font.name = "Microsoft YaHei"
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), "2E75B6")
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "E3F2FD")
    pPr.append(shd)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Microsoft YaHei"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, DARK_GRAY)
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(51, 51, 51)
                run.font.name = "Microsoft YaHei"
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph()

def add_timeline_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = "Microsoft YaHei"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, RED)
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            if cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(51, 51, 51)
                run.font.name = "Microsoft YaHei"
            if row_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
    doc.add_paragraph()

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# COVER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI时代批判思维与幻觉识别")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("讲师手册")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("讲师指南")
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 102, 102)
run.font.name = "Microsoft YaHei"

for t in ["课程类型：2天工作坊（每天6小时，共12小时）", "目标学员：知识工作者、管理者、决策者", "版本：v1.0"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(t)
    run.font.size = Pt(11)

add_tip_box(doc, "国际版权课标准 | 完整课程包 | 可直接授课")
doc.add_page_break()

# TOC
p = doc.add_paragraph()
run = p.add_run("目录")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

for ch in ["第一章  讲师资质与准备", "第二章  课程导入设计", "第三章  模块一：AI幻觉识别（2小时）", "第四章  模块二：批判思维四步法（4小时）", "第五章  模块三：验证工具箱（3小时）", "第六章  模块四：实战演练（3小时）", "第七章  课程收尾", "第八章  附录"]:
    p = doc.add_paragraph()
    run = p.add_run(ch)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 51, 51)
    run.font.name = "Microsoft YaHei"

doc.add_page_break()

# CHAPTER 1
p = doc.add_paragraph()
run = p.add_run("第一章 讲师资质与准备")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("1.1 讲师能力要求")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["能力维度", "具体要求", "评估方式"], [["专业知识", "深入理解AI幻觉的成因、类型及识别方法", "内部认证考核"], ["批判思维", "自身具备扎实的批判性思考能力", "案例分析评审"], ["授课技巧", "能够引导成人学习，善于提问与反馈", "试讲评估"], ["课堂把控", "灵活应对不同学员群体和突发情况", "实际授课观察"], ["AI工具使用", "熟练使用主流AI工具并能演示其局限性", "工具操作考核"]])

p = doc.add_paragraph()
run = p.add_run("1.2 课前准备清单")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["准备项目", "具体内容", "完成状态"], [["PPT课件", "全部模块课件已完成并备份", "□"], ["学员手册", "印刷完毕，数量足够", "□"], ["案例素材", "所有AI幻觉案例已准备就绪", "□"], ["验证工具", "演示用AI工具账号已登录", "□"], ["分组材料", "小组讨论道具已分类整理", "□"], ["评估工具", "评分标准表、反馈表已打印", "□"], ["设备检查", "投影、音响、白板等设备正常", "□"]])

p = doc.add_paragraph()
run = p.add_run("1.3 设备与材料检查")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["设备/材料", "检查要点", "紧急预案"], [["电脑/投影", "画面清晰度、HDMI连接正常", "备用笔记本电脑"], ["音响", "音量测试、无线麦克风电量", "备用麦克风"], ["白板/马克笔", "书写流畅、颜色充足", "备用马克笔"], ["网络", "AI工具可正常访问", "本地缓存案例"], ["备用电源", "接线板、UPS准备", "缩短演示内容"]])

p = doc.add_paragraph()
run = p.add_run("1.4 教室布置要求")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["布置要素", "标准要求"], [["座位安排", "小组围坐式，每组4-6人"], ["投影位置", "讲师正前方或侧边，不背对学员"], ["走动通道", "讲师可自由走动巡视全教室"], ["茶歇区域", "与授课区适当分离"], ["灯光空调", "光线充足、温度适宜"]])
doc.add_page_break()

# CHAPTER 2
p = doc.add_paragraph()
run = p.add_run("第二章 课程导入设计")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("2.1 开场破冰活动（完整话术和时间线）")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_timeline_table(doc, ["时间", "活动内容", "讲师话术", "物料"], [["0-2分钟", "问候与自我介绍", "「大家好，我是XXX。过去10年我一直从事AI研究，今天和大家一起探讨如何在AI时代保持清醒的思考。」", "无"], ["2-5分钟", "抛出震撼问题", "「请问在座各位，你们最近一次相信AI给出的答案，后来发现是错误的是什么情况？」", "PPT"], ["5-8分钟", "学员分享", "「好的，请分享一个案例。其他人也可以补充。」（等待3-4人分享）", "无"], ["8-12分钟", "引入主题", "「你们分享的这些案例，都指向同一个问题——AI幻觉。今天我们将系统性地学习如何识别和应对AI幻觉。」", "PPT"], ["12-15分钟", "课程框架说明", "「这门课分为四个模块：识别幻觉、批判思维四步法、验证工具、实战演练。两天的学习后，你们将具备独立验证AI输出真伪的能力。」", "PPT"]])

p = doc.add_paragraph()
run = p.add_run("2.2 学习氛围建立")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_tip_box(doc, "黄金法则：讲三练七——讲师讲解占30%，学员互动占70%")
for t in ["1. 安全感建立：强调「没有愚蠢的问题」，鼓励提问", "2. 平等氛围：讲师不是权威，是引导者，学员的经验同样宝贵", "3. 允许质疑：鼓励学员质疑讲师、质疑AI、质疑教材", "4. 隐私保护：分享个人AI失误经历时无需透露具体场景"]:
    p = doc.add_paragraph()
    p.add_run(t).font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("2.3 学员需求评估")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "3分钟快速调研：让学员匿名写下最想解决的AI幻觉场景")
add_script_box(doc, "「请每位学员拿出一张纸，写下你们最近一次被AI'骗'的经历——哪怕只是一个小错误。不用写名字，但要把具体情况记下来。」（等待3分钟，然后收集）")
p = doc.add_paragraph()
p.add_run("收集后可快速归类：1. 事实性错误（数据、日期） 2. 逻辑错误（推理过程） 3. 语境错误（理解偏差） 4. 创造性幻觉（编造不存在的引用）").font.size = Pt(11)
doc.add_page_break()

# CHAPTER 3
p = doc.add_paragraph()
run = p.add_run("第三章 模块一：AI幻觉识别")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
run = p.add_run("时间：2小时（120分钟）")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

p = doc.add_paragraph()
run = p.add_run("3.1 教学目标")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["目标层级", "描述", "可衡量标准"], [["记忆层", "识别AI幻觉的定义与分类", "能准确说出4种幻觉类型"], ["理解层", "解释各类幻觉的产生机制", "能用自己的话阐述AI幻觉成因"], ["应用层", "在给定的AI输出中识别幻觉", "准确率≥80%"], ["分析层", "分析特定场景下幻觉的风险等级", "能评估不同场景的潜在危害"]])

p = doc.add_paragraph()
run = p.add_run("3.2 详细时间线（分钟级）")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_timeline_table(doc, ["时间", "教学活动", "物料", "话术要点"], [["0-10min", "开场引入", "PPT/数据", "「AI越来越强大，但有一个问题始终没有解决——它会一本正经地胡说八道」"], ["10-25min", "什么是AI幻觉", "PPT", "定义讲解：AI幻觉是指AI生成的内容听起来合理但实际错误或虚构的现象"], ["25-45min", "四类幻觉分类", "PPT/案例", "1.事实性错误 2.逻辑错误 3.语境错误 4.创造性幻觉"], ["45-70min", "为什么AI会产生幻觉", "PPT/白板", "从训练数据、注意力机制、概率生成三个维度解释"], ["70-90min", "真实案例展示", "PPT/演示", "展示3-5个真实的AI幻觉案例，包括新闻、代码、医学建议等"], ["90-110min", "案例分析练习", "学员手册", "小组讨论：分析案例属于哪类幻觉，有何危害"], ["110-120min", "小结与过渡", "PPT", "预告下一模块：我们如何用批判思维来识别这些幻觉"]])

p = doc.add_paragraph()
run = p.add_run("3.3 讲解要点与话术")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("核心定义讲解")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「AI幻觉不是AI'撒谎'，而是AI在处理信息时产生的系统性错误。它产生的原因是AI本质上是一个概率模型，它在'猜'下一个词最可能是什么，而不一定是在'回忆'正确答案。」")

p = doc.add_paragraph()
run = p.add_run("四类幻觉分类讲解")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「第一类是事实性错误——AI说得像真的一样，但数字、日期、人名都是错的。比如AI说'2024年马斯克收购了推特'，但实际收购发生在2022年。第二类是逻辑错误——推理过程看起来完整，但结论经不起推敲。第三类是语境错误——AI没有理解问题的真正含义，答非所问。第四类是创造性幻觉——AI编造不存在的书籍、文章、引用，看起来非常权威。」")

p = doc.add_paragraph()
run = p.add_run("3.4 互动活动设计")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "找茬游戏（15分钟）：展示3个AI回答，让学员找出其中的幻觉")
add_script_box(doc, "「现在我给你们展示三个AI的回答。你们的任务是——找出其中哪些内容是'幻觉'。注意，有些回答看起来非常可信，但可能藏着错误。小组讨论5分钟后，每组派代表汇报。」")

p = doc.add_paragraph()
run = p.add_run("3.5 常见问题应答")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["学员问题", "标准回答"], [["AI越来越先进还会幻觉吗？", "会的。幻觉是LLM架构的固有问题，不是技术不够先进就能解决的。"], ["哪个AI幻觉最少？", "目前没有绝对不幻觉的AI。关键是用批判思维验证，不是依赖某一家的答案。"], ["AI自己知道吗？", "有时候能识别，但大多数时候AI会非常自信地坚持自己的错误答案。"], ["幻觉和偏见有什么区别？", "幻觉是无中生有，偏见是以偏概全。两者都需要批判思维来识别。"]])

p = doc.add_paragraph()
run = p.add_run("3.6 过渡到下一模块")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「今天上午的上半场结束了，我们学会了识别AI幻觉。但光识别还不够——我们需要一套方法来验证AI给出的答案是否可靠。下一模块，我将给大家介绍批判思维四步法——质疑假设、验证来源、检验证据、推演结论。这套方法将帮助你们在AI时代保持清醒的判断力。」")
doc.add_page_break()

# CHAPTER 4
p = doc.add_paragraph()
run = p.add_run("第四章 模块二：批判思维四步法")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
run = p.add_run("时间：4小时（240分钟）")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

p = doc.add_paragraph()
run = p.add_run("4.1 教学目标")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["目标层级", "描述", "可衡量标准"], [["记忆层", "掌握批判思维四步法的名称与步骤", "能完整背诵四步法"], ["理解层", "理解每一步的核心原理与作用", "能解释为什么要这样做"], ["应用层", "能够将四步法应用于实际AI输出验证", "完成独立验证练习"], ["分析层", "能够根据情境灵活运用四步法组合", "解决复杂验证场景"]])

p = doc.add_paragraph()
run = p.add_run("4.2 详细时间线（分钟级）")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_timeline_table(doc, ["时间", "教学活动", "物料", "话术要点"], [["0-10min", "模块引入", "PPT", "「批判思维不是批评，而是系统性思考的方法论」"], ["10-50min", "第一步：质疑假设（40分钟）", "PPT/案例", "讲解什么是假设、如何识别隐藏假设"], ["50-90min", "第一步练习（40分钟）", "学员手册", "小组练习：找出现有AI回答中的假设"], ["90-100min", "茶歇（10分钟）", "", ""], ["100-140min", "第二步：验证来源（40分钟）", "PPT/演示", "来源可靠性评估五要素"], ["140-180min", "第二步练习（40分钟）", "学员手册", "练习验证AI引用的来源真实性"], ["180-220min", "第三步：检验证据（40分钟）", "PPT/案例", "证据类型、证据质量评估"], ["220-240min", "第三步练习+过渡（20分钟）", "学员手册", "完成练习并预告第四步"]])

p = doc.add_paragraph()
run = p.add_run("4.3 第一步：质疑假设")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("讲解话术")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「什么是假设？假设是AI回答中没有明说但默认正确的事情。比如AI说'根据最新研究，早睡早起比晚睡晚起更健康'——它默认了这个研究是存在的、是真的被发表了的。再比如AI说'根据我们的讨论'——它假设你们之前确实讨论过这个话题。识别假设，是批判思维的第一步。」")

p = doc.add_paragraph()
run = p.add_run("练习设计：假设识别")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "练习时长：20分钟，小组进行")
add_script_box(doc, "「现在请阅读学员手册第15页的三个AI回答。每个回答都有2-3个隐藏假设，用红色标出来。小组讨论10分钟后，我们一起核对答案。」")
add_table(doc, ["AI回答", "隐藏假设"], [["「根据最新研究，早睡早起比晚睡晚起更健康」", "1.存在这样的研究 2.研究的结论是可靠的"], ["「这个问题医学界已经有定论了」", "1.确实有定论 2.定论是科学共识"], ["「大部分人认为...」", "1.大部分人的观点是正确的 2.样本有代表性"]])

p = doc.add_paragraph()
run = p.add_run("4.4 第二步：验证来源")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("讲解话术")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「第二步是验证来源。AI经常会说'根据某研究'、'数据显示'、'专家说'——但这些来源真的存在吗？可靠吗？我们需要问五个问题：这个研究是谁做的？发表在哪里？其他专家认可吗？样本量够吗？有没有利益冲突？任何一个问题回答'否'，这个来源就值得怀疑。」")

p = doc.add_paragraph()
run = p.add_run("来源评估五要素")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["要素", "问题", "评估标准"], [["权威性", "发布者是否权威？", "政府机构、学术期刊、知名媒体优先"], ["可验证性", "来源可以查到吗？", "DOI号、URL、可追溯"], ["时效性", "信息是最新的吗？", "是否过时、是否被更新"], ["客观性", "有没有偏见？", "是否有利益关联"], ["相关性", "来源和话题相关吗？", "是否直接支持论点"]])

p = doc.add_paragraph()
run = p.add_run("练习设计：来源验证")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "练习时长：25分钟，角色扮演")
add_script_box(doc, "「假设你是一个编辑，收到了一篇引用了5个来源的文章。请判断每个来源是否可信。小组扮演'研究员'和'怀疑者'两个角色，研究员负责辩护每个来源，怀疑者负责提出质疑。」")

p = doc.add_paragraph()
run = p.add_run("4.5 第三步：检验证据")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("讲解话术")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「第三步是检验证据。即使来源可靠，证据本身也需要检验。证据有强弱之分：个人经验 < 案例 < 相关性研究 < 因果性研究 < 随机对照实验。我们要看数据的样本量、是否控制了变量、结论是否被过度推广。特别注意——相关性不等于因果性。AI经常会把两个同时发生的事情说成因果关系。」")

p = doc.add_paragraph()
run = p.add_run("证据质量金字塔")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["证据级别", "类型", "可信度", "AI常见错误"], [["Level 5", "个人经验/案例", "低", "用个别案例证明普遍规律"], ["Level 4", "相关性研究", "中", "混淆相关性与因果性"], ["Level 3", "准实验研究", "中高", "过度推广适用范围"], ["Level 2", "随机对照实验", "高", "样本不具代表性"], ["Level 1", "元分析/系统综述", "最高", "引用过时或不完整的综述"]])

p = doc.add_paragraph()
run = p.add_run("4.6 第四步：推演结论")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("讲解话术")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「第四步是推演结论。即使前几步都没问题，结论本身也需要检验。推演结论包括三个方面：1.逻辑是否有效——从证据到结论的推理过程有没有漏洞？2.结论是否被过度推广——AI是否把一个小样本的结论说成普适规律？3.是否还有其他合理解释——同样的证据能否支持其他结论？」")

p = doc.add_paragraph()
run = p.add_run("常见逻辑谬误")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["谬误类型", "特征", "AI例子"], [["滑坡谬误", "假设一点点偏离会导致灾难性后果", "如果AI助手普及，医生都会失业"], ["稻草人谬误", "攻击一个弱化版的论点", "AI研究者认为AI有风险=AI不能发展"], ["诉诸权威", "权威说的一定对", "某知名科学家说AI有风险，所以..."], ["虚假两难", "只有两个选项，实际上有更多", "要么完全相信AI，要么完全不用AI"], ["循环论证", "用结论证明前提", "AI是可靠的，因为它的设计是可靠的"]])

p = doc.add_paragraph()
run = p.add_run("4.7 四步法整合练习")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "综合练习时长：45分钟")
add_script_box(doc, "「现在进入最关键的环节——综合练习。手册第25页有一个完整的AI分析报告，请运用四步法，每一步都写出你的分析过程：1.质疑假设——报告中有哪些假设？2.验证来源——引用的来源可靠吗？3.检验证据——证据质量如何？4.推演结论——结论是否合理？最终给这个AI报告打分，并说明理由。」")
add_tip_box(doc, "提示：这道练习是两天课程的核心，请确保每位学员都完成。可安排15分钟小组讨论后再个人独立完成书面报告。")

p = doc.add_paragraph()
run = p.add_run("4.8 过渡到下一模块")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「大家已经掌握了批判思维四步法。但光有方法还不够，我们还需要工具来提高验证效率。下一模块，我将介绍一套验证工具箱——从搜索引擎到学术数据库，从事实核查网站到逻辑检验清单。学完这个模块，你们就能把四步法落地到日常工作中了。」")
doc.add_page_break()

# CHAPTER 5
p = doc.add_paragraph()
run = p.add_run("第五章 模块三：验证工具箱")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
run = p.add_run("时间：3小时（180分钟）")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

p = doc.add_paragraph()
run = p.add_run("5.1 教学目标")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["目标层级", "描述", "可衡量标准"], [["应用层", "熟练使用至少5种验证工具", "能现场演示工具操作"], ["分析层", "根据验证对象选择合适工具", "能快速判断需要查什么"], ["评估层", "综合多种工具验证复杂信息", "完成综合验证任务"]])

p = doc.add_paragraph()
run = p.add_run("5.2 详细时间线")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_timeline_table(doc, ["时间", "教学活动", "物料", "话术要点"], [["0-20min", "工具箱概述", "PPT", "介绍验证工具的分类和使用场景"], ["20-60min", "搜索引擎验证技巧（40分钟）", "演示/实操", "高级搜索指令、site:、时间筛选"], ["60-100min", "学术资源验证（40分钟）", "演示", "Google Scholar、DOI验证、期刊分级"], ["100-110min", "茶歇", "", ""], ["110-150min", "事实核查工具（40分钟）", "演示", "Fact-check网站、AI检测工具"], ["150-180min", "批判性提示词（30分钟）", "演示", "如何向AI提问获得更可靠的答案"]])

p = doc.add_paragraph()
run = p.add_run("5.3 工具演示技巧")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("搜索引擎验证五步法")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「当AI引用了一个事实或数据，第一步不是相信或不相信，而是立刻打开搜索引擎验证。我的标准流程是：1.精确搜索AI的原话或关键词 2.加上site:gov或site:edu限制来源 3.加上时间范围看最新信息 4.查找原始来源而非二手引用 5.对比3个以上独立来源。这个流程熟练后，验证一个事实只需要2-3分钟。」")

p = doc.add_paragraph()
run = p.add_run("学术来源验证清单")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
for t in ["1. Google Scholar搜索论文标题", "2. 验证DOI号是否真实有效", "3. 检查发表期刊的影响因子和声誉", "4. 查看论文被引用次数和引用者", "5. 确认研究是否经过同行评审"]:
    p = doc.add_paragraph()
    p.add_run(t).font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("5.4 练习活动设计")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "工具实操练习：30分钟")
add_script_box(doc, "「手册第30页有10个待验证的事实陈述。请使用工具箱中的工具，至少验证其中5个。记录你用了什么工具、验证步骤是什么、最终结论是什么。这个练习的目的是让你们形成肌肉记忆——遇到AI输出，下意识就去验证。」")

p = doc.add_paragraph()
run = p.add_run("5.5 批判性提示词讲解")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("核心原则")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_tip_box(doc, "向AI提问时，明确要求它提供来源，并说明你打算验证")

p = doc.add_paragraph()
run = p.add_run("高效提示词模板")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["场景", "提示词模板", "效果"], [["查事实", "请提供这个信息的来源，包括原始链接或DOI", "强制溯源"], ["查数据", "这个数据来自哪个机构的什么调查？样本量是多少？", "要求元数据"], ["查引用", "请提供这个引用的原文，包括作者、标题、发表期刊、年份", "防止虚假引用"], ["不确定时", "如果你不确定这个信息，请明确说你不确定", "避免幻觉"], ["多重验证", "请用三个不同来源来支持这个观点，并说明哪个最可靠", "强制交叉验证"]])
doc.add_page_break()

# CHAPTER 6
p = doc.add_paragraph()
run = p.add_run("第六章 模块四：实战演练")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
run = p.add_run("时间：3小时（180分钟）")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

p = doc.add_paragraph()
run = p.add_run("6.1 教学目标")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["目标层级", "描述", "可衡量标准"], [["应用层", "综合运用批判思维四步法和工具箱", "完成完整验证流程"], ["分析层", "独立分析真实场景中的AI可靠性", "解决复杂案例"], ["评估层", "形成个人AI使用策略", "制定个人行动承诺"]])

p = doc.add_paragraph()
run = p.add_run("6.2 详细时间线")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_timeline_table(doc, ["时间", "教学活动", "物料", "话术要点"], [["0-15min", "案例选择说明", "PPT", "介绍可选案例和行业背景"], ["15-75min", "小组案例分析（60分钟）", "案例材料", "选择一个案例进行完整四步法分析"], ["75-85min", "茶歇", "", ""], ["85-135min", "小组汇报与点评（50分钟）", "汇报材料", "每组5分钟汇报，其他组补充提问"], ["135-165min", "个人反思与行动承诺（30分钟）", "学员手册", "写下个人应用计划"], ["165-180min", "课程总结与收尾（15分钟）", "PPT", "回顾核心要点，宣布结课"]])

p = doc.add_paragraph()
run = p.add_run("6.3 案例选择指南")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["案例类型", "适用对象", "案例简介", "难度"], [["AI新闻核实", "媒体从业者", "AI生成的一篇新闻报道，需要核实其中的数据和事实", "中"], ["AI医疗建议", "医疗从业者", "AI给出的患者咨询建议，需要评估其安全性", "高"], ["AI法律分析", "法务从业者", "AI对合同条款的分析，需要核实法律依据", "高"], ["AI投资建议", "金融从业者", "AI给出的投资分析，需要验证数据来源", "中"], ["AI教育内容", "教育从业者", "AI生成的教学材料，需要评估准确性", "低"], ["AI代码审查", "技术人员", "AI生成的代码片段，需要验证安全性和效率", "中"]])

p = doc.add_paragraph()
run = p.add_run("6.4 小组演练引导")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "小组讨论引导话术")
add_script_box(doc, "「各小组，现在你们有60分钟对一个案例进行深度分析。我要求：1.先独立思考10分钟，在手册上写下你的初步判断 2.然后小组讨论30分钟，尝试用四步法系统性地分析 3.最后小组达成共识，准备5分钟汇报。记住，没有标准答案——我要看的是你们的思考过程和分析方法。」")

p = doc.add_paragraph()
run = p.add_run("6.5 评估标准")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["维度", "权重", "优秀标准", "合格标准"], [["假设识别", "25%", "准确识别所有关键假设，包括隐藏假设", "识别出大部分主要假设"], ["来源验证", "25%", "使用多个工具验证来源，给出可信度评分", "验证了主要来源但工具使用不够全面"], ["证据评估", "25%", "准确判断证据质量，识别证据局限性", "基本判断正确但分析深度不足"], ["结论推演", "25%", "逻辑严密，考虑替代解释，结论谨慎", "结论基本合理但推演过程有漏洞"]])

p = doc.add_paragraph()
run = p.add_run("6.6 行动承诺引导")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_activity_box(doc, "个人行动承诺")
add_script_box(doc, "「课程的最后，请每位学员写下一个具体的行动承诺——回到工作岗位后，你将在什么时候、用什么方式把学到的批判思维用起来。比如：'从明天起，每次使用AI查资料后，用5分钟时间验证关键数据'。把这个承诺写在手册最后一页，课程结束后我们不收回来——这是你们给自己的承诺。」")
doc.add_page_break()

# CHAPTER 7
p = doc.add_paragraph()
run = p.add_run("第七章 课程收尾")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("7.1 总结与回顾")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「两天的学习即将结束，让我们回顾一下今天学到的核心内容：第一，AI幻觉是LLM架构的固有问题，每个使用AI的人都可能遇到。第二，批判思维四步法——质疑假设、验证来源、检验证据、推演结论——是应对AI幻觉的系统性方法。第三，工具箱能帮助我们高效验证信息，搜索引擎、学术数据库、事实核查网站都是有力的武器。」")
add_script_box(doc, "「但最重要的是——批判思维不是一次性的学习，而是一生的习惯。希望各位在AI时代，都能保持独立思考的能力，不盲目相信，也不盲目怀疑。」")

p = doc.add_paragraph()
run = p.add_run("7.2 学习成果确认")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["能力", "达成情况（自评）", "继续改进方向"], [["识别AI幻觉的四种类型", "□完全掌握 □部分掌握 □还需加强", ""], ["运用批判思维四步法", "□完全掌握 □部分掌握 □还需加强", ""], ["使用验证工具箱", "□完全掌握 □部分掌握 □还需加强", ""], ["综合分析AI输出可靠性", "□完全掌握 □部分掌握 □还需加强", ""]])

p = doc.add_paragraph()
run = p.add_run("7.3 资源推荐")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
run = p.add_run("推荐书籍：")
run.bold = True
run.font.size = Pt(11)
for t in ["1. 《Thinking, Fast and Slow》- Daniel Kahneman", "2. 《The Art of Thinking Clearly》- Rolf Dobelli", "3. 《Critical Thinking: A Very Short Introduction》- Tim Connex"]:
    p = doc.add_paragraph()
    p.add_run(t).font.size = Pt(11)
p = doc.add_paragraph()
run = p.add_run("推荐工具：")
run.bold = True
run.font.size = Pt(11)
for t in ["1. Google Scholar (scholar.google.com) - 学术文献验证", "2. Snopes (snopes.com) - 事实核查", "3. DOI.org (doi.org) - 学术DOI验证", "4. Semantic Scholar (semanticscholar.org) - AI辅助学术搜索"]:
    p = doc.add_paragraph()
    p.add_run(t).font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("7.4 结束语")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_script_box(doc, "「最后送大家一句话——AI时代最稀缺的，不是信息，而是判断信息真伪的能力。希望各位学成回去后，不仅自己用好批判思维，也能把这份能力传递给身边的人。谢谢大家两天的投入和参与！课程到此结束。」")
add_tip_box(doc, "提示：宣布结束后，留出10-15分钟让学员私下提问或交换联系方式建立学习小组")
doc.add_page_break()

# CHAPTER 8
p = doc.add_paragraph()
run = p.add_run("第八章 附录")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor.from_string(RED)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("附录A：时间管理技巧")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["场景", "问题", "应对策略"], [["练习超时", "小组讨论超时", "使用计时器，提前5分钟提醒，严格执行时间节点"], ["学员冷场", "提问无人响应", "先给2分钟独立思考时间，再从主动学员开始引导"], ["偏离主题", "讨论偏离原话题", "温和打断，用「这个问题很有趣，我们可以课后深入探讨」拉回"], ["内容太多", "时间不够用", "优先保证核心练习，删减补充内容，记住「讲三练七」"], ["学员疲惫", "下午场注意力下降", "安排站立活动、换话题、或短暂茶歇"]])

p = doc.add_paragraph()
run = p.add_run("附录B：困难学员应对策略")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["学员类型", "特征表现", "应对策略"], [["沉默型", "全程不说话", "分配具体任务降低参与门槛，课后私下交流"], ["主导型", "话太多主导讨论", "设立「发言人」角色轮流制，感谢其贡献但温和引导他人"], ["质疑一切型", "不断挑战讲师", "肯定其批判精神，将其问题引导回学员讨论"], ["消极抵触型", "不参与活动", "了解抵触原因，降低练习难度，强调学习收益"], ["炫耀型", "用专业术语显摆", "肯定其知识面，将话题引导到实践应用层面"]])

p = doc.add_paragraph()
run = p.add_run("附录C：课程调整方案")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"

p = doc.add_paragraph()
run = p.add_run("1天精华版调整（6小时）")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["原模块", "调整", "时间"], [["模块一", "压缩为30分钟概念引入", "30分钟"], ["模块二", "四步法精讲+1个练习", "120分钟"], ["模块三", "工具箱概述，学员自行练习", "60分钟"], ["模块四", "压缩为1个案例快速分析", "90分钟"]])

p = doc.add_paragraph()
run = p.add_run("半天入门版调整（3小时）")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
p = doc.add_paragraph()
p.add_run("聚焦模块一（AI幻觉识别）和模块二（四步法概述），删除工具箱和实战演练，增加更多互动案例演示").font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run("附录D：讲师反馈表")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
run.font.name = "Microsoft YaHei"
add_table(doc, ["评估维度", "1-5分", "具体建议"], [["课程内容实用性", "", ""], ["讲师授课清晰度", "", ""], ["互动活动有效性", "", ""], ["案例贴近程度", "", ""], ["时间节奏把控", "", ""], ["场地设备满意度", "", ""], ["整体满意度", "", ""]])
for t in ["其他意见或建议：", "______________________________________________________________________", "______________________________________________________________________", "______________________________________________________________________"]:
    p = doc.add_paragraph()
    p.add_run(t).font.size = Pt(11)

doc.save(outputPath)
print(f"Document created successfully: {outputPath}")

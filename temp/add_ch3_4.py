# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {})
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn("w:eastAsia"), "Microsoft YaHei")
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run._r.get_or_add_rPr().set(qn("w:eastAsia"), "Microsoft YaHei")
    return h

output_path = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册/讲师手册-变革成果固化机制.docx"
doc = Document(output_path)

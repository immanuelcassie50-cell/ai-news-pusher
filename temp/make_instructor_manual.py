"""Generate 讲师手册.docx - A4 portrait, ~90 pages.
Merges 10 MD files from 02_讲师手册 into a single Word document.
Features: TOC, time allocation tables, interaction design, instructor scripts.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC_DIR = r"D:\Downloads\利益相关方影响和干预\完整课程包\02_讲师手册"
OUT = r"D:\Downloads\利益相关方影响和干预\完整课程包\13_Office文档\讲师手册.docx"

FILES_ORDER = [
    ("00_讲师手册_总纲.md", "总纲：讲师手册使用说明"),
    ("01_M1_共同语言.md", "模块一  共同语言（讲师版）"),
    ("02_M2_会议准则.md", "模块二  会议准则（讲师版）"),
    ("03_M3_利益相关方分析与引导.md", "模块三  利益相关方分析与引导（讲师版）"),
    ("04_M4_三层目标框架.md", "模块四  三层目标框架（讲师版）"),
    ("05_M5_现象界定与根因追问.md", "模块五  现象界定与根因追问（讲师版）"),
    ("06_M6_回到正轨.md", "模块六  回到正轨（讲师版）"),
    ("07_M7_潜在问题预演.md", "模块七  潜在问题预演（讲师版）"),
    ("08_M8_综合演练.md", "模块八  综合演练（讲师版）"),
    ("09_差异化应对_常见干扰情景.md", "差异化应对  常见干扰情景"),
]

# Color palette
WINE = RGBColor(0x8b, 0x28, 0x28)
WINE_MID = RGBColor(0xc0, 0x39, 0x2b)
GOLD = RGBColor(0xc9, 0xa9, 0x6e)
GOLD_DEEP = RGBColor(0xa8, 0x88, 0x4a)
INK = RGBColor(0x1a, 0x1a, 0x1a)
INK_SOFT = RGBColor(0x3a, 0x3a, 0x3a)
INK_MID = RGBColor(0x6e, 0x6e, 0x6e)
PAPER = RGBColor(0xfa, 0xf6, 0xec)
TINT = RGBColor(0xf8, 0xe6, 0xe1)
SOFT = RGBColor(0xf0, 0xd5, 0xcf)


def set_run_font(run, font_name="Source Han Serif CN", size=10.5, color=INK, bold=False, italic=False):
    """Set both ASCII and East Asian font for a run."""
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:cs'), font_name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_para(doc, text, size=10.5, color=INK, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, indent=0, first_line=0,
             line_spacing=1.5, space_before=0, space_after=4,
             font="Source Han Serif CN"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(first_line)
    run = p.add_run(text)
    set_run_font(run, font_name=font, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    if level == 0:
        # Chapter (主标题)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(40)
        p.paragraph_format.space_after = Pt(16)
        p.paragraph_format.keep_with_next = True
        # Decorative top line
        run = p.add_run("━" * 12)
        set_run_font(run, size=12, color=GOLD, bold=True)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_after = Pt(8)
        run = p2.add_run(text)
        set_run_font(run, size=26, color=WINE, bold=True)
    elif level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        # Number prefix + text
        run = p.add_run("■ ")
        set_run_font(run, size=15, color=GOLD_DEEP, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=15, color=WINE, bold=True)
        # Bottom line
        p_line = doc.add_paragraph()
        p_line.paragraph_format.space_after = Pt(8)
        run = p_line.add_run("─" * 30)
        set_run_font(run, size=8, color=GOLD)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run("▸ ")
        set_run_font(run, size=12, color=GOLD_DEEP, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=12.5, color=INK, bold=True)
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run("· ")
        set_run_font(run, size=11, color=GOLD_DEEP, bold=True)
        run = p.add_run(text)
        set_run_font(run, size=11, color=INK_SOFT, bold=True)
    return p


def add_quote_block(doc, text, color=INK_SOFT):
    """Add a quote/notes block with left border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run("「 ")
    set_run_font(run, size=11, color=GOLD_DEEP, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=color, italic=True)
    run = p.add_run(" 」")
    set_run_font(run, size=11, color=GOLD_DEEP, bold=True)
    return p


def add_tip_block(doc, text, label="提示"):
    """Add a tip box with light background."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    # Shading
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'f8e6e1')
    pPr.append(shd)
    # Border left
    pBdr = OxmlElement('w:pBdr')
    leftBdr = OxmlElement('w:left')
    leftBdr.set(qn('w:val'), 'single')
    leftBdr.set(qn('w:sz'), '24')
    leftBdr.set(qn('w:color'), '8b2828')
    pBdr.append(leftBdr)
    pPr.append(pBdr)
    run = p.add_run(f"■ {label}  ")
    set_run_font(run, size=10, color=WINE, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10, color=INK_SOFT)
    return p


def add_handwrite_line(doc, label="", width_chars=40):
    """Add a hand-write line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0
    if label:
        run = p.add_run(label)
        set_run_font(run, size=10, color=INK, bold=True)
    # Underline empty space
    run = p.add_run("_" * width_chars)
    set_run_font(run, size=10, color=INK_MID)
    return p


def add_bullet(doc, text, level=0, marker="●"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{marker}  ")
    set_run_font(run, size=10.5, color=GOLD_DEEP, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_numbered(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"{n}.  ")
    set_run_font(run, size=10.5, color=WINE, bold=True)
    run = p.add_run(text)
    set_run_run_font(run, size=10.5, color=INK) if False else set_run_font(run, size=10.5, color=INK)
    return p


def set_run_run_font(*args, **kwargs): pass


def add_horizontal_rule(doc, color=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("─" * 60)
    set_run_font(run, size=8, color=color)
    return p


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def parse_markdown_line(line):
    """Parse a single markdown line into segments."""
    # Strip
    line = line.rstrip()
    return line


def add_table_from_rows(doc, rows, col_widths_cm=None):
    """Add a table from list-of-lists."""
    if not rows:
        return None
    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)
    t = doc.add_table(rows=n_rows, cols=n_cols)
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in t.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = t.cell(i, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                # Clear default paragraph
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.3
                run = p.add_run(str(cell_text))
                if i == 0:
                    set_run_font(run, size=10, color=PAPER, bold=True)
                    # Shading
                    tcPr = cell._element.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '8b2828')
                    tcPr.append(shd)
                else:
                    set_run_font(run, size=9.5, color=INK)
                    if i % 2 == 0:
                        tcPr = cell._element.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'faf6ec')
                        tcPr.append(shd)
    return t


def process_md_content(doc, content, is_module=False):
    """Process markdown content, handling common patterns."""
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []

    def flush_table():
        if not table_buffer:
            return
        # Skip separator row (|---|---|)
        cleaned = []
        for r in table_buffer:
            cells = [c.strip() for c in r.strip().strip("|").split("|")]
            if all(re.match(r"^[-:]+$", c) for c in cells):
                continue
            cleaned.append(cells)
        if cleaned:
            add_table_from_rows(doc, cleaned)
        table_buffer.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code block
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                pPr = p._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'f0d5cf')
                pPr.append(shd)
                for cl in code_buffer:
                    run = p.add_run(cl + "\n")
                    set_run_font(run, size=9.5, color=INK)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(stripped)
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.search(r"\|\s*[-:]+\s*\|", lines[i + 1]):
            in_table = True
            table_buffer.append(line)
            i += 1
            continue
        elif in_table and "|" in line:
            table_buffer.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        # Empty line
        if not stripped:
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip markdown formatting
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            if level == 1:
                add_heading(doc, text, level=1)
            elif level == 2:
                add_heading(doc, text, level=2)
            elif level == 3:
                add_heading(doc, text, level=3)
            else:
                add_heading(doc, text, level=2)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            if text.startswith("金句"):
                add_tip_block(doc, text, label="金句")
            elif text.startswith("使用说明") or text.startswith("核心原则"):
                add_tip_block(doc, text, label="提示")
            else:
                add_quote_block(doc, text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            text = m.group(2).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            marker = ["●", "○", "▪"][min(indent, 2)]
            add_bullet(doc, text, level=indent, marker=marker)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            num_match = re.match(r"^\s*(\d+)\.\s+", line)
            n = num_match.group(1) if num_match else "1"
            text = m.group(1).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            add_numbered(doc, n, text)
            i += 1
            continue

        # Checkbox list
        m = re.match(r"^\s*-\s+\[\s*[xX ]\s*\]\s+(.*)$", line)
        if m:
            text = m.group(1).strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            add_bullet(doc, "☐ " + text, level=0, marker="☐")
            i += 1
            continue

        # Normal paragraph
        text = stripped
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        text = re.sub(r"\*(.+?)\*", r"\1", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)

        # If line is just dashes/equals (sometimes heading underline)
        if re.match(r"^[-=]+\s*$", text):
            i += 1
            continue

        # Long paragraph - check if it should be hand-write
        if "_______" in text and len(text) < 100:
            add_handwrite_line(doc, text.replace("_______", "").strip() or "请填写：")
        elif "______" in text and len(text) < 100:
            add_handwrite_line(doc, text.replace("______", "").strip() or "请填写：")
        else:
            add_para(doc, text, size=10.5, first_line=0.74, line_spacing=1.55)
        i += 1

    flush_table()


def setup_doc_styles(doc):
    """Set up document-wide styles and properties."""
    # Page setup
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Source Han Serif CN"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Source Han Serif CN')
    rfonts.set(qn('w:hAnsi'), 'Source Han Serif CN')
    rfonts.set(qn('w:eastAsia'), 'Source Han Serif CN')
    rfonts.set(qn('w:cs'), 'Source Han Serif CN')
    # Core properties
    cp = doc.core_properties
    cp.title = "共同语言：高效项目执行与问题解决工作坊 - 讲师手册"
    cp.author = "罗宏伟"
    cp.subject = "讲师手册 v1.0"
    cp.keywords = "共同语言;项目执行;问题解决;工作坊"


# ==================== MAIN ====================
doc = Document()
setup_doc_styles(doc)

# Cover page
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("共同语言")
set_run_font(run, size=48, color=WINE, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Efficient Project Execution & Problem Solving")
set_run_font(run, size=11, color=GOLD_DEEP, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("━" * 18)
set_run_font(run, size=12, color=GOLD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
run = p.add_run("讲师手册")
set_run_font(run, size=32, color=INK, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("Instructor Handbook  ·  v1.0  ·  2026")
set_run_font(run, size=10, color=GOLD_DEEP)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run("高效项目执行与问题解决工作坊")
set_run_font(run, size=16, color=WINE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("2 天  ·  14 小时  ·  8 模块  ·  30 天跟进")
set_run_font(run, size=11, color=INK_SOFT)

# Decorative
for _ in range(2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 30)
    set_run_font(run, size=8, color=GOLD)

# Instructor
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("主  讲")
set_run_font(run, size=11, color=GOLD_DEEP, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("罗  宏  伟")
set_run_font(run, size=22, color=INK, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run("行动学习催化师")
set_run_font(run, size=11, color=INK_SOFT, italic=True)

# Quote
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.left_indent = Cm(2)
p.paragraph_format.right_indent = Cm(2)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("「")
set_run_font(run, size=14, color=GOLD_DEEP)
run = p.add_run("共同语言，是项目执行里被低估的第一生产力。")
set_run_font(run, size=12, color=INK, italic=True)
run = p.add_run("」")
set_run_font(run, size=14, color=GOLD_DEEP)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(80)
run = p.add_run("—— 让团队的会议更短、返工更少、共识更快、决策更清。")
set_run_font(run, size=10.5, color=INK_MID)

# Page break to TOC
add_page_break(doc)

# Table of Contents (manual)
add_heading(doc, "目  录", level=0)
add_horizontal_rule(doc)

toc_items = [
    ("总纲：讲师手册使用说明", "1"),
    ("模块一  共同语言", "5"),
    ("模块二  会议准则", "12"),
    ("模块三  利益相关方分析与引导", "20"),
    ("模块四  三层目标框架", "30"),
    ("模块五  现象界定与根因追问", "40"),
    ("模块六  回到正轨", "52"),
    ("模块七  潜在问题预演", "64"),
    ("模块八  综合演练", "76"),
    ("行动计划  30天跟进", "88"),
]
for i, (title, page) in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    # Use tab stops
    pPr = p._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '8500')
    tabs.append(tab)
    pPr.append(tabs)
    # Number
    num_run = p.add_run(f"{i:02d}  ")
    set_run_font(num_run, size=11, color=GOLD_DEEP, bold=True)
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, color=INK)
    tab_run = p.add_run("\t")
    set_run_font(tab_run, size=11)
    page_run = p.add_run(page)
    set_run_font(page_run, size=11, color=WINE, bold=True)

# Pre-content note
add_horizontal_rule(doc)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.3)
p.paragraph_format.right_indent = Cm(0.3)
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.line_spacing = 1.6
run = p.add_run("使用建议  ")
set_run_font(run, size=10.5, color=WINE, bold=True)
run = p.add_run("本手册与课程完全对应——上课时跟讲义,下课后用工具。建议在每个模块开始时先阅读「学习目标」,模块结束时填写「本模块反思」。空白横线与方框是手写区域,留作你自己的笔记与行动计划。")
set_run_font(run, size=10, color=INK_SOFT)

add_page_break(doc)

# Process each MD file
chapter_num = 0
for fname, chapter_title in FILES_ORDER:
    fp = os.path.join(SRC_DIR, fname)
    if not os.path.exists(fp):
        continue
    with open(fp, encoding="utf-8") as f:
        content = f.read()
    chapter_num += 1
    # Add chapter title
    if chapter_num > 1:
        add_page_break(doc)
    add_heading(doc, chapter_title, level=0)
    # Process content - skip the original H1 (we already added it)
    # Find first H1 and skip
    lines = content.split("\n")
    start_idx = 0
    for k, line in enumerate(lines):
        if re.match(r"^#\s+", line):
            start_idx = k + 1
            break
    # Skip the subtitle and initial blank lines
    while start_idx < len(lines) and (not lines[start_idx].strip() or re.match(r"^#+\s+", lines[start_idx])):
        if re.match(r"^#+\s+", lines[start_idx]):
            start_idx += 1
        elif not lines[start_idx].strip():
            start_idx += 1
        else:
            break
    new_content = "\n".join(lines[start_idx:])
    process_md_content(doc, new_content, is_module=(chapter_num > 1))

# Final page - signature
add_page_break(doc)
add_heading(doc, "结语  写给 30 天后的你", level=0)
add_horizontal_rule(doc)

p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.8
p.paragraph_format.first_line_indent = Cm(0.74)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("30 天后,当你翻回本手册的第一页,你会发现:")
set_run_font(run, size=11, color=INK)

reflections = [
    "□ 我能用 5 要素句式陈述一个真实的工作问题。",
    "□ 我开会的平均时长变短了,会议结束时有清晰的行动项。",
    "□ 我面对跨部门相关方时,知道先找谁、什么时候找、怎么说。",
    "□ 出现异常时,我能分清「恢复」「直接原因」「根本原因」三层。",
    "□ 我会用 IS/IS NOT 缩小问题范围,找到真正的根本原因。",
    "□ 项目偏离计划时,我会用纠偏对话 4 步处理,而不是「再坚持」或「大改」。",
    "□ 重要行动前,我会做潜在问题预演 + 回滚标准。",
    "□ 我会持续用 30 天行动表来跟进自己,而不是「学完就忘」。",
]
for r in reflections:
    add_bullet(doc, r, marker="□")

# Signature block
add_horizontal_rule(doc)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.line_spacing = 2.0
run = p.add_run("我的 30 天承诺")
set_run_font(run, size=14, color=WINE, bold=True)

add_handwrite_line(doc, "我承诺,今后 30 天,我会: ", width_chars=42)
add_handwrite_line(doc, "", width_chars=72)
add_handwrite_line(doc, "", width_chars=72)
add_handwrite_line(doc, "我的签名:_____________________  日期:_____________  见证人:_____________________", width_chars=24)

# Save
doc.save(OUT)
print(f"Created: {OUT}")
print(f"Size: {os.path.getsize(OUT)} bytes")
print(f"Chapters: {chapter_num}")

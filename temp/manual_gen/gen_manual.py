# -*- coding: utf-8 -*-
"""
生成《打造组织创新力：营造创新土壤》学员手册完整 Word 版本 (32 页)
+ 学员手册 HTML 可视化版
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = r'D:\2026年课程\竞越\打造组织创新力：营造创新土壤\完整课程包\04_学员手册\学员手册_完整版.docx'
OUT_HTML = r'D:\2026年课程\竞越\打造组织创新力：营造创新土壤\完整课程包\04_学员手册\学员手册_HTML可视化版.html'
os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)

# === 配色 ===
DEEP_BLUE = RGBColor(0x1A, 0x2E, 0x4C)
VERMILION = RGBColor(0xC8, 0x44, 0x2A)
PAPER = RGBColor(0xF5, 0xF0, 0xE6)
OCHRE = RGBColor(0xB8, 0x89, 0x3A)
GRAY_DARK = RGBColor(0x55, 0x55, 0x55)
GRAY_LIGHT = RGBColor(0xAA, 0xAA, 0xAA)

CN_FONT = '思源宋体'
CN_FONT_SANS = '思源黑体'
EN_FONT = 'Cambria'
EN_FONT_SANS = 'Calibri'


# === 工具函数 ===
def set_run_font(run, name_cn=CN_FONT, name_en=EN_FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)


def set_para_format(p, alignment=None, space_before=None, space_after=None, line_spacing=None, indent_first=None):
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if indent_first is not None:
        pf.first_line_indent = Cm(indent_first)


def add_para(doc, text, size=11, bold=False, italic=False, color=None, font=CN_FONT, align=None, indent=None, space_after=6, space_before=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.first_line_indent = Cm(indent)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_run_font(run, name_cn=font, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text, color=DEEP_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, name_cn=CN_FONT_SANS, size=20, bold=True, color=color)
    # 底部线条
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '1A2E4C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text, color=VERMILION):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, name_cn=CN_FONT_SANS, size=14, bold=True, color=color)
    return p


def add_h3(doc, text, color=DEEP_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name_cn=CN_FONT_SANS, size=12, bold=True, color=color)
    return p


def add_callout(doc, text, color=OCHRE, italic=True, bold=False):
    """添加引用块（左侧色条）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F0E6')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, size=10, italic=italic, bold=bold, color=GRAY_DARK)
    return p


def add_card_box(doc, title, content, fill_color='1A2E4C'):
    """添加核心概念卡（带标题栏的卡片）"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # 背景色（标题栏）
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(8)
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(f'【{title}】')
    set_run_font(r0, name_cn=CN_FONT_SANS, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 内容行
    for line in content:
        p = cell.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.2)
        for run in p.runs:
            run.font.name = EN_FONT
            run.font.size = Pt(10)
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), CN_FONT)
            rFonts.set(qn('w:ascii'), EN_FONT)
            rFonts.set(qn('w:hAnsi'), EN_FONT)
    # 设置表格列宽
    tbl.columns[0].width = Cm(15)
    cell.width = Cm(15)
    # 设置背景色（仅标题行/整个表格）
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    return tbl


def add_card_box_v2(doc, title, content):
    """带颜色标题栏 + 白底内容的概念卡"""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    # 标题行
    cell0 = tbl.cell(0, 0)
    cell0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(title)
    set_run_font(r0, name_cn=CN_FONT_SANS, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    tcPr0 = cell0._tc.get_or_add_tcPr()
    shd0 = OxmlElement('w:shd')
    shd0.set(qn('w:fill'), '1A2E4C')
    shd0.set(qn('w:val'), 'clear')
    tcPr0.append(shd0)
    # 内容行
    cell1 = tbl.cell(1, 0)
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tcPr1 = cell1._tc.get_or_add_tcPr()
    shd1 = OxmlElement('w:shd')
    shd1.set(qn('w:fill'), 'F5F0E6')
    shd1.set(qn('w:val'), 'clear')
    tcPr1.append(shd1)
    # 清空第一段
    p_first = cell1.paragraphs[0]
    p_first.paragraph_format.space_after = Pt(2)
    if content:
        run = p_first.add_run(content[0])
        set_run_font(run, size=10)
    for line in content[1:]:
        p = cell1.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.2)
        for run in p.runs:
            run.font.size = Pt(10)
    tbl.columns[0].width = Cm(15)
    cell0.width = Cm(15)
    cell1.width = Cm(15)
    return tbl


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_table_styled(doc, headers, rows, col_widths_cm=None, header_fill='1A2E4C'):
    """添加带样式的表格"""
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        set_run_font(r, name_cn=CN_FONT_SANS, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # 内容
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            set_run_font(r, size=10)
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            tbl.columns[i].width = Cm(w)
            for r in tbl.rows:
                r.cells[i].width = Cm(w)
    return tbl


def add_blank_lines(doc, n=3):
    for _ in range(n):
        doc.add_paragraph()


def add_signature_line(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'{label}：')
    set_run_font(run, name_cn=CN_FONT_SANS, size=11, bold=True)
    run2 = p.add_run('________________________________')
    set_run_font(run2, size=11)


# ====================================================================
# 1. 构建 Word 文档
# ====================================================================
def build_docx():
    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)
    # 页面设置
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # === 封面 ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('打造组织创新力')
    set_run_font(r, name_cn=CN_FONT_SANS, size=36, bold=True, color=DEEP_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('营造创新土壤')
    set_run_font(r, name_cn=CN_FONT_SANS, size=36, bold=True, color=VERMILION)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run('—— 拉开组织创新的序章 ——')
    set_run_font(r, name_cn=CN_FONT, size=18, italic=True, color=OCHRE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run('学 员 手 册')
    set_run_font(r, name_cn=CN_FONT_SANS, size=24, bold=True, color=DEEP_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run('STUDENT HANDBOOK · v1.0')
    set_run_font(r, name_en=EN_FONT, size=12, color=GRAY_DARK)

    add_blank_lines(doc, 4)
    info_lines = [
        '学员姓名：______________________',
        '公司/部门：______________________',
        '课程日期：______________________',
        '主讲老师：罗宏伟（竞越）',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line)
        set_run_font(r, name_cn=CN_FONT_SANS, size=12)

    add_page_break(doc)

    # === 手册使用说明 ===
    add_h1(doc, '📖 手册使用说明')
    add_para(doc, '这本手册不是讲义——它是你回到岗位后还要翻 30 次的工具书。', size=11, bold=True, color=DEEP_BLUE)
    add_para(doc, '课堂上：跟着讲师节奏，把空白处写满', size=11, indent=0.5)
    add_para(doc, '课后 1 周内：复习模块三、四、五的关键概念', size=11, indent=0.5)
    add_para(doc, '课后 30/60/90 天：对照自己的"承诺事项"翻阅', size=11, indent=0.5)

    add_callout(doc, '一句话预告今天的核心命题：')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run('"答案不在员工不努力，而在领导者把自己当成了\'种子\'。"')
    set_run_font(r, size=14, bold=True, italic=True, color=VERMILION)

    add_page_break(doc)

    # === 第 1 页：欢迎词 ===
    add_h1(doc, '第 1 页：欢迎词')
    add_callout(doc, '致学员：欢迎来到《打造组织创新力：营造创新土壤》——这不是一门教你"如何创新"的课程，这是一门教你"如何让你的组织能持续创新"的课程。')
    add_para(doc, '今天 6 小时结束后，你带走的不是 9 条原则，是 3 件具体要做的事。', size=11, bold=True, color=DEEP_BLUE, space_after=12)
    add_para(doc, '我们开始。', size=14, bold=True, color=VERMILION, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_page_break(doc)

    # === 第 2 页：阶梯自评表（课前） ===
    add_h1(doc, '第 2 页：阶梯自评表（课前）')
    add_callout(doc, '请用 30 秒/题回答，不要思考太久——第一直觉最重要。', italic=True)
    headers = ['#', '问题', '你的答案（1-5 分）']
    rows = [
        ['1', '创新是灵感事件，靠的是"天才员工"', '1 不同意 ← → 5 完全同意'],
        ['2', '创新失败的员工应该被追责', '1 不同意 ← → 5 完全同意'],
        ['3', '我们团队里员工敢提不成熟的想法', '1 不同意 ← → 5 完全同意'],
        ['4', '我对员工提的不成熟点子能给出具体回应', '1 不同意 ← → 5 完全同意'],
        ['5', '我们团队有固定的复盘节奏', '1 不同意 ← → 5 完全同意'],
        ['6', '我们的创新资源不会因为业务忙被挪用', '1 不同意 ← → 5 完全同意'],
        ['7', '我们团队正式分享过失败的创新案例', '1 不同意 ← → 5 完全同意'],
        ['8', '我最近一次承认"我错了"是在过去 30 天内', '1 不同意 ← → 5 完全同意'],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[1.0, 8.5, 6.0])
    add_para(doc, '分数越低的题目 = 你越需要补的位置；分数越高的题目 = 你已经具备的能力', size=10, italic=True, color=GRAY_DARK, space_before=8)
    add_callout(doc, '💡 不要现在就分析分数。今天结束时，我们会重做这 8 道题。前后对比，就是你今天最大的收获。')
    add_page_break(doc)

    # === 第 3 页：今天我最重要的 1 个期待 ===
    add_h1(doc, '第 3 页：今天我最重要的 1 个期待')
    add_para(doc, '请用一句话写出，你今天最想带走的 1 个东西：', size=11, space_after=12)
    for _ in range(3):
        add_signature_line(doc, '')
    add_callout(doc, '📌 这一页会贯穿全天。如果你中途觉得"这件事我已经得到了"，打个 ✅。如果没有，继续听课。')
    add_page_break(doc)

    # === 第 4 页：模块一学习目标 ===
    add_h1(doc, '第 4 页：模块一学习目标 · 创新与领导者')
    add_para(doc, '完成这个模块后，你将能够：', size=12, bold=True, color=DEEP_BLUE, space_after=8)
    add_para(doc, '1. 识别自己组织目前卡在哪个能力位置', size=11, indent=0.5)
    add_para(doc, '2. 区分"发明"与"创新"', size=11, indent=0.5)
    add_para(doc, '3. 建立"领导者 = 土壤"的核心立场', size=11, indent=0.5)
    add_page_break(doc)

    # === 第 5 页：核心概念卡 1 - 创新 ≠ 发明 ===
    add_h1(doc, '第 5 页：核心概念卡 1')
    add_card_box_v2(doc, '创新 ≠ 发明', [
        '发明 = 想出一个新点子',
        '创新 = 把一个新点子变成创造价值的现实',
        '',
        '发明可以是 1 个人的事；创新必须是 1 个组织的事',
        '发明可以是灵感的、偶然的；创新必须是过程的、可复盘的',
    ])
    add_callout(doc, '💡 判断标准：如果你团队里有人想出了好点子，但没人能把它变成现实——你团队有"发明"，但没有"创新"。')
    add_page_break(doc)

    # === 第 6 页：核心概念卡 2 - 领导者 = 土壤 ===
    add_h1(doc, '第 6 页：核心概念卡 2')
    add_card_box_v2(doc, '领导者 = 土壤（不是种子、不是船长、不是指挥家）', [
        '❌ 传统认知：领导者是"船长"，要指方向',
        '✅ 课程立场：领导者是"土壤"，要提供条件',
        '',
        '船长的角色：定方向、做决策、发号施令',
        '土壤的角色：提供心理安全、响应机制、容错文化',
        '',
        '船长 = 让员工按你的方向走',
        '土壤 = 让员工自己生长出方向',
    ])
    add_callout(doc, '💡 你的诊断：今天回到岗位，用 5 分钟写下你最近一次会议上，你说的最多的一句话是什么。这句话是"船长话"还是"土壤话"？')
    add_page_break(doc)

    # === 第 7 页：核心概念卡 3 - 土壤三要素 ===
    add_h1(doc, '第 7 页：核心概念卡 3')
    add_card_box_v2(doc, '土壤三要素（创新土壤）', [
        '1. 心理安全',
        '   员工敢不敢说出一个不成熟的想法？',
        '   ❌ 缺失表现：开会沉默 / 提建议"先看看别人"',
        '   ✅ 改进动作：每周固定 1 次"好点子分享会"',
        '',
        '2. 响应机制',
        '   员工说出来之后有没有人接住？',
        '   ❌ 缺失表现：提完石沉大海 / 等待数周无回应',
        '   ✅ 改进动作：24 小时内必须有具体回应',
        '',
        '3. 容错文化',
        '   想法尝试失败后会怎样？',
        '   ❌ 缺失表现：失败被追责 / 不敢尝试',
        '   ✅ 改进动作：月会有"这次没成，但我们学到了"分享',
    ])
    add_page_break(doc)

    # === 第 8 页：新视窗服务公司案例 ===
    add_h1(doc, '第 8 页：小组研讨 · 新视窗服务公司案例')
    add_callout(doc, '📌 研讨说明：4-6 人一组，10 分钟研讨，5 分钟汇报。', italic=True)
    add_h3(doc, '案例背景')
    add_para(doc, '售后维修服务公司一线维修技师小陈，发现保修条款描述与实际服务流程不一致，提出"工单系统增加保修条款快速对照表"建议。', size=11, space_after=4)
    add_para(doc, '主管表示"不错，但需总部 IT 排期三个月"。', size=11, space_after=4)
    add_para(doc, '三个月后小陈调走，想法消失。', size=11, space_after=4)
    add_para(doc, '半年后外部顾问重新梳理出同样问题。', size=11, space_after=12)
    add_h3(doc, '研讨问题')
    add_para(doc, '1. 这个案例里，土壤三要素各缺失了哪一个？', size=11, space_after=4)
    add_para(doc, '2. 小陈之后还有员工会提建议吗？', size=11, space_after=4)
    add_para(doc, '3. 如果你是小陈的主管，你会怎么回应？', size=11, space_after=12)
    add_h3(doc, '我的小组讨论记录')
    add_para(doc, '心理安全：', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '响应机制：', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '容错文化：', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_page_break(doc)

    # === 第 9 页：模块二学习目标 ===
    add_h1(doc, '第 9 页：模块二学习目标 · 创新是怎样的过程')
    add_para(doc, '完成这个模块后，你将能够：', size=12, bold=True, color=DEEP_BLUE, space_after=8)
    add_para(doc, '1. 描述创新过程的 5 个阶段', size=11, indent=0.5)
    add_para(doc, '2. 识别每个阶段最常见的卡点', size=11, indent=0.5)
    add_para(doc, '3. 列出每个卡点的应对方法', size=11, indent=0.5)
    add_page_break(doc)

    # === 第 10 页：核心概念卡 4 - 创新过程五阶段 ===
    add_h1(doc, '第 10 页：核心概念卡 4')
    add_card_box_v2(doc, '创新过程五阶段  发现 → 构思 → 验证 → 落地 → 固化', [
        '01 发现',
        '   "哪里不对劲？哪里可以更好？"',
        '   ⚠️ 最容易被忽视：员工埋头干活，没时间看外面',
        '',
        '02 构思',
        '   "针对这个问题，怎么解决？"',
        '   ⚠️ 最容易停在抽象：想到了但没具体化',
        '',
        '03 验证',
        '   "我们能不能小范围试一下？"',
        '   ⚠️ 最易卡住：不敢试 / 不会试',
        '',
        '04 落地',
        '   "怎么让大家都用这种方式？"',
        '   ⚠️ 易卡跨部门协调',
        '',
        '05 固化',
        '   "换个领导，这个方式还能持续吗？"',
        '   ⚠️ 最易被遗漏：做完不写 SOP / 不培训',
    ])
    add_callout(doc, '💡 核心立场：五阶段里最容易卡的是验证，最容易被遗漏的是固化。这两个阶段必须由领导者亲自推。')
    add_page_break(doc)

    # === 第 11 页：模块三学习目标 ===
    add_h1(doc, '第 11 页：模块三学习目标 · 领导者在创新过程中的作用')
    add_para(doc, '完成这个模块后，你将能够：', size=12, bold=True, color=DEEP_BLUE, space_after=8)
    add_para(doc, '1. 演练 IL-1 寻求新的点子的 3 个动作', size=11, indent=0.5)
    add_para(doc, '2. 演练 IL-2 领导创新的过程的 3 个动作', size=11, indent=0.5)
    add_para(doc, '3. 对照自己团队的现状，找到最该补的 1 个动作', size=11, indent=0.5)
    add_page_break(doc)

    # === 第 12 页：能力阶梯 ===
    add_h1(doc, '第 12 页：核心概念卡 5')
    add_card_box_v2(doc, '领导者管理创新的能力阶梯', [
        'IL-6  创新日常化、创新本能化     ← 目标态',
        '   ↑',
        'IL-5  培养设计文化               ← 上三阶之第 3 阶',
        '   ↑',
        'IL-4  维护创新纪律               ← 上三阶之第 2 阶',
        '   ↑',
        'IL-3  给员工赋能新的工作方式      ← 上三阶之第 1 阶',
        '   ↑',
        'IL-2  领导创新的过程             ← 下两阶之第 2 阶',
        '   ↑',
        'IL-1  寻求新的点子               ← 下两阶之第 1 阶',
        '   ↑',
        'IL-0  为点子而创新 / 因指令而创新 ← 起点',
    ])
    add_page_break(doc)

    # === 第 13 页：IL-1 三个动作 ===
    add_h1(doc, '第 13 页：核心概念卡 6')
    add_card_box_v2(doc, 'IL-1 寻求新的点子（3 个动作）', [
        '动作 1：主动开辟点子来源',
        '   └─ 你最近一次亲自听"用户/客户/一线员工"声音是何时？',
        '',
        '动作 2：降低提点子门槛',
        '   └─ 员工提一个点子到被看见，需要多长时间？',
        '',
        '动作 3：对不成熟点子给出回应',
        '   └─ 用公式：[肯定角度] + [具体疑问] + [下一步动作]',
    ])
    add_h3(doc, '回应公式（填写示例）')
    add_para(doc, '原话（员工的点子）：', size=10, bold=True, color=GRAY_DARK, space_after=2)
    add_para(doc, '"我觉得我们客户回访的流程太长了，能不能让客户用微信直接发语音回访？"', size=10, italic=True, color=GRAY_DARK, space_after=8)
    add_para(doc, 'IL-1 状态组的回应（按公式）：', size=10, bold=True, color=VERMILION, space_after=2)
    add_para(doc, '✅ "用语音代替文字，这个角度有意思。"', size=10, color=GRAY_DARK, indent=0.5, space_after=2)
    add_para(doc, '✅ "我担心的是 60 岁以上客户可能不习惯微信语音。"', size=10, color=GRAY_DARK, indent=0.5, space_after=2)
    add_para(doc, '✅ "咱们这样，下周找 3 个老客户先试一下，看接受度。"', size=10, color=GRAY_DARK, indent=0.5, space_after=2)
    add_para(doc, '✅ "下周三之前给我一个反馈。"', size=10, color=GRAY_DARK, indent=0.5, space_after=8)
    add_para(doc, '起点状态组的回应（反例）：', size=10, bold=True, color=VERMILION, space_after=2)
    add_para(doc, '❌ "嗯，我看看。"', size=10, color=GRAY_DARK, indent=0.5, space_after=2)
    add_para(doc, '❌ "这个想法挺有意思，不过……（10 个但是）"', size=10, color=GRAY_DARK, indent=0.5, space_after=2)
    add_para(doc, '❌ "这个我们以前试过，不行。"', size=10, color=GRAY_DARK, indent=0.5)
    add_page_break(doc)

    # === 第 14 页：情境演练 A 记录 ===
    add_h1(doc, '第 14 页：情境演练 A 记录')
    add_callout(doc, '📌 演练主题：起点状态组 vs IL-1 状态组的回应对比', italic=True)
    add_h3(doc, '演练情境')
    add_para(doc, '一位团队成员在周会提出："我觉得我们客户回访的流程太长了，能不能让客户用微信直接发语音回访？"', size=11, space_after=8)
    add_para(doc, '我的角色：A 组（起点状态组）/ B 组（IL-1 状态组）', size=11, bold=True, space_after=8)
    add_h3(doc, '演练后我的观察')
    add_para(doc, 'A 组的回应给我的感受：', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, 'B 组的回应给我的感受：', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '我更愿意在哪位领导者手下工作？为什么？', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '我平时的回应更像 A 组还是 B 组？', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_page_break(doc)

    # === 第 15 页：IL-2 三个动作 ===
    add_h1(doc, '第 15 页：核心概念卡 7')
    add_card_box_v2(doc, 'IL-2 领导创新的过程（3 个动作）', [
        '动作 4：为五阶段建立可见节点',
        '   └─ 你团队的创新项目有没有可视化看板？',
        '',
        '动作 5：明确每个阶段的卡点和应对',
        '   └─ 最易卡：验证（资源/方法）；最易漏：固化（SOP/培训）',
        '',
        '动作 6：复盘成为日常而非事件',
        '   └─ 每周小复盘 15 分钟，每月项目群复盘，结束完整复盘',
    ])
    add_page_break(doc)

    # === 第 16 页：五阶段卡点识别练习 ===
    add_h1(doc, '第 16 页：五阶段卡点识别练习')
    add_callout(doc, '📌 情境：某团队启动了一个"客户回访流程优化"项目，已经 3 个月了。现状如下：', italic=True)
    add_para(doc, '• 项目启动时领导说"你们去调研一下"', size=11, space_after=2)
    add_para(doc, '• 团队做了一份 30 页调研报告，但没人读过第二遍', size=11, space_after=2)
    add_para(doc, '• 团队想了 10 个"可能可行"的方案，但都说"先等等再决定"', size=11, space_after=2)
    add_para(doc, '• 没有人去客户那试过任何一个方案', size=11, space_after=2)
    add_para(doc, '• 3 个月过去了，没人说得清"这项目现在在哪个阶段"', size=11, space_after=12)
    add_h3(doc, '我的分析')
    add_para(doc, '这个项目卡在了哪个阶段？', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '这个阶段最常见的卡点是什么？', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_para(doc, '如果我是这位领导者，第 1 周我会做什么？', size=11, bold=True, space_after=4)
    add_signature_line(doc, '')
    add_page_break(doc)

    # === 第 17 页：模块四学习目标 ===
    add_h1(doc, '第 17 页：模块四学习目标 · 善于创新的公司，其领导者有什么不同')
    add_para(doc, '完成这个模块后，你将能够：', size=12, bold=True, color=DEEP_BLUE, space_after=8)
    add_para(doc, '1. 对照灯塔家电集团的 9 个动作，给自己的组织打分', size=11, indent=0.5)
    add_para(doc, '2. 识别自己组织最需要补的 3 个动作', size=11, indent=0.5)
    add_para(doc, '3. 制定 90 天改造路径图', size=11, indent=0.5)
    add_page_break(doc)

    # === 第 18 页：灯塔起点 vs 今天 ===
    add_h1(doc, '第 18 页：灯塔家电集团 · 起点 vs 今天')
    headers = ['维度', '2014 起点', '2022 今天']
    rows = [
        ['年营收', '80 亿', '180 亿'],
        ['员工', '12000', '15000'],
        ['创新 = 研发部占比', '100%', '18%'],
        ['年专利申报', '200 件', '450 件'],
        ['专利转化率', '12%', '38%'],
        ['部门自主微创新占比', '0%', '67%'],
        ['高频词', '"这不是我们部门的事"', '"这事儿我们部门能试一下"'],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[5.0, 5.0, 5.0])
    add_page_break(doc)

    # === 第 19 页：灯塔 8 年拆解 ===
    add_h1(doc, '第 19 页：灯塔家电集团 · 8 年拆解')
    headers = ['年份', '关键变化', '当前位置', '占比']
    rows = [
        ['2014', '起点：研发部独家创新', 'IL-0', '—'],
        ['2015', '方法显性化 + 决策边界下放', 'IL-3 起步', '25%'],
        ['2016', '月度复盘 + 受保护资源', 'IL-3 完成', '50%'],
        ['2017', '组合视角 + 让失败被看见', 'IL-4 起步', '60%'],
        ['2018', '领导者示范 + 反方意见机制', 'IL-4 完成', '70%'],
        ['2019-2020', '把所有动作编入 SOP/JD', 'IL-5 完成', '85%'],
        ['2021-2022', '不再单独讲"创新"', 'IL-6 起点', '95%'],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[2.5, 6.0, 4.0, 2.5])
    add_page_break(doc)

    # === 第 20 页：IL-3 三个动作 ===
    add_h1(doc, '第 20 页：核心概念卡 8')
    add_card_box_v2(doc, 'IL-3 给员工赋能新的工作方式（3 个动作）', [
        '动作 1：方法显性化',
        '   └─ 你团队最值钱的工作方法有没有被写下来？',
        '   └─ 发明人离职，下一个人能不能照着做？',
        '',
        '动作 2：重新划定决策边界',
        '   └─ 员工想试一个新方法，需要经过几道审批？',
        '   └─ 超过 3 道审批，创新基本不会发生',
        '',
        '动作 3：让方法成为共同语言',
        '   └─ 你团队描述"创新"，有几个不同的词？',
    ])
    add_page_break(doc)

    # === 第 21 页：IL-4 三个动作 ===
    add_h1(doc, '第 21 页：核心概念卡 9')
    add_card_box_v2(doc, 'IL-4 维护创新纪律（3 个动作）', [
        '动作 4：建立固定复盘节奏',
        '   └─ 复盘是约定俗成的习惯，还是每次重新组织？',
        '',
        '动作 5：预留受保护资源',
        '   └─ 创新资源会不会被业务挤压？',
        '',
        '动作 6：用组合视角看项目',
        '   └─ 你团队同时跑几个创新小项目？失败怎么算？',
    ])
    add_page_break(doc)

    # === 第 22 页：IL-5 三个动作 ===
    add_h1(doc, '第 22 页：核心概念卡 10')
    add_card_box_v2(doc, 'IL-5 培养设计文化（3 个动作）', [
        '动作 7：在日常决策中示范',
        '   └─ 最近一次会议前 10 分钟讲的是什么？',
        '   └─ 用户声音 vs 上级要求？',
        '',
        '动作 8：让尝试和学习被看见',
        '   └─ 你团队有没有正式分享过失败的创新？',
        '   └─ 你奖励的是成功还是学习？',
        '',
        '动作 9：领导者保持被影响的开放度',
        '   └─ 你最近一次说"我错了"是什么时候？',
    ])
    add_page_break(doc)

    # === 第 23 页：灯塔 9 个动作对照表（研讨用） ===
    add_h1(doc, '第 23 页：灯塔 9 个动作对照表（研讨用）')
    add_callout(doc, '📌 填表说明：4-6 人一组研讨，用 0-3 分给每个动作打分。', italic=True)
    headers = ['灯塔动作', '我的组织现状', '差距(0-3)', '90 天能做什么']
    rows = [
        ['方法显性化', '核心方法是否被写下来？', '', ''],
        ['决策边界', '5 万以下小创新谁能决策？', '', ''],
        ['共同语言', '描述"创新"有几个词？', '', ''],
        ['固定复盘', '多久复盘一次？谁负责？', '', ''],
        ['受保护资源', '创新资源会不会被挤压？', '', ''],
        ['组合视角', '同时跑几个项目？失败怎么算？', '', ''],
        ['日常示范', '会议前 10 分钟讲什么？', '', ''],
        ['让尝试被看见', '有没有正式分享过失败？', '', ''],
        ['保持被影响', '最近一次提"我不同意"是何时？', '', ''],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[3.0, 6.0, 2.5, 4.0])
    add_page_break(doc)

    # === 第 24 页：90 天改造路径图 ===
    add_h1(doc, '第 24 页：我的 90 天改造路径图')
    add_card_box_v2(doc, '我的 90 天改造路径图', [
        '第一阶段（第 1-30 天）：从 0 到 1',
        '改造动作 1：',
        '负责人：',
        '第 1 步：',
        '第 30 天标志：',
        '',
        '第二阶段（第 31-60 天）：从 1 到 3',
        '改造动作 2：',
        '负责人：',
        '第 31 天起点：',
        '第 60 天标志：',
        '',
        '第三阶段（第 61-90 天）：从 3 到 5',
        '改造动作 3：',
        '负责人：',
        '第 61 天起点：',
        '第 90 天标志：',
    ])
    add_page_break(doc)

    # === 第 25 页：模块三/四 课后自评 ===
    add_h1(doc, '第 25 页：模块三/四 课后自评')
    add_callout(doc, '📌 填写说明：5 分钟独立完成。', italic=True)
    add_para(doc, '在下两阶（IL-1/IL-2）6 个动作中：', size=11, bold=True, space_after=4)
    add_para(doc, '我目前做得最好的是动作 # ______', size=11, space_after=4)
    add_signature_line(doc, '理由')
    add_para(doc, '我最需要补的是动作 # ______', size=11, space_after=4)
    add_signature_line(doc, '理由')
    add_para(doc, '在上三阶（IL-3/IL-4/IL-5）9 个动作中：', size=11, bold=True, space_after=4)
    add_para(doc, '我最有感觉的是动作 # ______', size=11, space_after=4)
    add_signature_line(doc, '理由')
    add_para(doc, '我最担心做不到的是动作 # ______', size=11, space_after=4)
    add_signature_line(doc, '理由')
    add_page_break(doc)

    # === 第 26 页：上三阶 / 下两阶 关系图 ===
    add_h1(doc, '第 26 页：上三阶 / 下两阶 关系图')
    add_card_box_v2(doc, '上三阶 vs 下两阶 · 关系图', [
        '下两阶（IL-1 / IL-2）解决"效率问题"',
        '   └─ 点子进得来、出得去',
        '   └─ 周期：1-3 个月见效',
        '   └─ 关键：领导者的"动作"',
        '',
        '                 ↓',
        '',
        '上三阶（IL-3 / IL-4 / IL-5）解决"能力问题"',
        '   └─ 整个组织具备持续创新的能力',
        '   └─ 周期：1-3 年见效',
        '   └─ 关键：领导者的"示范"',
    ])
    add_callout(doc, '💡 金句：下两阶是"教员工做"，上三阶是"让员工自己会做"。前者是命令链，后者是文化场。')
    add_page_break(doc)

    # === 第 27 页：模块五学习目标 ===
    add_h1(doc, '第 27 页：模块五学习目标 · 让创新之河永不枯竭')
    add_para(doc, '完成这个模块后，你将能够：', size=12, bold=True, color=DEEP_BLUE, space_after=8)
    add_para(doc, '1. 描述创新日常化的 3 个标志', size=11, indent=0.5)
    add_para(doc, '2. 列出创新本能化的 4 个条件', size=11, indent=0.5)
    add_para(doc, '3. 对照自己组织，识别"距离 IL-6 还差多远"', size=11, indent=0.5)
    add_page_break(doc)

    # === 第 28 页：核心概念卡 11 - 创新日常化 ===
    add_h1(doc, '第 28 页：核心概念卡 11')
    add_card_box_v2(doc, '创新日常化的 3 个标志', [
        '标志 1：创新不再被"特别提起"',
        '   └─ IL-5：大会单独讲创新、月报单列创新',
        '   └─ IL-6：创新进展自然包含在业务进展中',
        '',
        '标志 2：改进是每个人的默认动作',
        '   └─ IL-5：改进需要被"要求"',
        '   └─ IL-6：改进不需要提醒，是工作的一部分',
        '',
        '标志 3：失败不被放大，成功不被神话',
        '   └─ IL-5：失败案例有辩解感、成功案例有英雄感',
        '   └─ IL-6：失败和成功用同一种方式被看见',
    ])
    add_callout(doc, '💡 金句：日常化的本质是"创新不再是一个 special 项目，而是 normal 工作"。当 special 变成 normal，IL-6 到了。')
    add_page_break(doc)

    # === 第 29 页：核心概念卡 12 - 创新本能化 ===
    add_h1(doc, '第 29 页：核心概念卡 12')
    add_card_box_v2(doc, '创新本能化的 4 个条件', [
        '条件 1：方法已经变成肌肉记忆',
        '   └─ 解决"会不会"',
        '',
        '条件 2：时间已经被自动预留',
        '   └─ 解决"有没有空"',
        '',
        '条件 3：失败已经被组织"消化"',
        '   └─ 解决"敢不敢试"',
        '',
        '条件 4：领导者已经把"被影响"作为默认姿态',
        '   └─ 解决"上同不同意"',
    ])
    add_callout(doc, '💡 金句：本能化不是"4 个条件都做到"，是"4 个条件都做到并且没人再提它"。')
    add_page_break(doc)

    # === 第 30 页：阶梯自评表（课后复评） ===
    add_h1(doc, '第 30 页：阶梯自评表（课后复评）')
    add_callout(doc, '📌 请用 30 秒/题回答——第一直觉最重要。', italic=True)
    headers = ['#', '问题', '你的答案（1-5 分）']
    rows = [
        ['1', '创新是灵感事件，靠的是"天才员工"', '1 不同意 ← → 5 完全同意'],
        ['2', '创新失败的员工应该被追责', '1 不同意 ← → 5 完全同意'],
        ['3', '我们团队里员工敢提不成熟的想法', '1 不同意 ← → 5 完全同意'],
        ['4', '我对员工提的不成熟点子能给出具体回应', '1 不同意 ← → 5 完全同意'],
        ['5', '我们团队有固定的复盘节奏', '1 不同意 ← → 5 完全同意'],
        ['6', '我们的创新资源不会因为业务忙被挪用', '1 不同意 ← → 5 完全同意'],
        ['7', '我们团队正式分享过失败的创新案例', '1 不同意 ← → 5 完全同意'],
        ['8', '我最近一次承认"我错了"是在过去 30 天内', '1 不同意 ← → 5 完全同意'],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[1.0, 8.5, 6.0])
    add_h3(doc, '前后对比')
    headers = ['维度', '开场自评', '课程结束自评', '变化']
    rows = [
        ['1. 创新认知', '', '', ''],
        ['2. 失败态度', '', '', ''],
        ['3. 心理安全', '', '', ''],
        ['4. 回应方式', '', '', ''],
        ['5. 复盘节奏', '', '', ''],
        ['6. 资源保护', '', '', ''],
        ['7. 失败分享', '', '', ''],
        ['8. 承认错误', '', '', ''],
    ]
    add_table_styled(doc, headers, rows, col_widths_cm=[3.5, 4.0, 4.0, 3.5])
    add_callout(doc, '💡 关键引导：分数变了的维度，就是你今天最大的收获。没变的维度，正是你下周一的第一件事。')
    add_page_break(doc)

    # === 第 31 页：个人行动承诺书 ===
    add_h1(doc, '第 31 页：个人行动承诺书')
    add_card_box_v2(doc, '个人行动承诺书 · 《打造组织创新力：营造创新土壤》', [
        '姓名：______________  公司/部门：______________',
        '日期：______________  课程版本：v1.0',
        '',
        '─── 第 1 周（7 天内开始） ───',
        '我要做的事（用动词开头）：',
        '我为什么要做这件事（与今天的哪个学习相关）：',
        '完成后我能看到的具体标志：',
        '',
        '─── 第 1 个月（30 天内完成） ───',
        '我要做的事：',
        '我为什么要做这件事：',
        '完成后我能看到的具体标志：',
        '',
        '─── 第 3 个月（90 天内完成） ───',
        '我要做的事：',
        '我为什么要做这件事：',
        '完成后我能看到的具体标志：',
        '',
        '─── 我的承诺 ───',
        '我承诺自己，在上述 3 个时间节点内完成这 3 件事。',
        '如果做不到，我会诚实面对，并在小组群说明原因。',
        '签名：______________   日期：______________',
    ])
    add_callout(doc, '📌 写承诺前，先想 30 秒：你今天最触动你的一句话是什么？那个触动，就对应着你最需要做的那件事。', italic=True)
    add_para(doc, '3 个筛选条件（每件事都要满足）：', size=11, bold=True, space_after=4)
    add_para(doc, '1. 这件事今天下班前就能开始吗？', size=10, indent=0.5)
    add_para(doc, '2. 这件事不依赖其他人吗？', size=10, indent=0.5)
    add_para(doc, '3. 这件事规定时间内能完成吗？', size=10, indent=0.5)
    add_para(doc, '用动词开头：', size=11, bold=True, space_before=8)
    add_para(doc, '❌ "我会更关注员工的反馈"', size=10, color=VERMILION)
    add_para(doc, '✅ "每周五下午 5 点，我会给团队做 15 分钟\'用户声音分享\'"', size=10, color=DEEP_BLUE)
    add_page_break(doc)

    # === 第 32 页：小组互评 + 跟进机制 ===
    add_h1(doc, '第 32 页：小组互评 + 课后 30/60/90 天跟进机制')
    add_h3(doc, '小组互评记录')
    add_callout(doc, '📌 4-6 人一组，每组围成圈', italic=True)
    add_para(doc, '互评规则：', size=11, bold=True, space_after=4)
    add_para(doc, '✅ "我觉得你这件事的盲点是 X"', size=10, color=DEEP_BLUE)
    add_para(doc, '❌ "你这件事做得很好"（这是表扬，不是反馈）', size=10, color=VERMILION)
    add_signature_line(doc, '我的承诺')
    add_signature_line(doc, '组员 1 的反馈（盲点）')
    add_signature_line(doc, '组员 2 的反馈（盲点）')
    add_signature_line(doc, '组员 3 的反馈（盲点）')
    add_signature_line(doc, '我修改后的承诺')

    add_h3(doc, '课后 30/60/90 天跟进机制卡')
    add_card_box_v2(doc, '课后 30/60/90 天跟进机制卡', [
        '第 30 天：复盘会议',
        '   └─ 形式：微信群语音会议，30 分钟',
        '   └─ 议程：每位学员 2 分钟汇报第 1 周+第 1 月承诺进展',
        '   └─ 目标：让承诺"被看见"——没被看见的承诺最容易被放弃',
        '',
        '第 60 天：小组互访',
        '   └─ 形式：同城学员线下互访 1 次（不同公司/部门）',
        '   └─ 议程：上午参观对方组织 + 下午 90 分钟对照研讨',
        '   └─ 目标：让"创新土壤"被对照、被拓展',
        '',
        '第 90 天：返校日',
        '   └─ 形式：线下 0.5 天返校',
        '   └─ 议程：90 分钟案例复盘 + 90 分钟新内容',
        '   └─ 目标：让"创新日常化"有节奏、有反馈、有进阶',
    ])

    add_h3(doc, '推荐阅读')
    add_para(doc, '《创新者的窘境》 克里斯滕森', size=11, indent=0.5)
    add_para(doc, '《创新与企业家精神》 德鲁克', size=11, indent=0.5)
    add_para(doc, '《组织能力的杨三角》 杨国安', size=11, indent=0.5, space_after=12)

    add_h3(doc, '离场前 3 件事')
    add_para(doc, '☐ 把承诺书签名后交给讲师（讲师扫描存档，课后 PDF 反馈给学员本人）', size=11)
    add_para(doc, '☐ 把工具卡 8 张夹进学员手册', size=11)
    add_para(doc, '☐ 找到至少 1 位组员交换微信（建立 30/60/90 天同学群）', size=11, space_after=12)

    add_callout(doc, '"拉开序章，不是写结局。创新这件事，没有结局，只有一个又一个的开始。今天，是你的开始。"', italic=False, bold=True)

    add_para(doc, '完整电子版：04_学员手册/学员手册_HTML可视化版.html', size=10, italic=True, color=GRAY_DARK, space_before=12)
    add_para(doc, '工具表单电子版：08_全流程工具表单/', size=10, italic=True, color=GRAY_DARK)
    add_para(doc, '推荐阅读与延伸：11_数字工具包/、13_管理者工具包/', size=10, italic=True, color=GRAY_DARK)
    add_para(doc, 'v1.0 / 2026-06-16 / 主讲 罗宏伟 / 竞越', size=9, italic=True, color=GRAY_LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

    doc.save(OUT_DOCX)
    print(f'OK: {OUT_DOCX}')


if __name__ == '__main__':
    build_docx()
    print('---WORD DOCX DONE---')

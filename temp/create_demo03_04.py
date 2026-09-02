# -*- coding: utf-8 -*-
"""
生成 demo03-隐藏考题诊断表示例.docx
demo04-类比模型识别清单示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_demo03():
    """demo03: 隐藏考题诊断表示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("隐藏考题诊断表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 识别表面问题背后的真实考题，避免答非所问")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    # 学员信息
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：陈战略", "团队：战略发展组", "日期：2026-07-22", "教练：张破局"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # 核心概念
    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】隐藏考题 = 对方话语中隐含的、真正希望你回答的问题")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    # 案例
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、典型案例分析")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 案例场景
    scene = doc.add_paragraph()
    run = scene.add_run("【场景】总裁在季度会上说：")
    run.font.size = Pt(11)
    run.font.bold = True

    quote = doc.add_paragraph()
    run = quote.add_run('"这个季度的市场份额下降了2个百分点，你们怎么看？"')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # 诊断表格
    diag_table = doc.add_table(rows=5, cols=2)
    diag_table.style = 'Table Grid'

    headers = ["诊断维度", "内容"]
    for j, h in enumerate(headers):
        diag_table.rows[0].cells[j].text = h
        para = diag_table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        tc = diag_table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    diag_data = [
        ["表面问题\n（总裁说的）", "市场份额下降2个百分点"],
        ["可能的隐藏考题\n（总裁真正想问的）", "1. 谁该为此负责？\n2. 你们有什么对策？\n3. 需要我做什么决定？\n4. 下季度能回来吗？"],
        ["破题关键", "总裁不是在问市场份额的数字，而是在问：\n- 团队有没有系统性分析？\n- 有没有应对方案？\n- 需不需要资源支持？"],
        ["学员回答示例", '"张总，我们分析了三方面原因：竞品新品上市、渠道铺货延迟、团队人力紧张。我们计划下月推出X方案，预计能夺回1.5个百分点。需要您批准招聘2名资深销售。"']
    ]

    for i, (label, content) in enumerate(diag_data):
        diag_table.rows[i+1].cells[0].text = label
        diag_table.rows[i+1].cells[1].text = content
        for j, cell in enumerate(diag_table.rows[i+1].cells):
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
                if j == 0:
                    para.runs[0].font.bold = True

    doc.add_paragraph()

    # 教练批注
    coach = doc.add_paragraph()
    run = coach.add_run("[教练批注：总裁的问题从来不只是问题本身。当你听到'市场份额下降'，要立刻意识到：他在考验你的系统思维和应对能力。这是一道综合题，不是计算题。]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    doc.add_paragraph()

    # 练习
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、隐藏考题识别练习")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    practice_items = [
        ('"这个项目为什么失败了？"', "表面：问原因\n隐藏考题：1.谁的责任？2.下次怎么避免？3.你从中学到了什么？"),
        ('"你觉得这个方案可行吗？"', "表面：问意见\n隐藏考题：1.你在说服我还是敷衍我？2.你有勇气坚持吗？3.你能承担后果吗？"),
        ('"你知道团队有人反映你什么吗？"', "表面：传话\n隐藏考题：1.我要敲打你 2.给你解释机会 3.看你会怎么反应")
    ]

    for i, (question, analysis) in enumerate(practice_items):
        p = doc.add_paragraph()
        run = p.add_run(f"练习{i+1}：")
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(question)
        run.font.size = Pt(11)
        run.font.italic = True

        a = doc.add_paragraph()
        run = a.add_run(f"   分析：{analysis}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo03-隐藏考题诊断表示例.docx"
    doc.save(output_path)
    print(f"demo03已生成: {output_path}")
    return output_path


def create_demo04():
    """demo04: 类比模型识别清单示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("类比模型识别清单示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 用经典类比模型快速定位问题本质")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    # 学员信息
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：王子明", "团队：组织变革组", "日期：2026-07-25", "教练：李破局"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # 概念说明
    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】类比模型是前人总结的经典问题原型，帮助你快速对号入座")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    # 模型清单表格
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、类比模型识别清单")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    models = [
        ["冰山模型", "表面现象 vs 深层原因", "员工流失 → 表面：钱少；深层：成长受限",
         "看到问题A时，问自己：A的B面是什么？"],
        ["GPS模型", "Goal-Pattern-Startup", "市场进入 → 先定目标(G)，再找规律(P)，最后起步(S)",
         "面对新领域时，按G→P→S的顺序推进"],
        ["乐队模型", "不同乐器协同演奏", "跨部门协作 → 谁是指挥？各自什么角色？",
         "组织问题时，问：指挥在哪儿？"],
        ["倒酒模型", "杯子满了要换瓶", "产品升级 → 旧瓶装新酒，还是新瓶装旧酒？",
         "变革时问：形式变了，内容变了吗？"],
        ["种子模型", "播种-发芽-开花-结果", "人才培养 → 什么时候播？什么时候收？",
         "长期项目问：周期各阶段是什么？"],
        ["拼图模型", "碎片完整后才清晰", "战略制定 → 每个碎片是什么？缺哪个？",
         "复杂问题问：完整图景是什么？"]
    ]

    model_table = doc.add_table(rows=len(models)+1, cols=4)
    model_table.style = 'Table Grid'

    headers = ["模型名称", "核心隐喻", "典型应用场景", "使用提示"]
    for j, h in enumerate(headers):
        model_table.rows[0].cells[j].text = h
        para = model_table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        tc = model_table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for i, row_data in enumerate(models):
        for j, text in enumerate(row_data):
            model_table.rows[i+1].cells[j].text = text
            for para in model_table.rows[i+1].cells[j].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # 案例应用
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、案例应用")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【问题】公司推行新系统，3个月了还是推行不下去，各部门抱怨连连")
    run.font.size = Pt(11)
    run.font.bold = True

    # 应用分析
    analysis_table = doc.add_table(rows=4, cols=2)
    analysis_table.style = 'Table Grid'

    analysis_data = [
        ["识别到的类比模型", "倒酒模型 + 乐队模型"],
        ["模型分析", "倒酒模型：新系统（新酒）用旧流程（旧瓶）装，自然装不下\n乐队模型：没有明确的指挥，各部门各奏各的调"],
        ["破题方向", "1. 明确：新系统的推行谁是总负责（指挥）\n2. 调整：先固化新流程，再推广新系统\n3. 激励：设立试点成功奖励机制"],
        ["学员感悟", '"用类比模型分析后，这个问题从"系统问题"变成了"组织+流程问题"，解决路径一下子清晰了"']
    ]

    for i, (label, content) in enumerate(analysis_data):
        analysis_table.rows[i].cells[0].text = label
        analysis_table.rows[i].cells[1].text = content
        analysis_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in analysis_table.rows[i].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo04-类比模型识别清单示例.docx"
    doc.save(output_path)
    print(f"demo04已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    create_demo03()
    create_demo04()

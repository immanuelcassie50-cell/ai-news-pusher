#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# Pre-class prep
pp = doc.add_paragraph()
r = pp.add_run('课前准备')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph()

tt = doc.add_paragraph()
r = tt.add_run('课前自测问卷'); r.bold = True; r.font.size = Pt(14)

for q in ['你过去做过的复盘，现在还能想起哪些具体内容？',
          '你曾经是否遇到过复盘做了很多但下次还是踩坑的情况？',
          '你认为复盘和决策卡最大的区别是什么？',
          '你是否曾经有过我就是觉得不对但说不清的经历？',
          '你觉得一个组织中最值得沉淀的是什么？']:
    doc.add_paragraph(q)

doc.add_paragraph()
th = doc.add_paragraph()
r = th.add_run('课前思考'); r.bold = True; r.font.size = Pt(14)

tcp = doc.add_paragraph()
tcp.add_run('你上一次做的复盘现在还记得什么？')
tcp.paragraph_format.left_indent = Cm(1)

note = doc.add_paragraph()
r = note.add_run('请在下方记录你的思考：')
r.italic = True; r.font.color.rgb = RGBColor(128,128,128)

for _ in range(6):
    line = doc.add_paragraph()
    line.add_run('_' * 70)
    line.paragraph_format.space_after = Pt(12)

doc.add_page_break()
print('S4: Pre-class prep done')
doc.save(r'D:\CC\temp\handbook_s4.docx')

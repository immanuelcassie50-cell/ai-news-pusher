# -*- coding: utf-8 -*-
"""
创建讲师手册_工业革命.docx
课程29：工业革命——为什么是英国率先起飞
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 输出路径
output_dir = r"D:/新课开发/经济学/29_工业革命"
output_path = os.path.join(output_dir, "讲师手册_工业革命.docx")

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {})
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    shading.append(shd)

def add_horizontal_line(doc):
    """添加水平线"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '666666')
    pBdr.append(bottom)
    return para

# 创建文档
doc = Document()

# 设置页面
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# 设置文档默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# ====== 封面 ======
title = doc.add_heading('', level=0)
run = title.add_run('工业革命')
run.font.size = Pt(44)
run.font.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——为什么是英国率先起飞')
run.font.size = Pt(28)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para.add_run('讲师手册')
run.font.size = Pt(22)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# 讲师信息表
info_table = doc.add_table(rows=4, cols=4)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

labels = ['讲师姓名', '所在机构', '课程日期', '课程时长']
for i, label in enumerate(labels):
    row = info_table.rows[i]
    cell = row.cells[0]
    cell.text = label
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(cell, 'E8E8E8')

doc.add_page_break()

# ====== 使用指南 ======
doc.add_heading('如何使用这本讲师手册', level=1)

guide_text = """这本讲师手册是课程的核心辅助工具，包含以下内容：

第一章：课程概览与设计思路——帮助讲师理解课程的整体定位和设计逻辑

第二章：每次课的授课要点——详细的知识点讲解和话术建议

第三章：互动设计说明——课堂活动的设计意图和操作指南

第四章：常见问题与应对——学员可能提出的问题和参考回答

第五章：案例讲解话术——核心案例的讲解要点和表达方式

第六章：时间把控建议——每个环节的推荐时长和节奏控制

第七章：评分标准与评估——学员表现的评估维度和建议"""

for para_text in guide_text.strip().split('\n\n'):
    doc.add_paragraph(para_text.strip())

add_horizontal_line(doc)

# 讲师须知
doc.add_heading('讲师须知', level=2)

notice_text = """1. 本课程面向管理者和内容创作者，强调历史比较方法的实际应用
2. 避免过于学术化的表述，用案例和故事来传达核心概念
3. 鼓励学员积极参与讨论，课程的核心价值在于思考框架的建立
4. 灵活调整案例：可根据学员的行业背景替换为中国企业的案例
5. 时间控制是本课程的难点，建议每个章节预留弹性时间"""

for para_text in notice_text.strip().split('\n'):
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(para_text.strip())

doc.add_page_break()

# ====== 第一章：课程概览 ======
doc.add_heading('第一章　课程概览与设计思路', level=1)

doc.add_heading('1.1 课程基本信息', level=2)

basic_info = [
    ('课程名称', '工业革命——为什么是英国率先起飞'),
    ('课程编号', '29'),
    ('课程时长', '3小时（180分钟）'),
    ('学员对象', '对长时段历史规律感兴趣的管理者和内容创作者'),
    ('前置要求', '无特殊要求，但建议学员对商业史或技术史有基础兴趣'),
    ('课程形式', '讲授+案例讨论+小组练习')
]

basic_table = doc.add_table(rows=len(basic_info)+1, cols=2)
basic_table.style = 'Table Grid'

basic_table.rows[0].cells[0].text = '项目'
basic_table.rows[0].cells[1].text = '内容'
for para in basic_table.rows[0].cells[0].paragraphs:
    for run in para.runs:
        run.bold = True
set_cell_shading(basic_table.rows[0].cells[0], '4472C4')
for para in basic_table.rows[0].cells[0].paragraphs:
    for run in para.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
set_cell_shading(basic_table.rows[0].cells[1], '4472C4')
for para in basic_table.rows[0].cells[1].paragraphs:
    for run in para.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

for i, (item, content) in enumerate(basic_info):
    row = basic_table.rows[i+1]
    row.cells[0].text = item
    row.cells[1].text = content
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(row.cells[0], 'F2F2F2')

doc.add_paragraph()

doc.add_heading('1.2 课程核心问题', level=2)

para = doc.add_paragraph()
run = para.add_run('核心问题：')
run.bold = True
run.font.size = Pt(14)
para.add_run('工业革命为什么在英国发生，而不是法国、德国或中国？')

para = doc.add_paragraph()
run = para.add_run('延伸问题：')
run.bold = True
para.add_run('这一历史现象对理解今天的技术革命（如AI浪潮）有什么启发？')

para = doc.add_paragraph()
run = para.add_run('课程产出：')
run.bold = True
para.add_run('一套"重大技术革命起飞条件清单"，可用于分析任何技术变革')

doc.add_heading('1.3 设计逻辑', level=2)

design_text = """本课程采用"问题溯源→要素分析→比较验证→框架提炼→类比应用"的五段式结构：

第一段（问题溯源）：从"工业革命=瓦特蒸汽机"这个常见误解切入，揭示问题的复杂性
第二段（要素分析）：引入六要素分析框架，系统拆解英国工业革命的条件
第三段（比较验证）：通过法国、中国、德国案例，验证六要素框架的解释力
第四段（框架提炼）：将六要素升华为"重大技术革命起飞条件清单"
第五段（类比应用）：用清单分析AI革命，实现从历史到现实的迁移学习"""

for para_text in design_text.strip().split('\n\n'):
    doc.add_paragraph(para_text.strip())

doc.add_heading('1.4 与学员手册的对应关系', level=2)

mapping_table = doc.add_table(rows=8, cols=2)
mapping_table.style = 'Table Grid'

mapping_data = [
    ('讲师手册章节', '对应学员手册章节'),
    ('第一章：课程概览', '引言、课程目标'),
    ('第二章：授课要点', '第一-四章'),
    ('第三章：互动设计', '练习与思考'),
    ('第四章：常见问题', '延伸阅读'),
    ('第五章：案例话术', '案例摘要'),
    ('第六章：时间控制', '附录：时间规划）'),
    ('第七章：评分标准', '附录：评估维度')
]

for i, (col1, col2) in enumerate(mapping_data):
    row = mapping_table.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    if i == 0:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(cell, '4472C4')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    else:
        for cell in row.cells:
            set_cell_shading(cell, 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_page_break()

# ====== 第二章：授课要点 ======
doc.add_heading('第二章　授课要点详解', level=1)

# 第一部分
doc.add_heading('第一部分：重新理解工业革命（建议时长：30分钟）', level=2)

doc.add_heading('1.1 常见误解（10分钟）', level=3)

points_11 = [
    '核心观点：工业革命不是单一技术发明，而是系统性转变',
    '误解一：蒸汽机发明=工业革命',
    '话术："瓦特确实伟大，但如果没有焦炭炼铁、运河运输、煤矿开采，蒸汽机可能只停留在纽卡斯尔的煤矿里抽水。为什么？"',
    '误解二：少数发明家的功劳',
    '话术："阿克莱特不是发明家，但他是现代工厂制度之父。工业革命需要的不只是技术，还需要将技术组织起来的方式。"',
    '误解三：可以复制粘贴',
    '话术："德国学英国、美国学英国、日本学英国，结果如何？每个国家都在模仿，但成功路径各有不同。说明技术可以模仿，制度难以复制。"'
]

for point in points_11:
    if point.startswith('核心观点：'):
        para = doc.add_paragraph()
        run = para.add_run(point)
        run.bold = True
    elif point.startswith('话术："'):
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('1.2 历史学家的新发现（10分钟）', level=3)

points_12 = [
    '制度学派：产权保护、司法独立、合同执行',
    '过渡语："为什么英国能借到钱，法国借不到？因为英国政府借了钱是真的要还的，而且有议会监督。光荣革命之后，英国国债成为全球最安全的资产。"',
    '地理学派：煤矿分布、岛屿地形、港口优势',
    '过渡语："宋朝的GDP占全球50%以上，为什么没有发生工业革命？地理因素是一个重要解释——中国的煤矿在北方，工业中心在南方，运输成本抵消了技术优势。"',
    '社会学派：新教伦理、劳动力商品化',
    '过渡语："圈地运动听起来很残酷，但从经济学角度看，它创造了工业革命所需的劳动力市场。"'
]

for point in points_12:
    if '过渡语' in point:
        para = doc.add_paragraph()
        run = para.add_run('【过渡语】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('1.3 六要素分析框架（10分钟）', level=3)

points_13 = [
    '介绍顺序：制度→资源→市场→技术→人力→金融',
    '每个要素讲一个英国案例：',
    '制度：普通法系 + 议会监督',
    '资源：煤矿与铁矿的地理重合',
    '市场：殖民地 + 国内消费',
    '技术：科学革命 + 工程师传统',
    '人力：识字率 + 技术学校',
    '金融：英格兰银行 + 股份公司',
    '总结："这六个要素不是孤立的，而是相互强化的。这就是为什么简单模仿某一个要素往往失败——你需要整套系统。"'
]

for point in points_13:
    doc.add_paragraph(point)

doc.add_page_break()

# 第二部分
doc.add_heading('第二部分：英国的特殊条件（建议时长：40分钟）', level=2)

doc.add_heading('2.1 制度优势（15分钟）', level=3)

points_21 = [
    '光荣革命的意义（1688）',
    '核心论点是王权受约束：议会是真正的权力中心',
    '三个制度创新：',
    '1. 议会财政权：政府花钱需要议会批准 → 财政纪律',
    '2. 司法独立：普通法传统 → 合同执行',
    '3. 政党轮替：避免权力垄断 → 政策可预期',
    '案例：英法对比',
    '英国国债利率：3-4%',
    '法国国债利率：6-8%（革命前）',
    '差异原因：英国有议会监督，法国君主可以赖账',
    '话术："一个能借到钱且利率低的国家，在工业化竞争中就赢了一半。"'
]

for point in points_21:
    if '话术' in point:
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('2.2 地理禀赋（10分钟）', level=3)

points_22 = [
    '三个地理优势：',
    '1. 煤炭分布与工业中心重合',
    '对比：中国煤矿在山西，工业在长三角，运输成本高',
    '2. 岛屿地形减少军事压力',
    '英国不需要维持大陆军队，可以将资源用于工业',
    '对比：普鲁士/俄国需要大量常备军',
    '3. 港口优势支撑贸易',
    '伦敦成为全球金融中心不是偶然',
    '讨论问题："地理决定论vs制度决定论——你怎么看？"'
]

for point in points_22:
    if '对比' in point:
        para = doc.add_paragraph()
        para.add_run('【对比】').bold = True
        para.add_run(point.replace('对比：', '：'))
    elif '讨论' in point:
        para = doc.add_paragraph()
        run = para.add_run('【讨论】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xC0)
    else:
        doc.add_paragraph(point)

doc.add_heading('2.3 市场发育（15分钟）', level=3)

points_23 = [
    '三个市场机制：',
    '1. 圈地运动：土地规模化 + 劳动力商品化',
    '2. 殖民地贸易：原材料进口 + 工业品出口',
    '棉花案例：印度棉花→英国纺织→美洲棉花→英国纺织→全球市场',
    '3. 消费分层：新兴中产 + 工人阶级消费',
    '棉布价格下降：1760年代到1830年代，价格下降80%',
    '话术："工业革命的本质是什么？是让普通人也能用上以前只有贵族才能用的东西。规模化生产+价格下降=市场扩大+更多投资=更多生产。"',
    '互动建议：问学员"你们行业有没有类似的'让普通人用得起'的案例？"'
]

for point in points_23:
    if '话术' in point:
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    elif '互动' in point:
        para = doc.add_paragraph()
        run = para.add_run('【互动】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xC0)
    else:
        doc.add_paragraph(point)

doc.add_page_break()

# 第三部分
doc.add_heading('第三部分：技术革命协同效应（建议时长：30分钟）', level=2)

doc.add_heading('3.1 技术簇群（15分钟）', level=3)

points_31 = [
    '核心概念：技术簇群（Technological Cluster）',
    '单个技术发明 vs 技术系统',
    '案例一：纺织机械进化',
    '1764 珍妮纺纱机：效率提升但质量不稳',
    '1769 水力纺纱机：质量提升但依赖水源',
    '1779 走锭纺纱机：两者兼顾但复杂度增加',
    '每一个发明都在解决上一个的问题，同时创造新的问题',
    '案例二：能源-运输-材料的三角关系',
    '煤炭→蒸汽机→钢铁→铁路→煤炭开采→更多煤炭',
    '话术："这不是设计出来的，是市场力量推动的。但为什么是英国？因为英国有煤矿、有铁矿、有港口、有愿意投资的企业家——缺一不可。"'
]

for point in points_31:
    if '话术' in point:
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('3.2 企业家精神（15分钟）', level=3)

points_32 = [
    '三位关键人物：',
    '1. 阿克莱特（Richard Arkwright）',
    '不是发明家，是企业家',
    '创办了第一个现代工厂（1771年）',
    '创新：工厂制度、工人管理、品质控制',
    '2. 博尔顿（Matthew Boulton）',
    '瓦特的合伙人',
    '建立了销售和服务网络',
    '3. 斯蒂芬森（George Stephenson）',
    '铁路之父',
    '将工程知识转化为商业应用',
    '讨论："发明家vs企业家——谁对工业革命更重要？"'
]

for point in points_32:
    if '讨论' in point:
        para = doc.add_paragraph()
        run = para.add_run('【讨论】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xC0)
    else:
        doc.add_paragraph(point)

doc.add_page_break()

# 第四部分
doc.add_heading('第四部分：比较视角（建议时长：30分钟）', level=2)

doc.add_heading('4.1 法国案例（10分钟）', level=3)

points_41 = [
    '法国的失败教训：',
    '1. 行会制度：保护既得利益，阻碍创新',
    '案例：新机器要获得行会认可，需要多年时间',
    '2. 金融体系：家族银行为主，中小企业融资困难',
    '3. 劳动力市场：封建残余，劳动力流动受阻',
    '4. 文化因素："学而优则仕"，轻视技术和商业',
    '对比数据：',
    '英国工厂数：1760年200家→1800年1000家',
    '法国工厂数：同期几乎没有增长',
    '话术："制度僵化的代价是失去竞争优势。法国有技术、有人才、有市场，但没有正确的制度安排。"'
]

for point in points_41:
    if '话术' in point:
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('4.2 中国比较（10分钟）', level=3)

points_42 = [
    '宋朝资本主义萌芽',
    '宋朝GDP占全球50%以上',
    '有发达的手工业、商业、银行（交子）',
    '为何没有发生工业革命？',
    '1. 明清闭关锁国：海禁政策阻断海外市场',
    '2. 政策导向：重农抑商，抑制商业资本',
    '3. 制度缺失：没有产权保护，技术扩散受阻',
    '4. 文化因素：科举制度虹吸人才',
    '讨论："郑和下西洋为什么没有变成殖民扩张？"'
]

for point in points_42:
    if '讨论' in point:
        para = doc.add_paragraph()
        run = para.add_run('【讨论】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x00, 0x00, 0xC0)
    else:
        doc.add_paragraph(point)

doc.add_heading('4.3 后发国家案例（10分钟）', level=3)

points_43 = [
    '德国（19世纪后期）：',
    '国家主导工业化战略',
    '建立完善职业教育体系（双元制）',
    '大学与工业结合（洪堡大学改革）',
    '技术引进+本土创新',
    '美国（19世纪末）：',
    '移民带来的劳动力红利',
    '广袤的国内市场',
    '大规模铁路建设',
    '托拉斯：资本集中加速工业化',
    '启示："后发优势是真实的，但需要主动建设制度环境"'
]

for point in points_43:
    doc.add_paragraph(point)

doc.add_page_break()

# 第五部分
doc.add_heading('第五部分：历史类比与框架提炼（建议时长：40分钟）', level=2)

doc.add_heading('5.1 AI革命类比（20分钟）', level=3)

points_51 = [
    '六要素框架对照：',
    '制度：数据确权、AI伦理规范、监管框架——各国正在建设中',
    '资源：算力成本下降，但高端芯片受限——部分满足',
    '市场：企业级应用，全球数十亿潜在用户——满足',
    '技术：深度学习十年积累，Transformer突破——满足',
    '人力：全球AI研究人员数十万，工程师红利——满足',
    '金融：VC/PE大量投入，科技巨头持续投资——满足',
    '关键问题："今天AI革命的主要短板是什么？"',
    '回答方向：制度（数据确权、隐私保护）和硬件（芯片）是主要瓶颈',
    '话术："看一个技术革命能不能起飞，不是看它有多炫酷，而是看它的短板在哪里。工业革命的短板是煤炭运输，AI革命的短板是算力和制度环境。"'
]

for point in points_51:
    if '话术' in point:
        para = doc.add_paragraph()
        run = para.add_run('【话术】' + point)
        run.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph(point)

doc.add_heading('5.2 起飞条件清单（20分钟）', level=3)

points_52 = [
    '重大技术革命起飞条件清单使用说明：',
    '1. 不是所有条件都必须完美满足',
    '2. 关键是短板不能太短',
    '3. 条件之间会相互强化',
    '4. 时机很重要——太早或太晚都可能失败',
    '学员练习引导：',
    '让学员用清单分析自己行业的某个新技术',
    '分组讨论：每组分析一个技术',
    '全班分享：各组发现的共同规律',
    '总结："这份清单的价值不是给你一个答案，而是让你问出正确的问题。"'
]

for point in points_52:
    doc.add_paragraph(point)

doc.add_page_break()

# ====== 第三章：互动设计 ======
doc.add_heading('第三章　互动设计说明', level=1)

doc.add_heading('3.1 小组讨论：六要素归类（15分钟）', level=2)

activity_31 = [
    ('活动目标', '让学员主动思考六要素的具体表现'),
    ('活动形式', '4-5人小组，每组领取一个国家案例（英国/法国/德国/美国/中国）'),
    ('讨论任务', '将该国工业革命/工业化的成败，归因到六要素中'),
    ('呈现方式', '每组用3分钟汇报，其他组可以提问'),
    ('时间分配', '讨论10分钟 + 汇报15分钟'),
    ('点评要点', '引导学员看到：不同国家可能在不同要素上有优势；单一要素不能决定成败')
]

act_table = doc.add_table(rows=len(activity_31)+1, cols=2)
act_table.style = 'Table Grid'
act_table.rows[0].cells[0].text = '项目'
act_table.rows[0].cells[1].text = '内容'
for cell in act_table.rows[0].cells:
    set_cell_shading(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (item, content) in enumerate(activity_31):
    row = act_table.rows[i+1]
    row.cells[0].text = item
    row.cells[1].text = content
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(row.cells[0], 'F2F2F2')

doc.add_paragraph()

doc.add_heading('3.2 角色扮演：历史决策者（20分钟）', level=2)

activity_32 = [
    ('活动目标', '让学员体验历史决策的复杂性'),
    ('情境设计', '1760年代的英国，企业家/银行家/政府官员/技术工匠四种角色'),
    ('决策任务', '面对新技术（蒸汽机/纺纱机），是否投资？如何投资？'),
    ('讨论问题', '资金从哪来？风险如何分担？市场在哪里？'),
    ('时间分配', '情境说明3分钟 + 角色准备5分钟 + 模拟辩论10分钟 + 总结2分钟'),
    ('点评要点', '引出金融体系的重要性；风险分担机制（股份公司）的作用')
]

act_table2 = doc.add_table(rows=len(activity_32)+1, cols=2)
act_table2.style = 'Table Grid'
act_table2.rows[0].cells[0].text = '项目'
act_table2.rows[0].cells[1].text = '内容'
for cell in act_table2.rows[0].cells:
    set_cell_shading(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (item, content) in enumerate(activity_32):
    row = act_table2.rows[i+1]
    row.cells[0].text = item
    row.cells[1].text = content
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(row.cells[0], 'F2F2F2')

doc.add_paragraph()

doc.add_heading('3.3 即兴辩论：AI革命vs工业革命（15分钟）', level=2)

activity_33 = [
    ('活动目标', '检验学员对类比框架的掌握程度'),
    ('辩论题目', '正方：AI革命与工业革命同样重要；反方：AI革命的影响更为深远'),
    ('辩论规则', '每方3人，各准备3分钟，辩论8分钟'),
    ('评判标准', '能否运用课程框架？有无新观点？论证是否有力？'),
    ('时间分配', '抽签分组5分钟 + 准备3分钟 + 辩论10分钟 + 点评2分钟'),
    ('点评要点', '肯定学员的洞见；指出类比的局限；强调框架的实用价值')
]

act_table3 = doc.add_table(rows=len(activity_33)+1, cols=2)
act_table3.style = 'Table Grid'
act_table3.rows[0].cells[0].text = '项目'
act_table3.rows[0].cells[1].text = '内容'
for cell in act_table3.rows[0].cells:
    set_cell_shading(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (item, content) in enumerate(activity_33):
    row = act_table3.rows[i+1]
    row.cells[0].text = item
    row.cells[1].text = content
    for para in row.cells[0].paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(row.cells[0], 'F2F2F2')

doc.add_page_break()

# ====== 第四章：常见问题 ======
doc.add_heading('第四章　常见问题与应对', level=1)

doc.add_heading('Q1：地理决定论 vs 制度决定论，哪个更重要？', level=2)

answer_1 = """这是一个经典的学术争论。作为讲师，我的建议是：

1. 两者都重要，但不是非此即彼
2. 地理提供了初始条件，制度决定了如何利用这些条件
3. 英国的成功在于地理优势和制度创新的结合
4. 可以这样类比：地理是牌，制度是打牌的技术

如果学员追问，可以补充：
- 地理学派的代表：贾雷德·戴蒙德《枪炮、病菌与钢铁》
- 制度学派的代表：道格拉斯·诺斯《西方世界的兴起》
- 两派都有解释力，但都不能单独解释全部现象"""

doc.add_paragraph(answer_1)

doc.add_heading('Q2：中国能不能绕过制度短板实现技术领先？', level=2)

answer_2 = """这个问题涉及对中国当下发展的讨论。作为讲师，建议：

1. 承认中国在某些技术领域（5G、新能源等）的进步
2. 指出制度短板在某些领域确实是瓶颈（如半导体）
3. 强调"后发优势"与"后发劣势"的区别
4. 避免政治敏感讨论，将重点放在分析框架上

可以用德国和韩国的案例来回应：
- 德国在19世纪后期通过国家主导的制度建设实现了工业化
- 韩国在20世纪后期通过出口导向和产业政策实现了技术跨越
- 关键在于：是否有意识地建设制度环境"""

doc.add_paragraph(answer_2)

doc.add_heading('Q3：AI革命真的可以和工业革命相比吗？', level=2)

answer_3 = """这个问题可以用来检验学员的批判性思维：

1. 类比的价值：帮助我们提出正确的问题
2. 类比的局限：两个革命在速度、范围、驱动力上有显著差异
3. 关键差异：
   - 工业革命主要影响物质生产
   - AI革命正在影响知识生产和决策
   - AI革命的监管挑战更为复杂

4. 建议用"起飞条件清单"来评估，而不是简单说"一样"或"不一样""

doc.add_paragraph(answer_3)

doc.add_heading('Q4：为什么瓦特成功了，而纽可门失败了？', level=2)

answer_4 = """这是技术史的经典问题：

1. 纽可门大气机（1712）：效率低，只能用于煤矿抽水
2. 瓦特改进（1769）：增加冷凝器，效率提升4倍
3. 商业成功的关键：
   - 博尔顿的商业网络
   - 专利保护（ patents）
   - 持续改进（从矿山到纺织到运输）

4. 教训：技术发明不等于商业成功。瓦特的成功是技术+商业+制度的胜利"""

doc.add_paragraph(answer_4)

doc.add_heading('Q5：普通学员学这个课有什么用？', level=2)

answer_5 = """这是课程价值的核心问题：

1. 建立历史比较思维：不只看技术，要看系统
2. 避免"技术决定论"的陷阱
3. 用六要素框架分析自己行业的变革
4. 理解制度环境对商业的影响
5. 培养从历史中学习的能力

可以用学员自己的行业举例，帮助他们看到框架的实用价值"""

doc.add_paragraph(answer_5)

doc.add_page_break()

# ====== 第五章：案例话术 ======
doc.add_heading('第五章　案例讲解话术', level=1)

doc.add_heading('案例一：光荣革命与国债制度', level=2)

script_1 = """【开场】
"1688年，英国发生了一场'不流血的革命'。这场革命听起来不如法国大革命轰轰烈烈，但它为英国带来的变化，可能比法国大革命还要深远。"

【核心数据】
光荣革命后，英国政府能够以3-4%的利率借到钱。
同期法国国王的借贷利率是6-8%。
这个利率差异，在几十年后会变成国力的差距。

【解释机制】
"为什么英国的信用更好？因为光荣革命确立了议会监督财政的权力。英国国王借钱，议会有权监督还款。如果国王赖账，议会可以拒绝批准新的税收法案。
法国呢？国王一个人说了算，今天可以借钱，明天可以宣布债务无效。"

【结论】
"国债制度让英国能够以低成本融资，这些钱用来建造舰队、修建道路、投资技术。当法国还在为军费发愁的时候，英国已经在为工业革命做准备了。"

【互动问题】
"如果你是1688年的英国议员，你会担心国债制度带来什么问题？" """

for para_text in script_1.strip().split('\n'):
    if para_text.startswith('【'):
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        doc.add_paragraph(para_text)

doc.add_heading('案例二：阿克莱特与工厂制度', level=2)

script_2 = """【开场】
"阿克莱特可能是工业革命中最重要的企业家，但你在历史书上可能找不到他的名字。瓦特是发明家，但阿克莱特是现代管理模式的开创者。"

【核心贡献】
1769年，阿克莱特获得水力纺纱机专利
1771年，在克罗姆福德建立第一个现代工厂
工厂规模：雇工300人，24小时运转

【创新之处】
"以前的纺织工匠在家工作，按件计酬。阿克莱特的工厂改变了这一切：
第一，工人必须到工厂上班
第二，工作时间由工厂规定
第三，工资按工作时间而非产品数量计算
这意味着什么？意味着劳动成了商品。"

【争议与评价】
"圈地运动把农民从土地上赶走，阿克莱特的工厂把农民变成工人。没有这个转变，工业革命就不可能发生。"

【结论】
"阿克莱特的创新不是技术发明，而是生产组织的创新。工厂制度让规模化生产成为可能。"

【互动问题】
"今天的远程工作是不是对工厂制度的某种'回归'？你怎么看？" """

for para_text in script_2.strip().split('\n'):
    if para_text.startswith('【'):
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        doc.add_paragraph(para_text)

doc.add_heading('案例三：英国与法国的金融体系对比', level=2)

script_3 = """【开场】
"我们来做一道数学题。如果你要创办一家铁路公司，需要100万英镑。你能从银行借到吗？1770年的英国，你可以。1770年的法国，你借不到。为什么？"

【对比分析】
英国：
- 英格兰银行（1694）提供贷款
- 股份公司可以向公众发行股票
- 债券市场让政府可以大规模融资
- 利率：3-4%

法国：
- 家族银行占主导，不愿意长期贷款
- 股份公司受限制
- 缺乏统一的债券市场
- 利率：6-8%（有时更高）

【背后的制度差异】
"为什么差距这么大？因为英国的制度让债权人放心。光荣革命之后，如果英国国王想赖账，议会不会答应。法国呢？国王一句话，债务可以一笔勾销。"

【结论】
"金融体系的核心是信任。制度决定了信任的成本。英国的低利率背后是低风险溢价，而低风险溢价来自制度的保障。"

【互动问题】
"今天，为什么硅谷的科技公司能拿到那么多投资？这和19世纪的英国有什么相似之处？" """

for para_text in script_3.strip().split('\n'):
    if para_text.startswith('【'):
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        doc.add_paragraph(para_text)

doc.add_page_break()

# ====== 第六章：时间把控 ======
doc.add_heading('第六章　时间把控建议', level=1)

doc.add_heading('6.1 课程总时间分配', level=2)

time_table = doc.add_table(rows=8, cols=3)
time_table.style = 'Table Grid'

time_data = [
    ('环节', '时长', '累计'),
    ('引言与破题', '10分钟', '10分钟'),
    ('第一部分：重新理解工业革命', '30分钟', '40分钟'),
    ('第二部分：英国的特殊条件', '40分钟', '80分钟'),
    ('第三部分：技术革命协同效应', '30分钟', '110分钟'),
    ('第四部分：比较视角', '30分钟', '140分钟'),
    ('第五部分：历史类比与框架', '40分钟', '180分钟'),
    ('总计', '180分钟', '3小时')
]

for i, row_data in enumerate(time_data):
    row = time_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0 or i == 7:
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(row.cells[j], '4472C4')
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_shading(row.cells[j], 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_paragraph()

doc.add_heading('6.2 时间超标处理预案', level=2)

time_tips = [
    ('如果某个案例讨论超时', '可以在最后的"学员练习"环节补时间，不要压缩其他案例的讲解'),
    ('如果学员问题太多', '可以记录下来，在茶歇时回答，或者引导到课后讨论群'),
    ('如果互动环节冷场', '不要强制点名，用开放式问题引导；必要时可以自己先做示范'),
    ('如果时间紧张', '第五部分（历史类比）不能省，这是课程的产出核心；前面的案例可以精简'),
    ('如果时间充裕', '可以增加学员分享自己行业案例的环节，让大家用框架分析更多案例')
]

for title, content in time_tips:
    para = doc.add_paragraph()
    para.add_run(title + '：').bold = True
    para.add_run(content)

doc.add_heading('6.3 茶歇建议', level=2)

break_text = """建议安排在课程第80分钟（第二部分结束后）。
茶歇时间：10分钟
茶歇地点：教室外走廊或休息区
茶歇时可以做的小事：
- 播放BBC工业革命纪录片片段（3-5分钟）
- 展示工业革命时期的机器图片
- 发放课程相关的补充阅读材料"""

doc.add_paragraph(break_text)

doc.add_page_break()

# ====== 第七章：评分标准 ======
doc.add_heading('第七章　评分标准与评估', level=1)

doc.add_heading('7.1 学员表现评估维度', level=2)

eval_table = doc.add_table(rows=6, cols=4)
eval_table.style = 'Table Grid'

eval_data = [
    ('评估维度', '权重', '优秀（90-100）', '良好（70-89）'),
    ('框架掌握', '30%', '能熟练运用六要素框架分析问题', '能说出六要素框架的主要内容'),
    ('案例分析', '25%', '能举出1-2个案例并进行分析', '能复述课程中的案例'),
    ('批判思维', '20%', '能提出独特见解，指出框架局限', '能思考框架的适用范围'),
    ('课堂参与', '15%', '积极发言，带动讨论氛围', '能够回应老师的问题'),
    ('实际应用', '10%', '能将框架应用于自己的工作场景', '知道如何在自己的场景中应用')
]

for i, row_data in enumerate(eval_data):
    row = eval_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(row.cells[j], '4472C4')
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_shading(row.cells[j], 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_paragraph()

doc.add_heading('7.2 课程满意度评估', level=2)

satisfaction_text = """课程结束时，可以发放简短的满意度问卷：

评估项目（5分制）：
1. 课程内容的实用性
2. 案例讲解的清晰度
3. 互动的有效性
4. 时间节奏的把控
5. 整体满意度

开放问题：
- 本课程对你最有价值的一点是？
- 你认为可以改进的地方是？
- 你会将本课程推荐给同事吗（是/否，原因）？"""

doc.add_paragraph(satisfaction_text)

doc.add_heading('7.3 课后跟进建议', level=2)

followup_text = """1. 课后24小时：
   - 发送课程PPT和补充阅读材料到学员群
   - 收集学员的课后感想和改进建议

2. 课后一周：
   - 邀请学员分享他们用框架分析自己行业的发现
   - 评选"最佳应用案例"并给予小额奖励

3. 课后一月：
   - 跟踪学员是否在工作中应用了课程框架
   - 收集应用案例，更新课程案例库

4. 持续改进：
   - 根据学员反馈，每年更新一次案例和互动设计
   - 保持课程内容与当下技术变革的关联性"""

doc.add_paragraph(followup_text)

doc.add_page_break()

# ====== 附录 ======
doc.add_heading('附录：参考文献与推荐阅读', level=1)

refs = [
    ('核心文献', [
        '艾伦：《工业革命》- 牛津大学公开课教材',
        '肯尼迪：《大国的兴衰》- 技术与经济的长期分析',
        '诺斯：《西方世界的兴起》- 制度学派代表作',
        '戴蒙德：《枪炮、病菌与钢铁》- 地理学派代表作'
    ]),
    ('补充阅读', [
        '弗格森：《帝国》- 大英帝国的兴衰',
        '霍布斯鲍姆：《工业与帝国》- 工业革命的社会影响',
        '克里季：《创新者们》- 数字革命的群像'
    ]),
    ('视频资源', [
        'BBC《工业革命》纪录片系列',
        '网易公开课：牛津大学《工业革命》'
    ])
]

for section_title, items in refs:
    para = doc.add_paragraph()
    run = para.add_run(section_title)
    run.bold = True
    run.font.size = Pt(14)

    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# 页脚
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = '工业革命讲师手册'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
doc.save(output_path)
print(f"讲师手册已保存至: {output_path}")

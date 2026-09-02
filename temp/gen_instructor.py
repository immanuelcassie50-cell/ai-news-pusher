# -*- coding: utf-8 -*-
"""Generate Instructor Handbook Word Document"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"D:\新课开发\行动学习2026\02-对事-教程\完整课程包\讲师手册\看清现实找到能动的缝隙_讲师手册_v1.0.docx"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def create_handbook():
    doc = Document()

    # Page setup - A4 landscape
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ===== Cover Page =====
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("看清现实，找到能动的缝隙")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0xB8, 0x35, 0x1C)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle_p.add_run("信息分析与突破口识别 · 讲师手册")
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x1F, 0x1B, 0x16)

    doc.add_paragraph()

    info_table = doc.add_table(rows=3, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("课程时长：", "180分钟", "学员人数：", "20-30人"),
        ("课程对象：", "中层管理者/骨干员工", "场地要求：", "分组教室"),
        ("版本号：", "v1.0", "发布日期：", "2026年"),
    ]
    for ri, row_data in enumerate(info_data):
        row = info_table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11)
                    if ci % 2 == 0:
                        run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    quote_p = doc.add_paragraph()
    quote_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote_p.add_run("*今天你不只是来听课的，你是来经历一个完整的分析过程的。*")
    quote_run.italic = True
    quote_run.font.size = Pt(12)
    quote_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== How to Use This Handbook =====
    add_heading(doc, "讲师使用指南", 1)

    p = doc.add_paragraph()
    p.add_run("这本讲师手册不是流程剧本，是地图。")
    p.add_run("\n\n课程是一个引导学员完成真实分析工作的过程。你需要根据现场学员的实际情况，决定什么时候展开讨论、什么时候给案例、什么时候让学员自己写。手册里的时间分配是参考值，不是硬性规定。")

    add_heading(doc, "三个关键原则", 2)

    doc.add_paragraph("原则一：真实任务优先于完美流程", style='List Bullet')
    p1 = doc.add_paragraph("如果学员带来的任务足够真实，即使讨论超时、环节没有走完，效果也比用虚构案例把流程走完要好。")

    doc.add_paragraph("原则二：让学员自己写，不要替他们总结", style='List Bullet')
    p2 = doc.add_paragraph('讲师最容易犯的错误是：学员还没写完，就给出"标准答案"。每一个环节，一定要给学员充分的写作时间，然后让他们自己分享。')

    doc.add_paragraph("原则三：提问比告知更有价值", style='List Bullet')
    p3 = doc.add_paragraph("当学员分享完毕后，不要急着评价，先问他们：你为什么这么判断？依据是什么？")

    doc.add_page_break()

    # ===== Course Overview =====
    add_heading(doc, "课程全局时间表", 1)

    time_table = doc.add_table(rows=8, cols=4)
    time_table.style = 'Table Grid'
    time_data = [
        ("部分", "内容", "时长", "关键产出"),
        ("开场", "课程导入+真实任务导入", "15分钟", "学员确定自己的真实任务"),
        ("第一部分", "体检思维", "20分钟", "识别自己的三个偏误"),
        ("第二部分", "体检清单", "25分钟", "基于框架重建体检清单"),
        ("第三部分", "调研设计", "25分钟", "完成调研计划"),
        ("第四部分", "四维分析", "30分钟", "完成四维分析表"),
        ("第五部分", "突破口识别", "30分钟", "确定2-4个突破口"),
        ("第六部分", "行动方案", "30分钟", "完成完整行动方案"),
    ]

    for ri, row_data in enumerate(time_data):
        row = time_table.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
                        set_cell_shading(cell, "1F3864")
                        run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_page_break()

    # ===== Part 1: 体检思维 =====
    add_heading(doc, "第一部分 体检思维（20分钟）", 1)

    add_heading(doc, "学习目标", 2)
    doc.add_paragraph("识别三种常见的认知偏误及其对分析的影响", style='List Bullet')
    doc.add_paragraph("理解体检思维与找原因思维的本质区别", style='List Bullet')

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "开场故事（5分钟）", 3)
    doc.add_paragraph("直接读轨道交通案例，不需要放PPT。读完之后问：这个团队的问题出在哪里？让学员自由回答。")

    add_heading(doc, "三个偏误讲解（8分钟）", 3)
    doc.add_paragraph("可及性偏误：不是能力问题，是注意力分配问题。最近最烦的≠最重要的。", style='List Bullet')
    doc.add_paragraph("归责偏误：找是谁的问题是本能，但系统性因素才是真正的杠杆点。", style='List Bullet')
    doc.add_paragraph("局部视角偏误：每个层级的人都只能看到自己那一角，这是正常的，但要意识到它的存在。", style='List Bullet')

    add_heading(doc, "学员练习（5分钟）", 3)
    doc.add_paragraph("给学员3分钟快速写下他们当前项目的因素，然后给每条打标签。")

    add_heading(doc, "讨论问题", 3)
    doc.add_paragraph('有多少条被打上了"可及"标签？这说明什么？', style='List Bullet')
    doc.add_paragraph("有多少条指向了人或部门？这和系统性因素的区别是什么？", style='List Bullet')

    add_heading(doc, "常见问题", 2)
    doc.add_paragraph('Q：学员说"我已经想得很全面了，不需要体检"', style='List Bullet')
    doc.add_paragraph('A：让他把那套"全面的分析"写下来，和体检清单的框架对照，看有没有遗漏的维度。', style='List Bullet')

    doc.add_page_break()

    # ===== Part 2: 体检清单 =====
    add_heading(doc, "第二部分 体检清单（25分钟）", 1)

    add_heading(doc, "学习目标", 2)
    doc.add_paragraph("理解体检清单的本质和格式要求", style='List Bullet')
    doc.add_paragraph("掌握用维度框架系统覆盖所有相关因素的方法", style='List Bullet')

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "讲解要点（5分钟）", 3)
    doc.add_paragraph("体检清单三个要点：名词短语、不加评价、穷举式覆盖。", style='List Bullet')
    doc.add_paragraph('强调"不加评价"是最难的——我们很容易一边列因素一边想"这个问题不大"。', style='List Bullet')

    add_heading(doc, "维度框架介绍（5分钟）", 3)
    doc.add_paragraph("7个维度：资源、流程、能力、系统与工具、管理机制、外部条件、历史遗留。", style='List Bullet')
    doc.add_paragraph('每个维度给1-2个例子，确保学员理解"什么叫这个维度的因素"。', style='List Bullet')

    add_heading(doc, "学员练习（12分钟）", 3)
    doc.add_paragraph("给学员10分钟，用框架重建自己的体检清单。讲师巡场，注意：", style='List Bullet')
    doc.add_paragraph("格式是不是名词短语？有没有出现评价性词汇？", style='List Bullet')
    doc.add_paragraph("维度覆盖是否完整？", style='List Bullet')

    add_heading(doc, "讨论与分享（3分钟）", 3)
    doc.add_paragraph("每组分享1条：你们认为哪个维度最容易漏掉？为什么？")

    add_heading(doc, "常见问题", 2)
    doc.add_paragraph("Q：历史遗留维度不知道怎么写", style='List Bullet')
    doc.add_paragraph('A：问学员"这个问题以前有人试图解决过吗？结果怎样？"', style='List Bullet')

    doc.add_page_break()

    # ===== Part 3: 调研设计 =====
    add_heading(doc, "第三部分 调研设计（25分钟）", 1)

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "为什么要调研（3分钟）", 3)
    doc.add_paragraph("不能跳过调研的两个原因：体检清单是假设的合集；信息质量决定分析质量。")

    add_heading(doc, "四个核心问题（7分钟）", 3)
    doc.add_paragraph("哪些信息已经有？库存盘点，避免重复收集。", style='List Bullet')
    doc.add_paragraph("用什么方式收集？访谈/观察/数据分析/文档研读各适用不同类型。", style='List Bullet')
    doc.add_paragraph("找谁收集？同一问题要向不同层级的人了解。", style='List Bullet')
    doc.add_paragraph("怎么确保可靠？交叉验证原则。", style='List Bullet')

    add_heading(doc, "特别维度：历史改善尝试（5分钟）", 3)
    doc.add_paragraph("这个维度是学员最容易跳过的，但往往也是最有价值的。", style='List Bullet')
    doc.add_paragraph("问：如果这个问题以前改善过，为什么没有彻底解决？", style='List Bullet')

    add_heading(doc, "学员练习（8分钟）", 3)
    doc.add_paragraph("针对自己的体检清单，设计调研计划。", style='List Bullet')

    add_heading(doc, "常见问题", 2)
    doc.add_paragraph('Q：学员说"调研太费时间，项目周期不允许"', style='List Bullet')
    doc.add_paragraph('A：问"不调研就做判断，代价是什么？"让学员自己算这笔账。', style='List Bullet')

    doc.add_page_break()

    # ===== Part 4: 四维分析 =====
    add_heading(doc, "第四部分 四维分析（30分钟）", 1)

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "四维框架讲解（5分钟）", 3)
    doc.add_paragraph("影响大小：影响有多显著？", style='List Bullet')
    doc.add_paragraph("影响范围：是局部还是系统性的？", style='List Bullet')
    doc.add_paragraph("可动性：实际上能被推动吗？", style='List Bullet')
    doc.add_paragraph("突破可能性：能推动到什么程度？", style='List Bullet')

    add_heading(doc, "可动性判断陷阱（5分钟）", 3)
    doc.add_paragraph('陷阱一：把"应该可动"当成"实际可动"。', style='List Bullet')
    doc.add_paragraph("陷阱二：只考虑自己能做什么，忘了需要其他人配合的情况。", style='List Bullet')
    doc.add_paragraph("陷阱三：可动性是动态的，今天不可动不代表明天不可动。", style='List Bullet')

    add_heading(doc, "学员练习（15分钟）", 3)
    doc.add_paragraph("用四维分析表对自己的关键因素打分。", style='List Bullet')
    doc.add_paragraph('讲师注意：学员容易只填"高/中/低"，要引导他们写出判断依据。', style='List Bullet')

    add_heading(doc, "讨论（5分钟）", 3)
    doc.add_paragraph('有没有因素你觉得"应该很重要但可动性很低"？这说明什么？')

    doc.add_page_break()

    # ===== Part 5: 突破口识别 =====
    add_heading(doc, "第五部分 突破口识别（30分钟）", 1)

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "突破口三个条件（5分钟）", 3)
    doc.add_paragraph("影响显著：改善之后目标指标有可感知的变化。", style='List Bullet')
    doc.add_paragraph("有实际撬动可能：在当前资源、权限、时机条件下可以实质性推动。", style='List Bullet')
    doc.add_paragraph("在项目周期内能看到变化：不是三年后见效。", style='List Bullet')

    add_heading(doc, "优先级矩阵（5分钟）", 3)
    doc.add_paragraph("直接可动+高影响 = 优先突破口。", style='List Bullet')
    doc.add_paragraph("间接可动+高影响 = 需要撬动策略的突破口。", style='List Bullet')
    doc.add_paragraph("当前不可动+高影响 = 单独标注，长期关注。", style='List Bullet')

    add_heading(doc, "深度验证5问（8分钟）", 3)
    doc.add_paragraph("这5个问题要学员认真回答，不能跳过。", style='List Bullet')
    doc.add_paragraph("重点问：最好的结果是什么？需要多长时间能看到？", style='List Bullet')

    add_heading(doc, "学员练习（10分钟）", 3)
    doc.add_paragraph("识别2-4个候选突破口，完成深度验证。", style='List Bullet')

    doc.add_page_break()

    # ===== Part 6: 行动方案 =====
    add_heading(doc, "第六部分 行动方案（30分钟）", 1)

    add_heading(doc, "关键引导点", 2)

    add_heading(doc, "三层结构（5分钟）", 3)
    doc.add_paragraph("What：解决方向，一句话说清楚我们要做什么。", style='List Bullet')
    doc.add_paragraph("How：具体举措，2-5个可执行的动作。", style='List Bullet')
    doc.add_paragraph("Pre-flight check：前置条件检查。", style='List Bullet')

    add_heading(doc, "学员练习（20分钟）", 3)
    doc.add_paragraph("把突破口转化为完整的行动方案。", style='List Bullet')
    doc.add_paragraph("时间分配：解决方向5分钟，具体举措10分钟，前置条件5分钟。", style='List Bullet')

    add_heading(doc, "最后复盘（5分钟）", 3)
    doc.add_paragraph("问学员：这个行动方案，和你一开始的分析有什么不同？", style='List Bullet')
    doc.add_paragraph("为什么要先做体检、再做调研、再四维分析？", style='List Bullet')

    doc.add_page_break()

    # ===== Appendix =====
    add_heading(doc, "附录", 1)

    add_heading(doc, "附录一：讲师工具清单", 2)
    doc.add_paragraph("白板/大白纸（每组至少2张）", style='List Bullet')
    doc.add_paragraph("彩色便签纸（至少4色）", style='List Bullet')
    doc.add_paragraph("计时器（小组讨论时使用）", style='List Bullet')
    doc.add_paragraph("学员手册（每人一本）", style='List Bullet')

    add_heading(doc, "附录二：分组建议", 2)
    doc.add_paragraph("最优分组：4-5人/组，混合部门分组效果更好。", style='List Bullet')
    doc.add_paragraph("每组需要有一人担任记录员，负责把讨论结果写在大纸上。", style='List Bullet')

    add_heading(doc, "附录三：时间调整参考", 2)
    doc.add_paragraph("如果时间不够，优先保证：学员写作时间 > 讨论环节 > 讲师讲解。", style='List Bullet')
    doc.add_paragraph("可以压缩的部分：开场案例时间、每个偏误的展开讨论。", style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*版权所有 · 罗宏伟 · 本手册仅供本课程讲师使用*")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUTPUT_PATH)
    print("Instructor handbook generated:", OUTPUT_PATH)

if __name__ == "__main__":
    create_handbook()

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'D:/新课开发/经营/系列/09_大客户开发与管理/课程大纲/02_对外大纲_大客户开发与管理.docx'

doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Color palette
DARK_RED = RGBColor(0xB8, 0x10, 0x25)
LIGHT_RED = RGBColor(0xF0, 0xE8, 0xE8)
CREAM = RGBColor(0xF6, 0xF3, 0xEF)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_TEXT = RGBColor(0x66, 0x66, 0x66)

def set_font(run, name='Microsoft YaHei', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), name)

def add_heading(doc, text, level=1, color=DARK_RED, size=16, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)
    return p

def add_para(doc, text, size=11, color=DARK_TEXT, space_before=0, space_after=6, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return p

def add_bullet(doc, text, size=11, color=DARK_TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size, color=color)
    return p

def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

# ============ COVER / HERO ============
# Top brand bar
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run('经营进阶')
set_font(run, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# Shade paragraph background
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'B81025')
pPr.append(shd)

doc.add_paragraph()

# Course title
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('大客户开发与管理')
set_font(run, size=32, bold=True, color=DARK_RED)

# Subtitle
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('策略驱动的大客户经营系统')
set_font(run, size=16, color=GRAY_TEXT)

# Metadata badges line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(24)
run = p.add_run('9小时  |  2天  |  7大模块  |  7项工具')
set_font(run, size=12, bold=True, color=DARK_RED)

# Thin separator line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), 'B81025')
pBdr.append(bottom)
pPr.append(pBdr)

# ============ COURSE OVERVIEW ============
add_heading(doc, '课程简介', size=14, space_before=12, space_after=8)

p = add_para(doc, '这是一门面向企业销售团队负责人、大客户经理、商务拓展人员的实战课程。课程聚焦于如何识别、开发、维系那些能够为企业带来战略性价值的大客户——不是泛泛的"客户关系管理"，而是系统性的战略级大客户经营方法论。', size=11, space_after=12)

# ============ TARGET AUDIENCE ============
add_heading(doc, '目标学员', size=14, space_before=6, space_after=8)

audience_items = [
    '大客户销售经理/主管',
    '商务拓展负责人',
    '企业销售团队管理层',
    '负责重要客户关系的业务骨干',
]
for item in audience_items:
    add_bullet(doc, item)

doc.add_paragraph()

# ============ COURSE VALUE ============
add_heading(doc, '课程价值', size=14, space_before=6, space_after=8)

value_points = [
    ('战略级大客户识别框架', '不是所有大客户都值得投入，这套评估体系帮助你识别真正的战略级客户'),
    ('深度需求洞察方法', '从表面需求到战略驱动，真正理解大客户决策逻辑'),
    ('系统化开发路径', '从初次接触到战略合作，设计完整的开发路径和关键动作'),
    ('决策链渗透策略', '穿透大客户复杂的组织结构，找到真正的决策者和推动者'),
    ('持续经营与价值升级', '从一次性交易到长期战略伙伴，建立持续增值的合作关系'),
]

for i, (title, desc) in enumerate(value_points, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('{}. '.format(i))
    set_font(run, size=11, bold=True, color=DARK_RED)
    run2 = p.add_run(title + '：')
    set_font(run2, size=11, bold=True, color=DARK_TEXT)
    run3 = p.add_run(desc)
    set_font(run3, size=11, color=DARK_TEXT)

doc.add_paragraph()

# ============ LEARNING OUTCOMES ============
add_heading(doc, '学习产出', size=14, space_before=6, space_after=8)

outcomes = [
    '掌握大客户价值评估的五个维度和方法',
    '能够独立完成目标大客户的深度画像',
    '设计完整的客户开发策略和路径图',
    '建立关键人关系网络和管理方案',
    '掌握双赢谈判的核心筹码和让步策略',
    '构建大客户服务体系和健康度监测机制',
    '制定客户持续经营计划，实现增购和续约',
]
for item in outcomes:
    add_bullet(doc, item)

doc.add_paragraph()

# ============ MODULES ============
add_heading(doc, '课程大纲', size=14, space_before=6, space_after=8)

modules = [
    ('模块一：大客户战略——识别真正的"大"客户', '90分钟', [
        '大客户评估五维度模型',
        '从收入导向到战略导向的思维转变',
        '工具：大客户价值评估矩阵',
    ]),
    ('模块二：大客户画像——需求冰山下的真实驱动', '120分钟', [
        '需求三层结构：表面/业务/战略',
        '大客户画像三维图',
        '工具：需求深挖话术树',
    ]),
    ('模块三：开发策略——从线索到战略合作', '120分钟', [
        'SPIN-PRO开发路径四阶段',
        '客户接触五大关键时机',
        '工具：开发策略书模板',
    ]),
    ('模块四：关系建构——高层关系与决策链渗透', '120分钟', [
        '关键人关系导航图绘制',
        '利益相关者矩阵分析',
        '工具：关键人关系突破方案',
    ]),
    ('模块五：谈判与成交——价值交换的艺术', '120分钟', [
        '双赢谈判框架四要素',
        '让步策略矩阵设计',
        '工具：核心谈判筹码清单',
    ]),
    ('模块六：大客户服务体系——从交易到共生', '90分钟', [
        '大客户服务蓝图设计',
        '客户健康度仪表盘建立',
        '工具：服务体系设计方案',
    ]),
    ('模块七：持续经营——增购、续约与生态共建', '90分钟', [
        '客户生命周期价值地图',
        '增购四大触发点识别',
        '工具：客户持续经营计划',
    ]),
]

for title, duration, items in modules:
    # Module title bar
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'B81025')
    pPr.append(shd)
    run = p.add_run('  {}  |  {}'.format(title, duration))
    set_font(run, size=11, bold=True, color=RGBColor(0xFF,0xFF,0xFF))

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run('• ' + item)
        set_font(run, size=10, color=DARK_TEXT)

doc.add_paragraph()

# ============ FOOTER ============
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(0)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'F6F3EF')
pPr.append(shd)
run = p.add_run('  经营进阶  |  大客户开发与管理  |  课程代码：KA-2024  |  联系方式：400-XXX-XXXX')
set_font(run, size=9, color=GRAY_TEXT)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print('Saved:', OUT)

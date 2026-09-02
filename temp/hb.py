# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = 'D:/新课开发/经验萃取/带教手册/完整课程包/05_学员手册/学员手册_组织经验传承_AI赋能岗位带教手册开发.docx'
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc = Document()

def set_cell_shading(cell, fill_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_headers(doc, headers, rows, header_color="4472C4"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_color)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
    return table

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 56, 100)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(46, 84, 150)
    return heading

def add_para(doc, text, bold=False, italic=False, color=None, size=10):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def add_form_table(doc, cells, cols):
    rows = len(cells) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            idx = r * cols + c
            cell = table.rows[r].cells[c]
            cell.text = cells[idx]
            if r == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, "4472C4")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    return table

# Cover page
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("组织经验传承")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("AI赋能岗位带教手册开发")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(46, 84, 150)

doc.add_paragraph()

label = doc.add_paragraph()
label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = label.add_run("学员手册")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()

add_table_with_headers(doc, [], [
    ["学员姑不", "________________________"],
    ["所在部门", "________________________"],
    ["课程日期", "________________________"],
    ["课程讲师", "________________________"]
])

doc.add_paragraph()

quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote.add_run(""带教不是把事情交代清楚，而是让你的弟弟真正学会。"")
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.save(OUTPUT_PATH)
print("Document created: " + OUTPUT_PATH)

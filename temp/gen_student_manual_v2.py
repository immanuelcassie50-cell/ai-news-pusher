"""
gen_student_manual_v2.py - 扩充版学员手册,目标 80-120 页
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\2026年课程\竞越\教练技术：新的管理方式\完整课程包\14_Word手册"

PURPLE = RGBColor(0x5B, 0x3A, 0x8C)
ORANGE = RGBColor(0xF2, 0xA0, 0x3D)
INK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x9A, 0x98, 0x90)
TERRACOTTA = RGBColor(0xB8, 0x5C, 0x3E)
GOLD = RGBColor(0xC9, 0xA9, 0x61)
LIGHT_GREY = RGBColor(0xE8, 0xE5, 0xE0)
RED = RGBColor(0xC0, 0x39, 0x2B)


def set_zh_font(run, font_name="思源黑体"):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")


def add_body_no_indent(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")


def add_quote(doc, text, source=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run('"' + text + '"')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run("—— " + source)
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = GREY
        set_zh_font(r2, "思源黑体")


def add_case(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("【案例】" + title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TERRACOTTA
    set_zh_font(r, "思源宋体")
    for line in body:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + line)
        r.font.size = Pt(11)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


def add_faq(doc, q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Q: " + q)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源黑体")
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.74)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("A: " + a)
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK
    set_zh_font(r2, "思源黑体")


def add_blank(doc, hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "bottom", "left", "right"]:
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "dashed")
        bdr.set(qn("w:sz"), "8")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "C0C0C0")
        pBdr.append(bdr)
    pPr.append(pBdr)
    if hint:
        r = p.add_run("【此处填写】 " + hint)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GREY
        set_zh_font(r, "思源黑体")
    for _ in range(3):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(" ")


def add_table_simple(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_zh_font(r, "思源黑体")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2A03D")
        tcPr.append(shd)
    for i, row in enumerate(rows):
        for j, cell_data in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_data) if cell_data else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = INK
                    set_zh_font(r, "思源黑体")


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def add_cover(doc, title, subtitle, kicker):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(48)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(subtitle)
    r.font.size = Pt(28)
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(kicker)
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    r = p.add_run("罗宏伟 · 竞越合作讲师")
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    set_zh_font(r, "思源黑体")


def add_toc(doc):
    add_h1(doc, "目录")
    items = [
        "第一部分 课程概览",
        "模块一 AI 时代的管理坍塌",
        "模块二 教练 vs 传统管理者的 6 个差异",
        "模块三 3C 信任模型 + 4 模式员工",
        "模块四 GROW+ 5 阶段对话框架",
        "模块五 七层倾听阶梯",
        "模块六 30 天落地实践",
        "第二部分 19 个工具使用说明",
        "第三部分 真实案例集(8 例)",
        "第四部分 30 天行动手册",
        "第五部分 反思日记模板",
        "第六部分 常见问题 FAQ",
        "附录:6 大金句卡片",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


def add_module_one(doc):
    add_h1(doc, "模块一 AI 时代的管理坍塌")
    add_body(doc, "为什么 20 年的老管理,突然不会管了?")
    add_body(doc, "不是你变差了,是你脚下的地,塌了。")
    add_body(doc, "本章我们一起看清 3 次坍塌,理解'为什么必须变'。")

    add_h2(doc, "1.1 第一次坍塌:信息差消失")
    add_body(doc, "过去 30 年,管理者是信息中枢。员工遇到问题,来问你。")
    add_body(doc, "现在,AI 3 秒出答案。员工比你先知道。")
    add_body(doc, "权威根基动摇——你凭什么管?")

    add_h3(doc, "1.1.1 案例:老王翻手册的尴尬")
    add_case(doc, "老王查参数,翻不过手机",
        [
            "60 后车间主任,30 年工龄",
            "新员工用手机扫了一下设备,3 秒读出参数",
            "老王翻了 5 分钟手册没找到",
            "他开始怀疑:我是不是该退了?",
            "30 年经验,被 3 秒打败",
        ])

    add_h3(doc, "1.1.2 信息坍塌的 3 个征兆")
    add_body_no_indent(doc, "1. 员工问你问题前,先问 AI")
    add_body_no_indent(doc, "2. 会议上,你最后才知道结论")
    add_body_no_indent(doc, "3. 你的经验,员工记不下来")

    add_h3(doc, "1.1.3 信息坍塌的 3 个应对")
    add_body_no_indent(doc, "1. 接受:你不是唯一的信息源,这是好事")
    add_body_no_indent(doc, "2. 升级:从'告诉你'变成'和你一起判断'")
    add_body_no_indent(doc, "3. 整合:把员工用 AI 找到的信息,组织成决策")

    add_h2(doc, "1.2 第二次坍塌:经验失效")
    add_body(doc, "过去 30 年,经验就是权威。10 年工龄,说了算。")
    add_body(doc, "现在,AI 看过 100 万案例,比你多 1 万倍。")
    add_body(doc, "经验贬值——30 年老师傅,3 个月被 AI 替代。")

    add_h3(doc, "1.2.1 案例:30 年老师傅被 AI 替代")
    add_case(doc, "上海某机械厂老师傅",
        [
            "30 年工龄,看裂纹 30 年,听声音就知道哪里有问题",
            "AI 音频分析 + 图像识别,准确率 98%",
            "半年后,老师傅被通知:你带新徒弟吧",
            "他不是被 AI 替代,是被'经验'替代",
        ])

    add_h3(doc, "1.2.2 经验坍塌的 3 个征兆")
    add_body_no_indent(doc, "1. 你的经验,AI 也有")
    add_body_no_indent(doc, "2. 你的直觉,AI 数据更准")
    add_body_no_indent(doc, "3. 你的口碑,新员工不再认")

    add_h3(doc, "1.2.3 经验坍塌的 3 个应对")
    add_body_no_indent(doc, "1. 接受:经验贬值是事实,不是攻击")
    add_body_no_indent(doc, "2. 转化:把经验变成方法论,教给团队")
    add_body_no_indent(doc, "3. 升级:从'我有经验'变成'我能教经验'")

    add_h2(doc, "1.3 第三次坍塌:决策权下沉")
    add_body(doc, "过去 30 年,你定方向,员工执行。")
    add_body(doc, "现在,Z 世代自己定 KPI,自己选项目。")
    add_body(doc, "决策权消失——你不再是决策者。")

    add_h3(doc, "1.3.1 案例:95 后员工自己决策")
    add_case(doc, "互联网公司 95 后产品经理",
        [
            "老板:下周上线这个功能",
            "95 后:我做了一份用户调研,数据不支持。先改",
            "老板沉默 3 秒:行,你说了算",
            "这个团队里,老板负责拍板,员工负责'替老板做决定'",
        ])

    add_h3(doc, "1.3.2 决策权坍塌的 3 个征兆")
    add_body_no_indent(doc, "1. 员工不等你,自己先动")
    add_body_no_indent(doc, "2. 你的方案,被员工挑战")
    add_body_no_indent(doc, "3. 你的权威,被员工朋友圈超越")

    add_h3(doc, "1.3.3 决策权坍塌的 3 个应对")
    add_body_no_indent(doc, "1. 接受:决策权下沉是趋势,不是失控")
    add_body_no_indent(doc, "2. 赋能:帮员工自己会决策")
    add_body_no_indent(doc, "3. 协作:从'我做决定'变成'你做决定,我陪'")

    add_h2(doc, "1.4 三次坍塌的累积效应")
    add_table_simple(doc, ["坍塌", "过去", "现在", "新角色"],
                    [
                        ["信息坍塌", "我是信息源", "我不是唯一信息源", "整合者"],
                        ["经验坍塌", "经验=权威", "经验贬值", "方法论沉淀者"],
                        ["决策权坍塌", "我是决策者", "决策权下沉", "教练 + 陪练"],
                    ])

    add_quote(doc, "三次坍塌 = 传统管理者的 3 个支柱,全部塌了。")

    add_h2(doc, "1.5 新角色:从指挥到教练")
    add_body(doc, "指挥者:我告诉你做什么")
    add_body(doc, "教练:我帮你找到为什么做")
    add_quote(doc, "延时摄影——和时间做朋友。", "—— 短期会慢,长期会快")

    add_h2(doc, "1.6 模块一常见问题 FAQ")
    add_faq(doc, "Q1:如果我团队都不爱用 AI 怎么办?",
            "A:不是不爱用,是没安全感。先做 1V1 心谈,问'你最担心什么?'。让团队感到'用了 AI 也不会被替代',他们才会开始试。")
    add_faq(doc, "Q2:我 30 年经验,真的没用了吗?",
            "A:经验不会没用,但'凭经验拍脑袋'会没用。把经验变成方法论,变成可教的东西,你的价值会从'我懂'变成'我会教'。")
    add_faq(doc, "Q3:年轻人不听我的,怎么办?",
            "A:这是好事。说明他们有主见。你要做的是从'替他定'变成'陪他看'。让他自己定,你帮他看清代价。")
    add_faq(doc, "Q4:如果我学 AI 学不会,怎么办?",
            "A:学不会是正常的,AI 也在变。关键是不要抗拒,每周学 1 个新动作。1 年后,你就超过 80% 的人。")
    add_faq(doc, "Q5:三次坍塌,真的会同时发生吗?",
            "A:不同行业,坍塌顺序不同。但 3 年内,3 次坍塌都会到。先看清,早准备。")

    add_h2(doc, "1.7 模块一思考题")
    add_blank(doc, "回顾你最近 3 次和下属的对话,你能看到'信息坍塌'的痕迹吗?")
    add_blank(doc, "你身边有没有类似'老王'或'老师傅'的人?他们现在的状态怎样?")
    add_blank(doc, "你最近 1 次被员工挑战决策,是怎样的?你的反应是什么?")


def add_module_two(doc):
    add_h1(doc, "模块二 教练 vs 传统管理者的 6 个差异")
    add_body(doc, "你不需要变成另一个人,你只需要看清 6 个差异,每个差异往前一步。")
    add_body(doc, "本章逐个拆解 6 个差异,每个差异配 1 个真实剧本。")

    add_h2(doc, "2.1 6 个差异总览")
    add_table_simple(doc, ["维度", "传统管理者", "教练"],
                    [
                        ["权威来源", "职位给我权威", "专业+信任给我权威"],
                        ["对话姿态", "我讲你听", "我问你答"],
                        ["决策方式", "我定方向", "你定方向,我陪"],
                        ["角色定位", "我是裁判", "我是教练"],
                        ["反馈风格", "你做错了,批评", "你做到了,再加什么"],
                        ["长期目标", "本季度业绩", "5 年后他还在成长"],
                    ])

    add_quote(doc, "讨喜而不是讨好。", "—— 别只说你爱听的话")

    add_h2(doc, "2.2 差异 1:权威来源")
    add_body(doc, "传统:职位给我权威,我说了算。")
    add_body(doc, "教练:专业+信任给我权威,你愿意听。")
    add_body(doc, "职位权威 1 年失效,信任权威 10 年有效。")

    add_h3(doc, "2.2.1 真实剧本:主任 vs 教练")
    add_case(doc, "车间主任老周(传统)",
        [
            "30 年车间主任,员工叫他'周总'",
            "员工表面尊重,背后叫他'老古董'",
            "新员工来了 3 个月,一次都没主动找他聊",
            "周总觉得自己被孤立了",
        ])
    add_case(doc, "技术总监王(教练)",
        [
            "互联网公司技术总监,30 岁",
            "员工叫他'王老师'或直接叫'王哥'",
            "新员工入职第 1 天,他都会约 1 次 8 分钟心谈",
            "6 个月后,团队流失率 5%(行业平均 25%)",
        ])

    add_h2(doc, "2.3 差异 2:对话姿态")
    add_body(doc, "传统:我讲你听,我命令你执行。")
    add_body(doc, "教练:我问你答,我帮你找到答案。")
    add_case(doc, "小李的两次老板",
        [
            "传统老板:小李,这件事明天必须做",
            "教练老板:小李,这件事你怎么看?为什么重要?",
            "结果:1 年后,小李从前者离职,从后者升职",
        ])

    add_h2(doc, "2.4 差异 3:决策方式")
    add_body(doc, "传统:我定方向,你执行。")
    add_body(doc, "教练:我帮你看清,你自己定。")
    add_body(doc, "下属自己定的方案,执行度 3 倍于你定的。")
    add_case(doc, "老陈定 vs 小林自己定",
        [
            "老陈:我给你定了 1 个方案,你去做",
            "小林:我想了 3 个方案,你看哪个更靠谱?",
            "老陈的方案:小林做了 50%,延期 2 周",
            "小林的方案:小林做了 100%,还做了 2 个客户拓展",
        ])

    add_h2(doc, "2.5 差异 4:角色定位")
    add_body(doc, "传统:我是裁判,定对错。")
    add_body(doc, "教练:我是教练,陪你上场。")
    add_quote(doc, "善借力打力。", "—— 让每个人成为自己的教练")

    add_h2(doc, "2.6 差异 5:反馈风格")
    add_body(doc, "传统:你做错了,再批评。")
    add_body(doc, "教练:你做到了,再加什么更好。")
    add_body(doc, "行为反馈 vs 人格评判——后者让人崩。")
    add_table_simple(doc, ["反馈类型", "传统话术", "教练话术"],
                    [
                        ["做错时", "你怎么这么笨", "这件事复盘一下,卡点在哪?"],
                        ["做好时", "不错", "你做到了 A,再加 B 会更好"],
                        ["拖延时", "你又拖了", "我看到你最近状态不好,需要帮什么?"],
                        ["挑战时", "不可能", "你说说你的理由,我也说说我的"],
                    ])

    add_h2(doc, "2.7 差异 6:长期目标")
    add_body(doc, "传统:本季度业绩,完成就 OK。")
    add_body(doc, "教练:5 年后他还在成长,是我赢。")
    add_quote(doc, "延时摄影——和时间做朋友。")

    add_h2(doc, "2.8 教练心态:4 个核心信念")
    add_body_no_indent(doc, "1. 每个人都比自己以为的更强")
    add_body_no_indent(doc, "2. 答案在他心里,不在我嘴里")
    add_body_no_indent(doc, "3. 我不是来救他,我是来陪他")
    add_body_no_indent(doc, "4. 短期会慢,长期会快")

    add_h2(doc, "2.9 教练红线:5 个不能")
    add_body_no_indent(doc, "1. 不能替他做决定")
    add_body_no_indent(doc, "2. 不能评判他的选择")
    add_body_no_indent(doc, "3. 不能忽视他的情绪")
    add_body_no_indent(doc, "4. 不能用'我是为你好'绑架")
    add_body_no_indent(doc, "5. 不能在他没准备好时硬推")

    add_h2(doc, "2.10 模块二常见问题 FAQ")
    add_faq(doc, "Q1:如果下属不想被教练,只想被告知呢?",
            "A:先尊重他的偏好。但同时引导:'你愿意试 1 次吗?'大多数时候,试了 1 次就会尝到甜头。")
    add_faq(doc, "Q2:教练是不是太慢了,任务急的时候怎么办?",
            "A:教练是'先慢后快'。第 1 次教练对话要 30 分钟,但后面下属自己会做,你反而更省时间。")
    add_faq(doc, "Q3:如果员工自己定的方案很差,怎么办?",
            "A:让他先试。如果失败,你陪他复盘:'下次你会怎么改?'失败的方案,是最有价值的教练场景。")
    add_faq(doc, "Q4:6 个差异我都做不到,先练哪个?",
            "A:先练'对话姿态'。从'我讲你听'变成'我问你答'。1 周就能见效。")
    add_faq(doc, "Q5:教练和放羊式管理有什么区别?",
            "A:教练=我陪你,陪你看见。放羊=我不管,你爱咋咋地。教练有连接、有承诺。放羊没有。")

    add_h2(doc, "2.11 模块二思考题")
    add_blank(doc, "回想本周 1 次对话,按 6 个差异打分(1-5)。哪个差异需要重点练?")
    add_blank(doc, "你最常违反的'教练红线'是哪一个?")
    add_blank(doc, "你身边有'老陈'这样的人吗?他在你的对话里是 6 个差异中的哪个?")


def add_module_three(doc):
    add_h1(doc, "模块三 3C 信任模型 + 4 模式员工")
    add_body(doc, "看见人,才能带人。3C 是建立信任的 3 个关键动作,4 模式是识别员工的 4 种工具。")
    add_body(doc, "本章分 3 部分:3C 信任模型、4 模式员工、跨模式协作。")

    add_h2(doc, "3.1 1 个数字:信任的价值")
    add_quote(doc, "高信任团队,绩效比低信任团队高 76%。", "Paul Zak 2017 NEJM 研究")
    add_body(doc, "信任 = 团队的隐形业绩杠杆。")

    add_h2(doc, "3.2 3C 信任模型")

    add_h3(doc, "3.2.1 Connection 连接")
    add_body(doc, "看见他:名字、状态、困难。")
    add_body(doc, "记住他:上次说的话、他的家庭。")
    add_body(doc, "在他需要时,出现在场。")
    add_body(doc, "连接 = 我看见你这个人,不是你的岗位。")
    add_case(doc, "小张记得下属生日",
        [
            "技术经理,30 人团队",
            "坚持 1 年:每位下属过生日,他写 1 张卡片",
            "1 年后,团队 NPS 90(行业平均 30)",
            "员工私下说:张哥记得我,我就愿意跟他",
        ])

    add_h3(doc, "3.2.2 Commitment 承诺")
    add_body(doc, "我说的,我做到。")
    add_body(doc, "小事也兑现(请假、加薪、建议)。")
    add_body(doc, "做不到的,提前说,不甩锅。")
    add_body(doc, "承诺 = 我信你,因为你一直说到做到。")
    add_case(doc, "老李 3 次没说",
        [
            "总监,30 年工龄",
            "答应员工 3 件事,1 件都没做",
            "1 年后,3 个下属全走了",
            "不是因为他能力差,是因为他'说话不算数'",
        ])

    add_h3(doc, "3.2.3 Caring 关怀")
    add_body(doc, "关心他的状态(累不累)。")
    add_body(doc, "关心他的成长(想学什么)。")
    add_body(doc, "关心他的生活(家里怎么样)。")
    add_body(doc, "关怀 = 我在乎你这个人,不只是你的产出。")
    add_case(doc, "小陈记得下属母亲住院",
        [
            "客户经理,5 人团队",
            "小张母亲住院,他主动发 1 个红包",
            "半年后小张业绩翻倍",
            "他说:陈哥记得我妈,我就跟他",
        ])

    add_h2(doc, "3.3 4 模式员工:AI 时代的全新分类")
    add_body(doc, "AI 时代,员工分 4 种。")
    add_body(doc, "看清每个人在哪个模式,才能给到对的动作。")
    add_table_simple(doc, ["模式", "特征", "管理动作", "风险"],
                    [
                        ["先行者", "抢着用 AI,跑得快", "授权+挑战+允许试错", "骄傲,看不起别人"],
                        ["整合者", "愿意尝试,先看效果", "桥梁+跨部门协作", "执行力强但慢"],
                        ["观望者", "等别人先试,看情况", "1V1 心谈+安全尝试", "被落下,焦虑"],
                        ["保守者", "担心被替代,不敢动", "尊重节奏+稳定本职", "被替代,情绪崩"],
                    ])

    add_h3(doc, "3.3.1 先行者深度画像")
    add_case(doc, "小张(先行者)怎么带",
        [
            "28 岁,产品经理,AI 工具用得比主管还溜",
            "用 AI 做用户调研报告,1 小时搞定",
            "团队其他人都不会,他成了'AI 翻译'",
            "但他开始骄傲,说主管是'老古董'",
        ])
    add_body(doc, "管理动作清单:")
    add_body_no_indent(doc, "1. 授权:让他带 AI 项目")
    add_body_no_indent(doc, "2. 挑战:让他做更难的事")
    add_body_no_indent(doc, "3. 允许试错:失败了不批评")
    add_body_no_indent(doc, "4. 关注骄傲:每月 1V1 谈'你怎么看团队'")

    add_h3(doc, "3.3.2 整合者深度画像")
    add_case(doc, "小林(整合者)怎么带",
        [
            "32 岁,运营经理,愿意尝试 AI 工具",
            "但每次先看团队谁用得好,他再学",
            "擅长把工具变成流程",
            "但慢半拍,机会来了会犹豫",
        ])
    add_body(doc, "管理动作清单:")
    add_body_no_indent(doc, "1. 桥梁:让他做先行者和观望者之间的桥")
    add_body_no_indent(doc, "2. 跨部门:让他牵头 AI 试点项目")
    add_body_no_indent(doc, "3. 速度感:每周 1V1 谈'哪些事你可以更快'")

    add_h3(doc, "3.3.3 观望者深度画像")
    add_case(doc, "小赵(观望者)怎么带",
        [
            "29 岁,设计师,AI 工具不会用",
            "但她不说,默默等",
            "主管安排她学,她答应,但没动",
            "3 个月后,她开始焦虑,怕被淘汰",
        ])
    add_body(doc, "管理动作清单:")
    add_body_no_indent(doc, "1. 1V1 心谈:先问'你担心什么'")
    add_body_no_indent(doc, "2. 安全尝试:从低风险任务开始")
    add_body_no_indent(doc, "3. Buddy 机制:配 1 个先行者带她")
    add_body_no_indent(doc, "4. 庆祝小胜利:每做 1 个,都认可")

    add_h3(doc, "3.3.4 保守者深度画像")
    add_case(doc, "老王(保守者)怎么带",
        [
            "58 岁,30 年工龄,AI 焦虑",
            "3 个月没睡好,失眠,易怒",
            "想提前退休",
        ])
    add_body(doc, "管理动作清单:")
    add_body_no_indent(doc, "1. 尊重节奏:不强迫学 AI")
    add_body_no_indent(doc, "2. 稳定本职:先做好他擅长的事")
    add_body_no_indent(doc, "3. AI 从低风险切入:从查资料开始")
    add_body_no_indent(doc, "4. 情绪关注:1V1 问'你最近睡得怎样'")
    add_body_no_indent(doc, "5. 找新角色:让他做'AI 教练'带新人")

    add_h2(doc, "3.4 跨模式协作:1+1>2 的搭配")
    add_table_simple(doc, ["搭配", "效果", "适用场景"],
                    [
                        ["先行者 + 整合者", "创新+落地,黄金组合", "新业务、创新项目"],
                        ["观望者 + 先行者", "Buddy 机制,先看后试", "新工具推广"],
                        ["保守者 + 整合者", "稳定+信任,小步前进", "老业务转型"],
                        ["4 模式混搭", "多样化,生态丰富", "稳定团队"],
                    ])

    add_h2(doc, "3.5 3C + 4 模式综合应用")
    add_body(doc, "3C 是动作,4 模式是对象。")
    add_body(doc, "对先行者:Connection 看见他,Commitment 兑现, Caring 关注骄傲")
    add_body(doc, "对保守者:Connection 看见他, Commitment 兑现, Caring 关注情绪")
    add_body(doc, "对观望者:Connection 看见他, Commitment 兑现, Caring 关注焦虑")
    add_body(doc, "对整合者:Connection 看见他, Commitment 兑现, Caring 关注成长")

    add_h2(doc, "3.6 模块三常见问题 FAQ")
    add_faq(doc, "Q1:如果下属不是单纯的 1 种模式,怎么办?",
            "A:每个人都是 4 模式的混合体。取主类型(打分最高的),其他模式是'次要风格'。")
    add_faq(doc, "Q2:3C 是不是要花很多时间?",
            "A:不用。每天 5 分钟 Connection,1 件小事 Commitment,1 句关心 Caring。30 天后,团队气质就变了。")
    add_faq(doc, "Q3:保守者是不是该被淘汰?",
            "A:不一定。保守者稳定,是团队的压舱石。给他们稳定本职 + AI 低风险切入,大多数会慢慢跟上来。")
    add_faq(doc, "Q4:先行者骄傲,怎么管?",
            "A:1V1 心谈 + 看见他的贡献 + 让他带人。当他开始'输出'时,骄傲会变成责任。")
    add_faq(doc, "Q5:3C 和 4 模式哪个先学?",
            "A:先学 3C。3C 是底层,4 模式是工具。先有连接,后有分类。")

    add_h2(doc, "3.7 模块三练习")
    add_blank(doc, "列出你的 5 名下属,按 4 模式打分(1-5),主类型取最高分")
    add_table_simple(doc, ["姓名", "先行者", "整合者", "观望者", "保守者", "主类型"],
                    [[f"下属{i + 1}", "", "", "", "", ""] for i in range(5)])
    add_blank(doc, "你的 3C 哪个最弱?下个月怎么练?")


def add_module_four(doc):
    add_h1(doc, "模块四 GROW+ 5 阶段对话框架")
    add_body(doc, "GROW+ = 教练对话的标准动作。")
    add_body(doc, "5 阶段 25 问,覆盖一次完整对话。")
    add_body(doc, "本章逐个拆解 5 阶段,每阶段配 1 个完整剧本。")

    add_h2(doc, "4.1 传统对话 vs GROW 对话")
    add_table_simple(doc, ["传统对话", "GROW 对话"],
                    [
                        ["我讲你听,我说你做", "我问你想,我帮你看清"],
                        ["我告诉你答案", "你自己找到答案"],
                        ["10 分钟结束", "30 分钟走完"],
                        ["下属听懂了 30%", "下属执行率 90%"],
                    ])

    add_h2(doc, "4.2 G 阶段:目标 Goal")
    add_body(doc, "整个对话的灯塔。G 阶段没做对,后面都白搭。")
    add_h3(doc, "G 阶段 5 个关键问题")
    for q in [
        "1. 你最想达成的是什么?",
        "2. 如果成功了,你会看到什么?",
        "3. 这个目标对你为什么重要?",
        "4. 你愿意为它付出什么?",
        "5. 什么时候必须达成?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "G 阶段常见错误:")
    add_body_no_indent(doc, "1. 跳过 G,直接给方案")
    add_body_no_indent(doc, "2. G 阶段问得太浅,只问'想做什么'")
    add_body_no_indent(doc, "3. G 阶段问得太深,变成'为什么活着'")
    add_body(doc, "G 阶段关键:问'为什么重要',让目标内化。")

    add_h2(doc, "4.3 R 阶段:现实 Reality")
    add_body(doc, "看清脚下,才能走对方向。R 阶段关键:从'我没机会'变成'我能做什么'。")
    add_h3(doc, "R 阶段 5 个关键问题")
    for q in [
        "1. 现在的情况是什么?",
        "2. 你做了哪些尝试?效果如何?",
        "3. 卡在哪里?",
        "4. 哪些是你能控制的,哪些不能?",
        "5. 如果不做任何改变,3 个月后会怎样?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "R 阶段常见错误:")
    add_body_no_indent(doc, "1. R 阶段变成'诉苦会',没聚焦")
    add_body_no_indent(doc, "2. R 阶段不给具体数据,只说'不太好'")
    add_body_no_indent(doc, "3. R 阶段急着给建议")
    add_body(doc, "R 阶段关键:让对方先看见自己的位置。")

    add_h2(doc, "4.4 O 阶段:选择 Options")
    add_body(doc, "让他自己看见 N 种可能。O 阶段关键:激发他'做不到'的假设。")
    add_h3(doc, "O 阶段 5 个关键问题")
    for q in [
        "1. 如果不受限,你会怎么做?",
        "2. 你身边谁做得最好?你能学他什么?",
        "3. 还有哪些方法你没试过?",
        "4. 哪个方法最让你兴奋?",
        "5. 如果只能选一个,你先做哪个?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "O 阶段常见错误:")
    add_body_no_indent(doc, "1. O 阶段只列 1-2 个方案")
    add_body_no_indent(doc, "2. O 阶段让主管的方案压过下属的")
    add_body_no_indent(doc, "3. O 阶段没追问'为什么'")
    add_body(doc, "O 阶段关键:列 5 个以上方案,每个都问'为什么这个好'。")

    add_h2(doc, "4.5 W 阶段:意愿 Will")
    add_body(doc, "激活内在动力。等他说'我愿意',不是你说。")
    add_h3(doc, "W 阶段 5 个关键问题")
    for q in [
        "1. 你最想做哪个?为什么?",
        "2. 做到后对你意味着什么?",
        "3. 你愿意为它放弃什么?",
        "4. 如果不做,你会有什么感觉?",
        "5. 你对自己的承诺是什么?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "W 阶段常见错误:")
    add_body_no_indent(doc, "1. W 阶段问'你能做到吗?'变成压力")
    add_body_no_indent(doc, "2. W 阶段跳过,直接进 A")
    add_body_no_indent(doc, "3. W 阶段主管说'我相信你',代替他自己的承诺")
    add_body(doc, "W 阶段关键:让他自己说'我愿意'。")

    add_h2(doc, "4.6 A 阶段:行动 Action")
    add_body(doc, "把对话变成真东西。A 阶段关键:具体的、可衡量的、有时限的。")
    add_h3(doc, "A 阶段 5 个关键问题")
    for q in [
        "1. 你第一步做什么?",
        "2. 什么时候做?",
        "3. 你需要的支持是什么?",
        "4. 我(主管)能帮你什么?",
        "5. 我们怎么知道你做到了?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "A 阶段常见错误:")
    add_body_no_indent(doc, "1. A 阶段太模糊:'我努力做'")
    add_body_no_indent(doc, "2. A 阶段没有时间:'尽快'")
    add_body_no_indent(doc, "3. A 阶段没有衡量:'做好了就行'")
    add_body(doc, "A 阶段关键:具体到今天/本周/这个月。")

    add_h2(doc, "4.7 GROW+ 完整剧本:小李延期")
    add_body(doc, "小李的项目延期 1 周,主管想骂他。")
    add_table_simple(doc, ["阶段", "问话", "小李的回答"],
                    [
                        ["G", "小李,你最想达成什么?", "这个项目准时交付"],
                        ["R", "现在延期了,情况怎样?", "技术上卡 1 个 bug,3 天没解"],
                        ["O", "如果不延期 1 周,你想怎么做?", "找人帮 debug,或换 1 个方案"],
                        ["W", "你愿意做哪个?", "找人帮 debug,我有人选"],
                        ["A", "第一步是什么?什么时候?", "今天下午找张工,2 天内解决"],
                    ])
    add_body(doc, "结果:小李主动加班赶回进度,1 周准时交付。")
    add_body(doc, "主管学到的 3 件事:")
    add_body_no_indent(doc, "1. 别骂,问")
    add_body_no_indent(doc, "2. 别告诉答案,让他自己找到")
    add_body_no_indent(doc, "3. 别替代他做,让他自己承诺")

    add_h2(doc, "4.8 GROW+ 第 2 个剧本:老张想转岗")
    add_body(doc, "老张 55 岁,工程师,3 年后退休,想转岗做培训。")
    add_table_simple(doc, ["阶段", "问话", "老张的回答"],
                    [
                        ["G", "你为什么想转岗?", "我想做培训,1 对 1 带新人"],
                        ["R", "现在你做这件事的卡点是什么?", "没有培训经验,不知道从哪开始"],
                        ["O", "如果不受限,你会怎么做?", "先旁听 1 个培训课,再带 1 个新员工"],
                        ["W", "你愿意做哪个?", "我先旁听,2 周后申请带新员工"],
                        ["A", "第一步是什么?什么时候?", "下周一找培训经理,谈旁听"],
                    ])
    add_body(doc, "结果:老张 3 个月后开始带新人,找到意义感。")

    add_h2(doc, "4.9 GROW+ 第 3 个剧本:小赵创新被否")
    add_body(doc, "小赵推 AI 项目被否 3 次,想放弃。")
    add_table_simple(doc, ["阶段", "问话", "小赵的回答"],
                    [
                        ["G", "你为什么推这个项目?", "我相信它能帮团队提效 30%"],
                        ["R", "前 3 次被否,卡在哪?", "我没给具体数据,只说'有前景'"],
                        ["O", "如果再推 1 次,你会怎么做?", "做 1 个试点,跑 2 周,给数据"],
                        ["W", "你愿意做吗?", "愿意,我想再试 1 次"],
                        ["A", "第一步是什么?", "这周选 1 个团队试点"],
                    ])
    add_body(doc, "结果:小赵重启提案,第 4 次通过。")

    add_h2(doc, "4.10 模块四常见问题 FAQ")
    add_faq(doc, "Q1:GROW+ 对话要多长时间?",
            "A:30-60 分钟。第 1 次可能 60 分钟,熟练后 30 分钟。")
    add_faq(doc, "Q2:GROW+ 5 阶段顺序能变吗?",
            "A:不能。G 必须在第 1 位,否则后面都乱。但每个阶段可以来回走。")
    add_faq(doc, "Q3:GROW+ 可以用在 1V1 吗?",
            "A:可以。每周 1V1 走 1 遍 GROW+,团队会很有方向。")
    add_faq(doc, "Q4:GROW+ 是不是太结构化了?",
            "A:结构化是手段,目的是让下属自己想清楚。可以灵活变,但 5 阶段不能少。")
    add_faq(doc, "Q5:GROW+ 对话,主管要做笔记吗?",
            "A:建议做。记下关键句子,后面 1V1 用得上。")
    add_faq(doc, "Q6:如果下属在 G 阶段就卡了,怎么办?",
            "A:问 1 个更深的问题:'如果成功了,你会看到什么?'让他看见目标的样子。")

    add_h2(doc, "4.11 模块四练习")
    add_blank(doc, "找一个下属,用 GROW+ 完整走一遍,记下每个阶段他说的话")
    add_blank(doc, "下次 1V1 之前,把 GROW+ 25 问背 3 遍,熟练到能脱口而出")


def add_module_five(doc):
    add_h1(doc, "模块五 七层倾听阶梯")
    add_body(doc, "从'我听见了'到'你被听见了'。")
    add_body(doc, "大多数管理者卡在第 3 层,我们要到第 5 层以上。")
    add_body(doc, "本章拆解 7 层倾听,3 个核心动作,8 分钟心谈操作指南。")

    add_h2(doc, "5.1 七层阶梯")
    add_table_simple(doc, ["层", "名称", "描述", "行为"],
                    [
                        ["7", "生成性倾听", "让对方找到他自己都没意识到的答案", "沉默 30 秒,等他想"],
                        ["6", "反思性倾听", "把他说的话翻译成深层含义", "'听上去,你感到...对吗?'"],
                        ["5", "共鸣性倾听", "分享我类似的感受,让他不孤单", "'我也遇到过类似的...'"],
                        ["4", "同理心倾听", "听到情绪,命名情绪,让他被看见", "'你感到...对吗?'"],
                        ["3", "理解性倾听", "听他说完,复述确认(大多数人在这)", "'我听你说...是这样吗?'"],
                        ["2", "选择性倾听", "只听我感兴趣的,过滤其他的", "立刻想到自己要说什么"],
                        ["1", "忽视", "心不在焉,人在心不在", "看手机,打断,转移话题"],
                    ])

    add_h2(doc, "5.2 命名情绪:5 个万能句式")
    for q in [
        "1. 你看起来感到___对吗?",
        "2. 听上去,你现在___对吗?",
        "3. 如果是我,我会感到___。你呢?",
        "4. 你愿意多说说这个___吗?",
        "5. 你希望我___吗?",
    ]:
        add_body_no_indent(doc, q)
    add_body(doc, "30 个常用情绪词:")
    add_table_simple(doc, ["积极情绪", "中性情绪", "消极情绪"],
                    [
                        ["开心、兴奋、满足、感激", "平静、好奇、若有所思", "焦虑、失落、委屈、愤怒"],
                        ["成就、被认可、温暖、连接", "迷茫、犹豫、期待", "疲惫、孤独、无力、失望"],
                        ["兴奋、惊喜、热爱、自豪", "担心、困惑、谨慎", "崩溃、绝望、嫉妒、恐惧"],
                    ])

    add_h2(doc, "5.3 8 分钟心谈:操作指南")
    for s in [
        "1. 找一个安静的地方(不被打扰)",
        "2. 说:'我有 8 分钟,你愿意聊聊吗?'",
        "3. 前 4 分钟:只听,不评判,不建议",
        "4. 后 4 分钟:问 1-2 个开放问题",
        "5. 结束:问'你还需要什么吗?'",
    ]:
        add_body_no_indent(doc, s)
    add_quote(doc, "8 分钟 = 不长不短,刚好够听见他。")

    add_h2(doc, "5.4 8 分钟心谈:5 个典型错误")
    for s in [
        "1. 边听边看手机",
        "2. 立刻给建议",
        "3. 打断他说话",
        "4. '我以前也是',变成我讲",
        "5. 急着解决他的问题",
    ]:
        add_body_no_indent(doc, s)
    add_body(doc, "正解:5 个错误 + 1 个'对的做法'")
    add_table_simple(doc, ["错误", "对的做法"],
                    [
                        ["看手机", "手机放包里,8 分钟不看"],
                        ["给建议", "前 4 分钟只听,后 4 分钟问 1-2 个问题"],
                        ["打断", "等他说完,等 3 秒,再说"],
                        ["变我讲", "忍住,让他说完"],
                        ["急着解决", "问'你想要我帮你什么?'"],
                    ])

    add_h2(doc, "5.5 反思倾听:翻译术")
    add_body(doc, "翻译 = 把话背后的意思说出来,让他'被深度理解'。")
    add_table_simple(doc, ["他说", "你翻译", "他被看见什么"],
                    [
                        ["老板没看见我的方案", "你感到被忽视,所以失落", "被看见:他的失落"],
                        ["我不想做这个项目了", "这个项目让你失去了意义感", "被看见:他需要意义感"],
                        ["同事不配合我", "你感到孤立,需要支持", "被看见:他需要连接"],
                    ])

    add_h2(doc, "5.6 共鸣倾听:我也有过")
    add_body(doc, "共鸣 = 我也有过类似的感受,让他不孤单。")
    add_table_simple(doc, ["他说的", "你的共鸣"],
                    [
                        ["我加班加到崩溃", "我 5 年前也有过这种感受,后来我...后来我才..."],
                        ["我和老板说不通", "我以前也遇到过说不通的老板,后来我学到...后来我..."],
                        ["团队都不配合", "我以前带团队也遇到过,后来我学到...后来我..."],
                    ])

    add_h2(doc, "5.7 生成倾听:沉默的力量")
    for s in [
        "1. 他说完,你不说话,等他",
        "2. 他沉默 10 秒,你也不说话",
        "3. 这时候,他会'想'出更深的东西",
        "4. 30 秒后,他可能说出他自己都没意识到的话",
        "5. 你只问:'你刚才在沉默里,在想什么?'",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "5.8 7 层倾听自评表")
    add_body(doc, "回想最近 5 次对话,你在哪一层?")
    add_table_simple(doc, ["对话", "层数", "关键证据"],
                    [[f"对话{i + 1}", "", ""] for i in range(5)])

    add_h2(doc, "5.9 模块五常见问题 FAQ")
    add_faq(doc, "Q1:如果我没什么情绪经验,怎么听?",
            "A:先学命名情绪。每天 1 个新词,1 个月后你就掌握 30 个。")
    add_faq(doc, "Q2:如果下属在第 1 层(忽视),我能在第 7 层听吗?",
            "A:不能。倾听是双向的。下属不打开,你再会听也没用。先用 3C 让他信任你。")
    add_faq(doc, "Q3:8 分钟太短,能不能 30 分钟?",
            "A:可以,但建议从 8 分钟开始。先做到 8 分钟,再延到 15 分钟、30 分钟。")
    add_faq(doc, "Q4:沉默时,我不说话,会不会冷场?",
            "A:不会。沉默是给对方'想'的时间。30 秒后,他会说更深的话。")
    add_faq(doc, "Q5:我能用'我以前也...'来共鸣吗?",
            "A:可以,但 1 句话就够了。说多了,就变成'我讲'了。")
    add_faq(doc, "Q6:如果我听了,下属没反应,是不是我听错了?",
            "A:不是。也许他还没准备好。第 2 次听,他可能就打开了。")

    add_h2(doc, "5.10 模块五练习")
    add_blank(doc, "找 1 个下属,做 1 次 8 分钟心谈,记下他的反应")
    add_blank(doc, "你的倾听自评在哪一层?下个月怎么升 1 层?")


def add_module_six(doc):
    add_h1(doc, "模块六 30 天落地实践")
    add_body(doc, "把方法变成肌肉记忆。")
    add_body(doc, "30 天,5 个维度,4 周。")
    add_body(doc, "本章给出 30 天详细计划。")

    add_h2(doc, "6.1 30 天 5 维度")
    add_table_simple(doc, ["维度", "内容", "频次"],
                    [
                        ["心谈", "每天 8 分钟心谈 1 次", "5 次/周"],
                        ["GROW+", "每天用 GROW+ 解决 1 个小问题", "3 次/周"],
                        ["3C", "每天 1 个 3C 动作", "7 次/周"],
                        ["团队分享", "每周 1 次内部分享", "1 次/周"],
                        ["反思", "每天写 1 段反思日记", "7 次/周"],
                    ])

    add_h2(doc, "6.2 第 1 周:心谈")
    for s in [
        "Day 1: 找一个下属,做 8 分钟心谈",
        "Day 2: 找另一个下属,同样",
        "Day 3-5: 每天 1 个,共 5 个下属",
        "Day 6: 复盘 - 哪个对话印象最深?",
        "Day 7: 把印象最深的写进反思日记",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "6.3 第 2 周:GROW+")
    for s in [
        "Day 8: 选 1 个下属,用 GROW+ 聊 20 分钟",
        "Day 9-10: 每天 1 次,共 3 个下属",
        "Day 11: 找最难的 1 个下属(观望者/保守者)",
        "Day 12: 找最熟的 1 个下属(先行者/整合者)",
        "Day 13: 复盘 + 写反思日记",
        "Day 14: 整理剧本,准备内部分享",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "6.4 第 3 周:3C")
    for s in [
        "Day 15: Connection - 记住下属生日",
        "Day 16: Commitment - 兑现 1 个小事",
        "Day 17: Caring - 关心 1 个下属的私事",
        "Day 18: Connection - 听下属说一件私事",
        "Day 19: Commitment - 提前说 1 件做不到的事",
        "Day 20: Caring - 给 1 个下属写卡片",
        "Day 21: 复盘 + 3C 评估",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "6.5 第 4 周:团队分享")
    for s in [
        "Day 22-24: 把 30 天实践整理成 1 页纸",
        "Day 25: 找 1 个下属做内部分享",
        "Day 26: 在周会上分享 1 个 GROW+ 剧本",
        "Day 27: 收集 3 个下属的反馈",
        "Day 28: 写完 30 天反思日记",
        "Day 29: 做 F1 训后评估",
        "Day 30: 30 天庆祝,加入校友群",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "6.6 30 天共学群规则")
    for s in [
        "1. 每天 1 个动作打卡(心谈/GROW+/3C/分享/反思 任选)",
        "2. 每周 1 次 8 分钟心谈实战复盘",
        "3. 每周 1 次团队分享",
        "4. 每周 1 个 30 天反思日记片段",
        "5. 30 天后全群庆祝 + 内部分享",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "6.7 30 天常见卡点 + 应对")
    add_table_simple(doc, ["卡点", "应对"],
                    [
                        ["没时间", "5 分钟也做。'不完美地开始',比'完美地不做'好"],
                        ["下属不愿意", "1V1 心谈 + 问'你担心什么'"],
                        ["做 1 周没效果", "GROW 是慢功夫,1 周不够,1 个月见"],
                        ["做不下去", "找 1 个 Buddy 互相打卡"],
                    ])

    add_h2(doc, "6.8 模块六常见问题 FAQ")
    add_faq(doc, "Q1:30 天一定要做完吗?",
            "A:建议做完。30 天不连续,效果会打 5 折。")
    add_faq(doc, "Q2:共学群是必须的吗?",
            "A:建议加入。1 个人走 30 天很孤独,1 群人走 30 天很容易。")
    add_faq(doc, "Q3:30 天后,还要继续吗?",
            "A:建议继续。但 30 天后,从'每天打卡'变成'每周打卡'。")
    add_faq(doc, "Q4:如果 30 天后,下属没变化,是不是我白做了?",
            "A:不是。下属变化是慢的。3 个月后,你会看到效果。")
    add_faq(doc, "Q5:30 天后,我能教别人吗?",
            "A:能。我们会邀请优秀的学员做'30 天校友分享'。")

    add_h2(doc, "6.9 模块六思考题")
    add_blank(doc, "你的 30 天,从哪天开始?写下具体日期:")
    add_blank(doc, "你 30 天后,希望看到什么变化?")
    add_blank(doc, "你 30 天的 Buddy 是谁?")


def add_tools(doc):
    add_h1(doc, "第二部分 19 个工具使用说明")
    add_body(doc, "19 个工具覆盖训前、训中、训后全流程。每个工具都是可操作的练习,不是空话。")

    add_h2(doc, "F1 预评量表")
    add_body(doc, "用途:训前训后对比,看 5 维度变化。")
    add_body(doc, "5 维度:自我认知、沟通能力、领导风格、信任度、教练技能。")
    add_body(doc, "使用:30 题 × 5 维度,1-5 分。")
    add_body(doc, "场景:训前 1 周填写,训后立即再填 1 次。")

    add_h2(doc, "F2 时间账单")
    add_body(doc, "用途:看清时间怎么用。")
    add_body(doc, "使用:12 小时 × 30 分钟,记录每个 30 分钟做了什么。")
    add_body(doc, "场景:第 1 天结束前填写,看清'刷手机多少小时'。")

    add_h2(doc, "F3 4 模式自评")
    add_body(doc, "用途:团队成员分类,管理动作分人。")
    add_body(doc, "使用:5 个下属,4 模式打分。")
    add_body(doc, "场景:模块三练习 2 用。")

    add_h2(doc, "F4 3C 信任度评估")
    add_body(doc, "用途:打分 1-10,找最弱维度。")
    add_body(doc, "使用:Connection、Commitment、Caring 3 项,各 1-10 分。")
    add_body(doc, "场景:模块三练习 3 用。")

    add_h2(doc, "F5 GROW+ 剧本卡")
    add_body(doc, "用途:5 阶段 + 25 问,实战对话。")
    add_body(doc, "使用:对话时拿在手上,边看边问。")
    add_body(doc, "场景:模块四练习用。")

    add_h2(doc, "F6 倾听层级表")
    add_body(doc, "用途:7 层 × 自评,看自己在第几层。")
    add_body(doc, "使用:回想 5 次对话,给每层打分。")
    add_body(doc, "场景:模块五自评用。")

    add_h2(doc, "F7 8 分钟心谈卡")
    add_body(doc, "用途:5 步操作 + 8 分钟计时器。")
    add_body(doc, "使用:8 分钟计时 + 5 步操作。")
    add_body(doc, "场景:每天 1 次,30 天打卡。")

    add_h2(doc, "F8 情绪命名卡")
    add_body(doc, "用途:30 个情绪词 + 5 个万能句式。")
    add_body(doc, "使用:听下属说话时,问'你感到___对吗?'")
    add_body(doc, "场景:8 分钟心谈 + 模块五练习用。")

    add_h2(doc, "F9 5 要素卡")
    add_body(doc, "用途:5 要素拆解行动。")
    add_body(doc, "使用:目标/意义/行动/卡点/衡量。")
    add_body(doc, "场景:模块六 30 天计划。")

    add_h2(doc, "F10 3C 每日动作")
    add_body(doc, "用途:连接/承诺/关怀 3 类 9 个动作。")
    add_body(doc, "使用:每天 1 个,30 天轮一遍。")
    add_body(doc, "场景:30 天第 3 周。")

    add_h2(doc, "F11 团队氛围雷达")
    add_body(doc, "用途:7 维度 × 12 月追踪。")
    add_body(doc, "使用:每月 1 次,30 个下属打分。")
    add_body(doc, "场景:训后 90 天调研。")

    add_h2(doc, "F12 GROW+ 观察卡")
    add_body(doc, "用途:教练 12 项观察。")
    add_body(doc, "使用:对话后给自己打分。")
    add_body(doc, "场景:GROW+ 对话后立即填。")

    add_h2(doc, "F13 角色扮演脚本")
    add_body(doc, "用途:8 个真实剧本。")
    add_body(doc, "使用:2 人一组,1 个主管 1 个下属。")
    add_body(doc, "场景:模块四练习用。")

    add_h2(doc, "F14 GROW+ 情景卡")
    add_body(doc, "用途:8 张卡牌(扑克牌尺寸)。")
    add_body(doc, "使用:抽 1 张,做 1 次 GROW+。")
    add_body(doc, "场景:30 天练习用。")

    add_h2(doc, "F15 教练观察卡")
    add_body(doc, "用途:12 项观察 + GROW+ 阶段。")
    add_body(doc, "使用:每阶段 3 项观察。")
    add_body(doc, "场景:GROW+ 对话打分用。")

    add_h2(doc, "F16 4 模式分布图")
    add_body(doc, "用途:4 模式人数 + 占比饼图。")
    add_body(doc, "使用:团队 5-10 人,统计 4 模式占比。")
    add_body(doc, "场景:模块三练习后汇总。")

    add_h2(doc, "F17 失败复盘模板")
    add_body(doc, "用途:3 种没效果处理。")
    add_body(doc, "使用:3 种卡点(没时间/没意愿/没效果)。")
    add_body(doc, "场景:30 天遇到卡点时用。")

    add_h2(doc, "F18 30 天打卡表")
    add_body(doc, "用途:5 维度 × 30 天追踪。")
    add_body(doc, "使用:每天 1 行,5 维度勾选。")
    add_body(doc, "场景:30 天主打卡表。")

    add_h2(doc, "F19 行动计划")
    add_body(doc, "用途:30 天 × 5 要素。")
    add_body(doc, "使用:30 天共学群第 1 天填,30 天后回看。")
    add_body(doc, "场景:30 天主输出。")

    add_h2(doc, "19 个工具汇总表")
    add_table_simple(doc, ["编号", "工具名", "用途", "使用场景"],
                    [
                        ["F1", "预评量表", "30 题 × 5 维度,训前训后对比", "训前 1 周 + 训后立即"],
                        ["F2", "时间账单", "12 小时 × 30 分钟,看清时间", "Day 1 结束前"],
                        ["F3", "4 模式自评", "团队成员分类,管理动作分人", "模块三练习"],
                        ["F4", "3C 信任度评估", "打分 1-10,找最弱维度", "模块三练习"],
                        ["F5", "GROW+ 剧本卡", "5 阶段 + 25 问,实战对话", "模块四练习"],
                        ["F6", "倾听层级表", "7 层 × 自评,看自己在第几层", "模块五自评"],
                        ["F7", "8 分钟心谈卡", "5 步操作 + 8 分钟计时", "每天 1 次"],
                        ["F8", "情绪命名卡", "30 个情绪词 + 5 句式", "8 分钟心谈"],
                        ["F9", "5 要素卡", "目标/意义/行动/卡点/衡量", "30 天计划"],
                        ["F10", "3C 每日动作", "3 类 9 个动作", "30 天第 3 周"],
                        ["F11", "团队氛围雷达", "7 维度 × 12 月追踪", "训后 90 天"],
                        ["F12", "GROW+ 观察卡", "教练 12 项观察", "GROW+ 后打分"],
                        ["F13", "角色扮演脚本", "8 个真实剧本", "模块四练习"],
                        ["F14", "GROW+ 情景卡", "8 张卡牌", "30 天练习"],
                        ["F15", "教练观察卡", "12 项观察 + GROW+ 阶段", "GROW+ 打分"],
                        ["F16", "4 模式分布图", "4 模式人数 + 占比", "模块三汇总"],
                        ["F17", "失败复盘模板", "3 种没效果处理", "30 天卡点"],
                        ["F18", "30 天打卡表", "5 维度 × 30 天追踪", "30 天主表"],
                        ["F19", "行动计划", "30 天 × 5 要素", "30 天主输出"],
                    ])


def add_cases(doc):
    add_h1(doc, "第三部分 真实案例集(8 例)")
    add_body(doc, "8 个真实案例,覆盖先行者/整合者/观望者/保守者 4 类员工。")
    add_body(doc, "每个案例包含:背景、关键动作、结果、3 个学到的。")

    cases = [
        ("案例 1:小李延期", "90 后项目骨干", [
            "小李,90 后,项目骨干,做事快。",
            "老板信任他,什么都给他。他咬牙坚持 1 个月,提了离职。",
            "GROW+ 解决方案:G 找到他真正想要的,R 厘清现状,O 拓宽可能,W 激活意愿,A 行动。",
            "结果:小李主动加班赶回进度,1 周准时交付。",
            "3 个学到的:别骂,问;别告诉答案,让他自己找;别替代他做,让他自己承诺。",
        ]),
        ("案例 2:老王转型", "60 后车间主任 30 年工龄", [
            "老王,60 后车间主任,30 年工龄,担心被 AI 替代。",
            "GROW+ 解决方案:不评判他的焦虑,问'你愿意为它做什么?'",
            "结果:老王主动申请做'AI 教练',带新员工。",
            "3 个学到的:对保守者先稳定本职,AI 低风险切入;找新角色比硬学新工具好;情绪关注比技能训练优先。",
        ]),
        ("案例 3:小张迷茫", "工作 3 年不知道往哪走", [
            "小张,工作 3 年,不知道往哪走。",
            "GROW+ 解决方案:问'3 年后你想成为什么?'",
            "结果:小张选了产品方向,3 个月后转岗成功。",
            "3 个学到的:迷茫的人缺的不是答案,是看见;GROW 的 G 阶段是给迷茫的人最好的礼物;不评判,只问。",
        ]),
        ("案例 4:老李瓶颈", "做了 5 年总监升不上去", [
            "老李,做了 5 年总监,升不上去,想跳不动。",
            "GROW+ 解决方案:问'如果不受限,你会怎么做?'",
            "结果:老李找到新赛道,开了副业。",
            "3 个学到的:升不上去的人,不是不行,是赛道没选对;问'如果不受限'能打开新空间;不替员工判断。",
        ]),
        ("案例 5:小陈加班", "95 后连续 1 个月 996", [
            "小陈,95 后,连续 1 个月 996,在茶水间哭。",
            "8 分钟心谈 + 命名情绪:'你感到撑不住了对吗?'",
            "结果:小陈说出来了,主管才知道他的极限。",
            "3 个学到的:情绪不命名,下属撑到崩溃;8 分钟心谈能救人;主管不知道的事,比知道的更多。",
        ]),
        ("案例 6:老张 AI 焦虑", "55 岁工程师被 AI 替代", [
            "老张,55 岁,工程师,被 AI 替代了一半工作。",
            "GROW+ 解决方案:不评判他的失落,问'你最想达成什么?'",
            "结果:老张开始学 AI 工具,3 个月后成了 AI 培训师。",
            "3 个学到的:被 AI 替代的人,不是没用,是没找到新角色;GROW+ 慢,但救人;55 岁也能学新东西。",
        ]),
        ("案例 7:小赵创新", "推 AI 项目被否 3 次", [
            "小赵,推 AI 项目被否 3 次,想放弃。",
            "GROW+ 解决方案:R 阶段问'你做了哪些尝试?'让他看见自己已经做了 3 次。",
            "结果:小赵重启提案,加上数据,第 4 次通过。",
            "3 个学到的:被否多次的下属,不是没努力,是看不见自己;GROW+ R 阶段能让人看见自己;放弃前再给 1 次机会。",
        ]),
        ("案例 8:老周临退休", "58 岁 5 年后退休", [
            "老周,58 岁,5 年后退休,突然空虚。",
            "8 分钟心谈 + 翻译:'听起来,你想留下点什么'",
            "结果:老周申请做导师,带新人,找到意义感。",
            "3 个学到的:临退休的人,缺的不是钱,是意义;8 分钟心谈 + 翻译,胜过 10 次建议;老员工是宝藏。",
        ]),
    ]
    for title, _, body in cases:
        add_h2(doc, title)
        for line in body:
            add_body_no_indent(doc, line)

    add_h2(doc, "8 个案例的共同模式")
    for s in [
        "1. 全部用 GROW+ 或 8 分钟心谈",
        "2. 全部先'问',后'说'",
        "3. 全部看到下属自己找到答案",
        "4. 全部强调'不是替他做,是陪他看清'",
        "5. 全部 30 天内出结果",
        "6. 全部不需要 1 次做对,需要多次迭代",
        "7. 全部从'不评判'开始",
        "8. 全部是真实的、有名字的、具体的",
    ]:
        add_body_no_indent(doc, s)


def add_action_30d(doc):
    add_h1(doc, "第四部分 30 天行动手册")
    add_body(doc, "把 30 天写下来,贴在桌上。每天 5 分钟打卡,30 天后看变化。")

    add_h2(doc, "F19 行动计划模板")
    add_table_simple(doc, ["要素", "你的填写"],
                    [
                        ["1. 我要达成什么(目标)", ""],
                        ["2. 我为什么在乎(意义)", ""],
                        ["3. 我第一步做什么(行动)", ""],
                        ["4. 我可能遇到什么卡点(障碍)", ""],
                        ["5. 我怎么知道自己做到了(衡量)", ""],
                    ])

    add_h2(doc, "30 天打卡表")
    for week in range(1, 5):
        add_h3(doc, f"第 {week} 周")
        add_table_simple(doc, ["Day", "心谈", "GROW+", "3C", "分享", "反思"],
                        [[f"Day {d}", "", "", "", "", ""] for d in range((week - 1) * 7 + 1, week * 7 + 1)])

    add_h2(doc, "30 天共学群打卡格式")
    add_body(doc, "Day 1 心谈 - 和小张")
    add_body(doc, "学到:他最近焦虑被 AI 替代")
    add_body(doc, "下一步:帮他做 1V1,问'你愿意为它做什么'")
    add_body(doc, "")
    add_body(doc, "Day 5 反思 - 这一周印象最深的事")
    add_body(doc, "学到:3C 不难,难的是持续做")
    add_body(doc, "下一步:把 Connection 加入每天 1V1")

    add_h2(doc, "30 天统计模板")
    add_table_simple(doc, ["维度", "完成次数", "占计划%"],
                    [
                        ["心谈", "", ""],
                        ["GROW+", "", ""],
                        ["3C", "", ""],
                        ["团队分享", "", ""],
                        ["反思日记", "", ""],
                    ])


def add_journal(doc):
    add_h1(doc, "第五部分 反思日记模板")
    add_body(doc, "反思 = 把经验变成能力。每天 5 分钟,问自己 5 个问题。")
    add_body(doc, "本章给出 5 个反思模板,30 天每天 1 个。")

    questions = [
        "今天我做了什么?为什么这样做?",
        "今天的对话中,哪个瞬间让我印象最深?",
        "我下次会做哪个不一样的动作?",
        "我今天最大的卡点是什么?",
        "我今天学到了什么?",
    ]

    for d in [1, 5, 8, 12, 15, 19, 22, 26, 29]:
        add_h2(doc, f"Day {d} 反思")
        for q in questions:
            add_body(doc, q, indent=False)
        add_blank(doc, "写下你的反思")

    add_h2(doc, "30 天反思日记合集")
    for week in range(1, 5):
        add_h3(doc, f"第 {week} 周反思汇总")
        add_blank(doc, f"第 {week} 周印象最深的 3 件事")
        add_blank(doc, f"第 {week} 周最大的卡点")
        add_blank(doc, f"第 {week} 周学到的 3 件事")
        add_blank(doc, f"第 {week} 周下周想做的 3 件事")


def add_faq_section(doc):
    add_h1(doc, "第六部分 常见问题 FAQ")
    add_body(doc, "30 个常见问题,覆盖 6 个模块。")

    add_h2(doc, "模块一 FAQ(5 问)")
    add_faq(doc, "Q1:如果我团队都不爱用 AI 怎么办?",
            "A:不是不爱用,是没安全感。先做 1V1 心谈,问'你最担心什么?'。让团队感到'用了 AI 也不会被替代',他们才会开始试。")
    add_faq(doc, "Q2:我 30 年经验,真的没用了吗?",
            "A:经验不会没用,但'凭经验拍脑袋'会没用。把经验变成方法论,变成可教的东西,你的价值会从'我懂'变成'我会教'。")
    add_faq(doc, "Q3:年轻人不听我的,怎么办?",
            "A:这是好事。说明他们有主见。你要做的是从'替他定'变成'陪他看'。让他自己定,你帮他看清代价。")
    add_faq(doc, "Q4:如果我学 AI 学不会,怎么办?",
            "A:学不会是正常的,AI 也在变。关键是不要抗拒,每周学 1 个新动作。1 年后,你就超过 80% 的人。")
    add_faq(doc, "Q5:三次坍塌,真的会同时发生吗?",
            "A:不同行业,坍塌顺序不同。但 3 年内,3 次坍塌都会到。先看清,早准备。")

    add_h2(doc, "模块二 FAQ(5 问)")
    add_faq(doc, "Q1:如果下属不想被教练,只想被告知呢?",
            "A:先尊重他的偏好。但同时引导:'你愿意试 1 次吗?'大多数时候,试了 1 次就会尝到甜头。")
    add_faq(doc, "Q2:教练是不是太慢了,任务急的时候怎么办?",
            "A:教练是'先慢后快'。第 1 次教练对话要 30 分钟,但后面下属自己会做,你反而更省时间。")
    add_faq(doc, "Q3:如果员工自己定的方案很差,怎么办?",
            "A:让他先试。如果失败,你陪他复盘:'下次你会怎么改?'失败的方案,是最有价值的教练场景。")
    add_faq(doc, "Q4:6 个差异我都做不到,先练哪个?",
            "A:先练'对话姿态'。从'我讲你听'变成'我问你答'。1 周就能见效。")
    add_faq(doc, "Q5:教练和放羊式管理有什么区别?",
            "A:教练=我陪你,陪你看见。放羊=我不管,你爱咋咋地。教练有连接、有承诺。放羊没有。")

    add_h2(doc, "模块三 FAQ(5 问)")
    add_faq(doc, "Q1:如果下属不是单纯的 1 种模式,怎么办?",
            "A:每个人都是 4 模式的混合体。取主类型(打分最高的),其他模式是'次要风格'。")
    add_faq(doc, "Q2:3C 是不是要花很多时间?",
            "A:不用。每天 5 分钟 Connection,1 件小事 Commitment,1 句关心 Caring。30 天后,团队气质就变了。")
    add_faq(doc, "Q3:保守者是不是该被淘汰?",
            "A:不一定。保守者稳定,是团队的压舱石。给他们稳定本职 + AI 低风险切入,大多数会慢慢跟上来。")
    add_faq(doc, "Q4:先行者骄傲,怎么管?",
            "A:1V1 心谈 + 看见他的贡献 + 让他带人。当他开始'输出'时,骄傲会变成责任。")
    add_faq(doc, "Q5:3C 和 4 模式哪个先学?",
            "A:先学 3C。3C 是底层,4 模式是工具。先有连接,后有分类。")

    add_h2(doc, "模块四 FAQ(5 问)")
    add_faq(doc, "Q1:GROW+ 对话要多长时间?",
            "A:30-60 分钟。第 1 次可能 60 分钟,熟练后 30 分钟。")
    add_faq(doc, "Q2:GROW+ 5 阶段顺序能变吗?",
            "A:不能。G 必须在第 1 位,否则后面都乱。但每个阶段可以来回走。")
    add_faq(doc, "Q3:GROW+ 可以用在 1V1 吗?",
            "A:可以。每周 1V1 走 1 遍 GROW+,团队会很有方向。")
    add_faq(doc, "Q4:GROW+ 是不是太结构化了?",
            "A:结构化是手段,目的是让下属自己想清楚。可以灵活变,但 5 阶段不能少。")
    add_faq(doc, "Q5:GROW+ 对话,主管要做笔记吗?",
            "A:建议做。记下关键句子,后面 1V1 用得上。")

    add_h2(doc, "模块五 FAQ(5 问)")
    add_faq(doc, "Q1:如果我没什么情绪经验,怎么听?",
            "A:先学命名情绪。每天 1 个新词,1 个月后你就掌握 30 个。")
    add_faq(doc, "Q2:如果下属在第 1 层(忽视),我能在第 7 层听吗?",
            "A:不能。倾听是双向的。下属不打开,你再会听也没用。先用 3C 让他信任你。")
    add_faq(doc, "Q3:8 分钟太短,能不能 30 分钟?",
            "A:可以,但建议从 8 分钟开始。先做到 8 分钟,再延到 15 分钟、30 分钟。")
    add_faq(doc, "Q4:沉默时,我不说话,会不会冷场?",
            "A:不会。沉默是给对方'想'的时间。30 秒后,他会说更深的话。")
    add_faq(doc, "Q5:如果我听了,下属没反应,是不是我听错了?",
            "A:不是。也许他还没准备好。第 2 次听,他可能就打开了。")

    add_h2(doc, "模块六 FAQ(5 问)")
    add_faq(doc, "Q1:30 天一定要做完吗?",
            "A:建议做完。30 天不连续,效果会打 5 折。")
    add_faq(doc, "Q2:共学群是必须的吗?",
            "A:建议加入。1 个人走 30 天很孤独,1 群人走 30 天很容易。")
    add_faq(doc, "Q3:30 天后,还要继续吗?",
            "A:建议继续。但 30 天后,从'每天打卡'变成'每周打卡'。")
    add_faq(doc, "Q4:如果 30 天后,下属没变化,是不是我白做了?",
            "A:不是。下属变化是慢的。3 个月后,你会看到效果。")
    add_faq(doc, "Q5:30 天后,我能教别人吗?",
            "A:能。我们会邀请优秀的学员做'30 天校友分享'。")


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "思源黑体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "思源黑体")
    rFonts.set(qn("w:ascii"), "思源黑体")
    rFonts.set(qn("w:hAnsi"), "思源黑体")

    add_cover(doc, "教练技术", "新的管理方式", "AI 时代的教练领导力 · 学员手册")
    add_page_break(doc)
    add_toc(doc)

    add_page_break(doc)
    add_h1(doc, "第一部分 课程概览")
    add_body(doc, "本课程为期 2 天,聚焦 AI 时代管理者从'指挥者'到'教练'的角色转变。")
    add_body(doc, "2 天 × 6 模块 × 19 工具,我们一起完成 3 件事:")
    add_body(doc, "1. 看清 3 次管理坍塌,理解'为什么必须变'")
    add_body(doc, "2. 掌握 3C 信任、GROW+ 对话、七层倾听 3 大工具")
    add_body(doc, "3. 用 30 天 5 维度打卡,把方法变成肌肉记忆")

    add_quote(doc, "谁能兼容谁,谁就能领导谁。", "罗宏伟 · 寄语")

    add_page_break(doc)
    add_module_one(doc)

    add_page_break(doc)
    add_module_two(doc)

    add_page_break(doc)
    add_module_three(doc)

    add_page_break(doc)
    add_module_four(doc)

    add_page_break(doc)
    add_module_five(doc)

    add_page_break(doc)
    add_module_six(doc)

    add_page_break(doc)
    add_tools(doc)

    add_page_break(doc)
    add_cases(doc)

    add_page_break(doc)
    add_action_30d(doc)

    add_page_break(doc)
    add_journal(doc)

    add_page_break(doc)
    add_faq_section(doc)

    add_page_break(doc)
    add_h1(doc, "附录:6 大金句卡片")
    add_body(doc, "这 6 句话,贴在你桌上,每天看一次。")
    add_quote(doc, "1. 谁能兼容谁,谁就能领导谁。", "—— 兼容")
    add_quote(doc, "2. 真实感而不是真实。", "—— 真实感")
    add_quote(doc, "3. 讨喜而不是讨好。", "—— 讨喜")
    add_quote(doc, "4. 领先半步,吃尽红利。", "—— 领先半步")
    add_quote(doc, "5. 延时摄影——和时间做朋友。", "—— 延时")
    add_quote(doc, "6. 善借力打力。", "—— 借力")

    add_body(doc, "")
    add_body(doc, "")
    add_quote(doc, "你不需要变成一个完美的教练,你只需要开始,每周多 1 个 8 分钟心谈。", "罗老师寄语")
    add_body(doc, "30 天后,你会感谢今天的自己。")

    output_path = os.path.join(OUT_DIR, "教练技术_学员手册.docx")
    doc.save(output_path)
    print(f"[OK] 学员手册生成完成: {output_path}")


if __name__ == "__main__":
    build()

# Chapter 2 content
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {})
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

output_path = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册/讲师手册-变革成果固化机制.docx"
doc = Document(output_path)

# ============ CHAPTER 2 ============
add_heading(doc, "第二章  模块一：固化机制认知（2小时）", 1)

add_heading(doc, "2.1 教学目标", 2)
add_para(doc, '1. 理解变革成果固化的本质是"将变革成果从个人能力转化为组织能力"')
add_para(doc, "2. 识别导致推动者一走，变革就倒退的根本原因")
add_para(doc, "3. 掌握固化与坚持的关键区别")
add_para(doc, "4. 理解固化的三个层次：文件层、流程层、文化层")

add_heading(doc, "2.2 时间分配", 2)
table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Table Grid'
for i, h in enumerate(["教学内容", "时长", "核心产出"]):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, "D9D9D9")
for row_idx, (content, duration, output) in enumerate([
    ("变革成果固化的本质", "45分钟", "固化现状诊断框架"),
    ("固化 vs 坚持", "45分钟", "固化成熟度模型"),
    ("固化现状诊断", "30分钟", "固化成熟度自评"),
]):
    table4.rows[row_idx + 1].cells[0].text = content
    table4.rows[row_idx + 1].cells[1].text = duration
    table4.rows[row_idx + 1].cells[2].text = output

add_heading(doc, "2.3 详细教学流程", 2)

add_heading(doc, "2.3.1 变革成果固化的本质（45分钟）", 3)
add_para(doc, "【0-5分钟】开场白", bold=True)
add_para(doc, "讲师话术：各位，今天我们进入模块一的学习。我想先问大家一个问题——你们有没有经历过这样的情况：变革项目在推动的时候运转良好，但推动者一离开，变革就悄悄倒退了？")
add_para(doc, "这个问题几乎是所有变革推动者的痛。今天，我们就来学习如何避免这种人走茶凉的现象。")

add_para(doc, "【5-15分钟】概念讲解", bold=True)
add_para(doc, "核心概念：固化不是保存而是转化")
add_para(doc, "变革成果的两种形态对比：")
add_para(doc, "• 初始形态：依附于人的知识和能力 → 依赖特定个人 → 人员变动即失效")
add_para(doc, "• 固化形态：脱离个人、嵌入组织的制度和习惯 → 依赖组织系统 → 风险可控")

add_para(doc, "【15-25分钟】案例分析", bold=True)
add_para(doc, "情境案例：消失的AI审批流程")
add_para(doc, "某制造企业在引入AI质检系统后，IT部门设计了完整的AI审批流程。三个月后项目经理小张调岗，新任项目经理小李发现AI审批流程名存实亡，变成了先上车后补票。")

add_para(doc, "讨论问题：")
add_para(doc, "1. 这个案例中，固化失败的原因是什么？")
add_para(doc, "2. AI审批流程目前处于哪个固化成熟度级别？")
add_para(doc, "3. 如果你是新任项目经理，你会怎么做？")

add_para(doc, "【25-40分钟】小组活动", bold=True)
add_para(doc, "活动：固化程度投票（15分钟）")
add_para(doc, "• 展示四个情景描述，让学员匿名投票")
add_para(doc, "• 情景A（Level 1）、情景B（Level 3）、情景C（Level 2）、情景D（Level 0）")
add_para(doc, "• 公布答案并讲解评分依据")

add_para(doc, "活动：小组诊断研讨（30分钟，提前布置）", bold=True)
add_para(doc, "• 学员分为4-5人小组")
add_para(doc, "• 每组选择一个本企业的变革案例")
add_para(doc, "• 使用固化成熟度评估表进行评估")

add_para(doc, "【40-45分钟】总结与过渡", bold=True)
add_para(doc, "• 回顾关键要点：固化是转化，不是保存；三个层次层层递进")
add_para(doc, "• 预告下一部分：固化vs坚持——为什么用意志力维持变革注定失败")

add_heading(doc, "2.3.2 固化 vs 坚持（45分钟）", 3)
add_para(doc, "【0-10分钟】概念辨析", bold=True)
add_para(doc, "坚持的定义：坚持是用意志力抵抗放弃的冲动，是一种被动防御状态。")
add_para(doc, "固化的定义：固化是将变革成果转化为组织能力，是主动建设状态。")

add_para(doc, "【10-25分钟】核心区别表格讲解", bold=True)
add_para(doc, "用表格对比坚持和固化的六个维度，重点强调：")
add_para(doc, "• 本质区别：坚持是人治，固化是法治+文治")
add_para(doc, "• 成本差异：人治的成本是无限的，法治+文治的成本是递减的")

add_para(doc, "【25-40分钟】坚持陷阱分析", bold=True)
add_para(doc, "很多变革推动者陷入坚持陷阱：")
add_para(doc, "• 不断提醒、不断催促、不断检查")
add_para(doc, "• 把大量精力消耗在维持现状上")
add_para(doc, "• 没有精力做新的变革")
add_para(doc, "固化是逃脱坚持陷阱的唯一出路。")

add_para(doc, "【40-45分钟】核心逻辑", bold=True)
add_para(doc, "固化的核心逻辑是机制替代人：")
add_para(doc, "• 不是让人记住，而是让系统记住")
add_para(doc, "• 不是让人执行，而是让流程自动触发")
add_para(doc, "• 不是让人监督，而是让数据自动预警")

add_heading(doc, "2.3.3 固化现状诊断（30分钟）", 3)
add_para(doc, "【0-10分钟】固化健康度四问", bold=True)
add_para(doc, "介绍固化成熟度四问模型：")
add_para(doc, "• 依赖问：如果推动者明天调岗，变革成果能持续吗？")
add_para(doc, "• 标准问：新人仅凭文件能否正确执行新流程？")
add_para(doc, "• 激励问：新流程被执行时，执行者得到正向反馈了吗？")
add_para(doc, "• 遗忘问：过去6个月，有没有人主动提醒不要退回旧做法？")

add_para(doc, "【10-20分钟】固化成熟度模型", bold=True)
add_para(doc, "Level 0（未固化）：推动者一走，变革立即停止")
add_para(doc, "Level 1（文件化）：有制度、有流程、有文档，但执行靠自觉")
add_para(doc, "Level 2（流程化）：有定期检查、自动提醒机制")
add_para(doc, "Level 3（文化化）：新流程成为默认做法，退回旧做法反而需要解释理由")

add_para(doc, "【20-30分钟】工具使用", bold=True)
add_para(doc, "介绍固化成熟度评估表的使用方法")
add_para(doc, "学员使用四问进行自评")

add_heading(doc, "2.4 核心产出", 2)
add_para(doc, "固化现状诊断框架")

add_heading(doc, "2.5 讲师注意事项", 2)
add_para(doc, "1. 避免文件层幻觉：要强调文件层只是起点，真正的固化是让变革成果自动运转")
add_para(doc, "2. 强调固化是系统建设：固化不是做一件大事，而是做好很多件小事")
add_para(doc, "3. 警惕坚持陷阱：帮助学员识别是否陷入维持现状的陷阱")
add_para(doc, "4. 固化需要时间：从Level 0到Level 3通常需要1-2年，要有耐心")

doc.save(output_path)
print("Chapter 2 added")

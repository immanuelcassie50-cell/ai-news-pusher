# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_run_font(run, font_name='微软雅黑', font_size=11, bold=False):
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), font_name)

def add_form_title(doc, number, title):
    """Add tool form title"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"工具{number}：{title}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_instruction(doc, text):
    """Add instruction text"""
    p = doc.add_paragraph()
    run = p.add_run(f"【填写说明】{text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')
    return p

def create_table_with_headers(doc, headers, rows, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], 'D9E2F3')
        for p in header_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                r = run._r
                rPr = r.get_or_add_rPr()
                rPr.set(qn('w:eastAsia'), '微软雅黑')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx+1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for p in row_cells[col_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    r = run._r
                    rPr = r.get_or_add_rPr()
                    rPr.set(qn('w:eastAsia'), '微软雅黑')

    return table

# ================== 工具01 ==================
def add_tool_01(doc):
    add_form_title(doc, '01', '社会阶层认知自测表')
    add_instruction(doc, '三维度测评：经济资本、文化资本、社会资本。每维度5题，共15题。每题A=4分，B=3分，C=2分，D=1分。')

    # Scoring table
    headers = ['维度', '题目', 'A(4分)', 'B(3分)', 'C(2分)', 'D(1分)', '得分']
    questions = [
        ['经济资本', '1. 您的家庭年收入水平属于？', '□', '□', '□', '□', ''],
        ['经济资本', '2. 您拥有的房产数量？', '□', '□', '□', '□', ''],
        ['经济资本', '3. 您的金融投资额度？', '□', '□', '□', '□', ''],
        ['经济资本', '4. 您家庭消费水平在社会哪个层次？', '□', '□', '□', '□', ''],
        ['经济资本', '5. 您拥有的固定资产（如车、贵重物品）？', '□', '□', '□', '□', ''],
        ['文化资本', '6. 您的最高学历？', '□', '□', '□', '□', ''],
        ['文化资本', '7. 您每年阅读的书籍数量？', '□', '□', '□', '□', ''],
        ['文化资本', '8. 您参加文化艺术活动的频率？', '□', '□', '□', '□', ''],
        ['文化资本', '9. 您家庭对教育的投入程度？', '□', '□', '□', '□', ''],
        ['文化资本', '10. 您具备哪些专业技能或资格证书？', '□', '□', '□', '□', ''],
        ['社会资本', '11. 您认识多少位可以提供帮助的朋友？', '□', '□', '□', '□', ''],
        ['社会资本', '12. 您加入了多少个社交组织/社团？', '□', '□', '□', '□', ''],
        ['社会资本', '13. 您家人/亲属的社会地位总体如何？', '□', '□', '□', '□', ''],
        ['社会资本', '14. 您在遇到困难时有多少人愿意帮助？', '□', '□', '□', '□', ''],
        ['社会资本', '15. 您通过社交网络获取信息的频率？', '□', '□', '□', '□', ''],
    ]

    table = create_table_with_headers(doc, headers, questions)

    # Result interpretation
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【结果解读】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    interpretations = [
        ('上层阶层', '总分45分以上，三项均≥13分', '拥有丰富的经济、文化和社会资本，在社会中处于优势地位。'),
        ('中上阶层', '总分36-44分，任一维度≥12分', '资本积累较好，有较强的社会流动潜力。'),
        ('中产阶层', '总分25-35分，三项较为均衡', '温饱有余但安全感不足，面临较大的向下流动压力。'),
        ('底层阶层', '总分25分以下，任一维度≤6分', '资本匮乏，社会流动困难，需要政策关注和帮扶。'),
    ]

    headers2 = ['类型', '判定标准', '特征描述']
    rows2 = [[t[0], t[1], t[2]] for t in interpretations]
    create_table_with_headers(doc, headers2, rows2)
    doc.add_paragraph()

# ================== 工具02 ==================
def add_tool_02(doc):
    add_form_title(doc, '02', '阶层流动路径分析表')
    add_instruction(doc, '纵向分析三代人的资本变化，横向对比三种资本类型。在每个格子中填写资源占有情况（丰富/一般/匮乏）和流动判断（↑上升/↓下降/→持平）。')

    headers = ['', '父辈', '自己', '子女', '变化趋势']
    rows = [
        ['经济资本', '资源：____\n判断：__', '资源：____\n判断：__', '资源：____\n判断：__', '□向上流动\n□向下流动\n□保持稳定'],
        ['文化资本', '资源：____\n判断：__', '资源：____\n判断：__', '资源：____\n判断：__', '□向上流动\n□向下流动\n□保持稳定'],
        ['社会资本', '资源：____\n判断：__', '资源：____\n判断：__', '资源：____\n判断：__', '□向上流动\n□向下流动\n□保持稳定'],
    ]

    table = create_table_with_headers(doc, headers, rows)
    doc.add_paragraph()

    # Analysis summary
    p = doc.add_paragraph()
    run = p.add_run('【综合分析】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('流动关键节点：________________________________________________________')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')

    p3 = doc.add_paragraph()
    run3 = p3.add_run('阻碍/促进因素：________________________________________________________')
    run3.font.size = Pt(10)
    run3.font.name = '微软雅黑'
    r3 = run3._r
    rPr3 = r3.get_or_add_rPr()
    rPr3.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具03 ==================
def add_tool_03(doc):
    add_form_title(doc, '03', '家族"拯救者"识别表')
    add_instruction(doc, '识别家族网络中资源贡献最大、角色定位关键、影响机制明显的成员。填写下表并分析其特征。')

    headers = ['识别维度', '具体表现', '是/否', '说明']
    rows = [
        ['资源贡献', '是否在家族中提供经济支持（如借钱、资助购房）？', '□是 □否', ''],
        ['资源贡献', '是否在家族中提供人脉资源（介绍工作、帮忙办事）？', '□是 □否', ''],
        ['资源贡献', '是否在家族中提供信息资源（传递消息、提供建议）？', '□是 □否', ''],
        ['角色定位', '是否在家族聚会中担任主持人/决策者角色？', '□是 □否', ''],
        ['角色定位', '是否经常被家族成员求助/依赖？', '□是 □否', ''],
        ['角色定位', '是否在家族冲突中充当调解人？', '□是 □否', ''],
        ['影响机制', '其意见是否对家族决策有重大影响？', '□是 □否', ''],
        ['影响机制', '其行为模式是否被家族其他成员模仿？', '□是 □否', ''],
        ['影响机制', '其价值观念是否主导家族氛围？', '□是 □否', ''],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【拯救者特征总结】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('您的家族拯救者是谁？其最突出的三项特征：\n1. ____________________________________\n2. ____________________________________\n3. ____________________________________')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具04 ==================
def add_tool_04(doc):
    add_form_title(doc, '04', '中产感觉评估表')
    add_instruction(doc, '评估中产阶层的安全感状态。四个层次由高到低，参照系分析帮助定位自己的位置。')

    p = doc.add_paragraph()
    run = p.add_run('一、四个层次自评')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    headers = ['层次', '状态描述', '自评(打√)']
    rows = [
        ['第一层\n（最稳定）', '有房有车无贷款，子女教育金充足，退休有保障，医疗无担忧', '□'],
        ['第二层\n（较稳定）', '有房有车略有贷款，子女教育基本覆盖，退休有一定储备', '□'],
        ['第三层\n（不稳定）', '租房或房贷压力大，子女教育支出勉强，退休准备不足', '□'],
        ['第四层\n（最不稳定）', '基本生存型消费，工作不稳定，担忧子女未来，随时可能返贫', '□'],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run2 = p2.add_run('二、参照系分析')
    run2.font.bold = True
    run2.font.size = Pt(11)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')

    headers2 = ['参照对象', '您的状况', '差距分析']
    rows2 = [
        ['与父辈同年龄时相比', '', '□更好 □相当 □更差'],
        ['与同龄同学/同事相比', '', '□更好 □相当 □更差'],
        ['与期望的生活状态相比', '', '□达标 □基本达标 □差距大'],
    ]
    create_table_with_headers(doc, headers2, rows2)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    run3 = p3.add_run('三、不安全感具体来源')
    run3.font.bold = True
    run3.font.size = Pt(11)
    run3.font.name = '微软雅黑'
    r3 = run3._r
    rPr3 = r3.get_or_add_rPr()
    rPr3.set(qn('w:eastAsia'), '微软雅黑')

    p4 = doc.add_paragraph()
    run4 = p4.add_run('请勾选您感受到的不安全感来源：\n□职业发展天花板 □行业衰退/裁员风险 □房产贬值 □子女教育压力\n□父母养老负担 □医疗费用担忧 □通货膨胀 □社会地位下滑 □其他：______')
    run4.font.size = Pt(10)
    run4.font.name = '微软雅黑'
    r4 = run4._r
    rPr4 = r4.get_or_add_rPr()
    rPr4.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具05 ==================
def add_tool_05(doc):
    add_form_title(doc, '05', '知识变现能力评估')
    add_instruction(doc, '评估四类知识分子的变现能力和市场化适应度。选择最符合您实际情况的选项。')

    headers = ['知识分子类型', '特征描述', '变现能力', '市场化适应度', '自评']
    rows = [
        ['学术型', '专注于理论研究，发表论文，参与学术项目', '★★☆', '较低', '□'],
        ['技术型', '具备专业技能，能够解决实际问题', '★★★', '较高', '□'],
        ['管理型', '具备组织协调能力，能够带领团队', '★★★', '高', '□'],
        ['创意型', '具备创新思维，能够创造新价值', '★★★', '高', '□'],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【变现能力详细评估】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    headers2 = ['评估维度', '具体问题', '评分(1-5)']
    rows2 = [
        ['专业技能', '您的专业技能在市场上的稀缺程度？', ''],
        ['变现渠道', '您有多少种将知识转化为收入的渠道？', ''],
        ['客户获取', '您获取目标客户的能力如何？', ''],
        ['品牌建设', '您个人品牌的知名度和美誉度？', ''],
        ['定价能力', '您对服务/产品的定价话语权？', ''],
    ]
    create_table_with_headers(doc, headers2, rows2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run('\n综合得分：___/25  变现能力评估：□强 □中 □弱')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具06 ==================
def add_tool_06(doc):
    add_form_title(doc, '06', '平民生存状况评估表')
    add_instruction(doc, '评估最起码生存状况指标。符合条件填"是"，不符合填"否"。')

    headers = ['评估指标', '具体标准', '当前状态', '风险等级']
    rows = [
        ['收入水平', '月收入是否达到当地最低工资标准的1.5倍？', '□是 □否', '□高 □中 □低'],
        ['居住条件', '是否有稳定的居所（租房也算）？', '□是 □否', '□高 □中 □低'],
        ['医疗保障', '是否有基本医疗保险？', '□是 □否', '□高 □中 □低'],
        ['子女教育', '子女是否能正常接受义务教育？', '□是 □否', '□高 □中 □低'],
        ['食物保障', '是否能保证每天三餐营养均衡？', '□是 □否', '□高 □中 □低'],
        ['衣物保障', '是否有足够的衣物应对四季变化？', '□是 □否', '□高 □中 □低'],
        ['交通出行', '是否能承担基本的交通费用？', '□是 □否', '□高 □中 □低'],
        ['通讯联络', '是否有手机且能正常缴费？', '□是 □否', '□高 □中 □低'],
        ['社交联系', '是否能维持基本的社会交往？', '□是 □否', '□高 □中 □低'],
        ['应急储备', '是否有相当于3个月生活费的储蓄？', '□是 □否', '□高 □中 □低'],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【风险预警汇总】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('高风险项：___________  中风险项：___________  低风险项：___________')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')

    p3 = doc.add_paragraph()
    run3 = p3.add_run('综合评估：□基本安全 □存在风险（需关注） □危机状态（需紧急干预）')
    run3.font.size = Pt(10)
    run3.font.name = '微软雅黑'
    r3 = run3._r
    rPr3 = r3.get_or_add_rPr()
    rPr3.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具07 ==================
def add_tool_07(doc):
    add_form_title(doc, '07', '农民工城市融入度测评')
    add_instruction(doc, '八维度评估农民工的城市融入程度。每个维度满分10分，结合自评和代际对比分析。')

    headers = ['评估维度', '具体指标', '自己评分(1-10)', '父辈评分(1-10)', '子女评分(1-10)']
    rows = [
        ['经济融入', '收入水平、工作稳定性、社会保障', '', '', ''],
        ['居住融入', '住房条件、居住稳定性、居住环境', '', '', ''],
        ['社会融入', '社交网络、社会参与、社区归属', '', '', ''],
        ['文化融入', '城市文化认同、价值观念适应', '', '', ''],
        ['心理融入', '城市认同感、归属感、未来预期', '', '', ''],
        ['制度融入', '户籍制度、公共服务获取平权', '', '', ''],
        ['代际融入', '子女教育融入、代际关系调适', '', '', ''],
        ['身份融入', '市民身份认同、标签认同程度', '', '', ''],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【融入障碍分析】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('最大障碍：____________________________________________________________\n次要障碍：____________________________________________________________\n促进因素：____________________________________________________________')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具08 ==================
def add_tool_08(doc):
    add_form_title(doc, '08', '权钱结网识别清单')
    add_instruction(doc, '识别权力与金钱的不正当勾连。以下清单用于自查和预警。')

    p = doc.add_paragraph()
    run = p.add_run('一、预警信号检测（符合的请打√）')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    headers = ['预警类别', '信号描述', '是否出现']
    rows = [
        ['权力寻租', '官员利用职务便利为特定企业谋取利益', '□是 □否'],
        ['权力寻租', '官员家属从事与父辈职务相关的经营活动', '□是 □否'],
        ['资本渗透', '企业通过行贿获取稀缺资源或行政许可', '□是 □否'],
        ['资本渗透', '企业通过政商旋转门获取内幕信息', '□是 □否'],
        ['利益输送', '通过关联交易向权力关系人输送利益', '□是 □否'],
        ['利益输送', '通过子女联姻实现政商联盟', '□是 □否'],
        ['保护伞', '黑恶势力背后有官员撑腰', '□是 □否'],
        ['保护伞', '企业违法行为长期未被查处', '□是 □否'],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run2 = p2.add_run('二、场景案例自查')
    run2.font.bold = True
    run2.font.size = Pt(11)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')

    p3 = doc.add_paragraph()
    run3 = p3.add_run('案例1：某企业主通过官员子女留学费用赞助获取行政许可，您的判断是：\n□权钱交易  □正常商业行为  □难以判断\n\n案例2：某官员退休后到曾审批项目的企业任高管，您的判断是：\n□期权腐败  □正常职业选择  □难以判断')
    run3.font.size = Pt(10)
    run3.font.name = '微软雅黑'
    r3 = run3._r
    rPr3 = r3.get_or_add_rPr()
    rPr3.set(qn('w:eastAsia'), '微软雅黑')

    p4 = doc.add_paragraph()
    run4 = p4.add_run('三、自查问题')
    run4.font.bold = True
    run4.font.size = Pt(11)
    run4.font.name = '微软雅黑'
    r4 = run4._r
    rPr4 = r4.get_or_add_rPr()
    rPr4.set(qn('w:eastAsia'), '微软雅黑')

    p5 = doc.add_paragraph()
    run5 = p5.add_run('1. 您或您的亲友是否曾被要求"找关系"办事？\n2. 您是否遇到过不合理的市场准入障碍？\n3. 您是否了解身边存在的政商不当关系？')
    run5.font.size = Pt(10)
    run5.font.name = '微软雅黑'
    r5 = run5._r
    rPr5 = r5.get_or_add_rPr()
    rPr5.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具09 ==================
def add_tool_09(doc):
    add_form_title(doc, '09', '三种资本转化路径表')
    add_instruction(doc, '分析经济资本、文化资本、社会资本之间的转化路径和具体方法。')

    headers = ['转化方向', '转化逻辑', '具体方法', '案例']
    rows = [
        ['经济→文化', '用金钱购买教育、文化体验，提升文化资本',
         '1. 购买优质教育资源\n2. 培养艺术爱好\n3. 参加高端培训\n4. 收藏艺术品',
         '中产阶层送子女上国际学校'],
        ['文化→社会', '用知识和人脉获取更多社会资源',
         '1. 通过学术圈拓展人脉\n2. 用专业能力换取信任\n3. 参与社会活动建立声誉',
         '律师通过专业能力建立客户网络'],
        ['社会→经济', '用人脉关系获取经济利益',
         '1. 介绍项目获取佣金\n2. 合作投资机会\n3. 政商资源对接',
         '掮客通过人脉撮合交易获利'],
        ['文化→经济', '将知识直接转化为产品或服务',
         '1. 出书、授课\n2. 咨询服务\n3. 技术转让\n4. 内容创业',
         '知识博主通过知识付费变现'],
        ['社会→文化', '通过社交圈获取文化资源',
         '1. 加入高端俱乐部\n2. 参与文化交流活动\n3. 进入文化圈层',
         '通过朋友介绍参观私人美术馆'],
        ['经济→社会', '通过消费建立社交身份',
         '1. 高端消费场所消费\n2. 参加商务社交活动\n3. 赞助社会活动',
         '通过高尔夫球会结识商界精英'],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【您的资本转化实践】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('最成功的资本转化：_______________________________________________________\n尚未实现的转化方向：___________________________________________________')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 工具10 ==================
def add_tool_10(doc):
    add_form_title(doc, '10', '五个一行动计划表')
    add_instruction(doc, '制定从即时到年度的行动计划。每个层面对应：目标、行动、资源、里程碑。')

    headers = ['时间维度', '目标', '行动', '资源需求', '里程碑/截止日期']
    rows = [
        ['一天\n(24小时内)', '', '', '', ''],
        ['一周\n(7天内)', '', '', '', ''],
        ['一月\n(30天内)', '', '', '', ''],
        ['一季\n(90天内)', '', '', '', ''],
        ['一年\n(365天内)', '', '', '', ''],
    ]
    create_table_with_headers(doc, headers, rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('【行动计划示例参考】')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    p2 = doc.add_paragraph()
    run2 = p2.add_run('一天：完成社会资本盘点，列出可依赖的10个人脉\n一周：约见一位能够提供职业建议的前辈\n一月：完成职业技能评估，确定需要提升的能力项\n一季：完成一门专业认证课程的学习\n一年：实现收入增长20%或完成职业转型')
    run2.font.size = Pt(10)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ================== 主函数 ==================
def main():
    doc = Document()

    # Set page margins
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_before = Pt(72)
    run = title.add_run('中国社会各阶层分析')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '微软雅黑'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run('课程工具表单集锦')
    run2.font.size = Pt(20)
    run2.font.name = '微软雅黑'
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), '微软雅黑')

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = author.add_run('\n\n基于梁晓声《中国社会各阶层分析》')
    run3.font.size = Pt(14)
    run3.font.name = '微软雅黑'
    r3 = run3._r
    rPr3 = r3.get_or_add_rPr()
    rPr3.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # Add all tools
    add_tool_01(doc)
    add_tool_02(doc)
    add_tool_03(doc)
    add_tool_04(doc)
    add_tool_05(doc)
    add_tool_06(doc)
    add_tool_07(doc)
    add_tool_08(doc)
    add_tool_09(doc)
    add_tool_10(doc)

    # Save
    output_path = 'D:/新课开发/工作手册/梁晓声社会阶层分析/中国社会各阶层分析-原始版/完整课程包/08-工具表单集锦/工具集锦-中国社会各阶层分析.docx'
    doc.save(output_path)
    print(f'Word document saved to: {output_path}')

if __name__ == '__main__':
    main()

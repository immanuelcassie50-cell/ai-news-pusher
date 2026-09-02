# -*- coding: utf-8 -*-
"""生成对外宣传文案 docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False, color=None):
    """设置run的字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, font_size=16, color=(43, 45, 66)):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=True, color=color)
    return p

def add_paragraph(doc, text, font_size=11, indent=False, space_before=0, space_after=8):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size)
    return p

def add_bullet_point(doc, text, font_size=11):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run_font(run, font_size=font_size)
    return p

def create_promotion_doc(output_path):
    """创建宣传文案文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ========== 封面区域 ==========
    # 主标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(60)
    p_title.paragraph_format.space_after = Pt(20)
    run = p_title.add_run("手册进化")
    set_run_font(run, font_size=36, bold=True, color=(43, 45, 66))

    # 副标题
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(40)
    run = p_sub.add_run('从"阅读版"到"执行手册"')
    set_run_font(run, font_size=24, bold=False, color=(141, 153, 174))

    # 分隔线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(40)
    run = p_line.add_run("_" * 40)
    set_run_font(run, font_size=12, color=(141, 153, 174))

    # 核心主张
    p_core = doc.add_paragraph()
    p_core.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_core.paragraph_format.space_after = Pt(60)
    run = p_core.add_run("读得懂 ≠ 做得对")
    set_run_font(run, font_size=28, bold=True, color=(239, 35, 60))

    # ========== 课程定位 ==========
    add_heading(doc, "课程定位", font_size=18)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("这是一门关于")
    set_run_font(run, font_size=11)
    run = p.add_run("手册升级转化")
    set_run_font(run, font_size=11, bold=True)
    run = p.add_run("的实操工作坊。")
    set_run_font(run, font_size=11)

    add_paragraph(doc, '很多企业已经积累了丰富的业务手册：场景写了、流程写了、案例和问答也都有了。但拿到现场用的时候，执行者还是频繁卡壳——不是因为内容不够，而是因为内容的组织方式不对。', indent=True)
    add_paragraph(doc, '本课程聚焦一个核心问题：如何把一份已经有了但不实用的手册，改造成任何人都能照着做对的执行手册。', indent=True)

    # ========== 核心主张 ==========
    add_heading(doc, "核心主张", font_size=18)

    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(6)
    p_box.paragraph_format.space_after = Pt(6)
    p_box.paragraph_format.left_indent = Pt(20)
    run = p_box.add_run("读得懂，和做得对，是两件事。")
    set_run_font(run, font_size=14, bold=True, color=(239, 35, 60))

    add_paragraph(doc, '阅读版手册解决的是"信息有没有被交代清楚"，执行手册解决的是"一个从没干过这件事的人，拿着它，能不能在没人带的情况下把事做对、做完整、不出事故"。这中间差的不是内容量，是内容的组织方式。', indent=True)

    # ========== 学员对象 ==========
    add_heading(doc, "学员对象", font_size=18)

    targets = [
        '内训师——需要把手册转化为人人能照着做的执行标准',
        '知识萃取专员——负责把业务经验转化为可复用的知识资产',
        '培训项目负责人——确保培训产出物真正能落地执行',
        '行动学习催化师——推动从"知道"到"做到"的转化'
    ]
    for target in targets:
        add_bullet_point(doc, target)

    # ========== 学员收获 ==========
    add_heading(doc, "学员收获", font_size=18)

    p_gain = doc.add_paragraph()
    p_gain.paragraph_format.space_after = Pt(12)
    run = p_gain.add_run('学完本课程，你将带走：')
    set_run_font(run, font_size=11)

    gains = [
        ('一套可复用的转化自检清单', '诊断现有手册缺什么，验收转化完成度'),
        ('六类内容要素的空白写作模板', '场景节点链、操作清单、判断标准卡、分级处置表、案例结构、问答分类框架'),
        ('判断标准卡的通用写作公式', '条件信号 + 判断结论 + 对应动作 + 关键变量 + 可忽略因素'),
        ('一套完整的工具包', '可直接用于实际手册升级项目的模板和清单')
    ]

    for title, desc in gains:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Pt(20)
        run = p.add_run('✓ ' + title)
        set_run_font(run, font_size=11, bold=True, color=(42, 45, 66))
        run = p.add_run('：' + desc)
        set_run_font(run, font_size=11)

    # ========== 课程内容 ==========
    add_heading(doc, "课程内容（六大模块）", font_size=18)

    modules = [
        ('01', '场景定位', '把泛泛的场景描述拆解成可识别的动作节点链'),
        ('02', '标准动作', '把模糊的流程写成可打卡、可核对的清单结构'),
        ('03', '判断标准卡', '在关键节点单独给出判断条件，不让读者自己提炼'),
        ('04', '分级处置', '按严重程度分级，每级对应明确的处置动作'),
        ('05', '情境案例', '用真实场景串起判断逻辑，建立代入感'),
        ('06', '高频自查', '覆盖真正会卡壳的现场问题，不是知识点罗列')
    ]

    for num, title, desc in modules:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(20)
        run = p.add_run(num + ' ' + title)
        set_run_font(run, font_size=11, bold=True, color=(42, 45, 66))
        run = p.add_run(' — ' + desc)
        set_run_font(run, font_size=11)

    # ========== 教学方式 ==========
    add_heading(doc, "教学方式", font_size=18)

    methods = [
        '开篇给真实冲突场景——让学员先认同问题存在',
        '每个模块配套动手练习——现场拿真实素材做一次转化',
        '练习后有明确对照标准——供学员自查或互查',
        '结尾给出可带走反复使用的自检工具'
    ]
    for method in methods:
        add_bullet_point(doc, method)

    # ========== 课程信息 ==========
    add_heading(doc, "课程信息", font_size=18)

    info_items = [
        ('课程时长', '半天工作坊，3-4小时'),
        ('建议人数', '20-30人'),
        ('目标学员', '内训师、知识萃取专员、培训项目负责人'),
        ('课程形式', '理论讲解 + 案例示范 + 小组练习 + 现场点评')
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label + '：')
        set_run_font(run, font_size=11, bold=True)
        run = p.add_run(value)
        set_run_font(run, font_size=11)

    # ========== 联系方式 ==========
    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(40)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_contact.add_run('—— 联系方式 ——')
    set_run_font(run, font_size=12, color=(141, 153, 174))

    # 保存文档
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')

if __name__ == '__main__':
    output_path = 'D:/新课开发/经验萃取/阅读手册转执行手册/完整课程包/004-对外宣传文案/004-对外宣传文案.docx'
    create_promotion_doc(output_path)
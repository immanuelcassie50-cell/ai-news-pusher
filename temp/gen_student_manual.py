"""
gen_student_manual.py - 生成《教练技术_学员手册.docx》
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\2026年课程\竞越\教练技术：新的管理方式\完整课程包\14_Word手册"
os.makedirs(OUT_DIR, exist_ok=True)

# 配色
PURPLE = RGBColor(0x5B, 0x3A, 0x8C)
ORANGE = RGBColor(0xF2, 0xA0, 0x3D)
INK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x9A, 0x98, 0x90)
LIGHT_GREY = RGBColor(0xE8, 0xE5, 0xE0)
TERRACOTTA = RGBColor(0xB8, 0x5C, 0x3E)
GOLD = RGBColor(0xC9, 0xA9, 0x61)

# ============================================================
# 工具函数
# ============================================================
def set_zh_font(run, font_name="思源黑体"):
    """设置中文字体"""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")
    return p


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")
    return p


def add_quote(doc, text, source=""):
    """金句"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run('"' + text + '"')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run("—— " + source)
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = GREY
        set_zh_font(r2, "思源黑体")
    return p


def add_case(doc, title, body):
    """案例方框"""
    # 标题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("【案例】" + title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TERRACOTTA
    set_zh_font(r, "思源宋体")
    # 正文
    for line in body:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + line)
        r.font.size = Pt(11)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


def add_blank(doc, hint=""):
    """学员可填写区"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "bottom", "left", "right"]:
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "dashed")
        bdr.set(qn("w:sz"), "8")
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "C0C0C0")
        pBdr.append(bdr)
    pPr.append(pBdr)
    if hint:
        r = p.add_run("【此处填写】 " + hint)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GREY
        set_zh_font(r, "思源黑体")
    # 加几个空行
    for _ in range(3):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(" ")


def add_table_simple(doc, headers, rows, widths=None):
    """简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_zh_font(r, "思源黑体")
        # 表头底色
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2A03D")
        tcPr.append(shd)
    # 数据
    for i, row in enumerate(rows):
        for j, cell_data in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_data) if cell_data else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = INK
                    set_zh_font(r, "思源黑体")
    if widths:
        for j, w in enumerate(widths):
            for row in table.rows:
                row.cells[j].width = w


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def add_cover(doc, title, subtitle, kicker):
    """封面"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(48)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(subtitle)
    r.font.size = Pt(28)
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(kicker)
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    r = p.add_run("罗宏伟 · 竞越合作讲师")
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    set_zh_font(r, "思源黑体")


def add_toc(doc):
    """目录"""
    add_h1(doc, "目录")
    items = [
        "第一部分 课程概览",
        "模块一 AI 时代的管理坍塌",
        "模块二 教练 vs 传统管理者的 6 个差异",
        "模块三 3C 信任模型 + 4 模式员工",
        "模块四 GROW+ 5 阶段对话框架",
        "模块五 七层倾听阶梯",
        "模块六 30 天落地实践",
        "第二部分 19 个工具使用说明",
        "第三部分 真实案例集",
        "第四部分 30 天行动手册",
        "第五部分 反思日记模板",
        "附录:6 大金句卡片",
    ]
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


# ============================================================
# 生成学员手册
# ============================================================
def build_student_manual():
    doc = Document()
    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 设置正文默认字体
    style = doc.styles["Normal"]
    style.font.name = "思源黑体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "思源黑体")
    rFonts.set(qn("w:ascii"), "思源黑体")
    rFonts.set(qn("w:hAnsi"), "思源黑体")

    # 封面
    add_cover(doc, "教练技术", "新的管理方式", "AI 时代的教练领导力 · 学员手册")

    add_page_break(doc)
    add_toc(doc)

    add_page_break(doc)

    # 第一部分:课程概览
    add_h1(doc, "第一部分 课程概览")
    add_body(doc, "本课程为期 2 天,聚焦 AI 时代管理者从'指挥者'到'教练'的角色转变。")
    add_body(doc, "2 天 × 6 模块 × 19 工具,我们一起完成 3 件事:")
    add_body(doc, "1. 看清 3 次管理坍塌,理解'为什么必须变'")
    add_body(doc, "2. 掌握 3C 信任、GROW+ 对话、七层倾听 3 大工具")
    add_body(doc, "3. 用 30 天 5 维度打卡,把方法变成肌肉记忆")

    add_quote(doc, "谁能兼容谁,谁就能领导谁。", "罗宏伟 · 寄语")

    # 模块一
    add_h1(doc, "模块一 AI 时代的管理坍塌")
    add_body(doc, "为什么 20 年的老管理,突然不会管了?不是你变差了,是你脚下的地,塌了。")

    add_h2(doc, "1.1 第一次坍塌:信息差消失")
    add_body(doc, "过去,管理者是信息中枢。员工遇到问题,来找你问。")
    add_body(doc, "现在,AI 3 秒出答案。员工比你先知道。")
    add_body(doc, "权威根基动摇:你凭什么管?")
    add_case(doc, "老王查资料的尴尬", [
        "60 后车间主任,30 年工龄",
        "新员工用手机扫了一下设备,3 秒读出参数",
        "老王翻了 5 分钟手册没找到",
        "他开始怀疑:我是不是该退了?",
    ])

    add_h2(doc, "1.2 第二次坍塌:经验失效")
    add_body(doc, "过去,经验就是权威。10 年工龄,说了算。")
    add_body(doc, "现在,AI 看过 100 万案例,比你多 1 万倍。")
    add_body(doc, "经验贬值:30 年老师傅,3 个月就被 AI 替代。")
    add_case(doc, "30 年老师傅被 AI 替代", [
        "上海某机械厂老师傅:看裂纹 30 年,听声音就知道",
        "AI 音频分析 + 图像识别,准确率 98%",
        "半年后:老师傅,你带新徒弟吧",
        "他不是被 AI 替代,是被'经验'替代",
    ])

    add_h2(doc, "1.3 第三次坍塌:决策权下沉")
    add_body(doc, "过去,你定方向,员工执行。")
    add_body(doc, "现在,Z 世代自己定 KPI,自己选项目。")
    add_body(doc, "决策权消失:你不再是决策者。")
    add_case(doc, "95 后员工自己决策", [
        "互联网公司 95 后产品经理",
        "老板:下周上线这个功能",
        "95 后:我做了一份用户调研,数据不支持。先改",
        "老板沉默 3 秒:行,你说了算",
    ])

    add_h2(doc, "1.4 三次坍塌的累积效应")
    add_body(doc, "信息坍塌:你不再是唯一的信息源")
    add_body(doc, "经验坍塌:你不再是唯一的判断者")
    add_body(doc, "决策权坍塌:你不再是唯一的决策者")
    add_quote(doc, "三次坍塌 = 传统管理者的 3 个支柱,全部塌了。")

    add_h2(doc, "1.5 新角色:从指挥到教练")
    add_body(doc, "指挥者:我告诉你做什么")
    add_body(doc, "教练:我帮你找到为什么做")
    add_quote(doc, "延时摄影——和时间做朋友。", "做教练短期会慢,长期会快")

    add_h2(doc, "1.6 模块一思考题")
    add_blank(doc, "回顾你最近 3 次和下属的对话,你能看到'信息坍塌'的痕迹吗?")

    # 模块二
    add_h1(doc, "模块二 教练 vs 传统管理者的 6 个差异")
    add_body(doc, "你不需要变成另一个人,你只需要看清 6 个差异,每个差异往前一步。")

    add_h2(doc, "2.1 6 个差异总览")
    add_table_simple(doc, ["维度", "传统管理者", "教练"],
                    [
                        ["权威来源", "职位给我权威", "专业+信任给我权威"],
                        ["对话姿态", "我讲你听", "我问你答"],
                        ["决策方式", "我定方向", "你定方向,我陪"],
                        ["角色定位", "我是裁判", "我是教练"],
                        ["反馈风格", "你做错了,批评", "你做到了,再加什么"],
                        ["长期目标", "本季度业绩", "5 年后他还在成长"],
                    ])

    add_quote(doc, "讨喜而不是讨好。", "—— 别只说你爱听的话")

    add_h2(doc, "2.2 差异 1:权威来源")
    add_body(doc, "传统:职位给我权威,我说了算")
    add_body(doc, "教练:专业+信任给我权威,你愿意听")
    add_body(doc, "职位权威 1 年失效,信任权威 10 年有效。")

    add_h2(doc, "2.3 差异 2:对话姿态")
    add_body(doc, "传统:我讲你听,我命令你执行")
    add_body(doc, "教练:我问你答,我帮你找到答案")
    add_case(doc, "小李的两次老板", [
        "传统老板:小李,这件事明天必须做",
        "教练老板:小李,这件事你怎么看?为什么重要?",
        "结果:1 年后,小李从前者离职,从后者升职",
    ])

    add_h2(doc, "2.4 差异 3:决策方式")
    add_body(doc, "传统:我定方向,你执行")
    add_body(doc, "教练:我帮你看清,你自己定")
    add_body(doc, "下属自己定的方案,执行度 3 倍于你定的。")

    add_h2(doc, "2.5 差异 4:角色定位")
    add_body(doc, "传统:我是裁判,定对错")
    add_body(doc, "教练:我是教练,陪你上场")
    add_quote(doc, "善借力打力。", "—— 让每个人成为自己的教练")

    add_h2(doc, "2.6 差异 5:反馈风格")
    add_body(doc, "传统:你做错了,再批评")
    add_body(doc, "教练:你做到了,再加什么更好")
    add_body(doc, "行为反馈 vs 人格评判——后者让人崩。")

    add_h2(doc, "2.7 差异 6:长期目标")
    add_body(doc, "传统:本季度业绩,完成就 OK")
    add_body(doc, "教练:5 年后他还在成长,是我赢")
    add_quote(doc, "延时摄影——和时间做朋友。")

    add_h2(doc, "2.8 教练心态:4 个核心信念")
    for s in [
        "1. 每个人都比自己以为的更强",
        "2. 答案在他心里,不在我嘴里",
        "3. 我不是来救他,我是来陪他",
        "4. 短期会慢,长期会快",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "2.9 教练红线:5 个不能")
    for s in [
        "1. 不能替他做决定",
        "2. 不能评判他的选择",
        "3. 不能忽视他的情绪",
        "4. 不能用'我是为你好'绑架",
        "5. 不能在他没准备好时硬推",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "2.10 模块二思考题")
    add_blank(doc, "回想本周 1 次对话,按 6 个差异打分(1-5)。哪个差异需要重点练?")

    add_page_break(doc)

    # 模块三
    add_h1(doc, "模块三 3C 信任模型 + 4 模式员工")
    add_body(doc, "看见人,才能带人。3C 是建立信任的 3 个关键动作,4 模式是识别员工的 4 种工具。")

    add_h2(doc, "3.1 1 个数字:信任的价值")
    add_quote(doc, "高信任团队,绩效比低信任团队高 76%。", "Paul Zak 2017 NEJM 研究")
    add_body(doc, "信任 = 团队的隐形业绩杠杆。")

    add_h2(doc, "3.2 3C 信任模型")
    add_h3(doc, "3.2.1 Connection 连接")
    add_body(doc, "看见他:名字、状态、困难")
    add_body(doc, "记住他:上次说的话、他的家庭")
    add_body(doc, "在他需要时,出现在场")
    add_body(doc, "连接 = 我看见你这个人,不是你的岗位。")

    add_h3(doc, "3.2.2 Commitment 承诺")
    add_body(doc, "我说的,我做到")
    add_body(doc, "小事也兑现(请假、加薪、建议)")
    add_body(doc, "做不到的,提前说,不甩锅")
    add_body(doc, "承诺 = 我信你,因为你一直说到做到。")

    add_h3(doc, "3.2.3 Caring 关怀")
    add_body(doc, "关心他的状态(累不累)")
    add_body(doc, "关心他的成长(想学什么)")
    add_body(doc, "关心他的生活(家里怎么样)")
    add_body(doc, "关怀 = 我在乎你这个人,不只是你的产出。")

    add_h2(doc, "3.3 4 模式员工:AI 时代的全新分类")
    add_table_simple(doc, ["模式", "特征", "管理动作"],
                    [
                        ["先行者", "抢着用 AI,跑得快", "授权+挑战+允许试错"],
                        ["整合者", "愿意尝试,先看效果", "桥梁+跨部门协作"],
                        ["观望者", "等别人先试,看情况", "1V1 心谈+安全尝试"],
                        ["保守者", "担心被替代,不敢动", "尊重节奏+稳定本职"],
                    ])

    add_case(doc, "小张(先行者)怎么带", [
        "28 岁,产品经理,AI 工具用得比主管还溜",
        "用 AI 做用户调研报告,1 小时搞定",
        "团队其他人都不会,他成了'AI 翻译'",
        "但他开始骄傲,说主管是'老古董'",
        "管理动作:授权 + 挑战性任务,允许试错,纳入创新项目",
    ])

    add_case(doc, "老王(保守者)怎么带", [
        "58 岁,30 年工龄,AI 焦虑,3 个月没睡好",
        "学不会,失眠,易怒,想提前退休",
        "管理动作:尊重节奏,先稳定本职工作,AI 工具从低风险场景切入",
    ])

    add_h2(doc, "3.4 跨模式协作:1+1>2 的搭配")
    add_table_simple(doc, ["搭配", "效果"],
                    [
                        ["先行者 + 整合者", "创新+落地,黄金组合"],
                        ["观望者 + 先行者", "Buddy 机制,先看后试"],
                        ["保守者 + 整合者", "稳定+信任,小步前进"],
                    ])

    add_h2(doc, "3.5 模块三练习")
    add_blank(doc, "列出你的 5 名下属,按 4 模式打分(1-5),主类型取最高分")
    add_table_simple(doc, ["姓名", "先行者", "整合者", "观望者", "保守者", "主类型"],
                    [[f"下属{i + 1}", "", "", "", "", ""] for i in range(5)])

    add_page_break(doc)

    # 模块四
    add_h1(doc, "模块四 GROW+ 5 阶段对话框架")
    add_body(doc, "GROW+ = 教练对话的标准动作。5 阶段 25 问,覆盖一次完整对话。")

    add_h2(doc, "4.1 传统对话 vs GROW 对话")
    add_table_simple(doc, ["传统对话", "GROW 对话"],
                    [
                        ["我讲你听,我说你做", "我问你想,我帮你看清"],
                        ["我告诉你答案", "你自己找到答案"],
                    ])

    add_h2(doc, "4.2 G 阶段:目标 Goal")
    add_body(doc, "整个对话的灯塔。G 阶段没做对,后面都白搭。")
    add_h3(doc, "G 阶段 5 个关键问题")
    for q in [
        "1. 你最想达成的是什么?",
        "2. 如果成功了,你会看到什么?",
        "3. 这个目标对你为什么重要?",
        "4. 你愿意为它付出什么?",
        "5. 什么时候必须达成?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "4.3 R 阶段:现实 Reality")
    add_body(doc, "看清脚下,才能走对方向。R 阶段关键:从'我没机会'变成'我能做什么'。")
    add_h3(doc, "R 阶段 5 个关键问题")
    for q in [
        "1. 现在的情况是什么?",
        "2. 你做了哪些尝试?效果如何?",
        "3. 卡在哪里?",
        "4. 哪些是你能控制的,哪些不能?",
        "5. 如果不做任何改变,3 个月后会怎样?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "4.4 O 阶段:选择 Options")
    add_body(doc, "让他自己看见 N 种可能。O 阶段关键:激发他'做不到'的假设。")
    add_h3(doc, "O 阶段 5 个关键问题")
    for q in [
        "1. 如果不受限,你会怎么做?",
        "2. 你身边谁做得最好?你能学他什么?",
        "3. 还有哪些方法你没试过?",
        "4. 哪个方法最让你兴奋?",
        "5. 如果只能选一个,你先做哪个?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "4.5 W 阶段:意愿 Will")
    add_body(doc, "激活内在动力。等他说'我愿意',不是你说。")
    add_h3(doc, "W 阶段 5 个关键问题")
    for q in [
        "1. 你最想做哪个?为什么?",
        "2. 做到后对你意味着什么?",
        "3. 你愿意为它放弃什么?",
        "4. 如果不做,你会有什么感觉?",
        "5. 你对自己的承诺是什么?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "4.6 A 阶段:行动 Action")
    add_body(doc, "把对话变成真东西。A 阶段关键:具体的、可衡量的、有时限的。")
    add_h3(doc, "A 阶段 5 个关键问题")
    for q in [
        "1. 你第一步做什么?",
        "2. 什么时候做?",
        "3. 你需要的支持是什么?",
        "4. 我(主管)能帮你什么?",
        "5. 我们怎么知道你做到了?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "4.7 GROW+ 完整剧本:小李延期")
    add_body(doc, "小李的项目延期 1 周,主管想骂他。")
    add_table_simple(doc, ["阶段", "问话"],
                    [
                        ["G", "小李,你最想达成什么?(问目标)"],
                        ["R", "现在延期了,情况怎样?(问现状)"],
                        ["O", "如果不延期 1 周,你想怎么做?(问方案)"],
                        ["W", "你愿意做哪个?(问意愿)"],
                        ["A", "第一步是什么?什么时候?(问行动)"],
                    ])
    add_body(doc, "结果:小李主动加班赶回进度,1 周准时交付。")

    add_h2(doc, "4.8 模块四练习")
    add_blank(doc, "找一个下属,用 GROW+ 完整走一遍,记下每个阶段他说的话")

    add_page_break(doc)

    # 模块五
    add_h1(doc, "模块五 七层倾听阶梯")
    add_body(doc, "从'我听见了'到'你被听见了'。大多数管理者卡在第 3 层,我们要到第 5 层以上。")

    add_h2(doc, "5.1 七层阶梯")
    add_table_simple(doc, ["层", "名称", "描述"],
                    [
                        ["7", "生成性倾听", "让对方找到他自己都没意识到的答案"],
                        ["6", "反思性倾听", "把他说的话翻译成深层含义"],
                        ["5", "共鸣性倾听", "分享我类似的感受,让他不孤单"],
                        ["4", "同理心倾听", "听到情绪,命名情绪,让他被看见"],
                        ["3", "理解性倾听", "听他说完,复述确认(大多数人在这)"],
                        ["2", "选择性倾听", "只听我感兴趣的,过滤其他的"],
                        ["1", "忽视", "心不在焉,人在心不在"],
                    ])

    add_h2(doc, "5.2 命名情绪:5 个万能句式")
    for q in [
        "1. 你看起来感到___对吗?",
        "2. 听上去,你现在___对吗?",
        "3. 如果是我,我会感到___。你呢?",
        "4. 你愿意多说说这个___吗?",
        "5. 你希望我___吗?",
    ]:
        add_body(doc, q, indent=False)

    add_h2(doc, "5.3 8 分钟心谈:操作指南")
    for s in [
        "1. 找一个安静的地方(不被打扰)",
        "2. 说:'我有 8 分钟,你愿意聊聊吗?'",
        "3. 前 4 分钟:只听,不评判,不建议",
        "4. 后 4 分钟:问 1-2 个开放问题",
        "5. 结束:问'你还需要什么吗?'",
    ]:
        add_body(doc, s, indent=False)

    add_quote(doc, "8 分钟 = 不长不短,刚好够听见他。")

    add_h2(doc, "5.4 8 分钟心谈:5 个典型错误")
    for s in [
        "1. 边听边看手机",
        "2. 立刻给建议",
        "3. 打断他说话",
        "4. '我以前也是',变成我讲",
        "5. 急着解决他的问题",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "5.5 反思倾听:翻译术")
    add_body(doc, "翻译 = 把话背后的意思说出来,让他'被深度理解'。")
    add_table_simple(doc, ["他说", "你翻译"],
                    [
                        ["老板没看见我的方案", "你感到被忽视,所以失落"],
                        ["我不想做这个项目了", "这个项目让你失去了意义感"],
                        ["同事不配合我", "你感到孤立,需要支持"],
                    ])

    add_h2(doc, "5.6 生成倾听:沉默的力量")
    for s in [
        "1. 他说完,你不说话,等他",
        "2. 他沉默 10 秒,你也不说话",
        "3. 这时候,他会'想'出更深的东西",
        "4. 30 秒后,他可能说出他自己都没意识到的话",
        "5. 你只问:'你刚才在沉默里,在想什么?'",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "5.7 模块五练习")
    add_blank(doc, "找 1 个下属,做 1 次 8 分钟心谈,记下他的反应")

    add_page_break(doc)

    # 模块六
    add_h1(doc, "模块六 30 天落地实践")
    add_body(doc, "把方法变成肌肉记忆。30 天,5 个维度,4 周。")

    add_h2(doc, "6.1 30 天 5 维度")
    add_table_simple(doc, ["维度", "内容"],
                    [
                        ["心谈", "每天 8 分钟心谈 1 次"],
                        ["GROW+", "每天用 GROW+ 解决 1 个小问题"],
                        ["3C", "每天 1 个 3C 动作"],
                        ["团队分享", "每周 1 次内部分享"],
                        ["反思", "每天写 1 段反思日记"],
                    ])

    add_h2(doc, "6.2 第 1 周:心谈")
    for s in [
        "Day 1: 找一个下属,做 8 分钟心谈",
        "Day 2: 找另一个下属,同样",
        "Day 3-5: 每天 1 个,共 5 个下属",
        "Day 6: 复盘 - 哪个对话印象最深?",
        "Day 7: 把印象最深的写进反思日记",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "6.3 第 2 周:GROW+")
    for s in [
        "Day 8: 选 1 个下属,用 GROW+ 聊 20 分钟",
        "Day 9-10: 每天 1 次,共 3 个下属",
        "Day 11: 找最难的 1 个下属(观望者/保守者)",
        "Day 12: 找最熟的 1 个下属(先行者/整合者)",
        "Day 13: 复盘 + 写反思日记",
        "Day 14: 整理剧本,准备内部分享",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "6.4 第 3 周:3C")
    for s in [
        "Day 15: Connection - 记住下属生日",
        "Day 16: Commitment - 兑现 1 个小事",
        "Day 17: Caring - 关心 1 个下属的私事",
        "Day 18: Connection - 听下属说一件私事",
        "Day 19: Commitment - 提前说 1 件做不到的事",
        "Day 20: Caring - 给 1 个下属写卡片",
        "Day 21: 复盘 + 3C 评估",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "6.5 第 4 周:团队分享")
    for s in [
        "Day 22-24: 把 30 天实践整理成 1 页纸",
        "Day 25: 找 1 个下属做内部分享",
        "Day 26: 在周会上分享 1 个 GROW+ 剧本",
        "Day 27: 收集 3 个下属的反馈",
        "Day 28: 写完 30 天反思日记",
        "Day 29: 做 F1 训后评估",
        "Day 30: 30 天庆祝,加入校友群",
    ]:
        add_body(doc, s, indent=False)

    add_h2(doc, "6.6 30 天共学群规则")
    for s in [
        "1. 每天 1 个动作打卡(心谈/GROW+/3C/分享/反思 任选)",
        "2. 每周 1 次 8 分钟心谈实战复盘",
        "3. 每周 1 次团队分享",
        "4. 每周 1 个 30 天反思日记片段",
        "5. 30 天后全群庆祝 + 内部分享",
    ]:
        add_body(doc, s, indent=False)

    add_page_break(doc)

    # 第二部分:19 个工具
    add_h1(doc, "第二部分 19 个工具使用说明")
    add_body(doc, "19 个工具覆盖训前、训中、训后全流程。每个工具都是可操作的练习,不是空话。")

    tools = [
        ("F1", "预评量表", "30 题 × 5 维度,训前训后对比"),
        ("F2", "时间账单", "12 小时 × 30 分钟,看清时间怎么用"),
        ("F3", "4 模式自评", "团队成员分类,管理动作分人"),
        ("F4", "3C 信任度评估", "打分 1-10,找最弱维度"),
        ("F5", "GROW+ 剧本卡", "5 阶段 + 25 问,实战对话"),
        ("F6", "倾听层级表", "7 层 × 自评,看自己在第几层"),
        ("F7", "8 分钟心谈卡", "5 步操作 + 8 分钟计时器"),
        ("F8", "情绪命名卡", "30 个情绪词 + 5 个万能句式"),
        ("F9", "F19 行动卡", "5 要素拆解行动"),
        ("F10", "3C 每日动作", "连接/承诺/关怀 3 类 9 个动作"),
        ("F11", "团队氛围雷达", "7 维度 × 12 月追踪"),
        ("F12", "GROW+ 观察卡", "教练 12 项观察"),
        ("F13", "角色扮演脚本", "8 个真实剧本"),
        ("F14", "GROW+ 情景卡", "8 张卡牌(扑克牌尺寸)"),
        ("F15", "教练观察卡", "12 项观察 + GROW+ 阶段"),
        ("F16", "F3 模式分布图", "4 模式人数 + 占比饼图"),
        ("F17", "失败复盘模板", "3 种没效果处理"),
        ("F18", "30 天打卡表", "5 维度 × 30 天追踪"),
        ("F19", "行动计划", "32 天 × 5 要素"),
    ]
    add_table_simple(doc, ["编号", "工具名", "用途"],
                    tools)

    add_page_break(doc)

    # 第三部分:案例集
    add_h1(doc, "第三部分 真实案例集")
    add_body(doc, "8 个真实案例,覆盖先行者/整合者/观望者/保守者 4 类员工。")

    add_h2(doc, "案例 1:小李延期")
    add_body(doc, "小李,90 后,项目骨干,做事快。")
    add_body(doc, "老板信任他,什么都给他。他咬牙坚持 1 个月,提了离职。")
    add_body(doc, "GROW+ 解决方案:G 找到他真正想要的,R 厘清现状,O 拓宽可能,W 激活意愿,A 行动。")
    add_body(doc, "结果:小李主动加班赶回进度,1 周准时交付。")

    add_h2(doc, "案例 2:老王转型")
    add_body(doc, "老王,60 后车间主任,30 年工龄,担心被 AI 替代。")
    add_body(doc, "GROW+ 解决方案:不评判他的焦虑,问'你愿意为它做什么?'")
    add_body(doc, "结果:老王主动申请做'AI 教练',带新员工。")

    add_h2(doc, "案例 3:小张迷茫")
    add_body(doc, "小张,工作 3 年,不知道往哪走。")
    add_body(doc, "GROW+ 解决方案:问'3 年后你想成为什么?'")
    add_body(doc, "结果:小张选了产品方向,3 个月后转岗成功。")

    add_h2(doc, "案例 4:老李瓶颈")
    add_body(doc, "老李,做了 5 年总监,升不上去,想跳不动。")
    add_body(doc, "GROW+ 解决方案:问'如果不受限,你会怎么做?'")
    add_body(doc, "结果:老李找到新赛道,开了副业。")

    add_h2(doc, "案例 5:小陈加班")
    add_body(doc, "小陈,95 后,连续 1 个月 996,在茶水间哭。")
    add_body(doc, "8 分钟心谈 + 命名情绪:'你感到撑不住了对吗?'")
    add_body(doc, "结果:小陈说出来了,主管才知道他的极限。")

    add_h2(doc, "案例 6:老张 AI 焦虑")
    add_body(doc, "老张,55 岁,工程师,被 AI 替代了一半工作。")
    add_body(doc, "GROW+ 解决方案:不评判他的失落,问'你最想达成什么?'")
    add_body(doc, "结果:老张开始学 AI 工具,3 个月后成了 AI 培训师。")

    add_h2(doc, "案例 7:小赵创新")
    add_body(doc, "小赵,推 AI 项目被否 3 次,想放弃。")
    add_body(doc, "GROW+ 解决方案:R 阶段问'你做了哪些尝试?'让他看见自己已经做了 3 次。")
    add_body(doc, "结果:小赵重启提案,加上数据,第 4 次通过。")

    add_h2(doc, "案例 8:老周临退休")
    add_body(doc, "老周,58 岁,5 年后退休,突然空虚。")
    add_body(doc, "8 分钟心谈 + 翻译:'听起来,你想留下点什么'")
    add_body(doc, "结果:老周申请做导师,带新人,找到意义感。")

    add_page_break(doc)

    # 第四部分:30 天行动手册
    add_h1(doc, "第四部分 30 天行动手册")
    add_body(doc, "把 30 天写下来,贴在桌上。每天 5 分钟打卡,30 天后看变化。")

    add_h2(doc, "F19 行动计划模板")
    add_table_simple(doc, ["要素", "你的填写"],
                    [
                        ["1. 我要达成什么(目标)", ""],
                        ["2. 我为什么在乎(意义)", ""],
                        ["3. 我第一步做什么(行动)", ""],
                        ["4. 我可能遇到什么卡点(障碍)", ""],
                        ["5. 我怎么知道自己做到了(衡量)", ""],
                    ])

    add_h2(doc, "30 天打卡表")
    for week in range(1, 5):
        add_h3(doc, f"第 {week} 周")
        add_table_simple(doc, ["Day", "心谈", "GROW+", "3C", "分享", "反思"],
                        [[f"Day {d}", "", "", "", "", ""] for d in range((week - 1) * 7 + 1, week * 7 + 1)])

    add_page_break(doc)

    # 第五部分:反思日记
    add_h1(doc, "第五部分 反思日记模板")
    add_body(doc, "反思 = 把经验变成能力。每天 5 分钟,问自己 5 个问题。")

    for d in [1, 8, 15, 22, 29]:
        add_h2(doc, f"Day {d} 反思")
        for q in [
            "今天我做了什么?为什么这样做?",
            "今天的对话中,哪个瞬间让我印象最深?",
            "我下次会做哪个不一样的动作?",
            "我今天最大的卡点是什么?",
            "我今天学到了什么?",
        ]:
            add_body(doc, q, indent=False)
        add_blank(doc, "写下你的反思")

    add_page_break(doc)

    # 附录:6 大金句
    add_h1(doc, "附录:6 大金句卡片")
    add_body(doc, "这 6 句话,贴在你桌上,每天看一次。")
    add_quote(doc, "1. 谁能兼容谁,谁就能领导谁。", "—— 兼容")
    add_quote(doc, "2. 真实感而不是真实。", "—— 真实感")
    add_quote(doc, "3. 讨喜而不是讨好。", "—— 讨喜")
    add_quote(doc, "4. 领先半步,吃尽红利。", "—— 领先半步")
    add_quote(doc, "5. 延时摄影——和时间做朋友。", "—— 延时")
    add_quote(doc, "6. 善借力打力。", "—— 借力")

    add_body(doc, "")
    add_body(doc, "")
    add_quote(doc, "你不需要变成一个完美的教练,你只需要开始,每周多 1 个 8 分钟心谈。", "罗老师寄语")
    add_body(doc, "30 天后,你会感谢今天的自己。")

    # 保存
    output_path = os.path.join(OUT_DIR, "教练技术_学员手册.docx")
    doc.save(output_path)
    print(f"[OK] 学员手册生成完成: {output_path}")


if __name__ == "__main__":
    build_student_manual()

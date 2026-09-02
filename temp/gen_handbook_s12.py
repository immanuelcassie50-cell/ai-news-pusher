#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 附录二：稽核表模板
app2 = doc.add_paragraph()
r = app2.add_run('附录二：稽核表模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

doc.add_paragraph()
it = doc.add_paragraph()
r = it.add_run('基本信息')
r.bold = True

t2 = doc.add_table(rows=5, cols=2)
t2.style = 'Table Grid'

for i, label in enumerate(['决策名称', '决策日期', '决策者', '稽核日期', '稽核人']):
    cl = t2.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True

doc.add_paragraph()
at = doc.add_paragraph()
r = at.add_run('稽核问题')
r.bold = True

aqs = [
    ('问题1：沉淀检查', '这次决策过程中，有没有出现过让决策者感到当时不确定的关键节点？\n□ 有 → 该节点是否已被记录？\n□ 没有\n备注：'),
    ('问题2：使用检查', '如果对应的场景已有决策卡覆盖：\n□ 已被使用，反馈：_______\n□ 未被使用，原因：_______\n□ 无对应卡片'),
    ('问题3：价值识别', '这次决策是否具备做成卡片的价值？\n□ 是 → 是否已启动访谈？\n□ 否 → 原因：_______\n备注：'),
    ('问题4：反馈迭代', '是否发现现有卡片存在漏洞或过时之处？\n□ 是 → 是否已反馈给认领人？\n□ 否\n备注：'),
]

for q_title, q_content in aqs:
    tq = doc.add_table(rows=2, cols=1)
    tq.style = 'Table Grid'
    ct = tq.cell(0, 0)
    set_cell_shading(ct, 'D9E2F3')
    r = ct.paragraphs[0].add_run(q_title)
    r.bold = True
    tq.cell(1, 0).paragraphs[0].add_run(q_content)
    doc.add_paragraph()

et = doc.add_paragraph()
r = et.add_run('额外问题')
r.bold = True
doc.add_paragraph('这次有没有基于自己的判断，偏离了卡片上的建议？为什么？\n（用于检测是否过度依赖卡片、失去自主判断力）')

doc.add_paragraph()
coc = doc.add_paragraph()
r = coc.add_run('稽核结论')
r.bold = True
doc.add_paragraph('□ 经验已有效沉淀\n□ 需要跟进：_______\n□ 需要迭代卡片：_______')

doc.add_page_break()
print('S12: Tool template 2 done')
doc.save(r'D:\CC\temp\handbook_s12.docx')

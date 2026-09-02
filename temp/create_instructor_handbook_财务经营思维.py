# -*- coding: utf-8 -*-
"""
《财务经营思维——非财务经理的报表解读与经营分析》讲师手册生成脚本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width = Inches(11.69)  # A4横向
section.page_height = Inches(8.27)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ===== 样式定义 =====
styles = doc.styles
style_normal = styles['Normal']
style_normal.font.name = 'Microsoft YaHei'
style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style_normal.font.size = Pt(11)

# ===== 辅助函数 =====
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 56, 100)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(46, 117, 182)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
    return heading

def add_para(doc, text, bold=False, indent=False, size=11):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.font.bold = bold
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    return para

def add_table_row(table, data, bold_first=False, size=10):
    row = table.add_row()
    for i, text in enumerate(data):
        cell = row.cells[i]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                run.font.size = Pt(size)
                if bold_first and i == 0:
                    run.font.bold = True
    return row

def add_warning_box(doc, text, title="注意事项"):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ""
    para = cell.paragraphs[0]
    run1 = para.add_run(f"⚠ {title}：")
    run1.font.name = 'Microsoft YaHei'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run1.font.size = Pt(11)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(192, 80, 77)
    run2 = para.add_run(text)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_key_point_box(doc, text, title="核心要点"):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ""
    para = cell.paragraphs[0]
    run1 = para.add_run(f"★ {title}：")
    run1.font.name = 'Microsoft YaHei'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run1.font.size = Pt(11)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0, 112, 192)
    run2 = para.add_run(text)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_time_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(f"⏱ {text}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    doc.add_paragraph()
    return table

def add_golden_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(f"📌 {text}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.italic = True
    doc.add_paragraph()
    return table

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run("财务经营思维")
title_run.font.name = 'Microsoft YaHei'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
title_run.font.size = Pt(40)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 56, 100)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run("——非财务经理的报表解读与经营分析")
subtitle_run.font.name = 'Microsoft YaHei'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
subtitle_run.font.size = Pt(24)
subtitle_run.font.color.rgb = RGBColor(46, 117, 182)

doc.add_paragraph()

course_info = doc.add_paragraph()
course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = course_info.add_run("FACILITATOR GUIDE\n\n")
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(14)
info_run = course_info.add_run("课程编号：30\n")
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(14)
info_run = course_info.add_run("培训时长：2天（每天6小时，共约12小时）\n")
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(14)
info_run = course_info.add_run("目标学员：非财务背景的业务负责人、部门经理、创业者\n")
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(14)
info_run = course_info.add_run("授课讲师：罗宏伟")
info_run.font.name = 'Microsoft YaHei'
info_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
info_run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

disclaimer = doc.add_paragraph()
disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
disclaimer_run = disclaimer.add_run("本手册为内部培训使用材料，请勿对外传播\n让每一位管理者都具备财务经营思维")
disclaimer_run.font.name = 'Microsoft YaHei'
disclaimer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
disclaimer_run.font.size = Pt(10)
disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ===== 目录 =====
add_heading(doc, "目 录", 1)
doc.add_paragraph()

toc_items = [
    ("第一部分", "讲师准备"),
    ("  一", "课程概述与目标"),
    ("  二", "学员画像分析"),
    ("  三", "教学准备清单"),
    ("  四", "场地与设备要求"),
    ("  五", "课前沟通要点"),
    ("第二部分", "教学设计详解"),
    ("  一", "整体教学框架"),
    ("  二", "各模块教学目标"),
    ("  三", "各模块时间分配"),
    ("  四", "知识点讲解要点"),
    ("  五", "案例使用指南"),
    ("  六", "互动环节设计说明"),
    ("  七", "常见问题应对"),
    ("第三部分", "讲师工具"),
    ("  一", "授课时间轴参考"),
    ("  二", "板书/投影建议"),
    ("  三", "道具/材料清单"),
    ("  四", "Q&A准备"),
    ("第四部分", "评估支持"),
    ("  一", "学员表现观察要点"),
    ("  二", "评分标准说明"),
    ("  三", "反馈话术参考"),
    ("第五部分", "课程迭代"),
    ("  一", "学员反馈收集方法"),
    ("  二", "课程改进建议"),
    ("  三", "常见问题记录表"),
]

for part, title in toc_items:
    para = doc.add_paragraph()
    para.add_run(f"{part}：{title}")
    para.paragraph_format.left_indent = Cm(0.5)
    for run in para.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(12)

doc.add_page_break()

# ===== 第一部分：讲师准备 =====
add_heading(doc, "第一部分 | 讲师准备", 1)

add_heading(doc, "一、课程概述与目标", 2)

add_heading(doc, "1. 课程定位", 3)
add_para(doc, """财务经营思维，是每一位管理者必备的商业语言。在当今竞争激烈的商业环境中，业务负责人、部门经理甚至创业者，都需要理解财务报表背后的经营逻辑，才能做出更明智的决策。""")

add_para(doc, """本课程的核心价值：""", bold=True)
add_bullet(doc, "让非财务背景的管理者读懂三张报表的核心逻辑")
add_bullet(doc, "掌握经营分析的核心指标体系（毛利率、周转率、现金转换周期、ROE）")
add_bullet(doc, "建立用财务视角验证业务决策的思维习惯")
add_bullet(doc, "实现财务与业务的协同语言统一")

add_heading(doc, "2. 课程背景", 3)
add_para(doc, """在企业实践中，我们发现许多业务负责人对财务知识存在明显短板：""")
add_bullet(doc, "看不懂财务报表，只能依靠财务部门的解读")
add_bullet(doc, "做业务决策时缺乏财务验证，拍脑袋决策")
add_bullet(doc, "对成本结构、盈亏平衡点等概念模糊不清")
add_bullet(doc, "无法将业务语言和财务语言统一沟通")

add_warning_box(doc, "课程设计原则：先建立全局观，再深入细节；先培养直觉，再精确定量。每讲一个财务概念，必配一个业务案例。", "核心设计理念")

add_heading(doc, "3. 教学理念", 3)
add_para(doc, """本课程遵循"认知升级三段式"：""")
add_bullet(doc, "第一阶段：建立财务感觉——理解报表的底层逻辑，不畏财务数据")
add_bullet(doc, "第二阶段：掌握分析工具——学会用指标说话，用数据验证决策")
add_bullet(doc, "第三阶段：形成经营思维——将财务视角内化为管理习惯")

add_heading(doc, "4. 课程目标", 3)

goals_table = doc.add_table(rows=6, cols=2)
goals_table.style = 'Table Grid'
add_table_row(goals_table, ["目标层次", "具体描述"], bold_first=True)
add_table_row(goals_table, ["知识目标", "理解资产负债表、利润表、现金流量表的底层逻辑及相互关系"])
add_table_row(goals_table, ["技能目标", "掌握毛利率、周转率、现金转换周期、ROE等核心指标的计算与分析方法"])
add_table_row(goals_table, ["应用目标", "能够用财务语言验证业务决策（定价、成本结构、盈亏平衡）"])
add_table_row(goals_table, ["思维目标", "形成\"用财务视角审视业务\"的思维习惯"])
add_table_row(goals_table, ["沟通目标", "实现与财务部门的高效协同，用统一语言沟通经营状况"])

doc.add_paragraph()

add_heading(doc, "二、学员画像分析", 2)

add_heading(doc, "1. 目标学员特征", 3)
add_para(doc, """本课程面向以下人群：""", bold=True)
add_bullet(doc, "业务负责人：销售、市场、运营等非财务背景的管理者")
add_bullet(doc, "部门经理：研发、生产、人力等需要理解经营数据的部门负责人")
add_bullet(doc, "创业者：需要自己掌控财务全局的创始人和联合创始人")
add_bullet(doc, "储备干部：有潜力走向管理岗位的业务骨干")

add_heading(doc, "2. 学员常见特征", 3)

char_table = doc.add_table(rows=6, cols=2)
char_table.style = 'Table Grid'
add_table_row(char_table, ["特征维度", "具体描述"], bold_first=True)
add_table_row(char_table, ["财务认知", "对财务报表有畏惧感，认为那是\"专业的财务人员才懂的东西\""])
add_table_row(char_table, ["学习动机", "渴望理解经营数据，在团队协作和业务决策中有实际需求"])
add_table_row(char_table, ["知识基础", "具备基本商业常识（如收入、成本、利润），但缺乏系统框架"])
add_table_row(char_table, ["学习风格", "偏好案例驱动、问题导向，厌恶纯粹的理论灌输"])
add_table_row(char_table, ["常见痛点", "看不懂报表提不出问题、做决策缺乏数据支撑、与财务沟通困难"])

doc.add_paragraph()

add_heading(doc, "3. 学员分层教学建议", 3)

layer_table = doc.add_table(rows=4, cols=3)
layer_table.style = 'Table Grid'
add_table_row(layer_table, ["学员类型", "特征", "教学策略"], bold_first=True)
add_table_row(layer_table, ["初学者", "从未接触过财务报表", "放慢节奏，多用生活类比，强调\"不怕财务\"的心态建立"])
add_table_row(layer_table, ["有基础者", "看过报表但理解不深", "强调框架和逻辑，重点突破关键概念的错误认知"])
add_table_row(layer_table, ["有经验者", "有一定财务知识", "提供深度案例和延伸思考题，满足高阶需求"])

doc.add_paragraph()

add_warning_box(doc, "课前通过问卷或访谈了解学员的财务背景，根据多数学员水平调整案例深度和讲解节奏。", "关键提醒")

add_heading(doc, "三、教学准备清单", 2)

add_heading(doc, "1. 内容熟悉度准备", 3)
add_bullet(doc, "完整阅读本手册全部内容，特别注意各模块的【开场引导】和【讲师注意事项】")
add_bullet(doc, "深入理解三张报表的底层逻辑——这不是教会计，而是教经营思维")
add_bullet(doc, "预演各模块的关键讲解点，确保流畅度")
add_bullet(doc, "准备2-3个自己在商业实践中的真实经历案例（最好是自己的企业或投资案例）")
add_bullet(doc, "熟悉所有案例数据，能够回答学员的延伸提问")

add_heading(doc, "2. 课件与材料准备", 3)
checklist_table = doc.add_table(rows=9, cols=2)
checklist_table.style = 'Table Grid'
add_table_row(checklist_table, ["材料名称", "准备标准"], bold_first=True)

checklist_data = [
    ["学员手册", "已打印，每人一份；确认印刷清晰，特别是表格和数据部分"],
    ["三张报表案例", "准备A3纸质案例（或投影），供小组研讨使用"],
    ["指标计算练习题", "准备不同难度的练习题，确保学员有足够的练习机会"],
    ["盈亏平衡分析表", "准备可操作的Excel模板，学员可带走复习"],
    ["金句卡片", "每模块结束后发给学员，作为课后复习材料"],
    ["白板/翻页纸", "用于板书关键框架：杜邦分析、现金转换周期等"],
    ["投影设备", "用于展示PPT、案例数据和课程框架图"],
    ["计时器", "用于控制各环节时间，特别是互动讨论环节"],
]

for row_data in checklist_data:
    add_table_row(checklist_table, row_data)

doc.add_paragraph()

add_heading(doc, "3. 课前自检清单（48小时前）", 3)
self_check = [
    "课件是否已全部制作完成并检查无误？",
    "案例数据是否已核实？特别注意数字的合理性",
    "是否已确认场地设备（投影、音响、白板）运转正常？",
    "是否已获取学员名单，了解学员背景和人数？",
    "是否准备了学员可能问到的延伸问题的答案？",
    "是否确认了助教或辅助人员的人选和分工？",
    "是否准备了备用方案（如投影故障的替代方案）？",
]

for item in self_check:
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(f"☐ {item}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)

add_heading(doc, "4. 课前自检清单（当天开场前）", 3)
day_check = [
    "再次确认投影、音响、白板等设备运转正常",
    "将案例打印材料按小组分配好，每组一份",
    "在白板上写好当天课程的整体框架（供学员了解整体结构）",
    "准备好签到表和反馈表",
    "确认茶水、休息区等后勤安排",
    "给自己预留5分钟的静心时间，调整状态",
]

for item in day_check:
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(f"☐ {item}")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)

add_heading(doc, "四、场地与设备要求", 2)

add_heading(doc, "1. 场地要求", 3)
venue_table = doc.add_table(rows=6, cols=2)
venue_table.style = 'Table Grid'
add_table_row(venue_table, ["项目", "标准"], bold_first=True)

venue_data = [
    ["座位安排", "岛型或圆桌分组（优先于剧院式）——便于小组讨论和互动"],
    ["空间大小", "人均不低于2平方米，避免拥挤感"],
    ["灯光", "明亮但柔和，避免昏暗（财务课程需要专注氛围）"],
    ["室温", "24-26摄氏度，避免过冷或过热影响学员注意力"],
    ["噪音", "选择安静场地，避免临近施工现场或嘈杂区域"],
]

for row_data in venue_data:
    add_table_row(venue_table, row_data)

doc.add_paragraph()

add_heading(doc, "2. 设备要求", 3)
equip_table = doc.add_table(rows=6, cols=2)
equip_table.style = 'Table Grid'
add_table_row(equip_table, ["设备", "准备标准"], bold_first=True)

equip_data = [
    ["投影仪", "亮度不低于3000流明，投影尺寸能看清数据表格（建议80寸以上）"],
    ["音响", "确保讲师麦克风声音清晰，无回音或杂音"],
    ["白板/翻页架", "至少2个翻页架，配红、蓝、黑三种颜色马克笔"],
    ["电脑", "备用电脑一台，防止主电脑故障"],
    ["网络", "确认网络畅通（用于在线案例或视频播放）"],
]

for row_data in equip_data:
    add_table_row(equip_table, row_data)

doc.add_paragraph()

add_warning_box(doc, "财务课程涉及大量数据展示，建议使用投影仪而非电视屏幕，以确保后排学员也能清晰看到数字。", "设备建议")

add_heading(doc, "五、课前沟通要点", 2)

add_heading(doc, "1. 与培训组织者的沟通", 3)
add_bullet(doc, "确认学员名单和背景：了解行业分布、职位层级、年龄结构")
add_bullet(doc, "确认课程时间：是否有调整、是否涉及用餐安排")
add_bullet(doc, "确认学员人数：15-30人为最佳互动规模，超过30人需增加助教")
add_bullet(doc, "了解往期类似课程的学员反馈：哪些内容反响好、哪些需要改进")
add_bullet(doc, "确认课后跟进安排：是否有作业、是否需要提供复习材料")

add_heading(doc, "2. 与学员的课前沟通（可选）", 3)
add_para(doc, """如果条件允许，可以在课前1周发送简短的预习邮件或群消息：""", bold=True)
add_golden_box(doc, "主题：课前预习 | 财务经营思维\n\n各位同学好！\n欢迎参加《财务经营思维》课程。\n\n为帮助大家更好地进入状态，建议预习以下内容：\n1. 回忆最近一次你看到的公司财务报表（哪怕只是瞟了一眼）\n2. 思考：你当时最想知道什么？最看不懂什么？\n\n带着这些问题来上课，效果会更好。\n\n课程中会用到大量真实商业案例，建议大家准备好计算器（手机计算器也行）。")

add_heading(doc, "3. 常见学员问题预判", 3)

preq_table = doc.add_table(rows=5, cols=2)
preq_table.style = 'Table Grid'
add_table_row(preq_table, ["问题类型", "参考回答方向"], bold_first=True)

preq_data = [
    ["\"我不是财务出身，能学会吗？\"", "\"本课程专为零基础设计，只需要基本的商业常识。我们的目标是建立经营思维，不是培养会计。\""],
    ["\"课程涉及的数学复杂吗？\"", "\"只需要基础的加减乘除。我们强调的是商业直觉和数据思维，不是计算能力。\""],
    ["\"学完能直接用吗？\"", "\"当然！每个模块都配有实际案例和工具，可以直接应用到工作中。\""],
    ["\"需要提前看什么书吗？\"", "\"不需要。本课程自成体系，课后会提供延伸阅读清单。\""],
]

for row_data in preq_data:
    add_table_row(preq_table, row_data)

doc.add_paragraph()

doc.add_page_break()

# ===== 第二部分：教学设计详解 =====
add_heading(doc, "第二部分 | 教学设计详解", 1)

add_heading(doc, "一、整体教学框架", 2)

add_heading(doc, "1. 课程结构总览", 3)

structure_table = doc.add_table(rows=8, cols=4)
structure_table.style = 'Table Grid'
add_table_row(structure_table, ["模块", "主题", "核心产出", "时长"], bold_first=True)

structure_data = [
    ["模块一", "经营思维导论", "建立\"财务为经营服务\"的认知框架", "1.5小时"],
    ["模块二", "三张报表的底层逻辑", "读懂资产负债表、利润表、现金流量表", "2.5小时"],
    ["模块三", "经营分析核心指标", "掌握毛利率、周转率、现金转换周期、ROE", "2小时"],
    ["模块四", "业务决策的财务验证", "用财务语言验证定价、成本、盈亏平衡决策", "2小时"],
    ["模块五", "现金为王的经营实践", "理解现金转换周期和经营预警信号", "1.5小时"],
    ["模块六", "财务与业务的协同语言", "统一财务与业务的沟通语言", "1.5小时"],
    ["合计", "—", "—", "约12小时（2天）"],
]

for row_data in structure_data:
    add_table_row(structure_table, row_data)

doc.add_paragraph()

add_heading(doc, "2. 教学流程图", 3)

add_para(doc, """【两天课程整体流程】""", bold=True)

flow_text = """
第一天（上午）
    ↓
模块一：经营思维导论（1.5小时）
    ↓ 建立"为什么管理者需要财务视角"的认知基础
模块二：三张报表的底层逻辑 - 上（1.5小时）
    ↓ 深入理解资产负债表和利润表
    ↓
第一天（下午）
    ↓
模块二：三张报表的底层逻辑 - 下（1小时）
    ↓ 现金流量表及三表关系
模块三：经营分析核心指标 - 上（1.5小时）
    ↓ 毛利率、周转率
模块四：业务决策的财务验证 - 上（1.5小时）
    ↓ 定价决策、成本结构
    ↓
第二天（上午）
    ↓
模块四：业务决策的财务验证 - 下（1小时）
    ↓ 盈亏平衡点分析
模块五：现金为王的经营实践（1.5小时）
    ↓ 现金转换周期、经营预警信号
    ↓
第二天（下午）
    ↓
模块六：财务与业务的协同语言（1.5小时）
    ↓ 用财务语言沟通业务
课程总结与行动计划（1小时）
    ↓
课程结束 → 学员带走实用工具和方法论
"""
para = doc.add_paragraph()
run = para.add_run(flow_text)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(10)

add_key_point_box(doc, "每模块的设计逻辑：开场（建立连接）→ 核心内容（讲解+讨论）→ 工具运用（练习）→ 收尾（总结+预告）", "模块通用结构")

add_heading(doc, "3. 教学方法论", 3)
add_para(doc, """本课程采用多元教学方法：""", bold=True)
add_bullet(doc, "案例教学法：用真实商业案例（华为、美的、拼多多等）讲解抽象概念")
add_bullet(doc, "行动学习法：每个模块配备实际练习，学员动手计算和分析")
add_bullet(doc, "翻转课堂：课前预习概念，课堂时间用于深度讨论和答疑")
add_bullet(doc, "同伴学习：小组讨论和互评，发挥集体智慧")

add_heading(doc, "二、各模块教学目标", 2)

add_heading(doc, "模块一：经营思维导论", 3)
mod1_table = doc.add_table(rows=5, cols=2)
mod1_table.style = 'Table Grid'
add_table_row(mod1_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod1_table, ["认知目标", "理解\"管理者为什么需要财务视角\"——财务是经营的语言，不是会计的技术"])
add_table_row(mod1_table, ["情感目标", "建立对财务数据的信心——\"我也能读懂财务报告\""])
add_table_row(mod1_table, ["技能目标", "能够区分\"财务思维\"与\"会计思维\"的本质差异"])
add_table_row(mod1_table, ["行动目标", "能够向团队成员解释\"财务经营思维\"的价值"])

doc.add_paragraph()

add_heading(doc, "模块二：三张报表的底层逻辑", 3)
mod2_table = doc.add_table(rows=5, cols=2)
mod2_table.style = 'Table Grid'
add_table_row(mod2_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod2_table, ["认知目标", "理解三张报表的核心构成要素——资产、负债、权益、收入、成本、费用、现金流"])
add_table_row(mod2_table, ["理解目标", "掌握三张报表之间的底层关联——资产负债表是基础，利润表是结果，现金流量表是验证"])
add_table_row(mod2_table, ["技能目标", "能够快速定位三张报表中的关键数据，并理解其业务含义"])
add_table_row(mod2_table, ["分析目标", "能够通过三张报表的交叉验证，判断企业的真实经营状况"])

doc.add_paragraph()

add_heading(doc, "模块三：经营分析核心指标", 3)
mod3_table = doc.add_table(rows=5, cols=2)
mod3_table.style = 'Table Grid'
add_table_row(mod3_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod3_table, ["知识目标", "掌握毛利率、周转率、现金转换周期、ROE的定义和计算方法"])
add_table_row(mod3_table, ["分析目标", "能够解读指标背后的业务含义——不是数字本身，而是数字反映的经营逻辑"])
add_table_row(mod3_table, ["应用目标", "能够用指标体系分析竞争态势，发现经营中的潜在风险和机会"])
add_table_row(mod3_table, ["决策目标", "能够将指标分析与业务决策结合，用数据验证业务选择"])

doc.add_paragraph()

add_heading(doc, "模块四：业务决策的财务验证", 3)
mod4_table = doc.add_table(rows=5, cols=2)
mod4_table.style = 'Table Grid'
add_table_row(mod4_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod4_table, ["定价目标", "理解定价决策的财务逻辑——不是\"成本加成\"，而是\"价值定价\""])
add_table_row(mod4_table, ["成本目标", "掌握成本结构分析方法——固定成本、变动成本、边际成本"])
add_table_row(mod4_table, ["盈亏目标", "能够计算盈亏平衡点，并用于业务决策验证"])
add_table_row(mod4_table, ["决策目标", "形成\"业务决策财务验证\"的习惯——任何业务决定，先问财务影响"])

doc.add_paragraph()

add_heading(doc, "模块五：现金为王的经营实践", 3)
mod5_table = doc.add_table(rows=5, cols=2)
mod5_table.style = 'Table Grid'
add_table_row(mod5_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod5_table, ["认知目标", "理解\"现金为王\"的真实含义——不是利润为王，是现金流为王"])
add_table_row(mod5_table, ["分析目标", "掌握现金转换周期（CCC）的计算和分析方法"])
add_table_row(mod5_table, ["预警目标", "能够识别经营预警信号——在危机爆发前发现端倪"])
add_table_row(mod5_table, ["行动目标", "能够提出改善现金流的具体建议"])

doc.add_paragraph()

add_heading(doc, "模块六：财务与业务的协同语言", 3)
mod6_table = doc.add_table(rows=5, cols=2)
mod6_table.style = 'Table Grid'
add_table_row(mod6_table, ["目标维度", "具体描述"], bold_first=True)
add_table_row(mod6_table, ["沟通目标", "掌握财务与业务之间的\"翻译\"能力——将财务语言转化为业务语言"])
add_table_row(mod6_table, ["协作目标", "理解财务部门和业务部门的协作逻辑——不是监督，是服务和支持"])
add_table_row(mod6_table, ["汇报目标", "能够用财务语言向不同利益相关者（老板、投资人、团队）汇报经营状况"])
add_table_row(mod6_table, ["系统目标", "建立全局经营观——将各模块知识融会贯通，形成完整的经营分析能力"])

doc.add_paragraph()

add_heading(doc, "三、各模块时间分配", 2)

time_alloc = doc.add_table(rows=9, cols=4)
time_alloc.style = 'Table Grid'
add_table_row(time_alloc, ["模块", "主题", "建议时长", "时间分配说明"], bold_first=True)

time_data = [
    ["模块一", "经营思维导论", "1.5小时", "概念建立20% + 案例讨论30% + 练习50%"],
    ["模块二", "三张报表底层逻辑", "2.5小时", "讲解40% + 案例30% + 小组研讨30%"],
    ["模块三", "经营分析核心指标", "2小时", "讲解35% + 计算练习30% + 案例分析35%"],
    ["模块四", "业务决策财务验证", "2小时", "讲解30% + 案例40% + 决策模拟30%"],
    ["模块五", "现金为王经营实践", "1.5小时", "讲解35% + 案例35% + 预警练习30%"],
    ["模块六", "财务业务协同语言", "1.5小时", "讲解30% + 角色扮演40% + 总结30%"],
    ["总结", "课程总结与行动计划", "1小时", "知识回顾30% + 行动计划40% + 答疑30%"],
    ["合计", "—", "约12小时", "2天完整课程"],
]

for row_data in time_data:
    add_table_row(time_alloc, row_data)

doc.add_paragraph()

add_warning_box(doc, "时间分配是参考值，可根据学员反应和现场情况灵活调整。如果某个模块的案例讨论特别热烈，可以适当延长；如果学员明显感到疲惫，要增加休息或调整节奏。", "时间弹性")

add_heading(doc, "四、知识点讲解要点", 2)

# 模块一
add_heading(doc, "模块一：经营思维导论——为什么管理者需要财务视角", 3)
add_time_box(doc, "建议时长：1.5小时")

add_para(doc, """【开场建议】（前10分钟）""", bold=True)
add_key_point_box(doc, "不开门见山讲报表，先用一个生活化的例子建立连接：问学员\"你们公司最让你头疼的三个数据指标是什么？\"", "核心开场策略")
add_para(doc, """然后揭示：管理者之所以需要财务视角，是因为：""")
add_bullet(doc, "财务语言是商业世界的通用语言——不懂财务，就像不懂英语却要在国际商战中谈判")
add_bullet(doc, "数据驱动决策的前提是读懂数据——财务数据是经营状况的\"仪表盘\"")
add_bullet(doc, "与财务部门有效协作的前提是理解财务逻辑——否则永远是\"听天书\"")

add_para(doc, """【核心论点一：财务思维vs会计思维】（约20分钟）""", bold=True)
add_key_point_box(doc, "会计思维是\"向后看\"——记录已经发生的交易；财务思维是\"向前看\"——用数据指导未来的决策。", "核心区分")
add_para(doc, """讲解要点：""")
add_bullet(doc, "用家庭财务做类比：会计思维是记账（记流水账），财务思维是理财（让钱生钱）")
add_bullet(doc, "管理者的财务思维：不是要学会做账，而是要学会用数据做判断")
add_bullet(doc, "强调：管理者不需要成为会计，但需要理解财务逻辑")

add_para(doc, """【核心论点二：管理者必备的三个财务视角】（约20分钟）""", bold=True)
add_bullet(doc, "视角一：结果视角——通过财务结果（利润、现金流）检验业务决策的正确性")
add_bullet(doc, "视角二：过程视角——通过财务过程（周转率、毛利率）发现经营改善机会")
add_bullet(doc, "视角三：风险视角——通过财务结构（负债率、流动性）预警企业生存风险")

add_para(doc, """【互动练习：你的决策有没有财务验证？】（约20分钟）""", bold=True)
add_para(doc, """让学员回忆最近自己做的一个业务决策，然后讨论：""")
add_bullet(doc, "这个决策有没有经过财务验证？")
add_bullet(doc, "如果验证了，用了哪些财务指标？")
add_bullet(doc, "如果没有验证，原因是什么？")

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq1_table = doc.add_table(rows=4, cols=2)
faq1_table.style = 'Table Grid'
add_table_row(faq1_table, ["问题", "参考回答"], bold_first=True)
faq1_data = [
    ["会计和财务有什么区别？", "会计是信息的记录和报告（后视镜），财务是信息的分析和决策（望远镜）。管理者需要的是财务思维，不是会计技能。"],
    ["我需要学做账吗？", "不需要。专业的财务人员会做账。你的任务是读懂账本、用数据分析决策、与财务部门有效沟通。"],
    ["财务思维能帮到我什么？", "三个层面：1）避免\"拍脑袋\"决策；2）提前发现经营风险；3）与财务部门高效协作。"],
]
for row_data in faq1_data:
    add_table_row(faq1_table, row_data)
doc.add_paragraph()

add_warning_box(doc, "模块一是整门课程的\"定调\"模块。不要急于讲报表，先让学员理解\"为什么要学\"，这比\"学什么\"更重要。", "关键提醒")

# 模块二
add_heading(doc, "模块二：三张报表的底层逻辑", 3)
add_time_box(doc, "建议时长：2.5小时")

add_para(doc, """【核心概念一：资产负债表】（约45分钟）""", bold=True)
add_key_point_box(doc, "资产负债表是企业的一张\"快照\"——在某一时点，企业拥有什么（资产），欠着什么（负债），净资产是多少（权益）。", "核心金句")
add_para(doc, """讲解要点：""")
add_bullet(doc, "资产：企业拥有并能产生经济利益的资源（现金、存货、设备、品牌）")
add_bullet(doc, "负债：企业欠别人的（银行贷款、供应商货款、员工工资）")
add_bullet(doc, "权益：资产减去负债后的\"家底\"（股东投资 + 累计盈利）")
add_bullet(doc, "恒等式：资产 = 负债 + 权益（永远平衡）")

add_para(doc, """【案例分析：华为的资产负债表解读】（约15分钟）""", bold=True)
add_para(doc, """要点解读：""")
add_bullet(doc, "华为的高研发投入形成无形资产（研发费用资本化）")
add_bullet(doc, "华为不上市，权益主要是内部创业者的贡献")
add_bullet(doc, "华为的\"现金为王\"理念体现在大量的货币资金")

add_para(doc, """【核心概念二：利润表】（约45分钟）""", bold=True)
add_key_point_box(doc, "利润表是企业的一部\"电影\"——在一段时间内，企业收入多少，花了多少费用，赚了多少利润。", "核心金句")
add_para(doc, """讲解要点：""")
add_bullet(doc, "收入：卖产品的钱（营业总收入）")
add_bullet(doc, "成本：生产产品的直接代价（营业成本）")
add_bullet(doc, "费用：运营管理的代价（销售费用、管理费用、研发费用、财务费用）")
add_bullet(doc, "利润 = 收入 - 成本 - 费用（但有多个层次的利润：毛利润、营业利润、净利润）")

add_para(doc, """【关键区分：收入vs利润、规模vs效率】（约15分钟）""", bold=True)
add_bullet(doc, "收入代表\"面子\"——做多大")
add_bullet(doc, "利润代表\"里子\"——赚多少")
add_bullet(doc, "高收入不一定高利润（案例：京东vs阿里巴巴）")
add_bullet(doc, "财务思维关注效率：每投入1元，能产生多少收入和利润")

add_para(doc, """【核心概念三：现金流量表】（约30分钟）""", bold=True)
add_key_point_box(doc, "现金流量表是企业的一本\"流水账\"——在一段时间内，现金从哪来，到哪去。", "核心金句")
add_para(doc, """讲解要点：""")
add_bullet(doc, "经营现金流：主营业务产生的现金流动——\"造血能力\"")
add_bullet(doc, "投资现金流：投资活动产生的现金流动——\"放血/补血\"")
add_bullet(doc, "融资现金流：融资活动产生的现金流动——\"输血能力\"")

add_warning_box(doc, "现金流量表是三张报表中最重要的一张！很多企业的死亡不是因为亏损，而是因为现金流断裂。", "关键提醒")

add_para(doc, """【三张报表的底层关系】（约20分钟）""", bold=True)
add_key_point_box(doc, "资产负债表是基础（财务状况），利润表是结果（经营成果），现金流量表是验证（钱从哪来、到哪去）。", "核心关系")

rel_table = doc.add_table(rows=4, cols=3)
rel_table.style = 'Table Grid'
add_table_row(rel_table, ["报表", "核心问题", "时间维度"], bold_first=True)
add_table_row(rel_table, ["资产负债表", "企业有什么？值多少？", "某一时点的\"快照\""])
add_table_row(rel_table, ["利润表", "企业赚了多少？", "某一时期的\"电影\""])
add_table_row(rel_table, ["现金流量表", "现金够用吗？", "某一时期的\"流水账\""])

doc.add_paragraph()

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq2_table = doc.add_table(rows=4, cols=2)
faq2_table.style = 'Table Grid'
add_table_row(faq2_table, ["问题", "参考回答"], bold_first=True)
faq2_data = [
    ["三张报表哪个最重要？", "都重要，但如果非要选，现金流量表最重要——企业死于现金流断裂，而非亏损。"],
    ["为什么利润表盈利但现金流为负？", "这是常见的\"纸面富贵\"——利润是权责发生制，现金流是收付实现制。比如签了合同但没收到钱，利润增加但现金流没有。"],
    ["三张报表如何互相验证？", "比如：净利润高但经营现金流为负，说明利润可能\"水分\"较大；应收账款大幅增加可能虚增收入。交叉验证是财务分析的基本功。"],
]
for row_data in faq2_data:
    add_table_row(faq2_table, row_data)
doc.add_paragraph()

# 模块三
add_heading(doc, "模块三：经营分析核心指标", 3)
add_time_box(doc, "建议时长：2小时")

add_para(doc, """【指标一：毛利率】（约25分钟）""", bold=True)
add_key_point_box(doc, "毛利率 = (收入 - 成本) / 收入。毛利率反映的是产品或服务本身的盈利能力，是竞争的起点。", "核心公式")
add_para(doc, """讲解要点：""")
add_bullet(doc, "高毛利率意味着什么？——产品有溢价能力、竞争壁垒高、商业模式好")
add_bullet(doc, "低毛利率意味着什么？——可能深陷价格战、护城河低、商业模式堪忧")
add_bullet(doc, "毛利率的横向对比：同一行业，毛利率高的企业通常更值得关注")
add_bullet(doc, "毛利率的时间趋势：持续下降可能是竞争加剧或成本上升的信号")

add_para(doc, """【案例：茅台vs洋河的毛利率对比】""", bold=True)
add_para(doc, """茅台毛利率常年在90%以上，洋河在70%左右。这个差距反映了什么？""")
add_bullet(doc, "茅台的品牌溢价能力和定价权远超洋河")
add_bullet(doc, "高端白酒的护城河（品牌、稀缺性）远高于中端白酒")
add_bullet(doc, "投资角度看，茅台的盈利质量更高")

add_para(doc, """【指标二：周转率】（约25分钟）""", bold=True)
add_key_point_box(doc, "周转率 = 收入 / 资产。周转率反映的是资产的使用效率——每投入1元资产，能产生多少收入。", "核心公式")
add_para(doc, """讲解要点：""")
add_bullet(doc, "总资产周转率：企业整体资产的使用效率")
add_bullet(doc, "存货周转率：存货从入库到卖出的速度（天数 = 365 / 周转次数）")
add_bullet(doc, "应收账款周转率：应收账款收回的速度（天数 = 365 / 周转次数）")
add_bullet(doc, "周转率高意味着什么？——资产效率高、商业模式轻、现金流好")

add_para(doc, """【指标三：现金转换周期（CCC）】（约25分钟）""", bold=True)
add_key_point_box(doc, "现金转换周期 = 存货周转天数 + 应收账款周转天数 - 应付账款周转天数。CCC是经营效率的综合体现。", "核心公式")
add_para(doc, """讲解要点：""")
add_bullet(doc, "CCC越短越好——意味着钱从投入到回收的速度快")
add_bullet(doc, "CCC为负数更牛——意味着先收钱后付款（如苹果、茅台）")
add_bullet(doc, "CCC大幅延长是危险的信号——可能意味着库存积压或回款困难")

add_para(doc, """【指标四：ROE（净资产收益率）】（约25分钟）""", bold=True)
add_key_point_box(doc, "ROE = 净利润 / 股东权益。ROE是衡量股东回报的核心指标，也是杜邦分析的核心。", "核心公式")
add_para(doc, """杜邦分析三因子：""")
add_bullet(doc, "净利率（净利润/收入）——盈利能力")
add_bullet(doc, "资产周转率（收入/资产）——运营效率")
add_bullet(doc, "权益乘数（资产/权益）——财务杠杆")
add_bullet(doc, "ROE = 净利率 × 资产周转率 × 权益乘数")

add_para(doc, """【综合练习：指标分析】（约20分钟）""", bold=True)
add_para(doc, """给出一家虚拟公司的财务数据，让学员计算并分析：""")
add_bullet(doc, "计算毛利率、周转率、CCC、ROE")
add_bullet(doc, "判断各项指标的优劣")
add_bullet(doc, "提出经营改善建议")

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq3_table = doc.add_table(rows=4, cols=2)
faq3_table.style = 'Table Grid'
add_table_row(faq3_table, ["问题", "参考回答"], bold_first=True)
faq3_data = [
    ["毛利率是越高越好吗？", "不一定。要结合行业特性和竞争策略。奢侈品毛利率高但周转率低；平价零售毛利率低但周转率高。关键是\"盈利能力 × 运营效率\"的综合结果。"],
    ["ROE高一定好吗？", "不一定。高ROE可能来自高杠杆（借钱炒股），风险很大。要用杜邦分析拆解，看高ROE是来自盈利、效率还是杠杆。"],
    ["CCC为负数怎么可能？", "完全可能，且是优秀商业模式的标志。比如苹果：先收消费者钱，后付供应商款，现金在手里停留时间极长。"],
]
for row_data in faq3_data:
    add_table_row(faq3_table, row_data)
doc.add_paragraph()

# 模块四
add_heading(doc, "模块四：业务决策的财务验证", 3)
add_time_box(doc, "建议时长：2小时")

add_para(doc, """【决策一：定价决策】（约30分钟）""", bold=True)
add_key_point_box(doc, "定价不是\"成本加成\"，而是\"价值定价\"——价格由客户愿意支付的价值决定，不是由成本决定。", "核心原则")
add_para(doc, """讲解要点：""")
add_bullet(doc, "成本定价的陷阱：忽略市场供需、扼杀创新激励、容易陷入价格战")
add_bullet(doc, "价值定价的逻辑：先确定客户价值，再倒推成本结构")
add_bullet(doc, "案例：奢侈品的定价逻辑——不是\"这件衣服面料多少钱\"，而是\"这个品牌值多少钱\"")

add_para(doc, """【定价工具：价值曲线分析】""", bold=True)
add_bullet(doc, "画出行业的价值曲线（各要素的竞争焦点）")
add_bullet(doc, "找到被行业\"过度提供\"和\"提供不足\"的要素")
add_bullet(doc, "重新定义竞争要素，开创蓝海市场")

add_para(doc, """【决策二：成本结构分析】（约30分钟）""", bold=True)
add_key_point_box(doc, "成本分析的关键是区分固定成本和变动成本——固定成本是\"不管卖多少都要付的\"，变动成本是\"每多卖一个就多一份的\"。", "核心区分")
add_para(doc, """讲解要点：""")
add_bullet(doc, "固定成本：租金、工资、设备折旧——规模小时负担重，规模大时摊薄低")
add_bullet(doc, "变动成本：原材料、包装、运输——随收入同比例变化")
add_bullet(doc, "边际成本：每多生产一件产品需要增加的成本——数字经济边际成本趋近于零")
add_bullet(doc, "经营杠杆：固定成本占比高的企业，收入增长时利润弹性更大（风险也更大）")

add_para(doc, """【决策三：盈亏平衡点分析】（约30分钟）""", bold=True)
add_key_point_box(doc, "盈亏平衡点 = 固定成本 / (单价 - 单位变动成本) = 固定成本 / 边际贡献", "核心公式")
add_para(doc, """讲解要点：""")
add_bullet(doc, "边际贡献 = 单价 - 单位变动成本——每卖一件产品，对固定成本的贡献")
add_bullet(doc, "超过盈亏平衡点的收入才是真正的\"赚到的钱\"")
add_bullet(doc, "盈亏平衡分析的应用：投资决策、定价策略、规模选择")

add_para(doc, """【决策验证练习：假如你要开一家奶茶店】""", bold=True)
add_bullet(doc, "给定租金、人工、原料成本，计算盈亏平衡点")
add_bullet(doc, "分析每天需要卖出多少杯才能盈利")
add_bullet(doc, "讨论：如何通过调整成本结构或定价来降低盈亏平衡点")

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq4_table = doc.add_table(rows=4, cols=2)
faq4_table.style = 'Table Grid'
add_table_row(faq4_table, ["问题", "参考回答"], bold_first=True)
faq4_data = [
    ["成本加成定价有什么问题？", "成本加成定价的逻辑是\"我收回成本加利润\"，但客户只关心\"我得到了什么价值\"。成本加成容易导致\"内部视角\"，忽视市场和客户。"],
    ["边际成本为零怎么定价？", "这是数字经济的经典难题。常见策略：1）免费+增值服务；2）订阅制；3）平台模式（多方博弈）。核心是找到\"愿意付费的那个价值\"。"],
    ["盈亏平衡分析有什么局限？", "盈亏平衡分析是静态的，假设单价、变动成本固定。但现实中这些都可能变化。更重要的是\"敏感性分析\"——各因素变化时，盈亏平衡点如何变化。"],
]
for row_data in faq4_data:
    add_table_row(faq4_table, row_data)
doc.add_paragraph()

# 模块五
add_heading(doc, "模块五：现金为王的经营实践", 3)
add_time_box(doc, "建议时长：1.5小时")

add_para(doc, """【核心论点一：为什么现金为王】（约20分钟）""", bold=True)
add_key_point_box(doc, "利润是账面游戏，现金是真实存在。一家盈利但现金流断裂的企业，照样会死。", "核心金句")
add_para(doc, """讲解要点：""")
add_bullet(doc, "案例：某连锁零售企业，年度报表显示盈利，但因扩张过快导致现金流断裂，最终破产")
add_bullet(doc, "教训：增长是最大的风险——没有现金流支撑的增长，是在\"自杀\"")
add_bullet(doc, "商业史上，死于现金流断裂的企业远多于死于亏损的企业")

add_para(doc, """【核心论点二：现金转换周期深度解读】（约25分钟）""", bold=True)
add_para(doc, """CCC的三个组成部分：""")
add_bullet(doc, "存货周转天数：原料→在产品→成品的平均时间")
add_bullet(doc, "应收账款周转天数：从发货到收款的时间")
add_bullet(doc, "应付账款周转天数：从收货到付款的时间")
add_para(doc, """优化CCC的策略：""")
add_bullet(doc, "缩短存货周转天数：JIT、订单驱动生产")
add_bullet(doc, "缩短应收账款：强化信用管理、缩短账期、提供早付折扣")
add_bullet(doc, "延长应付账款：供应链金融、谈判更好的付款条件")

add_para(doc, """【核心论点三：经营预警信号】（约25分钟）""", bold=True)
add_key_point_box(doc, "企业经营恶化从来不是一夜之间发生的。在危机爆发前，总有预警信号。", "核心意识")
add_para(doc, """红色预警信号：""")
add_bullet(doc, "信号一：应收账款天数突然大幅增加——客户回款困难，可能存在坏账风险")
add_bullet(doc, "信号二：存货周转天数突然大幅增加——产品滞销，市场需求可能萎缩")
add_bullet(doc, "信号三：毛利率持续下降——竞争加剧、定价能力丧失或成本上升")
add_bullet(doc, "信号四：经营现金流持续为负——商业模式可能存在根本性问题")
add_bullet(doc, "信号五：供应商要求现款结算——供应商对你的信用失去信心")

add_para(doc, """【案例讨论：某公司是如何在破产边缘被救回来的】""", bold=True)
add_para(doc, """讨论要点：""")
add_bullet(doc, "管理层是如何发现现金流危机的？")
add_bullet(doc, "采取了哪些措施来改善现金流？")
add_bullet(doc, "这个案例对你的企业有什么启示？")

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq5_table = doc.add_table(rows=4, cols=2)
faq5_table.style = 'Table Grid'
add_table_row(faq5_table, ["问题", "参考回答"], bold_first=True)
faq5_data = [
    ["现金越多越好吗？", "不一定。现金太多可能说明：1）没有好的投资机会（资本效率低）；2）商业模式不健康（不敢投资）。关键是\"现金的来源和使用\"是否合理。"],
    ["如何平衡增长和现金流？", "增长需要投入，但投入需要现金流支撑。原则是：\"现金为王\"——宁可慢一点，也要确保现金流安全。在扩张时，保留至少6个月的现金储备。"],
    ["小公司如何监控现金流？", "每周更新现金流量表（不是每月）；建立预警机制（现金低于X个月成本就报警）；与银行保持良好关系（信用额度备用）。"],
]
for row_data in faq5_data:
    add_table_row(faq5_table, row_data)
doc.add_paragraph()

# 模块六
add_heading(doc, "模块六：财务与业务的协同语言", 3)
add_time_box(doc, "建议时长：1.5小时")

add_para(doc, """【核心能力一：财务语言的业务翻译】（约25分钟）""", bold=True)
add_key_point_box(doc, "财务数据的价值不在于数字本身，而在于它背后的业务含义。管理者的核心能力是\"翻译\"——将财务语言转化为业务语言。", "核心能力")
add_para(doc, """常见翻译示例：""")

trans_table = doc.add_table(rows=6, cols=3)
trans_table.style = 'Table Grid'
add_table_row(trans_table, ["财务语言", "业务翻译", "管理行动"], bold_first=True)
trans_data = [
    ["应收账款周转天数增加20天", "回款变慢，客户可能经营困难", "加强信用管理，评估客户质量"],
    ["毛利率下降5个点", "产品溢价能力下降或成本上升", "分析原因，调整定价或优化供应链"],
    ["存货周转率下降", "产品滞销，市场需求疲软", "促销去库存，调整生产计划"],
    ["ROE持续高于30%", "股东回报优秀，护城河深", "关注能否持续，分析驱动因素"],
    ["经营现金流为负", "主营业务不造血，靠输血维持", "评估商业模式，警惕流动性风险"],
]
for row_data in trans_data:
    add_table_row(trans_table, row_data)
doc.add_paragraph()

add_para(doc, """【核心能力二：与财务部门的有效协作】（约25分钟）""", bold=True)
add_para(doc, """业务部门与财务部门的常见矛盾：""")
add_bullet(doc, "业务：\"财务审批太慢，错过市场机会！\"")
add_bullet(doc, "财务：\"业务的预算超支太离谱！\"")
add_bullet(doc, "根源：双方语言不通、视角不同、目标不一致")

add_para(doc, """解决方案：建立共同的\"经营语言\"：""")
add_bullet(doc, "共同语言一：KPI体系——用统一的指标衡量业务和财务")
add_bullet(doc, "共同语言二：预算流程——业务和财务共同参与预算制定")
add_bullet(doc, "共同语言三：决策会机制——重大决策必须财务和业务共同签字")

add_para(doc, """【核心能力三：向不同对象汇报经营状况】（约20分钟）""", bold=True)
add_para(doc, """不同对象的关注重点：""")

report_table = doc.add_table(rows=4, cols=3)
report_table.style = 'Table Grid'
add_table_row(report_table, ["汇报对象", "关注重点", "汇报策略"], bold_first=True)
report_data = [
    ["老板/董事会", "战略目标达成、股东回报、风险控制", "用ROE、现金流、战略KPI说话，重点突出战略执行情况"],
    ["投资人/股东", "投资回报、增长潜力、退出路径", "用增长率、毛利率、市场份额说话，突出想象空间"],
    ["团队成员", "目标达成、个人贡献、激励机制", "用业务KPI、利润率、奖金挂钩机制说话，激活团队动力"],
]
for row_data in report_data:
    add_table_row(report_table, row_data)
doc.add_paragraph()

add_para(doc, """【角色扮演练习：业务vs财务对话】""", bold=True)
add_para(doc, """场景：销售总监要求增加200万市场费用，财务总监质疑ROI。""")
add_para(doc, """让两组学员分别扮演销售总监和财务总监，进行对话。""")
add_bullet(doc, "销售总监如何用财务语言说服财务总监？")
add_bullet(doc, "财务总监如何用业务视角理解销售需求？")
add_bullet(doc, "双方如何达成共识？")

add_para(doc, """【常见问题与回答（FAQ）】""", bold=True)
faq6_table = doc.add_table(rows=4, cols=2)
faq6_table.style = 'Table Grid'
add_table_row(faq6_table, ["问题", "参考回答"], bold_first=True)
faq6_data = [
    ["业务和财务总是吵架怎么办？", "建立共同语言是关键。用\"经营指标\"（毛利率、周转率、ROE）取代\"会计术语\"，让双方在同一频道对话。定期的\"经营分析会\"可以培养共同语言。"],
    ["如何让财务理解业务？", "邀请财务同事参与业务会议；让财务了解市场一线的情况；建立\"业务影子\"计划，财务人员定期跟一线跑客户。"],
    ["如何让业务理解财务？", "用业务语言讲财务；每个财务概念都配上业务案例；定期进行\"财务经营思维\"培训。核心是\"翻译\"而不是\"灌输\"。"],
]
for row_data in faq6_data:
    add_table_row(faq6_table, row_data)
doc.add_paragraph()

doc.add_page_break()

# ===== 第三部分：讲师工具 =====
add_heading(doc, "第三部分 | 讲师工具", 1)

add_heading(doc, "一、授课时间轴参考", 2)

add_heading(doc, "第一天时间轴", 3)
day1_table = doc.add_table(rows=12, cols=4)
day1_table.style = 'Table Grid'
add_table_row(day1_table, ["时间段", "内容", "时长", "备注"], bold_first=True)
day1_data = [
    ["08:30-09:00", "签到与准备", "30分钟", "检查设备、发放材料"],
    ["09:00-09:15", "开场与破冰", "15分钟", "建立连接、明确目标"],
    ["09:15-10:30", "模块一：经营思维导论", "75分钟", "含互动练习"],
    ["10:30-10:45", "茶歇", "15分钟", "—"],
    ["10:45-12:15", "模块二：三张报表（上）", "90分钟", "资产负债表+利润表"],
    ["12:15-13:30", "午餐与休息", "75分钟", "—"],
    ["13:30-15:00", "模块二：三张报表（下）", "90分钟", "现金流量表+三表关系"],
    ["15:00-15:15", "茶歇", "15分钟", "—"],
    ["15:15-16:45", "模块三：核心指标（上）", "90分钟", "毛利率+周转率"],
    ["16:45-17:00", "茶歇", "15分钟", "—"],
    ["17:00-18:00", "模块三：核心指标（下）+今日总结", "60分钟", "CCC+ROE+总结"],
    ["18:00", "第一天结束", "—", "—"],
]
for row_data in day1_data:
    add_table_row(day1_table, row_data)
doc.add_paragraph()

add_heading(doc, "第二天时间轴", 3)
day2_table = doc.add_table(rows=12, cols=4)
day2_table.style = 'Table Grid'
add_table_row(day2_table, ["时间段", "内容", "时长", "备注"], bold_first=True)
day2_data = [
    ["08:30-09:00", "签到与答疑", "30分钟", "回答第一天遗留问题"],
    ["09:00-10:30", "模块四：业务决策财务验证（上）", "90分钟", "定价决策+成本结构"],
    ["10:30-10:45", "茶歇", "15分钟", "—"],
    ["10:45-12:15", "模块四：业务决策财务验证（下）", "90分钟", "盈亏平衡点"],
    ["12:15-13:30", "午餐与休息", "75分钟", "—"],
    ["13:30-15:00", "模块五：现金为王的经营实践", "90分钟", "含预警练习"],
    ["15:00-15:15", "茶歇", "15分钟", "—"],
    ["15:15-16:30", "模块六：财务与业务的协同语言", "75分钟", "含角色扮演"],
    ["16:30-17:30", "课程总结与行动计划", "60分钟", "知识回顾+行动制定"],
    ["17:30-18:00", "结业仪式与反馈", "30分钟", "颁发证书、收集反馈"],
    ["18:00", "课程结束", "—", "—"],
]
for row_data in day2_data:
    add_table_row(day2_table, row_data)
doc.add_paragraph()

add_heading(doc, "二、板书/投影建议", 2)

add_heading(doc, "1. 推荐板书框架", 3)
add_para(doc, """【模块一板书：财务思维框架图】""", bold=True)
add_golden_box(doc, "财务视角三问：\n1. 这个决策的结果如何？（结果视角 → 利润/现金流）\n2. 这个决策的效率如何？（过程视角 → 周转率/毛利率）\n3. 这个决策的风险如何？（风险视角 → 负债率/流动性）")

add_para(doc, """【模块二板书：三表关系图】""", bold=True)
add_golden_box(doc, "资产负债表（基础）\n      ↕ 经营活动\n利润表（结果）\n      ↕ 现金变动\n现金流量表（验证）")

add_para(doc, """【模块三板书：杜邦分析图】""", bold=True)
add_golden_box(doc, "ROE = 净利率 × 资产周转率 × 权益乘数\n      ↓           ↓           ↓\n   盈利能力      运营效率     财务杠杆\n   （产品）      （管理）     （风险）")

add_para(doc, """【模块四板书：盈亏平衡图】""", bold=True)
add_golden_box(doc, "盈亏平衡点 = 固定成本 / 边际贡献\n边际贡献 = 单价 - 单位变动成本")

add_para(doc, """【模块五板书：CCC计算公式】""", bold=True)
add_golden_box(doc, "现金转换周期 = 存货周转天数 + 应收账款周转天数 - 应付账款周转天数\nCCC越短 → 现金效率越高 → 经营越健康")

add_heading(doc, "2. 投影材料建议", 3)
proj_table = doc.add_table(rows=7, cols=2)
proj_table.style = 'Table Grid'
add_table_row(proj_table, ["模块", "建议投影内容"], bold_first=True)
proj_data = [
    ["模块一", "课程整体框架图（让学员了解全貌）"],
    ["模块二", "三张报表模板（空白报表供标注讲解）"],
    ["模块三", "指标计算公式汇总表（方便学员记忆）"],
    ["模块四", "盈亏平衡分析表（Excel模板现场演示）"],
    ["模块五", "企业危机案例视频或图文材料"],
    ["模块六", "不同汇报对象的关注重点对比表"],
]
for row_data in proj_data:
    add_table_row(proj_table, row_data)
doc.add_paragraph()

add_heading(doc, "三、道具/材料清单", 2)

props_table = doc.add_table(rows=10, cols=2)
props_table.style = 'Table Grid'
add_table_row(props_table, ["道具名称", "用途与说明"], bold_first=True)
props_data = [
    ["白板/翻页纸", "用于板书关键框架：杜邦分析、现金转换周期、盈亏平衡图等"],
    ["投影设备", "展示PPT、案例数据、课程框架图"],
    ["计时器", "控制各环节时间，特别是互动讨论环节"],
    ["马克笔（红、蓝、黑、绿）", "用于白板书写和框架图标注"],
    ["学员手册", "每人一份，用于记录要点和做练习"],
    ["三张报表案例打印材料", "A3纸质，每小组一份，用于小组研讨"],
    ["金句卡片", "每模块结束发给学员，作为课后复习材料"],
    ["A4纸（充足）", "用于学员做练习（指标计算、盈亏平衡分析等）"],
    ["便签纸（不同颜色）", "用于小组讨论时的观点汇总和展示"],
    ["计算器（学员自备或统一提供）", "用于指标计算练习"],
]
for row_data in props_data:
    add_table_row(props_table, row_data)
doc.add_paragraph()

add_heading(doc, "四、Q&A准备", 2)

add_heading(doc, "1. 高频问题汇总", 3)

qna_table = doc.add_table(rows=11, cols=2)
qna_table.style = 'Table Grid'
add_table_row(qna_table, ["问题", "参考回答要点"], bold_first=True)
qna_data = [
    ["财务思维和会计思维的区别是什么？", "会计是记录（后视镜），财务是决策（望远镜）。管理者需要的是后者。"],
    ["非财务背景能学会吗？", "本课程专为零基础设计，只需要基本的商业常识。"],
    ["数学基础不好怎么办？", "只需要加减乘除。我们强调的是商业直觉，不是计算能力。"],
    ["学完能直接用在工作中吗？", "当然！每个模块都有工具和模板，可以直接应用到工作中。"],
    ["如何快速读懂财务报表？", "先看现金流量表（最重要），再看资产负债表和利润表。交叉验证三张报表。"],
    ["毛利率多少算合理？", "不同行业差异很大。要进行横向对比（同行）和纵向对比（自身历史）。"],
    ["企业最关键的财务指标是什么？", "没有\"最关键\"，要看具体情况。但ROE、毛利率、现金转换周期是最常用的。"],
    ["盈亏平衡点有什么用？", "用于决策验证：投资决策、定价策略、规模选择。知道\"卖多少才能不亏\"。"],
    ["如何改善现金流？", "三个方向：加快回款、延长付款、优化库存。核心是提高现金转换效率。"],
    ["业务和财务总吵架怎么办？", "建立共同语言是关键。用经营指标取代会计术语，让双方在同一频道对话。"],
]
for row_data in qna_data:
    add_table_row(qna_table, row_data)
doc.add_paragraph()

add_heading(doc, "2. 延伸问题准备", 3)
add_para(doc, """以下问题可能在学员深入追问时出现，提前准备：""")
add_bullet(doc, "杜邦分析的局限性是什么？——ROE可以被财务杠杆扭曲，需要结合其他指标")
add_bullet(doc, "为什么有的企业现金流为负但盈利很好？——扩张期企业，现金流紧张但利润表漂亮")
add_bullet(doc, "如何评估互联网公司的财务数据？——互联网公司常用GMV、MAU等非财务指标")
add_bullet(doc, "并购中的财务尽调主要看什么？——资产负债表的\"黑匣子\"（或有负债、诉讼等）")

doc.add_page_break()

# ===== 第四部分：评估支持 =====
add_heading(doc, "第四部分 | 评估支持", 1)

add_heading(doc, "一、学员表现观察要点", 2)

add_heading(doc, "1. 课堂参与度观察", 3)
participation_table = doc.add_table(rows=6, cols=3)
participation_table.style = 'Table Grid'
add_table_row(participation_table, ["观察维度", "观察指标", "评估标准"], bold_first=True)
participation_data = [
    ["出勤守时", "是否全程参与、是否按时出席", "5分=全程参与；3分=偶尔迟到早退；1分=缺勤较多"],
    ["课堂互动", "举手发言、小组讨论参与度", "5分=主动积极；3分=被动回应；1分=基本不参与"],
    ["提问质量", "问题是否切中要害、是否有独立思考", "5分=问题深刻；3分=一般性提问；1分=无关问题"],
    ["练习投入度", "是否认真完成练习、是否深度思考", "5分=高质量完成；3分=基本完成；1分=敷衍了事"],
    ["小组贡献", "是否为小组讨论贡献有价值的观点", "5分=核心贡献者；3分=一般参与者；1分=旁观者"],
]
for row_data in participation_data:
    add_table_row(participation_table, row_data)
doc.add_paragraph()

add_heading(doc, "2. 学习效果即时评估", 3)
add_para(doc, """【每个模块结束时的快速检验】""", bold=True)
add_bullet(doc, "让学员用一句话总结本模块最重要的收获")
add_bullet(doc, "让学员举手回答一个与本模块相关的基础问题")
add_bullet(doc, "观察学员的表情和反应——是\"懂了\"还是\"懵了\"")

add_para(doc, """【两天课程的结业测试】""", bold=True)
test_table = doc.add_table(rows=4, cols=2)
test_table.style = 'Table Grid'
add_table_row(test_table, ["测试类型", "具体形式"], bold_first=True)
test_data = [
    ["知识测试", "10道选择题（概念+计算），检验核心知识掌握情况"],
    ["案例分析", "给出一家虚拟公司的财务数据，让学员进行分析并提出建议"],
    ["行动计划", "让学员写下课程后将在工作中实际应用的三个行动"],
]
for row_data in test_data:
    add_table_row(test_table, row_data)
doc.add_paragraph()

add_heading(doc, "二、评分标准说明", 2)

score_table = doc.add_table(rows=6, cols=4)
score_table.style = 'Table Grid'
add_table_row(score_table, ["评分维度", "权重", "评分标准", "备注"], bold_first=True)
score_data = [
    ["出勤与守时", "10%", "全程参与=10分；偶尔迟到=7分；缺勤较多=5分以下", "基础分"],
    ["课堂参与", "20%", "积极互动=20分；被动回应=15分；不参与=10分以下", "观察评估"],
    ["练习完成", "30%", "高质量完成=30分；基本完成=25分；敷衍=15分以下", "作业评估"],
    ["案例分析", "25%", "分析深入且有独到见解=25分；分析基本正确=20分；分析有误=15分以下", "作业评估"],
    ["行动计划", "15%", "计划具体可执行=15分；计划较模糊=10分；无计划=5分以下", "作业评估"],
]
for row_data in score_data:
    add_table_row(score_table, row_data)
doc.add_paragraph()

add_para(doc, """【评分等级参考】""", bold=True)
grade_table = doc.add_table(rows=5, cols=2)
grade_table.style = 'Table Grid'
add_table_row(grade_table, ["等级", "分数范围"], bold_first=True)
grade_data = [
    ["优秀", "90-100分"],
    ["良好", "75-89分"],
    ["合格", "60-74分"],
    ["待提升", "40-59分"],
    ["不及格", "40分以下"],
]
for row_data in grade_data:
    add_table_row(grade_table, row_data)
doc.add_paragraph()

add_heading(doc, "三、反馈话术参考", 2)

add_heading(doc, "1. 课堂反馈话术", 3)
feedback_phrases = [
    ("鼓励参与", "\"你的问题很有价值，说明你在深入思考\"", "\"这个角度很独特，能不能再详细说说？\""),
    ("肯定表现", "\"你对这个概念的理解很准确\"", "\"你的分析逻辑很清晰，抓住了关键点\""),
    ("引导思考", "\"如果从这个角度看会怎样？\"", "\"你有没有考虑过另一个因素？\""),
    ("纠正错误", "\"你的思路是对的，但这个计算可以再想想\"", "\"这个案例还有另一个角度，我们可以一起看看\""),
]

for phrase_type, good, better in feedback_phrases:
    para = doc.add_paragraph()
    run = para.add_run(f"【{phrase_type}】")
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(11)
    run.font.bold = True
    add_bullet(doc, f"肯定：{good}")
    add_bullet(doc, f"更好：{better}")

add_heading(doc, "2. 课程结束反馈话术", 3)
add_para(doc, """【结业时的总结话术】""", bold=True)
add_golden_box(doc, "\"两天的学习结束了，但财务经营思维的修炼才刚刚开始。记住今天的三个收获：\n1. 财务语言是商业的通用语言\n2. 用数据验证决策，而不是用直觉\n3. 持续关注现金流，它是企业的生命线\n\n期待大家在工作中实践这些方法，用财务思维做出更明智的决策。\"")

add_para(doc, """【反馈收集话术】""", bold=True)
add_golden_box(doc, "\"您的反馈对课程改进非常重要。无论好评还是建议，都是我们进步的动力。请填写反馈表，您的真实想法是我们最珍贵的礼物。\"")

doc.add_page_break()

# ===== 第五部分：课程迭代 =====
add_heading(doc, "第五部分 | 课程迭代", 1)

add_heading(doc, "一、学员反馈收集方法", 2)

feedback_method_table = doc.add_table(rows=5, cols=2)
feedback_method_table.style = 'Table Grid'
add_table_row(feedback_method_table, ["反馈类型", "收集方法"], bold_first=True)
method_data = [
    ["课程内容反馈", "课后问卷：内容深度是否合适、节奏是否合理、案例是否贴切"],
    ["讲师表现反馈", "课后问卷：讲解是否清晰、互动是否充分、时间控制是否得当"],
    ["整体满意度", "课后问卷：NPS评分（0-10分）、会推荐给同事吗、会再次参加吗"],
    ["行为改变跟踪", "1个月后跟进：学员是否在实际工作中使用了课程工具"],
]
for row_data in method_data:
    add_table_row(feedback_method_table, row_data)
doc.add_paragraph()

add_heading(doc, "二、课程改进建议", 2)

add_heading(doc, "1. 常见问题改进方向", 3)
improve_table = doc.add_table(rows=5, cols=3)
improve_table.style = 'Table Grid'
add_table_row(improve_table, ["问题类型", "可能原因", "改进方向"], bold_first=True)
improve_data = [
    ["学员反映节奏太快", "内容过多或讲解不够通俗", "精简内容、增加类比、放慢节奏"],
    ["学员反映案例不够贴切", "案例与学员行业不匹配", "收集学员背景，增加行业定制案例"],
    ["互动环节不够热烈", "问题设计或分组不当", "优化问题设计、调整分组策略"],
    ["部分概念难以理解", "讲解方式需要调整", "增加生活类比、提供更多练习机会"],
]
for row_data in improve_data:
    add_table_row(improve_table, row_data)
doc.add_paragraph()

add_heading(doc, "2. 课程迭代记录表", 3)
iteration_table = doc.add_table(rows=6, cols=5)
iteration_table.style = 'Table Grid'
add_table_row(iteration_table, ["迭代版本", "日期", "主要改进", "改进原因", "效果评估"], bold_first=True)
for i in range(5):
    iteration_table.add_row()
doc.add_paragraph()

add_heading(doc, "三、常见问题记录表", 2)

prob_table = doc.add_table(rows=8, cols=4)
prob_table.style = 'Table Grid'
add_table_row(prob_table, ["序号", "问题描述", "出现场景", "解决方案"], bold_first=True)
for i in range(7):
    prob_table.add_row()
doc.add_paragraph()

add_para(doc, """【记录说明】""", bold=True)
add_bullet(doc, "每期课程结束后，及时记录学员提出的典型问题")
add_bullet(doc, "分析问题的类型：是概念不清、案例不匹配、还是讲解方式有问题")
add_bullet(doc, "制定改进措施，并在下一期课程中实施")
add_bullet(doc, "定期回顾问题记录，发现共性规律，持续优化课程内容")

doc.add_page_break()

# ===== 附录 =====
add_heading(doc, "附录一：核心框架回顾", 1)

summary_table = doc.add_table(rows=7, cols=4)
summary_table.style = 'Table Grid'
add_table_row(summary_table, ["模块", "核心概念", "核心问题", "核心工具"], bold_first=True)
summary_data = [
    ["模块一", "财务思维vs会计思维", "为什么管理者需要财务视角？", "财务视角三问"],
    ["模块二", "三张报表底层逻辑", "如何读懂财务报表？", "资产负债表、利润表、现金流量表"],
    ["模块三", "经营分析核心指标", "如何用指标分析经营状况？", "毛利率、周转率、CCC、ROE"],
    ["模块四", "业务决策财务验证", "如何用财务验证业务决策？", "盈亏平衡分析、价值曲线"],
    ["模块五", "现金为王的经营实践", "如何预警经营风险？", "现金转换周期、预警信号"],
    ["模块六", "财务与业务的协同语言", "如何实现财务业务协同？", "翻译能力、共同语言"],
]
for row_data in summary_data:
    add_table_row(summary_table, row_data)
doc.add_paragraph()

add_para(doc, """【课程终极目标】""", bold=True)
goal_box = doc.add_table(rows=1, cols=1)
goal_box.style = 'Table Grid'
cell = goal_box.cell(0, 0)
cell.text = ""
para = cell.paragraphs[0]
run = para.add_run("学员能够理解并运用财务经营思维，来指导业务决策、预警经营风险、实现与财务部门的高效协同，成为具备全局经营视野的管理者。")
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
run.font.size = Pt(11)
doc.add_paragraph()

add_para(doc, """【讲师寄语】""", bold=True)
add_para(doc, """"财务思维不是财务人员的专利，而是每一位管理者的必备能力。

两天的学习，我们一起走过了从\"畏财务\"到\"用财务\"的旅程。希望大家带走的不只是知识，更是一种新的思维方式——用数据说话、用指标验证、用财务逻辑指导决策。

记住：财务不是业务的对面，而是业务的镜子。用财务思维看业务，你会看到更真实的商业世界。

感谢您成为这场课程的讲师。愿您带着这份方法论，帮助更多的管理者建立财务经营思维。

——罗宏伟""")

doc.add_paragraph()
doc.add_paragraph()

# ===== 结束语 =====
end_para = doc.add_paragraph()
end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end_para.add_run("— 完 —")
end_run.font.name = 'Microsoft YaHei'
end_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
end_run.font.size = Pt(14)
end_run.font.bold = True
end_run.font.color.rgb = RGBColor(128, 128, 128)

# ===== 保存文档 =====
output_dir = "D:/新课开发/管理学/30-财务经营思维/讲师手册"
import os
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

output_path = f"{output_dir}/讲师手册_财务经营思维.docx"
doc.save(output_path)
print(f"讲师手册已保存到: {output_path}")
print(f"文档生成完成！")
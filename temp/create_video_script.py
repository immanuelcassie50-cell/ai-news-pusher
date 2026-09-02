# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm

doc = Document()

# ---- Page setup ----
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ---- Styles helper ----
def set_run_font(run, font_name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = font_name
    run._r.get_or_add_rPr().set(
        qn('w:eastAsia'), font_name
    )
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, size=16, bold=True, color=(31, 56, 100)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_para(doc, text, size=11, bold=False, space_before=0, space_after=8, indent=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_label_value(doc, label, value, label_size=11, value_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_run_font(r1, size=label_size, bold=True, color=(31, 56, 100))
    r2 = p.add_run(value)
    set_run_font(r2, size=value_size)
    return p

def add_table_row(table, col1, col2, col3='', col4='',
                  bold1=False, bg_color=None):
    row = table.add_row()
    cells = [row.cells[0], row.cells[1]]
    if col3:
        cells.append(row.cells[2])
    if col4:
        cells.append(row.cells[3])

    texts = [col1, col2, col3, col4]
    for i, (cell, txt) in enumerate(zip(cells, texts)):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        is_bold = (i == 0 and bold1) or (i > 0 and False)
        set_run_font(run, size=9.5, bold=is_bold)
        cell.width = Inches(1.5) if i == 0 else Inches(3.5)
        if bg_color:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)

# ============================================================
# TITLE PAGE
# ============================================================
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(48)
title_para.paragraph_format.space_after = Pt(12)
tr = title_para.add_run('注意力管理')
set_run_font(tr, size=28, bold=True, color=(31, 56, 100))

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(6)
sr = sub_para.add_run('课程介绍视频脚本')
set_run_font(sr, size=16, bold=False, color=(68, 114, 196))

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.paragraph_format.space_after = Pt(6)
ir = info_para.add_run('时长：3分钟  |  分镜：10个  |  旁白：中文')
set_run_font(ir, size=11, color=(89, 89, 89))

doc.add_paragraph()

# ============================================================
# SECTION 1: 创作说明
# ============================================================
add_heading(doc, '一、创作说明', size=14)

add_para(doc,
    '本视频为《注意力管理》课程的概念介绍片，面向企业内训受众，'
    '旨在传递课程价值、激发学习兴趣。风格参考方太文化研究院对外宣传片调性：'
    '专业而有温度，接地气而不浅薄，有洞察而不说教。',
    size=10.5, space_after=6)

# Meta table
meta_table = doc.add_table(rows=5, cols=2)
meta_table.style = 'Table Grid'
meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

meta_data = [
    ('片名', '《注意力管理》课程介绍视频'),
    ('时长', '3分00秒'),
    ('调性', '专业、温暖、接地气、有洞察'),
    ('参考风格', '方太文化研究院对外宣传片'),
    ('目标受众', '企业管理者、知识工作者、职场新人'),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].text = ''
    row.cells[1].text = ''
    r1 = row.cells[0].paragraphs[0].add_run(k)
    set_run_font(r1, size=10, bold=True, color=(31, 56, 100))
    row.cells[0].width = Inches(1.5)
    r2 = row.cells[1].paragraphs[0].add_run(v)
    set_run_font(r2, size=10)
    row.cells[1].width = Inches(4.5)

doc.add_paragraph()

# ============================================================
# SECTION 2: 分镜脚本
# ============================================================
add_heading(doc, '二、分镜脚本', size=14)
add_para(doc,
    '以下分镜按时间线排列，每镜包含：镜号、画面描述、镜头类型、时长、旁白、视觉素材建议。',
    size=10, color=(89, 89, 89), space_after=10)

# Scene table
scene_headers = ['镜号', '时间', '画面描述', '镜头类型', '时长', '旁白', '视觉素材']
scene_table = doc.add_table(rows=1, cols=7)
scene_table.style = 'Table Grid'
hdr_row = scene_table.rows[0]
header_texts = ['镜号', '时间', '画面描述', '镜头类型', '时长', '旁白', '视觉素材']
for i, (cell, txt) in enumerate(zip(hdr_row.cells, header_texts)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_run_font(run, size=9.5, bold=True, color=(255, 255, 255))
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F3864')
    tcPr.append(shd)

# Scenes data
scenes = [
    {
        'no': '01',
        'time': '0:00-0:10',
        'description': '城市清晨，车流人流穿梭。上班族戴着耳机、低头看手机，表情疲惫而茫然。',
        'shot': '大远景航拍 + 推镜头',
        'duration': '10秒',
        'voiceover': '（低沉旁白）\n这个时代，不缺信息，缺的是专注。',
        'visual': '城市航拍视频素材，需后期调色偏冷调。快节奏剪辑。',
    },
    {
        'no': '02',
        'time': '0:10-0:20',
        'description': '特写：一个人坐在办公桌前，面前三块屏幕、七八个APP同时运转。微信弹窗、邮件提醒、新闻推送同时亮起。',
        'shot': '特写 + 中景切换',
        'duration': '10秒',
        'voiceover': '（冷静陈述）\n注意力，被切成碎片。',
        'visual': '屏幕录制素材 + 合成特效，突出"碎片感"。',
    },
    {
        'no': '03',
        'time': '0:20-0:35',
        'description': '一组对比画面：左边的职场人焦头烂额、加班到深夜；右边的职场人从容不迫、准时下班，生活充实。',
        'shot': '分屏对比',
        'duration': '15秒',
        'voiceover': '同样的8小时，为什么有人高效，有人疲惫？\n答案不在时间管理，而在于——注意力管理。',
        'visual': '情景短片素材（可版权授权或情景演绎），左右分屏对比效果。',
    },
    {
        'no': '04',
        'time': '0:35-0:55',
        'description': '一位讲师站在培训讲台上，背景是简洁的企业大学场景。讲师目光温和，语气笃定。PPT上显示课程名称"注意力管理"。',
        'shot': '中近景，平视',
        'duration': '20秒',
        'voiceover': '我叫[讲师名]，[10年企业管理经验]。\n过去8年，我为超过200家企业提供注意力管理培训。\n今天，用3分钟，帮你重建专注力。',
        'visual': '讲师出镜视频（或后期配音）。需企业培训场景背景画面。',
    },
    {
        'no': '05',
        'time': '0:55-1:20',
        'description': '动画演示：大脑接收信息的"漏斗模型"——海量信息涌入，只有10%进入短时记忆，1%最终转化为有效输出。配合轻快的节奏。',
        'shot': '动画/二维角色',
        'duration': '25秒',
        'voiceover': '大脑每天处理约74GB信息，但真正记住的，不到1%。\n不是你不努力，是大脑的带宽天生有限。\n注意力管理，是让有限带宽，用在刀刃上。',
        'visual': '二维动画（漏斗模型示意图），建议用Motion Graphics风格，明快配色。',
    },
    {
        'no': '06',
        'time': '1:20-1:50',
        'description': '三个真实学员案例（企业高管、中层管理者、一线员工），每人一句话总结收获。配合字幕卡。',
        'shot': '采访式镜头 + 字幕卡',
        'duration': '30秒',
        'voiceover': '（三位真实学员原声或旁白配音）\n"我以前一天回200封邮件，现在控制在20封。"\n"学会了单线程工作法，效率翻倍。"\n"睡眠好了，焦虑少了。"\n——他们做对了什么？注意力训练。',
        'visual': '真实学员采访视频（如无法获取，用图文人物形象代替）。字幕卡设计简洁。',
    },
    {
        'no': '07',
        'time': '1:50-2:20',
        'description': '课程内容四模块展示：认知重塑 → 场景训练 → 习惯养成 → 持续迭代。每个模块配一个核心金句。',
        'shot': '动态图表 + 字幕',
        'duration': '30秒',
        'voiceover': '这门课，帮你从四个维度重建注意力：\n认知重塑——理解大脑底层逻辑；\n场景训练——在真实工作场景中练习；\n习惯养成——让专注成为本能；\n持续迭代——建立长效注意力资产。',
        'visual': '四宫格/时间轴动态图表，简洁专业。每模块一个图标代表。',
    },
    {
        'no': '08',
        'time': '2:20-2:40',
        'description': '一位学员从上课前状态（焦虑、低效）到上课后状态（专注、从容）的转变过程，以时间流逝方式呈现。',
        'shot': '转场蒙太奇',
        'duration': '20秒',
        'voiceover': '改变，从不是一夜之间。\n但每天25分钟，持续21天，你会看到变化。\n专注力，是可以训练的。',
        'visual': '人物变化对比视频/图片素材，建议用时间流逝（Timelapse）风格呈现转变。',
    },
    {
        'no': '09',
        'time': '2:40-2:55',
        'description': '课程配套资源展示：视频课 + 练习册 + 21天训练营 + 社群陪伴。',
        'shot': '平铺展示',
        'duration': '15秒',
        'voiceover': '课程包含：12节高清视频课，配套练习册，\n以及21天训练营——有人陪你一起练。',
        'visual': '课程产品截图/界面设计展示，配简洁图标。',
    },
    {
        'no': '10',
        'time': '2:55-3:00',
        'description': '结束画面：课程名称 + 讲师头像 + 二维码 + 品牌标识。背景留白，简洁有力。',
        'shot': '定格/近景',
        'duration': '5秒',
        'voiceover': '注意力在哪里，成长就在哪里。\n从今天开始，训练你的注意力。',
        'visual': '结束页设计模板，含二维码、课程名、机构logo。',
    },
]

for scene in scenes:
    row = scene_table.add_row()
    vals = [
        scene['no'],
        scene['time'],
        scene['description'],
        scene['shot'],
        scene['duration'],
        scene['voiceover'],
        scene['visual'],
    ]
    widths = [0.4, 0.7, 1.8, 0.9, 0.5, 1.6, 1.6]
    for i, (cell, txt) in enumerate(zip(row.cells, vals)):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        is_hdr = (i == 0)
        set_run_font(run, size=8.5, bold=is_hdr)
        cell.width = Inches(widths[i])

doc.add_paragraph()

# ============================================================
# SECTION 3: 旁白稿全文
# ============================================================
add_heading(doc, '三、旁白稿全文', size=14)
add_para(doc, '（可直接用于配音录制，建议语速：约150字/分钟，全文约450字）', size=10, color=(89,89,89), space_after=8)

voiceover_sections = [
    ('开篇（0:00-0:20）',
     '这个时代，不缺信息，缺的是专注。\n同样8小时，为什么有人高效，有人疲惫？\n答案不在时间管理，而在于——注意力管理。'),
    ('讲师介绍（0:20-0:40）',
     '我叫[讲师名]，[10年企业管理经验]。\n过去8年，我为超过200家企业提供注意力管理培训。\n今天，用3分钟，帮你重建专注力。'),
    ('核心洞察（0:40-1:15）',
     '大脑每天处理约74GB信息，但真正记住的，不到1%。\n不是你不努力，是大脑的带宽天生有限。\n注意力管理，是让有限带宽，用在刀刃上。'),
    ('学员证言（1:15-1:45）',
     '"我以前一天回200封邮件，现在控制在20封。"\n"学会了\'单线程\'工作法，效率翻倍。"\n"睡眠好了，焦虑少了。"\n——他们做对了什么？注意力训练。'),
    ('课程内容（1:45-2:15）',
     '这门课，帮你从四个维度重建注意力：\n认知重塑——理解大脑底层逻辑；\n场景训练——在真实工作场景中练习；\n习惯养成——让专注成为本能；\n持续迭代——建立长效注意力资产。'),
    ('改变路径（2:15-2:35）',
     '改变，从不是一夜之间。\n但每天25分钟，持续21天，你会看到变化。\n专注力，是可以训练的。'),
    ('课程资源（2:35-2:50）',
     '课程包含：12节高清视频课，配套练习册，\n以及21天训练营——有人陪你一起练。'),
    ('结尾收束（2:50-3:00）',
     '注意力在哪里，成长就在哪里。\n从今天开始，训练你的注意力。'),
]

for section_title, section_text in voiceover_sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(section_title)
    set_run_font(r1, size=10.5, bold=True, color=(31, 56, 100))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Inches(0.3)
    r2 = p2.add_run(section_text)
    set_run_font(r2, size=10.5)

doc.add_paragraph()

# ============================================================
# SECTION 4: 视觉素材清单
# ============================================================
add_heading(doc, '四、视觉素材清单', size=14)
add_para(doc,
    '以下素材建议在视频制作前提前准备，部分可通过图库授权或自行拍摄获取。',
    size=10, color=(89,89,89), space_after=8)

material_headers = ['类别', '具体素材', '来源建议', '优先级']
mat_table = doc.add_table(rows=1, cols=4)
mat_table.style = 'Table Grid'
hdr_row = mat_table.rows[0]
for i, (cell, txt) in enumerate(zip(hdr_row.cells, material_headers)):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    set_run_font(run, size=9.5, bold=True, color=(255, 255, 255))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F3864')
    tcPr.append(shd)

materials = [
    ('开场视频', '城市清晨航拍镜头，30秒', 'Pexels / Pixabay 免费授权', '必选'),
    ('屏幕录制', '多APP弹窗同时出现画面', '自行录屏或AfterEffects合成', '必选'),
    ('情景短片', '高效vs低效对比情景剧，20秒', '版权授权或自行拍摄', '推荐'),
    ('讲师出镜', '讲师培训课堂中近景，20秒', '专业拍摄，3机位建议', '必选'),
    ('动画制作', '漏斗模型/四宫格图表动画', 'Motion Graphics设计师制作', '必选'),
    ('学员证言', '3位真实学员采访视频，各20秒', '学员授权录制', '推荐'),
    ('界面截图', '课程产品界面（视频课/练习册/社群）', '课程平台截图', '推荐'),
    ('结束页', '课程名+二维码+机构logo设计', '设计师设计，AI工具辅助', '必选'),
    ('背景音乐', '轻快节奏背景音乐，全程3分钟', 'Epidemic Sound / Artlist授权', '必选'),
    ('音效', '转场音效、强调音效', 'Logic Pro / AE音效库', '可选'),
]

for mat in materials:
    row = mat_table.add_row()
    for i, (cell, txt) in enumerate(zip(row.cells, mat)):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        set_run_font(run, size=9, bold=(i==0))
        cell.width = Inches(1.0) if i == 0 else Inches(1.8)

doc.add_paragraph()

# ============================================================
# SECTION 5: 制作注意事项
# ============================================================
add_heading(doc, '五、制作注意事项', size=14)

notes = [
    ('节奏把控', '整体节奏"前慢后快"——前30秒用真实场景引发共鸣，中段动画加速认知输出，结尾收束干脆，不拖沓。'),
    ('音乐设计', '背景音乐全程约65分贝，旁白出现时音乐淡出10%，旁白结束后淡入，避免BGM压人声。'),
    ('字体规范', '所有字幕字体使用思源黑体（Regular/Medium），避免使用宋体。字幕颜色：白色带1px描边。'),
    ('品牌一致性', '片头片尾使用统一品牌色（约#1F3864），视觉元素不超过3种颜色。'),
    ('审核要点', '发布前确认：讲师姓名/机构名称/二维码/版权音乐授权证明均无错误。'),
]

for title, content in notes:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f'◆ {title}：')
    set_run_font(r1, size=10.5, bold=True, color=(31, 56, 100))
    r2 = p.add_run(content)
    set_run_font(r2, size=10.5)

doc.add_paragraph()

# ---- Footer line ----
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.paragraph_format.space_before = Pt(24)
fr = footer_para.add_run('— 完 —')
set_run_font(fr, size=10, color=(150, 150, 150))

# ---- Save ----
output_path = 'D:/新课开发/工作手册/注意力管理/完整课程包/08-成果demo/VideoScript.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')

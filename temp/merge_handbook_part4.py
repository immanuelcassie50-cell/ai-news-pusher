#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# ========== Tool Templates ==========
tools = doc.add_paragraph()
r = tools.add_run('工具模板')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

# 附录一：决策卡标准模板
doc.add_paragraph()
app1 = doc.add_paragraph()
r = app1.add_run('附录一：决策卡标准模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

t1 = doc.add_table(rows=8, cols=2)
t1.style = 'Table Grid'

c = t1.cell(0, 0)
c.merge(t1.cell(0, 1))
set_cell_shading(c, 'D9E2F3')
r = c.paragraphs[0].add_run('【决策卡名称】')
r.bold = True

tmpl1 = [
    ('触发条件', '（条件句+动作指令，写在卡片最显眼位置）\n如果出现[具体可观测信号]，先做[具体核实动作]，再决定是否继续。'),
    ('检查表', '（5-8条具体可执行的核对项，每条用能直接对照的动词开头）\n1. [具体核对动作1]\n2. [具体核对动作2]\n3. [具体核对动作3]\n...'),
    ('应急方案', '（当触发条件满足且出现变体场景时的处理方式）\n• 如果[场景变体A]，则[处理方式]\n• 如果[场景变体B]，则[处理方式]'),
    ('适用场景', '（说明本卡适用的具体情境，以及不适用的情境）'),
    ('警示案例', '（本卡对应的失败岔路口描述，用于嵌入式警示对照）'),
    ('认领人', '姓名：      联系方式：\n最近更新日期：'),
]

for i, (label, content) in enumerate(tmpl1, 1):
    cl = t1.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True
    t1.cell(i, 1).paragraphs[0].add_run(content)

c = t1.cell(7, 0)
c.merge(t1.cell(7, 1))
set_cell_shading(c, 'FFF2CC')
r = c.paragraphs[0].add_run('定位说明：本卡列出的触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。')
r.italic = True

doc.add_page_break()
print('Template 1 done')

# 附录二：稽核表模板
app2 = doc.add_paragraph()
r = app2.add_run('附录二：稽核表模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

doc.add_paragraph()
it = doc.add_paragraph()
r = it.add_run('基本信息')
r.bold = True

t2 = doc.add_table(rows=5, cols=2)
t2.style = 'Table Grid'

for i, label in enumerate(['决策名称', '决策日期', '决策者', '稽核日期', '稽核人']):
    cl = t2.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True

doc.add_paragraph()
at = doc.add_paragraph()
r = at.add_run('稽核问题')
r.bold = True

aqs = [
    ('问题1：沉淀检查', '这次决策过程中，有没有出现过让决策者感到当时不确定的关键节点？\n□ 有 → 该节点是否已被记录？\n□ 没有\n备注：'),
    ('问题2：使用检查', '如果对应的场景已有决策卡覆盖：\n□ 已被使用，反馈：_______\n□ 未被使用，原因：_______\n□ 无对应卡片'),
    ('问题3：价值识别', '这次决策是否具备做成卡片的价值？\n□ 是 → 是否已启动访谈？\n□ 否 → 原因：_______\n备注：'),
    ('问题4：反馈迭代', '是否发现现有卡片存在漏洞或过时之处？\n□ 是 → 是否已反馈给认领人？\n□ 否\n备注：'),
]

for q_title, q_content in aqs:
    tq = doc.add_table(rows=2, cols=1)
    tq.style = 'Table Grid'
    ct = tq.cell(0, 0)
    set_cell_shading(ct, 'D9E2F3')
    r = ct.paragraphs[0].add_run(q_title)
    r.bold = True
    tq.cell(1, 0).paragraphs[0].add_run(q_content)
    doc.add_paragraph()

et = doc.add_paragraph()
r = et.add_run('额外问题')
r.bold = True
doc.add_paragraph('这次有没有基于自己的判断，偏离了卡片上的建议？为什么？\n（用于检测是否过度依赖卡片、失去自主判断力）')

doc.add_paragraph()
coc = doc.add_paragraph()
r = coc.add_run('稽核结论')
r.bold = True
doc.add_paragraph('□ 经验已有效沉淀\n□ 需要跟进：_______\n□ 需要迭代卡片：_______')

doc.add_page_break()
print('Template 2 done')

# 附录三：访谈提纲模板
app3 = doc.add_paragraph()
r = app3.add_run('附录三：访谈提纲模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

doc.add_paragraph()
intt = doc.add_paragraph()
r = intt.add_run('访谈基本信息')
r.bold = True

t3 = doc.add_table(rows=5, cols=2)
t3.style = 'Table Grid'
for i, label in enumerate(['决策名称', '决策者', '访谈日期', '访谈者', '访谈时长']):
    cl = t3.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True

doc.add_paragraph()
introt = doc.add_paragraph()
r = introt.add_run('访谈导入（约5分钟）')
r.bold = True
doc.add_paragraph('建立信任，说明访谈目的，不涉及具体判断对错的内容\n• 先聊一些相对轻松的内容\n• 等对方稍微放松下来，再切入正题')

dims = [
    ('第一维度：追因（约20-30分钟）', '核心问题：当时是什么信号让你觉得需要做决策，而不是按原计划走？\n\n追问要点：\n• 在那之前呢？（往前再退一步）\n• 不要满足于第一个触发点答案'),
    ('第二维度：权衡（约30-40分钟）', '核心问题：当时你能想到的选项有哪些，你是怎么排除掉其他选项的？\n\n追问要点：\n• 需要追问两到三轮\n• 具体化追问\n• 警惕事后合理化'),
    ('第三维度：未预见的假设（约15-20分钟）', '核心问题：如果当时那个前提不成立，你还会做同样的判断吗？\n\n追问要点：\n• 挖掘决策者无意识依赖的假设'),
]

for dimt, dimc in dims:
    doc.add_paragraph()
    dh = doc.add_paragraph()
    r = dh.add_run(dimt)
    r.bold = True
    doc.add_paragraph(dimc)

doc.add_paragraph()
hidt = doc.add_paragraph()
r = hidt.add_run('隐性判断挖掘（贯穿全程）')
r.bold = True
doc.add_paragraph('触发信号：决策者反复用我就是觉得的地方\n\n技法1-模式追问：把模糊感受放一起对比\n技法2-反事实追问：逼想象未发生情境\n技法3-角色互换追问：切换到教别人视角')

doc.add_paragraph()
endh = doc.add_paragraph()
r = endh.add_run('访谈收尾')
r.bold = True
doc.add_paragraph('• 感谢决策者的时间\n• 说明后续流程\n• 确认草稿确认方式')

doc.add_page_break()
print('Template 3 done')

# ========== Post-class Resources ==========
res = doc.add_paragraph()
r = res.add_run('课后资源')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

doc.add_paragraph()
rt = doc.add_paragraph()
r = rt.add_run('延伸阅读清单')
r.bold = True; r.font.size = Pt(16)

for book in ['《复盘+：把经验转化为能力》—— 邱昭良', '《U型理论》—— Otto Scharmer', '《行动学习实战录》', '《组织能力的杨三角》—— 杨国安', '《斯坦福大学最受欢迎的创意课》']:
    doc.add_paragraph(book, style='List Bullet')

doc.add_paragraph()
hwt = doc.add_paragraph()
r = hwt.add_run('实践作业')
r.bold = True; r.font.size = Pt(16)

for i, h in enumerate([
    '选择一个让你真正纠结过的决策，用三问法评估它是否值得做成决策卡',
    '找到那个决策的相关人员，进行一次完整的复盘访谈',
    '基于访谈结果，制作一张完整的决策卡',
    '为你制作的决策卡安排一个反对者评审，收集反馈并修改',
    '选择一个失败案例，写出它的警示清单三段式内容',
], 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. ')
    r.bold = True
    p.add_run(h)

doc.add_page_break()

# 90天行动计划
plan = doc.add_paragraph()
r = plan.add_run('90天行动计划模板')
r.bold = True; r.font.size = Pt(16)

t4 = doc.add_table(rows=10, cols=3)
t4.style = 'Table Grid'

for i, header in enumerate(['阶段', '重点行动', '里程碑成果']):
    c = t4.cell(0, i)
    set_cell_shading(c, '4472C4')
    r = c.paragraphs[0].add_run(header)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

phases = [
    ('第1-2周', '识别与访谈', '选定值得做卡的决策\n完成第一次访谈'),
    ('第3-4周', '挖掘与结构化', '完成全部三个维度的访谈\n提炼出隐性判断规则'),
    ('第5-6周', '卡片制作', '完成决策卡初稿\n内部评审修改'),
    ('第7-8周', '评审与定稿', '完成反对者评审\n卡片定稿并指定认领人'),
    ('第9-10周', '训练活动设计', '设计三环节微课脚本\n进行第一次训练'),
    ('第11-12周', '应用与反馈', '在真实场景中应用\n收集使用反馈'),
]

for i, (phase, action, result) in enumerate(phases, 1):
    t4.cell(i, 0).paragraphs[0].add_run(phase)
    t4.cell(i, 1).paragraphs[0].add_run(action)
    t4.cell(i, 2).paragraphs[0].add_run(result)

for i in range(7, 10):
    t4.cell(i, 0).paragraphs[0].add_run(f'第{i*2+1}-{i*2+2}周')

doc.add_page_break()
print('Post-class resources done')

# ========== Appendices ==========
appt = doc.add_paragraph()
r = appt.add_run('附录')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

doc.add_paragraph()
termt = doc.add_paragraph()
r = termt.add_run('附录A：术语表')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

glossary = [
    ('复盘', '对过去事件进行叙事性描述，总结经验教训'),
    ('决策卡', '将判断结构提炼成可现场调用的检查表和触发条件'),
    ('隐性判断', '决策者本人未曾语言化的经验规则'),
    ('判断难度', '决策的困难程度，取决于是否存在多个合理选项'),
    ('触发条件', '告诉使用者什么时候该打开决策卡的判断性描述'),
    ('检查表', '列出最容易被忽略但忽略会出大问题的信号点'),
    ('开关', '告诉使用者什么时候该停下来重新想的一句话'),
    ('场景映射矩阵', '把同一条底层判断逻辑，拆解成几个具体场景各自版本'),
    ('追因', '挖掘决策触发点的访谈维度'),
    ('权衡', '挖掘选项排除过程的访谈维度'),
    ('未预见的假设', '挖掘无意识依赖前提的访谈维度'),
    ('决策稽核', '检查判断逻辑有没有被沉淀成组织可复用东西的机制'),
    ('组织判断力', '团队学会识别、访谈、结构化、复用判断的能力'),
]

for term, defn in glossary:
    p = doc.add_paragraph()
    r = p.add_run(f'{term}：')
    r.bold = True
    p.add_run(defn)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# 金句合集
goldt = doc.add_paragraph()
r = goldt.add_run('附录B：金句合集')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

gold_quotes = [
    '教训只能让人不再犯错，决策卡能让人复制成功。',
    '值得复盘的从来不是结果好坏，是判断有没有难度。',
    '决策者自己讲不出他做对了什么，因为他当时没觉得那是个选择。',
    '每一次我就是觉得不对背后，都有一条没被说出来的经验规则。',
    '流程图告诉你按顺序做什么，决策卡告诉你什么时候该停下来想。',
    '决策卡最贵的三个字不是怎么做，是什么时候。',
    '卡片越通用，越没人用；卡片越具体，用的人越多。',
    '讲失败案例的目的不是让人害怕，是让人在自己身上找到那个岔路口。',
    '卡片评审不是找错别字，是找这句话在什么情况下会害死人。',
    '讲道理讲三小时，不如让他在模拟场景里做错一次决策。',
    '一份没有主人的工具，用一次就是最后一次。',
    '稽核表查的不是做对了没有，是这次的判断有没有被沉淀下来。',
    '一张三年没改过的决策卡，大概率已经没人真的在用。',
    '你替他说出来的道理，他转身就忘；他自己说出来的判断，才是他的。',
    '决策卡会被淘汰，做卡的能力不会。',
]

for i, quote in enumerate(gold_quotes, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'{i}. ')
    r.bold = True
    r = p.add_run(f'"{quote}"')
    r.italic = True
    p.paragraph_format.space_after = Pt(8)

print('Appendices done')

doc.save(r'D:\CC\temp\merged_part4.docx')
print('Part 4 saved')

# -*- coding: utf-8 -*-
"""生成 主持人主持词.docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"D:\2026年课程\顺造科技\AI\评审\04-评委与主持\主持人主持词.docx"

doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

styles = doc.styles
normal = styles['Normal']
normal.font.name = '宋体'
normal.font.size = Pt(11)
rpr = normal.element.rPr
if rpr is None:
    rpr = OxmlElement('w:rPr')
    normal.element.append(rpr)
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:eastAsia'), '宋体')
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_cn_font(run, name='宋体', size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 18, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3A68')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 14, bold=True, color=RGBColor(0x2E, 0x5C, 0x8A))
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run, '黑体', 12, bold=True, color=RGBColor(0x44, 0x44, 0x44))
    return p

def add_para(text, size=11, indent_first=False, justify=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(4)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(22)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_cn_font(run, '宋体', size)
    return p

def add_callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAF1F8')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        b = OxmlElement('w:'+side)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), '1F3A68')
        pBdr.append(b)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_cn_font(run, '楷体', 11, color=RGBColor(0x1F, 0x3A, 0x68))
    return p

def add_script_block(title, lines):
    """一个完整的主持词块：标题 + 多行主持词（楷体）+ 时长/动作提示"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run('【主持词】' + title)
    set_cn_font(run, '黑体', 11, bold=True, color=RGBColor(0x1F, 0x3A, 0x68))
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.7
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        set_cn_font(run, '楷体', 12)

def add_note(text):
    """主持词下方的【动作/时长】提示"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(8)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '999999')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_cn_font(run, '宋体', 9, color=RGBColor(0x99, 0x99, 0x99))
    return p

# ===== 封面 =====
for line, sz, color in [
    ('顺造科技 · AI 项目成果评审', 22, RGBColor(0x1F, 0x3A, 0x68)),
    ('主 持 人 主 持 词', 36, RGBColor(0x1F, 0x3A, 0x68)),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80 if line.startswith('主 持') else 200)
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(line)
    set_cn_font(run, '黑体', sz, bold=True, color=color)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(40)
run = p.add_run('— 主持人专用 · 现场剧本 —')
set_cn_font(run, '楷体', 14, color=RGBColor(0x66, 0x66, 0x66))

for line in ['版本：V1.0', '编制：AI 项目评审组', '适用场景：D-Day 下午 + D+1 全天']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(line)
    set_cn_font(run, '宋体', 12)

doc.add_page_break()

# ===== 写在前面 =====
add_heading_1('写在前面：主持人的三重身份')
add_para('作为本次评审的主持人，您不是简单的串场报幕员，而是评审现场的"节拍器 + 氛围官 + 危机处理员"。具体来说：', size=11)
add_para('节拍器：每位学员 16-18 分钟，每位评委点评 1.5-2 分钟，每个茶歇 10 分钟。主持人必须像节拍器一样精准，不能让任何环节拖过它的边界。', size=11)
add_para('氛围官：开场要热，路演要稳，点评要诚，颁奖要燃。每一个情绪节点都需要主持人用语言和动作去托起。', size=11)
add_para('危机处理员：设备故障、学员超时、评委分歧、突发状况——这些都不会提前通知您，必须有预案、有话术、有定力。', size=11)
add_callout('本主持词是"现场剧本"，不是逐字稿。主持人请根据现场气氛灵活调整——最重要的不是"念对"，而是"在场"。')

# ===========================================================
# 第一部分 D-Day 下午
# ===========================================================
doc.add_page_break()
add_heading_1('第一部分  D-Day 下午（14:00-18:00）')

# ---- 1.1 启动前预热 ----
add_heading_2('1.1  启动前预热（13:55-14:00）')
add_script_block('启动前 5 分钟预热（5 分钟）', [
    '各位评委、各位学员，欢迎大家来到顺造科技 AI 项目成果评审的现场。',
    '我是今天的主持人 XXX。',
    '在正式评审开始前，我们用 5 分钟时间做几件事：',
    '第一，请各位把手机调到静音模式，谢谢配合。',
    '第二，请评委到评委席就座，评分夹和笔已经放在您面前的位置上。',
    '第三，请各位学员在候场区就座，我会按顺序叫到您。',
    '第四，茶歇区在场地右侧，洗手间在出门左转。',
    '好，还有 1 分钟，我们马上开始。',
])
add_note('【动作/时长】5 分钟。配合 PPT 倒计时或 LED 屏。')

# ---- 1.2 启动会开场 ----
add_heading_2('1.2  评审启动会开场（14:00-14:15）')

add_script_block('开场白（2 分钟）', [
    '各位评委、各位同事：',
    '下午好。欢迎来到顺造科技第一期 AI 项目成果评审。',
    '过去 8 周，我们做了一件在顺造历史上从未做过的事——',
    '我们让质量、人力、财务、客服、生产、研发等部门的 12 位业务一线员工，',
    '亲手把自己部门的业务，',
    '用 AI 重新做了一遍。',
    '今天，他们将带着自己的方案、自己的代码、自己的小工具，',
    '站在你们面前，讲一讲他们怎么想、怎么做、做出了什么。',
    '这 12 个人，不是程序员，不是产品经理，是我们的同事——',
    '他们用 8 周的业余时间，从零开始，学会了用 AI 改造自己手里的活儿。',
    '我相信，你们今天听到的每一个方案，都是真实的业务、真实的痛点、真实的尝试。',
    '可能还有些粗糙，可能还不够完美——但这是顺造 AI 转型的第一步。',
    '好，下面有请我们的评委组组长，X 总，为我们做开场致辞。',
])
add_note('【动作/时长】2 分钟。语调平和、有温度。强调"真实"和"第一步"。')

add_script_block('评委介绍（3 分钟）', [
    '感谢 X 组的致辞。',
    '下面我为大家介绍今天的 7 位评委——',
    '他们分别从战略、研发、业务、培训、AI 专业五个角度，全方位把关每一个项目：',
])
# 评委介绍表
table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '评委'
hdr[1].text = '一句话定位（按此介绍）'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            set_cn_font(r, '黑体', 11, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F3A68')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

judges = [
    ('评委 1 · X 总（集团副总裁）', '战略与组织视角，关注方案能否在公司层面复制推广。'),
    ('评委 2 · X 总监（业务部门）', '业务影响力视角，关注与年度战略的契合度。'),
    ('评委 3 · X 主任（总裁办）', '对内激励、对外展示视角。'),
    ('评委 4 · X 经理（培训经理）', '培训目标达成视角，关注学员成长性和内容设计。'),
    ('评委 5 · X 老师（内训师）', '实操可落地性视角，关注是否便于复制到其他部门。'),
    ('评委 6 · X 总监（业务部门负责人）', '业务痛点还原度视角，关注上线后能否真解决问题。'),
    ('评委 7 · X 博士（AI 专家）', 'AI 介入级别视角，关注技术合理性和未来扩展性。'),
]
for i, (a, b) in enumerate(judges, start=1):
    row = table.rows[i].cells
    row[0].text = a
    row[1].text = b
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                set_cn_font(r, '宋体', 10)

add_para('  ', size=4)

add_script_block('规则说明（3 分钟）', [
    '感谢各位评委。下面我简单说一下今天的评审规则。',
    '今天下午有 4 位学员进行路演，每位学员的时间是 16-18 分钟，分配如下：',
    '—— 路演环节：10 分钟。学员介绍自己的方案，9 分 30 秒我会举黄牌，10 分钟我会举红牌，10 分 30 秒我会请您收尾。',
    '—— Q&A 环节：3-5 分钟。在场评委和学员可以向演讲者提问 1-2 个问题。',
    '—— 点评环节：3 分钟。每位评委 1.5-2 分钟，先肯定再建议。',
    '—— 亮分环节：1 分钟。主持人喊"请评委亮分"后，1 分钟提示 + 30 秒内必须亮分。',
    '7 位评委去掉一个最高分、去掉一个最低分，剩余 5 位取算术平均 = 该学员的最终得分。',
    '奖项分 6 个：最佳 AI 方案、最佳智能体、最佳 HTML、最佳路演、最佳团队、最佳新人。',
    '所有奖项将在明天下午结营仪式上公布。',
    '好，下面请 1 号学员准备上场。',
])
add_note('【动作/时长】3 分钟。配合幻灯片展示评分规则与奖项。')

# ---- 1.3 学员路演主持词 ----
add_heading_2('1.3  学员路演主持词（每位学员，循环使用）')

# 模板
add_heading_3('模板：路演开始前（30 秒）')
add_script_block('1 号学员上场前（示例）', [
    '好，规则大家都清楚了。',
    '下面有请今天的第一位学员——来自质量部门的 XXX，他/她带来的课题是《AI 辅助 8D 报告审核》。',
    '掌声有请。',
])
add_note('【动作】请评委准备评分夹。')
add_para('【通用模板】下一位上场的是 XX 部门 XXX，他/她带来的课题是《XX》。掌声有请。', size=11)

add_heading_3('模板：路演结束后 + Q&A 引导（30 秒）')
add_script_block('1 号学员路演结束后（示例）', [
    '感谢 XXX 的精彩分享。',
    '我们用掌声再次感谢 XXX。',
    '看得出，XXX 同学为了这个项目花了不少心思，把 8D 报告的"要素拆解 → 知识图谱 → 智能体"链路梳理得非常清楚。',
    '好，下面进入 Q&A 环节——',
    '关于这个课题，在场各位有 1-2 个问题吗？',
])
add_note('【动作】把话筒递给第一个举手的人。')
add_para('【兜底问题】如果没人提问，主持人按顺序问以下 3-5 个问题中的 1-2 个：', size=11)
add_para('  1) 这个方案上线后，预期能省多少时间？有没有真实数据？', size=11)
add_para('  2) 跨部门同事想用你这个方案，需要做什么？', size=11)
add_para('  3) 你在做的过程中遇到的最大困难是什么？怎么解决的？', size=11)
add_para('  4) 如果只能改一个地方让效果翻倍，你会改哪里？', size=11)
add_para('  5) 这个项目做完，你觉得自己最大的收获是什么？', size=11)

add_heading_3('模板：Q&A 结束 + 点评引导（30 秒）')
add_script_block('Q&A 结束后（示例）', [
    '好，感谢各位的提问与回应。',
    '下面进入点评环节——',
    '我们请 X 位评委分别用 1.5-2 分钟，先肯定再建议，给 XXX 一些具体的改进方向。',
    '有请第一位评委。',
])
add_note('【动作】按评委序号 1-7 依次邀请，最后一位评委点评完之前 30 秒举牌提醒。')

add_heading_3('模板：点评结束 + 亮分（1 分钟）')
add_script_block('点评结束后（示例）', [
    '好，感谢各位评委的中肯点评。',
    '我注意到 X 位评委都提到了"跨部门推广"这个方向，说明这是我们顺造 AI 接下来要重点突破的点。',
    'XXX，请回到座位上休息。下面请评委准备亮分。',
    '请各位评委在 1 分钟内，根据三轨评分（AI 方案价值 40% + AI 产出物质量 35% + 路演表现 25%）完成打分。',
    '（30 秒后）好，请亮分。',
    '1 号评委 X 分，2 号评委 X 分，3 号评委 X 分，4 号评委 X 分，5 号评委 X 分，6 号评委 X 分，7 号评委 X 分。',
    '去掉一个最高分 X 分，去掉一个最低分 X 分，最终得分：X.X 分。',
])
add_note('【动作】举评分牌 / 唱分 / 写在大白板上。')

# ---- 1.4 茶歇 ----
add_heading_2('1.4  茶歇主持词（每个 10 分钟）')
add_script_block('D-Day 第一次茶歇（15:30）', [
    '好，刚才 3 位学员的精彩分享，让我们对顺造 AI 的"工具化"和"团队化"路径有了非常具体的感知。',
    '我们用 10 分钟时间休息一下，茶歇区在场地右侧。',
    '请各位学员利用这 10 分钟稍作调整，下一位学员请做好准备。',
    '我们 15:40 准时开始下一位。',
])
add_note('【动作】音乐响起。评委可短暂交流整体观察（不可讨论具体分数）。')

add_script_block('D-Day 第二次茶歇（17:00）', [
    '好，下午的评审已经过半。',
    '我们用 10 分钟时间休息一下。',
    '提醒评委老师：还有 X 位学员的路演。',
    '我们 17:10 准时开始。',
])

# ---- 1.5 D-Day 总结 ----
add_heading_2('1.5  D-Day 总结主持词（18:00-18:30）')
add_script_block('D-Day 收工（18:00）', [
    '各位评委、各位同事：',
    '今天的评审告一段落。',
    '感谢 7 位评委的专业点评——',
    '我特别注意到，今天的点评中出现频率最高的两个词是"业务痛点"和"跨部门复用"。',
    '这两个词会写进我们明天的评委组总结中。',
    '明天上午我们将继续评审剩余 X 位学员，下午是结营仪式和颁奖。',
    '感谢各位学员——每一位站上来的同事，都是顺造 AI 转型路上的先行者。',
    '我们明天 8:50 评委碰头会，9:00 正式开始评审。',
    '今晚好好休息，我们明天见。',
])
add_note('【动作】评委 18:00 进入碰头评分（30 分钟），签字确认今日所有打分。')

# ===========================================================
# 第二部分 D+1
# ===========================================================
doc.add_page_break()
add_heading_1('第二部分  D+1 上午 + 下午（09:00-15:30）')

# ---- 2.1 上午开场 ----
add_heading_2('2.1  D+1 上午开场（09:00-09:05）')
add_script_block('D+1 上午开场（5 分钟）', [
    '各位评委、各位同事：',
    '早上好。',
    '欢迎来到顺造科技 AI 项目成果评审的第二天。',
    '今天上午我们将继续评审剩余 X 位学员，下午 13:30 正式开始结营仪式。',
    '在结营仪式上，我们将公布 6 个奖项的归属，颁发结业证书，听学员代表发言，'
    '听评委组组长做总结，听公司领导做收官发言。',
    '提醒评委老师：上午的评审标准和昨天一致——每学员 16-18 分钟，'
    '听完点评后请独立打分，不要受昨日分数影响。',
    '好，下面有请 X 号学员准备上场。',
])
add_note('【动作】参照昨日的"路演前 → 路演后 → Q&A → 点评 → 亮分"流程。')

# ---- 2.2 上午茶歇 ----
add_heading_2('2.2  上午茶歇主持词（10:30）')
add_script_block('D+1 上午茶歇（10:30）', [
    '好，上午的评审过半。',
    '我们用 10 分钟时间休息一下。',
    '提醒各位学员：下午 13:30 是结营仪式，请务必准时到场。',
    '我们 10:40 准时开始。',
])

# ---- 2.3 中午过渡 ----
add_heading_2('2.3  中午过渡主持词（12:00）')
add_script_block('D+1 中午过渡（12:00）', [
    '各位评委、各位同事：',
    '上午的评审告一段落。',
    '感谢今天上午 X 位学员的精彩分享。',
    '下午 13:30 准时开始结营仪式，请大家准时到场。',
    '午餐在二楼餐厅，评委老师和学员可以自由交流，氛围轻松。',
    '13:30 我们不见不散。',
])

# ---- 2.4 结营仪式 ----
add_heading_2('2.4  D+1 结营仪式（13:30-15:30）')

# 项目总结
add_script_block('环节 1：项目总结（13:30-13:40）开场', [
    '各位领导、各位评委、各位同事：',
    '下午好。',
    '欢迎来到顺造科技第一期 AI 项目成果评审的结营仪式。',
    '过去的两天里，我们见证了 X 位学员的精彩分享，'
    '看到了质量、人力、财务、客服、生产、研发等部门的业务一线，'
    '被 AI 重新想象的真实过程。',
    '下面，我们用 10 分钟时间，请本次 AI 项目的负责人 XXX，'
    '为我们做一个项目总结。',
    '有请。',
])
add_note('【动作】投影切到项目总结 PPT。')

# 分数公布
add_script_block('环节 2：评审分数公布（13:40-13:45）开场', [
    '感谢 XXX 的项目总结。',
    '下面，进入大家最期待的环节——',
    '评审分数公布。',
    '我们今天有 6 个奖项：',
    '最佳 AI 方案、最佳智能体、最佳 HTML、最佳路演、最佳团队、最佳新人。',
    '所有分数均由 7 位评委独立打分，去掉一个最高分、去掉一个最低分后取平均，'
    '并经评委组组长签字确认。',
    '下面，我们一个奖项一个奖项来公布。',
    '（停顿 3 秒）',
    '首先，是最佳 AI 方案奖——',
])
add_note('【动作】主持人 + 评委组组长共同宣读。每一项宣读前停顿 3 秒，制造悬念。')

# 颁奖
add_heading_3('环节 3：颁奖（13:45-14:05）')
add_para('颁奖是整个结营仪式的"高光 20 分钟"。请按以下 6 个奖项的顺序逐项颁奖。', size=11)
add_para('', size=4)

# 奖项 1
add_script_block('奖项 1 · 最佳 AI 方案奖（3 分钟）', [
    '首先，是最佳 AI 方案奖。',
    '这个奖项颁发给方案在"业务痛点还原度、AI 介入合理性、可推广性"三个维度上综合表现最突出的学员。',
    '（停顿 2 秒）',
    '获奖者是——来自 XX 部门的 XXX，他/她的课题是《XX》。',
    '请 XXX 上台领奖。',
    '（上台后）',
    '请 X 总为 XXX 颁奖。',
    '（颁奖动作进行中）',
    '恭喜 XXX！我们用掌声向他/她表示祝贺！',
    '请 X 总与获奖者合影。',
    '（合影后）',
    '请 XXX 下台就座，我们继续下一项。',
])
add_note('【动作】颁奖嘉宾 + 获奖者合影。可用颁奖音乐。')

# 奖项 2-6 通用模板
add_heading_3('奖项 2-6 通用模板（替换奖项名称和颁奖嘉宾）')
awards_template = [
    ('奖项 2 · 最佳智能体奖', '这个奖项颁发给在"AI 介入深度 + 工具集成能力"上表现最突出的学员。'),
    ('奖项 3 · 最佳 HTML 奖', '这个奖项颁发给做出了"可直接打开、可交互、UI 友好"HTML 工具的学员。'),
    ('奖项 4 · 最佳路演奖', '这个奖项颁发给"演讲结构清晰、Q&A 应答流畅、AI 思维体现"三方面最出色的学员。'),
    ('奖项 5 · 最佳团队奖', '这个奖项颁发给"全员参与、互相支持、整体产出"最突出的团队。'),
    ('奖项 6 · 最佳新人奖', '这个奖项颁发给"AI 启动基线最低、但成长最显著"的新人学员。'),
]
for name, desc in awards_template:
    add_script_block(name + '（3 分钟）', [
        '接下来，是' + name + '。',
        desc,
        '（停顿 2 秒）',
        '获奖者是——',
        '（念出获奖者姓名 + 部门 + 课题）',
        '请 XXX 上台领奖。',
        '（上台后）',
        '请 X 总/总监为 XXX 颁奖。',
        '（颁奖动作进行中）',
        '恭喜 XXX！',
        '请合影。',
        '（合影后）',
        '请 XXX 下台就座。',
    ])

# 学员发言
add_script_block('环节 4：学员代表发言（14:05-14:15）开场', [
    '感谢以上 X 位获奖者。',
    '下面，我们用 10 分钟时间，请两位学员代表发言。',
    '第一位是优秀学员代表——最佳 AI 方案奖获得者 XXX；',
    '第二位是新人学员代表——最佳新人奖获得者 XXX。',
    '有请 XXX 先发言。',
])
add_note('【动作】话筒准备。投影切到发言 PPT（如果有）。')

# 结业证书
add_script_block('环节 5：结业证书颁发（14:15-14:25）开场', [
    '感谢两位学员代表的精彩发言。',
    '接下来，是结业证书颁发环节。',
    '本次参与评审的 X 位学员，将获得顺造科技第一期 AI 项目结业证书。',
    '请所有学员上台列队。',
    '（学员列队中）',
    '我们请 X 总、X 总监、X 主任为学员们颁发结业证书。',
    '（颁发中）',
    '请合影留念。',
])
add_note('【动作】证书提前摆放在桌面上。学员按顺序领取。')

# 评委总结
add_script_block('环节 6：评委代表总结（14:25-14:35）开场', [
    '感谢所有学员。',
    '接下来是评委组总结环节。',
    '我们请评委组组长 X 总，代表 7 位评委做 10 分钟的总结发言。',
    '有请 X 总。',
])
add_note('【动作】X 总做评委组总结——肯定亮点 + 指出共性短板 + 公司 AI 战略建议。')

# 成果交付
add_script_block('环节 7：成果交付（14:35-14:45）开场', [
    '感谢 X 总的精彩总结。',
    '下面进入成果交付环节。',
    '本次项目的所有产出将打包交付给公司，包括：',
    '—— 优秀方案汇编（10 份 PDF）；',
    '—— HTML 模板库（可直接复用）；',
    '—— 智能体清单（在 Dify / 飞书 Aily 上的所有智能体）。',
    '我们请项目负责人 XXX 做 10 分钟的成果交付说明。',
    '有请。',
])
add_note('【动作】交付物摆放在主席台前方。')

# 领导致辞
add_script_block('环节 8：领导致辞（14:45-14:55）开场', [
    '感谢 XXX 的成果交付说明。',
    '下面，是本次结营仪式的最后一个环节——',
    '公司领导致辞。',
    '我们请公司 X 总做 10 分钟的收官发言。',
    '有请 X 总。',
])

# 散场
add_script_block('环节 9：合影 + 散场（14:55-15:30）开场', [
    '感谢 X 总的收官发言。',
    '—— 顺造科技第一期 AI 项目成果评审——',
    '到此，圆满结束。',
    '感谢 X 位学员的精彩呈现。',
    '感谢 X 位评委的专业点评。',
    '感谢在场的每一位同事。',
    '请评委、学员、领导到台前合影。',
    '（合影中）',
    '好，今天的结营仪式到此结束。',
    '请大家到一楼茶歇区，我们准备了小点心和饮料。',
    '我们合影后见。',
])
add_note('【动作】合影 → 散场 → 自由交流。')

# ===========================================================
# 第三部分 应急话术
# ===========================================================
doc.add_page_break()
add_heading_1('第三部分  应急话术（现场突发状况）')

add_heading_2('3.1  设备故障应急话术')

add_script_block('情况 1：投影 / 电脑故障（5 分钟内）', [
    '各位，请稍等——',
    '我们的设备出现了一个小问题，技术老师正在处理。',
    '请 XXX 学员先用手机投屏做备用方案。',
    '评委老师请利用这个时间翻一下学员手册。',
    '我们 3 分钟内恢复。',
])

add_script_block('情况 2：投影 / 电脑故障（5-15 分钟）', [
    '各位，因为设备问题，我们启动 B 计划。',
    '请 XXX 学员用手机投屏或口头介绍的方式继续。',
    '评委老师，我们先用学员手册 + 提问的方式完成评审。',
    '请大家理解。',
])

add_script_block('情况 3：设备长时间无法恢复（>15 分钟）', [
    '各位，因为设备问题长时间无法恢复，我们做以下调整：',
    '本轮学员的评审改为"手册讲解 + Q&A"模式，路演分由评委酌情处理（设备原因不扣学员分）。',
    '请评委老师在评分表"备注"栏注明"设备原因"。',
    '我们继续。',
])

add_heading_2('3.2  学员突发状况应急话术')

add_script_block('情况 4：学员明显紧张 / 卡壳', [
    'XXX 同学，请先喝口水，我们不着急。',
    '（停顿 5 秒）',
    '请继续——或者，你可以从你最熟悉的那个部分开始。',
    '（如果学员实在无法继续）',
    '我们看到 XXX 同学准备得非常用心，下面请评委基于现有材料做点评。',
])

add_script_block('情况 5：学员超时严重', [
    'XXX 同学，提示一下，我们还有 30 秒。',
    '（10 分钟后）',
    'XXX 同学，时间到了，请用 30 秒收尾。',
    '（超时 30 秒）',
    '感谢 XXX 的分享。请评委准备点评。',
])

add_script_block('情况 6：学员方案被质疑，氛围紧张', [
    '这个点我们记下来了，会后单独沟通。',
    'XXX 同学请先回座休息一下。',
    '（如学员情绪激动）',
    '各位，请允许我们的学员先回座。评委请继续完成打分。',
])

add_heading_2('3.3  评委突发状况应急话术')

add_script_block('情况 7：评委意见分歧巨大（现场）', [
    '各位评委老师，我们尊重每位评委的独立判断。',
    '本轮不公开讨论分数，最终得分按"去最高去最低取平均"计算。',
    '我们继续下一环节。',
])

add_script_block('情况 8：评委临时缺席', [
    '各位，今天的评委老师因为 [原因] 暂时无法到场。',
    '我们请 [候补评委] 顶替。',
    '最终得分按 5 位评委"去最高去最低取平均"计算。',
    '我们继续。',
])

add_heading_2('3.4  现场气氛应急话术')

add_script_block('情况 9：现场气氛过于沉闷', [
    '我注意到大家可能有点累了——',
    '我提议，我们用 30 秒做一个深呼吸。',
    '吸气——',
    '呼气——',
    '好，我们继续。',
])

add_script_block('情况 10：现场气氛过于嘈杂', [
    '各位，请允许我打断一下。',
    '请大家把手机调到静音，谢谢。',
    '我们继续。',
])

# ===========================================================
# 第四部分 附录
# ===========================================================
doc.add_page_break()
add_heading_1('附录 A  主持人开场 / 收场"金句"备选库')

add_heading_2('开场金句（10 选 1）')
quotes_open = [
    '今天我们不在听"AI 能做什么"，我们在听"AI 已经做了什么"。',
    '我们不评 PPT，我们评"真的把业务改了没有"。',
    '在座的每一位学员，都是顺造 AI 转型路上的"第一个吃螃蟹的人"。',
    '我们不是来挑刺的，我们是来帮忙的——帮忙让每一个项目都变得更好。',
    '评委不是来比谁更懂 AI，是来比谁更懂"业务 + AI"的化学反应。',
    'AI 不会取代人，但懂 AI 的人会取代不懂 AI 的人。今天这 12 位，就是"懂的人"。',
    '今天的评审不是"考试"，是"产品发布会"——你们的产品就是你们自己做的 AI 工具。',
    '评审的核心不是"打分"，是"对齐"——我们对齐什么是好的 AI 应用。',
    '如果你们今天被我们挑出 10 个问题，那是赚了——10 个问题就是 10 个改进点。',
    '我期待今天听到的每一个方案，都让我"哇"一声——不管"哇"是大还是小。',
]
for i, q in enumerate(quotes_open, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'  {i}. {q}')
    set_cn_font(run, '楷体', 11)

add_heading_2('收场金句（5 选 1）')
quotes_end = [
    '感谢 X 位学员的精彩呈现。AI 转型不是 8 周的事，是 8 年的事。我们在路上。',
    '今天不是结束，是开始——每个项目都还有 2.0、3.0。我们一起把它做完。',
    '各位，你们今天站上来的样子，就是顺造 AI 的样子。谢谢你们。',
    '在座的每一位，都是顺造的"AI 火种"。希望你们回到各自部门，把火带回去。',
    '评审结束，但思考不停。我期待三个月后再听到你们的迭代版本。',
]
for i, q in enumerate(quotes_end, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'  {i}. {q}')
    set_cn_font(run, '楷体', 11)

# 附录 B
add_heading_1('附录 B  主持人"快速反应"清单（30 条）')
add_para('以下是主持人在现场可能用到的"快速反应"短句，按场景分类：', size=11)
add_para('', size=4)

reactions = [
    '【破冰】"在我们正式开始前，我提议大家用 10 秒钟跟身边的人打个招呼——也许你们今天会成为一个跨部门 AI 项目组的成员。"',
    '【暖场】"我看到已经有评委老师在翻学员手册了——这说明大家都是认真的。"',
    '【过渡】"好的，刚才这位学员给我们带来的信息量很大，请大家用 3 秒钟消化一下。"',
    '【控时】"我们还有 1 分钟，XXX 同学请准备收尾。"',
    '【控时】"时间到——感谢 XXX 的分享。"',
    '【救场】"不着急——我们这里没有标准答案，只有思考过程。"',
    '【救场】"这个点很有价值，但因为时间关系我们记下来，会后单独聊。"',
    '【Q&A 引导】"哪位评委先来一个破冰问题？"',
    '【Q&A 引导】"我注意到这个问题可能在场很多人都想问——我代大家问一下。"',
    '【Q&A 救场】"如果大家都没问题，我有一个——"',
    '【点评过渡】"感谢 XXX 的分享，下面我们有请 X 位评委分别用 1.5-2 分钟做点评。"',
    '【点评过渡】"我特别想听 X 总监的点评——您是做这个业务的，您怎么看？"',
    '【点评提醒】"X 评委老师，还有 30 秒。"',
    '【点评救场】"X 评委老师点评得非常到位，我建议时间到了我们先收住——其他的我们记下来会后聊。"',
    '【亮分前】"请评委准备亮分。我帮大家读一下规则——去掉一个最高分，去掉一个最低分，取平均。"',
    '【亮分后】"X 分！恭喜 XXX！掌声再次送给他/她！"',
    '【茶歇前】"刚才 3 位学员的信息量很大，我提议我们用 10 分钟消化一下。"',
    '【茶歇前】"茶歇区在场地右侧，洗手间出门左转。"',
    '【茶歇后】"好，大家回到座位上，我们继续。"',
    '【D-Day 收工】"今天到这里。明天 8:50 评委碰头，9:00 正式开始。"',
    '【D-Day 收工】"今晚好好休息。明天还要打起精神。"',
    '【D+1 上午】"上午和昨天的标准一致——独立打分，不要锚定。"',
    '【结营开场】"过去的两天，我们见证了 X 位学员的精彩瞬间——这 12 个人，就是顺造 AI 的火种。"',
    '【颁奖悬念】"下面，我宣布——"',
    '【颁奖悬念】"（停顿 2 秒）……"',
    '【颁奖中】"请 X 总为 XXX 颁奖。"',
    '【颁奖后】"请合影留念。"',
    '【散场】"今天的结营仪式到此结束。一楼茶歇区有茶点，我们合影后见。"',
    '【设备故障】"设备出了一点小问题，请大家稍等——技术老师 3 分钟内解决。"',
    '【人员缺席】"X 评委因为 [原因] 暂时无法到场，我们请 [候补] 顶替。"',
]
for r in reactions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(r)
    set_cn_font(run, '宋体', 10.5)

# 封底
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(200)
run = p.add_run('— 主持人主持词结束 —')
set_cn_font(run, '楷体', 14, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('祝评审圆满！')
set_cn_font(run, '楷体', 12, color=RGBColor(0x66, 0x66, 0x66))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('顺造科技 · AI 项目评审组')
set_cn_font(run, '楷体', 12, color=RGBColor(0x66, 0x66, 0x66))

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print('OK:', OUT_PATH)
print('段落数:', len(doc.paragraphs))

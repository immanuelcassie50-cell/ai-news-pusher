"""参数提取核心模块：多格式文件解析、文本提取、同义词匹配参数提取"""
import pdfplumber
from docx import Document
import pandas as pd
from pathlib import Path
import re
import json
from rapidfuzz import fuzz


def extract_text_from_pdf(file_path):
    """从 PDF 提取文本，跳过图片型页面"""
    texts = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                texts.append(f"[第{page_num}页]\n{text}")
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row:
                        clean_row = [str(cell or "").strip() for cell in row]
                        texts.append(" | ".join(clean_row))
    return "\n".join(texts)


def extract_text_from_docx(file_path):
    """从 Word 文件提取文本和表格"""
    doc = Document(file_path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            texts.append(" | ".join(row_texts))
    return "\n".join(texts)


def extract_text_from_excel(file_path):
    """从 Excel 提取所有 Sheet 的文本内容"""
    texts = []
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
        texts.append(f"[Sheet: {sheet_name}]")
        for _, row in df.iterrows():
            row_text = " | ".join([str(v) for v in row if pd.notna(v) and str(v).strip()])
            if row_text:
                texts.append(row_text)
    return "\n".join(texts)


def extract_text(file_path):
    """统一入口：根据文件类型调用对应提取函数"""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix in (".xlsx", ".xls"):
        return extract_text_from_excel(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")


def get_supplier_name(file_path):
    """从文件名中提取供应商标识（取下划线或空格前的第一段）"""
    stem = Path(file_path).stem
    parts = re.split(r"[_\-\s]", stem)
    return parts[0] if parts else stem


def load_config(config_path):
    """加载参数提取配置"""
    with open(config_path, encoding="utf-8") as f:
        return json.load(f)


def find_value_near_keyword(text, keyword, window=150, config=None):
    """
    在文本中找到关键词，提取其前后 window 字符内的数值和单位。
    返回：(提取值, 原文片段, 置信度)
    """
    pattern = re.compile(re.escape(keyword), re.IGNORECASE)
    matches = list(pattern.finditer(text))
    if not matches:
        return None, None, None

    best_match = matches[0]
    start = max(0, best_match.start() - 20)
    end = min(len(text), best_match.end() + window)
    context = text[start:end]

    value_pattern = re.compile(
        r"[-−]?\d+\.?\d*\s*(?:~|～|to|-)\s*[-−]?\d+\.?\d*"
        r"|[-−]?\d+\.?\d*\s*[℃°CVAWkgmm%]"
        r"|\d+\.?\d*"
    , re.IGNORECASE)
    value_matches = value_pattern.findall(context)

    if not value_matches:
        return None, context.strip(), "low"

    extracted_value = value_matches[0].strip()

    confidence = "medium"
    if config:
        confidence_keywords = config.get("confidence_keywords", {})
        for kw in confidence_keywords.get("high", []):
            if kw.lower() in context.lower():
                confidence = "high"
                break
        for kw in confidence_keywords.get("low", []):
            if kw.lower() in context.lower():
                confidence = "low"
                break

    return extracted_value, context.strip()[:100], confidence


def extract_parameters(text, config):
    """对一份文件的文本内容，逐参数提取值"""
    results = {}
    for param_def in config.get("parameters", []):
        param_name = param_def["name"]
        synonyms = param_def.get("synonyms", [param_name])

        best_value = None
        best_context = None
        best_confidence = None

        for synonym in synonyms:
            value, context, confidence = find_value_near_keyword(text, synonym, config=config)
            if value:
                if best_confidence is None or \
                   (confidence == "high" and best_confidence != "high") or \
                   (confidence == "medium" and best_confidence == "low"):
                    best_value = value
                    best_context = context
                    best_confidence = confidence
                break

        results[param_name] = {
            "value": best_value if best_value else "⚠️ 文件未提及",
            "context": best_context or "",
            "confidence": best_confidence or "—"
        }
    return results


def extract_from_file(file_path, config):
    """从单个文件提取所有参数"""
    text = extract_text(file_path)
    return extract_parameters(text, config)
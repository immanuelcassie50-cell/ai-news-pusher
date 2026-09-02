#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert 学员手册.md to docx, and use Edge headless to print HTML to PDF."""
import os
import sys
import subprocess
import time
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

BASE = Path("D:/2026年课程/ai课2026整理/经营型企业大学构建：从培训部门到战略引擎/学员手册")
TASK = BASE.name  # "学员手册"
MD = BASE / f"{TASK}.md"
HTML_FILE = BASE / f"{TASK}.html"
DOCX = BASE / f"{TASK}.docx"
PDF = BASE / f"{TASK}.pdf"

EDGE = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"


def make_docx(md_text: str, out_path: Path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Noto Serif SC"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = [r for r in table_lines if r.strip().startswith("|")]
        if len(rows) < 2:
            table_lines = []
            return
        header = [c.strip() for c in rows[0].strip("|").split("|")]
        data_rows = []
        for r in rows[2:]:
            cells = [c.strip() for c in r.strip("|").split("|")]
            data_rows.append(cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
        table.style = "Light Grid Accent 1"
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(10)
        for i_r, row in enumerate(data_rows):
            for j_c, c in enumerate(row[:len(header)]):
                cell = table.rows[i_r + 1].cells[j_c]
                cell.text = c
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        table_lines = []

    in_table = False
    for line in lines:
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            try:
                p.style = doc.styles["Intense Quote"]
            except KeyError:
                pass
            run = p.add_run(line[2:].strip())
            run.italic = True
        elif line.strip() == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip() == "":
            doc.add_paragraph()
        elif line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            content = line.lstrip()[2:]
            doc.add_paragraph(content, style="List Bullet")
        else:
            stripped = line.strip()
            is_numbered = False
            num_prefix = ""
            for k in range(1, 10):
                pfx = f"{k}. "
                if stripped.startswith(pfx):
                    is_numbered = True
                    num_prefix = pfx
                    break
            if is_numbered:
                content = stripped[len(num_prefix):]
                doc.add_paragraph(content, style="List Number")
            else:
                doc.add_paragraph(stripped)

    if in_table:
        flush_table()
    doc.save(str(out_path))
    print(f"Created DOCX: {out_path}")


def html_to_pdf_with_edge(html_file: Path, pdf_file: Path):
    """Use Edge headless to print HTML to PDF."""
    # Edge headless PDF print
    # Build a file:// URL
    url = "file:///" + str(html_file).replace("\\", "/")
    cmd = [
        EDGE,
        "--headless=new",
        "--disable-gpu",
        "--no-sandbox",
        f"--print-to-pdf={pdf_file}",
        url,
    ]
    print(f"Running: {' '.join(cmd[:4])} ...")
    proc = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        timeout=180,
    )
    if proc.returncode != 0:
        print("STDERR:", proc.stderr[:500])
        raise RuntimeError(f"Edge headless failed: returncode={proc.returncode}")
    print(f"Created PDF: {pdf_file}")


def main():
    md_text = MD.read_text(encoding="utf-8")
    print(f"Reading: {MD} ({len(md_text)} chars)")

    # DOCX
    make_docx(md_text, DOCX)

    # PDF (use HTML for printing)
    if HTML_FILE.exists():
        html_to_pdf_with_edge(HTML_FILE, PDF)
    else:
        print(f"HTML not found: {HTML_FILE}")

    # Report
    print()
    print("=" * 60)
    for f in [MD, HTML_FILE, DOCX, PDF]:
        if f.exists():
            print(f"  {f.name}: {f.stat().st_size:,} bytes")
    print("=" * 60)


if __name__ == "__main__":
    main()

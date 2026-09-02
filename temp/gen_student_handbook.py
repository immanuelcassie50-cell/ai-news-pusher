#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "D:/新课开发/变革管理/08-向上管理与高层说服术：让决策层理解容错的成本逻辑/完整课程包/03-学员手册/学员手册-向上管理与高层说服术.docx"

# Colors
C_PRIMARY = RGBColor(0x2b, 0x2d, 0x42)   # deep blue-gray
C_SECONDARY = RGBColor(0x8d, 0x99, 0xae) # gray
C_ACCENT = RGBColor(0xef, 0x23, 0x3c)    # red
C_LIGHT = RGBColor(0xed, 0xf2, 0xf4)     # light gray

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_colored_heading(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_callout(doc, text, bg_hex="edf2f4"):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    cell.text = text
    for para in cell.paragraphs:
        para.runs[0].font.color.rgb = C_PRIMARY
        para.runs[0].font.size = Pt(11)
    doc.add_paragraph()

def docx_main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Default font
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = C_PRIMARY

    # ---- COVER ----
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading('学员手册', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = C_PRIMARY

    subtitle = doc.add_paragraph('向上管理与高层说服术')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)
    subtitle.runs[0].font.color.rgb = C_ACCENT
    subtitle.runs[0].font.bold = True

    sub2 = doc.add_paragraph('让决策层理解容错的成本逻辑')
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(16)
    sub2.runs[0].font.color.rgb = C_SECONDARY

    doc.add_paragraph()
    doc.add_paragraph()

    # Course info box
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ("课程编号", "08"),
        ("课程时长", "2天（每天6小时）"),
        ("目标学员", "中高层管理者、变革推动者"),
        ("前置要求", "无（建议先学《变革管理基础》）"),
    ]
    for i, (k, v) in enumerate(info_data):
        cell_k = info_table.cell(i, 0)
        cell_v = info_table.cell(i, 1)
        set_cell_bg(cell_k, "2b2d42")
        cell_k.text = k
        cell_k.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell_k.paragraphs[0].runs[0].font.bold = True
        cell_v.text = v
        cell_v.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- TOC ----
    add_colored_heading(doc, '目录', 1, C_PRIMARY)
    toc_items = [
        ("前言", "3"),
        ("模块一：理解容错成本逻辑", "4"),
        ("模块二：分析决策者心理", "6"),
        ("模块三：掌握说服策略", "8"),
        ("模块四：量化变革价值", "10"),
        ("模块五：练习与复盘", "12"),
        ("工具表单速查", "14"),
        ("课后行动清单", "16"),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.color.rgb = C_PRIMARY
        p.add_run('\t' + page)

    doc.add_page_break()

    # ---- 前言 ----
    add_colored_heading(doc, '前言', 1, C_PRIMARY)
    doc.add_paragraph(
        '变革管理者的核心竞争力，不是执行力，而是「让高层理解变革价值」的能力。'
        '大多数变革失败，不是因为方案不好，而是因为缺乏有效的向上说服。'
    )
    doc.add_paragraph(
        '本课程帮助你掌握一套系统的方法论：用数据说话、用框架说服、用策略争取资源。'
        '学完这门课，你将能够精准识别高层决策者类型，用对方能接受的语言讲清楚变革的成本逻辑，'
        '最终赢得持续的高层支持。'
    )

    add_callout(doc, "核心能力：不是让高层「批准」变革，而是让高层「相信」变革的价值", "ef233c")
    doc.add_paragraph("本课程的学习目标：")
    for obj in [
        "识别三类决策者类型（财务导向型/关系导向型/战略导向型）",
        "掌握向上说服四步法：定位受众→量化成本→设计方案→建立信任",
        "学会计算Error Cost（不变革的成本）和Inaction Cost（等待的成本）",
        "设计最小授权方案，降低高层决策门槛",
        "掌握常见高层质疑的应答策略",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj).font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 模块一 ----
    add_colored_heading(doc, '模块一：理解容错成本逻辑', 1, C_PRIMARY)

    add_colored_heading(doc, '1.1 为什么要学这门课', 2, C_SECONDARY)
    doc.add_paragraph(
        '研究表明，70%的变革失败是因为缺乏高层持续支持。而有高层支持的变革项目成功率是无支持项目的3倍。'
        '然而，60%的中层管理者不知道如何有效与高层沟通。'
    )
    add_callout(doc, "关键洞察：不是高层不愿意支持变革，而是没有看到足够的理由。")

    add_colored_heading(doc, '1.2 容错成本的定义', 2, C_SECONDARY)
    doc.add_paragraph(
        '容错成本（Error Cost）= 问题未及时解决导致的损失 × 等待解决的时间'
    )
    doc.add_paragraph('理解容错成本的意义：')
    for point in [
        "让高层看到「不行动」的真实代价",
        "将抽象的「问题」转化为具体的「金钱损失」",
        "创造变革的紧迫感，推动决策",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(point)

    add_colored_heading(doc, '1.3 决策者的三种类型', 2, C_SECONDARY)

    type_table = doc.add_table(rows=4, cols=3)
    type_table.style = 'Table Grid'
    headers = ["类型", "核心关切", "说服策略"]
    for j, h in enumerate(headers):
        cell = type_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True

    type_data = [
        ("财务导向型", "ROI、回收期、现金流", "用数字说话，展示投资回报"),
        ("关系导向型", "团队稳定、人心向背", "展示团队支持，强调平稳过渡"),
        ("战略导向型", "行业趋势、竞争优势", "引用行业报告，展示竞争价值"),
    ]
    for i, row in enumerate(type_data):
        for j, val in enumerate(row):
            cell = type_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 模块二 ----
    add_colored_heading(doc, '模块二：分析决策者心理', 1, C_PRIMARY)

    add_colored_heading(doc, '2.1 决策者画像工具', 2, C_SECONDARY)
    doc.add_paragraph('使用决策者画像工具，快速识别高层类型：')
    for q in [
        "他最关心的是财务指标还是团队稳定？",
        "他做决策的主要参考是数据还是直觉？",
        "他更关注短期收益还是长期战略？",
        "他上一次批准的大型项目是什么类型的？",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(q)

    add_colored_heading(doc, '2.2 向上说服四步法', 2, C_SECONDARY)
    step_table = doc.add_table(rows=5, cols=3)
    step_table.style = 'Table Grid'
    step_headers = ["步骤", "核心动作", "关键输出"]
    for j, h in enumerate(step_headers):
        cell = step_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True

    step_data = [
        ("第一步：定位受众", "识别决策者类型", "决策者画像表"),
        ("第二步：量化成本", "计算Error Cost和Inaction Cost", "成本量化报告"),
        ("第三步：设计方案", "设计最小授权方案", "最小授权路线图"),
        ("第四步：建立信任", "展示能力，预设应答", "高频应答卡"),
    ]
    for i, row in enumerate(step_data):
        for j, val in enumerate(row):
            cell = step_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 模块三 ----
    add_colored_heading(doc, '模块三：掌握说服策略', 1, C_PRIMARY)

    add_colored_heading(doc, '3.1 最小授权策略', 2, C_SECONDARY)
    doc.add_paragraph(
        '最小授权是一种降低高层决策风险的策略：通过限定变革的范围、时间、资源和决策权，'
        '让高层以最小的赌注开始，看到成果后再逐步扩大。'
    )

    phase_table = doc.add_table(rows=4, cols=4)
    phase_table.style = 'Table Grid'
    for j, h in enumerate(["阶段", "范围", "周期", "授权"]):
        cell = phase_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True

    phase_data = [
        ("第一阶段", "试点", "1-3个月", "执行方式自主，方向需汇报"),
        ("第二阶段", "扩展", "3-6个月", "小幅调整预算和范围"),
        ("第三阶段", "固化", "6-12个月", "全面授权，季度审核"),
    ]
    for i, row in enumerate(phase_data):
        for j, val in enumerate(row):
            cell = phase_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    add_colored_heading(doc, '3.2 汇报叙事结构', 2, C_SECONDARY)
    doc.add_paragraph('高效的变革汇报应遵循以下结构：')
    for s in [
        "开场：1句话点明问题和紧迫性",
        "成本分析：Error Cost + Inaction Cost双视角",
        "方案设计：最小授权，降低风险",
        "预期收益：量化ROI，展示竞争价值",
        "信任建立：主动说风险，展示应对预案",
    ]:
        p = doc.add_paragraph(style='List Number')
        p.add_run(s).font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 模块四 ----
    add_colored_heading(doc, '模块四：量化变革价值', 1, C_PRIMARY)

    add_colored_heading(doc, '4.1 Error Cost计算方法', 2, C_SECONDARY)
    doc.add_paragraph('Error Cost = 问题导致的损失 × 持续时间 × 影响范围')
    doc.add_paragraph('计算步骤：')
    for step in [
        "识别核心问题及其直接损失（如效率损失、质量损失）",
        "估算问题持续时间（已持续多久，还会持续多久）",
        "评估影响范围（哪些部门/流程受影响）",
        "汇总计算年度总损失",
    ]:
        p = doc.add_paragraph(style='List Number')
        p.add_run(step)

    add_colored_heading(doc, '4.2 ROI计算公式', 2, C_SECONDARY)
    doc.add_paragraph('ROI = (年化收益 - 投资成本) / 投资成本 × 100%')
    doc.add_paragraph('关键指标：')
    for metric in [
        "投资成本：一次性投入 + 持续运营成本",
        "年化收益：效率提升 + 成本节省 + 质量改善",
        "回收期：投资成本 / 月均净收益",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(metric)

    add_callout(doc, "提示：使用F01-F04工具表单计算你的项目ROI", "edf2f4")

    doc.add_page_break()

    # ---- 模块五 ----
    add_colored_heading(doc, '模块五：练习与复盘', 1, C_PRIMARY)

    add_colored_heading(doc, '5.1 情境模拟练习', 2, C_SECONDARY)
    doc.add_paragraph('练习一：决策者类型识别')
    doc.add_paragraph('阅读3个高管简介，判断其决策类型，并说明判断依据。')
    doc.add_paragraph()
    doc.add_paragraph('练习二：Error Cost计算')
    doc.add_paragraph('使用F01工具表，计算你所在部门的一个核心问题的Error Cost。')
    doc.add_paragraph()
    doc.add_paragraph('练习三：最小授权方案设计')
    doc.add_paragraph('将一个大型变革项目拆解为三阶段最小授权方案。')

    add_colored_heading(doc, '5.2 案例分析', 2, C_SECONDARY)
    for case in [
        "案例一：ERP升级提案被拒——财务导向型CEO如何说服",
        "案例二：组织变革被搁置——关系导向型CEO如何应对",
        "案例三：成功说服的关键转折——战略导向型CEO的沟通策略",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(case)

    doc.add_page_break()

    # ---- 工具表单速查 ----
    add_colored_heading(doc, '工具表单速查', 1, C_PRIMARY)

    tool_table = doc.add_table(rows=8, cols=3)
    tool_table.style = 'Table Grid'
    for j, h in enumerate(["表单编号", "表单名称", "核心用途"]):
        cell = tool_table.cell(0, j)
        set_cell_bg(cell, "2b2d42")
        cell.text = h
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cell.paragraphs[0].runs[0].font.bold = True

    tools = [
        ("F01", "Error Cost计算表", "量化不变革的年度损失"),
        ("F02", "Inaction Cost计算表", "量化等待机会的成本"),
        ("F03", "ROI计算表", "计算变革项目投资回报率"),
        ("F04", "决策者画像表", "分析高层决策者类型和关切"),
        ("F05", "最小授权方案设计表", "设计分阶段授权路线图"),
        ("F06", "汇报叙事检查表", "检查汇报结构和内容完整性"),
        ("F07", "高频应答卡", "准备高层常见质疑的应答策略"),
    ]
    for i, row in enumerate(tools):
        for j, val in enumerate(row):
            cell = tool_table.cell(i+1, j)
            if i % 2 == 0:
                set_cell_bg(cell, "edf2f4")
            cell.text = val
            cell.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY

    doc.add_page_break()

    # ---- 课后行动清单 ----
    add_colored_heading(doc, '课后行动清单', 1, C_PRIMARY)
    doc.add_paragraph('课程结束后30天内完成以下行动：')

    actions = [
        ("第一周", [
            "选择一个正在推动的变革项目",
            "识别项目的核心决策者，判断其类型",
            "使用F01工具表计算项目的Error Cost",
        ]),
        ("第二周", [
            "使用F04决策者画像表完成决策者分析",
            "设计最小授权方案（F05工具表）",
            "准备3个高频问题的应答策略（F07工具表）",
        ]),
        ("第三周", [
            "向高层进行一次改革汇报",
            "汇报后记录高层反馈，更新策略",
        ]),
        ("第四周", [
            "复盘汇报效果，总结经验教训",
            "与导师/同事分享学习心得",
        ]),
    ]

    for week, items in actions:
        add_colored_heading(doc, week, 2, C_SECONDARY)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)

    doc.add_paragraph()
    add_callout(doc, "记住：向上说服不是一次性的，而是一个持续建立信任的过程。", "ef233c")

    # Save
    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    docx_main()
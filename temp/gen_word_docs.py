# -*- coding: utf-8 -*-
"""
跨部门协作剧本杀 · 阶段 4A · 4 个 Word 模板
- 01-培训师手册.docx
- 02-学员手册.docx
- 03-角色剧本.docx
- 04-复盘报告.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\2026年课程\ai课2026整理\剧本杀\跨部门协作剧本杀\可视化输出\Office物料"
os.makedirs(OUT_DIR, exist_ok=True)

# 颜色（剧本厂牌 Editorial）
INK = RGBColor(0x1A, 0x1A, 0x1A)
RED = RGBColor(0x8B, 0x1A, 0x1A)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
BLUE = RGBColor(0x2E, 0x5B, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_run_font(run, size=10.5, bold=False, color=INK, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_para(doc, text, size=10.5, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT, font_name="Microsoft YaHei"):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, font_name=font_name)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: RED, 2: INK, 3: BLUE}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, INK))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, color=INK, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + indent)
    run = p.add_run("• " + text)
    set_run_font(run, size=10.5, color=color)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_callout(doc, text, color=RED, bg="FAF6F2"):
    """加底色的高亮块"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg)
    pPr.append(shd)
    run = p.add_run("「 " + text + " 」")
    set_run_font(run, size=10.5, color=color, bold=True)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # 标题行
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(255, 255, 255))
        # 黑底
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A1A1A")
        tcPr.append(shd)
    # 数据行
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[1 + r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10, color=INK)
            # 偶数行底色
            if r_i % 2 == 1:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F4F0E8")
                tcPr.append(shd)
    # 列宽
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


def setup_doc(doc, title, subtitle="跨部门协作剧本杀 · v1.0 · 2026-06-07"):
    """统一页面：横向 A4 + 页眉"""
    for section in doc.sections:
        section.page_height = Cm(21.0)
        section.page_width = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    # 封面
    add_para(doc, "跨部门协作剧本杀", size=11, bold=True, color=RED,
             align=WD_ALIGN_PARAGRAPH.CENTER, font_name="Microsoft YaHei")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CROSS-DEPARTMENT COLLABORATION SCRIPT-KILLER")
    set_run_font(r, size=9, color=GRAY, font_name="Courier New")
    add_para(doc, title, size=28, bold=True, color=INK,
             align=WD_ALIGN_PARAGRAPH.CENTER, font_name="SimHei")
    add_para(doc, subtitle, size=9, color=GRAY,
             align=WD_ALIGN_PARAGRAPH.CENTER, font_name="Courier New")
    add_para(doc, "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", size=10, color=INK,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


# ============================================================
# 01 培训师手册.docx
# ============================================================
def doc_trainer_manual():
    doc = Document()
    setup_doc(doc, "培训师手册")

    add_heading(doc, "00  培训师守则 · 10 条不可逾越的红线", 1)
    rules = [
        "绝不替学员下结论。觉察 = 学员自己的,不是你塞的。",
        "绝不泄露卡片信息。卡片 ≠ 答案。",
        "绝不惩罚学员。失分 = 学习机会。",
        "绝不简化流程。4 小时主流程 = 4 小时,不能缩。",
        "绝不暗示答案。引导 = 提问,不是给答案。",
        "绝不让 PM 替 VP 表态。PM = 桥梁,不是代言人。",
        "绝不让 CEO 跳过 PM 决断。CEO = 听者,不是拍脑袋者。",
        "绝不忽视 12:30 硬窗口。12:30 之后 = 协同失败。",
        "绝不使用 12 敏感词。差异化信息→独有信息。策略性→主动。",
        "绝不忽视 5-31 验证。5-31 = 真实水平 = 5 项必过。",
    ]
    for r in rules:
        add_bullet(doc, r, color=INK)

    add_heading(doc, "01  总时刻表 · 4 小时主流程 + 7 个不可撤回时间窗", 1)
    add_table(doc, ["时间", "阶段", "动作", "责任人"],
              [["09:00-09:30", "启动会", "5 VP 公开陈述 5 分钟 × 5", "7 角色全员"],
               ["09:30-10:00", "全员质询", "5 VP 互相质询 30 分钟", "7 角色全员"],
               ["11:00-12:00", "私下对账", "PM 5 VP × 5 分钟", "PM"],
               ["12:00-12:30", "升级邮件", "PM 4 段式风险汇总", "PM"],
               ["16:00-16:45", "CEO 决断", "5 VP 表态 + CEO 三选一", "CEO + 5 VP"],
               ["17:00-17:30", "兑现承诺", "5 VP 5/14 兑现承诺", "5 VP"],
               ["60 min", "复盘", "5 阶段 60 分钟", "培训师"]],
              col_widths=[3, 2.5, 7, 4])

    add_heading(doc, "02  7 角色矩阵", 1)
    add_table(doc, ["角色", "职级", "私利", "路线 1 立场"],
              [["林晓晨 PM", "项目集经理", "5 VP 桥梁", "支持（协调）"],
               ["陈伟 CTO", "首席技术官", "5/31 接 CEO 候补", "支持（91% 复评）"],
               ["王敏 CLO", "首席法务官", "5-31 必过 5 项", "支持（5 项清单）"],
               ["张建国 CFO", "首席财务官", "保住 CFO 职位", "支持（80 万分摊）"],
               ["赵丽 CMO", "首席市场官", "4-25 复评基准", "支持（减范围 30%）"],
               ["刘芳 CHO", "首席人事官", "9 月接班人", "支持（5 VP 副手）"],
               ["方远航 CEO", "创始人", "B 轮对赌", "听 PM + 5 VP"]],
              col_widths=[3, 3, 5, 5])

    add_heading(doc, "03  13 风险码全清单", 1)
    add_table(doc, ["风险码", "责任 VP", "风险描述", "必失分点"],
              [["R-01", "张建国", "物料 80 万缺口", "80 万分摊失败"],
               ["R-02", "陈伟", "AI 95% 不达标", "91% 35 天 + 误差 ≤3"],
               ["R-03", "赵丽", "客户流失 30%", "4-25 复评基准"],
               ["R-04", "王敏", "5-31 飞行检查", "5 项必过清单"],
               ["R-05", "刘芳", "客户合同 27 天", "5-10 培训达标"],
               ["R-06", "王敏", "监管罚款 500 万", "路线 2 必失分"],
               ["R-07", "PM", "团队协作断裂", "5 VP 协同失败"],
               ["R-08", "PM", "项目延期 4-30", "物料 27 天"],
               ["R-09 ★", "CEO", "★ 核心爆炸点 ★", "拍脑袋 27 天"],
               ["R-10", "张建国", "资源争夺", "80 万分摊失败"],
               ["R-11", "刘芳", "5-10 培训达标", "减范围 30%"],
               ["R-12", "PM", "接班人问题", "9 月副手"],
               ["R-13", "CEO", "战略转型", "B 轮对赌"]],
              col_widths=[2, 2.5, 5, 5])

    add_heading(doc, "04  三轮剧本 · 时长 + 关键动作", 1)
    add_table(doc, ["轮次", "时段", "动作", "硬窗口"],
              [["第 1 轮", "09:00-10:00", "启动会 + 全员质询", "10:00 结束"],
               ["第 2 轮", "11:00-12:30", "私下对账 + 升级邮件", "12:30 决断"],
               ["第 3 轮", "16:00-17:30", "CEO 决断 + 兑现承诺", "17:30 结束"]],
              col_widths=[2.5, 3, 6, 4])

    add_heading(doc, "05  25 问引导问题库", 2)
    add_para(doc, "5 阶段 × 5 问题 = 25 个。每个对应 1 个失分点。", size=10, color=GRAY)
    qs = [
        "01  启动会 · 5 VP 装不知道了吗？",
        "02  启动会 · 13 风险码提到了几个？",
        "03  启动会 · 谁察觉 R-09 拍脑袋？",
        "04  启动会 · 5 VP 私利 = 路线 1/2/3？",
        "05  启动会 · PM 第一反应是什么？",
        "06  升级邮件 · 风险码全吗？",
        "07  升级邮件 · 私下对账谈了什么？",
        "08  升级邮件 · 路线 1 都同意吗？",
        "09  升级邮件 · 12:30 前都表态了吗？",
        "10  升级邮件 · PM 怎么汇总 5 VP 冲突？",
        "11  CEO 决断 · CEO 听 PM 了吗？",
        "12  CEO 决断 · 5 VP 有人变脸吗？",
        "13  CEO 决断 · CEO 选了什么？",
        "14  CEO 决断 · 5/14 承诺具体吗？",
        "15  CEO 决断 · 你会怎么决断？",
        "16  物料承诺 · 5 VP 谁最具体？",
        "17  物料承诺 · 80 万怎么分？",
        "18  物料承诺 · 5-10 减范围具体动作？",
        "19  物料承诺 · 9 月 PM 副手是谁？",
        "20  物料承诺 · 5/14 兑现 = 5-25 必达,你信吗？",
        "21  5-31 验证 · 5 项过了几个？",
        "22  5-31 验证 · 4-30 失分还能补救吗？",
        "23  5-31 验证 · 5 VP 谁贡献最大？",
        "24  5-31 验证 · 重演一遍会改什么？",
        "25  5-31 验证 · 最大的觉察是什么？",
    ]
    for q in qs:
        add_bullet(doc, q, color=BLUE, indent=0.3)

    add_heading(doc, "06  评分标准 · 4 维度", 1)
    add_table(doc, ["维度", "权重", "评分点"],
              [["风险码识别", "40%", "13 个 × 3 轮 = 39 个表态点"],
               ["跨部门协同", "25%", "5 VP × 3 轮一致性"],
               ["决断力", "20%", "CEO 听 PM + 5 VP"],
               ["5/14 兑现", "15%", "5 VP 各自本部门承诺"]],
              col_widths=[3, 2, 8])
    add_callout(doc, "加分项：R-09 单独识别 +10；5 VP 三轮一致 +5；5/14 出现具体数字 +5。失分项：CEO 拍脑袋 -10；PM 12:30 迟到 -10；互相推诿 -10。", color=GOLD, bg="FAF6E8")

    add_heading(doc, "07  复盘剧本 · 60 分钟", 1)
    add_table(doc, ["时段", "阶段", "动作", "学员动作"],
              [["0-10 min", "回顾", "重放 16:00 决断片段", "5 VP 自评"],
               ["10-25 min", "诊断", "一致性矩阵展示", "VP 之间互评"],
               ["25-40 min", "觉察", "引导我以为 vs 原来", "个人觉察发言"],
               ["40-50 min", "承诺", "引导 5/14 我做什么", "5 VP 各写 1 句"],
               ["50-60 min", "落地", "引导 5-31 真实水平验证", "5 VP 各写 1 句"]],
              col_widths=[2.5, 2, 5, 5])

    add_heading(doc, "08  物料清单 · 24 件", 1)
    mats = [
        "13 风险码卡（独立塑封）", "18 信息卡（6 角色 × 3 轮）", "评分表 7 张 / 角色",
        "时间线挂图 1 张", "复盘模板 7 张", "引导问题卡 7 张",
        "5 VP 角色剧本 × 7", "学员手册 7 本", "培训师手册 1 本",
        "白板 2 块 + 白板笔 4 色", "桌签 × 7", "A3 草稿纸 × 7",
        "签到表", "5-31 验证表 × 7", "觉察点清单 × 7",
        "计时器 × 2", "录播设备 1 套", "投影 1 套",
        "便利贴 × 100", "签字笔 × 14", "打印 13 + 18 + 24 = 55 件",
    ]
    for m in mats:
        add_bullet(doc, m, color=INK, indent=0.2)

    add_heading(doc, "09  培训师 FAQ · 20 问", 2)
    faqs = [
        "Q1 启动会有人迟到 5 分钟？A: 准时开,迟到者失分。",
        "Q2 5 VP 装不知道？A: 培训师请其补充,不直接戳破。",
        "Q3 PM 替 VP 表态？A: 培训师请 VP 亲自说。",
        "Q4 CEO 拍脑袋？A: 16:00 决断前请 PM 补充升级邮件。",
        "Q5 路线 2 强烈支持？A: 培训师请王敏 CLO 报 5-31 必失分。",
        "Q6 12:30 硬窗口迟到？A: 失分 10 分,记入评分表。",
        "Q7 5 VP 三轮变脸？A: 一致性矩阵失分,记入。",
        "Q8 5/14 兑现空话？A: 培训师请 VP 给出具体数字。",
        "Q9 80 万分摊失败？A: 路线 1 共识失败,记入。",
        "Q10 9 月接班人不报？A: 失分 10 分,记入。",
        "Q11 复盘学员不说话？A: 培训师点名 1 分钟。",
        "Q12 复盘学员哭？A: 停 5 分钟,共情。",
        "Q13 5-31 飞行检查没准备好？A: 复盘推迟到 5-31 后。",
        "Q14 学员问 R-09 是什么？A: 不答,让 5 VP 自己发现。",
        "Q15 学员问怎么评 5 VP？A: 培训师按 4 维度评分,加 5 VP 互评。",
        "Q16 学员要求简化流程？A: 拒绝,4 小时 = 4 小时。",
        "Q17 学员要求看答案？A: 不给,卡片 ≠ 答案。",
        "Q18 学员之间吵架？A: 培训师叫停,记入复盘。",
        "Q19 学员要求延长？A: 不延长,60 分钟复盘 = 60 分钟。",
        "Q20 复盘没觉察？A: 培训师引导我以为 vs 原来,直到觉察。",
    ]
    for f in faqs:
        add_bullet(doc, f, color=INK, indent=0.2)

    add_heading(doc, "10  附录 · 12 敏感词 + 致谢", 1)
    add_callout(doc, "12 敏感词（绝不使用）：策略性 · 差异化信息 · 不可告人 · 嫁祸 · 攻击 · 防御 · 报复 · 甩锅 · 博弈 · 操控 · 信息窝藏 · 策略性隐瞒", color=RED, bg="FAF6F2")
    add_callout(doc, "安全替换：差异化信息→独有信息/特权信息；策略性→主动；防御→警觉/谨慎；甩锅→责任划分/归因；博弈→协商/权衡", color=GREEN if False else BLUE, bg="EEF1FA")
    add_para(doc, "—— 跨部门协作剧本杀 · 培训师手册 · v1.0 · 2026-06-07 编订 ——",
             size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(OUT_DIR, "01-培训师手册.docx")
    doc.save(out)
    print(f"[OK] {out}")


# ============================================================
# 02 学员手册.docx
# ============================================================
def doc_student_manual():
    doc = Document()
    setup_doc(doc, "学员手册")

    add_heading(doc, "00  欢迎 · 跨部门协作剧本杀", 1)
    add_para(doc, "你是 7 角色之一。4 小时主流程 + 60 分钟复盘 = 5 小时,体验 49 天跨部门协作。", size=11)
    add_callout(doc, "复盘不是找错,是觉察模式。觉察比共识更重要,承诺比分析更重要,验证比承诺更重要。", color=BLUE, bg="EEF1FA")

    add_heading(doc, "01  7 选 1 角色矩阵", 1)
    add_table(doc, ["角色", "一句话定位", "难度"],
              [["林晓晨 PM", "5 VP 桥梁,12:30 升级邮件", "★★★★★"],
               ["陈伟 CTO", "5/31 接 CEO 候补,AI 91% 复评", "★★★★"],
               ["王敏 CLO", "5-31 必过 5 项,5 项清单", "★★★"],
               ["张建国 CFO", "80 万分摊,保住 CFO", "★★★★"],
               ["赵丽 CMO", "4-25 复评基准,5-10 必过", "★★★"],
               ["刘芳 CHO", "5-10 培训达标,9 月接班", "★★★"],
               ["方远航 CEO", "16:00 决断,听 = 95", "★★★★★"]],
              col_widths=[3, 9, 3])

    add_heading(doc, "02  三轮时刻表", 1)
    add_table(doc, ["轮次", "时间", "必做", "硬窗口"],
              [["第 1 轮", "09:00-10:00", "陈述 + 质询", "10:00 结束"],
               ["第 2 轮", "11:00-12:30", "对账 + 升级邮件", "12:30 决断"],
               ["第 3 轮", "16:00-17:30", "表态 + 兑现承诺", "17:30 结束"]],
              col_widths=[2.5, 3, 6, 4])

    add_heading(doc, "03  13 风险码速查", 1)
    add_table(doc, ["风险码", "责任", "一句话", "失分点"],
              [["R-01", "CFO", "80 万缺口", "分摊失败"],
               ["R-02", "CTO", "AI 95% 不达", "91% 35 天"],
               ["R-03", "CMO", "客户流失 30%", "4-25 复评"],
               ["R-04", "CLO", "5-31 检查", "5 项必过"],
               ["R-05", "CHO", "合同 27 天", "5-10 培训"],
               ["R-06", "CLO", "罚款 500 万", "路线 2"],
               ["R-07", "PM", "团队断裂", "协同失败"],
               ["R-08", "PM", "延期 4-30", "物料 27 天"],
               ["R-09 ★", "CEO", "★ 核心爆炸点", "拍脑袋"],
               ["R-10", "CFO", "资源争夺", "80 万"],
               ["R-11", "CHO", "5-10 培训", "减范围"],
               ["R-12", "PM", "接班人", "9 月副手"],
               ["R-13", "CEO", "战略转型", "B 轮对赌"]],
              col_widths=[2, 2, 5, 5])

    add_heading(doc, "04  评分规则", 1)
    add_table(doc, ["维度", "权重", "怎么算"],
              [["风险码识别", "40%", "13 × 3 = 39 个表态点"],
               ["跨部门协同", "25%", "5 VP × 3 轮一致"],
               ["决断力", "20%", "CEO 听 PM + 5 VP"],
               ["5/14 兑现", "15%", "5 VP 具体承诺"]],
              col_widths=[3, 2, 8])

    add_heading(doc, "05  5 种必失分行为陷阱", 1)
    traps = [
        "陷阱 1：装不知道。12:00 必暴露,失分 5 分。",
        "陷阱 2：替 VP 表态。越权,失分 10 分。",
        "陷阱 3：12:30 硬窗口迟到。协同失败,失分 10 分。",
        "陷阱 4：5/14 兑现空话。失分 10 分。",
        "陷阱 5：CEO 拍脑袋。R-09 失分 10 分。",
    ]
    for t in traps:
        add_bullet(doc, t, color=RED, indent=0.2)

    add_heading(doc, "06  复盘指南", 1)
    add_para(doc, "5 阶段 60 分钟：回顾 → 诊断 → 觉察 → 承诺 → 落地", size=11)
    add_callout(doc, "我以为……,原来……。这是复盘的核心,不是找错。", color=GOLD, bg="FAF6E8")

    add_heading(doc, "07  致谢 + 12 敏感词自查", 1)
    add_callout(doc, "12 敏感词自查：策略性 · 差异化信息 · 不可告人 · 嫁祸 · 攻击 · 防御 · 报复 · 甩锅 · 博弈 · 操控 · 信息窝藏 · 策略性隐瞒", color=RED, bg="FAF6F2")
    add_para(doc, "—— 跨部门协作剧本杀 · 学员手册 · v1.0 ——", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(OUT_DIR, "02-学员手册.docx")
    doc.save(out)
    print(f"[OK] {out}")


# ============================================================
# 03 角色剧本.docx（7 角色合一）
# ============================================================
def doc_role_script():
    doc = Document()
    setup_doc(doc, "角色剧本 · 7 角色合一")

    roles = [
        {
            "name": "林晓晨", "title": "项目集经理 PM", "color": "蓝 #2E5BFF", "private": "5 VP 桥梁,CEO 耳朵,信息中枢",
            "script_09": "CEO 27 天,我理解压力——但 27 天物料 27 天 AI 91% 是不可能的。13 风险码已识别,12:00 升级邮件给 3 方案。",
            "script_12": "升级邮件 4 段式:风险码清单 + 3 方案 + 私下对账结果 + 决断请求。12:30 硬窗口。",
            "script_16": "听 CEO 听不听 5 VP。CEO 选路线 1 = 5 VP 协同 = 95;路线 3 = 张建国保位 = 9 月接班失分。",
            "must": ["5 VP 私利处理", "12:30 升级邮件", "不替 VP 表态", "5/14 兑现跟踪"]
        },
        {
            "name": "陈伟", "title": "首席技术官 CTO", "color": "绿 #1A8E5F", "private": "5/31 接 CEO 候补,AI 91% 复评",
            "script_09": "AI 95% 不可能,91% 35 天训练 + 14 天复评,误差 ≤3 个百分点。4-25 复评 = 减范围 30%。",
            "script_12": "AI 91% 复评数据:35 天训练(4-13→5-18)+ 14 天复评(5-18→5-31)。5-25 必达 35 天。",
            "script_16": "我强烈支持路线 1。5-25 必达 = 5/31 接 CEO 候补。",
            "must": ["AI 95% 不说不可能", "4-25 复评不说减范围", "误差 ≤3", "5-25 必达"]
        },
        {
            "name": "王敏", "title": "首席法务官 CLO", "color": "棕 #8B4513", "private": "5-31 必过 5 项,法务部门核心 KPI",
            "script_09": "5-31 上海网信办飞行检查 5 项必过。路线 2(推 14 天)= 5-31 必失分 = 罚款 500 万。我强烈反对。",
            "script_12": "5 项必过清单:数据来源 + 算法合规 + 合同条款 + 培训覆盖 + 复评报告。5-25 必达。",
            "script_16": "我强烈支持路线 1。5-25 必达 5 项 = 5-31 必过。",
            "must": ["5 项清单", "路线 2 强烈反对", "5-25 必达", "5-31 必过"]
        },
        {
            "name": "张建国", "title": "首席财务官 CFO", "color": "金 #B8860B", "private": "保住 CFO + 9 月接班人",
            "script_09": "80 万物料缺口 = 5 VP 分摊 16 万 × 5 = 80 万。4-30 物料 27 天 = CFO 功劳。",
            "script_12": "80 万分摊方案:IT 16 万 + 法务 16 万 + 财务 16 万 + 市场 16 万 + HR 16 万 = 80 万。",
            "script_16": "我支持路线 1。80 万分摊 = 路线 1 共识 = CFO 功劳。",
            "must": ["80 万分摊", "路线 3 警觉", "4-30 物料 27 天", "9 月接班副手"]
        },
        {
            "name": "赵丽", "title": "首席市场官 CMO", "color": "玫红 #C71585", "private": "4-25 复评基准,客户流失 30%",
            "script_09": "客户流失 30%。5-10 客户合同 27 天 = 必过。4-25 复评基准 = 减范围 30%。",
            "script_12": "4-25 复评方案:复评基准 + 客户合同 + 客户流失 <30% + 5/14 兑现。",
            "script_16": "我支持路线 1。4-25 复评基准 = 5-10 必过 = 客户合同 27 天。",
            "must": ["4-25 复评", "5-10 客户合同 27 天", "客户流失 <30%", "5/14 兑现"]
        },
        {
            "name": "刘芳", "title": "首席人事官 CHO", "color": "灰 #708090", "private": "9 月接班人 + 5 VP 部门副手",
            "script_09": "5-10 客户合同 27 天 = 培训 100% 覆盖。9 月接班 = 5 VP 部门副手。",
            "script_12": "5-10 培训达标方案:培训覆盖 + 减范围 30% + 5 VP 副手 + 5/14 兑现。",
            "script_16": "我支持路线 1。5-10 培训达标 100% 覆盖 = 9 月接班。",
            "must": ["5-10 培训达标", "减范围 30%", "5 VP 副手", "9 月接班"]
        },
        {
            "name": "方远航", "title": "创始人 CEO", "color": "黑 #1A1A1A", "private": "B 轮对赌 + 5-31 接 CEO 候补",
            "script_09": "4-30 B 轮对赌物料 27 天,是我定的目标,这是拍脑袋的。16:00 决断,我听 5 VP。",
            "script_12": "PM 升级邮件我全部听完。13 风险码已识别,3 方案,5 VP 私下对账结果。",
            "script_16": "★ 三选一 ★ 我选路线 1。4-30 物料 27 天必达,5-31 必过 5 项。",
            "must": ["听 = 95", "拍脑袋 = 失分", "路线 1 共识", "5 项必过"]
        }
    ]

    for i, r in enumerate(roles, 1):
        if i > 1:
            doc.add_page_break()
        add_heading(doc, f"角色 {i:02d}  {r['name']} · {r['title']}", 1)
        add_para(doc, f"代表色：{r['color']}", size=10, color=GRAY)
        add_callout(doc, f"私利：{r['private']}", color=BLUE, bg="EEF1FA")

        add_heading(doc, f"09:00 启动会 · 陈述 5 分钟", 2)
        add_para(doc, r["script_09"], size=11)

        add_heading(doc, f"12:00 升级邮件 · 配合 PM", 2)
        add_para(doc, r["script_12"], size=11)

        add_heading(doc, f"16:00 CEO 决断 · 最终表态", 2)
        add_para(doc, r["script_16"], size=11)

        add_heading(doc, f"必失分点 · {r['name']} 专属", 2)
        for m in r["must"]:
            add_bullet(doc, m, color=RED, indent=0.3)

    add_para(doc, "—— 跨部门协作剧本杀 · 角色剧本 · v1.0 ——", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(OUT_DIR, "03-角色剧本.docx")
    doc.save(out)
    print(f"[OK] {out}")


# ============================================================
# 04 复盘报告.docx
# ============================================================
def doc_debrief_report():
    doc = Document()
    setup_doc(doc, "复盘报告 · 5 阶段 60 分钟")

    add_heading(doc, "00  复盘报告封面", 1)
    add_table(doc, ["场次", "日期", "培训师", "地点", "学员数"],
              [["第 __ 场", "2026-__-__", "__________", "__________", "7 人"]],
              col_widths=[3, 3, 3, 4, 3])

    add_heading(doc, "01  5 VP 自评 · 回顾", 1)
    add_para(doc, "每位 VP 用 1 分钟回答：今天我做得好的 1 点。具体到 1 个动作。", size=11)
    vps = ["林晓晨 PM", "陈伟 CTO", "王敏 CLO", "张建国 CFO", "赵丽 CMO", "刘芳 CHO"]
    for vp in vps:
        add_para(doc, f"{vp} 自评：", size=11, bold=True)
        add_para(doc, "                                                  ", size=11, color=GRAY)
        add_para(doc, "                                                  ", size=11, color=GRAY)

    add_heading(doc, "02  一致性矩阵 · 诊断", 1)
    add_table(doc, ["VP / 轮", "09:00 启动会", "12:00 升级邮件", "16:00 CEO 决断", "一致性"],
              [[vp, "□", "□", "□", "□ 一致 □ 变脸"] for vp in vps],
              col_widths=[3, 3, 3, 3, 4])

    add_heading(doc, "03  觉察 · 我以为 vs 原来", 1)
    add_para(doc, "每位学员写：我以为……,原来……。这是复盘的核心。", size=11)
    for vp in vps:
        add_para(doc, f"{vp} 我以为：", size=11, bold=True)
        add_para(doc, "                                                  ", size=11, color=GRAY)
        add_para(doc, f"{vp} 原来：", size=11, bold=True)
        add_para(doc, "                                                  ", size=11, color=GRAY)

    add_heading(doc, "04  承诺 · 5/14 我做什么", 1)
    add_table(doc, ["VP", "5/14 我承诺", "具体数字"],
              [[vp, "", ""] for vp in vps],
              col_widths=[3, 8, 4])
    add_callout(doc, "好承诺：5/14 前完成 AI 模型 91% 复评,准确率误差 ≤3 个百分点。坏承诺：我会加强一下。", color=GOLD, bg="FAF6E8")

    add_heading(doc, "05  落地 · 5-31 真实水平验证", 1)
    add_table(doc, ["VP", "5-31 我验证", "5 项必过"],
              [[vp, "", ""] for vp in vps],
              col_widths=[3, 8, 4])
    add_callout(doc, "觉察比共识更重要,承诺比分析更重要,验证比承诺更重要。", color=BLUE, bg="EEF1FA")

    add_heading(doc, "06  培训师总评", 1)
    add_table(doc, ["项", "评语"],
              [["最大亮点", ""],
               ["最大风险", ""],
               ["改进建议", ""],
               ["5-31 预测", ""]],
              col_widths=[3, 12])

    add_heading(doc, "07  评分汇总", 1)
    add_table(doc, ["角色", "风险码 40%", "协同 25%", "决断 20%", "兑现 15%", "加权总分"],
              [[vp, "__/100", "__/100", "__/100", "__/100", "__分"] for vp in vps],
              col_widths=[3, 2.5, 2, 2, 2, 2])

    add_heading(doc, "08  等级 · 培训师印章", 1)
    add_table(doc, ["等级", "分数", "印章"],
              [["优秀", "90+", "★"],
               ["良好", "70-90", "○"],
               ["合格", "50-70", "△"],
               ["不合格", "<50", "×"]],
              col_widths=[3, 3, 4])

    add_para(doc, "")
    add_para(doc, "培训师签字：__________  日期：__________", size=11)
    add_para(doc, "PM 签字：__________  日期：__________", size=11)

    add_para(doc, "—— 跨部门协作剧本杀 · 复盘报告 · v1.0 ——", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = os.path.join(OUT_DIR, "04-复盘报告.docx")
    doc.save(out)
    print(f"[OK] {out}")


if __name__ == "__main__":
    doc_trainer_manual()
    doc_student_manual()
    doc_role_script()
    doc_debrief_report()
    print("\n=== 4 个 Word 文档生成完毕 ===")

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 课后资源
res = doc.add_paragraph()
r = res.add_run('课后资源')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

doc.add_paragraph()
rt = doc.add_paragraph()
r = rt.add_run('延伸阅读清单')
r.bold = True; r.font.size = Pt(16)

for book in ['《复盘+：把经验转化为能力》—— 邱昭良', '《U型理论》—— Otto Scharmer', '《行动学习实战录》', '《组织能力的杨三角》—— 杨国安', '《斯坦福大学最受欢迎的创意课》']:
    doc.add_paragraph(book, style='List Bullet')

doc.add_paragraph()
hwt = doc.add_paragraph()
r = hwt.add_run('实践作业')
r.bold = True; r.font.size = Pt(16)

for i, h in enumerate([
    '选择一个让你真正纠结过的决策，用三问法评估它是否值得做成决策卡',
    '找到那个决策的相关人员，进行一次完整的复盘访谈',
    '基于访谈结果，制作一张完整的决策卡',
    '为你制作的决策卡安排一个反对者评审，收集反馈并修改',
    '选择一个失败案例，写出它的警示清单三段式内容',
], 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. ')
    r.bold = True
    p.add_run(h)

doc.add_page_break()
print('S14: Post-class resources done')
doc.save(r'D:\CC\temp\handbook_s14.docx')

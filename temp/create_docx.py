# File: D:/CC/temp/create_docx.py
# Run: python D:/CC/temp/create_docx.py
import sys
import os

# Add the path for python-docx
sys.path.insert(0, 'C:/Users/Administrator/.claude/skills/Word文档处理/scripts/python')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Create document
doc = Document()

# Set default font for CJK
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# Helper function to add heading with Chinese font
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

# Helper function to add paragraph with formatting
def add_para(doc, text, bold=False, italic=False, font_size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    return para

# Helper function to add table
def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return table

# ============ DOCUMENT CONTENT ============

# Title
title = doc.add_heading('06 不确定性焦虑：在模糊局面中做出决策', 0)
for run in title.runs:
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_para(doc, '讲师手册')
doc.add_paragraph('─' * 50)

# Version info
add_para(doc, '版本：v1.0')
add_para(doc, '最后更新：2026-08-20')
add_para(doc, '本手册为内部培训使用材料，请勿对外传播')
doc.add_paragraph()

# ============ PART 1 ============
add_heading(doc, '第一部分：讲师指南', 1)

add_heading(doc, '课程基本信息', 2)

headers = ['属性', '内容']
rows = [
    ['课程代码', '06'],
    ['课程名称', '不确定性焦虑：在模糊局面中做出决策'],
    ['课程定位', '焦虑系列第3课'],
    ['目标学员', '面对未明确结果的情况时容易陷入反复纠结、迟迟无法行动的人'],
    ['课程时长', '两日课程（6小时/日，共12小时）'],
    ['班级规模', '建议20-30人'],
    ['核心价值', '解决不确定性本身带来的心理消耗'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# Course Background
add_heading(doc, '课程背景与问题起点', 2)

add_heading(doc, '什么是不确定性焦虑', 3)
add_para(doc, '不确定性焦虑（Intolerance of Uncertainty，简称IU）是一种对未知情境的强烈心理反应。当人们面对无法预测结果的情况时，大脑的杏仁核会被激活，触发"战斗-逃跑-冻结"反应。这种反应在进化上有其意义——人类祖先需要对潜在危险保持敏感——但在现代社会中，它往往导致我们过度准备、反复纠结、迟迟无法做出决定。')

add_heading(doc, '两种典型的焦虑反应', 3)
add_para(doc, '国际心理学研究（IDS, 2026）表明，人们面对不确定性时普遍存在两种极端反应：')
add_para(doc, '1. 过度准备型：认为必须收集全部信息才能行动，结果陷入"分析瘫痪"')
add_para(doc, '2. 冲动行动型：为了逃避不确定性带来的不适，随意做出决定然后后悔')
add_para(doc, '两种反应看似相反，本质上都是对不确定性缺乏健康应对方式的表现。')

add_heading(doc, '本课程要解决的三个核心问题', 3)
headers = ['问题', '内涵']
rows = [
    ['认知重启', '不确定性焦虑的生理和心理机制是什么'],
    ['识别模式', '我在不确定性面前的独特"签名"是什么'],
    ['行动策略', '如何在模糊局面中做出有效决策'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# Instructor Requirements
add_heading(doc, '讲师资质要求', 2)
headers = ['要求类别', '具体条件']
rows = [
    ['专业知识', '具备心理学或组织行为学基础，熟悉焦虑相关干预技术'],
    ['授课能力', '良好的引导和控场能力，能够处理课堂中的情绪触发情况'],
    ['行业理解', '了解成年人在工作/生活决策中的常见困境'],
    ['认证要求', '建议完成本课程授权讲师认证'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# Pre-class Preparation
add_heading(doc, '课前准备清单', 2)

add_heading(doc, '课程前48小时确认事项', 3)
add_para(doc, '内容熟悉度')
add_para(doc, '• 熟读本次讲授的模块对应的教学文档')
add_para(doc, '• 准备自己在每个模块上的1-2个真实亲身经历案例')
add_para(doc, '• 预演各练习的示范答案，确保能在现场展示')
add_para(doc, '• 熟悉附录中的参考答案，准备好应对学员的不同答案')

add_para(doc, '学员信息')
add_para(doc, '• 了解学员背景：行业、岗位层级、平均工作年限')
add_para(doc, '• 了解学员日常面临的主要不确定性情境（可通过报名问卷收集）')
add_para(doc, '• 识别可能对课程提出质疑或认知抵触的学员类型')
add_para(doc, '• 确认学员规模：20人以内（最佳）、20-30人（可行）、30人以上（需调整互动设计）')

add_para(doc, '材料准备')
add_para(doc, '• 每人一份学员手册（已打印，双面彩印）')
add_para(doc, '• 每人一套配套表单（空表版，已装订）')
add_para(doc, '• 补充用表：A4纸备用，用于额外练习书写')
add_para(doc, '• 白板/翻页纸板，用于现场板书和练习展示')
add_para(doc, '• 计时器（建议使用投影计时，或手机分屏显示）')
add_para(doc, '• 贴纸或卡片，用于练习成果展示（推荐A5卡片）')

add_para(doc, '室内布置')
add_para(doc, '• U型桌或岛型分组座位（4-6人一组）——优先于剧院式座位')
add_para(doc, '• 每桌配备彩色马克笔+白板纸/便签纸')
add_para(doc, '• 讲台可见但不孤立，讲师能方便地走到学员间')
add_para(doc, '• 投影屏幕可从室内所有位置清晰阅读')

add_heading(doc, '课程前1周准备事项', 3)
add_para(doc, '• 确认场地布置（U型桌或分组桌）')
add_para(doc, '• 测试投影/音响设备')
add_para(doc, '• 准备案例视频素材（3段，每段2-3分钟）')
add_para(doc, '• 打印工具表单（按人数，每套约12页）')
add_para(doc, '• 发送课前预习材料给学员')

add_heading(doc, '课程当天早上准备事项', 3)
add_para(doc, '• 提前30分钟到场')
add_para(doc, '• 检查设备正常运行（投影、音响、计时器）')
add_para(doc, '• 摆好桌椅分组，确认每组都有马克笔和白板纸')
add_para(doc, '• 准备好水和小点心')
add_para(doc, '• 将学员手册和工具表单按组摆放')

# How to use teaching documents
add_heading(doc, '如何使用教学文档', 2)
add_para(doc, '共读型教学文档的设计是：讲师打开文档带着讲，学员对着文档一起跟，课后翻开文档能复习。请避免将其当成PPT逐字阅读——关键是引导学员在练习时"停下来动手"。')
headers = ['使用方式', '说明']
rows = [
    ['共读模式', '讲师朗读或引导学员交替阅读场景/案例段落，增加代入感'],
    ['静读模式', '复杂工具讲解前给学员2分钟静读，让每人先形成独立理解再讨论'],
    ['练习模式', '到练习环节时，明确说"现在停下来，大家翻到第X页"'],
    ['参考模式', '课后学员可以作为工具手册反复翻阅，标记对自己有用的工具'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# ============ PART 2 ============
add_heading(doc, '第二部分：时间分配总表', 1)

add_heading(doc, '课程总体时间结构', 2)
headers = ['模块', '标题', '时长', '累计时间', '占比']
rows = [
    ['开场', '课程介绍与开场活动', '15分钟', '15分钟', '2%'],
    ['M1', '模块一：认知重启——不确定性到底是什么', '45分钟', '60分钟', '6%'],
    ['M2', '模块二：识别焦虑——你的不确定性模式', '60分钟', '120分钟', '8%'],
    ['演练一', '演练一：真实情境工作坊', '90分钟', '210分钟', '13%'],
    ['M3', '模块三：拆解行动——止损点设计', '60分钟', '270分钟', '8%'],
    ['M4', '模块四：与模糊共处——心理韧性训练', '60分钟', '330分钟', '8%'],
    ['复习', '第一日回顾与复习', '30分钟', '360分钟', '4%'],
    ['M5', '模块五：综合演练与日常应用', '120分钟', '480分钟', '17%'],
    ['演练二', '演练二：真实情境实践', '60分钟', '540分钟', '8%'],
    ['迁移', '日常应用与30天计划', '60分钟', '600分钟', '8%'],
    ['讲师示范', '关键环节示范', '60分钟', '660分钟', '8%'],
    ['综合演练', '全班参与综合演练', '45分钟', '705分钟', '6%'],
    ['收尾', '收尾：承诺、评估、Q&A', '35分钟', '740分钟', '5%'],
]
add_table(doc, headers, rows)
doc.add_paragraph()
add_para(doc, '总计：740分钟 ≈ 12小时（按两日课程安排）')

# Two-day schedule
add_heading(doc, '两天课程时间分配建议', 2)

add_heading(doc, '第一日（上午）', 3)
headers = ['时段', '时间', '模块', '内容']
rows = [
    ['上午', '09:00-09:15', '开场', '课程介绍、规则建立'],
    ['', '09:15-09:45', 'M1', '认知重启'],
    ['', '09:45-10:30', 'M2', '识别焦虑（前半）'],
    ['', '10:30-10:45', '休息', '茶歇'],
    ['', '10:45-11:45', 'M2（续）', '识别焦虑（后半）'],
    ['午休', '12:15-13:30', '午餐', ''],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '第一日（下午）', 3)
headers = ['时段', '时间', '模块', '内容']
rows = [
    ['下午', '13:30-15:00', '演练一', '真实情境工作坊'],
    ['', '15:00-15:15', '休息', '茶歇'],
    ['', '15:15-16:15', 'M3', '止损点设计'],
    ['', '16:15-16:30', '休息', '茶歇'],
    ['', '16:30-17:30', 'M4', '心理韧性训练'],
    ['', '17:30-18:00', '收尾', '第一日回顾、明日预告'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '第二日（上午）', 3)
headers = ['时段', '时间', '模块', '内容']
rows = [
    ['上午', '09:00-09:30', '复习', '第一日要点回顾'],
    ['', '09:30-10:30', 'M5', '综合演练（前半）'],
    ['', '10:30-10:45', '休息', '茶歇'],
    ['', '10:45-11:45', '演练二', '真实情境实践'],
    ['午休', '12:15-13:30', '午餐', ''],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '第二日（下午）', 3)
headers = ['时段', '时间', '模块', '内容']
rows = [
    ['下午', '13:30-14:30', '迁移', '日常应用与30天计划'],
    ['', '14:30-14:45', '休息', '茶歇'],
    ['', '14:45-15:45', '讲师示范', '关键环节示范'],
    ['', '15:45-16:30', '综合演练', '全班参与'],
    ['', '16:30-17:00', '收尾', '承诺、评估'],
    ['', '17:00-17:10', 'Q&A', '课程答疑'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# ============ PART 3 ============
add_heading(doc, '第三部分：模块教学指南', 1)

# Opening Module
add_heading(doc, '开场：课程介绍与开场活动', 2)
add_para(doc, '时长：15分钟')
add_para(doc, '核心问题：如何让学员立即感受到课程的实用价值？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 学员能理解课程的整体框架和核心价值')
add_para(doc, '2. 学员能建立对"不确定性焦虑"这一概念的基本认知')
add_para(doc, '3. 学员能带着具体问题进入学习状态')

add_heading(doc, '教学流程', 3)
headers = ['活动', '时长', '流程', '物料']
rows = [
    ['破冰活动', '5分钟', '两难选择游戏→统计选择比例→引出主题', '两难选择卡'],
    ['框架介绍', '5分钟', '课程整体框架讲解→学员期待收集', 'PPT/板书'],
    ['规则建立', '5分钟', '课堂约定→小组分组', '分组名单'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '破冰活动：两难选择', 3)
add_para(doc, '让学员选择并解释原因，引出"不确定性焦虑"的主题：')
add_para(doc, '问题一："你更愿意：知道自己会得到100万，还是有50%概率得到300万？"')
add_para(doc, '问题二："你更愿意：确定得到50元，还是有50%概率得到100元但50%概率什么都没有？"')
add_para(doc, '统计选择比例，让选不同选项的人解释原因。')

add_heading(doc, '过渡到主题', 3)
add_para(doc, '"刚才的选择里，有多少人选的是确定的100万/50元？如果是，你为什么会这样选？今天我们要探讨的，就是这个选择的背后——你对不确定性的态度。"')

# M1
add_heading(doc, 'M1：认知重启——不确定性到底是什么', 2)
add_para(doc, '时长：45分钟')
add_para(doc, '核心问题：为什么"不确定"这件事本身会让我们如此焦虑？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 理解不确定性焦虑的生理和心理机制')
add_para(doc, '2. 识别自己的不确定性容忍度水平')
add_para(doc, '3. 觉察焦虑驱动的三种行为模式')

add_heading(doc, '关键概念', 3)
add_para(doc, '• 自主神经系统反应：当大脑感知到不确定性时，杏仁核激活，触发"战斗-逃跑-冻结"反应')
add_para(doc, '• 不确定性容忍度（IU）：个体对不确定情境的敏感程度，存在个体差异')
add_para(doc, '• 三种焦虑行为模式：过度准备、冲动行动、彻底回避')
add_para(doc, '• 模糊耐受连续谱：从"极度回避"到"健康接纳"的光谱')

add_heading(doc, '学员常见困惑/应对', 3)
headers = ['学员问题', '讲师应对']
rows = [
    ['我觉得我很理性，不是在焦虑', '引导识别"焦虑驱动的准备"vs"真正的准备"——关键看行为是否有效'],
    ['我的焦虑是合理的，不是非理性的', '承认合理性，但探讨"焦虑反应是否有效"——合理的不适感≠有效的行动'],
    ['我知道应该做决定，但就是做不到', '转向讨论"做不到"背后的心理机制——往往是情绪在驱动而非理性'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '互动设计', 3)
headers = ['活动', '时长', '流程', '物料']
rows = [
    ['导入活动：两难选择', '10分钟', '破冰游戏→引出IU概念', '两难选择卡'],
    ['概念讲解', '15分钟', 'IU概念讲解→焦虑生理机制', 'PPT/板书'],
    ['小组讨论', '10分钟', '识别三种焦虑模式', '白板纸'],
    ['表单练习', '10分钟', 'IU自评', 'IU量表'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '讲师话术', 3)
add_para(doc, '开场："今天我们来解决一个问题，这个问题可能正在消耗你大量的心理能量——不是工作的压力，不是人际的冲突，而是"不知道结果会怎样"这件事本身带来的焦虑。"')
add_para(doc, '过渡："很多人以为"焦虑"是因为事情太难，或者结果可能会很糟糕。但今天我们要探讨的是：有时候，让你焦虑的不是事情本身，而是"不确定"这件事。"')
add_para(doc, '结尾："知道了焦虑的来源，接下来我们来看看——你的焦虑有什么独特的"签名"？"')

# M2
add_heading(doc, 'M2：识别焦虑——你的不确定性模式', 2)
add_para(doc, '时长：60分钟')
add_para(doc, '核心问题：我的焦虑有什么独特的"签名"？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 识别焦虑循环的四个阶段')
add_para(doc, '2. 找到自己的不确定性"签名"')
add_para(doc, '3. 初步掌握认知解离技术')

add_heading(doc, '关键概念', 3)
add_para(doc, '• 焦虑循环（触发→放大→回避→强化）：不确定性触发焦虑→焦虑放大不确定性感知→回避行为→回避强化焦虑')
add_para(doc, '• 四种典型不确定性情境：职业决策、财务决策、人际决策、健康决策')
add_para(doc, '• 不确定性签名：每个人在面对不确定性时的独特反应模式，包括身体反应、思维模式、行为倾向')
add_para(doc, '• 认知解离：将"我正在想..."与"我想的内容"分离的技术')

add_heading(doc, '学员常见困惑/应对', 3)
headers = ['学员问题', '讲师应对']
rows = [
    ['我能识别出焦虑，但就是控制不住', '强调"识别"本身就是改变的开始，不需要"控制"——觉察是第一步'],
    ['我每次都是同一种模式，太绝望了', '指出"模式可识别"意味着"模式可改变"——能识别就能干预'],
    ['认知解离听起来太玄了', '用具体练习演示，不做理论解释——做三次呼吸，体会"我在观察我的想法"'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '讲师话术', 3)
add_para(doc, '"你们有没有注意到，你的焦虑有一个独特的"签名"？就像指纹一样，每个人在面对不确定性时的反应模式是不同的。有的人主要是胃部反应，有的人是思维反刍，有的人是拖延不动。识别你的签名，是改变它的第一步。"')

add_heading(doc, '过渡话术', 3)
add_para(doc, '"识别焦虑是第一步，接下来要解决的是——当焦虑来了，我该如何有效行动？接下来我们学习一个强大的工具：止损点设计。"')

# Workshop 1
add_heading(doc, '演练一：真实情境工作坊', 2)
add_para(doc, '时长：90分钟')
add_para(doc, '核心问题：如何将识别技术应用于真实场景？')

add_heading(doc, '演练目标', 3)
add_para(doc, '1. 学员能识别自己当前面临的不确定性情境')
add_para(doc, '2. 学员能绘制自己的焦虑循环图')
add_para(doc, '3. 学员能初步运用认知解离技术')

add_heading(doc, '讲师重点', 3)
add_para(doc, '1. 鼓励学员选择真实的当前情境——假设情境的效果远不如真实情境')
add_para(doc, '2. 循环绘制要具体——不是泛泛的"我焦虑"，而是具体的情境-反应链条')
add_para(doc, '3. 关注在分享中可能情绪触发的学员——必要时课后单独跟进')

# M3
add_heading(doc, 'M3：拆解行动——止损点设计', 2)
add_para(doc, '时长：60分钟')
add_para(doc, '核心问题：如何在模糊局面中设定行动的边界？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 理解止损点的心理学原理')
add_para(doc, '2. 掌握三种止损点设计方法')
add_para(doc, '3. 为自己的不确定性情境设计止损点')

add_heading(doc, '关键概念', 3)
add_para(doc, '• 分析瘫痪：过度收集信息导致无法做出决策的状态')
add_para(doc, '• 时间止损点：设定决策的最后期限')
add_para(doc, '• 信息止损点：设定收集信息的停止点')
add_para(doc, '• 行动止损点：设定开始行动的触发点')
add_para(doc, '• 70%法则：当信息收集到70%时，往往已经足够做出有效决策')

add_heading(doc, '学员常见困惑/应对', 3)
headers = ['学员问题', '讲师应对']
rows = [
    ['设定截止日期不是会让我更焦虑吗？', '澄清：没有截止日期的焦虑是"无限焦虑"，有截止日期的焦虑是"有限焦虑"'],
    ['我怎么知道信息收集够了？', '介绍70%法则和"新信息能改变决定吗"测试'],
    ['我设了止损点但还是做不到', '探讨执行障碍，可能是目标设定不合理——需要足够小、足够具体'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '讲师话术', 3)
add_para(doc, '"止损点这个概念，来自投资领域。投资中，你会在亏损达到某个点时止损出场，否则会亏得更多。在人生决策中，这个概念同样适用——你需要在"损失可控"的时候做出决定，否则不确定性会无限消耗你。"')

add_heading(doc, '关键点评要点', 3)
add_para(doc, '1. 止损点不是限制，而是保护——它让你的决策能量聚焦在有效的行动上')
add_para(doc, '2. 三种止损点需要组合使用——时间+信息+行动的联动')
add_para(doc, '3. 止损点需要足够具体才能执行——"尽快"不是止损点，"本周五前"才是')

add_heading(doc, '过渡话术', 3)
add_para(doc, '"有了止损点，我们有了行动的边界。但有时候，即使知道该做，焦虑还是让我们动不了。接下来，我们需要建立内在的稳定性——这就是心理韧性训练。"')

# M4
add_heading(doc, 'M4：与模糊共处——心理韧性训练', 2)
add_para(doc, '时长：60分钟')
add_para(doc, '核心问题：如何在焦虑存在的情况下依然能有效行动？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 理解正念与不确定性觉察的关系')
add_para(doc, '2. 掌握ACT三步法')
add_para(doc, '3. 建立日常模糊耐受练习习惯')

add_heading(doc, '关键概念', 3)
add_para(doc, '• 正念：有意识地、不带评判地关注当下')
add_para(doc, '• 3分钟呼吸空间：ACT中的基础正念练习')
add_para(doc, '• ACT（接纳承诺疗法）：Acceptance and Commitment Therapy')
add_para(doc, '• "允许-观察-选择"技术：允许焦虑存在→观察焦虑的变化→选择有效行动')
add_para(doc, '• 模糊耐受练习：每天5分钟，逐步提高对不确定性的耐受度')

add_heading(doc, '学员常见困惑/应对', 3)
headers = ['学员问题', '讲师应对']
rows = [
    ['正念和"躺平"有什么区别？', '澄清：正念不是不行动，而是在觉察中行动——觉察让你更清楚什么该做'],
    ['接纳焦虑不等于放任不管吗？', '区分"接纳情绪"和"接纳不行动"——接纳是不对抗，选择权在你'],
    ['练习太简单了，我觉得没用', '强调：简单的事情重复做，才是基础功——高深的技术源于扎实的基础练习'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '讲师话术', 3)
add_para(doc, '"这一章可能是整个课程最"软"的部分，但它可能是最重要的。我们一直在讲技术、讲方法，但技术要有效，你需要有一个稳定的内在基础。正念不是让你的焦虑消失，而是让你在焦虑存在的情况下，依然能看清自己在做什么。"')

add_heading(doc, '过渡话术', 3)
add_para(doc, '"我们学了很多工具和技术，接下来要把它们整合起来，在真实情境中走一遍完整的流程。"')

# M5
add_heading(doc, 'M5：综合演练与日常应用', 2)
add_para(doc, '时长：120分钟')
add_para(doc, '核心问题：如何将所有工具整合应用于真实情境？')

add_heading(doc, '模块目标', 3)
add_para(doc, '1. 在真实情境中应用止损点设计')
add_para(doc, '2. 建立从课堂到日常的迁移路径')
add_para(doc, '3. 制定30天实践计划')

add_heading(doc, '关键概念', 3)
add_para(doc, '• 全流程五步法：识别情境→绘制循环→设计止损点→练习正念→承诺行动')
add_para(doc, '• 迁移三阶段：理解→应用→习惯')
add_para(doc, '• 30天实践计划：将课堂学习转化为日常习惯')
add_para(doc, '• 行动承诺：具体、可衡量、有时限的行动声明')

add_heading(doc, '学员常见困惑/应对', 3)
headers = ['学员问题', '讲师应对']
rows = [
    ['课程结束后我肯定会忘', '建立30天实践计划，让改变可持续——写下来才不会忘'],
    ['我一个人很难坚持', '建立支持系统（告知他人、使用工具、加入社群）——公开承诺更有效'],
    ['如果我退步了怎么办', '复盘是正常的，区别"放弃"和"调整"——退步是学习的一部分'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '讲师话术', 3)
add_para(doc, '"课程结束才是真正的开始。我们在这两天里学了很多技术，但技术要变成习惯，需要持续的练习。这不是"学完就好了"的课程，而是"学完才是开始"的课程。"')

# ============ PART 4 ============
add_heading(doc, '第四部分：关键环节讲师指南', 1)

add_heading(doc, '开场设计', 2)
add_heading(doc, '破冰活动：两难选择', 3)
add_para(doc, '让学员选择并解释原因，引出"不确定性焦虑"的主题：')
add_para(doc, '问题一："你更愿意：知道自己会得到100万，还是有50%概率得到300万？"')
add_para(doc, '问题二："你更愿意：确定得到50元，还是有50%概率得到100元但50%概率什么都没有？"')
add_para(doc, '引导步骤：')
add_para(doc, '1. 统计选择比例（举手或站立）')
add_para(doc, '2. 让选择不同选项的人解释原因')
add_para(doc, '3. 引导发现：选择"确定"选项的人往往不是因为理性计算，而是因为对不确定性的不适')

add_heading(doc, '过渡设计', 2)
headers = ['过渡节点', '过渡话术']
rows = [
    ['M1→M2', '"知道了焦虑的来源，接下来我们来看看——你的焦虑有什么独特的"签名"？"'],
    ['M2→演练一', '"理论说完了，该动手了。接下来我们用真实情境来练习识别自己的焦虑模式。"'],
    ['演练一→M3', '"识别了焦虑循环，接下来要解决的是——当焦虑来了，我该如何有效行动？接下来我们学习一个强大的工具：止损点设计。"'],
    ['M3→M4', '"有了止损点，我们有了行动的边界。但有时候，即使知道该做，焦虑还是让我们动不了。接下来，我们需要建立内在的稳定性——这就是心理韧性训练。"'],
    ['M4→M5', '"我们学了很多工具和技术，接下来要把它们整合起来，在真实情境中走一遍完整的流程。"'],
    ['M5→收尾', '"课程结束才是真正的开始。最后，让我们来做一个承诺——你带走的具体行动是什么？"'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '结尾设计', 2)
add_heading(doc, '收尾仪式', 3)
add_para(doc, '第一步：带走一句话（邀请学员分享）')
add_para(doc, '"这堂课你带走的一句话是什么？不需要长，一句话就好。"')
add_para(doc, '记录方式：可以请几位自愿分享的学员口头说，也可以让大家写在便签上贴到白板上。')
add_para(doc, '')
add_para(doc, '第二步：集体承诺（每人说出一个具体的行动承诺）')
add_para(doc, '"现在，让我们每个人说出一个具体的行动承诺。不是"我要改变"，而是"我回到公司/家里的第一件事是做什么"。"')
add_para(doc, '承诺要求：')
add_para(doc, '• 具体（不是"少焦虑"而是"每天花5分钟做呼吸练习"）')
add_para(doc, '• 可衡量（有时间节点）')
add_para(doc, '• 有人见证（小组见证）')
add_para(doc, '')
add_para(doc, '第三步：后续支持（告知后续支持渠道）')
add_para(doc, '"课程结束后，如果你需要支持，可以......（告知支持渠道）"')

# ============ PART 5 ============
add_heading(doc, '第五部分：常见挑战与应对', 1)
headers = ['挑战', '应对策略']
rows = [
    ['学员说"我知道这些，但没有用"', '引导识别"知道"和"做到"的差距——"你今天知道了你的焦虑签名，但你在刚才的练习中有没有真的识别出来？"'],
    ['学员分享过于深入，情绪触发', '及时收拢——"感谢你的分享，这个话题我们线下继续交流"；必要时提供一对一咨询转介'],
    ['学员质疑方法有效性', '用数据说话（IU量表前后测对比）——鼓励尝试，"有效没效，试了才知道"'],
    ['时间不够用', '准备好可跳过的补充内容——根据实际情况调整，核心流程不能省'],
    ['学员参与度低', '增加小组讨论——减少单人发言，增加互动形式'],
    ['学员说"我就是这种人，改不了"', '指出"模式是习惯，不是性格"——习惯是可以改变的，只是需要时间和练习'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# ============ PART 6 ============
add_heading(doc, '第六部分：评估工具', 1)

add_heading(doc, '课堂评估', 2)
headers = ['评估维度', '评估方式']
rows = [
    ['参与度观察', '讲师在互动环节观察学员的投入程度'],
    ['提问质量', '记录学员在Q&A环节提出的问题类型'],
    ['表单完成情况', '检查各练习表单的完成度'],
    ['小组互动质量', '巡回各组时观察小组讨论质量'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '课后评估', 2)
headers = ['评估方式', '时机', '用途']
rows = [
    ['课程满意度问卷', '课程结束时', '了解学员主观体验'],
    ['知识掌握测试', '课程结束时', '评估知识内化程度'],
    ['行为改变跟踪', '30天后', '评估长期效果'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '知识掌握测试（样题）', 2)
add_para(doc, '单选题')
add_para(doc, '1. 不确定性焦虑的核心特征是什么？')
add_para(doc, '   A. 对危险过度敏感')
add_para(doc, '   B. 对未知情境的强烈心理反应')
add_para(doc, '   C. 拖延行为')
add_para(doc, '   答案：B')
add_para(doc, '')
add_para(doc, '2. 焦虑循环的四个阶段不包括：')
add_para(doc, '   A. 触发')
add_para(doc, '   B. 放大')
add_para(doc, '   C. 压抑')
add_para(doc, '   答案：C')
add_para(doc, '')
add_para(doc, '3. 止损点设计的核心目的是：')
add_para(doc, '   A. 限制决策')
add_para(doc, '   B. 让焦虑变成有限焦虑')
add_para(doc, '   C. 消除不确定性')
add_para(doc, '   答案：B')
add_para(doc, '')
add_para(doc, '简答题')
add_para(doc, '1. 请描述你的"不确定性签名"中的一个关键特征，以及你计划如何干预它。')
add_para(doc, '   评估要点：是否能识别具体的签名特征；是否有可行的干预计划')
add_para(doc, '')
add_para(doc, '2. 你计划如何在接下来的30天里练习"与模糊共处"？')
add_para(doc, '   评估要点：是否有具体的练习计划；是否可执行')

# ============ PART 7 ============
add_heading(doc, '第七部分：讲师资源', 1)

add_heading(doc, '附件目录', 2)
headers = ['附件编号', '名称', '说明']
rows = [
    ['附件A', '两难选择卡（打印版）', '用于开场破冰活动'],
    ['附件B', 'IU量表完整版', '用于不确定性容忍度评估'],
    ['附件C', '焦虑循环挂图', '大尺寸打印用于课堂展示'],
    ['附件D', '止损点设计模板', '用于工作坊练习'],
    ['附件E', '认知解离练习音频', '用于正念练习引导'],
    ['附件F', '30天实践指南', '用于课后学员自主练习'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '参考书单', 2)
headers = ['书名', '作者', '推荐理由']
rows = [
    ['《ACT made simple》', 'Russ Harris', 'ACT入门经典，通俗易懂'],
    ['《正念：此刻是一枝花》', 'Jon Kabat-Zinn', '正念经典，适合初学者'],
    ['《思考，快与慢》', 'Daniel Kahneman', '决策心理学经典'],
    ['《噪声》', 'Daniel Kahneman', '决策中的"噪声"问题'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# ============ APPENDIX ============
add_heading(doc, '附录', 1)

add_heading(doc, '附录A：讲师自我检查清单', 2)

add_heading(doc, '课程前一天检查清单', 3)
add_para(doc, '• [ ] 再次确认场地布置（U型桌或分组桌）')
add_para(doc, '• [ ] 再次确认设备正常运行（投影、音响）')
add_para(doc, '• [ ] 打印物料齐全（学员手册、工具表单）')
add_para(doc, '• [ ] 案例素材备份（电脑+U盘）')
add_para(doc, '• [ ] 准备讲师自己的案例故事')
add_para(doc, '• [ ] 休息充足，避免熬夜')

add_heading(doc, '课程当天早上检查清单', 3)
add_para(doc, '• [ ] 提前30分钟到场')
add_para(doc, '• [ ] 检查设备正常运行')
add_para(doc, '• [ ] 摆好桌椅分组，确认每组有马克笔和白板纸')
add_para(doc, '• [ ] 将学员手册和工具表单按组摆放')
add_para(doc, '• [ ] 准备好水和小点心')
add_para(doc, '• [ ] 调整好状态，保持热情')

add_heading(doc, '课程进行中检查清单', 3)
add_para(doc, '• [ ] 关注学员状态，适时调整节奏')
add_para(doc, '• [ ] 记录学员精彩观点（用于后续引用）')
add_para(doc, '• [ ] 观察需要个别关注的学员')
add_para(doc, '• [ ] 记录需要改进的地方（课后复盘）')
add_para(doc, '• [ ] 确保每个活动环节时间可控')
add_para(doc, '• [ ] 留出Q&A时间')

add_heading(doc, '课程结束后检查清单', 3)
add_para(doc, '• [ ] 收集学员反馈（纸质或电子）')
add_para(doc, '• [ ] 与助教复盘当天表现')
add_para(doc, '• [ ] 发感谢邮件/消息给学员')
add_para(doc, '• [ ] 整理教室，归还设备')
add_para(doc, '• [ ] 整理学员产出（承诺书、方案等）')
add_para(doc, '• [ ] 记录需要跟进的学员')

add_heading(doc, '附录B：课堂能量管理', 2)
add_heading(doc, '课堂能量曲线', 3)
add_para(doc, '一天的培训，学员的能量遵循一条可预测的曲线。了解这条曲线，可以帮助讲师主动管理学员状态。')
headers = ['时间段', '典型状态与应对策略']
rows = [
    ['开场前10分钟', '学员处于社交模式，轻松聊天。讲师可以主动与学员交流，提前了解他们的期待。'],
    ['开场—上午第一个小时', '新鲜感最强，注意力最集中。放入最重要的认知颠覆内容（M1+M2核心）。'],
    ['午前最后30分钟', '能量开始下滑。引入互动练习，让学员动手，避免连续讲授。'],
    ['午饭后30分钟', '生理性低谷，注意力最难维持。轻松的小组讨论或案例分享，避免工具密集讲解。'],
    ['下午第一个小时', '重新进入节奏。放入第二批核心工具和实操练习（M3+M4）。'],
    ['全天最后60分钟', '学员期待"今天有没有真正带走东西"。留给综合练习+个人计划+全课回顾（M5+收尾）。'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

add_heading(doc, '附录C：讲师注意事项汇总', 2)
headers = ['模块', '注意事项']
rows = [
    ['开场', '不要跳过破冰活动——这是建立课堂氛围的关键'],
    ['M1', '用通俗语言讲神经科学，避免术语；用投资类比讲解止损点'],
    ['M2', '强调"识别"本身就是改变；签名不是性格是习惯'],
    ['演练一', '鼓励选择真实情境；关注情绪触发的学员'],
    ['M3', '止损点是保护不是限制；三种止损点要组合使用'],
    ['M4', '正念练习时强调"做不到也没关系"；简单的事情重复做'],
    ['M5', '给足演练时间；30天计划要足够小才能坚持'],
    ['收尾', '创造仪式感；承诺要具体、公开、有时限'],
]
add_table(doc, headers, rows)
doc.add_paragraph()

# Footer
doc.add_paragraph()
add_para(doc, '─' * 50)
add_para(doc, '版权所有 · 罗宏伟 · 本手册仅供本课程讲师使用')
add_para(doc, '版本：v1.0 | 最后更新：2026-08-20')

# Save document
output_path = 'D:/新课开发/心理学/06-不确定性焦虑：在模糊局面中做出决策/讲师手册/06-不确定性焦虑_讲师手册.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')

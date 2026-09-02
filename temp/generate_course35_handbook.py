# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY_BLUE = RGBColor(0x25, 0x63, 0xEB)
GROWTH_GREEN = RGBColor(0x10, 0xB9, 0x81)
WARNING_ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
DARK_TEXT = RGBColor(0x1F, 0x38, 0x64)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_knowledge_point(doc, title, content, example=None, highlight=None):
    para = doc.add_paragraph()
    run = para.add_run("◆ " + title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DARK_TEXT
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)

    para = doc.add_paragraph()
    run = para.add_run(content)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(6)

    if example:
        para = doc.add_paragraph()
        run = para.add_run("案例：")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WARNING_ORANGE
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run2 = para.add_run(example)
        run2.font.size = Pt(10)
        run2.font.italic = True
        run2.font.name = "微软雅黑"
        run2._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        para.paragraph_format.space_after = Pt(8)

    if highlight:
        para = doc.add_paragraph()
        run = para.add_run(highlight)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = GROWTH_GREEN
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(12)

def add_exercise(doc, exercises):
    para = doc.add_paragraph()
    run = para.add_run("练习题")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = WARNING_ORANGE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)
    for i, exercise in enumerate(exercises, 1):
        para = doc.add_paragraph()
        run = para.add_run(str(i) + ". " + exercise)
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        para.paragraph_format.space_after = Pt(4)

def add_form_header(doc, form_num, form_title, instructions):
    para = doc.add_paragraph()
    run = para.add_run("【表单" + form_num + "】" + form_title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(4)
    para = doc.add_paragraph()
    run = para.add_run(instructions)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(8)

def create_action_list_table(doc, items):
    table = doc.add_table(rows=len(items)+1, cols=3)
    table.style = "Table Grid"
    headers = ["序号", "行动项目", "执行记录"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_idx, item in enumerate(items, 1):
        row = table.rows[row_idx]
        row.cells[0].text = str(row_idx)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = item
        row.cells[2].text = ""
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table

def create_survey_table(doc, questions):
    table = doc.add_table(rows=len(questions)+1, cols=5)
    table.style = "Table Grid"
    headers = ["行为特征", "几乎从不", "偶尔如此", "经常这样", "基本如此"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row_idx, question in enumerate(questions):
        row = table.rows[row_idx + 1]
        row.cells[0].text = question
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j in range(1, 5):
            row.cells[j].text = "O"
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(10)
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Cover page
    para = doc.add_paragraph()
    run = para.add_run("课程学员手册")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run("《心理安全感与信任文化》")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)

    para = doc.add_paragraph()
    run = para.add_run("——谷歌氧气计划的启示")
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_TEXT
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(32)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("课程ID", "35"),
        ("课程名称", "心理安全感与信任文化——谷歌氧气计划的启示"),
        ("课程讲师", "________________"),
        ("课程日期", "________________"),
    ]
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_shading(row.cells[0], "E8F4FD")
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Manual usage instructions
    para = doc.add_paragraph()
    run = para.add_run("手册使用说明")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(12)

    usage_intro = "这本手册是你在课程中完成的工作台。每一个框架、每一张表单、每一道练习，都要在课堂上完成。你做完的每一页，都是你带走的真实成果，而不是别人告诉你的理论。"
    para = doc.add_paragraph()
    run = para.add_run(usage_intro)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(12)

    principles = [
        ("原则一：带着真实任务来", "手册里所有的练习都要用你自己工作中的真实场景。坐在课堂里想象一个场景，不如直接面对真实的挑战。"),
        ("原则二：写下来比记下来更有价值", "看懂了不等于会用，写下来才是真正内化的开始。空白处就是你的思考痕迹。"),
        ("原则三：这是起点，不是终点", "课程结束不是学习的终点，手册最后一章是你的30天行动计划。真正的改变发生在课程之外。"),
    ]
    for title, content in principles:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = GROWTH_GREEN
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)
        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.size = Pt(11)
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        para.paragraph_format.space_after = Pt(12)

    # Course overview
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("课程全景图")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(12)

    overview_text = '本课程分为四个模块，层层递进，帮助你从"知道"到"做到":\n\n模块一：认知篇 —— 理解心理安全感是什么，为什么它决定团队绩效\n\n模块二：诊断篇 —— 评估你的团队心理安全感现状，发现改进空间\n\n模块三：建设篇 —— 学习具体的管理行为，建立心理安全感的日常实践\n\n模块四：实践篇 —— 制定30天行动计划，从知道到做到\n\n核心洞见：心理安全感不是"你好我好大家好"，而是一种允许直言不讳的信任氛围。研究表明，谷歌最佳团队的秘诀不是明星员工，而是每个人都敢于说真话。'

    para = doc.add_paragraph()
    run = para.add_run(overview_text)
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(12)

    # Scene card
    para = doc.add_paragraph()
    run = para.add_run("场景卡：你的真实任务")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    scenarios = [
        ("任务背景", "在课程开始前，想一件你最近需要解决的团队挑战或管理难题。\n这个问题应该是：\n- 与团队成员互动相关\n- 让你感到有些棘手或不确定如何处理\n- 值得在课程中深入探索"),
        ("具体情境", "描述这个挑战发生的具体情境：\n- 什么时候？在哪里？\n- 涉及哪些人？\n- 发生了什么？"),
        ("你的角色", "在这个情境中，你扮演什么角色？\n- 领导者？协调者？参与者？\n- 你对结果有多大的影响力？"),
        ("期望成果", "课程结束时，你希望：\n- 对这个问题有什么新的认识？\n- 带走哪些可操作的改进方案？\n- 计划如何开始行动？"),
    ]
    for i, (label, content) in enumerate(scenarios):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(row.cells[0], "F2F2F2")
        row.cells[1].text = content
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    row = table.rows[4]
    row.cells[0].text = "我的场景"
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_shading(row.cells[0], "E8F4FD")
    row.cells[1].text = "（在此写下你的真实任务，这个任务将贯穿整个课程）"
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.italic = True

    # Module 1
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("模块一：认知篇——理解心理安全感")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("知识点1：什么是心理安全感")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "心理安全感的定义",
        "心理安全感（Psychological Safety）是指团队成员在面对人际风险时的一种共同信念——即相信自己是安全的，不会因为说出真实想法、承认错误或提出不同意见而受到惩罚或羞辱。\n\n这个概念由哈佛商学院教授Amy Edmondson在1999年首次提出，经过20多年的研究，已成为组织行为学最具影响力的概念之一。",
        "想象一个团队会议场景：一位年轻员工鼓起勇气指出方案中的一个潜在风险。如果团队心理安全感高，其他人会感谢他的坦诚；如果心理安全感低，他可能会担心被认为找麻烦而选择沉默。",
        "心理安全感不是让大家舒服，而是让大家敢于不舒服地说真话。"
    )

    add_knowledge_point(doc,
        "心理安全感的两个维度",
        "1. 个人层面：个体在团队中感到被尊重、被接纳相信自己不会因真实表现而受到负面评价。\n\n2. 团队层面：整个团队形成一种氛围，支持直言、容忍犯错、把不知道当作学习机会。",
        highlight="两个维度缺一不可：个人感到安全，但如果团队氛围不鼓励直言，依然难以建立真正的心理安全感。"
    )

    para = doc.add_paragraph()
    run = para.add_run("知识点2：为什么心理安全感决定团队绩效")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "谷歌氧气计划的重要发现",
        "2008年，谷歌启动了一个内部研究项目——氧气计划（Project Aristotle），旨在找出高效团队的关键特征。研究团队分析了180多个团队的200多个变量，包括绩效数据、员工调查、面对面访谈等。\n\n研究结论出人意料：最优秀的团队不是拥有最聪明的成员，不是最有经验的领导，而是心理安全感最高的团队。\n\n在心理安全感高的团队中：\n- 员工更愿意承认错误并从中学习\n- 团队更能应对挑战和压力\n- 成员之间更能相互协作\n- 创新想法更容易被提出和实施",
        "谷歌的研究发现，当团队成员在会议中发言时，如果发言者被打断或观点被忽视，高绩效团队与低绩效团队的反应截然不同。高绩效团队会主动邀请被打断者继续说完，而低绩效团队往往默认了这种行为。",
        "优秀团队的秘诀：不是没有冲突，而是有能力建设性地处理冲突。"
    )

    add_knowledge_point(doc,
        "心理安全感与团队绩效的关系",
        "研究表明，心理安全感通过以下机制影响团队绩效：\n\n1. 学习行为：高心理安全感让团队成员更愿意提问、试错、反馈，而这些学习行为直接提升团队能力。\n\n2. 知识共享：在安全的环境中，成员更愿意分享隐性知识和经验，避免重复犯错。\n\n3. 创新涌现：最有创新性的想法往往来自异端观点，心理安全感让这些想法能够浮出水面。\n\n4. 危机应对：面对突发问题，心理安全感高的团队能快速坦诚沟通，而不是相互指责或掩盖。",
        highlight="你可能在技术上很强，但如果团队心理安全感不足，你的团队永远无法发挥全部潜力。"
    )

    para = doc.add_paragraph()
    run = para.add_run("知识点3：心理安全感的四大支柱")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "支柱一：表达思想的自由",
        "团队成员相信他们可以提出与主流不同的观点，而不会被视为麻烦制造者。\n\n这不意味着每个观点都会被采纳，而是意味着每个声音都会被认真倾听。\n\n检验标准：团队会议上，当一个非主流观点被提出时，其他人的第一反应是什么？",
        "在3M公司有一条著名的政策：员工可以用15%的时间做任何自己想做的事情。正是这种心理安全感，让Post-it便签这样的创新产品得以诞生。",
        "允许异端存在，是创新的前提条件。"
    )

    add_knowledge_point(doc,
        "支柱二：承认无知的坦然",
        "团队成员不怕说我不知道，不怕向他人请教。\n\n在心理安全感低的组织里，不知道被视为弱点，人们倾向于假装懂得比实际多。在心理安全感高的组织里，不知道被当作学习的起点。\n\n检验标准：你的团队中，有多少次我不知道被诚实地表达出来？",
        "一位医院管理者发现，医护人员害怕报告接近失误（差点出错但没出错）——因为担心被追责。这导致医院错失了从差点出事中学习的机会。后来通过建立心理安全感，医院鼓励报告所有接近失误，反而显著减少了实际医疗事故。",
        "越是掩盖不知道，越会暴露更多的不知道。"
    )

    add_knowledge_point(doc,
        "支柱三：犯错的勇气",
        "团队成员相信错误会被当作学习机会，而不是被用来追责。\n\n这不意味着对错误漠不关心，而是意味着面对错误时的第一反应是我们能从中学到什么，而不是是谁的错。\n\n检验标准：当一个重要错误发生时，团队的第一个会议是在追责，还是在分析原因？",
        highlight="态度决定行为：如果你想知道一个组织的真实文化，就看他们如何对待错误。"
    )

    add_knowledge_point(doc,
        "支柱四：求助的常态",
        "团队成员不把求助视为软弱，而视为智慧和团队协作的体现。\n\n在心理安全感高的团队中，求助被理解为对团队的信任，也是建立关系的机会。\n\n检验标准：当你遇到问题时，你有多确定团队会支持你而不是评判你？",
        highlight="一个人走可能走得很快，但一群人走才能走得很远——而这种一起走的前提是相互支持的文化。"
    )

    # Form 1.1
    add_form_header(doc, "1.1", "心理安全感自评问卷",
        "目的：帮助你认识自己在心理安全感方面的现状\n要求：请根据过去一个月的情况，选择最符合你实际感受的选项\n时间：5分钟完成")

    questions_1_1 = [
        "在团队会议上，我会主动提出与主流意见不同的观点",
        "当我犯错误时，我相信团队会把它当作学习机会",
        "我敢于在团队面前承认我不知道或我需要帮助",
        "我相信我的意见会被认真倾听，即使与他人不同",
        "当我对某件事有疑虑时，我会直接表达，而不会保持沉默",
        "我感觉团队成员之间可以坦诚地相互反馈",
        "我不担心因为提出不同意见而被团队排斥或打击",
        "当团队遇到问题时，成员们能够开放地讨论而不是相互指责"
    ]
    create_survey_table(doc, questions_1_1)

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("得分说明：基本如此=4分，经常这样=3分，偶尔如此=2分，几乎从不=1分。总分越高，心理安全感体验越强。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Form 1.2
    add_form_header(doc, "1.2", "我的团队现状分析",
        "目的：分析你所在团队的心理安全感现状\n要求：结合自评结果和你的观察，填写以下分析\n时间：10分钟完成")

    analysis_items = [
        ("我观察到的团队优势", ""),
        ("我观察到的团队不足", ""),
        ("最让我担忧的一个现象", ""),
        ("如果我想提升团队心理安全感，第一件要做的事是", "")
    ]
    table = doc.add_table(rows=len(analysis_items), cols=2)
    table.style = "Table Grid"
    for i, (label, content) in enumerate(analysis_items):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_shading(table.rows[i].cells[0], "F2F2F2")
        table.rows[i].cells[1].text = content
        table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    add_exercise(doc, [
        "用自己的语言向一位同事解释心理安全感的概念，看看你是否能说得清楚",
        "回想一个你在团队中选择沉默的时刻，当时是什么让你没有说出真实想法？",
        "你认为在当前团队中，建立心理安全感最大的障碍是什么？"
    ])

    # Module 2
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("模块二：诊断篇——评估你的团队")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("知识点1：团队心理安全感诊断工具")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "诊断的重要性",
        "管理者的一个常见误区是我觉得团队有心理安全感——但感觉往往不准确。\n\n研究显示，管理者往往高估自己团队的心理安全感程度。这不是管理者在说谎，而是他们与团队成员的日常互动方式本身就在传递信号，而这些信号往往与成员的真实感受不同步。\n\n因此，系统性的诊断比直觉判断更可靠。",
        highlight="你认为是开放沟通的行为，在团队成员眼中可能恰恰相反。"
    )

    add_knowledge_point(doc,
        "七题版评估问卷说明",
        "以下是经过验证的团队心理安全感评估问卷（基于Edmondson的研究）。建议：\n\n1. 匿名填写：确保答案真实性\n2. 团队共识：让多位团队成员分别填写\n3. 持续追踪：每季度评估一次，观察趋势变化\n\n评分方式：\n- 几乎从不 = 1分\n- 偶尔如此 = 2分\n- 有时如此 = 3分\n- 经常这样 = 4分\n- 几乎总是 = 5分\n\n总分说明：\n- 35分以上：团队心理安全感很高\n- 25-34分：中等水平，有改进空间\n- 25分以下：需要立即关注和改进",
        "一家科技公司的团队在首次诊断时平均得分只有22分。经过三个月的干预建设后，团队得分提升到31分，同期项目成功率提升了40%。",
        "没有测量就没有管理：诊断是改进的第一步。"
    )

    # Form 2.1
    add_form_header(doc, "2.1", "团队心理安全感评估问卷（7题版）",
        "目的：评估团队整体的心理安全感水平\n要求：请根据过去一个月团队的普遍情况，选择最符合的选项\n时间：5分钟完成")

    table = doc.add_table(rows=8, cols=7)
    table.style = "Table Grid"
    headers = ["行为特征", "几乎从不(1分)", "偶尔(2分)", "有时(3分)", "经常(4分)", "几乎总是(5分)", "得分"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    questions_2_1 = [
        "在团队中，大家可以自由地表达自己的想法和意见",
        "团队成员敢于承认自己的错误",
        "团队欢迎并鼓励不同的观点和意见",
        "当团队遇到问题时，大家能够开放地讨论问题本身而不是相互指责",
        "团队成员相信领导者会支持他们的工作决策",
        "团队成员不怕向同事或领导请教问题",
        "在团队会议中，每个人的声音都能被听到"
    ]
    for row_idx, question in enumerate(questions_2_1):
        row = table.rows[row_idx + 1]
        row.cells[0].text = question
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        for j in range(1, 7):
            row.cells[j].text = "O"
            row.cells[j].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("知识点2：管理者行为自检")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "管理者是心理安全感的关键",
        "研究表明，管理者对团队心理安全感的建立起着决定性作用。你的行为每天都在向团队传递信号——什么是被鼓励的，什么是不被接受的。\n\n以下是管理者需要自检的关键行为维度：",
        highlight="作为管理者，你的一举一动都在塑造团队文化。你不是建立心理安全感的推动者，你就是心理安全感本身。"
    )

    add_knowledge_point(doc,
        "维度一：示范脆弱性",
        "你自己是否愿意在团队面前：\n- 承认自己不知道某些事情？\n- 承认自己犯了错误？\n- 寻求他人的帮助或建议？\n\n如果你从不示弱，团队会认为示弱是不被接受的。",
        "一位CTO在季度复盘会上主动承认自己在一个技术决策上犯了错，并详细分析了原因。团队成员反馈说，他们第一次感到承认错误是被允许的。之后一个季度内，团队主动报告的从错误中学习的案例数量增加了3倍。",
        "示弱不是软弱，示弱是勇气的表现。"
    )

    add_knowledge_point(doc,
        "维度二：包容异议",
        "当团队成员提出与你不同的意见时，你的反应是什么？\n\n- 立刻反驳或解释为什么不？\n- 还是先说谢谢你提出这个观点？\n\n即便你最终需要做出不同的决定，如何处理异议本身就传递着强烈的信号。",
        highlight="你可以不采纳建议，但你不能否定提出建议的权利。"
    )

    add_knowledge_point(doc,
        "维度三：惩罚告警",
        "当团队成员报告问题或潜在风险时，你的反应是什么？\n\n- 这个问题你怎么早没发现？\n- 还是感谢你及时报告，我们一起来看看怎么解决？\n\n前者传递的信号是：报告问题会招来麻烦。后者传递的信号是：报告问题是受欢迎的行为。",
        highlight="你的第一反应，决定了团队下次是否还会开口。"
    )

    # Form 2.2
    add_form_header(doc, "2.2", "管理者行为自检表",
        "目的：帮助管理者反思自身行为对团队心理安全感的影响\n要求：针对每个行为回想最近一次的具体表现，诚实评估\n时间：10分钟完成")

    behavior_items = [
        ("我是否在团队面前承认过自己的错误或不足？", "上次是什么时候？当时团队反应如何？"),
        ("当有人提出不同意见时，我是否先表示感谢再回应？", "最近一次是什么时候？"),
        ("当团队成员报告问题或失误时，我的第一反应是什么？", "我通常会说："),
        ("我是否主动邀请过团队成员提出反对意见或疑虑？", "我是如何邀请的："),
        ("在会议中，我是否注意到每个人都有机会发言？", "我通常如何确保这一点："),
    ]
    table = doc.add_table(rows=len(behavior_items)+1, cols=3)
    table.style = "Table Grid"
    headers = ["自检行为", "反思问题", "我的实际情况"]
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (item, question) in enumerate(behavior_items):
        row = table.rows[i + 1]
        row.cells[0].text = item
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = question
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.italic = True
        row.cells[2].text = ""

    add_exercise(doc, [
        "完成管理者行为自检表，特别关注你忽略或跳过的问题——这些可能就是你最需要改进的地方",
        "找一个信任的同事，请他/她给你最近一次在团队中处理不同意见的具体反馈",
        "如果你要给自己最近一周的管理行为打分（1-10分），你会打几分？你的团队会打几分？"
    ])

    # Module 3
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("模块三：建设篇——管理者能做什么")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("知识点1：日常行为改进清单")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "从小处开始，从现在开始",
        "建立心理安全感不需要轰轰烈烈的变革，而是需要在日常互动中持续践行的微小行为。\n\n以下是管理者每天都可以做到的行动清单：",
        highlight="改变不是一次事件，而是一个习惯。与其等待一个完美时机，不如从今天开始，从下一件事开始。"
    )

    add_form_header(doc, "3.1", "每日行动清单",
        "目的：帮助管理者建立每日践行心理安全感的行为习惯\n要求：每天结束时回顾完成情况，在执行记录栏打勾\n频率：建议连续坚持21天形成习惯")

    daily_items = [
        "在会议开始时，主动邀请还没发言的同事发表意见",
        "当有人提出不同观点时，先说谢谢你提出这个观点",
        "当团队成员报告问题或失误时，先说谢谢你告诉我再讨论解决方案",
        "在团队面前承认一件你自己不知道或不确定的事情",
        "对团队成员的工作表示感谢，具体说明哪一点做得好",
        "如果有决策需要改变，解释原因而不是只宣布结论",
        "询问团队成员你最近有什么困难需要帮助吗？"
    ]
    create_action_list_table(doc, daily_items)

    para = doc.add_paragraph()
    run = para.add_run("知识点2：具体的建设行动")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "行动一：建立安全信号仪式",
        "通过固定的仪式，让团队知道现在是可以说真话的时间。\n\n具体做法：\n- 每次重要会议最后增加反思环节：有人说我们今天讨论中有什么没敢说的吗？\n- 在项目启动时明确在这里，愚蠢的问题不被嘲笑\n- 在团队复盘时，第一个问题永远是我们从中学到了什么，而不是谁犯了错",
        "Spotify的squad模式中，每个sprint结束时都有一个retrospective环节。引导者会问三个问题：1）什么进展得好？2）什么进展得不好？3）我们承诺做什么改进？这种结构化的仪式让团队逐渐建立了心理安全感。",
        "仪式创造安全空间：让说真话变成一种习惯，而不是一次冒险。"
    )

    add_knowledge_point(doc,
        "行动二：示范脆弱性求助",
        "当你自己遇到困难或不确定时，主动向团队或某位成员请教。\n\n这不是示弱，而是在授权——你在告诉团队：不知道是被允许的，求助是被欢迎的。\n\n关键要点：\n- 要真诚，不要假装求助\n- 即便你其实有答案，也可以问如果是你们，会怎么处理这件事？\n- 感谢任何回应，不要评判答案的质量",
        highlight="你的每一次求助，都在向团队宣告：求助是安全的。"
    )

    add_knowledge_point(doc,
        "行动三：建立说真话的正向反馈机制",
        "让说出不同意见变得有吸引力，而不是有风险。\n\n具体做法：\n- 对提出不同意见的人给予具体表扬\n- 当有人因为说真话而受到外部压力时，公开支持他们\n- 记录并感谢那些挽救了团队的直言",
        "一家金融公司的CEO在一次全员大会上公开表扬了一位年轻分析师——这位分析师在公司重大决策会议上提出了与CEO不同的观点，并坚持了自己的判断。最终决策被证明是正确的。CEO说我希望每个人都像他一样勇敢地表达观点。这句话在公司内部被广泛传播，成为心理安全感的标志性故事。",
        "表扬一次直言，胜过奖励十个服从。"
    )

    add_form_header(doc, "3.2", "每周行动清单",
        "目的：帮助管理者建立每周深度的心理安全感建设习惯\n要求：每周至少选择3项完成，记录具体情境和观察到的反应\n时间：每周30分钟")

    weekly_items = [
        "安排一次一对一谈话，主动询问对方最近有什么事是你不敢在团队会议上说的？",
        "在团队会议上分享一个你自己最近犯的错误，以及你从中学到了什么",
        "主动联系一位你较少互动的团队成员，了解他的工作情况和想法",
        "当团队成员提出与你不同的意见时，不要立即回应，先请其他成员也发表看法",
        "在周报或团队通讯中，表扬一位敢于说真话的团队成员（具体说明事件）"
    ]
    create_action_list_table(doc, weekly_items)

    add_form_header(doc, "3.3", "每月行动清单",
        "目的：帮助管理者进行更深入的心理安全感建设\n要求：每月完成至少2项，并在团队中观察变化\n时间：每月1-2小时")

    monthly_items = [
        "组织一次说真话专题讨论会：问团队阻碍我们直言不讳的原因是什么？",
        "匿名调查团队成员对心理安全感现状的反馈（可使用表单2.1）",
        "与团队一起重新审视团队规范，明确什么是我们承诺的沟通方式",
        "与HR或外部教练讨论你在建立心理安全感方面的挑战和困惑",
        "回顾过去一个月，记录3个关键时刻——你是如何处理的？团队反应如何？"
    ]
    create_action_list_table(doc, monthly_items)

    add_exercise(doc, [
        "从每日行动清单中选择3项，明天就在工作中践行",
        "观察你今天与团队成员的互动，记录一次你错过了建立心理安全感机会的时刻",
        "如果你只能选择一项每日行动坚持21天，你会选择哪一项？为什么？"
    ])

    # Module 4
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("模块四：实践篇——从知道到做到")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("知识点1：30天行动计划")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "为什么是30天？",
        "行为科学研究表明，形成一个新习惯平均需要21天。这就是为什么30天是一个理想的行为改变周期——它给了你足够的时间形成习惯，又足够短以保持紧迫感。\n\n更重要的是，30天计划让你能够：\n- 设定清晰、可衡量的目标\n- 追踪进度并看到变化\n- 在结束后总结学习并规划下一阶段",
        highlight="知道做不到，等于不知道。真正的学习发生在行动之后。"
    )

    add_knowledge_point(doc,
        "制定有效的30天计划",
        "一个有效的30天计划应该包含：\n\n1. 具体目标：不是我要提升团队心理安全感，而是我要在每次会议中主动邀请至少一位还没发言的同事\n\n2. 可衡量：每周结束时评估是否达成\n3. 关联工作：目标应该与你的实际工作场景紧密结合\n4. 预留缓冲：留出时间应对计划外的紧急事务\n\n失败计划 vs 成功计划对比：\nX 失败：我要对团队更好一点（太模糊）\n对勾 成功：我每天要具体感谢至少一位团队成员的贡献（具体可衡量）",
        "一位项目经理发现自己团队的问题是会议中总是一两个人主导，其他人很少发言。他制定的30天计划是：每次团队会议，我都要在讨论环节点名邀请至少2位最近较少发言的成员发表意见，并在下周复盘时记录他们的参与情况。30天后，他发现这些成员的发言频率提升了60%。",
        "模糊的意图导致模糊的结果。具体的计划导致具体的变化。"
    )

    add_form_header(doc, "4.1", "30天行动计划模板",
        "目的：帮助你制定具体、可执行的30天行动计划\n要求：填写以下各栏，确保每一条都是具体的、可衡量的\n时间：20分钟完成")

    plan_table = doc.add_table(rows=8, cols=3)
    plan_table.style = "Table Grid"
    headers = ["要素", "具体内容", "说明示例"]
    header_row = plan_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    plan_items = [
        ("我的目标", "", "例：在每次团队会议中邀请至少2位较少发言的成员"),
        ("为什么这个目标重要", "", "例：团队决策需要全员的智慧，而不只是少数人的观点"),
        ("我将如何衡量是否达成", "", "例：记录每周会议上被点名发言的成员数量，目标是提升50%"),
        ("可能遇到的障碍", "", "例：时间紧张时可能会跳过邀请环节；有人可能不愿意发言"),
        ("我的应对方案", "", "例：设置会议提醒；先私下与不主动发言的成员沟通"),
        ("第一周行动计划", "", "例：周一团队会议开始试行，记录情况"),
        ("需要的支持或资源", "", "例：希望团队成员配合，当我邀请时愿意回应"),
    ]
    for i, (item, content, example) in enumerate(plan_items):
        row = plan_table.rows[i + 1]
        row.cells[0].text = item
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(row.cells[0], "F2F2F2")
        row.cells[1].text = content
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = example
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(8)
        row.cells[2].paragraphs[0].runs[0].font.italic = True
        row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    para = doc.add_paragraph()
    run = para.add_run("知识点2：常见陷阱与避免方法")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    add_knowledge_point(doc,
        "陷阱一：过度简化",
        "以为建立心理安全感只是对团队好一点或多表扬。\n\n实际上，心理安全感建设涉及深层的文化和工作方式改变。过度简化会让你在遇到挑战时轻易放弃。",
        highlight="心理安全感建设不是nice，而是有效。"
    )

    add_knowledge_point(doc,
        "陷阱二：急于求成",
        "期待几周内看到显著变化。\n\n心理安全感的建立需要时间，尤其是当团队已经习惯了某种互动模式时。你可能需要3-6个月才能看到可持续的变化。",
        "一家公司在推行心理安全感建设6周后做了一次评估，发现分数没有明显变化。团队负责人一度想放弃。但继续坚持到第12周时，分数开始显著提升，到第6个月时已经成为公司心理安全感最高的团队。",
        "变化是非线性的——往往是积累到某个临界点才会显现。"
    )

    add_knowledge_point(doc,
        "陷阱三：言行不一",
        "说了要建立心理安全感，但实际行动与承诺不符。\n\n这是最危险的陷阱——如果你一边说欢迎不同意见，一边在有人提出异议时脸色骤变，团队会记住你的行为而不是你的话语。",
        highlight="团队成员是敏锐的观察者。他们看的不是你说什么，而是你做什么。"
    )

    add_knowledge_point(doc,
        "陷阱四：单独行动",
        "试图一个人推动心理安全感建设，而不借助团队或组织的力量。\n\n实际上，心理安全感建设需要：\n- 上级的支持\n- HR或培训资源的配合\n- 团队成员的共同参与",
        highlight="你可以是发起者，但你不能是唯一的推动者。"
    )

    add_form_header(doc, "4.2", "效果追踪表",
        "目的：帮助你在30天行动计划期间追踪进展并及时调整\n要求：每周结束时填写一次，对照目标评估进展\n时间：每周10分钟")

    tracking_table = doc.add_table(rows=6, cols=5)
    tracking_table.style = "Table Grid"
    headers = ["周次", "目标完成情况", "遇到的挑战", "调整措施", "下一步行动"]
    header_row = tracking_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i in range(1, 6):
        tracking_table.rows[i].cells[0].text = "第" + str(i) + "周"
        tracking_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j in range(5):
            cell = tracking_table.rows[i].cells[j]
            if cell.text:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            else:
                run = cell.paragraphs[0].add_run("")
                run.font.size = Pt(9)

    add_exercise(doc, [
        "基于课程开始时填写的场景卡，制定你的30天行动计划",
        "识别你认为最可能的陷阱，并提前准备应对方案",
        "找一个问责伙伴（可以是同事或导师），每周分享你的进展"
    ])

    # Course closing
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("课程收尾")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("七习惯重测自评（与课前对比）")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    intro_text = "课程开始时，你填写了心理安全感自评问卷。现在，请再次完成同样的问卷，与课前对比：\n\n- 你的总分变化了多少？\n- 哪些维度提升最多？哪些提升最少？\n- 这些变化说明什么？"
    para = doc.add_paragraph()
    run = para.add_run(intro_text)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(12)

    comparison_table = doc.add_table(rows=9, cols=4)
    comparison_table.style = "Table Grid"
    headers = ["自评维度", "课前得分", "课后得分", "变化"]
    header_row = comparison_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    dimensions = [
        "在团队会议上，我会主动提出与主流意见不同的观点",
        "当我犯错误时，我相信团队会把它当作学习机会",
        "我敢于在团队面前承认我不知道或我需要帮助",
        "我相信我的意见会被认真倾听，即使与他人不同",
        "当我对某件事有疑虑时，我会直接表达，而不会保持沉默",
        "我感觉团队成员之间可以坦诚地相互反馈",
        "我不担心因为提出不同意见而被团队排斥或打击",
        "当团队遇到问题时，成员们能够开放地讨论而不是相互指责"
    ]
    for i, dim in enumerate(dimensions):
        row = comparison_table.rows[i + 1]
        row.cells[0].text = dim
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = ""
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].text = ""
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].text = ""
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("我的30天行动计划承诺")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    commitment_text = """今天我承诺，在接下来的30天里，我将坚持以下行动来建设团队心理安全感：

承诺行动：_______________________________________________________________

开始日期：____________  预计结束日期：____________

问责伙伴：____________  联系方式：____________

签名：________________"""

    para = doc.add_paragraph()
    run = para.add_run(commitment_text)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("致出发的你")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    closing_text = """这不是结束，而是开始。

你带着这本手册离开，但带走的不应该只是纸上的知识。真正的学习发生在明天早上你走进办公室的那一刻，发生在你下一次主持团队会议的时候，发生在你下一次面对不同意见时的反应中。

记住：

领先一步，枪打出头鸟；落后半步，别人牵牛我拔桩；领先半步，吃尽红利。

建立心理安全感这件事，你不需要等团队准备好，不需要等领导支持，不需要等合适的时机。

从今天开始，从下一件事开始，从你自己开始。

你能做到。

——课程团队"""

    para = doc.add_paragraph()
    run = para.add_run(closing_text)
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_after = Pt(16)

    # Appendix
    doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run("附录")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after = Pt(16)

    para = doc.add_paragraph()
    run = para.add_run("附录一：术语速查表")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    terms = [
        ("心理安全感", "Psychological Safety", "团队成员相信自己在面对人际风险时是安全的，不会因说真话、承认错误或提出不同意见而受到惩罚或羞辱。"),
        ("谷歌氧气计划", "Project Aristotle", "谷歌于2008年启动的内部研究项目，旨在找出高效团队的关键特征。研究发现心理安全感是区分高绩效团队的最重要因素。"),
        ("Amy Edmondson", "", "哈佛商学院教授，心理安全感概念的提出者。她的研究改变了人们对团队绩效和领导力的认知。"),
        ("人际风险", "Interpersonal Risk", "在社交互动中可能面临的风险，如被嘲笑、被排斥、被认为无能等。心理安全感让人们愿意承担这些人际风险。"),
        ("学习行为", "Learning Behavior", "包括提问、反馈、承认错误、寻求帮助等行为。在心理安全感高的环境中，这类行为更频繁。"),
        ("示弱", "Showing Vulnerability", "领导者公开承认自己的不足、不确定或需要帮助的行为。这是建立心理安全感的关键行为之一。"),
    ]
    term_table = doc.add_table(rows=len(terms)+1, cols=3)
    term_table.style = "Table Grid"
    headers = ["术语", "英文/来源", "定义或说明"]
    header_row = term_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (term, eng, definition) in enumerate(terms):
        row = term_table.rows[i + 1]
        row.cells[0].text = term
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].text = eng
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[1].paragraphs[0].runs[0].font.italic = True
        row.cells[2].text = definition
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("附录二：工具速查索引")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_BLUE
    run.font.name = "微软雅黑"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(8)

    tools = [
        ("表单1.1", "心理安全感自评问卷", "个人自我评估", "P. 6"),
        ("表单1.2", "我的团队现状分析", "团队现状分析", "P. 7"),
        ("表单2.1", "团队心理安全感评估问卷（7题版）", "团队评估工具", "P. 11"),
        ("表单2.2", "管理者行为自检表", "管理者自我反思", "P. 13"),
        ("表单3.1", "每日行动清单", "日常行为养成", "P. 16"),
        ("表单3.2", "每周行动清单", "周期性深度建设", "P. 18"),
        ("表单3.3", "每月行动清单", "长期文化建设", "P. 19"),
        ("表单4.1", "30天行动计划模板", "个人行动计划制定", "P. 22"),
        ("表单4.2", "效果追踪表", "进度追踪与调整", "P. 24"),
    ]
    tool_table = doc.add_table(rows=len(tools)+1, cols=4)
    tool_table.style = "Table Grid"
    headers = ["表单编号", "表单名称", "用途", "页码"]
    header_row = tool_table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "2563EB")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, (num, name, usage, page) in enumerate(tools):
        row = tool_table.rows[i + 1]
        row.cells[0].text = num
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = name
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[2].text = usage
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[3].text = page
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path = "D:/新课开发/管理学/35-心理安全感与信任文化/03_学员手册/35-心理安全感与信任文化_学员手册.docx"
    doc.save(output_path)
    print("Document saved to: " + output_path)
    return output_path

if __name__ == "__main__":
    build_document()

# -*- coding: utf-8 -*-
"""
生成【08_GDP通胀与失业】课程讲师手册 Word文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

OUTPUT_PATH = "D:/新课开发/经济学/08_GDP通胀与失业/讲师手册/讲师手册_GDP通胀与失业.docx"

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def add_styled_heading(doc, text, level=1, color=None):
    """添加带样式标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if color:
            run.font.color.rgb = RGBColor(*color)
    return heading

def add_normal_para(doc, text, bold=False, indent=False):
    """添加普通段落"""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    if indent:
        para.paragraph_format.first_line_indent = Pt(21)
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = bold
    return para

def add_bullet_point(doc, text, indent_level=0):
    """添加项目符号段落"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    return para

def create_table_with_header(doc, headers, rows, header_color='D9E2F3'):
    """创建带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)

    return table

def create_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('GDP通胀与失业')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('宏观经济仪表盘阅读能力')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()
    doc.add_paragraph()

    manual_title = doc.add_paragraph()
    manual_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = manual_title.add_run('讲师手册')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 课程信息表
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ('课程定位', '零基础财经素养课，聚焦宏观经济仪表盘阅读能力'),
        ('目标学员', '零基础成年人（25-50岁），需要理解经济数据做出日常决策'),
        ('课程时长', '约6小时（不含休息）'),
        ('学员人数', '建议20-40人'),
        ('课程难度', '入门级，无需经济学背景')
    ]

    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.rows[i].cells[0]
        cell_label.text = label
        set_cell_shading(cell_label, 'E2EFDA')
        for para in cell_label.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(11)

        cell_value = info_table.rows[i].cells[1]
        cell_value.text = value
        for para in cell_value.paragraphs:
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(11)

    doc.add_page_break()

    # ========== 目录 ==========
    add_styled_heading(doc, '目录', level=1, color=(31, 73, 125))

    toc_items = [
        '第一章：讲师准备指南（课前检查清单）',
        '第二章：授课流程指引',
        '第三章：学员常见问题处理（FAQ）',
        '第四章：点评话术参考',
        '第五章：时间控制建议',
        '第六章：讲师注意事项',
        '附录：课程框架与核心概念速查'
    ]

    for item in toc_items:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ========== 第一章：讲师准备指南 ==========
    add_styled_heading(doc, '第一章：讲师准备指南（课前检查清单）', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '1.1 课前30天准备事项', level=2)

    headers = ['准备事项', '具体内容', '完成状态']
    rows = [
        ('确认课程时间和地点', '与组织方确认教室/线上会议平台预约', '□'),
        ('发送课前通知', '向学员发送课程大纲、课前预习材料、学习目标', '□'),
        ('检查PPT和讲义', '确认所有课件可正常播放，备用U盘准备', '□'),
        ('准备互动道具', '白板/马克笔、便签纸、计时器、分组卡片', '□'),
        ('了解学员背景', '通过问卷或访谈了解学员行业、职位、学习期望', '□'),
        ('准备案例素材', '更新最新经济数据（2024年数据为主）', '□'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '1.2 课前7天准备事项', level=2)

    headers = ['准备事项', '具体内容', '完成状态']
    rows = [
        ('确认学员名单', '了解学员人数、行业分布、年龄层', '□'),
        ('检查设备', '投影仪、音响、麦克风、网络连接测试', '□'),
        ('打印教材', '学员手册、评估量表、工具表单（如适用）', '□'),
        ('准备水和小点心', '教室环境准备', '□'),
        ('准备自我介绍', '设计破冰自我介绍（约2分钟）', '□'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '1.3 课前1天准备事项', level=2)

    headers = ['准备事项', '具体内容', '完成状态']
    rows = [
        ('再次确认设备', 'PPT翻页笔、备用电池、投影测试', '□'),
        ('打印应急清单', '课程流程表、FAQ答案卡、突发情况处理方案', '□'),
        ('准备音乐', '课间休息背景音乐（轻松风格）', '□'),
        ('充足的休息', '前一晚不熬夜，保持精力充沛', '□'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '1.4 课前30分钟准备事项', level=2)

    headers = ['准备事项', '具体内容', '完成状态']
    rows = [
        ('布置教室', '座位摆放（建议U型或小组团坐）、投影角度', '□'),
        ('打开课件', '封面页展示，确认播放正常', '□'),
        ('发放材料', '学员手册、笔、便签', '□'),
        ('签到准备', '准备签到表或电子签到', '□'),
        ('开场预热', '播放轻音乐，迎接学员入场', '□'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '1.5 讲师自检清单（出发前确认）', level=2)

    checklist_items = [
        '□ 笔记本电脑和充电器',
        '□ PPT翻页笔和备用电池',
        '□ 讲师手册（电子版+打印版）',
        '□ 学员手册份数（按人数）',
        '□ 评估工具（前后测问卷）',
        '□ 白板笔（黑/红/蓝）和板擦',
        '□ 计时器（手机或专用）',
        '□ 名片（便于课后联络）',
        '□ 水杯和个人用品',
    ]

    for item in checklist_items:
        add_bullet_point(doc, item)

    doc.add_page_break()

    # ========== 第二章：授课流程指引 ==========
    add_styled_heading(doc, '第二章：授课流程指引', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '2.1 课程总体时间分配', level=2)

    headers = ['模块', '内容', '时长', '占比']
    rows = [
        ('破冰与导入', '课程介绍、学习目标设定、学员破冰', '20分钟', '5.5%'),
        ('第一章', '走进宏观经济仪表盘', '30分钟', '8.3%'),
        ('第二章', 'GDP经济总量怎么看', '50分钟', '13.9%'),
        ('第三章', '通胀——物价变化怎么读', '50分钟', '13.9%'),
        ('午餐/休息', '', '60分钟', '-'),
        ('第四章', '失业率——就业冷暖怎么判', '50分钟', '13.9%'),
        ('第五章', '三个指标的三角关系', '40分钟', '11.1%'),
        ('第六章', '经济周期春夏秋冬怎么看', '45分钟', '12.5%'),
        ('第七章', '政策应对——政府央行怎么出牌', '35分钟', '9.7%'),
        ('第八章', '总结与行动方案', '30分钟', '8.3%'),
        ('Q&A与结业', '问题解答、课程总结、行动承诺', '10分钟', '2.8%'),
        ('合计', '（不含休息）', '约360分钟', '100%'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '2.2 详细模块流程', level=2)

    # 模块1：破冰与导入
    add_styled_heading(doc, '模块1：破冰与导入（20分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '注意事项']
    rows = [
        ('迎接学员', '5分钟', '站在门口迎接，微笑点头，引导入座', '确认每位学员都收到课前通知'),
        ('破冰游戏', '8分钟', '两个真相一个谎言——每人写三个关于自己的经济相关事实，一个假的，其他学员猜哪个是假的', '话题要引导向经济话题，如我关注CPI、我买过股票'),
        ('课程介绍', '7分钟', '使用PPT封面，介绍课程目标：学完这门课，你能读懂财经新闻，为自己的钱包做决策', '不要照读PPT，用自己的话讲'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第一章
    add_styled_heading(doc, '第一章：走进宏观经济仪表盘（30分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('开场提问', '5分钟', '问学员：你们最近看到过哪些财经新闻？让3-4人分享', '开放式提问，营造讨论氛围'),
        ('概念讲解', '10分钟', '用汽车仪表盘类比解释GDP、CPI、失业率三大指标', '用生活化语言，避免学术定义'),
        ('案例分析', '8分钟', '引用2024年数据，展示三大指标的实际表现', '数据要新，不要用过时的例子'),
        ('表单练习', '5分钟', '让学员完成表单1.1宏观经济关联度自测', '允许低声讨论，不要给答案'),
        ('小结', '2分钟', '总结本章三个要点，说明下一章内容', '承上启下'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第二章
    add_styled_heading(doc, '第二章：GDP经济总量怎么看（50分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('引入', '5分钟', '展示2023年中国GDP数据，引出这个数字意味着什么的问题', '可以用问答形式'),
        ('GDP定义', '8分钟', '用一个国家一年创造的总财富解释GDP，区分名义GDP和实际GDP', '举例要具体，如126万亿约等于...'),
        ('三种计算方法', '15分钟', '生产法、收入法、支出法各用一个生活化例子说明', '用家庭收支做类比，让学员代入'),
        ('中国GDP解读', '10分钟', '展示2010-2024年GDP增速变化表，讨论增速换挡现象', '引导学员思考：增速放缓是好是坏？'),
        ('局限性讨论', '8分钟', '用GDP能告诉你什么、不能告诉你什么结构讲解', '可以提问：GDP增长，你的工资增长了吗？'),
        ('练习与小结', '4分钟', '完成练习2-A判断GDP增长的质量，快速对答案', '给3分钟思考，然后公布参考答案'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第三章
    add_styled_heading(doc, '第三章：通胀——物价变化怎么读（50分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('场景导入', '5分钟', '问学员：你们有没有感觉最近超市的东西涨价了？引入通胀概念', '用生活场景引发共鸣'),
        ('CPI/PPI讲解', '10分钟', '用菜市场账本和工厂进货账本类比解释CPI和PPI', '画图帮助理解传导关系'),
        ('通胀成因', '12分钟', '讲需求拉动、成本推动、货币超发三种类型，用2024年真实案例', '每个类型用一个生活例子'),
        ('对普通人的影响', '10分钟', '讲储蓄缩水、工资购买力、债务负担三个维度', '重点讲谁受益谁受损'),
        ('中国特色', '8分钟', '讲猪周期和房租统计差异', '这是中国独有的难点，要放慢速度'),
        ('表单练习', '5分钟', '让学员完成表单3.1通胀敏感度自测', '给学员2分钟填写，然后简短讨论'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第四章
    add_styled_heading(doc, '第四章：失业率——就业冷暖怎么判（50分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('开头提问', '5分钟', '问学员：如果明天你失业了，你能撑多久？引入主题', '问题要具体，让学员有代入感'),
        ('统计口径', '10分钟', '区分调查失业率和登记失业率，强调数据局限', '可以用真实体温vs挂号记录类比'),
        ('三种失业类型', '15分钟', '用三个真实案例讲解摩擦性、结构性、周期性失业', '案例要有画面感，如东北老工业基地'),
        ('充分就业', '5分钟', '解释自然失业率和奥肯定律', '可以用简单数字推演'),
        ('中国就业市场', '10分钟', '讲青年失业率、农民工、灵活用工三个特点', '这是中国特有的问题，要重点讲'),
        ('表单练习', '5分钟', '让学员完成表单4.1就业形势判断', '结合当前就业环境讨论'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第五章
    add_styled_heading(doc, '第五章：三个指标的三角关系（40分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('概念讲解', '15分钟', '讲菲利普斯曲线、奥肯定律，解释三个指标的联动关系', '用画图方式展示三角关系'),
        ('中美对比', '10分钟', '展示2020-2024年中美GDP/CPI/失业率数据对比', '让学员自己发现规律'),
        ('政策困境', '10分钟', '用滞胀、通缩螺旋、低通胀陷阱三个情境讲政策两难', '用跑步机的类比'),
        ('综合练习', '5分钟', '用练习5-A判断各国经济阶段', '给3分钟思考，然后讨论'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第六章
    add_styled_heading(doc, '第六章：经济周期春夏秋冬怎么看（45分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('季节比喻', '5分钟', '用春夏秋冬类比经济周期的四个阶段', '可以问学员：你们觉得现在是什么季节？'),
        ('四阶段详解', '20分钟', '复苏、过热、滞胀、衰退各阶段的特征、信号、应对', '每个阶段用表情包/生活场景描述'),
        ('美林时钟', '10分钟', '讲复苏期配置股票、过热期配置商品、滞胀期配置现金、衰退期配置债券', '重点讲现在（2024年）中国是什么季节'),
        ('表单练习', '10分钟', '让学员完成表单6.1经济周期自测，结合自己行业讨论', '这个练习可以让学员代入自己的行业'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第七章
    add_styled_heading(doc, '第七章：政策应对——政府央行怎么出牌（35分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('财政政策', '12分钟', '讲税收、政府支出、赤字与国债三个工具', '用家庭开支类比'),
        ('货币政策', '12分钟', '讲利率、存款准备金率、公开市场操作三个工具', '用水龙头类比货币供应'),
        ('传导机制', '6分钟', '讲政策从出台到影响个人的完整链条', '强调药到病除需要时间'),
        ('中美对比', '5分钟', '对比2020-2024年中美政策应对差异', '让学员理解不同国情，不同棋路'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # 第八章
    add_styled_heading(doc, '第八章：总结与行动方案（30分钟）', level=3)

    headers = ['环节', '时长', '操作指引', '互动设计']
    rows = [
        ('框架回顾', '10分钟', '用PPT框架图回顾整门课的知识结构', '快速串讲，不要展开细节'),
        ('七步法', '8分钟', '讲解如何解读一条宏观经济新闻的七步法', '用本章的完整案例示范'),
        ('行动承诺', '10分钟', '让学员填写行为承诺，并分享给同桌', '给5分钟写，然后2分钟分享'),
        ('结业', '2分钟', '总结课程，解答最后问题，发放课程完成证书', '微笑结语，感谢学员参与'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_page_break()

    # ========== 第三章：学员常见问题处理（FAQ） ==========
    add_styled_heading(doc, '第三章：学员常见问题处理（FAQ）', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '3.1 关于课程内容的FAQ', level=2)

    faq_items = [
        (
            'Q1: GDP增长这么快，为什么我的工资没涨？',
            'A: 这个问题触及GDP的分配结构。GDP增长是总量概念，但分配结构会影响每个人实际感受。可以用GDP含金量概念（可支配收入/人均GDP）来说明：中国约55-60%，发达国家约60-70%，说明分配还有优化空间。另外，平均数不代表中位数，高收入者拉高了平均数。',
            '重点：不要否认学员的感受，而是用数据解释这种感受的来源'
        ),
        (
            'Q2: CPI数据说通胀很低，但我感觉物价涨了很多，是数据有问题吗？',
            'A: 这是个好问题。CPI统计的是一篮子商品，而且权重是固定的。猪肉在中国CPI中权重约2.5-3%，如果猪价下跌，会拉低整体CPI，即便其他商品在涨。另外，CPI统计的是消费品价格，不包含房价（购房算投资）。可以用租房子和买房的区别来解释。',
            '重点：引导学员理解统计口径的局限性，而不是说数据是假的'
        ),
        (
            'Q3: 中国GDP超过日本了，这意味着什么？',
            'A: GDP总量超过日本是中国经济发展的重要里程碑，说明中国是全球第二大经济体。但要注意：1，人均GDP中国约1.27万美元，日本约4万美元，差距还很大；2，GDP是总量指标，不反映经济质量（创新能力、产业结构、人均福利）；3，可以引用中国用40年走完了发达国家上百年的工业化路程来说明成就，同时指出差距。',
            '重点：既肯定成就，也不回避差距，保持客观'
        ),
        (
            'Q4: 失业率5%，是不是意味着100个人里有5个失业？',
            'A: 基本正确，但要注意：1，这个数字是城镇调查失业率，不包含农村劳动力和灵活就业者；2，失业的定义是没有工作+正在寻找+可以工作，很多人放弃找工作后就不算失业了；3，实际就业压力可能比数字呈现的更大，尤其是青年失业率和农民工群体。',
            '重点：帮助学员理解统计口径，避免简单化'
        ),
        (
            'Q5: 经济周期春夏秋冬是固定的循环吗？能预测吗？',
            'A: 经济周期有规律，但不是精确的时钟。每个周期的长度和幅度都不同，受政策、技术、外部冲击（如疫情）影响。达利欧说过：经济像一台机器，但它不是精密仪器。我们的目标不是预测，而是判断当前大概率处于哪个阶段，然后做对应的决策。',
            '重点：强调概率思维和应对而非预测'
        ),
        (
            'Q6: 政府出那么多政策，为什么经济还是不好？',
            'A: 政策传导需要时间，通常有几周到几个月的时滞。另外，政策效果需要信心这个放大器——如果市场预期悲观，同样的政策效果会打折。还有，中国经济的问题往往是结构性的（房地产调整、人口老龄化、全球需求萎缩），不是靠短期政策能解决的。',
            '重点：用传导链条和结构性问题来解释'
        ),
        (
            'Q7: 我应该关注哪些经济数据？在哪里看？',
            'A: 入门级推荐：1，国家统计局官网和公众号（月度数据）；2，中国人民银行官网（货币政策）；3，Wind/同花顺/东方财富（综合数据平台）。进阶级：4，美联储FRED数据库（全球比较）；5，任泽平、券商宏观研报。关键是建立定期跟踪的习惯，而不是追求面面俱到。',
            '重点：给出具体可操作的数据来源'
        ),
        (
            'Q8: 通胀对我的存款有什么影响？我应该怎么应对？',
            'A: 如果通胀率是3%，你存银行利率是2%，你的实际购买力在缩水。应对方式：1，减少现金持有，适当配置抗通胀资产（黄金、房产、股票指数基金）；2，提升自己的技能，让工资涨幅跑赢通胀；3，如果有房贷，适度通胀其实在帮你减轻债务负担。',
            '重点：给出具体的行动建议'
        ),
    ]

    for q, a, note in faq_items:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.bold = True

        para2 = doc.add_paragraph()
        run2 = para2.add_run(a)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para2.paragraph_format.space_after = Pt(3)

        para3 = doc.add_paragraph()
        run3 = para3.add_run('讲师提示：' + note)
        run3.font.name = '微软雅黑'
        run3._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(128, 128, 128)
        para3.paragraph_format.space_after = Pt(12)

    add_styled_heading(doc, '3.2 关于课堂互动的FAQ', level=2)

    interaction_faqs = [
        (
            'Q: 学员不主动发言怎么办？',
            'A: 1，先用选择题而非开放式问题；2，给30秒独立思考时间；3，可以举手投票代替发言；4，对第一个发言的学员给予具体表扬；5，小组讨论后让代表汇报，降低发言压力。'
        ),
        (
            'Q: 学员提问超出课程范围或很难回答怎么办？',
            'A: 1，承认这个问题很好，我需要研究一下；2，可以说这个问题和我们今天的主题有关系的部分是...；3，记录问题，课后解答；4，不要不懂装懂。'
        ),
        (
            'Q: 有学员质疑课程内容的权威性怎么办？',
            'A: 1，承认经济学有很多学派和观点；2，可以说这个框架是主流学术共识，也有人持不同看法；3，展示数据来源（国家统计局、美联储等）；4，避免争论，聚焦课程目标。'
        ),
    ]

    for q, a in interaction_faqs:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.bold = True

        para2 = doc.add_paragraph()
        run2 = para2.add_run(a)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para2.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ========== 第四章：点评话术参考 ==========
    add_styled_heading(doc, '第四章：点评话术参考', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '4.1 表单/练习点评话术', level=2)

    headers = ['场景', '点评话术', '目的']
    rows = [
        ('学员完成自测表单后', '你这个答案说明你已经能联系自己的生活来思考了，这就是我们这门课要培养的能力', '鼓励学员建立经济数据与生活的联系'),
        ('学员回答GDP增长是好事', 'GDP增长确实是好事，但我们要看这个增长是怎么来的、花在哪儿、分配给了谁——这才完整', '引导学员深入思考'),
        ('学员回答通胀越低越好', '你的直觉很有道理，但通胀太低甚至通缩，反而可能是需求萎缩的信号，比如日本失去的二十年就是例子', '纠正常见误区'),
        ('学员回答失业率高一定是坏事', '你看2023年青年失业率高，但整体失业率只有5.1%——这说明不是没有人找工作，而是有些人找不到想要的岗位——这是结构性问题', '帮助学员理解结构性失业'),
        ('学员回答政府应该大量发钱刺激经济', '发钱可以短期刺激需求，但副作用是通胀——你看美国2020年的例子。政策制定者是在两难中找平衡，这也是为什么我们要理解这些逻辑', '用案例说明政策权衡'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '4.2 提问后的点评技巧', level=2)

    question_tips = [
        ('学员回答正确', '很好，你抓住了关键——这就是我们说的XXX', '确认正确理解'),
        ('学员回答部分正确', '你看到了XXX这一点，这很好。但还有一个维度——YYY，我们也要考虑进去', '补充遗漏信息'),
        ('学员回答偏离主题', '你的角度很有意思，这让我想到YYY——但今天这个话题我们先聚焦在XXX', '巧妙拉回主题'),
        ('学员回答很有深度', '你说到点子上了，这其实就是我们下一章要深入讲的XXX——你可以带着这个问题继续学习', '肯定+引导'),
        ('无人回答', '有没有人愿意分享自己的第一反应？不用完美，说错也没关系', '降低发言压力'),
    ]

    for scenario, tip, _ in question_tips:
        para = doc.add_paragraph()
        run1 = para.add_run(scenario + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(tip)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(8)

    add_styled_heading(doc, '4.3 过渡语参考', level=2)

    transitions = [
        ('第一章→第二章', '刚才我们认识了三个指标，现在我们来深入第一个——GDP。学习完这章，你能说出GDP是怎么算出来的，以及它有哪些局限'),
        ('第二章→第三章', 'GDP告诉我们经济蛋糕有多大，但蛋糕大了，每个人分到的购买力是变多了还是变少了？这就是通胀要回答的问题'),
        ('第三章→第四章', '通胀是物价问题，但找工作是另一个维度的民生——接下来我们看失业率'),
        ('第四章→第五章', '三个指标不是孤立的，它们之间有怎样的联动关系？这就是本章要回答的问题'),
        ('第五章→第六章', '理解了三角关系，我们就有了判断经济周期的工具——接下来看经济四季'),
        ('第六章→第七章', '知道经济在哪个季节，下一步就是理解政府和央行怎么调温——这章学政策应对'),
        ('第七章→第八章', '工具都学完了，最后我们来做总结，把知识变成行动'),
    ]

    for transition, words in transitions:
        para = doc.add_paragraph()
        run1 = para.add_run(transition + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(words)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ========== 第五章：时间控制建议 ==========
    add_styled_heading(doc, '第五章：时间控制建议', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '5.1 时间分配总览', level=2)

    headers = ['模块', '建议时长', '允许浮动', '超时后果']
    rows = [
        ('破冰与导入', '20分钟', '+5分钟', '压缩第一章时间'),
        ('第一章', '30分钟', '+5分钟', '影响第二章GDP的讲解深度'),
        ('第二章', '50分钟', '+5/-5分钟', '核心模块，可适当延长'),
        ('第三章', '50分钟', '+5/-5分钟', '核心模块，可适当延长'),
        ('午餐/休息', '60分钟', '+15分钟', '学员需要放松，不要压缩'),
        ('第四章', '50分钟', '+5/-5分钟', '核心模块，可适当延长'),
        ('第五章', '40分钟', '+5分钟', '传导关系是难点，可延'),
        ('第六章', '45分钟', '+5/-5分钟', '美林时钟是重点，可延'),
        ('第七章', '35分钟', '+5分钟', '政策工具讲解清楚即可'),
        ('第八章', '30分钟', '+5分钟', '行动承诺是核心，不能省'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '5.2 时间管理技巧', level=2)

    time_tips = [
        ('开场守时', '正式开课时间一到就开讲，不要等迟到学员。守时是对准时到达学员的尊重。'),
        ('手表/计时器', '在讲台放一个计时器（手机即可），每个环节开始时看一眼。'),
        ('模块截止', '如果某个环节超时超过5分钟，果断跳过或压缩，不要影响后续内容。'),
        ('休息准时', '午餐和下午休息时间要准时，学员需要上厕所、补充能量。'),
        ('留白意识', '不要讲太满，留出Q&A和过渡的时间。'),
        ('记录偏差', '课后记录本次实际时间分配，下次改进。'),
    ]

    for tip_title, tip_content in time_tips:
        para = doc.add_paragraph()
        run1 = para.add_run('• ' + tip_title + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(tip_content)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(6)

    add_styled_heading(doc, '5.3 突发时间不足的处理方案', level=2)

    shortage_solutions = [
        ('第一章可压缩', '走进宏观经济仪表盘可以缩短至20分钟，因为学员课前大概率已经预习过。'),
        ('第二章GDP计算方法可简化', '三种方法各讲一个例子即可，不用每个方法都展开。'),
        ('第三章中国特色可略讲', '猪周期和房租统计差异如果学员反应平淡，可以跳过。'),
        ('第五章传导机制可略讲', '传导链条如果学员理解吃力，用记住政策需要时间带过即可。'),
        ('第七章政策工具可合并', '财政政策和货币政策可以各只讲一个核心工具。'),
        ('第八章行动承诺不能省', '这是课程的核心产出，无论如何要给足时间。'),
    ]

    for scenario, solution in shortage_solutions:
        para = doc.add_paragraph()
        run1 = para.add_run('• ' + scenario)
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(solution)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 第六章：讲师注意事项 ==========
    add_styled_heading(doc, '第六章：讲师注意事项', level=1, color=(31, 73, 125))

    add_styled_heading(doc, '6.1 授课禁忌', level=2)

    taboos = [
        ('不要照读PPT', 'PPT是辅助，不是讲稿。用自己的话讲，用生活例子解释概念。学员手册上有文字，不需要你再念一遍。'),
        ('不要回避学员的真实困惑', '学员问工资没涨是好事，说明他在思考。不要说GDP增长总是好的来敷衍，而是承认复杂性。'),
        ('不要把观点说成事实', '可以说主流经济学认为...，但不要说事实是...。经济学有学派，有争议，要留有余地。'),
        ('不要使用过多专业术语', '说失业不说摩擦性失业；说钱不值钱不说货币购买力下降。'),
        ('不要忽略学员的反应', '如果看到学员皱眉、打哈欠、玩手机，尝试调整节奏或互动方式。'),
        ('不要超时', '守时是对学员的尊重。如果内容讲不完，果断收尾，不要拖堂。'),
    ]

    for title, content in taboos:
        para = doc.add_paragraph()
        run1 = para.add_run('⚠️ ' + title + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True
        run1.font.color.rgb = RGBColor(192, 0, 0)

        run2 = para.add_run(content)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(8)

    add_styled_heading(doc, '6.2 突发情况处理', level=2)

    headers = ['突发情况', '应对方案']
    rows = [
        ('设备故障（投影仪/电脑）', '立即切换到备用方案：1，如果有U盘，插入另一台电脑；2，如果U盘也无法使用，让学员看手册上的图表，用口述方式继续；3，实在不行，临时改为问答互动，延后讲新内容'),
        ('学员突然晕倒/身体不适', '1，保持冷静，宣布休息10分钟；2，请工作人员协助处理；3，评估是否需要叫急救；4，其他学员继续讨论，不影响课程进度'),
        ('学员情绪激动/质疑课程', '1，不与学员争论；2，说你的感受我理解，我们可以课后继续聊；3，课后单独沟通，不要在课堂上公开辩论'),
        ('学员提问无法回答', '1，承认这个问题我需要研究一下；2，记录问题，课后解答；3，可以说这不在今天范围内，但你的问题很好'),
        ('课堂氛围过于沉闷', '1，立即切换互动方式——从讲授改为小组讨论；2，加入一个30秒的站起来伸懒腰环节；3，用问题引导，而不是讲更多内容'),
        ('课堂过于兴奋/偏离主题', '1，用微笑和暂停让气氛平静；2，说你们的讨论很有意思，我们先收回来；3，用过渡语拉回主题'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '6.3 讲师自我状态管理', level=2)

    self_care = [
        ('课前准备', '前一天早睡，课前30分钟到场，熟悉环境，调整状态。'),
        ('水分补充', '上课带水，课间及时补充水分，避免声音沙哑。'),
        ('能量管理', '午餐要吃，但不要过饱，避免下午犯困。'),
        ('情绪稳定', '遇到质疑或突发情况，保持微笑和冷静。学员能感受到你的状态。'),
        ('走动与站姿', '不要一直站在讲台，适度在教室走动，与学员建立连接。'),
        ('课后复盘', '下课后花5分钟记录本次授课的亮点和不足，下次改进。'),
    ]

    for item, content in self_care:
        para = doc.add_paragraph()
        run1 = para.add_run('• ' + item + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(content)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 附录：课程框架与核心概念速查 ==========
    add_styled_heading(doc, '附录：课程框架与核心概念速查', level=1, color=(31, 73, 125))

    add_styled_heading(doc, 'A.1 课程知识框架', level=2)

    framework = [
        ('第一章', '走进宏观经济仪表盘', '宏观经济定义、三大指标（GDP/CPI/失业率）、经济周期四阶段'),
        ('第二章', 'GDP经济总量怎么看', 'GDP定义、三种计算方法（生产/收入/支出）、GDP局限'),
        ('第三章', '通胀——物价变化怎么读', 'CPI/PPI/核心通胀、三种成因（需求/成本/货币）、三种类型（温和/恶性/通缩）'),
        ('第四章', '失业率——就业冷暖怎么判', '两种统计口径、三种失业类型（摩擦/结构/周期）、自然失业率'),
        ('第五章', '三个指标的三角关系', '菲利普斯曲线、奥肯定律、政策困境'),
        ('第六章', '经济周期春夏秋冬怎么看', '四阶段特征、美林时钟、2024年中国经济诊断'),
        ('第七章', '政策应对——政府央行怎么出牌', '财政政策三工具、货币政策三工具、政策传导机制'),
        ('第八章', '总结与行动方案', '四层分析框架、七步解读法、12个日常习惯'),
    ]

    headers = ['章节', '标题', '核心概念']
    rows = framework
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, 'A.2 关键数据速查（2020-2024年）', level=2)

    headers = ['指标', '2020', '2021', '2022', '2023', '2024']
    rows = [
        ('中国GDP增速', '2.2%', '8.4%', '3.0%', '5.2%', '约5.0%'),
        ('中国CPI', '0.2%', '0.9%', '2.0%', '0.2%', '约0.1%'),
        ('中国城镇调查失业率', '5.6%', '5.1%', '5.5%', '5.1%', '5.1%'),
        ('美国GDP增速', '-2.8%', '5.9%', '2.1%', '2.5%', '约2.8%'),
        ('美国CPI', '1.2%', '4.7%', '8.0%', '4.1%', '约3.0%'),
        ('美国失业率', '6.3%', '3.9%', '3.6%', '3.6%', '4.0%'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, 'A.3 核心公式速查', level=2)

    formulas = [
        ('实际收益率', '实际收益率 = 名义利率 - 通胀率'),
        ('GDP支出法', 'GDP = 消费(C) + 投资(I) + 政府购买(G) + 净出口(NX)'),
        ('GDP收入法', 'GDP = 劳动者报酬 + 生产税净额 + 固定资产折旧 + 营业盈余'),
        ('自然失业率', '自然失业率 = 摩擦性失业率 + 结构性失业率'),
        ('奥肯定律', 'GDP增速每提高2个百分点，失业率下降约1个百分点'),
    ]

    for formula_name, formula in formulas:
        para = doc.add_paragraph()
        run1 = para.add_run('• ' + formula_name + '：')
        run1.font.name = '微软雅黑'
        run1._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run1.font.size = Pt(11)
        run1.bold = True

        run2 = para.add_run(formula)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(6)

    add_styled_heading(doc, 'A.4 关键阈值速查', level=2)

    headers = ['指标', '阈值', '含义']
    rows = [
        ('CPI', '>3%', '超过央行容忍度，可能触发收紧政策'),
        ('CPI', '<0%', '通缩信号，需求不足'),
        ('城镇调查失业率', '>5.5%', '超过年度目标线'),
        ('青年失业率', '>20%', '2023年中国曾达此水平'),
        ('GDP增速', '>6%', '持续偏高，经济过热风险'),
        ('GDP增速', '<3%', '低于潜在增速，经济偏冷'),
        ('PMI', '>50', '制造业扩张'),
        ('PMI', '<50', '制造业收缩'),
    ]
    create_table_with_header(doc, headers, rows)

    doc.add_paragraph()

    # ========== 封底 ==========
    doc.add_page_break()
    doc.add_paragraph()
    doc.add_paragraph()

    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run('祝授课顺利！')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('《GDP通胀与失业——宏观经济仪表盘阅读能力》')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(128, 128, 128)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = footer_para2.add_run('讲师手册 v1.0')
    run2.font.name = '微软雅黑'
    run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"讲师手册已生成：{OUTPUT_PATH}")

if __name__ == '__main__':
    create_document()

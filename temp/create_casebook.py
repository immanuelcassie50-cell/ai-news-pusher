# -*- coding: utf-8 -*-
"""
Create Case Collection for Party Course
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

OUTPUT_DIR = "D:/新课开发/党业融合/经营者讲党课/完整课程包/007-案例集"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False):
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading_with_style(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 1:
        set_run_font(run, 'Microsoft YaHei', 18, bold=True)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, 'Microsoft YaHei', 14, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(8)
    return para

def create_casebook():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Cover title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("案例集")
    set_run_font(run, 'Microsoft YaHei', 28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("《讲党课：业务干部的登台表达赋能工作坊》")
    set_run_font(run, 'Microsoft YaHei', 14)

    doc.add_paragraph()

    # Usage instructions
    inst = doc.add_paragraph()
    run = inst.add_run("使用说明：")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)

    p = doc.add_paragraph()
    inst_text = (
        "本案例集包含讲师示范案例和学员参考案例。"
        "每个案例都标注了所属的主题模块，"
        "方便学员在学习时参照。\n\n"
        "案例使用建议：\n"
        "1. 先阅读讲师示范案例，理解“好的党课故事”长什么样\n"
        "2. 再阅读学员参考案例，了解不同风格和主题的呈现方式\n"
        "3. 对照《党课素材转化卡》，尝试将自己的故事进行转化"
    )
    run = p.add_run(inst_text)
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    # ===== Part 1: Instructor Demo Cases =====
    add_heading_with_style(doc, "第一部分：讲师示范案例", 1)

    # Case 1
    add_heading_with_style(doc, "案例一：一次艰难的抉择", 2)

    case1_meta = doc.add_paragraph()
    run = case1_meta.add_run("【主题】坚守与放弃之间的抉择 | 【时长】约4分钟")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    case1_text = (
        "2019年年中，我接手了一个连续三个月业绩下滑的团队。\n\n"
        "第一次开全员会，我刚说完“大家好”，"
        "下面就开始有人交头接耳。我看到老张——"
        "我们团队资历最老的销售——直接收拾东西准备走人。"
        "那一刻，我知道情况毕我想象的还要糟。\n\n"
        "摆在面前的路有三条：一是新官上任三把火，雷厉风行地换人；"
        "二是无为而治，给大家时间慢慢调整；"
        "三是找出问题的真正原因，对症下药。\n\n"
        "我选择了第五条路——我没有急着做任何决定，"
        "而是花了整整两周，一对一地和每一个销售人员深谈。"
        "不是问他们业绩为什么下滑，而是问他们："
        "“你当初为什么选择做销售？”“你最辉煌的一次经历是什么？”\n\n"
        "两周后，老张主动来找我，说：“领导，我想明白了，"
        "不是团队不行了，是我带头躲平了。”\n\n"
        "那一刻我意识到了，管理者最重要的不是做决策，"
        "而是先让团队愿意跟你一起面对问题。\n\n"
        "这就是为什么我常说：好的管理者，不是让问题消失，"
        "而是让团队敢于面对问题。"
    )

    p = doc.add_paragraph()
    run = p.add_run(case1_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_paragraph()

    analysis1 = doc.add_paragraph()
    run = analysis1.add_run("【案例分析】")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)

    analysis1_text = (
        "- 开场悬念：用一个画面（老张收拾东西准备走人）开场，立刻制造簧张感\n"
        "- 抉择还原：展示了三条路，让听众思考“如果是我会怎么选”\n"
        "- 转折揭示：第五条路（两周深谈）出人意料，但又在情理之中\n"
        "- 感悟自然生长：通过老张自己的觉醇，让道理自然浮现\n"
        "- 回扭主题：用一句话简洁有力地收尾"
    )

    p = doc.add_paragraph()
    run = p.add_run(analysis1_text)
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_page_break()

    # Case 2
    add_heading_with_style(doc, "案例二：那个深夜的电话", 2)

    case2_meta = doc.add_paragraph()
    run = case2_meta.add_run("【主题】责任与担当 | 【时长】约3分钟")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    case2_text = (
        "那是2020年春节前的最后一个工作日，晚上十一点，我正准备下班回家。\n\n"
        "电话响了。是一个项目现场的技术工人打来的，他说：“领导，设备出了点问题，"
        "可能要影响明天的发货。”\n\n"
        "我问他具体情况，他说不是很严重，但要是我能过来看看，可能会更稳如。\n\n"
        "我老婉当时就在客厅着着看我，她说：“这么晚了，明天再去不行吗？”\n\n"
        "我没有回答她。我想起十年前自己刚参加工作的时候，"
        "有一次深夜设备出了问题，我的师傅二话不说就赶到了现场。"
        "那年他五十二岁。\n\n"
        "我抓起外套出了门。四十分钟后，我到了现场，"
        "和工人们一起检查、维修、调试。凌晨三点，问题解决了。\n\n"
        "第二天早上八点，准时发货。\n\n"
        "这事让我明白了一个道理：什么叫担当？担当不是挂在嘴边的大道理，"
        "而是在别人犹豫的时候，你已经迈出了第一步。"
    )

    p = doc.add_paragraph()
    run = p.add_run(case2_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_paragraph()

    analysis2 = doc.add_paragraph()
    run = analysis2.add_run("【案例分析】")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)

    analysis2_text = (
        "- 时间线紧凑：从晚上十一点到凌晨三点，完整的故事弧线\n"
        "- 画面感强：“我老婉当时就在客厅着着看我”这个细节特别真实\n"
        "- 对比手法：十年前师傅的榜样 vs 现在的自己\n"
        "- 情感共鸣：每个人都有过年三十或深夜被工作打断的经历\n"
        "- 金句结尾：“担当不是挂在嘴边的大道理，而是......”"
    )

    p = doc.add_paragraph()
    run = p.add_run(analysis2_text)
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_page_break()

    # ===== Part 2: Student Reference Cases =====
    add_heading_with_style(doc, "第二部分：学员参考案例", 1)

    intro = doc.add_paragraph()
    run = intro.add_run("以下案例来自往期学员的真实故事，经脱故改编后作为学习参考。每个案例都保留了原始的救事结构，供学员对照学习。")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    # Student Case 1
    add_heading_with_style(doc, "学员案例一：第一次跨部门协作", 2)

    s1_meta = doc.add_paragraph()
    run = s1_meta.add_run("【学员】某制造企业生产总监 | 【主题】团队协作")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    s1_text = (
        "我来讲一个我自己的故事。\n\n"
        "三年前，我被调到一个新工厂当责任人。第一次开协调会，我就跟采购部的老王哀起来了。\n\n"
        "我觉得他的交期太慢，他说我的需求太急。那次会议不欢而散。\n\n"
        "会后，我没有去找领导告状，也没有跟他继续吵。我做了两件事：第一，我请他吃了一顿饭；第二，我带着他参观了我的车间。\n\n"
        "吃饭的时候，我没有谈工作，就是聊天。参观的时候，我没有解释我的困难，而是让他自己看。\n\n"
        "后来老王跟我说：“我以前真不知道你们车间这么忙。”从那以后，我们的配合好了很多。\n\n"
        "这事让我明白，跨部门协作的秘诰不是开多少次会，而是让对方真正理解你的处境。"
    )

    p = doc.add_paragraph()
    run = p.add_run(s1_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_paragraph()

    # Student Case 2
    add_heading_with_style(doc, "学员案例二：那个想离职的年轻人", 2)

    s2_meta = doc.add_paragraph()
    run = s2_meta.add_run("【学员】某科技公司部门经理 | 【主题】人才培养")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    s2_text = (
        "去年，我们团队有个小伙子写了一份邮件给领导，说想要转岗。\n\n"
        "领导把邮件转给了我，说：“你们部门的人，你自己处理。”\n\n"
        "我找那个小伙子聊了两次。第一次，他跟我说觉得在部门学不到东西。"
        "第二次，他跟我说其实也不知道想去哪里，就觉得现在的工作没意思。\n\n"
        "我没有跟他讲道理，也没有说他不懂感恩。"
        "我跟他做了三个月的“项目对子”——每周一次，我带着他一起做项目，"
        "让他看到真实的工作是什么样的。\n\n"
        "三个月后，他自己跟我说：“领导，我不想转岗了。”\n\n"
        "我想说的是，有时候员工不是真的想走，他只是看不到路在哪里。"
        "管理者的责任，不是评判他的选择对不对，而是帮他找到那条路。"
    )

    p = doc.add_paragraph()
    run = p.add_run(s2_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_paragraph()

    # Student Case 3
    add_heading_with_style(doc, "学员案例三：客户的无理投诉", 2)

    s3_meta = doc.add_paragraph()
    run = s3_meta.add_run("【学员】某服务企业客服总监 | 【主题】客户至上")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    s3_text = (
        "我想讲一个让我当时很娽尴，但现在想来很感谢那个客户的故事。\n\n"
        "两年前，我们接到一个投诉，客户在电话里骂了我们整整四十分钟。"
        "我的下属接完电话后哭了。\n\n"
        "我接手处理这件事。第一次跟客户沟通，我解释了我们的难处。"
        "客户不接受。第二次，我遙歉，不管是不是我们的错，先遙歉。"
        "客户还是不接受。\n\n"
        "第三次，我直接上门拜访。我跟客户说：“我不是来解决问题的，"
        "我是来听您说话的。”\n\n"
        "那天下午，我听他讲了三个小时。\n\n"
        "最后，客户跟我说：“其实我不是要你们赔钱，我就是要一个态度。”\n\n"
        "这事让我学会了一个道理：有时候客户投诉的不是事情本身，而是没有人听他们说话。"
    )

    p = doc.add_paragraph()
    run = p.add_run(s3_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_page_break()

    # ===== Part 3: Case Transformation Comparison =====
    add_heading_with_style(doc, "第三部分：案例转化对照", 1)

    intro3 = doc.add_paragraph()
    intro3_text = (
        "以下是同一个故事在“原始版本”和“转化版本”中的对比，"
        "帮助理解如何将日常工作汇报转化为有感染力的党课故事。"
    )
    run = intro3.add_run(intro3_text)
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    # Comparison table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ["转化维度", "原始版本（汇报式）", "转化版本（故事式）"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2F5496')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)
                r.font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ("开头", "各位领导好，我今天汇报一下本季度的工作情况。", "那是2020年春节前的最后一个工作日，晚上十一点，电话响了......"),
        ("主体", "本季度我们完成了XX指标，同比增长XX%，主要做法有三点：一是......二是......三是......", "摆在面前的路有三条：一是......二是......三是......我选择了第五条路——......"),
        ("细节", "通过团队共同努力，取得较好效果。", "我没有急着做任何决定，而是花了整整两周，一对一地和每一个人深谈......"),
        ("结尾", "以上是我的汇报，请各位领寽批评指正。", "那一刻我意识到了，管理者最重要的不是......这就是为什么我常说......"),
        ("时长", "约8分钟", "约3-4分钟"),
    ]

    for i, (dim, orig, transformed) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = dim
        set_cell_shading(row.cells[0], 'D9E2F3')
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 9, bold=True)

        row.cells[1].text = orig
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 9)

        row.cells[2].text = transformed
        for para in row.cells[2].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 9)

    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—— 结束 ——")
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    output_path = os.path.join(OUTPUT_DIR, "案例集.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")

if __name__ == "__main__":
    create_casebook()
    print("Case collection created!")

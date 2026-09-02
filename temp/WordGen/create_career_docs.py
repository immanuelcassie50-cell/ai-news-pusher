# -*- coding: utf-8 -*-
"""
创建「破局・重启：用CEO思维重塑职业生涯」课程Word文档
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COURSE_NAME = "破局・重启：用CEO思维重塑职业生涯"
COPYRIGHT = "© 罗宏伟 2026 | 仅供课程内部使用"
OUTPUT_BASE = "D:/新课开发/职业生涯和画布/破局・重启：用 CEO 思维重塑职业生涯"

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:' + edge
            element = OxmlElement(tag)
            for key, val in edge_data.items():
                element.set(qn(key), val)
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)

def add_title_and_subtitle(doc, title, subtitle):
    """添加标题和副标题"""
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

def add_copyright(doc):
    """添加版权声明"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(COPYRIGHT)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

def add_formatted_para(doc, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加格式化段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return para

# ========== 文档1: 个人经营诊断表 ==========
def create_diagnosis_table():
    """创建01-个人经营诊断表.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "个人经营诊断表", COURSE_NAME + " | 工具表单")

    # 10道自检题表格
    add_formatted_para(doc, "一、自检题", bold=True, size=12)

    headers = ["诊断维度", "评估问题", "选项", "得分"]
    questions = [
        ["时间观念", "你如何安排每天的工作时间？", "A. 按他人要求完成任务  B. 按自己计划执行  C. 灵活调整但有目标", ""],
        ["责任承担", "当工作出现问题时，你的反应是？", "A. 等待领导指示  B. 主动承担并解决  C. 推卸给他人", ""],
        ["价值创造", "你是否清楚自己的核心价值？", "A. 不清楚  B. 大概知道  C. 非常清晰", ""],
        ["成本意识", "你是否计算过自己的时薪价值？", "A. 从未想过  B. 想过但没计算  C. 清楚自己的时薪", ""],
        ["客户思维", "你如何定义自己的客户？", "A. 只想到公司/领导  B. 包括所有相关方  C. 建立了客户清单", ""],
        ["资产积累", "你是否有意识地积累个人资产？", "A. 只关注工资  B. 有一些意识  C. 系统性地积累", ""],
        ["风险意识", "你是否定期评估职业风险？", "A. 从不评估  B. 偶尔想想  C. 定期系统评估", ""],
        ["决策模式", "面对重大决策时，你通常的做法是？", "A. 等待指令  B. 建议后等待批准  C. 分析后自主决策", ""],
        ["资源整合", "你是否主动整合资源解决问题？", "A. 只用自己资源  B. 偶尔整合  C. 系统整合资源", ""],
        ["价值变现", "你如何将能力转化为收入？", "A. 只靠工资  B. 有副业尝试  C. 多元变现渠道", ""],
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], 'D9E2F3')

    # 数据行
    for row_data in questions:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val

    doc.add_paragraph()

    # 评分标准表格
    add_formatted_para(doc, "二、评分标准", bold=True, size=12)

    score_headers = ["总分范围", "心态类型", "说明"]
    score_data = [
        ["25-40分", "强雇员心态", "以完成指令为导向，等待被经营"],
        ["15-24分", "中度雇员心态", "有一定经营者意识，但仍需引导"],
        ["8-14分", "轻度雇员心态", "开始觉醒，具备基础经营者思维"],
        ["0-7分", "经营者心态", "完全具备CEO思维，自主经营人生"],
    ]

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr = table2.rows[0].cells
    for i, h in enumerate(score_headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'E2EFDA')

    for row_data in score_data:
        row = table2.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    # 学员信息表格
    add_formatted_para(doc, "三、学员信息", bold=True, size=12)

    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'
    info_headers = ["姓名", "日期", "得分", "心态类型"]

    for i, h in enumerate(info_headers):
        info_table.rows[0].cells[i].text = h
        info_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(info_table.rows[0].cells[i], 'FFF2CC')

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档2: 价值定位图 ==========
def create_value_positioning():
    """创建02-价值定位图.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "价值定位图", COURSE_NAME + " | 工具表单")

    add_formatted_para(doc, "请根据以下框架，明确你的个人价值定位：", size=11)

    # 5行2列表格
    headers = ["维度", "内容"]
    rows = [
        ["身份定位", "你希望成为什么样的人？（如：首席增长官、个人CEO、行业专家）"],
        ["目标客户", "谁是你最重要的客户？（公司、领导、自己、合作伙伴等）"],
        ["价值主张", "你能为他们解决什么问题和创造什么价值？"],
        ["差异化", "你的独特优势是什么？与同行相比有何不同？"],
        ["可信证明", "什么证据能证明你的价值？（案例、证书、数据、推荐信）"],
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = headers[0]
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].text = headers[1]
    hdr[1].paragraphs[0].runs[0].bold = True
    set_cell_shading(hdr[0], 'D9E2F3')
    set_cell_shading(hdr[1], 'D9E2F3')

    for row_data in rows:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row[0], 'F2F2F2')
        row[1].text = row_data[1]

    doc.add_paragraph()

    # 个人定位声明模板
    add_formatted_para(doc, "个人定位声明模板：", bold=True, size=12)

    template = doc.add_table(rows=1, cols=1)
    template.style = 'Table Grid'
    cell = template.rows[0].cells[0]
    cell.text = """我是一名________________（身份定位）

我专注于帮助________________（目标客户）解决________________问题（价值主张）

我的独特优势是________________，这是因为________________（差异化+可信证明）

我的行动承诺：________________"""
    set_cell_shading(cell, 'FFF2CC')

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档3: 个人资产负债表 ==========
def create_personal_balance_sheet():
    """创建03-个人资产负债表.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "个人资产负债表", COURSE_NAME + " | 工具表单")

    add_formatted_para(doc, "全面盘点你的个人资产与负债，制定净资产增长策略：", size=11)

    def add_asset_table(doc, title, items):
        add_formatted_para(doc, title, bold=True, size=12)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "资产项目"
        hdr[1].text = "当前价值（1-10分）"
        hdr[2].text = "备注"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E2EFDA')

        for item in items:
            row = table.add_row().cells
            row[0].text = item
            row[1].text = ""
            row[2].text = ""

        doc.add_paragraph()

    # 时间资产
    add_asset_table(doc, "一、时间资产", [
        "可控时间（每天专注工作时间）",
        "学习时间（自我提升投入）",
        "社交时间（人脉建设投入）",
    ])

    # 人脉资产
    add_asset_table(doc, "二、人脉资产", [
        "核心人脉（关键时刻可求助的人）",
        "行业资源（行业信息和机会来源）",
        "支持网络（情感和专业支持系统）",
    ])

    # 能力资产
    add_asset_table(doc, "三、能力资产", [
        "专业技能（可变现的专业能力）",
        "可迁移能力（跨领域适用的能力）",
        "认知水平（思维方式和判断力）",
    ])

    # 信任资产
    add_asset_table(doc, "四、信任资产", [
        "个人信誉（他人对你的信任度）",
        "品牌背书（已有的专业背书）",
        "成功案例（可证明的能力证据）",
    ])

    # 负债
    add_formatted_para(doc, "五、负债（消耗资产的项）", bold=True, size=12)
    liability_table = doc.add_table(rows=1, cols=3)
    liability_table.style = 'Table Grid'
    hdr = liability_table.rows[0].cells
    hdr[0].text = "负债项目"
    hdr[1].text = "严重程度（1-10分）"
    hdr[2].text = "应对策略"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'FCE4D6')

    liabilities = [
        "无效社交（消耗时间无回报）",
        "消极关系（消耗能量的关系）",
        "坏习惯（吸烟、熬夜等）",
        "盲目忙碌（无价值产出）",
    ]
    for item in liabilities:
        row = liability_table.add_row().cells
        row[0].text = item
        row[1].text = ""
        row[2].text = ""

    doc.add_paragraph()

    # 净资产计算
    add_formatted_para(doc, "六、净资产计算", bold=True, size=12)
    calc_table = doc.add_table(rows=6, cols=2)
    calc_table.style = 'Table Grid'
    calc_data = [
        ["时间资产总分", ""],
        ["人脉资产总分", ""],
        ["能力资产总分", ""],
        ["信任资产总分", ""],
        ["负债扣减总分", ""],
        ["净资产（资产总分-负债）", ""],
    ]
    for i, (label, val) in enumerate(calc_data):
        calc_table.rows[i].cells[0].text = label
        calc_table.rows[i].cells[1].text = val
        if i == 5:
            set_cell_shading(calc_table.rows[i].cells[0], 'D9E2F3')
            set_cell_shading(calc_table.rows[i].cells[1], 'D9E2F3')
            calc_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            calc_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档4: 风险评估决策表 ==========
def create_risk_decision_table():
    """创建04-风险评估决策表.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "风险评估决策表", COURSE_NAME + " | 工具表单")

    add_formatted_para(doc, "系统评估当前职业状态的沉没成本与决策点：", size=11)

    # 评估表格
    add_formatted_para(doc, "一、风险评估", bold=True, size=12)

    eval_table = doc.add_table(rows=5, cols=2)
    eval_table.style = 'Table Grid'
    eval_data = [
        ["沉没成本识别", "你已经投入但无法收回的：时间、资金、精力、机会成本"],
        ["沉没成本评估", "这些投入的总成本是多少？继续下去的额外成本是多少？"],
        ["继续预期", "如果保持现状，3个月/6个月/1年后会怎样？"],
        ["重新开始评估", "重启需要什么？时间和资源成本？成功的可能性？"],
        ["止损点设定", "什么信号出现时，你应该立即止损而非继续坚持？"],
    ]
    for i, (label, desc) in enumerate(eval_data):
        eval_table.rows[i].cells[0].text = label
        eval_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(eval_table.rows[i].cells[0], 'FFF2CC')
        eval_table.rows[i].cells[1].text = desc

    doc.add_paragraph()

    # 决策类型选择
    add_formatted_para(doc, "二、决策类型（请选择）", bold=True, size=12)

    decision_types = [
        ["□", "继续坚持", "在现有轨道上继续，但调整策略和方法"],
        ["□", "止损退出", "承认沉没成本，有序退出寻找新方向"],
        ["□", "平行探索", "保持现有工作的同时，探索新可能性"],
        ["□", "全面转型", "全身心投入新方向，重新开始"],
    ]

    dt_table = doc.add_table(rows=1, cols=3)
    dt_table.style = 'Table Grid'
    hdr = dt_table.rows[0].cells
    hdr[0].text = "选择"
    hdr[1].text = "类型"
    hdr[2].text = "说明"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    for row_data in decision_types:
        row = dt_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()

    # 决策内容
    add_formatted_para(doc, "三、决策内容填写", bold=True, size=12)

    content_table = doc.add_table(rows=1, cols=1)
    content_table.style = 'Table Grid'
    cell = content_table.rows[0].cells[0]
    cell.text = """我的决策类型：________________

我的决策理由：________________

我的止损时间点：________________

我的行动计划：________________

第一步行动：________________"""
    set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档5: 个人品牌审计表 ==========
def create_personal_brand_audit():
    """创建05-个人品牌审计表.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "个人品牌审计表", COURSE_NAME + " | 工具表单")

    add_formatted_para(doc, "定期审计你的个人品牌建设情况：", size=11)

    # 审计表格
    add_formatted_para(doc, "一、品牌审计维度", bold=True, size=12)

    audit_table = doc.add_table(rows=6, cols=4)
    audit_table.style = 'Table Grid'
    headers = ["审计维度", "当前状态", "问题诊断", "改进方向"]
    hdr = audit_table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'D9E2F3')

    audit_dims = [
        ["品牌认知", "他人对你的印象是什么？", "", ""],
        ["差异化", "你的独特卖点是否清晰？", "", ""],
        ["一致性", "你在所有渠道的形象是否统一？", "", ""],
        ["传播", "你是否主动传播个人品牌？", "", ""],
        ["品牌资产", "你积累了哪些品牌资产？", "", ""],
    ]
    for i, row_data in enumerate(audit_dims):
        for j, val in enumerate(row_data):
            audit_table.rows[i+1].cells[j].text = val

    doc.add_paragraph()

    # 行动计划表格
    add_formatted_para(doc, "二、行动计划", bold=True, size=12)

    action_table = doc.add_table(rows=4, cols=4)
    action_table.style = 'Table Grid'
    action_headers = ["序号", "行动项", "具体措施", "时间"]
    hdr = action_table.rows[0].cells
    for i, h in enumerate(action_headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'E2EFDA')

    for i in range(1, 4):
        action_table.rows[i].cells[0].text = str(i)
        set_cell_shading(action_table.rows[i].cells[0], 'F2F2F2')

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档6: 90天行动计划表 ==========
def create_90day_action_plan():
    """创建06-90天行动计划表.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "90天行动计划表", COURSE_NAME + " | 工具表单")

    # Objective
    add_formatted_para(doc, "Objective（目标）", bold=True, size=12)

    obj_table = doc.add_table(rows=1, cols=1)
    obj_table.style = 'Table Grid'
    cell = obj_table.rows[0].cells[0]
    cell.text = "在90天内，我要实现的目标是：\n\n"
    set_cell_shading(cell, 'FFF2CC')

    doc.add_paragraph()

    # Key Results
    add_formatted_para(doc, "Key Results（关键结果）", bold=True, size=12)

    kr_table = doc.add_table(rows=4, cols=3)
    kr_table.style = 'Table Grid'
    kr_headers = ["KR", "关键结果", "完成标准"]
    hdr = kr_table.rows[0].cells
    for i, h in enumerate(kr_headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'D9E2F3')

    for i in range(1, 4):
        kr_table.rows[i].cells[0].text = "KR" + str(i)
        set_cell_shading(kr_table.rows[i].cells[0], 'F2F2F2')

    doc.add_paragraph()

    # 12周里程碑
    add_formatted_para(doc, "12周里程碑", bold=True, size=12)

    milestone_table = doc.add_table(rows=13, cols=3)
    milestone_table.style = 'Table Grid'
    ms_headers = ["周次", "里程碑目标", "完成情况"]
    hdr = milestone_table.rows[0].cells
    for i, h in enumerate(ms_headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'E2EFDA')

    for i in range(1, 13):
        milestone_table.rows[i].cells[0].text = "第" + str(i) + "周"
        set_cell_shading(milestone_table.rows[i].cells[0], 'F2F2F2')

    doc.add_paragraph()
    add_copyright(doc)

    return doc

# ========== 文档7: 案例集 ==========
def create_case_study_collection():
    """创建案例集.docx"""
    doc = Document()
    add_title_and_subtitle(doc, "案例集", COURSE_NAME + " | 案例集")

    cases = [
        {
            "title": "案例一：林晓的「被经营」困局（模块一）",
            "background": "林晓，32岁，在一家大型企业工作8年，担任部门经理。她每天工作10+小时，周末也经常加班，但始终觉得自己只是在完成别人的期望，从未真正掌控自己的职业生涯。",
            "conflict": "核心冲突：林晓发现自己在公司里承担了大部分重要项目，但升职加薪总是轮不到她。她开始意识到，自己一直在「被经营」——被领导的指令驱动，被KPI驱动，被deadline驱动，却从未停下来思考自己真正想要什么。",
            "insight": "案例启示：很多人像林晓一样，陷入了「勤劳的陷阱」——越努力工作，却越失去对人生的掌控权。关键不是更努力，而是转变思维：从「被经营者」变成「自我经营者」。当你开始用CEO的思维看待自己时，才能真正掌握主动权。",
        },
        {
            "title": "案例二：林晓的「个人定位」探索（模块二）",
            "background": "在职业顾问的引导下，林晓开始探索自己的个人定位。她回顾了过去8年的工作经历，试图找出自己的核心价值和独特优势。",
            "conflict": "核心冲突：在梳理过程中，林晓发现自己有很多「标签」——擅长沟通、项目管理、数据分析等，但这些标签组合在一起，却没有形成清晰的个人定位。她陷入了一个尴尬的境地：什么都懂一点，但什么都不精通，不知道该如何定义自己。",
            "insight": "案例启示：个人定位不是找「最擅长什么」，而是找「最独特的组合」。当你能够用一句话说清楚你是谁、你能为谁解决什么问题时，你的定位才算真正建立。林晓后来意识到，她的独特价值在于「既懂业务又懂数据」这个跨界组合。",
        },
        {
            "title": "案例三：林晓的「资产负债表」盘点（模块三）",
            "background": "林晓使用「个人资产负债表」工具对自己进行了全面盘点。她惊讶地发现，自己一直在「负债经营」——无效社交消耗了大量时间，消极的工作关系在悄悄侵蚀她的能量，而她却从未意识到这些「隐形的负债」。",
            "conflict": "核心冲突：在资产盘点时，林晓发现自己的时间资产严重不足——每天可控的专注工作时间不到3小时。人脉资产也让她震惊：虽然认识很多人，但关键时刻能真正求助的人不超过5个。更让她警醒的是，她发现自己几乎没有可迁移的个人品牌资产。",
            "insight": "案例启示：大多数人都高估了自己的「资产」，低估了自己的「负债」。个人资产负债表的价值在于让你用财务的视角审视自己的人生经营状况。只有清楚自己拥有什么、欠着什么，才能制定有效的「扭亏为盈」策略。",
        },
        {
            "title": "案例四：林晓的「90天行动计划」（模块四）",
            "background": "基于前三个模块的探索，林晓制定了她的90天行动计划。她明确了目标：90天后，她要拿到一个能够体现她个人价值的项目主导权。",
            "conflict": "核心冲突：制定计划容易，执行计划难。林晓在第一周就遇到了挑战：原有的工作节奏太难打破，总是被紧急的事情打断。后来她意识到，90天计划不是「待办清单」，而是「战略路线图」——需要每周复盘、动态调整。",
            "insight": "案例启示：90天行动计划的核心不是计划本身，而是「以终为始」的思维模式。每个90天都是一个战略周期，它强迫你思考：什么是最重要的？我要把有限的时间投入到哪里？当林晓学会用90天为单位规划人生时，她发现自己对时间的掌控感强了很多。",
        },
    ]

    for case in cases:
        add_formatted_para(doc, case["title"], bold=True, size=13)

        add_formatted_para(doc, "【案例背景】", bold=True, size=11)
        add_formatted_para(doc, case["background"], size=10)

        add_formatted_para(doc, "【核心冲突】", bold=True, size=11)
        add_formatted_para(doc, case["conflict"], size=10)

        add_formatted_para(doc, "【案例启示】", bold=True, size=11)
        add_formatted_para(doc, case["insight"], size=10)

        doc.add_paragraph()

    add_copyright(doc)

    return doc

# ========== 主程序 ==========
if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')

    base_tools = OUTPUT_BASE + "/08-工具集锦"
    base_cases = OUTPUT_BASE + "/07-案例集"

    docs = [
        (create_diagnosis_table(), base_tools + "/01-个人经营诊断表.docx"),
        (create_value_positioning(), base_tools + "/02-价值定位图.docx"),
        (create_personal_balance_sheet(), base_tools + "/03-个人资产负债表.docx"),
        (create_risk_decision_table(), base_tools + "/04-风险评估决策表.docx"),
        (create_personal_brand_audit(), base_tools + "/05-个人品牌审计表.docx"),
        (create_90day_action_plan(), base_tools + "/06-90天行动计划表.docx"),
        (create_case_study_collection(), base_cases + "/案例集.docx"),
    ]

    for doc, path in docs:
        doc.save(path)
        print("已创建: " + path)

    print("\n全部文档创建完成！")
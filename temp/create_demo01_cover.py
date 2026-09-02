# -*- coding: utf-8 -*-
"""
生成《破题力》课程成果demo - 封面文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_comment_to_paragraph(para, comment_text, author="课程助教", initials="助"):
    """为段落添加批注"""
    # 由于python-docx不直接支持批注，需要操作底层XML
    # 这里先在段落末尾添加批注标记文本
    run = para.add_run(f"[{author}: {comment_text}]")
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.size = Pt(9)
    run.font.italic = True
    return run

def create_cover_document():
    """创建封面文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4横向
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== 顶部标题区 =====
    # 课程名称标签
    title_label = doc.add_paragraph()
    title_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_label.add_run("━ 行动学习项目成果 ━")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "微软雅黑"

    # 主标题
    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = main_title.add_run("《破题力》")
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("行动学习者的四维问题定义训练营")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.font.name = "微软雅黑"
    subtitle.paragraph_format.space_after = Pt(30)

    # ===== 成果展示标题 =====
    result_title = doc.add_paragraph()
    result_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = result_title.add_run("学员成果 Demo 展示")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.name = "微软雅黑"

    # 分隔线
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 60)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    line.paragraph_format.space_after = Pt(20)

    # ===== 学员信息表格 =====
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表格宽度
    for row in info_table.rows:
        for cell in row.cells:
            cell.width = Inches(2.3)

    # 第一行：学员信息
    info_data = [
        ["学员姓名", "________________", "所属团队", "________________"],
        ["指导教练", "________________", "完成日期", "________________"],
        ["课题名称", "________________", "", ""]
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0] if para.runs else para.add_run(text)
            if j in [0, 2]:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            run.font.size = Pt(11)
            run.font.name = "微软雅黑"

    # 合并最后一行
    info_table.cell(2, 1).merge(info_table.cell(2, 3))
    info_table.cell(2, 1).text = "________________"

    # ===== 成果清单标题 =====
    list_title = doc.add_paragraph()
    list_title.paragraph_format.space_before = Pt(30)
    list_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = list_title.add_run("📋 成果文件清单")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    # ===== 成果清单表格 =====
    demo_list = [
        ["demo01", "课题三层定义表", "demo07", "因果归因分析示例"],
        ["demo02", "描述性定义练习示例", "demo08", "竞争性假说复盘表示例"],
        ["demo03", "隐藏考题诊断表示例", "demo09", "调研方案自查表示例"],
        ["demo04", "类比模型识别清单示例", "demo10", "判断标准梳理表示例"],
        ["demo05", "类比三问清单示例", "demo11", "隐藏标准发现清单示例"],
        ["demo06", "反面喻体法汇报示例", "demo12", "价值观冲突定位表示例"],
    ]

    list_table = doc.add_table(rows=len(demo_list), cols=4)
    list_table.style = 'Table Grid'
    list_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(demo_list):
        row = list_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
            run = para.runs[0]
            if j in [0, 2]:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(10)
            run.font.name = "微软雅黑"

    # ===== 底部说明 =====
    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(30)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("本 Demo 展示学员完成全部课程模块后的标准成果产出\n所有内容均为真实企业场景模拟练习")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "微软雅黑"

    # 保存文档
    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo成果封面.docx"
    doc.save(output_path)
    print(f"封面文档已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    create_cover_document()

# -*- coding: utf-8 -*-
import codecs
from docx import Document

doc = Document('D:/2026年课程/新课开发demo/讲师手册demo/讲师手册_五维表达.docx')

output = []

output.append('=== PARAGRAPH STYLES ===')
styles = {}
for para in doc.paragraphs:
    s = para.style.name if para.style else 'None'
    styles[s] = styles.get(s, 0) + 1
for style, count in sorted(styles.items(), key=lambda x: -x[1]):
    output.append(f'{style}: {count}')

output.append('\n=== SPECIAL MARKERS/FORMATTING ===')
special_markers = ['⚠', '?', '【', '】', '"', '"', '·']
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if text:
        for marker in special_markers:
            if marker in text:
                output.append(f'[{i}] {text[:80]}')
                break

output.append('\n=== KEY SECTIONS (HEADING 1 & 2) ===')
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name if para.style else 'None'
    text = para.text
    if 'Heading 1' in style_name or 'Heading 2' in style_name:
        output.append(f'[{style_name}] {text}')

output.append('\n=== TABLES SUMMARY ===')
for i, table in enumerate(doc.tables):
    output.append(f'Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols')
    for j, row in enumerate(table.rows[:2]):
        cells = [cell.text[:25] for cell in row.cells]
        output.append(f'  Row {j}: {cells}')

with codecs.open('D:/CC/temp/output.txt', 'w', 'utf-8') as f:
    f.write('\n'.join(output))

print('Done - output saved to D:/CC/temp/output.txt')
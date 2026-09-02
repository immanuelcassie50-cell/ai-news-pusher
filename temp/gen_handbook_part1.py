#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_chapter_content(doc, chapter_data):
    h = doc.add_paragraph()
    run = h.add_run(chapter_data["title"])
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    quote_p = doc.add_paragraph()
    quote_p.paragraph_format.left_indent = Cm(1)
    run = quote_p.add_run(f"金句：{chapter_data['gold_quote']}")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 0, 0)
    
    obj_h = doc.add_paragraph()
    run = obj_h.add_run("章节学习目标")
    run.bold = True
    for obj in chapter_data["objectives"]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
    
    concept_h = doc.add_paragraph()
    run = concept_h.add_run("核心概念")
    run.bold = True
    for term, definition in chapter_data["concepts"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"{term}：")
        run.bold = True
        p.add_run(definition)
    
    key_h = doc.add_paragraph()
    run = key_h.add_run("关键要点")
    run.bold = True
    for j, point in enumerate(chapter_data["key_points"], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f"{j}. ")
        run.bold = True
        p.add_run(point)
    
    think_h = doc.add_paragraph()
    run = think_h.add_run("思考题")
    run.bold = True
    for t in chapter_data["thinking"]:
        p = doc.add_paragraph(t, style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
    
    ex_h = doc.add_paragraph()
    run = ex_h.add_run("练习题")
    run.bold = True
    for e in chapter_data["exercise"]:
        p = doc.add_paragraph(e, style="List Bullet")
        p.paragraph_format.left_indent = Cm(1)
    
    sum_h = doc.add_paragraph()
    run = sum_h.add_run("章节小结")
    run.bold = True
    sum_p = doc.add_paragraph()
    sum_p.paragraph_format.left_indent = Cm(1)
    run = sum_p.add_run(chapter_data["summary"])
    run.italic = True

def create_student_handbook():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    
    # 封面
    for _ in range(8):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代决策工作手册")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("把个人判断变成组织可复用的决策资产")
    run.font.size = Pt(16)
    run.italic = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    aud = doc.add_paragraph()
    aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = aud.add_run("学员对象")
    run.bold = True
    run.font.size = Pt(14)
    
    aud2 = doc.add_paragraph()
    aud2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = aud2.add_run("企业中层管理者、项目负责人、HRBP、培训经理、行动学习催化师")
    run.font.size = Pt(12)
    
    for _ in range(6):
        doc.add_paragraph()
    
    cr = doc.add_paragraph()
    cr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cr.add_run("© 罗宏伟")
    run.bold = True
    run.font.size = Pt(14)
    
    cr2 = doc.add_paragraph()
    cr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cr2.add_run("版权所有 侵权必究")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    # 目录
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目 录")
    run.bold = True
    run.font.size = Pt(24)
    doc.add_paragraph()
    
    toc_items = [
        ("学习目标", "4"),
        ("课前准备", "5"),
        ("PART 1  认知升级：复盘与决策卡的本质区分", "6"),
        ("  第一章  复盘写的是过去，决策卡写的是未来", "6"),
        ("  第二章  不是每个决策都值得做成一张卡", "7"),
        ("  第三章  复盘访谈问的不是你做了什么，是你当时不确定什么", "8"),
        ("  第四章  隐性判断是问出来的，不是想出来的", "9"),
        ("  第五章  决策卡不是流程图，是一份检查表加一个开关", "10"),
        ("  第六章  触发条件写不清楚，这张卡就是废纸", "11"),
        ("  第七章  场景映射矩阵：一张卡管不住所有场景", "12"),
        ("  第八章  失败案例不是用来吓人的，是用来对照的", "13"),
        ("  第九章  你写的决策卡，第一个读者应该是反对你的人", "14"),
        ("  第十章  训练活动的目的不是讲透道理，是让人在场景里犯一次错", "15"),
        ("PART 1  Q&A", "16"),
        ("PART 2  组织落地：从个人经验到组织能力", "17"),
        ("  第一章  决策卡最大的敌人不是会不会用，是没人认领", "17"),
        ("  第二章  决策稽核不是查错，是替组织记住它学过什么", "18"),
        ("  第三章  好的决策卡会过时，这是它活着的证据", "19"),
        ("  第四章  复盘引导者最容易犯的错，是替决策者总结", "20"),
        ("  第五章  这份工作最终交付的不是文档，是组织的判断力", "21"),
        ("PART 2  Q&A", "22"),
        ("工具模板", "23"),
        ("  附录一  决策卡标准模板", "23"),
        ("  附录二  稽核表模板", "25"),
        ("  附录三  访谈提纲模板", "27"),
        ("课后资源", "30"),
        ("  延伸阅读清单", "30"),
        ("  实践作业", "30"),
        ("  90天行动计划模板", "31"),
        ("附录", "32"),
        ("  附录A  术语表", "32"),
        ("  附录B  金句合集", "33"),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run("	" + page)
    
    doc.add_page_break()
    
    # 学习目标
    obj_t = doc.add_paragraph()
    run = obj_t.add_run("学习目标")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()
    
    for label, content in [
        ("知识目标", "理解复盘与决策卡的本质区分"),
        ("技能目标", "掌握决策卡制作的全流程方法"),
        ("态度目标", "建立判断力是组织核心资产的意识"),
    ]:
        obj_h = doc.add_paragraph()
        run = obj_h.add_run(label)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph(content)
    
    doc.add_page_break()
    
    # 课前准备
    prep = doc.add_paragraph()
    run = prep.add_run("课前准备")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()
    
    test_t = doc.add_paragraph()
    run = test_t.add_run("课前自测问卷")
    run.bold = True
    run.font.size = Pt(14)
    
    for q in ["你过去做过的复盘，现在还能想起哪些具体内容？",
              "你曾经是否遇到过类似场景：复盘做了很多，但下次还是踩了同样的坑？",
              "你认为复盘和决策卡最大的区别是什么？",
              "你是否曾经有过当时我就是觉得不对，但说不清楚哪里不对的经历？",
              "你觉得一个组织中，最值得沉淀的是什么东西？"]:
        doc.add_paragraph(q)
    
    doc.add_paragraph()
    think_t = doc.add_paragraph()
    run = think_t.add_run("课前思考")
    run.bold = True
    run.font.size = Pt(14)
    
    think_c = doc.add_paragraph()
    think_c.add_run("你上一次做的复盘现在还记得什么？")
    think_c.paragraph_format.left_indent = Cm(1)
    
    note = doc.add_paragraph()
    run = note.add_run("请在下方记录你的思考：")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    for _ in range(6):
        line = doc.add_paragraph()
        line.add_run("_" * 70)
        line.paragraph_format.space_after = Pt(12)
    
    doc.add_page_break()
    
    # PART 1
    p1 = doc.add_paragraph()
    run = p1.add_run("PART 1")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p1s = doc.add_paragraph()
    run = p1s.add_run("认知升级：复盘与决策卡的本质区分")
    run.bold = True
    run.font.size = Pt(16)

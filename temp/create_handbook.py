# -*- coding: utf-8 -*-
"""
创建《人机协同权责边界与决策分级》学员手册
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "D:/新课开发/生态链/2.人机协同权责边界与决策分级：AI辅助生态协同谁该拍板/学员手册/学员手册-人机协同权责边界与决策分级.docx"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """添加水平分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def create_handbook():
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4横向
    section.page_height = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ========== 封面页 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("学员手册")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("人机协同权责边界与决策分级")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI辅助生态协同谁该拍板")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("链主企业管理者实战指南")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ========== 引言部分 ==========
    # 课程全景图
    p = doc.add_heading('课程全景图', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("本课程围绕人机协同权责边界这一核心命题，分为五个部分：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)

    # 五部分关系图
    parts_table = doc.add_table(rows=6, cols=2)
    parts_table.style = 'Table Grid'
    parts_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['部分', '核心内容']
    for i, header in enumerate(headers):
        cell = parts_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

    parts_data = [
        ('第一部分', '生态协同的决策复杂性'),
        ('第二部分', '决策分级体系设计'),
        ('第三部分', '权责边界划定'),
        ('第四部分', '典型场景与人机协作'),
        ('第五部分', '实施路径与治理保障'),
    ]

    for i, (part, content) in enumerate(parts_data, 1):
        parts_table.cell(i, 0).text = part
        parts_table.cell(i, 1).text = content

    doc.add_paragraph()

    # 关键洞见
    p = doc.add_paragraph()
    run = p.add_run("【关键洞见】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    p = doc.add_paragraph()
    insight_text = (
        "在生态协同中，AI辅助决策的核心挑战不是能否做，而是谁来做、谁负责。不同决策的风险系数、"
        "涉及主体利益、时效要求各不相同，必须建立科学的决策分级体系，明确人机各自的优势领域，才能实现真正的协同增效。"
    )
    run = p.add_run(insight_text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)

    doc.add_page_break()

    # 场景卡
    p = doc.add_heading('场景卡：智能汽车生态的质量事故决策', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("【场景背景】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    scenario_lines = [
        "某智能汽车链主企业收到供应商提供的电池管理系统（BMS）软件升级通知。在评估升级可行性的过程中，涉及多个利益相关方：",
        "",
        "- 芯片供应商：提供BMS主控芯片，对升级兼容性负责",
        "- 软件供应商：提供BMS基础算法，参与升级适配",
        "- 整车厂（链主）：最终决策是否批准升级",
        "- 终端用户：受影响的车主群体",
        "",
        "升级评估涉及：技术风险（功能稳定性）、商业风险（供应商责任）、法律风险（产品责任）、用户信任（OTA升级口碑）等多个维度。",
        "",
        "在做出最终决策前，需要回答以下问题：",
        "1. 这个升级决策应该由谁来主导？",
        "2. AI系统可以承担哪些评估工作？",
        "3. 哪些决策必须保留给人类管理者？",
        "4. 如果出现事故，责任如何界定？",
    ]

    for line in scenario_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(11)

    doc.add_paragraph()
    add_horizontal_line(doc)

    # 出发点自评表
    p = doc.add_heading('出发点自评表', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("请根据自己的实际情况，选择最符合的选项。这有助于您了解学习起点，更有针对性地完成课程。")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()

    # 自评表格
    assessment_table = doc.add_table(rows=6, cols=5)
    assessment_table.style = 'Table Grid'

    headers = ['评估维度', '1-完全不符合', '2-不太符合', '3-基本符合', '4-完全符合']
    for i, header in enumerate(headers):
        cell = assessment_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    assessment_items = [
        '我清楚生态协同中AI辅助决策的主要挑战',
        '我了解不同类型决策的风险差异',
        '我能准确判断哪些决策适合AI、哪些必须人类拍板',
        '我的组织已有明确的AI决策权责划分机制',
        '我具备推动人机协同落地的实操经验',
    ]

    for i, item in enumerate(assessment_items, 1):
        assessment_table.cell(i, 0).text = item
        for j in range(1, 5):
            assessment_table.cell(i, j).text = '○'

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("得分说明：15-20分=高起点，10-14分=中等起点，5-9分=需要从基础开始")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ========== 第一章 ==========
    p = doc.add_heading('第一章：生态协同的决策复杂性', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    # 学习目标
    p = doc.add_paragraph()
    run = p.add_run("【学习目标】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    objectives = [
        "理解生态协同决策与企业内部决策的本质差异",
        "掌握多方利益博弈下的决策风险识别方法",
        "能够使用决策风险自评表评估真实场景"
    ]
    for obj in objectives:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.runs[0].font.name = 'Microsoft YaHei'
        p.runs[0].font.size = Pt(11)

    # 核心概念
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("【核心概念】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    concept_table = doc.add_table(rows=3, cols=2)
    concept_table.style = 'Table Grid'

    concept_data = [
        ('生态协同决策', '涉及多个独立主体（供应商、链主、客户）利益平衡的决策，决策结果影响链条上多个参与方的权责分配。'),
        ('企业内部决策', '在单一组织边界内，由管理层或特定岗位做出的决策，决策结果由组织内部承担。'),
    ]

    for i, (term, definition) in enumerate(concept_data):
        concept_table.cell(i, 0).text = term
        set_cell_shading(concept_table.cell(i, 0), 'E7E6E6')
        concept_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        concept_table.cell(i, 1).text = definition

    doc.add_paragraph()

    # 关键洞见
    p = doc.add_paragraph()
    run = p.add_run("【关键洞见】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    p = doc.add_paragraph()
    insight_text = (
        "生态协同中的决策风险不是概率乘以损失那么简单。当AI的建议影响多个利益主体时，"
        "最大的风险是责任真空——每个人都觉得不是自己的事。"
    )
    run = p.add_run(insight_text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    # 决策风险自评表
    p = doc.add_heading('工具：决策风险自评表', level=2)

    risk_table = doc.add_table(rows=8, cols=4)
    risk_table.style = 'Table Grid'

    risk_headers = ['风险维度', '评估要点', '风险等级（低/中/高）', '备注']
    for i, header in enumerate(risk_headers):
        cell = risk_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    risk_items = [
        ('多方利益关联度', '决策是否影响2个以上独立主体的利益', '', ''),
        ('责任归属清晰度', '出现负面结果时，责任方是否明确', '', ''),
        ('可逆性程度', '决策结果是否可以撤回或修改', '', ''),
        ('信息透明度', 'AI决策依据的信息是否充分、可解释', '', ''),
        ('法律合规性', '是否符合行业法规和产品责任要求', '', ''),
        ('声誉影响度', '决策对品牌和用户信任的影响程度', '', ''),
        ('时间紧迫性', '决策的时间窗口和容错空间', '', ''),
    ]

    for i, (dim, point, risk, note) in enumerate(risk_items, 1):
        risk_table.cell(i, 0).text = dim
        risk_table.cell(i, 1).text = point
        risk_table.cell(i, 2).text = risk
        risk_table.cell(i, 3).text = note

    doc.add_paragraph()

    # 练习
    p = doc.add_heading('练习：我的生态决策经历', level=2)

    p = doc.add_paragraph()
    run = p.add_run("请回想一次您参与的生态协同决策经历，尝试用上面的自评表进行分析：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    exercise_items = [
        "决策背景：描述这次决策的场景和涉及的利益方",
        "决策过程：AI系统提供了什么建议？最终是如何做出的决策？",
        "风险评估：对照自评表，各维度风险等级是什么？",
        "经验教训：如果重来一次，会有什么不同？"
    ]

    for item in exercise_items:
        p = doc.add_paragraph(style='List Number')
        p.text = item
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)

    doc.add_paragraph()
    # 书写区域
    for _ in range(4):
        p = doc.add_paragraph("_" * 80)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 第二章 ==========
    p = doc.add_heading('第二章：决策分级体系设计', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("【学习目标】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    objectives_ch2 = [
        "掌握L1-L5五级决策模型的核心逻辑",
        "理解风险类型、决策权重与AI能力边界的关系",
        "能够为自己的生态决策进行分级"
    ]
    for obj in objectives_ch2:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.runs[0].font.name = 'Microsoft YaHei'
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("【核心概念：L1-L5五级决策模型】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    # 决策分级表格
    level_table = doc.add_table(rows=6, cols=4)
    level_table.style = 'Table Grid'

    level_headers = ['等级', '决策类型', 'AI参与程度', '人类角色']
    for i, header in enumerate(level_headers):
        cell = level_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    level_data = [
        ('L1', '例行操作', 'AI全自动执行', '监督与审计'),
        ('L2', '标准化决策', 'AI建议+人类确认', '审核与批准'),
        ('L3', '复杂分析决策', 'AI分析+人类判断', '主导决策'),
        ('L4', '战略性决策', 'AI辅助+人类主导', '最终决策者'),
        ('L5', '高风险/高价值', 'AI参考+人类全权', '全权负责'),
    ]

    for i, (level, dtype, ai_role, human_role) in enumerate(level_data, 1):
        level_table.cell(i, 0).text = level
        level_table.cell(i, 1).text = dtype
        level_table.cell(i, 2).text = ai_role
        level_table.cell(i, 3).text = human_role

    doc.add_paragraph()

    # 关键洞见
    p = doc.add_paragraph()
    run = p.add_run("【关键洞见】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    p = doc.add_paragraph()
    insight_text = (
        "决策分级的核心不是AI能不能做，而是AI做的结果谁来负责。L3以上的决策，因为涉及多方利益和法律责任，"
        "必须确保人类有真实的判断能力而非形式化审批。"
    )
    run = p.add_run(insight_text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    # 决策分级矩阵
    p = doc.add_heading('工具：决策分级矩阵', level=2)

    matrix_table = doc.add_table(rows=5, cols=5)
    matrix_table.style = 'Table Grid'

    matrix_headers = ['维度', 'L1', 'L2', 'L3', 'L4-L5']
    for i, header in enumerate(matrix_headers):
        cell = matrix_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    matrix_data = [
        ('信息完备性', '完整确定', '基本完整', '部分缺失', '高度不确定'),
        ('结果可逆性', '完全可逆', '较易恢复', '难以恢复', '不可逆'),
        ('利益关联度', '单一主体', '少数相关', '多方博弈', '全链条影响'),
        ('法律风险', '无责任', '有限责任', '明确责任', '重大责任'),
    ]

    for i, row_data in enumerate(matrix_data, 1):
        for j, cell_text in enumerate(row_data):
            matrix_table.cell(i, j).text = cell_text

    doc.add_paragraph()

    # 练习
    p = doc.add_heading('练习：为我的生态决策分级', level=2)

    p = doc.add_paragraph()
    run = p.add_run("请选择您组织中的一个典型生态协同决策，对照分级标准确定其等级：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    exercise_table = doc.add_table(rows=5, cols=2)
    exercise_table.style = 'Table Grid'

    exercise_items = [
        ('决策名称', ''),
        ('涉及利益方', ''),
        ('对照L1-L5标准评估', ''),
        ('初步分级结果', ''),
        ('分级理由', ''),
    ]

    for i, (label, value) in enumerate(exercise_items):
        exercise_table.cell(i, 0).text = label
        set_cell_shading(exercise_table.cell(i, 0), 'E7E6E6')
        exercise_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        exercise_table.cell(i, 1).text = value

    doc.add_page_break()

    # ========== 第三章 ==========
    p = doc.add_heading('第三章：权责边界划定', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("【学习目标】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    objectives_ch3 = [
        "明确AI可承担决策与人类必须保留决策的边界",
        "掌握法律责任边界的判断标准",
        "能够识别和规避权责真空风险"
    ]
    for obj in objectives_ch3:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.runs[0].font.name = 'Microsoft YaHei'
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # AI可承担 vs 人类必须保留
    p = doc.add_paragraph()
    run = p.add_run("【核心概念：AI可承担 vs 人类必须保留】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    boundary_table = doc.add_table(rows=2, cols=2)
    boundary_table.style = 'Table Grid'

    # 表头
    cell = boundary_table.cell(0, 0)
    cell.text = 'AI可承担'
    set_cell_shading(cell, '70AD47')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    cell = boundary_table.cell(0, 1)
    cell.text = '人类必须保留'
    set_cell_shading(cell, 'C00000')
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    ai_items = [
        '数据收集与整理',
        '模式识别与异常检测',
        '标准化流程优化',
        '风险量化与评估',
        '预案匹配与推荐'
    ]
    human_items = [
        '多方利益协调与平衡',
        '法律责任判定与承担',
        '商业价值判断与取舍',
        '危机情境下的非常规决策',
        '伦理与声誉相关判断'
    ]

    for i, item in enumerate(ai_items):
        p = boundary_table.cell(1, 0).paragraphs[0] if i == 0 else boundary_table.cell(1, 0).add_paragraph()
        p.text = "- " + item
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)

    for i, item in enumerate(human_items):
        p = boundary_table.cell(1, 1).paragraphs[0] if i == 0 else boundary_table.cell(1, 1).add_paragraph()
        p.text = "- " + item
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)

    doc.add_paragraph()

    # 关键洞见
    p = doc.add_paragraph()
    run = p.add_run("【关键洞见】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    p = doc.add_paragraph()
    insight_text = (
        "法律责任边界的核心判断标准是可追溯性——如果一个决策导致损害，能够清晰地追溯是谁做出的判断、"
        "基于什么信息、走了什么流程。AI可以提供分析，但法律主体必须是明确的自然人或法人。"
    )
    run = p.add_run(insight_text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    # 权责边界清单
    p = doc.add_heading('工具：权责边界清单', level=2)

    checklist_table = doc.add_table(rows=9, cols=3)
    checklist_table.style = 'Table Grid'

    checklist_headers = ['边界场景', '判断标准', '责任归属']
    for i, header in enumerate(checklist_headers):
        cell = checklist_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    checklist_items = [
        ('产品缺陷责任', '是否在AI建议范围内有明确的人类确认签字', ''),
        ('数据隐私泄露', '数据收集是否获得授权、存储是否符合规定', ''),
        ('供应商选择', '选择依据是否充分、审批流程是否完整', ''),
        ('价格调整决策', '调价权限是否在授权范围内、是否告知相关方', ''),
        ('违约处理', '处理方式是否符合合同约定、是否有裁量余地', ''),
        ('危机公关决策', '决策是否经过评估影响范围、是否在授权内', ''),
        ('技术路线选择', '选择依据是否有文档记录、是否经过评审', ''),
        ('战略合作伙伴', '引入前是否做尽职调查、决策是否有据可查', ''),
    ]

    for i, (scene, criteria, resp) in enumerate(checklist_items, 1):
        checklist_table.cell(i, 0).text = scene
        checklist_table.cell(i, 1).text = criteria
        checklist_table.cell(i, 2).text = resp

    doc.add_paragraph()

    # 练习
    p = doc.add_heading('练习：边界判断实战', level=2)

    p = doc.add_paragraph()
    run = p.add_run("请判断以下场景属于AI可承担还是人类必须保留，并说明理由：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    practice_scenarios = [
        '场景1：供应商季度绩效评分（AI系统根据数据自动评分）',
        '场景2：重要合作关系的终止决定（涉及多年合作积累）',
        '场景3：产品质量问题的责任归属判定',
        '场景4：常规库存补货建议（基于历史销售数据）',
        '场景5：危机情况下的紧急采购授权',
    ]

    for scenario in practice_scenarios:
        p = doc.add_paragraph(style='List Number')
        p.text = scenario
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)
        p = doc.add_paragraph('判断：______  理由：_________________________')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 第四章 ==========
    p = doc.add_heading('第四章：典型场景与人机协作', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("【学习目标】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    objectives_ch4 = [
        "掌握5个典型场景的人机协作模式",
        "能够根据场景特征选择合适的协作模式",
        "建立场景匹配的实战能力"
    ]
    for obj in objectives_ch4:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.runs[0].font.name = 'Microsoft YaHei'
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # 五个典型场景
    scenarios = [
        {
            'name': '场景一：供应商准入评估',
            'description': '新供应商加入生态前的综合评估',
            'ai_role': '数据收集、资质比对、风险初筛',
            'human_role': '实地考察、商业判断、最终批准',
            'mode': 'AI辅助-L3',
            'key_point': 'AI负责标准化评估，人类负责价值判断'
        },
        {
            'name': '场景二：质量事故责任判定',
            'description': '产品出现质量问题时的责任归属',
            'ai_role': '数据溯源、异常检测、模式匹配',
            'human_role': '法律判断、责任协商、处置决策',
            'mode': '人类主导-L4',
            'key_point': '法律责任必须由人类承担'
        },
        {
            'name': '场景三：供应链风险预警',
            'description': '对潜在供应风险进行预警和响应',
            'ai_role': '多源数据监控、风险识别、预警推送',
            'human_role': '风险评估、响应策略、资源调配',
            'mode': '人机协作-L3',
            'key_point': 'AI发现问题，人类决策响应'
        },
        {
            'name': '场景四：生态伙伴绩效评价',
            'description': '对生态伙伴进行周期性综合评价',
            'ai_role': '数据分析、指标计算、报告生成',
            'human_role': '综合评定、反馈沟通、改进要求',
            'mode': 'AI辅助-L2',
            'key_point': 'AI处理数据，人类做判断'
        },
        {
            'name': '场景五：战略合作谈判',
            'description': '重要战略合作机会的评估与谈判',
            'ai_role': '对方分析、方案模拟、风险预测',
            'human_role': '商业判断、价值创造、最终签约',
            'mode': '人类主导-L5',
            'key_point': '高价值决策必须人类全权负责'
        },
    ]

    for scenario in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(scenario['name'])
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(68, 114, 196)

        scenario_table = doc.add_table(rows=5, cols=2)
        scenario_table.style = 'Table Grid'

        rows = [
            ('场景描述', scenario['description']),
            ('AI角色', scenario['ai_role']),
            ('人类角色', scenario['human_role']),
            ('协作模式', scenario['mode']),
            ('核心要点', scenario['key_point']),
        ]

        for i, (label, value) in enumerate(rows):
            scenario_table.cell(i, 0).text = label
            set_cell_shading(scenario_table.cell(i, 0), 'E7E6E6')
            scenario_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            scenario_table.cell(i, 1).text = value

        doc.add_paragraph()

    # 人机协作模式选择器
    p = doc.add_heading('工具：人机协作模式选择器', level=2)

    p = doc.add_paragraph()
    run = p.add_run("根据以下问题快速确定适合的协作模式：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    selector_table = doc.add_table(rows=7, cols=3)
    selector_table.style = 'Table Grid'

    selector_headers = ['问题', '选项', '推荐模式']
    for i, header in enumerate(selector_headers):
        cell = selector_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    selector_data = [
        ('决策是否涉及法律责任？', '是', '人类主导（L4-L5）'),
        ('', '否，继续下一步', ''),
        ('决策影响是否涉及多方利益？', '是', '人类主导（L3+）'),
        ('', '否，继续下一步', ''),
        ('决策是否有明确标准可循？', '是', 'AI辅助确认（L1-L2）'),
        ('', '否', '人机协作（L3）'),
    ]

    for i, (q, opt, mode) in enumerate(selector_data, 1):
        selector_table.cell(i, 0).text = q
        selector_table.cell(i, 1).text = opt
        selector_table.cell(i, 2).text = mode

    doc.add_paragraph()

    # 练习
    p = doc.add_heading('练习：场景匹配训练', level=2)

    p = doc.add_paragraph()
    run = p.add_run("请为以下场景选择合适的协作模式，并说明理由：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    practice_scenes = [
        '1. 经销商库存异常预警后的补货决策',
        '2. 核心供应商财务状况恶化的应对决策',
        '3. 新产品定价策略的制定',
        '4. 常规采购订单的审批',
        '5. 生态合作伙伴年度考核',
    ]

    for scene in practice_scenes:
        p = doc.add_paragraph()
        p.text = scene
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)
        p = doc.add_paragraph('模式：______  理由：_________________________')
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ========== 第五章 ==========
    p = doc.add_heading('第五章：实施路径与治理保障', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    p = doc.add_paragraph()
    run = p.add_run("【学习目标】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(68, 114, 196)

    objectives_ch5 = [
        "掌握决策分级的推行步骤",
        "理解治理机制建设的关键要素",
        "能够设计符合自身组织的推行计划"
    ]
    for obj in objectives_ch5:
        p = doc.add_paragraph(style='List Bullet')
        p.text = obj
        p.runs[0].font.name = 'Microsoft YaHei'
        p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # 推行步骤
    p = doc.add_paragraph()
    run = p.add_run("【决策分级的推行步骤】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    steps = [
        ('第一步：现状盘点', '梳理现有决策类型，识别AI辅助场景，评估决策风险'),
        ('第二步：分级设计', '根据L1-L5模型，为每类决策确定合适的分级和协作模式'),
        ('第三步：权责明确', '制定权责清单，明确每类决策的AI角色和人类角色'),
        ('第四步：流程嵌入', '将决策分级嵌入现有业务流程，确保执行落地'),
        ('第五步：持续优化', '建立反馈机制，根据执行效果持续调整优化'),
    ]

    for step, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
        run.font.bold = True
        p = doc.add_paragraph()
        p.text = desc
        p.paragraph_format.left_indent = Inches(0.5)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(10)
        p = doc.add_paragraph()

    # 治理机制建设要点
    p = doc.add_paragraph()
    run = p.add_run("【治理机制建设要点】")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 80, 77)

    governance_table = doc.add_table(rows=5, cols=2)
    governance_table.style = 'Table Grid'

    governance_headers = ['治理要素', '建设要点']
    for i, header in enumerate(governance_headers):
        cell = governance_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    governance_data = [
        ('组织保障', '明确决策分级的负责部门/岗位，建立协调机制'),
        ('制度规范', '制定权责边界清单，嵌入业务操作流程'),
        ('技术支撑', '建设决策支持系统，记录决策过程，确保可追溯'),
        ('培训推广', '开展分层培训，确保相关人员理解并执行'),
    ]

    for i, (element, point) in enumerate(governance_data, 1):
        governance_table.cell(i, 0).text = element
        governance_table.cell(i, 1).text = point

    doc.add_paragraph()

    # 治理检查清单
    p = doc.add_heading('工具：治理检查清单', level=2)

    gov_checklist = doc.add_table(rows=9, cols=2)
    gov_checklist.style = 'Table Grid'

    gov_headers = ['检查项', '是否完成']
    for i, header in enumerate(gov_headers):
        cell = gov_checklist.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    gov_items = [
        '决策分级方案已获得管理层批准',
        '各类决策的权责边界已明确并文档化',
        '决策流程已嵌入现有业务系统',
        '相关岗位人员已完成培训',
        '决策记录机制已建立',
        '定期审查机制已建立',
        '异常处理流程已明确',
        '持续优化机制已建立',
    ]

    for i, item in enumerate(gov_items, 1):
        gov_checklist.cell(i, 0).text = item
        gov_checklist.cell(i, 1).text = '○ 是  ○ 否  ○ 部分'

    doc.add_paragraph()

    # 练习
    p = doc.add_heading('练习：设计我的推行计划', level=2)

    p = doc.add_paragraph()
    run = p.add_run("基于以上内容，制定您组织的决策分级推行计划：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    plan_table = doc.add_table(rows=6, cols=3)
    plan_table.style = 'Table Grid'

    plan_headers = ['阶段', '主要任务', '时间节点']
    for i, header in enumerate(plan_headers):
        cell = plan_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    plan_stages = [
        ('现状盘点', '', ''),
        ('分级设计', '', ''),
        ('权责明确', '', ''),
        ('流程嵌入', '', ''),
        ('上线运行', '', ''),
    ]

    for i, (stage, task, time) in enumerate(plan_stages, 1):
        plan_table.cell(i, 0).text = stage
        plan_table.cell(i, 1).text = task
        plan_table.cell(i, 2).text = time

    doc.add_page_break()

    # ========== 课程收尾 ==========
    p = doc.add_heading('课程收尾', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    # 七习惯重测自评
    p = doc.add_heading('七习惯重测自评', level=2)

    p = doc.add_paragraph()
    run = p.add_run("请重新完成以下评估，对比学习前的结果，评估您的成长：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    retest_table = doc.add_table(rows=6, cols=5)
    retest_table.style = 'Table Grid'

    headers = ['评估维度', '学习前得分', '学习后得分', '变化', '持续改进方向']
    for i, header in enumerate(headers):
        cell = retest_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    retest_items = [
        '我清楚生态协同中AI辅助决策的主要挑战',
        '我了解不同类型决策的风险差异',
        '我能准确判断哪些决策适合AI、哪些必须人类拍板',
        '我的组织已有明确的AI决策权责划分机制',
        '我具备推动人机协同落地的实操经验',
    ]

    for i, item in enumerate(retest_items, 1):
        retest_table.cell(i, 0).text = item
        retest_table.cell(i, 1).text = ''
        retest_table.cell(i, 2).text = ''
        retest_table.cell(i, 3).text = ''
        retest_table.cell(i, 4).text = ''

    doc.add_paragraph()

    # 30天行动计划
    p = doc.add_heading('30天行动计划', level=2)

    p = doc.add_paragraph()
    run = p.add_run("将您的学习成果转化为具体行动：")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.italic = True

    action_table = doc.add_table(rows=5, cols=4)
    action_table.style = 'Table Grid'

    action_headers = ['阶段', '具体行动', '责任人', '完成时间']
    for i, header in enumerate(action_headers):
        cell = action_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(9)

    action_stages = [
        ('第1-7天', '', '', ''),
        ('第8-14天', '', '', ''),
        ('第15-21天', '', '', ''),
        ('第22-30天', '', '', ''),
    ]

    for i, (stage, action, owner, time) in enumerate(action_stages, 1):
        action_table.cell(i, 0).text = stage
        action_table.cell(i, 1).text = action
        action_table.cell(i, 2).text = owner
        action_table.cell(i, 3).text = time

    doc.add_page_break()

    # ========== 附录 ==========
    p = doc.add_heading('附录', level=1)
    p.runs[0].font.color.rgb = RGBColor(31, 56, 100)

    # 术语速查表
    p = doc.add_heading('附录A：术语速查表', level=2)

    term_table = doc.add_table(rows=11, cols=2)
    term_table.style = 'Table Grid'

    term_headers = ['术语', '定义']
    for i, header in enumerate(term_headers):
        cell = term_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    terms = [
        ('生态协同', '多个独立主体（供应商、链主、客户等）通过协作创造价值的商业模式'),
        ('AI辅助决策', '利用人工智能技术提供数据分析、建议或自动化操作，由人类保持最终决策权'),
        ('决策分级', '根据决策的风险程度、利益关联度等维度，将决策划分为不同级别并匹配不同的人机协作模式'),
        ('权责边界', '明确AI系统和人类各自承担的责任范围和决策权限'),
        ('L1-L5模型', '从完全自动化到人类全权负责的五级决策分级模型'),
        ('链主企业', '在生态协同中占据核心地位、主导生态发展方向的企业'),
        ('责任真空', '多方参与决策但无人真正负责的风险状态'),
        ('可追溯性', '决策过程和依据能够被记录、追溯的特性'),
        ('治理机制', '为保障决策分级落地而建立的组织、制度、技术体系'),
        ('人机协作', '人类与AI系统在决策过程中各自发挥优势、协同工作的模式'),
    ]

    for i, (term, definition) in enumerate(terms, 1):
        term_table.cell(i, 0).text = term
        term_table.cell(i, 1).text = definition

    doc.add_paragraph()

    # 工具速查索引
    p = doc.add_heading('附录B：工具速查索引', level=2)

    tool_table = doc.add_table(rows=8, cols=3)
    tool_table.style = 'Table Grid'

    tool_headers = ['工具名称', '所在章节', '用途']
    for i, header in enumerate(tool_headers):
        cell = tool_table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    tools = [
        ('决策风险自评表', '第一章', '评估生态协同决策的多维风险'),
        ('决策分级矩阵', '第二章', '根据多维度特征确定决策等级'),
        ('权责边界清单', '第三章', '明确AI与人类的责任边界'),
        ('人机协作模式选择器', '第四章', '快速确定场景适用的协作模式'),
        ('治理检查清单', '第五章', '检查治理机制建设的完备性'),
        ('出发点自评表', '引言', '评估学习前的起点状态'),
        ('30天行动计划表', '课程收尾', '规划学习后的具体行动'),
    ]

    for i, (name, chapter, use) in enumerate(tools, 1):
        tool_table.cell(i, 0).text = name
        tool_table.cell(i, 1).text = chapter
        tool_table.cell(i, 2).text = use

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"学员手册已生成：{OUTPUT_PATH}")

if __name__ == "__main__":
    create_handbook()

#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "D:/新课开发/变革管理/08-向上管理与高层说服术：让决策层理解容错的成本逻辑/完整课程包/08-案例集/案例集-向上管理与高层说服术.docx"

C_PRIMARY = RGBColor(0x2b, 0x2d, 0x42)
C_SECONDARY = RGBColor(0x8d, 0x99, 0xae)
C_ACCENT = RGBColor(0xef, 0x23, 0x3c)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_colored_heading(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def docx_main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = C_PRIMARY

    # Cover
    doc.add_paragraph()
    title = doc.add_heading('案例集', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(36)
        run.font.color.rgb = C_PRIMARY

    subtitle = doc.add_heading('向上管理与高层说服术', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = C_ACCENT

    doc.add_paragraph()

    # Case 1
    add_colored_heading(doc, '案例一：ERP升级提案被拒', 1, C_PRIMARY)
    add_colored_heading(doc, '背景', 2, C_SECONDARY)
    doc.add_paragraph('某制造企业IT总监提议ERP系统升级，预算1200万，预计年效益350万。项目团队认为方案成熟，但CEO在评审会上以风险过高为由拒绝。')

    add_colored_heading(doc, '第一次汇报（失败案例）', 2, C_SECONDARY)
    doc.add_paragraph('李总监向张总汇报：「我们计划引入AI客服系统，可以提高效率、降低成本。预算1200万，请张总批准。」')
    doc.add_paragraph('张总的反应：「ROI是多少？回本周期多长？这1200万如果打了水漂怎么办？有没有分阶段投入的方案？」')
    doc.add_paragraph('张总最终没有批准，要求重新准备后再汇报。')

    add_colored_heading(doc, '失败原因分析', 2, C_SECONDARY)
    for cause in [
        "没有用财务语言（ROI、回本周期）来包装提案",
        "没有量化风险边界和止损机制",
        "没有提供分阶段投入的选项",
        "直接「请求许可」而非「争取授权」"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cause)

    add_colored_heading(doc, '第二次汇报（成功案例）', 2, C_SECONDARY)
    doc.add_paragraph('李总监重新准备后再次汇报...')

    add_colored_heading(doc, '关键教训', 2, C_ACCENT)
    doc.add_paragraph('好方案需要好包装。财务导向型CEO需要用数字说话，最小授权不是魄力不足，而是专业自信的体现。')

    doc.add_page_break()

    # Case 2
    add_colored_heading(doc, '案例二：组织变革被搁置', 1, C_PRIMARY)
    add_colored_heading(doc, '背景', 2, C_SECONDARY)
    doc.add_paragraph('HR总监提议组织架构调整，计划先证明效果再长期投入。CEO认为时机不对，担心影响团队稳定。')

    add_colored_heading(doc, '问题诊断', 2, C_SECONDARY)
    for cause in [
        "未识别CEO是关系导向型",
        "没有展示团队支持基础",
        "缺乏渐进式变革路径"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cause)

    add_colored_heading(doc, '改进策略', 2, C_SECONDARY)
    for strategy in [
        "先做小范围试点展示可行性",
        "收集团队支持声音和预期收益",
        "设计6个月渐进式推进路径"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(strategy)

    add_colored_heading(doc, '关键教训', 2, C_ACCENT)
    doc.add_paragraph('关系导向型CEO需要先解决人的问题。展示团队支持是建立信任的关键，渐进式变革比一步到位更易被接受。')

    doc.add_page_break()

    # Case 3
    add_colored_heading(doc, '案例三：成功说服的关键转折', 1, C_PRIMARY)
    add_colored_heading(doc, '背景', 2, C_SECONDARY)
    doc.add_paragraph('市场总监提议数字化营销系统，预算500万。CMO是战略导向型，关注行业趋势和竞争优势。前两次汇报均被质疑无法量化价值。')

    add_colored_heading(doc, '问题诊断', 2, C_SECONDARY)
    for cause in [
        "用功能性语言而非战略语言",
        "缺乏行业对标数据",
        "没有展示竞争价值"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cause)

    add_colored_heading(doc, '关键转折', 2, C_SECONDARY)
    for tactic in [
        "引入行业报告数据支撑",
        "展示竞争对手案例",
        "提供3个月试点ROI数据"
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(tactic)

    add_colored_heading(doc, '关键教训', 2, C_ACCENT)
    doc.add_paragraph('战略导向型CEO需要用行业趋势和竞争格局说话。引入第三方数据是建立可信度的好方法。')

    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    docx_main()
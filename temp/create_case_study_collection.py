# -*- coding: utf-8 -*-
"""
创建【廉政风险情景决策训练营】案例集
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "D:/新课开发/党业融合/廉政风险情景决策/完整课程包/007-案例集/案例集.docx"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {})
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    shading.append(shd)

def set_cell_text_style(cell, bold=False, font_size=12, font_name='微软雅黑'):
    """设置单元格文字样式"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(font_size)
            run.font.name = font_name
            r = run._r
            rPr = r.get_or_add_rPr()
            rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

def add_case_document(doc, case_data):
    """添加单个案例到文档"""
    # 案例标题
    doc.add_heading(case_data['title'], level=1)

    # 背景介绍
    doc.add_heading('一、背景介绍', level=2)
    doc.add_paragraph(case_data['background'])

    # 人物设定
    doc.add_heading('二、人物设定', level=2)
    for person in case_data['characters']:
        p = doc.add_paragraph()
        p.add_run(f"【{person['role']}】{person['name']}：").bold = True
        p.add_run(person['description'])
        p.paragraph_format.space_after = Pt(6)

    # 情景发展
    doc.add_heading('三、情景发展（5阶段）', level=2)
    for i, stage in enumerate(case_data['stages'], 1):
        p = doc.add_paragraph()
        p.add_run(f"第{i}阶段：{stage['title']}").bold = True
        doc.add_paragraph(stage['description'])
        if 'dialogue' in stage:
            p = doc.add_paragraph()
            p.add_run("关键对话：").bold = True
            p.add_run(stage['dialogue'])
            p.paragraph_format.left_indent = Inches(0.3)

    # 决策点
    doc.add_heading('四、决策点', level=2)
    doc.add_paragraph(case_data['decision_point'])

    # 后果推演
    doc.add_heading('五、后果推演', level=2)
    for outcome in case_data['outcomes']:
        p = doc.add_paragraph()
        p.add_run(f"【{outcome['type']}】").bold = True
        p.add_run(outcome['description'])

    # 追问要点
    doc.add_heading('六、追问要点', level=2)
    for i, question in enumerate(case_data['questions'], 1):
        doc.add_paragraph(f"{i}. {question}")

    # 讨论问题
    doc.add_heading('七、讨论问题', level=2)
    for i, discuss in enumerate(case_data['discuss'], 1):
        doc.add_paragraph(f"{i}. {discuss}")

    doc.add_page_break()

def create_guide_section(doc):
    """创建案例讨论指南"""
    doc.add_heading('案例讨论指南', level=0)
    doc.add_paragraph()

    doc.add_heading('一、使用说明', level=1)
    guide_items = [
        "本案例集共包含5个情景案例，涵盖采购/招投标、工程、财务报销、信贷审批、内部资源分配五个关键领域。",
        "每个案例均设计了5个递进式情景阶段，对应不同程度的诱惑升级链。",
        "案例仅供教学训练使用，所有人物、情节均为虚构，如有雷同纯属巧合。",
        "建议每个案例讨论时长约45-60分钟。",
    ]
    for item in guide_items:
        doc.add_paragraph(item)

    doc.add_heading('二、讨论流程', level=1)
    doc.add_paragraph("1. 阅读案例背景和人物设定（10分钟）")
    doc.add_paragraph("2. 分组讨论每个阶段应该如何应对（20分钟）")
    doc.add_paragraph("3. 汇报各组决策及理由（15分钟）")
    doc.add_paragraph("4. 讲师点评并揭示正确答案/最佳实践（10分钟）")
    doc.add_paragraph("5. 总结本案例的廉政风险点和防控措施（5分钟）")

    doc.add_heading('三、讨论规则', level=1)
    rules = [
        "保密原则：案例内容仅用于培训，不得外传。",
        "开放原则：鼓励畅所欲言，各种观点都可以表达。",
        "移情原则：设身处地理解当事人在特定情境下的处境。",
        "反思原则：结合自身岗位思考可能存在的类似风险。",
    ]
    for rule in rules:
        doc.add_paragraph(rule)

    doc.add_heading('四、引导技巧', level=1)
    tips = [
        "不要急于给出正确答案，让学员充分讨论。",
        "适时追问'为什么这样想'，挖掘深层原因。",
        "注意引导学员从制度、流程、道德多角度分析。",
        "对有争议的观点不要急于评判，鼓励继续讨论。",
        "最后一定要明确正确的做法是什么。",
    ]
    for i, tip in enumerate(tips, 1):
        doc.add_paragraph(f"{i}. {tip}")

    doc.add_heading('五、风险提示', level=1)
    doc.add_paragraph("在讨论过程中，请注意以下风险点：")
    risks = [
        "【利益冲突风险】是否存在个人利益与组织利益冲突的情况",
        "【程序违规风险】是否违反了相关规章制度或工作流程",
        "【廉洁底线风险】是否触碰了廉洁自律的红线",
        "【外部诱惑风险】面对请托、打招呼等外部压力如何应对",
    ]
    for risk in risks:
        doc.add_paragraph(risk)

def create_document():
    """创建完整文档"""
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(12)
    r = font._element
    rPr = r.get_or_add_rPr()
    rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

    # ========== 封面 ==========
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('廉政风险情景决策训练营')
    run.font.size = Pt(32)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('案例集')
    run.font.size = Pt(28)
    run.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('（仅供内部培训使用）')
    run.font.size = Pt(16)

    doc.add_paragraph()

    # 案例清单
    doc.add_heading('案例清单', level=1)
    cases_list = [
        ('案例一', '采购/招投标类', '张建国', '某集团采购部经理'),
        ('案例二', '工程类', '李明轩', '某单位基建科科长'),
        ('案例三', '财务报销类', '王秀芬', '某公司财务总监'),
        ('案例四', '信贷审批类', '陈志强', '某银行信贷部主管'),
        ('案例五', '内部资源分配类', '刘建业', '某机关后勤中心主任'),
    ]

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ['序号', '案例类型', '主人公', '职务']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9D9D9')
        set_cell_text_style(cell, bold=True)

    # 数据行
    for row_idx, (num, case_type, name, position) in enumerate(cases_list, 1):
        table.rows[row_idx].cells[0].text = str(row_idx)
        table.rows[row_idx].cells[1].text = case_type
        table.rows[row_idx].cells[2].text = name
        table.rows[row_idx].cells[3].text = position

    doc.add_page_break()

    # ========== 案例使用说明 ==========
    doc.add_heading('案例使用说明', level=1)
    doc.add_paragraph()
    usage_text = [
        "本案例集是《廉政风险情景决策训练营》的核心教学材料，共包含5个典型情景案例。",
        "每个案例均按照'背景介绍→人物设定→情景发展→决策点→后果推演→追问要点→讨论问题'的结构设计，",
        "形成完整的训练逻辑链条，帮助学员在沉浸式体验中提升廉政风险识别和应对能力。",
        "",
        "【案例特点】",
        "• 真实性：案例情节贴近实际工作场景，具有很强的代入感",
        "• 渐进性：每个案例的5个阶段呈现诱惑逐步升级的过程",
        "• 选择性：每个决策点都面临多个选择，没有标准答案",
        "• 反思性：后果推演让学员看到不同选择带来的不同结果",
        "",
        "【使用建议】",
        "• 建议在讲师引导下进行集体讨论",
        "• 每个案例可单独使用，也可系列使用",
        "• 讨论过程中讲师注意控制节奏和时间",
        "• 强调'所有案例均为虚构，仅供教学使用'",
    ]
    for text in usage_text:
        doc.add_paragraph(text)

    doc.add_page_break()

    # ========== 案例一：采购/招投标类 ==========
    case_1 = {
        'title': '案例一：采购/招投标类',
        'background': '''张建国，45岁，某大型国有企业集团采购部经理，在公司工作18年，业务能力突出，人脉广泛。近年来，其母亲患重病，妻子工资较低，儿子即将出国留学，家庭经济压力较大。

公司近期有一个金额达5000万元的设备采购项目，按照规定需要公开招标。作为采购部负责人，张建国负责招标文件的编制和投标单位的资格审查工作。

此时，张建国的大学同学、也是多年的好友李海明找到他。李海明是某设备制造商的副总经理，他告诉张建国，如果能够帮助其中标，可以给予张建国中标金额3%的好处费（150万元）。李海明表示这只是"行业惯例"，不会有什么问题。''',
        'characters': [
            {'role': '主人公', 'name': '张建国', 'description': '集团采购部经理，45岁，工作18年，业务能力强，家庭经济压力较大'},
            {'role': '请托人', 'name': '李海明', 'description': '张建国大学同学，某设备制造商副总经理，以"行业惯例"为名提出回扣'},
            {'role': '影响人', 'name': '张母', 'description': '张建国母亲，患重病需长期治疗，医疗费用高昂'},
        ],
        'stages': [
            {
                'title': '日常关怀型诱惑',
                'description': '李海明以老同学身份经常约张建国吃饭、打球，关心其母亲病情，送些土特产，表示"哥们之间的正常往来"。',
                'dialogue': '李海明："老张，伯母身体好点了吗？我托人从老家带了点儿灵芝，对提高免疫力有好处，回头给你送去。"'
            },
            {
                'title': '小恩小惠型诱惑',
                'description': '在一次聚会后，李海明悄悄给张建国一个信封，里面是2万元现金，说是"给伯母买点营养品"。张建国推辞一番后收下了。',
                'dialogue': '李海明："老张，别客气，咱俩谁跟谁啊？这点儿钱算啥，你就当帮我个忙，让我表达一下心意。"'
            },
            {
                'title': '请托办事型诱惑',
                'description': '李海明提出希望张建国在招标时"帮忙看看资质"，暗示可以"灵活掌握"。张建国犹豫后答应了，在资格审查时为李海明公司开了绿灯。',
                'dialogue': '李海明："老张，这次招标的资质要求能不能帮忙看看有没有可以通融的地方？放心，不会让你为难的。"'
            },
            {
                'title': '利益交换型诱惑',
                'description': '李海明正式提出合作意向，明确表示事成之后给予150万元好处费，并先支付了20万元"定金"。张建国看到账户里多出的数字，内心十分挣扎。',
                'dialogue': '李海明："老张，这是20万定金，事成之后还有130万。这在行业里太正常了，没人會发现的。"'
            },
            {
                'title': '威胁利诱型诱惑',
                'description': '张建国意识到风险想要退出，李海明态度大变，威胁说已经保留了之前的"合作"证据，如果退出就举报。张建国陷入困境。',
                'dialogue': '李海明："老张，你现在退出也来不及了，我可都留着证据呢。你要是不能帮我中标，那咱们就一起完蛋。"'
            }
        ],
        'decision_point': '面对老同学的巨额利益诱惑和威胁，以及家庭经济的现实压力，张建国应该如何应对？他是否应该在这次招标中为李海明公司提供帮助？',
        'outcomes': [
            {'type': '错误选择', 'description': '张建国接受利益，为李海明公司提供帮助。最终被发现查处，受开除处分并移送司法机关，家庭支离破碎。'},
            {'type': '被动接受', 'description': '张建国因担心被举报而被动配合，长期处于恐惧和压力中，最终也因其他问题被查处。'},
            {'type': '正确选择', 'description': '张建国及时向组织报告，保存证据，配合调查。虽然暂时面临压力，但最终得到组织保护，李海明因行贿被追究法律责任。'},
            {'type': '理想选择', 'description': '张建国在第一次收到信封时就果断拒绝，并及时提醒李海明不要越线。最终既维护了友情，又坚守了底线。'},
        ],
        'questions': [
            '张建国收受2万元"营养品费"是否构成违纪违法？',
            '如果你是张建国，在第一次收到信封时应该如何处理？',
            '李海明提出"行业惯例"的说法是否成立？为什么？',
            '当李海明威胁要举报时，张建国应该怎么办？',
            '从制度角度，如何预防采购环节的廉政风险？',
        ],
        'discuss': [
            '讨论：在面对老同学的人情和巨额利益时，如何平衡情感与原则？',
            '讨论：家庭经济困难是否可以是违纪违法的理由？我们应该如何对待这种"不得已"？',
            '讨论：如果发现同事或下属可能存在廉洁问题，应该如何处理？',
        ]
    }
    add_case_document(doc, case_1)

    # ========== 案例二：工程类 ==========
    case_2 = {
        'title': '案例二：工程类',
        'background': '''李明轩，38岁，某市政府机关基建科科长，负责单位办公楼重建项目，项目总预算1.2亿元。这是该市重点工程，也是李明轩工作以来负责的最大项目。

李明轩工作认真负责，业务能力较强，但性格耿直，不善于处理人际关系。单位领导王某（分管副局长）是李明轩的老上级，一直对他比较关照。

项目招标前，王某找到李明轩，暗示某家建筑公司（是其亲属开办）的施工质量好、价格合理，希望在评标时"多多关照"。王某还许诺，项目完成后会考虑李明轩的晋升。''',
        'characters': [
            {'role': '主人公', 'name': '李明轩', 'description': '基建科科长，38岁，工作认真但性格耿直，第一次负责如此重大的项目'},
            {'role': '施压人', 'name': '王某', 'description': '分管副局长，李明轩的老上级，以职务影响力干预招标'},
            {'role': '利益方', 'name': '赵总', 'description': '王某亲属，建筑公司负责人，通过王某关系获得内部信息'},
        ],
        'stages': [
            {
                'title': '隐性施压型诱惑',
                'description': '王某在非正式场合提到某家建筑公司"做得不错"，让李明轩"可以考虑"。李明轩当时没有在意，以为只是随口一说。',
                'dialogue': '王某："小李啊，最近有个叫华建的公司听说做得不错，你下次有项目可以多多关注。"'
            },
            {
                'title': '明示关照型诱惑',
                'description': '王某单独约李明轩吃饭，明确表示希望他在招标中"关照"某家公司，并暗示这关系到李明轩的前途。',
                'dialogue': '王某："小李啊，这次的项目很重要。华建公司是我多年的老朋友，你在中标上多多支持一下，以后有你好处。"'
            },
            {
                'title': '施压与利诱型诱惑',
                'description': '在项目开标前，王某再次找到李明轩，表示"组织上在考虑你的副处级，这次项目是个考验"。李明轩感到巨大压力。',
                'dialogue': '王某："小李，这次项目是组织对你的考验。你也知道，现在提拔干部不只看能力，还要看协调能力。"'
            },
            {
                'title': '要挟施压型诱惑',
                'description': '李明轩坚持原则，按正常程序招标。华建公司未中标。王某非常不满，在随后的工作中处处刁难李明轩，并在各种场合暗示李明轩"不懂事"。',
                'dialogue': '王某："小李啊，你这人有能力，但就是太死板了。不知道变通，这样下去很难有发展啊。"'
            },
            {
                'title': '后续报复型诱惑',
                'description': '半年后，单位进行人事调整，李明轩本应被提拔，但王某以"群众基础不够"为由否决了推荐。李明轩陷入迷茫和委屈中。',
                'dialogue': '人事部门："明轩，这次的推荐被否了，主要是因为领导觉得你还需要进一步考验。"'
            }
        ],
        'decision_point': '面对直接领导的施压和职业发展的威胁，李明轩应该如何应对？是坚持原则还是"识时务者为俊杰"？',
        'outcomes': [
            {'type': '错误选择', 'description': '李明轩屈从于压力，在招标中为华建公司提供帮助。最终项目出现严重质量问题，李明轩因失职渎职被处分。'},
            {'type': '消极抵抗', 'description': '李明轩表面答应但暗中抵制，结果既得罪了领导，又没有真正阻止违规行为，最终"里外不是人"。'},
            {'type': '正确选择', 'description': '李明轩坚持原则，按程序办事。在遭受不公待遇后，向上级纪委反映情况，最终得到组织核实和保护。'},
            {'type': '智慧选择', 'description': '李明轩提前向组织报备相关情况，请组织出面协调解决，既坚持了原则，又保护了自己。'},
        ],
        'questions': [
            '王某的行为是否构成违规违纪？如果构成，属于哪种类型？',
            '李明轩在第一次听到王某暗示时应该如何回应？',
            '当领导以"考验"为名施压时，应该如何识别和应对？',
            '面对职业发展受阻，应该如何正确处理？',
            '从制度角度，如何防范招标中的"打招呼"问题？',
        ],
        'discuss': [
            '讨论：当原则与上级压力冲突时，如何既坚持原则又保护自己？',
            '讨论：如何正确处理"知情不报"与"告密"的关系？',
            '讨论：如果发现领导干部存在违纪行为，应该如何举报才既能保护自己又能发挥作用？',
        ]
    }
    add_case_document(doc, case_2)

    # ========== 案例三：财务报销类 ==========
    case_3 = {
        'title': '案例三：财务报销类',
        'background': '''王秀芬，42岁，某民营上市公司财务总监，公司正在筹备上市。她工作严谨，财务业务精通，是公司老板张总最信任的下属之一。

公司为上市需要进行财务规范，但在规范过程中发现一笔300万元的"账外资金"需要处理。这笔钱是之前张总以"业务招待费"名义从公司支取但未入账的现金。张总希望王秀芬通过虚增成本的方式把这笔账"抹平"。''',
        'characters': [
            {'role': '主人公', 'name': '王秀芬', 'description': '财务总监，42岁，业务能力强，工作严谨，正在负责公司上市筹备工作'},
            {'role': '施压人', 'name': '张总', 'description': '公司老板，强势型领导，以"为公司好"的名义要求财务造假'},
            {'role': '旁观者', 'name': '小李', 'description': '财务部主管，王秀芬的下属，对这种做法有异议但不敢出声'},
        ],
        'stages': [
            {
                'title': '情感绑架型诱惑',
                'description': '张总以"公司发展大局"为由，希望王秀芬"灵活处理"这笔账目，暗示这是"为了公司好"，也是对王秀芬的"信任考验"。',
                'dialogue': '张总："秀芬啊，公司上市是大事，这点小问题你就帮帮忙解决了。我最信任的就是你，这个事只有你知道，你要是都不帮我，谁帮我？"'
            },
            {
                'title': '利益诱惑型',
                'description': '张总承诺，如果事情办成，上市后给王秀芬5%的股权激励（价值数百万元），并且"天知地知你知我知"，不会有任何风险。',
                'dialogue': '张总："这事办成了，上市后我给你5个点的股份，你看怎么样？这种事在私企太正常了，没人會查的。"'
            },
            {
                'title': '威逼施压型',
                'description': '王秀芬有所犹豫，张总态度大变，暗示"你的职位和薪酬都是我来定的"，并且提到"财务账的事，没有我点头你也出不了账"。',
                'dialogue': '张总："秀芬啊，你可要想清楚。在这个公司，谁是老板？谁给你发工资？你干的这些事，我要是较真起来，哪件不是问题？"'
            },
            {
                'title': '责任转嫁型',
                'description': '张总提供了"解决方案"：让王秀芬签一个字，说是"财务复核"就行，其他的不用她操心。张总强调"有事我负责"。',
                'dialogue': '张总："你就签个字，证明财务复核过就行。其他的你不用管，出了事我承担，我是法人代表，还能让你背锅不成？"'
            },
            {
                'title': '最后通牒型',
                'description': '张总下了最后通牒：要么帮忙处理这个账目，要么"另谋高就"。张总还提到，如果王秀芬拒绝，他会找别人来做，"反正财务部不是只有你一个人"。',
                'dialogue': '张总："我给你两天时间考虑。要么帮我解决这个问题，要么你就另找出路吧。这个公司离了你还就不转了？"'
            }
        ],
        'decision_point': '面对老板的命令和利益诱惑，以及可能失去工作的压力，王秀芬应该如何应对？她是应该坚持原则还是"人在屋檐下不得不低头"？',
        'outcomes': [
            {'type': '错误选择', 'description': '王秀芬屈从于压力，按要求做了虚假财务处理。公司上市时被审计发现，王秀芬被追究法律责任，面临牢狱之灾。'},
            {'type': '消极拖延', 'description': '王秀芬以"技术困难"为由一直拖延，结果张总找了别人来处理，王秀芬反而被以"不配合工作"为由辞退。'},
            {'type': '正确选择', 'description': '王秀芬明确拒绝，并向董事会和审计委员会报告。初期面临巨大压力，但最终得到董事会支持，张总因其他问题被调查。'},
            {'type': '智慧选择', 'description': '王秀芬提前咨询律师，保存相关证据，在拒绝的同时提出合规的解决方案，最终既坚持了原则，又没有完全激化矛盾。'},
        ],
        'questions': [
            '张总的行为是否构成犯罪？王秀芬如果配合可能涉及什么罪名？',
            '"老板让做的"是否可以作为免责理由？',
            '当领导的命令明显违法时，应该如何拒绝？',
            '在拒绝领导违法要求时，如何保护自己？',
            '从制度角度，上市公司如何防范财务造假？',
        ],
        'discuss': [
            '讨论：有人说"在中国做生意，不做点假账活不下去"，这种观点是否正确？',
            '讨论：面对失业压力和家庭责任，如何坚持原则？有没有两全其美的办法？',
            '讨论：财务人员发现领导有财务问题应该怎么办？如何既坚持原则又保护自己？',
        ]
    }
    add_case_document(doc, case_3)

    # ========== 案例四：信贷审批类 ==========
    case_4 = {
        'title': '案例四：信贷审批类',
        'background': '''陈志强，35岁，某城市商业银行信贷部主管，负责企业贷款审批。近年来银行信贷业务竞争激烈，完不成任务绩效会受严重影响。

某天，陈志强的高中同学、现在已是某房地产公司老板的周明找到他。周明的公司需要贷款1亿元用于新项目开发，但公司资产负债率偏高，按照正常审批很难通过。周明承诺，如果陈志强能帮忙"包装"一下，贷款批下来后给他贷款金额2%的好处费（200万元）。''',
        'characters': [
            {'role': '主人公', 'name': '陈志强', 'description': '信贷部主管，35岁，业务能力突出但绩效压力很大，父亲刚查出重病'},
            {'role': '请托人', 'name': '周明', 'description': '高中同学，房地产公司老板，以"老感情"和利益为诱饵'},
            {'role': '影响人', 'name': '陈父', 'description': '陈志强父亲，患癌症需靶向治疗，年治疗费用约50万元'},
        ],
        'stages': [
            {
                'title': '情感牌型诱惑',
                'description': '周明以老同学身份约陈志强吃饭，回忆当年同窗情谊，强调"咱俩什么关系"，表示只是想"请你帮个忙"，不会让他为难。',
                'dialogue': '周明："老陈，还记得当年咱俩一起打篮球的日子吗？那时候你是队里主力，我是啦啦队长。这么多年了，咱俩的交情可不能忘啊。"'
            },
            {
                'title': '利益诱惑型',
                'description': '在熟悉后，周明提出合作方案，表示"行业里都是这样操作的"，事成之后给200万好处费，并先给10万元"活动经费"。',
                'dialogue': '周明："老陈，这事要是成了，200万好处费先给你。10万块先拿着买点年货，不够再说。反正这事只有你知我知。"'
            },
            {
                'title': '信息威胁型',
                'description': '陈志强表示为难，周明暗示知道陈志强最近在银行的"一些小问题"（消费贷用于炒股），如果不能合作，可能会"聊聊"。',
                'dialogue': '周明："老陈啊，听说你最近手头有点紧？消费贷炒 股这事，说大不大说小不小，你要是觉得我不够意思，我也可以找别人聊聊。"'
            },
            {
                'title': '制度漏洞型',
                'description': '周明提供了一套"包装方案"：通过关联公司做虚假贸易合同，降低资产负债率，并表示"这种操作在业内很常见，出了事也有人担着"。',
                'dialogue': '周明："老陈，这是我们操作过的案例，你看，资产负债率从75%降到了55%，完全符合审批条件。这种事太正常了，业内都这么干。"'
            },
            {
                'title': '连环套型',
                'description': '陈志强犹豫再三后答应帮忙。贷款批下来后，周明如约支付了200万。但随后周明又提出新的"合作"：第二笔更大金额的贷款需要帮助。',
                'dialogue': '周明："老陈，第一笔贷款顺利批下来了，我们的合作很愉快。第二笔3个亿的项目还要继续合作啊，你放心，费用是上一笔的两倍。"'
            }
        ],
        'decision_point': '面对老同学的情谊、高额利益、以及可能被揭发的把柄，陈志强应该如何应对？这笔"第一笔"贷款是否应该帮忙审批？',
        'outcomes': [
            {'type': '错误选择', 'description': '陈志强越陷越深，成为周明的"工具人"。最终违规贷款问题暴雷，陈志强被开除并追究法律责任，周明也因骗贷被调查。'},
            {'type': '中途动摇', 'description': '陈志强做了第一笔后想要退出，但已经被周明套住，最终被迫继续合作，直到案发。'},
            {'type': '正确选择', 'description': '陈志强在发现周明威胁时就向银行监察部门报告，配合调查。虽然短期受影响，但最终得到保护，并因举报有功受到表彰。'},
            {'type': '理想选择', 'description': '陈志强在第一次被情感绑架时就保持警惕，明确拒绝，并及时向领导报备。虽然失去"老朋友"，但保护了自己。'},
        ],
        'questions': [
            '陈志强如果配合周明的"包装方案"，可能涉及哪些违法犯罪？',
            '周明威胁要举报陈志强的"小问题"，这是否构成敲诈勒索？',
            '当老同学以"感情"为名提出违法要求时，应该如何拒绝？',
            '金融行业如何防范信贷审批中的廉政风险？',
            '如果发现同事存在违规贷款行为，应该怎么办？',
        ],
        'discuss': [
            '讨论：有人说"银行是弱势群体"，在信贷审批中，银行员工面临哪些压力？如何应对？',
            '讨论：周明的行为是"行贿"还是"诈骗"？有什么区别？',
            '讨论：陈志强自己用消费贷炒股的行为是否违纪？这是否影响他举报的正当性？',
        ]
    }
    add_case_document(doc, case_4)

    # ========== 案例五：内部资源分配类 ==========
    case_5 = {
        'title': '案例五：内部资源分配类',
        'background': '''刘建业，48岁，某省直机关后勤中心主任，负责机关办公楼、车辆、办公用品等后勤资源的管理和分配。在这个岗位上工作10年，刘建业积累了广泛的人脉关系。

最近，机关要进行办公设备更新采购，预算200万元。同时，机关有3个下属单位也在申请办公设备，但预算有限，只能满足2个单位的要求。

刘建业的表弟开了一家办公设备公司，想拿下这个采购项目。而机关下属单位A的负责人老张也来找刘建业，希望在分配时"多多关照"。''',
        'characters': [
            {'role': '主人公', 'name': '刘建业', 'description': '后勤中心主任，48岁，工作10年，权力资源丰富但面临多重请托压力'},
            {'role': '亲情请托人', 'name': '刘表弟', 'description': '刘建业表弟，办公设备公司老板，以"血缘关系"为由请托'},
            {'role': '利益请托人', 'name': '老张', 'description': '机关下属单位负责人，以"老同事"关系请托，并暗示"不会让你吃亏"'},
        ],
        'stages': [
            {
                'title': '亲情绑架型诱惑',
                'description': '刘表弟以"血缘关系"为由，多次找刘建业，希望"肥水不流外人田"。刘建业表示会"在同等条件下优先考虑"。',
                'dialogue': '刘表弟："表哥，咱俩从小一块长大，我开了这个公司也不容易。这次的采购你帮帮忙，在同等条件下考虑一下我，我不会让你失望的。"'
            },
            {
                'title': '利益诱惑型',
                'description': '刘表弟在一次家宴后，悄悄给刘建业一个"信封"，里面有5万元现金，表示是"给姨妈的营养费"，与采购无关。',
                'dialogue': '刘表弟："表哥，这是5万块，给姨妈买点营养品。采购的事你就当帮忙了，成不成再说，反正这点心意你先收下。"'
            },
            {
                'title': '公私兼顾型诱惑',
                'description': '老张也来找刘建业，暗示如果能"优先保障"他们单位的设备，可以给刘建业个人"辛苦费"，并且"发票开成办公用品"很难查。',
                'dialogue': '老张："老刘，这次设备分配的事你多操心。放心，不会让你白忙的，辛苦费到时候给你弄成办公用品发票，神不知鬼不觉。"'
            },
            {
                'title': '资源分配决策点',
                'description': '设备采购和分配都需要刘建业签字。他面临多重压力：表弟要订单，老张要优先，还有其他请托人也在活动。刘建业开始思考如何在"不违规"的情况下"平衡"各方。',
                'dialogue': '刘建业内心："这也太难了，各方都得照顾到啊。要不我就稍微倾斜一下，反正都是在规则范围内..."'
            },
            {
                'title': '双重标准型诱惑',
                'description': '刘建业决定"灵活处理"：在采购中让表弟公司中标，在分配时给老张单位优先。但他没有想到，这其实已经构成了利益输送。',
                'dialogue': '刘建业："我就这么定了，反正表弟公司产品质量也可以，老张单位也确实需要帮助。这样谁都不伤和气..."'
            }
        ],
        'decision_point': '面对亲情、友情和利益的多重压力，刘建业应该如何处理内部资源分配问题？"在规则范围内灵活处理"是否可行？',
        'outcomes': [
            {'type': '错误选择', 'description': '刘建业自以为"灵活处理"没违规，但实际上已经构成利益输送。最终被举报查处，受到党纪政纪处分。'},
            {'type': '侥幸心理型', 'description': '刘建业认为自己做得隐蔽，但天网恢恢，最终因其他案件牵连被查出，身败名裂。'},
            {'type': '正确选择', 'description': '刘建业坚持原则，按照制度规定公开招标、公平分配。主动向组织报备各方请托情况，请组织监督。'},
            {'type': '智慧选择', 'description': '刘建业建立公开透明的分配机制，让程序在阳光下运行。让各方请托都无所遁形，既坚持了原则，又避免了尴尬。'},
        ],
        'questions': [
            '刘建业"在同等条件下优先考虑"表弟公司，是否违规？为什么？',
            '刘建业收取"营养费"后进行利益分配，是否构成受贿？',
            '"发票开成办公用品"这种操作是否合规？为什么？',
            '当亲情与原则冲突时，应该如何处理？',
            '从制度角度，如何防范内部资源分配中的廉政风险？',
        ],
        'discuss': [
            '讨论：有人认为"帮亲戚朋友一点忙不算什么"，这种观点为什么危险？',
            '讨论：如何正确处理人情世故与原则纪律的关系？',
            '讨论：如果发现资源分配中存在"潜规则"，应该怎么办？',
        ]
    }
    add_case_document(doc, case_5)

    # ========== 案例讨论指南 ==========
    doc.add_page_break()
    create_guide_section(doc)

    # 保存文档
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")

if __name__ == '__main__':
    create_document()
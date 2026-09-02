#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)

    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)

    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)

    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)

    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

# Create final merged document
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# ========== Cover ==========
for _ in range(8): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('AI时代决策工作手册')
r.bold = True; r.font.size = Pt(32); r.font.color.rgb = RGBColor(0,51,102)

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run('把个人判断变成组织可复用的决策资产')
r.font.size = Pt(16); r.italic = True

for _ in range(4): doc.add_paragraph()

a = doc.add_paragraph()
a.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a.add_run('学员对象'); r.bold = True; r.font.size = Pt(14)

a2 = doc.add_paragraph()
a2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = a2.add_run('企业中层管理者、项目负责人、HRBP、培训经理、行动学习催化师')
r.font.size = Pt(12)

for _ in range(6): doc.add_paragraph()

c = doc.add_paragraph()
c.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.add_run('© 罗宏伟'); r.bold = True; r.font.size = Pt(14)

c2 = doc.add_paragraph()
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c2.add_run('版权所有 侵权必究'); r.font.size = Pt(10); r.font.color.rgb = RGBColor(128,128,128)

doc.add_page_break()
print('Cover done')

# ========== TOC ==========
toct = doc.add_paragraph()
toct.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = toct.add_run('目 录')
r.bold = True; r.font.size = Pt(24)
doc.add_paragraph()

toc_items = [
    ('学习目标', '4'), ('课前准备', '5'),
    ('PART 1  认知升级：复盘与决策卡的本质区分', '6'),
    ('  第一章  复盘写的是过去，决策卡写的是未来', '6'),
    ('  第二章  不是每个决策都值得做成一张卡', '7'),
    ('  第三章  复盘访谈问的是你当时不确定什么', '8'),
    ('  第四章  隐性判断是问出来的，不是想出来的', '9'),
    ('  第五章  决策卡是检查表加开关', '10'),
    ('  第六章  触发条件写不清楚这张卡就是废纸', '11'),
    ('  第七章  场景映射矩阵', '12'),
    ('  第八章  失败案例是用来对照的', '13'),
    ('  第九章  第一个读者应该是反对你的人', '14'),
    ('  第十章  训练是让人在场景里犯一次错', '15'),
    ('PART 1  Q&A', '16'),
    ('PART 2  组织落地', '17'),
    ('  第一章  没人认领是最大的敌人', '17'),
    ('  第二章  稽核不是查错是记住学过什么', '18'),
    ('  第三章  会过时是活着的证据', '19'),
    ('  第四章  最易犯的错是替决策者总结', '20'),
    ('  第五章  交付的是组织判断力', '21'),
    ('PART 2  Q&A', '22'),
    ('工具模板', '23'),
    ('课后资源', '30'),
    ('附录', '32'),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.add_run('\t' + page)

doc.add_page_break()
print('TOC done')

# ========== Learning Objectives ==========
ot = doc.add_paragraph()
r = ot.add_run('学习目标')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph()

for label, content in [
    ('知识目标', '理解复盘与决策卡的本质区分'),
    ('技能目标', '掌握决策卡制作的全流程方法'),
    ('态度目标', '建立判断力是组织核心资产的意识'),
]:
    oh = doc.add_paragraph()
    r = oh.add_run(label); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0,102,204)
    doc.add_paragraph(content)

doc.add_page_break()
print('Objectives done')

# ========== Pre-class Prep ==========
pp = doc.add_paragraph()
r = pp.add_run('课前准备')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0,51,102)
doc.add_paragraph()

tt = doc.add_paragraph()
r = tt.add_run('课前自测问卷'); r.bold = True; r.font.size = Pt(14)

for q in ['你过去做过的复盘，现在还能想起哪些具体内容？',
          '你曾经是否遇到过复盘做了很多但下次还是踩坑的情况？',
          '你认为复盘和决策卡最大的区别是什么？',
          '你是否曾经有过我就是觉得不对但说不清的经历？',
          '你觉得一个组织中最值得沉淀的是什么？']:
    doc.add_paragraph(q)

doc.add_paragraph()
th = doc.add_paragraph()
r = th.add_run('课前思考'); r.bold = True; r.font.size = Pt(14)

tcp = doc.add_paragraph()
tcp.add_run('你上一次做的复盘现在还记得什么？')

note = doc.add_paragraph()
r = note.add_run('请在下方记录你的思考：')
r.italic = True; r.font.color.rgb = RGBColor(128,128,128)

for _ in range(6):
    line = doc.add_paragraph()
    line.add_run('_' * 70)
    line.paragraph_format.space_after = Pt(12)

doc.add_page_break()
print('Pre-class done')

doc.save(r'D:\CC\temp\merged_part1.docx')
print('Part 1 saved')

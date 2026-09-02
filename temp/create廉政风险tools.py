# -*- coding: utf-8 -*-
"""
Create all 8 Word documents for 廉政风险情景决策训练营
- 5 工具集锦 files
- 3 练习材料 files
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_TOOLS = "D:/新课开发/党业融合/廉政风险情景决策/完整课程包/008-课堂工具集锦"
OUTPUT_PRACTICE = "D:/新课开发/党业融合/廉政风险情景决策/完整课程包/012-练习材料"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Add borders to all cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '999999')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_title_para(doc, text, size=22, bold=True, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def add_heading_para(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        run.font.size = Pt(14)
        set_cell_bg  # not used here
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    return para

def add_body_para(doc, text, bold=False, indent=False):
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(0.75)
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return para

def add_checkbox_row(table, row_idx, text, col_span=True):
    """Add a row with checkbox symbol."""
    cell = table.cell(row_idx, 0)
    cell.text = "☐"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(11)

    if col_span and len(table.columns) > 1:
        # merge rest of row
        for col in range(1, len(table.columns)):
            table.cell(row_idx, col).text = ""
        merged = table.cell(row_idx, 0).merge(table.cell(row_idx, len(table.columns)-1))
        p = merged.paragraphs[0]
        p.clear()
        run2 = p.add_run("☐  " + text)
        run2.font.size = Pt(11)
        run2.font.name = '微软雅黑'
        run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_section_title(doc, text):
    """Add a section title bar."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("  " + text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # Add bottom border via paragraph properties
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2F5496')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# FILE 1: 工具集锦-个人风险自检工具.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file1(doc):
    doc = Document()
    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    add_title_para(doc, "廉政风险情景决策训练营", size=16, color=(47, 84, 150))
    add_title_para(doc, "个人风险自检工具", size=22, bold=True, color=(47, 84, 150))
    add_body_para(doc, "——四种合理化模式自查表")
    doc.add_paragraph()

    # ── Section 1: 四种合理化模式自查表 ──────────────────────────────────
    add_section_title(doc, "一、四种合理化模式自查表")

    patterns = [
        {
            "title": "模式一："这只是人之常情"",
            "desc": "把利益输送包装成正常人情往来",
            "items": [
                "逢年过节收受管理服务对象的礼品礼金",
                "以"人情往来"为由接受宴请",
                "在婚丧嫁娶时收受超出正常范围的礼金",
                "以"帮忙办事"为名收取好处费",
            ]
        },
        {
            "title": "模式二："我又没主动要，是别人塞给我的"",
            "desc": "被动接受减轻责任感",
            "items": [
                "对方趁你不注意时留下现金或礼品",
                "以"代购""代付"名义掩饰利益输送",
                "对方坚持要给，你碍于情面没有拒绝",
                "对方说"只是小心意，不收就是不给面子"",
            ]
        },
        {
            "title": "模式三："这次帮个忙，以后再也不会有下次"",
            "desc": ""仅此一次"的承诺",
            "items": [
                "答应为对方"就这一次"提供便利",
                "认为偶尔一次"不算什么"",
                "相信对方说的"下不为例"",
                "以"特例"为由突破原则底线",
            ]
        },
        {
            "title": "模式四："反正大家都是这么做的"",
            "desc": "群体行为分摊责任感",
            "items": [
                "看到周围人都在做，觉得自己不做是"异类"",
                "认为"法不责众"，随大流不会有事",
                "以"行业惯例"为由自我安慰",
                "觉得"别人能收，我为什么不能收"",
            ]
        },
    ]

    for p_idx, pat in enumerate(patterns, 1):
        # Pattern header table (2 cols: number + title)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = 'Table Grid'
        tbl.columns[0].width = Cm(1.5)
        tbl.columns[1].width = Cm(14)
        hdr_cell = tbl.cell(0, 0)
        hdr_cell.text = f"模式{p_idx}"
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cell, '2F5496')
        for run in hdr_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True

        title_cell = tbl.cell(0, 1)
        p = title_cell.paragraphs[0]
        p.clear()
        run = p.add_run(pat["title"])
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        p2 = title_cell.add_paragraph(pat["desc"])
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.italic = True
        p2.runs[0].font.name = '微软雅黑'
        p2.runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_borders(tbl)

        # Items table
        item_tbl = doc.add_table(rows=len(pat["items"])+1, cols=1)
        item_tbl.style = 'Table Grid'
        set_cell_borders(item_tbl)
        # Header row
        hdr = item_tbl.cell(0, 0)
        hdr.text = "自查项目"
        set_cell_bg(hdr, 'D6DCE5')
        hdr.paragraphs[0].runs[0].font.bold = True
        hdr.paragraphs[0].runs[0].font.size = Pt(10.5)

        for i, item in enumerate(pat["items"], 1):
            cell = item_tbl.cell(i, 0)
            cell.text = f"☐  {item}"
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            cell.paragraphs[0].runs[0].font.name = '微软雅黑'
            cell.paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph()

    # ── Section 2: 个人风险警报语清单 ───────────────────────────────────
    add_section_title(doc, "二、个人风险警报语清单")

    add_body_para(doc, "当以下任意一条在你脑海中出现时，请立即警惕——这是合理化模式的预警信号：", bold=True)
    doc.add_paragraph()

    alert_items = [
        (""这只是人之常情，没关系的"", "模式一"),
        (""他非要给我，我也没办法"", "模式二"),
        (""就这一次，不会有人知道的"", "模式三"),
        (""大家都在做，我不过是随大流"", "模式四"),
        (""这件事天知地知你知我知"", "侥幸心理"),
        (""我帮了他，他以后也会帮我"", "利益交换"),
        (""这点小钱算什么，根本不算贿赂"", "轻视行为"),
        (""我有分寸的，不会有问题"", "过度自信"),
    ]

    alert_tbl = doc.add_table(rows=len(alert_items)+1, cols=3)
    alert_tbl.style = 'Table Grid'
    set_cell_borders(alert_tbl)
    headers = ["预警话语", "所属模式", "风险等级"]
    for col_idx, hdr_text in enumerate(headers):
        cell = alert_tbl.cell(0, col_idx)
        cell.text = hdr_text
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10.5)

    risk_levels = ["高风险", "高风险", "高风险", "高风险", "中风险", "高风险", "中风险", "中风险"]
    risk_colors = ["FF9999", "FF9999", "FF9999", "FF9999", "FFDD99", "FF9999", "FFDD99", "FFDD99"]

    for row_idx, (item, mode) in enumerate(zip(alert_items, risk_levels), 1):
        alert_tbl.cell(row_idx, 0).text = item[0]
        alert_tbl.cell(row_idx, 0).paragraphs[0].runs[0].font.size = Pt(10.5)
        alert_tbl.cell(row_idx, 1).text = item[1]
        alert_tbl.cell(row_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        alert_tbl.cell(row_idx, 2).text = risk_levels[row_idx-1]
        alert_tbl.cell(row_idx, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(alert_tbl.cell(row_idx, 2), risk_colors[row_idx-1])
        for col in range(3):
            alert_tbl.cell(row_idx, col).paragraphs[0].runs[0].font.name = '微软雅黑'
            alert_tbl.cell(row_idx, col).paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # ── Section 3: 决策前自问清单 ───────────────────────────────────────
    add_section_title(doc, "三、决策前自问清单")

    add_body_para(doc, "在做出任何涉及利益往来的决策前，请逐项问自己：")
    doc.add_paragraph()

    questions = [
        ("这件事合法吗？", "是否违反党纪国法、单位的各项规章制度？"),
        ("这件事正当吗？", "是否超出正常人情往来、工作职责范围？"),
        ("这件事安全吗？", "如果被公开、被监督，是否能经得起审查？"),
        ("这件事值得吗？", "眼前的利益与长期职业发展、声誉相比，孰轻孰重？"),
        ("我能告诉家人吗？", "如果家人知道这件事，他们会为我感到骄傲还是担忧？"),
        ("我能向组织坦白吗？", "如果主动向组织报告，是否能减轻或免于处罚？"),
    ]

    q_tbl = doc.add_table(rows=len(questions)+1, cols=3)
    q_tbl.style = 'Table Grid'
    set_cell_borders(q_tbl)
    hdrs = ["自问问题", "深度追问", "回答（是/否）"]
    for ci, h in enumerate(hdrs):
        cell = q_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10.5)

    for ri, (q, sub) in enumerate(questions, 1):
        q_tbl.cell(ri, 0).text = q
        q_tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        q_tbl.cell(ri, 1).text = sub
        q_tbl.cell(ri, 2).text = "□ 是    □ 否"
        for col in range(3):
            for run in q_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    # Footer note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【使用说明】每日上岗前、工作决策前、情绪波动时各使用一次，逐项勾选，做好记录。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 2: 工具集锦01-合理化模式识别量表.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file2(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "合理化模式识别量表", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "指导语：以下每道题描述一种情境或想法，请根据实际情况选择最符合的选项。")
    doc.add_paragraph()

    patterns = [
        {
            "name": "模式一："这只是人之常情"",
            "desc": "把利益输送包装成正常人情往来",
            "questions": [
                "我认为逢年过节收受管理服务对象的礼品，是人之常情，无可厚非",
                "只要金额不大、不影响公正，收点"小意思"不算什么",
                "同事之间互相送点礼物，是正常的工作关系维护",
                "对方是多年老朋友/老同事，拒绝礼物反而伤感情",
                "办喜事时收受管理服务对象的礼金，是主客双方的礼节",
                "对方送礼时说是"顺便""随手"，我便不再推辞",
                "我认为在公务活动中适当接受宴请，有助于联络感情、协调工作",
                "只要不超标、不影响公正，接受一点土特产不算违规",
            ]
        },
        {
            "name": "模式二："我又没主动要，是别人塞给我的"",
            "desc": "被动接受减轻责任感",
            "questions": [
                "对方趁我不注意时放下现金/礼品，我事后发现但没有主动退还",
                "我曾以"代购""代付"等名义，实际上收受了对方的"差价"",
                "对方坚持要给，我说不要但对方非给，我就收下了",
                "对方说"只是一点心意，不收就是看不起我"，我碍于情面收下",
                "我认为"不是我要的，是对方硬塞的"，主要责任在对方",
                "对方借着"还钱""还情"的名义给我好处，我没有拒绝",
                "我在不知情的情况下收到了对方通过第三方转交的好处",
                "我认为"只要不主动索取，被动接受也不算太严重"",
            ]
        },
        {
            "name": "模式三："这次帮个忙，以后再也不会有下次"",
            "desc": ""仅此一次"的承诺",
            "questions": [
                "我曾在某次特定情况下答应为他人提供便利，当时认为"就这一次"",
                "我曾以"下不为例"为由，在某次事件中突破了原则",
                "我曾认为偶尔一次违规"不算什么"，不会造成严重后果",
                "我曾相信对方"只此一次"的承诺，继续保持不正当往来",
                "我曾在帮助他人后，收受了对方"感谢性"的好处费",
                "我认为在紧急/特殊情况下，可以暂时放宽要求",
                "我曾对自己说"这是特例"，为不当行为寻找借口",
                "我曾认为"偶尔为之不会被发现"，存在侥幸心理",
            ]
        },
        {
            "name": "模式四："反正大家都是这么做的"",
            "desc": "群体行为分摊责任感",
            "questions": [
                "我看到周围同事都有类似行为，认为自己不做是"吃亏"",
                "我认为"法不责众"，大家都做的事被查处的概率很低",
                "我以"行业惯例"或"惯例做法"来说服自己接受某种不当利益",
                "我曾觉得"别人能收，我为什么不能收"，产生了攀比心理",
                "我认为只要随大流、不出风头，就不会成为"出头鸟"",
                "在群体行为中，我曾减少对自己行为的道德约束",
                "我曾认为"大家都这样干，肯定没问题"，放松了警惕",
                "我曾用"别人都这么做"来为自己的不当行为辩护",
            ]
        },
    ]

    # Scoring guide
    scoring_note = doc.add_paragraph()
    scoring_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = scoring_note.add_run("评分标准：从不这样=1分  偶尔这样=2分  有时这样=3分  经常这样=4分  总是这样=5分")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for p_idx, pat in enumerate(patterns, 1):
        page_break(doc)

        # Pattern header
        tbl_hdr = doc.add_table(rows=1, cols=1)
        tbl_hdr.style = 'Table Grid'
        hcell = tbl_hdr.cell(0, 0)
        hcell.text = f"第{p_idx}部分：{pat['name']}"
        set_cell_bg(hcell, '2F5496')
        for run in hcell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(12)

        p_desc = doc.add_paragraph()
        run = p_desc.add_run(f"「{pat['desc']}」")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # Questions table
        q_tbl = doc.add_table(rows=len(pat["questions"])+1, cols=6)
        q_tbl.style = 'Table Grid'
        set_cell_borders(q_tbl)

        # Header
        hdrs = ["题号", "情境描述", "从不这样", "偶尔这样", "有时这样", "经常这样"]
        col_widths = [Cm(1.2), Cm(9), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.8)]
        for ci, h in enumerate(hdrs):
            cell = q_tbl.cell(0, ci)
            cell.text = h
            set_cell_bg(cell, 'D6DCE5')
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for qi, q in enumerate(pat["questions"], 1):
            q_tbl.cell(qi, 0).text = str(qi)
            q_tbl.cell(qi, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            q_tbl.cell(qi, 1).text = q
            for col in range(6):
                for run in q_tbl.cell(qi, col).paragraphs[0].runs:
                    run.font.size = Pt(9.5)
                    run.font.name = '微软雅黑'
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # Checkboxes in score columns
            for col in range(2, 6):
                q_tbl.cell(qi, col).text = "☐"
                q_tbl.cell(qi, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Scoring interpretation ─────────────────────────────────────────────
    page_break(doc)
    add_section_title(doc, "综合评分与解读")

    interp_tbl = doc.add_table(rows=5, cols=3)
    interp_tbl.style = 'Table Grid'
    set_cell_borders(interp_tbl)

    interp_data = [
        ["总分范围", "风险等级", "解读与建议"],
        ["32–40分", "红色（极高风险）", "合理化模式已深度内化，必须立即采取干预措施。建议寻求组织帮助，进行专项廉政谈话。"],
        ["24–31分", "橙色（高风险）", "存在明显的自我欺骗倾向，需要系统学习廉政风险防控知识，建立严格的自我监督机制。"],
        ["16–23分", "黄色（中等风险）", "有轻微的合理化倾向，保持警惕。建议建立定期自检习惯，主动接受组织和群众监督。"],
        ["8–15分", "绿色（低风险）", "风险意识较强，继续保持。如能帮助他人提高廉政意识，可发挥示范作用。"],
    ]

    risk_colors_interp = ["FFFFFF", "FFCCCC", "FFDD99", "DDFFDD"]
    for row_idx, row_data in enumerate(interp_data):
        for col_idx, val in enumerate(row_data):
            cell = interp_tbl.cell(row_idx, col_idx)
            cell.text = val
            if row_idx == 0:
                set_cell_bg(cell, '2F5496')
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
            else:
                if row_idx > 0:
                    set_cell_bg(cell, risk_colors_interp[row_idx-1])
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    # Pattern breakdown table
    add_body_para(doc, "各模式得分统计：", bold=True)
    pat_score_tbl = doc.add_table(rows=5, cols=3)
    pat_score_tbl.style = 'Table Grid'
    set_cell_borders(pat_score_tbl)
    pat_hdrs = ["模式", "得分（8–40分）", "风险判断"]
    for ci, h in enumerate(pat_hdrs):
        cell = pat_score_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10.5)

    pat_names = ["模式一：人之常情", "模式二：被动接受", "模式三：仅此一次", "模式四：随大流"]
    for ri, pn in enumerate(pat_names, 1):
        pat_score_tbl.cell(ri, 0).text = pn
        pat_score_tbl.cell(ri, 1).text = "______分"
        pat_score_tbl.cell(ri, 2).text = "□ 正常  □ 警惕  □ 危险"
        for col in range(3):
            for run in pat_score_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    note2 = doc.add_paragraph()
    note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note2.add_run("【使用建议】课程前做前测，课程结束后做后测，对比分数变化，检验学习成效。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 3: 工具集锦02-决策情境回顾表.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file3(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "决策情境回顾表", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "指导语：回想一个你曾经面临廉政风险的决策情境，填写下表，帮助你识别合理化模式。")
    doc.add_paragraph()

    # Single回顾 table
    fields = [
        ("一、情境基本信息", None),
        ("情境编号", "例：案例A"),
        ("发生时间", "____年____月____日"),
        ("发生地点", ""),
        ("涉及人员", ""),
        ("情境类型", "□ 利益输送    □ 请托办事    □ 违规宴请    □ 其他：______"),
    ]

    # Section 1: basic info
    add_section_title(doc, "一、情境基本信息")
    info_tbl = doc.add_table(rows=len(fields)-1, cols=2)
    info_tbl.style = 'Table Grid'
    set_cell_borders(info_tbl)
    for ri, (label, hint) in enumerate(fields[1:], 0):
        cell_label = info_tbl.cell(ri, 0)
        cell_label.text = label
        set_cell_bg(cell_label, 'EBF0F9')
        cell_label.paragraphs[0].runs[0].font.bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(10.5)
        cell_hint = info_tbl.cell(ri, 1)
        cell_hint.text = hint if hint else ""
        cell_hint.paragraphs[0].runs[0].font.size = Pt(10.5) if hint else None
        for run in cell_hint.paragraphs[0].runs:
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    add_section_title(doc, "二、当时决策回顾")

    decision_fields = [
        ("当时的决策是什么？", 3),
        ("当时有哪些人在场？", 2),
        ("对方提出了什么要求/提供了什么利益？", 2),
        ("你当时是如何回应的？", 3),
    ]

    for label, rows in decision_fields:
        lbl_tbl = doc.add_table(rows=1, cols=1)
        lbl_tbl.style = 'Table Grid'
        lc = lbl_tbl.cell(0, 0)
        lc.text = label
        set_cell_bg(lc, 'D6DCE5')
        for run in lc.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10.5)

        val_tbl = doc.add_table(rows=rows, cols=1)
        val_tbl.style = 'Table Grid'
        set_cell_borders(val_tbl)
        for r in range(rows):
            val_tbl.cell(r, 0).text = ""
        doc.add_paragraph()

    add_section_title(doc, "三、当时说服自己的话术（请尽量回忆原话）")
    talk_tbl = doc.add_table(rows=4, cols=2)
    talk_tbl.style = 'Table Grid'
    set_cell_borders(talk_tbl)
    talk_labels = ["当时对自己说了什么？", "这些话术属于哪种合理化模式？", "现在回看，你有什么新认识？", "如果重来，你会怎么做？"]
    for ri, lbl in enumerate(talk_labels):
        talk_tbl.cell(ri, 0).text = lbl
        set_cell_bg(talk_tbl.cell(ri, 0), 'EBF0F9')
        talk_tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        talk_tbl.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(10.5)
        talk_tbl.cell(ri, 1).text = ""
        for run in talk_tbl.cell(ri, 0).paragraphs[0].runs:
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    add_section_title(doc, "四、模式识别与重做选择")

    mode_id_tbl = doc.add_table(rows=6, cols=2)
    mode_id_tbl.style = 'Table Grid'
    set_cell_borders(mode_id_tbl)

    mode_labels = [
        "这个情境涉及哪种合理化模式？（可多选）",
        "模式一："这只是人之常情"  □ 是    □ 否",
        "模式二："我又没主动要，是别人塞给我的"  □ 是    □ 否",
        "模式三："这次帮个忙，以后再也不会有下次"  □ 是    □ 否",
        "模式四："反正大家都是这么做的"  □ 是    □ 否",
        "如果重来，你会做出什么不同选择？",  # merged row
    ]

    mode_id_tbl.cell(0, 0).text = mode_labels[0]
    set_cell_bg(mode_id_tbl.cell(0, 0), 'D6DCE5')
    mode_id_tbl.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    merged = mode_id_tbl.cell(0, 0).merge(mode_id_tbl.cell(0, 1))

    for ri, lbl in enumerate(mode_labels[1:5], 1):
        mode_id_tbl.cell(ri, 0).text = lbl
        merged2 = mode_id_tbl.cell(ri, 0).merge(mode_id_tbl.cell(ri, 1))

    mode_id_tbl.cell(5, 0).text = mode_labels[5]
    set_cell_bg(mode_id_tbl.cell(5, 0), 'EBF0F9')
    merged3 = mode_id_tbl.cell(5, 0).merge(mode_id_tbl.cell(5, 1))
    mode_id_tbl.cell(5, 1).text = ""

    for ri in range(6):
        for run in mode_id_tbl.cell(ri, 0).paragraphs[0].runs:
            run.font.size = Pt(10.5)
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【使用建议】每周回顾一次近期决策情境，持续训练识别合理化模式的能力。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 4: 工具集锦03-承诺与监督表.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file4(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "承诺与监督表", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "以书面形式做出承诺，并指定监督人，借助外部力量巩固廉政意识。")
    doc.add_paragraph()

    # Part 1: Personal commitment
    add_section_title(doc, "一、个人廉政承诺书")

    commitment_text = """本人郑重承诺：

一、严格遵守党纪国法，坚守廉政底线，不以任何形式收受管理服务对象的财物、不当利益。

二、主动识别并抵制四种合理化模式：
  1. 不以"人之常情"为借口进行利益输送；
  2. 不以"被动接受"为由减轻自身廉政责任；
  3. 不以"仅此一次"为由突破原则底线；
  4. 不以"随大流"为由分摊个人责任。

三、如遇廉政风险情境，主动及时向组织和监督人报告，不隐瞒、不拖延。

四、自愿接受组织监督和群众监督，如违反承诺，愿意接受组织处理。

五、积极参加廉政学习，定期进行自我检视，做到警钟长鸣。

本人承诺以上内容真实可信，并愿意承担相应责任。"""

    comm_tbl = doc.add_table(rows=1, cols=1)
    comm_tbl.style = 'Table Grid'
    set_cell_borders(comm_tbl)
    comm_cell = comm_tbl.cell(0, 0)
    comm_cell.text = commitment_text
    for para in comm_cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Signature fields
    sig_tbl = doc.add_table(rows=2, cols=3)
    sig_tbl.style = 'Table Grid'
    set_cell_borders(sig_tbl)
    sig_labels = ["承诺人（签名）：", "监督人（签名）：", "日期："]
    for ci, lbl in enumerate(sig_labels):
        sig_tbl.cell(0, ci).text = lbl
        set_cell_bg(sig_tbl.cell(0, ci), 'EBF0F9')
        sig_tbl.cell(0, ci).paragraphs[0].runs[0].font.bold = True
        sig_tbl.cell(1, ci).text = ""

    doc.add_paragraph()

    # Part 2: Supervision arrangement
    add_section_title(doc, "二、监督人信息")
    sup_tbl = doc.add_table(rows=5, cols=2)
    sup_tbl.style = 'Table Grid'
    set_cell_borders(sup_tbl)
    sup_fields = [
        ("监督人姓名", ""),
        ("监督人职务", ""),
        ("监督人联系方式", ""),
        ("与承诺人关系", "□ 上级领导  □ 同事  □ 家人  □ 其他：______"),
        ("监督方式", "□ 定期谈话  □ 随时咨询  □ 每月检视  □ 其他：______"),
    ]
    for ri, (label, hint) in enumerate(sup_fields):
        sup_tbl.cell(ri, 0).text = label
        set_cell_bg(sup_tbl.cell(ri, 0), 'D6DCE5')
        sup_tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        sup_tbl.cell(ri, 1).text = hint
        for col in range(2):
            for run in sup_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Part 3: Regular review
    add_section_title(doc, "三、定期自检回顾提醒")

    review_tbl = doc.add_table(rows=8, cols=3)
    review_tbl.style = 'Table Grid'
    set_cell_borders(review_tbl)

    rev_headers = ["自检周期", "自检日期", "自检结果（正常/异常）及说明"]
    for ci, h in enumerate(rev_headers):
        cell = review_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10.5)

    periods = ["第1次", "第2次", "第3次", "第4次", "第5次", "第6次", "第7次"]
    for ri, period in enumerate(periods, 1):
        review_tbl.cell(ri, 0).text = period
        review_tbl.cell(ri, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        review_tbl.cell(ri, 1).text = "____年__月__日"
        review_tbl.cell(ri, 2).text = ""
        for col in range(3):
            for run in review_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【使用说明】本承诺书一式三份，承诺人保留一份，监督人保留一份，交组织备案一份。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 5: 工具集锦合集.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file5():
    # Combine content from files 1-4 into one document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "课堂工具集锦（合集本）", size=22, bold=True, color=(47, 84, 150))
    doc.add_paragraph()

    intro = doc.add_paragraph()
    run = intro.add_run("本合集收录训练营全部课堂工具，共包含四个独立工具，建议配合课程进度使用。")
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # TOC-like table
    toc_tbl = doc.add_table(rows=5, cols=3)
    toc_tbl.style = 'Table Grid'
    set_cell_borders(toc_tbl)
    toc_headers = ["序号", "工具名称", "适用场景"]
    for ci, h in enumerate(toc_headers):
        cell = toc_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)

    toc_data = [
        ("1", "个人风险自检工具", "课前自测 / 每日上岗前 / 决策前"),
        ("2", "合理化模式识别量表", "课前测评 / 课后复测，对比学习成效"),
        ("3", "决策情境回顾表", "课后作业 / 每周复盘"),
        ("4", "承诺与监督表", "课程结束时签订，持久约束"),
    ]
    toc_colors = ["FFFFFF", "F2F7FF", "F2F7FF", "F2F7FF"]

    for ri, (num, name, scene) in enumerate(toc_data, 1):
        toc_tbl.cell(ri, 0).text = num
        toc_tbl.cell(ri, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_tbl.cell(ri, 1).text = name
        toc_tbl.cell(ri, 2).text = scene
        set_cell_bg(toc_tbl.cell(ri, 0), toc_colors[ri-1])
        for col in range(3):
            for run in toc_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Insert each tool
    page_break(doc)
    p = doc.add_paragraph()
    run = p.add_run("工具一：个人风险自检工具")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(47, 84, 150)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Create file1 and copy its content (we recreate sections)
    doc1 = create_file1(Document())
    # Copy body from doc1 - we'll just note that this is the combined version
    # and include summary content for each tool

    combined_tools_content = [
        {
            "title": "工具一：个人风险自检工具",
            "sections": [
                {
                    "name": "四种合理化模式自查表",
                    "desc": "针对每种合理化模式设计的自查项目清单，通过逐项勾选帮助识别自身存在的风险点。包含模式一定价核查项目（4项）、模式二被动接受核查项目（4项）、模式三仅此一次核查项目（4项）、模式四随大流核查项目（4项），共16项。",
                },
                {
                    "name": "个人风险警报语清单",
                    "desc": "8条典型预警话语，对应各自所属的合理化模式，并标注风险等级（高/中），帮助在第一时间识别自我欺骗倾向。",
                },
                {
                    "name": "决策前自问清单",
                    "desc": "6个深度追问问题，涵盖合法性、正当性、安全性、价值判断、家庭影响、组织坦白等维度，引导全面思考决策后果。",
                }
            ]
        },
        {
            "title": "工具二：合理化模式识别量表",
            "sections": [
                {
                    "name": "32道自测题",
                    "desc": "4个模式各8道题，采用Likert五级评分（从不这样-总是这样），全面评估个人在四种合理化模式上的表现程度。",
                },
                {
                    "name": "综合评分与解读",
                    "desc": "总分32-40分为红色（极高风险）；24-31分为橙色（高风险）；16-23分为黄色（中等风险）；8-15分为绿色（低风险）。各模式单独计分，帮助定向干预。",
                }
            ]
        },
        {
            "title": "工具三：决策情境回顾表",
            "sections": [
                {
                    "name": "情境基本信息",
                    "desc": "记录决策情境的时间、地点、涉及人员、情境类型，便于后续追踪分析。",
                },
                {
                    "name": "当时决策回顾",
                    "desc": "还原决策过程：决策内容、在场人员、利益交换内容、自身回应方式。",
                },
                {
                    "name": "当时说服话术分析",
                    "desc": "回忆当时的内心独白，识别属于哪种合理化模式，并进行重新审视。",
                },
                {
                    "name": "模式识别与重做选择",
                    "desc": "确认涉及的合理化模式，重新思考如果重来会如何选择，强化正确决策模式。",
                }
            ]
        },
        {
            "title": "工具四：承诺与监督表",
            "sections": [
                {
                    "name": "个人廉政承诺书",
                    "desc": "书面承诺遵守廉政底线、主动抵制四种合理化模式、如实报告廉政风险、接受监督，具有约束效力。",
                },
                {
                    "name": "监督人信息",
                    "desc": "指定监督人（上级/同事/家人），明确监督方式和联系渠道，建立外部监督机制。",
                },
                {
                    "name": "定期自检回顾提醒",
                    "desc": "7次自检记录表，每周一次，持续巩固廉政意识，形成长期习惯。",
                }
            ]
        },
    ]

    for tool in combined_tools_content:
        page_break(doc)
        p_title = doc.add_paragraph()
        run = p_title.add_run(tool["title"])
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(47, 84, 150)
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        for sec in tool["sections"]:
            sec_tbl = doc.add_table(rows=1, cols=1)
            sec_tbl.style = 'Table Grid'
            set_cell_borders(sec_tbl)
            hdr_cell = sec_tbl.cell(0, 0)
            hdr_cell.text = sec["name"]
            set_cell_bg(hdr_cell, 'D6DCE5')
            for run in hdr_cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)

            desc_tbl = doc.add_table(rows=1, cols=1)
            desc_tbl.style = 'Table Grid'
            set_cell_borders(desc_tbl)
            desc_cell = desc_tbl.cell(0, 0)
            desc_cell.text = sec["desc"]
            for run in desc_cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            doc.add_paragraph()

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 6: 练习材料01-情景决策练习册.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file6(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "情景决策练习册", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "课后练习 · 5个情景 · 决策点设计 · 追问提示 · 自我反思")
    doc.add_paragraph()

    scenarios = [
        {
            "num": "情景一",
            "title": "老同学的情意",
            "situation": """张科长在财政局工作，负责项目审批工作。他的大学同学李某经营的建筑公司正好有一个项目需要申请审批。李某得知消息后来找张科长叙旧，临走时留下一个信封，里面装有一张购物卡（价值5000元）和一张纸条："老同学，辛苦帮忙看看，多关照。"

张科长当时不在家，他的妻子收到了信封。张科长回家后，妻子问他怎么办。""",
            "decision_points": [
                "张科长是否应该收下这个信封？",
                "如果不应该收下，张科长应该怎么处理？",
                "张科长是否应该帮李某的申请加快审批？",
            ],
            "follow_ups": [
                "妻子说"就一张购物卡，又不是现金，老同学之间别太计较了"，张科长该如何回应？",
                "如果张科长把购物卡退回去，李某说"不给面子，以后没法做朋友了"，张科长该怎么办？",
                "这件事如果被审计部门发现，张科长该如何解释？",
            ],
            "reflection": "这个情景涉及哪种合理化模式？张科长应该如何打破这种合理化思维？",
            "pattern": "模式一："这只是人之常情"",
        },
        {
            "num": "情景二",
            "title": ""帮忙"的诱惑",
            "situation": """王主任是某市重点中学的教务主任，每年学校都有一些择校名额。他的老邻居张总的孩子正好要升学，张总多次找到王主任，说："王哥，帮帮忙，让孩子进你们学校，我给你准备了五万块作为感谢费。"

王主任知道这个择校名额确实还有少量剩余，但按规定需要经过统一摇号程序。王主任对张总说"我尽量想办法"，张总当场把一个黑色塑料袋塞进王主任的汽车后备箱。""",
            "decision_points": [
                "王主任是否应该接受这个"感谢费"？",
                "王主任说"我尽量想办法"算不算答应了？",
                "张总强行塞进后备箱，王主任没有及时退还，算不算"被动接受"？",
            ],
            "follow_ups": [
                "王主任回家打开塑料袋发现是五万元现金，他应该怎么办？",
                "如果王主任把钱退回去，张总说"你不收，就是不愿意帮忙"，王主任该怎么应对？",
                "如果王主任帮张总的孩子获得了择校名额，后来被查处，他会面临什么后果？",
            ],
            "reflection": "请分析王主任在处理这件事过程中的心理变化，他运用了哪些合理化模式？",
            "pattern": "模式二："我又没主动要，是别人塞给我的"",
        },
        {
            "num": "情景三",
            "title": "供应商的"心意"",
            "situation": """刘经理是某大型国企的采购部经理，负责办公设备采购。某供应商王经理在一次业务洽谈后说："刘总，你们公司是我们最大的客户，这是一点心意，感谢你们多年支持。"说着送上了一个高档手提包。

刘经理打开一看，里面有一条金项链和一张价值2万元的购物卡。刘经理说"这不合适"，王经理说："刘总放心，这是我私人感谢，跟生意没关系，你们继续用我们货就行。""",
            "decision_points": [
                "王经理说这是"私人感谢"，跟生意没关系，刘经理应该相信吗？",
                "刘经理应该收下这个手提包吗？",
                "刘经理如果拒绝，会不会影响以后的业务合作关系？",
            ],
            "follow_ups": [
                "刘经理回家后，妻子看到手提包说"供应商的一点心意，不收白不收"，刘经理该怎么办？",
                "如果刘经理最终收下了这个手提包，这个行为是否构成受贿？",
                "刘经理应该如何既不得罪供应商，又不违反廉政规定？",
            ],
            "reflection": "分析"私人感谢"这个说法的本质，它属于哪种合理化模式？",
            "pattern": "模式三："这次帮个忙，以后再也不会有下次"",
        },
        {
            "num": "情景四",
            "title": "年底的"红包雨"",
            "situation": """某县交通局在年底召开座谈会，参会的有局领导、各科室负责人，以及多家参与县里交通项目的施工企业代表。座谈会后，局办公室给每位参会人员发放了一个红包，里面有2000元现金，说是"会议补贴"。

该局赵科长发现，其他科室的同事都收下了红包，有些同事还说"这是惯例了，每年都有"。赵科长平时对廉政纪律比较注意，但当时场面有些尴尬，大家都收下了，他也不好意思拒绝。""",
            "decision_points": [
                "赵科长是否应该收下这个红包？",
                ""大家都在收，我不好意思拒绝"是不是合理的理由？",
                "如果赵科长不收，他会面临什么压力？",
            ],
            "follow_ups": [
                "如果赵科长当场拒绝收红包，会产生什么后果？",
                "如果赵科长先收下，事后悄悄退还，可行吗？",
                "这种以"会议补贴"名义发放的红包，是否属于违规收受礼金？",
            ],
            "reflection": "分析赵科长的心理过程，他受到了哪些合理化模式的影响？如果是你，你会怎么做？",
            "pattern": "模式四："反正大家都是这么做的"",
        },
        {
            "num": "情景五",
            "title": "换届前夜的"最后机会"",
            "situation": """某县教育局局长陈局长即将调任，在调任前一周，一位曾受他关照过的培训机构负责人钱老板来到他办公室，说："陈局长，您要高升了，我一直想感谢您，这张银行卡里有10万块，密码是您生日，没别的意思，就是一点心意。"

陈局长知道钱老板的培训机构在自己任期内获得了不少政策支持。陈局长有些犹豫，说："这个……不太合适吧。"钱老板说："陈局长，您要走了，以后想帮也帮不了了，这钱您不收，迟早也是别人的。"陈局长陷入沉思。""",
            "decision_points": [
                "钱老板说"您不收，迟早也是别人的"，这句话有没有道理？",
                "陈局长即将调任，收下这笔钱是否就"安全"了？",
                "陈局长是否应该收下这张银行卡？",
            ],
            "follow_ups": [
                "如果陈局长以"马上就走了，不会有人查"为由收下这笔钱，后果会怎样？",
                "钱老板说"以后想帮也帮不了了"，这句话是否构成威胁或利益交换？",
                "陈局长应该如何妥善处理这件事，既不得罪钱老板，又不违反纪律？",
            ],
            "reflection": "分析陈局长的心理过程中运用了哪些合理化模式。如果你是陈局长，你会如何应对？",
            "pattern": "综合（多种合理化模式叠加）",
        },
    ]

    for s_idx, sc in enumerate(scenarios, 1):
        if s_idx > 1:
            page_break(doc)

        # Scenario header
        sc_tbl = doc.add_table(rows=1, cols=1)
        sc_tbl.style = 'Table Grid'
        sc_cell = sc_tbl.cell(0, 0)
        sc_cell.text = f"{sc['num']}：{sc['title']}"
        set_cell_bg(sc_cell, '2F5496')
        for run in sc_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(13)

        # Situation
        sit_tbl = doc.add_table(rows=1, cols=1)
        sit_tbl.style = 'Table Grid'
        set_cell_borders(sit_tbl)
        sit_cell = sit_tbl.cell(0, 0)
        sit_cell.text = sc["situation"]
        for para in sit_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        doc.add_paragraph()

        # Decision points
        dp_tbl = doc.add_table(rows=1, cols=1)
        dp_tbl.style = 'Table Grid'
        set_cell_borders(dp_tbl)
        dp_hdr = dp_tbl.cell(0, 0)
        dp_hdr.text = "决策点（请认真思考）"
        set_cell_bg(dp_hdr, 'D6DCE5')
        dp_hdr.paragraphs[0].runs[0].font.bold = True

        dp_items_tbl = doc.add_table(rows=len(sc["decision_points"]), cols=1)
        dp_items_tbl.style = 'Table Grid'
        set_cell_borders(dp_items_tbl)
        for ri, dp in enumerate(sc["decision_points"]):
            dp_items_tbl.cell(ri, 0).text = f"☐  {dp}"
            dp_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(11)
            dp_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.name = '微软雅黑'
            dp_items_tbl.cell(ri, 0).paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        doc.add_paragraph()

        # Follow-up questions
        fu_tbl = doc.add_table(rows=1, cols=1)
        fu_tbl.style = 'Table Grid'
        set_cell_borders(fu_tbl)
        fu_hdr = fu_tbl.cell(0, 0)
        fu_hdr.text = "追问提示（深入思考）"
        set_cell_bg(fu_hdr, 'D6DCE5')
        fu_hdr.paragraphs[0].runs[0].font.bold = True

        fu_items_tbl = doc.add_table(rows=len(sc["follow_ups"]), cols=1)
        fu_items_tbl.style = 'Table Grid'
        set_cell_borders(fu_items_tbl)
        for ri, fu in enumerate(sc["follow_ups"]):
            fu_items_tbl.cell(ri, 0).text = f"{ri+1}. {fu}"
            fu_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(11)
            fu_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.name = '微软雅黑'
            fu_items_tbl.cell(ri, 0).paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        doc.add_paragraph()

        # Pattern & Reflection
        pr_tbl = doc.add_table(rows=2, cols=1)
        pr_tbl.style = 'Table Grid'
        set_cell_borders(pr_tbl)
        pr_tbl.cell(0, 0).text = f"涉及模式：{sc['pattern']}"
        set_cell_bg(pr_tbl.cell(0, 0), 'FFF2CC')
        pr_tbl.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)
        pr_tbl.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        pr_tbl.cell(1, 0).text = f"自我反思：{sc['reflection']}"
        pr_tbl.cell(1, 0).paragraphs[0].runs[0].font.size = Pt(11)
        for run in pr_tbl.cell(0, 0).paragraphs[0].runs:
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        for run in pr_tbl.cell(1, 0).paragraphs[0].runs:
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【使用建议】每个情景小组讨论20分钟，个人思考10分钟，然后全班分享。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 7: 练习材料02-小组讨论指南.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file7(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "小组讨论指南", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "讨论规则 · 引导问题 · 共识形成方法")
    doc.add_paragraph()

    # Part 1: Rules
    add_section_title(doc, "一、小组讨论规则")

    rules = [
        ("规则一：安全空间", "讨论内容仅限本小组内分享，不外传。每个人分享的真实经历会受到保护，不得嘲笑或评判。"),
        ("规则二：真实分享", "鼓励分享真实经历和真实想法，哪怕是不太光彩的想法。真实是成长的前提。"),
        ("规则三：保密原则", "小组成员分享的个人经历，讨论结束后不得对外透露。如有违规，取消继续参与资格。"),
        ("规则四：平等尊重", "无论职务高低、资历深浅，每人享有同等的发言权。禁止打断他人发言。"),
        ("规则五：建设性反馈", "反馈聚焦于情境和决策本身，不针对个人进行人身攻击或道德审判。"),
        ("规则六：限时发言", "每次发言不超过2分钟，超时提醒，不打断，让每个人都有机会表达。"),
    ]

    rules_tbl = doc.add_table(rows=len(rules)+1, cols=2)
    rules_tbl.style = 'Table Grid'
    set_cell_borders(rules_tbl)
    rules_hdrs = ["规则", "具体要求"]
    for ci, h in enumerate(rules_hdrs):
        cell = rules_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)

    for ri, (rule, req) in enumerate(rules, 1):
        rules_tbl.cell(ri, 0).text = rule
        rules_tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        rules_tbl.cell(ri, 1).text = req
        for col in range(2):
            for run in rules_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Part 2: Guiding questions
    add_section_title(doc, "二、情景讨论引导问题")

    add_body_para(doc, "每个情景讨论时，按照以下层次递进提问：")
    doc.add_paragraph()

    question_stages = [
        ("第一层：情境还原（5分钟）", [
            "这个情景中，主要涉及哪些人物？",
            "当事人面临的选择是什么？",
            "如果是你，你会怎么选？",
        ]),
        ("第二层：模式识别（8分钟）", [
            "当事人的决策过程中，可能出现了哪种合理化模式？",
            "当事人可能对自己说了哪些"合理化"的话？",
            "这种合理化模式在现实中常见吗？",
        ]),
        ("第三层：后果分析（5分钟）", [
            "如果当事人按情景中的方式继续，会产生什么后果？",
            "这些后果对个人、家庭、组织、社会分别有什么影响？",
            "从长远来看，这种"收益"和"代价"相比，值得吗？",
        ]),
        ("第四层：对策设计（7分钟）", [
            "当事人应该如何打破这种合理化思维？",
            "如果你是当事人的朋友或同事，你会怎么劝他？",
            "除了拒绝，当事人还有哪些更好的选择？",
        ]),
    ]

    for stage_title, questions in question_stages:
        stage_tbl = doc.add_table(rows=1, cols=2)
        stage_tbl.style = 'Table Grid'
        set_cell_borders(stage_tbl)
        stage_tbl.cell(0, 0).text = stage_title
        set_cell_bg(stage_tbl.cell(0, 0), 'D6DCE5')
        stage_tbl.cell(0, 0).paragraphs[0].runs[0].font.bold = True

        q_text = "\n".join(f"{i+1}. {q}" for i, q in enumerate(questions))
        stage_tbl.cell(0, 1).text = q_text
        for col in range(2):
            for run in stage_tbl.cell(0, col).paragraphs[0].runs:
                run.font.size = Pt(10.5)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph()

    # Part 3: Consensus
    add_section_title(doc, "三、小组共识形成方法")

    add_body_para(doc, "当小组成员出现分歧时，按以下步骤形成共识：")
    doc.add_paragraph()

    steps = [
        ("步骤一：各自陈述（每人2分钟）", "分歧双方各用2分钟陈述自己的观点和理由，其他人认真聆听，不打断。"),
        ("步骤二：澄清提问（每人1分钟）", "每方针对对方观点提问一个澄清性问题，确保理解对方真实意思。"),
        ("步骤三：寻找共同点（3分钟）", "小组共同找出双方观点的共同之处，作为共识基础。"),
        ("步骤四：差异分析（3分钟）", "小组讨论双方观点的差异，分析各自合理性和局限性。"),
        ("步骤五：综合共识（2分钟）", "小组形成综合性共识意见，可以是"我们认为最合理的做法是……""),
        ("步骤六：记录呈现", "书记员将共识记录在讨论记录表上，并向全班汇报。"),
    ]

    cons_tbl = doc.add_table(rows=len(steps)+1, cols=2)
    cons_tbl.style = 'Table Grid'
    set_cell_borders(cons_tbl)
    cons_hdrs = ["步骤", "内容"]
    for ci, h in enumerate(cons_hdrs):
        cell = cons_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)

    for ri, (step, content) in enumerate(steps, 1):
        cons_tbl.cell(ri, 0).text = step
        cons_tbl.cell(ri, 0).paragraphs[0].runs[0].font.bold = True
        cons_tbl.cell(ri, 1).text = content
        for col in range(2):
            for run in cons_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Discussion record
    add_section_title(doc, "四、小组讨论记录表")

    record_tbl = doc.add_table(rows=7, cols=4)
    record_tbl.style = 'Table Grid'
    set_cell_borders(record_tbl)
    rec_hdrs = ["情景", "主要讨论观点", "共识结论", "未达成共识的分歧点"]
    for ci, h in enumerate(rec_hdrs):
        cell = record_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    for ri in range(1, 7):
        record_tbl.cell(ri, 0).text = f"情景{ri}"
        for col in range(4):
            record_tbl.cell(ri, col).paragraphs[0].runs[0].font.size = Pt(10) if record_tbl.cell(ri, col).paragraphs[0].runs else None
            for run in record_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【使用建议】每组4-6人，设书记员1名。小组长确保每人发言，讨论时间控制在40分钟内。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# FILE 8: 练习材料03-课后行动作业.docx
# ════════════════════════════════════════════════════════════════════════════

def create_file8(doc):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_title_para(doc, "廉政风险情景决策训练营", size=14, color=(47, 84, 150))
    add_title_para(doc, "课后行动作业", size=20, bold=True, color=(47, 84, 150))
    add_body_para(doc, "每日自检7天 · 每周复盘模板 · 成果汇报模板")
    doc.add_paragraph()

    # Part 1: 7-day self-check
    add_section_title(doc, "一、每日廉政风险自检（7天）")

    add_body_para(doc, "使用《个人风险自检工具》，每日上岗前、午餐后、下班前各完成一次自检，在下表记录异常情况。")
    doc.add_paragraph()

    daily_tbl = doc.add_table(rows=9, cols=5)
    daily_tbl.style = 'Table Grid'
    set_cell_borders(daily_tbl)

    daily_hdrs = ["日期", "时段（早/午/晚）", "是否有风险提示语出现", "具体情况描述", "处理方式"]
    for ci, h in enumerate(daily_hdrs):
        cell = daily_tbl.cell(0, ci)
        cell.text = h
        set_cell_bg(cell, '2F5496')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(10)

    for ri in range(1, 9):
        daily_tbl.cell(ri, 0).text = f"第{ri}天"
        for col in range(5):
            daily_tbl.cell(ri, col).paragraphs[0].runs[0].font.size = Pt(10) if daily_tbl.cell(ri, col).paragraphs[0].runs else None
            for run in daily_tbl.cell(ri, col).paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()

    # Part 2: Weekly review
    add_section_title(doc, "二、每周复盘模板（第____周）")

    weekly_fields = [
        ("本周风险情境回顾", [
            "本周是否遇到涉及廉政风险的决策情境？  □ 是    □ 否",
            "如果"是"，请简要描述情境：",
            "当时你运用了哪种合理化模式？",
        ]),
        ("模式识别练习", [
            "本周你最常出现的合理化模式是：□ 模式一  □ 模式二  □ 模式三  □ 模式四",
            "这种模式在你过去的工作/生活中出现过几次？",
            "你是否成功识别并抵制了这种合理化模式？  □ 是    □ 否",
        ]),
        ("学习收获", [
            "本周学习对你影响最大的一点是什么？",
            "你计划在工作中做出什么改变？",
            "你希望获得什么支持（组织/同事/家人）？",
        ]),
        ("下周计划", [
            "下周你将如何应用课程所学？",
            "下周你计划进行一次决策情境回顾吗？  □ 是    □ 否",
            "下周你打算与谁分享你的学习收获？",
        ]),
    ]

    for sec_idx, (sec_title, fields) in enumerate(weekly_fields, 1):
        sec_tbl = doc.add_table(rows=1, cols=1)
        sec_tbl.style = 'Table Grid'
        set_cell_borders(sec_tbl)
        sec_hdr = sec_tbl.cell(0, 0)
        sec_hdr.text = sec_title
        set_cell_bg(sec_hdr, 'D6DCE5')
        sec_hdr.paragraphs[0].runs[0].font.bold = True
        sec_hdr.paragraphs[0].runs[0].font.size = Pt(11)

        field_tbl = doc.add_table(rows=len(fields), cols=1)
        field_tbl.style = 'Table Grid'
        set_cell_borders(field_tbl)
        for ri, f in enumerate(fields):
            field_tbl.cell(ri, 0).text = f
            field_tbl.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(11)
            field_tbl.cell(ri, 0).paragraphs[0].runs[0].font.name = '微软雅黑'
            field_tbl.cell(ri, 0).paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # Answer space
        ans_tbl = doc.add_table(rows=3, cols=1)
        ans_tbl.style = 'Table Grid'
        set_cell_borders(ans_tbl)
        for r in range(3):
            ans_tbl.cell(r, 0).text = ""
        doc.add_paragraph()

    # Part 3: Outcome report
    add_section_title(doc, "三、成果汇报模板")

    add_body_para(doc, "以下模板用于课程结束时的成果汇报，每项限时3分钟：")
    doc.add_paragraph()

    report_sections = [
        {
            "title": "汇报结构（总分总）",
            "fields": [
                "【开篇】一句话概括你最大的改变或收获",
                "【背景】课程前你在廉政风险方面面临的最大挑战是什么？",
                "【过程】课程中哪个情景/工具对你触动最大？为什么？",
                "【行动】课程结束后你做了哪些具体行动？（举1-2个实例）",
                "【数据】这些行动带来了什么可衡量的变化？（如：自检次数、拒绝利益输送次数等）",
                "【感悟】一句话总结你的成长感悟",
                "【致谢】感谢谁对你的帮助和支持",
            ]
        },
        {
            "title": "汇报评分标准",
            "fields": [
                "内容真实性（30%）：是否基于真实经历和行动",
                "模式识别能力（25%）：能否准确识别自身合理化模式",
                "行动具体性（25%）：行动是否具体、可落地",
                "表达清晰性（20%）：语言是否简洁、有条理",
            ]
        },
        {
            "title": "同伴互评表",
            "fields": [
                "姓名：__________  汇报人：__________  日期：__________",
                "评估维度          得分（1-5）  具体说明",
                "内容真实性：__________  ____________________",
                "模式识别能力：__________  ____________________",
                "行动具体性：__________  ____________________",
                "表达清晰性：__________  ____________________",
                "综合建议：____________________",
            ]
        },
    ]

    for rep_sec in report_sections:
        rep_tbl = doc.add_table(rows=1, cols=1)
        rep_tbl.style = 'Table Grid'
        set_cell_borders(rep_tbl)
        rep_hdr = rep_tbl.cell(0, 0)
        rep_hdr.text = rep_sec["title"]
        set_cell_bg(rep_hdr, 'D6DCE5')
        rep_hdr.paragraphs[0].runs[0].font.bold = True
        rep_hdr.paragraphs[0].runs[0].font.size = Pt(11)

        rep_items_tbl = doc.add_table(rows=len(rep_sec["fields"]), cols=1)
        rep_items_tbl.style = 'Table Grid'
        set_cell_borders(rep_items_tbl)
        for ri, f in enumerate(rep_sec["fields"]):
            rep_items_tbl.cell(ri, 0).text = f
            rep_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.size = Pt(10.5)
            rep_items_tbl.cell(ri, 0).paragraphs[0].runs[0].font.name = '微软雅黑'
            rep_items_tbl.cell(ri, 0).paragraphs[0].runs[0]._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        doc.add_paragraph()

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("【提交要求】7天自检记录于课程结束后3天内提交；每周复盘于每周日晚提交；成果汇报在课程最后半天进行。")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    return doc


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import os
    os.makedirs(OUTPUT_TOOLS, exist_ok=True)
    os.makedirs(OUTPUT_PRACTICE, exist_ok=True)

    files = [
        (os.path.join(OUTPUT_TOOLS, "工具集锦-个人风险自检工具.docx"), create_file1),
        (os.path.join(OUTPUT_TOOLS, "工具集锦01-合理化模式识别量表.docx"), create_file2),
        (os.path.join(OUTPUT_TOOLS, "工具集锦02-决策情境回顾表.docx"), create_file3),
        (os.path.join(OUTPUT_TOOLS, "工具集锦03-承诺与监督表.docx"), create_file4),
        (os.path.join(OUTPUT_TOOLS, "工具集锦合集.docx"), create_file5),
        (os.path.join(OUTPUT_PRACTICE, "练习材料01-情景决策练习册.docx"), create_file6),
        (os.path.join(OUTPUT_PRACTICE, "练习材料02-小组讨论指南.docx"), create_file7),
        (os.path.join(OUTPUT_PRACTICE, "练习材料03-课后行动作业.docx"), create_file8),
    ]

    for path, creator in files:
        print(f"Creating: {path}")
        d = creator(Document()) if creator != create_file5 else creator()
        d.save(path)
        print(f"  Saved: {path}")

    print("\nAll files created successfully!")

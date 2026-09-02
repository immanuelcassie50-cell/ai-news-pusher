#!/usr/bin/env python3
"""Extract text from all student materials for grading."""

import os
import sys
from pathlib import Path
from docx import Document
import zipfile
import re

def extract_docx(path):
    """Extract text from docx file."""
    try:
        doc = Document(path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return '\n'.join(texts)
    except Exception as e:
        return f"ERROR: {e}"

def extract_pptx(path):
    """Extract text from pptx file."""
    try:
        text_content = []
        with zipfile.ZipFile(path, 'r') as z:
            for name in z.namelist():
                if name.startswith('ppt/slides/slide') and name.endswith('.xml'):
                    content = z.read(name).decode('utf-8', errors='replace')
                    # Extract text between <a:t> tags
                    texts = re.findall(r'<a:t[^>]*>([^<]+)</a:t>', content)
                    text_content.extend(texts)
        return '\n'.join(text_content)
    except Exception as e:
        return f"ERROR: {e}"

def extract_pptx_notes(path):
    """Extract notes from pptx file."""
    try:
        notes_content = []
        with zipfile.ZipFile(path, 'r') as z:
            for name in z.namelist():
                if name.startswith('ppt/notesSlides/notesSlide') and name.endswith('.xml'):
                    content = z.read(name).decode('utf-8', errors='replace')
                    texts = re.findall(r'<a:t[^>]*>([^<]+)</a:t>', content)
                    notes_content.extend(texts)
        return '\n'.join(notes_content)
    except Exception as e:
        return f"ERROR: {e}"

base = Path("D:/2026年课程/云南磷化/第一阶段作业/第一组作业/第一组作业")

students = {
    "赵鹤然": "第一期作业-第一组-赵鹤然/第一期作业—第一组—赵鹤然",
    "钱开瑞": "第一期作业-第一组-钱开瑞/第一期作业，第一组，钱开瑞",
    "刘思杨": "第一期作业-第一组-刘思杨",
    "张迎吕": "第一期作业-第一组-张迎吕",
    "杨春": "第一期作业-第一组-杨春",
    "殷石才": "第一期作业-第一组-殷石才",
    "马崇伟": "第一期作业-第一组-马崇伟",
    "高鹏江": "第一期作业-第一组-高鹏江",
}

output_dir = Path("D:/CC/temp/grading/texts")
output_dir.mkdir(parents=True, exist_ok=True)

for student, subdir in students.items():
    student_dir = base / subdir
    student_out = output_dir / student
    student_out.mkdir(parents=True, exist_ok=True)

    print(f"\n=== {student} ===")

    # Find all docx and pptx files
    for fpath in sorted(student_dir.rglob("*")):
        if fpath.is_file():
            fname = fpath.name
            if fname.endswith('.docx') or fname.endswith('.pptx'):
                # Create short name based on file type
                if '定位' in fname:
                    short = "定位表"
                elif '大纲' in fname:
                    short = "大纲"
                elif 'PPT' in fname or 'pptx' in fname:
                    short = "PPT"
                elif '百问' in fname:
                    short = "百问百答"
                elif '测试' in fname:
                    short = "测试题"
                elif '录音' in fname:
                    short = "录音文稿"
                else:
                    short = fname[:20]

                out_file = student_out / f"{short}.txt"

                if fname.endswith('.docx'):
                    content = extract_docx(str(fpath))
                else:
                    content = extract_pptx(str(fpath))

                with open(out_file, 'w', encoding='utf-8') as out:
                    out.write(content)
                print(f"  {short}: {len(content)} chars")

print("\n\nDone! All files extracted.")

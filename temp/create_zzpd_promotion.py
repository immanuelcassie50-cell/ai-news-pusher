# -*- coding: utf-8 -*-
"""
创建《政治判断力情景决策训练营》对外宣传文案
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, cn_font='微软雅黑', en_font='Arial', size=12, bold=False, color=None):
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, cn_font='微软雅黑', size=16, bold=True, color=(0, 51, 102), space_before=12, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold, color=color)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para

def add_body_para(doc, text, cn_font='微软雅黑', size=11, bold=False, indent=False, space_before=3, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size, bold=bold)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent:
        para.paragraph_format.first_line_indent = Pt(22)
    return para

def add_bullet(doc, text, cn_font='微软雅黑', size=11):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_run_font(run, cn_font=cn_font, size=size)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    return para

def create_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header cell background to dark blue
        tc = header_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '003366')
        tcPr.append(shd)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for para in row_cells[col_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10.5)

    return table

def create_docx():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ==================== 封面部分 ====================
    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_para.paragraph_format.space_after = Pt(12)
    run = title_para.add_run('《政治判断力情景决策训练营》')
    set_run_font(run, cn_font='微软雅黑', size=28, bold=True, color=(0, 51, 102))

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(6)
    subtitle_para.paragraph_format.space_after = Pt(36)
    run = subtitle_para.add_run('——在两难抉择中显性化你的判断力')
    set_run_font(run, cn_font='微软雅黑', size=18, bold=False, color=(80, 80, 80))

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(12)
    sub_para.paragraph_format.space_after = Pt(48)
    run = sub_para.add_run('精品课程认证班')
    set_run_font(run, cn_font='微软雅黑', size=22, bold=True, color=(0, 102, 153))

    # 分隔线
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.paragraph_format.space_before = Pt(12)
    line_para.paragraph_format.space_after = Pt(48)
    run = line_para.add_run('─' * 40)
    set_run_font(run, cn_font='Arial', size=12, color=(0, 51, 102))

    # ==================== 课程定位 ====================
    add_heading(doc, '课程定位语', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    loc_para = doc.add_paragraph()
    loc_para.paragraph_format.space_before = Pt(6)
    loc_para.paragraph_format.space_after = Pt(12)
    run = loc_para.add_run(
        '不讲"如何提高政治判断力"，只做一件事：\n'
        '设计一系列真实两难场景，把干部的判断力逼出来、暴露出来、然后当场纠偏。'
    )
    set_run_font(run, cn_font='微软雅黑', size=14, bold=True, color=(0, 0, 0))
    loc_para.paragraph_format.left_indent = Pt(24)

    # ==================== 目标受众痛点 ====================
    add_heading(doc, '目标受众痛点', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    pain_points = [
        '业务能力普遍不弱，缺的不是知识，而是"遇到真实两难时的取舍直觉"',
        '这种直觉此前从未被训练过，只能靠自己"悟"，悟错了才知道',
        '新提拔或后备中层干部，正处于"凡提必训"节点上',
        '已在岗但组织认为其判断力有待观察/培养的中层干部',
        '国有企业、金融机构、大型制造业等有明确干部任职培训需求的单位'
    ]
    for point in pain_points:
        add_bullet(doc, point)

    # ==================== 课程特色 ====================
    add_heading(doc, '课程特色', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    features = [
        ('没有标准答案', '每个情景关卡都经过精心设计，选项之间无明显对错之分，训练的是判断力而非记忆力'),
        ('压力叠加设计', '每个关卡叠加至少两种压力源（时间/人际/信息不全/利益诱惑），还原真实决策环境'),
        ('代价可见', '学员做出选择后，讲师揭示该选择可能引发的后续连锁反应，让后果"提前发生"'),
        ('追问技术', '通过三层递进追问（还原逻辑→暴露盲区→迁移内化），让学员自己看见自己判断逻辑的盲区')
    ]

    for feat_title, feat_desc in features:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Pt(24)
        run1 = para.add_run('◆ ' + feat_title + '：')
        set_run_font(run1, cn_font='微软雅黑', size=12, bold=True, color=(0, 102, 153))
        run2 = para.add_run(feat_desc)
        set_run_font(run2, cn_font='微软雅黑', size=11, bold=False, color=(0, 0, 0))

    # ==================== 五维能力模型介绍 ====================
    add_heading(doc, '五维能力模型', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    ability_headers = ['能力维度', '定义', '典型训练场景']
    ability_rows = [
        ['全局意识', '能否跳出局部利益看到组织整体影响', '局部业绩 vs 组织声誉/长远利益的冲突场景'],
        ['底线意识', '能否在压力下守住不能碰的红线', '业绩诱惑/领导施压/时间紧迫三重压力叠加场景'],
        ['分寸感', '能否在灵活处理和坚持原则之间找到恰当的度', '合规边缘的"擦边球"请求如何恰当应对'],
        ['舆情敏感度', '能否预判一个决策可能引发的舆论/群体反应', '业务合理但传播出去可能引发误解的决策场景'],
        ['请示汇报时机', '能否判断什么事该自己扛、什么事必须及时上报', '信息不完整、时间压力大的紧急决策场景']
    ]
    create_table_with_header(doc, ability_headers, ability_rows)

    # ==================== 学员收获 ====================
    add_heading(doc, '学员收获', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    benefits = [
        '经历一场真实的两难决策训练，让判断力在压力下自然显形',
        '被追问技术逼出自己判断逻辑的盲区，当场获得纠偏',
        '提炼一份个人化的"判断力自检清单"，带走可迁移的判断原则',
        '理解政治判断力五个可训练维度，知道自己在哪个维度上有优势、在哪个维度上容易犯迷糊',
        '体验"没有标准答案"的决策场景，建立对复杂政治情境的心理预期'
    ]
    for benefit in benefits:
        add_bullet(doc, benefit)

    # ==================== 适合场景 ====================
    add_heading(doc, '适合场景', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    scenarios = [
        ('干部任职培训', '"凡提必训"节点上的新提拔或后备中层干部'),
        ('判断力专项训练', '组织认为判断力有待观察/培养的中层干部'),
        ('党政干部能力提升', '国有企业、金融机构、大型制造业等单位的干部培训')
    ]

    for scenario_title, scenario_desc in scenarios:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Pt(24)
        run1 = para.add_run('◆ ' + scenario_title + '：')
        set_run_font(run1, cn_font='微软雅黑', size=12, bold=True, color=(0, 102, 153))
        run2 = para.add_run(scenario_desc)
        set_run_font(run2, cn_font='微软雅黑', size=11, bold=False, color=(0, 0, 0))

    # ==================== 报名信息 ====================
    add_heading(doc, '报名信息', level=1, size=16, color=(0, 51, 102), space_before=24, space_after=12)

    info_items = [
        ('课程时长', '1天（6-7课时）'),
        ('建议人数', '每场24人以内（4-6组，每组4-5人）'),
        ('课程形式', '情景沉浸 + 追问复盘 + 行为观察'),
        ('核心人群', '新提拔或后备中层干部'),
        ('次要人群', '已在岗但组织认为其判断力有待观察/培养的中层干部')
    ]

    info_table = doc.add_table(rows=len(info_items), cols=2)
    info_table.style = 'Table Grid'
    for i, (label, value) in enumerate(info_items):
        row_cells = info_table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        for j, cell in enumerate(row_cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(11)
                    if j == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 51, 102)
            if i % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F5FA')
                tcPr.append(shd)

    # Save
    output_path = 'D:/新课开发/党业融合/政治判断力/完整课程包/004-对外宣传文案/认证班宣传文案-政治判断力.docx'
    doc.save(output_path)
    print(f'Created: {output_path}')

if __name__ == '__main__':
    create_docx()

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)
    
    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)
    
    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)
    
    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)
    
    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)
    
    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

chapters = [
    {
        'title': '第七章  场景映射矩阵',
        'gold_quote': '卡片越通用，越没人用；卡片越具体，用的人越多。',
        'objectives': ['理解通用性与具体性的悖论', '掌握场景映射矩阵的设计方法', '学会找出真正影响判断的场景变量'],
        'concepts': [('场景映射矩阵', '把同一条底层判断逻辑，拆解成几个具体场景各自版本'), ('场景变量', '影响决策卡关键内容是否变化的场景维度'), ('维护者地图 vs 使用者工具', '矩阵是给维护者用的地图，具体场景的卡才是使用者真正会打开的')],
        'key_points': ['越往通用方向改，卡片话越空洞', '不做一张覆盖所有场景的卡，而是做场景映射矩阵', '判断变量该不该进矩阵：换掉它，关键内容会不会变？', '每个场景格子应该是可独立使用的小卡片', '先做粗粒度版本，投入使用后收集反馈再细分'],
        'thinking': ['你见过的那种万能决策流程或模板，实际使用效果如何？', '为什么越是通用的东西，实际用起来越不好用？'],
        'exercise': ['选择一个业务场景，列出可能影响判断的所有变量，筛选出真正需要区分的场景变量'],
        'summary': '一张卡管不住所有场景，硬要它管，卡片就会变得空洞到没人愿意用。场景映射矩阵解决的是让同一条底层判断逻辑，在不同场景里长出各自具体的样子。'
    },
    {
        'title': '第八章  失败案例是用来对照的',
        'gold_quote': '讲失败案例的目的不是让人害怕，是让人在自己身上找到那个岔路口。',
        'objectives': ['理解失败案例的正确用法是对照而非警示', '掌握失败案例警示清单的三段式写法', '学会把警示分散嵌入到对应的决策卡里'],
        'concepts': [('警示清单', '描述失败案例的关键岔路口、被忽略的信号，供使用者对照自查'), ('岔路口', '决策者面临选择的关键节点'), ('嵌入式警示', '将警示内容分散嵌入对应决策卡，不集中成册')],
        'key_points': ['讲失败案例的初衷是让人看到自己可能踩中的岔路口', '不是让人变得畏手畏脚', '三段式：岔路口描述→本可注意到的信号→对照当下的自查问题', '损失中等的平淡案例比极端大损失案例更有效', '警示清单不要单独成册，分散嵌入对应决策卡'],
        'thinking': ['你听过或经历过的失败案例，有没有让你产生这种事不会发生在我身上的想法？', '什么样的失败案例讲述方式最容易让人产生代入感？'],
        'exercise': ['找一个中等损失的失败案例，写出它的警示清单三段式内容'],
        'summary': '失败案例不是用来吓人的，是用来对照的。讲清楚岔路口在哪、被忽略的信号是什么，再逼使用者问自己一句我现在是不是也站在类似的地方，这份清单才真正转化成了行动。'
    },
    {
        'title': '第九章  第一个读者应该是反对你的人',
        'gold_quote': '卡片评审不是找错别字，是找这句话在什么情况下会害死人。',
        'objectives': ['理解评审姿态的重要性：找漏洞而非找认同', '掌握决策卡评审的三个检验维度', '学会设计双轮评审制度'],
        'concepts': [('双轮评审制', '第一轮由熟悉领域但没参与原始决策的人挑漏洞，第二轮由完全不了解领域的人检验表述'), ('反对者评审', '专门去找这句话什么时候会出错'), ('边界情况检验', '检查触发条件阈值卡在线上等边界情况')],
        'key_points': ['让认同你的人看，只能收获掌声，堵不住漏洞', '请有经验但没参与原始决策的人来评审', '三个检验维度：检查表有无可能变成误导？触发条件边界情况？应急方案在变体场景是否成立？', '修改完的版本要拿回给提反对意见的人再看一遍', '第二轮由完全不熟悉领域的人来检验表述是否清晰'],
        'thinking': ['你写完一份重要文档后，通常找谁来评审？', '为什么找认同的人评审很难发现问题？'],
        'exercise': ['为你的一张决策卡组织一次反对者评审，记录提出的问题并修改'],
        'summary': '你写的决策卡，第一个读者应该是反对你的人。让认同你的人看，只能收获掌声，掌声堵不住那个会在特殊场景下害人的漏洞。'
    },
    {
        'title': '第十章  训练是让人在场景里犯一次错',
        'gold_quote': '讲道理讲三小时，不如让他在模拟场景里做错一次决策。',
        'objectives': ['理解听懂和在压力情境下能想起来用是两种完全不同的能力', '掌握微课脚本的三环节设计', '理解延迟追踪的重要性'],
        'concepts': [('条件反射式调用', '在压力情境下能想起来使用决策卡的能力'), ('微课脚本', '用于训练使用者决策卡的三环节情境模拟教学设计'), ('延迟追踪', '在训练活动结束后数月跟进，检验真实决策现场是否真正使用卡片')],
        'key_points': ['测验成绩好不代表能在真实场景用上', '人对本来可以避免这个错的记忆比原则深刻得多', '三环节设计：情境代入→揭示与对照→二次决策', '第一环节不能提前暴露这是决策卡测试', '训练结束后要延迟追踪'],
        'thinking': ['你过去参加过的培训，有哪次是讲的时候觉得懂了，但后来遇到类似场景想不起来的？', '为什么单纯的讲解无法建立在压力下想起来用的能力？'],
        'exercise': ['设计一个三环节的微课脚本，训练他人使用某张决策卡'],
        'summary': '训练活动的目的不是讲透道理，是让人在场景里犯一次错。讲解让人知道，模拟让人经历，只有经历过那种如果早点看这张卡就好了的懊恼，这张卡才会真正嵌进使用者做判断的本能里。'
    },
]

for ch in chapters:
    add_chapter_content(doc, ch)
    doc.add_page_break()

print('S7: Chapters 7-10 done')
doc.save(r'D:\CC\temp\handbook_s7.docx')

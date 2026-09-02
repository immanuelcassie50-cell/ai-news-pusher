# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {})
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

output_path = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册/讲师手册-变革成果固化机制.docx"
doc = Document(output_path)

# ============ CHAPTER 1 ============
add_heading(doc, "第一章  课程整体设计理念", 1)

add_heading(doc, "1.1 课程背景", 2)
add_para(doc, "很多变革项目在推动者还在场的时候运转良好，推动者一调岗或离职，组织就悄悄退回了老做法。这说明变革的成果从来没有真正被固化进组织的制度和习惯里。")

add_heading(doc, "1.2 变革成果固化的本质", 2)
add_para(doc, "变革成果固化是指将变革项目产生的新的工作方式、流程、行为模式，从依赖特定个人或短期项目推动的状态，转化为不依赖特定个人、可长期持续运作的组织能力的过程。")

add_para(doc, "【关键认知】", bold=True)
add_para(doc, '固化不是"保存"，不是把东西放进冰箱。固化是"转化"——把变革成果从一种形态转化为另一种形态。')
add_para(doc, "变革成果的初始形态：依附于人的知识和能力")
add_para(doc, "变革成果的固化形态：脱离个人、嵌入组织的制度和习惯")

add_heading(doc, "1.3 固化的三个层次", 2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
headers = ['层次', '定义', '特点', '固化标志']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, 'D9D9D9')

data = [
    ['文件层', '制度、流程、标准、考核指标等形成文字的规范', '看得见、摸得着，但可能"写归写、做归做"', '新人入职后能通过阅读文件知道"应该怎么做"'],
    ['流程层', '日常运营中自动执行、无需每次重新决策的机制', '融入组织的运营节奏，不需要每次都重新推动', '即使没有人提醒，流程也会在正确的时间被执行'],
    ['文化层', '成为组织默认的行为方式和思维习惯', '不需要外部监督，新人会自动模仿', '新员工说"我们这里就是这样做的"，而不是"公司要求这样做"'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()
add_para(doc, "三个层次由浅入深，文件层是基础，流程层是保障，文化层是最高境界。真正的固化需要三个层次协同作用。")

add_heading(doc, "1.4 固化与坚持的本质区别", 2)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
headers2 = ['维度', '坚持', '固化']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, 'D9D9D9')

data2 = [
    ['动力来源', '个人意志', '组织系统'],
    ['持续性', '依赖个人状态', '依赖机制自动运转'],
    ['可扩展性', '线性衰减', '可以复制放大'],
    ['抗风险能力', '低（一人离职即失效）', '高（不依赖特定个人）'],
    ['成本', '高（持续消耗能量）', '低（一次性投入，长期回报）'],
    ['本质', '人治', '法治+文治'],
]
for row_idx, row_data in enumerate(data2):
    for col_idx, text in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = text

doc.add_paragraph()
add_para(doc, '【核心结论】固化是逃脱"坚持陷阱"的唯一出路。当固化机制建立后，推动者可以从"维持变革"中解放出来，去做真正有价值的新的变革。', bold=True)

add_heading(doc, "1.5 课程核心能力目标", 1)
add_para(doc, "1. 如何把变革成果写进制度、考核和流程文档而不只是停留在口头共识")
add_para(doc, "2. 如何设计不依赖某一个关键推动者也能持续运转的机制")
add_para(doc, "3. 如何定期检验固化效果，及时发现组织正在悄悄退回旧模式的苗头")

add_heading(doc, "1.6 课程教学方法", 1)

table3 = doc.add_table(rows=5, cols=2)
table3.style = 'Table Grid'
for i, h in enumerate(['教学方法', '占比']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(cell, 'D9D9D9')
for row_idx, (method, ratio) in enumerate([('理论讲授', '30%'), ('案例分析', '25%'), ('小组讨论', '20%'), ('实战演练', '25%')]):
    table3.rows[row_idx + 1].cells[0].text = method
    table3.rows[row_idx + 1].cells[1].text = ratio

doc.save(output_path)
print("Chapter 1 added")

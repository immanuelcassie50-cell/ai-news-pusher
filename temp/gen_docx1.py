# -*- coding: utf-8 -*-
"""
生成《学员路演指引.docx》
风格：浅色背景 + 深色头部标题 + 表格密集 + 紧凑排版
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# 颜色定义（深色头部用深蓝灰，浅色背景）
HEADER_COLOR = "1F3A5F"  # 深蓝灰（头部）
ACCENT_COLOR = "C8102E"  # 红色强调
TEXT_DARK = "1A1A1A"     # 主文字
TEXT_GRAY = "555555"     # 次要文字
BG_LIGHT = "F5F7FA"      # 浅色背景
TABLE_HEADER_BG = "1F3A5F"  # 表头背景
TABLE_ALT_BG = "EEF1F5"     # 交替行背景

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置页面
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_header_para(text, level=1, color=HEADER_COLOR, size=None):
    """添加带颜色的标题"""
    if size is None:
        size = {1: 22, 2: 16, 3: 13, 4: 11}.get(level, 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        # 加下边框线
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:color'), HEADER_COLOR)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_text(text, size=10.5, bold=False, color=TEXT_DARK, indent=0, space_after=4):
    """添加正文"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def add_bullet(text, size=10.5):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p


def shade_cell(cell, color_hex):
    """给单元格上背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color="FFFFFF", size=10, align='left'):
    """设置单元格文字"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(headers, rows, col_widths=None, header_bg=TABLE_HEADER_BG, alt_bg=TABLE_ALT_BG):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        shade_cell(cell, header_bg)
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=10, align='center')
    # 数据行
    for i, row in enumerate(rows):
        bg = alt_bg if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            shade_cell(cell, bg)
            set_cell_text(cell, val, color=TEXT_DARK, size=9.5, align='left')
    # 列宽
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


# ============== 标题页 ==============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(40)
title.paragraph_format.space_after = Pt(8)
r = title.add_run('顺造科技AI项目成果评审')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(TEXT_GRAY)
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2.paragraph_format.space_after = Pt(8)
r = title2.add_run('学员路演指引')
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(HEADER_COLOR)
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(40)
r = sub.add_run('（小白版）')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor.from_string(ACCENT_COLOR)
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = quote.add_run('写给每位准备上台的学员：你做的这件事值得被看见。这份指南，帮你把它说清楚。')
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor.from_string(TEXT_GRAY)
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ============== 第一部分：开篇寄语 ==============
add_header_para('第一部分：开篇寄语', level=1)

add_header_para('写给你的一段话', level=3)

add_text('你已经走到这一步了。')
add_text('从对AI几乎一无所知，到亲手做出一个能在自己部门落地的工具——这件事本身就是成果。但成果不会自己被看见。')
add_text('接下来这10分钟的路演 + 3-5分钟Q&A，是你把"我做的事情"变成"领导愿意支持的事情"的关键节点。')
add_text('你可能会紧张。这很正常。但请记住三件事：')

add_text('第一，你不是去汇报工作的，是去分享一个成果的。', bold=True, color=ACCENT_COLOR)
add_text('路演不是答辩。领导坐在那里，不是为了挑你毛病，是为了看清楚——这件事值不值得他投入更多关注。')

add_text('第二，你已经做完了最难的部分。', bold=True, color=ACCENT_COLOR)
add_text('做一个针对自己业务的AI工具，比你以为的难得多。你已经做出来了。剩下10分钟的表达，比起你过去几周的动手量，是轻活。')

add_text('第三，领导不是你的评委，是你的潜在资源。', bold=True, color=ACCENT_COLOR)
add_text('他不是来给你打分的，他是来评估"这件事和我有什么关系"的。你回答了这个问题，他就成了你往后做这件事的助力。')

add_text('接下来这份指南，会陪你走完这10分钟 + 5分钟Q&A的全部准备过程。', bold=True)
add_text('不需要背。只需要顺着读、跟着做。')

# ============== 第二部分：路演全流程 ==============
add_header_para('第二部分：路演全流程（10分钟）', level=1)

add_header_para('整体框架：痛 · 做 · 效 · 求', level=3)

add_table(
    headers=['步骤', '任务', '时间', '占比', '核心目标'],
    rows=[
        ['痛', '亮出问题', '90秒', '15%', '让领导"痛感同身受"'],
        ['做', '说你做了什么', '120秒', '20%', '让领导知道"我们做了什么"'],
        ['效', '展示真正的变化', '270秒', '45%', '让领导看到"效果和价值"'],
        ['求', '提出你的请求', '120秒', '20%', '让领导"愿意给资源/支持"'],
        ['合计', '—', '600秒', '100%', '—'],
    ],
    col_widths=[1.8, 3.0, 2.0, 1.5, 8.7]
)

add_text('最关键的一个分配原则：效果（45%）是重点，过程（20%）是辅助。', bold=True, color=ACCENT_COLOR)
add_text('大多数人容易把"做"讲得太详细、把"效"讲得太简单——反过来才对。')

# ---- 第1步：痛 ----
add_header_para('第1步：痛（15%，90秒）', level=2)

add_header_para('目标', level=3)
add_text('让领导感受到"这个问题我们真的有，它真的有影响"。')
add_text('不追溯历史，不讲行业大背景，只让领导一听就感受到：哦，这件事确实存在，确实造成了影响。')

add_header_para('必含内容', level=3)
add_text('① 一个具体数字（时间、次数、人·时、金额，任选其一）', bold=True)
add_bullet('"上个月，质量部门在报告审核上加班了超过30个人·时"')
add_bullet('"我们部门上季度返工了40份报销单"')
add_bullet('"这个月，客服日均处理工单270件，人工分类要花6小时"')
add_text('数字要真实、具体。"很多""大量"不算数字。', color=TEXT_GRAY)

add_text('② 一个场景描述（让领导能脑补画面）', bold=True)
add_bullet('"上个月X部门加班到凌晨2点，整层楼就剩两盏灯"')
add_bullet('"上周五下午4点，小李的报销单被退回来，她不知道超了哪条规定"')
add_bullet('"每个季度末，财务的桌面都堆着两百多份待核查单据"')

add_header_para('参考话术：3个开场白模板', level=3)
add_table(
    headers=['开场方式', '模板', '适用场景'],
    rows=[
        ['数字开场', '"领导，上个月我们在[某事]上花了[具体数字]。"', '成果有量化对比'],
        ['场景还原开场', '"[具体时间]，[谁]在做[什么事]，遇到了[什么情况]……"', '有生动的具体案例'],
        ['对比开场', '"之前，[某事]需要[时间/成本]。现在，[某事]需要[新时间/成本]。"', '想用最简洁方式直击主题'],
    ],
    col_widths=[2.8, 8.5, 5.7]
)

add_header_para('注意事项', level=3)
add_bullet('不要说"AI能解决"——领导此刻还没问"你怎么解决"')
add_bullet('要说"具体卡在哪"——是数量问题、流程问题、还是质量问题')
add_bullet('不要追溯"3年前部门扩张时期"——历史背景在前90秒不重要')
add_bullet('要说"上个月、上周、上周五"——时间越具体，越像真的')
add_bullet('不要"行业里类似挑战普遍存在"——领导不关心行业，只关心你')

add_header_para('一个错误示范 vs 正确示范', level=3)
add_text('错误示范（领导会走神）：', bold=True, color=ACCENT_COLOR)
add_text('"我们部门面临一个长期存在的工作效率问题，可以追溯到三年前的部门扩张期。根据行业研究，类似挑战在制造业普遍存在，主要表现为……"')

add_text('正确示范（领导会被钩住）：', bold=True, color="00733B")
add_text('"领导，我们质量部门上个月有11份供应商报告积压，两个工程师加班了4个晚上。这不是偶发——上个季度类似情况出现了3次。"')

# ---- 第2步：做 ----
add_header_para('第2步：做（20%，120秒）', level=2)

add_header_para('目标', level=3)
add_text('让领导知道你采取了什么行动——不需要搞清楚技术细节，需要看见真实的一个动作。')

add_header_para('必含内容', level=3)
add_text('① 一句话说明你做了什么', bold=True)
add_text('"我做了一个[工具/提示词/智能体/Skill/流程]，叫[名字]，用来[解决什么问题]。"')

add_text('② 用到的工具/技术清单（一句话带过）', bold=True)
add_text('"用到了[大模型/某工具/某平台]，搭配[提示词/智能体/知识库]。"')

add_text('③ 一个最具体的案例（60-90秒的60秒故事，见下文）', bold=True)

add_text('④ AI介入级别（一句话标注）', bold=True)
add_table(
    headers=['级别', '描述', '典型场景'],
    rows=[
        ['一级', '单点提示词', '让AI写一段邮件、改一段话'],
        ['二级', '多步结构化提示词', '一组提示词完成一个完整任务'],
        ['三级', '智能体', 'AI按预设流程自动完成多步操作'],
        ['四级', '工具调用+API', 'AI调用外部系统/数据'],
        ['五级', '完整Skill/知识库', '形成可复用的能力包'],
    ],
    col_widths=[1.5, 4.5, 11.0]
)

add_header_para('业务流程图（建议在HTML网页中可视化展示）', level=3)
add_text('至少包含两张图：')
add_bullet('当前流程图：原来的步骤、卡点、人·时消耗')
add_bullet('AI介入后流程图：AI替代/辅助了哪些步骤')

add_header_para('参考话术：3个介绍模板', level=3)
add_text('模板1：简洁版', bold=True)
add_text('"我做了一个[工具名]，专门处理[某类问题]。上周有一份[具体问题]的[单据/报告/工单]，以前要花[时间]，现在[时间]出结果。"')

add_text('模板2：案例+工具版', bold=True)
add_text('"针对[某痛点]，我设计了一个[智能体/Skill]。上周小王提交了一份[具体问题]的[材料]，我让他先用工具跑一遍，[具体输出]。原来要3天的事，当天就闭环了。"')

add_text('模板3：协作版', bold=True)
add_text('"我先看了部门手册里[某段流程]，发现[具体卡点]。然后用[工具+方法]做了一套[提示词/智能体]，组里现在每周用一次，已经用了一个月。"')

add_header_para('注意事项', level=3)
add_bullet('不要讲技术细节——"提示词怎么写""模型什么参数"')
add_bullet('要讲"用了哪些AI，做了什么"——是一句话，不是技术报告')
add_bullet('不要详细介绍所有功能——选一个最有冲击力的案例')
add_bullet('要让领导"看见"——"我把报告内容粘进去，90秒后屏幕上出现了一份清单"')
add_bullet('不要说"还在完善中"——先把你做到的最好的状态说清楚')

add_text('一句话：领导要的是什么', bold=True, color=ACCENT_COLOR)
add_text('领导要看到的不是"你造了一个工具"，而是"你用AI替代/加速了某段真实工作"。')

# ---- 第3步：效 ----
add_header_para('第3步：效（45%，270秒）', level=2)

add_header_para('目标', level=3)
add_text('让领导看见"之前是什么样，之后是什么样，变化对团队意味着什么"。')
add_text('这是整个10分钟里最重要的部分——时间给足。', bold=True, color=ACCENT_COLOR)

add_header_para('必含内容', level=3)
add_text('① 三层成果数字（基线→新状态→对领导的意义）', bold=True)
add_table(
    headers=['层级', '内容', '示例'],
    rows=[
        ['第一层', '之前的数字（基线）', '"以前每批次审核需要32人·时"'],
        ['第二层', '之后的数字（新状态）', '"现在需要8人·时"'],
        ['第三层', '意义（对领导的翻译）', '"这个季度，我们没出现一次因报告积压导致的加班"'],
    ],
    col_widths=[2.0, 5.0, 10.0]
)

add_text('② 1-2个真实故事案例', bold=True)
add_text('每个案例用"60秒故事结构"（背景10秒+转折10秒+行动20秒+结果20秒）讲完。')

add_text('③ HTML可视化展示', bold=True)
add_text('路演时必须打开HTML网页（或截图），让领导"看见"你的成果：')
add_bullet('A3海报风格的可视化')
add_bullet('流程图、对比图、关键数字')
add_bullet('不要纯文字堆砌')

add_header_para('参考话术：数字呈现模板', level=3)
add_text('模板1：节省时间', bold=True)
add_text('"每批次审核，之前两个工程师要花两整天，也就是32人·时。现在一个人半天搞定，大约8人·时。"')

add_text('模板2：提高准确率', bold=True)
add_text('"出错率从上个月的0.8%降到了0.2%。更重要的是，再没有出现不同工程师标准不一样的情况。"')

add_text('模板3：行为变化', bold=True)
add_text('"以前这种情况出现后，通常要来回3-4封邮件才确认完。上周同样的情况，当天就确认了。"')

add_text('模板4：累计数字', bold=True)
add_text('"这个季度到今天，我们累计处理了47份报告，节省了大约1200人·时的工作量。"')

add_header_para('60秒故事模板', level=3)
add_text('背景（10秒）——时间、角色、情境', bold=True)
add_text('"三周前，我们收到一份新供应商提交的8D报告……"')

add_text('转折（10秒）——具体问题，不解决会怎样', bold=True)
add_text('"报告的根因分析那一栏只写了四个字"操作不规范"，按以前的做法，工程师需要逐字审阅，整理问题，写反馈邮件，差不多要花一个多小时……"')

add_text('行动（20秒）——用了什么，怎么操作', bold=True)
add_text('"我把报告内容粘进工具，90秒之后，屏幕上出现了一份清单：D4的根因没有追溯到设备层，D5的临时措施和根因不在同一个逻辑层……"')

add_text('结果（20秒）——具体结果，最好有对比', bold=True)
add_text('"把清单发给供应商，他们当天就回复了修改版。以前这种来回，通常要3-5天。"')

add_header_para('注意事项', level=3)
add_bullet('不要说"效率显著提升"——这是空话')
add_bullet('要说"从3天缩到半天"——具体的对比数字')
add_bullet('不要说"用户反馈良好"——领导不知道"用户"是谁')
add_bullet('要说"同事小王说：' + chr(0x201C) + '以前月末最头疼，现在基本正常下班了' + chr(0x201D) + '"——具体的人+具体的感受')
add_bullet('不要说"当然这只是初步结果"——不要在台上主动给自己打折扣')
add_bullet('要"可视化展示HTML网页"——打开、截图、互动演示')
add_bullet('必须用估算/初步数字，但诚实说明"目前是估算，下个月会做更系统统计"——这比没有数字强100倍')

# ---- 第4步：求 ----
add_header_para('第4步：求（20%，120秒）', level=2)

add_header_para('目标', level=3)
add_text('让领导知道你需要他做一件具体的事——这件事要小、有理由、有时限。')
add_text('领导听完，必须知道他要做什么。', bold=True, color=ACCENT_COLOR)

add_header_para('必含内容', level=3)
add_text('① 一个具体请求', bold=True)
add_text('不是"请领导支持"，不是"希望多多关注"，是一件具体到领导知道怎么做的事。')

add_text('② 一句"因为"——说明这件事为什么值得做', bold=True)
add_text('让领导知道"支持你这件小事，对他也有什么好处"。')

add_text('③ 后续承诺', bold=True)
add_text('"我会在[X时间]给您反馈结果"——让领导知道这不是一个"撒手没"的请求。')

add_header_para('参考话术：3个请求模板', level=3)
add_table(
    headers=['模板', '适用场景', '示例'],
    rows=[
        ['时间型', '想要一次曝光', '"我需要您在下周部门例会上，给我5分钟让其他组看一下这个工具"'],
        ['资源型', '想要一个具体资源', '"我需要您批准一个账号，让财务部同事也能访问这个工具"'],
        ['试跑型', '想要一次小范围试验', '"我需要您允许我们组下个月先试跑一次，跑完了把结果给您看"'],
    ],
    col_widths=[2.0, 4.0, 11.0]
)

add_header_para('请求的"三要素"清单', level=3)
add_table(
    headers=['要素', '检查', '示例'],
    rows=[
        ['具体', '一件事，不是一个方向', '"下周二例会5分钟" ✓'],
        ['小', '领导容易答"是"', '"发一条群通知" ✓'],
        ['有理由', '领导知道为什么值得做', '"因为其他组工程师也遇到类似情况" ✓'],
    ],
    col_widths=[2.0, 5.0, 10.0]
)

add_header_para('注意事项', level=3)
add_bullet('不要"画大饼"——"希望领导把这个项目作为明年重点"（太大）')
add_bullet('要"小步快跑"——"先让其他组来看15分钟，看有没有用"（小）')
add_bullet('不要"请领导多多支持"——这句话领导听完不知道要做什么')
add_bullet('要"我需要您帮我做一件事：[具体描述]"——有动词、有对象')
add_bullet('不要"希望以后能有更多机会"——这是感慨，不是请求')
add_bullet('要"这件事如果跑通了，下个月我整理成数据给您看"——后续承诺')

# ============== 第三部分：Q&A准备 ==============
add_header_para('第三部分：Q&A准备（3-5分钟）', level=1)

add_text('路演结束后，领导会有3-5分钟提问。这是好事——问题越多，说明他越感兴趣。', bold=True)

add_text('应对原则：', bold=True)
add_bullet('不懂就承认，不要现编')
add_bullet('用"我目前掌握的数据是……更准确的我去核实后回复您"')
add_bullet('提问是表达兴趣，不是刁难')

add_header_para('5类常见问题速查', level=3)
add_table(
    headers=['问题类型', '一句话应对'],
    rows=[
        ['追问数字来源', '"这个数字我们这批次手工统计的，下个月做更系统统计"'],
        ['问推广可行性', '"适配一个新场景大概需要1-2周"'],
        ['问数据安全', '"我们只用文字内容，没涉及客户信息"'],
        ['问"为什么以前没做"', '"之前没这样便捷的工具，这次学习才发现可以这样优化"'],
        ['完全没想到的问题', '"这是个好问题，我明天给您书面回复"'],
    ],
    col_widths=[5.0, 12.0]
)

add_text('详细话术和"问题本质分析 + 错误示范 + 正确示范"对照表，见《Q&A应答指南》。', color=TEXT_GRAY)

# ============== 第四部分：彩排建议 ==============
add_header_para('第四部分：彩排建议', level=1)

add_header_para('6条核心建议', level=3)

add_text('① 聚焦一个最痛的点', bold=True, color=ACCENT_COLOR)
add_text('10分钟里只讲一件事。不要"我做了ABCDE五个事"——领导记不住。')
add_text('挑一个最痛的部门场景，狠狠讲透。其他成果一笔带过即可。')

add_text('② 形式多样（数字/故事/对比/可视化）', bold=True, color=ACCENT_COLOR)
add_text('不要全是文字。让领导看到：')
add_bullet('数字（"32人·时 → 8人·时"）')
add_bullet('故事（小王/小李的真实经历）')
add_bullet('对比（"之前需要3天，现在90秒"）')
add_bullet('可视化（HTML网页、流程图、截图）')

add_text('③ 有亮点（与别人不一样的地方）', bold=True, color=ACCENT_COLOR)
add_text('你有什么是别人没做的？')
add_bullet('你做了一个完整的Skill？（别人是简单提示词）')
add_bullet('你跑通了真实业务流程？（别人只是demo）')
add_bullet('你有真实数字？（别人是估算）')
add_bullet('你做了一个可复用的工具？（别人是单点使用）')
add_text('找到你的"不一样"，放在"做"或"效"的高潮处说出来。')

add_text('④ 开场结尾精心设计', bold=True, color=ACCENT_COLOR)
add_bullet('第一句话 = 钩子（让领导"被拉进场景"）')
add_bullet('最后一句话 = 落点（让领导知道"他要做什么"）')
add_bullet('这两句话改了再改，演练到能脱口而出')

add_text('⑤ 完全进入"AI应用专家"角色', bold=True, color=ACCENT_COLOR)
add_bullet('不是"我刚学的"——是"我做的"')
add_bullet('不是"我试试看"——是"我已经用了一个月"')
add_bullet('不是"还需要完善"——是"下一步会扩展到XX场景"')
add_text('路演这10分钟，你不是"刚接触AI的小白"，你是"已经用AI解决本部门问题的专家"。', bold=True)

add_text('⑥ 部门内/同主题的人互相衔接', bold=True, color=ACCENT_COLOR)
add_text('如果你们部门/同一主题有多人参与，提前沟通：')
add_bullet('谁主讲、谁补充')
add_bullet('数字不要打架')
add_bullet('案例不重复')
add_bullet('互相帮对方过一遍"开场句"和"请求"')

add_header_para('演练流程', level=3)
add_table(
    headers=['步骤', '做什么', '多久'],
    rows=[
        ['1. 一页纸框架', '把四步内容压到一页A4', '30分钟'],
        ['2. 大声说一遍', '对着镜子/手机录音说一遍', '15分钟'],
        ['3. 找搭档演练', '找一个人听你说，他填反馈卡', '30分钟'],
        ['4. 改一处再说', '根据反馈改一处，再练一遍', '15分钟'],
        ['5. 应急版', '准备一个2分钟极简版（防止被压缩时间）', '15分钟'],
    ],
    col_widths=[3.5, 9.0, 4.5]
)

add_header_para('应急版（2分钟极简版）', level=3)
add_text('如果到时候时间被压缩，保留这三件事：')

add_text('一个数字：', bold=True)
add_text('"我们在[某事]上，从[X]变成了[Y]。"')

add_text('一个案例细节（最有冲击力的一句）：', bold=True)
add_text('"举个例子：[最有画面感的10-15秒]"')

add_text('一个请求：', bold=True)
add_text('"我需要您做一件事：[具体]。"')

add_text('实例：', bold=True, color=ACCENT_COLOR)
add_text('"领导，我们的报告审核时间从每批次32人·时降到了8人·时。具体什么感受呢——上周那份根因写了4个字的报告，以前要花一个多小时整理，现在90秒有结果。我需要您下周让其他组来看15分钟，看他们是不是也有这个需求。"')
add_text('整个内容不到1分钟，但包含了数字、故事、请求。', color=TEXT_GRAY)

# ============== 第五部分：评审那些事儿 ==============
add_header_para('第五部分：评审那些事儿', level=1)

add_header_para('着装', level=3)
add_text('正装或商务休闲（不要运动装、不要拖鞋、不要太过随意的T恤）。', bold=True)
add_bullet('男士：衬衫/Polo衫 + 西裤')
add_bullet('女士：商务休闲或职业装')
add_bullet('颜色：深色系为主，干净整洁')
add_text('第一印象在7秒内形成。着装是最简单可控的第一印象。', color=TEXT_GRAY)

add_header_para('评审结构', level=3)
add_text('每人不超过15分钟，分两部分：', bold=True)
add_table(
    headers=['部分', '时长', '评价侧重'],
    rows=[
        ['路演', '10分钟', '表达清晰度、成果展示、问题思考深度'],
        ['Q&A', '3-5分钟', '应变能力、专业度、对业务的理解'],
    ],
    col_widths=[3.0, 3.0, 11.0]
)
add_text('评委和领导主要看：', bold=True)
add_bullet('你的痛点是不是真的（真实性）')
add_bullet('你的成果能不能推广（可复制性）')
add_bullet('你是不是用AI解决了本部门的实际问题（贴合度）')
add_bullet('你的下一步计划清不清楚（可期待性）')

add_header_para('放松心态', level=3)
add_text('几个事实帮你放松：', bold=True)
add_bullet('评委听了一整天，他比你更累')
add_bullet('你已经做完了最难的部分（做工具）')
add_bullet('路演不是考试，是分享')
add_bullet('领导问问题 = 他感兴趣，不是在刁难')

add_text('几个动作帮你放松：', bold=True)
add_bullet('进会议室前深呼吸3次')
add_bullet('第一句话前停顿1秒（不要急）')
add_bullet('说完请求后微笑，等领导反应')
add_bullet('如果忘词，看你的一页纸框架（不是演讲稿，是锚点）')

add_text('几个绝对不要：', bold=True, color=ACCENT_COLOR)
add_bullet('不要念稿（领导一眼能看出来）')
add_bullet('不要"我讲的不好请多多包涵"（这是在替自己的成果道歉）')
add_bullet('不要"我们还有很多不足"（不要主动打折扣）')
add_bullet('不要超时（10分钟到就停，请求还没说完就一句话结束）')

add_header_para('进会议室之前，最后一遍清单', level=3)
add_text('内容类：', bold=True)
add_text('□ 我知道我要说的四件事（痛/做/效/求）')
add_text('□ 我有一个具体的开场句（不以"我"开头）')
add_text('□ 我有一个60秒的具体故事（有时间、角色、行动、结果）')
add_text('□ 我有两个以上的对比数字（有"之前"有"之后"有"意味着"）')
add_text('□ 我有一个具体的请求（具体、小、有理由）')

add_text('节奏类：', bold=True)
add_text('□ 我演练过，知道自己大概需要多少分钟')
add_text('□ 如果被压缩到5分钟，我知道优先说哪三件事')
add_text('□ "做"的部分，我没有解释技术细节，只讲了案例')

add_text('心态类：', bold=True)
add_text('□ 我不是去"汇报工作"，是去帮领导解决一个他关心的问题')
add_text('□ 如果领导问了我没想到的问题，我知道怎么回应')

# ============== 写在最后 ==============
add_header_para('写在最后', level=1)

add_text('你做这份手册 / 这个工具 / 这个智能体，花了时间，动了脑子，解决了一个真实的问题。')
add_text('这是真的。', bold=True)

add_text('今天这场路演，不是在向领导交答卷，证明自己是个好学生。是在告诉他：')
add_text('你在他关注的方向上，做了一件有意义的事——而且你需要他帮你让这件事走得更远。', bold=True, color=ACCENT_COLOR)

add_text('带着这个心态走进那间会议室。')
add_text('你不是去被评审的，你是去分享一个成果的。')

add_text('> 你的工具已经够好了。今天做的这些，是帮它被看见——不是重新证明它值不值得。', bold=True, color=HEADER_COLOR)

# ============== 附录 ==============
add_header_para('附录：一页纸框架（建议打印带进会议室）', level=2)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('【我的一页纸呈现框架】')
r.font.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(HEADER_COLOR)
r.font.name = 'Microsoft YaHei'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_text('开场第一句话（不超过2句）：', bold=True)
add_text('■ 痛（一个问题，一个数字）：', bold=True)
add_text('    问题：')
add_text('    数字/影响：')
add_text('■ 做（一句话 + 一个故事关键词）：', bold=True)
add_text('    一句话：')
add_text('    案例关键词（提醒自己用）：')
add_text('■ 效（之前→之后→意味着什么）：', bold=True)
add_text('    之前：        之后：')
add_text('    意味着：')
add_text('■ 求（一件具体的事）：', bold=True)
add_text('预计总时长：        分钟')

add_text('【应急版 - 2分钟】', bold=True, color=ACCENT_COLOR)
add_text('    数字：')
add_text('    画面：')
add_text('    请求：')

# ============== 配套材料 ==============
add_header_para('本指引配套材料', level=3)
add_bullet('《10分钟逐字稿模板》（.md + .html）：5个场景的完整逐字稿')
add_bullet('《Q&A应答指南》（.md + .html）：6类问题的详细应对话术')
add_bullet('《学员交付物清单》（.xlsx）：4件必交 + 4类选交')

# 保存
out = r'D:\2026年课程\顺造科技\AI\评审\02-学员指南\学员路演指引.docx'
doc.save(out)
print(f'OK: {out}')

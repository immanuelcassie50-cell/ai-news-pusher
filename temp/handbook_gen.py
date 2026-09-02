# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"D:\新课开发\工作手册\AI时代决策工作手册\完整课程包-学员手册\AI时代决策工作手册-学员手册.docx"

def set_cell_shading(cell, color):
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_quote_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F5DC")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(139, 69, 19)
    doc.add_paragraph()

doc = Document()


doc.add_paragraph()

# 封面
for _ in range(8):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI时代决策工作手册")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("——从复盘到决策卡的系统方法论")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(80, 80, 80)

for _ in range(3):
    doc.add_paragraph()

audience = doc.add_paragraph()
audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = audience.add_run("学员对象：企业管理者、项目负责人、行动学习引导师")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

copyright_text = doc.add_paragraph()
copyright_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = copyright_text.add_run("© 版权课程内部资料 | 仅限学员使用")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# 目录
add_heading_with_style(doc, "目  录", level=1, color=RGBColor(0, 51, 102))

toc_items = [
    ("前言", "3"),
    ("学习目标", "4"),
    ("课前准备", "5"),
    ("PART 1 · 认知升级：复盘与决策卡的本质区分", "6"),
    ("  第一章  复盘写的是过去，决策卡写的是未来", "6"),
    ("  第二章  不是每个决策都值得做成一张卡", "7"),
    ("  第三章  复盘访谈问的不是你做了什么，是你当时不确定什么", "8"),
    ("  第四章  隐性判断是问出来的，不是想出来的", "9"),
    ("  第五章  决策卡不是流程图，是一份检查表加一个开关", "10"),
    ("  第六章  触发条件写不清楚，这张卡就是废纸", "11"),
    ("  第七章  场景映射矩阵：一张卡管不住所有场景", "12"),
    ("  第八章  失败案例不是用来吓人的，是用来对照的", "13"),
    ("  第九章  你写的决策卡，第一个读者应该是反对你的人", "14"),
    ("  第十章  训练活动的目的不是讲透道理，是让人在场景里犯一次错", "15"),
    ("PART 1 章节小结与练习", "16"),
    ("PART 2 · 组织落地：从个人经验到组织能力", "18"),
    ("  第一章  决策卡最大的敌人不是会不会用，是没人认领", "18"),
    ("  第二章  决策稽核不是查错，是替组织记住它学过什么", "19"),
    ("  第三章  好的决策卡会过时，这是它活着的证据", "20"),
    ("  第四章  复盘引导者最容易犯的错，是替决策者总结", "21"),
    ("  第五章  这份工作最终交付的不是文档，是组织的判断力", "22"),
    ("PART 2 章节小结与练习", "23"),
    ("工具模板", "25"),
    ("  附录一  决策卡标准模板", "25"),
    ("  附录二  稽核表模板", "26"),
    ("  附录三  访谈提纲模板", "27"),
    ("课后资源", "29"),
    ("附录  术语表", "30"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(item)
    run.font.size = Pt(11)
    if item.startswith("  "):
        run.font.color.rgb = RGBColor(80, 80, 80)
    else:
        run.bold = True

doc.add_page_break()
print("Cover and TOC done")

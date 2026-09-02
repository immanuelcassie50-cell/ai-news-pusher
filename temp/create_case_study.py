from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

output_path = "D:/新课开发/职业生涯和画布/破局・重启：用CEO思维重塑职业生涯/07-案例集/案例集.docx"

doc = Document()

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Helper functions
def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    if italic:
        for run in p.runs:
            run.italic = True
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_info_table(data):
    """data is list of (label, value) tuples"""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the label
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    return table

def add_data_table(headers, rows):
    """headers is list of column headers, rows is list of row data lists"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
        for paragraph in header_row.cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = cell_text

    return table

# ========== COVER ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("破局・重启：用CEO思维重塑职业生涯")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("案例集")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
version.add_run("版本：1.0")

copyright = doc.add_paragraph()
copyright.alignment = WD_ALIGN_PARAGRAPH.CENTER
copyright.add_run("版权所有：罗宏伟")

doc.add_page_break()

# ========== 案例集说明 ==========
add_heading("案例集说明", 1)

add_heading("本案例集简介", 2)
add_paragraph("本案例集是「破局・重启：用CEO思维重塑职业生涯」课程的核心教学材料之一。案例以虚构人物林晓的职业生涯经历为主线，贯穿课程四大模块，帮助学员在真实情境中理解并应用CEO思维框架。")

add_heading("案例结构", 2)
add_paragraph("每个案例均包含以下七个组成部分：")
add_bullet("案例基本信息：人物背景与场景设定")
add_bullet("案例背景：事件发展的完整脉络")
add_bullet("核心冲突：推动故事发展的关键矛盾")
add_bullet("决策两难：没有标准答案的战略抉择")
add_bullet("案例分析要点：讲师引导讨论的关键维度")
add_bullet("讨论问题：激发学员深度思考的开放性问题")
add_bullet("案例启示：从个案抽象出的普适性洞察")

add_heading("教学目标", 2)
add_bullet("模块一：帮助学员看清自己当前\"被经营\"的被动处境，意识到个人价值不等于岗位价值")
add_bullet("模块二：引导学员探索个人定位，找到差异化竞争优势")
add_bullet("模块三：教会学员用资产负债表工具全面盘点人生资产")
add_bullet("模块四：推动学员制定可落地的90天行动计划")

add_heading("使用建议", 2)
add_paragraph("建议授课时长：每个案例30-45分钟，其中案例阅读10分钟、引导讨论15-20分钟、总结提炼5-10分钟。")
add_paragraph("讨论环节设计：建议采用\"苏格拉底式提问\"，通过递进式问题引导学员自己得出结论，而非直接给出答案。")

# ========== 案例一 ==========
add_heading("模块一案例：看清局——林晓的\"被经营\"困局", 1)

add_heading("案例基本信息", 2)
add_info_table([
    ("人物", "林晓"),
    ("性别", "女"),
    ("年龄", "34岁"),
    ("职业", "某互联网公司产品运营经理"),
    ("工作年限", "8年"),
    ("教育背景", "二本毕业，校招进入现公司"),
    ("晋升轨迹", "从基层做起，经历过2次晋升，1次调岗"),
])

add_heading("案例背景", 2)
add_paragraph("林晓是一家互联网公司产品运营部的经理，在这家公司已经工作了整整8年。")
add_paragraph("当年她以应届生身份校招进入公司，从最基础的数据运营专员做起。凭着踏实肯干的性格和出色的业绩表现，她先后获得两次晋升，从专员一路做到经理。她所负责的用户运营工作也一直是公司的标杆项目，连续三年绩效考核获得A级评定。")
add_paragraph("然而，一切在三个月前发生了转折。")
add_paragraph("公司新任CEO上任后启动了战略调整，核心方向从\"用户增长\"转向\"成本控制\"。作为调整的一部分，林晓所在的产品运营部被裁撤——整个团队被拆分到两个不同的部门。林晓被安排到了一个边缘岗位：数据整理组，负责日常数据的汇总和报表制作，不再参与任何核心项目的策划和执行。")
add_paragraph("更让她感到不安的是，她发现过去那些对她\"重用\"的领导，在部门调整后都选择了沉默，没有人主动跟她沟通职业发展的问题。她感觉自己从一颗\"明星\"瞬间变成了一颗\"弃子\"。")

add_heading("核心冲突", 2)
add_paragraph("林晓一直信奉一条职场铁律：\"好好工作就会被认可\"。她八年如一日地坚守这条信念，勤奋、踏实、从不抱怨。然而这次组织调整让她第一次深刻地意识到：")
p = add_paragraph("一个人的价值，未必是由她的能力决定的，而是由她所在的\"位置\"决定的。")
p.runs[0].bold = True
add_paragraph("当她还在为绩效考核A级的荣誉沾沾自喜时，公司已经在战略层面做出了将她边缘化的决定。那些她引以为傲的\"能力\"和\"业绩\"，在公司层面的决策中，似乎并没有她想象中那么重要。")
add_paragraph("她陷入了深深的困惑：")
add_bullet("是自己的能力真的不行，还是从一开始她就误解了职场成功的法则？")
add_bullet("如果能力不是决定因素，那什么才是？")
add_bullet("她是应该继续等待——等待公司重新发现她的价值，还是应该主动出击——掌握自己职业发展的主动权？")

add_heading("决策两难", 2)
add_paragraph("选择留下")
add_paragraph("优势：")
add_bullet("8年的积累让她熟悉公司的一切，包括很多隐藏的资源和关系")
add_bullet("当前经济环境不确定，跳槽风险较高")
add_bullet("公司正在进行战略调整，可能孕育着新的机会")
add_paragraph("劣势：")
add_bullet("被边缘化的岗位很难做出成绩，容易陷入\"展示价值\"的困境")
add_bullet("长期来看，重复性工作会侵蚀她的核心竞争力")
add_bullet("其他部门的领导可能对她有\"被裁撤\"的印象")

add_paragraph("选择离开")
add_paragraph("优势：")
add_bullet("可以重新选择能够发挥自己能力的平台")
add_bullet("离开现有环境可能带来新的思路和机会")
add_bullet("主动跳槽通常能获得更好的薪资和职位")
add_paragraph("劣势：")
add_bullet("34岁的年龄在就业市场上面临一定歧视")
add_bullet("8年专注于一家公司，换赛道意味着重新开始")
add_bullet("林晓还有房贷压力，不允许太长的空档期")

add_heading("案例分析要点", 2)
add_bullet("职场价值的本质：讨论\"岗位价值\"与\"个人能力\"的区别。多数人把时间卖给了岗位，而不是在经营自己这个人。")
add_bullet("职场生存法则：除了做好本职工作，还需要思考如何让自己\"不可替代\"或\"难以替代\"。")
add_bullet("被动 vs 主动：林晓的困境在于她一直在\"被动等待\"，而不是\"主动经营\"。")
add_bullet("沉没成本陷阱：8年的投入是继续坚持的理由，还是应该果断止损的信号？")

add_heading("讨论问题", 2)
add_bullet("如果你是林晓，在得知部门被裁撤的消息后，你做的第一件事是什么？")
add_bullet("林晓的8年\"好员工\"经历，是她的资本还是负担？为什么？")
add_bullet("你觉得\"好好工作就会被认可\"这句话，在今天还适用吗？为什么？")
add_bullet("如果留下，林晓应该如何\"经营\"自己，增加重新被重用的可能性？")
add_bullet("如果离开，林晓可能面临哪些风险？她应该如何准备？")

add_heading("案例启示", 2)
add_paragraph("林晓的故事并非个例。在职业生涯的长河中，每个人都可能面临类似的\"被调整\"。真正决定一个人职业高度的，不是某一次选择的对错，而是他是否具备CEO思维——把职业生涯当作一家公司来经营。")
add_paragraph("当你把选择权交给别人，你就成为了别人棋盘上的一颗棋子。只有主动经营自己，才能掌握命运的主动权。")

# ========== 案例二 ==========
add_heading("模块二案例：立战略——林晓的\"个人定位\"探索", 1)

add_heading("案例基本信息", 2)
add_info_table([
    ("所处阶段", "职业转型探索期"),
    ("核心任务", "重新定义个人定位，找到差异化价值"),
    ("关键工具", "个人商业画布、个人定位三问"),
    ("时间背景", "部门调整后第2个月，开始反思职业方向"),
])

add_heading("案例背景", 2)
add_paragraph("在朋友的推荐下，林晓参加了一个职业发展规划工作坊。第一次接触\"CEO思维\"这个概念时，她感到既陌生又新奇。工作坊的老师说了一句让她印象深刻的话：")
p = add_paragraph("\"你的职业生涯就是一家公司，你是这家公司的唯一创始人兼CEO。你的能力是产品，你的时间是资本，你的每一次选择都是战略决策。\"")
p.runs[0].italic = True
add_paragraph("这句话像一道闪电，照亮了她过去8年从未思考过的盲区。")
add_paragraph("带着这个全新的视角，林晓开始用\"个人商业画布\"工具重新审视自己。在填写的过程中，她发现自己从来没有认真思考过三个问题：")
add_bullet("我的核心能力是什么？")
add_bullet("我的独特价值在哪里？")
add_bullet("谁是我的\"目标客户\"——谁愿意为我的能力买单？")

add_heading("核心冲突", 2)
add_paragraph("能力盘点后的迷茫")
add_paragraph("通过系统梳理，林晓列出了自己的：")
add_paragraph("核心能力：用户运营、活动策划、数据分析")
add_paragraph("独特优势：跨部门协调能力强、懂业务逻辑、执行力强")
add_paragraph("失败经验：曾尝试负责一个新项目但因资源不足而失败")
add_paragraph("然而，当她把这些能力写下来后，一个更深层的困惑浮现了：")
add_paragraph("她发现自己过去8年一直在做一个\"执行者\"——执行领导的决策、执行公司的战略、执行既定的KPI。她从来不是一个\"经营者\"——她不曾自己定战略、不曾自己找方向、不曾自己为结果负责。")
add_paragraph("这种\"执行者\"的定位，让她的能力看起来很全面，但缺乏一个足够尖锐的\"尖刀\"——那种让人一想到某个领域就会想到她的核心标签。")

add_paragraph("定位的三条路")
add_paragraph("在导师的引导下，林晓看到了三条可能的定位路径：")
add_paragraph("路径A——专家路线：继续深耕运营能力，成为某个细分领域的专家，比如\"用户增长专家\"或\"活动运营专家\"。这条路最稳妥，但可能需要较长时间积累。")
add_paragraph("路径B——管理路线：转型做管理者，带领团队。从管自己到管别人，从做执行到做决策。这条路能快速提升影响力，但管理能力是她的短板。")
add_paragraph("路径C——创造路线：不要在现有赛道上竞争，而是创造一个全新的定位。比如\"互联网运营+职业规划\"的跨界定位，或者\"运营方法论布道者\"的内容创业方向。这条路风险最大，但潜在回报也最高。")

add_heading("决策两难", 2)
add_paragraph("三条路摆在她面前，每一条都有道理，每一条也都有风险：")
add_bullet("选择A，可能在一家新公司继续做执行者，但有机会成为某个领域的专家")
add_bullet("选择B，管理岗位可能向她招手，但管理经验的缺失让她心里没底")
add_bullet("选择C，看起来很诱人，但创业的不确定性和现有的经济压力形成了一对矛盾")
add_paragraph("更让她纠结的是，她不知道哪个选择能让她\"不可替代\"，哪个选择可能在若干年后再次让她陷入同样的困境。")

add_heading("案例分析要点", 2)
add_bullet("执行者 vs 经营者：大多数人的职业困境，根源在于一直是\"执行者思维\"而非\"经营者思维\"。")
add_bullet("差异化定位：在这个高度竞争的时代，\"全面\"可能意味着\"没有特色\"。找到足够尖锐的标签，比罗列所有能力更重要。")
add_bullet("定位三问：我的目标客户是谁？我提供什么价值？为什么客户选择我而不是别人？")
add_bullet("定位的动态性：定位不是一成不变的，而是随着能力和市场的变化不断迭代的。")

add_heading("讨论问题", 2)
add_bullet("如果你是林晓，你会选择哪条路径？为什么？")
add_bullet("你觉得\"执行者\"和\"经营者\"最大的区别是什么？")
add_bullet("在当前的市场环境下，你认为哪种定位路线最具长期价值？")
add_bullet("有人说\"选择比努力更重要\"，你认同吗？为什么？")
add_bullet("如果林晓选择跨界创造新定位，她可能面临哪些挑战？")

add_heading("案例启示", 2)
add_paragraph("个人定位不是\"你想成为谁\"，而是\"市场需要什么\"和\"你能提供什么\"的交集。一个好的定位，需要同时满足三个条件：")
add_bullet("你擅长做这件事（能力）")
add_bullet("你热爱做这件事（动力）")
add_bullet("有人愿意为这件事付费（市场）")
add_paragraph("只有三者重叠的区域，才是真正的\"甜蜜点\"。")

# ========== 案例三 ==========
add_heading("模块三案例：调资产——林晓的\"资产负债表\"盘点", 1)

add_heading("案例基本信息", 2)
add_info_table([
    ("所处阶段", "职业重启准备期"),
    ("核心任务", "全面盘点个人资产与负债"),
    ("关键工具", "个人资产负债表"),
    ("时间背景", "个人定位确定后，开始系统梳理资源"),
])

add_heading("案例背景", 2)
add_paragraph("在确定了\"运营专家\"的定位方向后，导师建议林晓用\"个人资产负债表\"工具对自己进行一次全面的盘点。")
add_paragraph("这个工具的核心理念是：把你自己当作一家公司，你的职业生涯就是这家公司的经营。用资产负债表的逻辑来审视你的人生资源——哪些是资产，哪些是负债。")
add_paragraph("林晓花了一周时间，认真填写了这张表。当她把结果写出来的时候，她被自己吓了一跳。")

add_heading("核心冲突", 2)
add_paragraph("资产盘点：发现一直在\"消耗\"而非\"增值\"")

add_data_table(
    ["资产类别", "具体内容", "当前价值", "增值潜力"],
    [
        ["时间资产", "每天有效工作时间约4小时", "低", "中等"],
        ["人脉资产", "业内人脉较少，主要集中在本公司", "低", "高"],
        ["能力资产", "运营能力、数据分析能力、项目管理能力", "中等", "高"],
        ["信任资产", "公司内部口碑较好，但外界认知度低", "中等", "高"],
        ["财务资产", "有房贷，储蓄有限", "低", "中等"],
        ["健康资产", "身体状况良好，精力充沛", "高", "高"],
    ]
)

add_paragraph("林晓惊讶地发现，她的时间资产远没有她想象的那么充裕。过去她一直以为自己很忙，但仔细追踪后发现：每天真正产生价值的工作时间只有大约4小时。其余时间被低效会议、重复性工作和无意义的\"等待\"消耗掉了。")
add_paragraph("更让她意外的是人脉资产。工作了8年，她竟然没有积累多少真正有价值的业内人脉。这让她在求职时会非常被动。")

add_paragraph("负债盘点：看清那些消耗你的东西")

add_data_table(
    ["负债类型", "具体表现", "严重程度", "改善难度"],
    [
        ["沉没成本", "8年投入在现公司，换赛道意味着重新开始", "高", "心理层面"],
        ["能力短板", "缺乏战略规划能力、商业思维弱", "中", "可学习"],
        ["负面情绪", "最近的低落情绪影响了工作状态", "中", "可调节"],
        ["路径依赖", "习惯了执行者思维，难以快速转变", "中", "需要刻意练习"],
        ["房贷压力", "每月固定支出限制了我的选择空间", "中", "短期难以改变"],
    ]
)

add_paragraph("通过盘点，林晓意识到她背负着不少\"隐性负债\"。最大的负债是\"沉没成本\"——那8年的投入像一根无形的绳索，让她很难下决心离开。但同时她也意识到，沉没成本其实是\"伪负债\"——它只是心理上的障碍，而不是真正的限制。")

add_paragraph("核心洞察")
add_paragraph("经过这番盘点，林晓得出了一个让她心惊的结论：")
p = add_paragraph("她过去8年一直在\"消耗资产\"而不是\"增值资产\"。")
p.runs[0].bold = True
add_paragraph("她把最宝贵的青春年华花在了重复性的工作上，没有刻意投资自己的能力资产和人脉资产。她像一家只顾着日常运营、却从不考虑长期发展的公司——账面数字看似稳定，实则在悄悄贬值。")

add_heading("决策两难", 2)
add_paragraph("盘点结果让林晓面临一个根本性的问题：")
add_bullet("是继续维持现状，慢慢修复资产负债表——这条路更安全，但可能需要3-5年才能看到明显改变")
add_bullet("还是现在就采取激进行动——主动切换到高增值赛道，用行动倒逼成长——这条路风险更大，但可能缩短转型周期")
add_paragraph("无论选择哪条路，她都必须开始真正地\"经营\"自己，而不是继续\"消耗\"自己。")

add_heading("案例分析要点", 2)
add_bullet("隐性资产：很多人低估了自己的人脉资产和潜力资产。一周的时间投入，可能会打开完全不同的机会之门。")
add_bullet("伪负债：沉没成本是最典型的\"伪负债\"。它只在心理上存在，并没有真正消耗你的未来。")
add_bullet("增值 vs 消耗：每一年、每个月、每一天，你是在增值还是在消耗？这是CEO每天都要问自己的问题。")
add_bullet("资产配置：把你的时间、精力、金钱投向哪里，决定了你的资产负债表是改善还是恶化。")

add_heading("讨论问题", 2)
add_bullet("你觉得林晓最大的\"资产\"是什么？最大的\"负债\"是什么？")
add_bullet("你有多久没有\"投资\"自己的能力资产了？最近一次学习新技能是什么时候？")
add_bullet("沉没成本是如何影响你当前的选择的？你打算如何处理它？")
add_bullet("如果你来做个人资产负债表，你会发现什么令你惊讶的事实？")
add_bullet("你现在的每一天，是在\"增值\"还是在\"消耗\"？")

add_heading("案例启示", 2)
add_paragraph("个人资产负债表是一面镜子，照出你真实的财务状况。更重要的是，它是一张导航图，告诉你下一步应该把资源投向哪里。")
add_paragraph("真正的聪明人，不是想办法节省消耗，而是想办法加速增值。当你的资产增长速度快于消耗速度时，你就开始走向富足。反之，你就是在悄然走向破产。")

# ========== 案例四 ==========
add_heading("模块四案例：起新局——林晓的\"90天行动计划\"", 1)

add_heading("案例基本信息", 2)
add_info_table([
    ("所处阶段", "职业转型执行期"),
    ("核心任务", "制定并执行90天转型行动计划"),
    ("关键工具", "90天行动计划表、每周复盘机制"),
    ("时间背景", "完成前三个模块学习后，正式启动转型"),
])

add_heading("案例背景", 2)
add_paragraph("完成前三个模块的学习后，林晓感觉自己像被升级了一台新的操作系统。她第一次用CEO的视角来看待自己的职业生涯，第一次清晰地看到了自己的资产与负债，第一次认真地思考了自己的定位。")
add_paragraph("但她知道，认知升级只是第一步，真正的考验是执行。")
add_paragraph("在导师的指导下，林晓制定了一份\"90天转型行动计划\"。这份计划的核心是一个决策：")
p = add_paragraph("止损决策：不再等待内部调岗机会，立即开始外部探索。")
p.runs[0].bold = True
add_paragraph("她的理由是：等待是最贵的成本。与其把命运交给别人，不如主动出击。哪怕失败了，至少知道自己败在哪里。")

add_heading("核心冲突", 2)
add_paragraph("第一个可执行决策")
add_paragraph("在制定计划的过程中，林晓遇到了一个经典的两难：")
add_paragraph("她可以选择\"骑驴找马\"——在职找工作，边工作边寻觅新机会。这样更安全，但时间精力分散，可能拉长转型周期。")
add_paragraph("她也可以选择\"破釜沉舟\"——辞职后全力找工作。这样更专注，但经济压力和心理压力会成倍增加。")
add_paragraph("最终，林晓选择了中间路线：")
add_bullet("第一阶段（1-30天）：在职探索，更新简历，投递目标岗位")
add_bullet("第二阶段（31-60天）：如果有了面试机会，再考虑是否辞职")
add_bullet("第三阶段（61-90天）：集中精力拿到offer，做出最终选择")

add_paragraph("90天行动计划")

add_data_table(
    ["阶段", "时间", "核心任务", "具体行动", "预期成果"],
    [
        ["第一阶段", "第1-30天", "个人品牌重塑", "更新简历，突出可量化成果；重新包装LinkedIn和职业社交形象；开始撰写行业洞察文章", "简历通过率提升，获得首批面试机会"],
        ["第二阶段", "第31-60天", "深度networking", "每周至少认识2个新朋友；参加3场行业线下活动；主动联系目标公司内部人员", "建立有效人脉圈，获取内部推荐"],
        ["第三阶段", "第61-90天", "offer冲刺", "拿到至少2个offer；全面评估每个offer的优劣势；做出最终选择", "成功转型，薪资提升30%+"],
    ]
)

add_heading("成果", 2)
add_paragraph("林晓严格执行了这份计划。")
add_paragraph("在第30天，她收到了第一个面试邀请。")
add_paragraph("在第45天，她同时有3个面试在进行中。")
add_paragraph("在第60天，她拿到了第一个offer——某成长型公司的运营总监岗位，薪资提升25%。")
add_paragraph("在第75天，她又收到了第二个offer——一家知名互联网公司的高级运营经理岗位，薪资提升30%。")
add_paragraph("经过慎重比较，林晓选择了第二个offer。她说理由有三：")
add_bullet("公司品牌更有利于未来发展")
add_bullet("岗位更符合\"运营专家\"的定位")
add_bullet("直接汇报给业务负责人，成长空间更大")

add_heading("案例分析要点", 2)
add_bullet("止损思维：等待是最贵的成本。当机会成本大于行动成本时，就应该果断行动。")
add_bullet("分阶段决策：不是所有决定都要一次性做完。有些决定可以\"等等看\"，但要设定明确的\"等等看\"的期限。")
add_bullet("可量化的目标：\"拿到2个offer\"比\"找到好工作\"更清晰、更可衡量、也更容易执行。")
add_bullet("执行闭环：计划-执行-复盘-迭代。没有复盘的计划很容易走偏，没有迭代的执行很难持续改进。")

add_heading("讨论问题", 2)
add_bullet("林晓的90天计划有哪些值得借鉴的地方？有哪些可以改进的地方？")
add_bullet("你觉得\"在职找工作\"和\"辞职找工作\"各有什么利弊？你会选择哪种方式？")
add_bullet("为什么林晓最终选择了第二个offer而不是第一个？如果是你，你会怎么选？")
add_bullet("90天计划执行过程中，林晓可能会遇到哪些阻力？她应该如何应对？")
add_bullet("你是否有类似的\"90天计划\"？如果没有，你现在愿意开始制定一个吗？")

add_heading("案例启示", 2)
add_paragraph("任何伟大的计划，都必须落实到每一天的具体行动中。")
add_paragraph("90天看起来不长，但足够改变很多事。关键不是你有多少时间，而是你把时间用在哪里。")
add_paragraph("当你用CEO思维来经营自己时，你会发现：机会不是等来的，是创造出来的。命运不是被动接受的，是主动争取的。")

# ========== 使用说明 ==========
add_heading("案例使用说明", 1)

add_heading("各模块案例如何配合教学", 2)

add_paragraph("建议教学顺序")
add_paragraph("建议按照案例集的自然顺序进行教学：模块一 → 模块二 → 模块三 → 模块四。每个案例都建立在前一个案例的认知基础上，顺序颠倒可能会影响教学效果。")

add_paragraph("模块间的逻辑关系")
add_bullet("模块一（看清局）：帮助学员建立危机意识，意识到\"被动等待\"的危险性")
add_bullet("模块二（立战略）：在危机意识基础上，引导学员思考\"我要往哪里走\"")
add_bullet("模块三（调资产）：在明确方向后，帮助学员看清\"我现在有什么\"")
add_bullet("模块四（起新局）：在前三个模块基础上，推动学员制定\"我现在要做什么\"")

add_paragraph("时间分配建议")
add_paragraph("全套案例教学建议总时长：3-4小时")
add_bullet("模块一：45-60分钟（案例阅读10分钟，讨论30分钟，总结10分钟）")
add_bullet("模块二：45-60分钟")
add_bullet("模块三：50-70分钟（因为包含表格工具讲解）")
add_bullet("模块四：45-60分钟")

add_heading("讨论环节设计建议", 2)

add_paragraph("苏格拉底式提问法")
add_paragraph("建议采用\"苏格拉底式提问\"进行案例讨论，而非直接给出答案。核心原则是：通过递进式问题，引导学员自己得出结论。")
add_paragraph("提问层次：")
add_bullet("第一层（事实层）：案例中发生了什么？林晓面临的情况是什么？")
add_bullet("第二层（分析层）：为什么会发生这种情况？林晓的问题出在哪里？")
add_bullet("第三层（决策层）：如果是你，你会怎么选择？你会问自己什么问题？")
add_bullet("第四层（行动层）：从现在开始，你要做什么具体的行动？")

add_paragraph("小组讨论设计")
add_paragraph("建议将学员分成3-4人小组，每个小组讨论一个特定问题，5-8分钟后各组分享结论。讲师进行总结和点评。")

add_paragraph("角色扮演设计")
add_paragraph("针对模块四的90天行动计划，可以设计角色扮演环节：让学员扮演林晓，向\"面试官\"（由其他学员扮演）推销自己。")

add_heading("后续跟进建议", 2)
add_bullet("课后作业：让学员为自己制定一份类似的\"个人资产负债表\"或\"90天行动计划\"")
add_bullet("一个月后跟进：让学员分享自己在执行计划过程中的进展和困难")
add_bullet("三个月后跟进：让学员复盘自己的转型经历，无论成功还是失败都是宝贵的学习素材")

# Save
doc.save(output_path)
print("Document created successfully!")

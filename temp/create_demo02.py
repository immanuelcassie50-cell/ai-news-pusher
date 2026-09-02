# -*- coding: utf-8 -*-
"""
生成 demo02-描述性定义练习示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_demo02():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("描述性定义练习示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 描述性定义帮你精确描述问题现状，而非模糊感受")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    # 学员信息
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：林晓风", "团队：营销创新组", "日期：2026-07-18", "教练：王行动"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # 练习说明
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、练习任务")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    task_para = doc.add_paragraph()
    run = task_para.add_run("请用描述性定义的方式，精确定义以下问题现象：")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 左侧问题，右侧定义对照
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、练习内容")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 问题1
    q1 = doc.add_paragraph()
    run = q1.add_run("【问题A】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run = q1.add_run('  "客户对我们不满意"')
    run.font.size = Pt(11)

    # 表格：模糊描述 vs 描述性定义
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'

    # 表头
    table1.rows[0].cells[0].text = "模糊描述（听到的）"
    table1.rows[0].cells[1].text = "描述性定义（精准描述）"
    for cell in table1.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml import OxmlElement
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    table1.rows[1].cells[0].text = "客户对我们不满意"
    table1.rows[1].cells[1].text = "在过去6个月里，我们收到了23起客户投诉，其中：\n- 12起因交付延期（平均延迟5.3天）\n- 7起因产品质量问题（退货率4.2%）\n- 4起因服务响应慢（平均响应时间超过48小时）\n客户好评率从92%下降至86%，净推荐值（NPS）从45降至38"
    for cell in table1.rows[1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    # 教练批注
    coach_comment = doc.add_paragraph()
    run = coach_comment.add_run("[教练批注：描述性定义的关键是'可衡量、可追溯'。当你说'客户不满意'，这是一个感受；当你说'投诉23起、NPS 38'，这是一个事实。你的解决方案将完全不同。]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    doc.add_paragraph()

    # 问题2
    q2 = doc.add_paragraph()
    run = q2.add_run("【问题B】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run = q2.add_run('  "团队士气低落"')
    run.font.size = Pt(11)

    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'

    table2.rows[0].cells[0].text = "模糊描述（听到的）"
    table2.rows[0].cells[1].text = "描述性定义（精准描述）"
    for cell in table2.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    table2.rows[1].cells[0].text = "团队士气低落"
    table2.rows[1].cells[1].text = '本季度团队有以下表现：\n- 出勤率：92%，较上季度下降5%\n- 主动加班率：8%，较上季度下降12%\n- 合理化建议提交：3条（去年同期15条）\n- 员工满意度调研：68分（低于基准线75分）\n- 主动离职意向：32%表示"有想法"'
    for cell in table2.rows[1].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()

    coach_comment2 = doc.add_paragraph()
    run = coach_comment2.add_run("[教练批注：从模糊的'士气低落'到5个可量化的指标，这才构成一个完整的问题诊断。每个数字背后都有原因，这才是真正的起点。]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    doc.add_paragraph()

    # 练习要点总结
    summary_title = doc.add_paragraph()
    run = summary_title.add_run("三、描述性定义三要素")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    elements = [
        ("1. 具体数据", "用数字说话，而非形容词", "不用'很高'，用'85%'；不用'很久'，用'平均7.2天'"),
        ("2. 时间边界", "明确时间范围", "不是'一直以来'，而是'过去6个月'或'本季度'"),
        ("3. 维度拆分", "从多角度描述", "将'客户不满意'拆解为质量、服务、交付等维度")
    ]

    elem_table = doc.add_table(rows=4, cols=3)
    elem_table.style = 'Table Grid'
    elem_table.rows[0].cells[0].text = "要素"
    elem_table.rows[0].cells[1].text = "要求"
    elem_table.rows[0].cells[2].text = "错误示例 vs 正确示例"
    for cell in elem_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, (elem, req, example) in enumerate(elements):
        elem_table.rows[i+1].cells[0].text = elem
        elem_table.rows[i+1].cells[1].text = req
        elem_table.rows[i+1].cells[2].text = example
        for j, cell in enumerate(elem_table.rows[i+1].cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if j == 0:
                        run.font.bold = True

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo02-描述性定义练习示例.docx"
    doc.save(output_path)
    print(f"demo02已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    create_demo02()

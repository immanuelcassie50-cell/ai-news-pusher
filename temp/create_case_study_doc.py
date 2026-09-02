# -*- coding: utf-8 -*-
"""生成案例集 docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False, color=None):
    """设置run的字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, font_size=18, color=(43, 45, 66)):
    """添加标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=True, color=color)
    return p

def add_subheading(doc, text, font_size=14, color=(42, 45, 66)):
    """添加副标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size, bold=True, color=color)
    return p

def add_paragraph(doc, text, font_size=11, indent=False, space_after=8):
    """添加正文段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, font_size=font_size)
    return p

def add_card(doc, title, content_items, accent_color=(239, 35, 60)):
    """添加卡片式内容块"""
    # 卡片背景用段落边框效果实现
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(title)
    set_run_font(run, font_size=12, bold=True, color=accent_color)

    for item in content_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(4)
        p_item.paragraph_format.left_indent = Pt(20)
        run = p_item.add_run(item)
        set_run_font(run, font_size=10)
    return p

def create_case_study_doc(output_path):
    """创建案例集文档"""
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ========== 封面 ==========
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(40)
    p_title.paragraph_format.space_after = Pt(20)
    run = p_title.add_run("手册进化 案例集")
    set_run_font(run, font_size=32, bold=True, color=(43, 45, 66))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(30)
    run = p_sub.add_run('从"阅读版"到"执行手册" 教学案例')
    set_run_font(run, font_size=16, color=(141, 153, 174))

    p_desc = doc.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_desc.paragraph_format.space_after = Pt(60)
    run = p_desc.add_run('本案例集包含：化工反应釜操作案例 + 客户服务投诉处理案例\n展示六大模块的转化方法与对照效果')
    set_run_font(run, font_size=11, color=(100, 100, 100))

    # ========== 目录 ==========
    add_heading(doc, "案例目录", font_size=18)

    toc_items = [
        ('案例一', '化工反应釜升温异常处置', '标准案例，展示完整转化流程'),
        ('案例二', '客户服务投诉处理', '拓展案例，验证跨行业适用性'),
        ('转化对照表', '阅读版 vs 执行版', '直观展示转化前后的差异')
    ]

    for num, title, desc in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(num + '：' + title)
        set_run_font(run, font_size=12, bold=True, color=(42, 45, 66))
        run = p.add_run(' — ' + desc)
        set_run_font(run, font_size=11)

    doc.add_page_break()

    # =============================================
    # 案例一：化工反应釜升温异常处置
    # =============================================
    add_heading(doc, "案例一：化工反应釜升温异常处置", font_size=20, color=(239, 35, 60))

    add_paragraph(doc, '本案例展示一份完整的执行手册转化过程。原始素材来自化工企业实操手册，经系统化转化后，形成可直接照做的执行版本。', space_after=12)

    # 场景定位
    add_subheading(doc, "【场景定位】阅读版 vs 执行版")

    add_card(doc, '阅读版（原始素材）', [
        '操作人员应该严格按照操作规程进行日常巡检，',
        '发现异常情况及时上报，确保生产安全。',
        '巡检内容包括设备运行状态、参数指标、安全设施等。'
    ])

    add_card(doc, '执行版（转化后）', [
        '触发条件：到达反应釜区域准备巡检',
        '动作1：确认巡检时间（在交接班后30分钟内）',
        '动作2：检查控制室DCS画面，确认塔顶温度≤85℃',
        '动作3：确认回流比控制在1.2-1.5范围内',
        '动作4：现场确认各阀门处于自动状态',
        '动作5：填写巡检记录表（电子签名）',
        '异常判断：如参数异常立即进入【异常处置】流程'
    ], accent_color=(42, 163, 143))

    # 标准动作
    add_subheading(doc, "【标准动作】判断标准卡")

    add_card(doc, '判断标准卡示例：塔顶温升异常', [
        '条件信号1：塔顶温度 > 85℃',
        '条件信号2：回流比 < 1.2',
        '同时满足 → 判断结论：冷凝效果下降，需立即处置',
        '对应动作：',
        '  ① 降低原料进料量30%',
        '  ② 开大循环冷却水流量',
        '  ③ 观察10分钟，如温度持续上升立即停料',
        '关键变量：温度阈值（85℃）、时间窗口（10分钟）',
        '可忽略因素：环境温度变化（不影响判断）'
    ], accent_color=(42, 163, 143))

    # 分级处置
    add_subheading(doc, "【分级处置】异常三级分类")

    levels = [
        ('轻微', '参数轻微偏移，自行调整可恢复', '5分钟内处置', '降低进料量+加大冷却'),
        ('中度', '参数明显异常，需技术支持', '立即联系班长', '远程指导+现场配合'),
        ('严重', '参数失控，有安全风险', '立即停机上报', '启动应急流程+疏散人员')
    ]

    for level, desc, timing, action in levels:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(10)
        run = p.add_run('[' + level + '] ' + desc)
        set_run_font(run, font_size=11, bold=True)
        run = p.add_run(' | 时限：' + timing)
        set_run_font(run, font_size=10, color=(100, 100, 100))
        run = p.add_run(' | 处置：' + action)
        set_run_font(run, font_size=10)

    doc.add_page_break()

    # =============================================
    # 案例二：客户服务投诉处理
    # =============================================
    add_heading(doc, "案例二：客户服务投诉处理", font_size=20, color=(239, 35, 60))

    add_paragraph(doc, '本案例展示跨行业适用性——将化工案例的转化框架应用于客户服务场景。', space_after=12)

    # 场景定位
    add_subheading(doc, "【场景定位】客户投诉处理场景")

    add_card(doc, '阅读版（原始素材）', [
        '客服人员应认真处理客户投诉，',
        '对于客户反映的问题及时响应，',
        '不断提升客户满意度。'
    ])

    add_card(doc, '执行版（转化后）', [
        '触发条件：收到客户投诉（电话/在线/邮件）',
        '动作1：确认投诉类型（产品质量/服务态度/配送问题/其他）',
        '动作2：在CRM系统录入投诉工单（30分钟内）',
        '动作3：根据投诉类型判断处理权限',
        '动作4：联系相关责任人，约定回复时间',
        '动作5：按【分级处置表】执行对应处置',
        '动作6：24小时内首次回复客户，48小时内给出解决方案'
    ], accent_color=(42, 163, 143))

    # 判断标准卡
    add_subheading(doc, "【判断标准卡】投诉分级")

    add_card(doc, '判断标准卡示例：紧急投诉识别', [
        '条件信号1：客户明确表示将升级至媒体/监管机构',
        '条件信号2：涉及产品质量安全事故',
        '条件信号3：客户为重点大客户（年度消费≥10万）',
        '满足任一条件 → 判断结论：紧急投诉，启动快速响应',
        '对应动作：',
        '  ① 立即上报客服总监',
        '  ② 2小时内给出初步回应',
        '  ③ 24小时内提供完整解决方案',
        '关键变量：客户等级、工单来源、投诉历史',
        '可忽略因素：客户语气激动程度（不影响判断）'
    ], accent_color=(42, 163, 143))

    # 分级处置
    add_subheading(doc, "【分级处置】投诉处理三级分类")

    levels2 = [
        ('普通', '常规投诉，处理流程清晰', '48小时', '按标准流程处理+回复'),
        ('重要', '涉及多个部门，需协调', '24小时', '成立专项小组+升级处理'),
        ('紧急', '有升级风险，需快速响应', '4小时', '高管介入+快速处置+客户安抚')
    ]

    for level, desc, timing, action in levels2:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(10)
        run = p.add_run('[' + level + '] ' + desc)
        set_run_font(run, font_size=11, bold=True)
        run = p.add_run(' | 时限：' + timing)
        set_run_font(run, font_size=10, color=(100, 100, 100))
        run = p.add_run(' | 处置：' + action)
        set_run_font(run, font_size=10)

    doc.add_page_break()

    # =============================================
    # 转化对照表
    # =============================================
    add_heading(doc, "转化对照表：阅读版 vs 执行版", font_size=18)

    add_paragraph(doc, '以下表格直观展示阅读版与执行版的核心差异：', space_after=12)

    # 创建对照表
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    # 表头
    headers = ['对比维度', '阅读版手册', '执行版手册']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        # 设置单元格背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2b2d42')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # 对照内容
    comparisons = [
        ('内容组织', '按信息完整度组织（背景→流程→注意事项）', '按决策点组织（条件→动作→判断→处置）'),
        ('颗粒度', '模糊表述（"注意安全"、"必要时"）', '可量化条件（温度>85℃、回流比<1.2）'),
        ('结构', '段落式叙述，逻辑隐藏在文字中', '卡片式结构，逻辑显性化'),
        ('断点', '一带而过（"视情况处理"）', '明确触发条件和决策路径'),
        ('异常处理', '笼统描述（"异常情况报IT"）', '分级分类（轻微→远程协助，严重→现场处理）'),
        ('学习效果', '读得懂，但不知道怎么判断和执行', '知道什么时候该做什么判断，怎么处置')
    ]

    for row_idx, (dim, before, after) in enumerate(comparisons, start=1):
        table.rows[row_idx].cells[0].text = dim
        table.rows[row_idx].cells[1].text = before
        table.rows[row_idx].cells[2].text = after
        set_run_font(table.rows[row_idx].cells[0].paragraphs[0].runs[0], font_size=10, bold=True)
        set_run_font(table.rows[row_idx].cells[1].paragraphs[0].runs[0], font_size=10)
        set_run_font(table.rows[row_idx].cells[2].paragraphs[0].runs[0], font_size=10)

    # ========== 结语 ==========
    add_heading(doc, "案例使用说明", font_size=18)

    usage = [
        '本案例集用于课程教学，每个案例配套【转化前素材】+【转化后成品】+【对照分析】三部分',
        '建议教学顺序：先展示化工案例的完整转化过程，再让学员用客户服务案例进行练习',
        '案例中的数值和参数为教学虚构，请勿直接套用实际生产场景',
        '学员练习时，可使用自己手头的真实手册素材进行对照练习'
    ]

    for item in usage:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('• ' + item)
        set_run_font(run, font_size=11)

    # 保存文档
    doc.save(output_path)
    print(f'案例集已保存至: {output_path}')

if __name__ == '__main__':
    output_path = 'D:/新课开发/经验萃取/阅读手册转执行手册/完整课程包/007-案例集/007-案例集.docx'
    create_case_study_doc(output_path)
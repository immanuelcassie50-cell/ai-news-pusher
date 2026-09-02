# -*- coding: utf-8 -*-
"""
创建企业大学品牌资产模块1教学文档
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 输出路径
output_dir = r"D:\新课开发\企业大学\对外\4.企业大学品牌资产设计：把理想愿景转化为对外可辨识的品牌\教学文档"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "01-模块1_品牌本质.docx")

doc = Document()

# ==================== 页面设置 ====================
section = doc.sections[0]
section.page_width = Inches(11.69)  # A4横向
section.page_height = Inches(8.27)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ==================== 辅助函数 ====================
def set_run_font(run, font_name="微软雅黑", size=12, bold=False, color=None):
    """设置run的字体和大小"""
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    """添加标题"""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        run.font.name = "微软雅黑"
        run._r.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    return para

def add_body_para(doc, text, first_line_indent=True):
    """添加正文段落"""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(8)
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # 两个字符
    run = para.add_run(text)
    set_run_font(run)
    return para

def add_bullet_para(doc, text, level=0):
    """添加要点"""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    set_run_font(run)
    return para

def add_table_row(table, cells_data, bold_first=False):
    """添加表格行"""
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = data
        if bold_first and i == 0:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
    return row

# ==================== 元信息头 ====================
# 标题
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(12)
title_para.paragraph_format.space_after = Pt(6)
run = title_para.add_run("模块一：品牌本质——企业大学为什么需要品牌资产")
set_run_font(run, size=22, bold=True)

# 版本信息表格
meta_table = doc.add_table(rows=4, cols=2)
meta_table.style = 'Table Grid'
meta_data = [
    ("版本号", "V1.0"),
    ("保密级别", "内部使用"),
    ("模块时长", "90分钟"),
    ("最后更新", "2026年8月")
]
for i, (label, value) in enumerate(meta_data):
    meta_table.cell(i, 0).text = label
    meta_table.cell(i, 1).text = value
    # 加粗标签
    for para in meta_table.cell(i, 0).paragraphs:
        for run in para.runs:
            run.font.bold = True
    for para in meta_table.cell(i, 1).paragraphs:
        for run in para.runs:
            set_run_font(run)

doc.add_paragraph()  # 空行

# ==================== 使用说明 ====================
add_heading(doc, "使用说明", level=1)

instructions = [
    ("填答说明", "学员在引导下填写个人反思"),
    ("目的", "帮助学员理解品牌对企业大学的战略价值"),
    ("时长建议", "90分钟")
]
for title, desc in instructions:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run1 = para.add_run(f"• {title}：")
    set_run_font(run1, bold=True)
    run2 = para.add_run(desc)
    set_run_font(run2)

doc.add_paragraph()

# ==================== 内容导航 ====================
add_heading(doc, "内容导航", level=1)

add_heading(doc, "学习目标", level=2)
objectives = [
    "理解品牌的本质定义与核心内涵",
    "认识企业大学品牌的战略价值",
    "掌握品牌资产的构成要素",
    "分析企业大学品牌建设的现状与挑战"
]
for i, obj in enumerate(objectives, 1):
    para = doc.add_paragraph()
    run1 = para.add_run(f"{i}. ")
    set_run_font(run1, bold=True)
    run2 = para.add_run(obj)
    set_run_font(run2)
    para.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

add_heading(doc, "模块概览", level=2)
overview_table = doc.add_table(rows=5, cols=3)
overview_table.style = 'Table Grid'
overview_data = [
    ("知识点", "核心概念", "时长"),
    ("1.1", "品牌的本质定义", "20分钟"),
    ("1.2", "企业大学品牌的战略价值", "25分钟"),
    ("1.3", "品牌资产的构成要素", "25分钟"),
    ("1.4", "企业大学品牌建设现状分析", "20分钟")
]
for i, row_data in enumerate(overview_data):
    for j, cell_data in enumerate(row_data):
        overview_table.cell(i, j).text = cell_data
        for para in overview_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

# ==================== 知识点模块 ====================

# --- 知识点1.1 ---
add_heading(doc, "知识点1.1：品牌的本质定义", level=1)

core_concept_para = doc.add_paragraph()
run1 = core_concept_para.add_run("核心概念：")
set_run_font(run1, bold=True, color=RGBColor(0, 51, 102))
run2 = core_concept_para.add_run("品牌是一种承诺，是利益相关者对企业价值主张的认知总和")
set_run_font(run2, color=RGBColor(0, 51, 102))
core_concept_para.paragraph_format.space_after = Pt(12)

add_heading(doc, "一、品牌的经典定义", level=2)

definitions = [
    ("美国营销学会（AMA）", "品牌是一种名称、术语、标记、符号或设计，或是它们之间的组合运用，其目的是借以辨认某个销售者或某群销售者的产品或服务，并使之与竞争对手的产品或服务区别开来。"),
    ("大卫·奥格威（David Ogilvy）", "品牌是消费者对产品或服务的感知总和，包括名称、包装、价格、广告风格等一切有形和无形要素。"),
    ("戴维·阿克（David Aaker）", "品牌是超越产品或服务本身的有形资产，它代表了名称、声誉和消费者与品牌之间的情感联系。"),
    ("凯文·凯勒（Kevin Keller）", "品牌是消费者头脑中与品牌名称、标识等元素相关的所有品牌知识的总和。")
]

for scholar, definition in definitions:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    run1 = para.add_run(f"• {scholar}：")
    set_run_font(run1, bold=True)
    para.add_run("\n")
    run2 = para.add_run(f"  {definition}")
    set_run_font(run2)

doc.add_paragraph()

add_heading(doc, "二、品牌与产品/服务的区别", level=2)

distinction_table = doc.add_table(rows=5, cols=3)
distinction_table.style = 'Table Grid'
distinction_data = [
    ("维度", "产品/服务", "品牌"),
    ("本质", "功能性利益", "情感性利益+功能性利益"),
    ("可持续性", "容易被模仿", "难以被复制"),
    ("价值来源", "技术创新或成本优势", "消费者认知和情感连接"),
    ("生命周期", "有限", "可以无限延续")
]
for i, row_data in enumerate(distinction_data):
    for j, cell_data in enumerate(row_data):
        distinction_table.cell(i, j).text = cell_data
        for para in distinction_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "三、品牌的三个层次", level=2)

brand_levels = [
    ("第一层：标识层", "这是品牌最基础的功能，包括品牌名称、标志、色彩、字体等视觉识别元素。", "企业大学的校徽、校名、标准色都属于这一层次。"),
    ("第二层：承诺层", "品牌向消费者做出的价值承诺，包括产品或服务的质量保证、功能特性、利益主张等。", "企业大学对学员承诺的培养效果、课程质量、认证价值等。"),
    ("第三层：关系层", "品牌与消费者之间建立的情感联系和信任关系，这是品牌资产的真正来源。", " alumni对母校的归属感、雇主对人才的认可度等。")
]

for level_name, description, example in brand_levels:
    para = doc.add_paragraph()
    run1 = para.add_run(f"• {level_name}：")
    set_run_font(run1, bold=True)
    para.add_run("\n")
    run2 = para.add_run(f"  {description}")
    set_run_font(run2)
    para.add_run("\n")
    run3 = para.add_run(f"  【案例】{example}")
    set_run_font(run3, color=RGBColor(100, 100, 100))
    para.paragraph_format.space_after = Pt(10)

doc.add_paragraph()

add_heading(doc, "四、企业大学品牌的特殊性", level=2)

specialty_points = [
    "双重属性：企业大学品牌既具有企业品牌的属性（如企业文化、战略导向），又具有教育品牌的属性（如学术声誉、教学质量）。",
    "利益相关者多元：学员、员工、合作伙伴、社会公众、政府监管机构等都是企业大学的利益相关者。",
    "价值交付滞后：品牌承诺的兑现往往需要较长时间才能验证，如人才培养效果的显现。",
    "内部与外部的统一：对外品牌建设必须以内部品牌认同为基础，只有内部员工认可的品牌理念才能真正传递出去。"
]
for point in specialty_points:
    add_bullet_para(doc, point)

doc.add_paragraph()

# --- 知识点1.2 ---
add_heading(doc, "知识点1.2：企业大学品牌的战略价值", level=1)

core_concept_para2 = doc.add_paragraph()
run1 = core_concept_para2.add_run("核心概念：")
set_run_font(run1, bold=True, color=RGBColor(0, 51, 102))
run2 = core_concept_para2.add_run("品牌是企业大学最重要的无形资产")
set_run_font(run2, color=RGBColor(0, 51, 102))
core_concept_para2.paragraph_format.space_after = Pt(12)

add_heading(doc, "一、战略价值三角形：愿景、使命、价值观", level=2)

triangle_para = doc.add_paragraph()
triangle_para.add_run("战略价值三角形是品牌建设的根基，它回答了企业大学最根本的三个问题：")
triangle_para.paragraph_format.space_after = Pt(8)

triangle_table = doc.add_table(rows=4, cols=2)
triangle_table.style = 'Table Grid'
triangle_data = [
    ("维度", "核心问题"),
    ("愿景（Vision）", "我们要成为什么？——长远目标和社会价值"),
    ("使命（Mission）", "我们做什么？——核心业务和存在理由"),
    ("价值观（Values）", "我们怎么做？——行为准则和文化基因")
]
for i, row_data in enumerate(triangle_data):
    for j, cell_data in enumerate(row_data):
        triangle_table.cell(i, j).text = cell_data
        for para in triangle_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "二、品牌对内凝聚共识、对外建立信任", level=2)

trust_points = [
    ("对内：凝聚共识", [
        "统一认知：品牌是企业和员工共同认可的价值观和行为准则",
        "激励员工：强势品牌能够激发员工的自豪感和归属感",
        "指导决策：在面临选择时，品牌理念提供判断标准"
    ]),
    ("对外：建立信任", [
        "降低选择成本：强势品牌减少利益相关者的决策风险感知",
        "创造溢价：品牌信任带来更高的价格接受度",
        "形成壁垒：品牌忠诚度构建起竞争对手难以逾越的护城河"
    ])
]

for category, points in trust_points:
    para = doc.add_paragraph()
    run = para.add_run(category)
    set_run_font(run, bold=True)
    para.paragraph_format.space_after = Pt(6)
    for point in points:
        add_bullet_para(doc, point, level=1)

doc.add_paragraph()

add_heading(doc, "三、企业大学品牌的利益相关者分析", level=2)

stakeholder_para = doc.add_paragraph()
stakeholder_para.add_run("企业大学的利益相关者可分为三大群体，每个群体对品牌有不同的期望和感知：")
stakeholder_para.paragraph_format.space_after = Pt(8)

stakeholder_table = doc.add_table(rows=4, cols=4)
stakeholder_table.style = 'Table Grid'
stakeholder_data = [
    ("利益相关者", "核心期望", "品牌接触点", "关键诉求"),
    ("内部员工", "职业发展、归属感", "培训经历、文化活动", "专业性、认同感"),
    ("学员及其家属", "能力提升、就业保障", "课程体验、证书认证", "质量、声誉"),
    ("企业客户/雇主", "人才适配、ROI", "人才输送、合作项目", "效果、效率")
]
for i, row_data in enumerate(stakeholder_data):
    for j, cell_data in enumerate(row_data):
        stakeholder_table.cell(i, j).text = cell_data
        for para in stakeholder_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "四、品牌资产的财务价值体现", level=2)

financial_points = [
    "人才吸引成本降低：强势品牌使招聘更容易，降低猎头费用和招聘成本。",
    "员工保留率提升：品牌认同感高的员工流失率更低，减少重置成本。",
    "合作溢价：合作伙伴愿意为强势品牌支付更高的合作费用。",
    "校友捐赠：品牌情感连接转化为校友和社会捐赠意愿。",
    "估值溢价：在并购或上市时，品牌资产计入企业估值。"
]
for point in financial_points:
    add_bullet_para(doc, point)

doc.add_paragraph()

# --- 知识点1.3 ---
add_heading(doc, "知识点1.3：品牌资产的构成要素", level=1)

core_concept_para3 = doc.add_paragraph()
run1 = core_concept_para3.add_run("核心概念：")
set_run_font(run1, bold=True, color=RGBColor(0, 51, 102))
run2 = core_concept_para3.add_run("品牌资产 = 品牌知名度 + 品牌联想 + 品牌忠诚度 + 感知质量 + 品牌专有资产")
set_run_font(run2, color=RGBColor(0, 51, 102))
core_concept_para3.paragraph_format.space_after = Pt(12)

add_heading(doc, "一、Aaker品牌资产模型解读", level=2)

aaker_para = doc.add_paragraph()
aaker_para.add_run("大卫·阿克的品牌资产模型包含五个核心维度：")
aaker_para.paragraph_format.space_after = Pt(8)

aaker_elements = [
    ("品牌知名度", "消费者对品牌的认知和记忆程度。从无知名度→提示知名度→未提示知名度→第一提及知名度。"),
    ("品牌联想", "与品牌相关的所有记忆和概念。包括产品属性、使用场景、企业个性、用户类型、情感连接等。"),
    ("感知质量", "消费者对产品或服务整体质量的主观判断。并非客观质量，而是相对于期望的感知。"),
    ("品牌忠诚度", "消费者对品牌的持续偏好和重复购买意愿。表现为拒绝竞争品牌、价格容忍度、品牌推荐意愿。"),
    ("品牌专有资产", "可被法律保护的知识产权，如商标、专利、渠道关系等。")
]

aaker_table = doc.add_table(rows=6, cols=2)
aaker_table.style = 'Table Grid'
aaker_table.cell(0, 0).text = "要素"
aaker_table.cell(0, 1).text = "定义与说明"
for para in aaker_table.cell(0, 0).paragraphs:
    for run in para.runs:
        run.font.bold = True
        set_run_font(run)
for para in aaker_table.cell(0, 1).paragraphs:
    for run in para.runs:
        run.font.bold = True
        set_run_font(run)

for i, (element, description) in enumerate(aaker_elements, 1):
    aaker_table.cell(i, 0).text = element
    aaker_table.cell(i, 1).text = description
    for j in range(2):
        for para in aaker_table.cell(i, j).paragraphs:
            for run in para.runs:
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "二、Keller品牌知识模型解读", level=2)

keller_para = doc.add_paragraph()
keller_para.add_run("凯文·凯勒的品牌知识模型从消费者视角出发，强调品牌知识的两个维度：")
keller_para.paragraph_format.space_after = Pt(8)

keller_table = doc.add_table(rows=3, cols=2)
keller_table.style = 'Table Grid'
keller_data = [
    ("维度", "内容"),
    ("品牌认知（Brand Recall）", "品牌在消费者记忆中的强度。包括品牌识别（认出品牌）和品牌记忆（在相关类别中想起品牌）。"),
    ("品牌形象（Brand Image）", "消费者对品牌的整体感知。包括品牌联想的类型、 favorability（喜爱程度）、strength（强度）和uniqueness（独特性）。")
]
for i, row_data in enumerate(keller_data):
    for j, cell_data in enumerate(row_data):
        keller_table.cell(i, j).text = cell_data
        for para in keller_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "三、Interbrand品牌价值评估方法", level=2)

interbrand_intro = doc.add_paragraph()
interbrand_intro.add_run("Interbrand是全球最权威的品牌价值评估机构之一，其方法论包含三个核心要素：")
interbrand_intro.paragraph_format.space_after = Pt(8)

interbrand_formula = doc.add_paragraph()
interbrand_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = interbrand_formula.add_run("品牌价值 = 品牌带来的利润 × 品牌指数")
set_run_font(run, bold=True, size=14)
interbrand_formula.paragraph_format.space_after = Pt(12)

interbrand_elements = [
    ("品牌带来的利润", "通过计算品牌对产品或服务溢价的贡献，剥离非品牌因素（如渠道、成本等）后的利润增量。"),
    ("品牌指数", "反映品牌对未来利润贡献能力的系数。基于品牌角色（理性vs感性）、品牌一致性、品牌支持、品牌反应性、品牌共鸣等维度评估。")
]
for element, description in interbrand_elements:
    para = doc.add_paragraph()
    run1 = para.add_run(f"• {element}：")
    set_run_font(run1, bold=True)
    para.add_run("\n")
    run2 = para.add_run(f"  {description}")
    set_run_font(run2)
    para.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

add_heading(doc, "四、企业大学品牌资产的特殊性", level=2)

cu_specialty = [
    ("教育属性的权重更高", "相比商业品牌，企业大学的感知质量和品牌联想中，教育价值（学术声誉、师资水平、课程体系）占更大权重。"),
    ("关系资产的核心地位", "企业大学与学员之间的时间跨度长、连接点密集，品牌忠诚度更多基于长期关系而非交易。"),
    ("难以量化的社会资本", "校友网络、行业影响力、思想领导力等社会资本难以用传统财务模型衡量，但价值巨大。"),
    ("内部品牌的镜像效应", "内部员工对品牌的认知和认同直接影响外部品牌形象，需要内外统一的品牌建设策略。")
]
for title, description in cu_specialty:
    para = doc.add_paragraph()
    run1 = para.add_run(f"• {title}：")
    set_run_font(run1, bold=True)
    run2 = para.add_run(description)
    set_run_font(run2)
    para.paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# --- 知识点1.4 ---
add_heading(doc, "知识点1.4：企业大学品牌建设现状分析", level=1)

core_concept_para4 = doc.add_paragraph()
run1 = core_concept_para4.add_run("核心概念：")
set_run_font(run1, bold=True, color=RGBColor(0, 51, 102))
run2 = core_concept_para4.add_run("现状分析是品牌建设的起点")
set_run_font(run2, color=RGBColor(0, 51, 102))
core_concept_para4.paragraph_format.space_after = Pt(12)

add_heading(doc, "一、企业大学品牌建设的常见误区", level=2)

mistakes = [
    ("重硬件轻软件", "投入大量资金建设校园、购买设备，但忽视品牌建设和文化塑造。硬件可以快速复制，软实力才是差异化竞争的核心。"),
    ("将VI等同于品牌", "认为设计一套视觉识别系统就是完成了品牌建设，忽视了品牌战略和品牌体验的系统性。"),
    ("内部视角主导", "从企业视角出发定义品牌价值，而没有真正了解目标受众（学员、雇主）的真实需求和感知。"),
    ("短期效果导向", "追求立竿见影的品牌效果，忽视品牌资产的长期积累特性。品牌建设是马拉松，不是短跑。"),
    ("品牌建设孤立化", "将品牌部门视为独立运作的单元，没有将品牌理念融入课程设计、师资管理、学员服务等各个环节。")
]

for i, (mistake, description) in enumerate(mistakes, 1):
    para = doc.add_paragraph()
    run1 = para.add_run(f"误区{i}：{mistake}")
    set_run_font(run1, bold=True)
    para.add_run("\n")
    run2 = para.add_run(f"   {description}")
    set_run_font(run2)
    para.paragraph_format.space_after = Pt(10)

doc.add_paragraph()

add_heading(doc, "二、内部视角vs外部视角的差异", level=2)

perspective_table = doc.add_table(rows=4, cols=3)
perspective_table.style = 'Table Grid'
perspective_data = [
    ("维度", "内部视角", "外部视角"),
    ("品牌定义", "我们是什么", "我感知到什么"),
    ("价值主张", "我们提供什么", "我能获得什么"),
    ("成功标准", "完成既定目标", "满足甚至超越期望")
]
for i, row_data in enumerate(perspective_data):
    for j, cell_data in enumerate(row_data):
        perspective_table.cell(i, j).text = cell_data
        for para in perspective_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

perspective_note = doc.add_paragraph()
run = perspective_note.add_run("关键洞察：优秀的品牌战略必须同时满足内部一致性和外部差异化。品牌建设的本质是在内外部之间建立一座桥梁，让内部价值能够被外部感知和认可。")
set_run_font(run)
perspective_note.paragraph_format.space_after = Pt(12)

add_heading(doc, "三、品牌成熟度评估模型", level=2)

maturity_intro = doc.add_paragraph()
maturity_intro.add_run("企业大学品牌成熟度可分为五个阶段：")
maturity_intro.paragraph_format.space_after = Pt(8)

maturity_table = doc.add_table(rows=6, cols=3)
maturity_table.style = 'Table Grid'
maturity_data = [
    ("阶段", "特征", "典型表现"),
    ("1. 初始级", "无意识品牌建设", "品牌元素不统一，无专门团队管理"),
    ("2. 职能级", "职能部门驱动", "市场部或培训部各自为政"),
    ("3. 协调级", "跨部门协作", "建立品牌管理委员会，但缺战略顶层设计"),
    ("4. 战略级", "品牌战略清晰", "品牌与业务战略对齐，有系统评估体系"),
    ("5. 卓越级", "品牌驱动发展", "品牌成为核心竞争力，持续创新引领")
]
for i, row_data in enumerate(maturity_data):
    for j, cell_data in enumerate(row_data):
        maturity_table.cell(i, j).text = cell_data
        for para in maturity_table.cell(i, j).paragraphs:
            for run in para.runs:
                if i == 0:
                    run.font.bold = True
                set_run_font(run)

doc.add_paragraph()

add_heading(doc, "四、案例：标杆企业大学品牌建设实践", level=2)

case_intro = doc.add_paragraph()
case_intro.add_run("以下案例展示了几家知名企业大学的品牌建设实践：")
case_intro.paragraph_format.space_after = Pt(8)

cases = [
    ("GE克劳顿维尔管理学院", "作为企业大学的先驱，GE将克劳顿维尔定位为'领导力培养的圣地'。品牌核心价值：诚信、业绩、变革。通过持续输出管理思想和培养全球商业领袖，建立了强大的品牌影响力。"),
    ("华为培训中心", "华为将培训品牌与'狼性文化'和'以奋斗者为本'的企业精神深度绑定。品牌特色：实战导向、训战结合。培养体系与业务需求紧密耦合，形成独特的培训品牌资产。"),
    ("腾讯学院", "腾讯学院强调'连接一切'的互联网思维，在品牌传播上善用内部社交平台和标杆人物打造。品牌特色：创新、开放、连接。建立了技术类人才培训的专业权威。")
]

for company, description in cases:
    para = doc.add_paragraph()
    run1 = para.add_run(f"【{company}】")
    set_run_font(run1, bold=True)
    para.add_run("\n")
    run2 = para.add_run(f"   {description}")
    set_run_font(run2)
    para.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# ==================== 互动表单 ====================

add_heading(doc, "互动表单", level=1)

add_heading(doc, "表单1.1：品牌认知自评", level=2)

self_eval_intro = doc.add_paragraph()
self_eval_intro.add_run("填写说明：请根据您的实际情况填写，1=完全不符合，5=完全符合")
self_eval_intro.paragraph_format.space_after = Pt(8)

self_eval_questions = [
    "我认为品牌就是企业的名称和标志",
    "我相信强势品牌能为企业带来溢价",
    "我清楚知道我们企业大学的品牌定位",
    "我认为品牌建设是每个人的责任",
    "我相信品牌资产可以转化为财务价值"
]

for i, question in enumerate(self_eval_questions, 1):
    para = doc.add_paragraph()
    run = para.add_run(f"{i}. {question}")
    set_run_font(run)
    para.paragraph_format.space_after = Pt(6)
    # 添加评分框提示
    score_hint = doc.add_paragraph()
    score_run = score_hint.add_run("   评分：1 —— 2 —— 3 —— 4 —— 5")
    set_run_font(score_run, color=RGBColor(150, 150, 150))
    score_hint.paragraph_format.space_after = Pt(10)

doc.add_paragraph()

add_heading(doc, "表单1.2：企业大学品牌现状诊断", level=2)

diagnosis_intro = doc.add_paragraph()
diagnosis_intro.add_run("填写说明：对照检查您所在企业的品牌建设现状，在符合的选项上打勾")
diagnosis_intro.paragraph_format.space_after = Pt(8)

diagnosis_items = [
    ("品牌知名度", ["学员主动申请", "同行知晓", "行业媒体曝光"]),
    ("品牌一致性", ["视觉识别统一", "信息传播一致", "员工行为符合品牌价值观"]),
    ("品牌体验", ["课程质量稳定", "服务流程标准化", "学员反馈机制健全"]),
    ("品牌关系", ["学员满意度高", "校友活跃度高", "雇主认可度强"])
]

for category, options in diagnosis_items:
    para = doc.add_paragraph()
    run = para.add_run(f"• {category}")
    set_run_font(run, bold=True)
    para.paragraph_format.space_after = Pt(4)
    for option in options:
        option_para = doc.add_paragraph()
        option_run = option_para.add_run(f"  □ {option}")
        set_run_font(option_run)
        option_para.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# ==================== 练习题 ====================

add_heading(doc, "练习题", level=1)

add_heading(doc, "练习1.1：品牌定义辨析", level=2)

exercise1_intro = doc.add_paragraph()
exercise1_intro_run = exercise1_intro.add_run("案例背景：")
set_run_font(exercise1_intro_run, bold=True)
exercise1_case = doc.add_paragraph()
exercise1_case_run = exercise1_case.add_run('某科技企业大学A定位为"技术人才培养专家"，投入大量资源开发技术课程，校名中包含"技术"二字，招生时重点宣传其技术课程项目。然而，调研发现：学员普遍反映课程内容偏理论、与业务脱节；HR部门反馈招聘时候选人不了解该企业大学；市场上存在多家提供类似技术培训的非学历教育机构，学员认为替代性很强。')
set_run_font(exercise1_case_run)
exercise1_case.paragraph_format.space_after = Pt(8)

exercise1_questions = [
    "企业大学A的品牌建设存在哪些问题？",
    "从品牌资产模型角度分析，该企业大学的品牌资产状况如何？",
    "如果你是品牌顾问，你会提出哪些改进建议？"
]

exercise1_para = doc.add_paragraph()
exercise1_para_run = exercise1_para.add_run("问题：")
set_run_font(exercise1_para_run, bold=True)
for i, q in enumerate(exercise1_questions, 1):
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(f"{i}. {q}")
    set_run_font(q_run)
    q_para.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

answer1_label = doc.add_paragraph()
answer1_label_run = answer1_label.add_run("参考答案：")
set_run_font(answer1_label_run, bold=True)

answer1_content = """
1. 问题分析：
   • 品牌承诺与实际体验不符（承诺"技术专家"但内容偏理论）
   • 品牌识别薄弱（仅靠校名，缺乏独特视觉和传播）
   • 品牌联想单一（仅限技术课程，缺乏情感连接）
   • 品牌差异化不足（可替代性强）

2. 品牌资产分析：
   • 品牌知名度：低（候选人不了解）
   • 品牌联想：弱（联想单一且与负面感知关联）
   • 感知质量：中等（课程质量反馈一般）
   • 品牌忠诚度：低（缺乏情感连接）

3. 改进建议：
   • 重新定义品牌定位，聚焦差异化优势
   • 建立品牌承诺与实际体验的一致性
   • 强化品牌识别系统的差异化设计
   • 打造品牌特色项目和标杆案例
   • 建立雇主品牌联动机制提升知名度
"""

answer1_para = doc.add_paragraph()
answer1_para_run = answer1_para.add_run(answer1_content)
set_run_font(answer1_para_run)

doc.add_paragraph()

add_heading(doc, "练习1.2：品牌资产计算", level=2)

exercise2_intro = doc.add_paragraph()
exercise2_intro_run = exercise2_intro.add_run("案例背景：")
set_run_font(exercise2_intro_run, bold=True)
exercise2_case = doc.add_paragraph()
exercise2_case_run = exercise2_case.add_run("某企业大学B年营收5000万元，假设品牌贡献系数为0.3，行业品牌指数为1.5。请计算该企业大学的品牌价值。")
set_run_font(exercise2_case_run)
exercise2_case.paragraph_format.space_after = Pt(8)

exercise2_questions = [
    "计算该企业大学的品牌价值",
    "如果品牌贡献系数提升到0.4，品牌价值增加多少？",
    "分析影响品牌贡献系数的关键因素"
]

exercise2_para = doc.add_paragraph()
exercise2_para_run = exercise2_para.add_run("问题：")
set_run_font(exercise2_para_run, bold=True)
for i, q in enumerate(exercise2_questions, 1):
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(f"{i}. {q}")
    set_run_font(q_run)
    q_para.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

answer2_label = doc.add_paragraph()
answer2_label_run = answer2_label.add_run("参考答案：")
set_run_font(answer2_label_run, bold=True)

answer2_content = """
1. 品牌价值计算：
   品牌价值 = 5000万 × 0.3 × 1.5 = 2250万元

2. 品牌贡献系数提升后的计算：
   新品牌价值 = 5000万 × 0.4 × 1.5 = 3000万元
   品牌价值增加 = 3000万 - 2250万 = 750万元

3. 影响品牌贡献系数的关键因素：
   • 品牌在购买决策中的角色（高介入度vs低介入度）
   • 品牌的情感vs理性定位（情感型品牌通常系数更高）
   • 品牌的市场份额和竞争地位
   • 品牌的一致性和清晰度
   • 品牌的历史积累和认知深度
"""

answer2_para = doc.add_paragraph()
answer2_para_run = answer2_para.add_run(answer2_content)
set_run_font(answer2_para_run)

doc.add_paragraph()

# ==================== 知识框架 ====================

add_heading(doc, "知识框架", level=1)

framework_intro = doc.add_paragraph()
framework_intro.add_run("本章知识结构可概括为以下框架：")
framework_intro.paragraph_format.space_after = Pt(8)

framework_diagram = doc.add_paragraph()
framework_diagram.add_run("""
                    ┌─────────────────────────────────────────────┐
                    │         模块一：品牌本质                     │
                    └─────────────────────────────────────────────┘
                                      │
            ┌─────────────────────────┼─────────────────────────┐
            │                         │                         │
            ▼                         ▼                         ▼
    ┌───────────────┐         ┌───────────────┐         ┌───────────────┐
    │ 1.1 品牌定义   │         │ 1.2 战略价值   │         │ 1.3 品牌资产   │
    │               │         │               │         │               │
    │ • 经典定义    │         │ • 战略三角形   │         │ • Aaker模型   │
    │ • 三个层次    │         │ • 内外价值    │         │ • Keller模型  │
    │ • 特殊性     │         │ • 利益相关者  │         │ • Interbrand  │
    └───────────────┘         │ • 财务价值    │         └───────────────┘
                              └───────────────┘                  │
                                    │                            │
                                    └──────────┬─────────────────┘
                                               │
                                               ▼
                                    ┌─────────────────┐
                                    │ 1.4 现状分析    │
                                    │                 │
                                    │ • 常见误区     │
                                    │ • 内外视角     │
                                    │ • 成熟度模型   │
                                    │ • 标杆案例     │
                                    └─────────────────┘
""")
framework_diagram.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# ==================== 行为承诺 ====================

add_heading(doc, "行为承诺", level=1)

commitment_intro = doc.add_paragraph()
commitment_intro.add_run("我承诺在课程结束后一周内完成以下行动：")
commitment_intro.paragraph_format.space_after = Pt(8)

commitment_items = [
    "完成本机构品牌现状的自我诊断，识别至少3个品牌建设改进点",
    "与部门负责人沟通品牌建设现状，制定初步品牌提升计划",
    "阅读推荐书目：《战略品牌管理》（凯文·凯勒著）相关章节",
    "收集3个以上行业标杆企业大学品牌建设案例"
]

for i, item in enumerate(commitment_items, 1):
    para = doc.add_paragraph()
    run = para.add_run(f"□ {i}. {item}")
    set_run_font(run)
    para.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# 签名栏
signature_table = doc.add_table(rows=2, cols=2)
signature_table.style = 'Table Grid'
signature_table.cell(0, 0).text = "学员签名："
signature_table.cell(0, 1).text = "日期："
signature_table.cell(1, 0).text = "讲师签名："
signature_table.cell(1, 1).text = "日期："

for row in signature_table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run)
                run.font.size = Pt(12)

# ==================== 保存文档 ====================
doc.save(output_path)
print(f"文档已保存至: {output_path}")

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

def add_chapter_content(doc, ch):
    h = doc.add_paragraph()
    r = h.add_run(ch['title'])
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,51,102)

    qp = doc.add_paragraph()
    qp.paragraph_format.left_indent = Cm(1)
    r = qp.add_run(f"金句：{ch['gold_quote']}")
    r.italic = True; r.font.color.rgb = RGBColor(128,0,0)

    oh = doc.add_paragraph()
    r = oh.add_run('章节学习目标'); r.bold = True
    for obj in ch['objectives']:
        p = doc.add_paragraph(obj, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    cph = doc.add_paragraph()
    r = cph.add_run('核心概念'); r.bold = True
    for term, defn in ch['concepts']:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{term}：'); r.bold = True
        p.add_run(defn)

    kh = doc.add_paragraph()
    r = kh.add_run('关键要点'); r.bold = True
    for j, pt in enumerate(ch['key_points'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f'{j}. '); r.bold = True
        p.add_run(pt)

    th = doc.add_paragraph()
    r = th.add_run('思考题'); r.bold = True
    for t in ch['thinking']:
        p = doc.add_paragraph(t, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    eh = doc.add_paragraph()
    r = eh.add_run('练习题'); r.bold = True
    for e in ch['exercise']:
        p = doc.add_paragraph(e, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1)

    sh = doc.add_paragraph()
    r = sh.add_run('章节小结'); r.bold = True
    sp = doc.add_paragraph()
    sp.paragraph_format.left_indent = Cm(1)
    r = sp.add_run(ch['summary']); r.italic = True

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# ========== PART 2 ==========
p2 = doc.add_paragraph()
r = p2.add_run('PART 2')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

p2s = doc.add_paragraph()
r = p2s.add_run('组织落地：从个人经验到组织能力')
r.bold = True; r.font.size = Pt(16)

chapters_p2 = [
    {
        'title': '第一章  没人认领是最大的敌人',
        'gold_quote': '一份没有主人的工具，用一次就是最后一次。',
        'objectives': ['理解认领不是保管文档，而是对持续使用和更新负责', '掌握如何指定认领人和备份认领人', '理解复盘引导者不适合做认领人的原因'],
        'concepts': [('认领', '对决策卡负有三个具体责任：定期回访是否被使用、收集反馈、主动判断是否需要更新'), ('主认领人', '最贴近使用场景的岗位负责人'), ('备份认领人', '主认领人离岗时自然接手')],
        'key_points': ['做完一张卡，真正的工作才刚开始一半', '认领不是指定部门或岗位，是具体到一个真实的人', '认领人的三件具体事：定期回访使用情况、收集反馈、主动判断更新', '复盘引导者不适合做认领人', '明确写清楚主责和备份各自做什么'],
        'thinking': ['你组织做过的那些工具、方法论，后来还有人在用吗？', '为什么岗位负责人比项目负责人更适合做认领人？'],
        'exercise': ['为你的一张决策卡指定认领人和备份认领人，明确各自职责'],
        'summary': '决策卡最大的敌人不是会不会用，是没人认领。做完一张卡，真正的工作才刚开始一半，另一半是让它找到自己的主人。'
    },
    {
        'title': '第二章  稽核不是查错是记住学过什么',
        'gold_quote': '稽核表查的不是做对了没有，是这次的判断有没有被沉淀下来。',
        'objectives': ['理解稽核不是绩效考核，是知识管理', '掌握决策稽核的四个核心问题', '学会把书面稽核变成面对面对话'],
        'concepts': [('决策稽核', '检查判断逻辑有没有被沉淀成组织可复用东西的机制'), ('沉淀检查', '问的是沉淀这个动作有没有发生，而不是这次决策对不对'), ('稽核对话', '和决策者面对面聊十几分钟，把值得记的东西挖出来')],
        'key_points': ['把稽核做成考核，关注结果，使用者天然进入防御姿态', '稽核表查的是沉淀有没有发生，不是这次决策对不对', '四个核心问题：有没有不确定节点被记录？现有卡片是否被使用？是否具备做卡价值？卡片是否有漏洞需要反馈？', '稽核不适合套用在每一次决策', '不要做成书面填表，做成面对面简短的对话'],
        'thinking': ['你公司现有的检查机制，是让人产生防御还是真正在积累知识？', '为什么问对不对会让人防御，但问有没有沉淀不会？'],
        'exercise': ['设计一次稽核对话，用四个核心问题和一个决策者聊'],
        'summary': '决策稽核不是查错，是替组织记住它学过什么。查沉淀有没有发生，才能让每一次决策，不管结果如何，都变成组织真正拿到手里的东西。'
    },
    {
        'title': '第三章  会过时是活着的证据',
        'gold_quote': '一张三年没改过的决策卡，大概率已经没人真的在用。',
        'objectives': ['理解迭代是决策卡活着的证明', '掌握持续迭代的两层机制', '学会判断卡片是否该废止而非删除'],
        'concepts': [('迭代记录', '判断决策卡是否还活着的信号'), ('固定复审周期', '通常半年或一年一次，由认领人主动召集复审'), ('废止归档', '不再适用的卡标注废止但保留归档')],
        'key_points': ['长期不变、没有任何修改记录的卡，大概率已经被遗忘', '迭代频繁的卡，说明有人在真实场景里频繁打开它、对照它', '两层机制：稽核对话时问有没有发现内容和实际对不上 + 设定固定复审周期', '有些卡对应业务场景已不存在，与其留着占据位置，不如标注废止', '废止不等于删除——业务模式可能轮回式重现'],
        'thinking': ['你电脑里那个三年前做的文档，现在打开来还能用吗？', '为什么经得起时间考验反而可能是已经没人用了的意思？'],
        'exercise': ['检查你手头的决策卡，有哪些超过一年没更新过？评估它们是否还活着'],
        'summary': '好的决策卡会过时，这是它活着的证据。真正被使用的卡，会随着业务环境的变化不断被人发现漏洞、提出更新，这种持续的磨损和修补，才是一张卡真正在为组织工作的证明。'
    },
    {
        'title': '第四章  最易犯的错是替决策者总结',
        'gold_quote': '你替他说出来的道理，他转身就忘；他自己说出来的判断，才是他的。',
        'objectives': ['理解替决策者总结为什么是错的', '学会在关键时刻克制住帮他说出来的冲动', '掌握容忍沉默的技巧'],
        'concepts': [('facilitator vs 讲师', 'facilitator创造环境让对方自己挖出来，讲师准确传递知识'), ('沉默五秒', '遇到决策者陷入思考时，默数到五再决定要不要开口'), ('两次角色互换', '第一次问你会提醒新人什么，追问第二次知道了还可能在哪判断错')],
        'key_points': ['替决策者说出隐性判断规则，剥夺了他自己第一次亲口讲清楚这个过程', 'facilitator不是讲师', '在访谈接近挖出隐性规则的关键时刻，要克制住帮他说出来的冲动', '容忍沉默——默数到五，是决策者自己在脑子里组织语言的关键时刻'],
        'thinking': ['你有没有过这样的经历——别人替你总结了一个观点，你后来记不住？', '为什么教比告诉更难忘记？'],
        'exercise': ['在下一次访谈中，练习在关键问题后容忍五秒沉默，不急着补充'],
        'summary': '复盘引导者最容易犯的错，是替决策者总结。你替他说出来的道理，说得再准确，也只是你的归纳，他转身就可能忘记；他自己一句一句摸索着讲出来的判断，才真正是他的东西。'
    },
    {
        'title': '第五章  交付的是组织判断力',
        'gold_quote': '决策卡会被淘汰，做卡的能力不会。',
        'objectives': ['理解会过时的不是真正资产的深层含义', '掌握判断组织是否真正内化了方法论的标准', '理解项目节奏：前期主导，后期交还给团队'],
        'concepts': [('组织判断力', '团队学会识别值得复盘的决策、访谈挖隐性判断、把判断变成可复用触发条件的能力'), ('内化标准', '合作结束后，团队能在没有外部引导者的情况下主动提出做决策卡'), ('能力传递', '让团队自己持续产出新的决策卡，不需要每次都靠外部引导者介入')],
        'key_points': ['卡片会过时，会过时的不是真正的资产', '真正的资产是团队学会了这套方法论', '项目前半段以主导访谈和撰写为主，让团队看到完整方法论跑一遍', '项目后半段逐步把主导权交还给团队内部指定人选，退到旁观和纠偏', '内化标准：团队主动提出这次值得做一张卡'],
        'thinking': ['你带过的团队，现在能自己独立做这件事了吗？', '如果现在派你去另一个项目，你能留下的最重要的东西是什么？'],
        'exercise': ['评估你所在团队的内化程度——他们现在能自己识别值得做卡的决策吗？'],
        'summary': '这份工作最终交付的不是文档，是组织的判断力。重要的是那套挖掘判断、结构化判断、让判断可以被复制的能力，有没有真正留在了团队自己手里。这才是这份工作真正想要留下的东西。'
    },
]

for i, ch in enumerate(chapters_p2):
    add_chapter_content(doc, ch)
    if i < len(chapters_p2) - 1:
        doc.add_page_break()

print('Part 2 chapters done')

doc.add_page_break()

# PART 2 Q&A
qa2 = doc.add_paragraph()
r = qa2.add_run('PART 2  Q&A')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0,51,102)

qas2 = [
    ('问：决策卡推行半年后使用率下降，是卡片本身的问题还是推行方式的问题？', '先做简单排查——去问几个不再使用的人，具体是这张卡内容不准了还是忘了这回事，也没人提醒。'),
    ('问：高层觉得决策卡太麻烦，只想要一页纸的总结，怎么办？', '做两个版本——给高层看的确实可以是一页纸摘要，浓缩触发条件和核心判断逻辑；但真正给一线使用者用的完整版，该有的细节不能被砍掉。'),
    ('问：有些决策卡涉及的判断带有强烈的个人风格，别人真的能复制吗？', '决策卡复制的从来不是这个人的风格，是这个人在特定信号出现时的判断结构。风格没法复制也不需要复制；但看到这个信号该往哪个方向想这个结构性的东西，是可以脱离具体的人被传递下去的。'),
    ('问：稽核表会不会让团队变得只敢按卡片行事，反而不敢做真正的判断？', '稽核对话里可以主动加一个问题——这次你有没有基于自己的判断，偏离了卡片上的建议？为什么？如果答案一直是没有，这本身就是一个信号。'),
    ('问：我在做这套方法论时，有没有哪个环节到现在还没想透？', '有。场景细分该到什么颗粒度、组织什么时候算真正内化了这套能力，这两个问题目前没有精确的判断标准。'),
]

for q, a in qas2:
    qp = doc.add_paragraph()
    r = qp.add_run(q); r.bold = True
    ap = doc.add_paragraph()
    ap.paragraph_format.left_indent = Cm(1)
    ap.add_run(a)

doc.add_page_break()
print('Part 2 Q&A done')

doc.save(r'D:\CC\temp\merged_part3.docx')
print('Part 3 saved')

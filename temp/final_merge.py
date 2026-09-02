#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
import os
import sys

def merge_docx(source_doc, target_doc):
    """Merge source_doc into target_doc"""
    for element in source_doc.element.body:
        target_doc.element.body.append(element)

# Load all parts
print('Loading parts...')
doc1 = Document(r'D:\CC\temp\merged_part1.docx')
doc2 = Document(r'D:\CC\temp\merged_part2.docx')
doc3 = Document(r'D:\CC\temp\merged_part3.docx')
doc4 = Document(r'D:\CC\temp\merged_part4.docx')

# Create final document
print('Creating final document...')
final = Document()
style = final.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# Merge all parts
print('Merging part 1...')
merge_docx(doc1, final)
print('Merging part 2...')
merge_docx(doc2, final)
print('Merging part 3...')
merge_docx(doc3, final)
print('Merging part 4...')
merge_docx(doc4, final)

# Save to final path
outpath = 'D:\\新课开发\\工作手册\\AI时代决策工作手册\\完整课程包\\04-学员手册\\AI时代决策工作手册-学员手册.docx'
outdir = os.path.dirname(outpath)
print(f'Output directory: {outdir}')

# Try to create directory with unicode path
try:
    os.makedirs(outdir, exist_ok=True)
except Exception as e:
    print(f'Note: {e}')

print(f'Saving to: {outpath}')
final.save(outpath)

# Verify
if os.path.exists(outpath):
    size = os.path.getsize(outpath)
    print(f'SUCCESS! File saved: {outpath}')
    print(f'File size: {size} bytes')
else:
    print('ERROR: File was not created')

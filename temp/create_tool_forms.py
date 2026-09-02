# -*- coding: utf-8 -*-
"""
生成变革管理工具表单Word文档
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 输出目录
OUTPUT_DIR = r"D:\新课开发\变革管理\12-抵抗信号的早期识别：变革失败之前，组织早就发出过警告\完整课程包\06-工具表单"

# 确保目录存在
os.makedirs(OUTPUT_DIR, exist_ok=True)

# 颜色定义
NAVY = RGBColor(0x1F, 0x38, 0x64)      # 深蓝色 - 主标题
BLUE = RGBColor(0x2E, 0x75, 0xB6)       # 中蓝色 - 副标题
LIGHT_BLUE = RGBColor(0x44, 0x72, 0xC4)  # 亮蓝色 - 表头
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)   # 深灰色 - 正文
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66) # 中灰色 - 说明文字
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# 风险等级颜色
RED = RGBColor(0xC0, 0x00, 0x00)         # 高风险
ORANGE = RGBColor(0xFF, 0x6C, 0x0B)      # 中风险
GREEN = RGBColor(0x00, 0x70, 0x00)      # 低风险

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, border_size=4):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), '4472C4')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY if level == 1 else (BLUE if level == 2 else DARK_GRAY)
        run.font.bold = True
    return heading

def add_para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    """添加段落"""
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color:
        run.font.color.rgb = color
    return para

def add_instruction_box(doc, text):
    """添加填写说明框"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("【填写说明】")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    para2 = doc.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = para2.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = MEDIUM_GRAY
    run2.font.name = '微软雅黑'
    run2._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return para

def create_table_with_header(doc, headers, rows, col_widths=None):
    """创建带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 设置表头
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, '4472C4')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    return table

# ============================================================
# 文档1: 预警指标检查表
# ============================================================
def create_01_early_warning_checklist():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4横向
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('预警指标检查表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '四维度抵抗信号识别工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_instruction_box(doc, "使用说明：定期检查以下四个维度的指标，根据观察到的信号进行风险等级评估，并在备注栏记录具体情况。")

    # 四个维度的检查表
    dimensions = [
        {
            'name': '一、行为维度',
            'items': [
                ('会议参与度下降', '迟到、缺席增加，发言减少'),
                ('工作效率变化', '完成时间延长，质量下滑'),
                ('主动性降低', '等待指令才行动'),
                ('团队协作减少', '减少与他人的互动交流'),
                ('出错率上升', '犯错误频率明显增加'),
            ]
        },
        {
            'name': '二、沟通维度',
            'items': [
                ('信息反馈减少', '不再主动汇报工作进展'),
                ('负面言论增加', '私下表达不满情绪'),
                ('谣言传播', '小道消息变得活跃'),
                ('沉默文化', '开会时无人发言或只有少数人说话'),
                ('反馈渠道减少', '减少使用正式反馈渠道'),
            ]
        },
        {
            'name': '三、流程维度',
            'items': [
                ('流程遵守度下降', '开始绕过既定流程'),
                ('制度执行敷衍', '表面应付检查'),
                ('变更抵触', '对新流程的消极抵制'),
                ('资源浪费', '不珍惜变革资源'),
                ('创新意愿降低', '不愿尝试新方法'),
            ]
        },
        {
            'name': '四、文化维度',
            'items': [
                ('归属感减弱', '表达"与我无关"的态度'),
                ('信任度下降', '对管理层决策的质疑增加'),
                ('价值观冲突', '公开质疑变革方向'),
                ('帮派现象', '形成小团体对抗变革'),
                ('离职倾向', '员工开始寻求外部机会'),
            ]
        },
    ]

    # 表头
    headers = ['观察指标', '具体表现', '低风险', '中风险', '高风险', '应对建议', '备注']

    for dim in dimensions:
        add_heading(doc, dim['name'], level=2)

        rows = []
        for item_name, description in dim['items']:
            row = [
                item_name,
                description,
                '☐ 偶尔观察到此信号',
                '☐ 频繁观察到此信号',
                '☐ 已成为普遍现象',
                '',
                ''
            ]
            rows.append(row)

        table = create_table_with_header(doc, headers, rows, [1.2, 1.8, 0.8, 0.8, 0.8, 1.5, 1.3])

    # 底部汇总
    doc.add_paragraph()
    add_heading(doc, '风险汇总与行动计划', level=2)

    summary_headers = ['风险等级', '指标数量', '主要问题', '优先级', '建议介入层级']
    summary_rows = [
        ['高风险（红）', '___项', '', 'P0 - 立即处理', '高层管理者'],
        ['中风险（橙）', '___项', '', 'P1 - 本周处理', '中层管理者'],
        ['低风险（绿）', '___项', '', 'P2 - 持续关注', '直接主管'],
    ]
    create_table_with_header(doc, summary_headers, summary_rows, [1.2, 1.0, 2.5, 1.3, 1.5])

    # 页脚
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 预警指标检查表 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '01-预警指标检查表.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 文档2: 关键人物访谈提纲
# ============================================================
def create_02_key_person_interview():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('关键人物访谈提纲', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '变革管理深度访谈工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 访谈准备清单
    add_heading(doc, '一、访谈准备清单', level=1)
    prep_items = [
        '确定访谈对象（高管/中层/基层/关键意见领袖）',
        '收集访谈对象的基本背景信息',
        '了解访谈对象在变革中的角色和利益关系',
        '准备访谈提纲和记录模板',
        '预约访谈时间（建议30-60分钟）',
        '选择私密、安静的访谈环境',
        '准备录音设备（如需）',
    ]
    for item in prep_items:
        para = doc.add_paragraph()
        run = para.add_run('☐ ' + item)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 不同角色的访谈问题
    add_heading(doc, '二、分角色访谈问题', level=1)

    roles = [
        {
            'title': '2.1 高层管理者访谈',
            'questions': [
                '您认为本次变革的核心目标是什么？',
                '您如何看待变革过程中遇到的阻力？',
                '您计划如何支持变革的推进？',
                '您对变革成功的预期时间线是什么？',
                '您认为组织需要哪些资源来支持变革？',
            ]
        },
        {
            'title': '2.2 中层管理者访谈',
            'questions': [
                '您的团队对变革有什么反应？',
                '您观察到哪些具体的抵抗信号？',
                '您在变革中遇到了哪些困难？',
                '您需要上级提供什么支持？',
                '您对变革计划有什么建议？',
            ]
        },
        {
            'title': '2.3 基层员工访谈',
            'questions': [
                '您理解为什么要进行这次变革吗？',
                '变革对您的日常工作有什么影响？',
                '您对变革有什么担忧或疑问？',
                '您希望从管理层获得什么信息？',
                '您愿意为变革做些什么？',
            ]
        },
        {
            'title': '2.4 关键意见领袖访谈',
            'questions': [
                '您认为变革会给组织带来什么变化？',
                '您的同事对变革有什么看法？',
                '您认为变革的最大障碍是什么？',
                '您愿意帮助传播变革信息吗？',
                '您对管理层有什么建议？',
            ]
        },
    ]

    for role in roles:
        add_heading(doc, role['title'], level=2)
        for i, q in enumerate(role['questions'], 1):
            para = doc.add_paragraph()
            run = para.add_run(f'{i}. {q}')
            run.font.size = Pt(10.5)
            run.font.name = '微软雅黑'
            run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 访谈记录模板
    add_heading(doc, '三、访谈记录模板', level=1)

    record_headers = ['访谈对象', '', '职位/部门', '', '访谈日期', '']
    record_table = doc.add_table(rows=2, cols=6)
    record_table.style = 'Table Grid'

    record_data = [
        ['访谈对象：', '________________', '职位/部门：', '________________', '访谈日期：', '________________'],
        ['访谈时长：', '________________', '访谈地点：', '________________', '访谈人：', '________________'],
    ]
    for row_idx, row_data in enumerate(record_data):
        row = record_table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.bold = (col_idx % 2 == 0)
                    run.font.name = '微软雅黑'
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '\n核心发现记录：')
    findings_table = doc.add_table(rows=5, cols=1)
    findings_table.style = 'Table Grid'
    for row in findings_table.rows:
        row.height = Cm(1)
        for cell in row.cells:
            cell.text = '  '
            set_cell_shading(cell, 'F2F2F2')

    add_para(doc, '\n关键引述（直接引用）：')
    quote_table = doc.add_table(rows=3, cols=1)
    quote_table.style = 'Table Grid'
    for row in quote_table.rows:
        row.height = Cm(1.2)
        for cell in row.cells:
            cell.text = '  '
            set_cell_shading(cell, 'F2F2F2')

    # 访谈分析框架
    add_heading(doc, '四、访谈分析框架', level=1)

    analysis_headers = ['分析维度', '主要发现', '抵抗信号识别', '建议行动']
    analysis_rows = [
        ['信息认知', '', '☐ 存在误解  ☐ 信息不足', ''],
        ['情感态度', '', '☐ 消极抵触  ☐ 中立观望', ''],
        ['行为意愿', '', '☐ 公开反对  ☐ 暗中抵制', ''],
        ['利益影响', '', '☐ 利益受损  ☐ 担心风险', ''],
        ['信任程度', '', '☐ 信任不足  ☐ 存有疑虑', ''],
    ]
    create_table_with_header(doc, analysis_headers, analysis_rows, [1.2, 2.5, 1.8, 2.0])

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 关键人物访谈提纲 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '02-关键人物访谈提纲.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 文档3: 抵抗信号分类矩阵
# ============================================================
def create_03_resistance_classification_matrix():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('抵抗信号分类矩阵', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '组织变革抵抗信号系统化识别与分析工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_instruction_box(doc, "使用方法：根据观察到的抵抗信号，在矩阵中找到对应的分类位置，评估风险等级，并按优先级采取相应措施。")

    # 分类矩阵
    add_heading(doc, '一、信号类型分类表', level=1)

    headers = ['信号类型', '具体表现', '出现场景', '潜在根源', '风险等级']
    rows = [
        ['显性言语抵抗', '直接表达不满、质疑决策', '会议、正式沟通场合', '信息不对称/利益冲突', '高'],
        ['隐性言语抵抗', '冷嘲热讽、散布消极言论', '非正式交流、茶水间', '信任缺失/恐惧', '中高'],
        ['显性行为抵抗', '公开拒绝、罢工、集体行动', '工作场所、公开场合', '强烈反对/组织动员', '高'],
        ['隐性行为抵抗', '磨洋工、破坏设备、降低质量', '日常工作、个人行为', '不满累积/报复心理', '中高'],
        ['政治性抵抗', '拉帮结派、向上告状、操纵信息', '组织内部、权力结构', '权力斗争/自保', '中'],
        ['文化性抵抗', '坚守旧习惯、抵制新价值观', '长期习惯、价值观念', '文化惯性/身份认同', '中'],
        ['技术性抵抗', '不会用/不愿用新系统/工具', '技术应用场景', '能力不足/心理抗拒', '低中'],
        ['程序性抵抗', '拖延、绕过流程、制造障碍', '流程执行环节', '效率优先/控制欲', '低中'],
    ]
    create_table_with_header(doc, headers, rows, [1.3, 1.8, 1.5, 1.5, 0.8])

    # 风险等级评估
    add_heading(doc, '二、风险等级评估标准', level=1)

    risk_headers = ['风险等级', '定义', '影响范围', '紧急程度', '颜色标识']
    risk_rows = [
        ['红色 - 紧急高危', '已造成实质性损害或即将爆发', '涉及多个部门或核心岗位', '需立即介入', '🔴'],
        ['橙色 - 高度关注', '信号明确且持续增强', '影响一个或多个团队', '本周内处理', '🟠'],
        ['黄色 - 中度预警', '存在明显信号但未扩大', '局部或个人层面', '两周内处理', '🟡'],
        ['绿色 - 轻度观察', '偶发信号或误解造成', '个别情况', '持续监测', '🟢'],
    ]
    create_table_with_header(doc, risk_headers, risk_rows, [1.5, 2.0, 1.5, 1.2, 0.8])

    # 优先级判定
    add_heading(doc, '三、优先级判定矩阵', level=1)

    priority_headers = ['组合类型', '情况描述', '优先级', '建议行动']
    priority_rows = [
        ['高频+高强度', '多次观察且反应激烈', 'P0', '立即介入，高层关注'],
        ['高频+低强度', '经常发生但反应温和', 'P1', '深入了解原因，重点关注'],
        ['低频+高强度', '偶尔发生但反应强烈', 'P2', '查明动机，防止升级'],
        ['低频+低强度', '偶发且反应平淡', 'P3', '持续观察，记录变化'],
    ]
    create_table_with_header(doc, priority_headers, priority_rows, [1.5, 2.5, 0.8, 2.2])

    # 建议响应方式
    add_heading(doc, '四、建议响应方式对照', level=1)

    response_headers = ['抵抗类型', '沟通策略', '管理策略', '关键成功因素']
    response_rows = [
        ['言语抵抗', '开放对话、澄清误解', '倾听理解、反馈渠道', '及时回应、真诚沟通'],
        ['行为抵抗', '明确规则、后果告知', '辅导支持、绩效管理', '一致性、公正执行'],
        ['集体抵抗', '代表对话、协商解决', '利益协调、渐进变革', '尊重参与、共享决策'],
        ['权力抵抗', '借力使力、联盟构建', '政治游戏、变革倡导', '耐心策略、长期规划'],
    ]
    create_table_with_header(doc, response_headers, response_rows, [1.2, 2.0, 2.0, 2.0])

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 抵抗信号分类矩阵 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '03-抵抗信号分类矩阵.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 文档4: 组织情绪监测日志
# ============================================================
def create_04_organization_mood_log():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('组织情绪监测日志', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '变革期间组织情绪跟踪与分析工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 日常监测记录表
    add_heading(doc, '一、日常监测记录表', level=1)

    add_instruction_box(doc, "使用说明：每日填写一次，重点记录当天的关键情绪事件和观察到的信号。")

    daily_headers = ['日期', '监测区域/部门', '整体情绪状态', '关键事件', '信号强度(1-5)', '记录人', '备注']
    daily_rows = [
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
        ['____年__月__日', '', '☐积极 ☐中性 ☐消极', '', '☐1 ☐2 ☐3 ☐4 ☐5', '', ''],
    ]
    create_table_with_header(doc, daily_headers, daily_rows, [0.8, 1.0, 0.8, 1.5, 0.7, 0.6, 0.8])

    # 周报模板
    add_heading(doc, '二、周报模板', level=1)

    week_headers = ['周次', '起止日期', '本周情绪趋势', '主要变化', '风险信号', '已采取措施', '效果评估', '下周重点']
    week_rows = [
        ['第__周', '__月__日-__月__日', '☐上升 ☐平稳 ☐下降', '', '', '', '', ''],
        ['第__周', '__月__日-__月__日', '☐上升 ☐平稳 ☐下降', '', '', '', '', ''],
        ['第__周', '__月__日-__月__日', '☐上升 ☐平稳 ☐下降', '', '', '', '', ''],
        ['第__周', '__月__日-__月__日', '☐上升 ☐平稳 ☐下降', '', '', '', '', ''],
    ]
    create_table_with_header(doc, week_headers, week_rows, [0.5, 0.9, 0.8, 1.2, 1.2, 1.2, 0.7, 0.9])

    # 月度分析模板
    add_heading(doc, '三、月度分析模板', level=1)

    month_headers = ['月份', '总体情绪指数', '同比变化', '环比变化', '主要风险点', '成功经验', '下月预测', '建议措施']
    month_rows = [
        ['____年__月', '/100', '☐上升 ☐持平 ☐下降', '☐上升 ☐持平 ☐下降', '', '', '', ''],
        ['____年__月', '/100', '☐上升 ☐持平 ☐下降', '☐上升 ☐持平 ☐下降', '', '', '', ''],
        ['____年__月', '/100', '☐上升 ☐持平 ☐下降', '☐上升 ☐持平 ☐下降', '', '', '', ''],
    ]
    create_table_with_header(doc, month_headers, month_rows, [0.7, 0.8, 0.8, 0.8, 1.3, 1.3, 0.8, 1.1])

    # 情绪指标说明
    add_heading(doc, '四、情绪指标说明', level=1)

    indicator_headers = ['指标名称', '计算方法', '阈值说明', '预警级别']
    indicator_rows = [
        ['情绪指数', '(积极样本数/总样本数)×100', '80+优秀/60-79正常/<60预警', '🔴<40红色预警'],
        ['变化率', '(本期-上期)/上期×100%', '±5%内正常/超出为异常', '⚠️变化>15%需关注'],
        ['信号强度', '1-5级评分均值', '1-2低/3中/4-5高', '均值>3.5需介入'],
        ['参与度', '主动参与人数/总人数', '70%+优秀/40-69%正常/<40%低', '低于30%危险'],
    ]
    create_table_with_header(doc, indicator_headers, indicator_rows, [1.0, 2.0, 1.8, 1.2])

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 组织情绪监测日志 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '04-组织情绪监测日志.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 文档5: 响应策略选择器
# ============================================================
def create_05_response_strategy_selector():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('响应策略选择器', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '变革抵抗响应策略匹配与实施工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_instruction_box(doc, "使用方法：根据抵抗信号的类型和强度，在左侧找到对应的策略类型，参考中间列的判断标准选择合适的响应策略。")

    # 不同抵抗类型的策略对照
    add_heading(doc, '一、抵抗类型与策略对照表', level=1)

    strategy_headers = ['抵抗类型', '表现特征', '策略选择', '具体措施', '不推荐做法']
    strategy_rows = [
        ['信息型抵抗', '因不了解而产生疑虑', '教育沟通策略', '及时透明沟通/培训解释/问答互动', '隐瞒信息/强制命令'],
        ['利益型抵抗', '担心个人利益受损', '谈判协商策略', '利益补偿/调整方案/参与决策', '忽视利益/强硬压制'],
        ['信任型抵抗', '对管理层缺乏信任', '关系修复策略', '高层出面/言行一致/兑现承诺', '空头承诺/推卸责任'],
        ['能力型抵抗', '担心无法适应变化', '支持辅导策略', '技能培训/过渡期支持/导师制', '放任自流/立即考核'],
        ['价值观型抵抗', '与个人理念冲突', '愿景引导策略', '阐述长远意义/找到共鸣点/尊重选择', '嘲笑讽刺/强行灌输'],
        ['惯性型抵抗', '习惯原有方式', '渐进适应策略', '分阶段推进/逐步替代/正向激励', '一步到位/惩罚威胁'],
    ]
    create_table_with_header(doc, strategy_headers, strategy_rows, [1.0, 1.3, 1.0, 2.5, 1.5])

    # 介入时机判断
    add_heading(doc, '二、介入时机判断标准', level=1)

    timing_headers = ['信号阶段', '识别特征', '最佳介入时机', '响应时限', '介入层级']
    timing_rows = [
        ['潜伏期', '个别负面言论/轻微抵触', '信号出现的第一时间', '48小时内', '直接主管'],
        ['显现期', '明确的不满表达/集体抱怨', '形成小范围共识前', '24小时内', '中层管理者'],
        ['爆发期', '公开反对/集体行动', '情绪升级前/事件发酵前', '即时介入', '高层管理者'],
        ['持续期', '长期消极抵制/士气低落', '成为组织文化前', '一周内', 'HR+高层'],
        ['消退期', '表面接受实则消极', '重新激活前', '两周内', '多方协作'],
    ]
    create_table_with_header(doc, timing_headers, timing_rows, [0.8, 1.8, 1.3, 0.9, 1.0])

    # 效果评估表
    add_heading(doc, '三、响应效果评估表', level=1)

    effect_headers = ['评估维度', '评估指标', '优秀(5分)', '良好(4分)', '一般(3分)', '较差(1-2分)', '权重']
    effect_rows = [
        ['即时效果', '抵抗行为是否停止', '完全停止', '明显减少', '部分减少', '未减少或加剧', '20%'],
        ['短期效果', '情绪是否改善', '明显积极', '有所改善', '基本持平', '更加消极', '25%'],
        ['中期效果', '态度是否转变', '完全接受', '理解接受', '勉强配合', '继续抵制', '30%'],
        ['长期效果', '行为是否持续', '持续积极', '基本保持', '有所波动', '回归原状', '25%'],
    ]
    create_table_with_header(doc, effect_headers, effect_rows, [0.8, 1.3, 1.2, 1.2, 1.2, 1.2, 0.6])

    # 综合评估
    add_heading(doc, '四、综合评估与调整', level=1)

    adjust_headers = ['综合得分', '效果等级', '策略调整建议', '是否升级处理']
    adjust_rows = [
        ['4.5-5.0', '非常成功', '总结经验，形成最佳实践', '否'],
        ['3.5-4.4', '比较成功', '保持现有策略，微调优化', '否'],
        ['2.5-3.4', '基本有效', '分析不足，调整具体措施', '考虑升级'],
        ['1.5-2.4', '效果不佳', '重新评估，策略转向', '建议升级'],
        ['1.0-1.4', '完全失败', '立即停止，全面重新规划', '必须升级'],
    ]
    create_table_with_header(doc, adjust_headers, adjust_rows, [1.0, 1.0, 3.0, 1.0])

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 响应策略选择器 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '05-响应策略选择器.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 文档6: 行动计划模板
# ============================================================
def create_06_action_plan_template():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # 标题
    title = doc.add_heading('行动计划模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = NAVY
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_para(doc, '变革抵抗应对行动计划制定工具', size=12, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # 目标设定表
    add_heading(doc, '一、目标设定表', level=1)

    add_instruction_box(doc, "SMART原则：目标应具体(Specific)、可衡量(Measurable)、可达成(Achievable)、相关性(Relevant)、时限性(Time-bound)")

    goal_headers = ['目标编号', '目标描述', '衡量指标', '目标值', '完成时间', '责任人', '相关方']
    goal_rows = [
        ['G-01', '', '', '', '____年__月__日', '', ''],
        ['G-02', '', '', '', '____年__月__日', '', ''],
        ['G-03', '', '', '', '____年__月__日', '', ''],
    ]
    create_table_with_header(doc, goal_headers, goal_rows, [0.6, 2.0, 1.5, 0.7, 0.8, 0.7, 0.7])

    # 行动项目清单
    add_heading(doc, '二、行动项目清单', level=1)

    action_headers = ['行动编号', '对应目标', '具体行动', '开始日期', '结束日期', '所需资源', '状态', '备注']
    action_rows = [
        ['A-01', 'G-01', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
        ['A-02', 'G-01', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
        ['A-03', 'G-02', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
        ['A-04', 'G-02', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
        ['A-05', 'G-03', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
        ['A-06', 'G-03', '', '__月__日', '__月__日', '', '☐待启动 ☐进行中 ☐已完成', ''],
    ]
    create_table_with_header(doc, action_headers, action_rows, [0.5, 0.6, 1.8, 0.6, 0.6, 1.0, 1.2, 0.8])

    # 时间线规划
    add_heading(doc, '三、时间线规划', level=1)

    timeline_headers = ['阶段', '时间范围', '主要任务', '里程碑', '交付成果']
    timeline_rows = [
        ['准备阶段', '__月__日至__月__日', '调研分析/方案制定', '', ''],
        ['启动阶段', '__月__日至__月__日', '宣贯沟通/资源到位', '', ''],
        ['实施阶段', '__月__日至__月__日', '措施落地/持续沟通', '', ''],
        ['评估阶段', '__月__日至__月__日', '效果评估/方案优化', '', ''],
        ['固化阶段', '__月__日至__月__日', '经验总结/标准化', '', ''],
    ]
    create_table_with_header(doc, timeline_headers, timeline_rows, [0.7, 1.0, 1.8, 1.0, 1.5])

    # 甘特图区域
    add_heading(doc, '四、甘特图', level=1)

    gantt_table = doc.add_table(rows=8, cols=13)
    gantt_table.style = 'Table Grid'

    # 设置表头
    header_row = gantt_table.rows[0]
    header_row.cells[0].text = '行动项'
    set_cell_shading(header_row.cells[0], '4472C4')
    for i in range(1, 13):
        header_row.cells[i].text = f'W{i}'
        set_cell_shading(header_row.cells[i], '4472C4')
        for para in header_row.cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = WHITE

    # 添加行动项行
    for i in range(1, 8):
        row = gantt_table.rows[i]
        row.cells[0].text = f'A-0{i}'
        set_cell_shading(row.cells[0], 'F2F2F2')
        for j in range(1, 13):
            row.cells[j].text = '  '
            row.cells[j].width = Cm(0.8)

    # 资源需求表
    add_heading(doc, '五、资源需求表', level=1)

    resource_headers = ['资源类型', '具体需求', '数量', '预算', '到位时间', '负责人', '备注']
    resource_rows = [
        ['人力资源', '', '', '____元', '__月__日', '', ''],
        ['财务资源', '', '', '____元', '__月__日', '', ''],
        ['技术资源', '', '', '____元', '__月__日', '', ''],
        ['时间资源', '', '', '-', '__月__日', '', ''],
        ['外部支持', '', '', '____元', '__月__日', '', ''],
    ]
    create_table_with_header(doc, resource_headers, resource_rows, [0.8, 1.5, 0.6, 0.7, 0.7, 0.6, 1.0])

    # 效果评估标准
    add_heading(doc, '六、效果评估标准', level=1)

    eval_headers = ['评估指标', '基线值', '目标值', '评估方法', '评估周期', '负责人']
    eval_rows = [
        ['抵抗信号减少比例', '', '', '观察统计', '每周', ''],
        ['员工满意度提升', '', '', '问卷调查', '每月', ''],
        ['关键行为改变率', '', '', '行为观察', '每两周', ''],
        ['沟通参与度', '', '', '数据统计', '每周', ''],
        ['目标达成率', '', '', '结果对比', '每月', ''],
    ]
    create_table_with_header(doc, eval_headers, eval_rows, [1.3, 0.7, 0.7, 0.9, 0.7, 0.7])

    # 风险与应对
    add_heading(doc, '七、风险识别与应对', level=1)

    risk_headers = ['风险描述', '发生可能性', '影响程度', '应对策略', '预警指标', '预案负责人']
    risk_rows = [
        ['', '☐高 ☐中 ☐低', '☐高 ☐中 ☐低', '', '', ''],
        ['', '☐高 ☐中 ☐低', '☐高 ☐中 ☐低', '', '', ''],
        ['', '☐高 ☐中 ☐低', '☐高 ☐中 ☐低', '', '', ''],
    ]
    create_table_with_header(doc, risk_headers, risk_rows, [1.5, 0.8, 0.8, 1.5, 1.0, 0.8])

    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '变革管理 - 行动计划模板 | 机密文件'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filepath = os.path.join(OUTPUT_DIR, '06-行动计划模板.docx')
    doc.save(filepath)
    print(f"已生成: {filepath}")
    return filepath

# ============================================================
# 主函数 - 生成所有文档
# ============================================================
if __name__ == '__main__':
    print("=" * 60)
    print("开始生成变革管理工具表单...")
    print("=" * 60)

    create_01_early_warning_checklist()
    create_02_key_person_interview()
    create_03_resistance_classification_matrix()
    create_04_organization_mood_log()
    create_05_response_strategy_selector()
    create_06_action_plan_template()

    print("=" * 60)
    print("所有工具表单已生成完成！")
    print(f"输出目录: {OUTPUT_DIR}")
    print("=" * 60)

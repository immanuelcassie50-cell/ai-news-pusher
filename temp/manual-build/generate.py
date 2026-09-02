"""
学员手册主生成器
"""
import sys
import os
sys.path.insert(0, r'D:\CC\temp\manual-build')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_helpers import (
    set_page_setup, add_page_numbers, add_running_header,
    COLOR_PRIMARY, COLOR_ACCENT, COLOR_TEXT, COLOR_MUTED,
    set_run_font, set_paragraph_spacing, add_p
)

# Import all part builders
from gen_part1 import build_cover, build_copyright, build_toc, build_part1
from gen_part2 import build_part2_opening, build_part3_part1
from gen_part3 import build_part4_part2
from gen_part4 import build_part5_part3
from gen_part5 import build_part6_checklist
from gen_part6 import build_part7_tools
from gen_part7 import build_part8_appendix
from gen_backcover import build_backcover


OUTPUT_PATH = r'D:\2026年课程\竞越\绩效管理和绩效面谈：通过绩效面谈让员工更加胜任\完整课程包\04_学员手册\学员手册_完整版.docx'


def main():
    print("=" * 60)
    print("开始生成学员手册...")
    print("=" * 60)

    # 1. 创建文档
    doc = Document()
    print("[1/3] 文档创建成功")

    # 2. 页面设置
    set_page_setup(doc)
    add_page_numbers(doc)
    add_running_header(doc, '对话驱动 · 学员手册')
    print("[2/3] 页面设置、页码、页眉已添加")

    # 3. 依次构建各部分
    print("开始构建各部分内容...")

    print("  - 构建封面")
    build_cover(doc)

    print("  - 构建使用说明")
    build_copyright(doc)

    print("  - 构建目录")
    build_toc(doc)

    print("  - 构建第一部分：课前准备")
    build_part1(doc)

    print("  - 构建第二/三部分：开场 + Part 1 工作空间")
    build_part2_opening(doc)
    build_part3_part1(doc)

    print("  - 构建第四部分：Part 2 艰难面谈")
    build_part4_part2(doc)

    print("  - 构建第五部分：Part 3 发展面谈")
    build_part5_part3(doc)

    print("  - 构建第六部分：核心交付物（4 大区）")
    build_part6_checklist(doc)

    print("  - 构建第七部分：工具索引")
    build_part7_tools(doc)

    print("  - 构建第八部分：附录")
    build_part8_appendix(doc)

    print("  - 构建封底")
    build_backcover(doc)

    print("[3/3] 所有部分构建完成")

    # 4. 保存
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)

    # 5. 统计
    file_size = os.path.getsize(OUTPUT_PATH)
    file_size_kb = file_size / 1024
    file_size_mb = file_size / (1024 * 1024)

    print()
    print("=" * 60)
    print("生成完成！")
    print("=" * 60)
    print(f"输出路径: {OUTPUT_PATH}")
    print(f"文件大小: {file_size:,} 字节 ({file_size_kb:.1f} KB / {file_size_mb:.2f} MB)")
    print()
    print("注意：Word 中显示的页数取决于字体、边距、实际渲染。")
    print("本脚本无法预知 Word 打开后的实际页数，请在 Word 中按 F5 检查。")

    return OUTPUT_PATH, file_size


if __name__ == '__main__':
    main()

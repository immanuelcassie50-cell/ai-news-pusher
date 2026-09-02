"""
学员手册生成器 - 共用工具与样式
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

# ===== 配色（活泼而专业）=====
COLOR_PRIMARY = "1F4E79"      # 深蓝
COLOR_ACCENT = "C00000"       # 重点红
COLOR_WARM = "ED7D31"         # 暖橙
COLOR_GREEN = "548235"        # 绿
COLOR_TEAL = "2E9D8E"         # 青
COLOR_PURPLE = "7030A0"       # 紫
COLOR_PINK = "C2185B"         # 粉

COLOR_LIGHT_BLUE = "DDEBF7"
COLOR_LIGHT_YELLOW = "FFF2CC"
COLOR_LIGHT_GREEN = "E2EFDA"
COLOR_LIGHT_PINK = "FCE4D6"
COLOR_LIGHT_GRAY = "F2F2F2"
COLOR_LIGHT_PURPLE = "F4E5F7"
COLOR_BORDER = "BFBFBF"
COLOR_TEXT = "262626"
COLOR_MUTED = "595959"

CN_FONT = "Microsoft YaHei"
EN_FONT = "Calibri"


def set_run_font(run, size=10.5, bold=False, italic=False, color=None, cn_font=CN_FONT, en_font=EN_FONT):
    """设置 run 字体"""
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), en_font)
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(p, before=0, after=0, line=None, line_rule='auto'):
    """设置段落间距（磅）"""
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        if line_rule == 'exact':
            pf.line_spacing = Pt(line)
        else:
            pf.line_spacing = line  # 倍数


def set_paragraph_borders(p, color=COLOR_BORDER, size=6, sides=None):
    """设置段落边框"""
    if sides is None:
        sides = ['top', 'left', 'bottom', 'right']
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.appendChild(pBdr)
    for side in sides:
        b = pBdr.find(qn(f'w:{side}'))
        if b is None:
            b = OxmlElement(f'w:{side}')
            pBdr.appendChild(b)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), color)


def set_paragraph_shading(p, color):
    """设置段落底色"""
    pPr = p._element.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        pPr.appendChild(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)


def set_paragraph_indent(p, left=None, right=None, hanging=None, first_line=None):
    """设置段落缩进"""
    pf = p.paragraph_format
    if left is not None:
        pf.left_indent = Cm(left)
    if right is not None:
        pf.right_indent = Cm(right)
    if hanging is not None:
        pf.first_line_indent = -Cm(hanging)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def set_paragraph_keep(p, keep_next=True, keep_lines=True):
    """设置段后不拆页"""
    pPr = p._element.get_or_add_pPr()
    if keep_next:
        e = OxmlElement('w:keepNext')
        pPr.appendChild(e)
    if keep_lines:
        e = OxmlElement('w:keepLines')
        pPr.appendChild(e)


def add_p(doc_or_cell, text="", size=10.5, bold=False, italic=False, color=None,
          align=None, before=0, after=4, line=None, line_rule='auto',
          style=None, font=CN_FONT):
    """通用段落添加"""
    p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color, cn_font=font)
    set_paragraph_spacing(p, before=before, after=after, line=line, line_rule=line_rule)
    return p


def add_h1(doc, text):
    p = add_p(doc, text, size=22, bold=True, color=COLOR_PRIMARY,
              before=24, after=12, line=1.3, font=CN_FONT)
    set_paragraph_keep(p, True, True)
    return p


def add_h2(doc, text):
    p = add_p(doc, text, size=16, bold=True, color=COLOR_PRIMARY,
              before=18, after=9, line=1.3, font=CN_FONT)
    set_paragraph_keep(p, True, True)
    return p


def add_h3(doc, text):
    p = add_p(doc, text, size=13, bold=True, color=COLOR_ACCENT,
              before=12, after=6, line=1.3, font=CN_FONT)
    set_paragraph_keep(p, True, True)
    return p


def add_h4(doc, text):
    p = add_p(doc, text, size=11, bold=True, color=COLOR_TEXT,
              before=9, after=5, line=1.3, font=CN_FONT)
    set_paragraph_keep(p, True, True)
    return p


def add_pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._element.appendChild(br)
    set_paragraph_spacing(p, 0, 0)
    return p


def add_callout(doc, label, text, style='tip'):
    """添加提示框（tip/warn/idea/quote/ai/action）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 6, 6, 1.5)
    set_paragraph_borders(p, color=COLOR_BORDER, size=8)
    if style == 'tip':
        bg = COLOR_LIGHT_BLUE
        border = COLOR_PRIMARY
    elif style == 'warn':
        bg = COLOR_LIGHT_YELLOW
        border = COLOR_ACCENT
    elif style == 'idea':
        bg = COLOR_LIGHT_GREEN
        border = COLOR_GREEN
    elif style == 'quote':
        bg = COLOR_LIGHT_PINK
        border = COLOR_WARM
    elif style == 'ai':
        bg = COLOR_LIGHT_PURPLE
        border = COLOR_PURPLE
    elif style == 'action':
        bg = "D5E3F0"
        border = "2E5C8A"
    else:
        bg = COLOR_LIGHT_GRAY
        border = COLOR_MUTED
    set_paragraph_borders(p, color=border, size=8)
    set_paragraph_shading(p, bg)
    set_paragraph_indent(p, left=0.3, right=0.3)
    if label:
        r1 = p.add_run(f"【{label}】 ")
        set_run_font(r1, size=10.5, bold=True, color=COLOR_ACCENT)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5)
    return p


def add_quote(doc, text, author=None):
    """添加金句块"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 8, 8, 1.4)
    set_paragraph_indent(p, left=1.2, right=1.2)
    set_paragraph_borders(p, color=COLOR_WARM, size=24, sides=['left'])
    r1 = p.add_run("“" + text + "”")
    set_run_font(r1, size=12, italic=True, color=COLOR_WARM)
    if author:
        p.add_run().add_break()
        r2 = p.add_run("—— " + author)
        set_run_font(r2, size=9, color=COLOR_MUTED)
    return p


def add_write_area(doc, lines=4, hint=None, border_color=COLOR_BORDER):
    """添加手写区域（带边框的空白行）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 6, 6, 28, 'exact')
    set_paragraph_borders(p, color=border_color, size=6)
    set_paragraph_indent(p, left=0.3, right=0.3)
    for i in range(lines):
        if i > 0:
            p.add_run().add_break()
        p.add_run(" ")
    if hint:
        r = p.add_run(f"  （{hint}）")
        set_run_font(r, size=9, color=COLOR_MUTED, italic=True)
    return p


def add_checkbox_line(doc, label, hint=None):
    """添加带勾选框的条目"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 5, 5, 1.5)
    set_paragraph_indent(p, left=0.8, hanging=0.4)
    r1 = p.add_run("☐  ")
    set_run_font(r1, size=11)
    r2 = p.add_run(label)
    set_run_font(r2, size=10.5)
    if hint:
        r3 = p.add_run(f"  （{hint}）")
        set_run_font(r3, size=9, color=COLOR_MUTED, italic=True)
    return p


def add_chapter_header(doc, part_label, title, subtitle, color=None):
    """添加章节大标题（带底色）"""
    color = color or COLOR_PRIMARY
    p = doc.add_paragraph()
    set_paragraph_shading(p, color)
    set_paragraph_spacing(p, 18, 12, 1.4)
    set_paragraph_indent(p, left=0.6, right=0.6, hanging=None)
    set_paragraph_borders(p, color=color, size=4)
    r1 = p.add_run(f"{part_label}  {title}")
    set_run_font(r1, size=16, bold=True, color="FFFFFF")
    p.add_run().add_break()
    r2 = p.add_run(subtitle)
    set_run_font(r2, size=10, color="FFFFFF")
    set_paragraph_keep(p, True, True)
    return p


def add_section_banner(doc, text, color=None):
    """小节横幅"""
    color = color or COLOR_TEAL
    p = doc.add_paragraph()
    set_paragraph_shading(p, color)
    set_paragraph_spacing(p, 10, 6, 1.3)
    set_paragraph_indent(p, left=0.4, right=0.4)
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True, color="FFFFFF")
    return p


def shade_cell(cell, color):
    """设置单元格底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.appendChild(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)


def set_table_borders(table, color=COLOR_BORDER, size=6, inside_color=None, inside_size=4):
    inside_color = inside_color or color
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.appendChild(b)
    for side in ['insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(inside_size))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), inside_color)
        tblBorders.appendChild(b)
    tblPr.appendChild(tblBorders)


def set_cell_text(cell, text, size=10, bold=False, color=None, align=None, cn_font=CN_FONT):
    """设置单元格文字"""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, 3, 3, 1.3)
    if text is not None and text != "":
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, cn_font=cn_font)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def make_table(doc, headers, rows, col_widths=None, header_fill=None, header_color="FFFFFF",
               header_size=11, body_size=10, row_height=None, header_align=None):
    """创建标准表格"""
    n_cols = len(headers)
    if col_widths is None:
        total_w = 16.0  # A4 减去 2cm margin ≈ 17cm
        col_widths = [total_w / n_cols] * n_cols
    else:
        col_widths = [Cm(w) for w in col_widths]

    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = w

    # Header
    header_fill = header_fill or COLOR_PRIMARY
    header_align = header_align or WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = col_widths[i]
        shade_cell(cell, header_fill)
        set_cell_text(cell, h, size=header_size, bold=True, color=header_color, align=header_align)

    # Body
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        if row_height is not None:
            trPr = tr._tr.get_or_add_trPr()
            h_el = OxmlElement('w:trHeight')
            h_el.set(qn('w:val'), str(int(row_height * 20)))  # twips
            h_el.set(qn('w:hRule'), 'atLeast')
            trPr.appendChild(h_el)
        row_fill = "FFFFFF" if ri % 2 == 0 else COLOR_LIGHT_GRAY
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.width = col_widths[ci]
            shade_cell(cell, row_fill)
            set_cell_text(cell, val, size=body_size, align=WD_ALIGN_PARAGRAPH.LEFT)

    set_table_borders(table, color=COLOR_BORDER, size=6, inside_color="D9D9D9", inside_size=4)
    set_paragraph_spacing(doc.paragraphs[-1], 6, 6)
    return table


def make_info_card(doc, label, content, color=None, label_w=4.0):
    """行式信息卡（标签 + 内容）"""
    color = color or COLOR_PRIMARY
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(label_w)
    table.columns[1].width = Cm(17 - label_w)
    cell1 = table.rows[0].cells[0]
    cell1.width = Cm(label_w)
    shade_cell(cell1, color)
    set_cell_text(cell1, label, size=11, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    cell2 = table.rows[0].cells[1]
    cell2.width = Cm(17 - label_w)
    set_cell_text(cell2, content, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_table_borders(table, color=color, size=8, inside_size=8)
    set_paragraph_spacing(doc.paragraphs[-1], 4, 4)
    return table


def make_dialogue_compare(doc, scenario, wrong, right, label_a="A · 容易掉进去的版本", label_b="B · 更有效的版本", reason=None):
    """话术对比表（场景 + A + B + 解释）"""
    headers = ["维度", label_a, label_b]
    rows = [
        ["场景", scenario, scenario],
        ["对话", wrong, right],
    ]
    table = make_table(doc, headers, rows, col_widths=[2.5, 7.0, 7.0], header_size=10, body_size=9.5, row_height=1.2)
    if reason:
        # 添加"为什么 B 更好"行
        row = table.add_row()
        tr = table.rows[-1]
        # 合并三列
        cell0 = tr.cells[0]
        cell1 = tr.cells[1]
        cell2 = tr.cells[2]
        merged = cell0.merge(cell1).merge(cell2)
        shade_cell(merged, COLOR_LIGHT_YELLOW)
        merged.text = ""
        p = merged.paragraphs[0]
        set_paragraph_spacing(p, 4, 4, 1.4)
        r1 = p.add_run("为什么 B 更好：")
        set_run_font(r1, size=10, bold=True, color=COLOR_ACCENT)
        p.add_run(" ")
        r2 = p.add_run(reason)
        set_run_font(r2, size=10)
        merged.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


def add_divider(doc, color=None):
    color = color or COLOR_PRIMARY
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 6, 6)
    set_paragraph_borders(p, color=color, size=12, sides=['bottom'])
    return p


def set_page_setup(doc):
    """页面设置 A4，适中边距"""
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)


def add_page_numbers(doc):
    """添加页码到页脚"""
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加 PAGE 域
    r = p.add_run()
    set_run_font(r, size=9, color=COLOR_MUTED)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE   \\* MERGEFORMAT"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r._element.appendChild(fldChar1)
    r._element.appendChild(instrText)
    r._element.appendChild(fldChar2)
    r._element.appendChild(fldChar3)


def add_running_header(doc, text):
    """页眉"""
    sec = doc.sections[0]
    header = sec.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size=9, color=COLOR_MUTED)
    set_paragraph_borders(p, color=COLOR_PRIMARY, size=4, sides=['bottom'])

"""
学员手册《对话驱动》生成器 - 第一部分
"""
import sys
sys.path.insert(0, r'D:\CC\temp\manual-build')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_helpers import *


def build_cover(doc):
    """封面页"""
    # 顶部色块
    p = doc.add_paragraph()
    set_paragraph_shading(p, COLOR_PRIMARY)
    set_paragraph_spacing(p, 80, 80)
    doc.add_paragraph()
    doc.add_paragraph()

    # 主标题
    add_p(doc, '对 话 驱 动', size=42, bold=True, color=COLOR_PRIMARY,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=12, line=1.3)
    add_p(doc, 'AI 时代的绩效面谈与能力发展', size=22, color=COLOR_ACCENT, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=24, line=1.3)
    add_p(doc, '— 学员手册 —', size=20, bold=True, color=COLOR_PRIMARY,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12, line=1.3)
    add_p(doc, 'Performance Dialogue in the AI Era · Participant Workbook',
          size=11, color=COLOR_MUTED, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=48, line=1.3)

    add_p(doc, '')
    add_p(doc, '')

    make_info_card(doc, '姓    名', '________________________________________________')
    add_p(doc, '', before=4, after=4)
    make_info_card(doc, '部    门', '________________________________________________')
    add_p(doc, '', before=4, after=4)
    make_info_card(doc, '课程日期', '________________________________________________')
    add_p(doc, '', before=4, after=4)
    make_info_card(doc, '今天的分析对象', '员工代号：________________________（用代号，不用实名）', color=COLOR_WARM)
    add_p(doc, '', before=4, after=4)
    make_info_card(doc, '我的问责伙伴', '____________________________________________ （下午填写）', color=COLOR_GREEN)

    add_p(doc, '')
    add_p(doc, '')
    add_p(doc, '开发者：罗宏伟', size=10, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=4)
    add_p(doc, '课程版本：完整版 · 2026', size=9, color=COLOR_MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)

    add_pagebreak(doc)


def build_copyright(doc):
    """使用说明"""
    add_h1(doc, '使用说明')
    add_p(doc, '这本手册是全天课程的工作空间，不是讲义。内容框架在现场讲解，这里是你记录真实洞见、完成演练、规划下一步的空间。',
          size=11, before=4, after=8, line=1.5)
    add_quote(doc, '绩效面谈不是管理者给员工的"评分仪式"，而是员工一年里能感受到的"我被认真对待了"的少数几个时刻之一。',
              author='——本课程核心信念')
    add_p(doc, '', before=4, after=4)

    add_h3(doc, '这本手册怎么用')
    add_checkbox_line(doc, '在每个核心概念之后，停下来，写下你的版本（不要事后补写）')
    add_checkbox_line(doc, '配对演练时，按"我扮演的角色 + 最难做到 + 伙伴反馈"三栏填写')
    add_checkbox_line(doc, '下午的"下次面谈准备清单"——这是今天课程对你下周工作最直接的产出')
    add_checkbox_line(doc, '选一名具体的员工（不要泛泛而想），填写所有的发展分析')
    add_checkbox_line(doc, '找一个问责伙伴——30 天后会真问你的那个人')
    add_p(doc, '', before=4, after=4)

    add_h3(doc, '手册结构')
    add_p(doc, '第一部分 课前准备  ·  第二部分 开场模块  ·  第三部分 Part 1 工作空间  ·  第四部分 Part 2 工作空间  ·  第五部分 Part 3 工作空间  ·  第六部分 下次面谈准备清单（核心交付物）  ·  第七部分 工具索引  ·  第八部分 附录',
          size=10, color=COLOR_MUTED, before=4, after=8, line=1.5)
    add_p(doc, '', before=2, after=2)

    add_callout(doc, '提示',
                '学员手册是给你写的工作空间。讲师手册的内容不在这里——讲师手册的内容由讲师掌握。这里只有你需要写、你想写、你能写的地方。',
                style='tip')
    add_pagebreak(doc)


def build_toc(doc):
    """目录"""
    add_h1(doc, '目录')
    add_p(doc, '', before=2, after=2)

    items = [
        ('第一部分  课前准备', 'P03', True),
        ('  1. 进入课堂前准备的两件事', 'P03', False),
        ('  2. 你带来的一个困境', 'P03', False),
        ('  3. 你带来的一个困境（详细记录）', 'P04', False),
        ('  4. 课前一周倒计时（每天的任务）', 'P05', False),
        ('  5. 行前清单', 'P06', False),
        ('', '', False),
        ('第二部分  开场模块：八个失效场景', 'P07', True),
        ('  1. 八个失效场景识别记录', 'P07', False),
        ('  2. 真实背景记录', 'P10', False),
        ('  3. AI 时代三个新场景出现频率', 'P11', False),
        ('', '', False),
        ('第三部分  Part 1：面谈的价值与四步面谈法', 'P12', True),
        ('  1. 面谈的真实价值（我的笔记）', 'P12', False),
        ('  2. 五个前提条件', 'P13', False),
        ('  3. 四步面谈法详解', 'P16', False),
        ('  4. AI 时代关键：探寻归因', 'P20', False),
        ('  5. 三个缺口层次 + AI 时代新增的第四层', 'P21', False),
        ('  6. 配对演练记录（3 轮）', 'P22', False),
        ('  7. 全班复盘洞见', 'P25', False),
        ('', '', False),
        ('第四部分  Part 2：艰难面谈——说真话的技术', 'P27', True),
        ('  1. 艰难面谈的三类成因', 'P27', False),
        ('  2. 我最常遇到的成因', 'P28', False),
        ('  3. 四原则详解（正面·全面·情面·事面）', 'P29', False),
        ('  4. AI 时代五类艰难场景速查', 'P38', False),
        ('  5. 场景讨论活动记录', 'P42', False),
        ('  6. 角色扮演记录（2 轮）', 'P43', False),
        ('  7. AI 时代品行问题笔记', 'P45', False),
        ('', '', False),
        ('第五部分  Part 3：发展面谈与双轨成长', 'P47', True),
        ('  1. 发展面谈 vs 评估面谈（表格）', 'P47', False),
        ('  2. 我在发展面谈上的误区', 'P48', False),
        ('  3. 双轨胜任度框架详解', 'P49', False),
        ('  4. 我的分析对象的双轨现状', 'P51', False),
        ('  5. 从目标缺口到发展需要（分析练习）', 'P52', False),
        ('  6. 发展对话三个启动问题', 'P55', False),
        ('  7. 配对演练记录（发展对话开场）', 'P58', False),
        ('', '', False),
        ('第六部分  下次面谈准备清单（核心交付物）', 'P60', True),
        ('  第一区：事实与归因准备', 'P61', False),
        ('  第二区：预估难点与准备', 'P63', False),
        ('  第三区：发展对话规划', 'P65', False),
        ('  第四区：四步面谈预演（关键词版）', 'P67', False),
        ('  问责伙伴', 'P69', False),
        ('', '', False),
        ('第七部分  工具索引', 'P70', True),
        ('', '', False),
        ('第八部分  附录', 'P72', True),
        ('  附录 A：面谈前后关键动作清单', 'P72', False),
        ('  附录 B：话术对比速查', 'P74', False),
        ('  附录 C：自我反思问题（30/60/90 天）', 'P77', False),
    ]

    for label, page, is_chapter in items:
        if not label:
            add_p(doc, '', before=2, after=2)
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 4, 4, 1.4)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '9000')
        tabs.appendChild(tab)
        pPr.appendChild(tabs)
        r = p.add_run(label)
        set_run_font(r, size=(12 if is_chapter else 10), bold=is_chapter,
                      color=(COLOR_PRIMARY if is_chapter else COLOR_TEXT))
        p.add_run().add_tab()
        r2 = p.add_run(page)
        set_run_font(r2, size=(12 if is_chapter else 10), bold=is_chapter,
                      color=(COLOR_PRIMARY if is_chapter else COLOR_MUTED))

    add_pagebreak(doc)


def build_part1(doc):
    """第一部分：课前准备"""
    add_chapter_header(doc, '第一部分', '课前准备', 'Pre-Work · 进入课堂前，让今天更有收获')
    add_p(doc, '', before=4, after=4)

    add_h2(doc, '1. 进入课堂前准备的两件事')
    add_p(doc, '进入课堂前，请在脑子里准备两件事。', size=11, before=2, after=4, line=1.5)
    add_p(doc, '', before=2, after=2)

    add_h3(doc, '1) 你的分析对象')
    add_p(doc, '一名你即将（或应该）进行绩效面谈的员工，或一名你在绩效面谈上感到困难的员工。', size=10.5, line=1.5, after=4)
    add_callout(doc, '注意', '用代号，不用实名。保护员工的隐私，也让你能更自由地写下真实情况。', style='warn')
    add_p(doc, '', before=2, after=2)
    add_p(doc, '我的分析对象是（员工代号）：__________________________', size=11, bold=True, color=COLOR_PRIMARY, after=2)
    add_p(doc, '这个人的职位：__________________________________________', size=10.5, after=2)
    add_p(doc, '我和他/她合作的时长：__________________________________', size=10.5, after=2)
    add_p(doc, '他/她这个周期里，我最想谈的一件事是：', size=10.5, after=2)
    add_write_area(doc, lines=2)
    add_p(doc, '', before=2, after=2)

    add_h3(doc, '2) 你带来的一个困境')
    add_p(doc, '在绩效面谈上，你最真实的困难是什么？用一句话写下来：', size=11, before=2, after=2, line=1.5)
    add_write_area(doc, lines=3, hint='用一句话写下来——不要写"还没想好"，先写再说')
    add_p(doc, '', before=2, after=2)
    add_callout(doc, '提示',
                '这一天我们最想帮到的，就是带着真实困境来的人。你的困境越具体，你能从今天拿走的东西越多。',
                style='tip')
    add_pagebreak(doc)

    add_h2(doc, '2. 你带来的一个困境（详细记录）')
    add_p(doc, '展开写——把场景、当事人、你最卡的地方、你的担心，都写下来。',
          size=10.5, before=2, after=4, line=1.5, italic=True, color=COLOR_MUTED)
    add_p(doc, '', before=2, after=2)

    add_h4(doc, '场景背景')
    add_write_area(doc, lines=4, hint='什么时候发生的？什么情境？')
    add_p(doc, '', before=2, after=2)

    add_h4(doc, '你当时怎么做的')
    add_write_area(doc, lines=4, hint='你具体说了什么、做了什么？')
    add_p(doc, '', before=2, after=2)

    add_h4(doc, '你的反应和感受')
    add_write_area(doc, lines=4, hint='你当时心里在想什么？现在回看有什么感受？')
    add_p(doc, '', before=2, after=2)

    add_h4(doc, '你卡住的地方')
    add_write_area(doc, lines=4, hint='具体是什么让你难以推进？')
    add_p(doc, '', before=2, after=2)

    add_h4(doc, '你希望这次课能帮你解决什么')
    add_write_area(doc, lines=4, hint='一句话或一段话——你离开教室时最想带走的一个收获')
    add_pagebreak(doc)

    add_h2(doc, '3. 课前一周倒计时（每天的任务）')
    add_p(doc, '如果你现在距离一场面谈还有一周时间，下面这个倒计时表能帮你准备好自己。', size=10.5, before=2, after=2, line=1.5)
    add_p(doc, '每完成一项打勾——打勾本身就是承诺。', size=10, color=COLOR_MUTED, italic=True, after=4)
    add_p(doc, '', before=2, after=2)

    headers = ['倒计时', '任务', '完成打勾', '我的备注']
    rows = [
        ['D-7', '通知员工面谈时间和大致方向（让他/她有时间准备）', '☐', ''],
        ['D-6', '回顾这名员工这个周期里 3-5 个关键绩效事实（具体事件，不是印象）', '☐', ''],
        ['D-5', '对每个事实做归因预判（哪些归因清晰，哪些需要面谈里共同探索）', '☐', ''],
        ['D-4', '识别主要绩效缺口（技能/行为/认知/AI 工具协作）', '☐', ''],
        ['D-3', '预判可能出现的面谈困难，准备应对方案（参考四原则和五类场景）', '☐', ''],
        ['D-2', '用双轨框架初步分析这名员工的发展方向', '☐', ''],
        ['D-1', '完成"下次面谈准备清单"四个区', '☐', ''],
        ['D-Day', '找问责伙伴配对，约定 30 天后的回顾时间', '☐', ''],
    ]
    make_table(doc, headers, rows, col_widths=[1.8, 8.5, 1.5, 5.0], body_size=9.5, row_height=0.7)
    add_p(doc, '', before=2, after=2)
    add_callout(doc, '提示',
                '面谈前 1-2 周告知员工是面谈能"做对"的前提条件之一。突然袭击式的面谈，员工进入的是防御模式，不是反思模式。',
                style='tip')
    add_pagebreak(doc)

    add_h2(doc, '4. 行前清单')
    add_p(doc, '进入教室前，最后做一次自检。', size=11, before=2, after=4, line=1.5)

    add_h3(doc, '我准备好了吗？')
    add_checkbox_line(doc, '我已经选好了一个具体的分析对象（员工代号）')
    add_checkbox_line(doc, '我已经写下了一个真实的困境')
    add_checkbox_line(doc, '我能在心里讲出这个员工这个周期里 3 个具体的事实')
    add_checkbox_line(doc, '我对这个员工的归因做过初步判断')
    add_checkbox_line(doc, '我知道这场课之后我准备面谈的是哪一位员工')
    add_p(doc, '', before=2, after=2)

    add_h3(doc, '我的心态')
    add_checkbox_line(doc, '我愿意承认我在绩效面谈上有些地方做得不够好')
    add_checkbox_line(doc, '我愿意用一整天的时间认真学这件事')
    add_checkbox_line(doc, '我愿意让我的真实工作场景进入学习，而不是用抽象案例')
    add_checkbox_line(doc, '我愿意找一个问责伙伴（30 天后会真问我的那个人）')
    add_p(doc, '', before=2, after=2)

    add_h3(doc, '今天结束时的承诺')
    add_p(doc, '我希望离开教室时，手上能有（用一句话或一段话写）：', size=10.5, before=2, after=2, line=1.5)
    add_write_area(doc, lines=4, hint='具体到下周能开始的一个动作，不是"我会做得更好"')

    add_pagebreak(doc)

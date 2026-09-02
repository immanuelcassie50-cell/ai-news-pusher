# -*- coding: utf-8 -*-
"""
AI时代决策工作手册 - 工具表单生成脚本
生成9个专业表单文件
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from datetime import datetime

# 输出目录
OUTPUT_DIR = r"D:\新课开发\工作手册\AI时代决策工作手册\完整课程包\06-工具表单"

# 红色强调色
RED_COLOR = RGBColor(0xC0, 0x00, 0x00)
DARK_COLOR = RGBColor(0x00, 0x00, 0x00)

def set_cell_style(cell, text, bold=False, red=False, align_center=False, font_size=10):
    """设置单元格样式"""
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if red:
        run.font.color.rgb = RED_COLOR
    else:
        run.font.color.rgb = DARK_COLOR

def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement('w:' + border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def add_title_paragraph(doc, text, level=1):
    """添加标题段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(12)
    run.font.color.rgb = DARK_COLOR
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_normal_paragraph(doc, text, bold=False, red=False):
    """添加普通段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    if red:
        run.font.color.rgb = RED_COLOR
    else:
        run.font.color.rgb = DARK_COLOR
    return p

def create_form_header(doc, form_name):
    """创建表单通用头部"""
    add_title_paragraph(doc, form_name)
    doc.add_paragraph()

# ==================== 表单1: 决策识别评估表 ====================
def create_01_decision_evaluation_form():
    """决策识别评估表"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "决策识别评估表")

    # 基本信息表
    table = doc.add_table(rows=3, cols=4)
    set_table_border(table)

    set_cell_style(table.cell(0, 0), "评估人", bold=True, red=True)
    set_cell_style(table.cell(0, 1), "")
    set_cell_style(table.cell(0, 2), "评估日期", bold=True, red=True)
    set_cell_style(table.cell(0, 3), "")

    set_cell_style(table.cell(1, 0), "决策名称", bold=True, red=True)
    set_cell_style(table.cell(1, 1), "", align_center=True)
    set_cell_style(table.cell(1, 2), "决策场景", bold=True, red=True)
    set_cell_style(table.cell(1, 3), "")

    set_cell_style(table.cell(2, 0), "备注", bold=True, red=True)
    # 合并单元格
    table.cell(2, 1).merge(table.cell(2, 3))
    set_cell_style(table.cell(2, 1), "")

    doc.add_paragraph()

    # 评估说明
    p = doc.add_paragraph()
    run = p.add_run("【评估说明】")
    run.font.bold = True
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run("请对以下三个核心问题进行评估，每个问题有三个选项。综合判断结果确定该决策是否值得做成决策卡。")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # 问题1: 多选项评估
    p = doc.add_paragraph()
    run = p.add_run("问题一：当时是否存在至少两个看起来都合理的选项？")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED_COLOR

    table1 = doc.add_table(rows=4, cols=2)
    set_table_border(table1)
    set_cell_style(table1.cell(0, 0), "选项", bold=True)
    set_cell_style(table1.cell(0, 1), "评估标准", bold=True)
    set_cell_style(table1.cell(1, 0), "A. 是的，有多个合理选项", bold=True)
    set_cell_style(table1.cell(1, 1), "存在真正的选择空间，需要权衡判断")
    set_cell_style(table1.cell(2, 0), "B. 不确定", bold=True)
    set_cell_style(table1.cell(2, 1), "看起来只有一个方向，但可能有隐藏选项")
    set_cell_style(table1.cell(3, 0), "C. 不是", bold=True)
    set_cell_style(table1.cell(3, 1), "明摆着的路，不需要判断，是执行")

    doc.add_paragraph()

    # 问题2: 信号拆解
    p = doc.add_paragraph()
    run = p.add_run("问题二：决策者能否清楚拆解出当时影响判断的具体信号？")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED_COLOR

    table2 = doc.add_table(rows=4, cols=2)
    set_table_border(table2)
    set_cell_style(table2.cell(0, 0), "选项", bold=True)
    set_cell_style(table2.cell(0, 1), "评估标准", bold=True)
    set_cell_style(table2.cell(1, 0), "A. 能清楚拆解", bold=True)
    set_cell_style(table2.cell(1, 1), "有可以被结构化的判断依据")
    set_cell_style(table2.cell(2, 0), "B. 部分能说清", bold=True)
    set_cell_style(table2.cell(2, 1), "有直觉反应，但难以语言化，需要深度挖掘")
    set_cell_style(table2.cell(3, 0), "C. 说不清楚", bold=True)
    set_cell_style(table2.cell(3, 1), "可能是直觉主导，或判断依据未被意识化")

    doc.add_paragraph()

    # 问题3: 情境复现
    p = doc.add_paragraph()
    run = p.add_run("问题三：这个决策的情境，未来有没有可能被别人重新遇到？")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED_COLOR

    table3 = doc.add_table(rows=4, cols=2)
    set_table_border(table3)
    set_cell_style(table3.cell(0, 0), "选项", bold=True)
    set_cell_style(table3.cell(0, 1), "评估标准", bold=True)
    set_cell_style(table3.cell(1, 0), "A. 很可能会复现", bold=True)
    set_cell_style(table3.cell(1, 1), "属于常见决策情境，有复用价值")
    set_cell_style(table3.cell(2, 0), "B. 有可能", bold=True)
    set_cell_style(table3.cell(2, 1), "有一定特殊性，但底层逻辑可能有参考价值")
    set_cell_style(table3.cell(3, 0), "C. 几乎不可能", bold=True)
    set_cell_style(table3.cell(3, 1), "极端个案，写叙事性记录留档即可，性价比低")

    doc.add_paragraph()

    # 综合判断结果
    p = doc.add_paragraph()
    run = p.add_run("【综合判断结果】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table4 = doc.add_table(rows=4, cols=2)
    set_table_border(table4)
    set_cell_style(table4.cell(0, 0), "A. 值得做成决策卡", bold=True, red=True)
    set_cell_style(table4.cell(0, 1), "三个问题中A选项居多，且B选项不超过一个")
    set_cell_style(table4.cell(1, 0), "B. 不确定", bold=True)
    set_cell_style(table4.cell(1, 1), "存在模糊地带，建议先做叙事性记录，条件成熟时再转决策卡")
    set_cell_style(table4.cell(2, 0), "C. 不值得", bold=True)
    set_cell_style(table4.cell(2, 1), "判断难度低或情境不可复用，做决策卡性价比不高")
    set_cell_style(table4.cell(3, 0), "判断说明", bold=True)
    set_cell_style(table4.cell(3, 1), "")

    doc.add_paragraph()

    # 特别警示
    p = doc.add_paragraph()
    run = p.add_run("【特别警示】")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("结果好的决策不一定是好决策卡：警惕把运气当方法论。如果判断简单、只是外部条件凑巧，做成卡片可能误导他人。")
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    run = p.add_run("结果不好的决策也可能是好决策卡：结果差但判断有难度的决策，恰恰是价值最高的警示素材。")
    run.font.size = Pt(9)

    doc.save(OUTPUT_DIR + "\\01-决策识别评估表.docx")
    print("已生成: 01-决策识别评估表.docx")

# ==================== 表单2: 复盘访谈记录表 ====================
def create_02_interview_record_form():
    """复盘访谈记录表"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "复盘访谈记录表")

    # 访谈基本信息
    table = doc.add_table(rows=4, cols=4)
    set_table_border(table)

    set_cell_style(table.cell(0, 0), "决策名称", bold=True, red=True)
    table.cell(0, 1).merge(table.cell(0, 3))
    set_cell_style(table.cell(0, 1), "")

    set_cell_style(table.cell(1, 0), "决策者", bold=True, red=True)
    set_cell_style(table.cell(1, 1), "")
    set_cell_style(table.cell(1, 2), "访谈日期", bold=True, red=True)
    set_cell_style(table.cell(1, 3), "")

    set_cell_style(table.cell(2, 0), "访谈者", bold=True, red=True)
    set_cell_style(table.cell(2, 1), "")
    set_cell_style(table.cell(2, 2), "访谈时长", bold=True, red=True)
    set_cell_style(table.cell(2, 3), "")

    set_cell_style(table.cell(3, 0), "决策背景", bold=True, red=True)
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_style(table.cell(3, 1), "")

    doc.add_paragraph()

    # 追因维度
    p = doc.add_paragraph()
    run = p.add_run("【追因维度】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("核心问题：当时是什么信号让你觉得需要做决策，而不是按原计划走？")
    run.font.italic = True
    run.font.size = Pt(10)

    table1 = doc.add_table(rows=5, cols=2)
    set_table_border(table1)
    set_cell_style(table1.cell(0, 0), "最原始的触发信号", bold=True)
    set_cell_style(table1.cell(0, 1), "")
    set_cell_style(table1.cell(1, 0), "信号出现的时间点", bold=True)
    set_cell_style(table1.cell(1, 1), "")
    set_cell_style(table1.cell(2, 0), "当时同时出现的背景因素", bold=True)
    set_cell_style(table1.cell(2, 1), "")
    set_cell_style(table1.cell(3, 0), "之前是否有类似信号（模式追问）", bold=True)
    set_cell_style(table1.cell(3, 1), "")
    set_cell_style(table1.cell(4, 0), "追问记录（在那之前呢？）", bold=True)
    set_cell_style(table1.cell(4, 1), "")

    doc.add_paragraph()

    # 权衡维度
    p = doc.add_paragraph()
    run = p.add_run("【权衡维度】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("核心问题：当时你能想到的选项有哪些，你是怎么排除掉其他选项的？")
    run.font.italic = True
    run.font.size = Pt(10)

    table2 = doc.add_table(rows=6, cols=2)
    set_table_border(table2)
    set_cell_style(table2.cell(0, 0), "被考虑过的所有选项（含最终没选的）", bold=True)
    set_cell_style(table2.cell(0, 1), "")
    set_cell_style(table2.cell(1, 0), "最终选择的选项", bold=True)
    set_cell_style(table2.cell(1, 1), "")
    set_cell_style(table2.cell(2, 0), "排除其他选项的具体理由", bold=True)
    set_cell_style(table2.cell(2, 1), "")
    set_cell_style(table2.cell(3, 0), "权衡时依赖的关键信息", bold=True)
    set_cell_style(table2.cell(3, 1), "")
    set_cell_style(table2.cell(4, 0), "追问：除了这个选项，有没有想过其他方案？", bold=True)
    set_cell_style(table2.cell(4, 1), "")
    set_cell_style(table2.cell(5, 0), "警惕事后合理化：当时算过还是事后回顾？", bold=True)
    set_cell_style(table2.cell(5, 1), "")

    doc.add_paragraph()

    # 未预见的假设
    p = doc.add_paragraph()
    run = p.add_run("【未预见的假设】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("核心问题：如果当时那个前提不成立，你还会做同样的判断吗？")
    run.font.italic = True
    run.font.size = Pt(10)

    table3 = doc.add_table(rows=3, cols=2)
    set_table_border(table3)
    set_cell_style(table3.cell(0, 0), "决策依赖的关键假设（无意识）", bold=True)
    set_cell_style(table3.cell(0, 1), "")
    set_cell_style(table3.cell(1, 0), "如果假设不成立，判断链条会如何松动", bold=True)
    set_cell_style(table3.cell(1, 1), "")
    set_cell_style(table3.cell(2, 0), "该假设是否被验证过？", bold=True)
    set_cell_style(table3.cell(2, 1), "")

    doc.add_paragraph()

    # 隐性判断挖掘
    p = doc.add_paragraph()
    run = p.add_run("【隐性判断挖掘】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("触发信号：决策者反复用模糊表达（如我觉得不对劲）的地方")
    run.font.italic = True
    run.font.size = Pt(10)

    table4 = doc.add_table(rows=4, cols=2)
    set_table_border(table4)
    set_cell_style(table4.cell(0, 0), "模式追问：上一次有类似感觉是什么时候？", bold=True)
    set_cell_style(table4.cell(0, 1), "")
    set_cell_style(table4.cell(1, 0), "反事实追问：如果信号没出现，还会做同样判断吗？", bold=True)
    set_cell_style(table4.cell(1, 1), "")
    set_cell_style(table4.cell(2, 0), "角色互换追问：会提醒新人注意什么？", bold=True)
    set_cell_style(table4.cell(2, 1), "")
    set_cell_style(table4.cell(3, 0), "提炼出的隐性规则", bold=True)
    set_cell_style(table4.cell(3, 1), "")

    doc.add_paragraph()

    # 访谈小结
    p = doc.add_paragraph()
    run = p.add_run("【访谈小结】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table5 = doc.add_table(rows=3, cols=2)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "最核心的判断依据（一句话总结）", bold=True)
    set_cell_style(table5.cell(0, 1), "")
    set_cell_style(table5.cell(1, 0), "是否值得做成决策卡？", bold=True)
    set_cell_style(table5.cell(1, 1), "A. 是   B. 不确定   C. 否")
    set_cell_style(table5.cell(2, 0), "后续跟进事项", bold=True)
    set_cell_style(table5.cell(2, 1), "")

    doc.save(OUTPUT_DIR + "\\02-复盘访谈记录表.docx")
    print("已生成: 02-复盘访谈记录表.docx")

# ==================== 表单3: 决策卡模板 ====================
def create_03_decision_card_template():
    """决策卡模板（可填写版）"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # 标题
    p = doc.add_paragraph()
    run = p.add_run("【决策卡名称】")
    run.font.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RED_COLOR
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 触发条件（最显眼位置）
    p = doc.add_paragraph()
    run = p.add_run("触发条件")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（条件句+动作指令，写在卡片最显眼位置）")
    run.font.size = Pt(9)
    run.font.italic = True

    table1 = doc.add_table(rows=2, cols=1)
    set_table_border(table1)
    set_cell_style(table1.cell(0, 0), "如果出现【具体可观测信号】，先做【具体核实动作】，再决定是否继续。", bold=True)
    set_cell_style(table1.cell(1, 0), "填写区：")

    doc.add_paragraph()

    # 检查表
    p = doc.add_paragraph()
    run = p.add_run("检查表")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（5-8条具体可执行的核对项，每条用能直接对照的动词开头）")
    run.font.size = Pt(9)
    run.font.italic = True

    table2 = doc.add_table(rows=8, cols=2)
    set_table_border(table2)
    for i in range(8):
        set_cell_style(table2.cell(i, 0), str(i+1) + ".", bold=True)
        set_cell_style(table2.cell(i, 1), "")

    doc.add_paragraph()

    # 应急方案
    p = doc.add_paragraph()
    run = p.add_run("应急方案")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（当触发条件满足且出现变体场景时的处理方式）")
    run.font.size = Pt(9)
    run.font.italic = True

    table3 = doc.add_table(rows=3, cols=2)
    set_table_border(table3)
    set_cell_style(table3.cell(0, 0), "如果【场景变体A】", bold=True)
    set_cell_style(table3.cell(0, 1), "则【处理方式】")
    set_cell_style(table3.cell(1, 0), "如果【场景变体B】", bold=True)
    set_cell_style(table3.cell(1, 1), "则【处理方式】")
    set_cell_style(table3.cell(2, 0), "如果【场景变体C】", bold=True)
    set_cell_style(table3.cell(2, 1), "则【处理方式】")

    doc.add_paragraph()

    # 适用场景
    p = doc.add_paragraph()
    run = p.add_run("适用场景")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table4 = doc.add_table(rows=2, cols=1)
    set_table_border(table4)
    set_cell_style(table4.cell(0, 0), "本卡适用：")
    set_cell_style(table4.cell(1, 0), "本卡不适用：")

    doc.add_paragraph()

    # 警示案例
    p = doc.add_paragraph()
    run = p.add_run("警示案例")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（本卡对应的失败岔路口描述，用于嵌入式警示对照）")
    run.font.size = Pt(9)
    run.font.italic = True

    table5 = doc.add_table(rows=2, cols=1)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "岔路口描述：")
    set_cell_style(table5.cell(1, 0), "本可注意到的信号：")

    doc.add_paragraph()

    # 认领人信息
    p = doc.add_paragraph()
    run = p.add_run("认领人信息")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table6 = doc.add_table(rows=3, cols=4)
    set_table_border(table6)
    set_cell_style(table6.cell(0, 0), "主认领人", bold=True)
    set_cell_style(table6.cell(0, 1), "")
    set_cell_style(table6.cell(0, 2), "备份认领人", bold=True)
    set_cell_style(table6.cell(0, 3), "")
    set_cell_style(table6.cell(1, 0), "联系方式", bold=True)
    set_cell_style(table6.cell(1, 1), "")
    set_cell_style(table6.cell(1, 2), "联系方式", bold=True)
    set_cell_style(table6.cell(1, 3), "")
    set_cell_style(table6.cell(2, 0), "最近更新日期", bold=True)
    set_cell_style(table6.cell(2, 1), "")
    set_cell_style(table6.cell(2, 2), "复审周期", bold=True)
    set_cell_style(table6.cell(2, 3), "")

    doc.add_paragraph()

    # 定位说明（固定内容）
    p = doc.add_paragraph()
    run = p.add_run("定位说明")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table7 = doc.add_table(rows=1, cols=1)
    set_table_border(table7)
    cell = table7.cell(0, 0)
    p = cell.paragraphs[0]
    run = p.add_run("本卡列出的触发条件是已知的高风险信号，不是判断的全部，任何时候你的直觉认为需要暂停，都应该优先于卡片。")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.save(OUTPUT_DIR + "\\03-决策卡模板.docx")
    print("已生成: 03-决策卡模板.docx")

# ==================== 表单4: 场景映射矩阵 ====================
def create_04_scenario_matrix():
    """场景映射矩阵模板（Excel）"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "场景映射矩阵"

    # 定义样式
    header_font = Font(bold=True, size=12, color="FFFFFF")
    header_fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    title_font = Font(bold=True, size=14)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # 标题行
    ws.merge_cells('A1:F1')
    ws['A1'] = '场景映射矩阵'
    ws['A1'].font = Font(bold=True, size=16)
    ws['A1'].alignment = Alignment(horizontal='center')

    # 说明行
    ws.merge_cells('A2:F2')
    ws['A2'] = '矩阵设计：找出真正影响判断的场景变量，给每个场景格子填入具体版本（可独立打印使用）'
    ws['A2'].font = Font(size=9, italic=True)
    ws['A2'].alignment = Alignment(horizontal='center')

    # 维度说明区
    ws['A4'] = '维度变量说明'
    ws['A4'].font = Font(bold=True, size=11)

    ws['A5'] = '维度1（行）：'
    ws['B5'] = ''
    ws['A6'] = '维度2（列）：'
    ws['B6'] = ''

    # 设置列宽
    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 20
    ws.column_dimensions['F'].width = 20

    # 主矩阵区域 - 从第8行开始
    start_row = 8

    # 行表头（场景变量1）
    row_labels = ['场景A', '场景B', '场景C', '场景D', '场景E']
    col_labels = ['场景1', '场景2', '场景3', '场景4']

    # 表头行
    ws.cell(row=start_row, column=1, value='场景变量1 \\ 场景变量2')
    ws.cell(row=start_row, column=1).font = header_font
    ws.cell(row=start_row, column=1).fill = header_fill
    ws.cell(row=start_row, column=1).border = thin_border
    ws.cell(row=start_row, column=1).alignment = Alignment(horizontal='center')

    for i, label in enumerate(col_labels):
        cell = ws.cell(row=start_row, column=i+2, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = Alignment(horizontal='center')

    # 数据区域
    for row_idx, row_label in enumerate(row_labels):
        cell = ws.cell(row=start_row+row_idx+1, column=1, value=row_label)
        cell.font = Font(bold=True)
        cell.border = thin_border
        cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        cell.alignment = Alignment(horizontal='center')

        for col_idx in range(len(col_labels)):
            cell = ws.cell(row=start_row+row_idx+1, column=col_idx+2)
            cell.border = thin_border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # 矩阵内容填写说明
    note_row = start_row + len(row_labels) + 3
    ws.merge_cells(f'A{note_row}:F{note_row}')
    ws[f'A{note_row}'] = '【填写说明】每个格子对应一张独立的小卡片，包含：触发条件、检查表要点、应急方案'
    ws[f'A{note_row}'].font = Font(bold=True, size=10, color="C00000")

    note_row += 1
    ws.merge_cells(f'A{note_row}:F{note_row}')
    ws[f'A{note_row}'] = '使用者只需要找到自己对应的格子，打开对应的卡即可使用'
    ws[f'A{note_row}'].font = Font(size=9)

    note_row += 1
    ws.merge_cells(f'A{note_row}:F{note_row}')
    ws[f'A{note_row}'] = '矩阵本身是维护者的地图，不是给一线使用者用的操作工具'
    ws[f'A{note_row}'].font = Font(size=9)

    note_row += 1
    ws.merge_cells(f'A{note_row}:F{note_row}')
    ws[f'A{note_row}'] = '先做粗粒度版本投入使用后，根据反馈再针对性细分'
    ws[f'A{note_row}'].font = Font(size=9)

    wb.save(OUTPUT_DIR + "\\04-场景映射矩阵.xlsx")
    print("已生成: 04-场景映射矩阵.xlsx")

# ==================== 表单5: 决策行为稽核表 ====================
def create_05_audit_form():
    """决策行为稽核表"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "决策行为稽核表")

    # 基本信息
    table = doc.add_table(rows=4, cols=4)
    set_table_border(table)

    set_cell_style(table.cell(0, 0), "决策名称", bold=True, red=True)
    table.cell(0, 1).merge(table.cell(0, 3))
    set_cell_style(table.cell(0, 1), "")

    set_cell_style(table.cell(1, 0), "决策日期", bold=True, red=True)
    set_cell_style(table.cell(1, 1), "")
    set_cell_style(table.cell(1, 2), "决策者", bold=True, red=True)
    set_cell_style(table.cell(1, 3), "")

    set_cell_style(table.cell(2, 0), "稽核日期", bold=True, red=True)
    set_cell_style(table.cell(2, 1), "")
    set_cell_style(table.cell(2, 2), "稽核人", bold=True, red=True)
    set_cell_style(table.cell(2, 3), "")

    set_cell_style(table.cell(3, 0), "决策背景", bold=True, red=True)
    table.cell(3, 1).merge(table.cell(3, 3))
    set_cell_style(table.cell(3, 1), "")

    doc.add_paragraph()

    # 稽核说明
    p = doc.add_paragraph()
    run = p.add_run("【稽核说明】")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("稽核不是查错，是替组织记住它学过什么。问的是沉淀这个动作有没有发生，而不是这次决策对不对。")
    run.font.size = Pt(9)
    run.font.italic = True

    doc.add_paragraph()

    # 四个稽核问题
    questions = [
        ("问题一：沉淀检查", "这次决策过程中，有没有出现过让决策者感到当时不确定的关键节点？如果有，这个节点有没有被记录下来？"),
        ("问题二：使用检查", "如果这次决策对应的场景之前已经有决策卡覆盖，这次是否被使用了，使用后是否有反馈？"),
        ("问题三：价值识别", "如果这次决策不具备做成卡片的价值，是否已经启动了识别流程？"),
        ("问题四：反馈迭代", "如果这次决策暴露出现有卡片的漏洞或过时之处，是否已经反馈给认领人进行更新？")
    ]

    for title, question in questions:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RED_COLOR

        p = doc.add_paragraph()
        run = p.add_run(question)
        run.font.size = Pt(10)

        table_q = doc.add_table(rows=2, cols=2)
        set_table_border(table_q)
        set_cell_style(table_q.cell(0, 0), "A. 有/是", bold=True)
        set_cell_style(table_q.cell(0, 1), "B. 没有/否")
        set_cell_style(table_q.cell(1, 0), "详情/备注：", bold=True)
        set_cell_style(table_q.cell(1, 1), "")

        doc.add_paragraph()

    # 额外问题
    p = doc.add_paragraph()
    run = p.add_run("【额外问题】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("这次有没有基于自己的判断，偏离了卡片上的建议？为什么？")
    run.font.size = Pt(10)
    run.font.italic = True

    p = doc.add_paragraph()
    run = p.add_run("（用于检测是否过度依赖卡片、失去自主判断力。如果答案一直是没有，需要警惕稽核机制正在制造僵化）")
    run.font.size = Pt(9)

    table_extra = doc.add_table(rows=2, cols=1)
    set_table_border(table_extra)
    set_cell_style(table_extra.cell(0, 0), "A. 有偏离（说明：")
    set_cell_style(table_extra.cell(1, 0), "B. 没有偏离")

    doc.add_paragraph()

    # 稽核结论
    p = doc.add_paragraph()
    run = p.add_run("【稽核结论】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table_conclusion = doc.add_table(rows=4, cols=1)
    set_table_border(table_conclusion)
    set_cell_style(table_conclusion.cell(0, 0), "A. 经验已有效沉淀")
    set_cell_style(table_conclusion.cell(1, 0), "B. 需要跟进：")
    set_cell_style(table_conclusion.cell(2, 0), "C. 需要迭代卡片：")
    set_cell_style(table_conclusion.cell(3, 0), "D. 需要废止旧卡（原因）：")

    doc.add_paragraph()

    # 跟进事项
    p = doc.add_paragraph()
    run = p.add_run("【跟进事项】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table_follow = doc.add_table(rows=4, cols=2)
    set_table_border(table_follow)
    set_cell_style(table_follow.cell(0, 0), "待办事项", bold=True)
    set_cell_style(table_follow.cell(0, 1), "负责人/截止日期")
    set_cell_style(table_follow.cell(1, 0), "")
    set_cell_style(table_follow.cell(1, 1), "")
    set_cell_style(table_follow.cell(2, 0), "")
    set_cell_style(table_follow.cell(2, 1), "")
    set_cell_style(table_follow.cell(3, 0), "")
    set_cell_style(table_follow.cell(3, 1), "")

    doc.save(OUTPUT_DIR + "\\05-决策行为稽核表.docx")
    print("已生成: 05-决策行为稽核表.docx")

# ==================== 表单6: 决策卡认领确认书 ====================
def create_06_claim_confirmation():
    """决策卡认领确认书"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "决策卡认领确认书")

    p = doc.add_paragraph()
    run = p.add_run("【说明】决策卡做完不是终点，找到认领人、让它被持续使用才是目标。一份没有主人的工具，用一次就是最后一次。")
    run.font.size = Pt(9)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 决策卡基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、决策卡基本信息")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table1 = doc.add_table(rows=5, cols=4)
    set_table_border(table1)

    set_cell_style(table1.cell(0, 0), "卡片名称", bold=True, red=True)
    table1.cell(0, 1).merge(table1.cell(0, 3))
    set_cell_style(table1.cell(0, 1), "")

    set_cell_style(table1.cell(1, 0), "对应决策描述", bold=True, red=True)
    table1.cell(1, 1).merge(table1.cell(1, 3))
    set_cell_style(table1.cell(1, 1), "")

    set_cell_style(table1.cell(2, 0), "主要使用场景", bold=True, red=True)
    table1.cell(2, 1).merge(table1.cell(2, 3))
    set_cell_style(table1.cell(2, 1), "")

    set_cell_style(table1.cell(3, 0), "关联的其他卡片", bold=True, red=True)
    table1.cell(3, 1).merge(table1.cell(3, 3))
    set_cell_style(table1.cell(3, 1), "")

    set_cell_style(table1.cell(4, 0), "卡片版本/日期", bold=True, red=True)
    set_cell_style(table1.cell(4, 1), "")
    set_cell_style(table1.cell(4, 2), "复审周期", bold=True, red=True)
    set_cell_style(table1.cell(4, 3), "")

    doc.add_paragraph()

    # 认领人职责
    p = doc.add_paragraph()
    run = p.add_run("二、认领人职责说明")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("认领不是简单地指定一个部门或岗位，而是要具体到一个真实的人。认领人对这张卡负三件具体责任：")
    run.font.size = Pt(10)

    duties = [
        "1. 定期回访这张卡是否还在被使用",
        "2. 收集使用中反馈的问题，主动推动迭代更新",
        "3. 发现业务变化时，主动判断卡片内容是否需要调整"
    ]

    table2 = doc.add_table(rows=4, cols=1)
    set_table_border(table2)
    set_cell_style(table2.cell(0, 0), "认领人核心职责：", bold=True, red=True)
    for i, duty in enumerate(duties):
        set_cell_style(table2.cell(i+1, 0), duty)

    doc.add_paragraph()

    table3 = doc.add_table(rows=4, cols=2)
    set_table_border(table3)

    set_cell_style(table3.cell(0, 0), "主认领人", bold=True, red=True)
    set_cell_style(table3.cell(0, 1), "")
    set_cell_style(table3.cell(1, 0), "备份认领人", bold=True, red=True)
    set_cell_style(table3.cell(1, 1), "")
    set_cell_style(table3.cell(2, 0), "主责说明", bold=True, red=True)
    set_cell_style(table3.cell(2, 1), "负责日常维护、收集反馈、推动迭代")
    set_cell_style(table3.cell(3, 0), "备份说明", bold=True, red=True)
    set_cell_style(table3.cell(3, 1), "主认领人离岗时自然接手，避免两张卡片都没人管")

    doc.add_paragraph()

    # 认领人确认签字
    p = doc.add_paragraph()
    run = p.add_run("三、认领人确认签字")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table4 = doc.add_table(rows=4, cols=3)
    set_table_border(table4)

    set_cell_style(table4.cell(0, 0), "", bold=True)
    set_cell_style(table4.cell(0, 1), "主认领人", bold=True)
    set_cell_style(table4.cell(0, 2), "备份认领人", bold=True)

    set_cell_style(table4.cell(1, 0), "签字", bold=True, red=True)
    set_cell_style(table4.cell(1, 1), "")
    set_cell_style(table4.cell(1, 2), "")

    set_cell_style(table4.cell(2, 0), "日期", bold=True, red=True)
    set_cell_style(table4.cell(2, 1), "")
    set_cell_style(table4.cell(2, 2), "")

    set_cell_style(table4.cell(3, 0), "联系方式", bold=True, red=True)
    set_cell_style(table4.cell(3, 1), "")
    set_cell_style(table4.cell(3, 2), "")

    doc.add_paragraph()

    # 生效日期
    p = doc.add_paragraph()
    run = p.add_run("四、生效日期")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table5 = doc.add_table(rows=2, cols=2)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "生效日期", bold=True, red=True)
    set_cell_style(table5.cell(0, 1), "")
    set_cell_style(table5.cell(1, 0), "失效日期（如有）", bold=True, red=True)
    set_cell_style(table5.cell(1, 1), "")

    doc.add_paragraph()

    # 注意事项
    p = doc.add_paragraph()
    run = p.add_run("【注意事项】")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED_COLOR

    notes = [
        "复盘引导者通常不适合做认领人（他们往往在项目结束后就转向下一个项目）",
        "认领人应该是最贴近这张卡使用场景的岗位负责人",
        "明确写清楚主责和备份各自做什么，避免两人都以为对方在负责",
        "一张三年没改过的决策卡，大概率已经没人真的在用"
    ]

    for note in notes:
        p = doc.add_paragraph()
        run = p.add_run(note)
        run.font.size = Pt(9)

    doc.save(OUTPUT_DIR + "\\06-决策卡认领确认书.docx")
    print("已生成: 06-决策卡认领确认书.docx")

# ==================== 表单7: 失败案例警示清单 ====================
def create_07_failure_case_checklist():
    """失败案例警示清单模板"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "失败案例警示清单")

    p = doc.add_paragraph()
    run = p.add_run("【目的】讲失败案例的目的不是让人害怕，是让人对照——在自己身上找到那个岔路口。")
    run.font.size = Pt(9)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 基本信息
    table1 = doc.add_table(rows=4, cols=4)
    set_table_border(table1)

    set_cell_style(table1.cell(0, 0), "案例编号", bold=True, red=True)
    set_cell_style(table1.cell(0, 1), "")
    set_cell_style(table1.cell(0, 2), "录入日期", bold=True, red=True)
    set_cell_style(table1.cell(0, 3), "")

    set_cell_style(table1.cell(1, 0), "关联决策卡", bold=True, red=True)
    table1.cell(1, 1).merge(table1.cell(1, 3))
    set_cell_style(table1.cell(1, 1), "")

    set_cell_style(table1.cell(2, 0), "审核人", bold=True, red=True)
    set_cell_style(table1.cell(2, 1), "")
    set_cell_style(table1.cell(2, 2), "审核状态", bold=True, red=True)
    set_cell_style(table1.cell(2, 3), "待审核 / 已通过 / 已归档")

    set_cell_style(table1.cell(3, 0), "匿名处理", bold=True, red=True)
    set_cell_style(table1.cell(3, 1), "已匿名 / 无法匿名需特殊处理")
    set_cell_style(table1.cell(3, 2), "当事人知情", bold=True, red=True)
    set_cell_style(table1.cell(3, 3), "是 / 否")

    doc.add_paragraph()

    # 岔路口描述
    p = doc.add_paragraph()
    run = p.add_run("一、岔路口描述")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（用一两句话说清楚当时存在哪几个选项，决策者最后选了哪一个）")
    run.font.size = Pt(9)
    run.font.italic = True

    table2 = doc.add_table(rows=4, cols=1)
    set_table_border(table2)
    set_cell_style(table2.cell(0, 0), "当时的选项：", bold=True)
    set_cell_style(table2.cell(1, 0), "")
    set_cell_style(table2.cell(2, 0), "最终选择：", bold=True)
    set_cell_style(table2.cell(3, 0), "")

    doc.add_paragraph()

    # 本可注意到的信号
    p = doc.add_paragraph()
    run = p.add_run("二、本可注意到的信号")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（这是警示清单里最关键的部分——当时有哪个具体的、可观测的信号已经出现，只是没被当回事）")
    run.font.size = Pt(9)
    run.font.italic = True

    table3 = doc.add_table(rows=5, cols=2)
    set_table_border(table3)
    set_cell_style(table3.cell(0, 0), "信号1", bold=True)
    set_cell_style(table3.cell(0, 1), "")
    set_cell_style(table3.cell(1, 0), "信号2", bold=True)
    set_cell_style(table3.cell(1, 1), "")
    set_cell_style(table3.cell(2, 0), "信号3", bold=True)
    set_cell_style(table3.cell(2, 1), "")
    set_cell_style(table3.cell(3, 0), "为什么当时被忽略了？", bold=True)
    set_cell_style(table3.cell(3, 1), "")
    set_cell_style(table3.cell(4, 0), "是否有合理的外部解释？", bold=True)
    set_cell_style(table3.cell(4, 1), "")

    doc.add_paragraph()

    # 对照问题
    p = doc.add_paragraph()
    run = p.add_run("三、对照问题（把案例转化成一个自查问题）")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（把失败案例转化成一个自查问题，让使用者映射到自己当下正在做的决策）")
    run.font.size = Pt(9)
    run.font.italic = True

    table4 = doc.add_table(rows=2, cols=1)
    set_table_border(table4)
    set_cell_style(table4.cell(0, 0), "自查问题：如果你现在正在推进的项目里，也出现了类似的信号，你打算怎么处理？", bold=True)
    set_cell_style(table4.cell(1, 0), "")

    doc.add_paragraph()

    # 案例背景
    p = doc.add_paragraph()
    run = p.add_run("四、案例背景（匿名处理）")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（不点名、不复述当事人的岗位和具体身份细节，只保留岔路口和信号）")
    run.font.size = Pt(9)
    run.font.italic = True

    table5 = doc.add_table(rows=3, cols=1)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "时间背景（不涉及具体日期）：")
    set_cell_style(table5.cell(1, 0), "行业/业务场景（泛化描述）：")
    set_cell_style(table5.cell(2, 0), "关键背景因素（不影响匿名的必要信息）：")

    doc.add_paragraph()

    # 警示提醒
    p = doc.add_paragraph()
    run = p.add_run("【警示提醒】")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RED_COLOR

    reminders = [
        "损失特别大的极端案例，读者容易产生这种事不会发生在我身上的心理距离，反而降低对照效果",
        "真正有效的往往是损失中等、事后回看当时那个信号其实很明显的案例",
        "警示清单最好分散嵌入到对应的决策卡里，而不是单独成册",
        "嵌入式呈现时：不点名、不展开损失细节，只留岔路口和信号"
    ]

    for reminder in reminders:
        p = doc.add_paragraph()
        run = p.add_run(reminder)
        run.font.size = Pt(9)

    doc.save(OUTPUT_DIR + "\\07-失败案例警示清单.docx")
    print("已生成: 07-失败案例警示清单.docx")

# ==================== 表单8: 训练活动设计模板 ====================
def create_08_training_activity_template():
    """训练活动设计模板"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "训练活动设计模板")

    p = doc.add_paragraph()
    run = p.add_run("【设计理念】讲道理讲三小时，不如让他在模拟场景里做错一次决策。训练的目的不是讲透道理，是让人在场景里犯一次错。")
    run.font.size = Pt(9)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 活动基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、活动基本信息")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table1 = doc.add_table(rows=4, cols=4)
    set_table_border(table1)

    set_cell_style(table1.cell(0, 0), "活动名称", bold=True, red=True)
    table1.cell(0, 1).merge(table1.cell(0, 3))
    set_cell_style(table1.cell(0, 1), "")

    set_cell_style(table1.cell(1, 0), "对应决策卡", bold=True, red=True)
    table1.cell(1, 1).merge(table1.cell(1, 3))
    set_cell_style(table1.cell(1, 1), "")

    set_cell_style(table1.cell(2, 0), "活动时长", bold=True, red=True)
    set_cell_style(table1.cell(2, 1), "")
    set_cell_style(table1.cell(2, 2), "参与人数", bold=True, red=True)
    set_cell_style(table1.cell(2, 3), "")

    set_cell_style(table1.cell(3, 0), "活动类型", bold=True, red=True)
    set_cell_style(table1.cell(3, 1), "小组讨论 / 角色扮演 / 案例推演 / 其他")
    set_cell_style(table1.cell(3, 2), "所需物料", bold=True, red=True)
    set_cell_style(table1.cell(3, 3), "")

    doc.add_paragraph()

    # 三个环节设计
    p = doc.add_paragraph()
    run = p.add_run("二、三个环节设计")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    # 环节1: 情境代入
    p = doc.add_paragraph()
    run = p.add_run("环节一：情境代入（约30%时间）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("【关键原则：不要提前暴露这是在考验决策卡的使用】")
    run.font.size = Pt(9)
    run.font.italic = True

    table2 = doc.add_table(rows=5, cols=2)
    set_table_border(table2)
    set_cell_style(table2.cell(0, 0), "模拟场景描述", bold=True, red=True)
    set_cell_style(table2.cell(0, 1), "")
    set_cell_style(table2.cell(1, 0), "包含的干扰信息", bold=True, red=True)
    set_cell_style(table2.cell(1, 1), "")
    set_cell_style(table2.cell(2, 0), "时间压力设定", bold=True, red=True)
    set_cell_style(table2.cell(2, 1), "")
    set_cell_style(table2.cell(3, 0), "让学员做出的初始判断", bold=True, red=True)
    set_cell_style(table2.cell(3, 1), "")
    set_cell_style(table2.cell(4, 0), "引导话术（如何导入情境不暴露意图）", bold=True, red=True)
    set_cell_style(table2.cell(4, 1), "")

    doc.add_paragraph()

    # 环节2: 揭示与对照
    p = doc.add_paragraph()
    run = p.add_run("环节二：揭示与对照（约40%时间）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table3 = doc.add_table(rows=5, cols=2)
    set_table_border(table3)
    set_cell_style(table3.cell(0, 0), "揭示的决策卡内容", bold=True, red=True)
    set_cell_style(table3.cell(0, 1), "")
    set_cell_style(table3.cell(1, 0), "让学员对照的检查项", bold=True, red=True)
    set_cell_style(table3.cell(1, 1), "")
    set_cell_style(table3.cell(2, 0), "学员判断与卡片的差距点", bold=True, red=True)
    set_cell_style(table3.cell(2, 1), "")
    set_cell_style(table3.cell(3, 0), "引导话术（如何引发懊恼体验）", bold=True, red=True)
    set_cell_style(table3.cell(3, 1), "")
    set_cell_style(table3.cell(4, 0), "关键提问：如果早点打开这张卡？", bold=True, red=True)
    set_cell_style(table3.cell(4, 1), "")

    doc.add_paragraph()

    # 环节3: 二次决策
    p = doc.add_paragraph()
    run = p.add_run("环节三：二次决策（约30%时间）")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RED_COLOR

    table4 = doc.add_table(rows=5, cols=2)
    set_table_border(table4)
    set_cell_style(table4.cell(0, 0), "变体场景描述", bold=True, red=True)
    set_cell_style(table4.cell(0, 1), "")
    set_cell_style(table4.cell(1, 0), "与原场景的关键差异", bold=True, red=True)
    set_cell_style(table4.cell(1, 1), "")
    set_cell_style(table4.cell(2, 0), "提供的工具（手边有卡可查）", bold=True, red=True)
    set_cell_style(table4.cell(2, 1), "")
    set_cell_style(table4.cell(3, 0), "预期正向体验", bold=True, red=True)
    set_cell_style(table4.cell(3, 1), "")
    set_cell_style(table4.cell(4, 0), "引导话术（强化这次我提前发现了）", bold=True, red=True)
    set_cell_style(table4.cell(4, 1), "")

    doc.add_paragraph()

    # 延迟追踪
    p = doc.add_paragraph()
    run = p.add_run("三、延迟追踪设计")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（真正检验训练有没有作用的，是几个月后他们在真实决策现场会不会想起来用这张卡）")
    run.font.size = Pt(9)
    run.font.italic = True

    table5 = doc.add_table(rows=4, cols=2)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "追踪时间点", bold=True, red=True)
    set_cell_style(table5.cell(0, 1), "1个月 / 3个月 / 6个月")
    set_cell_style(table5.cell(1, 0), "追踪方式", bold=True, red=True)
    set_cell_style(table5.cell(1, 1), "")
    set_cell_style(table5.cell(2, 0), "追踪问题", bold=True, red=True)
    set_cell_style(table5.cell(2, 1), "遇到类似信号时，有没有想起用这张卡？")
    set_cell_style(table5.cell(3, 0), "预期结果", bold=True, red=True)
    set_cell_style(table5.cell(3, 1), "大于50%使用率 / 30-50% / 小于30%")

    doc.save(OUTPUT_DIR + "\\08-训练活动设计模板.docx")
    print("已生成: 08-训练活动设计模板.docx")

# ==================== 表单9: 课程评估问卷 ====================
def create_09_course_evaluation_form():
    """课程评估问卷"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    create_form_header(doc, "课程评估问卷")

    p = doc.add_paragraph()
    run = p.add_run("【填写说明】请根据您的真实感受选择相应选项。您的反馈将帮助我们持续改进课程质量。")
    run.font.size = Pt(9)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 学员基本信息
    p = doc.add_paragraph()
    run = p.add_run("一、学员基本信息")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table1 = doc.add_table(rows=4, cols=4)
    set_table_border(table1)

    set_cell_style(table1.cell(0, 0), "姓名", bold=True, red=True)
    set_cell_style(table1.cell(0, 1), "")
    set_cell_style(table1.cell(0, 2), "部门/岗位", bold=True, red=True)
    set_cell_style(table1.cell(0, 3), "")

    set_cell_style(table1.cell(1, 0), "课程名称", bold=True, red=True)
    table1.cell(1, 1).merge(table1.cell(1, 3))
    set_cell_style(table1.cell(1, 1), "AI时代决策工作手册")

    set_cell_style(table1.cell(2, 0), "培训日期", bold=True, red=True)
    set_cell_style(table1.cell(2, 1), "")
    set_cell_style(table1.cell(2, 2), "培训讲师", bold=True, red=True)
    set_cell_style(table1.cell(2, 3), "")

    set_cell_style(table1.cell(3, 0), "工作年限", bold=True, red=True)
    set_cell_style(table1.cell(3, 1), "小于3年 / 3-5年 / 5-10年 / 大于10年")
    set_cell_style(table1.cell(3, 2), "是否首次接触决策卡", bold=True, red=True)
    set_cell_style(table1.cell(3, 3), "是 / 否")

    doc.add_paragraph()

    # 课程内容评估（5分制）
    p = doc.add_paragraph()
    run = p.add_run("二、课程内容评估（5分制：5=非常满意，4=满意，3=一般，2=不满意，1=非常不满意）")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table2 = doc.add_table(rows=9, cols=6)
    set_table_border(table2)

    # 表头
    headers = ["评估项目", "5", "4", "3", "2", "1"]
    for i, h in enumerate(headers):
        set_cell_style(table2.cell(0, i), h, bold=True)

    items = [
        "课程内容的实用性和可操作性",
        "复盘与决策卡的区别讲解清晰度",
        "决策识别标准的可理解性",
        "触发条件设计的合理性",
        "检查表设计的可执行性",
        "场景映射矩阵的适用性",
        "失败案例警示的教育价值",
        "整体内容架构的逻辑性"
    ]

    for i, item in enumerate(items):
        set_cell_style(table2.cell(i+1, 0), item)
        for j in range(1, 6):
            set_cell_style(table2.cell(i+1, j), "O")

    doc.add_paragraph()

    # 教学方法评估
    p = doc.add_paragraph()
    run = p.add_run("三、教学方法评估（5分制）")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table3 = doc.add_table(rows=6, cols=6)
    set_table_border(table3)

    for i, h in enumerate(headers):
        set_cell_style(table3.cell(0, i), h, bold=True)

    methods = [
        "访谈技巧的讲解清晰度",
        "角色互换追问法的实用性",
        "训练活动设计的有效性",
        "案例选择的代表性",
        "课堂互动和参与感"
    ]

    for i, method in enumerate(methods):
        set_cell_style(table3.cell(i+1, 0), method)
        for j in range(1, 6):
            set_cell_style(table3.cell(i+1, j), "O")

    doc.add_paragraph()

    # 讲师表现评估
    p = doc.add_paragraph()
    run = p.add_run("四、讲师表现评估（5分制）")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table4 = doc.add_table(rows=5, cols=6)
    set_table_border(table4)

    for i, h in enumerate(headers):
        set_cell_style(table4.cell(0, i), h, bold=True)

    lecturer_items = [
        "专业知识和经验",
        "表达清晰度",
        "案例讲解的生动性",
        "回答问题的针对性"
    ]

    for i, item in enumerate(lecturer_items):
        set_cell_style(table4.cell(i+1, 0), item)
        for j in range(1, 6):
            set_cell_style(table4.cell(i+1, j), "O")

    doc.add_paragraph()

    # 总体满意度
    p = doc.add_paragraph()
    run = p.add_run("五、总体满意度")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table5 = doc.add_table(rows=3, cols=1)
    set_table_border(table5)
    set_cell_style(table5.cell(0, 0), "您对本次课程的总体满意度：非常满意 / 满意 / 一般 / 不满意 / 非常不满意")
    set_cell_style(table5.cell(1, 0), "您愿意向同事推荐这门课程吗？非常愿意 / 愿意 / 一般 / 不愿意")
    set_cell_style(table5.cell(2, 0), "您认为这门课程的最大价值在于：")

    doc.add_paragraph()

    # 改进建议
    p = doc.add_paragraph()
    run = p.add_run("六、改进建议")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    table6 = doc.add_table(rows=4, cols=1)
    set_table_border(table6)
    set_cell_style(table6.cell(0, 0), "内容方面最需要改进的地方：")
    set_cell_style(table6.cell(1, 0), "")
    set_cell_style(table6.cell(2, 0), "教学方法方面最需要改进的地方：")
    set_cell_style(table6.cell(3, 0), "")

    doc.add_paragraph()

    # 行动计划
    p = doc.add_paragraph()
    run = p.add_run("七、行动计划")
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RED_COLOR

    p = doc.add_paragraph()
    run = p.add_run("（请列出您回去后计划应用的一项具体实践）")
    run.font.size = Pt(9)
    run.font.italic = True

    table7 = doc.add_table(rows=4, cols=2)
    set_table_border(table7)
    set_cell_style(table7.cell(0, 0), "计划应用的决策卡/工具", bold=True, red=True)
    set_cell_style(table7.cell(0, 1), "")
    set_cell_style(table7.cell(1, 0), "计划应用到的工作场景", bold=True, red=True)
    set_cell_style(table7.cell(1, 1), "")
    set_cell_style(table7.cell(2, 0), "预期达成的效果", bold=True, red=True)
    set_cell_style(table7.cell(2, 1), "")
    set_cell_style(table7.cell(3, 0), "应用时间节点", bold=True, red=True)
    set_cell_style(table7.cell(3, 1), "")

    doc.add_paragraph()

    # 感谢语
    p = doc.add_paragraph()
    run = p.add_run("感谢您完成本次评估！您的反馈将帮助我们持续改进课程质量。")
    run.font.size = Pt(10)
    run.font.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT_DIR + "\\09-课程评估问卷.docx")
    print("已生成: 09-课程评估问卷.docx")

# ==================== 主函数 ====================
def main():
    print("=" * 60)
    print("开始生成【AI时代决策工作手册】工具表单...")
    print("=" * 60)

    create_01_decision_evaluation_form()
    create_02_interview_record_form()
    create_03_decision_card_template()
    create_04_scenario_matrix()
    create_05_audit_form()
    create_06_claim_confirmation()
    create_07_failure_case_checklist()
    create_08_training_activity_template()
    create_09_course_evaluation_form()

    print("=" * 60)
    print("所有表单生成完成！")
    print("输出目录：" + OUTPUT_DIR)
    print("=" * 60)

if __name__ == "__main__":
    main()

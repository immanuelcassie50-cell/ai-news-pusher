# -*- coding: utf-8 -*-
"""
德赛西威 AI 赋能课程评审全流程 - Word 批量生成（主控接管）
批次 2：评审实施 6 份（D-06 / D-07 / D-09 / D-10 / D-11 / D-13）
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === 配色 ===
COLOR_MAIN = "003D7A"
COLOR_AI = "00A0E9"
COLOR_WARN = "F37021"
COLOR_SAFE = "3CB878"
COLOR_RED = "D0021B"
COLOR_YELLOW = "F5A623"
COLOR_TEXT = "333333"
COLOR_BG = "F4F6F9"
COLOR_WHITE = "FFFFFF"

OUT_DIR_2 = "D:/Downloads/xinjian/德赛西威评审全流程PRD/产出物/02-评审实施"
os.makedirs(OUT_DIR_2, exist_ok=True)

PROJECT_NAME = "德赛西威 AI 赋能课程评审全流程"

# 5 档评分表头
SCORING_HEADERS = ["序号", "评估项", "评估要点", "很好 24-25", "较好 21-23", "一般 18-20", "较差 15-17", "差 0-14", "评委评分"]

# ==================== 工具函数 ====================

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tc_pr.append(tcBorders)

def set_run_font(run, name="思源黑体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text, size=10.5, bold=False, color=COLOR_TEXT, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_title_block(doc, title, subtitle=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLOR_MAIN)
    cell.width = Cm(16)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run(title)
    set_run_font(run1, size=18, bold=True, color=COLOR_WHITE)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=10, bold=False, color=COLOR_WHITE)

def add_section_header(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("■ " + text)
        set_run_font(run, size=13, bold=True, color=COLOR_MAIN)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("▶ " + text)
        set_run_font(run, size=11.5, bold=True, color=COLOR_MAIN)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("· " + text)
        set_run_font(run, size=10.5, bold=True, color=COLOR_AI)

def add_table_with_header(doc, headers, rows, col_widths=None, header_color=COLOR_MAIN):
    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=COLOR_WHITE)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, COLOR_BG)
            set_cell_borders(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val) if val is not None else "")
            set_run_font(run, size=9.5, color=COLOR_TEXT)

def setup_page(doc, file_id="D-XX"):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        header = section.header
        h_table = header.add_table(rows=1, cols=2, width=Cm(15))
        h_table.autofit = False
        h_left = h_table.rows[0].cells[0]
        h_right = h_table.rows[0].cells[1]
        h_left.text = ""
        h_right.text = ""
        p_l = h_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_l = p_l.add_run(PROJECT_NAME)
        set_run_font(run_l, size=8.5, color=COLOR_MAIN, bold=True)
        p_r = h_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_r = p_r.add_run(f"文件编号 {file_id}")
        set_run_font(run_r, size=8.5, color=COLOR_MAIN, bold=True)
        for cell in [h_left, h_right]:
            tc_pr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top','left','bottom','right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tc_pr.append(tcBorders)
        footer = section.footer
        p_f = footer.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_f.add_run("德赛西威 AI 赋能课程评审全流程 | 第 ")
        set_run_font(run_f, size=8, color=COLOR_TEXT)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_xml = run_f._element
        run_xml.append(fldChar1)
        run_xml.append(instrText)
        run_xml.append(fldChar2)
        run_g = p_f.add_run(" 页 / 共 ")
        set_run_font(run_g, size=8, color=COLOR_TEXT)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run_xml2 = run_g._element
        run_xml2.append(fldChar3)
        run_xml2.append(instrText2)
        run_xml2.append(fldChar4)
        run_h = p_f.add_run(" 页")
        set_run_font(run_h, size=8, color=COLOR_TEXT)

def add_signature_row(doc, label="评委签名：", date_label="日期："):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label}________________________     {date_label}________年____月____日")
    set_run_font(run, size=10.5, color=COLOR_TEXT)

def add_info_box(doc, text, color=COLOR_SAFE, icon="●"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon} {text}")
    set_run_font(run, size=10, bold=True, color=color)


# ==================== D-06 基础班·提示词模板评分卡 ====================

def make_D06():
    doc = Document()
    setup_page(doc, "D-06")
    add_title_block(doc, "基础班·提示词模板评分卡", "10 条评分项 × 5 档打分（24-25/21-23/18-20/15-17/0-14）")

    # 基本信息
    add_para(doc, "学员姓名：________________     学员岗位：________________     学员部门：________________", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "提示词模板名称：________________     提交日期：________年____月____日", size=10)
    add_para(doc, "评委姓名：________________     评委类型：□业务方 □AI 方法论 □大众评审", size=10)

    add_section_header(doc, "评分标准", level=1)
    add_para(doc, "· 24-25 分（很好）：超出预期，可作标杆，全维度无明显短板", size=10)
    add_para(doc, "· 21-23 分（较好）：达到预期，有亮点，1-2 项可提升", size=10)
    add_para(doc, "· 18-20 分（一般）：达到基本要求，3 项左右待优化", size=10)
    add_para(doc, "· 15-17 分（较差）：未达预期，需大幅改进", size=10)
    add_para(doc, "· 0-14 分（差）：严重不达标，不建议继续推进", size=10)

    add_section_header(doc, "评分表（10 条）", level=1)

    rows = [
        ["01", "业务场景还原度", "是否针对真实工作场景、痛点描述具体", "", "", "", "", "", "____"],
        ["02", "四段式结构完整性", "角色 / 背景 / 目标 / 约束四要素是否齐全", "", "", "", "", "", "____"],
        ["03", "角色定位精准度", "AI 角色是否专业、有边界", "", "", "", "", "", "____"],
        ["04", "约束条件合理性", "安全 / 格式 / 风格约束是否到位", "", "", "", "", "", "____"],
        ["05", "提示词可复用性", "换场景 / 换人能否用", "", "", "", "", "", "____"],
        ["06", "业务价值可衡量", "能否算出节省时间 / 提升质量", "", "", "", "", "", "____"],
        ["07", "提示词迭代次数", "至少测试过 3 次并优化", "", "", "", "", "", "____"],
        ["08", "同事复用情况", "是否被同岗位其他人用", "", "", "", "", "", "____"],
        ["09", "信息安全合规", "是否遵守红黄绿灯 + 脱敏", "", "", "", "", "", "____"],
        ["10", "与工具地图一致性", "提示词所用工具是否与个人工具地图匹配", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "总评分", level=1)
    add_para(doc, "10 项平均分 = _____ 分    对应档位：□很好 □较好 □一般 □较差 □差", size=11, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "评委评语", level=1)
    add_para(doc, "亮点（值得推广的部分）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "改进建议：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "____________________________________________________", size=10)

    add_info_box(doc, "重要：本表仅为「提示词模板」单点评分。学员整体得分 = 提示词 × 50% + 工具地图 × 10% + 场景化 × 20% + AI 陪跑 × 20%", color=COLOR_WARN)
    add_signature_row(doc, "评委签名：", "评分日期：")


# ==================== D-07 基础班·AI 工具地图评分表 ====================

def make_D07():
    doc = Document()
    setup_page(doc, "D-07")
    add_title_block(doc, "基础班·AI 工具地图评分表", "8 条评分项 × 5 档打分")

    add_para(doc, "学员姓名：________________     学员岗位：________________     学员部门：________________", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "工具地图提交日期：________年____月____日", size=10)
    add_para(doc, "评委姓名：________________     评委类型：□业务方 □AI 方法论 □大众评审", size=10)

    add_section_header(doc, "评分标准", level=1)
    add_para(doc, "· 24-25 分（很好）：工具组合覆盖度+定位讲解+迭代预判全部到位", size=10)
    add_para(doc, "· 21-23 分（较好）：覆盖度够，有 1-2 项工具理解可加深", size=10)
    add_para(doc, "· 18-20 分（一般）：基础工具已掌握，缺乏深度判断", size=10)
    add_para(doc, "· 15-17 分（较差）：工具选择不清晰", size=10)
    add_para(doc, "· 0-14 分（差）：工具地图明显缺漏", size=10)

    add_section_header(doc, "评分表（8 条）", level=1)
    rows = [
        ["01", "内部平台判断力", "是否清楚数智小西 / 内部系统的能力边界", "", "", "", "", "", "____"],
        ["02", "外部工具熟悉度", "是否熟练使用 3 种以上外部 AI 工具", "", "", "", "", "", "____"],
        ["03", "工具组合能力", "能否把多个工具组合解决复杂问题", "", "", "", "", "", "____"],
        ["04", "工具局限说明", "能否准确说出每个工具的局限", "", "", "", "", "", "____"],
        ["05", "数据安全分级", "能否按红黄绿灯选择工具", "", "", "", "", "", "____"],
        ["06", "工具迭代预判", "能否预判工具升级对方法的影响", "", "", "", "", "", "____"],
        ["07", "推广复用情况", "工具地图是否被同岗位其他人引用", "", "", "", "", "", "____"],
        ["08", "工具地图可视化", "地图呈现是否清晰易读（流程图 / 矩阵）", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "总评分", level=1)
    add_para(doc, "8 项平均分 = _____ 分    对应档位：□很好 □较好 □一般 □较差 □差", size=11, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "评委评语", level=1)
    add_para(doc, "亮点：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "改进建议：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)

    add_info_box(doc, "参考：建议每张工具地图至少包含「内部平台定位 + 3 种外部工具 + 1 个组合方案 + 红黄绿灯使用规则」", color=COLOR_SAFE)
    add_signature_row(doc, "评委签名：", "评分日期：")
    doc.save(os.path.join(OUT_DIR_2, "D-07-基础班-AI工具地图评分表（每人一份）.docx"))
    print("[OK] D-07")


# ==================== D-09 内训师班·AI 教练技能评估表 ====================

def make_D09():
    doc = Document()
    setup_page(doc, "D-09")
    add_title_block(doc, "内训师班·AI 教练技能评估表", "20 条 × 4 大类 × 5 档打分")
    add_para(doc, "内训师姓名：________________     课题名称：________________     评审日期：________年____月____日", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "评委姓名：________________     评委类型：□业务方 □AI 方法论 □大众评审", size=10)

    # 4 大类提示
    add_section_header(doc, "评估框架（4 大类 × 20 项）", level=1)
    add_para(doc, "A. 工具判断力（01-05）", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "B. 场景植入能力（06-10）", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "C. 学员引导能力（11-15）", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "D. 教学法多样性（16-20）", size=10, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "A 类·工具判断力（01-05）", level=1)
    rows_a = [
        ["01", "工具选择判断力", "能否根据场景精准选择 AI 工具", "", "", "", "", "", "____"],
        ["02", "工具组合能力", "能否把多个工具组合解决复杂问题", "", "", "", "", "", "____"],
        ["03", "数智小西/内部平台定位讲解", "能否讲清内部平台能力边界", "", "", "", "", "", "____"],
        ["04", "工具局限说明", "能否准确说出每个工具的局限", "", "", "", "", "", "____"],
        ["05", "工具迭代预判", "能否预判工具升级对方法的影响", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows_a,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "B 类·场景植入能力（06-10）", level=1)
    rows_b = [
        ["06", "业务问题诊断", "能否快速诊断学员业务痛点", "", "", "", "", "", "____"],
        ["07", "场景化教学设计", "能否把 AI 嵌入到真实业务场景", "", "", "", "", "", "____"],
        ["08", "41 个场景清单的活用", "能否灵活引用 41 场景清单", "", "", "", "", "", "____"],
        ["09", "跨岗位场景迁移", "能否让学员把场景迁移到自己岗位", "", "", "", "", "", "____"],
        ["10", "场景化课后行动", "能否让学员在课后真正用起来", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows_b,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "C 类·学员引导能力（11-15）", level=1)
    rows_c = [
        ["11", "0 基础学员引导", "能否让 0 基础学员听懂", "", "", "", "", "", "____"],
        ["12", "跟练节奏把控", "能否把控跟练的 5 步节拍", "", "", "", "", "", "____"],
        ["13", "模板提炼引导", "能否引导学员自己提炼模板", "", "", "", "", "", "____"],
        ["14", "错误示范应对", "学员做错时能否及时纠偏", "", "", "", "", "", "____"],
        ["15", "学员成果反馈", "能否给学员具体可改进的反馈", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows_c,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "D 类·教学法多样性（16-20）", level=1)
    rows_d = [
        ["16", "5 步教学节拍执行度", "引发兴趣→植入工具→跟练→模板→交付，是否执行到位", "", "", "", "", "", "____"],
        ["17", "案例/故事/游戏/视频运用", "形式是否多样，是否有亮点", "", "", "", "", "", "____"],
        ["18", "互动设计", "是否能调动学员参与", "", "", "", "", "", "____"],
        ["19", "控场能力", "时间 / 节奏 / 氛围把控", "", "", "", "", "", "____"],
        ["20", "安全合规意识持续植入", "是否全程持续强调信息安全", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows_d,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "总评分（4 类分项 + 总均分）", level=1)
    add_table_with_header(doc, ["分类", "项目数", "小计平均分", "对应档位"],
                          [
                              ["A 工具判断力", "5", "____", "□很好 □较好 □一般 □较差 □差"],
                              ["B 场景植入能力", "5", "____", "□很好 □较好 □一般 □较差 □差"],
                              ["C 学员引导能力", "5", "____", "□很好 □较好 □一般 □较差 □差"],
                              ["D 教学法多样性", "5", "____", "□很好 □较好 □一般 □较差 □差"],
                              ["总均分（20 项）", "20", "____", "□很好 □较好 □一般 □较差 □差"],
                          ],
                          col_widths=[Cm(4), Cm(2), Cm(3), Cm(7)])

    add_section_header(doc, "评委评语", level=1)
    add_para(doc, "亮点（最强 1-2 项）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "待改进（最弱 1-2 项）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)

    add_info_box(doc, "重要：单项 ≤ 14 分（差）将触发该内训师复评；连续 2 次不达标需重修基础班", color=COLOR_WARN)
    add_signature_row(doc, "评委签名：", "评分日期：")
    doc.save(os.path.join(OUT_DIR_2, "D-09-内训师班-AI教练技能评估表（每人一份）.docx"))
    print("[OK] D-09")


# ==================== D-10 内训师班·10 项课程包交付物检查表 ====================

def make_D10():
    doc = Document()
    setup_page(doc, "D-10")
    add_title_block(doc, "内训师班·10 项课程包交付物检查表", "10 项交付物 × 4 维 × 5 档")
    add_para(doc, "内训师姓名：________________     课题名称：________________     评审日期：________年____月____日", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "评委姓名：________________     评委类型：□业务方 □AI 方法论", size=10)

    add_section_header(doc, "4 维评估说明", level=1)
    add_para(doc, "· 完整度（25%）：交付物是否齐全、要素是否完整", size=10)
    add_para(doc, "· 质量（30%）：内容专业度、可读性、逻辑性", size=10)
    add_para(doc, "· AI 嵌入度（25%）：AI 工具 + 提示词 + 工具地图是否深度嵌入", size=10)
    add_para(doc, "· 可推广性（20%）：换讲师 / 换学员 / 换场景能否复用", size=10)

    add_section_header(doc, "10 项课程包交付物清单", level=1)

    headers = ["#", "交付物", "完整度 25%", "质量 30%", "AI 嵌入度 25%", "可推广性 20%", "加权得分"]
    items = [
        ["01", "课程定位表（明确业务痛点 + 学员对象 + 教学目标）", "____", "____", "____", "____", "____"],
        ["02", "三级大纲（一级模块 / 二级单元 / 三级知识点）", "____", "____", "____", "____", "____"],
        ["03", "课程 PPT（≥ 20 页，案例 + 工具 + 提示词嵌入）", "____", "____", "____", "____", "____"],
        ["04", "典型案例库（≥ 3 个真实业务场景案例）", "____", "____", "____", "____", "____"],
        ["05", "提示词说明书（含 LangGPT 四段式示例）", "____", "____", "____", "____", "____"],
        ["06", "教学进度表（精确到分钟的节奏表）", "____", "____", "____", "____", "____"],
        ["07", "22 题课后题库（含答案 + 评分标准）", "____", "____", "____", "____", "____"],
        ["08", "3-5 任务行动改善计划（学员课后应用任务）", "____", "____", "____", "____", "____"],
        ["09", "讲师手册（含 5 步教学节拍 + AI 追问应对）", "____", "____", "____", "____", "____"],
        ["10", "学员手册（含跟练 + 模板 + 工具地图填空）", "____", "____", "____", "____", "____"],
    ]
    add_table_with_header(doc, headers, items,
                           col_widths=[Cm(0.8), Cm(6.2), Cm(2), Cm(2), Cm(2), Cm(2), Cm(1.5)])

    add_section_header(doc, "加权得分公式", level=1)
    add_para(doc, "每项加权得分 = 完整度 × 25% + 质量 × 30% + AI 嵌入度 × 25% + 可推广性 × 20%", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "10 项平均分 = _____ 分    对应档位：□很好 □较好 □一般 □较差 □差", size=11, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "关键质量门槛（一票否决）", level=1)
    add_para(doc, "· 课程定位表必须有：业务痛点 + 学员对象 + 教学目标（缺一即不通过）", size=10)
    add_para(doc, "· 课程 PPT 必须嵌入 AI 工具图 + 提示词模板（缺一即不通过）", size=10)
    add_para(doc, "· 提示词说明书必须使用 LangGPT 四段式（否则 0 分）", size=10)
    add_para(doc, "· 讲师手册 / 学员手册必须含红黄绿灯合规说明（否则 0 分）", size=10)

    add_section_header(doc, "评委评语", level=1)
    add_para(doc, "亮点（最强 1-2 项）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "待改进（最弱 1-2 项）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)

    add_info_box(doc, "提示：单项平均分 < 15 分（较差）需返工重做；不通过项需在 D6 结营前修订", color=COLOR_WARN)
    add_signature_row(doc, "评委签名：", "评分日期：")
    doc.save(os.path.join(OUT_DIR_2, "D-10-内训师班-10项课程包交付物检查表（每人一份）.docx"))
    print("[OK] D-10")


# ==================== D-11 信息安全合规一票否决检查表 ====================

def make_D11():
    doc = Document()
    setup_page(doc, "D-11")
    add_title_block(doc, "信息安全合规一票否决检查表", "12 类岗位敏感清单 + 5 类脱敏标准 + 4 条一票否决")
    add_para(doc, "学员姓名：________________     学员岗位：________________     学员部门：________________", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "提示词模板数：________     工具地图数：________     评审日期：________年____月____日", size=10)
    add_para(doc, "评委姓名：________________     评委类型：□IT 安全 □AI 方法论", size=10)

    add_section_header(doc, "一、4 条一票否决（任一触发即 0 分）", level=1)
    veto_items = [
        "1. 不用 AI 验证 AI —— 任何 AI 生成的「合规」「安全」「无风险」结论必须由人复核，否则一票否决",
        "2. AI 说「没有」 ≠ 事实不存在 —— AI 未列出的风险不能视为不存在，必须人工逐条排查，否则一票否决",
        "3. AI 推测的内部规定全部作废 —— AI 推测的「公司内部规定」「部门规章」若未在官方文档中验证，一律作废，否则一票否决",
        "4. 正式输出必须人工核实后发出 —— 任何发给外部 / 领导的 AI 辅助材料必须人工核实，否则一票否决",
    ]
    for i, v in enumerate(veto_items):
        add_para(doc, v, size=10, color=COLOR_RED, bold=True)
        add_para(doc, f"自评：□未触发 □已触发（触发即终止评审）    评委确认：□未触发 □已触发", size=10)

    add_section_header(doc, "二、5 类脱敏标准（必查）", level=1)
    headers = ["#", "敏感信息类型", "脱敏标准", "自评", "评委确认"]
    desensitize_rows = [
        ["01", "公司名", "→ XX 公司", "□合规 □不合规", "□合规 □不合规"],
        ["02", "人名", "→ 员工 A / 张三", "□合规 □不合规", "□合规 □不合规"],
        ["03", "产品代号", "→ 项目 X", "□合规 □不合规", "□合规 □不合规"],
        ["04", "金额", "→ XX 万元", "□合规 □不合规", "□合规 □不合规"],
        ["05", "日期", "→ 保留格式替换年份（2024 → 2023）", "□合规 □不合规", "□合规 □不合规"],
    ]
    add_table_with_header(doc, headers, desensitize_rows,
                           col_widths=[Cm(0.8), Cm(3), Cm(4.5), Cm(3.5), Cm(3.5)])

    add_section_header(doc, "三、12 类岗位敏感信息清单（按岗位勾选）", level=1)
    sensitive_headers = ["#", "敏感信息类型", "本岗位是否涉及", "学员勾选", "评委确认"]
    sensitive_rows = [
        ["01", "客户个人信息（身份证 / 手机 / 住址）", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["02", "客户商业信息（合同 / 报价 / 折扣）", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["03", "员工薪酬 / 绩效 / 人事档案", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["04", "财务报表 / 经营数据", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["05", "产品技术参数 / 源代码 / 算法", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["06", "供应链信息（供应商 / 采购价）", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["07", "战略规划 / 商业机密", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["08", "未公开的并购 / 投融资", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["09", "内部审批流程 / 印章", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["10", "法务纠纷 / 合规风险", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["11", "项目代号 / 客户代号", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
        ["12", "其他（请注明）____________", "□是 □否", "□已脱敏 □未脱敏", "□合规 □不合规"],
    ]
    add_table_with_header(doc, sensitive_headers, sensitive_rows,
                           col_widths=[Cm(0.8), Cm(5), Cm(2.5), Cm(3.5), Cm(3.5)])

    add_section_header(doc, "四、红黄绿灯使用情况", level=1)
    add_para(doc, "🟢 绿灯（公开信息可直接输入）：", size=10, bold=True, color=COLOR_SAFE)
    add_para(doc, "· 是否已识别？□是 □否", size=10)
    add_para(doc, "🟡 黄灯（公司内部信息脱敏后输入）：", size=10, bold=True, color=COLOR_YELLOW)
    add_para(doc, "· 是否已脱敏？□是 □否（脱敏后需保留业务可识别度）", size=10)
    add_para(doc, "🔴 红灯（保密信息禁止输入）：", size=10, bold=True, color=COLOR_RED)
    add_para(doc, "· 是否有红灯项出现在提示词 / 工具地图？□否（合规）□是（违规）", size=10)

    add_section_header(doc, "五、综合判定", level=1)
    add_para(doc, "□ 全部合规，准予通过", size=11, bold=True, color=COLOR_SAFE)
    add_para(doc, "□ 存在不合规项，需 24 小时内修订重审", size=11, bold=True, color=COLOR_WARN)
    add_para(doc, "□ 触发一票否决项，本次评审作废，需重修课程", size=11, bold=True, color=COLOR_RED)

    add_info_box(doc, "本表为信息安全一票否决：触发 4 条一票否决任意一条，本次评审直接作废", color=COLOR_RED, icon="⛔")
    add_signature_row(doc, "学员自评签名：", "自评日期：")
    add_signature_row(doc, "评委签名：", "评审日期：")
    doc.save(os.path.join(OUT_DIR_2, "D-11-信息安全合规一票否决检查表（每人一份）.docx"))
    print("[OK] D-11")


# ==================== D-13 课程评审指引 ====================

def make_D13():
    doc = Document()
    setup_page(doc, "D-13")
    add_title_block(doc, "课程评审指引（AI 课专用）", "说课 5 步 + 试讲 6 条 + AI 追问 5 问")

    add_para(doc, "亲爱的伙伴：", size=11, bold=True, color=COLOR_MAIN)
    add_para(doc, "当看到这份文件时，恭喜你在本次 AI 赋能内训师项目中的表现和成果已经获得认可，即将进入项目最终环节——课程评审。", size=10)
    add_para(doc, "本课程评审与传统课程评审有 3 大区别，请务必提前了解：", size=10)
    add_para(doc, "1. 评审对象不只是「课程」，更是「AI 方案 + 提示词 + 工具地图」三件套", size=10)
    add_para(doc, "2. 评审维度不只是「讲得好」，更要「业务问题被解得怎么样」", size=10)
    add_para(doc, "3. 评审纪律加了一条：信息安全合规一票否决", size=10)
    add_para(doc, "祝你在评审中展现真正的 AI 内训师实力！", size=10, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "第一部分：AI 课说课（5 步 · 5 分钟）", level=1)
    add_para(doc, "针对对象：评委（不是学员）    时长：5 分钟", size=10, bold=True, color=COLOR_AI)

    add_section_header(doc, "步骤 1：业务问题诊断（1 分钟）", level=2)
    add_para(doc, "要点：基于目前存在什么问题，你针对什么对象开发了这样一门 AI 课程", size=10)
    add_para(doc, "参考话术：", size=10, bold=True)
    add_para(doc, "「各位评委老师好，我是 XX 部门 XXX，在这次项目中，我主导开发的课程是《XXX》，学员对象是 XX。", size=10)
    add_para(doc, "为什么会开发这个课程呢？因为我们部门/我们岗位，目前存在 XXX 痛点：", size=10)
    add_para(doc, "  · 痛点 1：XXX 数据（耗时长 / 出错率高 / 重复劳动）", size=10)
    add_para(doc, "  · 痛点 2：XXX 数据（AI 可以解决但没人会用）", size=10)
    add_para(doc, "针对以上痛点，我开发了这门 AI 课程。」", size=10)

    add_section_header(doc, "步骤 2：AI 方案设计（1 分钟）", level=2)
    add_para(doc, "要点：讲清你打算用什么 AI 工具 + 什么提示词方法解决问题", size=10)
    add_para(doc, "参考话术：", size=10, bold=True)
    add_para(doc, "「我设计的 AI 方案分 3 步：", size=10)
    add_para(doc, "  · 第 1 步：用 XX 工具（数智小西 / 外部 AI）收集信息", size=10)
    add_para(doc, "  · 第 2 步：用 LangGPT 四段式提示词（角色/背景/目标/约束）让 AI 生成 XXX", size=10)
    add_para(doc, "  · 第 3 步：人工 + AI 双重核验（信息安全合规）」", size=10)

    add_section_header(doc, "步骤 3：提示词模板展示（1 分钟）", level=2)
    add_para(doc, "要点：直接展示你打磨好的 1-2 个提示词模板（脱敏后）", size=10)
    add_para(doc, "参考话术：", size=10, bold=True)
    add_para(doc, "「我打磨了 3 个提示词模板，其中最经典的是『合同审查提示词』：", size=10)
    add_para(doc, "  · 角色：你是资深法务专家", size=10)
    add_para(doc, "  · 背景：审查一份采购合同", size=10)
    add_para(doc, "  · 目标：找出 5 类风险条款", size=10)
    add_para(doc, "  · 约束：仅输出风险条款 + 引用原文位置，不输出修改建议", size=10)
    add_para(doc, "这个提示词已经在我们部门复用了 23 次，节省 70% 审查时间。」", size=10)

    add_section_header(doc, "步骤 4：效果数据对比（1 分钟）", level=2)
    add_para(doc, "要点：用真实数据对比「用 AI 前」 vs 「用 AI 后」", size=10)
    add_para(doc, "参考话术：", size=10, bold=True)
    add_para(doc, "「实施 1 个月后，我们部门统计的效果数据：", size=10)
    add_para(doc, "  · 提示词应用次数：人均 12 次 / 周", size=10)
    add_para(doc, "  · 业务节省时间：合同审查 8 小时 → 1.5 小时（节省 80%）", size=10)
    add_para(doc, "  · 业务产出：部门周报自动生成率 95%", size=10)
    add_para(doc, "  · 同事复用：5 个同岗位同事引入使用」", size=10)

    add_section_header(doc, "步骤 5：可复制性论证（1 分钟）", level=2)
    add_para(doc, "要点：换讲师 / 换学员 / 换场景能否复用", size=10)
    add_para(doc, "参考话术：", size=10, bold=True)
    add_para(doc, "「这套方法的复制门槛很低：", size=10)
    add_para(doc, "  · 换讲师：配套讲师手册 + 学员手册 + 22 题题库，新讲师 2 天即可上手", size=10)
    add_para(doc, "  · 换学员：0 基础学员 1 天可完成跟练", size=10)
    add_para(doc, "  · 换场景：提示词模板按业务类型可适配，已在 3 个部门验证」", size=10)

    add_section_header(doc, "第二部分：试讲片断（8 分钟）", level=1)
    add_para(doc, "针对对象：假设的学员（不是评委）    时长：8 分钟（开场 1 分钟 + 主体 6 分钟 + 结尾 1 分钟）", size=10, bold=True, color=COLOR_AI)

    add_section_header(doc, "试讲 6 条", level=2)
    trial_rules = [
        "1. 只聚焦一个 AI 场景 —— 讲透讲精彩就好，不要贪多",
        "2. 形式多样 —— 实操 + 案例 + 视频 + 讨论，不能硬讲",
        "3. 有亮点 —— 让人记忆深刻的一个金句 / 演示",
        "4. 开场结尾精心设计 —— 不能太平淡",
        "5. 完全按培训师角色走 —— 不要受评委干扰",
        "6. 同组之间互相过渡衔接 —— 不要突兀切换",
    ]
    for r in trial_rules:
        add_para(doc, r, size=10)

    add_section_header(doc, "第三部分：AI 追问应对（5 分钟）", level=1)
    add_para(doc, "评委在试讲后会问 5 个高频 AI 追问，请提前准备答案。", size=10)

    questions = [
        ("Q1：这个提示词在公司内部平台（数智小西）能跑吗？",
         "答：能 / 不能（如实回答）。如果不能，说明迁移方案；如果能，说明已在内部平台验证几次。"),
        ("Q2：如果学员不告诉你他具体做什么岗位，这套提示词还成立吗？",
         "答：成立，因为提示词是「业务类型 + 操作步骤」结构，不依赖具体岗位。可现场演示换岗位后的效果。"),
        ("Q3：数据脱敏你具体做了哪些处理？",
         "答：按 5 类脱敏标准（公司名 / 人名 / 产品代号 / 金额 / 日期），并使用红黄绿灯分级。详见 D-11。"),
        ("Q4：你这个 AI 方案的成本和传统方案比，节省在哪里？",
         "答：用「时间成本 + 错误成本 + 培训成本」3 维对比，引用真实数据。"),
        ("Q5：如果 AI 平台下周升级了，你这套方法还能用吗？",
         "答：能，因为方法论不变（LangGPT 四段式），只迭代提示词细节。可演示 3 次迭代记录。"),
    ]
    for q, a in questions:
        add_para(doc, q, size=10, bold=True, color=COLOR_MAIN)
        add_para(doc, a, size=10)
        add_para(doc, "", size=8)

    add_section_header(doc, "第四部分：评审纪律提醒", level=1)
    add_para(doc, "1. 任何环节若出现信息安全违规（一票否决 4 条中任意一条），立即终止评审", size=10)
    add_para(doc, "2. 评审过程中禁止使用未脱敏的真实数据，违者视为不合规", size=10)
    add_para(doc, "3. 评审结果将录入学习地图，作为后续晋级的依据", size=10)
    add_para(doc, "4. 对评审结果有异议，可在 24 小时内书面申诉", size=10)

    add_info_box(doc, "评审只是手段，成长才是目的。以轻松心态面对，会有更好发挥！", color=COLOR_SAFE)
    add_signature_row(doc, "学员签名：", "领取日期：")
    doc.save(os.path.join(OUT_DIR_2, "D-13-课程评审指引（每人一份）.docx"))
    print("[OK] D-13")


# ==================== D-06 主函数（最后调用）====================

# D-06 工具函数已定义在 make_D06，但代码缺保存 — 修复
def save_D06(doc):
    doc.save(os.path.join(OUT_DIR_2, "D-06-基础班-提示词模板评分卡（每人一份）.docx"))
    print("[OK] D-06")


# 补全 D-06 的结尾
def make_D06_full():
    doc = Document()
    setup_page(doc, "D-06")
    add_title_block(doc, "基础班·提示词模板评分卡", "10 条评分项 × 5 档打分（24-25/21-23/18-20/15-17/0-14）")

    add_para(doc, "学员姓名：________________     学员岗位：________________     学员部门：________________", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "提示词模板名称：________________     提交日期：________年____月____日", size=10)
    add_para(doc, "评委姓名：________________     评委类型：□业务方 □AI 方法论 □大众评审", size=10)

    add_section_header(doc, "评分标准", level=1)
    add_para(doc, "· 24-25 分（很好）：超出预期，可作标杆，全维度无明显短板", size=10)
    add_para(doc, "· 21-23 分（较好）：达到预期，有亮点，1-2 项可提升", size=10)
    add_para(doc, "· 18-20 分（一般）：达到基本要求，3 项左右待优化", size=10)
    add_para(doc, "· 15-17 分（较差）：未达预期，需大幅改进", size=10)
    add_para(doc, "· 0-14 分（差）：严重不达标，不建议继续推进", size=10)

    add_section_header(doc, "评分表（10 条）", level=1)
    rows = [
        ["01", "业务场景还原度", "是否针对真实工作场景、痛点描述具体", "", "", "", "", "", "____"],
        ["02", "四段式结构完整性", "角色 / 背景 / 目标 / 约束四要素是否齐全", "", "", "", "", "", "____"],
        ["03", "角色定位精准度", "AI 角色是否专业、有边界", "", "", "", "", "", "____"],
        ["04", "约束条件合理性", "安全 / 格式 / 风格约束是否到位", "", "", "", "", "", "____"],
        ["05", "提示词可复用性", "换场景 / 换人能否用", "", "", "", "", "", "____"],
        ["06", "业务价值可衡量", "能否算出节省时间 / 提升质量", "", "", "", "", "", "____"],
        ["07", "提示词迭代次数", "至少测试过 3 次并优化", "", "", "", "", "", "____"],
        ["08", "同事复用情况", "是否被同岗位其他人用", "", "", "", "", "", "____"],
        ["09", "信息安全合规", "是否遵守红黄绿灯 + 脱敏", "", "", "", "", "", "____"],
        ["10", "与工具地图一致性", "提示词所用工具是否与个人工具地图匹配", "", "", "", "", "", "____"],
    ]
    add_table_with_header(doc, SCORING_HEADERS, rows,
                           col_widths=[Cm(0.8), Cm(2.5), Cm(4.2), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.3), Cm(1.5)])

    add_section_header(doc, "总评分", level=1)
    add_para(doc, "10 项平均分 = _____ 分    对应档位：□很好 □较好 □一般 □较差 □差", size=11, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "评委评语", level=1)
    add_para(doc, "亮点（值得推广的部分）：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "改进建议：", size=10, bold=True)
    add_para(doc, "____________________________________________________", size=10)
    add_para(doc, "____________________________________________________", size=10)

    add_info_box(doc, "重要：本表仅为「提示词模板」单点评分。学员整体得分 = 提示词 × 50% + 工具地图 × 10% + 场景化 × 20% + AI 陪跑 × 20%", color=COLOR_WARN)
    add_signature_row(doc, "评委签名：", "评分日期：")
    save_D06(doc)


# ==================== 执行 ====================

if __name__ == "__main__":
    print("=== 开始生成 Word 批次 2（评审实施 6 份）===")
    make_D06_full()
    make_D07()
    make_D09()
    make_D10()
    make_D11()
    make_D13()
    print("=== 批次 2 完成 ===")

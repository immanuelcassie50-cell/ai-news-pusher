#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UYLP Student Handbook Generator
Creates a comprehensive Word document for the UYLP (Unleash Your Leadership Potential) course
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Color scheme
DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
MEDIUM_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE = RGBColor(0x44, 0x72, 0xC4)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)
ORANGE = RGBColor(0xED, 0x7D, 0x31)

def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_table_with_header(doc, headers, rows, col_widths=None):
    """Create a formatted table with header row"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], '2E75B6')
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data else ""
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[col_idx], 'F2F2F2')

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def add_form_title(doc, form_num, form_name, description=""):
    """Add form title with special styling"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"【表单 {form_num}】{form_name}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE

    if description:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        run2 = p2.add_run(description)
        run2.italic = True
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK_GRAY

def add_module_header(doc, module_num, module_name, subtitle=""):
    """Add module header"""
    # Module title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"模块 {module_num}：{module_name}")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = DARK_BLUE

    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run(subtitle)
        run2.italic = True
        run2.font.size = Pt(12)
        run2.font.color.rgb = MEDIUM_BLUE

    add_horizontal_line(doc)

def add_learning_objectives(doc, objectives):
    """Add learning objectives section"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("学习目标")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = MEDIUM_BLUE

    for obj in objectives:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.paragraph_format.left_indent = Inches(0.3)
        bullet.paragraph_format.space_after = Pt(2)
        bullet.add_run(obj)

def add_knowledge_point(doc, title, content_items, is_key_point=False):
    """Add knowledge point section"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

    prefix = "【核心洞见】" if is_key_point else "【知识点】"
    run = p.add_run(f"{prefix}{title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK_BLUE if is_key_point else MEDIUM_BLUE

    for item in content_items:
        if isinstance(item, tuple):
            # Sub-item with description
            p_item = doc.add_paragraph(style='List Bullet')
            p_item.paragraph_format.left_indent = Inches(0.3)
            p_item.paragraph_format.space_after = Pt(2)
            run_title = p_item.add_run(item[0])
            run_title.bold = True
            if len(item) > 1:
                p_item.add_run(item[1])
        else:
            p_item = doc.add_paragraph(style='List Bullet')
            p_item.paragraph_format.left_indent = Inches(0.3)
            p_item.paragraph_format.space_after = Pt(2)
            p_item.add_run(item)

def add_exercise(doc, exercise_num, title, description, tasks):
    """Add exercise section"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"练习 {exercise_num}：{title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ORANGE

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after = Pt(8)
    p_desc.add_run(description)

    for task in tasks:
        if isinstance(task, tuple):
            p_task = doc.add_paragraph()
            p_task.paragraph_format.left_indent = Inches(0.3)
            p_task.paragraph_format.space_after = Pt(4)
            run_label = p_task.add_run(f"{task[0]}：")
            run_label.bold = True
            p_task.add_run(task[1])
        else:
            p_task = doc.add_paragraph(style='List Bullet')
            p_task.paragraph_format.left_indent = Inches(0.3)
            p_task.paragraph_format.space_after = Pt(4)
            p_task.add_run(task)

def add_role_play_card(doc, scenario_num, situation, role_a, role_b, focus_points):
    """Add role-play card that can be cut out"""
    # Add page break for cut-out card
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"【角色扮演卡 {scenario_num}】")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ORANGE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Situation box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF2CC')
    p_sit = cell.paragraphs[0]
    p_sit.add_run("情境设定：").bold = True
    p_sit.add_run(situation)
    p_sit.paragraph_format.space_after = Pt(0)

    # Roles
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'

    cell_a = table2.rows[0].cells[0]
    set_cell_shading(cell_a, 'E2EFDA')
    p_a = cell_a.paragraphs[0]
    p_a.add_run(f"角色A（{role_a[0]}）：").bold = True
    p_a.add_run(role_a[1])

    cell_b = table2.rows[0].cells[1]
    set_cell_shading(cell_b, 'E2EFDA')
    p_b = cell_b.paragraphs[0]
    p_b.add_run(f"角色B（{role_b[0]}）：").bold = True
    p_b.add_run(role_b[1])

    # Focus points
    cell_fp = table2.rows[1].cells[0]
    cell_fp.merge(table2.rows[1].cells[1])
    set_cell_shading(cell_fp, 'DEEBF7')
    p_fp = cell_fp.paragraphs[0]
    p_fp.add_run("观察重点：").bold = True
    for point in focus_points:
        p_fp.add_run(f"\n• {point}")

    doc.add_paragraph()

def create_handbook():
    """Create the complete UYLP student handbook"""
    doc = Document()

    # Set document properties
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(11.69)  # A4 width
        section.page_height = Inches(16.54)  # A4 height
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ========== COVER PAGE ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("UYLP")
    run.bold = True
    run.font.size = Pt(60)
    run.font.color.rgb = DARK_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("释放你的领导潜能")
    run2.bold = True
    run2.font.size = Pt(36)
    run2.font.color.rgb = MEDIUM_BLUE

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(60)
    run3 = p3.add_run("学员手册")
    run3.bold = True
    run3.font.size = Pt(28)
    run3.font.color.rgb = DARK_GRAY

    # Course info
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(80)
    run4 = p4.add_run("Unleash Your Leadership Potential")
    run4.italic = True
    run4.font.size = Pt(16)
    run4.font.color.rgb = DARK_GRAY

    # Student info section
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_before = Pt(120)
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("学员姓名：__________________________    所在部门：__________________________")
    run5.font.size = Pt(12)

    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.paragraph_format.space_before = Pt(12)
    run6 = p6.add_run("课程日期：__________________________    课程讲师：__________________________")
    run6.font.size = Pt(12)

    # Quote
    p7 = doc.add_paragraph()
    p7.paragraph_format.space_before = Pt(100)
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run7 = p7.add_run('"领导力不是关于职位，而是关于影响。"')
    run7.italic = True
    run7.font.size = Pt(14)
    run7.font.color.rgb = MEDIUM_BLUE

    # Page break to TOC
    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    p_toc = doc.add_paragraph()
    p_toc.paragraph_format.space_before = Pt(0)
    p_toc.paragraph_format.space_after = Pt(20)
    run_toc = p_toc.add_run("目 录")
    run_toc.bold = True
    run_toc.font.size = Pt(24)
    run_toc.font.color.rgb = DARK_BLUE
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_horizontal_line(doc)

    toc_items = [
        ("课程介绍", "4"),
        ("模块1：领导者角色与绩效管理体系", "6"),
        ("模块2：辅导入门——WHEN/HOW/Y", "8"),
        ("模块3：有效反馈——CAIR模型", "10"),
        ("模块4：困难谈话", "12"),
        ("模块5：联结沟通与工作关系", "14"),
        ("模块6：高级辅导、有效授权与MAP", "16"),
        ("工具表单汇编", "18"),
        ("练习工作坊", "22"),
        ("课后资源", "24"),
    ]

    for item, page in toc_items:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.space_after = Pt(8)
        run_item = p_item.add_run(item)
        run_item.font.size = Pt(12)
        # Add tab
        p_item.add_run("\t" * 8 + page)

    doc.add_page_break()

    # ========== COURSE INTRODUCTION ==========
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(0)
    p_intro.paragraph_format.space_after = Pt(12)
    run_intro = p_intro.add_run("课程介绍")
    run_intro.bold = True
    run_intro.font.size = Pt(22)
    run_intro.font.color.rgb = DARK_BLUE

    add_horizontal_line(doc)

    # Course Overview
    p_overview = doc.add_paragraph()
    p_overview.paragraph_format.space_before = Pt(16)
    run_overview = p_overview.add_run("UYLP课程概述")
    run_overview.bold = True
    run_overview.font.size = Pt(14)
    run_overview.font.color.rgb = MEDIUM_BLUE

    overview_text = """
UYLP（Unleash Your Leadership Potential，释放你的领导潜能）是一门专为中高层管理者设计的管理能力提升课程。本课程通过系统的理论框架和实用的管理工具，帮助管理者掌握新时期下领导团队的核心能力。

课程聚焦于管理者的核心职责——通过他人的绩效达成组织目标。管理者不再需要亲自完成所有工作，而是通过有效的辅导、反馈、授权等干预手段，激发团队成员的潜能，提升整体绩效。
"""
    p_overview_body = doc.add_paragraph()
    p_overview_body.paragraph_format.space_before = Pt(8)
    p_overview_body.add_run(overview_text.strip())

    # Course Architecture
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.space_before = Pt(20)
    run_arch = p_arch.add_run("课程架构")
    run_arch.bold = True
    run_arch.font.size = Pt(14)
    run_arch.font.color.rgb = MEDIUM_BLUE

    # Architecture diagram (text-based)
    arch_table = doc.add_table(rows=4, cols=3)
    arch_table.style = 'Table Grid'
    arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 1: Header
    set_cell_shading(arch_table.rows[0].cells[0], '1F3864')
    arch_table.rows[0].cells[0].paragraphs[0].add_run("领导力基础").bold = True
    arch_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    arch_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_cell_shading(arch_table.rows[0].cells[1], '1F3864')
    arch_table.rows[0].cells[1].paragraphs[0].add_run("核心干预技能").bold = True
    arch_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    arch_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_cell_shading(arch_table.rows[0].cells[2], '1F3864')
    arch_table.rows[0].cells[2].paragraphs[0].add_run("整合应用").bold = True
    arch_table.rows[0].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    arch_table.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 2
    arch_table.rows[1].cells[0].paragraphs[0].add_run("模块1\n领导者角色与\n绩效管理体系")
    arch_table.rows[1].cells[1].paragraphs[0].add_run("模块2-3\n辅导入门\n有效反馈CAIR")
    arch_table.rows[1].cells[2].paragraphs[0].add_run("模块4-5\n困难谈话\n联结沟通")

    # Row 3
    arch_table.rows[2].cells[0].paragraphs[0].add_run("")
    arch_table.rows[2].cells[1].paragraphs[0].add_run("")
    arch_table.rows[2].cells[2].paragraphs[0].add_run("")

    # Row 4: Integration
    arch_table.rows[3].cells[0].merge(arch_table.rows[3].cells[1])
    set_cell_shading(arch_table.rows[3].cells[0], 'ED7D31')
    arch_table.rows[3].cells[0].paragraphs[0].add_run("模块6：高级辅导、有效授权与MAP").bold = True
    arch_table.rows[3].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    arch_table.rows[3].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    set_cell_shading(arch_table.rows[3].cells[2], 'ED7D31')
    arch_table.rows[3].cells[2].paragraphs[0].add_run("综合实战").bold = True
    arch_table.rows[3].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    arch_table.rows[3].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    # Learning Objectives
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.space_before = Pt(20)
    run_obj = p_obj.add_run("学习目标")
    run_obj.bold = True
    run_obj.font.size = Pt(14)
    run_obj.font.color.rgb = MEDIUM_BLUE

    objectives = [
        "理解管理者的角色定位和绩效管理系统",
        "掌握辅导的基本框架WHEN/HOW/Y",
        "运用CAIR模型提供有效反馈",
        "处理困难谈话的情境和技巧",
        "建立积极的联结沟通与工作关系",
        "实施高级辅导技能和有效授权",
        "制定个人MAP（管理行动规划）"
    ]

    for obj in objectives:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.left_indent = Inches(0.3)
        p_bullet.paragraph_format.space_after = Pt(4)
        p_bullet.add_run(obj)

    doc.add_page_break()

    # ========== MODULE 1 ==========
    add_module_header(doc, 1, "领导者角色与绩效管理体系", "理解管理者的核心角色与绩效管理循环")

    add_learning_objectives(doc, [
        "识别管理者的三大角色",
        "理解绩效管理四步循环",
        "掌握绩效分析的基本方法"
    ])

    # Section 1.1
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.space_before = Pt(14)
    p_s1.paragraph_format.space_after = Pt(6)
    run_s1 = p_s1.add_run("1.1 管理者的三大角色")
    run_s1.bold = True
    run_s1.font.size = Pt(12)
    run_s1.font.color.rgb = MEDIUM_BLUE

    roles_table = doc.add_table(rows=4, cols=2)
    roles_table.style = 'Table Grid'
    roles_headers = ["角色类型", "核心职责", "日常行为"]
    for i, h in enumerate(["角色类型", "核心职责"]):
        roles_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        set_cell_shading(roles_table.rows[0].cells[i], '2E75B6')
        roles_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    roles_data = [
        ("战略执行者", "将组织战略转化为团队行动", "分解目标、配置资源、监控执行"),
        ("团队开发者", "培养团队成员能力", "辅导、反馈、授权、激励"),
        ("绩效管理者", "对绩效结果负责", "设定标准、评估表现、处理问题")
    ]
    for row_idx, (role, duty, action) in enumerate(roles_data):
        roles_table.rows[row_idx + 1].cells[0].paragraphs[0].add_run(role)
        roles_table.rows[row_idx + 1].cells[1].paragraphs[0].add_run(f"{duty}；{action}")

    # Section 1.2
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.space_before = Pt(14)
    p_s2.paragraph_format.space_after = Pt(6)
    run_s2 = p_s2.add_run("1.2 绩效管理四步循环")
    run_s2.bold = True
    run_s2.font.size = Pt(12)
    run_s2.font.color.rgb = MEDIUM_BLUE

    cycle_items = [
        ("计划制定", "设定绩效目标、明确成功标准、确定资源需求"),
        ("过程辅导", "持续沟通、提供支持、解决问题、调整方向"),
        ("绩效评估", "收集数据、评估成果、识别差距、给出反馈"),
        ("结果应用", "奖励优秀、改进不足、制定发展计划")
    ]

    for i, (step, desc) in enumerate(cycle_items):
        p_step = doc.add_paragraph()
        p_step.paragraph_format.space_after = Pt(6)
        run_num = p_step.add_run(f"步骤{i+1}：{step}")
        run_num.bold = True
        p_step.add_run(f" —— {desc}")

    # Form 1.1
    add_form_title(doc, "1.1", "绩效分析表",
                   "填写说明：针对一个具体的团队绩效问题，完成以下分析")

    perf_headers = ["分析维度", "具体描述"]
    perf_data = [
        ("绩效差距", ""),
        ("根本原因", ""),
        ("影响程度", ""),
        ("改进机会", ""),
        ("所需资源", ""),
        ("时间框架", "")
    ]
    create_table_with_header(doc, perf_headers, perf_data, [2, 4])

    doc.add_page_break()

    # ========== MODULE 2 ==========
    add_module_header(doc, 2, "辅导入门——WHEN/HOW/Y", "掌握辅导的时机、方法和原因")

    add_learning_objectives(doc, [
        "判断何时应该进行辅导",
        "掌握HOW：辅导的基本步骤",
        "理解为什么辅导对团队成功至关重要"
    ])

    # Section 2.1 WHEN
    p_w = doc.add_paragraph()
    p_w.paragraph_format.space_before = Pt(14)
    p_w.paragraph_format.space_after = Pt(6)
    run_w = p_w.add_run("2.1 WHEN——何时辅导")
    run_w.bold = True
    run_w.font.size = Pt(12)
    run_w.font.color.rgb = MEDIUM_BLUE

    when_table = doc.add_table(rows=5, cols=2)
    when_table.style = 'Table Grid'
    when_headers = ["辅导时机", "具体表现"]
    for i, h in enumerate(when_headers):
        when_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        set_cell_shading(when_table.rows[0].cells[i], '2E75B6')
        when_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    when_data = [
        ("新任务开始", "成员开始承担新职责时"),
        ("绩效问题出现", "表现低于预期或出现下滑时"),
        ("能力提升需求", "成员需要发展新技能时"),
        ("职业发展讨论", "进行绩效面谈或职业规划时")
    ]
    for row_idx, (when, desc) in enumerate(when_data):
        when_table.rows[row_idx + 1].cells[0].paragraphs[0].add_run(when)
        when_table.rows[row_idx + 1].cells[1].paragraphs[0].add_run(desc)

    # Section 2.2 HOW
    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(14)
    p_h.paragraph_format.space_after = Pt(6)
    run_h = p_h.add_run("2.2 HOW——如何辅导")
    run_h.bold = True
    run_h.font.size = Pt(12)
    run_h.font.color.rgb = MEDIUM_BLUE

    how_steps = [
        ("准备", "收集信息、明确目标、选择合适的时间和地点"),
        ("开场", "建立 rapport、说明目的、营造开放氛围"),
        ("探索", "倾听理解、提问澄清、确认共识"),
        ("发展", "共同讨论、探索方案、获得承诺"),
        ("结束", "总结要点、确认行动、安排跟进")
    ]

    for i, (step, desc) in enumerate(how_steps):
        p_how = doc.add_paragraph(style='List Number')
        p_how.paragraph_format.left_indent = Inches(0.3)
        p_how.paragraph_format.space_after = Pt(4)
        run_step = p_how.add_run(f"{step}：")
        run_step.bold = True
        p_how.add_run(desc)

    # Section 2.3 WHY
    p_y = doc.add_paragraph()
    p_y.paragraph_format.space_before = Pt(14)
    p_y.paragraph_format.space_after = Pt(6)
    run_y = p_y.add_run("2.3 Y——为什么辅导")
    run_y.bold = True
    run_y.font.size = Pt(12)
    run_y.font.color.rgb = MEDIUM_BLUE

    why_points = [
        "辅导是提升团队绩效最有效的干预手段",
        "通过辅导，管理者可以从\"救火\"转向\"防火\"",
        "辅导帮助成员成长，建立持续绩效的文化",
        "有效的辅导能够降低员工流失率"
    ]

    for point in why_points:
        p_why = doc.add_paragraph(style='List Bullet')
        p_why.paragraph_format.left_indent = Inches(0.3)
        p_why.paragraph_format.space_after = Pt(4)
        p_why.add_run(point)

    # Form 2.1
    add_form_title(doc, "2.1", "辅导谈话准备表",
                   "填写说明：在进行辅导谈话前，完成以下准备")

    coaching_headers = ["准备项目", "内容"]
    coaching_data = [
        ("谈话目标", ""),
        ("成员情况分析", ""),
        ("关键信息/数据", ""),
        ("可能的解决方案", ""),
        ("谈话时间安排", ""),
        ("预期结果", "")
    ]
    create_table_with_header(doc, coaching_headers, coaching_data, [2, 4])

    # Exercise 2-A
    add_exercise(doc, "2-A", "判断辅导时机",
                 "以下情境是否需要进行辅导？为什么？你会如何处理？", [
        ("情境1", "小李刚入职三个月，第一次独立完成项目，但报告中有几个明显错误"),
        ("情境2", "张经理团队的王姐连续两周迟到早退，工作效率明显下降"),
        ("情境3", "技术骨干小刘提出辞职，表示找到了更高薪资的工作"),
        ("情境4", "季度末，团队整体业绩完成了目标的105%")
    ])

    doc.add_page_break()

    # ========== MODULE 3 ==========
    add_module_header(doc, 3, "有效反馈——CAIR模型", "掌握建设性反馈的核心框架")

    add_learning_objectives(doc, [
        "理解CAIR模型的四个步骤",
        "掌握建设性反馈的表达技巧",
        "识别反馈中的常见陷阱"
    ])

    # CAIR Model
    p_cair = doc.add_paragraph()
    p_cair.paragraph_format.space_before = Pt(14)
    p_cair.paragraph_format.space_after = Pt(10)
    run_cair = p_cair.add_run("CAIR反馈模型")
    run_cair.bold = True
    run_cair.font.size = Pt(14)
    run_cair.font.color.rgb = DARK_BLUE

    cair_table = doc.add_table(rows=5, cols=2)
    cair_table.style = 'Table Grid'

    cair_data = [
        ("C - Context\n情境", "描述具体的时间、地点、情境，让对方清楚你在反馈什么"),
        ("A - Action\n行为", "描述具体的行为，而不是对人的判断或推测"),
        ("I - Impact\n影响", "说明这个行为对团队、项目或他人的具体影响"),
        ("R - Request\n请求", "提出具体的改进请求或期望")
    ]

    for i, (key, desc) in enumerate(cair_data):
        set_cell_shading(cair_table.rows[i].cells[0], '2E75B6')
        cair_table.rows[i].cells[0].paragraphs[0].add_run(key).bold = True
        cair_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
        cair_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cair_table.rows[i].cells[1].paragraphs[0].add_run(desc)

    # Feedback Tips
    p_tips = doc.add_paragraph()
    p_tips.paragraph_format.space_before = Pt(14)
    p_tips.paragraph_format.space_after = Pt(6)
    run_tips = p_tips.add_run("有效反馈的关键原则")
    run_tips.bold = True
    run_tips.font.size = Pt(12)
    run_tips.font.color.rgb = MEDIUM_BLUE

    tips = [
        "及时性：尽快反馈，不要等到问题积累",
        "具体性：避免模糊的表述，提供具体例子",
        "双向性：鼓励对方回应和提问",
        "平衡性：指出问题的同时，也认可做得好的地方",
        "建设性：focus在改进，而非批评"
    ]

    for tip in tips:
        p_tip = doc.add_paragraph(style='List Bullet')
        p_tip.paragraph_format.left_indent = Inches(0.3)
        p_tip.paragraph_format.space_after = Pt(4)
        p_tip.add_run(tip)

    # Form 3.1
    add_form_title(doc, "3.1", "CAIR反馈记录表",
                   "填写说明：使用CAIR框架记录一次反馈谈话")

    cair_form_headers = ["CAIR要素", "记录内容"]
    cair_form_data = [
        ("C - 情境", ""),
        ("A - 行为", ""),
        ("I - 影响", ""),
        ("R - 请求", ""),
        ("对方反应", ""),
        ("后续跟进", "")
    ]
    create_table_with_header(doc, cair_form_headers, cair_form_data, [2, 4])

    # Exercise 3-A
    add_exercise(doc, "3-A", "CAIR应用练习",
                 "使用CAIR框架，为以下情境设计反馈：", [
        ("情境", "团队成员小陈在客户会议上打断了客户三次，让客户很不高兴"),
        ("你的CAIR反馈", "\nC（情境）：\nA（行为）：\nI（影响）：\nR（请求）：")
    ])

    doc.add_page_break()

    # ========== MODULE 4 ==========
    add_module_header(doc, 4, "困难谈话", "处理敏感和复杂人际情境的技巧")

    add_learning_objectives(doc, [
        "识别困难谈话的常见类型",
        "掌握困难谈话的准备框架",
        "运用有效的谈话技巧"
    ])

    # Section 4.1
    p_d1 = doc.add_paragraph()
    p_d1.paragraph_format.space_before = Pt(14)
    p_d1.paragraph_format.space_after = Pt(6)
    run_d1 = p_d1.add_run("4.1 困难谈话的类型")
    run_d1.bold = True
    run_d1.font.size = Pt(12)
    run_d1.font.color.rgb = MEDIUM_BLUE

    difficult_types = [
        ("绩效问题谈话", "处理表现不佳的情况"),
        ("行为纪律谈话", "处理违反政策或行为准则的问题"),
        ("职业发展谈话", "讨论晋升、调岗或离职"),
        ("人际冲突谈话", "调解团队成员之间的矛盾"),
        ("敏感信息谈话", "传达裁员、降职等敏感消息")
    ]

    for dtype, desc in difficult_types:
        p_type = doc.add_paragraph(style='List Bullet')
        p_type.paragraph_format.left_indent = Inches(0.3)
        p_type.paragraph_format.space_after = Pt(4)
        run_dtype = p_type.add_run(f"{dtype}：")
        run_dtype.bold = True
        p_type.add_run(desc)

    # Section 4.2
    p_d2 = doc.add_paragraph()
    p_d2.paragraph_format.space_before = Pt(14)
    p_d2.paragraph_format.space_after = Pt(6)
    run_d2 = p_d2.add_run("4.2 困难谈话准备框架")
    run_d2.bold = True
    run_d2.font.size = Pt(12)
    run_d2.font.color.rgb = MEDIUM_BLUE

    prep_framework = [
        ("目标清晰", "明确谈话的核心目的和期望结果"),
        ("事实准备", "收集具体的事实和数据支持"),
        ("情感准备", "预判对方的情绪反应并准备应对"),
        ("方案准备", "准备可能的解决方案或替代路径"),
        ("环境选择", "选择私密、舒适、无干扰的谈话环境")
    ]

    for step, desc in prep_framework:
        p_prep = doc.add_paragraph(style='List Number')
        p_prep.paragraph_format.left_indent = Inches(0.3)
        p_prep.paragraph_format.space_after = Pt(4)
        run_prep = p_prep.add_run(f"{step}：")
        run_prep.bold = True
        p_prep.add_run(desc)

    # Section 4.3
    p_d3 = doc.add_paragraph()
    p_d3.paragraph_format.space_before = Pt(14)
    p_d3.paragraph_format.space_after = Pt(6)
    run_d3 = p_d3.add_run("4.3 困难谈话技巧")
    run_d3.bold = True
    run_d3.font.size = Pt(12)
    run_d3.font.color.rgb = MEDIUM_BLUE

    techniques = [
        "使用\"我\"开头的陈述，而非\"你\"开头",
        "保持冷静，避免情绪化的语言",
        "积极倾听，给对方表达的机会",
        "聚焦行为和结果，而非人格攻击",
        "寻找共同利益点",
        "保持开放的态度，愿意调整"
    ]

    for tech in techniques:
        p_tech = doc.add_paragraph(style='List Bullet')
        p_tech.paragraph_format.left_indent = Inches(0.3)
        p_tech.paragraph_format.space_after = Pt(4)
        p_tech.add_run(tech)

    # Form 4.1
    add_form_title(doc, "4.1", "困难谈话脚本表",
                   "填写说明：为即将进行的困难谈话准备脚本")

    script_headers = ["谈话结构", "内容要点"]
    script_data = [
        ("开场白", "说明谈话目的，营造开放氛围"),
        ("事实陈述", "描述具体事件和数据"),
        ("影响说明", "说明行为的后果和影响"),
        ("对方回应", "（记录对方的反应）"),
        ("共同讨论", "探索解决方案"),
        ("达成共识", "总结决定和下一步"),
        ("后续跟进", "安排跟进和检查")
    ]
    create_table_with_header(doc, script_headers, script_data, [2, 4])

    # Role-play card
    add_role_play_card(doc, "4-1",
        "团队成员小王连续一个月迟到早退，影响了团队整体士气和工作进度。作为他的经理，你需要和他进行一次困难谈话。",
        ("经理", "需要指出问题、了解原因、寻求改进承诺"),
        ("小王", "可能有各种原因：个人问题、工作压力、对管理方式不满等"),
        ["是否使用CAIR框架", "是否避免指责和情绪化", "是否倾听对方解释", "是否达成具体改进计划"]
    )

    doc.add_page_break()

    # ========== MODULE 5 ==========
    add_module_header(doc, 5, "联结沟通与工作关系", "建立信任、开放的工作关系")

    add_learning_objectives(doc, [
        "理解联结沟通的重要性",
        "掌握建立信任的关键行为",
        "识别和修复工作关系中的裂痕"
    ])

    # Section 5.1
    p_c1 = doc.add_paragraph()
    p_c1.paragraph_format.space_before = Pt(14)
    p_c1.paragraph_format.space_after = Pt(6)
    run_c1 = p_c1.add_run("5.1 联结沟通的要素")
    run_c1.bold = True
    run_c1.font.size = Pt(12)
    run_c1.font.color.rgb = MEDIUM_BLUE

    connection_elements = [
        ("信任", "诚实、可靠、言行一致"),
        ("尊重", "认可他人的价值和观点"),
        ("理解", "积极倾听，试图理解对方立场"),
        ("开放", "愿意分享信息，接受反馈")
    ]

    for elem, desc in connection_elements:
        p_elem = doc.add_paragraph(style='List Bullet')
        p_elem.paragraph_format.left_indent = Inches(0.3)
        p_elem.paragraph_format.space_after = Pt(4)
        run_elem = p_elem.add_run(f"{elem}：")
        run_elem.bold = True
        p_elem.add_run(desc)

    # Section 5.2
    p_c2 = doc.add_paragraph()
    p_c2.paragraph_format.space_before = Pt(14)
    p_c2.paragraph_format.space_after = Pt(6)
    run_c2 = p_c2.add_run("5.2 建立信任的关键行为")
    run_c2.bold = True
    run_c2.font.size = Pt(12)
    run_c2.font.color.rgb = MEDIUM_BLUE

    trust_behaviors = [
        "说到做到，言出必行",
        "承认错误，不推卸责任",
        "保守秘密，保护隐私",
        "公开分享信息，不过度过滤",
        "给予信任，也赢得信任",
        "在小事上坚持原则"
    ]

    for behavior in trust_behaviors:
        p_b = doc.add_paragraph(style='List Bullet')
        p_b.paragraph_format.left_indent = Inches(0.3)
        p_b.paragraph_format.space_after = Pt(4)
        p_b.add_run(behavior)

    # Section 5.3
    p_c3 = doc.add_paragraph()
    p_c3.paragraph_format.space_before = Pt(14)
    p_c3.paragraph_format.space_after = Pt(6)
    run_c3 = p_c3.add_run("5.3 修复工作关系")
    run_c3.bold = True
    run_c3.font.size = Pt(12)
    run_c3.font.color.rgb = MEDIUM_BLUE

    repair_steps = [
        "识别问题：正视关系中的紧张或裂痕",
        "主动沟通：不等待，主动迈出第一步",
        "诚恳道歉：如果是自己的责任，真诚道歉",
        "寻求理解：倾听对方的感受和观点",
        "共同修复：一起讨论如何改善关系",
        "持续跟进：保持关注，防止问题复发"
    ]

    for i, step in enumerate(repair_steps):
        p_repair = doc.add_paragraph(style='List Number')
        p_repair.paragraph_format.left_indent = Inches(0.3)
        p_repair.paragraph_format.space_after = Pt(4)
        p_repair.add_run(step)

    # Exercise 5-A
    add_exercise(doc, "5-A", "联结沟通反思",
                 "反思你当前的工作关系，完成以下评估：", [
        ("信任评估", "与团队成员的信任程度（1-10分）"),
        ("沟通开放度", "团队沟通的开放程度如何？"),
        ("关系改进点", "哪些关系需要重点关注和改进？"),
        ("行动计划", "你会采取什么具体行动？")
    ])

    doc.add_page_break()

    # ========== MODULE 6 ==========
    add_module_header(doc, 6, "高级辅导、有效授权与MAP", "整合领导技能，制定管理行动规划")

    add_learning_objectives(doc, [
        "运用高级辅导技术解决复杂问题",
        "掌握有效授权的原则和步骤",
        "制定个人MAP（管理行动规划）"
    ])

    # Section 6.1
    p_a1 = doc.add_paragraph()
    p_a1.paragraph_format.space_before = Pt(14)
    p_a1.paragraph_format.space_after = Pt(6)
    run_a1 = p_a1.add_run("6.1 高级辅导技术")
    run_a1.bold = True
    run_a1.font.size = Pt(12)
    run_a1.font.color.rgb = MEDIUM_BLUE

    advanced_techniques = [
        ("深度倾听", "不仅听内容，更关注情感和需求"),
        ("强力提问", "使用开放式问题引导思考"),
        ("情感映射", "识别和命名对方的情绪"),
        ("重构框架", "帮助对方从新角度看待问题"),
        ("挑战假设", "质疑限制性信念")
    ]

    for tech, desc in advanced_techniques:
        p_tech = doc.add_paragraph(style='List Bullet')
        p_tech.paragraph_format.left_indent = Inches(0.3)
        p_tech.paragraph_format.space_after = Pt(4)
        run_tech = p_tech.add_run(f"{tech}：")
        run_tech.bold = True
        p_tech.add_run(desc)

    # Section 6.2
    p_a2 = doc.add_paragraph()
    p_a2.paragraph_format.space_before = Pt(14)
    p_a2.paragraph_format.space_after = Pt(6)
    run_a2 = p_a2.add_run("6.2 有效授权")
    run_a2.bold = True
    run_a2.font.size = Pt(12)
    run_a2.font.color.rgb = MEDIUM_BLUE

    delegation_table = doc.add_table(rows=5, cols=2)
    delegation_table.style = 'Table Grid'

    del_headers = ["授权步骤", "关键要点"]
    for i, h in enumerate(del_headers):
        delegation_table.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        set_cell_shading(delegation_table.rows[0].cells[i], '2E75B6')
        delegation_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    del_data = [
        ("选择任务", "识别适合授权的任务"),
        ("选择人选", "评估成员能力和意愿"),
        ("明确期望", "说明目标、标准、资源"),
        ("保持跟进", "监控进度，适时支持")
    ]

    for row_idx, (step, point) in enumerate(del_data):
        delegation_table.rows[row_idx + 1].cells[0].paragraphs[0].add_run(step)
        delegation_table.rows[row_idx + 1].cells[1].paragraphs[0].add_run(point)

    # Section 6.3 MAP
    p_a3 = doc.add_paragraph()
    p_a3.paragraph_format.space_before = Pt(14)
    p_a3.paragraph_format.space_after = Pt(6)
    run_a3 = p_a3.add_run("6.3 MAP——管理行动规划")
    run_a3.bold = True
    run_a3.font.size = Pt(12)
    run_a3.font.color.rgb = MEDIUM_BLUE

    p_map_desc = doc.add_paragraph()
    p_map_desc.paragraph_format.space_before = Pt(4)
    p_map_desc.add_run("MAP（Management Action Plan）是帮助管理者将学习转化为行动的工具。")

    # Form 6.1
    add_form_title(doc, "6.1", "授权评估表",
                   "填写说明：评估任务授权的适合性")

    auth_headers = ["评估维度", "评估内容"]
    auth_data = [
        ("任务性质", "这项任务是常规还是创新？"),
        ("时间紧迫性", "时间是否允许充分的授权？"),
        ("成员能力", "成员是否具备完成任务的能力？"),
        ("成员意愿", "成员是否愿意承担这项任务？"),
        ("风险程度", "授权失败的后果是什么？"),
        ("所需资源", "成员需要什么资源支持？")
    ]
    create_table_with_header(doc, auth_headers, auth_data, [2, 4])

    doc.add_page_break()

    # ========== TOOL FORMS COLLECTION ==========
    doc.add_page_break()
    p_tools = doc.add_paragraph()
    p_tools.paragraph_format.space_before = Pt(0)
    p_tools.paragraph_format.space_after = Pt(12)
    run_tools = p_tools.add_run("工具表单汇编")
    run_tools.bold = True
    run_tools.font.size = Pt(22)
    run_tools.font.color.rgb = DARK_BLUE

    add_horizontal_line(doc)

    p_tools_intro = doc.add_paragraph()
    p_tools_intro.paragraph_format.space_before = Pt(12)
    p_tools_intro.add_run("以下工具表单可作为日常管理的实用参考，建议打印后随时使用。")

    # Form collection
    forms_list = [
        ("表单1.1", "绩效分析表", "模块1"),
        ("表单2.1", "辅导谈话准备表", "模块2"),
        ("表单3.1", "CAIR反馈记录表", "模块3"),
        ("表单4.1", "困难谈话脚本表", "模块4"),
        ("表单5.1", "工作关系自检表", "模块5"),
        ("表单6.1", "授权评估表", "模块6"),
        ("表单6.2", "MAP行动计划表", "模块6")
    ]

    for form_num, form_name, module in forms_list:
        p_form = doc.add_paragraph()
        p_form.paragraph_format.space_before = Pt(8)
        p_form.paragraph_format.space_after = Pt(4)
        run_fnum = p_form.add_run(f"{form_num} {form_name}")
        run_fnum.bold = True
        run_fnum.font.size = Pt(11)
        run_fnum.font.color.rgb = MEDIUM_BLUE
        p_form.add_run(f" （{module}）")

    doc.add_page_break()

    # MAP Form
    add_form_title(doc, "6.2", "MAP行动计划表",
                   "填写说明：制定你的90天管理行动规划")

    map_headers = ["时间阶段", "重点领域", "具体行动", "成功指标", "时间节点"]
    map_data = [
        ("第1-30天", "", "", "", ""),
        ("第31-60天", "", "", "", ""),
        ("第61-90天", "", "", "", "")
    ]
    create_table_with_header(doc, map_headers, map_data, [1.2, 1.5, 2.5, 1.5, 1])

    p_map_sign = doc.add_paragraph()
    p_map_sign.paragraph_format.space_before = Pt(20)
    p_map_sign.add_run("我承诺执行以上行动计划：")
    p_map_sign.add_run("\n\n签名：________________    日期：________________")

    doc.add_page_break()

    # ========== PRACTICE WORKSHOP ==========
    doc.add_page_break()
    p_ws = doc.add_paragraph()
    p_ws.paragraph_format.space_before = Pt(0)
    p_ws.paragraph_format.space_after = Pt(12)
    run_ws = p_ws.add_run("练习工作坊")
    run_ws.bold = True
    run_ws.font.size = Pt(22)
    run_ws.font.color.rgb = DARK_BLUE

    add_horizontal_line(doc)

    # Workshop 1: Coaching Practice
    p_ws1 = doc.add_paragraph()
    p_ws1.paragraph_format.space_before = Pt(16)
    run_ws1 = p_ws1.add_run("工作坊1：辅导技能练习")
    run_ws1.bold = True
    run_ws1.font.size = Pt(14)
    run_ws1.font.color.rgb = MEDIUM_BLUE

    p_ws1_desc = doc.add_paragraph()
    p_ws1_desc.add_run("练习目标：运用WHEN/HOW/Y框架进行有效的辅导谈话")

    ws1_steps = [
        "准备阶段（10分钟）：选择练习情境，准备辅导谈话",
        "角色扮演（15分钟）：两人一组，分别扮演辅导员和被辅导者",
        "观察反馈（10分钟）：使用观察表记录并给予反馈",
        "讨论总结（5分钟）：小组分享心得和改进点"
    ]

    for step in ws1_steps:
        p_step = doc.add_paragraph(style='List Number')
        p_step.paragraph_format.left_indent = Inches(0.3)
        p_step.paragraph_format.space_after = Pt(4)
        p_step.add_run(step)

    # Workshop 2: Feedback Practice
    p_ws2 = doc.add_paragraph()
    p_ws2.paragraph_format.space_before = Pt(16)
    run_ws2 = p_ws2.add_run("工作坊2：CAIR反馈练习")
    run_ws2.bold = True
    run_ws2.font.size = Pt(14)
    run_ws2.font.color.rgb = MEDIUM_BLUE

    p_ws2_desc = doc.add_paragraph()
    p_ws2_desc.add_run("练习目标：使用CAIR模型提供具体、有效的建设性反馈")

    ws2_steps = [
        "案例分析（10分钟）：分析给定的反馈案例",
        "CAIR撰写（10分钟）：使用CAIR框架撰写反馈脚本",
        "角色扮演（15分钟）：实践反馈谈话",
        "反思改进（5分钟）：根据反馈改进表达方式"
    ]

    for step in ws2_steps:
        p_step = doc.add_paragraph(style='List Number')
        p_step.paragraph_format.left_indent = Inches(0.3)
        p_step.paragraph_format.space_after = Pt(4)
        p_step.add_run(step)

    # Observation Form
    add_form_title(doc, "W-1", "观察记录表",
                   "填写说明：观察角色扮演，记录关键行为和反馈")

    obs_headers = ["观察维度", "具体表现", "改进建议"]
    obs_data = [
        ("开场建立 rapport", "", ""),
        ("使用CAIR/WHY框架", "", ""),
        ("倾听和提问", "", ""),
        ("情感关注", "", ""),
        ("达成共识/承诺", "", ""),
        ("整体效果（1-10分）", "", "")
    ]
    create_table_with_header(doc, obs_headers, obs_data, [1.8, 2.5, 2.5])

    # Role-play cards section
    doc.add_page_break()
    p_rp = doc.add_paragraph()
    p_rp.paragraph_format.space_before = Pt(0)
    p_rp.paragraph_format.space_after = Pt(12)
    run_rp = p_rp.add_run("角色扮演卡")
    run_rp.bold = True
    run_rp.font.size = Pt(18)
    run_rp.font.color.rgb = DARK_BLUE

    p_rp_desc = doc.add_paragraph()
    p_rp_desc.paragraph_format.space_after = Pt(12)
    p_rp_desc.add_run("以下角色扮演卡可裁剪使用，每张卡片设计为独立情境。")

    add_role_play_card(doc, "RP-1",
        "新员工小张入职两个月，第一次独立完成用户调研报告，但报告结构混乱，数据分析也比较浅。作为他的导师，你需要进行辅导。",
        ("导师角色", "运用辅导框架，帮助小张认识问题并找到改进方法"),
        ("小张角色", "可能会有防御心理，需要引导其开放接受反馈"),
        ["是否使用WHEN确认辅导时机", "是否运用HOW的步骤", "是否达成具体改进行动"]
    )

    add_role_play_card(doc, "RP-2",
        "团队成员老王最近一周连续迟到，工作效率明显下降。在部门会议上，他与同事发生了激烈争执。",
        ("经理角色", "需要了解情况，提供反馈并帮助老王改进"),
        ("老王角色", "可能有个人困难或对工作环境有不满"),
        ["是否及时关注问题", "是否采用CAIR反馈", "是否倾听并理解原因"]
    )

    add_role_play_card(doc, "RP-3",
        "技术骨干小李提出离职，表示收到了一家竞争对手的offer，薪资比现在高30%。作为他的直接经理，你需要进行一次留人谈话。",
        ("经理角色", "了解真实原因，探索可能的解决方案"),
        ("小李角色", "可能有多重考虑，不仅仅是薪资问题"),
        ["是否建立信任和开放氛围", "是否倾听真实原因", "是否讨论可能的解决方案"]
    )

    doc.add_page_break()

    # ========== AFTER-CLASS RESOURCES ==========
    doc.add_page_break()
    p_ac = doc.add_paragraph()
    p_ac.paragraph_format.space_before = Pt(0)
    p_ac.paragraph_format.space_after = Pt(12)
    run_ac = p_ac.add_run("课后资源")
    run_ac.bold = True
    run_ac.font.size = Pt(22)
    run_ac.font.color.rgb = DARK_BLUE

    add_horizontal_line(doc)

    # Self-reflection
    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.space_before = Pt(16)
    run_ref = p_ref.add_run("自我复盘问题")
    run_ref.bold = True
    run_ref.font.size = Pt(14)
    run_ref.font.color.rgb = MEDIUM_BLUE

    reflection_questions = [
        "在过去一周的管理实践中，你使用了哪些辅导和反馈技巧？效果如何？",
        "你遇到了哪些困难谈话？你是如何处理的？有什么可以改进的地方？",
        "你与团队成员的工作关系如何？有哪些关系需要重点关注和改善？",
        "你是否有效地授权了任务？授权过程中遇到了什么挑战？",
        "基于课程学习，你计划在哪些方面优先改进自己的管理方式？"
    ]

    for i, q in enumerate(reflection_questions):
        p_q = doc.add_paragraph(style='List Number')
        p_q.paragraph_format.left_indent = Inches(0.3)
        p_q.paragraph_format.space_after = Pt(8)
        p_q.add_run(q)

    # 90-day Action Plan
    p_90 = doc.add_paragraph()
    p_90.paragraph_format.space_before = Pt(16)
    run_90 = p_90.add_run("90天行动计划模板")
    run_90.bold = True
    run_90.font.size = Pt(14)
    run_90.font.color.rgb = MEDIUM_BLUE

    plan_headers = ["阶段", "时间", "重点行动", "预期成果", "检验方式"]
    plan_data = [
        ("第一阶段", "第1-30天", "建立基础\n• 每日辅导微习惯\n• 记录反馈案例", "", ""),
        ("第二阶段", "第31-60天", "技能提升\n• 处理复杂情境\n• 授权实践", "", ""),
        ("第三阶段", "第61-90天", "整合优化\n• MAP执行\n• 总结反思", "", "")
    ]
    create_table_with_header(doc, plan_headers, plan_data, [1, 1, 2.5, 1.5, 1])

    # Recommended Reading
    p_read = doc.add_paragraph()
    p_read.paragraph_format.space_before = Pt(16)
    run_read = p_read.add_run("推荐阅读")
    run_read.bold = True
    run_read.font.size = Pt(14)
    run_read.font.color.rgb = MEDIUM_BLUE

    reading_list = [
        ("《非暴力沟通》", "马歇尔·卢森堡"),
        ("《关键对话》", "科里·帕特森等"),
        ("《高效能人士的七个习惯》", "史蒂芬·柯维"),
        ("《教练的智慧》", "约翰·惠特默"),
        ("《横向领导力》", "罗杰·费希尔")
    ]

    for title, author in reading_list:
        p_book = doc.add_paragraph(style='List Bullet')
        p_book.paragraph_format.left_indent = Inches(0.3)
        p_book.paragraph_format.space_after = Pt(4)
        run_book = p_book.add_run(f"{title}")
        run_book.bold = True
        p_book.add_run(f" —— {author}")

    # Closing
    doc.add_page_break()
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_before = Pt(100)
    p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_close = p_close.add_run("致出发")
    run_close.bold = True
    run_close.font.size = Pt(24)
    run_close.font.color.rgb = DARK_BLUE

    p_quote = doc.add_paragraph()
    p_quote.paragraph_format.space_before = Pt(40)
    p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_quote = p_quote.add_run('"领导力不是关于职位，而是关于影响。\n当你开始用新的方式思考，你就开始用新的方式行动。\n当你开始用新的方式行动，你就开始创造新的结果。"')
    run_quote.italic = True
    run_quote.font.size = Pt(14)
    run_quote.font.color.rgb = MEDIUM_BLUE

    p_final = doc.add_paragraph()
    p_final.paragraph_format.space_before = Pt(60)
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = p_final.add_run("从今天，从现在，从这一刻开始。\n\nUYLP——释放你的领导潜能")
    run_final.bold = True
    run_final.font.size = Pt(16)
    run_final.font.color.rgb = DARK_GRAY

    p_copyright = doc.add_paragraph()
    p_copyright.paragraph_format.space_before = Pt(120)
    p_copyright.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cp = p_copyright.add_run("© 竞越人才发展咨询    本手册仅供UYLP课程学员使用")
    run_cp.font.size = Pt(10)
    run_cp.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save document
    output_path = "D:/2026年课程/竞越/释潜：UYLP释放你的领导潜能/完整课程包/04_学员手册/UYLP_学员手册_完整版.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_handbook()

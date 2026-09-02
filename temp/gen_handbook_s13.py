#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 附录三：访谈提纲模板
app3 = doc.add_paragraph()
r = app3.add_run('附录三：访谈提纲模板')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

doc.add_paragraph()
intt = doc.add_paragraph()
r = intt.add_run('访谈基本信息')
r.bold = True

t3 = doc.add_table(rows=5, cols=2)
t3.style = 'Table Grid'
for i, label in enumerate(['决策名称', '决策者', '访谈日期', '访谈者', '访谈时长']):
    cl = t3.cell(i, 0)
    set_cell_shading(cl, 'E8F0FC')
    r = cl.paragraphs[0].add_run(label)
    r.bold = True

doc.add_paragraph()
introt = doc.add_paragraph()
r = introt.add_run('访谈导入（约5分钟）')
r.bold = True
doc.add_paragraph('建立信任，说明访谈目的，不涉及具体判断对错的内容\n• 先聊一些相对轻松的内容\n• 等对方稍微放松下来，再切入正题')

dims = [
    ('第一维度：追因（约20-30分钟）', '核心问题：当时是什么信号让你觉得需要做决策，而不是按原计划走？\n\n追问要点：\n• 在那之前呢？（往前再退一步）\n• 不要满足于第一个触发点答案'),
    ('第二维度：权衡（约30-40分钟）', '核心问题：当时你能想到的选项有哪些，你是怎么排除掉其他选项的？\n\n追问要点：\n• 需要追问两到三轮\n• 具体化追问\n• 警惕事后合理化'),
    ('第三维度：未预见的假设（约15-20分钟）', '核心问题：如果当时那个前提不成立，你还会做同样的判断吗？\n\n追问要点：\n• 挖掘决策者无意识依赖的假设'),
]

for dimt, dimc in dims:
    doc.add_paragraph()
    dh = doc.add_paragraph()
    r = dh.add_run(dimt)
    r.bold = True
    doc.add_paragraph(dimc)

doc.add_paragraph()
hidt = doc.add_paragraph()
r = hidt.add_run('隐性判断挖掘（贯穿全程）')
r.bold = True
doc.add_paragraph('触发信号：决策者反复用我就是觉得的地方\n\n技法1-模式追问：把模糊感受放一起对比\n技法2-反事实追问：逼想象未发生情境\n技法3-角色互换追问：切换到教别人视角')

doc.add_paragraph()
endh = doc.add_paragraph()
r = endh.add_run('访谈收尾')
r.bold = True
doc.add_paragraph('• 感谢决策者的时间\n• 说明后续流程\n• 确认草稿确认方式')

doc.add_page_break()
print('S13: Tool template 3 done')
doc.save(r'D:\CC\temp\handbook_s13.docx')

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert 学员手册.md to docx and pdf."""
import os
import sys
from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from weasyprint import HTML, CSS

BASE = Path("D:/2026年课程/ai课2026整理/经营型企业大学构建：从培训部门到战略引擎/学员手册")
MD = BASE / "学员手册.md"
HTML_FILE = BASE / "学员手册.html"
DOCX = BASE / "学员手册.docx"
PDF = BASE / "学员手册.pdf"


def md_to_html_for_pdf(md_text: str) -> str:
    """Convert markdown to styled HTML for PDF."""
    html_body = markdown.markdown(
        md_text,
        extensions=["extra", "tables", "fenced_code", "toc"],
    )
    css = """
    @page { size: A4; margin: 2cm; }
    body { font-family: 'Noto Serif SC', serif; line-height: 1.8; color: #2c2c2c; }
    h1 { font-size: 28pt; color: #1a1a1a; border-bottom: 3px solid #c4a35a; padding-bottom: 8px; margin-top: 32pt; }
    h2 { font-size: 20pt; color: #1a1a1a; margin-top: 28pt; }
    h3 { font-size: 16pt; color: #5d3a1a; margin-top: 20pt; }
    h4 { font-size: 13pt; color: #8b7355; margin-top: 16pt; border-left: 3px solid #c4a35a; padding-left: 10px; }
    table { border-collapse: collapse; width: 100%; margin: 16pt 0; font-size: 10pt; }
    th { background: #1a1a1a; color: white; padding: 8pt; text-align: left; }
    td { border-bottom: 1px solid #e8e3db; padding: 8pt; vertical-align: top; }
    tr:nth-child(odd) td { background: #f7f4ef; }
    blockquote { border-left: 4px solid #a0522d; padding: 8pt 16pt; background: #f7f4ef; font-style: italic; margin: 16pt 0; }
    code { background: #f7f4ef; padding: 2pt 4pt; border-radius: 2pt; font-family: 'Source Code Pro', monospace; font-size: 9pt; }
    pre { background: #f7f4ef; padding: 12pt; border-radius: 3pt; overflow-x: auto; font-size: 9pt; }
    hr { border: none; border-top: 1px solid #e8e3db; margin: 24pt 0; }
    strong { color: #a0522d; }
    """
    return f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"><style>{css}</style></head>
<body>{html_body}</body></html>"""


def make_docx(md_text: str, out_path: Path):
    """Convert markdown to docx using python-docx with a custom parser."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Noto Serif SC"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        # Parse markdown table
        rows = [r for r in table_lines if r.strip().startswith("|")]
        if len(rows) < 2:
            table_lines = []
            return
        # First row is header, second row is separator
        header = [c.strip() for c in rows[0].strip("|").split("|")]
        data_rows = []
        for r in rows[2:]:
            cells = [c.strip() for c in r.strip("|").split("|")]
            data_rows.append(cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
        table.style = "Light Grid Accent 1"
        # Header
        for j, h in enumerate(header):
            cell = table.rows[0].cells[j]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(10)
        # Data
        for i_r, row in enumerate(data_rows):
            for j_c, c in enumerate(row[:len(header)]):
                cell = table.rows[i_r + 1].cells[j_c]
                cell.text = c
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        table_lines = []

    for line in lines:
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_lines.append(line)
            in_table = True
            continue
        else:
            if in_table:
                flush_table()
                in_table = False

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.style = doc.styles["Intense Quote"]
            run = p.add_run(line[2:].strip())
            run.italic = True
        elif line.strip() == "---":
            doc.add_paragraph("─" * 40)
        elif line.strip() == "":
            doc.add_paragraph()
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            content = line.strip()[2:]
            doc.add_paragraph(content, style="List Bullet")
        elif line.strip().startswith(tuple(f"{i}. " for i in range(1, 10))):
            content = line.strip().split(". ", 1)[1] if ". " in line.strip() else line.strip()
            doc.add_paragraph(content, style="List Number")
        else:
            p = doc.add_paragraph(line.strip())

    if in_table:
        flush_table()

    doc.save(str(out_path))
    print(f"Created: {out_path}")


def make_pdf(html_str: str, out_path: Path):
    HTML(string=html_str).write_pdf(str(out_path))
    print(f"Created: {out_path}")


def main():
    md_text = MD.read_text(encoding="utf-8")
    print(f"Reading: {MD} ({len(md_text)} chars)")

    # DOCX
    make_docx(md_text, DOCX)

    # PDF (from markdown)
    html_for_pdf = md_to_html_for_pdf(md_text)
    make_pdf(html_for_pdf, PDF)

    # Report sizes
    for f in [MD, HTML_FILE, DOCX, PDF]:
        if f.exists():
            print(f"  {f.name}: {f.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()

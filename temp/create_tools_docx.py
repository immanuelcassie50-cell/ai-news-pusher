# -*- coding: utf-8 -*-
"""
战略解码双螺旋引擎 - 工具表单全集
创建13个工具表单的Word文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ========== 工具函数 ==========

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), val)
            border.set(qn('w:sz'), '4')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
    tcPr.append(tcBorders)

def add_form_title(doc, form_name, module, date_str):
    """添加表单标题栏"""
    # 标题行
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(form_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # 信息行
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = info_para.add_run(f"适用模块：{module}    ")
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2 = info_para.add_run(f"日期：{date_str}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

def add_section_header(doc, text, color='1F497D'):
    """添加小节标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

def add_guide_text(doc, text):
    """添加填写指引"""
    para = doc.add_paragraph()
    run = para.add_run(f"📌 {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
    run.italic = True
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)

def add_example_text(doc, text):
    """添加示例文本"""
    para = doc.add_paragraph()
    run = para.add_run(f"示例：{text}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    para.paragraph_format.left_indent = Cm(0.5)

def create_table_with_header(doc, headers, rows, col_widths=None, header_color='1F497D'):
    """创建带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = str(cell_text)
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            # 斑马纹
            if row_idx % 2 == 1:
                set_cell_shading(row_cells[col_idx], 'F2F2F2')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

    return table

def add_reflection_section(doc, questions):
    """添加反思区"""
    add_section_header(doc, "反思区", 'C55A11')
    for q in questions:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.font.size = Pt(10)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(6)

def add_action_section(doc, actions):
    """添加行动计划区"""
    add_section_header(doc, "行动计划区", '538135')
    for action in actions:
        para = doc.add_paragraph()
        run = para.add_run(f"☐ {action}")
        run.font.size = Pt(10)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(4)

# ========== 13个工具表单 ==========

def create_tool_1(doc):
    """工具1：战略落地障碍诊断表"""
    add_form_title(doc, "战略落地障碍诊断表", "第一部分·第一节 | 五大障碍诊断", "2024年____月____日")

    add_guide_text(doc, "诊断说明：这不是客观测评，是你基于目前认知的主观判断。评分1-5分（1=基本不存在这个障碍，5=这是我们最严重的问题之一）")

    headers = ["障碍类型", "具体表现", "评分(1-5)", "主要体现"]
    rows = [
        ["障碍一：战略转型与短期业绩的矛盾", "战略要求投入但短期损益压力大；推战略的人被绩效指标\"困死\"了", "", ""],
        ["障碍二：缺少Believe，缺少All-in", "中层管理者和执行层没有真正相信这个战略能成", "", ""],
        ["障碍三：无法打赢攻坚战役", "跨部门项目推进缓慢；关键突破迟迟不来", "", ""],
        ["障碍四：战略与日常管理无法协同", "周会、月会和战略目标没有关系；部门KPI和公司战略方向不一致", "", ""],
        ["障碍五：资源投入错位", "钱和人没有向战略重点倾斜；新业务争不过老业务的资源", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(4), Cm(6), Cm(2), Cm(4)])

    doc.add_paragraph()
    add_section_header(doc, "诊断结果")

    para = doc.add_paragraph()
    run = para.add_run("我的前两个最高分障碍是：")
    run.bold = True
    run.font.size = Pt(10)

    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    headers2 = ["障碍", "具体问题", "评分", "对我的影响"]
    for i, h in enumerate(headers2):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], '1F497D')
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)

    add_reflection_section(doc, [
        "这两个障碍，对我的战略推进具体表现在哪些方面？",
        "如果只能先解决一个障碍，我会选择哪个？为什么？"
    ])

    add_action_section(doc, [
        "本周内与直接上级沟通我对障碍的判断，争取共识",
        "针对评分最高的障碍，初步思考一个可能的解决方向"
    ])

    doc.add_page_break()

def create_tool_2(doc):
    """工具2：引擎紧迫度评估表"""
    add_form_title(doc, "引擎紧迫度评估表", "第一部分·第四节 | 引擎选择判断", "2024年____月____日")

    add_guide_text(doc, "练习说明：根据你对自己公司/事业部现状的了解，判断下面两组问题。这是帮助你判断哪个引擎对你最迫切")

    add_section_header(doc, "引擎一（必赢战役）的紧迫度")
    add_guide_text(doc, "回答2个以上\"是\"，说明引擎一是你的优先课题")

    headers = ["问题", "是 / 部分是 / 否", "备注"]
    rows = [
        ["公司战略里有1-3个关键突破点，是决定战略成败的\"关键战役\"", "", ""],
        ["这些突破点目前没有专属的跨部门团队、专属资源、专属负责人", "", ""],
        ["这些突破点的进展，最高管理层目前没有每2-4周进行一次系统回顾", "", ""],
        ["这些突破点当前推进缓慢，主要原因是跨部门协同不够、资源优先级不够", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(9), Cm(3), Cm(4)])

    doc.add_paragraph()
    add_section_header(doc, "引擎二（部门绩效协同）的紧迫度")
    add_guide_text(doc, "回答2个以上\"是\"，说明引擎二是你的优先课题")

    rows2 = [
        ["各部门的年度目标，与公司战略目标之间的关联，现在说不清楚", "", ""],
        ["各部门负责人的绩效考核指标，主要是财务和运营效率，缺少战略转型维度", "", ""],
        ["部门之间经常出现\"各干各的\"、协同成本高的问题", "", ""],
        ["公司战略发布后，各部门没有系统地更新过自己的工作重心和绩效指标", "", ""],
    ]
    create_table_with_header(doc, headers, rows2, col_widths=[Cm(9), Cm(3), Cm(4)])

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("我的判断：")
    run.bold = True

    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    data = [
        ["引擎一紧迫度评分（1-5）", ""],
        ["引擎二紧迫度评分（1-5）", ""],
        ["对我最迫切的是引擎_____，因为", ""]
    ]
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_text
            for p in table2.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if i == 2 and j == 0:
                        r.bold = True

    add_reflection_section(doc, [
        "我的判断依据是什么？是主观感受还是客观事实？",
        "如果两个引擎都紧迫，我应该如何分配精力？"
    ])

    add_action_section(doc, [
        "根据判断，确定我在本次课程中最需要重点学习的内容",
        "课后第一周，与战略负责人确认引擎选择的判断"
    ])

    doc.add_page_break()

def create_tool_3(doc):
    """工具3：一致力诊断表"""
    add_form_title(doc, "一致力诊断表", "第二部分·下·维度四 | 战略组织保障", "2024年____月____日")

    add_guide_text(doc, "诊断说明：基于你所在的公司/事业部，对照五个维度进行评估。评分1-5分（1=严重不对齐，3=部分对齐，5=高度对齐）")

    headers = ["维度", "评分(1-5)", "对齐的地方（做得好的）", "最大的对齐差距", "如果改善，影响会是"]
    rows = [
        ["架构：组织结构是否支撑战略", "", "", "", ""],
        ["流程：关键业务流程是否支持战略行为", "", "", "", ""],
        ["人才：战略所需的关键岗位，有没有具备能力的人", "", "", "", ""],
        ["绩效：激励和考核是否驱动战略行为", "", "", "", ""],
        ["价值观：组织文化是否支持战略所需的行为和决策方式", "", "", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(2), Cm(3.5), Cm(4), Cm(3.5)])

    add_reflection_section(doc, [
        "我评分最低的维度是哪个？这个不对齐，对战略推进具体体现为什么影响？",
        "一致力模型的核心，是\"战略拉力 vs 组织惯性\"。我所在组织的最大\"重力\"来自哪个维度？",
        "如果我们明年只能在一个维度上做出改变，我会选择哪个？为什么？"
    ])

    add_action_section(doc, [
        "选择一个最需要改善的维度，初步思考改善方向",
        "找到在这个维度上有话语权的关键人物，列入课后沟通计划"
    ])

    doc.add_page_break()

def create_tool_4(doc):
    """工具4：必赢战役识别表"""
    add_form_title(doc, "必赢战役识别表", "第二部分·上·维度一 | 聚焦关键领域", "2024年____月____日")

    add_guide_text(doc, "必赢战役的四个特征：①杠杆性（打赢了对全局决定性）②跨越性（现有能力无法轻松实现）③跨部门性（需要跨团队协同）④可界定性（赢的标准是可量化的）")

    add_section_header(doc, "战略背景")
    para = doc.add_paragraph()
    run = para.add_run("公司/事业部的战略目标（用一句话描述）：_______________________________________________")
    run.font.size = Pt(10)

    add_section_header(doc, "战略突破点识别")
    add_guide_text(doc, "从战略目标出发，问一个问题：\"如果这个目标3年后实现了，我们必须在哪2-3个关键领域取得突破？\"")

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "关键突破领域"
    table.rows[0].cells[1].text = "为什么这个突破对战略是决定性的"
    set_cell_shading(table.rows[0].cells[0], '1F497D')
    set_cell_shading(table.rows[0].cells[1], '1F497D')
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for p in table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i in range(1, 4):
        table.rows[i].cells[0].text = f"领域{i}："
        table.rows[i].cells[1].text = ""

    doc.add_paragraph()
    add_section_header(doc, "必赢战役候选")

    headers = ["候选", "战役名称", "是否符合四个特征", "风险点"]
    rows = [
        ["战役候选一", "", "□杠杆性 □跨越性 □跨部门性 □可界定性\n全部满足才能算必赢战役", ""],
        ["战役候选二", "", "□杠杆性 □跨越性 □跨部门性 □可界定性", ""],
        ["战役候选三（可选）", "", "□杠杆性 □跨越性 □跨部门性 □可界定性", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(2), Cm(4), Cm(6), Cm(4)])

    add_reflection_section(doc, [
        "我的战役候选，真正符合\"跨越性\"的是哪一个？现有能力是否真的无法轻松实现？",
        "如果识别出了超过3个候选战役，我应该如何筛选？"
    ])

    add_action_section(doc, [
        "确定我的1-2个必赢战役候选，列入下一张表单详细定义",
        "思考每个战役候选需要哪些部门协同，准备在战役定义时明确"
    ])

    doc.add_page_break()

def create_tool_5(doc):
    """工具5：战役定义标准表"""
    add_form_title(doc, "战役定义标准表", "第二部分·上·维度一 | 必赢战役定义", "2024年____月____日")

    add_guide_text(doc, "一个正式的必赢战役，需要用这个结构来定义。战役名称要简洁、有力、清楚；赢的标准必须量化、有时限")

    headers = ["战役", "定义内容"]
    rows = [
        ["战役名称", ""],
        ["赢的标准\n（量化+时限）", ""],
        ["战役负责人\n（一个具体的人）", ""],
        ["战役时限\n（开始-结束）", ""],
        ["战略意义\n（为什么这场战役对整体战略是决定性的）", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(13)])

    doc.add_paragraph()
    add_section_header(doc, "好与不好的战役名称对比")

    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Table Grid'
    comparison = [
        ["好的战役名称", "不好的战役名称"],
        ["解决方案销售能力建设战役", "提升竞争力"],
        ["核心标杆客户深度渗透战役", "加强团队建设"],
    ]
    for i, row_data in enumerate(comparison):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text
            if i == 0:
                set_cell_shading(table2.rows[i].cells[j], '1F497D')
                for p in table2.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for p in table2.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_reflection_section(doc, [
        "我的战役名称是否能让人一眼看出\"要打赢什么\"？",
        "赢的标准是否能用一句话说清楚\"打赢是什么样子\"？"
    ])

    add_action_section(doc, [
        "用上面的格式，正式定义我的必赢战役",
        "与战役负责人（如果已确定）确认战役定义是否准确"
    ])

    doc.add_page_break()

def create_tool_6(doc):
    """工具6：战役拆解表"""
    add_form_title(doc, "战役拆解表", "第二部分·上·维度二 | 梳理战术策略", "2024年____月____日")

    add_guide_text(doc, "战役拆解的层级：必赢战役（顶层目标）→ 战术策略（3-5个方向）→ 关键战斗（具体行动：动词+时限+成果标志）")

    para = doc.add_paragraph()
    run = para.add_run("战役名称：______________________    赢的标准：______________________    负责人：______________________")
    run.font.size = Pt(10)

    add_section_header(doc, "战术策略一：______________________")

    headers = ["关键战斗", "完成时间", "成果标志", "负责人"]
    rows = [
        ["关键战斗1", "", "", ""],
        ["关键战斗2", "", "", ""],
        ["关键战斗3", "", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(5), Cm(2.5), Cm(4.5), Cm(2)])

    doc.add_paragraph()
    add_section_header(doc, "战术策略二：______________________")
    create_table_with_header(doc, headers, rows, col_widths=[Cm(5), Cm(2.5), Cm(4.5), Cm(2)])

    doc.add_paragraph()
    add_section_header(doc, "战术策略三（可选）：______________________")
    create_table_with_header(doc, headers, rows, col_widths=[Cm(5), Cm(2.5), Cm(4.5), Cm(2)])

    doc.add_paragraph()
    add_section_header(doc, "跨部门协同需求")
    headers2 = ["部门", "需要提供的资源或配合", "时间要求", "确认状态"]
    rows2 = [
        ["", "", "", "□已确认 □待确认"],
        ["", "", "", "□已确认 □待确认"],
    ]
    create_table_with_header(doc, headers2, rows2, col_widths=[Cm(3), Cm(5), Cm(3), Cm(3)])

    doc.add_paragraph()
    add_section_header(doc, "关键里程碑")
    headers3 = ["时间节点", "里程碑描述", "判断标准"]
    rows3 = [
        ["", "", ""],
        ["", "", ""],
        ["", "", ""],
    ]
    create_table_with_header(doc, headers3, rows3, col_widths=[Cm(3), Cm(5), Cm(6)])

    add_reflection_section(doc, [
        "每个关键战斗是否都有具体动词（而非\"加强\"\"提升\"\"优化\"）？",
        "每个关键战斗是否有明确时限和成果标志？",
        "战术策略之间是否相互独立，没有大量重叠？"
    ])

    add_action_section(doc, [
        "完善每个战术策略下的关键战斗，确保质量达标",
        "与跨部门协同方确认资源配合承诺"
    ])

    doc.add_page_break()

def create_tool_7(doc):
    """工具7：战役OKR设计表"""
    add_form_title(doc, "战役OKR设计表", "第二部分·下·维度三 | OKR与KPI战略应用", "2024年____月____日")

    add_guide_text(doc, "OKR的精髓不是\"完成KR\"，而是\"在追求O的过程中，不断学习什么有效、什么无效，并快速调整\"。好的O应该有雄心格局，KR是结果而非任务")

    add_section_header(doc, "O（Objective）")
    para = doc.add_paragraph()
    run = para.add_run("一句话，描述你想要达到的雄心状态——有格局，能激励，不能直接用数字测量")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    table = doc.add_table(rows=3, cols=1)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "O："
    set_cell_shading(table.rows[0].cells[0], '1F497D')
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    table.rows[1].cells[0].text = ""
    table.rows[2].cells[0].text = ""

    doc.add_paragraph()
    add_section_header(doc, "KR（Key Results）")
    add_guide_text(doc, "KR格式：[可测量的结果] + [时间节点]。好的KR应该能通过：①是结果而非任务 ②可测量 ③实现KR后O令人信服 ④KR之间相互独立")

    headers = ["KR", "内容", "自检"]
    rows = [
        ["KR1", "", "□是结果 □可测量 □O令人信服 □相互独立"],
        ["KR2", "", "□是结果 □可测量 □O令人信服 □相互独立"],
        ["KR3（可选）", "", "□是结果 □可测量 □O令人信服 □相互独立"],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(2), Cm(10), Cm(4)])

    add_section_header(doc, "KR常见错误检查")
    para = doc.add_paragraph()
    run = para.add_run("❌ 把KPI包装成OKR：\"Q4销售额达到5000万\"是KPI，不是O\n"
                       "❌ KR写成了任务清单：\"完成解决方案培训\"不是KR，是任务\n"
                       "✅ 好的KR是结果：\"60%的大客户销售完成解决方案认证（Q4）\"")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_reflection_section(doc, [
        "我的O是否有足够的雄心格局？还是只是一个可以测量的数字？",
        "我的KR里是否还在用\"完成\"\"推进\"\"开展\"等行动词？",
        "如果我每次都能100%完成KR，说明什么？"
    ])

    add_action_section(doc, [
        "重新审视OKR，修改不符合要求的地方",
        "与战役指挥官（如果已确定）确认OKR是否合适"
    ])

    doc.add_page_break()

def create_tool_8(doc):
    """工具8：战役执行机制设计表"""
    add_form_title(doc, "战役执行机制设计表", "第二部分·下·维度三 | 内生机制设计", "2024年____月____日")

    add_guide_text(doc, "内生机制的五要素：①战役指挥官 ②战役指挥部（固定节奏） ③可视化看板 ④快速决策通道 ⑤核心战役团队")

    headers = ["机制要素", "设计内容", "我们的现状", "差距分析"]
    rows = [
        ["①战役指挥官\n（强授权、单一责任人）", "姓名：__________\n跨部门调动权：□有 □无\n向谁汇报：__________", "", ""],
        ["②战役指挥部\n（固定节奏）", "战役级回顾频率：□每周 □每2周 □每月\n单次时长：______分钟\n成员范围：__________", "", ""],
        ["③可视化看板\n（进展透明）", "形式：□物理白板 □在线看板\n更新频率：__________\n放置位置：__________", "", ""],
        ["④快速决策通道\n（阻碍24-48小时内拍板）", "什么问题找谁在多长时间内答复：\n__________\n战役指挥官的直通最高管理层通道：□有 □无", "", ""],
        ["⑤核心战役团队\n（专属+跨部门+保护时间）", "成员人数：______\n来自哪些部门：__________\n每人保护时间（%）：__________", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(5.5), Cm(3.5), Cm(4)])

    add_section_header(doc, "战役指挥官授权清单（建议）")
    para = doc.add_paragraph()
    run = para.add_run("□ 有权要求各部门提供承诺的资源\n"
                       "□ 有权在战役推进受阻时直接向最高管理层上报\n"
                       "□ 有权在24-48小时内获得最高管理层的决策反馈\n"
                       "□ 每周直接向最高管理层汇报战役进展（不是书面报告，是面对面）")
    run.font.size = Pt(9)

    add_reflection_section(doc, [
        "我们的战役执行目前依赖什么？是\"外生推动\"（领导持续催促）还是\"内生机制\"（系统自动运转）？",
        "五个要素中，哪个是我们当前最大的漏洞？",
        "如果没有固定节奏的战役回顾会，战役进展会如何？"
    ])

    add_action_section(doc, [
        "确定战役指挥官人选，与他/她确认授权范围",
        "确定战役指挥部的固定节奏，在日历上锁定这段时间",
        "建立可视化看板，明确更新责任"
    ])

    doc.add_page_break()

def create_tool_9(doc):
    """工具9：BSC战略图绘制表"""
    add_form_title(doc, "BSC战略图绘制表", "第三部分·第一节 | BSC战略图绘制", "2024年____月____日")

    add_guide_text(doc, "BSC核心逻辑（因果链）：学习与成长（基础能力）→ 内部流程（能力运营）→ 客户（价值创造）→ 财务（最终结果）。战略主题：______________________")

    add_section_header(doc, "学习与成长层面（基础能力）")
    headers = ["战略目标", "衡量指标", "2024年目标值"]
    rows = [
        ["L1：", "", ""],
        ["L2：", "", ""],
        ["L3：", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(4), Cm(4), Cm(3)])

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("         ↓ 使能")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    add_section_header(doc, "内部流程层面（关键流程）")
    rows2 = [
        ["I1：", "", ""],
        ["I2：", "", ""],
        ["I3：", "", ""],
    ]
    create_table_with_header(doc, headers, rows2, col_widths=[Cm(4), Cm(4), Cm(3)])

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("         ↓ 支撑")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    add_section_header(doc, "客户层面（客户价值）")
    rows3 = [
        ["C1：", "", ""],
        ["C2：", "", ""],
    ]
    create_table_with_header(doc, headers, rows3, col_widths=[Cm(4), Cm(4), Cm(3)])

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("         ↓ 驱动")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    add_section_header(doc, "财务层面（最终结果）")
    rows4 = [
        ["F1：", "", ""],
        ["F2：", "", ""],
    ]
    create_table_with_header(doc, headers, rows4, col_widths=[Cm(4), Cm(4), Cm(3)])

    doc.add_paragraph()
    add_section_header(doc, "因果链梳理")
    para = doc.add_paragraph()
    run = para.add_run("L__ → I__ → C__ → F__：_____________________________________________\n"
                       "L__ → I__ → C__ → F__：_____________________________________________\n"
                       "L__ → I__ → C__ → F__：_____________________________________________")
    run.font.size = Pt(9)

    add_reflection_section(doc, [
        "我的因果链最薄弱的地方是哪一步（哪个箭头背后的逻辑最不确定）？",
        "我的学习层和流程层之间，有没有明显缺失的基础能力？",
        "财务目标的驱动因素，是否都能追溯到客户层面的价值主张？"
    ])

    add_action_section(doc, [
        "完善BSC战略图的每个空格，确保因果链清晰",
        "与上级或战略负责人确认战略图的逻辑是否正确",
        "把战略图中的战略举措转化为部门行动"
    ])

    doc.add_page_break()

def create_tool_10(doc):
    """工具10：部门绩效指标设计表"""
    add_form_title(doc, "部门绩效指标设计表", "第三部分·第二节 | 部门绩效设计", "2024年____月____日")

    add_guide_text(doc, "部门绩效设计三步骤：①从战略图识别本部门最相关的战略目标 ②翻译成部门层面的绩效指标 ③确保战略与运营平衡（战略指标30-50%，运营指标50-70%）")

    add_section_header(doc, "基本信息")
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    info = [
        ["部门名称", ""],
        ["对应的战略图主要层面", "□学习与成长 □内部流程 □客户 □财务"]
    ]
    for i, row_data in enumerate(info):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    add_section_header(doc, "战略级绩效指标设计")
    headers = ["指标名称", "来自战略图哪个目标", "衡量方式", "2024年目标值", "建议权重"]
    rows = [
        ["", "", "", "", ""],
        ["", "", "", "", ""],
        ["", "", "", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(3), Cm(3), Cm(2.5), Cm(1.5)])

    doc.add_paragraph()
    add_section_header(doc, "运营级绩效指标（日常运营基线）")
    headers2 = ["指标名称", "衡量方式", "基线目标", "建议权重"]
    rows2 = [
        ["", "", "", ""],
        ["", "", "", ""],
    ]
    create_table_with_header(doc, headers2, rows2, col_widths=[Cm(3), Cm(3), Cm(2.5), Cm(1.5)])

    doc.add_paragraph()
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = "战略级指标总权重：___%"
    table3.rows[0].cells[1].text = "运营级指标总权重：___%"
    for p in table3.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
            r.bold = True
    for p in table3.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
            r.bold = True

    add_reflection_section(doc, [
        "我的战略级指标里，最难衡量的是哪一个？可能的衡量方法是什么？",
        "现有绩效体系中，与战略最不对齐的1个指标是什么？",
        "指标是否太多了（超过7个）？如果是，应该精简哪些？"
    ])

    add_action_section(doc, [
        "与HRBP或绩效负责人确认绩效指标设计的可行性",
        "识别需要新增或修改的绩效指标，列入沟通计划"
    ])

    doc.add_page_break()

def create_tool_11(doc):
    """工具11：关键岗位绩效设计表"""
    add_form_title(doc, "关键岗位绩效设计表", "第三部分·第三节 | 关键岗位绩效设计", "2024年____月____日")

    add_guide_text(doc, "关键岗位是\"这个岗位的行为，对战略的成败有直接影响\"的岗位。设计原则：①个人绩效可追溯到部门绩效和战略图 ②指标数量控制在3-5个 ③结果指标+行为指标组合")

    add_section_header(doc, "岗位基本信息")
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    info = [
        ["岗位名称", ""],
        ["所属部门", ""],
        ["对应的部门绩效指标", ""]
    ]
    for i, row_data in enumerate(info):
        for j, text in enumerate(row_data):
            table.rows[i].cells[j].text = text
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    add_section_header(doc, "关键绩效指标设计")
    headers = ["指标类型", "指标名称", "衡量方式", "目标值", "权重"]
    rows = [
        ["结果指标", "", "", "", ""],
        ["结果指标", "", "", "", ""],
        ["结果指标", "", "", "", ""],
        ["行为指标", "", "", "", ""],
        ["行为指标", "", "", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(2), Cm(3.5), Cm(3), Cm(2), Cm(1.5)])

    doc.add_paragraph()
    para = doc.add_paragraph()
    run = para.add_run("说明：转型期行为指标权重可以适当提高，因为行为指标比结果指标更早反映\"是否在往正确方向走\"")
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_reflection_section(doc, [
        "这个岗位的指标，是否能追溯到战略图的某个具体目标？",
        "指标之间是否有重叠或重复？",
        "行为指标是否真正对应了推动战略转型所需的关键行为？"
    ])

    add_action_section(doc, [
        "与岗位本人沟通绩效指标设计，争取认同",
        "确定每个指标的衡量方式，确保可操作性"
    ])

    doc.add_page_break()

def create_tool_12(doc):
    """工具12：战略推动行动计划表"""
    add_form_title(doc, "战略推动行动计划表", "第三部分·综合收尾 | 战略推动整合", "2024年____月____日")

    add_guide_text(doc, "这是你离开课程后回到公司的战略推动路线图。建议在回到公司的第一周内，和直接上级或团队把这份初稿过一遍")

    add_section_header(doc, "第一部分：引擎一——必赢战役")

    headers = ["项目", "战役一", "战役二"]
    rows = [
        ["战役名称", "", ""],
        ["赢的标准", "", ""],
        ["战役指挥官", "", ""],
        ["前三个关键战斗", "①\n②\n③", "①\n②\n③"],
        ["战役OKR的O", "", ""],
        ["现有执行机制的最大缺口", "", ""],
        ["一致力诊断：最需解决的对齐差距", "", ""],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3.5), Cm(5.5), Cm(5.5)])

    doc.add_paragraph()
    add_section_header(doc, "第二部分：引擎二——部门绩效协同")

    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    data = [
        ["我绘制的BSC战略图的战略主题", ""],
        ["学习层最关键的1个目标", ""],
        ["流程层最关键的1个目标", ""],
        ["部门核心战略级绩效指标（3个）", "① ② ③"],
        ["部门需要更新/新增的绩效指标", ""],
        ["现有绩效体系中，与战略最不对齐的1个指标", ""],
    ]
    for i, row_data in enumerate(data):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text
            if j == 0:
                set_cell_shading(table2.rows[i].cells[j], 'F2F2F2')
            for p in table2.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()
    add_section_header(doc, "第三部分：回去之后的第一件事")

    table3 = doc.add_table(rows=3, cols=1)
    table3.style = 'Table Grid'
    table3.rows[0].cells[0].text = "回到公司的第一周，我要做的最重要的一件事是："
    table3.rows[1].cells[0].text = ""
    table3.rows[2].cells[0].text = ""
    set_cell_shading(table3.rows[0].cells[0], '1F497D')
    for p in table3.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)

    doc.add_paragraph()
    headers3 = ["我要找的人是", "我要用的工具是", "预计完成时间"]
    rows3 = [["", "", ""], ["", "", ""]]
    create_table_with_header(doc, headers3, rows3, col_widths=[Cm(4), Cm(4), Cm(3)])

    add_action_section(doc, [
        "把这份计划存入可执行的地方（不要让它停在培训文件夹里）",
        "回到公司的第一周内，与直接上级或团队过一遍这份计划",
        "设定一个90天后回顾的时间节点"
    ])

    doc.add_page_break()

def create_tool_13(doc):
    """工具13：课后90天行动计划表"""
    add_form_title(doc, "课后90天行动计划表", "课后持续 | 90天落地计划", "2024年____月____日")

    add_guide_text(doc, "战略解码不是一个项目，而是一种持续的组织能力。这份90天计划是让你把课程里的工具真正用到工作中的路线图")

    add_section_header(doc, "第1-30天：夯实基础（必赢战役识别与定义）")

    headers = ["行动项", "具体动作", "责任人", "完成时限", "状态"]
    rows = [
        ["完成必赢战役识别", "基于战略目标，识别2-3个必赢战役候选", "", "第2周", "☐ 未开始"],
        ["确定战役定义", "用战役定义标准表，正式定义每个战役", "", "第3周", "☐ 未开始"],
        ["确认战役指挥官", "与战役指挥官确认授权和责任", "", "第3周", "☐ 未开始"],
        ["启动战役拆解", "完成至少一个战役的完整拆解", "", "第4周", "☐ 未开始"],
    ]
    create_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(5), Cm(2), Cm(2), Cm(1.5)])

    doc.add_paragraph()
    add_section_header(doc, "第31-60天：建立机制（内生机制与BSC战略图）")

    rows2 = [
        ["设计战役执行机制", "完成战役执行机制设计五要素", "", "第5周", "☐ 未开始"],
        ["绘制BSC战略图", "完成事业部/部门级BSC战略图初稿", "", "第6周", "☐ 未开始"],
        ["部门绩效指标梳理", "从战略图推导部门绩效指标", "", "第7周", "☐ 未开始"],
        ["关键岗位绩效对接", "完成关键岗位绩效指标设计", "", "第8周", "☐ 未开始"],
    ]
    create_table_with_header(doc, headers, rows2, col_widths=[Cm(3), Cm(5), Cm(2), Cm(2), Cm(1.5)])

    doc.add_paragraph()
    add_section_header(doc, "第61-90天：整合对齐（组织保障与战略合力）")

    rows3 = [
        ["一致力诊断回顾", "评估五个维度的对齐度变化", "", "第10周", "☐ 未开始"],
        ["战役中期回顾", "检视战役进展，调整战术", "", "第10周", "☐ 未开始"],
        ["战略推动整合", "整合双螺旋引擎完整工具到日常管理", "", "第12周", "☐ 未开始"],
        ["90天复盘", "回顾90天成果，制定下一阶段计划", "", "第12周", "☐ 未开始"],
    ]
    create_table_with_header(doc, headers, rows3, col_widths=[Cm(3), Cm(5), Cm(2), Cm(2), Cm(1.5)])

    doc.add_paragraph()
    add_section_header(doc, "90天关键检视点")
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    checkpoints = [
        ["检视时间", "检视内容", "需要支持"],
        ["第30天", "战役是否已经清晰定义并启动？", ""],
        ["第60天", "内生机制是否在运转？战略图是否在指导部门工作？", ""],
        ["第90天", "战略是否在日常运营中更可见了？", ""],
    ]
    for i, row_data in enumerate(checkpoints):
        for j, text in enumerate(row_data):
            table2.rows[i].cells[j].text = text
            if i == 0:
                set_cell_shading(table2.rows[i].cells[j], '1F497D')
                for p in table2.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(10)
            for p in table2.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_reflection_section(doc, [
        "90天后，我希望看到的最大改变是什么？",
        "在这90天里，最大的障碍可能是什么？我如何克服？",
        "谁是我最重要的支持者？我需要他/她什么样的帮助？"
    ])

    add_action_section(doc, [
        "把90天计划存入你的工作系统，设定检视提醒",
        "找一个\"问责伙伴\"（可以是同事或下属），每两周相互检视进展",
        "在第30天、第60天、第90天，固定做这三件事的回顾"
    ])

# ========== 主程序 ==========

def main():
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # 文档标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("战略解码·双螺旋引擎")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("工具表单全集")
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = intro.add_run("让大象跳舞 | 课程配套工具手册")
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 目录
    toc_para = doc.add_paragraph()
    run_toc = toc_para.add_run("工具表单目录")
    run_toc.bold = True
    run_toc.font.size = Pt(14)
    run_toc.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    toc_list = [
        "工具1  战略落地障碍诊断表",
        "工具2  引擎紧迫度评估表",
        "工具3  一致力诊断表",
        "工具4  必赢战役识别表",
        "工具5  战役定义标准表",
        "工具6  战役拆解表",
        "工具7  战役OKR设计表",
        "工具8  战役执行机制设计表",
        "工具9  BSC战略图绘制表",
        "工具10 部门绩效指标设计表",
        "工具11 关键岗位绩效设计表",
        "工具12 战略推动行动计划表",
        "工具13 课后90天行动计划表",
    ]

    for item in toc_list:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # 创建13个工具表单
    create_tool_1(doc)
    create_tool_2(doc)
    create_tool_3(doc)
    create_tool_4(doc)
    create_tool_5(doc)
    create_tool_6(doc)
    create_tool_7(doc)
    create_tool_8(doc)
    create_tool_9(doc)
    create_tool_10(doc)
    create_tool_11(doc)
    create_tool_12(doc)
    create_tool_13(doc)

    # 保存文档
    output_path = "D:/新课开发/战略和领导力/战略解码双螺旋引擎让大象跳舞/完整课程包/007-工具表单/001-工具表单全集.docx"
    doc.save(output_path)
    print(f"文档已生成：{output_path}")

if __name__ == "__main__":
    main()

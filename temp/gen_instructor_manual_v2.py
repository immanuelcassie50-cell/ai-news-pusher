"""
gen_instructor_manual_v2.py - 扩充版讲师手册,目标 120-160 页
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r"D:\2026年课程\竞越\教练技术：新的管理方式\完整课程包\14_Word手册"

PURPLE = RGBColor(0x5B, 0x3A, 0x8C)
ORANGE = RGBColor(0xF2, 0xA0, 0x3D)
INK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x9A, 0x98, 0x90)
TERRACOTTA = RGBColor(0xB8, 0x5C, 0x3E)
GOLD = RGBColor(0xC9, 0xA9, 0x61)
RED = RGBColor(0xC0, 0x39, 0x2B)


def set_zh_font(run, font_name="思源黑体"):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")


def add_body_no_indent(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")


def add_quote(doc, text, source=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run('"' + text + '"')
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run("—— " + source)
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = GREY
        set_zh_font(r2, "思源黑体")


def add_script(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("【讲师剧本】" + text)
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源黑体")


def add_action(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:left")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "24")
    bdr.set(qn("w:space"), "8")
    bdr.set(qn("w:color"), "F2A03D")
    pBdr.append(bdr)
    pPr.append(pBdr)
    r = p.add_run("【动作】 " + text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源黑体")


def add_warning(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:left")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "24")
    bdr.set(qn("w:space"), "8")
    bdr.set(qn("w:color"), "C0392B")
    pBdr.append(bdr)
    pPr.append(pBdr)
    r = p.add_run("【应急】 " + text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RED
    set_zh_font(r, "思源黑体")


def add_case(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("【案例】" + title)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TERRACOTTA
    set_zh_font(r, "思源宋体")
    for line in body:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run("• " + line)
        r.font.size = Pt(11)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


def add_faq(doc, q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Q: " + q)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源黑体")
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.74)
    p2.paragraph_format.line_spacing = 1.5
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run("A: " + a)
    r2.font.size = Pt(11)
    r2.font.color.rgb = INK
    set_zh_font(r2, "思源黑体")


def add_table_simple(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_zh_font(r, "思源黑体")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2A03D")
        tcPr.append(shd)
    for i, row in enumerate(rows):
        for j, cell_data in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_data) if cell_data else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = INK
                    set_zh_font(r, "思源黑体")


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._element.append(br)


def add_cover(doc, title, subtitle, kicker):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(48)
    r.font.bold = True
    r.font.color.rgb = PURPLE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(subtitle)
    r.font.size = Pt(28)
    r.font.color.rgb = ORANGE
    set_zh_font(r, "思源宋体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(kicker)
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    r = p.add_run("罗宏伟 · 内部讲师手册")
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    set_zh_font(r, "思源黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("竞越 × 教练技术 课程组 · 2026")
    r.font.size = Pt(12)
    r.font.color.rgb = GREY
    set_zh_font(r, "思源黑体")


def add_toc(doc):
    add_h1(doc, "目录")
    items = [
        "讲师手册使用说明",
        "第一部分 课程总览",
        "第二部分 时间控制表",
        "第三部分 训前准备清单",
        "第四部分 Day 1 完整剧本",
        "    模块一 AI 时代的管理坍塌",
        "    模块二 教练 vs 传统管理者",
        "    模块三 3C 信任 + 4 模式员工",
        "    第一天复盘",
        "第五部分 Day 2 完整剧本",
        "    模块四 GROW+ 对话框架",
        "    模块五 七层倾听阶梯",
        "    模块六 30 天落地实践",
        "    第二天收尾",
        "第六部分 案例讲解要点",
        "第七部分 应急方案",
        "第八部分 评估标准",
        "第九部分 训后跟进指南",
        "第十部分 讲师自评表",
        "第十一部分 学员互动剧本(高频问答)",
        "第十二部分 8 个案例的完整剧本",
        "附录:6 大金句卡片",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.6
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(12)
        r.font.color.rgb = INK
        set_zh_font(r, "思源黑体")


def add_lecture_block(doc, duration, topic, script_lines, actions, warns=None):
    add_h3(doc, f"⏱ {duration} · {topic}")
    if script_lines:
        for line in script_lines:
            add_script(doc, line)
    if actions:
        for a in actions:
            add_action(doc, a)
    if warns:
        for w in warns:
            add_warning(doc, w)


# ============================================================
# 完整剧本
# ============================================================
def build_instructor_manual():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "思源黑体"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "思源黑体")
    rFonts.set(qn("w:ascii"), "思源黑体")
    rFonts.set(qn("w:hAnsi"), "思源黑体")

    add_cover(doc, "教练技术", "讲师手册", "AI 时代的教练领导力 · 内部讲师专用")
    add_page_break(doc)
    add_toc(doc)
    add_page_break(doc)

    # 讲师手册使用说明
    add_h1(doc, "讲师手册使用说明")
    add_body(doc, "本手册是《教练技术:新的管理方式》课程的内部讲师专用文档。")
    add_body(doc, "完整剧本按 2 天 × 7 小时设计,共 14 段授课模块,每段标注时间、剧本、动作。")
    add_body(doc, "使用建议:")
    for s in [
        "1. 训前 1 周:读完剧本,标记你的卡点",
        "2. 训前 1 天:彩排 1 次,带计时器",
        "3. 训前 1 小时:检查设备、教室、学员名单",
        "4. 每天开课前:翻一遍当天剧本",
        "5. 训后当天:写讲师自评表",
    ]:
        add_body_no_indent(doc, s)
    add_quote(doc, "延迟摄影——和时间做朋友。", "—— 第一次讲,允许自己不完美")

    add_h2(doc, "讲师手册符号说明")
    add_table_simple(doc, ["符号", "含义"],
                    [
                        ["⏱", "时间(分钟)"],
                        ["【讲师剧本】", "逐字稿"],
                        ["【动作】", "讲师的具体动作"],
                        ["【应急】", "特殊情况处理"],
                        ["【案例】", "真实案例"],
                        ["【投影】", "PPT 翻到哪一页"],
                        ["【金句】", "需要写/说出的金句"],
                    ])

    add_h2(doc, "讲课 5 个原则")
    for s in [
        "1. 慢就是快:每个模块留 30% 缓冲",
        "2. 案例驱动:每个概念配 1 个真实案例",
        "3. 学员动手:每模块至少 1 个练习",
        "4. 不评判:学员说什么都先认同,后引导",
        "5. 金句密度:每个模块 1-2 个金句,贴教室",
    ]:
        add_body_no_indent(doc, s)

    add_page_break(doc)

    # 第一部分
    add_h1(doc, "第一部分 课程总览")
    add_h2(doc, "1.1 课程定位")
    add_body(doc, "AI 时代管理者从'指挥者'到'教练'的角色转变。")
    add_body(doc, "目标:2 天,3 件事:看清坍塌、掌握 3 大工具、30 天落地。")
    add_h2(doc, "1.2 课程结构")
    add_table_simple(doc, ["天", "模块", "时间", "核心工具"],
                    [
                        ["Day 1", "模块一 管理坍塌", "60min", "3 次坍塌分析"],
                        ["Day 1", "模块二 6 个差异", "60min", "6 维差异表"],
                        ["Day 1", "模块三 3C + 4 模式", "90min", "F1/F3/F4"],
                        ["Day 1", "复盘", "30min", "F2 时间账单"],
                        ["Day 2", "模块四 GROW+", "120min", "F5/F12/F13"],
                        ["Day 2", "模块五 七层倾听", "90min", "F6/F7/F8"],
                        ["Day 2", "模块六 30 天落地", "90min", "F19/F18"],
                        ["Day 2", "收尾", "30min", "F1 训后评估"],
                    ])
    add_h2(doc, "1.3 教学理念")
    add_body(doc, "1. 三层结构:Why→What→How,每模块都过这一遍")
    add_body(doc, "2. 案例驱动:8 个真实案例,老王/小李/小张/老周等")
    add_body(doc, "3. 即时练习:每模块 1 个练习,30% 时间学员动手")
    add_body(doc, "4. 慢就是快:5 分钟 1 个 8 分钟心谈,胜过 1 小时说教")
    add_quote(doc, "成关键节点,善借力打力。", "—— 找到学员的关键转换时刻")

    add_page_break(doc)

    # 第二部分
    add_h1(doc, "第二部分 时间控制表")
    add_body(doc, "下表是 2 天的完整时间分配,允许 ±10% 浮动。")
    add_h2(doc, "2.1 Day 1 时间表")
    add_table_simple(doc, ["时段", "内容", "时长", "总时"],
                    [
                        ["08:30-09:00", "签到 + 预评量表(F1)", "30min", "30"],
                        ["09:00-09:30", "开场 + 自我介绍", "30min", "60"],
                        ["09:30-10:30", "模块一:管理坍塌", "60min", "120"],
                        ["10:30-10:45", "茶歇", "15min", "135"],
                        ["10:45-11:45", "模块二:6 个差异", "60min", "195"],
                        ["11:45-12:00", "练习 1:6 维差异自我评估", "15min", "210"],
                        ["12:00-13:30", "午餐", "90min", "300"],
                        ["13:30-15:00", "模块三:3C + 4 模式", "90min", "390"],
                        ["15:00-15:15", "茶歇", "15min", "405"],
                        ["15:15-16:00", "练习 2-3:4 模式评估 + 3C 评估", "45min", "450"],
                        ["16:00-17:00", "第一天复盘 + F2 时间账单", "60min", "510"],
                    ])
    add_h2(doc, "2.2 Day 2 时间表")
    add_table_simple(doc, ["时段", "内容", "时长", "总时"],
                    [
                        ["08:30-09:00", "签到 + Day 1 复盘", "30min", "30"],
                        ["09:00-11:00", "模块四:GROW+(含 2 个练习)", "120min", "150"],
                        ["11:00-11:15", "茶歇", "15min", "165"],
                        ["11:15-12:45", "模块五:七层倾听(含 2 个练习)", "90min", "255"],
                        ["12:45-14:00", "午餐", "75min", "330"],
                        ["14:00-15:30", "模块六:30 天落地", "90min", "420"],
                        ["15:30-15:45", "茶歇", "15min", "435"],
                        ["15:45-16:45", "F19 行动计划 + 共学群", "60min", "495"],
                        ["16:45-17:15", "F1 训后评估 + 庆祝", "30min", "525"],
                    ])
    add_h2(doc, "2.3 时间控制 5 个技巧")
    for s in [
        "1. 永远带计时器,5 分钟响 1 次",
        "2. 超时 10 分钟,跳过 1 个练习",
        "3. 提前 10 分钟,加快 1 个练习",
        "4. 学员累了,直接进下 1 个模块",
        "5. 茶歇 15 分钟,卡死 1 次到点就开",
    ]:
        add_body_no_indent(doc, s)

    add_page_break(doc)

    # 第三部分
    add_h1(doc, "第三部分 训前准备清单")
    add_h2(doc, "3.1 教室准备")
    for s in [
        "投影仪 + 备用投影仪",
        "白板 + 3 色白板笔",
        "便利贴 5 包(黄、粉、绿、蓝、白)",
        "彩笔 30 支",
        "学员签到表",
        "F1 预评量表 30 份",
        "时钟/计时器(讲师台前)",
        "茶歇食物 + 矿泉水",
    ]:
        add_body_no_indent(doc, s)
    add_h2(doc, "3.2 设备检查")
    for s in [
        "1. 投影清晰度",
        "2. 翻页笔电池",
        "3. 麦克风(无线 2 套 + 备用 1 套)",
        "4. 音箱",
        "5. 教室空调",
        "6. 教室光线(可调暗)",
    ]:
        add_body_no_indent(doc, s)
    add_h2(doc, "3.3 个人准备")
    for s in [
        "1. 穿深色衣服(主色一致)",
        "2. 不带贵重物品上台",
        "3. 提前 30 分钟到场",
        "4. 提醒:把手机调静音,远离讲台",
        "5. 准备一瓶水",
    ]:
        add_body_no_indent(doc, s)
    add_h2(doc, "3.4 学员预热(训前 3 天)")
    add_body(doc, "发短信/微信给学员:")
    add_body_no_indent(doc, "1. 自我介绍 + 部门 + 1 个管理挑战")
    add_body_no_indent(doc, "2. F1 预评量表链接,训前 1 周填完")
    add_body_no_indent(doc, "3. 准备 1 个自己想解决的真实问题")
    add_body_no_indent(doc, "4. 30 天共学群二维码")

    add_page_break(doc)

    # 第四部分:Day 1
    add_h1(doc, "第四部分 Day 1 完整剧本")

    # 开场
    add_h2(doc, "Day 1 开场 09:00-09:30")
    add_lecture_block(doc, "30min", "开场 + 自我介绍",
        [
            "各位伙伴,早上好。欢迎来到《教练技术:新的管理方式》。",
            "我是罗宏伟,今天和明天,我们一起完成 3 件事。",
            "第一,看清 3 次管理坍塌。",
            "第二,掌握 3 大工具。",
            "第三,用 30 天,把方法变成肌肉记忆。",
            "先做个暖场:请大家用 1 分钟,和你左边和右边的伙伴互相介绍:叫什么、哪个部门、最近遇到的最大管理挑战。",
        ],
        [
            "走到学员中间,观察 2-3 组对话",
            "1 分钟到,拍 2 下掌,示意停止",
            "邀请 2-3 个学员分享他们的'最大管理挑战'",
            "金句板:谁能兼容谁,谁就能领导谁。",
        ],
        [
            "如果冷场:你举手先说一个自己的小挑战,降低门槛",
            "如果 1 个学员霸占话题:礼貌打断:'我们听下一个,3 个就好'",
        ])

    # Day 1 模块一
    add_h2(doc, "Day 1 模块一 09:30-10:30")
    add_lecture_block(doc, "60min", "AI 时代的管理坍塌",
        [
            "我们先问 1 个问题:20 年前,你凭什么当管理者?",
            "答案:我有信息、经验、决策权。",
            "20 年后,这 3 样东西,全塌了。",
            "今天我们一起看清这 3 次坍塌。",
        ],
        [
            "投影展示:3 次坍塌的时间线",
            "金句板:延时摄影——和时间做朋友。",
        ])

    add_lecture_block(doc, "20min", "第一次坍塌:信息差消失",
        [
            "过去,管理者是信息中枢。员工遇到问题,来找你问。",
            "现在,AI 3 秒出答案。员工比你先知道。",
            "我举个例子。60 后车间主任老王,30 年工龄。新员工用手机扫了一下设备,3 秒读出参数。老王翻了 5 分钟手册没找到。",
            "他开始怀疑:我是不是该退了?",
            "信息坍塌 = 你不再是唯一的信息源。",
            "应对:接受、升级、整合。",
        ],
        [
            "投影:老王翻手册的对比图",
            "问 1 个学员:你身边有没有这样的'老王'?",
            "金句板:谁能兼容谁,谁就能领导谁。",
        ],
        [
            "如果学员沉默:你等 5 秒,自己分享 1 个类似案例",
        ])

    add_lecture_block(doc, "20min", "第二次坍塌:经验失效",
        [
            "过去,经验就是权威。10 年工龄,说了算。",
            "现在,AI 看过 100 万案例,比你多 1 万倍。",
            "上海某机械厂,30 年老师傅,看裂纹 30 年,听声音就知道哪里有问题。",
            "AI 音频分析 + 图像识别,准确率 98%。半年后,老师傅,你带新徒弟吧。",
            "他不是被 AI 替代,是被'经验'替代。",
            "经验坍塌 = 你不再是唯一的判断者。",
            "应对:接受、转化、升级。",
        ],
        [
            "投影:老师傅 vs AI 的对比",
            "邀请 1 个学员分享:你身边有'老师傅'吗?",
        ])

    add_lecture_block(doc, "20min", "第三次坍塌:决策权下沉",
        [
            "过去,你定方向,员工执行。",
            "现在,Z 世代自己定 KPI,自己选项目。",
            "我有个客户,95 后产品经理。老板说:下周上线这个功能。",
            "95 后说:我做了一份用户调研,数据不支持。先改。",
            "老板沉默 3 秒:行,你说了算。",
            "决策权坍塌 = 你不再是唯一的决策者。",
            "应对:接受、赋能、协作。",
        ],
        [
            "投影:决策权金字塔倒置图",
            "金句板:真实感而不是真实。",
        ])

    add_lecture_block(doc, "20min", "三次坍塌的累积效应",
        [
            "信息坍塌:你不再是唯一的信息源。",
            "经验坍塌:你不再是唯一的判断者。",
            "决策权坍塌:你不再是唯一的决策者。",
            "3 个支柱,全塌了。所以,不是我们变差了,是我们脚下的地,塌了。",
            "怎么办?换一种站姿。从指挥者,变成教练。",
        ],
        [
            "投影:3 次坍塌的累积图",
            "金句板:谁能兼容谁,谁就能领导谁。",
        ])

    add_lecture_block(doc, "10min", "模块一 FAQ 解答",
        [
            "Q:如果团队都不爱用 AI 怎么办?",
            "A:不是不爱用,是没安全感。先做 1V1 心谈,问'你最担心什么?'。",
            "Q:30 年经验,真的没用了吗?",
            "A:经验不会没用,但'凭经验拍脑袋'会没用。把经验变成方法论。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # 模块二
    add_h2(doc, "Day 1 模块二 10:45-11:45")
    add_lecture_block(doc, "60min", "教练 vs 传统管理者 6 个差异",
        [
            "你不需要变成另一个人。你只需要看清 6 个差异,每个差异往前一步。",
        ],
        [
            "投影:6 维差异表(传统 vs 教练)",
        ])

    for diff, t_script, t_actions in [
        ("权威来源",
         [
             "传统:职位给我权威,我说了算。",
             "教练:专业+信任给我权威,你愿意听。",
             "职位权威 1 年失效,信任权威 10 年有效。",
         ],
         [
             "投影:6 维表 + 当前差异的高亮",
             "邀请 1 个学员分享他自己的例子",
         ]),
        ("对话姿态",
         [
             "传统:我讲你听,我命令你执行。",
             "教练:我问你答,我帮你找到答案。",
         ],
         [
             "投影:小李的两次老板剧本",
             "金句板:讨喜而不是讨好。",
         ]),
        ("决策方式",
         [
             "传统:我定方向,你执行。",
             "教练:我帮你看清,你自己定。",
             "下属自己定的方案,执行度 3 倍于你定的。",
         ],
         [
             "投影:老陈定 vs 小林自己定",
         ]),
        ("角色定位",
         [
             "传统:我是裁判,定对错。",
             "教练:我是教练,陪你上场。",
         ],
         [
             "金句板:善借力打力。",
         ]),
        ("反馈风格",
         [
             "传统:你做错了,再批评。",
             "教练:你做到了,再加什么更好。",
             "行为反馈 vs 人格评判——后者让人崩。",
         ],
         [
             "投影:4 种反馈话术对比表",
         ]),
        ("长期目标",
         [
             "传统:本季度业绩,完成就 OK。",
             "教练:5 年后他还在成长,是我赢。",
         ],
         [
             "金句板:延时摄影——和时间做朋友。",
         ]),
    ]:
        add_lecture_block(doc, "8min", f"差异: {diff}",
            t_script, t_actions)

    add_lecture_block(doc, "5min", "教练心态 4 个核心信念",
        [
            "1. 每个人都比自己以为的更强。",
            "2. 答案在他心里,不在我嘴里。",
            "3. 我不是来救他,我是来陪他。",
            "4. 短期会慢,长期会快。",
        ],
        [
            "投影:4 信念列表",
            "金句板:延时摄影——和时间做朋友。",
        ])

    add_lecture_block(doc, "5min", "教练红线:5 个不能",
        [
            "1. 不能替他做决定",
            "2. 不能评判他的选择",
            "3. 不能忽视他的情绪",
            "4. 不能用'我是为你好'绑架",
            "5. 不能在他没准备好时硬推",
        ],
        [
            "投影:5 红线",
        ])

    add_lecture_block(doc, "5min", "练习 1:6 维差异自我评估",
        [
            "现在请大家拿出 F1 评估表,在'6 维差异'那一栏,给现在的自己打分。1 分 = 还是传统,5 分 = 已经教练。",
            "3 分钟打分,3 分钟分享。",
        ],
        [
            "3 分钟计时",
            "3 分钟到,邀请 3 个学员分享",
        ],
        [
            "如果学员打分都很高(都 4-5):你问'哪个维度最难做到?'引导反思",
            "如果学员打分都很低(都 1-2):你问'哪个维度最容易先动?'引导起点",
        ])

    add_lecture_block(doc, "5min", "模块二 FAQ 解答",
        [
            "Q:如果下属不想被教练,只想被告知呢?",
            "A:先尊重他的偏好。同时引导:'你愿意试 1 次吗?'",
            "Q:教练是不是太慢了,任务急的时候怎么办?",
            "A:教练是'先慢后快'。第 1 次教练对话要 30 分钟,后面下属自己会做,你反而更省时间。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # 模块三
    add_h2(doc, "Day 1 模块三 13:30-15:00")
    add_lecture_block(doc, "90min", "3C 信任模型 + 4 模式员工",
        [
            "看见人,才能带人。3C 是建立信任的 3 个关键动作,4 模式是识别员工的 4 种工具。",
        ],
        [
            "投影:1 个数字(高信任团队绩效 +76%)",
        ])

    add_lecture_block(doc, "10min", "1 个数字:信任的价值",
        [
            "Paul Zak 2017 年 NEJM 研究:高信任团队,绩效比低信任团队高 76%。",
            "信任 = 团队的隐形业绩杠杆。",
        ],
        [
            "投影:研究数据图",
        ])

    add_lecture_block(doc, "20min", "3C 信任模型",
        [
            "3C = Connection 连接 + Commitment 承诺 + Caring 关怀。",
            "Connection 连接:看见他。名字、状态、困难。",
            "Commitment 承诺:我说的,我做到。",
            "Caring 关怀:关心他的状态、成长、生活。",
            "3C 不是'我对他做了什么',而是'他感受到什么'。",
        ],
        [
            "投影:3C 关系图",
            "案例:小张记得下属生日",
            "案例:老李 3 次没说",
            "案例:小陈记得下属母亲住院",
        ])

    add_lecture_block(doc, "20min", "4 模式员工:AI 时代的全新分类",
        [
            "AI 时代,员工分 4 种。",
            "先行者:抢着用 AI,跑得快。你要授权+挑战+允许试错。",
            "整合者:愿意尝试,先看效果。你要桥梁+跨部门协作。",
            "观望者:等别人先试,看情况。你要 1V1 心谈+安全尝试。",
            "保守者:担心被替代,不敢动。你要尊重节奏+稳定本职。",
        ],
        [
            "投影:4 模式象限图",
            "案例:小张(先行者)用 AI 做调研报告,1 小时搞定",
            "案例:老王(保守者)学不会,失眠,易怒",
        ])

    add_lecture_block(doc, "20min", "跨模式协作:1+1>2 的搭配",
        [
            "先行者 + 整合者:创新+落地,黄金组合。",
            "观望者 + 先行者:Buddy 机制,先看后试。",
            "保守者 + 整合者:稳定+信任,小步前进。",
        ],
        [
            "投影:3 种搭配示意",
        ])

    add_lecture_block(doc, "20min", "练习 2-3:4 模式 + 3C 评估",
        [
            "请大家拿出 F3(4 模式自评)和 F4(3C 信任度评估)。",
            "先做 F3:列出 5 个下属,按 4 模式打分,主类型取最高分。10 分钟。",
            "再做 F4:给自己的 3C 打分,找最弱维度。10 分钟。",
        ],
        [
            "10 分钟计时",
            "10 分钟计时",
            "邀请 2-3 个学员分享",
        ],
        [
            "如果学员纠结打分:你提醒'主类型取最高分就行,不需要精确'",
        ])

    add_lecture_block(doc, "10min", "模块三 FAQ 解答",
        [
            "Q:3C 是不是要花很多时间?",
            "A:不用。每天 5 分钟 Connection,1 件小事 Commitment,1 句关心 Caring。",
            "Q:保守者是不是该被淘汰?",
            "A:不一定。保守者稳定,是团队的压舱石。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # Day 1 复盘
    add_h2(doc, "Day 1 复盘 16:00-17:00")
    add_lecture_block(doc, "60min", "第一天复盘 + F2 时间账单",
        [
            "今天我们走了 3 个模块:管理坍塌、6 个差异、3C + 4 模式。",
            "在结束之前,我想让大家做 1 件事:看清你的时间。",
            "请大家拿出 F2(时间账单)。回忆过去 12 小时(昨晚 18:00 到今天 06:00),每 30 分钟做了什么。",
        ],
        [
            "投影:F2 时间账单样本",
            "15 分钟填写",
        ])

    add_lecture_block(doc, "20min", "F2 时间账单复盘",
        [
            "现在我们看统计。",
            "12 小时里,睡眠多少小时?",
            "刷手机多少小时?",
            "和家人说话多少小时?",
            "和下属 1V1 多少小时?",
        ],
        [
            "投影:统计饼图",
            "邀请 3-4 个学员分享",
        ],
        [
            "如果时间账单看起来很惨(刷手机 4 小时):你笑着说'我也是,先接纳'",
        ])

    add_lecture_block(doc, "20min", "第一天文案作业",
        [
            "回去之后,做 1 件事:找 1 个下属,做 1 次 8 分钟心谈。",
            "明天上午,我们会花 30 分钟复盘这个心谈。",
        ],
        [
            "投影:Day 1 文案",
            "金句板:讨喜而不是讨好。",
        ])

    add_lecture_block(doc, "20min", "Day 1 总结 + 寄语",
        [
            "今天 3 个模块,我们学了 3 件事:",
            "看清坍塌,理解为什么变。",
            "看清 6 个差异,理解变成什么样。",
            "看清 3C + 4 模式,理解怎么变。",
            "明天我们学 3 件事:GROW+ 对话、七层倾听、30 天落地。",
        ],
        [
            "投影:Day 1 总结页",
        ])

    add_page_break(doc)

    # 第五部分:Day 2
    add_h1(doc, "第五部分 Day 2 完整剧本")

    add_h2(doc, "Day 2 开场 08:30-09:00")
    add_lecture_block(doc, "30min", "Day 1 复盘 + 8 分钟心谈分享",
        [
            "各位伙伴,早上好。昨天回去做了 8 分钟心谈的请举手。",
            "邀请 3-4 个学员分享。",
            "1 个学员 5 分钟,先说 8 分钟里发生什么,再说 1 个发现。",
        ],
        [
            "5 分钟计时",
            "金句板:真实感而不是真实。",
        ],
        [
            "如果没人做:你问'为什么没做?'然后说'我也是,所以今天我们一起做'",
        ])

    # 模块四
    add_h2(doc, "Day 2 模块四 09:00-11:00")
    add_lecture_block(doc, "120min", "GROW+ 5 阶段对话框架",
        [
            "GROW+ 是教练对话的标准动作。5 阶段 25 问,覆盖一次完整对话。",
        ],
        [
            "投影:GROW+ 5 阶段示意",
        ])

    for stage, q_lines, stage_actions in [
        ("G 目标",
         [
             "G = Goal。整个对话的灯塔。G 阶段没做对,后面都白搭。",
             "5 个关键问题:",
             "1. 你最想达成的是什么?",
             "2. 如果成功了,你会看到什么?",
             "3. 这个目标对你为什么重要?",
             "4. 你愿意为它付出什么?",
             "5. 什么时候必须达成?",
             "G 阶段关键:问'为什么重要',让目标内化。",
         ],
         [
             "投影:GROW+ G 阶段 + 5 个问题",
             "1 个 5 分钟迷你练习:2 人一组,1 个问 G,1 个答",
         ]),
        ("R 现实",
         [
             "R = Reality。看清脚下,才能走对方向。",
             "5 个关键问题:",
             "1. 现在的情况是什么?",
             "2. 你做了哪些尝试?效果如何?",
             "3. 卡在哪里?",
             "4. 哪些是你能控制的,哪些不能?",
             "5. 如果不做任何改变,3 个月后会怎样?",
             "R 阶段关键:让对方先看见自己的位置。",
         ],
         [
             "投影:GROW+ R 阶段 + 5 个问题",
             "现场演练:1 个人扮演下属,1 个人扮演教练,走 R 阶段",
         ]),
        ("O 选择",
         [
             "O = Options。让他自己看见 N 种可能。",
             "5 个关键问题:",
             "1. 如果不受限,你会怎么做?",
             "2. 你身边谁做得最好?你能学他什么?",
             "3. 还有哪些方法你没试过?",
             "4. 哪个方法最让你兴奋?",
             "5. 如果只能选一个,你先做哪个?",
             "O 阶段关键:列 5 个以上方案,每个都问'为什么这个好'。",
         ],
         [
             "投影:GROW+ O 阶段 + 5 个问题",
         ]),
        ("W 意愿",
         [
             "W = Will。激活内在动力。",
             "5 个关键问题:",
             "1. 你最想做哪个?为什么?",
             "2. 做到后对你意味着什么?",
             "3. 你愿意为它放弃什么?",
             "4. 如果不做,你会有什么感觉?",
             "5. 你对自己的承诺是什么?",
             "W 阶段关键:让他自己说'我愿意'。",
         ],
         [
             "投影:GROW+ W 阶段 + 5 个问题",
         ]),
        ("A 行动",
         [
             "A = Action。把对话变成真东西。",
             "5 个关键问题:",
             "1. 你第一步做什么?",
             "2. 什么时候做?",
             "3. 你需要的支持是什么?",
             "4. 我(主管)能帮你什么?",
             "5. 我们怎么知道你做到了?",
             "A 阶段关键:具体到今天/本周/这个月。",
         ],
         [
             "投影:GROW+ A 阶段 + 5 个问题",
         ]),
    ]:
        add_lecture_block(doc, "20min", f"{stage}", q_lines, stage_actions)

    add_lecture_block(doc, "20min", "GROW+ 完整剧本:小李延期",
        [
            "现在我们看 1 个完整剧本。",
            "小李,90 后,项目骨干,做事快。老板信任他,什么都给他。",
            "项目延期 1 周。老板想骂他。",
            "GROW+ 怎么走?",
            "G: 小李,你最想达成什么?",
            "R: 现在延期了,情况怎样?",
            "O: 如果不延期 1 周,你想怎么做?",
            "W: 你愿意做哪个?",
            "A: 第一步是什么?什么时候?",
            "结果:小李主动加班赶回进度,1 周准时交付。",
        ],
        [
            "投影:小李剧本",
            "邀请 1 个学员角色扮演'老板',1 个扮演'小李',现场走 1 遍",
        ],
        [
            "如果学员紧张:你亲自示范一次,然后再换学员",
        ])

    add_lecture_block(doc, "20min", "GROW+ 第 2 个剧本:老张想转岗",
        [
            "老张 55 岁,工程师,3 年后退休,想转岗做培训。",
            "G: 你为什么想转岗?",
            "R: 现在你做这件事的卡点是什么?",
            "O: 如果不受限,你会怎么做?",
            "W: 你愿意做哪个?",
            "A: 第一步是什么?什么时候?",
            "结果:老张 3 个月后开始带新人,找到意义感。",
        ],
        [
            "投影:老张剧本",
        ])

    add_lecture_block(doc, "10min", "模块四 FAQ 解答",
        [
            "Q:GROW+ 对话要多长时间?",
            "A:30-60 分钟。第 1 次可能 60 分钟,熟练后 30 分钟。",
            "Q:GROW+ 5 阶段顺序能变吗?",
            "A:不能。G 必须在第 1 位,否则后面都乱。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # 模块五
    add_h2(doc, "Day 2 模块五 11:15-12:45")
    add_lecture_block(doc, "90min", "七层倾听阶梯",
        [
            "从'我听见了'到'你被听见了'。",
            "大多数管理者卡在第 3 层,我们要到第 5 层以上。",
        ],
        [
            "投影:7 层倾听阶梯",
        ])

    add_lecture_block(doc, "20min", "七层阶梯总览",
        [
            "7 = 生成性倾听:让对方找到他自己都没意识到的答案。",
            "6 = 反思性倾听:把他说的话翻译成深层含义。",
            "5 = 共鸣性倾听:分享我类似的感受,让他不孤单。",
            "4 = 同理心倾听:听到情绪,命名情绪,让他被看见。",
            "3 = 理解性倾听:听他说完,复述确认(大多数人在这)。",
            "2 = 选择性倾听:只听我感兴趣的,过滤其他的。",
            "1 = 忽视:心不在焉,人在心不在。",
        ],
        [
            "投影:7 层 + 每层的关键词",
        ])

    add_lecture_block(doc, "20min", "命名情绪:5 个万能句式",
        [
            "1. 你看起来感到___对吗?",
            "2. 听上去,你现在___对吗?",
            "3. 如果是我,我会感到___。你呢?",
            "4. 你愿意多说说这个___吗?",
            "5. 你希望我___吗?",
        ],
        [
            "投影:5 句式",
            "现场演练:2 人一组,1 个说 1 句话,另 1 个用 5 句式回应",
        ])

    add_lecture_block(doc, "20min", "8 分钟心谈操作指南",
        [
            "1. 找一个安静的地方(不被打扰)。",
            "2. 说:'我有 8 分钟,你愿意聊聊吗?'",
            "3. 前 4 分钟:只听,不评判,不建议。",
            "4. 后 4 分钟:问 1-2 个开放问题。",
            "5. 结束:问'你还需要什么吗?'",
            "8 分钟 = 不长不短,刚好够听见他。",
        ],
        [
            "投影:8 分钟心谈步骤",
        ])

    add_lecture_block(doc, "30min", "练习:8 分钟心谈 + 8 分钟心谈错误",
        [
            "现在 2 人一组,做 1 次 8 分钟心谈。",
            "1 个做主管,1 个做下属,8 分钟,只看步骤。",
            "5 个典型错误:",
            "1. 边听边看手机。",
            "2. 立刻给建议。",
            "3. 打断他说话。",
            "4. '我以前也是',变成我讲。",
            "5. 急着解决他的问题。",
        ],
        [
            "8 分钟计时",
            "8 分钟到,拍 2 下掌",
            "5 个错误投屏",
        ],
        [
            "如果学员时间超时:你打断'时间到,下次做得更好'",
        ])

    add_lecture_block(doc, "10min", "模块五 FAQ 解答",
        [
            "Q:如果我没什么情绪经验,怎么听?",
            "A:先学命名情绪。每天 1 个新词,1 个月后你就掌握 30 个。",
            "Q:沉默时,我不说话,会不会冷场?",
            "A:不会。沉默是给对方'想'的时间。30 秒后,他会说更深的话。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # 模块六
    add_h2(doc, "Day 2 模块六 14:00-15:30")
    add_lecture_block(doc, "90min", "30 天落地实践",
        [
            "把方法变成肌肉记忆。30 天,5 个维度,4 周。",
        ],
        [
            "投影:30 天 5 维度",
        ])

    add_lecture_block(doc, "15min", "5 维度讲解",
        [
            "心谈:每天 8 分钟心谈 1 次。",
            "GROW+:每天用 GROW+ 解决 1 个小问题。",
            "3C:每天 1 个 3C 动作。",
            "团队分享:每周 1 次内部分享。",
            "反思:每天写 1 段反思日记。",
        ],
        [
            "投影:5 维度表",
        ])

    add_lecture_block(doc, "30min", "4 周计划详解",
        [
            "第 1 周:心谈。每天 1 个下属,5 天 5 个。",
            "第 2 周:GROW+。选 4-5 个下属,各 1 次。",
            "第 3 周:3C。每天 1 个动作(Connection/Commitment/Caring)。",
            "第 4 周:团队分享 + 30 天反思日记 + 庆祝。",
        ],
        [
            "投影:4 周日历",
        ])

    add_lecture_block(doc, "20min", "30 天共学群规则",
        [
            "1. 每天 1 个动作打卡。",
            "2. 每周 1 次 8 分钟心谈复盘。",
            "3. 每周 1 次团队分享。",
            "4. 每周 1 个 30 天反思日记片段。",
            "5. 30 天后全群庆祝。",
        ],
        [
            "投影:共学群规则",
            "金句板:领先半步,吃尽红利。",
        ])

    add_lecture_block(doc, "20min", "F19 行动计划",
        [
            "请大家拿出 F19。5 个要素:目标、意义、行动、卡点、衡量。",
            "5 分钟写,10 分钟分享。",
        ],
        [
            "5 分钟计时",
            "邀请 2-3 个学员分享",
        ])

    add_lecture_block(doc, "5min", "模块六 FAQ 解答",
        [
            "Q:30 天一定要做完吗?",
            "A:建议做完。30 天不连续,效果会打 5 折。",
            "Q:30 天后,还要继续吗?",
            "A:建议继续。但 30 天后,从'每天打卡'变成'每周打卡'。",
        ],
        [
            "投影:5 个常见问题",
        ])

    add_page_break(doc)

    # Day 2 收尾
    add_h2(doc, "Day 2 收尾 15:45-17:15")
    add_lecture_block(doc, "30min", "F19 行动计划 + 共学群",
        [
            "30 天不是 30 天学完,30 天是把方法变成肌肉记忆。",
            "我们建 1 个 30 天共学群。",
        ],
        [
            "投影:共学群二维码",
        ])

    add_lecture_block(doc, "30min", "F1 训后评估 + 庆祝",
        [
            "请大家再填一次 F1 训后评估。",
            "对比训前训后,看变化。",
        ],
        [
            "10 分钟填写",
            "金句板:成关键节点,善借力打力。",
        ],
        [
            "如果学员感觉'没什么变化':你提醒'2 天 30 分钟心谈不可能变很多,30 天后才能看到'",
        ])

    add_lecture_block(doc, "20min", "庆祝 + 寄语",
        [
            "各位伙伴,2 天结束了。",
            "你不需要变成一个完美的教练,你只需要开始。",
            "每周多 1 个 8 分钟心谈,30 天后,你会感谢今天的自己。",
            "恭喜大家完成 2 天学习。",
        ],
        [
            "金句板:谁能兼容谁,谁就能领导谁。",
            "播放庆祝音乐",
        ])

    add_page_break(doc)

    # 第六部分
    add_h1(doc, "第六部分 案例讲解要点")
    add_body(doc, "8 个真实案例,4 类员工覆盖。下面是每个案例的讲解要点。")

    cases = [
        ("小李延期", "90 后,项目骨干", "GROW+ G 找到他真正想要的", "1 周准时交付,主动加班"),
        ("老王转型", "60 后车间主任,30 年工龄", "不评判他的焦虑,问'你愿意为它做什么?'", "主动申请做 AI 教练"),
        ("小张迷茫", "工作 3 年,不知道往哪走", "问'3 年后你想成为什么?'", "3 个月后转岗成功"),
        ("老李瓶颈", "做了 5 年总监,升不上去", "问'如果不受限,你会怎么做?'", "找到新赛道,开副业"),
        ("小陈加班", "95 后,连续 1 个月 996", "8 分钟心谈 + 命名情绪", "主管才知道他的极限"),
        ("老张 AI 焦虑", "55 岁,工程师,被 AI 替代", "不评判他的失落,问'你最想达成什么?'", "3 个月后成 AI 培训师"),
        ("小赵创新", "推 AI 项目被否 3 次", "R 阶段问'你做了哪些尝试?'", "第 4 次提案通过"),
        ("老周临退休", "58 岁,5 年后退休", "8 分钟心谈 + 翻译", "申请做导师,找到意义感"),
    ]
    add_table_simple(doc, ["案例", "背景", "关键动作", "结果"], cases)

    add_h2(doc, "案例讲解 5 个原则")
    for s in [
        "1. 永远先说背景,再说动作,再说结果",
        "2. 关键动作要具体(谁问的、问什么、怎么问)",
        "3. 结果要可衡量(数据、时间、可观察的行为)",
        "4. 案例讲解不超过 5 分钟,讲完立刻让学员说感受",
        "5. 案例讲完,问 1 个问题:'你身边有这样的人吗?'",
    ]:
        add_body_no_indent(doc, s)

    add_page_break(doc)

    # 第七部分
    add_h1(doc, "第七部分 应急方案")

    add_h2(doc, "7.1 冷场应急")
    for s in [
        "场景:你问 1 个问题,没人回答。",
        "动作:沉默 5 秒。然后说'我先说 1 个,降低门槛。'",
        "剧本:我自己 1 个小挑战,5 秒。",
        "然后:再问 1 个简单问题:'你身边有没有这样的'老王'?'",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.2 超时应急")
    for s in [
        "场景:模块超时 10 分钟以上。",
        "动作:跳过 1 个练习,直接进入下 1 个模块。",
        "原则:宁可快 10 分钟,不可慢 10 分钟。学员累了就什么都学不进去。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.3 学员挑战应急")
    add_body(doc, "场景:学员说'这套东西在我公司行不通'。")
    for s in [
        "动作 1:认同他的感受。'是的,在我客户那里也遇到过类似的。'",
        "动作 2:问 1 个问题。'如果可以解决 1 个小问题,你愿意试吗?'",
        "动作 3:不在公开场合争论。课后 1V1 沟通。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.4 学员情绪应急")
    add_body(doc, "场景:学员被案例触动了,哭出来。")
    for s in [
        "动作 1:走过去,递 1 张纸巾。",
        "动作 2:不立刻说话,等 10 秒。",
        "动作 3:问'你需要 1 分钟吗?'",
        "动作 4:课后 1V1 跟进。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.5 技术故障应急")
    add_body(doc, "场景:投影没反应、麦克风没声、电脑死机。")
    for s in [
        "动作 1:白板代替。先画,再补投影。",
        "动作 2:让助理去解决,不耽误学员。",
        "动作 3:5 分钟内修不好,直接进下 1 个模块。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.6 极端情况应急")
    add_body(doc, "场景:学员说'我不想学了,我要走'。")
    for s in [
        "动作 1:不拦,不说服。",
        "动作 2:课间 1V1 沟通。",
        "动作 3:问 1 个问题:'如果有什么方法能让你学到东西,你会试吗?'",
        "动作 4:如果他还是要走,送走,课后 1V1 跟进。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.7 学员抢话应急")
    add_body(doc, "场景:1 个学员抢话太多,其他学员没机会。")
    for s in [
        "动作 1:礼貌打断:'谢谢分享,我们也听听其他伙伴的想法。'",
        "动作 2:点其他学员:右边那位伙伴,你说说?",
        "动作 3:课后 1V1 跟这位学员沟通,了解他为什么抢话。",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "7.8 学员睡着应急")
    add_body(doc, "场景:1 个学员在睡觉。")
    for s in [
        "动作 1:不直接叫醒,继续讲课。",
        "动作 2:走到他身边,轻轻拍下肩膀。",
        "动作 3:课间 1V1 沟通,问他是否需要咖啡/休息。",
    ]:
        add_body_no_indent(doc, s)

    add_page_break(doc)

    # 第八部分
    add_h1(doc, "第八部分 评估标准")
    add_h2(doc, "8.1 课程级别评估")
    add_table_simple(doc, ["级别", "表现"],
                    [
                        ["L1 反应层", "课程满意度 ≥ 4.5/5"],
                        ["L2 学习层", "F1 训后 ≥ 训前 30%"],
                        ["L3 行为层", "30 天后动作打卡 ≥ 80%"],
                        ["L4 成果层", "90 天后下属绩效改善"],
                    ])
    add_h2(doc, "8.2 学员考核标准")
    add_table_simple(doc, ["维度", "合格"],
                    [
                        ["出勤", "2 天全程参与,迟到 ≤ 1 次"],
                        ["F1 训后", "≥ 训前 30%"],
                        ["练习", "至少 3 个练习完成"],
                        ["F19", "5 要素全部填写"],
                    ])
    add_h2(doc, "8.3 评估数据收集")
    for s in [
        "1. 训后立即:F1 训后 + 课程满意度",
        "2. 训后 30 天:动作打卡数据 + 行为改变调研",
        "3. 训后 90 天:下属绩效 + 团队氛围调研",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "8.4 课程满意度收集表")
    add_body(doc, "5 分制,1=非常不满意,5=非常满意")
    add_table_simple(doc, ["维度", "1", "2", "3", "4", "5"],
                    [
                        ["内容", "", "", "", "", ""],
                        ["讲师", "", "", "", "", ""],
                        ["案例", "", "", "", "", ""],
                        ["练习", "", "", "", "", ""],
                        ["时长", "", "", "", "", ""],
                        ["环境", "", "", "", "", ""],
                        ["总体", "", "", "", "", ""],
                    ])

    add_page_break(doc)

    # 第九部分
    add_h1(doc, "第九部分 训后跟进指南")
    add_h2(doc, "9.1 30 天跟进(讲师)")
    for s in [
        "1. 训后 1 周:发 F18 打卡表到共学群",
        "2. 训后 2 周:在共学群做 1 次 8 分钟心谈复盘",
        "3. 训后 3 周:在共学群做 1 次 GROW+ 剧本分享",
        "4. 训后 4 周:30 天庆祝 + 颁发结业证书",
    ]:
        add_body_no_indent(doc, s)
    add_h2(doc, "9.2 90 天跟进(讲师)")
    for s in [
        "1. 训后 60 天:发行为改变调研",
        "2. 训后 90 天:发成果层调研 + 团队氛围调研",
        "3. 训后 90 天:邀请 3-5 个学员做案例分享",
    ]:
        add_body_no_indent(doc, s)
    add_h2(doc, "9.3 学员跟进清单")
    add_table_simple(doc, ["时间", "动作", "讲师"],
                    [
                        ["训后当天", "发 30 天打卡表", "群发"],
                        ["训后 1 周", "提醒 Day 1-7 心谈", "群发"],
                        ["训后 2 周", "提醒 Day 8-14 GROW+", "群发"],
                        ["训后 3 周", "提醒 Day 15-21 3C", "群发"],
                        ["训后 4 周", "30 天庆祝", "群发"],
                        ["训后 60 天", "行为改变调研", "1V1"],
                        ["训后 90 天", "成果层调研", "1V1"],
                    ])

    add_page_break(doc)

    # 第十部分
    add_h1(doc, "第十部分 讲师自评表")
    add_body(doc, "训后当天填写。每项 1-5 分。")

    self_eval = [
        ("内容完成度", "6 模块全部完成"),
        ("时间控制", "总时长 ±10% 之内"),
        ("金句密度", "≥ 6 大金句"),
        ("案例数量", "≥ 8 个真实案例"),
        ("练习数量", "≥ 6 个练习,每个 ≥ 5 分钟"),
        ("学员参与度", "≥ 70% 学员发言"),
        ("F1 训后提升", "≥ 30%"),
        ("课程满意度", "≥ 4.5/5"),
        ("讲师剧本执行", "≥ 90% 剧本到位"),
        ("应急方案", "无重大事故"),
    ]
    add_table_simple(doc, ["维度", "标准", "自评分(1-5)"], [(*row, "") for row in self_eval])

    add_h2(doc, "自评反思(3 个问题)")
    for s in [
        "1. 哪个模块讲得最好?为什么?",
        "2. 哪个模块讲得最差?下次怎么改?",
        "3. 学员最深的反应是什么?",
    ]:
        add_body_no_indent(doc, s)

    add_h2(doc, "同行反馈(请 1 位讲师同事填写)")
    add_table_simple(doc, ["维度", "评分(1-5)", "建议"],
                    [
                        ["内容深度", "", ""],
                        ["讲师台风", "", ""],
                        ["学员互动", "", ""],
                        ["案例质量", "", ""],
                        ["时间控制", "", ""],
                    ])

    add_page_break(doc)

    # 第十一部分:学员互动剧本
    add_h1(doc, "第十一部分 学员互动剧本(高频问答)")
    add_body(doc, "本部分收录 30 个学员高频问题的标准答案。讲师在课堂上遇到类似问题,可直接引用。")

    add_h2(doc, "模块一常见问题(5 问)")
    add_faq(doc, "Q1:如果我团队都不爱用 AI 怎么办?",
            "A:不是不爱用,是没安全感。先做 1V1 心谈,问'你最担心什么?'。让团队感到'用了 AI 也不会被替代',他们才会开始试。")
    add_faq(doc, "Q2:我 30 年经验,真的没用了吗?",
            "A:经验不会没用,但'凭经验拍脑袋'会没用。把经验变成方法论,变成可教的东西,你的价值会从'我懂'变成'我会教'。")
    add_faq(doc, "Q3:年轻人不听我的,怎么办?",
            "A:这是好事。说明他们有主见。你要做的是从'替他定'变成'陪他看'。让他自己定,你帮他看清代价。")
    add_faq(doc, "Q4:如果我学 AI 学不会,怎么办?",
            "A:学不会是正常的,AI 也在变。关键是不要抗拒,每周学 1 个新动作。1 年后,你就超过 80% 的人。")
    add_faq(doc, "Q5:三次坍塌,真的会同时发生吗?",
            "A:不同行业,坍塌顺序不同。但 3 年内,3 次坍塌都会到。先看清,早准备。")

    add_h2(doc, "模块二常见问题(5 问)")
    add_faq(doc, "Q1:如果下属不想被教练,只想被告知呢?",
            "A:先尊重他的偏好。但同时引导:'你愿意试 1 次吗?'大多数时候,试了 1 次就会尝到甜头。")
    add_faq(doc, "Q2:教练是不是太慢了,任务急的时候怎么办?",
            "A:教练是'先慢后快'。第 1 次教练对话要 30 分钟,但后面下属自己会做,你反而更省时间。")
    add_faq(doc, "Q3:如果员工自己定的方案很差,怎么办?",
            "A:让他先试。如果失败,你陪他复盘:'下次你会怎么改?'失败的方案,是最有价值的教练场景。")
    add_faq(doc, "Q4:6 个差异我都做不到,先练哪个?",
            "A:先练'对话姿态'。从'我讲你听'变成'我问你答'。1 周就能见效。")
    add_faq(doc, "Q5:教练和放羊式管理有什么区别?",
            "A:教练=我陪你,陪你看见。放羊=我不管,你爱咋咋地。教练有连接、有承诺。放羊没有。")

    add_h2(doc, "模块三常见问题(5 问)")
    add_faq(doc, "Q1:如果下属不是单纯的 1 种模式,怎么办?",
            "A:每个人都是 4 模式的混合体。取主类型(打分最高的),其他模式是'次要风格'。")
    add_faq(doc, "Q2:3C 是不是要花很多时间?",
            "A:不用。每天 5 分钟 Connection,1 件小事 Commitment,1 句关心 Caring。30 天后,团队气质就变了。")
    add_faq(doc, "Q3:保守者是不是该被淘汰?",
            "A:不一定。保守者稳定,是团队的压舱石。给他们稳定本职 + AI 低风险切入,大多数会慢慢跟上来。")
    add_faq(doc, "Q4:先行者骄傲,怎么管?",
            "A:1V1 心谈 + 看见他的贡献 + 让他带人。当他开始'输出'时,骄傲会变成责任。")
    add_faq(doc, "Q5:3C 和 4 模式哪个先学?",
            "A:先学 3C。3C 是底层,4 模式是工具。先有连接,后有分类。")

    add_h2(doc, "模块四常见问题(5 问)")
    add_faq(doc, "Q1:GROW+ 对话要多长时间?",
            "A:30-60 分钟。第 1 次可能 60 分钟,熟练后 30 分钟。")
    add_faq(doc, "Q2:GROW+ 5 阶段顺序能变吗?",
            "A:不能。G 必须在第 1 位,否则后面都乱。但每个阶段可以来回走。")
    add_faq(doc, "Q3:GROW+ 可以用在 1V1 吗?",
            "A:可以。每周 1V1 走 1 遍 GROW+,团队会很有方向。")
    add_faq(doc, "Q4:GROW+ 是不是太结构化了?",
            "A:结构化是手段,目的是让下属自己想清楚。可以灵活变,但 5 阶段不能少。")
    add_faq(doc, "Q5:GROW+ 对话,主管要做笔记吗?",
            "A:建议做。记下关键句子,后面 1V1 用得上。")

    add_h2(doc, "模块五常见问题(5 问)")
    add_faq(doc, "Q1:如果我没什么情绪经验,怎么听?",
            "A:先学命名情绪。每天 1 个新词,1 个月后你就掌握 30 个。")
    add_faq(doc, "Q2:如果下属在第 1 层(忽视),我能在第 7 层听吗?",
            "A:不能。倾听是双向的。下属不打开,你再会听也没用。先用 3C 让他信任你。")
    add_faq(doc, "Q3:8 分钟太短,能不能 30 分钟?",
            "A:可以,但建议从 8 分钟开始。先做到 8 分钟,再延到 15 分钟、30 分钟。")
    add_faq(doc, "Q4:沉默时,我不说话,会不会冷场?",
            "A:不会。沉默是给对方'想'的时间。30 秒后,他会说更深的话。")
    add_faq(doc, "Q5:如果我听了,下属没反应,是不是我听错了?",
            "A:不是。也许他还没准备好。第 2 次听,他可能就打开了。")

    add_h2(doc, "模块六常见问题(5 问)")
    add_faq(doc, "Q1:30 天一定要做完吗?",
            "A:建议做完。30 天不连续,效果会打 5 折。")
    add_faq(doc, "Q2:共学群是必须的吗?",
            "A:建议加入。1 个人走 30 天很孤独,1 群人走 30 天很容易。")
    add_faq(doc, "Q3:30 天后,还要继续吗?",
            "A:建议继续。但 30 天后,从'每天打卡'变成'每周打卡'。")
    add_faq(doc, "Q4:如果 30 天后,下属没变化,是不是我白做了?",
            "A:不是。下属变化是慢的。3 个月后,你会看到效果。")
    add_faq(doc, "Q5:30 天后,我能教别人吗?",
            "A:能。我们会邀请优秀的学员做'30 天校友分享'。")

    add_page_break(doc)

    # 第十二部分
    add_h1(doc, "第十二部分 8 个案例的完整剧本")
    add_body(doc, "8 个案例的完整剧本,每个剧本 3-5 分钟讲完。")

    case_scripts = [
        ("案例 1:小李延期(完整剧本)",
         [
             "背景:小李,90 后,项目骨干,做事快。老板信任他,什么都给他。",
             "问题:项目延期 1 周,老板想骂他。",
             "GROW+:",
             "G: 小李,你最想达成什么?",
             "  '准时交付'",
             "R: 现在延期了,情况怎样?",
             "  '技术上卡 1 个 bug,3 天没解'",
             "O: 如果不延期 1 周,你想怎么做?",
             "  '找人帮 debug,或换 1 个方案'",
             "W: 你愿意做哪个?",
             "  '找人帮 debug,我有人选'",
             "A: 第一步是什么?什么时候?",
             "  '今天下午找张工,2 天内解决'",
             "结果:小李主动加班赶回进度,1 周准时交付。",
             "讲师学到的 3 件事:",
             "1. 别骂,问",
             "2. 别告诉答案,让他自己找",
             "3. 别替代他做,让他自己承诺",
         ]),
        ("案例 2:老王转型(完整剧本)",
         [
             "背景:老王,60 后车间主任,30 年工龄,担心被 AI 替代。",
             "问题:学不会,失眠,易怒,想提前退休。",
             "GROW+:",
             "G: 你为什么担心被 AI 替代?",
             "  '我看不懂 AI,担心下岗'",
             "R: 现在你做这件事的卡点是什么?",
             "  '没有人教我,没有场景'",
             "O: 如果不受限,你会怎么做?",
             "  '有人带 1 次,我会快很多'",
             "W: 你愿意做哪个?",
             "  '带新人吧,我能教 30 年经验'",
             "A: 第一步是什么?",
             "  '这周找老板,申请做 AI 教练'",
             "结果:老王主动申请做'AI 教练',带新员工。",
             "讲师学到的 3 件事:",
             "1. 对保守者先稳定本职,AI 低风险切入",
             "2. 找新角色比硬学新工具好",
             "3. 情绪关注比技能训练优先",
         ]),
        ("案例 3:小张迷茫(完整剧本)",
         [
             "背景:小张,工作 3 年,不知道往哪走。",
             "问题:每天机械干活,没动力,想换工作。",
             "GROW+:",
             "G: 你为什么想换工作?",
             "  '没意思,没成长'",
             "R: 现在你做这件事的卡点是什么?",
             "  '不知道往哪走'",
             "O: 如果不受限,你会怎么做?",
             "  '去做产品经理,我喜欢从 0 到 1'",
             "W: 你愿意做哪个?",
             "  '先内部转岗,再想外部'",
             "A: 第一步是什么?",
             "  '这周找产品总监聊 1 次'",
             "结果:小张选了产品方向,3 个月后转岗成功。",
             "讲师学到的 3 件事:",
             "1. 迷茫的人缺的不是答案,是看见",
             "2. GROW+ 的 G 阶段是给迷茫的人最好的礼物",
             "3. 不评判,只问",
         ]),
        ("案例 4:老李瓶颈(完整剧本)",
         [
             "背景:老李,做了 5 年总监,升不上去,想跳不动。",
             "问题:总监做了 5 年,升不上 VP,跳也跳不动。",
             "GROW+:",
             "G: 你为什么升不上去?",
             "  '没新业绩'",
             "R: 现在你做这件事的卡点是什么?",
             "  '赛道没选对'",
             "O: 如果不受限,你会怎么做?",
             "  '做 1 个副业,新赛道'",
             "W: 你愿意做哪个?",
             "  '我开 1 个付费咨询'",
             "A: 第一步是什么?",
             "  '这周末写 1 篇公众号文章'",
             "结果:老李找到新赛道,开了副业,3 个月后副业收入超工资。",
             "讲师学到的 3 件事:",
             "1. 升不上去的人,不是不行,是赛道没选对",
             "2. 问'如果不受限'能打开新空间",
             "3. 不替员工判断",
         ]),
        ("案例 5:小陈加班(完整剧本)",
         [
             "背景:小陈,95 后,连续 1 个月 996,在茶水间哭。",
             "问题:加班到崩溃,主管不知道。",
             "8 分钟心谈:",
             "1. 找一个安静的地方",
             "2. 说:'我有 8 分钟,你愿意聊聊吗?'",
             "3. 前 4 分钟:只听,不评判",
             "4. 后 4 分钟:命名情绪:'你感到撑不住了对吗?'",
             "5. 结束:问'你需要我帮你什么?'",
             "结果:小陈说出来了,主管才知道他的极限。",
             "讲师学到的 3 件事:",
             "1. 情绪不命名,下属撑到崩溃",
             "2. 8 分钟心谈能救人",
             "3. 主管不知道的事,比知道的更多",
         ]),
        ("案例 6:老张 AI 焦虑(完整剧本)",
         [
             "背景:老张,55 岁,工程师,被 AI 替代了一半工作。",
             "问题:失落,易怒,想提前退休。",
             "GROW+:",
             "G: 你为什么失落?",
             "  '我做了 30 年的事,被 AI 替代了'",
             "R: 现在你做这件事的卡点是什么?",
             "  '没人告诉我该做什么'",
             "O: 如果不受限,你会怎么做?",
             "  '我懂机器,AI 我也懂'",
             "W: 你愿意做哪个?",
             "  '做 AI 培训师'",
             "A: 第一步是什么?",
             "  '下周一找培训经理'",
             "结果:老张开始学 AI 工具,3 个月后成了 AI 培训师。",
             "讲师学到的 3 件事:",
             "1. 被 AI 替代的人,不是没用,是没找到新角色",
             "2. GROW+ 慢,但救人",
             "3. 55 岁也能学新东西",
         ]),
        ("案例 7:小赵创新(完整剧本)",
         [
             "背景:小赵,推 AI 项目被否 3 次,想放弃。",
             "问题:被否多次,觉得老板不信任,想离职。",
             "GROW+:",
             "G: 你为什么推这个项目?",
             "  '我相信它能帮团队提效 30%'",
             "R: 前 3 次被否,卡在哪?",
             "  '我没给具体数据,只说有前景'",
             "O: 如果再推 1 次,你会怎么做?",
             "  '做 1 个试点,跑 2 周,给数据'",
             "W: 你愿意做吗?",
             "  '愿意,我想再试 1 次'",
             "A: 第一步是什么?",
             "  '这周选 1 个团队试点'",
             "结果:小赵重启提案,第 4 次通过。",
             "讲师学到的 3 件事:",
             "1. 被否多次的下属,不是没努力,是看不见自己",
             "2. GROW+ R 阶段能让人看见自己",
             "3. 放弃前再给 1 次机会",
         ]),
        ("案例 8:老周临退休(完整剧本)",
         [
             "背景:老周,58 岁,5 年后退休,突然空虚。",
             "问题:工作没意思,想提前退。",
             "8 分钟心谈:",
             "1. 找一个安静的地方",
             "2. 说:'我有 8 分钟,你愿意聊聊吗?'",
             "3. 前 4 分钟:只听",
             "4. 后 4 分钟:翻译:'听起来,你想留下点什么'",
             "5. 结束:问'你需要我帮你什么?'",
             "结果:老周申请做导师,带新人,找到意义感。",
             "讲师学到的 3 件事:",
             "1. 临退休的人,缺的不是钱,是意义",
             "2. 8 分钟心谈 + 翻译,胜过 10 次建议",
             "3. 老员工是宝藏",
         ]),
    ]

    for title, lines in case_scripts:
        add_h2(doc, title)
        for line in lines:
            add_script(doc, line)

    add_page_break(doc)

    # 附录
    add_h1(doc, "附录:6 大金句卡片")
    add_body(doc, "贴在教室四周,每天看 1 次。")
    add_quote(doc, "1. 谁能兼容谁,谁就能领导谁。", "—— 兼容")
    add_quote(doc, "2. 真实感而不是真实。", "—— 真实感")
    add_quote(doc, "3. 讨喜而不是讨好。", "—— 讨喜")
    add_quote(doc, "4. 领先半步,吃尽红利。", "—— 领先半步")
    add_quote(doc, "5. 延时摄影——和时间做朋友。", "—— 延时")
    add_quote(doc, "6. 善借力打力。", "—— 借力")

    add_body(doc, "")
    add_body(doc, "")
    add_quote(doc, "延迟摄影——和时间做朋友。", "—— 第一次讲,允许自己不完美")
    add_quote(doc, "30 天后,你会感谢今天坚持的自己。", "—— 罗老师寄语")

    output_path = os.path.join(OUT_DIR, "教练技术_讲师手册.docx")
    doc.save(output_path)
    print(f"[OK] 讲师手册生成完成: {output_path}")


if __name__ == "__main__":
    build_instructor_manual()

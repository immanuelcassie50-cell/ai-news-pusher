# -*- coding: utf-8 -*-
"""
创建变革成果固化机制课程练习材料
10个Word文档生成脚本
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配色方案
PRIMARY = RGBColor(0, 51, 102)      # 深蓝色
SECONDARY = RGBColor(51, 51, 51)    # 深灰色
ACCENT = RGBColor(0, 112, 192)      # 亮蓝色
LIGHT_BG = RGBColor(240, 248, 255)  # 浅蓝背景

OUTPUT_DIR = r'D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/07-练习材料'

def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level):
    """添加带样式标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = PRIMARY
        run.font.name = 'Microsoft YaHei'
        if level == 0:
            run.font.size = Pt(26)
        elif level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
    return heading

def add_info_box(doc, title, content_lines):
    """添加信息框"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'E6F2FF')

    p = cell.paragraphs[0]
    run = p.add_run(f'{title}：')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    for line in content_lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Microsoft YaHei'
    return table

def add_table_with_headers(doc, headers, rows, header_color='003366'):
    """添加带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 表头
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'

    # 数据行
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
    return table

def add_page_header_footer(doc, title):
    """添加页眉页脚"""
    section = doc.sections[0]

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.text = title
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(9)
        run.font.name = 'Microsoft YaHei'

    # 页脚
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer_para.add_run('第 ')
    run.font.size = Pt(9)
    run.font.name = 'Microsoft YaHei'

    # 页码字段
    run = footer_para.add_run()
    fldChar = run._r.makeelement(qn('w:fldChar'), {})
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = run._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    run._r.append(instrText)

    fldChar2 = run._r.makeelement(qn('w:fldChar'), {})
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)

    run2 = footer_para.add_run('1')
    run2.font.size = Pt(9)

    fldChar3 = run2._r.makeelement(qn('w:fldChar'), {})
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar3)

    run = footer_para.add_run(' 页')
    run.font.size = Pt(9)

def create_cover_page(doc, title, subtitle, module_name, duration, exercise_type):
    """创建封面页"""
    # 空白行
    for _ in range(6):
        doc.add_paragraph()

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(subtitle)
    run.font.size = Pt(20)
    run.font.color.rgb = SECONDARY
    run.font.name = 'Microsoft YaHei'

    for _ in range(4):
        doc.add_paragraph()

    # 模块信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f'【{module_name}】')
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 时长和类型
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_para.add_run(f'练习时长：{duration}  |  练习类型：{exercise_type}')
    run.font.size = Pt(12)
    run.font.color.rgb = SECONDARY
    run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

def add_exercise_info_section(doc, objectives, prep, duration, fill_guide):
    """添加练习信息部分"""
    add_styled_heading(doc, '练习信息', 1)

    # 练习目标
    p = doc.add_paragraph()
    run = p.add_run('练习目标')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    for obj in objectives:
        p = doc.add_paragraph()
        p.add_run(f'• {obj}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 练习准备
    p = doc.add_paragraph()
    run = p.add_run('练习准备')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    for item in prep:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 时长
    p = doc.add_paragraph()
    run = p.add_run('练习时长')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'
    run = p.add_run(f'：{duration}')
    run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 填写说明
    p = doc.add_paragraph()
    run = p.add_run('详细填写说明')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    for guide in fill_guide:
        p = doc.add_paragraph()
        p.add_run(guide)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

# ============== 模块一练习 ==============

def create_exercise_1():
    """练习-变革倒退案例分析"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-变革倒退案例分析',
        '识别变革固化的关键失败因素',
        '模块一：变革成果固化认知',
        '45分钟',
        '案例分析'

    )

    add_exercise_info_section(doc,
        objectives=[
            '识别导致变革倒退的关键因素',
            '分析组织变革固化的常见误区',
            '掌握变革成果流失的预警信号',
            '理解固化机制缺失的后果',
        ],
        prep=[
            '提前阅读案例背景材料',
            '准备变革管理相关理论笔记',
            '打印或准备讨论记录表',
        ],
        duration='45分钟',
        fill_guide=[
            '1. 先通读案例背景，理解变革的完整过程',
            '2. 逐一分析讨论问题，结合案例细节进行思考',
            '3. 小组讨论时，尝试从不同利益相关方角度分析',
            '4. 每道讨论题建议用10-15分钟深入讨论',
            '5. 记录讨论要点时，注意捕捉关键洞察',
            '6. 答案要点提示仅供参考，鼓励提出创新观点',
        ]
    )

    # 案例背景
    add_styled_heading(doc, '案例背景', 1)

    p = doc.add_paragraph()
    run = p.add_run('某企业精益生产变革案例')
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    case_text = """
某制造业企业（华新制造有限公司）在2019年启动了精益生产变革项目。项目启动初期，效果显著：生产效率提升25%，库存周转率提高40%，质量问题减少60%。然而，6个月后，这些成果逐渐开始退化。到项目启动一周年时，大部分改善措施已经名存实亡。

具体情况如下：
• 生产线的"5S"管理退回变革前状态，工具和物料又开始随意摆放
• 每日站会从最初的准时召开变成偶尔召开，再到基本取消
• 问题升级机制被绕过，基层问题直接被压下来不向上反馈
• 关键岗位的变革骨干相继离职或被调离
• 新入职员工不清楚新的生产流程，仍按老方式工作
• 管理层将注意力转移到其他项目，变革项目组合几乎被遗忘

变革推进小组在复盘时发现以下关键事件：
1. 项目启动第3个月，变革小组核心成员张经理被调往新业务线
2. 项目启动第5个月，公司面临紧急订单压力，部分改善措施被暂时搁置
3. 项目启动第8个月，负责精益生产的李总监提出离职
4. 公司在第10个月启动了另一个战略项目，资源被大幅转移
    """
    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('•'):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 讨论问题
    add_styled_heading(doc, '讨论问题', 1)

    questions = [
        ('问题1：变革倒退原因分析', [
            '导致华新公司精益生产变革成果倒退的根本原因是什么？',
            '从人员、流程、制度、文化四个维度进行分析',
            '哪些因素是管理层可以控制的？哪些是外部因素？',
        ]),
        ('问题2：识别固化缺失点', [
            '在案例中，哪些时间节点是固化机制缺失的关键点？',
            '如果你是变革负责人，在第3个月张经理调离时，你会采取什么措施？',
            '继任者培养计划在防止变革倒退中扮演什么角色？',
        ]),
        ('问题3：预警信号识别', [
            '案例中出现了哪些预警信号表明变革成果可能无法持续？',
            '这些预警信号在早期有哪些表现？',
            '建立预警机制对于变革成果固化有何重要意义？',
        ]),
        ('问题4：资源分配问题', [
            '当公司启动新战略项目后，原有变革项目的资源被转移，这反映了什么问题？',
            '如何在组织层面确保变革项目获得持续资源支持？',
            '变革成果固化需要什么样的资源保障机制？',
        ]),
        ('问题5：提出改进建议', [
            '基于案例分析，如果你是华新公司的变革顾问，你会提出哪些具体的改进建议？',
            '这些建议如何体现"固化机制设计"的核心原则？',
            '请为华新公司设计一套防止变革倒退的行动计划。',
        ]),
    ]

    for q_title, q_items in questions:
        add_styled_heading(doc, q_title, 2)
        for i, item in enumerate(q_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 小组讨论指引
    add_styled_heading(doc, '小组讨论指引', 1)

    guide_items = [
        '讨论时间分配：每个问题建议10-15分钟，总计不超过60分钟',
        '角色分工：建议指定主持人、记录员、汇报人各一名',
        '讨论原则：鼓励不同观点碰撞，避免过早达成共识',
        '记录要求：使用提供的讨论记录表，记录关键洞察和分歧点',
        '汇报准备：讨论结束后，汇报人准备3分钟的口头汇报',
    ]

    for item in guide_items:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 讨论记录表
    add_styled_heading(doc, '讨论记录表', 2)

    headers = ['问题', '关键洞察', '分歧点', '结论']
    rows = [
        ['问题1', '', '', ''],
        ['问题2', '', '', ''],
        ['问题3', '', '', ''],
        ['问题4', '', '', ''],
        ['问题5', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 答案要点提示
    add_styled_heading(doc, '答案要点提示', 1)

    answers = [
        ('问题1：变革倒退原因分析', [
            '根本原因：缺乏系统性的固化机制，仅依赖变革小组的持续推动',
            '人员维度：核心人员调离，继任者培养缺失，知识传承断裂',
            '流程维度：新流程未与绩效考核挂钩，缺乏刚性约束',
            '制度维度：精益生产规范未固化为公司制度',
            '文化维度：问题透明文化未建立，基层不敢上报问题',
        ]),
        ('问题2：固化缺失点识别', [
            '第3个月张经理调离：未建立AB角机制，关键知识未沉淀',
            '第5个月紧急订单：未建立"改善与运营平衡"机制',
            '第8个月李总监离职：继任者培养计划缺失',
            '第10个月资源转移：未建立变革成果保障的长效机制',
        ]),
        ('问题3：预警信号', [
            '人员流失：核心骨干相继离开',
            '活动减少：站会频率降低直至取消',
            '规范退化：5S管理倒退',
            '沟通减少：问题反馈机制失灵',
            '关注度下降：管理层注意力转移',
        ]),
        ('问题4：资源分配问题', [
            '反映了组织缺乏对变革成果的战略性重视',
            '需要建立"变革成果保障"纳入组织战略规划',
            '建议设立专职的持续改善部门或岗位',
        ]),
        ('问题5：改进建议', [
            '建立固化机制：制度化、考核化、文化化',
            '完善继任计划：关键岗位AB角、知识传承机制',
            '设置预警指标：定期评估固化度',
            '资源保障机制：变革成果维护预算',
        ]),
    ]

    for a_title, a_items in answers:
        add_styled_heading(doc, a_title, 2)
        for item in a_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(11)
        doc.add_paragraph()

    output_path = f'{OUTPUT_DIR}/练习-变革倒退案例分析.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def create_exercise_2():
    """练习-固化度自诊断实操"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-固化度自诊断实操',
        '掌握固化度自诊断工具的使用方法',
        '模块一：变革成果固化认知',
        '60分钟',
        '实操练习'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解固化度自诊断工具的设计原理',
            '掌握变革成果固化度评估的方法论',
            '能够独立使用自诊断表进行组织评估',
            '能够根据诊断结果制定改进计划',
        ],
        prep=[
            '准备《变革成果固化度自诊断表》',
            '选择一个熟悉的变革项目作为分析对象',
            '准备变革项目的相关数据和文档',
        ],
        duration='60分钟',
        fill_guide=[
            '1. 练习前先了解自诊断表的评估维度和评分标准',
            '2. 结合实际变革项目，认真评估每个指标的实际状态',
            '3. 评分时注意：1=完全未固化，5=完全固化',
            '4. 撰写分析报告时，要透过分数看到本质问题',
            '5. 改进建议要具体可执行，避免空洞的口号',
        ]
    )

    # 自诊断工具说明
    add_styled_heading(doc, '固化度自诊断工具说明', 1)

    p = doc.add_paragraph()
    run = p.add_run('工具概述')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    p = doc.add_paragraph()
    p.add_run('变革成果固化度自诊断表是一套用于评估组织变革成果持续性的工具。它从五个核心维度对变革成果的固化程度进行评估，帮助组织识别固化的薄弱环节，制定针对性的改进措施。')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 评估维度
    add_styled_heading(doc, '评估维度', 2)

    dimensions = [
        ('维度一：制度固化', '评估新流程、新规范是否已纳入组织正式制度体系'),
        ('维度二：人员固化', '评估关键岗位是否有人掌握新流程，继任者培养是否到位'),
        ('维度三：行为固化', '评估组织成员是否将新流程转化为日常行为习惯'),
        ('维度四：文化固化', '评估新理念是否成为组织文化的一部分'),
        ('维度五：保障机制', '评估是否有持续的資源投入和监督机制'),
    ]

    headers = ['维度', '说明']
    rows = [[d[0], d[1]] for d in dimensions]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 评分标准
    add_styled_heading(doc, '评分标准', 2)

    score_table = [
        ['1分', '完全未固化', '新流程未有效执行或已废弃'],
        ['2分', '初步尝试', '有尝试但执行不稳定，经常倒退'],
        ['3分', '基本固化', '在正常情况下能执行，但存在压力下倒退风险'],
        ['4分', '较好固化', '执行稳定，具备一定的抗干扰能力'],
        ['5分', '完全固化', '成为理所当然的习惯和文化，无需刻意维持'],
    ]

    headers = ['评分', '等级', '说明']
    add_table_with_headers(doc, headers, score_table)

    doc.add_paragraph()

    # 实操步骤
    add_styled_heading(doc, '实操步骤', 1)

    steps = [
        ('第一步：选择评估对象', [
            '选择一个你亲身参与或深入了解的变革项目',
            '明确该变革项目的核心成果是什么',
            '确认评估的范围和时间节点（如：项目结束6个月后的固化度）',
        ]),
        ('第二步：逐项评估', [
            '打开《变革成果固化度自诊断表》',
            '针对每个评估指标，根据实际情况打分（1-5分）',
            '每个维度选择2-3个代表性指标进行深度评估',
            '记录评分理由和支撑证据',
        ]),
        ('第三步：分析评分结果', [
            '计算各维度的平均分',
            '识别得分最低的维度（固化最薄弱环节）',
            '分析各维度之间的关联性（如：文化固化影响行为固化）',
            '撰写维度分析报告，说明得分背后的原因',
        ]),
        ('第四步：制定改进计划', [
            '针对每个低分维度，提出具体的改进措施',
            '明确改进措施的优先级（紧急/重要）',
            '设定改进目标（3个月后预期达到的分数）',
            '指定改进责任人',
        ]),
        ('第五步：汇报与反馈', [
            '准备3-5分钟的汇报，向小组成员说明诊断结果',
            '接受他人的质疑和建议',
            '根据反馈修正分析结论和改进计划',
        ]),
    ]

    for step_title, step_items in steps:
        add_styled_heading(doc, step_title, 2)
        for i, item in enumerate(step_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 注意事项
    add_styled_heading(doc, '注意事项', 1)

    cautions = [
        '评分应基于客观事实而非主观感受，避免美化现状',
        '如果某项变革从未真正推行过，不能因为"计划过"就给高分',
        '固化度评估需要有一定的时间跨度，项目刚结束时评估意义有限',
        '跨部门变革项目需要多方印证，避免单一视角的偏差',
        '低分不是耻辱，而是改进的机会，要以成长心态面对诊断结果',
    ]

    for item in cautions:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 练习成果模板
    add_styled_heading(doc, '练习成果模板', 1)

    # 成果1：自诊断表
    add_styled_heading(doc, '成果1：固化度自诊断表', 2)

    headers = ['评估维度', '具体指标', '评分(1-5)', '评分理由']
    rows = [
        ['制度固化', '流程规范已文件化', '', ''],
        ['制度固化', '制度与绩效考核挂钩', '', ''],
        ['人员固化', '关键岗位有人掌握新流程', '', ''],
        ['人员固化', '已建立继任者培养机制', '', ''],
        ['行为固化', '日常工作按新流程执行', '', ''],
        ['行为固化', '问题出现时能自我纠偏', '', ''],
        ['文化固化', '新理念获得广泛认同', '', ''],
        ['文化固化', '成为招聘和培训内容', '', ''],
        ['保障机制', '有持续的资源投入', '', ''],
        ['保障机制', '有定期评估和改进机制', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    # 成果2：分析报告
    add_styled_heading(doc, '成果2：固化度分析报告', 2)

    report_sections = [
        ('一、变革项目概况', [
            '项目名称：________________',
            '核心变革成果：_____________',
            '评估时间节点：_____________',
        ]),
        ('二、总体评分', [
            '制度固化平均分：___________',
            '人员固化平均分：___________',
            '行为固化平均分：___________',
            '文化固化平均分：___________',
            '保障机制平均分：___________',
            '综合固化度得分：___________',
        ]),
        ('三、低分维度分析', [
            '得分最低的维度：___________',
            '主要原因分析：_____________',
            '对整体固化的影响：_________',
        ]),
        ('四、改进计划', [
            '改进措施1：________________ 负责人：____ 时间：____',
            '改进措施2：________________ 负责人：____ 时间：____',
            '改进措施3：________________ 负责人：____ 时间：____',
        ]),
    ]

    for sec_title, sec_items in report_sections:
        p = doc.add_paragraph()
        run = p.add_run(sec_title)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        run.font.name = 'Microsoft YaHei'

        for item in sec_items:
            p = doc.add_paragraph()
            p.add_run(item)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    output_path = f'{OUTPUT_DIR}/练习-固化度自诊断实操.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

# ============== 模块二练习 ==============

def create_exercise_3():
    """练习-制度文本编写实战"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-制度文本编写实战',
        '掌握将流程固化为制度文本的技巧',
        '模块二：制度固化机制',
        '90分钟',
        '实战演练'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解制度文本的结构和要素',
            '掌握将流程转化为制度文本的方法',
            '能够编写规范的公司制度文本',
            '能够设计配套的考核指标',
        ],
        prep=[
            '阅读某企业客服流程优化案例',
            '了解制度文本的基本格式和规范',
            '准备制度文本编写模板',
        ],
        duration='90分钟',
        fill_guide=[
            '1. 先仔细阅读案例材料，理解客服流程优化的背景和内容',
            '2. 按照任务顺序逐步完成：先大纲、后正文、再考核',
            '3. 制度大纲要覆盖完整，体现层次结构',
            '4. 制度正文要用词准确、表述清晰、可操作性强',
            '5. 考核指标要量化、可测量、与制度条款对应',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业客服流程优化', 1)

    case_text = """
【公司背景】
华通电商是一家年销售额超过10亿元的电商平台，拥有超过500人的客服团队，日均处理客户咨询量超过10,000次。客服部门长期存在响应速度慢、问题解决率低、客户满意度不高等问题。

【问题诊断】
经过流程诊断，发现以下关键问题：
1. 客户咨询需要多次转接，平均转接次数达2.5次
2. 客服代表各自为战，缺乏标准化的处理流程
3. 复杂问题缺乏升级机制，导致客户等待时间长
4. 服务质量缺乏量化考核，客服代表缺乏改进动力
5. 新员工培训周期长，技能参差不齐

【优化方案】
1. 建立标准化的"首问负责制"流程
   - 首次接单的客服代表负责全程跟进直至问题解决
   - 需要转接时，由客服代表发起升级申请
   - 升级后仍由原客服代表负责闭环确认

2. 建立三级问题升级机制
   - 一线客服：处理简单咨询（响应时间<2分钟）
   - 二线客服：处理复杂问题（响应时间<15分钟）
   - 专家团队：处理特殊问题（响应时间<1小时）

3. 建立量化考核指标
   - 首次响应时长
   - 问题解决率
   - 一次性解决率
   - 客户满意度评分
   - 平均处理时长

4. 建立知识共享平台
   - 常见问题标准答案库
   - 复杂案例分享机制
   - 每周案例复盘会
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line and not line.startswith('-'):
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操任务
    add_styled_heading(doc, '实操任务', 1)

    tasks = [
        ('任务1：编写制度大纲', [
            '根据案例材料，编写《客户服务标准流程管理办法》的大纲',
            '大纲应包含：总则、组织职责、流程规范、考核办法、附则等章节',
            '注意章节之间的逻辑关系和层次结构',
            '用表格形式呈现大纲结构',
        ], 20),
        ('任务2：编写制度正文', [
            '从大纲中选择3个核心章节，编写详细的制度正文',
            '重点编写"首问负责制"和"问题升级机制"两节',
            '正文要包含：目的、适用范围、具体流程、禁止事项等',
            '语言要准确、简洁、可操作，避免模糊表述',
        ], 40),
        ('任务3：设计考核指标', [
            '为《客户服务标准流程管理办法》设计配套考核指标',
            '考核指标应包含：指标名称、计算公式、目标值、数据来源、考核周期',
            '至少设计5个核心考核指标',
            '确保指标可量化、可测量、与制度条款对应',
        ], 30),
    ]

    for task_title, task_items, time_min in tasks:
        add_styled_heading(doc, f'{task_title}（建议时间：{time_min}分钟）', 2)
        for i, item in enumerate(task_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 评分标准
    add_styled_heading(doc, '评分标准', 1)

    criteria = [
        ('结构完整性 (20分)', [
            '大纲结构完整，层次清晰',
            '包含总则、组织职责、流程规范、考核办法、附则等必要章节',
            '章节之间逻辑关系合理',
        ]),
        ('内容规范性 (30分)', [
            '制度正文格式规范，用词准确',
            '流程描述清晰，可操作性强',
            '覆盖正常流程和异常情况处理',
        ]),
        ('考核指标设计 (30分)', [
            '指标量化、可测量',
            '与制度条款对应',
            '目标值设定合理',
        ]),
        ('创新与实用 (20分)', [
            '有独特的改进思路或优化建议',
            '方案具有实际可行性',
            '考虑到了实施中可能遇到的阻力',
        ]),
    ]

    for crit_title, crit_items in criteria:
        add_styled_heading(doc, crit_title, 2)
        for item in crit_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    # 优秀范例
    add_styled_heading(doc, '优秀范例（制度大纲示例）', 1)

    headers = ['章节', '主要内容', '关键条款']
    rows = [
        ['第一章 总则', '目的、适用范围、名词定义', '首问负责制定义'],
        ['第二章 组织职责', '客服部、投诉处理组、专家团队职责', '各层级职责边界'],
        ['第三章 流程规范', '首问负责制流程、升级机制、服务标准', '流程图+文字描述'],
        ['第四章 考核办法', '考核指标、权重、数据来源', '5个核心指标'],
        ['第五章 附则', '解释权、生效日期、修订记录', '版本管理'],
    ]
    add_table_with_headers(doc, headers, rows)

    output_path = f'{OUTPUT_DIR}/练习-制度文本编写实战.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def create_exercise_4():
    """练习-考核指标重新设计"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-考核指标重新设计',
        '掌握将新流程转化为考核指标的方法',
        '模块二：制度固化机制',
        '60分钟',
        '实战演练'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解考核指标设计的基本原则',
            '掌握将新流程转化为量化考核指标的方法',
            '能够设计合理的指标权重和目标值',
            '能够建立考核数据收集机制',
        ],
        prep=[
            '阅读某企业供应链变革案例',
            '了解平衡计分卡等考核工具',
            '准备考核指标设计模板',
        ],
        duration='60分钟',
        fill_guide=[
            '1. 先分析案例中供应链变革的核心内容和新流程',
            '2. 梳理关键流程节点，识别需要考核的关键行为',
            '3. 设计指标时注意：量化优先、可测量、与流程对应',
            '4. 权重设计要考虑流程优先级和改善空间',
            '5. 每项指标都要明确数据来源，确保可执行',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业供应链变革', 1)

    case_text = """
【公司背景】
华泰制造是一家年产值30亿元的汽车零部件制造商，主要为整车厂提供发动机零部件、底盘零部件等产品。公司拥有完整的供应链体系，但近年来面临交付及时率下降、库存成本上升、质量问题频发等挑战。

【变革背景】
2023年，公司启动供应链数字化变革项目，核心目标是：
1. 提升交付及时率从85%到98%
2. 降低库存周转天数从45天到30天
3. 减少供应链质量事故50%
4. 建立供应链协同平台，实现信息实时共享

【新流程设计】
1. 需求预测流程优化
   - 从月度预测改为周度预测
   - 引入AI需求预测模型
   - 建立预测准确性评估机制

2. 采购管理流程优化
   - 实施JIT采购模式
   - 建立供应商分级管理机制
   - 核心供应商纳入VMI（供应商管理库存）

3. 生产计划流程优化
   - 建立柔性生产机制
   - 实施智能排产系统
   - 建立生产异常快速响应机制

4. 仓储物流流程优化
   - 实施智能化仓储管理
   - 建立物流跟踪系统
   - 优化配送路线算法

【配套变革】
- 组织调整：成立供应链控制塔（Supply Chain Control Tower）
- 人员能力：培养供应链数据分析能力
- 绩效考核：重新设计供应链相关岗位的考核指标
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line and not line.startswith('-'):
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操任务
    add_styled_heading(doc, '实操任务', 1)

    tasks = [
        ('任务1：梳理关键流程', [
            '根据案例，梳理供应链变革涉及的核心流程',
            '识别每个流程中的关键行为节点',
            '明确每个关键节点需要考核的行为标准',
        ], 15),
        ('任务2：设计考核维度', [
            '从效率、质量、成本、协同四个维度设计考核框架',
            '每个维度选择2-3个核心指标',
            '考虑使用平衡计分卡的思路设计维度',
        ], 15),
        ('任务3：设定指标值', [
            '为每个考核指标设定目标值和及格线',
            '目标值参考案例中的变革目标',
            '考虑历史数据和行业标杆',
        ], 15),
        ('任务4：制定考核办法', [
            '明确每个指标的数据来源和计算方法',
            '设定考核周期和权重',
            '设计数据收集和汇总流程',
        ], 15),
    ]

    for task_title, task_items, time_min in tasks:
        add_styled_heading(doc, f'{task_title}（建议时间：{time_min}分钟）', 2)
        for i, item in enumerate(task_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 评分标准
    add_styled_heading(doc, '评分标准', 1)

    criteria = [
        ('维度完整性 (25分)', [
            '覆盖效率、质量、成本、协同等核心维度',
            '各维度指标数量分布合理',
        ]),
        ('指标科学性 (35分)', [
            '指标量化、可测量',
            '与流程关键行为对应',
            '计算方法清晰准确',
        ]),
        ('目标合理性 (25分)', [
            '目标值具有挑战性但可达成',
            '参考了案例中的变革目标',
            '及格线设定合理',
        ]),
        ('可执行性 (15分)', [
            '数据来源明确',
            '考核周期适当',
            '权重分配合理',
        ]),
    ]

    for crit_title, crit_items in criteria:
        add_styled_heading(doc, crit_title, 2)
        for item in crit_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    # 参考答案
    add_styled_heading(doc, '参考答案（考核指标框架示例）', 1)

    headers = ['维度', '考核指标', '目标值', '数据来源']
    rows = [
        ['效率', '交付及时率', '≥98%', 'ERP系统'],
        ['效率', '订单响应时长', '≤4小时', '供应链平台'],
        ['效率', '库存周转天数', '≤30天', '财务系统'],
        ['质量', '来料合格率', '≥99.5%', 'QM系统'],
        ['质量', '质量问题响应时长', '≤24小时', 'QMS系统'],
        ['成本', '采购成本下降率', '≥5%', '财务系统'],
        ['成本', '仓储成本占比', '≤3%', '财务系统'],
        ['协同', '信息共享及时率', '≥95%', '供应链平台'],
        ['协同', '预测准确率', '≥85%', '数据分析平台'],
    ]
    add_table_with_headers(doc, headers, rows)

    output_path = f'{OUTPUT_DIR}/练习-考核指标重新设计.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

# ============== 模块三练习 ==============

def create_exercise_5():
    """练习-机制设计工作坊"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-机制设计工作坊',
        '掌握固化机制设计的方法',
        '模块三：保障机制设计',
        '120分钟',
        '团队工作坊'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解固化机制设计的核心要素',
            '掌握利益相关方分析的方法',
            '能够设计多层次的固化机制',
            '体验团队协作进行机制设计的过程',
        ],
        prep=[
            '提前分组，每组4-6人',
            '各组选择一个变革项目作为练习对象',
            '准备大白纸、彩笔等工具',
        ],
        duration='120分钟',
        fill_guide=[
            '1. 工作坊角色：组长（协调）、记录员（记录）、计时员（控制时间）、汇报人（代表发言）',
            '2. 研讨流程按5步进行，每步结束后小组内快速确认再进入下一步',
            '3. 固化机制设计表是核心产出，每组需完成完整填写',
            '4. 汇报时间控制在5分钟内，突出重点和创新点',
            '5. 其他小组可以提问或补充，每个问题控制在1分钟内',
        ]
    )

    # 角色分工
    add_styled_heading(doc, '角色分工说明', 1)

    roles = [
        ('组长', '负责协调讨论、把握方向、时间控制、确保达成共识'),
        ('记录员', '负责在白纸/模板上记录关键观点，使用便签纸管理分歧点'),
        ('计时员', '负责提醒各环节时间节点，控制讨论节奏'),
        ('汇报人', '负责整理讨论成果，准备5分钟口头汇报'),
        ('促进者（可选）', '负责提问挑战，推动深入思考，但不参与决策'),
    ]

    headers = ['角色', '职责']
    rows = [[r[0], r[1]] for r in roles]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    # 研讨流程
    add_styled_heading(doc, '研讨流程（5步法）', 1)

    steps = [
        ('第一步：变革回顾（20分钟）', [
            '回顾本组选择的变革项目背景',
            '明确变革的核心成果是什么',
            '识别当前固化面临的主要挑战',
            '使用问题树分析方法追根溯源',
        ]),
        ('第二步：利益相关方分析（20分钟）', [
            '列出变革涉及的所有利益相关方',
            '分析各方的核心关切和利益诉求',
            '识别可能支持或反对固化机制的力量',
            '思考如何扩大支持力量、转化阻力',
        ]),
        ('第三步：固化机制设计（30分钟）', [
            '针对每个关键挑战，设计固化机制',
            '考虑：制度机制、激励机制、监督机制、文化机制',
            '明确每种机制的设计要点和实施路径',
            '评估各机制的成本和效果',
        ]),
        ('第四步：风险与对策（20分钟）', [
            '识别固化机制可能遇到的阻力',
            '分析每种阻力的严重程度和发生概率',
            '制定针对性的应对策略',
            '建立预警指标和快速响应机制',
        ]),
        ('第五步：行动计划制定（30分钟）', [
            '制定具体的固化机制实施计划',
            '明确责任人、时间节点、资源需求',
            '设定阶段性里程碑和评估标准',
            '准备5分钟汇报材料',
        ]),
    ]

    for step_title, step_items in steps:
        add_styled_heading(doc, step_title, 2)
        for item in step_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 工具模板
    add_styled_heading(doc, '工具模板：固化机制设计表', 1)

    p = doc.add_paragraph()
    p.add_run('核心产出：每组需完成以下固化机制设计表的填写')
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = SECONDARY

    doc.add_paragraph()

    headers = ['挑战描述', '固化机制类型', '机制设计要点', '责任人', '时间节点']
    rows = [
        ['挑战1：...', '', '', '', ''],
        ['挑战2：...', '', '', '', ''],
        ['挑战3：...', '', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '风险与对策表', 2)

    headers = ['风险描述', '可能性', '影响程度', '应对策略']
    rows = [
        ['', '', '', ''],
        ['', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '行动计划表', 2)

    headers = ['行动项', '具体措施', '负责人', '资源需求', '完成时间', '评估标准']
    rows = [
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
        ['', '', '', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 成果输出要求
    add_styled_heading(doc, '成果输出要求', 1)

    outputs = [
        '完成《固化机制设计表》的完整填写',
        '完成《风险与对策表》的填写',
        '完成《行动计划表》的填写',
        '准备5分钟的汇报PPT或口头汇报大纲',
        '每组提交一份电子版工作坊成果文件',
    ]

    for item in outputs:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    # 汇报指引
    add_styled_heading(doc, '汇报指引', 1)

    report_guide = [
        '汇报时间：每组5分钟，超时将被打断',
        '汇报内容：变革项目背景 → 主要挑战 → 固化机制设计 → 创新亮点',
        '汇报顺序：由抽签决定',
        '提问环节：其他小组可提问（每个问题不超过1分钟）',
        '评分标准：机制设计完整性30%、创新性20%、可执行性30%、汇报表现20%',
    ]

    for item in report_guide:
        p = doc.add_paragraph()
        p.add_run(f'• {item}')
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'

    output_path = f'{OUTPUT_DIR}/练习-机制设计工作坊.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def create_exercise_6():
    """练习-继任者培养计划制定"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-继任者培养计划制定',
        '掌握关键岗位继任者培养计划制定方法',
        '模块三：保障机制设计',
        '60分钟',
        '实战演练'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解继任者培养对变革成果固化的重要性',
            '掌握关键岗位继任者培养计划制定的方法',
            '能够识别需要培养继任者的关键岗位',
            '能够设计继任者培养的内容和路径',
        ],
        prep=[
            '阅读某企业技术总监继任案例',
            '准备继任者培养计划模板',
            '选择一个自己组织中的关键岗位作为分析对象',
        ],
        duration='60分钟',
        fill_guide=[
            '1. 先仔细阅读案例，理解技术总监岗位的关键能力和继任挑战',
            '2. 继任者画像要具体，区分"必备条件"和"加分条件"',
            '3. 培养内容要与岗位关键能力对应，避免泛泛而谈',
            '4. 时间规划要合理，考虑加速培养的途径',
            '5. 评审要点用于自检，确保计划完整可执行',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业技术总监继任', 1)

    case_text = """
【公司背景】
华信科技是一家年营收20亿元的软件企业，专注于企业级SaaS产品的研发和服务。公司技术团队超过300人，其中核心技术骨干50余人。

【变革背景】
2022年，公司启动技术平台升级项目，目标是从传统的单体架构转向微服务架构。这是一个为期2年的大型技术变革项目。

【关键人物】
技术总监王强是这次变革的核心推动者：
- 主导了技术架构的设计和选型
- 建立了微服务开发的最佳实践
- 培养了一支能够实施转型的技术团队
- 与业务部门建立了良好的沟通机制

【继任挑战】
王强在2024年初向公司提出离职，计划6个月后离开。公司面临以下挑战：

1. 微服务转型项目只完成了60%，还有40%未完成
2. 关键的技术架构决策需要有人能够接手
3. 核心技术团队中有多名成员是王强亲自培养的
4. 供应商关系和重要的技术合作伙伴需要交接
5. 新技术平台上线后还有大量的运维和优化工作

【内部候选人】
公司内部有两名潜在候选人：

候选人A：张明
- 当前职位：架构师
- 在公司年限：5年
- 优势：技术能力强，对微服务架构有深入理解，参与了架构设计
- 劣势：缺乏团队管理经验，与业务部门的沟通较少

候选人B：李华
- 当前职位：技术经理
- 在公司年限：7年
- 优势：团队管理经验丰富，与业务部门关系好，善于推动跨部门协作
- 劣势：对微服务架构的理解不如张明深入，技术视野有一定局限
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('候选人', '王强在', '张明', '李华')):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操步骤
    add_styled_heading(doc, '实操步骤（4步法）', 1)

    steps = [
        ('第一步：关键岗位识别与继任需求分析（15分钟）', [
            '识别需要继任计划的关键岗位（不只是技术岗位）',
            '明确该岗位对变革成果固化的关键影响',
            '评估继任的紧迫程度和时间窗口',
            '分析继任者需要继承的"知识资产"',
        ]),
        ('第二步：继任者画像与候选人评估（15分钟）', [
            '明确继任者需要具备的关键能力素质',
            '区分"必备条件"和"加分条件"',
            '评估内部候选人的匹配度',
            '识别候选人的能力差距和培养重点',
        ]),
        ('第三步：继任者培养计划制定（20分钟）', [
            '制定针对性的培养内容和学习路径',
            '设计加速培养的措施（如：导师辅导、轮岗、项目历练）',
            '设定培养里程碑和评估标准',
            '考虑AB角机制，降低单点风险',
        ]),
        ('第四步：交接与过渡计划（10分钟）', [
            '制定知识转移计划，确保隐性知识显性化',
            '设计交接清单和过渡期安排',
            '建立过渡期的支持和辅导机制',
            '设定继任成功的评估标准',
        ]),
    ]

    for step_title, step_items in steps:
        add_styled_heading(doc, step_title, 2)
        for i, item in enumerate(step_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 计划模板
    add_styled_heading(doc, '计划模板', 1)

    add_styled_heading(doc, '1. 关键岗位继任需求分析', 2)

    headers = ['分析维度', '内容']
    rows = [
        ['岗位名称', ''],
        ['对变革成果的影响', ''],
        ['继任紧迫程度', ''],
        ['需要继承的知识资产', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '2. 继任者能力要求', 2)

    headers = ['能力维度', '具体要求', '优先级']
    rows = [
        ['技术能力', '', ''],
        ['管理能力', '', ''],
        ['业务理解', '', ''],
        ['人际网络', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '3. 培养计划', 2)

    headers = ['培养内容', '培养方式', '时间安排', '评估方式']
    rows = [
        ['', '', '', ''],
        ['', '', '', ''],
        ['', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '4. 交接计划', 2)

    headers = ['交接事项', '交接内容', '时间节点', '确认人']
    rows = [
        ['', '', '', ''],
        ['', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 评审要点
    add_styled_heading(doc, '评审要点', 1)

    review_points = [
        ('完整性', [
            '是否覆盖了岗位识别、能力要求、培养计划、交接计划等全部要素',
            '是否有明确的时间节点和责任人',
        ]),
        ('针对性', [
            '能力要求是否与岗位实际需求匹配',
            '培养内容是否针对候选人的能力差距设计',
        ]),
        ('可执行性', [
            '培养方式是否可行（如导师是否愿意配合）',
            '时间安排是否与工作安排冲突',
            '资源需求是否在可控范围内',
        ]),
        ('风险考虑', [
            '是否考虑了AB角机制',
            '是否有应急方案',
            '是否识别了关键知识流失风险',
        ]),
    ]

    for rp_title, rp_items in review_points:
        add_styled_heading(doc, rp_title, 2)
        for item in rp_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    output_path = f'{OUTPUT_DIR}/练习-继任者培养计划制定.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

# ============== 模块四练习 ==============

def create_exercise_7():
    """练习-文化固化策略设计"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-文化固化策略设计',
        '掌握文化固化策略设计方法',
        '模块四：文化固化机制',
        '90分钟',
        '策略设计'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解文化固化的深层机制',
            '掌握从行为到文化的固化路径',
            '能够识别推动文化固化的关键杠杆',
            '能够设计有效的文化固化策略',
        ],
        prep=[
            '阅读某企业服务文化固化案例',
            '了解组织文化研究的经典模型',
            '准备文化固化策略设计模板',
        ],
        duration='90分钟',
        fill_guide=[
            '1. 行为识别要具体，是可观察、可测量的行为而非抽象概念',
            '2. 固化机制要形成闭环：激励机制→行为发生→正向反馈→习惯形成',
            '3. 激励方案要考虑物质与精神的平衡，短期与长期的结合',
            '4. 不同利益相关方可能需要不同的激励方式',
            '5. 评分标准用于自检，确保策略完整有效',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业服务文化固化', 1)

    case_text = """
【公司背景】
华悦酒店是一家拥有30家连锁门店的中高端酒店品牌。公司成立10年来，以优质的服务著称，但近年来随着规模扩张，服务质量参差不齐的问题日益突出。

【文化变革目标】
公司提出"让每一次服务都成为美好回忆"的服务理念，目标是在一年内实现：
1. 客户满意度从85%提升到95%
2. 投诉率下降50%
3. 打造3-5个标杆门店，形成可复制的服务模式

【服务文化内涵】
经过提炼，华悦的服务文化包含以下核心要素：
- 主动服务：预见客户需求，在客户开口前提供服务
- 真诚关怀：像对待家人一样对待每一位客户
- 专业细节：关注服务的每一个细节，追求零失误
- 快速响应：对客户需求第一时间响应，不过度承诺

【文化推广困境】
在推广服务文化的过程中，管理层遇到以下挑战：

1. 服务行为难以标准化
   - 员工认为"服务是艺术，无法标准化"
   - 不同员工对"好服务"的理解不同
   - 客户需求多样，难以用统一标准应对

2. 激励效果不明显
   - 现有激励以物质奖励为主，但效果递减
   - 服务标兵的评选沦为轮流坐庄
   - 负面行为缺乏有效约束

3. 管理层言行不一
   - 部分店长对服务文化重视不足
   - 忙时服务质量下降被视为正常
   - 考核压力导致短期行为

4. 新员工融入困难
   - 服务文化培训流于形式
   - 老员工的行为与文化宣导存在差距
   - 缺乏持续的 文化强化机制
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操任务
    add_styled_heading(doc, '实操任务', 1)

    tasks = [
        ('任务1：识别目标行为（20分钟）', [
            '将服务文化理念转化为可观察的具体行为',
            '为每个文化要素识别3-5个代表性行为',
            '区分"做对了"和"做错了"的行为表现',
            '确保行为描述具体、可测量、可评估',
        ]),
        ('任务2：设计固化机制（35分钟）', [
            '针对每个核心行为，设计固化机制',
            '固化机制包括：示范机制、学习机制、强化机制、约束机制',
            '确保机制形成闭环，能够持续推动行为发生',
            '考虑不同场景下的机制适配（正常/忙时/突发）',
        ]),
        ('任务3：制定激励方案（35分钟）', [
            '设计物质激励方案（与考核指标挂钩）',
            '设计精神激励方案（荣誉、认可、成长）',
            '考虑短期激励与长期激励的结合',
            '设计负向约束机制（对负面行为的惩罚）',
        ]),
    ]

    for task_title, task_items in tasks:
        add_styled_heading(doc, task_title, 2)
        for i, item in enumerate(task_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 策略设计模板
    add_styled_heading(doc, '策略设计模板', 1)

    add_styled_heading(doc, '文化要素：主动服务', 2)

    headers = ['目标行为', '固化机制', '激励方案', '评估指标']
    rows = [
        ['在客户开口前提供服务', '示范机制：每月服务标兵分享\n强化机制：客户好评及时表彰', '物质：好评积分兑换\n精神：月度服务之星评选', '主动服务次数\n客户表扬信数量'],
        ['预见客户需求并提前准备', '学习机制：新员工导师制\n约束机制：未预见的投诉纳入考核', '物质：预防问题奖励\n精神：最佳预见奖', '需求预见准确率\n预防问题数量'],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '策略设计总览表', 2)

    headers = ['文化要素', '核心行为1', '核心行为2', '核心行为3', '主要固化机制']
    rows = [
        ['主动服务', '', '', '', ''],
        ['真诚关怀', '', '', '', ''],
        ['专业细节', '', '', '', ''],
        ['快速响应', '', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 评分标准
    add_styled_heading(doc, '评分标准', 1)

    criteria = [
        ('行为识别准确性 (25分)', [
            '行为描述具体、可观察、可测量',
            '行为与文化要素对应关系清晰',
            '覆盖正向和负向行为表现',
        ]),
        ('机制设计有效性 (35分)', [
            '固化机制形成闭环',
            '机制考虑了不同场景的适配',
            '机制之间相互协同，形成系统',
        ]),
        ('激励方案吸引力 (25分)', [
            '物质激励与精神激励结合',
            '短期激励与长期激励平衡',
            '负向约束机制明确有效',
        ]),
        ('整体系统性 (15分)', [
            '策略与企业实际匹配',
            '资源投入与效果预期合理',
            '评估指标可操作',
        ]),
    ]

    for crit_title, crit_items in criteria:
        add_styled_heading(doc, crit_title, 2)
        for item in crit_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    output_path = f'{OUTPUT_DIR}/练习-文化固化策略设计.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def create_exercise_8():
    """练习-故事挖掘与传播"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-故事挖掘与传播',
        '掌握故事收集和传播的方法',
        '模块四：文化固化机制',
        '60分钟',
        '实操练习'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解故事在文化固化中的独特作用',
            '掌握挖掘变革故事的方法和技巧',
            '能够撰写有感染力的变革故事',
            '能够设计有效的故事传播方案',
        ],
        prep=[
            '准备故事收集的访谈提纲',
            '选择一个变革项目作为故事挖掘对象',
            '了解不同传播渠道的特点',
        ],
        duration='60分钟',
        fill_guide=[
            '1. 故事挖掘的关键是找到"有血有肉"的细节，而非抽象总结',
            '2. 好的故事要有冲突、有转折、有情感，不只是成绩罗列',
            '3. 撰写故事时注意：从业人员的真实感受为切入点，而非事件概述',
            '4. 传播渠道选择要考虑受众特点和信息接收习惯',
            '5. 一个好故事可以重复使用，但要有不同的讲述版本',
        ]
    )

    # 故事挖掘技巧
    add_styled_heading(doc, '故事挖掘技巧说明', 1)

    tips = [
        ('技巧一：寻找"关键时刻"', [
            '变革中的关键时刻是故事的黄金素材',
            '关键时刻包括：第一次尝试、遇到阻力、突破困境、获得认可',
            '关注那些让当事人情绪波动的事件',
        ]),
        ('技巧二：挖掘"反差细节"', [
            '好的故事往往有反差：前后对比、预期与现实对比',
            '细节包括：具体的数字、具体的场景、具体的人物对话',
            '关注那些"不一样"的细节，而非大而化之的总结',
        ]),
        ('技巧三：关注"情感真实"', [
            '故事要能够引发共鸣，需要有真实的情感',
            '关注当事人的感受：困难时的焦虑、突破时的喜悦、获得认可时的自豪',
            '避免只讲"应该讲的"，要挖掘"真实发生的"',
        ]),
    ]

    for tip_title, tip_items in tips:
        add_styled_heading(doc, tip_title, 2)
        for item in tip_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 故事收集模板
    add_styled_heading(doc, '故事收集模板使用指引', 1)

    p = doc.add_paragraph()
    p.add_run('在挖掘故事时，使用以下模板进行记录：')
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    headers = ['模板要素', '引导问题', '记录区域']
    rows = [
        ['基本信息', '什么时候？在哪里？谁参与？', ''],
        ['事件背景', '为什么这件事会发生？之前是什么状态？', ''],
        ['核心冲突', '遇到了什么困难或阻力？当时最大的挑战是什么？', ''],
        ['关键转折', '是什么让情况发生变化？有没有"顿悟时刻"或关键帮助？', ''],
        ['当事人感受', '当时你的感受是什么？有没有想过放弃？', ''],
        ['成果与影响', '最后的结果是什么？对你和周围人有什么影响？', ''],
        ['意义提炼', '这个故事告诉我们什么道理？有什么可以分享的经验？', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    # 实操任务
    add_styled_heading(doc, '实操任务', 1)

    tasks = [
        ('任务1：挖掘故事（20分钟）', [
            '选择一个你亲身经历的变革故事',
            '使用故事收集模板进行深度挖掘',
            '尝试找到至少3个有感染力的细节',
            '如果自己经历的故事不够丰富，可以访谈他人',
        ]),
        ('任务2：撰写故事（25分钟）', [
            '将挖掘到的素材整理成完整的故事',
            '故事结构：背景→冲突→转折→结局→启示',
            '语言风格：生动、有画面感、情感真实',
            '字数控制在500-800字',
        ]),
        ('任务3：规划传播（15分钟）', [
            '确定故事的目标受众',
            '选择适合的传播渠道',
            '设计传播形式（文字、视频、现场分享等）',
            '制定传播时间表',
        ]),
    ]

    for task_title, task_items in tasks:
        add_styled_heading(doc, task_title, 2)
        for i, item in enumerate(task_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 传播渠道选择指南
    add_styled_heading(doc, '传播渠道选择指南', 1)

    headers = ['渠道类型', '特点', '适用场景', '注意事项']
    rows = [
        ['公司内刊/公众号', '覆盖面广，保存时间长', '重要故事、里程碑事件', '需要审核，内容要规范'],
        ['内部培训分享', '互动性强，可以问答', '可复制的经验、文化传承', '需要好的讲述者'],
        ['年会/大型活动', '影响大，传播广', '标杆故事、激励人心', '时间有限，需要精炼'],
        ['工作群/即时通讯', '传播快，即时性强', '小故事、日常点滴', '容易刷屏，需要控制频率'],
        ['视频/短视频', '生动直观，易于传播', '可视化故事、人物专访', '制作成本高，需要专业支持'],
        ['故事墙/文化墙', '持续可见，营造氛围', '日常文化渗透', '需要定期更新'],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    # 优秀故事范例
    add_styled_heading(doc, '优秀故事范例', 1)

    p = doc.add_paragraph()
    run = p.add_run('《那个通宵的夜晚》')
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Microsoft YaHei'

    sample_story = """
背景：2023年3月，公司ERP系统上线前夕，项目组面临巨大压力。

冲突：上线前两天，核心模块突然出现严重bug，可能导致数据丢失。整个团队陷入焦虑。

转折：项目经理老张没有责备任何人，而是带领大家通宵排查。从晚上8点到第二天早上6点，终于找到了问题根源——是一个隐藏的数据兼容性问题。

细节：凌晨3点，有人提议放弃，回滚到旧系统。老张说："再给我两个小时，如果找不到，我请大家吃早餐，然后我们一起做决定。"凌晨5点，当问题终于定位时，办公室里响起了欢呼声。

结局：系统按时上线，之后运行稳定。这个故事后来成为公司"永不放弃"精神的代表。

启示：有时候，团队的凝聚力就是在困难时刻体现的。
    """

    p = doc.add_paragraph()
    for line in sample_story.strip().split('\n'):
        line = line.strip()
        if line.startswith(('背景', '冲突', '转折', '细节', '结局', '启示')):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = ACCENT
            run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    output_path = f'{OUTPUT_DIR}/练习-故事挖掘与传播.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

# ============== 模块五练习 ==============

def create_exercise_9():
    """练习-固化效果评估实操"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-固化效果评估实操',
        '掌握固化效果评估的方法',
        '模块五：固化效果评估与预警',
        '60分钟',
        '实操练习'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解固化效果评估的核心维度',
            '掌握固化效果评估的方法和工具',
            '能够设计和实施固化效果评估方案',
            '能够根据评估结果提出改进建议',
        ],
        prep=[
            '阅读某企业精益变革固化评估案例',
            '准备固化效果评估报告模板',
            '选择一个变革项目作为评估对象',
        ],
        duration='60分钟',
        fill_guide=[
            '1. 评估前先明确评估目的和评估范围',
            '2. 选择评估指标时要考虑数据可得性',
            '3. 数据收集要客观，避免美化现状',
            '4. 分析要深入，透过数据看到本质问题',
            '5. 改进建议要具体可执行，避免空洞',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业精益变革固化评估', 1)

    case_text = """
【公司背景】
华力工厂是一家拥有2000名员工的汽车零部件制造企业。2022年，公司实施了精益生产变革项目，项目于2023年6月结束。

【评估背景】
变革项目结束6个月后，公司决定对精益变革成果的固化情况进行评估。

【评估范围】
- 评估对象：精益生产变革的核心成果
- 评估时间：2023年12月（项目结束后6个月）
- 评估团队：运营管理部牵头，人力资源部配合

【评估方法】
1. 数据分析
   - 收集精益生产核心指标的变化数据
   - 对比变革期间的高点和当前状态
   - 分析关键指标的走势

2. 现场观察
   - 对2条生产线进行为期2天的现场观察
   - 记录5S管理、标准化作业、持续改善活动的执行情况
   - 与一线员工访谈，了解真实执行状态

3. 问卷调查
   - 对50名相关岗位员工进行问卷调查
   - 调查内容包括：精益工具掌握程度、变革认知、改进参与度等

【评估发现】
1. 核心指标变化
   - 生产效率：变革期间提升25%，目前维持18%的提升（部分倒退）
   - 库存周转：变革期间提升40%，目前维持35%的提升
   - 质量不良率：变革期间下降60%，目前下降45%（部分倒退）

2. 现场观察发现
   - 5S管理：最初3个月执行较好，目前有所松懈
   - 每日站会：频率从每天下降到每周2-3次
   - 问题看板：更新不及时，部分问题长期挂起

3. 调查问卷发现
   - 85%的员工了解精益生产的基本理念
   - 但只有45%的员工能够熟练使用精益工具
   - 只有30%的员工表示经常参与持续改善活动

【根本原因分析】
1. 激励机制未及时调整：精益改善成果未与绩效考核挂钩
2. 骨干人员流失：3名变革核心成员在项目结束后离职
3. 管理层关注度下降：项目结束后，管理层精力的重心转移
4. 新员工培训不足：新入职员工对精益生产了解不足
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操步骤
    add_styled_heading(doc, '实操步骤（5步法）', 1)

    steps = [
        ('第一步：明确评估目的和范围（10分钟）', [
            '确定要评估的变革项目和核心成果',
            '明确评估的时间节点和范围',
            '确定评估团队成员和分工',
            '准备评估所需的资源和工具',
        ]),
        ('第二步：设计评估方案（15分钟）', [
            '选择评估维度和指标',
            '确定数据收集方法',
            '设计问卷或访谈提纲',
            '制定评估时间表',
        ]),
        ('第三步：实施数据收集（15分钟）', [
            '收集定量数据（指标数据、系统数据）',
            '进行现场观察和访谈',
            '发放和回收问卷',
            '记录关键发现和证据',
        ]),
        ('第四步：分析与报告（15分钟）', [
            '整理和分析收集到的数据',
            '识别主要发现和根本原因',
            '与历史数据或基准对比',
            '撰写评估报告',
        ]),
        ('第五步：提出改进建议（5分钟）', [
            '针对发现的问题提出改进建议',
            '明确改进措施的优先级',
            '指定改进责任人',
            '设定改进目标和时间节点',
        ]),
    ]

    for step_title, step_items in steps:
        add_styled_heading(doc, step_title, 2)
        for i, item in enumerate(step_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 评估报告模板
    add_styled_heading(doc, '评估报告模板', 1)

    sections = [
        ('一、评估概述', [
            '评估目的：________________',
            '评估对象：________________',
            '评估时间：________________',
            '评估方法：________________',
        ]),
        ('二、核心指标评估', [
            '【填写说明】使用下表记录核心指标的评估结果',
        ]),
        ('三、深度分析', [
            '主要发现：________________',
            '根本原因分析：____________',
            '与预期的差距：____________',
        ]),
        ('四、改进建议', [
            '建议1：________________ 负责人：____ 时间：____',
            '建议2：________________ 负责人：____ 时间：____',
            '建议3：________________ 负责人：____ 时间：____',
        ]),
    ]

    for sec_title, sec_items in sections:
        p = doc.add_paragraph()
        run = p.add_run(sec_title)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        run.font.name = 'Microsoft YaHei'

        for item in sec_items:
            p = doc.add_paragraph()
            p.add_run(item)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_paragraph()

    add_styled_heading(doc, '核心指标评估表', 2)

    headers = ['指标名称', '变革期间峰值', '当前值', '变化幅度', '评估结论']
    rows = [
        ['', '', '', '', ''],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_page_break()

    # 评审标准
    add_styled_heading(doc, '评审标准', 1)

    criteria = [
        ('评估方案完整性 (20分)', [
            '评估目的和范围明确',
            '评估方法选择适当',
            '评估指标覆盖全面',
        ]),
        ('数据收集有效性 (30分)', [
            '数据来源可靠',
            '样本量足够',
            '证据链完整',
        ]),
        ('分析深度 (30分)', [
            '数据解读准确',
            '原因分析深入',
            '不回避敏感问题',
        ]),
        ('建议可行性 (20分)', [
            '建议具体可执行',
            '责任人和时间明确',
            '与问题对应',
        ]),
    ]

    for crit_title, crit_items in criteria:
        add_styled_heading(doc, crit_title, 2)
        for item in crit_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    output_path = f'{OUTPUT_DIR}/练习-固化效果评估实操.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def create_exercise_10():
    """练习-预警指标设计与干预"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    add_page_header_footer(doc, '变革管理练习材料')

    create_cover_page(doc,
        '练习-预警指标设计与干预',
        '掌握预警指标设计和干预策略制定的方法',
        '模块五：固化效果评估与预警',
        '90分钟',
        '实战演练'
    )

    add_exercise_info_section(doc,
        objectives=[
            '理解预警指标在固化保障中的作用',
            '掌握预警指标设计的方法',
            '能够设定合理的预警阈值',
            '能够制定有效的干预策略',
        ],
        prep=[
            '阅读某企业流程固化预警案例',
            '了解预警指标设计的基本原则',
            '准备预警指标设计模板',
        ],
        duration='90分钟',
        fill_guide=[
            '1. 预警指标要灵敏，能够早期发现问题信号',
            '2. 阈值设定要基于历史数据和行业标杆，避免主观随意',
            '3. 干预策略要分级，不同预警级别对应不同的干预强度',
            '4. 追踪机制要明确，确保干预措施落实到位',
            '5. 定期回顾预警指标的有效性，根据实际情况调整',
        ]
    )

    # 案例材料
    add_styled_heading(doc, '案例材料：某企业流程固化预警', 1)

    case_text = """
【公司背景】
华新集团是一家多元化企业集团，业务涵盖制造、地产、金融三大板块。集团于2022年启动了财务共享服务中心建设项目，目标是通过流程标准化和系统集成，提升财务运营效率。

【项目概况】
财务共享服务中心于2023年1月正式上线，主要内容包括：
1. 统一费用报销流程
2. 集中应收应付管理
3. 自动化财务报表生成
4. 资金集中管理

【预警机制建设】
项目组在设计阶段就考虑了预警机制，但上线后发现预警指标的设计存在以下问题：

1. 预警指标选择不当
   - 选择的"系统使用率"指标过于滞后，无法早期预警
   - 缺乏对流程执行质量的监控指标
   - 财务指标的异常往往在问题发生很久后才显现

2. 预警阈值设定不合理
   - 部分阈值设定过于宽松，问题已经很严重才报警
   - 部分阈值设定过于严格，频繁报警导致"狼来了"效应
   - 缺乏对业务周期性的考虑（如月末年末的正常波动）

3. 干预机制不完善
   - 报警后的干预流程不清晰
   - 缺乏分级响应机制，大小问题都上报到同一层级
   - 干预措施执行后缺乏跟踪确认机制

【案例教训】
经过反思，项目组总结出预警机制建设的关键要点：
- 预警指标要选择"先行指标"而非"滞后指标"
- 阈值设定要基于数据分析和业务理解
- 要建立分级干预机制，明确各级响应责任
- 要有追踪确认机制，确保干预措施落地
    """

    p = doc.add_paragraph()
    for line in case_text.strip().split('\n'):
        line = line.strip()
        if line.startswith('【'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.bold = True
            run.font.color.rgb = PRIMARY
            run.font.name = 'Microsoft YaHei'
        elif line.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph()
            p.add_run(line)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        elif line:
            p = doc.add_paragraph()
            p.add_run(line)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'

    doc.add_page_break()

    # 实操任务
    add_styled_heading(doc, '实操任务', 1)

    tasks = [
        ('任务1：设计预警指标（25分钟）', [
            '根据案例，分析原有预警指标的问题',
            '设计新的预警指标体系，选择先行指标',
            '指标应覆盖：流程执行、问题发生、人员变动、资源保障等维度',
            '每个维度选择2-3个核心指标',
        ]),
        ('任务2：设定预警阈值（20分钟）', [
            '为每个预警指标设定绿、黄、红三级阈值',
            '阈值设定要说明依据（历史数据、行业标杆、业务理解）',
            '考虑业务周期性因素，设置动态调整机制',
            '避免阈值过于宽松或过于严格',
        ]),
        ('任务3：制定干预策略（25分钟）', [
            '针对不同预警级别，设计对应的干预策略',
            '明确各级干预的责任人和流程',
            '设计升级机制：低级报警未解决时的升级路径',
            '考虑干预的成本和效果平衡',
        ]),
        ('任务4：设计追踪机制（20分钟）', [
            '建立预警处理追踪表，记录每一起预警的处理过程',
            '明确追踪责任人，确保干预措施落实',
            '设定追踪的时间节点和确认方式',
            '建立预警处理的复盘机制，持续优化预警体系',
        ]),
    ]

    for task_title, task_items in tasks:
        add_styled_heading(doc, task_title, 2)
        for i, item in enumerate(task_items, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 评分标准
    add_styled_heading(doc, '评分标准', 1)

    criteria = [
        ('指标设计科学性 (30分)', [
            '选择先行指标而非滞后指标',
            '指标覆盖全面，维度完整',
            '指标定义清晰，计算方法明确',
        ]),
        ('阈值设定合理性 (25分)', [
            '阈值依据充分',
            '考虑业务周期性因素',
            '避免过松或过严',
        ]),
        ('干预策略有效性 (25分)', [
            '分级干预机制清晰',
            '责任人明确',
            '升级机制合理',
        ]),
        ('追踪机制完整性 (20分)', [
            '追踪流程明确',
            '时间节点清晰',
            '有复盘优化机制',
        ]),
    ]

    for crit_title, crit_items in criteria:
        add_styled_heading(doc, crit_title, 2)
        for item in crit_items:
            p = doc.add_paragraph()
            p.add_run(f'• {item}')
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
        doc.add_paragraph()

    doc.add_page_break()

    # 参考答案
    add_styled_heading(doc, '参考答案（预警指标框架示例）', 1)

    add_styled_heading(doc, '预警指标体系', 2)

    headers = ['维度', '预警指标', '指标性质', '数据来源']
    rows = [
        ['流程执行', '流程执行率', '先行指标', '流程系统'],
        ['流程执行', '流程偏差率', '先行指标', '流程系统'],
        ['问题发生', '问题发生频率', '先行指标', '问题管理系统'],
        ['问题发生', '问题平均解决时长', '先行指标', '工单系统'],
        ['人员变动', '关键岗位人员流失', '先行指标', 'HR系统'],
        ['人员变动', '培训完成率', '先行指标', '培训系统'],
        ['资源保障', '系统使用率', '滞后指标', '系统日志'],
        ['资源保障', '预算执行率', '滞后指标', '财务系统'],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '预警阈值示例', 2)

    headers = ['预警指标', '绿色（正常）', '黄色（关注）', '红色（警告）']
    rows = [
        ['流程执行率', '≥95%', '85%-95%', '<85%'],
        ['流程偏差率', '<5%', '5%-15%', '>15%'],
        ['问题发生频率', '<3次/周', '3-5次/周', '>5次/周'],
        ['问题平均解决时长', '<4小时', '4-8小时', '>8小时'],
        ['关键岗位人员流失', '0人/月', '1人/季度', '≥2人/季度'],
    ]
    add_table_with_headers(doc, headers, rows)

    doc.add_paragraph()

    add_styled_heading(doc, '干预策略示例', 2)

    headers = ['预警级别', '响应时间', '干预措施', '责任人']
    rows = [
        ['黄色预警', '24小时内', '分析原因，制定改进计划', '流程负责人'],
        ['红色预警', '4小时内', '立即召开专题会议，启动应急响应', '部门负责人'],
        ['升级预警', '2小时内', '上报分管领导，调动资源解决', '分管领导'],
    ]
    add_table_with_headers(doc, headers, rows)

    output_path = f'{OUTPUT_DIR}/练习-预警指标设计与干预.docx'
    doc.save(output_path)
    print(f'已创建: {output_path}')

def main():
    """主函数：生成所有练习材料"""
    import os

    # 确保输出目录存在
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("开始生成变革成果固化机制练习材料...")
    print(f"输出目录: {OUTPUT_DIR}")
    print("-" * 50)

    # 生成所有练习材料
    create_exercise_1()  # 变革倒退案例分析
    create_exercise_2()  # 固化度自诊断实操
    create_exercise_3()  # 制度文本编写实战
    create_exercise_4()  # 考核指标重新设计
    create_exercise_5()  # 机制设计工作坊
    create_exercise_6()  # 继任者培养计划制定
    create_exercise_7()  # 文化固化策略设计
    create_exercise_8()  # 故事挖掘与传播
    create_exercise_9()  # 固化效果评估实操
    create_exercise_10() # 预警指标设计与干预

    print("-" * 50)
    print("所有练习材料生成完成！")

if __name__ == '__main__':
    main()

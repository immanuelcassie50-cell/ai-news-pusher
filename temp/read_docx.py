# -*- coding: utf-8 -*-
from docx import Document

# 文件路径 - 使用中文引号 "
path = r'D:\2026年课程\云南磷化\第一阶段作业\第六组作业\第六组作业\第一期作业—第六组—雷利\1—定位表—沃尔沃A40F铰接卡车"保命排故"实战巡检法—第六组—雷利.docx'

doc = Document(path)

print('=' * 60)
print('文档段落')
print('=' * 60)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'段落 {i}: {para.text}')

print()
print('=' * 60)
print('文档表格')
print('=' * 60)
for t_idx, table in enumerate(doc.tables):
    print(f'\n表格 {t_idx}: {len(table.rows)} 行 x {len(table.columns)} 列')
    for r_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        print(f'  行 {r_idx}: {row_data}')

#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# PART 2 Q&A
qa2 = doc.add_paragraph()
r = qa2.add_run('PART 2  Q&A')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0,51,102)

qas2 = [
    ('问：决策卡推行半年后使用率下降，是卡片本身的问题还是推行方式的问题？', '先做简单排查——去问几个不再使用的人，具体是这张卡内容不准了还是忘了这回事，也没人提醒。'),
    ('问：高层觉得决策卡太麻烦，只想要一页纸的总结，怎么办？', '做两个版本——给高层看的确实可以是一页纸摘要，浓缩触发条件和核心判断逻辑；但真正给一线使用者用的完整版，该有的细节不能被砍掉。'),
    ('问：有些决策卡涉及的判断带有强烈的个人风格，别人真的能复制吗？', '决策卡复制的从来不是这个人的风格，是这个人在特定信号出现时的判断结构。风格没法复制也不需要复制；但看到这个信号该往哪个方向想这个结构性的东西，是可以脱离具体的人被传递下去的。'),
    ('问：稽核表会不会让团队变得只敢按卡片行事，反而不敢做真正的判断？', '稽核对话里可以主动加一个问题——这次你有没有基于自己的判断，偏离了卡片上的建议？为什么？如果答案一直是没有，这本身就是一个信号。'),
    ('问：我在做这套方法论时，有没有哪个环节到现在还没想透？', '有。场景细分该到什么颗粒度、组织什么时候算真正内化了这套能力，这两个问题目前没有精确的判断标准。'),
]

for q, a in qas2:
    qp = doc.add_paragraph()
    r = qp.add_run(q); r.bold = True
    ap = doc.add_paragraph()
    ap.paragraph_format.left_indent = Cm(1)
    ap.add_run(a)

doc.add_page_break()
print('S10: Part 2 Q&A done')
doc.save(r'D:\CC\temp\handbook_s10.docx')

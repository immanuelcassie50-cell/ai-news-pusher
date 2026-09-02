# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = 'D:/新课开发/经验萃取/带教手册/完整课程包/05_学员手册/学员手册_组织经验传承_AI赋能岗位带教手册开发.docx'
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc = Document()

def set_cell_shading(cell, fill_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_headers(doc, headers, rows, header_color="4472C4"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(hdr_cells[i], header_color)
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
    return table

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(31, 56, 100)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(46, 84, 150)
    return heading

def add_para(doc, text, bold=False, italic=False, color=None, size=10):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def add_form_table(doc, cells, cols):
    rows = len(cells) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for r in range(rows):
        for c in range(cols):
            idx = r * cols + c
            cell = table.rows[r].cells[c]
            cell.text = cells[idx]
            if r == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                set_cell_shading(cell, "4472C4")
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    return table

# Cover page
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("组织经验传承")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("AI赋能岗位带教手册开发")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(46, 84, 150)

doc.add_paragraph()

label = doc.add_paragraph()
label.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = label.add_run("学员手册")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(68, 114, 196)

doc.add_paragraph()

add_table_with_headers(doc, ["信息项", "填写内容"], [
    ["学员姓名", "________________________"],
    ["所在部门", "________________________"],
    ["课程日期", "________________________"],
    ["课程讲师", "________________________"]
])

doc.add_paragraph()

quote = doc.add_paragraph()
quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote.add_run('“带教不是把事情交代清楚，而是让你的徒弟真正学会。”')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# How to use
add_heading(doc, "如何使用这本手册", 1)
add_para(doc, "这本手册不是讲义，不是笔记，不是课后读物。")
add_para(doc, "它是你在课程中完成的工作台。每一个框架、每一张表单、每一道练习，都要在课堂上完成。")
add_para(doc, "三个使用原则：", bold=True, color=(31, 56, 100))
add_para(doc, "原则一：带着真实任务来", bold=True)
add_para(doc, "手册里所有的练习都要用你自己工作中的真实场景，不要虚构例子。")
add_para(doc, "原则二：写下来比记下来更有价值", bold=True)
add_para(doc, "看懂了不等于会用，写下来才是真正内化的开始。每一道练习，都请认真完成。")
add_para(doc, "原则三：这是起点，不是终点", bold=True)
add_para(doc, "课程结束不是学习的终点。手册最后一章是你的行动计划。")

# Course overview
add_heading(doc, "课程全景图", 1)
add_para(doc, "三天课程，两大核心能力，一个方向：", bold=True, color=(31, 56, 100))
add_table_with_headers(doc, ["时间", "核心任务", "你会带走什么"], [
    ["第一天", "理解带教手册是什么 + 锁定方向 + 挖掘原料", "定位表 + 访谈原料"],
    ["第二天", "AI辅助生成手册骨架 + 开发工具包与案例", "手册框架 + 工具包"],
    ["第三天", "五步优化 + 整合编排 + 成果展示", "完整可用的带教手册"]
])
doc.add_paragraph()
add_para(doc, "这本手册最终会包含以下核心模块：")
add_table_with_headers(doc, ["模块", "解决什么问题", "使用时机"], [
    ["学员画像与差异化策略", "要带的是什么人？不同类型怎么区别对待？", "接到带教任务时"],
    ["带教内容全景图", "要教哪些内容？先教什么后教什么？", "制定带教计划时"],
    ["分阶段带教计划", "按周/月划分，每个阶段教什么、验什么", "每个阶段开始时"],
    ["教学方法与带教技巧", "不同内容用什么方式教最有效", "准备教学时"],
    ["关键节点评估工具", "怎么判断学员真的会了", "阶段结束时"],
    ["带教话术与沟通指南", "怎么说不伤人、还有效", "遇到沟通困难时"],
    ["带教工具包", "评估表、日志、速查表等配套工具", "全程使用"]
])

# Learning preparation
add_heading(doc, "学习准备", 1)
add_heading(doc, "课前自评", 2)
add_para(doc, "填写说明：目的：了解你现在的起点，课程结束后用同一张表重测，看清变化。要求：如实打勾。时间：5分钟", italic=True, size=9, color=(102, 102, 102))
add_form_table(doc, [
    "行为特征", "几乎\n从不", "偶尔\n如此", "经常\n这样", "基本\n如此",
    "1. 我清楚带教手册和岗位操作手册的本质区别", "", "", "", "",
    "2. 我能说清楚带教手册必须回答的六个核心问题", "", "", "", "",
    "3. 我了解不同类型学员的差异化带教策略", "", "", "", "",
    "4. 我掌握结构化访谈技巧，能从优秀带教人身上萃取经验", "", "", "", "",
    "5. 我能设计分阶段的带教计划，并制定可衡量的验收标准", "", "", "", "",
    "6. 我会使用AI工具辅助生成手册内容", "", "", "", ""
], 5)
doc.add_paragraph()
add_para(doc, "我的几乎从不或偶尔如此共有 _____ 行。这些就是你在这门课里的重点方向。")

add_heading(doc, "我的场景卡", 2)
add_para(doc, "填写说明：目的：选定一个贯穿整个课程的真实工作任务，后续所有练习都围绕它展开。", italic=True, size=9, color=(102, 102, 102))
add_form_table(doc, [
    "场景卡", "你的填写",
    "我的岗位/角色", "",
    "我选定的任务名称", "",
    "这个任务通常怎么做？（简述主要步骤）", "",
    "目前这个任务最让我头疼的地方", "",
    "如果AI能帮上忙，我最希望改善什么", "",
    "这个任务最终产出是什么形式", ""
], 2)
doc.add_paragraph()
add_para(doc, "关键提示：填好这张卡，收好。课程中所有练习都会用到它。", bold=True, color=(31, 56, 100))

# Module 1
add_heading(doc, "模块一：理解带教手册的本质", 1)
add_heading(doc, "学习目标", 2)
for obj in ["1. 识别伪带教手册与真正的带教手册的本质差异", "2. 理解带教手册与岗位操作手册的核心区别", "3. 掌握带教手册必须回答的六个核心问题"]:
    add_para(doc, obj)

add_heading(doc, "核心概念：三种伪带教手册", 2)
add_para(doc, "第一种：岗位说明书改装型", bold=True, color=(46, 84, 150))
add_para(doc, "把岗位职责和技能要求列了一遍。内容是对的，但读完之后带教人还是不知道我该怎么教。")
add_para(doc, "第二种：培训课件搬运型", bold=True, color=(46, 84, 150))
add_para(doc, "把新员工培训PPT变成Word。信息量很大，但没有教学设计。")
add_para(doc, "第三种：带教制度型", bold=True, color=(46, 84, 150))
add_para(doc, "规定了带教周期、考核方式、导师签字流程。但带教人翻开后还是不知道明天具体要做什么。")

add_heading(doc, "带教手册的准确定位", 2)
add_para(doc, "带教手册不是培训课件，不是制度文件，也不是新员工入职指南。", bold=True)
add_para(doc, "它是带教人工位上随手能翻的教学工具书。")
add_para(doc, "核心判断标准：一个没带过人的员工，拿到这本手册后，能不能知道第一天该做什么？", italic=True)

add_table_with_headers(doc, ["维度", "岗位操作手册", "带教手册"], [
    ["核心问题", "这件事怎么做", "怎么教会别人做这件事"],
    ["读者", "操作者本人", "带教人（师傅、导师）"],
    ["内容重心", "操作步骤、判断标准", "教什么、按什么顺序教、怎么教、怎么验收"],
    ["成功标准", "照着做能做对", "照着带能把人带出来"],
    ["隐性知识重点", "操作中的判断逻辑", "教学节奏、因材施教、纠错技巧"]
])

add_heading(doc, "六个核心问题", 2)
add_para(doc, "第一个问题：教什么？", bold=True, color=(46, 84, 150))
add_para(doc, "在带教周期内，必须让学员掌握哪些核心技能？优先级怎么排？")
add_para(doc, "第二个问题：按什么顺序教？", bold=True, color=(46, 84, 150))
add_para(doc, "哪些内容必须先掌握才能学后面的？哪些可以并行推进？")
add_para(doc, "第三个问题：怎么教？", bold=True, color=(46, 84, 150))
add_para(doc, "知识类（讲解+举例+提问确认）、技能类（演示→陪练→独立操作）、判断力类（案例分析+场景模拟）。")
add_para(doc, "第四个问题：怎么知道学员学会了？", bold=True, color=(46, 84, 150))
add_para(doc, "差的验收标准：能做报表了。好的验收标准：能独立完成日常运营报表，完成时间不超过45分钟。")
add_para(doc, "第五个问题：遇到问题怎么办？", bold=True, color=(46, 84, 150))
add_para(doc, "学员学不进去怎么办？犯了错怎么纠正才有效果？")
add_para(doc, "第六个问题：带教过程怎么管？", bold=True, color=(46, 84, 150))
add_para(doc, "带教周期怎么规划？关键节点怎么把控？带教记录怎么留？")

add_heading(doc, "练习 1-1：快速辨认练习", 2)
add_form_table(doc, [
    "检查项", "有/没有/不确定",
    "有没有说带教人应该怎么教（而不只是教什么）", "",
    "有没有说什么时候放手，判断标准是什么", "",
    "有没有说学员不同类型，带法怎么调", "",
    "有没有说这个阶段结束，怎么知道学员学会了", ""
], 2)

add_heading(doc, "练习 1-2：六个问题检验练习", 2)
add_form_table(doc, [
    "核心问题", "覆盖程度", "最大缺口",
    "教什么（核心技能清单+优先级）", "完全覆盖/部分覆盖/完全没有", "",
    "按什么顺序教（学习路径）", "完全覆盖/部分覆盖/完全没有", "",
    "怎么教（教法建议）", "完全覆盖/部分覆盖/完全没有", "",
    "怎么知道学会了（验收标准）", "完全覆盖/部分覆盖/完全没有", "",
    "遇到问题怎么办（应对策略）", "完全覆盖/部分覆盖/完全没有", "",
    "过程怎么管（记录与工具）", "完全覆盖/部分覆盖/完全没有", ""
], 3)

# Module 2
add_heading(doc, "模块二：带教手册定位表", 1)
add_heading(doc, "学习目标", 2)
for obj in ["1. 理解带教手册定位的五个维度", "2. 能够区分不同类型学员的带教策略差异", "3. 掌握可衡量的验收标准写法"]:
    add_para(doc, obj)

add_heading(doc, "定位的五个维度", 2)
add_para(doc, "维度一：带教对象画像", bold=True, color=(46, 84, 150))
add_table_with_headers(doc, ["学员类型", "典型特征", "带教重点"], [
    ["应届生", "零基础或接近零基础，学习意愿强但职业习惯待培养", "从岗位认知开始，重点在基础技能和职业习惯"],
    ["社招有相关经验者", "有基础但可能有旧习惯，可能高估自己的适应速度", "先识别已有能力，重点在差异项和公司特有要求"],
    ["转岗员工", "有职场经验但跨领域，学习能力强但容易用旧框架套新业务", "重点在思维转换和关键差异"]
])

add_para(doc, "维度二：带教内容范围", bold=True, color=(46, 84, 150))
add_para(doc, "判断标准：带教结束时，如果学员没掌握这项内容，会不会直接影响他的独立上岗能力？")
add_para(doc, "维度三：带教周期与阶段划分", bold=True, color=(46, 84, 150))
add_table_with_headers(doc, ["阶段", "典型时长", "核心任务"], [
    ["认知期", "第1周左右", "了解岗位全貌、熟悉环境和基本流程"],
    ["跟学期", "第2～4周", "在师傅指导下完成核心操作"],
    ["试手期", "第5～8周", "独立完成常规操作，师傅在旁监督和纠偏"],
    ["独立期", "第9～12周", "独立处理常见场景，师傅仅在异常情况介入"]
])

add_para(doc, "维度四：验收标准", bold=True, color=(46, 84, 150))
add_para(doc, "差的验收标准：基本掌握了核心技能。好的验收标准：能独立完成日常数据核查，口径准确率100%，完成时间不超过30分钟。")
add_para(doc, "维度五：约束条件", bold=True, color=(46, 84, 150))
add_para(doc, "带教人每天能投入多少时间？学员是否有固定学习时间？")

add_heading(doc, "练习 2-1：填写带教手册定位表", 2)
add_para(doc, "第 ___ 组 · 带教手册定位表", bold=True, color=(31, 56, 100))
add_para(doc, "基本信息", bold=True, color=(46, 84, 150))
add_form_table(doc, ["信息项", "填写内容", "手册工作标题", "", "所属岗位/业务线", "", "开发组成员（姓名+角色）", ""], 2)

add_para(doc, "带教对象画像", bold=True, color=(46, 84, 150))
add_form_table(doc, ["信息项", "填写内容", "本手册主要面向", "应届生 / 社招有相关经验者 / 转岗员工", "各类学员的关键特征描述", ""], 2)

add_para(doc, "验收标准", bold=True, color=(46, 84, 150))
add_form_table(doc, ["标准类型", "填写内容", "带教结束时学员应该能独立完成什么", "", "带出来了的判定标准（至少3条）", "1. 2. 3."], 2)

# Module 3
add_heading(doc, "模块三：带教经验萃取", 1)
add_heading(doc, "学习目标", 2)
for obj in ["1. 理解岗位操作知识与带教知识的本质区别", "2. 掌握三类人群的结构化访谈技巧", "3. 能够把做了什么转化为怎么教的"]:
    add_para(doc, obj)

add_heading(doc, "两类知识的区别", 2)
add_table_with_headers(doc, ["岗位操作手册萃取的是", "带教手册萃取的是"], [
    ["我怎么做这件事", "我怎么教会别人做这件事"],
    ["操作步骤、判断标准、注意事项", "教学步骤、节奏把控、因材施教、纠错技巧"]
])

add_para(doc, "关键追问习惯：当访谈对象说我们这个岗位要做XX操作，具体分三步……的时候，要追问的是你教这三步的时候，是一次性教完还是拆开教？", italic=True)

add_heading(doc, "三类人群，三个视角", 2)
add_table_with_headers(doc, ["角色", "主要贡献", "访谈重点"], [
    ["被带教过的新人代表", "学员视角——还原被带教过程中的真实体验、困惑和期待", "当时最迷茫的是什么？什么做法最有效？"],
    ["优秀带教人", "带教经验视角——分享被验证有效的教法、节奏和因材施教策略", "你怎么开始带？怎么判断学员的底子？"],
    ["管理者", "组织标准视角——界定带教的范围、质量标准和验收标准", "你期望带教结束后学员达到什么水平？"]
])

add_heading(doc, "访谈追问技巧", 2)
add_para(doc, "把做了什么变成怎么教的：", bold=True, color=(31, 56, 100))
add_table_with_headers(doc, ["被访者说的", "你要追问的"], [
    ["我会先演示一遍", "演示的时候你会特别强调什么？会不会特意把容易错的地方放慢？"],
    ["然后让他自己做", "看他做的时候你关注什么？什么表现说明他真会了？"],
    ["这个人悟性差", "你具体观察到什么行为？你怎么调整带法的？"],
    ["一般两周后放手", "两周这个时间是怎么来的？有没有提前或延后的情况？"]
])

add_heading(doc, "练习 3-1：三轮结构化访谈", 2)
add_para(doc, "第一轮：访谈被带教过的新人代表", bold=True, color=(46, 84, 150))
add_form_table(doc, ["#", "问题", "追问方向", "1", "你刚来这个岗位的时候，最迷茫的是什么？", "是哪个具体的环节让你不知道从哪里开始？", "2", "师傅教你的时候，哪个做法让你进步最快？", "那个做法具体是怎么做的？", "3", "哪个阶段你觉得最难熬？什么帮你撑过来了？", "那段时间最大的障碍是什么？", "4", "你觉得师傅什么时候放手太早了/太晚了？", "如果那个时候他多帮一把，结果会怎么不同？"], 3)

add_para(doc, "第二轮：访谈优秀带教人", bold=True, color=(46, 84, 150))
add_form_table(doc, ["#", "问题", "追问方向", "1", "你带新人一般怎么开始？第一天做什么？", "为什么这样开始？", "2", "你怎么判断这个新人的基础和接受能力？", "你看什么？问什么？做什么来判断？", "3", "你教一个复杂操作的时候，一般分几步教？", "第一步教完怎么判断可以教第二步了？", "4", "你怎么判断他真的学会了，不是看着会了？", "有没有被假会骗过？那次怎么发现的？"], 3)

add_para(doc, "第三轮：访谈管理者", bold=True, color=(46, 84, 150))
add_form_table(doc, ["#", "问题", "追问方向", "1", "你对带教人的期望是什么？带教结束时你怎么判断带得好不好？", "你看什么指标？", "2", "你见过的最好的带教做法是什么？最差的呢？", "好在哪里/差在哪里？", "3", "带教期内，哪些内容是必须教到位的硬杠杠？", "如果这些没教到，会出现什么后果？"], 3)

# Module 4
add_heading(doc, "模块四：AI辅助生成手册", 1)
add_heading(doc, "学习目标", 2)
for obj in ["1. 掌握带教内容全景图的梳理方法", "2. 能够设计分阶段带教计划", "3. 理解三级大纲的标准结构"]:
    add_para(doc, obj)

add_heading(doc, "带教内容全景图", 2)
add_para(doc, "全景图回答三个问题：", bold=True, color=(31, 56, 100))
for obj in ["1. 教什么：这个岗位在带教周期内，必须让学员掌握的核心技能有哪些？", "2. 怎么分类：哪些是知识类、技能类、判断力类？", "3. 按什么顺序：哪些必须先学会才能学后面的？"]:
    add_para(doc, obj)

add_para(doc, "知识技能分类的三个类型：", bold=True, color=(46, 84, 150))
add_table_with_headers(doc, ["类型", "特征", "适合的教法", "典型内容举例"], [
    ["知识类", "需要理解、记忆", "讲解+举例+提问确认", "业务流程、制度规则"],
    ["技能类", "需要反复练习，形成肌肉记忆", "演示→陪练→独立操作", "操作系统操作、表单填写"],
    ["判断力类", "需要经验积累", "案例分析+场景模拟", "异常情况处理、客户投诉应对"]
])

add_heading(doc, "分阶段带教计划", 2)
add_para(doc, "每个阶段按以下结构填写：")
for obj in ["本阶段教学重点——重点教什么？用什么方式教？", "里程碑验收——阶段结束时确认达标的标准", "差异化标注——应届生/社招/转岗员工各怎么调", "带教人投入时间——每天大概投入多少时间"]:
    add_para(doc, obj)

add_heading(doc, "三级大纲的标准结构", 2)
add_para(doc, "每个带教内容模块，内部都要包含以下三层：")
add_table_with_headers(doc, ["层次", "内容", "回答的问题"], [
    ["教学目标", "本模块教完后，学员应该能做到什么", "教完之后我要达到什么效果"],
    ["教法设计", "推荐用什么方式教", "这个东西我怎么教"],
    ["验收标准", "怎么判断学员在这个模块学会了", "我怎么知道他真的掌握了"]
])

add_heading(doc, "练习 4-1：梳理带教内容全景图", 2)
add_form_table(doc, ["序号", "技能/内容名称", "类型（知识/技能/判断力）", "优先级", "1", "", "", "", "2", "", "", "", "3", "", "", "", "4", "", "", "", "5", "", "", ""], 4)

# Module 5
add_heading(doc, "模块五：五步优化法", 1)
add_heading(doc, "学习目标", 2)
for obj in ["1. 理解内容好不等于手册好的优化思维", "2. 掌握共鸣唤醒、场景还原、原理透传、行动锚点、价值升华五步法"]:
    add_para(doc, obj)

add_heading(doc, "第一步：共鸣唤醒——让人愿意读", 2)
add_para(doc, "共鸣不是写得感人，而是让读者在翻开手册的一瞬间，感觉到这本手册懂我。")
add_para(doc, "核心句式：描述真实困境——你是不是也遇到过这种情况……")

add_heading(doc, "第二步：场景还原——让人建立连接", 2)
add_para(doc, "场景还原是把步骤变成画面的过程。")
add_para(doc, "三个要素：具体动作（我要做什么）、数量/条件（做多少才够）、判断标准（什么时候该停了）。")

add_heading(doc, "第三步：原理透传——让人深度认同", 2)
add_para(doc, "不只告诉带教人怎么教，还要讲清楚为什么这样教有效。原理说明：说人话、讲因果、有对比。")

add_heading(doc, "第四步：行动锚点——让人转化行动", 2)
add_para(doc, "各章节末尾设计3个自查问题：")
for obj in ["我今天/这周教了哪几项内容？学员分别掌握到什么程度？", "在这几项里，学员最容易出问题的是哪个地方？", "我现在的带教节奏对这个学员来说是快了还是慢了？"]:
    add_para(doc, obj)

add_para(doc, "带教口诀设计（8字以内，朗朗上口）：", bold=True, color=(46, 84, 150))
add_table_with_headers(doc, ["核心原则", "好口诀示例"], [
    ["先拆后教", "先拆后教，不贪多"],
    ["看会不算会", "我做一遍不算，自己做对三遍才算"],
    ["纠错要及时", "错完就纠，趁热打铁"],
    ["放手看表现", "放手看表现，不看时间"]
])

add_heading(doc, "第五步：价值升华——让人铭记在心", 2)
add_para(doc, "结尾方式一：画面结尾——描述三个月后的场景。")
add_para(doc, "结尾方式二：过来人寄语——用真实感受打动新任带教人。")
add_para(doc, "结尾方式三：自我对话结尾——邀请带教人思考。")

# Comprehensive exercises
add_heading(doc, "综合练习与行动计划", 1)
add_heading(doc, "综合实战：带教手册开发全链路", 2)
add_table_with_headers(doc, ["阶段", "核心任务", "我的产出"], [
    ["第一步：理解手册定位", "区分真伪手册、回答六个核心问题", ""],
    ["第二步：方向定位", "填写定位表、确定带教对象和周期", ""],
    ["第三步：经验萃取", "完成三轮访谈、填写原料汇总表", ""],
    ["第四步：内容生成", "全景图+分阶段计划+三级大纲", ""],
    ["第五步：工具开发", "六种工具+典型案例", ""],
    ["第六步：五步优化", "共鸣+场景+原理+行动+价值", ""]
])

add_heading(doc, "我的30天行动计划", 2)
add_table_with_headers(doc, ["阶段", "目标", "我要做的一件事", "怎么知道自己做到了"], [
    ["第1-10天", "建立意识", "每次带教前过一遍六个核心问题", ""],
    ["第11-20天", "建立技能", "用全景图方法梳理本岗位带教内容", ""],
    ["第21-30天", "建立系统", "完成第一版带教手册初稿", ""]
])

doc.add_paragraph()
add_para(doc, "30天后可以检验的一个具体指标：", bold=True, color=(46, 84, 150))
add_para(doc, "（不是学到了什么，而是具体的变化）", italic=True, size=9, color=(102, 102, 102))

# Appendix 1
add_heading(doc, "附录一：工具速查", 1)
add_table_with_headers(doc, ["工具名称", "工具用途", "使用时机", "填/查需要的时间"], [
    ["学员起点评估工具", "快速了解学员的底子，制定个性化带教计划", "接到带教任务后，第一次见面时", "10～15分钟"],
    ["分阶段带教检查表", "每个阶段结束时，检查该教的教了没", "每个阶段末尾", "3～5分钟"],
    ["带教话术卡", "高频带教场景的标准话术，拿来就能用", "布置任务、纠错反馈时", "1～2分钟/条"],
    ["带教日志模板", "每日/每周记录带教进展，留痕可查", "每天带教结束后", "3～5分钟/天"],
    ["阶段评估表", "每个阶段验收时的结构化评估工具", "阶段切换前使用", "10～15分钟"],
    ["常见问题速查表", "按学员出现XX情况→可能原因→应对策略组织", "遇到问题时", "1～3分钟/条"]
])

# Appendix 2
add_heading(doc, "附录二：术语表", 1)
add_table_with_headers(doc, ["术语", "定义"], [
    ["带教手册", "带教人工位上随手能翻的教学工具书，回答怎么教会别人做这件事"],
    ["六个核心问题", "教什么/按什么顺序教/怎么教/怎么知道学会了/遇到问题怎么办/过程怎么管"],
    ["三类人群", "新人代表（学员视角）、优秀带教人（经验视角）、管理者（组织标准视角）"],
    ["全景图", "带教内容清单+分类+优先级+依赖关系"],
    ["分阶段带教计划", "按阶段划分，每个阶段有目标、方法和里程碑验收"],
    ["五步优化法", "共鸣唤醒→场景还原→原理透传→行动锚点→价值升华"],
    ["人类溢价", "AI无法替代的人类独特价值：判断力、情境感、创造性、责任感"]
])

# Ending
doc.add_page_break()
end_title = doc.add_paragraph()
end_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end_title.add_run("致出发的你")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 56, 100)

doc.add_paragraph()
add_para(doc, "你用三天时间，走完了这本手册的开发之旅。")
add_para(doc, "这不是一套理论，而是一套可以从明天起就开始用的工作方式。")
doc.add_paragraph()
add_para(doc, "带教是一件需要情感投入的事。这本手册的真正价值，不在于它的内容有多完整，而在于你——愿意把带得好从一个人的经验，变成任何人都能照着用的工具。")
doc.add_paragraph()

quote2 = doc.add_paragraph()
quote2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = quote2.add_run("领先一步，枪打出头鸟；落后半步，别人牵牛我拔桩；领先半步，吃尽红利。")
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(46, 84, 150)

doc.add_paragraph()
add_para(doc, "这就是组织经验传承的意义。", bold=True, color=(31, 56, 100))
add_para(doc, "恭喜你完成了本次工作坊！")

doc.save(OUTPUT_PATH)
print("Document created: " + OUTPUT_PATH)

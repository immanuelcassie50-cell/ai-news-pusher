# -*- coding: utf-8 -*-
"""
创建《经营者讲党课》工具集锦模板 - 完整版
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os

OUTPUT_DIR = "D:/新课开发/党业融合/经营者讲党课/完整课程包/008-工具集锦"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='Microsoft YaHei', font_size=11, bold=False):
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading_with_style(doc, text, level=1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if level == 1:
        set_run_font(run, 'Microsoft YaHei', 16, bold=True)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, 'Microsoft YaHei', 14, bold=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(8)
    return para

def set_cell_shading_para(para, color):
    pPr = para._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    pPr.append(shading)

def create_narrative_template():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("党课叙事结构模板")
    set_run_font(run, 'Microsoft YaHei', 22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——五段式叙事结构")
    set_run_font(run, 'Microsoft YaHei', 12)

    doc.add_paragraph()

    inst = doc.add_paragraph()
    run = inst.add_run("【五段式叙事结构】")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)

    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    headers = ["段落", "核心任务", "操作要点"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2F5496')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 11, bold=True)
                r.font.color.rgb = RGBColor(255, 255, 255)

    rows_data = [
        ("一、开场悬念\n（30秒内抓住注意力）", "设置情境", "用一个画面、一个问题或一个困境开场。避免\"各位领导好，今天我来讲...\"这类开场白。"),
        ("二、抉择还原\n（让听众代入）", "制造代入感", "展示当时面临的几条路，让听众思考\"如果是我会怎么选\"。这是最关键的共鸣点。"),
        ("三、转折揭示\n（真实选择）", "展现决策过程", "揭示你当时的真实选择，以及过程中的关键转折。不要只说结果，要说过程。"),
        ("四、感悟自然生长\n（不替听众总结）", "引发思考", "通过故事的结局，让道理自然浮现。不能说\"通过这件事我明白了...\"，而要让听众自己明白。"),
        ("五、回扣主题\n（简短有力）", "升华收尾", "用一句话回应开头，最快速度收尾。参考句式：\"这就是为什么...\"、\"这让我想到...\""),
    ]

    colors = ['D9E2F3', 'E2EFDA', 'FFF2CC', 'FCE4D6', 'DDEBF7']

    for i, (seg, task, points) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = seg
        set_cell_shading(row.cells[0], colors[i])
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)

        row.cells[1].text = task
        set_cell_shading(row.cells[1], 'F2F2F2')
        for para in row.cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

        row.cells[2].text = points
        for para in row.cells[2].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

    doc.add_paragraph()
    add_heading_with_style(doc, "示范案例：一次艰难的抉择", 1)

    case_para = doc.add_paragraph()
    case_text = (
        "2019年年中，我接手了一个连续三个月业绩下滑的团队。\n\n"
        "【开场悬念】\n"
        "第一次开全员会，我刚说完\"大家好\"，下面就开始有人交头接耳。我看到老张（我们团队资历最老的销售）直接收拾东西准备走人。那一刻，我知道情况比我想象的还要糟。\n\n"
        "【抉择还原】\n"
        "摆在面前的路有三条：一是烧新官上任三把火，雷厉风行地换人；二是无为而治，给大家时间慢慢调整；三是找出问题的真正原因，对症下药。\n\n"
        "【转折揭示】\n"
        "我选择了第四条路——我没有急着做任何决定，而是花了整整两周，一对一地和每一个销售人员深谈。不是问他们业绩为什么下滑，而是问他们：\"你当初为什么选择做销售？\"\"你最辉煌的一次经历是什么？\"\n\n"
        "【感悟自然生长】\n"
        "两周后，老张主动来找我，说：\"领导，我想明白了，不是团队不行了，是我带头躺平了。\"那一刻我意识到，管理者最重要的不是做决策，而是先让团队愿意跟你一起面对问题。\n\n"
        "【回扣主题】\n"
        "这就是为什么我常说：好的管理者，不是让问题消失，而是让团队敢于面对问题。"
    )
    run = case_para.add_run(case_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    output_path = os.path.join(OUTPUT_DIR, "党课叙事结构模板.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_stage_checklist():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("登台自查清单")
    set_run_font(run, 'Microsoft YaHei', 22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——党课登台表达一页纸速查卡")
    set_run_font(run, 'Microsoft YaHei', 12)

    doc.add_paragraph()

    inst = doc.add_paragraph()
    run = inst.add_run("使用说明：")
    set_run_font(run, 'Microsoft YaHei', 10, bold=True)
    run = inst.add_run("上台前快速浏览，确保每个要点都已准备到位。")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    sections_data = [
        ("开场设计（前3句话）", "D9E2F3", [
            "开场是否在30秒内抓住听众注意力？",
            "是否避免\"各位领导好，今天我来讲...\"这类死板开场？",
            "是否用了一个画面、问题或困境开场？"
        ]),
        ("停顿使用", "E2EFDA", [
            "关键转折处是否有意识停顿？",
            "停顿是否让听众有时间消化重要信息？",
            "是否避免从头到尾一个语速？"
        ]),
        ("眼神交流", "FFF2CC", [
            "是否有意识地扫视不同区域的听众？",
            "是否避免了死盯PPT或稿子？",
            "与听众是否有短暂的目光接触？"
        ]),
        ("手势状态", "FCE4D6", [
            "手势是否自然放松？",
            "是否避免了汇报PPT时的指点手势？",
            "手势是否与讲述内容节奏匹配？"
        ])
    ]

    for sec_title, color, items in sections_data:
        para = doc.add_paragraph()
        run = para.add_run(sec_title)
        set_run_font(run, 'Microsoft YaHei', 14, bold=True)
        set_cell_shading_para(para, color)

        table = doc.add_table(rows=len(items) + 1, cols=2)
        table.style = 'Table Grid'

        # Header row
        table.rows[0].cells[0].text = ""
        table.rows[0].cells[1].text = "检查项"
        set_cell_shading(table.rows[0].cells[0], 'F2F2F2')
        set_cell_shading(table.rows[0].cells[1], 'F2F2F2')
        for para in table.rows[0].cells[1].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)

        for i, item in enumerate(items):
            cell0 = table.rows[i + 1].cells[0]
            cell0.text = "□"
            for para in cell0.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell1 = table.rows[i + 1].cells[1]
            cell1.text = item
            for para in cell1.paragraphs:
                for r in para.runs:
                    set_run_font(r, 'Microsoft YaHei', 10)

        doc.add_paragraph()

    add_heading_with_style(doc, "状态自检问题", 2)

    questions = [
        "我今天讲的主题，对我自己有什么触动？",
        "我最想让听众记住的一句话是什么？",
        "如果只能讲3分钟，我会讲什么？"
    ]

    for q in questions:
        para = doc.add_paragraph()
        run = para.add_run("○ " + q)
        set_run_font(run, 'Microsoft YaHei', 11)

    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("自信登台，从容表达")
    set_run_font(run, 'Microsoft YaHei', 12, bold=True)
    run.font.color.rgb = RGBColor(47, 84, 150)

    output_path = os.path.join(OUTPUT_DIR, "登台自查清单.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_peer_feedback_card():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("同伴反馈卡")
    set_run_font(run, 'Microsoft YaHei', 22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("——只反馈感受，不评判对错")
    set_run_font(run, 'Microsoft YaHei', 12)

    doc.add_paragraph()

    principles = doc.add_paragraph()
    run = principles.add_run("【反馈原则】")
    set_run_font(run, 'Microsoft YaHei', 12, bold=True)

    principle_text = (
        "1. 只说感受，不评对错：我们是来帮助彼此进步的，不是来判断谁讲得好不好。\n"
        "2. 聚焦具体瞬间：反馈要具体到\"哪个瞬间让我有感觉\"，而不是\"整体讲得不错\"。\n"
        "3. 温和而真诚：有建设性的反馈不等于尖锐的批评。"
    )

    p = doc.add_paragraph()
    run = p.add_run(principle_text)
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ["反馈维度", "具体内容"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2F5496')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 11, bold=True)
                r.font.color.rgb = RGBColor(255, 255, 255)

    feedback_rows = [
        ("让我有感觉的瞬间", "请描述：在听的过程中，哪个具体的时刻打动了你？\n\n"),
        ("让我走神的瞬间", "请描述：哪个时刻你的注意力飘走了？为什么？\n\n"),
        ("我的一点建议", "请用\"如果......会不会更好？\"的句式来表达建设性意见。\n\n"),
        ("讲师风格亮点", "请描述：你觉得讲师哪些方面做得特别好，值得学习？\n\n")
    ]

    for i, (dim, content) in enumerate(feedback_rows):
        row = table.rows[i + 1]
        row.cells[0].text = dim
        set_cell_shading(row.cells[0], 'D9E2F3')
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)

        row.cells[1].text = content
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    add_heading_with_style(doc, "反馈示例", 2)

    example_text = (
        "• \"你在讲老张主动来找你那一刻，我突然起了一身鸡皮疙瘩。那个细节特别打动人。\"\n\n"
        "• \"如果你在停顿的时候，眼神能再慢一点扫过全场，会不会让那个瞬间更有力量？\"\n\n"
        "• \"我觉得你讲自己'带头躺平了'那个部分特别真实，这种自我揭短很需要勇气。\""
    )

    ex_para = doc.add_paragraph()
    run = ex_para.add_run(example_text)
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    footer = doc.add_table(rows=1, cols=3)
    footer.style = 'Table Grid'
    footer_items = ["反馈人（可选）：", "日期：", "被反馈人："]
    for i, text in enumerate(footer_items):
        cell = footer.rows[0].cells[i]
        cell.text = text
        for para in cell.paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

    output_path = os.path.join(OUTPUT_DIR, "同伴反馈卡.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_tools_collection():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("《讲党课：业务干部的登台表达赋能工作坊》")
    set_run_font(run, 'Microsoft YaHei', 20, bold=True)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title2.add_run("工具集锦合集")
    set_run_font(run, 'Microsoft YaHei', 26, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("课堂用工具完整版")
    set_run_font(run, 'Microsoft YaHei', 14)

    doc.add_paragraph()
    doc.add_paragraph()

    toc_title = doc.add_paragraph()
    run = toc_title.add_run("工具清单")
    set_run_font(run, 'Microsoft YaHei', 16, bold=True)

    tools = [
        "1. 党课素材转化卡 —— 从原始故事到讲稿框架的转化记录表",
        "2. 党课叙事结构模板 —— 五段式叙事结构及示范案例",
        "3. 登台自查清单 —— 一页纸速查卡",
        "4. 同伴反馈卡 —— 规范反馈表"
    ]

    for tool in tools:
        para = doc.add_paragraph()
        run = para.add_run(tool)
        set_run_font(run, 'Microsoft YaHei', 12)
        para.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    doc.add_paragraph()

    intro_title = doc.add_paragraph()
    run = intro_title.add_run("课程简介")
    set_run_font(run, 'Microsoft YaHei', 14, bold=True)

    intro_text = (
        "本课程聚焦于将业务干部真实的管理故事转化为有结构、有共鸣点的党课内容，"
        "再通过反复的登台练习，把这份内容内化成敢讲、能讲、讲得自然的能力。\n\n"
        "核心方法论：案例转化四步法\n"
        "• 第一步：故事盘点——从真实经历中挖掘有戏剧张力的素材\n"
        "• 第二步：主题锚定——找到故事与党课主题的真实连接\n"
        "• 第三步：结构搭建——用讲故事的方式重新组织内容\n"
        "• 第四步：语言转译——让文字有画面感，适合讲台表达\n\n"
        "核心课堂活动：分层递进登台练习\n"
        "• 第一层：小组内讲述练习（低压力环境）\n"
        "• 第二层：录制与回看（看到自己的讲述状态）\n"
        "• 第三层：全班展示与点评（实战检验）"
    )

    intro_para = doc.add_paragraph()
    run = intro_para.add_run(intro_text)
    set_run_font(run, 'Microsoft YaHei', 10.5)

    doc.add_page_break()

    # ===== 工具1：党课素材转化卡 =====
    h1 = doc.add_paragraph()
    run = h1.add_run("工具一：党课素材转化卡")
    set_run_font(run, 'Microsoft YaHei', 18, bold=True)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc1 = doc.add_paragraph()
    run = desc1.add_run("从原始故事到讲稿框架的转化记录表。本表用于记录将原始管理故事转化为党课讲稿的全过程。")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    table1 = doc.add_table(rows=6, cols=4)
    table1.style = 'Table Grid'

    headers1 = ["第一步\n原始故事", "第二步\n故事盘点", "第三步\n主题锚定", "第四步\n结构搭建"]
    for i, h in enumerate(headers1):
        cell = table1.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2F5496')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)
                r.font.color.rgb = RGBColor(255, 255, 255)

    steps_content = [
        ["何时", "冲突点", "本次党课主题", "开场悬念"],
        ["何地", "抉择时刻", "故事给我的启发", "抉择还原"],
        ["何人", "最纠结的瞬间", "与主题的连接", "转折揭示"],
        ["何事", "突破点", "", "感悟自然生长"],
        ["", "", "", "回扣主题"]
    ]

    for row_idx, row_data in enumerate(steps_content):
        for col_idx, text in enumerate(row_data):
            cell = table1.rows[row_idx + 1].cells[col_idx]
            cell.text = text
            for para in cell.paragraphs:
                for r in para.runs:
                    set_run_font(r, 'Microsoft YaHei', 9)

    doc.add_page_break()

    # ===== 工具2：党课叙事结构模板 =====
    h2 = doc.add_paragraph()
    run = h2.add_run("工具二：党课叙事结构模板")
    set_run_font(run, 'Microsoft YaHei', 18, bold=True)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc2 = doc.add_paragraph()
    run = desc2.add_run("五段式叙事结构：开场悬念 - 抉择还原 - 转折揭示 - 感悟自然生长 - 回扣主题")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'

    rows2 = [
        ("开场悬念（30秒内）", "用一个画面、一个问题或一个困境开场，让听众立刻进入情境"),
        ("抉择还原", "展示当时面临的几条路，引发\"如果是我会怎么选\"的思考"),
        ("转折揭示", "揭示你当时的真实选择，以及过程中的关键转折"),
        ("感悟自然生长", "通过故事的结局，让道理自然浮现，不直接说教"),
        ("回扣主题", "用一句话回应开头，最快速度收尾，不要拖泥带水")
    ]

    colors2 = ['D9E2F3', 'E2EFDA', 'FFF2CC', 'FCE4D6', 'DDEBF7']

    for i, (seg, desc) in enumerate(rows2):
        row = table2.rows[i]
        row.cells[0].text = seg
        set_cell_shading(row.cells[0], colors2[i])
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)

        row.cells[1].text = desc
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

    doc.add_page_break()

    # ===== 工具3：登台自查清单 =====
    h3 = doc.add_paragraph()
    run = h3.add_run("工具三：登台自查清单")
    set_run_font(run, 'Microsoft YaHei', 18, bold=True)
    h3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc3 = doc.add_paragraph()
    run = desc3.add_run("一页纸速查卡，上台前快速浏览，确保每个要点都已准备到位。")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    check_data = [
        ("开场设计（前3句话）", ["开场是否在30秒内抓住听众注意力？", "是否避免死板开场？", "是否用画面/问题/困境开场？"]),
        ("停顿使用", ["关键转折处是否有意识停顿？", "停顿是否让听众有时间消化？", "是否避免从头到尾一个语速？"]),
        ("眼神交流", ["是否有意识地扫视不同区域？", "是否避免死盯PPT或稿子？", "与听众是否有短暂的目光接触？"]),
        ("手势状态", ["手势是否自然放松？", "是否避免汇报PPT时的指点手势？", "手势是否与内容节奏匹配？"])
    ]

    for sec_title, items in check_data:
        para = doc.add_paragraph()
        run = para.add_run(sec_title)
        set_run_font(run, 'Microsoft YaHei', 12, bold=True)

        for item in items:
            p = doc.add_paragraph()
            run = p.add_run("[ ] " + item)
            set_run_font(run, 'Microsoft YaHei', 10)
            p.paragraph_format.left_indent = Inches(0.5)

        doc.add_paragraph()

    doc.add_page_break()

    # ===== 工具4：同伴反馈卡 =====
    h4 = doc.add_paragraph()
    run = h4.add_run("工具四：同伴反馈卡")
    set_run_font(run, 'Microsoft YaHei', 18, bold=True)
    h4.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc4 = doc.add_paragraph()
    run = desc4.add_run("规范反馈只聚焦\"感受层面\"而非内容对错，避免互相打击信心。")
    set_run_font(run, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    feedback_table = doc.add_table(rows=5, cols=2)
    feedback_table.style = 'Table Grid'

    fb_headers = ["反馈维度", "具体内容"]
    for i, h in enumerate(fb_headers):
        cell = feedback_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '2F5496')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 11, bold=True)
                r.font.color.rgb = RGBColor(255, 255, 255)

    feedback_rows = [
        ("让我有感觉的瞬间", ""),
        ("让我走神的瞬间", ""),
        ("我的一点建议", ""),
        ("讲师风格亮点", "")
    ]

    for i, (dim, content) in enumerate(feedback_rows):
        row = feedback_table.rows[i + 1]
        row.cells[0].text = dim
        set_cell_shading(row.cells[0], 'D9E2F3')
        for para in row.cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10, bold=True)

        row.cells[1].text = content
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                set_run_font(r, 'Microsoft YaHei', 10)

    doc.add_paragraph()

    principles_para = doc.add_paragraph()
    run = principles_para.add_run("反馈原则：")
    set_run_font(run, 'Microsoft YaHei', 11, bold=True)

    principles_text = "1. 只说感受，不评对错\n2. 聚焦具体瞬间\n3. 温和而真诚"

    p = doc.add_paragraph()
    run = p.add_run(principles_text)
    set_run_font(run, 'Microsoft YaHei', 10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    output_path = os.path.join(OUTPUT_DIR, "工具集锦合集.docx")
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    print("Creating templates...")
    create_narrative_template()
    create_stage_checklist()
    create_peer_feedback_card()
    create_tools_collection()
    print("All Word templates created!")

# -*- coding: utf-8 -*-
"""
生成 demo11-隐藏标准发现清单示例.docx
demo12-价值观冲突定位表示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_demo11():
    """demo11: 隐藏标准发现清单示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("隐藏标准发现清单示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 识别决策中未被明说的隐性标准")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：王发现", "团队：决策优化组", "日期：2026-08-12", "教练：刘洞察"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】隐藏标准 = 决策中真实起作用但未被明确说出的判断标准")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    from docx.oxml import OxmlElement

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、隐藏标准发现三步法")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    steps = [
        ("第一步：观察决策结果", "当决策出现争议时，追问：为什么最终是这个方案？\n什么因素在起作用？"),
        ("第二步：追问'为什么'", "多问几次'为什么'，直到找到真正起作用的标准\n通常3-5层就能挖到隐藏标准"),
        ("第三步：验证假设", "把隐藏标准说出来，看当事人是否认同\n如果对方说'对，就是这个意思'，就找对了")
    ]

    step_table = doc.add_table(rows=len(steps)+1, cols=2)
    step_table.style = 'Table Grid'

    step_table.rows[0].cells[0].text = "步骤"
    step_table.rows[0].cells[1].text = "操作要点"
    for cell in step_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for i, (step, points) in enumerate(steps):
        step_table.rows[i+1].cells[0].text = step
        step_table.rows[i+1].cells[1].text = points
        step_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in step_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、案例应用：供应商评选")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【场景】三家公司参与供应商评选，各项评分如下：")
    run.font.size = Pt(11)

    case_table = doc.add_table(rows=4, cols=5)
    case_table.style = 'Table Grid'

    headers = ["供应商", "价格", "质量", "交期", "服务"]
    for j, h in enumerate(headers):
        case_table.rows[0].cells[j].text = h
        para = case_table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    case_data = [
        ["A公司", "85分", "90分", "78分", "80分"],
        ["B公司", "90分", "82分", "85分", "75分"],
        ["C公司", "80分", "88分", "92分", "85分"]
    ]

    for i, row_data in enumerate(case_data):
        for j, text in enumerate(row_data):
            case_table.rows[i+1].cells[j].text = text
            case_table.rows[i+1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    question = doc.add_paragraph()
    run = question.add_run("【决策】最终选择了A公司，为什么？")
    run.font.size = Pt(11)
    run.font.bold = True

    decision_table = doc.add_table(rows=4, cols=2)
    decision_table.style = 'Table Grid'

    decision_data = [
        ["表面理由", "A公司综合评分最高（各项均衡）"],
        ["隐藏标准追问", "为什么不是交期最快的C公司？\n为什么不是价格最低的B公司？\n决策者说：'我们最看重质量稳定性'"],
        ["发现的隐藏标准", "质量稳定性 > 单项最优\n宁可多花钱，也要确保质量不出问题"],
        ["隐藏原因", "公司曾因供应商质量问题导致客户投诉，所以对质量'零容忍'"]
    ]

    for i, (label, content) in enumerate(decision_data):
        decision_table.rows[i].cells[0].text = label
        decision_table.rows[i].cells[1].text = content
        decision_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in decision_table.rows[i].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    coach = doc.add_paragraph()
    run = coach.add_run("[教练提示] 隐藏标准发现后，应该将其显性化，写入正式评估标准。这样未来的决策会更透明、更高效。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    run.font.italic = True

    doc.add_paragraph()

    section_title3 = doc.add_paragraph()
    run = section_title3.add_run("三、常见隐藏标准清单")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    common_table = doc.add_table(rows=6, cols=2)
    common_table.style = 'Table Grid'

    common_table.rows[0].cells[0].text = "场景"
    common_table.rows[0].cells[1].text = "常见隐藏标准"
    for cell in common_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    common_items = [
        ["供应商评选", "关系户优先 / 央企加分 / 老板推荐"],
        ["人员晋升", "年限优先 / 听话照做 / 避免冲突"],
        ["项目立项", "政治正确 / 抢占资源 / 领导意志"],
        ["方案评审", "新技术不用 / 竞争者已用 / 权威推荐"],
        ["预算分配", "不患寡而患不均 / 切块保护 / 历史惯例"]
    ]

    for i, (scene, standard) in enumerate(common_items):
        common_table.rows[i+1].cells[0].text = scene
        common_table.rows[i+1].cells[1].text = standard
        for cell in common_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo11-隐藏标准发现清单示例.docx"
    doc.save(output_path)
    print("demo11已生成:", output_path)
    return output_path


def create_demo12():
    """demo12: 价值观冲突定位表示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("价值观冲突定位表示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 找到冲突背后的价值立场")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：赵冲突", "团队：组织发展组", "日期：2026-08-15", "教练：陈协调"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】价值观冲突 = 双方在表面问题上意见不同，但本质是价值观层面的分歧")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    from docx.oxml import OxmlElement

    section_title = doc.add_paragraph()
    run = section_title.add_run("一、案例分析：团队扩张策略之争")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【背景】营销总监张强主张大量招聘新人抢占市场；财务总监李谨主张控制成本稳健发展")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 冲突分析表
    conflict_table = doc.add_table(rows=5, cols=2)
    conflict_table.style = 'Table Grid'

    headers = ["分析层次", "内容"]
    for j, h in enumerate(headers):
        conflict_table.rows[0].cells[j].text = h
        para = conflict_table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = conflict_table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    conflict_data = [
        ["表面冲突", "张强：要大规模招聘，目标是实现年营收翻番\n李谨：要严格控人，人均效能必须维持当前水平"],
        ["立场分析", "张强的立场：增长优先，市场份额是核心竞争力\n李谨的立场：稳健优先，现金流是企业生命线"],
        ["价值观层", "张强：进攻型价值观，相信'不进则退'，要抓住窗口期\n李谨：防守型价值观，相信'活下来比什么都重要'"],
        ["深层原因", "张强曾在创业公司因激进扩张失败而被裁员\n李谨曾在国企经历盲目投资后的重组"]
    ]

    for i, (label, content) in enumerate(conflict_data):
        conflict_table.rows[i+1].cells[0].text = label
        conflict_table.rows[i+1].cells[1].text = content
        conflict_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in conflict_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、冲突定位三步法")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    steps = [
        ("第一步：识别冲突", "确认这不是利益冲突，而是价值观冲突\n（如果是利益冲突，可以通过谈判解决）"),
        ("第二步：定位价值观", "理解对方的价值观从何而来\n（往往是过去的经历塑造的，不易改变）"),
        ("第三步：寻找共同上层目标", "找到双方都认同的更高层目标\n（通常是企业生存、长期发展等）")
    ]

    step_table = doc.add_table(rows=len(steps)+1, cols=2)
    step_table.style = 'Table Grid'

    step_table.rows[0].cells[0].text = "步骤"
    step_table.rows[0].cells[1].text = "操作要点"
    for cell in step_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    for i, (step, points) in enumerate(steps):
        step_table.rows[i+1].cells[0].text = step
        step_table.rows[i+1].cells[1].text = points
        step_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in step_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    section_title3 = doc.add_paragraph()
    run = section_title3.add_run("三、解决方案")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    solution_table = doc.add_table(rows=3, cols=2)
    solution_table.style = 'Table Grid'

    solution_data = [
        ["共同上层目标", "企业可持续增长 —— 双方都认同这是最终目标"],
        ["妥协方案", "1. 招聘与业绩挂钩（每增长X%才能招聘Y人）\n2. 设立招聘观察期（3个月内不达标则优化）\n3. 张强负责增长指标，李谨负责成本红线"],
        ["机制保障", "建立月度经营分析会，用数据说话\n避免价值观之争变成意气之争"]
    ]

    for i, (label, content) in enumerate(solution_data):
        solution_table.rows[i].cells[0].text = label
        solution_table.rows[i].cells[1].text = content
        solution_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in solution_table.rows[i].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    insight = doc.add_paragraph()
    run = insight.add_run("【破题关键】价值观冲突没有对错，只有差异。解决方案不是让一方说服另一方，而是设计一个机制，让不同价值观都能发挥作用。")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    run.font.italic = True

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo12-价值观冲突定位表示例.docx"
    doc.save(output_path)
    print("demo12已生成:", output_path)
    return output_path


if __name__ == "__main__":
    create_demo11()
    create_demo12()

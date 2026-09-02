# -*- coding: utf-8 -*-
"""
课程35《心理安全感与信任文化》讲师手册生成脚本
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 配色方案
PRIMARY_COLOR = RGBColor(0x25, 0x63, 0xEB)  # 信任蓝 #2563EB
SECONDARY_COLOR = RGBColor(0x10, 0xB9, 0x81)  # 成长绿 #10B981
DARK_TEXT = RGBColor(0x1F, 0x38, 0x64)  # 深蓝色
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)  # 深灰色

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_colorful_heading(doc, text, level=1, color=None):
    if color is None:
        color = PRIMARY_COLOR
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
    return heading

def add_time_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, '2563EB')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(10)
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, 'F0F9FF')
    return table

def add_script_box(doc, label, content):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(label)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    run.font.size = Pt(10)
    run2 = para.add_run(content)
    run2.font.size = Pt(10)

def add_qa_box(doc, question, answer):
    para1 = doc.add_paragraph()
    para1.paragraph_format.left_indent = Cm(0.5)
    para1.paragraph_format.space_before = Pt(6)
    run1 = para1.add_run(question)
    run1.font.bold = True
    run1.font.color.rgb = SECONDARY_COLOR
    run1.font.size = Pt(10)
    para2 = doc.add_paragraph()
    para2.paragraph_format.left_indent = Cm(0.5)
    para2.paragraph_format.space_after = Pt(6)
    run2 = para2.add_run(answer)
    run2.font.size = Pt(10)

def add_para(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)
    return para

def create_instructor_manual():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("课程35")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("心理安全感与信任文化")
    run2.font.size = Pt(32)
    run2.font.bold = True
    run2.font.color.rgb = DARK_TEXT

    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = title3.add_run("——谷歌氧气计划的启示")
    run3.font.size = Pt(20)
    run3.font.color.rgb = BODY_TEXT

    doc.add_paragraph()
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle.add_run("讲师手册")
    run_sub.font.size = Pt(28)
    run_sub.font.bold = True
    run_sub.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = info.add_run("课程时长：建议6小时\n适用对象：中高层管理者、团队负责人、人力资源工作者")
    run_info.font.size = Pt(12)
    run_info.font.color.rgb = BODY_TEXT

    doc.add_page_break()

    # 目录
    toc_title = doc.add_heading('目 录', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_items = [
        "第一章 讲师手册使用指南",
        "第二章 课程概览",
        "第三章 模块一：认知篇——理解心理安全感（90分钟）",
        "第四章 模块二：诊断篇——评估你的团队（90分钟）",
        "第五章 模块三：建设篇——管理者能做什么（90分钟）",
        "第六章 模块四：实践篇——从知道到做到（60分钟）",
        "第七章 课程收尾",
        "第八章 附录"
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ========== 第一章 ==========
    add_colorful_heading(doc, "第一章 讲师手册使用指南", level=1)
    add_colorful_heading(doc, "1.1 手册结构说明", level=2)
    add_para(doc, "本手册共分为八个章节，涵盖从课程设计理念到课后跟进的完整内容。每个模块都包含：")
    add_para(doc, "• 教学目标与关键知识点")
    add_para(doc, "• 案例库与讲解要点")
    add_para(doc, "• 互动设计（含详细操作指令）")
    add_para(doc, "• 时间分配表（精确到分钟）")
    add_para(doc, "• 讲师话术模板")
    add_para(doc, "• 常见问题应对")

    add_colorful_heading(doc, "1.2 课程设计理念", level=2)
    add_para(doc, "本课程基于以下核心设计理念：")
    add_para(doc, "（1）从谷歌氧气计划切入：谷歌用长达10年的数据研究证明，心理安全感是团队绩效的最强预测因子。这一\"大厂实证\"能快速建立学员的信任与兴趣。")
    add_para(doc, "（2）从认知到行动的学习路径：先理解\"是什么\"（认知），再评估\"在哪里\"（诊断），最后解决\"怎么做\"（建设）。符合成人学习的循序渐进规律。")
    add_para(doc, "（3）体验式学习：大量使用案例分析、角色扮演、自我评估工具，让学员在\"做中学\"。")
    add_para(doc, "（4）工具化输出：每个模块都配备实用的工具和模板，确保学习成果可落地。")

    add_colorful_heading(doc, "1.3 讲师角色定位", level=2)
    add_para(doc, "讲师在本课程中扮演以下三重角色：")
    add_para(doc, "• 促进者（Facilitator）：引导讨论、激发思考、创造安全的讨论环境")
    add_para(doc, "• 专家（Expert）：提供专业知识、解答疑惑、分享洞见")
    add_para(doc, "• 教练（Coach）：推动行动计划制定、促进行为改变")

    add_colorful_heading(doc, "1.4 课前准备清单", level=2)
    prep_table = doc.add_table(rows=8, cols=2)
    prep_table.style = 'Table Grid'
    prep_data = [
        ("□", "PPT课件（已确认投影设备正常）"),
        ("□", "学员手册每人一本"),
        ("□", "白板/大白纸、马克笔"),
        ("□", "心理安全感评估问卷（打印版，每位学员1份）"),
        ("□", "案例讨论分组名单"),
        ("□", "计时器/闹钟"),
        ("□", "名片盒（用于随机分组）"),
        ("□", "茶水/点心（建议安排）")
    ]
    for i, (checkbox, item) in enumerate(prep_data):
        row = prep_table.rows[i]
        row.cells[0].text = checkbox
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[1].text = item

    doc.add_page_break()

    # ========== 第二章 ==========
    add_colorful_heading(doc, "第二章 课程概览", level=1)

    add_colorful_heading(doc, "2.1 课程目标", level=2)
    add_para(doc, "知识目标：")
    add_para(doc, "• 理解心理安全感的定义与核心内涵")
    add_para(doc, "• 掌握谷歌氧气计划的研究发现与启示")
    add_para(doc, "• 识别高心理安全感团队的标志性特征")
    add_para(doc, "• 了解管理者行为对团队心理安全感的影响机制")
    add_para(doc, "技能目标：")
    add_para(doc, "• 能够使用心理安全感评估工具诊断团队现状")
    add_para(doc, "• 能够识别团队中破坏心理安全感的行为模式")
    add_para(doc, "• 能够制定并实施提升团队心理安全感的行动计划")
    add_para(doc, "态度目标：")
    add_para(doc, "• 认识到心理安全感对团队绩效的重要性")
    add_para(doc, "• 建立\"管理者是团队心理安全感第一责任人\"的意识")
    add_para(doc, "• 愿意反思并改变自己的管理行为")

    add_colorful_heading(doc, "2.2 学员画像分析", level=2)
    add_para(doc, "本课程面向以下学员群体：")
    add_para(doc, "• 中高层管理者：需要理解如何打造高绩效团队文化")
    add_para(doc, "• 团队负责人/项目经理：直接带团队，需要实操工具")
    add_para(doc, "• 人力资源工作者：推动组织文化建设，需要评估方法")
    add_para(doc, "• 潜在预设：学员已有一定管理经验，对\"团队建设\"话题不陌生")

    add_colorful_heading(doc, "2.3 课程时长分配", level=2)
    total_table = doc.add_table(rows=6, cols=4)
    total_table.style = 'Table Grid'
    total_headers = ["模块", "主题", "时长", "累计"]
    total_rows = [
        ("模块一", "认知篇——理解心理安全感", "90分钟", "90分钟"),
        ("模块二", "诊断篇——评估你的团队", "90分钟", "180分钟"),
        ("模块三", "建设篇——管理者能做什么", "90分钟", "270分钟"),
        ("模块四", "实践篇——从知道到做到", "60分钟", "330分钟"),
        ("合计", "（约5.5小时，含休息）", "330分钟", "约6小时")
    ]
    add_time_table(doc, total_headers, total_rows)

    add_colorful_heading(doc, "2.4 教学方法说明", level=2)
    add_para(doc, "本课程采用多元教学方法：")
    add_para(doc, "• 讲授法（30%）：核心概念、理论框架、研究发现")
    add_para(doc, "• 案例分析法（25%）：谷歌氧气计划案例、管理情境案例")
    add_para(doc, "• 小组讨论法（20%）：问题探讨、经验分享")
    add_para(doc, "• 工具演练法（15%）：评估工具使用、行动计划制定")
    add_para(doc, "• 角色扮演法（10%）：行为演练、情景模拟")

    doc.add_page_break()

    # ========== 第三章 模块一 ==========
    add_colorful_heading(doc, "第三章 模块一：认知篇——理解心理安全感（90分钟）", level=1)

    add_colorful_heading(doc, "3.1 教学目标", level=2)
    add_para(doc, "• 理解心理安全感的定义与核心内涵")
    add_para(doc, "• 认识谷歌氧气计划的研究价值")
    add_para(doc, "• 区分心理安全感与\"舒适区\"的概念边界")

    add_colorful_heading(doc, "3.2 关键知识点", level=2)
    add_para(doc, "知识点1：心理安全感的定义")
    add_para(doc, "心理安全感（Psychological Safety）是一种团队成员共同持有的信念——在这个团队中，你可以敞开心扉、冒险发声，而不用担心会被嘲笑、惩罚或排斥。")
    add_para(doc, "关键词：敞开心扉、冒险发声、不被惩罚")

    add_para(doc, "知识点2：谷歌氧气计划")
    add_para(doc, "2008年，谷歌启动\"氧气计划\"（Project Oxygen），用大数据分析识别优秀管理者的8项核心行为。历经10年、覆盖数万人次的研究发现：心理安全感是团队绩效的最强预测因子。")

    add_para(doc, "知识点3：高心理安全感团队的标志")
    add_para(doc, "• 成员敢于在会议中提出不同意见")
    add_para(doc, "• 错误被视为学习机会而非惩罚理由")
    add_para(doc, "• 新想法受到鼓励而非打压")
    add_para(doc, "• 团队成员相互支持而非相互推诿")

    add_colorful_heading(doc, "3.3 案例库", level=2)

    add_para(doc, "【案例一】谷歌氧气计划的故事")
    case1 = doc.add_paragraph()
    case1.add_run("案例背景：").bold = True
    case1.add_run("2008年，谷歌CEO施密特担心公司的管理出了问题。他让HR用大数据来找出答案——什么是优秀管理者的共同特质？")
    case1_b = doc.add_paragraph()
    case1_b.add_run("关键冲突点：").bold = True
    case1_b.add_run("研究团队原本预期会找到\"技术能力\"或\"业务经验\"等硬技能，但数据告诉他们：最重要的因素是——管理者是否善于帮助团队成员成长、是否创造了一个让人们敢于冒险的环境。")
    case1_c = doc.add_paragraph()
    case1_c.add_run("讨论问题：").bold = True
    case1_c.add_run("1. 如果你是当时的谷歌HR，你会如何向CEO汇报这个意外发现？\n2. 这个发现对你的管理工作有什么启发？")

    doc.add_paragraph()

    add_para(doc, "【案例二】Amy Edmondson的医院研究")
    case2 = doc.add_paragraph()
    case2.add_run("案例背景：").bold = True
    case2.add_run("哈佛商学院教授Amy Edmondson在一家医院研究医疗错误。她发现：不同科室的医护人员报告的错误数量差异巨大——但这不是因为他们犯的错误不同，而是因为他们\"感觉到安全\"的程度不同。")
    case2_b = doc.add_paragraph()
    case2_b.add_run("关键洞察：").bold = True
    case2_b.add_run("高心理安全感的科室，错误报告率更高——因为他们知道，报告错误是改进的机会，而不是被惩罚的理由。这反而让他们在真正的患者安全上表现更好。")
    case2_c = doc.add_paragraph()
    case2_c.add_run("讨论问题：").bold = True
    case2_c.add_run("1. 为什么\"报告错误最多\"的团队反而是\"最安全\"的团队？\n2. 你的团队目前处于什么状态？")

    add_colorful_heading(doc, "3.4 互动设计", level=2)

    add_para(doc, "【互动1】快速投票：你的团队安全吗？")
    add_script_box(doc, "【指令话术】", "现在我想请大家做一个快速投票。请根据你的真实感受，选择你所在团队的心理安全感程度：\n\n1分 = 完全不同意：我在团队中感到不安全，不敢说出真实想法\n2分 = 有点不同意\n3分 = 中立\n4分 = 有点同意\n5分 = 完全同意：我在团队中可以完全敞开心扉\n\n（等待30秒，让学员在心中选择）")

    add_para(doc, "【互动2】小组讨论：你见过的高心理安全感场景")
    add_script_box(doc, "【指令话术】", "请每个人花2分钟回忆一个场景——在这个场景中，你感到团队是真正安全的，你可以自由地表达想法而不用担心后果。\n\n然后请小组分享，每人2分钟。我会给大家8分钟小组讨论时间。\n\n（计时器设定8分钟）")

    add_colorful_heading(doc, "3.5 时间分配", level=2)

    time_table1 = doc.add_table(rows=6, cols=4)
    time_table1.style = 'Table Grid'
    headers1 = ["环节", "时长", "教学方法", "物料"]
    rows1 = [
        ("导入：破冰故事", "10分钟", "故事开场", "PPT第3页"),
        ("知识点讲解：心理安全感定义", "15分钟", "讲授+提问", "PPT第4-5页"),
        ("案例一：谷歌氧气计划", "20分钟", "案例分析", "PPT第6-8页"),
        ("案例二：医院研究", "15分钟", "案例分析+讨论", "PPT第9-10页"),
        ("互动：快速投票+小组讨论", "20分钟", "体验式学习", "无"),
        ("知识点总结+过渡", "10分钟", "讲授", "PPT第11页")
    ]
    add_time_table(doc, headers1, rows1)

    add_colorful_heading(doc, "3.6 讲师话术", level=2)

    add_script_box(doc, "【导入语】", "我想先给大家讲一个真实的故事。2008年，谷歌的CEO埃里克·施密特找到HR负责人，说了一句话——\"我们的管理一定哪里出了问题。\"\n\n当时谷歌已经是一家万亿美元市值的公司，但施密特敏锐地感觉到：有些东西不对了。\n\n于是，他做了一件谷歌最擅长的事——用大数据来找答案。这个项目叫\"氧气计划\"，后来成为管理学史上最著名的研究之一。\n\n今天的第一节课，我们就从这故事开始。")

    add_script_box(doc, "【过渡语】", "好，我们刚刚了解了什么是心理安全感。现在我想请大家思考一个问题——这些听起来很美好，但在中国企业的实际管理中，真的能做到吗？\n\n接下来，我们来看一个更深入的研究。Amy Edmondson教授在一家医院的研究，可能给我们更多启发。")

    add_script_box(doc, "【收尾语】", "在进入下一节之前，我想留给大家一个问题：\n\n今天课程结束后，请问问自己——我上一次在团队中\"冒险发声\"是什么时候？是因为什么原因让你犹豫或放弃？\n\n这个问题，我们会在最后一个模块再次回顾。")

    add_colorful_heading(doc, "3.7 常见问题应对", level=2)

    add_qa_box(doc, "Q1：心理安全感是不是就是\"一团和气\"？", "不是。心理安全感不是要消灭分歧，而是创造一种环境——让不同的声音能被听到，让人敢于说出\"我认为这样不对\"。适度的冲突往往来自高度的心理安全。")

    add_qa_box(doc, "Q2：强调心理安全感会不会让团队变得\"玻璃心\"？", "恰恰相反。真正高心理安全感的团队，更能接受建设性的批评和反馈，因为他们知道这些反馈是出于善意，而非恶意攻击。这让团队更\"皮实\"而非更\"玻璃心\"。")

    add_qa_box(doc, "Q3：这个概念适合国企/传统行业吗？", "非常适合。实际上，传统行业往往更需要关注心理安全感——因为这些组织的等级文化更强，\"说了会不会被穿小鞋\"的顾虑更深。作为管理者，我们的责任就是打破这种顾虑。")

    doc.add_page_break()

    # ========== 第四章 模块二 ==========
    add_colorful_heading(doc, "第四章 模块二：诊断篇——评估你的团队（90分钟）", level=1)

    add_colorful_heading(doc, "4.1 教学目标", level=2)
    add_para(doc, "• 掌握心理安全感评估工具的使用方法")
    add_para(doc, "• 能够识别团队中破坏心理安全感的行为信号")
    add_para(doc, "• 理解评估结果背后的深层原因")

    add_colorful_heading(doc, "4.2 关键知识点", level=2)
    add_para(doc, "知识点1：心理安全感评估维度")
    add_para(doc, "根据Amy Edmondson的研究，心理安全感可分为4个维度：")
    add_para(doc, "• 表达想法的自由度")
    add_para(doc, "• 犯错误的容错度")
    add_para(doc, "• 寻求帮助的舒适度")
    add_para(doc, "• 提出异议的安全性")

    add_para(doc, "知识点2：破坏心理安全感的行为信号")
    add_para(doc, "以下行为会悄悄破坏团队的心理安全感：")
    add_para(doc, "• 打断别人的发言或否定别人的观点")
    add_para(doc, "• \"这个问题很蠢\"——公开嘲讽的想法")
    add_para(doc, "• 决策时不让相关人员参与")
    add_para(doc, "• 出现问题时第一时间找\"责任人\"而非\"原因\"")

    add_colorful_heading(doc, "4.3 案例库", level=2)

    add_para(doc, "【案例三】某互联网公司的\"沉默会议室\"")
    case3 = doc.add_paragraph()
    case3.add_run("案例背景：").bold = True
    case3.add_run("某中型互联网公司的产品负责人李明很困惑：为什么每次产品评审会议，大家都不说话？明明在会前、会后私下沟通时，团队成员都能提出很多好想法。")
    case3_b = doc.add_paragraph()
    case3_b.add_run("关键冲突点：").bold = True
    case3_b.add_run("李明仔细观察后发现：当他表达某个想法后，如果有人提出不同意见，他会不自觉地说\"这个想法其实我们早就讨论过\"，或者\"这个方案技术上周明之前提过，有问题\"。这种回应方式让团队成员感到——自己的想法\"不够好\"，于是选择沉默。")
    case3_c = doc.add_paragraph()
    case3_c.add_run("讨论问题：").bold = True
    case3_c.add_run("1. 李明的问题出在哪里？\n2. 如果你是李明的教练，你会给他什么建议？")

    add_colorful_heading(doc, "4.4 互动设计", level=2)

    add_para(doc, "【互动3】心理安全感评估问卷")
    add_script_box(doc, "【指令话术】", "现在我要发给大家一份问卷——\"团队心理安全感自评问卷\"。\n\n请大家根据自己所在团队的实际情况，按照1-5分进行打分。\n\n（分发问卷，约3分钟完成）\n\n请注意：这是一份匿名问卷，你的答案不会被任何人看到。我们只需要总体数据，不需要你透露个人信息。\n\n（等待学员完成，约3-5分钟）")

    add_script_box(doc, "【分析引导话术】", "问卷结果汇总后，我会展示一个\"团队心理安全感画像\"图。\n\n请大家重点关注：\n1. 哪个维度的分数最低？\n2. 这个最低分符合你的预期吗？\n3. 如果你是团队负责人，你觉得为什么这个维度分数最低？")

    add_para(doc, "【互动4】行为信号识别练习")
    add_script_box(doc, "【指令话术】", "接下来，我会给大家展示6条管理者行为描述。请判断：每条行为是\"破坏\"还是\"建设\"团队心理安全感？\n\n（展示PPT第15页，等待30秒）\n\n现在请小组讨论：你们判断的依据是什么？有没有争议性的选项？")

    add_colorful_heading(doc, "4.5 时间分配", level=2)

    time_table2 = doc.add_table(rows=6, cols=4)
    time_table2.style = 'Table Grid'
    headers2 = ["环节", "时长", "教学方法", "物料"]
    rows2 = [
        ("回顾+导入", "5分钟", "讲授", "PPT第12页"),
        ("知识点：评估维度", "15分钟", "讲授+互动", "PPT第13-14页"),
        ("互动：评估问卷", "25分钟", "工具演练", "评估问卷"),
        ("案例三：沉默会议室", "15分钟", "案例分析", "PPT第15-16页"),
        ("互动：行为信号识别", "15分钟", "体验式学习", "PPT第17页"),
        ("知识点总结+过渡", "15分钟", "讲授+提问", "PPT第18页")
    ]
    add_time_table(doc, headers2, rows2)

    add_colorful_heading(doc, "4.6 讲师话术", level=2)

    add_script_box(doc, "【导入语】", "上一节我们理解了什么是心理安全感。这一节，我们要做一件更具体的事——评估你所在团队的现状。\n\n有一句话叫\"你无法管理你无法衡量的东西\"。所以在讨论\"如何提升\"之前，我们先要回答一个问题：我的团队，目前的心理安全感到底怎么样？")

    add_script_box(doc, "【问卷引导话术】", "请大家拿到问卷后，先不要急着填写。让我先解释一下题目：\n\n每个题目都是关于团队的一种感受。请按照你过去3个月在团队中的真实感受来打分，而不是你\"希望\"的状态。\n\n记住，这没有对错之分，只有真实与否的区别。")

    add_script_box(doc, "【过渡语】", "评估完了现状，我们来聊一个有趣的现象——为什么很多团队明明想要开放、想要创新，但大家的实际行为却恰恰相反？\n\n下一个案例，可能会让你找到答案。")

    add_colorful_heading(doc, "4.7 常见问题应对", level=2)

    add_qa_box(doc, "Q1：如果学员担心问卷结果被泄露怎么办？", "明确强调：这是匿名问卷，讲师不会收集任何个人信息。问卷结果只用于课堂集体分析，不会上报到任何部门。如果学员仍然担心，可以让他们把问卷收好，课后自行分析。")

    add_qa_box(doc, "Q2：如果团队整体分数很高，学员觉得\"不需要学这个\"怎么办？", "请他们思考：分数高是\"真实的高\"，还是\"我以为的高\"？很多时候，我们对自己的盲区一无所知。另外，高分数也可能意味着还有提升空间——从80分到90分往往比从50分到60分更有价值。")

    add_qa_box(doc, "Q3：如果学员说\"我们团队就是有问题，但老板不承认\"怎么办？", "这是一个很常见的困境。我的建议是：先从自己开始改变。没有人能改变老板的行为，但每个人可以先改变自己的行为。当你开始创造一个小小的\"安全空间\"，涟漪效应会慢慢扩散。")

    doc.add_page_break()

    # ========== 第五章 模块三 ==========
    add_colorful_heading(doc, "第五章 模块三：建设篇——管理者能做什么（90分钟）", level=1)

    add_colorful_heading(doc, "5.1 教学目标", level=2)
    add_para(doc, "• 掌握谷歌氧气计划提出的8项管理者核心行为")
    add_para(doc, "• 能够识别并改进自己的管理行为模式")
    add_para(doc, "• 学会具体的行为建设方法")

    add_colorful_heading(doc, "5.2 关键知识点", level=2)
    add_para(doc, "知识点1：谷歌氧气计划的8项管理者行为")
    add_para(doc, "根据谷歌10年研究，优秀管理者具备以下8项行为（按重要性排序）：")
    add_para(doc, "1. 做一名教练（Be a Coach）")
    add_para(doc, "2. 赋权团队，不要微观管理（Empower the Team）")
    add_para(doc, "3. 关心团队成员的成功与个人福祉（Create a Team Vision）")
    add_para(doc, "4. 工作高效且有战略眼光（Be Productive）")
    add_para(doc, "5. 积极倾听，分享信息（Communicate）")
    add_para(doc, "6. 帮助职业发展（Career Development）")
    add_para(doc, "7. 清晰的愿景/战略（Clear Vision）")
    add_para(doc, "8. 具备关键技能（Technical Skills）")

    add_para(doc, "知识点2：管理者日常行为检查清单")
    add_para(doc, "以下是建设心理安全感的\"DO\"和\"DON'T\"清单：")
    add_para(doc, "DO（要做）：")
    add_para(doc, "• 主动邀请沉默者发言：\"王老师，你有什么看法？\"")
    add_para(doc, "• 对错误表示好奇而非责备：\"发生了什么？我们能学到什么？\"")
    add_para(doc, "• 公开表扬提出不同意见的人")
    add_para(doc, "DON'T（不要做）：")
    add_para(doc, "• 说\"这个想法不成熟\"（暗示只有\"成熟\"的想法才值得表达）")
    add_para(doc, "• 在别人发言时低头看手机（传递\"你的话不重要\"的信号）")
    add_para(doc, "• 决策后说\"这是我的决定，你们执行就好\"（关闭对话）")

    add_colorful_heading(doc, "5.3 案例库", level=2)

    add_para(doc, "【案例四】张总的两场会议")
    case4 = doc.add_paragraph()
    case4.add_run("案例背景：").bold = True
    case4.add_run("某传统企业总监张总，两周内开了两场类似的会议，但效果截然不同。")
    case4_b = doc.add_paragraph()
    case4_b.add_run("会议A（破坏心理安全感）：").bold = True
    case4_b.add_run("张总在会议一开始就说：\"我召集大家是想听听你们的想法，但时间有限，每人最多说2分钟。\"然后在团队成员发言时，他不断看手机，当有人提出与他不同的意见时，他说：\"你的想法我理解，但这个方案我们去年试过，不可行。\"")
    case4_c = doc.add_paragraph()
    case4_c.add_run("会议B（建设心理安全感）：").bold = True
    case4_c.add_run("张总同样召集了会议，但这次他说：\"今天我想请大家挑战我的想法。如果你们发现我的逻辑有漏洞，请直接指出来——这对我帮助很大。\"当团队成员小李提出一个\"不成熟\"的想法时，张总没有否定，而是追问：\"这个想法很有意思，能再展开说说吗？\"")
    case4_d = doc.add_paragraph()
    case4_d.add_run("讨论问题：").bold = True
    case4_d.add_run("1. 两场会议的本质区别在哪里？\n2. 如果你是小李，你分别在两场会议中会怎么想？\n3. 张总在两场会议中的哪些具体行为最值得你学习或避免？")

    add_colorful_heading(doc, "5.4 互动设计", level=2)

    add_para(doc, "【互动5】行为演练——会议主持技巧")
    add_script_box(doc, "【指令话术】", "现在我们要做一个角色扮演。请大家从小组中选出两个人：一人扮演\"管理者\"，一人扮演\"提出异议的团队成员\"。\n\n情境设定：团队正在讨论一个方案，管理者提出了方案A，团队成员认为方案B更好。\n\n（等待2分钟准备）\n\n现在请开始演练。其他人请观察：管理者的哪些行为是在建设心理安全感？哪些是在破坏？\n\n（演练5分钟，然后集体反馈）")

    add_para(doc, "【互动6】个人行为反思——我的\"信任账户\"")
    add_script_box(doc, "【指令话术】", "请每个人拿出一张纸，画一个简单的表格：\n\n左列：过去一个月中，你对团队成员说过的\"破坏心理安全感\"的话或做过的事\n右列：你打算怎么弥补或改变\n\n（给5分钟时间书写）\n\n这个练习不需要交给任何人。但我建议你在课后把这张纸放在你能看到的地方，作为自己的\"行为改进提醒\"。")

    add_colorful_heading(doc, "5.5 时间分配", level=2)

    time_table3 = doc.add_table(rows=7, cols=4)
    time_table3.style = 'Table Grid'
    headers3 = ["环节", "时长", "教学方法", "物料"]
    rows3 = [
        ("导入：上节回顾", "5分钟", "讲授", "PPT第19页"),
        ("知识点：8项管理者行为", "20分钟", "讲授+讨论", "PPT第20-24页"),
        ("案例四：张总的两场会议", "15分钟", "案例分析", "PPT第25-26页"),
        ("互动：行为演练", "20分钟", "角色扮演", "无"),
        ("知识点：DO/DON'T清单", "10分钟", "讲授", "PPT第27页"),
        ("互动：我的\"信任账户\"", "10分钟", "反思工具", "白纸"),
        ("知识点总结+过渡", "10分钟", "讲授", "PPT第28页")
    ]
    add_time_table(doc, headers3, rows3)

    add_colorful_heading(doc, "5.6 讲师话术", level=2)

    add_script_box(doc, "【导入语】", "前面两节，我们学了\"是什么\"和\"在哪里\"。这一节，我们进入最关键的部分——\"怎么做\"。\n\n作为管理者，我们每天的行为都在塑造团队的空气。你是创造安全感的那个人，也是破坏安全感的那个人。\n\n接下来的内容，会给你一些具体的\"工具\"。但工具的价值在于使用，不在于收藏。")

    add_script_box(doc, "【行为清单讲解话术】", "我刚才给出了8项行为，但我知道你们可能记不住。所以我给你们一个更简单的方法——记住两个关键词：\n\n第一个是\"邀请\"。每次开会时，主动邀请那些不太说话的人发言。\n第二个是\"好奇\"。每次听到不同意见时，先问\"为什么\"，而不是立刻说\"但是\"。\n\n做到这两点，你的团队心理安全感至少能提升一个台阶。")

    add_script_box(doc, "【过渡语】", "好，我们已经学完了理论。但我知道，从\"知道\"到\"做到\"之间，往往有一条巨大的鸿沟。\n\n最后一节，我们要做的，就是帮大家跨过这条鸿沟。")

    add_colorful_heading(doc, "5.7 常见问题应对", level=2)

    add_qa_box(doc, "Q1：是不是只要管理者改变了，团队心理安全感就会提升？", "管理者是核心因素，但不是唯一因素。团队文化是系统性的——需要时间沉淀，也需要团队成员的共同参与。但毫无疑问，管理者的行为是最大变量。改变管理者，是投入产出比最高的干预方式。")

    add_qa_box(doc, "Q2：如果团队成员本身就不愿意表达呢？", "这种情况更需要耐心。可以用\"书面+口头\"的混合方式——先让每个人书面写下想法，再开会讨论。另外，持续地示范\"安全的发言方式\"，会慢慢改变团队的文化。")

    add_qa_box(doc, "Q3：谷歌的8项行为是否适合所有类型的管理者？", "框架是通用的，但具体行为需要根据情境调整。比如\"做教练\"这个行为，对技术型管理者和创意型管理者来说，具体表现会不同。重点不是模仿形式，而是理解背后的原则。")

    doc.add_page_break()

    # ========== 第六章 模块四 ==========
    add_colorful_heading(doc, "第六章 模块四：实践篇——从知道到做到（60分钟）", level=1)

    add_colorful_heading(doc, "6.1 教学目标", level=2)
    add_para(doc, "• 制定个人心理安全感提升行动计划")
    add_para(doc, "• 掌握行为改变的具体方法和工具")
    add_para(doc, "• 建立持续改进的自我监督机制")

    add_colorful_heading(doc, "6.2 关键知识点", level=2)
    add_para(doc, "知识点1：行为改变的\"小步快跑\"法")
    add_para(doc, "行为改变不需要\"大爆炸\"，而是需要\"小步快跑\"。建议方法是：")
    add_para(doc, "• 选1个最需要改变的行为（不要贪多）")
    add_para(doc, "• 设计1个\"触发场景\"（比如每次开会的第一个议题）")
    add_para(doc, "• 找1个\"反馈伙伴\"（可以是非直接下属）")
    add_para(doc, "• 坚持21天（形成习惯的最小周期）")

    add_para(doc, "知识点2：3个立即可用的\"开场白\"")
    add_para(doc, "以下3句话可以帮助你在日常管理中创造心理安全感：")
    add_para(doc, "1. \"这个决定最终是我来做，但我特别想听听你们的真实想法。\"——邀请参与")
    add_para(doc, "2. \"如果我们换个角度看这个问题，会不会有不同的答案？\"——鼓励挑战")
    add_para(doc, "3. \"最近有没有什么事情让你感到不舒服，哪怕很小的事情？\"——开放反馈")

    add_colorful_heading(doc, "6.3 互动设计", level=2)

    add_para(doc, "【互动7】个人行动计划制定")
    add_script_box(doc, "【指令话术】", "现在是今天最重要的一个练习——制定你自己的行动计划。\n\n请拿出你的行动计划表，按照以下步骤填写：\n\n第一步（5分钟）：回顾今天的内容，选择1个你认为最重要的行为改变点\n第二步（5分钟）：写下你打算在哪个具体场景中实践这个行为\n第三步（5分钟）：设计你的\"开场白\"——你打算怎么在那个场景中说第一句话\n\n（等待15分钟，让学员充分思考和书写）\n\n第四步（10分钟）：两人一组，互相分享你的计划，并请对方给出1条建议")

    add_script_box(doc, "【行动计划表说明】", "行动计划表包含以下字段：\n• 我选择改变的行为：_______\n• 我计划实践的场景：_______\n• 我的开场白：_______\n• 我需要的支持/资源：_______\n• 我承诺的完成时间：_______\n• 我将如何评估自己的进步：_______")

    add_colorful_heading(doc, "6.4 时间分配", level=2)

    time_table4 = doc.add_table(rows=5, cols=4)
    time_table4.style = 'Table Grid'
    headers4 = ["环节", "时长", "教学方法", "物料"]
    rows4 = [
        ("回顾+导入", "5分钟", "讲授", "PPT第29页"),
        ("知识点：行为改变方法", "10分钟", "讲授", "PPT第30-31页"),
        ("知识点：3个开场白", "5分钟", "讲授", "PPT第32页"),
        ("互动：行动计划制定", "30分钟", "工具演练", "行动计划表"),
        ("总结+收尾", "10分钟", "讲授", "PPT第33页")
    ]
    add_time_table(doc, headers4, rows4)

    add_colorful_heading(doc, "6.5 讲师话术", level=2)

    add_script_box(doc, "【导入语】", "今天课程的最后，我们要做一个很多人觉得\"太虚\"但其实\"最实\"的事情——制定行动计划。\n\n我带过很多管理培训课程，发现一个规律：听的时候激动，听完感动，课后不动。\n\n为什么会这样？因为没有把\"想法\"变成\"行动\"。没有截止日期的计划，等于没有计划。\n\n所以接下来的30分钟，请大家认真地做这件事。这可能是你今天最有价值的30分钟。")

    add_script_box(doc, "【收尾语】", "在结束之前，我想请大家记住一个数字：21。\n\n行为科学研究表明，形成一个新习惯需要至少21天。\n\n从今天起，请你选择一个行为，坚持21天。21天后，你会看到一个不一样的自己，不一样的团队。\n\n最后，送给大家一句话：\n\n\"你不必等到完美才开始，你只需要开始才能完美。\"\n\n谢谢大家今天的学习。")

    add_colorful_heading(doc, "6.6 常见问题应对", level=2)

    add_qa_box(doc, "Q1：选1个行为太少了，能多选几个吗？", "理解你想改变的心情。但研究表明，同时改变多个行为的成功率很低。请只选1个——那个对你来说最重要、或者最容易做到的行为。做到1个之后，再换下一个。")

    add_qa_box(doc, "Q2：找不到\"反馈伙伴\"怎么办？", "如果实在找不到，可以用一个替代方法：用手机录下你自己在模拟场景中的讲话，然后回放给自己听。这也是一种\"自我反馈\"。或者，可以找一位你信任的、愿意说真话的朋友定期喝咖啡聊聊。")

    add_qa_box(doc, "Q3：如果21天后发现没效果怎么办？", "没效果的原因可能是：选择的场景不对、行为太难落地、反馈机制不健全。建议：重新评估你的计划，找到障碍点，调整策略再次尝试。关键是：不要放弃，而是迭代。")

    doc.add_page_break()

    # ========== 第七章 课程收尾 ==========
    add_colorful_heading(doc, "第七章 课程收尾", level=1)

    add_colorful_heading(doc, "7.1 总结要点", level=2)
    add_para(doc, "今天我们一起学习了心理安全感与信任文化。回顾一下核心要点：")
    add_para(doc, "1. 什么是心理安全感？——一种团队成员共同持有的信念：可以敞开心扉、冒险发声，而不用担心被惩罚。")
    add_para(doc, "2. 谷歌氧气计划的发现——心理安全感是团队绩效的最强预测因子，重要性超过团队成员的个人能力。")
    add_para(doc, "3. 如何评估？——从4个维度（表达想法、犯错误、寻求帮助、提出异议）来诊断团队现状。")
    add_para(doc, "4. 管理者能做什么？——做教练而非老板，赋权而非管控，关心成员成功。")
    add_para(doc, "5. 如何做到？——小步快跑，选1个行为，找1个场景，坚持21天。")

    add_colorful_heading(doc, "7.2 行动计划制定指导", level=2)
    add_para(doc, "课后行动计划的制定要点：")
    add_para(doc, "• 不要贪多：只选1个最需要改变的行为")
    add_para(doc, "• 具体场景：明确在什么时间、什么会议、什么情境下实践")
    add_para(doc, "• 开口第一句话：想好第一句话怎么说，这会大大降低行动门槛")
    add_para(doc, "• 设定截止日期：没有截止日期的计划是空想")
    add_para(doc, "• 找反馈机制：有人监督你，比独自努力的成功率高3倍")

    add_colorful_heading(doc, "7.3 课后跟进建议", level=2)
    add_para(doc, "为巩固学习效果，建议以下课后跟进措施：")
    add_para(doc, "（1）7天后：发邮件/微信提醒学员回顾自己的行动计划")
    add_para(doc, "（2）21天后：邀请学员填写\"行为改变自评问卷\"")
    add_para(doc, "（3）30天后：组织1次30分钟的线上/线下回顾会，分享实践心得")
    add_para(doc, "（4）90天后：进行第2轮团队心理安全感评估，对比变化")

    doc.add_page_break()

    # ========== 第八章 附录 ==========
    add_colorful_heading(doc, "第八章 附录", level=1)

    add_colorful_heading(doc, "附录A：讲师资源库", level=2)

    add_para(doc, "【更多案例】")
    add_para(doc, "案例5：Netflix的文化自由与责任")
    add_para(doc, "Netflix前CHO Patty McCord在《告诉给你的力量》中描述：Netflix的成功很大程度上归功于\"绝对透明\"的文化——任何人可以批评任何人的想法，只要是为了把事情做得更好。这种文化的基础，正是高度的心理安全感。")

    add_para(doc, "案例6：IDEO的\"快速失败\"文化")
    add_para(doc, "设计公司IDEO有一个著名的做法：鼓励设计师快速做出原型，然后公开接受批评。他们发现，当失败被正常化、批评被欢迎时，创新反而会加速。")

    add_para(doc, "【更多互动游戏】")
    game1 = doc.add_paragraph()
    game1.add_run("游戏1：优点轰炸（Strength Spotting）\n").bold = True
    game1.add_run("形式：小组进行\n目标：建立正向反馈的习惯\n步骤：每人轮流被\"轰炸\"——其他组员轮流说出这个人的1个优点/贡献，不许重复。被打\"轰炸\"的人只能说\"谢谢\"。\n时长：每人不超过3分钟。")

    game2 = doc.add_paragraph()
    game2.add_run("游戏2：沉默风暴（Silent Brainstorm）\n").bold = True
    game2.add_run("形式：个人思考+集体分享\n目标：让内向者也能充分表达\n步骤：给每个人发一张大白纸，要求在5分钟内写下对某个问题的所有想法（不许说话）。然后每人用2分钟时间向全组展示自己的\"风暴成果\"。\n时长：15-20分钟。")

    add_colorful_heading(doc, "附录B：参考文献", level=2)
    add_para(doc, "1. Edmondson, A. C. (1999). Psychological Safety and Learning Behavior in Work Teams. Administrative Science Quarterly, 44(2), 350-383.")
    add_para(doc, "2. Edmondson, A. C. (2018). The Fearless Organization: Creating Psychological Safety in the Workplace for Learning, Innovation, and Growth. Wiley.")
    add_para(doc, "3. Project Oxygen: The ten things the best managers do. (Google re:Work).")
    add_para(doc, "4. Cialdini, R. B. (2006). Influence: The Psychology of Persuasion. Harper Business.")
    add_para(doc, "5. Patterson, K., et al. (2012). Crucial Conversations: Tools for Talking When Stakes Are High. McGraw-Hill.")

    add_colorful_heading(doc, "附录C：推荐阅读", level=2)
    add_para(doc, "1. 《无畏的组织》（The Fearless Organization）—— Amy Edmondson\n推荐理由：作者是心理安全感概念的提出者，本书是理解这个概念的最佳入门读物。")
    add_para(doc, "2. 《关键对话》（Crucial Conversations）—— Patterson等\n推荐理由：提供了在高压环境下如何进行有效对话的实用工具。")
    add_para(doc, "3. 《团队协作的五大障碍》（The Five Dysfunctions of a Team）—— Patrick Lencioni\n推荐理由：用小说的形式讲述团队建设的常见误区和解决方法，可读性极强。")
    add_para(doc, "4. 《非暴力沟通》（Nonviolent Communication）—— Marshall Rosenberg\n推荐理由：提供了一种建立心理安全感的底层语言框架。")

    doc.add_paragraph()
    doc.add_paragraph()

    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_end = end_p.add_run("—— 课程完 ——")
    run_end.font.size = Pt(14)
    run_end.font.color.rgb = PRIMARY_COLOR
    run_end.font.bold = True

    output_path = "D:/新课开发/管理学/35-心理安全感与信任文化/04_讲师手册/35-心理安全感与信任文化_讲师手册.docx"
    doc.save(output_path)
    print("File saved to: " + output_path)
    return output_path

if __name__ == "__main__":
    create_instructor_manual()

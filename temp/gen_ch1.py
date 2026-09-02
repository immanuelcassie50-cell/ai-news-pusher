# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

OUTPUT_DIR = "D:/新课开发/工作手册/AI时代决策工作手册/完整课程包/07-练习材料/"

def sf(run, bold=False, size=12):
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def h(doc, text, level=1):
    hp = doc.add_heading(text, level)
    for r in hp.runs:
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return hp

def p(doc, text, bold=False, size=12):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    sf(r, bold=bold, size=size)
    return pp

def box(doc, title, content):
    pp = doc.add_paragraph()
    pp.paragraph_format.left_indent = Cm(1)
    r = pp.add_run(f"【{title}】{content}")
    sf(r, size=11)
    return pp

def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hc = t.rows[0].cells
    for i, h_text in enumerate(headers):
        hc[i].text = h_text
        for pp in hc[i].paragraphs:
            for r in pp.runs:
                r.font.bold = True
                r.font.name = 'Microsoft YaHei'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                r.font.size = Pt(10)
    for ri, rd in enumerate(rows):
        cells = t.rows[ri+1].cells
        for ci, cd in enumerate(rd):
            cells[ci].text = str(cd)
            for pp in cells[ci].paragraphs:
                for r in pp.runs:
                    r.font.name = 'Microsoft YaHei'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    r.font.size = Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def ans(doc):
    doc.add_paragraph()
    p(doc, "—" * 30, size=10)
    h(doc, "参考答案", 2)

# Chapter 1
doc = Document()
h(doc, '第一章练习：教训与决策卡辨别', 0)
p(doc, "练习目标：学会区分『教训』和『决策卡』的不同用途，正确判断何时产出哪种沉淀物。", bold=True, size=12)
p(doc, "判断标准：教训是一句提醒，解决『别再犯这个错』；决策卡是一份检查表+开关，解决『下一次遇到类似情况该怎么判断』。", size=11)
doc.add_paragraph()

h(doc, "一、案例判断练习", 1)
p(doc, "请阅读以下3个案例描述，判断每个案例应该产出『教训』还是『决策卡』，并说明理由。", size=11)

h(doc, "案例一：供应商紧急切换", 2)
box(doc, "背景", """
某科技公司供应链部门在一次关键原材料供应危机中，临时切换了备选供应商。
切换过程顺利，最终保证了项目进度。但事后复盘时发现，决策者当时判断
切换的原因非常模糊，只是『觉得原供应商不太靠谱』。

复盘会上，决策者说：『以后遇到这种情况要多比较几家供应商。』
""")
p(doc, "问题1：这个案例的产出物类型应该是？", size=11)
p(doc, "□ 教训          □ 决策卡", size=11)
p(doc, "问题2：说明你的判断理由：", size=11)
doc.add_paragraph("_" * 60)

h(doc, "案例二：项目提前完成的决策", 2)
box(doc, "背景", """
某项目团队在一个大型项目中，因为客户临时缩减了验收标准，项目提前一个月完成。
项目负责人复盘时认为，项目的提前完成是因为自己『快速拍板』做了几个关键决策。

另一位项目负责人学习了这次经验后，在另一个类似项目中模仿『快速拍板』的风格，
结果因为信息核实不足导致判断失误。
""")
p(doc, "问题1：这个案例的产出物类型应该是？", size=11)
p(doc, "□ 教训          □ 决策卡", size=11)
p(doc, "问题2：说明你的判断理由：", size=11)
doc.add_paragraph("_" * 60)

h(doc, "案例三：连续三次质量异常后的供应商切换", 2)
box(doc, "背景", """
某制造企业的质量总监在面对供应商连续三次出现质量异常时，做出了切换供应商的决策。
经过访谈，发现他当时的判断依据是：
- 信号1：某项关键指标连续两次超出历史波动区间
- 信号2：异常持续出现不止一次
- 信号3：没有明显的外部解释可以说明这次偏离

而且他还明确说出：『如果当时那个前提（供应商数据是准的）不成立，
我可能不会做同样的判断。』

这个场景未来可能被其他质量负责人重新遇到。
""")
p(doc, "问题1：这个案例的产出物类型应该是？", size=11)
p(doc, "□ 教训          □ 决策卡", size=11)
p(doc, "问题2：说明你的判断理由：", size=11)
doc.add_paragraph("_" * 60)

doc.add_paragraph()
h(doc, "二、判断依据说明", 1)
p(doc, "请完成下面的判断标准对照表：", size=11)
make_table(doc, ["判断维度", "教训适用", "决策卡适用"], [
    ["目的", "告诉人『别怎么做』", ""],
    ["适用情境", "结果不好的简单失误", ""],
    ["判断难度", "判断过程简单", ""],
    ["可复制性", "结果好但无法识别判断结构", ""],
    ["触发条件", "无需明确触发条件", ""]
], [3, 5, 5])

doc.add_paragraph()
ans(doc)
h(doc, "参考答案", 2)

h(doc, "案例一判断：教训（警告型）", 2)
p(doc, "理由：决策者的判断依据模糊（『觉得不太靠谱』），无法提炼出可结构化的判断规则。教训可以是：『判断切换供应商时，不能仅凭直觉，需要具体信号支撑』。但由于判断依据本身不清晰，不适合做决策卡。", size=11)

h(doc, "案例二判断：教训（警示型）", 2)
p(doc, "理由：这个案例恰恰说明，把『运气』当成『方法论』做成卡片最危险。真正导致项目提前完成的原因是客户缩减验收标准，与『快速拍板』无关。做决策卡会误导后续使用者。适合做的是警示案例：『避免把运气型成功归因于个人决策风格』。", size=11)

h(doc, "案例三判断：决策卡", 2)
p(doc, "理由：判断过程清晰，包含多个可识别的具体信号，有明确的触发条件，存在可检验的假设，且场景可复制。适合做成决策卡：触发条件可以是『关键指标连续两次超出历史波动区间+异常持续出现+无外部合理解释』。", size=11)

h(doc, "判断标准对照表参考答案：", 2)
make_table(doc, ["判断维度", "教训适用", "决策卡适用"], [
    ["目的", "告诉人『别怎么做』", "告诉人『遇到X信号时该怎么判断』"],
    ["适用情境", "结果不好的简单失误", "判断有难度的复杂决策"],
    ["判断难度", "判断过程简单", "存在多个选项，真正纠结过"],
    ["可复制性", "结果好但无法识别判断结构", "能提炼出可复用的判断结构"],
    ["触发条件", "无需明确触发条件", "必须有明确的可观测触发条件"]
], [3, 5, 5])

doc.save(OUTPUT_DIR + "01-第一章练习-教训与决策卡辨别.docx")
print("01 done")

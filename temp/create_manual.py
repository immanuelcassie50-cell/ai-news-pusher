# -*- coding: utf-8 -*-
"""
创建《人机协同权责边界与决策分级》讲师手册
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_para(doc, text, bold=False, indent=False):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = bold
    if indent:
        para.paragraph_format.left_indent = Inches(0.3)
    return para

def add_list_item(doc, text, level=0):
    """添加列表项"""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    return para

def add_warning_box(doc, text):
    """添加提示框"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run("⚠️ " + text)
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 100, 0)
    return para

def create_table_with_header(doc, headers, rows, col_widths=None):
    """创建带表头的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = '微软雅黑'
                run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
        set_cell_shading(header_cells[i], 'D9E2F3')

    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            for para in row_cells[col_idx].paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)

    return table

def main():
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('讲师手册', 0)
    for run in title.runs:
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(28)

    subtitle = doc.add_heading('人机协同权责边界与决策分级', 1)
    for run in subtitle.runs:
        run.font.name = '微软雅黑'
        run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(22)

    doc.add_heading('AI出内容谁把关效果与合规', 2)

    doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run('内部培训材料  |  请勿对外传播')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ========== 第一部分：讲师准备 ==========
    add_heading(doc, '第一部分：讲师准备', 1)

    # 课程概述与目标
    add_heading(doc, '课程概述与目标', 2)

    add_para(doc, '课程定位')
    add_para(doc, '人机协同权责边界与决策分级是一门帮助培训管理者和内训师建立AI内容审核意识的实操课程。本课程的核心命题是：')

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.5)
    run = quote.add_run('AI生成培训内容，谁来把关？把什么关？')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)

    add_para(doc, '课程区别于市面上大多数AI使用课或提示词课，本课程聚焦于"内容质量控制"这一关键环节——AI可以生成内容，但最终的质量把关和合规审核必须由人类负责。本课程帮助学员建立清晰的人机分工意识，掌握内容风险分类方法，明确各类内容的审核标准。')

    add_heading(doc, '课程核心问题', 3)
    add_list_item(doc, 'AI能做什么？——AI生成内容的能力边界与常见幻觉')
    add_list_item(doc, '人类要管什么？——内容风险的四维分类（效果、合规、伦理、法律）')
    add_list_item(doc, '谁来管？——组织内不同角色的权责分配')
    add_list_item(doc, '怎么管？——分级审核流程与决策标准')

    add_heading(doc, '学习目标', 3)
    add_list_item(doc, '学员能准确区分"AI生成内容"与"人类审核内容"的责任边界')
    add_list_item(doc, '学员能对培训内容进行风险等级分类（高/中/低）')
    add_list_item(doc, '学员掌握四级审核机制的核心要点')
    add_list_item(doc, '学员能独立设计一套适用于本企业的人机协同内容审核流程')

    # 学员画像分析
    add_heading(doc, '学员画像分析', 2)

    add_para(doc, '本课程适合以下类型的学员：')
    add_list_item(doc, '培训经理、培训总监，需要对AI生成的培训内容进行质量把控')
    add_list_item(doc, '内训师，想了解如何与AI协同工作同时保持内容质量')
    add_list_item(doc, 'HR负责人或学习发展负责人，正在引入AI工具到培训体系')
    add_list_item(doc, '企业大学或培训部门管理者，需要建立AI内容审核机制')
    add_list_item(doc, '对"AI生成内容谁来负责"有困惑的一线培训工作者')

    add_heading(doc, '学员层次建议', 3)
    add_para(doc, '本课程建议混合以下层次的学员，效果更佳：')
    add_list_item(doc, '有3年以上培训管理经验的中高层培训管理者（提供深度案例）')
    add_list_item(doc, '有1-3年经验的一线内训师或培训执行者（提供实操视角）')
    add_list_item(doc, '对AI工具已有初步接触但缺乏系统方法论的学员（接受新知）')

    add_warning_box(doc, '注意：如果学员中对"AI持完全否定态度"或"AI完全信任"的比例较高，需要在课前增加15分钟的观念引导环节。')

    # 讲师角色定位
    add_heading(doc, '讲师角色定位', 2)

    add_para(doc, '本课程中，讲师扮演以下三个核心角色：')

    add_para(doc, '1. 场景引导者', bold=True)
    add_para(doc, '不是传授知识，而是引导学员通过真实场景的讨论，自己得出结论。课程中的很多答案都是开放性的——"谁来管"没有标准答案，但有清晰的思考框架。', indent=True)

    add_para(doc, '2. 风险提醒者', bold=True)
    add_para(doc, '讲师最重要的价值不是讲清楚"怎么做"，而是提醒"哪里有坑"。AI内容审核最大的风险是"想当然"——觉得AI生成的就是对的。讲师要不断制造这种警觉感。', indent=True)

    add_para(doc, '3. 流程设计教练', bold=True)
    add_para(doc, '课程的产出是学员带回企业的审核流程设计。讲师不提供标准答案，而是提供模板和框架，让学员根据自己企业的情况设计适合的流程。', indent=True)

    # 课前准备清单
    add_heading(doc, '课前准备清单', 2)
    add_para(doc, '【建议】请在课前48小时内逐项确认以下事项')

    add_heading(doc, '内容熟悉度', 3)
    add_list_item(doc, '熟读本次讲授的所有模块内容，确保理解每个知识点')
    add_list_item(doc, '准备2-3个自己在AI内容审核方面遇到的真实挑战或案例')
    add_list_item(doc, '预演各练习的参考答案，确保能引导学员讨论')
    add_list_item(doc, '熟悉附录中的审核流程模板和风险分类表')

    add_heading(doc, '学员信息', 3)
    add_list_item(doc, '了解学员背景：行业、岗位层级、平均工作年限')
    add_list_item(doc, '了解学员企业是否已有AI内容使用经验（可通过报名问卷收集）')
    add_list_item(doc, '识别可能对课程提出质疑或认知抵触的学员类型')
    add_list_item(doc, '确认学员规模：20人以内（最佳）、20-30人（可行）、30人以上（需调整互动设计）')

    add_heading(doc, '材料准备', 3)
    add_list_item(doc, '每人一份学员手册（已打印，双面彩印）')
    add_list_item(doc, '每人一套配套练习表单（空表版，已装订）')
    add_list_item(doc, '白板/翻页纸板，用于现场板书和练习展示')
    add_list_item(doc, '计时器（建议使用投影计时，或手机分屏显示）')
    add_list_item(doc, '彩色马克笔每组一套（用于流程设计练习）')
    add_list_item(doc, 'A4纸备用，用于额外书写')

    add_heading(doc, '室内布置', 3)
    add_list_item(doc, '圆桌或岛型分组座位（4-6人一组）——优先于剧院式座位')
    add_list_item(doc, '每桌配备彩色马克笔 + 白板纸/便签纸')
    add_list_item(doc, '讲台可见但不孤立，讲师能方便地走到学员间')
    add_list_item(doc, '投影屏幕可从室内所有位置清晰阅读')

    # 场地与设备要求
    add_heading(doc, '场地与设备要求', 2)

    create_table_with_header(doc,
        ['设备/场地', '要求', '备注'],
        [
            ['投影设备', '分辨率1080P以上，屏幕可显示A3尺寸', '提前测试PPT在投影上的显示效果'],
            ['音响设备', '有线麦克风 + 无线麦克风各一', '确保讲师走动时声音清晰'],
            ['白板/电子白板', '至少一块白板，或电子白板设备', '用于现场绘制审核流程图'],
            ['分组座位', '4-6人为一组，U型或岛型布置', '避免剧院式（不利于讨论）'],
            ['照明', '可调节亮度，光线均匀', '避免投影区域有过强光线'],
            ['电源', '每组附近有电源插座', '用于学员笔记本充电'],
        ]
    )

    doc.add_paragraph()

    # 时间安排建议
    add_heading(doc, '时间安排建议', 2)

    add_para(doc, '本课程标准时长为7小时（不含午休），建议日程安排如下：')

    create_table_with_header(doc,
        ['时段', '内容', '时长'],
        [
            ['08:30-09:00', '签到与开场准备', '30分钟'],
            ['09:00-10:30', '第一部分：AI能做什么不能做什么 + 练习一', '90分钟'],
            ['10:30-10:45', '茶歇', '15分钟'],
            ['10:45-12:00', '第二部分：内容风险四维分类 + 练习二', '75分钟'],
            ['12:00-13:30', '午餐休息', '90分钟'],
            ['13:30-15:00', '第三部分：权责边界与角色分工 + 练习三', '90分钟'],
            ['15:00-15:15', '茶歇', '15分钟'],
            ['15:15-16:30', '第四部分：四级审核流程设计 + 练习四', '75分钟'],
            ['16:30-17:15', '第五部分：实操演练与案例复盘', '45分钟'],
            ['17:15-17:30', '总结与Q&A', '15分钟'],
        ]
    )

    doc.add_paragraph()
    add_warning_box(doc, '时间弹性：如果课程被压缩为6小时版本，建议将第五部分（实操演练）调整为课后作业。如果延长为两天版本，建议在第二天上午增加学员企业真实案例的现场设计辅导。')

    doc.add_page_break()

    # ========== 第二部分：教学方法论 ==========
    add_heading(doc, '第二部分：教学方法论', 1)

    # 体验式学习设计
    add_heading(doc, '体验式学习设计', 2)

    add_para(doc, '本课程采用"体验式学习"（Experiential Learning）模式，核心逻辑是：')

    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.5)
    run = quote.add_run('真实场景 → 冲突体验 → 反思提炼 → 工具习得 → 应用实践')
    run.font.name = '微软雅黑'
    run._r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.bold = True

    add_para(doc, '每个模块的设计都遵循这个循环：')
    add_list_item(doc, '【体验】先呈现一个真实的AI内容审核困境（或由学员分享）')
    add_list_item(doc, '【冲突】让学员感受到"直觉反应"和"专业判断"之间的差距')
    add_list_item(doc, '【反思】通过引导提问，让学员自己发现思维盲区')
    add_list_item(doc, '【工具】提供框架和工具，而不是标准答案')
    add_list_item(doc, '【实践】用自己企业的真实场景来应用工具')

    # 案例教学法应用
    add_heading(doc, '案例教学法应用', 2)

    add_para(doc, '本课程的案例分为三种类型：')

    add_para(doc, '1. 拿来即用的经典案例', bold=True)
    add_para(doc, '每个模块都配有精心设计的标准案例，如"AI生成的合规培训内容踩了哪些坑"。这些案例有完整的背景、冲突和高潮，适合在课堂上统一讨论。', indent=True)

    add_para(doc, '2. 学员贡献的真实案例', bold=True)
    add_para(doc, '鼓励学员分享自己在工作中遇到的真实困境（可匿名处理）。这是课程中最有价值的部分——因为真实，所以印象最深。', indent=True)

    add_para(doc, '3. 虚构的极端案例', bold=True)
    add_para(doc, '用于测试学员的判断边界，如"如果AI生成的内容99%正确但1%涉及隐私泄露，该怎么办"。极端案例的价值不在于答案，而在于引发思考。', indent=True)

    # 小组讨论引导技巧
    add_heading(doc, '小组讨论引导技巧', 2)

    add_para(doc, '小组讨论是本课程的核心教学方法。以下是关键引导技巧：')

    add_para(doc, '提问技巧', bold=True)
    add_list_item(doc, '【发散性提问】"你觉得这个问题出在哪里？"——鼓励学员表达')
    add_list_item(doc, '【聚焦性提问】"如果要归类，这个案例属于哪种风险类型？"——引导学员做判断')
    add_list_item(doc, '【挑战性提问】"你有没有不同的看法？"——避免群体思维')
    add_list_item(doc, '【行动性提问】"回到你的企业，你会怎么设计这个流程？"——促进应用')

    add_para(doc, '讨论节奏把控', bold=True)
    add_list_item(doc, '小组讨论时间建议：每题5-8分钟')
    add_list_item(doc, '讲师在讨论中巡视，但不介入过早')
    add_list_item(doc, '讨论结束后，邀请1-2组分享，不求全面但求深度')

    add_warning_box(doc, '常见问题：如果某个小组讨论偏离主题太远，讲师可以说"你们讨论得很深入，我们先把这个问题放一放，回到今天的核心问题——"')

    # 角色扮演与情景模拟
    add_heading(doc, '角色扮演与情景模拟', 2)

    add_para(doc, '本课程设计了三个关键角色扮演环节：')

    add_para(doc, '练习一：内容审核官（第三部分）', bold=True)
    add_para(doc, '学员扮演"内容审核官"，对AI生成的三类内容进行快速判断。时间压力：每题30秒。目的是让学员感受"直觉判断"和"系统思考"的差异。', indent=True)

    add_para(doc, '练习二：审核流程设计师（第四部分）', bold=True)
    add_para(doc, '学员分组设计一套完整的审核流程，然后交换方案、互相挑刺。目的是让学员从"被执行者"的角度思考"设计者"的难处。', indent=True)

    add_para(doc, '练习三：冲突谈判（第五部分）', bold=True)
    add_para(doc, '模拟"内训师认为内容没问题，但合规部门要求大幅修改"的冲突场景。学员扮演内训师，培训经理扮演合规审核官。目的是理解不同角色的立场和痛点。', indent=True)

    # 反思与总结方法
    add_heading(doc, '反思与总结方法', 2)

    add_para(doc, '每个模块结束时，用以下方式引导反思：')
    add_list_item(doc, '【三句话总结】请每位学员用三句话总结今天学到的最重要的内容')
    add_list_item(doc, '【一个行动】请每位学员写一个课后马上要做的行动（不是"以后要做"，而是"明天要做"）')
    add_list_item(doc, '【一个困惑】请每位学员写一个仍然存在的困惑，匿名收集后统一解答')

    doc.add_page_break()

    # ========== 第三部分：每个模块的讲师指引 ==========
    add_heading(doc, '第三部分：每个模块的讲师指引', 1)

    # ========== 模块一：导入 ==========
    add_heading(doc, '模块一 | AI时代培训内容的新问题', 2)
    add_para(doc, '——为什么AI生成的内容需要人类把关')
    add_para(doc, '建议时长：45-60分钟')

    add_heading(doc, '教学目标', 3)
    add_list_item(doc, '学员能说出AI在内容生成方面的三个核心能力边界')
    add_list_item(doc, '学员能识别AI生成内容的三种典型风险类型')
    add_list_item(doc, '学员理解"人机协同"不是"AI干活、人类收尾"，而是全程协作')

    add_heading(doc, '关键讲解点', 3)

    add_para(doc, '开场引导（约10分钟）', bold=True)
    add_para(doc, '不要直接说"今天学什么"。先呈现三个真实困境：', indent=True)
    add_list_item(doc, '困境A：AI生成的情境模拟脚本，让新员工学会了他不该学的"潜规则"')
    add_list_item(doc, '困境B：AI生成的销售话术培训材料，包含了对竞品的虚假陈述')
    add_list_item(doc, '困境C：AI生成的合规培训视频，出现了一个穿品牌竞争对手服装的演员')
    add_para(doc, '然后问学员："这些问题出在哪里？是AI的错吗？"让学员先在头脑里形成疑问——"AI不是号称很智能吗，为什么还会出这些问题？"然后再进入理论框架。', indent=True)

    add_para(doc, '核心概念讲解（约25分钟）', bold=True)
    add_list_item(doc, 'AI内容生成的三大能力边界：')
    add_list_item(doc, '边界1：AI擅长"组合已知"，但不擅长"发现未知"——AI生成的内容都是已知的组合，不会有真正的创新或发现')
    add_list_item(doc, '边界2：AI的"自信"不等于"正确"——AI生成的内容看起来正确，但可能包含事实性错误、逻辑谬误或版权问题')
    add_list_item(doc, '边界3：AI无法理解"组织上下文"——AI不了解你公司的文化、流程、潜规则，生成的内容可能"合法但不合情"')

    add_heading(doc, '学员可能的困难点', 3)
    add_list_item(doc, '学员可能过度神话AI能力，认为"AI说的都是对的"')
    add_list_item(doc, '学员可能过度妖魔化AI，认为"AI生成的内容都不能用"')
    add_list_item(doc, '学员可能混淆"AI生成内容"和"AI辅助搜索"的概念')

    add_heading(doc, '应对策略', 3)
    add_warning_box(doc, '如果学员认为"AI说的都是对的"——让他举例"你用AI查过一个你熟悉的领域的专业问题，AI的回答完全准确吗？"几乎所有有过此经历的学员都会承认"不完全准确"。')
    add_warning_box(doc, '如果学员认为"AI内容都不能用"——告诉他：问题的核心不是AI能不能用，而是"谁来把关"。AI是工具，工具没有责任，人才有责任。')

    add_heading(doc, '互动环节引导', 3)
    add_para(doc, '【互动设计】看完三个困境后，请2-3位学员分享他们自己遇到的类似困境。让全班一起分析：这是AI的什么问题？是能力边界问题，还是使用方式问题？')

    add_heading(doc, '进度把控建议', 3)
    add_list_item(doc, '如果开场讨论过于热烈（超过20分钟），适当打断："这个话题我们后面会深入探讨，先继续。"')
    add_list_item(doc, '如果学员对这个模块的基础概念有较大分歧，可以临时增加15分钟的补充讲解。')

    # ========== 模块二 ==========
    add_heading(doc, '模块二 | 内容风险四维分类', 2)
    add_para(doc, '——培训内容有哪些风险？谁来识别这些风险？')
    add_para(doc, '建议时长：60-75分钟')

    add_heading(doc, '教学目标', 3)
    add_list_item(doc, '学员能准确说出内容风险的四个维度（效果风险、合规风险、伦理风险、法律风险）')
    add_list_item(doc, '学员能对给定内容案例进行四维风险评估')
    add_list_item(doc, '学员理解不同风险维度可能相互交织，需要综合判断')

    add_heading(doc, '关键讲解点', 3)

    add_para(doc, '四维风险框架讲解', bold=True)

    create_table_with_header(doc,
        ['风险维度', '核心问题', '典型案例', '审核关注点'],
        [
            ['效果风险', '内容能达到培训目标吗？', 'AI生成的话术与实际业务流程不符', '准确性、适用性、时效性'],
            ['合规风险', '内容符合法规和公司政策吗？', '培训内容涉及未披露的敏感信息', '数据隐私、版权授权、言论合规'],
            ['伦理风险', '内容传递的价值观是否恰当？', '案例中隐含了对某群体的偏见', '多样性、公平性、文化敏感性'],
            ['法律风险', '内容是否可能引发法律纠纷？', '培训中的操作演示可能导致安全事故', '免责声明、知识产权、责任边界'],
        ]
    )

    doc.add_paragraph()

    add_para(doc, '⚠️ 关键认知点', bold=True)
    add_para(doc, '四个维度不是独立的——一个合规的内容可能有伦理问题（如使用了文化刻板印象），一个有效果的内容可能有法律风险（如操作演示不完整）。审核时需要综合判断。')

    add_heading(doc, '练习引导说明', 3)

    add_para(doc, '练习一：风险四维评估（第一级）', bold=True)
    add_para(doc, '引导方式：给学员10分钟独立填写风险评估表，然后两两对答案，再全班讨论分歧最大的1-2题。重点关注"伦理风险"这一列——这是最容易被学员忽略的维度。')

    add_para(doc, '练习二：风险归类辩论（第二级）', bold=True)
    add_para(doc, '引导方式：给出一个模糊案例，让学员辩论"这主要属于哪个维度的风险"。辩论的价值不在于结论，而在于让学员理解风险分类的复杂性。')

    add_heading(doc, '过渡与衔接', 3)
    add_para(doc, '这个模块结束后，学员通常会有一个感受："原来内容风险不只是对不对的问题，还有这么多维度需要考虑。"这个感受是这个模块最好的结尾。')
    add_para(doc, '过渡到第三部分时：可以说"识别风险是第一步，但知道谁来管才是关键。接下来我们来讨论组织内的权责分工。"')

    # ========== 模块三 ==========
    add_heading(doc, '模块三 | 权责边界与角色分工', 2)
    add_para(doc, '——AI出内容，谁来把关？')
    add_para(doc, '建议时长：60-75分钟')

    add_heading(doc, '教学目标', 3)
    add_list_item(doc, '学员能画出人机协同的内容生产流程图')
    add_list_item(doc, '学员能说出内容审核中"人类必须做"和"AI可以做"的边界')
    add_list_item(doc, '学员能根据自己企业的情况设计权责分配方案')

    add_heading(doc, '关键讲解点', 3)

    add_para(doc, '人机协同三角模型', bold=True)
    add_para(doc, '内容生产涉及三个核心角色：')
    add_list_item(doc, '【内容生产者】使用AI工具生成内容的人——通常是内训师或培训经理')
    add_list_item(doc, '【内容审核者】对内容质量负责的人——通常是直属领导或培训负责人')
    add_list_item(doc, '【合规顾问】提供专业合规意见的人——通常是合规/法务/HR政策部门')

    add_para(doc, '三个角色的权责边界：')

    create_table_with_header(doc,
        ['角色', '主要职责', '必须人类做', 'AI可以辅助'],
        [
            ['内容生产者', '生成初稿、收集需求、修改迭代', '理解业务背景、做出最终判断', '生成初稿、检查格式、查找资料'],
            ['内容审核者', '判断内容是否达到培训目标', '判断价值导向、审核最终版本', '检查事实错误、提出修改建议'],
            ['合规顾问', '提供合规意见、识别潜在风险', '出具合规意见、处理申诉', '初步筛查、提示潜在风险点'],
        ]
    )

    doc.add_paragraph()

    add_para(doc, '⚠️ 常见误区', bold=True)
    add_list_item(doc, '误区1："AI生成的内容，AI负责"——AI是工具，工具不负责，人才负责')
    add_list_item(doc, '误区2："审核只是最后一关"——审核应该贯穿全过程，不是收尾工作')
    add_list_item(doc, '误区3："有了AI，我们就不需要审核者了"——AI生成内容反而更需要审核，因为AI不知道自己错')

    add_heading(doc, '练习引导说明', 3)

    add_para(doc, '练习三：权责分配设计（第三级）', bold=True)
    add_para(doc, '引导方式：给学员15分钟，用自己企业的真实场景设计一套权责分配方案。完成后，邀请2-3位学员分享方案，全班一起挑刺：哪些地方可能出问题？')

    add_heading(doc, '进度把控建议', 3)
    add_list_item(doc, '这个模块的讨论容易发散——如果学员开始讨论"我们公司就是这样分工的，但有问题"，可以先记录下来，在第五部分的案例复盘环节统一处理。')

    # ========== 模块四 ==========
    add_heading(doc, '模块四 | 四级审核流程设计', 2)
    add_para(doc, '——内容分级，审核分级')
    add_para(doc, '建议时长：75-90分钟')

    add_heading(doc, '教学目标', 3)
    add_list_item(doc, '学员能说出内容风险的三级分类标准（高/中/低）')
    add_list_item(doc, '学员能根据内容风险等级匹配相应的审核流程')
    add_list_item(doc, '学员能独立设计一套四级审核流程')

    add_heading(doc, '关键讲解点', 3)

    add_para(doc, '内容风险分级标准', bold=True)

    create_table_with_header(doc,
        ['风险等级', '判断标准', '典型内容类型', '审核要求'],
        [
            ['高风险', '涉及合规要求、可能有法律风险', '合规培训、政策解读、操作安全类', '必须人工审核 + 合规顾问复核'],
            ['中风险', '涉及业务知识、可能有效果偏差', '技能培训、流程培训、产品知识类', 'AI初筛 + 人工审核'],
            ['低风险', '一般性知识、通用技能', '通识课程、入职须知、企业文化类', 'AI审核为主、人工抽检'],
        ]
    )

    doc.add_paragraph()

    add_para(doc, '四级审核流程', bold=True)

    create_table_with_header(doc,
        ['审核级别', '触发条件', '审核主体', '时间要求', '产出物'],
        [
            ['L1 自动审核', '低风险内容首次生成', 'AI系统自动执行', '实时', 'AI审核报告'],
            ['L2 人工抽检', '低风险内容修改版 / 中风险内容', '内容生产者自检 + 直属领导抽检', '24小时内', '抽检记录'],
            ['L3 专业审核', '中风险内容完整版 / 高风险内容', '培训负责人 + 相关业务专家', '48小时内', '审核意见书'],
            ['L4 最高审核', '高风险内容涉及合规/法律问题', '合规/法务 + 高管审批', '72小时内', '审批签字单'],
        ]
    )

    doc.add_paragraph()

    add_para(doc, '⚠️ 关键原则', bold=True)
    add_list_item(doc, '审核不是"卡流程"，而是"保质量"——审核者的角色是帮助内容变得更好，不是挑刺')
    add_list_item(doc, '审核级别越高，耗时越长——要在内容生成阶段就把风险控制住，而不是依赖审核来补救')
    add_list_item(doc, 'AI可以作为L1审核工具，但L2以上必须有人类判断——因为风险的核心是"价值判断"，AI不具备价值判断能力')

    add_heading(doc, '练习引导说明', 3)

    add_para(doc, '练习四：审核流程设计（第三级）', bold=True)
    add_para(doc, '引导方式：给学员20分钟，用自己企业的真实内容类型设计一套分级审核流程。完成后，让学员两两交换方案，互相模拟一次"高风险内容过审"的过程。')

    add_heading(doc, '过渡与衔接', 3)
    add_para(doc, '这个模块结束后，学员最常见的感受是："原来我们公司的审核要么太松（没有分级）要么太死板（所有内容都走最高审核）。"这个认知本身比任何技巧都重要。')

    # ========== 模块五 ==========
    add_heading(doc, '模块五 | 实操演练与案例复盘', 2)
    add_para(doc, '——用真实场景检验学习成果')
    add_para(doc, '建议时长：45-60分钟')

    add_heading(doc, '教学目标', 3)
    add_list_item(doc, '学员能对一个真实的AI内容生产案例进行全流程复盘')
    add_list_item(doc, '学员能识别案例中权责边界模糊的环节')
    add_list_item(doc, '学员能提出具体的改进方案')

    add_heading(doc, '案例复盘框架', 3)

    add_para(doc, '复盘维度（每个维度约10分钟）：')
    add_list_item(doc, '【发生了什么】描述AI生成内容的实际过程和结果')
    add_list_item(doc, '【哪里出了问题】识别内容风险的具体表现')
    add_list_item(doc, '【谁来负责】分析权责边界是否清晰')
    add_list_item(doc, '【可以怎么改】提出具体的改进措施')
    add_list_item(doc, '【如何避免下次】总结系统性改进建议')

    add_heading(doc, '互动环节引导', 3)

    add_para(doc, '【角色扮演】安排学员进行"冲突谈判"：')
    add_para(doc, '情境：内训师小王认为AI生成的销售话术培训材料非常好，但合规审核员要求删除30%的内容（因为涉及竞品对比表述）。小王觉得删了之后内容就不完整了。', indent=True)
    add_para(doc, '角色分配：内训师小王（学员A）、合规审核员（学员B）、培训经理（学员C，作为调解人）', indent=True)
    add_para(doc, '讨论问题：如何达成共识？内训师的"效果优先"和合规的"安全优先"如何平衡？', indent=True)

    add_heading(doc, '总结与收尾', 3)
    add_para(doc, '课程结束时，用以下方式收尾：')
    add_list_item(doc, '让每位学员写下"我离开这个教室之后，第一个要做的改变是什么"')
    add_list_item(doc, '邀请2-3位学员分享')
    add_list_item(doc, '讲师做最后的总结：强调"人机协同的核心不是分工，而是协作——AI是工具，人类是主人，工具好用不好用，取决于主人会不会用"')

    doc.add_page_break()

    # ========== 第四部分：工具与资源 ==========
    add_heading(doc, '第四部分：工具与资源', 1)

    # 课程辅助工具
    add_heading(doc, '课程辅助工具', 2)

    create_table_with_header(doc,
        ['工具名称', '用途', '获取方式'],
        [
            ['内容风险四维评估表', '对任意培训内容进行风险分类', '学员手册附录'],
            ['权责分配矩阵模板', '设计组织内人机协同的权责分工', '学员手册附录'],
            ['四级审核流程图', '可视化内容审核的流程和节点', '学员手册附录'],
            ['AI内容自检清单', '内容生产者在提交审核前的自检', '学员手册附录'],
            ['审核意见反馈表', '记录审核意见和改进要求', '学员手册附录'],
        ]
    )

    doc.add_paragraph()

    # PPT使用建议
    add_heading(doc, 'PPT使用建议', 2)

    add_list_item(doc, 'PPT不是讲师稿——每张PPT只呈现核心概念或案例，关键讲解在口述')
    add_list_item(doc, '案例类页面先展示案例再公布分析——让学员先思考再对照')
    add_list_item(doc, '工具类页面配合白板演示——在白板上画出流程图，边讲边画')
    add_list_item(doc, '避免大段文字——如果PPT上文字太多，学员会低头看PPT而不是听讲')
    add_list_item(doc, '关键数据可以放大突出——如"AI生成内容的准确率通常在70-80%"')

    # 投影与设备
    add_heading(doc, '投影与设备', 2)

    add_list_item(doc, '投影屏幕尺寸：能清晰显示A3尺寸内容为佳')
    add_list_item(doc, '翻页笔：建议使用带激光点功能的翻页笔，方便指向白板内容')
    add_list_item(doc, '音响设备：测试麦克风时走到教室后方确认声音清晰')
    add_list_item(doc, '备用设备：带一份PPT打印版在U盘里，以备投影故障')

    # 计时技巧
    add_heading(doc, '计时技巧', 2)

    add_list_item(doc, '每个模块开始前在白板上写下"本模块结束时间"')
    add_list_item(doc, '讨论环节使用投影计时器，让学员看到剩余时间')
    add_list_item(doc, '提前3分钟提醒："还有3分钟，我们来做个小结"')
    add_list_item(doc, '如果超时，可以跳过某些练习的分享环节，但不要压缩练习时间')

    # 现场应急处理
    add_heading(doc, '现场应急处理', 2)

    create_table_with_header(doc,
        ['场景', '应对方式'],
        [
            ['投影设备故障', '切换到白板模式，讲师在白板上画出原本要展示的内容框架'],
            ['学员提出超纲问题', '记录问题并在Q&A环节统一处理，或课后单独交流'],
            ['学员讨论过于激烈', '用"我们先暂停讨论，休息5分钟"来缓和气氛'],
            ['学员分享的内容涉及敏感信息', '提醒学员"这个案例我们课后单独讨论"'],
            ['时间不够', '优先保证练习三和练习四的时间，缩减讲解时间'],
        ]
    )

    doc.add_page_break()

    # ========== 第五部分：评估与反馈 ==========
    add_heading(doc, '第五部分：评估与反馈', 1)

    # 学员评估方法
    add_heading(doc, '学员评估方法', 2)

    add_para(doc, '本课程采用"过程性评估"，不设笔试，而是通过以下方式评估学员的学习效果：')

    add_para(doc, '1. 课堂参与度评估', bold=True)
    add_list_item(doc, '观察学员在小组讨论中的参与程度')
    add_list_item(doc, '记录学员在Q&A环节的提问质量')
    add_list_item(doc, '关注学员能否用自己的语言复述核心概念')

    add_para(doc, '2. 练习成果评估', bold=True)
    add_list_item(doc, '练习一（风险四维评估）：能否正确识别各维度的风险')
    add_list_item(doc, '练习三（权责分配设计）：设计的权责分工是否清晰合理')
    add_list_item(doc, '练习四（审核流程设计）：流程是否有完整的分级和明确的审核主体')

    add_para(doc, '3. 课后行动评估', bold=True)
    add_list_item(doc, '课后一周内收集学员的"第一个改变"实施情况')
    add_list_item(doc, '课后一个月进行回访，了解审核流程的落地情况')

    # 课程效果评估
    add_heading(doc, '课程效果评估', 2)

    add_para(doc, '讲师的课后评估（每次课程结束后填写）：')

    create_table_with_header(doc,
        ['评估维度', '评估问题', '评分（1-5）'],
        [
            ['内容匹配度', '课程内容是否符合学员需求？', ''],
            ['互动质量', '学员讨论是否深入有效？', ''],
            ['时间把控', '各模块时间安排是否合理？', ''],
            ['学员反馈', '学员对课程的整体评价如何？', ''],
            ['待改进点', '下次需要调整哪些内容或方式？', ''],
        ]
    )

    doc.add_paragraph()

    # 讲师自我反思
    add_heading(doc, '讲师自我反思', 2)

    add_para(doc, '每次课程结束后，建议讲师进行以下自我反思：')
    add_list_item(doc, '【最有效的环节】哪个练习或讨论让学员收获最大？')
    add_list_item(doc, '【最具挑战的环节】哪个环节的节奏最难把控？')
    add_list_item(doc, '【学员的共性困惑】学员问得最多的问题是什么？')
    add_list_item(doc, '【内容的时效性】哪些内容可能需要根据最新的AI能力更新？')
    add_list_item(doc, '【个人成长】作为讲师，我在这次课程中有什么成长？')

    # 持续改进建议
    add_heading(doc, '持续改进建议', 2)

    add_list_item(doc, '每半年更新一次案例库，替换过时的案例')
    add_list_item(doc, '每年根据AI能力变化，更新"AI能做什么不能做什么"的章节')
    add_list_item(doc, '收集学员的企业审核流程设计，作为内部最佳实践分享')
    add_list_item(doc, '建立"AI内容风险案例库"——收集内部真实的AI内容风险事件（脱敏处理）')
    add_list_item(doc, '开发进阶课程：如"AI辅助课程设计"、"AI内容合规专项"等')

    doc.add_page_break()

    # ========== 附录 ==========
    add_heading(doc, '附录', 1)

    add_heading(doc, '附录一：内容风险四维评估表模板', 2)

    create_table_with_header(doc,
        ['评估维度', '评估要点', '是/否', '风险说明'],
        [
            ['效果风险', '内容与业务目标一致', '', ''],
            ['效果风险', '内容准确无误', '', ''],
            ['效果风险', '内容适合目标受众', '', ''],
            ['合规风险', '不涉及隐私数据', '', ''],
            ['合规风险', '版权使用合规', '', ''],
            ['合规风险', '符合公司政策', '', ''],
            ['伦理风险', '传递正确价值观', '', ''],
            ['伦理风险', '避免文化偏见', '', ''],
            ['伦理风险', '尊重多样性', '', ''],
            ['法律风险', '操作安全提示完整', '', ''],
            ['法律风险', '免责声明清晰', '', ''],
            ['法律风险', '知识产权归属明确', '', ''],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, '附录二：权责分配矩阵模板', 2)

    create_table_with_header(doc,
        ['内容类型', '内容生产者', '直属领导', '合规审核', '最终审批'],
        [
            ['合规培训类', '生成初稿', '审核效果', '审核合规', '培训负责人'],
            ['技能培训类', '生成初稿', '审核效果', '抽检合规', '培训负责人'],
            ['通识课程类', '生成初稿', '审核效果', '抽检合规', '内容生产者自检'],
            ['高风险内容', '生成初稿', '审核效果', '审核合规', '高管 + 法务'],
        ]
    )

    doc.add_paragraph()

    add_heading(doc, '附录三：AI内容自检清单（内训师使用）', 2)

    add_list_item(doc, '□ 我已通读AI生成的全部内容，确认与业务需求一致')
    add_list_item(doc, '□ 我已核实内容中的事实、数据、案例来源可靠')
    add_list_item(doc, '□ 我已确认内容不涉及任何隐私、机密或版权问题')
    add_list_item(doc, '□ 我已检查内容中是否有任何可能导致误解或偏见的表述')
    add_list_item(doc, '□ 我已确认操作类内容包含完整的安全提示和免责说明')
    add_list_item(doc, '□ 我已完成上述自检，并将内容提交审核')

    add_para(doc, '自检人：__________  日期：__________')

    add_heading(doc, '附录四：术语表', 2)

    create_table_with_header(doc,
        ['术语', '定义'],
        [
            ['人机协同', '人类与AI系统共同完成任务的模式，强调两者的协作而非替代'],
            ['内容风险', '培训内容可能带来的负面影响，包括效果、合规、伦理、法律四个维度'],
            ['L1-L4审核', '四个级别的内容审核流程，从自动审核到最高级别审批'],
            ['权责边界', '不同角色在内容生产与审核中的职责边界'],
            ['AI幻觉', 'AI生成看似合理但实际错误的内容的现象'],
        ]
    )

    # 保存文档
    output_path = r'D:/新课开发/HR/培训/02_人机协同权责边界与决策分级：AI出内容谁把关效果与合规/讲师手册/讲师手册_人机协同权责.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
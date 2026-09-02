# -*- coding: utf-8 -*-
"""
创建学员手册_工业革命.docx
课程29：工业革命——为什么是英国率先起飞
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 输出路径
output_dir = r"D:/新课开发/经济学/29_工业革命"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "学员手册_工业革命.docx")

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {})
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    shading.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """设置单元格边距"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.makeelement(qn('w:tcMar'), {})
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        mar = tcMar.makeelement(qn(f'w:{side}'), {})
        mar.set(qn('w:w'), str(val))
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)

def add_horizontal_line(doc):
    """添加水平线"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {})
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '666666')
    pBdr.append(bottom)
    return para

# 创建文档
doc = Document()

# 设置页面
section = doc.sections[0]
section.page_width = Inches(8.27)  # A4
section.page_height = Inches(11.69)
section.left_margin = Inches(1.18)
section.right_margin = Inches(1.18)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# 设置文档默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# ====== 封面 ======
title = doc.add_heading('', level=0)
run = title.add_run('工业革命')
run.font.size = Pt(44)
run.font.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——为什么是英国率先起飞')
run.font.size = Pt(28)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_para.add_run('学员手册')
run.font.size = Pt(22)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# 学员信息表
info_table = doc.add_table(rows=4, cols=4)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

labels = ['学员姓名', '所在部门', '课程日期', '课程讲师']
for i, label in enumerate(labels):
    row = info_table.rows[i]
    cell = row.cells[0]
    cell.text = label
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(cell, 'E8E8E8')

doc.add_page_break()

# ====== 引言 ======
doc.add_heading('引言：在我们开始之前', level=1)

intro_text = """很多人把工业革命简单归因于"蒸汽机的发明"。这种理解不能说错，但过于表面。蒸汽机只是结果，不是原因。

工业革命为什么发生在英国，而不是法国、德国或中国？这背后是制度、资源禀赋、市场规模、技术积累等多因素长期积累的结果。这个问题看似是历史问题，实则是管理问题——理解一个国家如何在关键历史节点上实现跨越式发展，对分析今天的技术革命（如AI浪潮）有极其重要的类比价值。

这门课要给你的，不是又一个"工业革命的故事"，而是一套分析框架：重大技术革命起飞需要哪些条件？今天的AI浪潮是否具备类似的条件？"""

for para_text in intro_text.split('\n\n'):
    para = doc.add_paragraph(para_text)
    para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.space_after = Pt(12)

add_horizontal_line(doc)

# 课程全景图
doc.add_heading('课程全景图', level=2)

# 全景图表格
map_table = doc.add_table(rows=3, cols=1)
map_table.style = 'Table Grid'
map_table.alignment = WD_TABLE_ALIGNMENT.CENTER

row_data = [
    ('第三层：历史比较框架', '用历史比较方法分析当下技术变革，理解制度与技术协同演化'),
    ('第二层：起飞条件清单', '工业革命成功的六大关键要素：制度/资源/市场/技术/人力/金融'),
    ('第一层：问题溯源', '为什么是英国？摆脱"技术发明决定论"的简单史观')
]

for i, (title, desc) in enumerate(row_data):
    row = map_table.rows[i]
    cell = row.cells[0]
    cell.text = ''

    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p1.add_run('：' + desc)

    colors = ['FFF2CC', 'D9EAD3', 'D9D9D9']
    set_cell_shading(cell, colors[i])

doc.add_paragraph()

# ====== 课程目标 ======
doc.add_heading('课程目标', level=1)

goals = [
    '理解工业革命为什么在英国发生，摆脱单一因素解释的局限',
    '掌握"重大技术革命起飞条件清单"的六要素框架',
    '学会用历史比较方法分析今天的技术变革（AI革命）',
    '理解制度与技术协同演化的重要性，而非单纯的"技术决定论"'
]

for goal in goals:
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(goal)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ====== 第一章 ======
doc.add_heading('第一章　重新理解工业革命', level=1)

# 1.1
doc.add_heading('1.1 常见的误解', level=2)

para = doc.add_paragraph()
para.add_run('误解一：工业革命是蒸汽机发明的直接结果').bold = True
para = doc.add_paragraph('蒸汽机确实是工业革命的重要标志，但蒸汽机的大规模应用需要其他条件的配合：钢铁技术的进步、煤矿的开采、运输系统的完善。蒸汽机只是工业革命技术体系的最后一环。')

para = doc.add_paragraph()
para.add_run('误解二：工业革命是少数发明家的功劳').bold = True
para = doc.add_paragraph('瓦特、阿克莱特等人的发明固然重要，但工业革命是一个系统性的转变，涉及生产组织方式、资本市场、劳动制度等多个维度的变化。')

para = doc.add_paragraph()
para.add_run('误解三：工业革命可以复制粘贴').bold = True
para = doc.add_paragraph('许多国家试图复制英国的工业化路径，但效果参差。原因在于：工业革命不是单一技术的突破，而是整套制度安排的涌现性结果。')

# 1.2
doc.add_heading('1.2 历史学家的新发现', level=2)

para = doc.add_paragraph('过去30年，历史学研究有了重大转向：从关注"发明"转向关注"条件"。主要发现包括：')

findings = [
    ('制度学派观点', '产权保护、司法体系、合同执行能力——英国在这些"软件"上的优势，可能是更关键的因素。'),
    ('地理学派观点', '英国煤炭分布与工业中心的重合、岛屿地形带来的运输优势——地理禀赋不可忽视。'),
    ('社会学派观点', '新教伦理、圈地运动带来的劳动力商品化——社会结构的转变是重要推手。'),
    ('比较史学观点', '为什么法国没有发生工业革命？制度僵化、金融体系落后、劳动力市场分割——法国案例提供了反面教材。')
]

for title, content in findings:
    para = doc.add_paragraph()
    para.add_run('● ' + title + '：').bold = True
    para.add_run(content)

doc.add_page_break()

# 1.3
doc.add_heading('1.3 六要素分析框架', level=2)

para = doc.add_paragraph('本课程采用六要素框架来分析工业革命的发生条件。这个框架既适用于历史研究，也适用于分析今天的技术革命。')

# 六要素表格
factors_table = doc.add_table(rows=7, cols=2)
factors_table.style = 'Table Grid'

headers = ['要素', '说明']
for i, h in enumerate(headers):
    cell = factors_table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

factors_data = [
    ('1. 制度安排', '产权保护、合同执行、稳定的政治环境'),
    ('2. 资源禀赋', '煤炭、钢铁等关键原材料的可获得性'),
    ('3. 市场规模', '国内消费能力、国际贸易网络'),
    ('4. 技术积累', '科学革命、工程师传统、手工业基础'),
    ('5. 人力资本', '教育水平、识字率、技术工人储备'),
    ('6. 金融支持', '银行体系、资本市场、风险投资机制')
]

for i, (factor, desc) in enumerate(factors_data):
    row = factors_table.rows[i+1]
    row.cells[0].text = factor
    row.cells[1].text = desc
    for j, cell in enumerate(row.cells):
        if j == 0:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        set_cell_shading(cell, 'F2F2F2' if i % 2 == 0 else 'FFFFFF')

doc.add_paragraph()

# ====== 第二章 ======
doc.add_heading('第二章　英国的特殊条件', level=1)

# 2.1
doc.add_heading('2.1 制度优势：光荣革命带来了什么', level=2)

para = doc.add_paragraph('光荣革命（1688年）通常被视为英国崛起的真正起点。这场"不流血的革命"建立了几个关键制度：')

institutions = [
    ('议会主权', '王权受到限制，财政权归属议会。这意味着政府的权力被关进了笼子，私人产权得到保护。'),
    ('司法独立', '普通法传统得到强化，法律面前人人平等原则逐步确立。合同纠纷可以得到公正裁决。'),
    ('国债制度', '政府借贷机制的创新，使英国能够以低利率筹集大量资金用于战争和建设。'),
    ('东印度公司', '股份公司的出现，使得大规模集资成为可能，风险由众多投资者分担。')
]

for title, content in institutions:
    para = doc.add_paragraph()
    para.add_run('● ' + title + '：').bold = True
    para.add_run(content)

doc.add_heading('2.2 地理禀赋：岛国的意外优势', level=2)

geo_text = """英国的地理条件常被忽视，但实际上非常重要：

煤炭分布：英国煤矿主要分布在工业中心附近，运输成本极低。相比之下，中国的煤矿主要在北方，而工业中心在南方，运输成本高昂。

岛屿地形：相对封闭的地理环境，减少了军事威胁，可以将更多资源用于经济建设。不必维持庞大的常备军。

港口优势：众多的天然良港，为对外贸易提供了便利。伦敦成为世界金融中心，与此不无关系。

海洋法传统：海洋法体系保护了英国的航运利益，海上霸权为贸易扩张提供了保障。"""

for para_text in geo_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_heading('2.3 市场发育：消费社会的兴起', level=2)

market_text = """工业革命不仅是生产革命，也是消费革命。

圈地运动：一方面造成了农民的流离失所，另一方面也为农业规模化经营创造了条件。失地农民成为工业劳动力的来源。

海外市场：北美殖民地的存在，为英国工业品提供了广阔的市场。糖、烟草、棉花等原材料的进口，支撑了制造业的发展。

消费分层：新兴的中产阶级和工人阶级的消费需求，创造了国内市场的规模效应。棉纺织品的普及就是典型案例。"""

for para_text in market_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_page_break()

# ====== 第三章 ======
doc.add_heading('第三章　技术革命的协同效应', level=1)

doc.add_heading('3.1 不是一个发明，是一堆发明', level=2)

para = doc.add_paragraph('工业革命不是由一个天才发明家创造的，而是由一群发明家、技术工匠、企业家共同推动的。关键技术发明包括：')

tech_list = [
    ('纺织机械', '珍妮纺纱机（1764）、水力纺纱机（1769）、走锭纺纱机（1779）——效率不断提升'),
    ('蒸汽机', '纽可门大气机（1712）→瓦特改进型（1769）→高压蒸汽机（1800）——应用范围不断扩大'),
    ('钢铁技术', '焦炭炼铁（1709）、轧钢技术（1783）——材料革命'),
    ('运输革命', '运河（1757-1830）、铁路（1825）——运输成本大幅下降')
]

for title, content in tech_list:
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(title + '：').bold = True
    para.add_run(content)

doc.add_heading('3.2 技术簇群效应', level=2)

para = doc.add_paragraph('理解工业革命的关键，不是某个单一发明，而是"技术簇群"——多个相关技术相互促进、相互强化，形成一个整体大于部分之和的系统。')

# 技术簇群表格
cluster_table = doc.add_table(rows=5, cols=2)
cluster_table.style = 'Table Grid'

cluster_data = [
    ('技术领域', '协同关系'),
    ('煤炭→蒸汽机→钢铁', '煤炭开采支撑蒸汽机应用，蒸汽机驱动钢铁生产'),
    ('钢铁→铁路→煤矿', '钢铁用于铁路建设，铁路扩大煤矿市场'),
    ('纺织→蒸汽机→棉花', '蒸汽机驱动纺织机械，美国棉花供应支持英国纺织业'),
    ('运河→铁路→全国市场', '运输革命将地区市场整合为全国统一市场')
]

for i, (col1, col2) in enumerate(cluster_data):
    row = cluster_table.rows[i]
    row.cells[0].text = col1
    row.cells[1].text = col2
    if i == 0:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(cell, '4472C4')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    else:
        for cell in row.cells:
            set_cell_shading(cell, 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_paragraph()

doc.add_heading('3.3 企业家精神', level=2)

entre_text = """工业革命不仅是技术革命，也是企业家精神的胜利。

阿克莱特（Richard Arkwright）：不是发明家，但他是企业家精神的代表。他将分散的纺织工匠组织成工厂体系，创建了现代工厂制度。

博尔顿（Matthew Boulton）：瓦特的合伙人，提供了商业眼光和市场网络，将实验室发明转化为商业产品。

斯蒂芬森（George Stephenson）：铁路之父，既是工程师又是企业家，推动了铁路技术的实用化。"""

for para_text in entre_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_page_break()

# ====== 第四章 ======
doc.add_heading('第四章　比较视角：为什么不是其他国家', level=1)

doc.add_heading('4.1 法国的教训', level=2)

france_text = """法国拥有许多类似的条件：也有煤矿，也有技术创新传统，也有殖民地。但法国为什么没有发生工业革命？

制度僵化：法国的行会制度保护既得利益者，阻碍了技术扩散。新技术要获得行会认可，困难重重。

金融体系落后：法国的金融体系以家族银行为主，没有形成现代银行体系。中小企业融资困难。

劳动力市场分割：法国的封建残余阻碍了劳动力流动。农民被束缚在土地上，无法成为自由的雇佣劳动者。

文化因素：法国社会重视文学艺术，轻视技术与商业。"学而优则仕"的传统观念影响了工业投资。"""

for para_text in france_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_heading('4.2 中国的比较', level=2)

china_text = """从比较史学的视角，中国提供了一个有趣的对照。

辉煌的古代：中国在宋朝就出现了资本主义萌芽，有发达的手工业和商业网络。

为何停滞？：明清两代的闭关锁国、政策导向的压抑、新教伦理的缺失——这些因素的综合作用，使中国错过了工业革命的早班车。

历史的教训：技术创新需要制度环境的配合。没有产权保护、技术扩散渠道、劳动力市场，技术发明只能停留在实验室。"""

for para_text in china_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_heading('4.3 后发国家的启示', level=2)

late_text = """德国和美国的工业革命是"后发优势"的典型案例。

德国（19世纪后期）：通过国家主导的工业化战略，建立了一套完善的职业教育和大学体系。技术引进+本土创新，实现了跨越式发展。

美国（19世纪末）：移民国家带来的劳动力红利、广袤的国内市场、大规模的铁路建设——美国在19世纪末成为世界第一工业大国。"""

for para_text in late_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_page_break()

# ====== 第五章 ======
doc.add_heading('第五章　历史类比：AI革命与工业革命', level=1)

doc.add_heading('5.1 类比框架', level=2)

para = doc.add_paragraph('用六要素框架来分析今天的AI革命，会得到一些有趣的洞见：')

# 类比表格
analogy_table = doc.add_table(rows=7, cols=3)
analogy_table.style = 'Table Grid'

analogy_data = [
    ('要素', '工业革命（英国）', 'AI革命（当代）'),
    ('制度', '普通法系、产权保护、议会监督', '数据确权、AI伦理规范、监管框架'),
    ('资源', '煤矿+铁矿+港口', '数据+算力+算法人才'),
    ('市场', '殖民地贸易网络+国内消费', '全球互联网+企业级应用'),
    ('技术', '科学革命+工程师传统', '深度学习+开源生态'),
    ('人力', '技术工人+识字率提升', 'AI研究人员+码农红利'),
    ('金融', '银行+股份公司+债券市场', 'VC+PE+二级市场科技股')
]

for i, row_data in enumerate(analogy_data):
    row = analogy_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(row.cells[j], '4472C4')
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_shading(row.cells[j], 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_paragraph()

doc.add_heading('5.2 关键差异', level=2)

diff_text = """类比不是等同。AI革命与工业革命有几个关键差异：

速度：工业革命用了近百年才完成全面渗透。AI革命在更短时间内席卷全球，影响更为深远。

范围：工业革命主要影响物质生产领域。AI革命正在渗透知识生产、行政管理、医疗教育等几乎所有领域。

驱动力：工业革命由私人企业家主导。AI革命由大企业和政府共同推动，研发成本更高。

监管：工业革命时期缺乏系统的环境和劳工监管。AI革命面临的伦理和监管挑战更为复杂。"""

for para_text in diff_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_heading('5.3 类比的价值与局限', level=2)

value_text = """类比的价值：

历史类比帮助我们看到当下技术革命的"结构性位置"——类似的条件是否具备？哪些条件是短板？类比帮助决策者避免"技术决定论"的陷阱，关注制度环境的建设。

类比的局限：

每个时代的技术革命都有其独特性。简单套用历史案例，可能导致误判。类比是思考的工具，不是预测的公式。"""

for para_text in value_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_page_break()

# ====== 第六章 ======
doc.add_heading('第六章　重大技术革命起飞条件清单', level=1)

doc.add_heading('6.1 清单使用指南', level=2)

para = doc.add_paragraph('这份清单是本课程的核心产出。请用这份清单分析你关心的任何技术革命，包括AI、区块链、生物技术等。')

# 起飞条件清单
checklist_table = doc.add_table(rows=8, cols=3)
checklist_table.style = 'Table Grid'

checklist_data = [
    ('序号', '条件', '评估要点'),
    ('1', '制度安排', '是否有清晰的产权保护？政府是鼓励创新还是限制竞争？'),
    ('2', '资源禀赋', '关键技术资源是否可获得？成本是否可控？'),
    ('3', '市场规模', '是否有足够大的初始市场支撑规模化生产？'),
    ('4', '技术积累', '基础科学是否有足够储备？工程师红利是否具备？'),
    ('5', '人力资本', '是否有足够的研发人才和应用人才储备？'),
    ('6', '金融支持', '是否有多元化融资渠道？VC/PE/银行是否愿意投资？'),
    ('7', '应用场景', '是否有清晰的第一批付费用户和应用场景？')
]

for i, row_data in enumerate(checklist_data):
    row = checklist_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
            set_cell_shading(row.cells[j], '4472C4')
            for para in row.cells[j].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            set_cell_shading(row.cells[j], 'F2F2F2' if i % 2 == 1 else 'FFFFFF')

doc.add_paragraph()

doc.add_heading('6.2 清单应用案例', level=2)

case_text = """案例：判断AI大模型革命是否具备起飞条件

制度安排：数据确权法规、AI伦理准则——各国正在建设中，但尚不完善。
资源禀赋：算力成本下降，但高端芯片受限——部分满足。
市场规模：企业级应用市场广阔，全球潜在用户数十亿——满足。
技术积累：深度学习十年积累，Transformer架构突破——满足。
人力资本：全球AI研究人员约数十万，中国AI工程师红利——满足。
金融支持：VC/PE大量投入，科技巨头持续投资——满足。
应用场景：客服、内容生成、代码辅助——第一批场景已验证。

综合评估：AI大模型革命已具备大部分起飞条件，但制度建设和芯片供应是主要短板。"""

for para_text in case_text.strip().split('\n\n'):
    p = doc.add_paragraph(para_text.strip())
    p.paragraph_format.first_line_indent = Cm(0.74)

doc.add_page_break()

# ====== 练习 ======
doc.add_heading('练习与思考', level=1)

doc.add_heading('练习一：六要素分析', level=2)

ex1_text = """选取你熟悉的一个技术变革案例（如新能源汽车、电商平台、在线教育等），用六要素框架进行分析：

要素一：制度安排
这项技术变革需要哪些制度支持？当前制度环境是否具备？

要素二：资源禀赋
关键技术资源是什么？中国在这项资源上的禀赋如何？

要素三：市场规模
目标市场有多大？是否足够支撑规模化？

要素四：技术积累
技术基础是否成熟？还需要哪些突破？

要素五：人力资本
需要哪些人才？供给是否充足？

要素六：金融支持
需要多少资金？融资渠道是否畅通？"""

for para_text in ex1_text.strip().split('\n\n'):
    doc.add_paragraph(para_text.strip())

doc.add_heading('练习二：历史比较', level=2)

ex2_text = """比较工业革命时期的英国和今天的中国，填写下表：

比较维度	工业革命英国	今日中国
制度优势
资源禀赋
市场规模
技术短板
金融体系
人才储备
"""

doc.add_paragraph(ex2_text.strip())

doc.add_heading('练习三：起飞条件评估', level=2)

ex3_text = """运用"重大技术革命起飞条件清单"，评估你所在行业的一个新技术变革：

我的行业：
我选择的技术变革：

评估结果：
"""

doc.add_paragraph(ex3_text.strip())

doc.add_page_break()

# ====== 延伸阅读 ======
doc.add_heading('延伸阅读推荐', level=1)

readings = [
    ('《工业革命》', '罗伯特·艾伦著——牛津大学公开课教材，简洁权威'),
    ('《大国的兴衰》', '保罗·肯尼迪著——从技术、经济角度分析大国崛起'),
    ('《蒸汽机驱动世界》', 'BBC纪录片——工业革命技术史'),
    ('《AI 2041》', '李开复等著——AI革命展望'),
    ('《创新的演化》', '沃尔玛·布里克曼著——技术创新的历史分析')
]

for title, desc in readings:
    para = doc.add_paragraph(style='List Bullet')
    para.add_run(title).bold = True
    para.add_run('：' + desc)

# ====== 笔记区域 ======
doc.add_heading('课程笔记', level=1)

notes_intro = doc.add_paragraph('在这里记录你在课程中的思考和发现：')

# 笔记表格
notes_table = doc.add_table(rows=10, cols=1)
notes_table.style = 'Table Grid'

for i in range(10):
    cell = notes_table.rows[i].cells[0]
    cell.text = f'笔记 {i+1}：'
    set_cell_shading(cell, 'FFFFCC' if i % 2 == 0 else 'FFFFFF')

doc.add_paragraph()
doc.add_paragraph()

# ====== 附录 ======
doc.add_heading('附录一：术语表', level=1)

terms = [
    ('工业革命', '18世纪后期至19世纪中期，以机器取代手工劳动为标志的生产方式根本性变革'),
    ('起飞条件', '经济学家罗斯托提出的概念，指经济实现持续增长所需的关键条件'),
    ('技术簇群', '多个相关技术相互促进、相互强化，形成整体大于部分之和的效应'),
    ('产权保护', '保护私人财产不受侵犯的法律和制度安排'),
    ('企业家精神', '识别机会、承担风险、整合资源、创造价值的综合能力'),
    ('比较史学', '通过对不同国家/地区的历史比较，发现历史规律的研究方法'),
    ('后发优势', '发展中国家可以借鉴先进国家的技术和经验，实现跨越式发展')
]

terms_table = doc.add_table(rows=len(terms)+1, cols=2)
terms_table.style = 'Table Grid'

term_headers = ['术语', '定义']
for i, h in enumerate(term_headers):
    cell = terms_table.rows[0].cells[i]
    cell.text = h
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    set_cell_shading(cell, '4472C4')
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

for i, (term, definition) in enumerate(terms):
    row = terms_table.rows[i+1]
    row.cells[0].text = term
    row.cells[1].text = definition
    for cell in row.cells:
        set_cell_shading(cell, 'F2F2F2' if i % 2 == 0 else 'FFFFFF')

# 页脚
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = '工业革命学员手册'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
doc.save(output_path)
print(f"学员手册已保存至: {output_path}")

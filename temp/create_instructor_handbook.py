# -*- coding: utf-8 -*-
"""
创建创新管理体系讲师手册
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_module_section(doc, module_num, module_title, duration, objectives, opening, key_points, exercises, transitions, notes):
    """添加模块完整内容"""
    # 模块标题
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(f"第{module_num}部分 | {module_title}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    run_sub = subtitle.add_run(f"——{module_title.split('——')[1] if '——' in module_title else ''}")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # 建议时长
    duration_para = doc.add_paragraph()
    duration_para.paragraph_format.space_before = Pt(0)
    duration_para.paragraph_format.space_after = Pt(12)
    run_dur_label = duration_para.add_run("建议时长：")
    run_dur_label.font.bold = True
    run_dur_label.font.size = Pt(11)
    run_dur = duration_para.add_run(duration)
    run_dur.font.size = Pt(11)

    # 模块目标
    obj_heading = doc.add_paragraph()
    obj_heading.paragraph_format.space_before = Pt(6)
    obj_heading.paragraph_format.space_after = Pt(3)
    run_obj_h = obj_heading.add_run("模块目标")
    run_obj_h.font.bold = True
    run_obj_h.font.size = Pt(11)
    run_obj_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    for obj in objectives:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.paragraph_format.space_before = Pt(0)
        bullet.paragraph_format.space_after = Pt(2)
        bullet.paragraph_format.left_indent = Inches(0.3)
        run = bullet.add_run(obj)
        run.font.size = Pt(10.5)

    # 开场引导
    opening_heading = doc.add_paragraph()
    opening_heading.paragraph_format.space_before = Pt(9)
    opening_heading.paragraph_format.space_after = Pt(3)
    run_open_h = opening_heading.add_run("开场引导")
    run_open_h.font.bold = True
    run_open_h.font.size = Pt(11)
    run_open_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    open_para = doc.add_paragraph()
    open_para.paragraph_format.space_before = Pt(0)
    open_para.paragraph_format.space_after = Pt(8)
    run_open = open_para.add_run(opening)
    run_open.font.size = Pt(10.5)

    # 核心工具讲解要点
    if key_points:
        kp_heading = doc.add_paragraph()
        kp_heading.paragraph_format.space_before = Pt(6)
        kp_heading.paragraph_format.space_after = Pt(3)
        run_kp_h = kp_heading.add_run("核心工具讲解要点")
        run_kp_h.font.bold = True
        run_kp_h.font.size = Pt(11)
        run_kp_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for kp in key_points:
            kp_para = doc.add_paragraph(style='List Bullet')
            kp_para.paragraph_format.space_before = Pt(2)
            kp_para.paragraph_format.space_after = Pt(2)
            kp_para.paragraph_format.left_indent = Inches(0.3)
            run = kp_para.add_run(kp)
            run.font.size = Pt(10.5)

    # 练习引导说明
    if exercises:
        ex_heading = doc.add_paragraph()
        ex_heading.paragraph_format.space_before = Pt(9)
        ex_heading.paragraph_format.space_after = Pt(3)
        run_ex_h = ex_heading.add_run("练习引导说明")
        run_ex_h.font.bold = True
        run_ex_h.font.size = Pt(11)
        run_ex_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        for ex in exercises:
            ex_para = doc.add_paragraph(style='List Bullet')
            ex_para.paragraph_format.space_before = Pt(2)
            ex_para.paragraph_format.space_after = Pt(2)
            ex_para.paragraph_format.left_indent = Inches(0.3)
            run = ex_para.add_run(ex)
            run.font.size = Pt(10.5)

    # 过渡与衔接
    if transitions:
        trans_heading = doc.add_paragraph()
        trans_heading.paragraph_format.space_before = Pt(9)
        trans_heading.paragraph_format.space_after = Pt(3)
        run_trans_h = trans_heading.add_run("过渡与衔接")
        run_trans_h.font.bold = True
        run_trans_h.font.size = Pt(11)
        run_trans_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        trans_para = doc.add_paragraph()
        trans_para.paragraph_format.space_before = Pt(0)
        trans_para.paragraph_format.space_after = Pt(8)
        run_trans = trans_para.add_run(transitions)
        run_trans.font.size = Pt(10.5)
        run_trans.font.italic = True

    # 讲师注意事项
    if notes:
        notes_heading = doc.add_paragraph()
        notes_heading.paragraph_format.space_before = Pt(6)
        notes_heading.paragraph_format.space_after = Pt(3)
        run_notes_h = notes_heading.add_run("讲师注意事项")
        run_notes_h.font.bold = True
        run_notes_h.font.size = Pt(11)
        run_notes_h.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        for note in notes:
            note_para = doc.add_paragraph(style='List Bullet')
            note_para.paragraph_format.space_before = Pt(2)
            note_para.paragraph_format.space_after = Pt(2)
            note_para.paragraph_format.left_indent = Inches(0.3)
            run = note_para.add_run(note)
            run.font.size = Pt(10.5)

    return doc

def create_instructor_handbook():
    doc = Document()

    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    # ========== 封面页 ==========
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(120)
    title_para.paragraph_format.space_after = Pt(20)
    run = title_para.add_run("创新管理体系")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(0)
    subtitle_para.paragraph_format.space_after = Pt(40)
    run_sub = subtitle_para.add_run("从创意到商业化的完整路径")
    run_sub.font.size = Pt(18)
    run_sub.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(40)
    run_line = line_para.add_run("━" * 50)
    run_line.font.size = Pt(12)
    run_line.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    handbook_para = doc.add_paragraph()
    handbook_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    handbook_para.paragraph_format.space_before = Pt(0)
    handbook_para.paragraph_format.space_after = Pt(10)
    run_hb = handbook_para.add_run("讲 师 手 册")
    run_hb.font.size = Pt(28)
    run_hb.font.bold = True
    run_hb.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    fac_para = doc.add_paragraph()
    fac_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fac_para.paragraph_format.space_before = Pt(0)
    fac_para.paragraph_format.space_after = Pt(60)
    run_fac = fac_para.add_run("FACILITATOR GUIDE")
    run_fac.font.size = Pt(14)
    run_fac.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    notice_para = doc.add_paragraph()
    notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_para.paragraph_format.space_before = Pt(60)
    run_notice = notice_para.add_run("本手册为内部培训使用材料，请勿对外传播")
    run_notice.font.size = Pt(10)
    run_notice.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
    run_notice.font.italic = True

    # ========== 目录页 ==========
    doc.add_page_break()
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_before = Pt(20)
    toc_title.paragraph_format.space_after = Pt(20)
    run_toc = toc_title.add_run("目 录")
    run_toc.font.size = Pt(20)
    run_toc.font.bold = True
    run_toc.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    toc_items = [
        ("课程概述", "课程定位 / 设计理念 / 适用学员"),
        ("课程结构总览", "七部分核心主题与工具总表"),
        ("课时建议", "标准版 / 精简版 / 系列课 / 工作坊"),
        ("课前准备", "讲师准备清单 / 材料准备 / 室内布置"),
        ("如何使用学员手册", "四种使用模式说明"),
        ("第零部分", "创新全链路导论：建立元框架"),
        ("第一部分", "创新战略与组织氛围"),
        ("第二部分", "创意产生与筛选机制"),
        ("第三部分", "创新项目立项与Stage-Gate"),
        ("第四部分", "创新团队组建与管理"),
        ("第五部分", "创新成果转化与商业化"),
        ("第六部分", "创新评估与持续改进"),
        ("第七部分", "综合演练与行动规划"),
    ]

    for i, (main, sub) in enumerate(toc_items, 1):
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_before = Pt(8)
        toc_para.paragraph_format.space_after = Pt(4)
        run_num = toc_para.add_run(f"{i:2d}. {main}")
        run_num.font.size = Pt(11)
        run_num.font.bold = True
        run_sub = toc_para.add_run(f"  —  {sub}")
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # ========== 关于本课程 ==========
    doc.add_page_break()
    about_title = doc.add_paragraph()
    about_title.paragraph_format.space_before = Pt(20)
    about_title.paragraph_format.space_after = Pt(16)
    run_about = about_title.add_run("关于本课程")
    run_about.font.size = Pt(18)
    run_about.font.bold = True
    run_about.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    pos_heading = doc.add_paragraph()
    pos_heading.paragraph_format.space_before = Pt(12)
    pos_heading.paragraph_format.space_after = Pt(6)
    run_pos = pos_heading.add_run("课程定位")
    run_pos.font.bold = True
    run_pos.font.size = Pt(12)
    run_pos.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    pos_para = doc.add_paragraph()
    pos_para.paragraph_format.space_before = Pt(0)
    pos_para.paragraph_format.space_after = Pt(8)
    run_pos_text = pos_para.add_run(
        '创新管理体系是一门帮助企业管理者建立"完整创新管理认知通路"的综合课程。'
        '区别于市面上大多数创新培训或创业课程，本课程的核心命题是：'
    )
    run_pos_text.font.size = Pt(11)

    highlight_para = doc.add_paragraph()
    highlight_para.paragraph_format.space_before = Pt(4)
    highlight_para.paragraph_format.space_after = Pt(12)
    highlight_para.paragraph_format.left_indent = Inches(0.5)
    run_high = highlight_para.add_run(
        '创新失败，大多数时候不是创意问题，而是体系问题——'
        '你有了好想法，但没有一套从创意到商业化的完整体系来支撑它落地。'
    )
    run_high.font.size = Pt(11)
    run_high.font.bold = True
    run_high.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    pos_desc = doc.add_paragraph()
    pos_desc.paragraph_format.space_before = Pt(0)
    pos_desc.paragraph_format.space_after = Pt(8)
    run_pos_desc = pos_desc.add_run(
        '课程将创新管理体系拆解为七个部分，从战略制定到最终评估，覆盖创新全生命周期。'
        '学员通过识别创新阶段→选择正确工具→建立完整流程，完成从"零散创新活动"到"系统化创新体系"的跃迁。'
    )
    run_pos_desc.font.size = Pt(11)

    # 设计理念
    design_heading = doc.add_paragraph()
    design_heading.paragraph_format.space_before = Pt(12)
    design_heading.paragraph_format.space_after = Pt(6)
    run_des = design_heading.add_run("设计理念")
    run_des.font.bold = True
    run_des.font.size = Pt(12)
    run_des.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    design_table = doc.add_table(rows=6, cols=2)
    design_table.style = 'Table Grid'

    hdr_cells = design_table.rows[0].cells
    hdr_cells[0].text = "设计原则"
    hdr_cells[1].text = "在课程中的体现"
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        set_cell_background(cell, "D6E3F8")

    design_items = [
        ("问题驱动，非知识点驱动", "每个模块以一个真实失败案例开场，工具在解决问题的过程中自然涌现"),
        ("先震撼后解释", "开场先呈现反直觉的案例对比，制造认知冲击，然后再解释背后逻辑"),
        ("练习是骨架", "三级难度递进：识别→应用→创造，每个知识点后都有配套练习"),
        ("内容可超，节奏不拖", "模块内容多于实际讲授时间，讲师可灵活取舍；每30-40分钟有切换"),
        ("链路贯通机制", "每个模块的产出为下一模块的素材，最终流入整合模块形成完整方案"),
    ]

    for i, (principle, manifestation) in enumerate(design_items, 1):
        row_cells = design_table.rows[i].cells
        row_cells[0].text = principle
        row_cells[1].text = manifestation
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # 适用学员
    audience_heading = doc.add_paragraph()
    audience_heading.paragraph_format.space_before = Pt(16)
    audience_heading.paragraph_format.space_after = Pt(6)
    run_aud = audience_heading.add_run("适用学员")
    run_aud.font.bold = True
    run_aud.font.size = Pt(12)
    run_aud.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    audience_intro = doc.add_paragraph()
    audience_intro.paragraph_format.space_before = Pt(0)
    audience_intro.paragraph_format.space_after = Pt(6)
    run_aud_intro = audience_intro.add_run(
        "本课程适合以下类型的学员，学员层次可混合，效果更佳："
    )
    run_aud_intro.font.size = Pt(11)

    audiences = [
        "企业中高层管理者，需要系统化推动组织创新",
        "战略规划、研发、培训等部门负责人",
        "产品经理、项目经理，需要创新方法论支撑",
        "创业者和创新业务负责人",
        '有意识到"创新靠零散活动不成体系"的学员',
    ]

    for aud in audiences:
        aud_para = doc.add_paragraph(style='List Bullet')
        aud_para.paragraph_format.space_before = Pt(2)
        aud_para.paragraph_format.space_after = Pt(2)
        run = aud_para.add_run(aud)
        run.font.size = Pt(10.5)

    # ========== 课程结构总览 ==========
    doc.add_page_break()
    structure_title = doc.add_paragraph()
    structure_title.paragraph_format.space_before = Pt(20)
    structure_title.paragraph_format.space_after = Pt(16)
    run_struct = structure_title.add_run("课程结构总览")
    run_struct.font.size = Pt(18)
    run_struct.font.bold = True
    run_struct.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    struct_table = doc.add_table(rows=9, cols=3)
    struct_table.style = 'Table Grid'

    hdr = struct_table.rows[0].cells
    hdr[0].text = "部分"
    hdr[1].text = "核心主题"
    hdr[2].text = "核心工具"
    for cell in hdr:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        set_cell_background(cell, "D6E3F8")

    structures = [
        ("第零部分", "创新全链路导论：建立元框架", "场景诊断三问、创新成熟度自测表、全链路地图"),
        ("第一部分", "创新战略与组织氛围", "四象限创新类型、3M创新文化、70-20-10模型、障碍诊断"),
        ("第二部分", "创意产生与筛选机制", "设计思维五阶段、德鲁克七来源、筛选漏斗四层、头脑风暴反模式"),
        ("第三部分", "创新项目立项与Stage-Gate", "Stage-Gate五阶段、Gate决策标准、商业论证四支柱"),
        ("第四部分", "创新团队组建与管理", "四角色模型、Ambidextrous组织、授权光谱、创造性张力管理"),
        ("第五部分", "创新成果转化与商业化", "四种商业化路径、IP保护策略、商业模式画布、GTM策略"),
        ("第六部分", "创新评估与持续改进", "创新ROI四层框架、创新BCG矩阵、失败复盘四步法、反馈闭环"),
        ("第七部分", "综合演练与行动规划", "设计思维Capstone、创新行动方案、路演技巧、知识地图"),
    ]

    for i, (part, theme, tools) in enumerate(structures, 1):
        row = struct_table.rows[i].cells
        row[0].text = part
        row[1].text = theme
        row[2].text = tools
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # ========== 课时建议 ==========
    doc.add_page_break()
    time_title = doc.add_paragraph()
    time_title.paragraph_format.space_before = Pt(20)
    time_title.paragraph_format.space_after = Pt(16)
    run_time = time_title.add_run("课时建议")
    run_time.font.size = Pt(18)
    run_time.font.bold = True
    run_time.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    time_table = doc.add_table(rows=5, cols=2)
    time_table.style = 'Table Grid'

    hdr_time = time_table.rows[0].cells
    hdr_time[0].text = "版本"
    hdr_time[1].text = "说明"
    for cell in hdr_time:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        set_cell_background(cell, "D6E3F8")

    time_items = [
        ("标准版（全天7小时）", "完整呈现全部7个部分，含所有练习和讨论。适合整天研修。"),
        ("精简版（半天3.5小时）", "保留第零部分（导论）+ 选2-3个最相关部分。每部分约45分钟。"),
        ("系列课（7次×90分钟）", "每次聚焦一个部分，适合定期培训嵌入形式。每次有明确的课后行动作业。"),
        ("主题工作坊（单维深化）", "针对特定岗位或问题，深度展开特定模块。如：战略创新工作坊、团队创新工作坊。"),
    ]

    for i, (version, desc) in enumerate(time_items, 1):
        row = time_table.rows[i].cells
        row[0].text = version
        row[1].text = desc
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # ========== 课前准备 ==========
    doc.add_page_break()
    prep_title = doc.add_paragraph()
    prep_title.paragraph_format.space_before = Pt(20)
    prep_title.paragraph_format.space_after = Pt(16)
    run_prep = prep_title.add_run("课前准备")
    run_prep.font.size = Pt(18)
    run_prep.font.bold = True
    run_prep.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    prep_heading = doc.add_paragraph()
    prep_heading.paragraph_format.space_before = Pt(12)
    prep_heading.paragraph_format.space_after = Pt(8)
    run_prep_h = prep_heading.add_run("讲师准备清单")
    run_prep_h.font.bold = True
    run_prep_h.font.size = Pt(12)
    run_prep_h.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    prep_notice = doc.add_paragraph()
    prep_notice.paragraph_format.space_before = Pt(0)
    prep_notice.paragraph_format.space_after = Pt(8)
    run_prep_n = prep_notice.add_run("【建议】请在课前48小时内逐项确认以下事项")
    run_prep_n.font.size = Pt(10)
    run_prep_n.font.bold = True
    run_prep_n.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    content_heading = doc.add_paragraph()
    content_heading.paragraph_format.space_before = Pt(8)
    content_heading.paragraph_format.space_after = Pt(4)
    run_cont = content_heading.add_run("内容熟悉度")
    run_cont.font.bold = True
    run_cont.font.size = Pt(11)

    content_items = [
        "熟读本次讲授的部分对应的学员手册内容",
        "准备自己在每个模块上的1-2个真实亲身经历案例（替换或补充课程案例）",
        "预演各练习的示范答案，确保能在现场展示第一级和第二级难度示范",
        "熟悉附录中的参考答案，准备好应对学员的不同答案",
    ]

    for item in content_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    audience_info_heading = doc.add_paragraph()
    audience_info_heading.paragraph_format.space_before = Pt(8)
    audience_info_heading.paragraph_format.space_after = Pt(4)
    run_aud_info = audience_info_heading.add_run("学员信息")
    run_aud_info.font.bold = True
    run_aud_info.font.size = Pt(11)

    audience_info_items = [
        "了解学员背景：行业、岗位层级、平均工作年限",
        "了解学员现有创新管理水平（可通过报名问卷收集）",
        "识别可能对课程提出质疑或认知抵触的学员类型",
        "确认学员规模：20人以内（最佳）、20-30人（可行）、30人以上（需调整互动设计）",
    ]

    for item in audience_info_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    materials_heading = doc.add_paragraph()
    materials_heading.paragraph_format.space_before = Pt(8)
    materials_heading.paragraph_format.space_after = Pt(4)
    run_mat = materials_heading.add_run("材料准备")
    run_mat.font.bold = True
    run_mat.font.size = Pt(11)

    materials_items = [
        "每人一份学员手册（已打印，双面彩印）",
        "每人一套配套表单（空表版，已装订）",
        "补充用表：A4纸备用，用于额外练习书写",
        "白板/翻页纸板，用于现场板书和练习展示",
        "计时器（建议使用投影计时，或手机分屏显示）",
        "贴纸或卡片，用于练习成果展示（推荐A5卡片）",
    ]

    for item in materials_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    room_heading = doc.add_paragraph()
    room_heading.paragraph_format.space_before = Pt(8)
    room_heading.paragraph_format.space_after = Pt(4)
    run_room = room_heading.add_run("室内布置")
    run_room.font.bold = True
    run_room.font.size = Pt(11)

    room_items = [
        "圆桌或岛型分组座位（4-6人一组）——优先于剧院式座位",
        "每桌配备彩色马克笔 + 白板纸/便签纸",
        "讲台可见但不孤立，讲师能方便地走到学员间",
        "投影屏幕可从室内所有位置清晰阅读",
    ]

    for item in room_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    # ========== 如何使用学员手册 ==========
    doc.add_page_break()
    howto_title = doc.add_paragraph()
    howto_title.paragraph_format.space_before = Pt(20)
    howto_title.paragraph_format.space_after = Pt(16)
    run_howto = howto_title.add_run("如何使用学员手册")
    run_howto.font.size = Pt(18)
    run_howto.font.bold = True
    run_howto.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    howto_intro = doc.add_paragraph()
    howto_intro.paragraph_format.space_before = Pt(0)
    howto_intro.paragraph_format.space_after = Pt(12)
    run_howto_intro = howto_intro.add_run(
        '学员手册的设计理念是：讲师打开手册带着讲，学员对着手册一起跟，'
        '课后翻开手册能复习。请避免将其当成PPT逐字阅读——'
        '关键是引导学员在练习时"停下来动手"。'
    )
    run_howto_intro.font.size = Pt(11)

    howto_table = doc.add_table(rows=5, cols=2)
    howto_table.style = 'Table Grid'

    hdr_howto = howto_table.rows[0].cells
    hdr_howto[0].text = "使用方式"
    hdr_howto[1].text = "说明"
    for cell in hdr_howto:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        set_cell_background(cell, "D6E3F8")

    howto_items = [
        ("共读模式", "讲师朗读或引导学员交替阅读场景/案例段落，增加代入感"),
        ("静读模式", "复杂工具讲解前给学员2-3分钟静读，让每人先形成独立理解再讨论"),
        ("练习模式", '到练习环节时，明确说"现在停下来，大家翻到第X页"'),
        ("参考模式", "课后学员可以作为工具手册反复翻阅，标记对自己有用的工具"),
    ]

    for i, (mode, desc) in enumerate(howto_items, 1):
        row = howto_table.rows[i].cells
        row[0].text = mode
        row[1].text = desc
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # ========== 第零部分 ==========
    doc.add_page_break()

    add_module_section(
        doc,
        module_num="零",
        module_title="创新全链路导论——建立元框架",
        duration="45—60分钟",
        objectives=[
            "学员能说出创新管理的完整链路（七个阶段）",
            '学员能用"创新成熟度自测表"对所在组织进行初步诊断',
            "学员完成个人场景卡初稿，明确自己在创新管理体系中的薄弱环节",
            "学员理解七部分课程的整体逻辑和相互关系",
        ],
        opening=(
            '开场不要介绍"今天学什么"。先带学员对比两个创新负责人的案例—— '
            '负责人A和B的做法对比。让学员先在头脑里形成疑问——'
            '"为什么同样的资源投入，结果天差地别？"然后再回答这个问题。'
        ),
        key_points=[
            '创新全链路地图（第一部分）：带着学员逐行阅读全链路地图，重点强调最后一列"你将带走什么"。',
            '在每个部分前，让学员先猜测"这个部分解决什么问题"，然后再揭示。这个小互动会制造惊喜感。',
            '场景诊断三问（第二部分）：先完整呈现三问框架，然后强调"这三个问题是有顺序的"—— '
            '一问完毕才进第二问，第二问完毕才进第三问。',
            '创新成熟度自测（第三部分）：给学员10分钟独立完成，然后邀请2-3位分享——不做点评，只做正向强化。',
        ],
        exercises=[
            "练习一：创新现状诊断（第一级）——给学员5分钟独立填写创新成熟度自测表，然后两两对答案",
            "练习二：场景卡填写（第二级）——给学员10分钟静默独立完成，完成后邀请3-4位分享",
            "练习三：找出最薄弱环节（第三级）——让学员互相讨论各自最需要加强的部分",
        ],
        transitions=(
            '这个部分是整门课的元框架。结束时用一句话总结：'
            '【创新全链路地图】是贯穿全课的唯一元工具。后面每个部分，都是对这个工具的一次具体化和深化。'
            '过渡到第一部分时：可以问学员"哪个部分你觉得最陌生、最需要加强"——自然引入下一模块。'
        ),
        notes=[
            "告诉学员保留好场景卡。课程结束时（第七部分）会回来对比——这是他们自我成长的参照点。",
            '有些学员会在导论阶段就开始问"我应该从哪个部分开始做"。告诉他们：这个问题在第七部分（综合演练）会有系统答案，现在先把每个部分的工具学好。',
        ]
    )

    # ========== 第一部分 ==========
    add_module_section(
        doc,
        module_num="一",
        module_title='创新战略与组织氛围——从"不得不创新"到"善于创新"',
        duration="55—70分钟",
        objectives=[
            "学员能准确区分四种创新类型（渐进式/激进式/探索性/利用性）",
            "学员能用创新组合视角审视自己组织的创新布局",
            "学员掌握3M案例的核心理念（允许失败、15%规则、跨部门协作）",
            "学员能识别自己组织面临的创新障碍类型并给出突破方向",
        ],
        opening=(
            '带学员读施乐PARC案例（或国内某制造企业的创新困境案例）。先停在"问题出在哪里"那里，'
            '让学员先猜。等学员确认"创新方向/组织文化/资源分配"这几个可能后，再揭示答案——三个问题都有。'
        ),
        key_points=[
            '四种创新类型（第一部分）：用四象限图展示，重点强调"企业最常见的错误是创新类型错配"。',
            "可以举诺基亚的例子——渐进式创新做到极致，但被颠覆式创新击败。",
            '3M案例（第二部分）：重点不是"15%规则"，而是"允许失败"的机制设计。强调"失败复盘和绩效考核必须分开"。',
            '70-20-10模型（第三部分）：强调这不是预算分配，而是注意力分配。可以让学员算算自己花在核心业务上的时间占比。',
            "创新障碍诊断（第四部分）：先让学员用自己的组织做诊断，再给解法。识别障碍类型比给出解法更重要。",
        ],
        exercises=[
            '练习一：创新类型分类（第一级）——给企业创新动作分类，重点关注"并购"属于哪种类型（探索性）',
            '练习二：70-20-10分配演练（第二级）——让学员实际分配自己的精力，重点避免"把所有重要的事都塞进70%"',
            "练习三：障碍诊断（第三级）——回想自己组织，最严重的障碍是哪种？给出具体例子和突破方案",
        ],
        transitions=(
            '这个模块结束后，学员通常会有一个感受："原来我们公司不是创新太少，而是创新类型可能错配了"。'
            '过渡到第二部分时：可以说"知道了创新类型，下一个问题是——好创意从哪里来？"'
        ),
        notes=[
            '学员有时会问"华为算探索性还是利用性"。答：华为同时在做多种类型，需要分开看。',
            '70-20-10模型有学员会质疑"我们的资源不允许10%做探索"。解释：10%不只是钱，是注意力。可以从5%开始。',
        ]
    )

    # ========== 第二部分 ==========
    add_module_section(
        doc,
        module_num="二",
        module_title="创意产生与筛选机制——从发散到收敛的完整体系",
        duration="70—85分钟",
        objectives=[
            "学员能完整走完设计思维五阶段（Empathize→Define→Ideate→Prototype→Test）",
            "学员能用德鲁克七来源系统识别自己业务中的创新机会",
            "学员掌握四层筛选漏斗，能对给定的多个创意做优先级排序",
            "学员能识别并规避六种常见的头脑风暴反模式",
        ],
        opening=(
            '带学员做一个现场演示——让学员体验一次"糟糕的头脑风暴"。先做一次有隐形评判、有锚定效应的 brainstorm，'
            '然后问："刚才有多少想法是真正疯狂的？"再揭示头脑风暴的六大反模式。'
        ),
        key_points=[
            '设计思维五阶段（第一部分）：重点是"共情"和"定义"——这两个阶段最容易被跳过，但ROI最高。',
            '建议做一个15分钟的现场演示：让学员对"工作中最烦的那个流程"走一遍共情和定义。',
            '德鲁克七来源（第二部分）：重点不是"记住七个来源"，而是"建立扫描习惯"。可以带学员扫描一下自己行业的意外成功/失败。',
            '筛选漏斗（第三部分）：强调"筛选不是选最好的，是排除最差的"。筛选的目标是最小化"把好想法枪毙了"的风险。',
            "头脑风暴反模式（第四部分）：可以做角色扮演——让一组人表演某种反模式，其他组识别。",
        ],
        exercises=[
            '练习一：设计思维工作坊（第二级）——用15分钟对"工作中最烦的那个流程"走完共情+定义阶段',
            '练习二：德鲁克七来源扫描（第二级）——每人带一个"意外数据"，用七来源框架分类',
            '练习三：筛选漏斗实战（第三级）——用评分矩阵给20个创意排序，重点练习"排除法"思维',
        ],
        transitions=(
            '这个模块结束后，学员通常会有一个感受："原来创意产生是有方法的，不是靠灵感"。'
            '过渡到第三部分时：可以说"有了好创意，下一步是——怎么让它变成一个正式项目？"'
        ),
        notes=[
            '学员容易把"头脑风暴"当成"开会讨论"。要反复强调：头脑风暴的核心是"先发散，不判断"。',
            '德鲁克七来源有学员觉得"太理论"。可以让他们立刻用自己的业务举例，这样更直观。',
        ]
    )

    # ========== 第三部分 ==========
    add_module_section(
        doc,
        module_num="三",
        module_title="创新项目立项与Stage-Gate——让创新从赌博变成工程",
        duration="65—80分钟",
        objectives=[
            "学员能完整描述Stage-Gate模型的五个阶段和四个Gate",
            "学员能区分Go/Kill/Hold三种决策结果及适用场景",
            "学员掌握商业论证四支柱（客户价值、市场潜力、商业模式、竞争优势）",
            "学员能用TAM/SAM/SOM框架评估市场规模",
        ],
        opening=(
            '带学员读施乐"达芬奇"项目的案例（5亿美元打水漂）。先停在"问题出在哪里"——'
            '让学员意识到：不是技术不行，是立项流程有问题。从这里引入Stage-Gate的价值。'
        ),
        key_points=[
            'Stage-Gate五阶段（第一部分）：重点是"阶段不是固定的"——根据项目类型和行业可以调整。'
            '但Gate的逻辑是固定的：每阶段结束必须有明确的Go/Kill/Hold决策。',
            'Gate决策标准（第二部分）：强调"标准要提前定，不能在Gate的时候临时发明"。'
            '可以展示一个真实的Gate评审打分表，让学员知道"决策是有据可依的"。',
            "商业论证四支柱（第三部分）：客户价值最难验证，也最重要。建议用JTBD框架深入讲。",
            '资源配置原则（第四部分）：强调"资源要富足不要勉强"——勉强凑齐资源的项目十有八九会失败。',
        ],
        exercises=[
            "练习一：Stage-Gate流程设计（第二级）——为一个创新项目设计Stage-Gate流程，包括每个Gate的决策标准",
            "练习二：商业论证四支柱（第二级）——用四支柱框架评估一个假想项目的商业论证完整性",
            '练习三：市场规模估算（第三级）——练习TAM/SAM/SOM三层估算，重点是"不要把TAM当成SOM"',
        ],
        transitions=(
            '这个模块结束后，学员通常会问："我们公司根本没有Gate机制，怎么办？"——'
            '告诉他们：这就是第四部分要解决的——谁来当Gatekeeper，怎么组建创新团队。'
        ),
        notes=[
            'Stage-Gate容易变成"过度设计"。提醒学员：早期项目用简化版（两阶段一Gate）也可以。',
            '有学员会担心"Gate机制会不会扼杀创新"。答：扼杀创新的是"没有结构的决策"，不是Gate本身。',
        ]
    )

    # ========== 第四部分 ==========
    add_module_section(
        doc,
        module_num="四",
        module_title="创新团队组建与管理——让对的人做对的事",
        duration="55—70分钟",
        objectives=[
            "学员能识别创新团队需要的四种关键角色（Champion/Gatekeeper/Connector/Executor）",
            "学员能用授权光谱评估自己给创新团队的真实授权层级",
            "学员掌握Amazon两个披萨团队的核心设计原则",
            '学员能运用失败判断框架区分"好的失败"和"坏的失败"',
        ],
        opening=(
            '问学员一个问题："过去一年，你们公司有没有一个创新项目，技术很先进但最终失败了？'
            '项目失败后，那个团队怎么样了？"——让学员意识到：创新失败后，团队的下场往往决定了下一个创新会不会发生。'
        ),
        key_points=[
            "四角色模型（第一部分）：四种角色不必四个人分别承担——在小团队中一人可能兼任两到三个角色。"
            '关键是这四种功能必须被覆盖。可以让学员画一张自己团队的"角色-功能矩阵"。',
            'Ambidextrous组织（第二部分）：重点不是"结构"，而是"如何在同一个人身上平衡探索和利用"。'
            "Amazon两个披萨团队是极致实践，但不是所有公司都适用。",
            '授权光谱（第三部分）：创新团队通常需要落在第4到第6层级，但大多数组织默认把创新团队放在第1到第2层级。'
            '这个差距是创新失败的主要原因之一。',
            '失败管理（第四部分）：强调"失败的成本计算公式"——容忍失败的价值=f(学习价值,可逆性)/f(失败成本,不可逆性)。',
        ],
        exercises=[
            "练习一：角色-功能矩阵（第二级）——绘制自己团队/项目的角色-功能矩阵，找出缺失的角色",
            "练习二：授权层级评估（第二级）——分别从财务、技术、人员、战略四个维度评估自己给团队的授权层级",
            '练习三：失败复盘（第三级）——用一个自己经历过的创新失败案例，走一遍"四步失败复盘法"',
        ],
        transitions=(
            '这个模块结束后，学员通常会有一个感受："原来我们公司不是没有创新人才，是没有让创新人才发挥作用的机制"。'
            '过渡到第五部分时：可以说"好创意有了，团队也组建好了——接下来是怎么让它变成钱？"'
        ),
        notes=[
            'Amazon两个披萨团队有学员会问"我们公司没有那种技术架构怎么办"。答：两个披萨是结果，不是原因。先看能不能给团队端到端所有权。',
            '有学员担心"给团队太多授权，会不会失控"。解释：授权是给边界，不是给无限空间。关键是设定清晰的目标和边界。',
        ]
    )

    # ========== 第五部分 ==========
    add_module_section(
        doc,
        module_num="五",
        module_title="创新成果转化与商业化——从实验室到市场",
        duration="65—80分钟",
        objectives=[
            "学员能区分四种商业化路径（内部孵化/技术许可/创业融资/生态嵌入）并选择适合的路径",
            "学员掌握专利/商标/商业秘密/著作权四种IP保护策略的适用场景",
            "学员能用商业模式画布九大构建块分析创新项目的商业可行性",
            "学员能制定创新产品的市场进入策略（GTM）",
        ],
        opening=(
            '带学员看一个反直觉的数据：70%的技术创新项目无法成功商业化。'
            '问学员："你们觉得这些失败的技术创新，问题出在哪里？"——'
            '让学员意识到：技术成功≠商业成功。从这里引入第五部分的核心命题。'
        ),
        key_points=[
            '商业化四种路径（第一部分）：重点是"路径选择取决于创新类型和自身能力"。'
            '可以带学员分析：大疆是哪条路径？微信红包是哪条路径？',
            '知识产权保护（第二部分）：重点不是"申请什么专利"，而是"如何形成组合保护"。'
            '可以讲一个真实的IP诉讼案例，说明"防守和进攻同样重要"。',
            "商业模式画布（第三部分）：建议让学员当场用画布分析自己的一个创新项目。"
            '九大构建块中，"客户关系"和"渠道通道"最容易被中国学员忽视。',
            'GTM策略（第四部分）：重点是"MVP客户选择三标准"——痛点最强烈、决策链条最短、口碑传播力最强。',
        ],
        exercises=[
            "练习一：商业化路径选择（第二级）——为一个创新项目选择最合适的商业化路径，给出理由",
            "练习二：IP保护策略制定（第二级）——为一个技术型创新项目制定IP保护策略组合",
            "练习三：商业模式画布实战（第三级）——用画布分析自己负责的一个创新业务",
            "练习四：GTM策略设计（第三级）——为一个创新产品设计前12个月的市场进入计划",
        ],
        transitions=(
            '这个模块结束后，学员通常会有一个感受："原来技术领先不等于商业成功，中间还有很远距离"。'
            '过渡到第六部分时：可以说"商业化成功了，接下来——怎么知道这次创新是不是真的成功了？"'
        ),
        notes=[
            '商业化路径选择没有标准答案，关键是"适合"。让学员充分讨论不同路径的利弊。',
            'IP保护有学员会陷入"过度保护"。提醒：保护是有成本的，要抓重点。',
        ]
    )

    # ========== 第六部分 ==========
    add_module_section(
        doc,
        module_num="六",
        module_title="创新评估与持续改进——让创新从赌博变成系统",
        duration="60—75分钟",
        objectives=[
            "学员能用创新ROI四层框架评估自己组织的创新投入",
            "学员能用创新BCG矩阵分析自己组织的创新组合健康度",
            "学员掌握创新失败复盘四步法，能组织一次有效的失败复盘",
            "学员能设计一个最小可行的创新反馈闭环",
        ],
        opening=(
            '带学员做一个思想实验："如果你是某家电巨头创新负责人，50亿ALL IN智能家居，两年后项目失败——'
            '你怎么向董事会解释？"——让学员意识到：没有评估标准，创新投入就是赌博。'
        ),
        key_points=[
            '创新ROI四层框架（第一部分）：重点是"战略价值和能力价值最难量化，但最重要"。'
            '可以用自己公司的实际数据做练习，让学员算出自己的创新ROI。',
            '创新BCG矩阵（第二部分）：四象限的资源配置建议是参考值，关键是"四类创新都要有"。'
            '可以让学员对照自己的组织，看哪个象限是空的。',
            '失败复盘四步法（第三部分）："情绪脱敏"最难，但最重要。复盘会议要快（项目结束2周内），'
            '但参与人不包括项目负责人（避免权力影响畅所欲言）。',
            '反馈闭环（第四部分）：强调"闭环不是一次性建成的，是逐步演进的"。'
            '先从一个机制开始（比如季度创新评估例会），再逐步完善。',
        ],
        exercises=[
            "练习一：创新ROI自测（第二级）——用ROI计算器评估自己组织的创新投入效率",
            "练习二：创新组合诊断（第二级）——用BCG矩阵评估自己组织的创新组合健康度",
            "练习三：失败复盘模板（第三级）——用一个自己经历过的创新失败案例，走一遍复盘四步法",
        ],
        transitions=(
            '这个模块结束后，学员通常会有一个感受："原来创新是可以被衡量和管理的"。'
            '过渡到第七部分时：可以说"前六部分我们学了创新的各个模块——第七部分是把它们串起来，做一次完整的演练。"'
        ),
        notes=[
            '创新ROI的计算经常遇到"数据不全"的问题。告诉学员：先用现有数据填，有数据的部分认真填，没数据的做合理假设。',
            '失败复盘有学员担心"会不会变成追责会"。强调：复盘要快，2周内；参与人不包括负责人；目的是学习不是追责。',
        ]
    )

    # ========== 第七部分 ==========
    add_module_section(
        doc,
        module_num="七",
        module_title="综合演练与行动规划——从框架到落地",
        duration="90—120分钟（结业项目）",
        objectives=[
            "学员能完整运用设计思维六步法解决一个真实的创新挑战",
            "学员能制定一份个人创新行动方案（含目标、路径、里程碑）",
            "学员能进行一次5分钟创新提案路演并接受质询",
            "学员完成课程知识地图，明确自己在创新管理体系中的收获和下一步",
        ],
        opening=(
            '这个部分是结业项目，不需要新的知识输入。开场直接说："今天我们做一次完整的创新演练，'
            '把前六天学的东西全部用上。"——然后直接进入Capstone Project。'
            '如果时间紧张，可以用"创新行动方案开发"替代完整的Capstone Project，'
            '让学员直接制定回到工作岗位后的创新行动计划。'
        ),
        key_points=[
            "Capstone Project（第一部分）：选择一个真实的创新挑战，用设计思维六步法完整走一遍。"
            '关键是"真实"——不是假设案例，而是学员自己真正在思考的问题。',
            '创新行动方案（第二部分）：强调"可执行"——不是学习心得，是行动计划。'
            '包含：目标、路径、里程碑、资源需求、成功指标。',
            '路演技巧（第三部分）：5分钟版本的路演结构——'
            '问题导入（1分钟）+ 解决方案（2分钟）+ 商业模式（1分钟）+ 呼吁行动（1分钟）。',
        ],
        exercises=[
            "练习一：设计思维Capstone（第三级）——完整走完六步法，输出可演示的原型",
            "练习二：创新行动方案（第二级）——制定一份回到工作岗位后可以立即执行的创新行动计划",
            "练习三：路演展示（第三级）——每组5分钟展示，接受其他组员和讲师的质询",
        ],
        transitions=[
            '课程结束时的总结：创新管理体系不是"学完就结束"，而是"学完才刚开始"。'
            '希望学员回到工作岗位后，能用这套框架分析真实问题、设计创新方案、推动组织变革。',
        ],
        notes=[
            '如果时间不够，优先保证"创新行动方案"环节——这是学员带回去最直接有用的产出。',
            "路演环节可以邀请学员上级或HR旁听，增加真实感。",
            "结业时告诉学员：课程群不会解散，大家可以继续交流创新实践。",
        ]
    )

    # ========== FAQ ==========
    doc.add_page_break()
    faq_title = doc.add_paragraph()
    faq_title.paragraph_format.space_before = Pt(20)
    faq_title.paragraph_format.space_after = Pt(16)
    run_faq = faq_title.add_run("常见问题FAQ")
    run_faq.font.size = Pt(18)
    run_faq.font.bold = True
    run_faq.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    faqs = [
        ('Q1：学员普遍觉得某个模块太理论，和自己公司情况不相关怎么办？',
         'A：首先承认这个模块确实更偏方法论。然后问："如果抛开你们公司的实际情况，单纯看这个工具——'
         '它解决的是什么类型的问题？"让学员跳出来看工具的价值，而不是急着找"直接能用的"。'),

        ('Q2：学员在练习环节讨论过于发散，收不回来怎么办？',
         'A：在练习开始前明确宣布时间限制（建议用投影计时）。'
         '在时间到达前5分钟给预警，到时间立刻收。讨论质量不是时间堆出来的，是目标驱动的。'),

        ('Q3：有学员质疑"这些工具在我们行业不适用"怎么办？',
         'A：让他举一个具体的例子，说明这个工具为什么不适用。'
         '大多数情况下，学员会发现不是工具不适用，而是自己还没有找到正确的应用方式。'
         '如果确实有不适用的地方，这也是有价值的学习——工具都有边界。'),

        ('Q4：学员提出的问题超出课程范围，怎么处理？',
         'A：简单问题当场回答，复杂问题可以记下来，在茶歇或课后单独讨论。'
         '不要让一个问题拖住全班进度，但也不要完全无视——这体现了对学员的尊重。'),

        ('Q5：两天的课程，时间总是不够用，哪些内容可以精简？',
         'A：建议优先保证：开场案例（震撼感必须保证）、练习环节（动手比听讲重要）、结业项目（收尾必须有完整性）。'
         '可以精简的部分：部分知识点的展开讲解、某些案例的深度分析。记住：内容可超，节奏不拖。'),

        ('Q6：学员规模超过30人，互动设计怎么调整？',
         'A：超过30人时，分组数量增加，每组人数可以是6-8人。'
         '减少需要全员讨论的环节，增加小组代表发言的环节。'
         '练习环节可以先小组讨论再派代表分享，而不是每人轮流发言。'),

        ('Q7：如何应对"老板派来听课但本人不太积极"的学员？',
         'A：给这类学员一个特殊任务——比如担任某组的记录员或计时员。'
         '赋予责任可以激活参与。同时，在小组讨论时特意问一下他的看法，让他感到被重视。'),

        ('Q8：课程中间有学员中途离开/加入，怎么处理？',
         'A：中途离开的学员，确保有人在他离开前把他参与的小组讨论结果记录下来。'
         '中途加入的学员，让组长在茶歇时做简要介绍，并安排他承担一个具体的任务（而不是从头跟）。'),
    ]

    for q, a in faqs:
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(10)
        q_para.paragraph_format.space_after = Pt(3)
        run_q = q_para.add_run(q)
        run_q.font.bold = True
        run_q.font.size = Pt(10.5)
        run_q.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        a_para = doc.add_paragraph()
        a_para.paragraph_format.space_before = Pt(0)
        a_para.paragraph_format.space_after = Pt(8)
        a_para.paragraph_format.left_indent = Inches(0.3)
        run_a = a_para.add_run(a)
        run_a.font.size = Pt(10.5)

    # ========== 附录：课程产出汇总 ==========
    doc.add_page_break()
    appendix_title = doc.add_paragraph()
    appendix_title.paragraph_format.space_before = Pt(20)
    appendix_title.paragraph_format.space_after = Pt(16)
    run_app = appendix_title.add_run("附录：学员带走的核心产出")
    run_app.font.size = Pt(18)
    run_app.font.bold = True
    run_app.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    output_intro = doc.add_paragraph()
    output_intro.paragraph_format.space_before = Pt(0)
    output_intro.paragraph_format.space_after = Pt(12)
    run_out_intro = output_intro.add_run(
        "两天课程结束后，每位学员应该带走以下核心产出："
    )
    run_out_intro.font.size = Pt(11)

    outputs = [
        ("场景卡（个人版）", "选择一个自己实际在推进的创新项目/方向，作为贯穿全程的练习素材"),
        ("创新机会地图", "第一部分产出：用四维框架（市场/技术/组织/资源）分析自己创新方向的机会"),
        ("创意筛选漏斗", "第二部分产出：自己创新流程的创意筛选标准和方法"),
        ("Stage-Gate计划", "第三部分产出：为一个创新项目设计的Stage-Gate流程和Gate决策标准"),
        ("角色-功能矩阵", "第四部分产出：自己团队的四种角色覆盖情况分析"),
        ("商业化路径图", "第五部分产出：自己创新业务的商业化路径和GTM策略"),
        ("创新评估仪表盘", "第六部分产出：自己组织的创新ROI四层框架和指标体系"),
        ("创新行动方案", "第七部分产出：回到工作岗位后可以立即执行的创新行动计划"),
    ]

    for i, (name, desc) in enumerate(outputs, 1):
        out_para = doc.add_paragraph()
        out_para.paragraph_format.space_before = Pt(8)
        out_para.paragraph_format.space_after = Pt(2)
        run_num = out_para.add_run(f"{i}. {name}")
        run_num.font.bold = True
        run_num.font.size = Pt(11)
        run_num.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.space_before = Pt(0)
        desc_para.paragraph_format.space_after = Pt(4)
        desc_para.paragraph_format.left_indent = Inches(0.4)
        run_desc = desc_para.add_run(desc)
        run_desc.font.size = Pt(10.5)

    # 保存文档
    output_path = "D:/新课开发/管理学/39-创新管理体系/讲师手册/讲师手册_创新管理体系.docx"
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    return output_path

if __name__ == "__main__":
    create_instructor_handbook()

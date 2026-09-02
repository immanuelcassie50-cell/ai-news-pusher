# -*- coding: utf-8 -*-
import sys
import os
os.chdir(r'C:\Users\Administrator\Desktop\标书')

from docx import Document

files = [f for f in os.listdir('.') if f.endswith('.docx')]
for fname in files:
    print(f"\n{'='*60}")
    print(f"FILE: {fname}")
    print('='*60)
    doc = Document(fname)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] ({para.style.name}): {para.text[:200]}")

    # Also print tables
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {t_idx} ---")
        for row in table.rows:
            cells = [cell.text.strip()[:50] for cell in row.cells]
            print(" | ".join(cells))
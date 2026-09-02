#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
学员手册 .docx 生成脚本
使用 python-docx 生成
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = r'D:\2026年课程\竞越\创新领导力：打造创新型团队\完整课程表\03-学员手册\学员手册_创新领导力.docx'

doc = Document()

# ============== 全局样式 ==============
# 设置默认字体（中文）
def set_zh_font(run, font_name='Microsoft YaHei', size=10.5):
    run.font.name = font_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

# 页面设置：A4
for section in doc.sections:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============== 页眉页脚 ==============
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr = header_para.add_run('创新领导力：打造创新型团队  ·  学员手册 v1.0')
set_zh_font(hr, size=9, font_name='Microsoft YaHei')
hr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_para.add_run('第 ')
set_zh_font(fr, size=9)
# 页码字段
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run = footer_para.add_run()
set_zh_font(run, size=9)
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

fr2 = footer_para.add_run(' 页 / 共 ')
set_zh_font(fr2, size=9)
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'begin')
instrText2 = OxmlElement('w:instrText')
instrText2.set(qn('xml:space'), 'preserve')
instrText2.text = 'NUMPAGES'
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')
run2 = footer_para.add_run()
set_zh_font(run2, size=9)
run2._r.append(fldChar3)
run2._r.append(instrText2)
run2._r.append(fldChar4)

fr3 = footer_para.add_run(' 页')
set_zh_font(fr3, size=9)


# ============== 工具函数 ==============
def add_title(text, level=1):
    """添加标题"""
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    sizes = {0: 26, 1: 18, 2: 15, 3: 13, 4: 12}
    set_zh_font(run, size=sizes.get(level, 12))
    run.font.bold = True
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    return h

def add_para(text, bold=False, italic=False, size=10.5, indent_first=True, align=None):
    """添加段落"""
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_zh_font(run, size=size)
    run.bold = bold
    run.italic = italic
    return p

def add_quote(text, source=None):
    """添加引用块（用表格+背景色）"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    # 设置背景色
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E7F0FA')
    tcPr.append(shd)
    # 设置左边框
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:color'), '2E74B5')
    tcBorders.append(left)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run('"' + text + '"')
    set_zh_font(run, size=11)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    if source:
        p2 = cell.add_paragraph()
        r2 = p2.add_run('—— ' + source)
        set_zh_font(r2, size=9.5, font_name='Microsoft YaHei')
        r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        r2.italic = True
    return table

def add_blank_line(n=1):
    """空行（练习留白）"""
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run('　' * 40)
        set_zh_font(run, size=10.5)

def add_tip(text, tip_type='💡'):
    """添加提示框"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF4CE')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_zh_font(run, size=10.5)
    run.font.color.rgb = RGBColor(0x7F, 0x60, 0x00)
    return table

def add_warning(text):
    """添加警告框"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FCE4E4')
    tcPr.append(shd)
    p = cell.paragraphs[0]
    run = p.add_run('⚠  ' + text)
    set_zh_font(run, size=10.5)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return table

def add_exercise_box(title, content=None):
    """添加练习框"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F8F0')
    tcPr.append(shd)
    # 边框
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:color'), '70AD47')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    p = cell.paragraphs[0]
    run = p.add_run('✋  练习：' + title)
    set_zh_font(run, size=11, font_name='Microsoft YaHei')
    run.bold = True
    run.font.color.rgb = RGBColor(0x4A, 0x7A, 0x2A)
    if content:
        for line in content:
            p2 = cell.add_paragraph()
            r2 = p2.add_run(line)
            set_zh_font(r2, size=10.5)
        # 加留白
        for _ in range(3):
            p3 = cell.add_paragraph()
            p3.add_run('　')
    return table

def add_table(headers, rows, col_widths=None, first_col_bold=False):
    """添加表格"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_zh_font(run, size=10.5, font_name='Microsoft YaHei')
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # 表头底色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
    # 数据行
    for ridx, row in enumerate(rows):
        for cidx, val in enumerate(row):
            cell = table.rows[ridx+1].cells[cidx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_zh_font(run, size=10)
            if first_col_bold and cidx == 0:
                run.bold = True
            # 隔行底色
            if ridx % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F6FC')
                tcPr.append(shd)
    if col_widths:
        for col, w in zip(table.columns, col_widths):
            for cell in col.cells:
                cell.width = w
    return table


# ============== 封面 ==============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('创新领导力')
set_zh_font(run, size=36, font_name='Microsoft YaHei')
run.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('打造创新型团队')
set_zh_font(run, size=28, font_name='Microsoft YaHei')
run.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('学员手册')
set_zh_font(run, size=22, font_name='Microsoft YaHei')
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('v1.0 · 国际版权课标准 · 第二版')
set_zh_font(run, size=12, font_name='Microsoft YaHei')
run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)

# 封面底部署名
for _ in range(8):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('国际版权课标准学员手册')
set_zh_font(run, size=11, font_name='Microsoft YaHei')
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('中高层管理者 · 两天工作坊')
set_zh_font(run, size=11, font_name='Microsoft YaHei')
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ============== 致学员 ==============
add_title('致学员', level=1)
add_para('这份手册是课程全程的"操作手册"，不是会议资料。')
add_para('它的写法有一个特点：每一处练习都对应着正在进行的课堂内容。请你在每一个练习处停下笔，认真填写它们——它们不是课后作业，而是你在这两天里最重要的思考产出，也是你带回团队的直接素材。')
add_para('两天之后，这份手册上有两类东西：')
add_para('1. 框架与案例 —— 可以反复翻阅的工具', indent_first=False)
add_para('2. 你亲手填写的练习与承诺 —— 可以拿出来和团队对照的承诺', indent_first=False)
add_para('后者比前者重要。', bold=True)

add_para('另外有几个使用建议：')
add_para('1. 课上跟着内容走，课后可以直接当教材用。每个章节都是独立的：课上没讲到的部分，课后也可以自己读懂。', indent_first=False)
add_para('2. 练习页留白足够大。请用笔写——手写比打字更容易让思考慢下来。', indent_first=False)
add_para('3. 不要急着翻到下一页。每个练习框都是经过设计的"思考减速带"，跳过它们，框架就只是框架。', indent_first=False)
add_para('4. 两天结束后，把这份手册和你的行动承诺卡一起带回团队。它不是纪念品，是工具。', indent_first=False)

add_para('祝你这两天，遇见不一样的自己。', bold=True, italic=True)
doc.add_page_break()

# ============== 目录（手工写） ==============
add_title('目录', level=1)
add_para('第一部分  课程导览', bold=True)
add_para('  1.1  课程目标：你将带走什么', indent_first=False)
add_para('  1.2  学习地图', indent_first=False)
add_para('  1.3  课前准备清单', indent_first=False)
add_para('  1.4  学员公约', indent_first=False)
add_para('  1.5  学习方法说明', indent_first=False)

add_para('第二部分  第一天上午：诊断你的团队', bold=True)
add_para('  2.1  开场：赵建设和那十几秒的沉默', indent_first=False)
add_para('  2.2  第一部分：创新型团队的八个真相测试', indent_first=False)
add_para('  2.3  第二部分：五个关键影响因素（含 5 个快速诊断）', indent_first=False)
add_para('  2.4  团队创新健康度快照', indent_first=False)
add_para('  2.5  第三部分：创新型领导者 vs. 运营管理者', indent_first=False)
add_para('  2.6  我的领导行为盘点', indent_first=False)

add_para('第三部分  第一天下午：客户洞察与交互涌现', bold=True)
add_para('  3.1  三大要素概览', indent_first=False)
add_para('  3.2  第四部分：客户洞察', indent_first=False)
add_para('  3.3  客户洞察练习', indent_first=False)
add_para('  3.4  第五部分：交互涌现', indent_first=False)
add_para('  3.5  知识流通审计', indent_first=False)
add_para('  3.6  今日收尾 + 反思卡', indent_first=False)
add_para('  3.7  今晚作业：创新挑战卡', indent_first=False)

add_para('第四部分  第二天上午：敏捷迭代', bold=True)
add_para('  4.1  开场：昨晚发生了什么', indent_first=False)
add_para('  4.2  第六部分：敏捷迭代', indent_first=False)
add_para('  4.3  最小可学习实验练习', indent_first=False)
add_para('  4.4  第七部分：管理者在迭代中的五种角色转换', indent_first=False)
add_para('  4.5  第八部分：三要素自诊', indent_first=False)
add_para('  4.6  情景导入：亮界科技', indent_first=False)

add_para('第五部分  第二天下午：情景模拟与行动承诺', bold=True)
add_para('  5.1  第九部分：诊断亮界科技', indent_first=False)
add_para('  5.2  第十部分：如果你是张力', indent_first=False)
add_para('  5.3  第十一部分：这些，也发生在我的团队里吗', indent_first=False)
add_para('  5.4  第十二部分：创新领导力行动承诺', indent_first=False)
add_para('  5.5  两天知识框架回顾', indent_first=False)

add_para('第六部分  工具表单', bold=True)
add_para('第七部分  行动承诺与跟进', bold=True)
add_para('第八部分  反思日志', bold=True)
add_para('第九部分  推荐阅读与延伸资源', bold=True)
doc.add_page_break()

# ============== 第一部分 课程导览 ==============
add_title('第一部分  课程导览', level=1)
add_title('1.1  课程目标：你将带走什么', level=2)
add_para('这不是一个关于"创新思维技巧"的课程。这是一个关于"如何让团队持续产生创新"的课程。')
add_para('两天的学习，围绕着三个层次的目标展开：')
add_para('第一层：诊断能力', bold=True)
add_para('你能清楚地判断自己的团队在创新维度上"卡在哪里"。你将掌握一套经过验证的诊断框架——五个关键影响因素（心理安全感、认知多样性、探索空间、学习速度、领导者信号），以及三大要素（客户洞察、交互涌现、敏捷迭代）。')
add_para('第二层：认知转变', bold=True)
add_para('你能在"执行思维"和"创新思维"之间做出有意识的切换。创新型领导者和运营管理者的根本差异，不在于"谁更聪明"，而在于"在面对不同任务时，能否调用不同的行为模式"。')
add_para('第三层：行动能力', bold=True)
add_para('你能设计出带回团队立即可用的机制。每个要素都配有可操作的工作机制。')

add_title('1.2  学习地图', level=2)
add_table(
    ['阶段', '核心问题', '学习形式'],
    [
        ['第一天上午', '我的团队为什么不创新？', '案例 + 框架 + 个人诊断'],
        ['第一天下午', '如何让客户洞察和知识碰撞成为常态？', '框架 + 工具 + 练习'],
        ['第二天上午', '如何用最小成本验证最重要的假设？', '案例 + 改写 + 自我诊断'],
        ['第二天下午', '这些问题也发生在我的团队里吗？我要做什么？', '模拟 + 对照 + 承诺'],
    ]
)
doc.add_paragraph()

add_title('1.3  课前准备清单', level=2)
add_para('请在开课前 3 天完成以下准备。这不是仪式感，是让你的学习效果最大化的必要条件。')
add_para('必做项（30 分钟）', bold=True)
add_para('• 阅读开课前发的学员手册电子版一遍', indent_first=False)
add_para('• 准备一份你当前最想解决的一个创新挑战的简述（不超过 100 字）', indent_first=False)
add_para('• 准备一支你喜欢的笔', indent_first=False)
add_para('• 选一个不被打扰的两整天', indent_first=False)
add_para('强烈建议（60 分钟）', bold=True)
add_para('• 阅读课程提供的先导阅读材料', indent_first=False)
add_para('• 和你的直接上级做一次 10 分钟的沟通', indent_first=False)
add_para('• 翻阅你团队最近半年的会议纪要、提案记录', indent_first=False)
add_para('• 想一想：过去 6 个月里，你的团队有没有一个"让你感到意外惊喜"的想法？', indent_first=False)
add_para('不要做的事：', bold=True)
add_para('• 不要把这两天当作"换个地方办公"——关掉邮箱、关掉钉钉/企微/飞书', indent_first=False)
add_para('• 不要带笔记本电脑到课堂上', indent_first=False)
add_para('• 不要带"学习"的心态来', indent_first=False)

add_title('1.4  学员公约', level=2)
add_para('我们承诺：真实、在场、保密、好奇、行动。', bold=True)
add_para('我们不做什么：评判其他学员的团队、决策、风格；不在课堂上做"完美管理者"；不用"我认识的一个朋友"作为案例来回避自己的真实问题。', bold=True)

add_title('1.5  学习方法说明', level=2)
add_para('这门课有几个特别之处：')
add_para('1. 它不是一个"听"的课程。整个两天，你将完成至少 12 个不同形式的个人或小组练习。', indent_first=False)
add_para('2. 案例会在两天里反复出现。开场的赵建设、第四部分的林总、第二天的亮界科技——这些案例不是孤立的故事。', indent_first=False)
add_para('3. 作业是在课间完成的。第一天晚上有 15-20 分钟的"创新挑战卡"作业。', indent_first=False)
add_para('4. 第二天结束时，你不会收到"完成证书"——你会收到一份你自己填写的"行动承诺卡"。', indent_first=False)
doc.add_page_break()

# ============== 第二部分 第一天上午 ==============
add_title('第二部分  第一天上午：诊断你的团队', level=1)
add_para('> 使用说明  本部分是第一天的核心内容。你将完成"五个关键影响因素"的诊断，并开始识别自己的"无意识创新抑制行为"。请认真填写每一个练习——这些是你带回去最重要的工具。', italic=True)

add_title('2.1  开场：赵建设和那十几秒的沉默', level=2)
add_para('赵建设在这家公司做产品总监已经七年了。')
add_para('他不是一个不在乎创新的人。他订阅了三份行业通讯，年会参加产品峰会，书架上有十几本讲创新的书。他真心相信：只要给团队足够的资源和空间，好的想法就会冒出来。')
add_para('过去半年，行业里有几家新公司凭借创新产品迅速抢占市场，他的老板开始明确施压：必须有突破。')
add_para('赵建设立刻行动。他做了所有"应该做"的事：申请了创新专项预算，在公司内网搭建了"创意提案平台"，组建了专项小组，还把几个骨干送去参加了为期三天的创新思维工作坊。那几个骨干回来说感觉"思维打开了很多"。')
add_para('半年后，平台上有 51 条提案。成功落地的：0 条。')
add_para('大多数提案的搁置理由是"资源不足""时机不对""这个方向我们没有积累"。每一条理由单独拿出来，都很合理。')
add_para('在最后一次季度复盘会上，赵建设忍不住问大家："我们哪里出了问题？"')
add_para('会议室里，沉默了十几秒。')
add_para('然后有人说："可能是时机不对。"有人说："行业竞争太激烈了。"')
add_para('没有人说出自己真实的想法——包括那个三个月前曾经私下告诉赵建设"我觉得我们的提案平台没有人真的在认真看"的骨干员工。这一次，他也没有说。')

add_warning('在继续阅读之前，请先回答这个问题：')
add_para('你认为，赵建设团队最核心的问题是什么？用一句话写下你的第一反应：', bold=True)
add_blank_line(2)
add_para('（请先写下来——等今天结束再回来对照。你的答案大概率会变。）', italic=True)

add_tip('这个沉默，才是真正需要被解决的问题。一个管理者可以为团队提供创新所需的一切资源——但如果团队成员不愿意说出真实的想法，所有资源都只是道具。')
add_para('这两天，我们要一起搞清楚：那个沉默是怎么来的，以及你能做什么来改变它。')

add_title('2.2  第一部分：创新型团队的八个真相测试', level=2)
add_para('在进入核心内容之前，先检验你现在的认知基准。')
add_exercise_box('认知自测（2 分钟独立完成）', [
    '目的：找出你对"创新型团队"的认知盲区。没有对错之分，只有真实和不真实。每题在括号内标注你的判断：✓（认为正确）或 ✗（认为错误）。'
])
add_table(
    ['#', '判断', '你的答案'],
    [
        ['1', '团队的创新能力，主要取决于成员的聪明程度和创意天赋', '（    ）'],
        ['2', '鼓励员工犯错，是打造创新文化的有效方式', '（    ）'],
        ['3', '一个高绩效的执行型团队，通常也更容易在创新上有突破', '（    ）'],
        ['4', '管理者主动提出创意方向，有助于引导团队创新', '（    ）'],
        ['5', '团队对失败越宽容，创新产出的质量就越高', '（    ）'],
        ['6', '要提升团队创新力，首先要给大家更多"自由时间"', '（    ）'],
        ['7', '真正的客户需求，通过用户访谈就能有效识别', '（    ）'],
        ['8', '创新型领导者的核心任务，是筛选好想法并给予资源支持', '（    ）'],
    ]
)
doc.add_paragraph()

add_para('答案与解析（这部分解析将在课堂上由讲师展开讨论）：', bold=True)
add_para('第 1 题 ✗  创新能力与智力或创意天赋的关联远低于预期。哈佛商学院 Amy Edmondson 的研究显示，影响团队创新产出最大的单一因素是心理安全感。')
add_para('第 2 题 ✗  "允许犯错"和"鼓励犯错"是两件不同的事。真正有效的是建立学习型失败文化。')
add_para('第 3 题 ✗  高绩效执行型团队恰恰常常更难创新——执行优秀意味着流程稳定，而创新需要拥抱模糊。')
add_para('第 4 题 ✗  当管理者主动提出方向，团队会从"独立思考"切换到"解读老板意图"模式。')
add_para('第 5 题 ✗  宽容失败本身不直接提升创新质量。真正能提升的是从失败中提取可用知识的速度。')
add_para('第 6 题 ✗  给"自由时间"是必要条件，但不是充分条件。3M 的 15% 自由时间之所以有效，是因为有配套的知识流通机制。')
add_para('第 7 题 ✗  用户访谈能告诉你用户"说他们想要什么"，但用户经常无法清晰描述自己真正的未满足需求。')
add_para('第 8 题 ✗  "筛选好想法并分配资源"是评审者的角色。真正的创新型领导者的核心任务，是创造让好想法能够冒出来、流通、被快速检验的条件和环境。')

add_tip('关键认知  创新型团队的天花板，不是成员有多聪明，而是管理者给他们创造了多少"敢说、能碰、快学"的空间。')
doc.add_page_break()

add_title('2.3  第二部分：五个关键影响因素', level=2)
add_para('哪五个因素决定了一个团队能否持续创新？')
add_para('这五个因素不是凭空总结的清单，它们来自可追溯的研究成果：Google 内部的 Project Aristotle 研究、Amy Edmondson 的团队心理安全感研究、以及大量对创新型组织的实证观察。')

# 因素 1
add_title('2.3.1  因素一：心理安全感', level=3)
add_para('案例：林总的那次周会', bold=True)
add_para('林总是某零售公司的区域总监，团队业绩在全公司排名第一。')
add_para('一次周会上，一个入职半年的年轻员工提出了一个想法："我们在某个门店试过一种新的陈列方式，感觉比现在的方法转化率高，想跟大家分享一下。"')
add_para('林总微微皱眉，礼貌地说："这个不在今天的议程里，先推进主要工作吧。"')
add_para('会议继续了。')
add_para('两个月后，一家竞争对手推出了类似的陈列优化方案，取得了显著效果。')
add_para('林总的那个年轻员工，此后再也没有在周会上提过任何想法。')

add_para('心理安全感，', bold=True, indent_first=False)
add_para('是 Amy Edmondson 提出的概念，指团队成员相信：在这个团队里，说出不成熟的想法、提出质疑、承认自己不知道某件事，不会让自己被惩罚或被嘲笑。')
add_para('它不是"大家关系好"，也不是"没有绩效压力"。它是：当一个人冒险说出真实想法时，他相信不会为此付出代价。')
add_para('Edmondson 的研究有一个出乎意料的发现：心理安全感高的团队，上报的"错误"和"问题"反而更多。这听起来像坏事，但恰恰相反——他们愿意暴露问题，而不是掩盖它，这才是真正能改进和创新的团队。')
add_para('心理安全感是所有创新的地基。地基不稳时，给团队再多创新工具、培训、预算，都只是在沙上建楼。')

add_tip('关键认知  你团队的心理安全感高不高，不是看大家关系好不好，而是看这一条：上一次，谁在会议上说了你不想听的话？你当时的反应是什么？')

add_exercise_box('快速诊断：心理安全感（独立完成，1 分钟）', [
    '根据你对自己团队实际情况的感知，在每项后面的数字上圈选（1=几乎不存在，5=非常普遍）：',
])
add_table(
    ['诊断项目', '评分'],
    [
        ['团队成员会在我面前说出还不成熟、不确定的想法', '1　2　3　4　5'],
        ['有人提出与我不同的意见时，讨论会继续而不是终止', '1　2　3　4　5'],
        ['成员敢于在公开场合承认自己不懂、不确定某件事', '1　2　3　4　5'],
        ['过去 6 个月，有失败或坏消息被当作学习机会认真对待', '1　2　3　4　5'],
    ]
)
doc.add_paragraph()
add_para('心理安全感小计：____ 分（满分 20 分）', bold=True)
add_para('关于这个维度的具体观察：', bold=True)
add_para('（写下 1-2 件你团队最近发生的、和这个维度相关的具体事件）')
add_blank_line(2)

# 因素 2
add_title('2.3.2  因素二：认知多样性', level=3)
add_para('Google 在 2012 年启动了名为 Project Aristotle 的内部研究，目标是找出"什么让一支团队表现卓越"。他们分析了公司内 180 多个团队，几乎考察了所有可能的变量：成员学历、工作经验、性格类型、甚至团队聚餐频率。')
add_para('结论出乎所有人意料：成员的个人素质，对团队表现的影响远远小于团队成员的互动方式。')
add_para('而对创新能力影响最显著的特征，不是"每个人都很聪明"，而是：团队中存在不同的思维方式，并且这些思维方式都能被听到。')
add_para('认知多样性，', bold=True, indent_first=False)
add_para('不等于性别或背景的多样性（尽管有时能带来认知多样性），而是指：有人擅长宏观思考，有人习惯扣细节；有人从用户情感出发，有人从数据出发；有人天然发现可能性，有人天然发现风险。')
add_para('这种多样性是创新的燃料——不同的思维模式碰在一起，才会产生真正意想不到的连接。')
add_tip('关键认知  认知多样性不是让团队充满分歧。它的意思是：确保团队里有人能看到你看不到的东西，并且那个人敢于说出来。')

add_exercise_box('快速诊断：认知多样性', [])
add_table(
    ['诊断项目', '评分'],
    [
        ['团队里有明显不同思维风格的成员', '1　2　3　4　5'],
        ['讨论中，与主流意见不同的声音不会被快速淹没', '1　2　3　4　5'],
        ['过去半年，有一个"来自意想不到的人"的想法被认真对待过', '1　2　3　4　5'],
        ['我能清楚说出团队里每个人与我"思维方式最不同"的地方', '1　2　3　4　5'],
    ]
)
doc.add_paragraph()
add_para('认知多样性小计：____ 分（满分 20 分）', bold=True)
add_para('关于这个维度的具体观察：')
add_blank_line(2)

# 因素 3
add_title('2.3.3  因素三：探索空间', level=3)
add_para('某消费品公司的研发团队，一年 365 天，日程排得满满当当：周一全员会，周二项目同步，周三客户对接，周四跨部门协调，周五冲刺……')
add_para('产品负责人林峰有时候隐隐感到一种焦虑："我们团队的人都很聪明，但好像大家只是在处理眼前的事情，很少有人在认真思考三年后这个产品应该是什么。"')
add_para('当被问到团队有没有时间做一些不直接产出的探索时，他沉默了一下说："这种话说出去，老板会觉得我们在混日子。"')
add_para('探索空间，', bold=True, indent_first=False)
add_para('是指团队成员能够把一部分时间和精力，用于非直接产出的探索活动——学习新知识，尝试没有明确方向的假设，研究那些"不确定有没有用"的问题。')
add_para('这不是奢侈品。这是创新的必要条件。')
add_para('一个团队如果 100% 的时间都在执行已经确定的任务，它在做的是已知范围内的事情。创新，永远发生在已知的边界处。没有探索空间，团队永远不会靠近那个边界。')
add_tip('关键认知  探索空间不是效率的敌人，而是创新的苗圃。')

add_exercise_box('快速诊断：探索空间', [])
add_table(
    ['诊断项目', '评分'],
    [
        ['团队成员有定期的非任务性学习和探索时间', '1　2　3　4　5'],
        ['探索活动不会因为"项目紧急"而被反复取消', '1　2　3　4　5'],
        ['成员能够对"还没有明确答案"的问题进行持续深入的研究', '1　2　3　4　5'],
        ['探索中产生的洞察、灵感、新视角，有渠道进入正式讨论和决策', '1　2　3　4　5'],
    ]
)
doc.add_paragraph()
add_para('探索空间小计：____ 分（满分 20 分）', bold=True)
add_para('关于这个维度的具体观察：')
add_blank_line(2)

# 因素 4
add_title('2.3.4  因素四：学习速度', level=3)
add_para('有两支销售团队，同时试用了一个新的客户沟通策略。')
add_para('A 团队在季度结束时做了一次正式复盘，得出结论："策略整体有效，继续推进。"')
add_para('B 团队每两周举行一次 15 分钟的"快速学习会"：什么有用？什么没用？下两周要调整什么？')
add_para('三个月后，A 团队的沟通策略还是最初的版本。B 团队的策略已经迭代了六轮，转化率提升了 40%。')
add_para('学习速度，', bold=True, indent_first=False)
add_para('指的是团队从行动中提取洞察、更新认知、调整方向的速度。受以下几个要素影响：')
add_para('• 反思的频率 —— 多久做一次复盘？', indent_first=False)
add_para('• 反思的深度 —— 是总结发生了什么，还是分析为什么发生？', indent_first=False)
add_para('• 从失败中学习的意愿', indent_first=False)
add_para('• 知识的流通', indent_first=False)
add_tip('关键认知  执行快 ≠ 学习快。')

add_exercise_box('快速诊断：学习速度', [])
add_table(
    ['诊断项目', '评分'],
    [
        ['团队有定期（至少每月一次）的复盘机制，且有具体的行动结论', '1　2　3　4　5'],
        ['失败项目或错误决策会被作为学习案例认真讨论', '1　2　3　4　5'],
        ['一个人在工作中获得的洞察，有机制让其他成员知道并受益', '1　2　3　4　5'],
        ['团队在执行过程中会持续微调方向', '1　2　3　4　5'],
    ]
)
doc.add_paragraph()
add_para('学习速度小计：____ 分（满分 20 分）', bold=True)
add_para('关于这个维度的具体观察：')
add_blank_line(2)

# 因素 5
add_title('2.3.5  因素五：领导者信号', level=3)
add_para('某公司 CEO 在年会上宣布："我们鼓励创新，欢迎大家大胆试错。"')
add_para('三个月后的季度总结大会上，他表彰了五个成功落地的创新项目，颁发了奖金和证书。')
add_para('没有一个"失败但带来重要学习"的案例被提及。没有一个"勇敢实验却没有成功"的团队被认可。')
add_para('一年后的员工调研显示：大多数人认为，公司"口头支持创新，实际还是只看结果"。')
add_para('领导者信号，', bold=True, indent_first=False)
add_para('是五个因素里最隐性、也最强大的一个。它指的是：管理者的日常行为——他关注什么，在会议上如何反应，表彰什么，追问什么，对什么沉默——向团队持续传递着"什么才是真正被重视的"的信号。')
add_para('这些信号比管理者说的话更有力量。')
add_tip('关键认知  管理者的行为对团队文化的影响，比语言大 10 倍。')

add_exercise_box('快速诊断：领导者信号', [])
add_table(
    ['诊断项目', '评分'],
    [
        ['我在过去 6 个月，在公开场合表彰过"有价值的失败学习"案例', '1　2　3　4　5'],
        ['当有人提出与我不同的观点，我的第一反应是好奇而不是评判', '1　2　3　4　5'],
        ['我给团队的日程安排，留有非执行性的"探索时间"', '1　2　3　4　5'],
        ['团队成员能举出具体例子，说明"我觉得安全表达"', '1　2　3　4　5'],
    ]
)
doc.add_paragraph()
add_para('领导者信号小计：____ 分（满分 20 分）', bold=True)
add_para('关于这个维度的具体观察：')
add_blank_line(2)
doc.add_page_break()

# 健康度快照
add_title('2.4  团队创新健康度快照', level=2)
add_exercise_box('综合诊断练习（10 分钟：独立完成 8 分钟 + 小组分享 2 分钟）', [
    '步骤一：汇总你五个维度的得分，填入下表。',
    '步骤二：找出得分最低的 1-2 个维度，写下你的观察。',
    '步骤三：在小组里用 2 分钟分享你得分最低的那个维度。'
])
add_table(
    ['因素', '你的得分', '满分'],
    [
        ['心理安全感', '', '20'],
        ['认知多样性', '', '20'],
        ['探索空间', '', '20'],
        ['学习速度', '', '20'],
        ['领导者信号', '', '20'],
        ['总计', '', '100'],
    ],
    first_col_bold=True
)
doc.add_paragraph()

add_para('快速可视化（把你的分数对应的数字圈起来）：', bold=True)
add_table(
    ['因素', '分数', '分布'],
    [
        ['心理安全感', '__/20', '0 2 4 6 8 [10] 12 14 16 18 20'],
        ['认知多样性', '__/20', '0 2 4 6 8 [10] 12 14 16 18 20'],
        ['探索空间', '__/20', '0 2 4 6 8 [10] 12 14 16 18 20'],
        ['学习速度', '__/20', '0 2 4 6 8 [10] 12 14 16 18 20'],
        ['领导者信号', '__/20', '0 2 4 6 8 [10] 12 14 16 18 20'],
    ],
    first_col_bold=True
)
doc.add_paragraph()

add_para('我得分最低的维度是：', bold=True)
add_blank_line(1)
add_para('它在我的团队里的具体表现是：', bold=True)
add_blank_line(2)
add_para('如果这个维度的问题持续下去，未来一年我最担心会发生的是：', bold=True)
add_blank_line(2)
add_tip('给你的提示  五个因素之间有内在关联：心理安全感是地基，领导者信号是土壤。')

doc.add_page_break()

# 第三部分
add_title('2.5  第三部分：创新型领导者 vs. 运营管理者', level=2)
add_para('在上半场，你拿到了你团队的创新健康度诊断。现在进入一个更难面对的问题：')
add_para('在那份诊断里，有多少是你的管理方式造成的？', bold=True)
add_para('这不是责难，而是解放——因为如果问题的根源在你，改变的钥匙也在你。')

add_para('来看两个场景，找出两位管理者的差异在哪里：', bold=True)
add_para('场景 A：', bold=True)
add_para('产品经理小王在周会上说："我想探索一下用 AR 技术优化我们的用户引导流程，现在还不知道能不能行。"')
add_para('管理者回应："这个方向很好，你先做一个可行性报告，下周给我看。"')
add_para('场景 B：', bold=True)
add_para('同样的情况，另一位管理者的回应："有意思。你觉得这个想法里最不确定的地方是什么？我们下午花 20 分钟先聊聊。"')
add_para('两个回应都是"支持"，但结果完全不同。')
add_para('场景 A 的小王需要写报告。他会把大量精力用于让想法显得更成熟、更可行，而不是让它更有探索性。')
add_para('场景 B 的小王被邀请进入了对话。他会更愿意继续思考这个想法里的不确定性。')
add_para('这就是创新型领导者和运营管理者在面对同一个场景时的差异。不是有没有支持，而是以什么方式回应，会把团队引向哪个方向。')

add_title('两种管理角色的行为差异', level=3)
add_table(
    ['场景', '运营管理者的典型反应', '创新型领导者的典型反应'],
    [
        ['有人提出不成熟的想法', '"想清楚了再来找我"', '"这个想法最有意思的地方是什么？说说"'],
        ['实验失败了', '"为什么会失败？怎么避免下次再发生"', '"我们从这个失败里学到了什么新的东西"'],
        ['团队方向出现分歧', '"我来拍板：方向是 X，大家执行"', '"我听到了两个不同方向，我们先把各自的假设摆出来"'],
        ['安排工作', '"这周优先级是完成 ABC 三个任务"', '"核心任务是 AB，留 20% 时间给 C 方向的探索"'],
        ['有人提出异见', '"你说的有道理，但现在不是时候"', '"你说的和我的判断不一样，帮我理解你的理由"'],
        ['季度总结', '表彰成功项目，分析成功经验', '表彰成功项目，同时复盘一个"有价值的失败"'],
        ['招募团队成员', '倾向于找"能力强、方向匹配"的人', '刻意寻找"与团队现有思维方式不同"的人'],
        ['资源分配', '把资源给有明确 ROI 预期的项目', '留一部分资源给"还不确定 ROI 但值得探索"的方向'],
    ],
    first_col_bold=True
)
doc.add_paragraph()

add_warning('运营管理者的做法在执行密集的任务中是正确的。问题不是哪一套更好，而是：你是否能在需要的时刻切换，还是在所有情况下只有一个默认模式？')

add_title('最难的部分：无意识的创新抑制', level=3)
add_para('大多数管理者并不是有意压制团队创新的。他们只是在用"有效管理"的逻辑——这套逻辑在日常运营中运转良好，但在需要创新的时刻，它会无意间发出错误的信号。')

add_para('① 要求想法"成熟"后再提出', bold=True)
add_para('听起来是在帮助团队提升质量，实际效果是：大家开始自我审查，把还在萌芽状态的想法扼杀在说出口之前。')

add_para('② 只在成功时给予可见的认可', bold=True)
add_para('听起来是在强调结果导向，实际效果是：团队学会了只做"有把握成功"的事。')

add_para('③ 遇到问题第一反应是找"谁的责任"', bold=True)
add_para('听起来是在强调责任感，实际效果是：团队成员在实验遇挫时选择隐瞒或拖延暴露。')

add_para('④ 在会议上首先表达自己的看法', bold=True)
add_para('听起来是在高效引领讨论，实际效果是：团队成员在你发言后会把后续发言方向对准"支持或延伸你的观点"。')

add_para('⑤ 把"探索时间"列为低优先级，遇到项目压力就取消', bold=True)
add_para('听起来是在正确管理优先级，实际效果是：团队理解了"探索是可以被牺牲的"。')

doc.add_page_break()

# 行为盘点
add_title('2.6  我的领导行为盘点', level=2)
add_exercise_box('我的领导行为盘点（20 分钟：个人盘点 12 分钟 + 两人互访 8 分钟）', [
    '目的：找出你在日常工作中无意间抑制创新的行为模式。',
    '不是评判自己，而是看清楚，才能改变。'
])

add_para('步骤一：回顾过去 30 天（12 分钟）', bold=True)
add_para('在下面的表格里，回忆过去 30 天内实际发生过的具体事件，写下你实际的反应：')
add_table(
    ['场景', '发生的具体事件（用一句话描述）', '你当时的实际反应'],
    [
        ['有人提出一个不成熟的新想法', '', ''],
        ['有项目或实验遇到挫折', '', ''],
        ['团队开始讨论一个"不确定能不能行"的方向', '', ''],
        ['有人在会议上提出了和你不同的判断', '', ''],
        ['季度总结或汇报场合', '', ''],
    ],
    first_col_bold=True
)
doc.add_paragraph()

add_para('步骤二：对照"两种管理角色"的表格，判断你的行为偏向', bold=True)
add_table(
    ['场景', '你的倾向是'],
    [
        ['有人提出不成熟的想法', '运营管理者 / 创新型领导者 / 视情况而定'],
        ['有项目遇到挫折', '运营管理者 / 创新型领导者 / 视情况而定'],
        ['讨论不确定的方向', '运营管理者 / 创新型领导者 / 视情况而定'],
        ['有人提出异见', '运营管理者 / 创新型领导者 / 视情况而定'],
        ['季度总结/汇报', '运营管理者 / 创新型领导者 / 视情况而定'],
    ],
    first_col_bold=True
)
doc.add_paragraph()

add_para('步骤三：找出你最想改变的一个行为（2 分钟）', bold=True)
add_para('如果你下周只能改变一件事，你会选择什么？')
add_blank_line(1)
add_para('我想改变的行为是：____________________________________________')
add_blank_line(1)
add_para('改变后，我希望团队感受到的不同是：________________________________')

add_para('步骤四：两人互访（8 分钟，每人 4 分钟）', bold=True)
add_para('和你旁边的一位同学分享你步骤三的答案——只说"我想改变什么"和"我希望带来什么不同"。')
add_blank_line(1)
add_para('听到的共鸣点：________________________________________________')

add_tip('上半场关键认知小结  创新型团队的天花板，不是成员有多聪明，而是管理者给他们创造了多少"敢说、能碰、快学"的空间。')
doc.add_page_break()

print('学员手册前半部分生成完成...')
# 暂时存盘
doc.save(OUTPUT)
print(f'已保存到: {OUTPUT}')
print(f'文件大小: {os.path.getsize(OUTPUT)/1024:.1f} KB')

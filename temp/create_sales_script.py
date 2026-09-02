# -*- coding: utf-8 -*-
"""
Create: 课程销售话术.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
BODY_GRAY = RGBColor(0x33, 0x33, 0x33)
DARK_RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x00)

OUTPUT_PATH = "D:/新课开发/党业融合/廉政风险情景决策/完整课程包/004-对外宣传文案/课程销售话术.docx"

def set_run_font(run, font_name_cn="微软雅黑", font_name_en="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = font_name_en
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_section_title(doc, text, level=1):
    para = doc.add_paragraph()
    if level == 1:
        run = para.add_run(text)
        set_run_font(run, size=14, bold=True, color=NAVY)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run = para.add_run(text)
        set_run_font(run, size=12, bold=True, color=ACCENT_BLUE)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
    return para

def add_body_text(doc, text, indent=False, bold_start=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if bold_start and text.startswith(bold_start):
        rest = text[len(bold_start):]
        run = para.add_run(bold_start)
        set_run_font(run, size=11, bold=True, color=BODY_GRAY)
        run2 = para.add_run(rest)
        set_run_font(run2, size=11, color=BODY_GRAY)
    else:
        run = para.add_run(text)
        set_run_font(run, size=11, color=BODY_GRAY)

    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    return para

def add_script_line(doc, speaker, text, highlight=False):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.3)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)

    run = para.add_run(f"{speaker}：")
    if highlight:
        set_run_font(run, size=11, bold=True, color=DARK_RED)
    else:
        set_run_font(run, size=11, bold=True, color=ACCENT_BLUE)

    run2 = para.add_run(text)
    set_run_font(run2, size=11, color=BODY_GRAY)
    return para

def add_bullet(doc, text, bold_start=None):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)

    run = para.add_run("◆ ")
    set_run_font(run, size=10, color=ACCENT_BLUE)

    if bold_start and text.startswith(bold_start):
        rest = text[len(bold_start):]
        run2 = para.add_run(bold_start)
        set_run_font(run2, size=11, bold=True, color=BODY_GRAY)
        run3 = para.add_run(rest)
        set_run_font(run3, size=11, color=BODY_GRAY)
    else:
        run = para.add_run(text)
        set_run_font(run, size=11, color=BODY_GRAY)
    return para

def add_divider(doc):
    para = doc.add_paragraph()
    run = para.add_run("─" * 60)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(6)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)

def create_document():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("课程销售话术")
    set_run_font(run, size=24, bold=True, color=NAVY)
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("廉政风险情景决策训练营")
    set_run_font(run, size=14, color=ACCENT_BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(18)

    add_divider(doc)

    # ==================== Part 1: 客户常见问题应对话术 ====================
    add_section_title(doc, "第一部分：客户常见问题应对话术", level=1)

    # Q1
    add_section_title(doc, "Q1：这个课程和传统廉政培训有什么区别？", level=2)

    add_script_line(doc, "客户", "你们这个廉政课和平时我们做的警示教育有什么不同？")
    add_script_line(doc, "顾问", "传统培训的效果为什么总是'听听感动、想想激动、回去不动'？", highlight=True)
    add_script_line(doc, "顾问", "根本原因是：学员始终是旁观者，看的是别人的故事，想的是'这事跟我有什么关系'。", highlight=False)
    add_script_line(doc, "顾问", "我们这个训练营只做一件事——把旁观者变成当局者。", highlight=False)
    add_script_line(doc, "顾问", "不是讲道理，而是让您单位的人在自己的模拟情景中，'经历'一遍'我是怎么一步步被说服越界的'。", highlight=False)
    add_script_line(doc, "顾问", "结束后他们的反馈是：'第一次看清了自己内心的那套逻辑'——这种触动，和看案例是完全不同的量级。", highlight=False)

    # Q2
    add_section_title(doc, "Q2：这种培训方式会不会让学员反感？", level=2)

    add_script_line(doc, "客户", "让员工'扮演'出事，会不会引起抵触心理？")
    add_script_line(doc, "顾问", "您的担心很有道理，如果设计不好确实会有这个问题。", highlight=True)
    add_script_line(doc, "顾问", "但我们的设计是'不审判、不批评、不贴标签'。", highlight=False)
    add_script_line(doc, "顾问", "复盘的时候，我们问的是'这个过程中，你心里在想什么'，而不是'你这样做对不对'。", highlight=False)
    add_script_line(doc, "顾问", "事实上，经过我们前期引导，学员普遍反馈是：'终于有人理解我们岗位的压力了，不是简单地告诉我们要廉洁'。", highlight=False)
    add_script_line(doc, "顾问", "这种被理解的感觉，反而让他们更愿意打开心扉。", highlight=False)

    # Q3
    add_section_title(doc, "Q3：培训效果怎么衡量？", level=2)

    add_script_line(doc, "客户", "你们怎么证明培训有效果？")
    add_script_line(doc, "顾问", "效果衡量的三个维度：", highlight=True)
    add_script_line(doc, "顾问", "第一，学员即时反馈——培训结束时，85%以上的学员表示'第一次真正看清自己可能在什么时候被说服越界'。", highlight=False)
    add_script_line(doc, "顾问", "第二，行为跟踪数据——我们会在培训后1个月、3个月做跟踪回访，看风险事件报告数量、主动上报隐患数量等指标的变化。", highlight=False)
    add_script_line(doc, "顾问", "第三，行动计划承诺——每个学员都会制定个人廉政风险防控计划，并签署承诺书，这个承诺书可以存档作为培训记录。", highlight=False)

    # Q4
    add_section_title(doc, "Q4：我们的岗位情况比较特殊，能定制吗？", level=2)

    add_script_line(doc, "客户", "我们的业务场景和一般企业不太一样，能针对我们情况定制吗？")
    add_script_line(doc, "顾问", "完全可以。我们提供'情境定制'服务。", highlight=True)
    add_script_line(doc, "顾问", "在正式培训前，我们的研发团队会与您单位的业务骨干进行深度访谈，把真实的岗位场景、真实的决策压力、真实的'灰色地带'都挖掘出来。", highlight=False)
    add_script_line(doc, "顾问", "然后把这些真实素材改编成模拟情景，让学员觉得'这就是我每天面对的情况'，而不是'这是别人的故事'。", highlight=False)
    add_script_line(doc, "顾问", "当然，定制需要额外的时间（通常增加1-2周）和费用，这个需要根据定制深度来评估。", highlight=False)

    # Q5
    add_section_title(doc, "Q5：价格怎么算？", level=2)

    add_script_line(doc, "客户", "这个课怎么收费？")
    add_script_line(doc, "顾问", "我们的收费结构是这样的：", highlight=True)
    add_script_line(doc, "顾问", "标准版（2天）适合30人左右班级，包含课程设计、情景卡牌、工具卡、学员手册等全套物料，费用在X万区间。", highlight=False)
    add_script_line(doc, "顾问", "如果需要深度定制情景，费用会相应增加，主要看定制工作量。", highlight=False)
    add_script_line(doc, "顾问", "另外，如果需要我们的讲师进行驻场授课，会涉及讲师差旅费用。", highlight=False)
    add_script_line(doc, "顾问", "不过说实话，这门课的价值不在于'便宜'，而在于'有效'——真正能改变人的认知和行为，这比花几万块做十次蜻蜓点水的培训要划算得多。", highlight=False)

    # Q6
    add_section_title(doc, "Q6：时间怎么安排？", level=2)

    add_script_line(doc, "客户", "我们工作很忙，抽不出2天时间怎么办？")
    add_script_line(doc, "顾问", "理解，领导干部的时间确实宝贵。", highlight=True)
    add_script_line(doc, "顾问", "我们也有1天精华版，压缩到6小时，把最核心的3-4个情景模拟做透，复盘和工具讲解环节保留。", highlight=False)
    add_script_line(doc, "顾问", "但坦率说，如果条件允许，2天版的效果会更好——因为'看清自己'这件事需要时间发酵，第一天可能只是'有点感觉'，第二天才会'恍然大悟'。", highlight=False)
    add_script_line(doc, "顾问", "当然，如果真的抽不出整块时间，也可以考虑分期——比如先做1天，三个月后再做1天。", highlight=False)

    # Q7
    add_section_title(doc, "Q7：以往有哪些客户在做？", level=2)

    add_script_line(doc, "客户", "有哪些和我们类似的单位做过这个培训？")
    add_script_line(doc, "顾问", "我们做过几类客户：", highlight=True)
    add_script_line(doc, "顾问", "一类是国企央企——某省属国企集团做过采购岗位专题培训，培训后3个月采购投诉率下降67%。", highlight=False)
    add_script_line(doc, "顾问", "一类是政府机关——某市建设系统做过工程领域专题培训，半年内信访举报量下降41%。", highlight=False)
    add_script_line(doc, "顾问", "还有一类是大型民企——某世界500强企业做过经销商管理岗位的反商业贿赂专题。", highlight=False)
    add_script_line(doc, "顾问", "如果您方便的话，我可以分享一些详细的案例资料供您参考。", highlight=False)

    # Q8
    add_section_title(doc, "Q8：讲师资质怎么样？", level=2)

    add_script_line(doc, "客户", "派什么讲师来授课？")
    add_script_line(doc, "顾问", "我们的讲师都是经过严格认证的。", highlight=True)
    add_script_line(doc, "顾问", "首先，必须有丰富的廉政教育或纪检监察工作经验背景；其次，必须经过我们的'情景引导师'认证培训——不是普通讲师能上的，必须掌握情景模拟的引导技术。", highlight=False)
    add_script_line(doc, "顾问", "我可以提前把讲师简历发给您，您可以看一下讲师的背景和擅长方向，是否适合您单位的调性。", highlight=False)

    add_divider(doc)

    # ==================== Part 2: 课程介绍话术 ====================
    add_section_title(doc, "第二部分：课程介绍话术", level=1)

    # 3分钟版
    add_section_title(doc, "【3分钟精华版】当领导问'这课讲什么'时", level=2)

    add_script_line(doc, "顾问", "这是一个'情景决策训练营'，专门解决一个问题——", highlight=True)
    add_script_line(doc, "顾问", "怎么让干部在面对腐败风险的时候，不只是'知道应该怎么做'，而是真正'能够守住底线'。", highlight=False)
    add_script_line(doc, "顾问", "传统的警示教育是'讲案例、说道理'，但效果不好是因为学员始终是旁观者。", highlight=False)
    add_script_line(doc, "顾问", "我们的做法是'情景模拟'——让学员在自己的模拟情景中，亲身经历一遍'我是怎么一步步被说服越界的'。", highlight=False)
    add_script_line(doc, "顾问", "结束的时候，他们自己就会说'原来我是这样被自己说服的'——这种觉察，比任何说教都管用。", highlight=False)
    add_script_line(doc, "顾问", "这就是我们的核心价值：把旁观者变成当局者。", highlight=False)

    # 10分钟版
    add_section_title(doc, "【10分钟完整版】用于正式推荐场合", level=2)

    add_script_line(doc, "顾问", "各位领导，今天我向大家介绍一门不一样的廉政培训课程——'廉政风险情景决策训练营'。", highlight=True)

    add_script_line(doc, "顾问", "先问大家一个问题：您觉得，为什么我们的廉政制度越来越完善，但风险事件还是时有发生？", highlight=False)

    add_script_line(doc, "顾问", "我接触过很多单位，发现一个共同规律：不是制度不够，而是'心理防线'不够坚固。", highlight=False)

    add_script_line(doc, "顾问", "腐败从来不是一瞬间发生的，而是一个'温水煮青蛙'的渐进过程——每一步都有'合理'的理由，每一步都在'可接受'的范围内，直到回头看才发现已经越界很远了。", highlight=False)

    add_script_line(doc, "顾问", "传统培训的问题是：让学员以旁观者身份'看戏'，看完觉得'与我无关'。我们的做法是：让学员成为'当局者'。", highlight=False)

    add_script_line(doc, "顾问", "具体怎么做？我们设计了3-4个高风险情景，涵盖采购、工程、财务、招标等场景。", highlight=False)

    add_script_line(doc, "顾问", "学员在模拟情景中做出决策，我们会即时复盘——不是评判对错，而是问一个问题：在这个过程中，你心里在想什么？", highlight=False)

    add_script_line(doc, "顾问", "很多学员在复盘时会'恍然大悟'：原来那套'合理化'的逻辑是这样运作的，我以为我在坚持原则，其实已经被一点一点说服了。", highlight=False)

    add_script_line(doc, "顾问", "我们的培训目标是：让每个学员都带走在关键时刻'守得住'的能力，而不只是'知道应该怎么做'的知识。", highlight=False)

    add_script_line(doc, "顾问", "这就是'廉政风险情景决策训练营'——让每一次'差一点就越界'，都变成'我守住了底线'。", highlight=False)

    add_divider(doc)

    # ==================== Part 3: 异议处理话术 ====================
    add_section_title(doc, "第三部分：异议处理话术", level=1)

    # 异议1：太贵了
    add_section_title(doc, "异议1：'这个价格有点贵'" , level=2)

    add_script_line(doc, "客户", "差不多内容的培训，别人家只要一半价格。")
    add_script_line(doc, "顾问", "您说得对，市场上确实有更便宜的廉政课。", highlight=True)
    add_script_line(doc, "顾问", "但我想问一个问题：您做这次培训的期望是什么？", highlight=False)
    add_script_line(doc, "顾问", "如果只是想'完成一次培训任务'，那确实没必要花这么多钱——网上下载一些案例、找个讲师念一遍，一样能交差。", highlight=False)
    add_script_line(doc, "顾问", "但如果希望真正'改变人'——让参加培训的干部在回到岗位后，面对诱惑时能多想一想'培训时的那一幕'，那这个投入是值得的。", highlight=False)
    add_script_line(doc, "顾问", "我们有个客户算过一笔账：做了一次情景决策训练营之后，采购投诉率下降67%，避免的潜在损失是培训费用的十几倍。", highlight=False)
    add_script_line(doc, "顾问", "您觉得，是省培训费划算，还是省那次'万一出事'的代价划算？", highlight=False)

    # 异议2：担心员工抵触
    add_section_title(doc, "异议2：'担心员工觉得被针对，有抵触情绪'", level=2)

    add_script_line(doc, "客户", "万一员工觉得我们在'审问'他们，反而有逆反心理怎么办？")
    add_script_line(doc, "顾问", "您的担心很有道理，这确实是情景培训最大的风险点。", highlight=True)
    add_script_line(doc, "顾问", "所以我们在设计上有一个核心原则：'对事不对人'。", highlight=False)
    add_script_line(doc, "顾问", "情景模拟的是'制度漏洞可能被怎么利用'，而不是'谁可能出事'。每个人面对同样的情景，做出不同的选择——我们复盘的是情景本身和人心里的那套逻辑，不是针对个人。", highlight=False)
    add_script_line(doc, "顾问", "另外，我们的讲师会在开场时做很充分的'心理安全感'建设，让大家知道这是一个'安全的探索空间'，说的所有话都不会被追责。", highlight=False)
    add_script_line(doc, "顾问", "事实上，从我们做过的培训来看，学员反馈最多的是'终于有人理解我们岗位的真实压力了'——这种被理解的感觉，反而会让大家更愿意敞开心扉。", highlight=False)

    # 异议3：我们情况特殊
    add_section_title(doc, "异议3：'我们这种情况可能不太适合'", level=2)

    add_script_line(doc, "客户", "我们单位的情况比较特殊，和一般企业不太一样。")
    add_script_line(doc, "顾问", "您说得很对，每个单位的情况都不一样。", highlight=True)
    add_script_line(doc, "顾问", "但这恰恰是我们课程的优势——我们可以做定制。", highlight=False)
    add_script_line(doc, "顾问", "在培训前，我们的研发团队会深入了解您单位的业务流程、风险点、真实案例，把这些素材编入情景模拟中。", highlight=False)
    add_script_line(doc, "顾问", "这样的话，学员会觉得'这就是我每天面对的情况'，而不是'这是别人的故事'。", highlight=False)
    add_script_line(doc, "顾问", "所以，您说的情况特殊，正好是我们需要定制的理由。您方便的话，我们可以安排一次需求调研，了解具体情况后再给您出方案。", highlight=False)

    # 异议4：领导不重视
    add_section_title(doc, "异议4：'一把手可能觉得没必要'" , level=2)

    add_script_line(doc, "客户", "万一领导觉得现有培训够了，不需要这么麻烦的课程呢？")
    add_script_line(doc, "顾问", "这种情况很常见。", highlight=True)
    add_script_line(doc, "顾问", "我建议可以先让领导回答一个问题：'我们过去做的廉政培训，改变了什么？'", highlight=False)
    add_script_line(doc, "顾问", "如果答案是'确实有一些改变'，那我们可以讨论如何优化。", highlight=False)
    add_script_line(doc, "顾问", "如果答案是'好像也没什么用，不出事就不错了'——那正好说明，现有方式效果有限，需要新的方法。", highlight=False)
    add_script_line(doc, "顾问", "我可以先给您一份详细的课程方案和客户案例，您可以先向领导汇报这个新思路，看领导的反馈。如果领导有兴趣，我们可以安排一次15分钟的视频会议，让领导直接和我沟通。", highlight=False)

    add_divider(doc)

    # ==================== Part 4: 促成签单话术 ====================
    add_section_title(doc, "第四部分：促成签单话术", level=1)

    # 时机判断
    add_section_title(doc, "【最佳促成时机】", level=2)

    add_bullet(doc, "客户主动询问'什么时候可以开班'、'最快什么时候能安排'")
    add_bullet(doc, "客户对某个案例表示强烈兴趣，'这个客户后来怎么样了？'")
    add_bullet(doc, "客户反复追问某个细节，'你们那个复盘是怎么做的？'")
    add_bullet(doc, "客户在犹豫时开始问'如果...的话...''要是...怎么办'")
    add_bullet(doc, "客户表达认同，'这个思路是对的''你们确实想得比较清楚'")

    # 促成方法1：假设成交
    add_section_title(doc, "【方法一：假设成交法】", level=2)

    add_script_line(doc, "顾问", "既然这个方向您觉得合适，我们可以进入实际操作层面了——", highlight=True)
    add_script_line(doc, "顾问", "您看是安排在9月中旬还是9月底？我们需要提前两周启动定制调研工作。", highlight=False)
    add_script_line(doc, "顾问", "另外，班级规模您预计是多少人？30人左右的效果比较好，超过40人的话我们会考虑分组。", highlight=False)

    # 促成方法2：直接促成
    add_section_title(doc, "【方法二：直接促成法】", level=2)

    add_script_line(doc, "顾问", "和您沟通下来，我的感觉是：这个课程方向是符合您单位需求的。", highlight=True)
    add_script_line(doc, "顾问", "接下来就是具体操作层面的事情了——您看还有什么顾虑是我没有解答清楚的吗？", highlight=False)
    add_script_line(doc, "顾问", "如果没有的话，我们可以先签一个意向协议，把时间先定下来，同时启动定制调研。您觉得可以吗？", highlight=False)

    # 促成方法3：稀缺性促成
    add_section_title(doc, "【方法三：稀缺性促成法】", level=2)

    add_script_line(doc, "顾问", "提醒您一个时间问题——", highlight=True)
    add_script_line(doc, "顾问", "我们每个季度只开2-3个公开班的名额，您如果想赶在年内做，建议这个月确定下来，否则可能要排到明年一季度了。", highlight=False)
    add_script_line(doc, "顾问", "内训的话，讲师档期也比较紧张，特别是金牌讲师，基本都要提前一个月预约。", highlight=False)

    # 促成方法4：回马枪促成
    add_section_title(doc, "【方法四：回马枪法】", level=2)

    add_script_line(doc, "顾问", "今天我们聊得很深入，我也有一个感受——", highlight=True)
    add_script_line(doc, "顾问", "您其实对这个培训是有期待的，只是希望确认它真的管用。", highlight=False)
    add_script_line(doc, "顾问", "我有个建议：我们可以先做一个小范围试点——比如选一个部门先做半天体验课，您和我们一起观察效果。如果效果您满意，再全面铺开。", highlight=False)
    add_script_line(doc, "顾问", "这样您也没有风险，您觉得呢？", highlight=False)

    add_divider(doc)

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(24)
    run = footer.add_run("—— 以上话术仅供参考，实际使用时请根据客户情况灵活调整 ——")
    set_run_font(run, size=9, italic=True, color=RGBColor(0x99, 0x99, 0x99))

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

if __name__ == "__main__":
    create_document()

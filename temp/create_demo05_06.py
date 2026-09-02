# -*- coding: utf-8 -*-
"""
生成 demo05-类比三问清单示例.docx
demo06-反面喻体法汇报示例.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_demo05():
    """demo05: 类比三问清单示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("类比三问清单示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 用类比三问深化问题理解")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    # 学员信息
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：刘明德", "团队：流程优化组", "日期：2026-07-28", "教练：赵破题"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # 核心概念
    concept = doc.add_paragraph()
    run = concept.add_run("【类比三问】")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run = concept.add_run("看到这个类比时，追问三步：")
    run.font.size = Pt(11)

    three_questions = [
        ("第一问：这个类比哪里像？", "找出表面相似的点"),
        ("第二问：这个类比哪里不像？", "找出本质差异点"),
        ("第三问：这个类比对我有什么启发？", "提炼可迁移的行动")
    ]

    for q, desc in three_questions:
        p = doc.add_paragraph()
        run = p.add_run(f"  {q}")
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(f" → {desc}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    doc.add_paragraph()

    # 案例
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、案例应用")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    case_para = doc.add_paragraph()
    run = case_para.add_run("【类比】领导说：\"流程审批就像堵塞的血管，要打通它\"")
    run.font.size = Pt(11)
    run.font.italic = True

    # 三问分析表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    from docx.oxml import OxmlElement

    headers = ["类比三问", "分析内容"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        para = table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    q1_content = "【像】\n- 堵塞：信息流动不畅\n- 血管：通道有其特定结构\n- 打通：需要疏通而非绕过\n\n【启发】审批流程确实像管道，有其固定路径"
    q2_content = "【不像】\n- 血管是自然演化，流程是人为设计\n- 血管不能随意改，流程可以优化\n- 血管堵塞会死，流程慢只会效率低\n\n【关键差异】流程是人设计的，可以重新设计"
    q3_content = "【可迁移行动】\n1. 不要只想着\"绕过\"堵点，而是重新设计流程\n2. 做流程审计：哪些节点是必要的？哪些是冗余的？\n3. 类比到组织：真正的\"血管\"是决策权在哪里"

    table.rows[1].cells[0].text = "第一问\n哪里像？"
    table.rows[1].cells[1].text = q1_content
    table.rows[2].cells[0].text = "第二问\n哪里不像？"
    table.rows[2].cells[1].text = q2_content
    table.rows[3].cells[0].text = "第三问\n启发什么？"
    table.rows[3].cells[1].text = q3_content

    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # 练习
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、练习题")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    practices = [
        ("\"团队像一支乐队\"", "第一问：哪里像？→ 各有分工、协同演奏\n第二问：哪里不像？→ 乐队有指挥，团队可能没有明确领导\n第三问：启发？→ 找到团队的\"指挥\"是关键"),
        ("\"客户是上帝\"", "第一问：哪里像？→ 要敬畏、要服务\n第二问：哪里不像？→ 上帝不需要解释，客户的信任需要积累\n第三问：启发？→ 把\"上帝\"改成\"伙伴\"更准确")
    ]

    for i, (analogy, analysis) in enumerate(practices):
        p = doc.add_paragraph()
        run = p.add_run(f"练习{i+1}：{analogy}")
        run.font.size = Pt(11)
        run.font.bold = True

        a = doc.add_paragraph()
        run = a.add_run(f"   {analysis}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo05-类比三问清单示例.docx"
    doc.save(output_path)
    print(f"demo05已生成: {output_path}")
    return output_path


def create_demo06():
    """demo06: 反面喻体法汇报示例"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("反面喻体法汇报示例")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run.font.name = "微软雅黑"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 用反面喻体增强说服力")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.italic = True
    subtitle.paragraph_format.space_after = Pt(15)

    # 学员信息
    info_table = doc.add_table(rows=1, cols=4)
    info_table.style = 'Table Grid'
    info_data = ["学员姓名：周汇报", "团队：战略市场组", "日期：2026-08-01", "教练：陈表达"]
    for j, text in enumerate(info_data):
        cell = info_table.rows[0].cells[j]
        cell.text = text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # 核心概念
    concept = doc.add_paragraph()
    run = concept.add_run("【核心概念】反面喻体法 = 先说一个错误的/反面的类比，再说正确的类比，通过对比强化理解")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    doc.add_paragraph()

    # 汇报模板
    section_title = doc.add_paragraph()
    run = section_title.add_run("一、反面喻体法汇报模板")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    template = doc.add_paragraph()
    run = template.add_run("【错误说法】")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    wrong = doc.add_paragraph()
    run = wrong.add_run("\"我们应该把客户当成上帝，提供全方位服务\"")
    run.font.size = Pt(11)
    run.font.italic = True

    vs = doc.add_paragraph()
    run = vs.add_run("【对比】")
    run.font.size = Pt(11)
    run.font.bold = True

    correct = doc.add_paragraph()
    run = correct.add_run("【正确说法】")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
    run = correct.add_run("\"把客户当成'队友'而不是'上帝'——上帝只需要敬畏，队友需要协同作战、相互成就\"")
    run.font.size = Pt(11)
    run.font.italic = True

    doc.add_paragraph()

    # 案例应用
    section_title2 = doc.add_paragraph()
    run = section_title2.add_run("二、案例：季度汇报中的应用")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # 汇报内容
    report_table = doc.add_table(rows=5, cols=2)
    report_table.style = 'Table Grid'

    from docx.oxml import OxmlElement

    headers = ["对比维度", "汇报内容"]
    for j, h in enumerate(headers):
        report_table.rows[0].cells[j].text = h
        para = report_table.rows[0].cells[j].paragraphs[0]
        para.runs[0].font.bold = True
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = report_table.rows[0].cells[j]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    report_data = [
        ["主题：渠道管理", "【反】渠道商是'合作伙伴' → 太模糊，谁主导？\n【正】渠道商是'价值放大器' → 我们的品牌是种子，渠道是阳光和水"],
        ["主题：产品定位", "【反】产品要'满足客户需求' → 正确但废话\n【正】产品不是'满足需求'而是'创造惊喜' → 需求是被满足的惊喜是超越的"],
        ["主题：团队激励", "【反】团队需要'狼性文化' → 狼性过度会破坏生态\n【正】团队需要'豹变精神' → 敏捷转型、优雅进化"],
        ["主题：变革管理", "【反】变革要'循序渐进' → 太慢！\n【正】变革要'破茧而出' → 该快时快，但尊重蝴蝶破茧的规律"]
    ]

    for i, (label, content) in enumerate(report_data):
        report_table.rows[i+1].cells[0].text = label
        report_table.rows[i+1].cells[1].text = content
        report_table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in report_table.rows[i+1].cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()

    # 使用技巧
    section_title3 = doc.add_paragraph()
    run = section_title3.add_run("三、使用技巧")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    tips = [
        ("技巧一：先破后立", "先指出常见说法的局限，再提出新的类比"),
        ("技巧二：保留核心相似性", "新类比要保留原问题的核心特征，不能为了新颖而偏离"),
        ("技巧三：控制对比篇幅", "反、正对比的比例建议 3:7，不要花太多时间在批判上")
    ]

    for title, desc in tips:
        p = doc.add_paragraph()
        run = p.add_run(f"  {title}：")
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    output_path = r"D:/新课开发/行动学习2026/破题力-思维课/课程包2-智慧版/09-成果demo/demo06-反面喻体法汇报示例.docx"
    doc.save(output_path)
    print(f"demo06已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    create_demo05()
    create_demo06()

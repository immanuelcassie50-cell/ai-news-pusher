#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build Word document templates for 双带头人破局 course.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = "D:/新课开发/党业融合/双带头人破局/完整课程包/06-工具表单"

def set_run_font(run, font_name="微软雅黑", font_size=12):
    """Set CJK font for a run"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), font_name)

def add_heading_with_font(doc, text, level=1, font_size=16, bold=True):
    """Add a heading with CJK font"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")
    return heading

def add_paragraph_with_font(doc, text="", font_size=12, bold=False):
    """Add a paragraph with CJK font"""
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.font.name = "微软雅黑"
        run.font.size = Pt(font_size)
        run.font.bold = bold
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")
    return para

def add_filled_paragraph(doc, label, value="", font_size=12):
    """Add a paragraph with label and fill-in value"""
    para = doc.add_paragraph()
    run1 = para.add_run(label)
    run1.font.name = "微软雅黑"
    run1.font.size = Pt(font_size)
    r1 = run1._r
    rPr1 = r1.get_or_add_rPr()
    rPr1.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run(value)
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(font_size)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")
    return para

def set_cell_font(cell, text, font_size=12, bold=False):
    """Set cell text with CJK font"""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

def set_table_style(table):
    """Set table borders"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def set_cell_shading(cell, fill_color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


# ===================== 文档1: 打法卡片模板 =====================
def create_da_fa_ka_pian():
    output = OUT_DIR + "/打法卡片模板.docx"
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== 封面页 =====
    doc.add_paragraph()  # spacing
    doc.add_paragraph()

    title = doc.add_heading("打法卡片", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(36)
        run.font.bold = True
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")

    subtitle = add_paragraph_with_font(doc, "从党建动作到业务成果的转化方法", font_size=18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Info fields
    para1 = add_filled_paragraph(doc, "【支部名称】________________", font_size=14)
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2 = add_filled_paragraph(doc, "【填写日期】________________", font_size=14)
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Page break
    doc.add_page_break()

    # ===== 填写说明页 =====
    heading = add_heading_with_font(doc, "如何填写打法卡片", level=1, font_size=18)

    doc.add_paragraph()

    # 打法名称
    add_filled_paragraph(doc, "打法名称: ", "", font_size=12)
    para = doc.add_paragraph()
    run = para.add_run("给这个方法取一个好记、有画面感的名字（示例：三色预警谈心法、班前十分钟微党课、党员先锋岗设备包机制）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 适用场景
    add_filled_paragraph(doc, "适用场景: ", "", font_size=12)
    para = doc.add_paragraph()
    run = para.add_run("在什么情况下使用这个打法？（示例：设备故障频发时期、新员工入职初期、项目攻坚关键阶段）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 具体步骤
    add_heading_with_font(doc, "具体步骤（3-5步，动词开头，可执行）:", level=2, font_size=14)
    steps = ["步骤1：_______________", "步骤2：_______________", "步骤3：_______________",
             "步骤4：_______________", "步骤5：_______________"]
    for step in steps:
        add_filled_paragraph(doc, step, font_size=12)

    doc.add_paragraph()

    # 关键注意事项
    add_heading_with_font(doc, "关键注意事项/踩坑提醒:", level=2, font_size=14)
    items = ["1. _______________", "2. _______________", "3. _______________"]
    for item in items:
        add_filled_paragraph(doc, item, font_size=12)

    doc.add_paragraph()

    # 效果验证方式
    add_filled_paragraph(doc, "效果验证方式: ", "", font_size=12)
    para = doc.add_paragraph()
    run = para.add_run("如何验证这个打法有效？（示例：设备故障率下降比例、客户满意度评分变化）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    # Page break
    doc.add_page_break()

    # ===== 填写示例页（制造业场景）=====
    heading = add_heading_with_font(doc, "打法卡片填写示例", level=1, font_size=18)

    doc.add_paragraph()

    # 打法名称
    add_filled_paragraph(doc, "打法名称: 党员先锋设备包机制", font_size=12)

    doc.add_paragraph()

    # 适用场景
    add_filled_paragraph(doc, "适用场景: 车间设备故障频发，非计划停机率高，严重影响产能", font_size=12)

    doc.add_paragraph()

    # 具体步骤
    add_heading_with_font(doc, "具体步骤:", level=2, font_size=14)
    steps = [
        '步骤1：党支部召开专题会议，确定"党员先锋设备包机"活动方案',
        '步骤2：划分党员设备责任区，每名党员承包2-3台关键设备',
        '步骤3：制定点检标准卡，明确点检周期和点检内容',
        '步骤4：每周主题党日进行点检情况通报，表扬先进、督促落后',
        '步骤5：每月进行设备故障率对比分析，形成台账记录'
    ]
    for step in steps:
        add_filled_paragraph(doc, step, font_size=12)

    doc.add_paragraph()

    # 关键注意事项
    add_heading_with_font(doc, "关键注意事项:", level=2, font_size=14)
    items = [
        "1. 包机责任区划分要结合党员岗位就近原则，避免形式主义",
        "2. 点检标准要简明扼要，一线员工能看懂能执行",
        "3. 通报要实事求是，不走过场"
    ]
    for item in items:
        add_filled_paragraph(doc, item, font_size=12)

    doc.add_paragraph()

    # 效果验证方式
    add_filled_paragraph(doc, "效果验证方式: ", "", font_size=12)
    para = doc.add_paragraph()
    run = para.add_run("设备故障响应时间从40分钟降至25分钟；非计划停机次数从月均8次降至3次（设备科月度报表可查）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    # Page break
    doc.add_page_break()

    # ===== 空白模板页 =====
    heading = add_heading_with_font(doc, "打法卡片（空白模板）", level=1, font_size=18)

    doc.add_paragraph()

    add_filled_paragraph(doc, "打法名称: _______________", font_size=12)

    doc.add_paragraph()

    add_filled_paragraph(doc, "适用场景: _______________", font_size=12)

    doc.add_paragraph()

    add_heading_with_font(doc, "具体步骤:", level=2, font_size=14)
    steps = ["步骤1：_______________", "步骤2：_______________", "步骤3：_______________",
             "步骤4：_______________", "步骤5：_______________"]
    for step in steps:
        add_filled_paragraph(doc, step, font_size=12)

    doc.add_paragraph()

    add_heading_with_font(doc, "关键注意事项:", level=2, font_size=14)
    items = ["1. _______________", "2. _______________", "3. _______________"]
    for item in items:
        add_filled_paragraph(doc, item, font_size=12)

    doc.add_paragraph()

    add_filled_paragraph(doc, "效果验证方式: _______________", font_size=12)

    doc.save(output)
    print(f"Created: {output}")


# ===================== 文档2: 一支部一品牌建设手册 =====================
def create_yi_zhi_bu_yi_pin_pai():
    output = OUT_DIR + "/一支部一品牌建设手册.docx"
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== 封面页 =====
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading("一支部一品牌建设手册", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(32)
        run.font.bold = True
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")

    subtitle = add_paragraph_with_font(doc, "让支部品牌成为可持续的组织资产", font_size=18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    para1 = add_filled_paragraph(doc, "【支部名称】________________", font_size=14)
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2 = add_filled_paragraph(doc, "【品牌名称】________________", font_size=14)
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== 填写说明页 =====
    heading = add_heading_with_font(doc, "品牌四件套框架", level=1, font_size=18)

    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run('品牌不是取个名字就完事，而是要形成"名称+方法论+案例库+数据支撑"四件套。本手册帮助您系统性地建设支部品牌。')
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 第一件
    add_heading_with_font(doc, "第一件：品牌名称与释义", level=2, font_size=14)
    add_filled_paragraph(doc, "品牌名称: _______________", font_size=12)
    doc.add_paragraph()
    add_filled_paragraph(doc, "品牌释义: （为什么叫这个名字？它承载了什么理念？）_______________", font_size=12)

    doc.add_paragraph()

    # 第二件
    add_heading_with_font(doc, "第二件：方法论总结（300字以内）", level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("（用简洁的语言总结这个品牌背后的方法逻辑，让别人听了能学会、能复制）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")
    doc.add_paragraph()

    # Add lines for writing
    for _ in range(5):
        doc.add_paragraph("_" * 60)

    doc.add_paragraph()

    # 第三件
    add_heading_with_font(doc, "第三件：案例库（至少3个）", level=2, font_size=14)

    # 案例1
    add_filled_paragraph(doc, "案例1标题: _______________", font_size=12)
    add_filled_paragraph(doc, "案例背景: _______________", font_size=12)
    add_filled_paragraph(doc, "具体做法: _______________", font_size=12)
    add_filled_paragraph(doc, "效果数据: _______________", font_size=12)

    doc.add_paragraph()

    # 案例2
    add_filled_paragraph(doc, "案例2标题: _______________", font_size=12)
    add_filled_paragraph(doc, "案例背景: _______________", font_size=12)
    add_filled_paragraph(doc, "具体做法: _______________", font_size=12)
    add_filled_paragraph(doc, "效果数据: _______________", font_size=12)

    doc.add_paragraph()

    # 案例3
    add_filled_paragraph(doc, "案例3标题: _______________", font_size=12)
    add_filled_paragraph(doc, "案例背景: _______________", font_size=12)
    add_filled_paragraph(doc, "具体做法: _______________", font_size=12)
    add_filled_paragraph(doc, "效果数据: _______________", font_size=12)

    doc.add_paragraph()

    # 第四件
    add_heading_with_font(doc, "第四件：数据支撑（前后对比）", level=2, font_size=14)

    # Create table
    table = doc.add_table(rows=4, cols=4)
    set_table_style(table)

    # Header row
    headers = ["指标名称", "品牌建设前", "品牌建设后", "变化幅度"]
    for i, header in enumerate(headers):
        set_cell_font(table.cell(0, i), header, font_size=12, bold=True)
        set_cell_shading(table.cell(0, i), "D9E2F3")

    # Data rows
    for row_idx in range(1, 4):
        for col_idx in range(4):
            set_cell_font(table.cell(row_idx, col_idx), "", font_size=12)

    doc.add_paragraph()

    # 品牌可持续发展设计
    add_heading_with_font(doc, "品牌可持续发展设计", level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("（如何让品牌在书记换届、人员流动后依然能延续？）")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()
    add_filled_paragraph(doc, "传承机制: _______________", font_size=12)
    add_filled_paragraph(doc, "核心人员备份: _______________", font_size=12)

    para = doc.add_paragraph()
    run1 = para.add_run("文档化程度: ")
    run1.font.name = "微软雅黑"
    run1.font.size = Pt(12)
    r1 = run1._r
    rPr1 = r1.get_or_add_rPr()
    rPr1.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("□完善  □基本完善  □待完善")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.save(output)
    print(f"Created: {output}")


# ===================== 文档3: 行动承诺书 =====================
def create_xing_dong_cheng_nuo_shu():
    output = OUT_DIR + "/行动承诺书.docx"
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== 封面页 =====
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading("30天行动承诺书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(32)
        run.font.bold = True
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")

    subtitle = add_paragraph_with_font(doc, "双带头人破局实战工作坊 · 结营仪式", font_size=16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== 正文 =====
    para = doc.add_paragraph()
    run = para.add_run("我郑重承诺，在课程结束后的30天内，我将完成以下具体行动：")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 我的基本信息
    add_heading_with_font(doc, "我的基本信息", level=2, font_size=14)
    add_filled_paragraph(doc, "姓名: _______________", font_size=12)
    add_filled_paragraph(doc, "支部名称: _______________", font_size=12)
    add_filled_paragraph(doc, "担任职务: _______________", font_size=12)

    doc.add_paragraph()

    # 我承诺的30天动作
    add_heading_with_font(doc, "我承诺的30天动作（至少填写1项，最多填写3项）", level=2, font_size=14)

    # 动作1
    doc.add_paragraph()
    add_filled_paragraph(doc, "动作1:", font_size=12)
    add_filled_paragraph(doc, "具体描述: _______________", font_size=12)
    add_filled_paragraph(doc, "预期成果: _______________", font_size=12)
    add_filled_paragraph(doc, "完成时间: _______________", font_size=12)

    # 动作2
    doc.add_paragraph()
    add_filled_paragraph(doc, "动作2:", font_size=12)
    add_filled_paragraph(doc, "具体描述: _______________", font_size=12)
    add_filled_paragraph(doc, "预期成果: _______________", font_size=12)
    add_filled_paragraph(doc, "完成时间: _______________", font_size=12)

    # 动作3
    doc.add_paragraph()
    add_filled_paragraph(doc, "动作3:", font_size=12)
    add_filled_paragraph(doc, "具体描述: _______________", font_size=12)
    add_filled_paragraph(doc, "预期成果: _______________", font_size=12)
    add_filled_paragraph(doc, "完成时间: _______________", font_size=12)

    doc.add_paragraph()

    # 我已知晓
    add_heading_with_font(doc, "我已知晓", level=2, font_size=14)
    items = [
        "1. 30天后将进行一次线上回访，核实行动完成情况",
        "2. 如因客观原因无法完成，我将主动联系课程组说明情况",
        "3. 本承诺书将作为课程学习成果的重要依据"
    ]
    for item in items:
        add_filled_paragraph(doc, item, font_size=12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名
    para1 = add_filled_paragraph(doc, "签名: _______________", font_size=12)
    para2 = add_filled_paragraph(doc, "日期: _______________", font_size=12)

    doc.save(output)
    print(f"Created: {output}")


# ===================== 文档4: 讲师巡场话术手册 =====================
def create_jiang_shi_xun_chang_hua_shu():
    output = OUT_DIR + "/讲师巡场话术手册.docx"
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ===== 封面页 =====
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading("讲师巡场话术手册", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(32)
        run.font.bold = True
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")

    subtitle = add_paragraph_with_font(doc, "双带头人破局课程 · 讲师辅助材料", font_size=16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ===== 目录 =====
    add_heading_with_font(doc, "目录", level=1, font_size=18)

    toc_items = [
        "模块一：动作盘点环节追问话术",
        "模块二：业务映射环节追问话术",
        "模块三：经验萃取环节追问话术",
        "模块四：品牌固化环节追问话术",
        "共性问题应答",
        "敏感话题处理"
    ]
    for item in toc_items:
        para = doc.add_paragraph()
        run = para.add_run(item)
        run.font.name = "微软雅黑"
        run.font.size = Pt(12)
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_page_break()

    # ===== 模块一 =====
    add_heading_with_font(doc, "模块一：动作盘点环节追问话术", level=1, font_size=16)

    para = doc.add_paragraph()
    run = para.add_run("适用场景: 学员填写党建动作清单表时，讲师巡场追问")
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    add_heading_with_font(doc, "开场话术:", level=2, font_size=14)
    opening_lines = [
        '"我看您填了三会一课，能具体说说这个月开了几次会、讨论了什么议题吗？"',
        '"您提到的谈心谈话，是主动找员工谈，还是员工找您谈比较多？"',
        '"您觉得目前党建工作中最花时间的是哪一块？"'
    ]
    for i, line in enumerate(opening_lines, 1):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_paragraph()

    add_heading_with_font(doc, "深挖话术:", level=2, font_size=14)
    digging_lines = [
        '"那次会议之后，有什么具体的后续动作吗？"',
        '"这次谈话之后，对方有什么变化吗？您观察到了什么？"',
        '"您刚才说的这个活动，有没有留下什么记录或者照片？"'
    ]
    for i, line in enumerate(digging_lines, 4):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_paragraph()

    add_heading_with_font(doc, "转化话术:", level=2, font_size=14)
    transform_lines = [
        '"如果让您用一句话总结这个月党建工作的亮点，您会怎么说？"',
        '"您觉得党建工作中哪一块对业务帮助最直接？"'
    ]
    for i, line in enumerate(transform_lines, 7):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_page_break()

    # ===== 模块二 =====
    add_heading_with_font(doc, "模块二：业务映射环节追问话术", level=1, font_size=16)

    para = doc.add_paragraph()
    run = para.add_run("适用场景: 学员进行党建动作与业务指标映射时")
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    add_heading_with_font(doc, "启发话术:", level=2, font_size=14)
    inspire_lines = [
        '"您刚才做的这个党建动作，实际解决了业务上的什么问题？"',
        '"这个活动开展之后，有没有什么可量化的变化？"',
        '"您觉得如果没有做这个动作，可能会有什么不同？"'
    ]
    for i, line in enumerate(inspire_lines, 1):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_paragraph()

    add_heading_with_font(doc, "追问话术:", level=2, font_size=14)
    followup_lines = [
        '"您提到的这个指标，有没有具体的数据可以支撑？"',
        '"这个变化是您观察到的，还是有文字记录可以证明的？"',
        '"还有没有其他的业务场景也受到了这个党建动作的影响？"'
    ]
    for i, line in enumerate(followup_lines, 4):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_page_break()

    # ===== 模块三 =====
    add_heading_with_font(doc, "模块三：经验萃取环节追问话术", level=1, font_size=16)

    para = doc.add_paragraph()
    run = para.add_run("适用场景: 学员提炼打法卡片时")
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    add_heading_with_font(doc, "引导话术:", level=2, font_size=14)
    guide_lines = [
        '"您觉得这次成功的关键是什么？是因为做了什么事情？"',
        '"如果让您把这个方法教给别的支部书记，您会怎么讲？"',
        '"在做的过程中，有没有遇到什么困难？后来是怎么解决的？"'
    ]
    for i, line in enumerate(guide_lines, 1):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_paragraph()

    add_heading_with_font(doc, "深化话术:", level=2, font_size=14)
    deepen_lines = [
        '"有没有什么需要注意的坑？其他支部书记如果学您这个方法，您会提醒他们什么？"',
        '"这个方法在什么情况下最有效？什么情况下可能效果不太好？"'
    ]
    for i, line in enumerate(deepen_lines, 4):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_page_break()

    # ===== 模块四 =====
    add_heading_with_font(doc, "模块四：品牌固化环节追问话术", level=1, font_size=16)

    para = doc.add_paragraph()
    run = para.add_run("适用场景: 学员设计支部品牌时")
    run.font.name = "微软雅黑"
    run.font.size = Pt(11)
    run.font.italic = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    add_heading_with_font(doc, "命名话术:", level=2, font_size=14)
    naming_lines = [
        '"您想给这个品牌取什么名字？这个名字有什么含义？"',
        '"这个名字好记吗？别人一听能理解你们在做什么吗？"'
    ]
    for i, line in enumerate(naming_lines, 1):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_paragraph()

    add_heading_with_font(doc, "深化话术:", level=2, font_size=14)
    deepen_lines2 = [
        '"支撑这个品牌的核心方法论是什么？能不能用一句话概括？"',
        '"有没有具体的案例可以证明这个品牌有效？"',
        '"如果换了新的书记，这个品牌还能继续吗？靠什么延续？"'
    ]
    for i, line in enumerate(deepen_lines2, 3):
        add_filled_paragraph(doc, f"{i}. {line}", font_size=11)

    doc.add_page_break()

    # ===== 共性问题应答 =====
    add_heading_with_font(doc, "共性问题应答", level=1, font_size=16)

    # 问题1
    add_heading_with_font(doc, '问题1: "我们支部人手少，事情多，实在忙不过来"', level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应答: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("理解您的压力。我们课程要解决的，不是增加工作量，而是找到党建和业务之间天然的结合点，让党建动作同时也是业务动作。请您先列出您每周花时间最多的3件事，我们一起来看看有没有可以转化的。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 问题2
    add_heading_with_font(doc, '问题2: "我们书记不重视党建，只抓业务"', level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应答: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("这种情况确实存在。我们的思路是，用业务语言来汇报党建成果，让书记看到党建是可以直接支持业务的。您觉得你们书记最关心的业务指标是什么？我们可以从那里倒推回来。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 问题3
    add_heading_with_font(doc, '问题3: "我们做过品牌，但是年年换，没有连续性"', level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应答: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("这正是我们课程要解决的核心问题。品牌不能只靠个人，必须要有机制保障。我们会教您如何设计一个可持续的传承机制，让品牌不依赖于某个人。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_page_break()

    # ===== 敏感话题处理 =====
    add_heading_with_font(doc, "敏感话题处理", level=1, font_size=16)

    para = doc.add_paragraph()
    run = para.add_run("原则: 先共情、后引导、不争论")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 话题1
    add_heading_with_font(doc, "话题1: 上级考核不合理", level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应对: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("我理解您的感受。考核确实有时候会带来压力。我们课程能帮您的，是让您有更多可以被量化、被看见的成果，让您在汇报时有话可说。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 话题2
    add_heading_with_font(doc, "话题2: 具体案例是否合规", level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应对: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("这个问题建议您直接咨询你们的组织部门，每个单位的口径可能不太一样。我们课程聚焦在方法论层面，具体执行以你们单位的正式口径为准。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.add_paragraph()

    # 话题3
    add_heading_with_font(doc, "话题3: 学员情绪发泄", level=2, font_size=14)
    para = doc.add_paragraph()
    run = para.add_run("应对: ")
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.font.bold = True
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), "微软雅黑")

    run2 = para.add_run("我能感受到您的沮丧。其实您愿意来上这个课，说明您内心还是希望把党建和业务都做好的。我们先把注意力放在方法论上，看看能实际帮到您什么。")
    run2.font.name = "微软雅黑"
    run2.font.size = Pt(12)
    r2 = run2._r
    rPr2 = r2.get_or_add_rPr()
    rPr2.set(qn('w:eastAsia'), "微软雅黑")

    doc.save(output)
    print(f"Created: {output}")


if __name__ == "__main__":
    create_da_fa_ka_pian()
    create_yi_zhi_bu_yi_pin_pai()
    create_xing_dong_cheng_nuo_shu()
    create_jiang_shi_xun_chang_hua_shu()
    print("\nAll 4 Word documents created successfully!")

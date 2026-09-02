# -*- coding: utf-8 -*-
"""
德赛西威 AI 赋能课程评审全流程 - Word 批量生成（主控接管）
批次 1：评审筹备 5 份（D-01 ~ D-05）
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === 配色（设计系统）===
COLOR_MAIN = "003D7A"
COLOR_AI = "00A0E9"
COLOR_WARN = "F37021"
COLOR_SAFE = "3CB878"
COLOR_RED = "D0021B"
COLOR_TEXT = "333333"
COLOR_BG = "F4F6F9"
COLOR_WHITE = "FFFFFF"

OUT_DIR_1 = "D:/Downloads/xinjian/德赛西威评审全流程PRD/产出物/01-评审筹备"
OUT_DIR_2 = "D:/Downloads/xinjian/德赛西威评审全流程PRD/产出物/02-评审实施"
os.makedirs(OUT_DIR_1, exist_ok=True)
os.makedirs(OUT_DIR_2, exist_ok=True)

PROJECT_NAME = "德赛西威 AI 赋能课程评审全流程"

# ==================== 工具函数 ====================

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_borders(cell, color="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{b}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tc_pr.append(tcBorders)

def set_run_font(run, name="思源黑体", size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_para(doc, text, size=10.5, bold=False, color=COLOR_TEXT, align=None, space_after=4):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p

def add_title_block(doc, title, subtitle=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLOR_MAIN)
    cell.width = Cm(16)
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    run1 = p1.add_run(title)
    set_run_font(run1, size=18, bold=True, color=COLOR_WHITE)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=10, bold=False, color=COLOR_WHITE)

def add_section_header(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("■ " + text)
        set_run_font(run, size=13, bold=True, color=COLOR_MAIN)
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("▶ " + text)
        set_run_font(run, size=11.5, bold=True, color=COLOR_MAIN)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("· " + text)
        set_run_font(run, size=10.5, bold=True, color=COLOR_AI)

def add_table_with_header(doc, headers, rows, col_widths=None, header_color=COLOR_MAIN):
    n_cols = len(headers)
    n_rows = len(rows) + 1
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_idx].width = w
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=COLOR_WHITE)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx % 2 == 0:
                set_cell_shading(cell, COLOR_BG)
            set_cell_borders(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val) if val is not None else "")
            set_run_font(run, size=9.5, color=COLOR_TEXT)

def setup_page(doc, file_id="D-XX"):
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        header = section.header
        h_table = header.add_table(rows=1, cols=2, width=Cm(15))
        h_table.autofit = False
        h_left = h_table.rows[0].cells[0]
        h_right = h_table.rows[0].cells[1]
        h_left.text = ""
        h_right.text = ""
        p_l = h_left.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_l = p_l.add_run(PROJECT_NAME)
        set_run_font(run_l, size=8.5, color=COLOR_MAIN, bold=True)
        p_r = h_right.paragraphs[0]
        p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_r = p_r.add_run(f"文件编号 {file_id}")
        set_run_font(run_r, size=8.5, color=COLOR_MAIN, bold=True)
        for cell in [h_left, h_right]:
            tc_pr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top','left','bottom','right']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tc_pr.append(tcBorders)
        footer = section.footer
        p_f = footer.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_f.add_run("德赛西威 AI 赋能课程评审全流程 | 第 ")
        set_run_font(run_f, size=8, color=COLOR_TEXT)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run_xml = run_f._element
        run_xml.append(fldChar1)
        run_xml.append(instrText)
        run_xml.append(fldChar2)
        run_g = p_f.add_run(" 页 / 共 ")
        set_run_font(run_g, size=8, color=COLOR_TEXT)
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run_xml2 = run_g._element
        run_xml2.append(fldChar3)
        run_xml2.append(instrText2)
        run_xml2.append(fldChar4)
        run_h = p_f.add_run(" 页")
        set_run_font(run_h, size=8, color=COLOR_TEXT)

def add_signature_row(doc, label="评委签名：", date_label="日期：" ):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{label}________________________     {date_label}________年____月____日")
    set_run_font(run, size=10.5, color=COLOR_TEXT)

def add_info_box(doc, text, color=COLOR_SAFE, icon="●"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{icon} {text}")
    set_run_font(run, size=10, bold=True, color=color)


# ==================== D-01 评审项目推进计划 ====================

def make_D01():
    doc = Document()
    setup_page(doc, "D-01")
    add_title_block(doc, "评审项目推进计划", "项目阶段：项目成果评审（第六阶段）")
    add_para(doc, "项目名称：德赛西威 AI 赋能课程评审全流程", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "项目周期：2026 年 X 月 X 日 - 2026 年 X 月 X 日", size=10)
    add_para(doc, "创建人：项目组    审核人：________    批准人：________", size=10)

    add_section_header(doc, "一、阶段说明", level=1)
    add_para(doc, "本阶段属于「岗位经验内化-内部课程开发与培训师培养」项目的第六阶段：项目成果评审。", size=10)
    add_para(doc, "全流程七大阶段：1.项目筹备 → 2.课程体系梳理 → 3.隐性经验显性化 → 4.显性成果标准化 → 5.标准课程生动化 → 6.项目成果评审（本阶段）→ 7.项目成果落地", size=10)
    add_para(doc, "本项目特殊之处：评价对象是「AI 赋能内训师」——评审标准必须区分「业务问题被解得怎么样」+「AI 方法论是否正确」+「信息安全合规」三个维度，区别于传统业务课。", size=10)

    add_section_header(doc, "二、阶段要点", level=1)
    pts = [
        "1. 确定成果评审的流程（上午内训师试讲+下午成果 PK+晚上结营）",
        "2. 确定成果评审专家团成员（业务方+AI 方法论+大众评审）",
        "3. 确定成果评审标准（提示词质量+AI 教练技能+陪跑数据）",
        "4. 确定项目优秀选拔标准（6 大奖项+评奖比例）",
    ]
    for pt in pts:
        add_para(doc, pt, size=10, color=COLOR_TEXT)

    add_section_header(doc, "三、具体任务（6 项）", level=1)
    headers = ["编号", "任务内容", "关键点", "负责方", "所需资源", "完成标准"]
    rows = [
        ["01", "组建评审团", "业务方 4-5 人+AI 方法论 2-3 人+大众评审 5-10 人/班+AI 陪跑", "HRBP", "见 D-03", "X 月 X 日前定名单并沟通评审标准"],
        ["02", "组建大众评审团", "同方向学员互评，作为大众评审团参与打分及代表发言", "HRBP", "学员名单", "X 月 X 日前下发通知"],
        ["03", "评选项目优秀奖", "6 大奖项：最具业务价值提示词/最佳场景化应用/最佳工具地图/最佳 AI 内训师/AI 推广卓越团队/AI 安全合规标兵", "评审组", "见 D-04", "X 月 X 号现场确定"],
        ["04", "准备评审材料", "学员提示词模板+工具地图+场景化作业+陪跑数据", "项目组", "见 D-06~D-15", "X 月 X 日前完成"],
        ["05", "准备《项目汇报 PPT》", "公司项目负责人做汇报，覆盖度+提示词库+节省时间+推广案例+优秀个人团队", "项目组", "D-18", "X 月 X 日前完成"],
        ["06", "现场物资准备", "奖品+证书+提示词打印件+工具地图打印+信息安全操作卡+AI 案例视频", "项目组", "见 D-05", "X 月 X 日前完成"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(1.2), Cm(2.5), Cm(4.5), Cm(1.5), Cm(2.5), Cm(3.5)])

    add_section_header(doc, "四、附件清单（20 份文件）", level=1)
    attachments = [
        "D-01 评审项目推进计划（本文件）",
        "D-02 评审及结营仪式流程表",
        "D-03 评审团组建方案与职责分工",
        "D-04 评审奖项设置与评奖标准",
        "D-05 评审物料清单",
        "D-06 基础班·提示词模板评分卡（每人一份）",
        "D-07 基础班·AI 工具地图评分表（每人一份）",
        "D-08 基础班·场景化应用成果 PK 表",
        "D-09 内训师班·AI 教练技能评估表（每人一份）",
        "D-10 内训师班·10 项课程包交付物检查表（每人一份）",
        "D-11 信息安全合规一票否决检查表（每人一份）",
        "D-12 评审准备自我检查表（每人一份）",
        "D-13 课程评审指引（每人一份）",
        "D-14 AI 陪跑结果评估表（每人一份）",
        "D-15 成果评审得分汇总表（电子版统计）",
        "D-16 评审日开场介绍 PPT",
        "D-17 内训师优秀课程示范 PPT",
        "D-18 项目成果汇报 PPT",
        "D-19 评审日信息门户 HTML",
        "D-20 AI 应用案例视频墙 HTML",
    ]
    for a in attachments:
        add_para(doc, "· " + a, size=9.5, color=COLOR_TEXT)

    add_signature_row(doc, "项目负责人：", "编制日期：")
    add_signature_row(doc, "审核人：", "审核日期：")

    doc.save(os.path.join(OUT_DIR_1, "D-01-评审项目推进计划.docx"))
    print("[OK] D-01")


# ==================== D-02 评审及结营仪式流程表 ====================

def make_D02():
    doc = Document()
    setup_page(doc, "D-02")
    add_title_block(doc, "评审及结营仪式流程表", "基础班当日评审 + 内训师班 D5-D6 评审 + 结营仪式")
    add_para(doc, "适用范围：德赛西威 AI 赋能课程评审全流程", size=10, bold=True, color=COLOR_MAIN)

    add_section_header(doc, "一、基础班（1 天全员课当日）评审节奏", level=1)
    add_para(doc, "基础班无单独评审日，采用「现场打分 + 课后 48 小时互评」模式。", size=10)
    headers = ["时段", "内容", "负责人", "资源需求"]
    rows = [
        ["课中 17:00-17:30", "离场三件套现场打分（提示词模板 + 工具地图）", "讲师", "评分卡 D-06 / D-07"],
        ["课后 48h 内", "同班学员互评提示词模板", "HRBP", "互评表"],
        ["课后 1 周内", "AI 陪跑组对工具地图打分", "AI 陪跑", "D-07"],
        ["课后 2-4 周", "场景化应用成果提交 + PK 评审", "评审组", "D-08 / D-14"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(3), Cm(7), Cm(2), Cm(4)])

    add_section_header(doc, "二、内训师班 D5 评审日（课程评审日）", level=1)
    add_para(doc, "评审时间：2026 年 X 月 X 日 09:00-17:00", size=10, bold=True)
    headers = ["时段", "内容", "评审团", "资源"]
    rows = [
        ["09:00-09:15", "播放项目花絮，主持人开场，介绍评审团，介绍评审规则", "全评审团", "花絮视频 + 评分表"],
        ["09:15-10:30", "内训师说课 × 5-6 人（每人 5 分钟：业务诊断/AI 方案/提示词/效果对比/可复制性）", "业务方 + AI 方法论", "课程 PPT + 评分表"],
        ["10:30-10:40", "休息 10 分钟", "/", "点心 + 饮料"],
        ["10:40-12:00", "内训师说课 × 5-6 人（续）", "业务方 + AI 方法论", "同上"],
        ["12:00-13:30", "中午休息", "/", "/"],
        ["13:30-15:00", "内训师试讲 × 5-6 人（每人 10 分钟 + AI 追问 5 分钟）", "全评审团", "PPT + 试讲设备"],
        ["15:00-15:10", "休息 10 分钟", "/", "点心 + 饮料"],
        ["15:10-16:00", "内训师试讲 × 5-6 人（续）", "全评审团", "同上"],
        ["16:00-17:00", "评委闭门评议 + 奖项初评", "闭门会议", "汇总表 D-15"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(2.5), Cm(7.5), Cm(3), Cm(3)])

    add_section_header(doc, "三、内训师班 D6 评审日（认证 + 结营）", level=1)
    add_para(doc, "评审时间：2026 年 X 月 X 日 09:00-19:30", size=10, bold=True)
    headers = ["时段", "内容", "评审团", "资源"]
    rows = [
        ["09:00-12:00", "10 项课程包文档评审（分组交叉审）", "业务方 + AI 方法论", "D-10 检查表"],
        ["12:00-13:30", "中午休息", "/", "/"],
        ["13:30-15:30", "综合评议 + 奖项评定（闭门）", "评审组", "D-15 汇总表"],
        ["15:30-16:00", "评审反馈 + 改进建议（学员旁听）", "全体内训师", "反馈表"],
        ["16:00-16:30", "休息 + 准备结营", "/", "/"],
        ["16:30-17:30", "内训师代表试讲示范 + 点评", "全评审团", "示范设备"],
        ["18:00-19:30", "结营仪式（详见下表）", "全员", "见 D-04 / D-05"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(2.5), Cm(7.5), Cm(3), Cm(3)])

    add_section_header(doc, "四、结营仪式（1.5 小时）", level=1)
    headers = ["工作流程", "负责人", "配合人", "资源需求", "时长"]
    rows = [
        ["1. 项目总结（含 5 阶段数据回顾）", "项目组", "HRBP", "D-18 汇报 PPT", "10 分钟"],
        ["2. AI 教练技能 Top 3 示范讲解", "获奖内训师", "项目组", "示范设备", "30 分钟"],
        ["3. 颁奖（6 大奖项）", "公司领导", "项目组", "D-04 奖项方案", "30 分钟"],
        ["4. 内训师代表发言", "内训师代表", "/", "话筒", "5 分钟"],
        ["5. 颁发电子证书 + 学习地图积分", "项目组", "IT 部", "证书系统", "10 分钟"],
        ["6. 领导致辞", "公司领导", "/", "/", "10 分钟"],
        ["7. 成果交付仪式（数据移交 IT 部）", "项目组", "IT 部", "成果数据", "10 分钟"],
        ["8. 合影留念", "/", "/", "摄像", "5 分钟"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(5), Cm(2), Cm(2), Cm(4.5), Cm(2)])

    add_section_header(doc, "五、应急预案", level=1)
    add_para(doc, "1. 设备故障：备用 U 盘 + 备用电脑 + 移动投影", size=10)
    add_para(doc, "2. 评委迟到：HRBP 立即联系候补评委", size=10)
    add_para(doc, "3. 学员缺席：评分表标注缺席 + 大众评审替补", size=10)
    add_para(doc, "4. 突发舆情：仅允许内部群传播，禁止外发", size=10)
    add_para(doc, "5. AI 平台故障：备用手写方案 + 截图存档", size=10)

    add_signature_row(doc, "流程编制人：")
    add_signature_row(doc, "审核人：")

    doc.save(os.path.join(OUT_DIR_1, "D-02-评审及结营仪式流程表.docx"))
    print("[OK] D-02")


# ==================== D-03 评审团组建方案与职责分工 ====================

def make_D03():
    doc = Document()
    setup_page(doc, "D-03")
    add_title_block(doc, "评审团组建方案与职责分工", "4 类评审团 + 权重设计 + 评委对齐会")

    add_section_header(doc, "一、评审团总体构成", level=1)
    add_para(doc, "本项目评审团由 4 类构成：业务方专家 + AI 方法论专家 + 大众评审 + AI 陪跑数据。", size=10)

    headers = ["类别", "人数", "来源", "主要评什么", "打分权重"]
    rows = [
        ["业务方专家", "4-5 人", "项目管理/通用管理/专业职能/测试/开发 5 方向各 1", "业务问题被解得怎么样", "40%"],
        ["AI 方法论专家", "2-3 人", "外部 AI 讲师 + HRBP + IT 安全", "AI 方法论是否正确 + 安全合规", "30%"],
        ["大众评审（同事）", "5-10 人/班", "学员同方向同事", "这个提示词我能不能用", "20%"],
        ["AI 陪跑数据", "系统", "课后 2-4 周应用数据", "实际产生了多少业务价值", "10%"],
        ["合计", "≥ 12 人/班", "—", "—", "100%"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(2.5), Cm(1.8), Cm(4.5), Cm(4.5), Cm(2)])

    add_section_header(doc, "二、专家评审团-业务方（4-5 人）", level=1)
    add_para(doc, "职责：", size=10, bold=True, color=COLOR_MAIN)
    for r in [
        "· 评审提示词模板的业务还原度、价值可衡量性、可推广性",
        "· 评审场景化应用成果是否真解决了业务痛点",
        "· 评审内训师课程包对业务部门的实际价值",
        "· 在大众评审中代表业务方立场发声",
    ]:
        add_para(doc, r, size=10)
    add_para(doc, "推荐来源：", size=10, bold=True, color=COLOR_MAIN)
    for s in [
        "· 项目管理部门负责人（1 人）",
        "· 通用管理 / 部门经理代表（1 人）",
        "· 专业职能总监（财务/HR/法务/客服 任选 1 人）",
        "· 测试 / 开发技术总监（任选 1 人）",
        "· 业务一把手或 HR 一把手（必选 1 人作为组长）",
    ]:
        add_para(doc, s, size=10)

    add_section_header(doc, "三、专家评审团-AI 方法论（2-3 人）", level=1)
    add_para(doc, "职责：", size=10, bold=True, color=COLOR_MAIN)
    for r in [
        "· 评审提示词结构（四段式）的完整性、迭代次数、复用人次",
        "· 评审 AI 工具地图的判断力、组合能力、平台定位讲解",
        "· 评审内训师的 AI 教练技能 20 项（4 大类）",
        "· 评审信息安全合规：是否触碰红灯项",
        "· 给出 AI 追问问题（试讲环节的 5 分钟）",
    ]:
        add_para(doc, r, size=10)
    add_para(doc, "推荐来源：", size=10, bold=True, color=COLOR_MAIN)
    for s in [
        "· 外部 AI 讲师（项目总教练）",
        "· 内部 HRBP（懂业务 + 懂 AI）",
        "· IT 安全负责人（信息安全合规把关）",
    ]:
        add_para(doc, s, size=10)

    add_section_header(doc, "四、大众评审团（5-10 人/班）", level=1)
    add_para(doc, "职责：", size=10, bold=True, color=COLOR_MAIN)
    for r in [
        "· 对基础班学员的提示词模板打分（评分卡 D-06）",
        "· 对内训师的试讲打分（试讲环节大众评审只看不评）",
        "· 推选 1 名代表在结营仪式发言",
        "· 反馈「这个提示词我能不能用」的真实体验",
    ]:
        add_para(doc, r, size=10)
    add_para(doc, "产生方式：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· 基础班：在课前由同班学员推选 5-10 人/班", size=10)
    add_para(doc, "· 内训师班：邀请同岗位同事旁听 D5 下午试讲", size=10)

    add_section_header(doc, "五、AI 陪跑数据（系统）", level=1)
    add_para(doc, "数据来源：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· 课后 2-4 周，IT 部门调取学员在数智小西 / 外部工具的使用数据", size=10)
    add_para(doc, "· 学员周报：每周提交「提示词应用次数 + 业务节省时间 + 业务产出数量」", size=10)
    add_para(doc, "· 同事复用：每邀请同事复用 1 次 +1 分", size=10)
    add_info_box(doc, "重要：陪跑数据需 IT 部门配合调取，避免学员自报虚高", color=COLOR_WARN)

    add_section_header(doc, "六、评委对齐会（D5 评审日前 30 分钟）", level=1)
    add_para(doc, "评审标准对齐是确保打分公平的关键。", size=10, bold=True)
    for a in [
        "1. 评委逐条过 5 档评分标准（24-25/21-23/18-20/15-17/0-14）",
        "2. 评委试评 1 份样卷，对齐打分尺度",
        "3. 评委确认一票否决项的判定标准（信息安全红灯）",
        "4. 评委确认 AI 追问话术（5 问清单）",
        "5. 评委确认大众评审的统计方式（去掉最高最低取平均）",
    ]:
        add_para(doc, a, size=10)

    add_section_header(doc, "七、评委纪律", level=1)
    add_para(doc, "1. 评委不得私下与学员交流打分倾向", size=10)
    add_para(doc, "2. 评委打分需签名，签过的评分表不得修改", size=10)
    add_para(doc, "3. 评委需准时参加 30 分钟对齐会，无故缺席取消资格", size=10)
    add_para(doc, "4. 评委打分若超出合理区间（如全打 25 或全打 0），HRBP 提醒", size=10)

    add_signature_row(doc, "评审团组建负责人：")
    add_signature_row(doc, "HRBP 审核：")

    doc.save(os.path.join(OUT_DIR_1, "D-03-评审团组建方案与职责分工.docx"))
    print("[OK] D-03")


# ==================== D-04 评审奖项设置与评奖标准 ====================

def make_D04():
    doc = Document()
    setup_page(doc, "D-04")
    add_title_block(doc, "评审奖项设置与评奖标准", "6 大奖项 + 评奖依据 + 颁奖流程")

    add_section_header(doc, "一、奖项设置总览", level=1)
    headers = ["奖项", "数量", "评选依据", "价值"]
    rows = [
        ["最具业务价值提示词奖", "1-2 个", "提示词评分 Top 1-2", "入选公司级 AI 应用案例库"],
        ["最佳场景化应用奖", "1-2 个", "场景化作业综合 Top", "公司内网首页推荐 + 周报专栏"],
        ["最具推广价值工具地图奖", "1 个", "工具地图 + 同事复用次数综合", "入选新员工入职必修"],
        ["最佳 AI 内训师奖", "若干", "内训师班综合评分 Top", "颁发证书 + 优先外派学习"],
        ["AI 推广卓越团队奖", "1 组", "部门整体参与度 + 应用率", "部门绩效加分"],
        ["AI 安全合规标兵奖", "1-2 个", "连续 4 周零违规 + 主动发现风险", "颁发证书 + 安全积分"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(4.5), Cm(1.5), Cm(4.5), Cm(5)])

    add_section_header(doc, "二、各奖项详细评奖标准", level=1)

    add_section_header(doc, "1. 最具业务价值提示词奖", level=2)
    add_para(doc, "评奖对象：基础班学员", size=10)
    add_para(doc, "评奖公式：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "总得分 = 提示词模板评分 × 50% + 工具地图评分 × 10% + 场景化作业 × 20% + AI 陪跑 × 20%", size=10)
    add_para(doc, "评奖门槛：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· 提示词模板得分 ≥ 21（较好及以上）", size=10)
    add_para(doc, "· 信息安全合规：无红灯", size=10)
    add_para(doc, "· 至少 1 个提示词被同岗位 ≥ 2 人复用", size=10)

    add_section_header(doc, "2. 最佳场景化应用奖", level=1)
    add_para(doc, "评奖对象：基础班学员（课后 2-4 周提交）", size=10)
    add_para(doc, "评奖标准（5 维 × 25 分制）：", size=10, bold=True, color=COLOR_MAIN)
    headers2 = ["维度", "权重", "评分要点"]
    rows2 = [
        ["业务问题真不真", "30%", "是否是真实工作场景的痛点"],
        ["AI 方案巧不巧", "25%", "用 AI 解决的方式是否巧妙"],
        ["效果数据硬不硬", "20%", "是否有节省时间 / 提升质量的量化数据"],
        ["可复制性强不强", "15%", "换场景 / 换人能否用"],
        ["同事口碑好不好", "10%", "被同事复用的频次"],
    ]
    add_table_with_header(doc, headers2, rows2, col_widths=[Cm(4), Cm(2), Cm(10)])

    add_section_header(doc, "3. 最具推广价值工具地图奖", level=1)
    add_para(doc, "评奖对象：基础班学员", size=10)
    add_para(doc, "评奖标准：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· 工具地图评分 ≥ 21", size=10)
    add_para(doc, "· 至少 5 名同岗位同事引用其工具组合", size=10)
    add_para(doc, "· 工具组合的「内部平台 + 外部工具」配比合理", size=10)

    add_section_header(doc, "4. 最佳 AI 内训师奖", level=1)
    add_para(doc, "评奖对象：内训师班学员", size=10)
    add_para(doc, "评奖公式：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "总得分 = 10 项课程包 × 40% + AI 教练技能 × 40% + AI 陪跑 × 20%", size=10)
    add_para(doc, "评奖门槛：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· AI 教练技能 20 项平均分 ≥ 21", size=10)
    add_para(doc, "· 10 项课程包全部完成（无空白项）", size=10)
    add_para(doc, "· 试讲环节大众评审打分 ≥ 21", size=10)

    add_section_header(doc, "5. AI 推广卓越团队奖", level=1)
    add_para(doc, "评奖对象：以部门 / 项目组为单位的团队", size=10)
    add_para(doc, "评奖标准（3 维）：", size=10, bold=True, color=COLOR_MAIN)
    headers3 = ["维度", "权重", "数据来源"]
    rows3 = [
        ["参与度", "40%", "部门学员参训率（基础班 + 内训师班）"],
        ["应用率", "40%", "课后 2-4 周部门内提示词应用总次数"],
        ["贡献度", "20%", "部门成员贡献的提示词被外部门引用次数"],
    ]
    add_table_with_header(doc, headers3, rows3, col_widths=[Cm(3), Cm(2), Cm(11)])

    add_section_header(doc, "6. AI 安全合规标兵奖", level=1)
    add_para(doc, "评奖对象：个人", size=10)
    add_para(doc, "评奖标准：", size=10, bold=True, color=COLOR_MAIN)
    add_para(doc, "· 连续 4 周（陪跑期）信息安全零违规", size=10)
    add_para(doc, "· 主动发现并上报 ≥ 1 次潜在安全风险", size=10)
    add_para(doc, "· 信息安全操作卡使用率 100%", size=10)

    add_section_header(doc, "三、颁奖流程", level=1)
    flow = [
        "1. 评审组闭门会议（13:30-15:30 D6）确定获奖名单",
        "2. HRBP 审核名单合规性（含一票否决项排查）",
        "3. 公司领导颁奖（18:30 结营仪式）",
        "4. 获奖者感言（每奖 1-2 人，30 秒 / 人）",
        "5. 颁发电子证书（同步到学习地图积分系统）",
        "6. 合影留念",
    ]
    for f in flow:
        add_para(doc, f, size=10)

    add_section_header(doc, "四、奖品配置", level=1)
    headers4 = ["奖项等级", "奖品", "颁奖人"]
    rows4 = [
        ["一等奖（最具业务价值提示词 / 最佳场景化应用）", "证书 + 智能音箱 + 学习地图 2000 积分", "公司一把手"],
        ["二等奖（最佳工具地图 / 最佳 AI 内训师）", "证书 + AI 课程年卡 + 学习地图 1500 积分", "HR 一把手"],
        ["三等奖（AI 安全合规标兵）", "证书 + 安全积分 500 + 学习地图 800 积分", "IT 负责人"],
        ["团队奖（AI 推广卓越团队）", "团队锦旗 + 团队建设基金 5000 元", "公司一把手"],
    ]
    add_table_with_header(doc, headers4, rows4, col_widths=[Cm(6), Cm(7), Cm(3)])

    add_signature_row(doc, "奖项方案编制人：")
    add_signature_row(doc, "审核人：")

    doc.save(os.path.join(OUT_DIR_1, "D-04-评审奖项设置与评奖标准.docx"))
    print("[OK] D-04")


# ==================== D-05 评审物料清单 ====================

def make_D05():
    doc = Document()
    setup_page(doc, "D-05")
    add_title_block(doc, "评审物料清单", "26 项物料 × 数量 × 负责方 × 完成情况")
    add_para(doc, "编制人：项目组    编制日期：2026 年 X 月 X 日", size=10)

    add_section_header(doc, "一、基础物料（评审现场通用）", level=1)
    headers = ["序号", "物料", "数量", "负责方", "负责人", "完成情况"]
    rows = [
        ["1", "横幅（「德赛西威 AI 赋能课程评审会」）", "1 条", "项目组", "_____", "[ ]"],
        ["2", "评审及结营仪式流程表（D-02）", "按学员 + 嘉宾人数", "项目组", "_____", "[ ]"],
        ["3", "签到表（学员 + 嘉宾 + 评委三栏）", "各 1 份", "HRBP", "_____", "[ ]"],
        ["4", "学员席卡", "按学员人数", "HRBP", "_____", "[ ]"],
        ["5", "嘉宾席卡", "按嘉宾人数", "HRBP", "_____", "[ ]"],
        ["6", "评委席卡（标注：业务方 / AI 方法论 / 大众评审）", "按评委人数", "HRBP", "_____", "[ ]"],
        ["7", "白板", "1", "项目组", "_____", "[ ]"],
        ["8", "白板纸", "每组 8 份", "项目组", "_____", "[ ]"],
        ["9", "白板笔（黑 + 红 + 蓝三色）", "每组 2 套", "项目组", "_____", "[ ]"],
        ["10", "投影仪 + 备用机", "1 + 1", "IT 部", "_____", "[ ]"],
        ["11", "摄像机 + 三脚架", "1", "项目组", "_____", "[ ]"],
        ["12", "照相机", "1", "项目组", "_____", "[ ]"],
        ["13", "翻页器（备用）", "2", "IT 部", "_____", "[ ]"],
        ["14", "激光笔", "2", "IT 部", "_____", "[ ]"],
    ]
    add_table_with_header(doc, headers, rows, col_widths=[Cm(1.2), Cm(5.5), Cm(2.5), Cm(1.8), Cm(2.5), Cm(2.5)])

    add_section_header(doc, "二、AI 课程专属物料（德赛西威项目特有）", level=1)
    rows2 = [
        ["15", "学员提示词模板打印件（脱敏后）", "每人 5 张", "项目组", "_____", "[ ]"],
        ["16", "AI 工具地图打印件", "每人 1 张", "项目组", "_____", "[ ]"],
        ["17", "信息安全操作卡（红黄绿灯）", "每人 1 张", "项目组", "_____", "[ ]"],
        ["18", "LangGPT 提示词模板卡（角色 / 背景 / 目标 / 约束）", "每人 1 张", "项目组", "_____", "[ ]"],
        ["19", "41 个场景清单手册", "每人 1 本", "项目组", "_____", "[ ]"],
        ["20", "AI 案例视频（陪跑期优秀案例）", "3-5 段", "项目组", "_____", "[ ]"],
        ["21", "评分表 D-06~D-15 全部打印", "按评委人数 × 课题数", "项目组", "_____", "[ ]"],
    ]
    add_table_with_header(doc, headers, rows2, col_widths=[Cm(1.2), Cm(5.5), Cm(2.5), Cm(1.8), Cm(2.5), Cm(2.5)])

    add_section_header(doc, "三、颁奖物料", level=1)
    rows3 = [
        ["22", "获奖证书（电子 + 纸质）", "按奖项数量", "HRBP", "_____", "[ ]"],
        ["23", "奖杯（团队奖）", "1", "HRBP", "_____", "[ ]"],
        ["24", "奖品（智能音箱 / AI 课程年卡 / 学习地图积分）", "按奖项等级", "HRBP", "_____", "[ ]"],
    ]
    add_table_with_header(doc, headers, rows3, col_widths=[Cm(1.2), Cm(5.5), Cm(2.5), Cm(1.8), Cm(2.5), Cm(2.5)])

    add_section_header(doc, "四、应急物料", level=1)
    rows4 = [
        ["25", "备用 U 盘（所有 PPT + 评分表备份）", "2", "IT 部", "_____", "[ ]"],
        ["26", "备用电脑 + 电源线 + 转接头", "1 套", "IT 部", "_____", "[ ]"],
    ]
    add_table_with_header(doc, headers, rows4, col_widths=[Cm(1.2), Cm(5.5), Cm(2.5), Cm(1.8), Cm(2.5), Cm(2.5)])

    add_section_header(doc, "五、物料检查清单（评审日前 1 天）", level=1)
    for c in [
        "1. 所有打印件到位（提示词 / 工具地图 / 操作卡 / 模板卡）",
        "2. 所有电子设备测试通过（投影 / 电脑 / 翻页器 / 激光笔）",
        "3. 所有评委到位确认（短信 + 邮件双通知）",
        "4. 学员座位图张贴完成",
        "5. 应急 U 盘 + 备用电脑到位",
        "6. 奖品 + 证书到位并分类摆放",
    ]:
        add_para(doc, c, size=10)

    add_signature_row(doc, "物料总负责人：")
    add_signature_row(doc, "HRBP 验收：")

    doc.save(os.path.join(OUT_DIR_1, "D-05-评审物料清单.docx"))
    print("[OK] D-05")


# ==================== 执行 ====================

if __name__ == "__main__":
    print("=== 开始生成 Word 批次 1（评审筹备 5 份）===")
    make_D01()
    make_D02()
    make_D03()
    make_D04()
    make_D05()
    print("=== 批次 1 完成 ===")

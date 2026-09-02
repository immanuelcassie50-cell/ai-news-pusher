#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Build 5 Word document templates for Joint Delivery Mode Design course.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = "D:/新课开发/变革管理/15-联合交付模式设计：技术公司与管理咨询公司怎么配合/完整课程包/06-工具表单"

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_header_row(table, headers, bg_color="1F3864"):
    """Create a header row with dark blue background"""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

def set_table_style(table):
    """Apply professional styling to table"""
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def add_title(doc, title, subtitle=""):
    """Add a styled title"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)

    if subtitle:
        para2 = doc.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(128, 128, 128)

def add_section_heading(doc, text, level=1):
    """Add a section heading"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(31, 56, 100)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(47, 84, 150)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

def add_info_field(doc, label, value=""):
    """Add an info field row"""
    para = doc.add_paragraph()
    run_label = para.add_run(label + ": ")
    run_label.font.bold = True
    run_label.font.size = Pt(10)
    run_value = para.add_run(value)
    run_value.font.size = Pt(10)

# ===================== W01: 联合交付合作协议模板 =====================
def create_w01():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title
    add_title(doc, "联合交付合作协议", "Joint Delivery Cooperation Agreement")
    doc.add_paragraph()

    # Info fields
    add_info_field(doc, "甲方（技术公司）", "______________________________")
    add_info_field(doc, "乙方（管理咨询公司）", "______________________________")
    add_info_field(doc, "项目名称", "______________________________")
    add_info_field(doc, "协议签订日期", "______________________________")
    doc.add_paragraph()

    # Article 1: Scope of Cooperation
    add_section_heading(doc, "第一条 合作范围 (Article 1: Scope of Cooperation)")
    para = doc.add_paragraph()
    para.add_run("1.1 双方同意按照约定的联合交付模式，共同为客户提供服务。")
    para = doc.add_paragraph()
    para.add_run("1.2 具体服务范围、交付内容和质量标准见附件《接口设计规划表》。")
    para = doc.add_paragraph()
    para.add_run("1.3 双方各自负责的技术/咨询工作范围如下：")

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    create_header_row(table, ["工作类别", "技术公司负责", "咨询公司负责"])
    row1 = table.rows[1].cells
    row1[0].text = "核心工作"
    row1[1].text = ""
    row1[2].text = ""
    row2 = table.rows[2].cells
    row2[0].text = "辅助工作"
    row2[1].text = ""
    row2[2].text = ""
    set_table_style(table)
    doc.add_paragraph()

    # Article 2: Responsibility Boundaries
    add_section_heading(doc, "第二条 责任边界 (Article 2: Responsibility Boundaries)")
    para = doc.add_paragraph()
    para.add_run("2.1 技术公司对技术方案的实现和质量负责。")
    para = doc.add_paragraph()
    para.add_run("2.2 咨询公司对变革管理的方案设计和推进效果负责。")
    para = doc.add_paragraph()
    para.add_run("2.3 因一方原因导致的损失，由责任方承担相应责任。")
    para = doc.add_paragraph()
    para.add_run("2.4 双方共同造成的问题，由双方协商承担相应责任。")

    # Article 3: Handoff Mechanisms
    add_section_heading(doc, "第三条 交接机制 (Article 3: Handoff Mechanisms)")
    para = doc.add_paragraph()
    para.add_run("3.1 双方应按照《接口设计规划表》中定义的里程碑节点进行工作交接。")
    para = doc.add_paragraph()
    para.add_run("3.2 每个交接点应完成《交接文档模板》规定的交付物和验收标准。")
    para = doc.add_paragraph()
    para.add_run("3.3 交接过程中发现的问题，应在48小时内书面反馈。")

    # Article 4: Language Conventions
    add_section_heading(doc, "第四条 话术约定 (Article 4: Language Conventions)")
    para = doc.add_paragraph()
    para.add_run("4.1 双方应统一使用《话术统一对照表》中定义的术语和表达方式。")
    para = doc.add_paragraph()
    para.add_run("4.2 对客户沟通时，双方应使用客户熟悉的语言体系，避免技术术语与管理术语混用。")
    para = doc.add_paragraph()
    para.add_run("4.3 对外发言应按照《角色分工说明书》规定的范围进行。")

    # Article 5: Profit Sharing
    add_section_heading(doc, "第五条 利益分配 (Article 5: Profit Sharing)")
    para = doc.add_paragraph()
    para.add_run("5.1 项目收益分配比例按照《利益分配协议模板》的规定执行。")
    para = doc.add_paragraph()
    para.add_run("5.2 收益分配应在每笔款项到账后30个工作日内完成结算。")
    para = doc.add_paragraph()
    para.add_run("5.3 双方应各自承担己方的税费和其他法定费用。")

    # Article 6: Dispute Resolution
    add_section_heading(doc, "第六条 争议解决 (Article 6: Dispute Resolution)")
    para = doc.add_paragraph()
    para.add_run("6.1 因本协议引起的或与本协议有关的任何争议，双方应首先通过友好协商解决。")
    para = doc.add_paragraph()
    para.add_run("6.2 协商不成的，任何一方可向双方共同认可的有管辖权的人民法院提起诉讼。")
    para = doc.add_paragraph()
    para.add_run("6.3 诉讼过程中，双方应继续履行本协议中无争议的部分。")

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("甲方（技术公司）：______________________________    日期：________________")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("乙方（管理咨询公司）：________________________    日期：________________")

    output_path = OUTPUT_DIR + "/W01 - 联合交付合作协议模板.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

# ===================== W02: 交接文档模板 =====================
def create_w02():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_title(doc, "交接文档", "Handoff Document")
    doc.add_paragraph()

    # Header info
    add_info_field(doc, "交接项目名称", "______________________________")
    add_info_field(doc, "交接日期", "______________________________")
    add_info_field(doc, "交接方（转出方）", "______________________________")
    add_info_field(doc, "接收方（转入方）", "______________________________")
    doc.add_paragraph()

    # Section 1: Deliverables List
    add_section_heading(doc, "一、交付物清单 (Deliverables List)")

    table = doc.add_table(rows=7, cols=5)
    table.style = 'Table Grid'
    create_header_row(table, ["编号", "交付物名称", "交接状态", "交接数量", "备注"])
    for i in range(1, 7):
        row = table.rows[i].cells
        row[0].text = f"D{i}"
        row[2].text = "□已完成 □未完成 □部分"
    set_table_style(table)
    doc.add_paragraph()

    # Section 2: Acceptance Criteria
    add_section_heading(doc, "二、验收标准 (Acceptance Criteria)")

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    create_header_row(table2, ["编号", "验收标准", "验收方式", "验收结果"])
    for i in range(1, 6):
        row = table2.rows[i].cells
        row[0].text = f"A{i}"
        row[2].text = "□文件审核 □现场验收 □演示"
        row[3].text = "□通过 □未通过"
    set_table_style(table2)
    doc.add_paragraph()

    # Section 3: Precautions
    add_section_heading(doc, "三、注意事项 (Precautions)")
    items = [
        "1. 交接时应确认所有文档的完整性和准确性",
        "2. 如有问题，应在交接时立即提出",
        "3. 交接完成后，转出方仍需保留相关文档副本至少6个月",
        "4. 涉及客户敏感信息的内容应按照保密协议处理",
        "5. 其他需要特别说明的事项：______________________________"
    ]
    for item in items:
        para = doc.add_paragraph(item)
        para.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()

    # Section 4: History Records
    add_section_heading(doc, "四、历史记录 (History Records)")

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'
    create_header_row(table3, ["日期", "事项", "处理人", "结果"])
    for i in range(1, 4):
        row = table3.rows[i].cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""
    set_table_style(table3)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    para = doc.add_paragraph()
    para.add_run("交接方签字：________________________    日期：________________")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("接收方签字：________________________    日期：________________")

    output_path = OUTPUT_DIR + "/W02 - 交接文档模板.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

# ===================== W03: 每周同步报告模板 =====================
def create_w03():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_title(doc, "每周同步报告", "Weekly Sync Report")
    doc.add_paragraph()

    # Header info
    add_info_field(doc, "报告周期", "____年____月____日至____月____日")
    add_info_field(doc, "报告人", "______________________________")
    add_info_field(doc, "所属公司", "______________________________")
    doc.add_paragraph()

    # Section 1: This Week's Progress
    add_section_heading(doc, "一、本周进展 (This Week's Progress)")

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    create_header_row(table, ["工作项", "计划完成", "实际完成", "状态"])
    for i in range(1, 6):
        row = table.rows[i].cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        row[3].text = "□已完成 □进行中 □延期"
    set_table_style(table)
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run("本周主要成果：")
    para = doc.add_paragraph()
    para.add_run("____________________________________________________________")
    doc.add_paragraph()

    # Section 2: Next Week's Plan
    add_section_heading(doc, "二、下周计划 (Next Week's Plan)")

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    create_header_row(table2, ["工作项", "计划完成日期", "责任人", "配合需求"])
    for i in range(1, 6):
        row = table2.rows[i].cells
        row[0].text = ""
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""
    set_table_style(table2)
    doc.add_paragraph()

    # Section 3: Risk Warning
    add_section_heading(doc, "三、风险预警 (Risk Warning)")

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'
    create_header_row(table3, ["风险编号", "风险描述", "影响程度", "应对措施"])
    for i in range(1, 4):
        row = table3.rows[i].cells
        row[0].text = f"R{i}"
        row[1].text = ""
        row[2].text = "□高 □中 □低"
        row[3].text = ""
    set_table_style(table3)
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run("无风险预警事项：□ 是（本期无新增风险）    □ 否（见上表）")
    doc.add_paragraph()

    # Section 4: Resource Needs
    add_section_heading(doc, "四、资源需求 (Resource Needs)")

    table4 = doc.add_table(rows=4, cols=4)
    table4.style = 'Table Grid'
    create_header_row(table4, ["资源类型", "需求描述", "紧急程度", "申请理由"])
    for i in range(1, 4):
        row = table4.rows[i].cells
        row[0].text = "□人力 □技术 □其他"
        row[1].text = ""
        row[2].text = "□高 □中 □低"
        row[3].text = ""
    set_table_style(table4)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    para = doc.add_paragraph()
    para.add_run("报告人：________________________    日期：________________")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("审阅人：________________________    日期：________________")

    output_path = OUTPUT_DIR + "/W03 - 每周同步报告模板.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

# ===================== W04: 角色分工说明书模板 =====================
def create_w04():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_title(doc, "角色分工说明书", "Roles and Responsibilities Document")
    doc.add_paragraph()

    # Header info
    add_info_field(doc, "项目名称", "______________________________")
    add_info_field(doc, "文件版本", "______________________________")
    add_info_field(doc, "生效日期", "______________________________")
    doc.add_paragraph()

    # RACI Matrix
    add_section_heading(doc, "一、RACI矩阵 (RACI Matrix)")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("R = Responsible（负责）  A = Accountable（问责）  C = Consulted（咨询）  I = Informed（知会）")
    para.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=9, cols=6)
    table.style = 'Table Grid'
    create_header_row(table, ["工作项", "技术公司\n项目经理", "咨询公司\n项目经理", "技术\n接口人", "管理\n接口人", "联合\n项目总监"])

    tasks = [
        ("技术方案设计", "R/A", "C", "R", "I", "A"),
        ("变革方案设计", "C", "R/A", "I", "R", "A"),
        ("客户沟通-技术", "R/A", "I", "R", "C", "A"),
        ("客户沟通-管理", "I", "R/A", "C", "R", "A"),
        ("里程碑评审", "R", "R", "C", "C", "A"),
        ("风险管理", "R", "R", "I", "I", "A"),
        ("进度汇报", "R", "R", "C", "C", "A"),
        ("争议处理", "C", "C", "I", "I", "A")
    ]

    for i, task in enumerate(tasks):
        row = table.rows[i+1]
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F2F2F2")
        row.cells[0].text = task[0]
        row.cells[1].text = task[1]
        row.cells[2].text = task[2]
        row.cells[3].text = task[3]
        row.cells[4].text = task[4]
        row.cells[5].text = task[5]
    set_table_style(table)
    doc.add_paragraph()

    # Responsibilities Detail
    add_section_heading(doc, "二、职责详细说明 (Responsibilities Detail)")

    # Tech Company PM
    add_section_heading(doc, "技术公司项目经理", level=2)
    items_tech = [
        "负责技术方案的完整性和可实现性",
        "管理技术团队的工作进度和质量",
        "与咨询公司项目经理协调技术接口事宜",
        "代表技术方参与客户技术层面的沟通",
        "对技术交付物负有直接责任"
    ]
    for item in items_tech:
        para = doc.add_paragraph("• " + item)
        para.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    # Consulting Company PM
    add_section_heading(doc, "咨询公司项目经理", level=2)
    items_consult = [
        "负责变革管理方案的设计和有效性",
        "管理咨询团队的工作进度和效果",
        "与技术公司项目经理协调管理接口事宜",
        "代表咨询方参与客户管理层面的沟通",
        "对变革交付物负有直接责任"
    ]
    for item in items_consult:
        para = doc.add_paragraph("• " + item)
        para.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    # Joint Project Director
    add_section_heading(doc, "联合项目总监", level=2)
    items_joint = [
        "对项目整体成功负有最终责任",
        "裁决双方之间的争议和分歧",
        "负责重大客户关系的维护",
        "审批超出双方项目经理权限的决策",
        "监督项目整体进度和质量"
    ]
    for item in items_joint:
        para = doc.add_paragraph("• " + item)
        para.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

    # Interface Contacts
    add_section_heading(doc, "三、接口人指定 (Interface Contacts)")

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    create_header_row(table2, ["角色", "姓名", "联系电话", "Email"])
    roles = ["技术公司接口人", "咨询公司接口人", "技术公司backup", "咨询公司backup"]
    for i, role in enumerate(roles):
        table2.rows[i+1].cells[0].text = role
    set_table_style(table2)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    para = doc.add_paragraph()
    para.add_run("技术公司项目经理：________________    日期：________________")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("咨询公司项目经理：________________    日期：________________")

    output_path = OUTPUT_DIR + "/W04 - 角色分工说明书模板.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

# ===================== W05: 冲突处理记录表 =====================
def create_w05():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    add_title(doc, "冲突处理记录表", "Conflict Handling Record")
    doc.add_paragraph()

    # Header info
    add_info_field(doc, "记录编号", "CHR-____-____")
    add_info_field(doc, "发生日期", "______________________________")
    add_info_field(doc, "记录人", "______________________________")
    doc.add_paragraph()

    # Section 1: Conflict Description
    add_section_heading(doc, "一、冲突描述 (Conflict Description)")

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "冲突标题"
    table.rows[0].cells[1].text = ""
    table.rows[1].cells[0].text = "冲突类型"
    table.rows[1].cells[1].text = "□技术争议 □管理争议 □资源争议 □进度争议 □责任争议 □其他"
    table.rows[2].cells[0].text = "涉及方"
    table.rows[2].cells[1].text = "□技术公司 □咨询公司 □双方"
    table.rows[3].cells[0].text = "冲突描述"
    table.rows[3].cells[1].text = ""
    for i in range(3, 6):
        table.rows[i].cells[1].text = ""
    set_cell_shading(table.rows[0].cells[0], "D6E4F7")
    set_cell_shading(table.rows[1].cells[0], "D6E4F7")
    set_cell_shading(table.rows[2].cells[0], "D6E4F7")
    set_cell_shading(table.rows[3].cells[0], "D6E4F7")
    set_table_style(table)
    doc.add_paragraph()

    # Section 2: Root Cause Analysis
    add_section_heading(doc, "二、原因分析 (Root Cause Analysis)")

    para = doc.add_paragraph()
    para.add_run("根本原因：")
    for i in range(3):
        para = doc.add_paragraph()
        para.add_run(f"原因{i+1}：________________________________________________")
        para.paragraph_format.space_after = Pt(3)

    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Table Grid'
    create_header_row(table2, ["分析维度", "分析内容", "结论"])
    dims = ["沟通层面", "流程层面", "利益层面"]
    for i, dim in enumerate(dims):
        table2.rows[i+1].cells[0].text = dim
    set_table_style(table2)
    doc.add_paragraph()

    # Section 3: Solution
    add_section_heading(doc, "三、解决方案 (Solution)")

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Table Grid'
    create_header_row(table3, ["方案选项", "具体措施", "责任方", "完成期限"])
    for i in range(1, 4):
        table3.rows[i].cells[0].text = f"方案{i}"
    set_table_style(table3)
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run("最终选择方案：□方案1 □方案2 □方案3 □综合方案")
    para = doc.add_paragraph()
    para.add_run("方案说明：________________________________________________")
    doc.add_paragraph()

    # Section 4: Lessons Learned
    add_section_heading(doc, "四、经验教训 (Lessons Learned)")

    table4 = doc.add_table(rows=4, cols=2)
    table4.style = 'Table Grid'
    create_header_row(table4, ["类别", "内容"])
    lessons = ["技术层面预防措施", "管理层面预防措施", "沟通层面预防措施"]
    for i, lesson in enumerate(lessons):
        table4.rows[i+1].cells[0].text = lesson
    set_table_style(table4)
    doc.add_paragraph()

    para = doc.add_paragraph()
    para.add_run("可复用的经验：________________________________________________")
    doc.add_paragraph()

    # Section 5: Follow-up
    add_section_heading(doc, "五、后续跟进 (Follow-up)")

    table5 = doc.add_table(rows=3, cols=4)
    table5.style = 'Table Grid'
    create_header_row(table5, ["跟进事项", "负责人", "完成期限", "状态"])
    for i in range(1, 3):
        table5.rows[i].cells[3].text = "□待完成 □进行中 □已完成"
    set_table_style(table5)
    doc.add_paragraph()
    doc.add_paragraph()

    # Signatures
    para = doc.add_paragraph()
    para.add_run("记录人：________________________    日期：________________")
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run("审阅人：________________________    日期：________________")

    output_path = OUTPUT_DIR + "/W05 - 冲突处理记录表.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

# Run all
if __name__ == "__main__":
    print("Building W01 - 联合交付合作协议模板...")
    create_w01()
    print("Building W02 - 交接文档模板...")
    create_w02()
    print("Building W03 - 每周同步报告模板...")
    create_w03()
    print("Building W04 - 角色分工说明书模板...")
    create_w04()
    print("Building W05 - 冲突处理记录表...")
    create_w05()
    print("\nAll 5 Word documents created successfully!")

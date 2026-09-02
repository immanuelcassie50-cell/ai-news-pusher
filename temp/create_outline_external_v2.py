# File: create_outline_external_v2.py
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Helper functions
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(4)
    return p

def set_run_font(run):
    run.font.name = 'Microsoft YaHei'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}lang', 'zh-CN')

def add_page_break(doc):
    doc.add_page_break()

# ================== Cover ==================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('\n\n\n\n')
run = title.add_run('Content Tools Confidence Trilogy')
run.font.size = Pt(36)
run.font.bold = True
set_run_font(run)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('\n\n')
run = subtitle.add_run('Building Persuasive Expression')
run.font.size = Pt(28)
run.font.bold = True
set_run_font(run)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('\n\n')
run = subtitle2.add_run('—— Course Outline ——')
run.font.size = Pt(20)
set_run_font(run)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('\n\n\n\n\n\n\n\n')
run = info.add_run('Version: V1.0')
run.font.size = Pt(14)
set_run_font(run)

add_page_break(doc)

# ================== Course Introduction ==================
add_heading(doc, 'Course Introduction', 1)

add_para(doc, 'Have you encountered these situations?')
add_bullet(doc, 'You were well prepared, but felt nervous when speaking')
add_bullet(doc, 'Everything you said was correct, but the audience did not believe')
add_bullet(doc, 'Your report was logically confusing, and even you felt it was not clear enough')

add_para(doc, 'Content Tools Confidence Trilogy solves these problems.')
add_para(doc, 'This is a practical course focused on workplace expression confidence. Not about theories, not about listening to cases - you bring your own real work task and practice until you can apply what you learn.')

add_para(doc, 'The core framework is the "Expression Confidence Trilogy":')
add_bullet(doc, 'Toolkit 1: Content Foundation - Make your arguments credible')
add_bullet(doc, 'Toolkit 2: Structure Tools - Make complex content clear and orderly')
add_bullet(doc, 'Toolkit 3: Confident Expression - Let your attitude enhance your content')

add_para(doc, 'The three toolkits progress layer by layer, from content to structure to confidence, helping you become a persuasive communicator.')

add_page_break(doc)

# ================== Target Audience ==================
add_heading(doc, 'Target Audience', 1)

add_heading(doc, 'Suitable For', 2)
add_bullet(doc, 'Working professionals who need to frequently report, present, or propose')
add_bullet(doc, 'Professionals who often need to persuade clients, colleagues, or leaders')
add_bullet(doc, 'Professionals who feel "I actually do well, but cannot express it"')
add_bullet(doc, 'All workplace individuals who want to improve public expression confidence')

add_heading(doc, 'Prerequisites', 2)
add_bullet(doc, 'Some work experience (2+ years recommended)')
add_bullet(doc, 'Real expression task scenarios (reporting, proposing, speaking, etc.)')
add_bullet(doc, 'Willingness to practice with real tasks')

add_heading(doc, 'Not Suitable For', 2)
add_bullet(doc, 'Students with absolutely no expression scenarios')
add_bullet(doc, 'Students who expect "I will automatically improve after listening"')
add_bullet(doc, 'Students whose communication skills are already excellent')

add_page_break(doc)

# ================== Learning Outcomes ==================
add_heading(doc, 'Learning Outcomes', 1)

add_para(doc, 'After completing this course, you will be able to:')
add_bullet(doc, 'Build persuasive content: Learn to support your arguments with data, cases, and emotions, turning listeners from "hearing you" to "believing what you say"')
add_bullet(doc, 'Design clear structure: Master the Pyramid Principle and three basic structures, making complex content organized and easy for listeners to follow')
add_bullet(doc, 'Deliver confident attitude: Learn confident delivery methods in three areas - voice, body, and attitude, letting your attitude enhance your content')
add_bullet(doc, 'Respond to pressure scenarios: Master the "Receive-Transfer-Answer" framework to handle sudden questions and challenges calmly')
add_bullet(doc, 'Build emotional resonance: Learn to use "we", authentic stories, and warm expression to build connection with your audience')

add_page_break(doc)

# ================== Course Outline ==================
add_heading(doc, 'Course Outline', 1)

add_heading(doc, 'Module 1: Content Foundation (Morning)', 2)

module1 = [
    ('Opening', 'Select your real task', '15 min'),
    ('Content Authenticity', 'Three types of content support (data/case/emotion)', '45 min'),
    ('Content Specificity', 'Four dimensions of specificity: time/number/behavior/standard', '30 min'),
    ('Content Focus', 'Less is more: Three-Point Principle and focus courage', '30 min'),
    ('Integration', 'Toolkit 1 comprehensive exercise', '30 min'),
]
table1 = doc.add_table(rows=len(module1)+1, cols=3)
table1.style = 'Table Grid'
table1.rows[0].cells[0].text = 'Session'
table1.rows[0].cells[1].text = 'Content'
table1.rows[0].cells[2].text = 'Duration'
for i, m in enumerate(module1):
    table1.rows[i+1].cells[0].text = m[0]
    table1.rows[i+1].cells[1].text = m[1]
    table1.rows[i+1].cells[2].text = m[2]

add_para(doc, 'Student Output: Complete content support plan for scenario card task')

add_heading(doc, 'Module 2: Structure Tools (Afternoon)', 2)

module2 = [
    ('Structure Clarity', 'Three basic structures (time/space/importance)', '45 min'),
    ('Structure Layering', 'Pyramid Principle: The power of conclusion-first', '45 min'),
    ('Structure Transition', 'Three transition techniques for smooth expression', '30 min'),
    ('Integration', 'Toolkit 2 comprehensive exercise', '30 min'),
]
table2 = doc.add_table(rows=len(module2)+1, cols=3)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = 'Session'
table2.rows[0].cells[1].text = 'Content'
table2.rows[0].cells[2].text = 'Duration'
for i, m in enumerate(module2):
    table2.rows[i+1].cells[0].text = m[0]
    table2.rows[i+1].cells[1].text = m[1]
    table2.rows[i+1].cells[2].text = m[2]

add_para(doc, 'Student Output: Structured expression framework for scenario card task')

add_heading(doc, 'Module 3: Confident Expression', 2)

module3 = [
    ('Confidence Delivery', 'Four methods to let attitude enhance content', '45 min'),
    ('Emotional Resonance', 'Three methods to build connection with audience', '30 min'),
    ('Pressure Response', 'Receive-Transfer-Answer: Standard framework for handling challenges', '30 min'),
    ('Comprehensive Exercise', 'Scenario card complete presentation and feedback', '60 min'),
]
table3 = doc.add_table(rows=len(module3)+1, cols=3)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = 'Session'
table3.rows[0].cells[1].text = 'Content'
table3.rows[0].cells[2].text = 'Duration'
for i, m in enumerate(module3):
    table3.rows[i+1].cells[0].text = m[0]
    table3.rows[i+1].cells[1].text = m[1]
    table3.rows[i+1].cells[2].text = m[2]

add_para(doc, 'Student Output: Deliverable complete expression work + 30-day action plan')

add_page_break(doc)

# ================== Schedule ==================
add_heading(doc, 'Schedule', 1)

schedule = [
    ('Course Duration', '6 hours (1 day, or 2 half-days)'),
    ('Recommended Time', 'Morning 9:00-12:00, Afternoon 13:30-18:00'),
    ('Break Schedule', '10-15 min break every 90 minutes'),
    ('Class Size', 'Recommended 12-30 students'),
    ('Teaching Method', '40% lecture + 50% practice + 10% feedback'),
]
table4 = doc.add_table(rows=len(schedule)+1, cols=2)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = 'Item'
table4.rows[0].cells[1].text = 'Details'
for i, s in enumerate(schedule):
    table4.rows[i+1].cells[0].text = s[0]
    table4.rows[i+1].cells[1].text = s[1]

add_page_break(doc)

# ================== Pricing and Registration ==================
add_heading(doc, 'Pricing and Registration', 1)

pricing = [
    ('Course Fee', '[TBD] yuan/person'),
    ('Early Bird', '[TBD] yuan/person (register 15 days in advance)'),
    ('Group Discount', '[TBD] yuan/person (3+ people group registration)'),
    ('Corporate Training', '[TBD] yuan/day (within 30 people)'),
]
table5 = doc.add_table(rows=len(pricing)+1, cols=2)
table5.style = 'Table Grid'
table5.rows[0].cells[0].text = 'Item'
table5.rows[0].cells[1].text = 'Details'
for i, p in enumerate(pricing):
    table5.rows[i+1].cells[0].text = p[0]
    table5.rows[i+1].cells[1].text = p[1]

add_para(doc, '')
add_heading(doc, 'Registration Contact', 2)
add_bullet(doc, 'Phone: [TBD]')
add_bullet(doc, 'Email: [TBD]')
add_bullet(doc, 'WeChat: [TBD]')

add_heading(doc, 'Class Location', 2)
add_para(doc, '[TBD]')

add_heading(doc, 'Notes', 2)
add_bullet(doc, 'Fee includes course handbook, tools forms, and coffee breaks')
add_bullet(doc, 'Fee does not include transportation or accommodation')
add_bullet(doc, 'Cancellation within 7 days before class, 30% fee deducted')
add_bullet(doc, 'Cancellation within 3 days before class, 50% fee deducted')

add_page_break(doc)

# ================== Student Testimonials ==================
add_heading(doc, 'Student Testimonials', 1)

add_para(doc, '"I only realized after taking this course that my expression problem was "lack of content support", not "poor eloquence". Now before reporting, I first ask myself: What is the support for this statement? The effect is completely different."')
add_para(doc, '— Project Manager, Tech Company')

add_para(doc, '"The trilogy sounds simple, but in practice I found that each toolkit has many details to pay attention to. Especially the Pyramid Principle - I used to think conclusion-first was too arbitrary, now I realize the opposite, the audience actually prefers conclusion-first."')
add_para(doc, '— Consultant, Consulting Company')

add_para(doc, '"The most useful is the Receive-Transfer-Answer framework. Last time in a meeting my leader questioned me, I used this method, and my leader immediately said "you have clearly thought this through". This framework is so practical."')
add_para(doc, '— Department Director, Financial Institution')

add_page_break(doc)

# ================== FAQ ==================
add_heading(doc, 'Frequently Asked Questions', 1)

faqs = [
    ('Q: I am an introverted person, is this course suitable for me?',
     'A: Absolutely suitable. The course is not about turning you into a "talkative person", but helping you clearly and persuasively express "what you want to say". Many introverted people actually progress faster in such courses because they are more used to thinking.'),
    ('Q: How is AI used in the course?',
     'A: AI tools are used to assist content preparation. For example, using AI to help build content support and check expression structure. But the core confident delivery and human interaction require you to complete personally.'),
    ('Q: Is one day enough? Will it be too rushed?',
     'A: The 6-hour course design is optimized, each module has clear learning objectives and outputs. If your task scenario is complex, it is recommended to continue using course tools for deeper practice after the class.'),
    ('Q: What follow-up support is there after the course?',
     'A: After the course, you will receive a course tools quick-reference card and 30-day action plan template. The student group will also exist long-term for everyone to exchange practice insights.'),
]
for qa in faqs:
    add_para(doc, qa[0])
    add_para(doc, qa[1])
    add_para(doc, '')

# Save
output_path = 'D:/NewCourse/PublicExpression/05_ContentToolsConfidence/Course_Outline_External.docx'
doc.save(output_path)
print(f'Saved to {output_path}')

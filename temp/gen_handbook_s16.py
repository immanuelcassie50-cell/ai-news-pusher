#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    se = OxmlElement('w:shd')
    se.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(se)

doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

# 附录
appt = doc.add_paragraph()
r = appt.add_run('附录')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor(0,51,102)

doc.add_paragraph()
termt = doc.add_paragraph()
r = termt.add_run('附录A：术语表')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

glossary = [
    ('复盘', '对过去事件进行叙事性描述，总结经验教训'),
    ('决策卡', '将判断结构提炼成可现场调用的检查表和触发条件'),
    ('隐性判断', '决策者本人未曾语言化的经验规则'),
    ('判断难度', '决策的困难程度，取决于是否存在多个合理选项'),
    ('触发条件', '告诉使用者什么时候该打开决策卡的判断性描述'),
    ('检查表', '列出最容易被忽略但忽略会出大问题的信号点'),
    ('开关', '告诉使用者什么时候该停下来重新想的一句话'),
    ('场景映射矩阵', '把同一条底层判断逻辑，拆解成几个具体场景各自版本'),
    ('追因', '挖掘决策触发点的访谈维度'),
    ('权衡', '挖掘选项排除过程的访谈维度'),
    ('未预见的假设', '挖掘无意识依赖前提的访谈维度'),
    ('决策稽核', '检查判断逻辑有没有被沉淀成组织可复用东西的机制'),
    ('组织判断力', '团队学会识别、访谈、结构化、复用判断的能力'),
]

for term, defn in glossary:
    p = doc.add_paragraph()
    r = p.add_run(f'{term}：')
    r.bold = True
    p.add_run(defn)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# 金句合集
goldt = doc.add_paragraph()
r = goldt.add_run('附录B：金句合集')
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0,102,204)

gold_quotes = [
    '教训只能让人不再犯错，决策卡能让人复制成功。',
    '值得复盘的从来不是结果好坏，是判断有没有难度。',
    '决策者自己讲不出他做对了什么，因为他当时没觉得那是个选择。',
    '每一次我就是觉得不对背后，都有一条没被说出来的经验规则。',
    '流程图告诉你按顺序做什么，决策卡告诉你什么时候该停下来想。',
    '决策卡最贵的三个字不是怎么做，是什么时候。',
    '卡片越通用，越没人用；卡片越具体，用的人越多。',
    '讲失败案例的目的不是让人害怕，是让人在自己身上找到那个岔路口。',
    '卡片评审不是找错别字，是找这句话在什么情况下会害死人。',
    '讲道理讲三小时，不如让他在模拟场景里做错一次决策。',
    '一份没有主人的工具，用一次就是最后一次。',
    '稽核表查的不是做对了没有，是这次的判断有没有被沉淀下来。',
    '一张三年没改过的决策卡，大概率已经没人真的在用。',
    '你替他说出来的道理，他转身就忘；他自己说出来的判断，才是他的。',
    '决策卡会被淘汰，做卡的能力不会。',
]

for i, quote in enumerate(gold_quotes, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f'{i}. ')
    r.bold = True
    r = p.add_run(f'"{quote}"')
    r.italic = True
    p.paragraph_format.space_after = Pt(8)

print('S16: Appendices done')
doc.save(r'D:\CC\temp\handbook_s16.docx')

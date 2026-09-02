# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._tc.get_or_add_tcPr().makeelement(qn('w:shd'), {})
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Microsoft YaHei'
    r.font.size = Pt(size)
    r.bold = bold
    r._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.paragraph_format.space_after = Pt(6)
    return p

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._r.get_or_add_rPr().set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h

output_path = r"D:/新课开发/变革管理/16-变革成果固化机制：防止新流程人走茶凉/完整课程包/05-讲师手册/讲师手册-变革成果固化机制.docx"
doc = Document(output_path)

# ============ CHAPTER 7 ============
add_heading(doc, "第七章  案例使用说明和讨论引导问题", 1)

add_heading(doc, "7.1 情境案例：消失的AI审批流程", 2)
add_para(doc, "案例背景：某制造企业在引入AI质检系统后，IT部门设计了完整的AI审批流程。三个月后项目经理小张调岗，新任项目经理小李发现AI审批流程名存实亡，变成了先上车后补票。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 这个案例中，固化失败的原因是什么？")
add_para(doc, "2. AI审批流程目前处于哪个固化成熟度级别？")
add_para(doc, "3. 如果你是新任项目经理，你会怎么做？")
add_para(doc, "4. 推动者调岗前，应该做什么才能防止这种情况？")
add_para(doc, "5. 如何设计预警指标来提前发现这种问题？")
add_para(doc, "讲师引导技巧：", bold=True)
add_para(doc, "先让学员自由发言，不要急于给答案。引导学员从不同角度分析（制度层、机制层、文化层）。最后总结：固化需要三个层次协同，不能只靠制度")

add_heading(doc, "7.2 案例：某企业无声无息的SOP革命", 2)
add_para(doc, "案例背景：某科技公司在推行新的代码审查SOP时，没有大张旗鼓的启动会，而是通过系统化的方式让新流程自然运转。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 为什么这种无声无息的方式能成功？")
add_para(doc, "2. 这种方式和传统的变革推动方式有什么不同？")
add_para(doc, "3. 这种方式适用于所有变革吗？什么情况下不适用？")

add_heading(doc, "7.3 案例：某企业将AI流程固化进SOP的完整过程", 2)
add_para(doc, "案例背景：某企业在引入AI客服系统后，如何将AI应答流程固化进SOP，确保系统升级后流程不倒退。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 制度固化在这个案例中是如何体现的？")
add_para(doc, "2. 岗位说明书是如何修改的？")
add_para(doc, "3. 考核指标是如何设计的？")

add_heading(doc, "7.4 案例：某企业设计的三棱镜固化体系", 2)
add_para(doc, "案例背景：某企业在推行精益生产时，设计了一套三棱镜固化体系，包括检查机制、激励机制和传承机制。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 三棱镜模型的三个面分别解决什么问题？")
add_para(doc, "2. 为什么三棱镜模型比单一机制更有效？")
add_para(doc, "3. 你能设计一个适用于本企业的三棱镜模型吗？")

add_heading(doc, "7.5 案例：某企业用户至上文化的故事传承", 2)
add_para(doc, "案例背景：某互联网公司在推行用户至上文化后，建立了完整的故事沉淀与传播机制，5年后文化成为公司的基因。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 故事传承在文化固化中起什么作用？")
add_para(doc, "2. 如何收集和筛选有效的故事？")
add_para(doc, "3. 你所在企业有哪些值得传承的变革故事？")

add_heading(doc, "7.6 案例：固化预警机制成功阻止了一次回潮", 2)
add_para(doc, "案例背景：某制造企业通过预警机制发现合理化建议提交率下降，及时干预阻止了一次回潮。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 预警机制在这个案例中是如何发挥作用的？")
add_para(doc, "2. 为什么预警响应要如此迅速？")
add_para(doc, "3. 干预措施为什么能有效？")

add_heading(doc, "7.7 案例：一次失败的固化复盘", 2)
add_para(doc, "案例背景：某企业在数字化转型两年后发现CRM系统使用率持续下降，复盘结论是员工不会用，决定加强培训。但三个月后使用率继续下降。")
add_para(doc, "讨论问题：", bold=True)
add_para(doc, "1. 这个复盘失败的原因是什么？")
add_para(doc, "2. 正确的复盘应该得出什么结论？")
add_para(doc, "3. 如何避免类似的复盘错误？")

# ============ CHAPTER 8 ============
add_heading(doc, "第八章  常见问题和应对方案", 1)

add_heading(doc, "问题1：学员认为写了制度就是固化了", 2)
add_para(doc, "应对：强调文件层只是起点，通过案例说明写归写、做归做的现象。")
add_para(doc, "引导话术：制度文本只是固化的第一步。真正的固化是让制度从写在纸上变成执行在行动中。案例延伸：可用消失的AI审批流程案例说明")

add_heading(doc, "问题2：学员陷入坚持陷阱", 2)
add_para(doc, "应对：帮助识别是否把太多精力花在维持现状上，固化才是出路。")
add_para(doc, "引导话术：如果你发现自己每天都在提醒、催促、检查，那你可能已经陷入了坚持陷阱。固化是逃脱陷阱的唯一出路。")

add_heading(doc, "问题3：如何处理奖励错位？", 2)
add_para(doc, "应对：奖励要明确标准，要与真正需要强化的行为挂钩。")
add_para(doc, "引导话术：奖励错位比没有奖励更糟糕，因为它会误导行为方向。案例延伸：用一次失败的固化复盘案例说明")

add_heading(doc, "问题4：小企业需要完整的固化体系吗？", 2)
add_para(doc, "应对：核心原则不变，具体可简化，五要素中至少四要素必须有。")
add_para(doc, "引导话术：小企业可能不需要复杂的制度文本，但检查机制、激励机制、传承机制仍然需要。讨论：学员讨论本企业可以简化的部分")

add_heading(doc, "问题5：固化失效的早期信号是什么？", 2)
add_para(doc, "应对：执行数据下滑、反馈声音变化、人员变动影响、配套机制退化。")
add_para(doc, "引导话术：固化失效往往不是突然发生的，而是有苗头的。关键是建立预警机制，及时发现苗头。")

add_heading(doc, "问题6：如何处理老员工的抵触？", 2)
add_para(doc, "应对：通过社会认同和正向激励，让新做法成为大多数人的选择。")
add_para(doc, "引导话术：老员工不是敌人，他们只是在用旧习惯工作。关键是用新习惯的收益来影响他们。")

add_heading(doc, "问题7：固化需要多长时间？", 2)
add_para(doc, "应对：从Level 0到Level 3通常需要1-2年，要有耐心。")
add_para(doc, "引导话术：固化不是一次性项目，而是一个持续的过程。快速固化往往不牢固。")

add_heading(doc, "问题8：如何让高层支持固化工作？", 2)
add_para(doc, "应对：用数据和案例说明固化失效的成本。")
add_para(doc, "引导话术：高层的支持不是等来的，而是争取来的。关键是用数据说话。")

# ============ CHAPTER 9 ============
add_heading(doc, "第九章  教室布置与道具准备清单", 1)

add_heading(doc, "9.1 教室布置要求", 2)
add_para(doc, "教室面积：能容纳25-35人")
add_para(doc, "桌椅布局：小组讨论布局（4-5人/组），每组配白板纸")
add_para(doc, "投影设备：高清投影仪，投影幕布或白墙")
add_para(doc, "白板设备：每组配1块白板或大白纸")
add_para(doc, "音响设备：用于案例视频播放")
add_para(doc, "灯光控制：便于调节投影时的灯光")

add_heading(doc, "9.2 道具准备清单（通用）", 2)
items = [
    "学员手册（每人1本）",
    "课程大纲（每人1份）",
    "工具模板打印本（每组1套）",
    "投票卡（每人1套）",
    "案例打印材料（每组1份）",
    "白板笔（每组2支）",
    "便签纸（用于小组讨论记录）",
    "计时器（讲师用）",
    "奖品/奖励道具（用于表彰活动）",
]
for item in items:
    add_para(doc, "□  " + item)

add_heading(doc, "9.3 模块二专项道具", 2)
add_para(doc, "□ 制度诊断工作坊：提前收集的3-5份脱敏制度文本")
add_para(doc, "□ 岗位说明书编写演练：打印的《变革职责嵌入岗位说明书》模板")
add_para(doc, "□ 考核指标设计挑战：6个场景卡片")

add_heading(doc, "9.4 模块三专项道具", 2)
add_para(doc, "□ 三棱镜体系设计工作坊：打印的设计模板")
add_para(doc, "□ 检查节点设计实战：打印的检查节点设计模板")

add_heading(doc, "9.5 模块四专项道具", 2)
add_para(doc, "□ 故事收集模板（每组1份）")
add_para(doc, "□ 新人文化适应90天计划模板（每组1份）")
add_para(doc, "□ 表彰体系设计模板（每组1份）")

add_heading(doc, "9.6 模块五专项道具", 2)
add_para(doc, "□ 固化效果预警指标体系模板（每组1份）")
add_para(doc, "□ 固化失效干预策略选择检查表（每组1份）")
add_para(doc, "□ 固化体系季度复盘报告模板（每组1份）")

# ============ CHAPTER 10 ============
add_heading(doc, "第十章  效果评估方法", 1)

add_heading(doc, "10.1 知识评估", 2)
add_para(doc, "概念辨析题：", bold=True)
add_para(doc, "固化和坚持的本质区别是什么？")
add_para(doc, "制度固化五要素是什么？")
add_para(doc, "三棱镜模型是哪三棱？")
add_para(doc, "判断题：", bold=True)
add_para(doc, "识别给定情景描述属于哪个固化成熟度级别")
add_para(doc, "简答题：", bold=True)
add_para(doc, "固化成熟度四问是什么？")
add_para(doc, "早期预警指标的设计原则是什么？")
add_para(doc, "论述题：", bold=True)
add_para(doc, "为什么关键岗位继任者培养是固化体系中不可或缺的环节？")
add_para(doc, "为什么说表彰是文化固化的重要手段？")

add_heading(doc, "10.2 技能评估", 2)
add_para(doc, "能否准确判断案例中变革的固化成熟度")
add_para(doc, "能否识别制度缺陷并提出修改建议")
add_para(doc, "能否设计无人驾驶的检查节点")
add_para(doc, "能否设计完整的三棱镜固化体系")
add_para(doc, "能否根据失效程度选择合适的干预策略")

add_heading(doc, "10.3 产出评估", 2)
add_para(doc, "课程期间产出评估：", bold=True)
add_para(doc, "固化成熟度自评（4问）")
add_para(doc, "本企业变革成果固化现状诊断报告")
add_para(doc, "三层固化体系设计案（制度层/机制层/文化层）")
add_para(doc, "固化效果评估指标体系")
add_para(doc, "后续90天固化行动方案")

add_heading(doc, "10.4 行动评估（课后30天）", 2)
add_para(doc, "跟踪学员是否在团队中建立了固化检验机制")
add_para(doc, "是否开始执行固化行动方案")
add_para(doc, "固化效果评估指标是否在实际工作中被使用")

# ============ CHAPTER 11 ============
add_heading(doc, "第十一章  讲师注意事项", 1)

add_heading(doc, "11.1 通用注意事项", 2)
add_para(doc, "1. 避免文件层幻觉：要强调文件层只是起点，真正的固化是让变革成果自动运转")
add_para(doc, "2. 强调固化是系统建设：固化不是做一件大事，而是做好很多件小事")
add_para(doc, "3. 警惕坚持陷阱：帮助学员识别是否陷入维持现状的陷阱")
add_para(doc, "4. 固化需要时间：从Level 0到Level 3通常需要1-2年，要有耐心")
add_para(doc, "5. 文化层是最高境界：文件可以抄，流程可以复制，但文化无法移植")

add_heading(doc, "11.2 各模块注意事项", 2)

add_heading(doc, "模块一注意事项", 3)
add_para(doc, "学员可能混淆固化和坚持，要用对比表格清晰区分")
add_para(doc, "固化成熟度模型要讲透，这是后续模块的基础")
add_para(doc, "投票活动要让学员参与，活跃课堂气氛")

add_heading(doc, "模块二注意事项", 3)
add_para(doc, "五要素框架是本模块核心，要让学员真正掌握")
add_para(doc, "制度编写规范要有大量实例，让学员看到模糊vs具体的对比")
add_para(doc, "工作坊环节要给足时间，让学员真正动手写")

add_heading(doc, "模块三注意事项", 3)
add_para(doc, "三棱镜模型是本模块亮点，要讲得生动")
add_para(doc, "检查节点设计要有实操，让学员设计自己企业的检查节点")
add_para(doc, "继任者培养的话题可能敏感，要以案例引导而非说教")

add_heading(doc, "模块四注意事项", 3)
add_para(doc, "习惯形成的67次规律要讲，这是行为改变的科学依据")
add_para(doc, "故事沉淀机制要结合案例，让学员感受到故事的力量")
add_para(doc, "表彰体系设计要有互动，学员要能看到真实案例的利弊")

add_heading(doc, "模块五注意事项", 3)
add_para(doc, "预警指标设计要有实战，学员要设计自己企业的指标")
add_para(doc, "干预策略选择矩阵要讲透，这是应对固化失效的关键工具")
add_para(doc, "复盘报告要结合学员自己企业的情况来练习")

add_heading(doc, "11.3 时间管理技巧", 2)
add_para(doc, "每个模块开始前回顾上一个模块的关键要点（2-3分钟）")
add_para(doc, "案例讨论要控制时间，避免过度延展")
add_para(doc, "工作坊环节要给足时间，但也要适时收尾")
add_para(doc, "留出Q&A时间，但控制问题数量")

add_heading(doc, "11.4 课堂氛围管理", 2)
add_para(doc, "开场破冰：用 人走茶凉 的现象引发共鸣")
add_para(doc, "小组竞赛：适当时机引入小组竞争，提高参与度")
add_para(doc, "及时表扬：对学员的优秀回答要当场表扬")
add_para(doc, "处理沉默：遇到沉默时，用提问打破僵局")

add_heading(doc, "11.5 处理学员问题", 2)
add_para(doc, "专业问题：给出清晰答案，不要模糊其辞")
add_para(doc, "延伸问题：可简短回答，提示课后深入探讨")
add_para(doc, "挑战性问题：肯定学员的思考，共同探讨")

# ============ HEADER/FOOTER ============
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = '变革成果固化机制——讲师手册'
header_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = '第 '
run = footer_para.add_run()
fldChar = run._r.makeelement(qn('w:fldChar'), {})
fldChar.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar)
instrText = run._r.makeelement(qn('w:instrText'), {})
instrText.text = ' PAGE '
run._r.append(instrText)
fldChar2 = run._r.makeelement(qn('w:fldChar'), {})
fldChar2.set(qn('w:fldCharType'), 'separate')
run._r.append(fldChar2)
run2 = footer_para.add_run('1')
run2.font.size = Pt(10)
fldChar3 = run2._r.makeelement(qn('w:fldChar'), {})
fldChar3.set(qn('w:fldCharType'), 'end')
run2._r.append(fldChar3)
footer_para.add_run(' 页')

doc.save(output_path)
print("Chapters 7-11 added successfully")
print("Document complete!")

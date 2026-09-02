# -*- coding: utf-8 -*-
"""
企业大学品牌资产设计课程 - 讲师手册生成脚本
Instructor Handbook Generator for Enterprise University Brand Asset Design Course
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================
# 配置参数
# ============================================================

OUTPUT_PATH = r"D:/新课开发/企业大学/对外/4.企业大学品牌资产设计：把理想愿景转化为对外可辨识的品牌/讲师手册/讲师手册.docx"

# 颜色定义
PRIMARY_COLOR = RGBColor(0x1F, 0x38, 0x64)      # 深蓝色 - 主色
SECONDARY_COLOR = RGBColor(0x2E, 0x75, 0xB6)   # 中蓝色 - 辅助色
ACCENT_COLOR = RGBColor(0xC0, 0x00, 0x00)       # 深红色 - 强调色
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)          # 深灰色 - 正文
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)         # 浅灰色 - 背景
TABLE_HEADER_BG = RGBColor(0x1F, 0x38, 0x64)   # 表格标题背景
TABLE_ALT_BG = RGBColor(0xE9, 0xED, 0xF4)      # 表格交替行背景

# ============================================================
# 辅助函数
# ============================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    # RGBColor stores values as integer, need to extract properly
    if isinstance(color, RGBColor):
        r = int(color[0]) if len(color) > 0 else 0
        g = int(color[1]) if len(color) > 1 else 0
        b = int(color[2]) if len(color) > 2 else 0
    else:
        r, g, b = 0, 0, 0
    shading.set(qn('w:fill'), '%02X%02X%02X' % (r, g, b))
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(table):
    """设置表格边框"""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), '4472C4')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_table_with_header(doc, headers, rows, header_bg=TABLE_HEADER_BG):
    """创建带标题行的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 标题行
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, header_bg)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    set_cell_borders(table)
    return table

def add_info_box(doc, title, content, box_type='note'):
    """添加信息框"""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)

    colors = {
        'note': ((0xE9, 0xED, 0xF4), SECONDARY_COLOR),
        'tip': ((0xE2, 0xEF, 0xDA), RGBColor(0x70, 0xAD, 0x47)),
        'warning': ((0xFF, 0xEB, 0x9C), ACCENT_COLOR),
    }

    bg_color, border_color = colors.get(box_type, colors['note'])
    set_cell_shading(cell, RGBColor(*bg_color))

    p = cell.paragraphs[0]
    run = p.add_run(f"【{title}】")
    run.font.bold = True
    run.font.color.rgb = border_color
    run.font.size = Pt(10)

    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(9)

    doc.add_paragraph()

def add_time_distribution_table(doc, activities):
    """添加时间分配表格"""
    headers = ["环节", "时长", "形式", "要点"]
    rows = []
    for activity in activities:
        rows.append([
            activity.get('name', ''),
            activity.get('duration', ''),
            activity.get('format', ''),
            activity.get('key_points', '')
        ])
    return add_table_with_header(doc, headers, rows)

def add_module_section(doc, module_num, module_title, module_data):
    """添加模块章节"""
    # 模块标题
    heading = doc.add_heading(f"模块{module_num}：{module_title}", level=1)
    for run in heading.runs:
        run.font.color.rgb = PRIMARY_COLOR

    # 基本信息表格
    info_data = [
        ["教学目标", module_data.get('objective', '')],
        ["建议时长", module_data.get('duration', '')],
        ["学员准备", module_data.get('prerequisites', '')],
    ]
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for i, row_data in enumerate(info_data):
        row = table.rows[i]
        row.cells[0].text = row_data[0]
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
        row.cells[1].text = row_data[1]
        for paragraph in row.cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    set_cell_borders(table)
    doc.add_paragraph()

    # 关键知识点
    doc.add_heading("关键知识点", level=2)
    for point in module_data.get('key_points', []):
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    # 互动设计
    doc.add_heading("互动设计", level=2)
    for interaction in module_data.get('interactions', []):
        p = doc.add_paragraph()
        run = p.add_run(f"【{interaction.get('name', '')}】")
        run.font.bold = True
        run.font.color.rgb = SECONDARY_COLOR
        run2 = p.add_run(f" - {interaction.get('description', '')}")
        run2.font.size = Pt(10)

    # 时间分配
    doc.add_heading("时间分配", level=2)
    if 'time_distribution' in module_data:
        add_time_distribution_table(doc, module_data['time_distribution'])

    # 常见问题与应对
    doc.add_heading("常见问题与应对", level=2)
    qa_data = []
    for qa in module_data.get('faqs', []):
        qa_data.append([qa.get('question', ''), qa.get('answer', '')])
    if qa_data:
        qa_table = doc.add_table(rows=len(qa_data), cols=2)
        qa_table.style = 'Table Grid'
        for i, (q, a) in enumerate(qa_data):
            row = qa_table.rows[i]
            row.cells[0].text = q
            set_cell_shading(row.cells[0], RGBColor(0xFF, 0xF2, 0xCC))
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
            row.cells[1].text = a
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        set_cell_borders(qa_table)

    # 物料清单
    doc.add_heading("物料清单", level=2)
    materials = module_data.get('materials', [])
    if materials:
        for mat in materials:
            p = doc.add_paragraph(mat, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.3)
    else:
        p = doc.add_paragraph("无特殊物料需求")
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_page_break()

# ============================================================
# 创建文档
# ============================================================

def create_instructor_handbook():
    """创建完整的讲师手册"""

    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("\n\n\n\n")
    run.font.size = Pt(36)

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("企业大学品牌资产设计")
    run2.font.size = Pt(32)
    run2.font.bold = True
    run2.font.color.rgb = PRIMARY_COLOR

    title3 = doc.add_paragraph()
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = title3.add_run("把理想愿景转化为对外可辨识的品牌")
    run3.font.size = Pt(24)
    run3.font.color.rgb = SECONDARY_COLOR

    title4 = doc.add_paragraph()
    title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = title4.add_run("\n\n\n\n讲师手册")
    run4.font.size = Pt(28)
    run4.font.bold = True
    run4.font.color.rgb = ACCENT_COLOR

    # 课程信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run("\n\n\n\n\n课程时长：2天（12小时）\n目标学员：企业大学负责人、品牌负责人、培训管理者")
    info_run.font.size = Pt(14)
    info_run.font.color.rgb = DARK_TEXT

    doc.add_page_break()

    # ========== 目录 ==========
    toc_title = doc.add_heading("目 录", level=0)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        ("第一章", "讲师资格与要求"),
        ("第二章", "课程目标与学员画像"),
        ("第三章", "教学方法论"),
        ("第四章", "模块1：品牌战略与定位（90分钟）"),
        ("第五章", "模块2：品牌识别系统设计（180分钟）"),
        ("第六章", "模块3：品牌传播与体验（150分钟）"),
        ("第七章", "模块4：品牌评估与优化（120分钟）"),
        ("第八章", "模块5：品牌资产管理体系（150分钟）"),
        ("第九章", "模块6：企业大学品牌建设实战（210分钟）"),
        ("第十章", "评估标准与评分指引"),
        ("附录", "参考文献、延伸阅读、工具清单"),
    ]

    for chapter, content in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"{chapter}  {content}")
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(12)

    doc.add_page_break()

    # ========== 第一章：讲师资格与要求 ==========
    doc.add_heading("第一章  讲师资格与要求", level=1)
    doc.add_paragraph("本章旨在明确企业大学品牌资产设计课程的讲师资质标准，确保授课质量与专业性。")

    doc.add_heading("1.1 专业背景要求", level=2)

    requirements_1 = [
        "学历要求：本科及以上学历，管理学、市场营销、品牌管理、人力资源开发等相关专业优先",
        "专业知识：深入理解品牌资产理论（Keller品牌资产模型、Aaker品牌资产理论等）",
        "行业经验：具有5年以上企业品牌管理或企业大学运营管理经验",
        "教学经验：具有3年以上培训课程开发或授课经验优先",
        "工具掌握：熟悉品牌调研工具、品牌资产评估方法、品牌传播渠道分析等"
    ]
    for req in requirements_1:
        p = doc.add_paragraph(req, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    doc.add_heading("1.2 授课经验要求", level=2)

    experience_items = [
        ["经验层级", "具体要求"],
        ["基础要求", "能够清晰讲解品牌资产基本概念，能够引导学员参与案例讨论"],
        ["进阶要求", "能够运用多种教学方法激发学员思考，能够处理课堂中的突发状况"],
        ["高级要求", "能够根据学员反馈灵活调整教学策略，能够引导高阶讨论和深度反思"],
    ]
    add_table_with_header(doc, experience_items[0], experience_items[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("1.3 认证资格要求", level=2)

    cert_items = [
        "必修培训：完成本课程内部讲师认证培训（16学时）",
        "试讲评估：完成至少2次试讲并获得评估通过（评分≥85分）",
        "持续学习：每年完成至少8学时的品牌管理或培训方法论相关继续教育",
        "认证有效期：认证有效期为2年，到期需重新认证"
    ]
    for item in cert_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    add_info_box(doc, "讲师发展路径",
        "初级讲师 → 中级讲师 → 高级讲师 → 首席讲师\n"
        "每级晋升需满足相应的授课时长、学员评价和课程开发贡献要求。", "tip")

    doc.add_page_break()

    # ========== 第二章：课程目标与学员画像 ==========
    doc.add_heading("第二章  课程目标与学员画像", level=1)

    doc.add_heading("2.1 课程总体目标", level=2)

    overall_goal = """通过本课程的学习，学员将能够：
1. 理解企业大学品牌资产的内涵与价值，明确品牌对组织发展的战略意义
2. 掌握品牌资产设计的核心方法论，能够将组织理想愿景转化为可辨识的品牌要素
3. 学会品牌识别系统的构建技巧，能够设计符合企业战略的品牌架构
4. 掌握品牌传播与体验管理的关键策略，能够提升品牌影响力
5. 建立品牌评估与优化体系，能够持续提升品牌资产价值
6. 了解企业大学品牌资产管理的最佳实践，能够结合自身企业实际情况制定品牌发展战略"""

    for line in overall_goal.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.left_indent = Inches(0.3)

    doc.add_heading("2.2 三维目标（知识/技能/态度）", level=2)

    obj_table_data = [
        ["维度", "目标描述", "具体表现"],
        ["知识目标", "掌握品牌资产的理论框架和设计方法", "能够准确阐述品牌资产模型，能够识别品牌资产的构成要素"],
        ["技能目标", "具备品牌资产设计和管理的实操能力", "能够独立完成品牌诊断，能够设计品牌识别系统，能够制定品牌传播策略"],
        ["态度目标", "认识到品牌资产的战略价值", "能够向他人有效传达品牌价值，能够主动推动品牌建设工作"],
    ]
    add_table_with_header(doc, obj_table_data[0], obj_table_data[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("2.3 学员画像分析", level=2)

    profile_data = [
        ["维度", "典型特征"],
        ["职业背景", "企业大学负责人、培训总监、学习与发展经理、品牌经理、市场经理等"],
        ["行业分布", "制造、金融、互联网、医疗健康、教育等专业服务行业居多"],
        ["职级层次", "中高层管理者为主，通常是决策参与者或影响者"],
        ["年龄分布", "35-50岁为主，具有丰富的管理经验"],
        ["学习特点", "注重实战性和可操作性，喜欢案例讨论和经验分享"],
    ]
    add_table_with_header(doc, profile_data[0], profile_data[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("2.4 学员痛点与需求", level=2)

    pain_points = [
        ["痛点类型", "具体表现", "课程回应方式"],
        ["战略认知不足", "对品牌资产的理解停留在表面，缺乏系统认知", "系统讲解品牌资产理论框架，提供全景视图"],
        ["方法工具缺乏", "有品牌意识但不知道如何落地实施", "提供详细的操作工具和方法论，配套练习"],
        ["跨部门协作难", "品牌建设涉及多部门，难以协调推动", "通过沙盘模拟和实战案例，帮助理解协作机制"],
        ["效果评估困难", "品牌效果难以量化，投资回报不清晰", "介绍品牌资产评估工具和方法，提供评估框架"],
        ["资源投入有限", "预算和人力有限，不知道如何高效建设", "分享低成本高效益的品牌建设策略和最佳实践"],
    ]
    add_table_with_header(doc, pain_points[0], pain_points[1:], TABLE_HEADER_BG)

    doc.add_page_break()

    # ========== 第三章：教学方法论 ==========
    doc.add_heading("第三章  教学方法论", level=1)

    doc.add_heading("3.1 教学方法概述", level=2)

    methods_intro = """本课程采用"多元融合"的教学方法体系，根据不同模块的内容特点和学员需求，灵活运用以下教学方法："""
    doc.add_paragraph(methods_intro)

    methods = [
        ["教学方法", "特点", "适用场景", "时间占比"],
        ["讲授法", "系统传授知识，效率高", "理论框架讲解、方法论介绍", "30%"],
        ["案例教学法", "贴近实际，启发思考", "品牌实战案例分析、讨论", "25%"],
        ["工作坊法", "动手实践，深度参与", "品牌设计练习、工具演练", "20%"],
        ["引导式教学", "学员为主，激发潜能", "问题探讨、方案共创", "15%"],
        ["沙盘模拟法", "沉浸体验，检验学习", "品牌战略推演、决策演练", "10%"],
    ]
    add_table_with_header(doc, methods[0], methods[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("3.2 各类教学方法的使用时机", level=2)

    usage_guide = [
        ["教学方法", "最佳使用时机", "操作要点"],
        ["讲授法", "模块开场引入新概念时\n知识体系庞大需要系统梳理时", "控制在15分钟以内\n配合视觉化呈现\n穿插提问互动"],
        ["案例教学法", "知识点的应用场景说明\n引发学员思考和讨论\n验证方法的有效性", "选择贴近学员行业的案例\n案例要有冲突和挑战\n留足讨论时间"],
        ["工作坊法", "方法工具的实践演练\n学员需要动手操作的技能", "提供清晰的指引\n分步骤实施\n及时给予反馈"],
        ["引导式教学", "开放性问题探讨\n学员经验分享\n方案共创", "提出好问题\n控制讨论方向\n总结提炼要点"],
        ["沙盘模拟法", "综合能力的检验\n决策能力的培养\n团队协作的锻炼", "设置清晰的规则\n营造紧迫感\n及时复盘反思"],
    ]
    add_table_with_header(doc, usage_guide[0], usage_guide[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("3.3 课堂管理要点", level=2)

    management_tips = [
        ["管理维度", "要点", "注意事项"],
        ["时间管理", "严格按照时间分配执行\n设置各环节的时间提醒", "提前5分钟提醒\n灵活调整但不过度超时"],
        ["节奏把控", "理论讲解与互动练习穿插\n关注学员状态调整节奏", "观察学员注意力变化\n适时休息和活动"],
        ["参与度管理", "关注不同类型学员的参与\n鼓励沉默学员发言", "避免少数人垄断讨论\n认可各种观点"],
        ["环境营造", "营造开放安全的氛围\n鼓励提问和质疑", "对幼稚问题也要正面回应\n示范尊重和倾听"],
        ["设备管理", "提前测试设备\n准备备份方案", "确保投影、音响正常\n打印材料提前准备"],
    ]
    add_table_with_header(doc, management_tips[0], management_tips[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("3.4 突发情况处理", level=2)

    emergency_handling = [
        ["突发情况", "应对策略", "预防措施"],
        ["学员提问无法回答", "坦诚承认不了解，表示课后研究回复\n引导学员一起思考解决方案", "充分准备，熟悉相关领域\n提前告知学员边界范围"],
        ["学员质疑课程内容", "认真倾听，表示理解\n用案例或数据回应\n避免正面冲突", "课程内容要经得起推敲\n引用权威来源"],
        ["讨论偏离主题", "温和拉回，邀请学员聚焦\n肯定参与但指出方向", "设置明确的讨论边界\n及时总结归纳"],
        ["学员之间的冲突", "暂停讨论，分别沟通\n强调课堂规范\n寻求共同点", "事先说明课堂规则\n营造包容氛围"],
        ["设备故障", "切换到备用方案\n利用白板或纸张\n调整教学方式", "提前测试所有设备\n准备纸质备份"],
        ["时间紧张", "优先保证核心内容\n压缩讨论时间\n提供课后延伸材料", "留出缓冲时间\n识别核心内容"],
    ]
    add_table_with_header(doc, emergency_handling[0], emergency_handling[1:], TABLE_HEADER_BG)

    add_info_box(doc, "讲师心态",
        "遇到突发情况时，保持冷静和开放的心态。\n"
        "记住：很多'突发'情况都是学习机会，关键是如何转化为教学资源。", "tip")

    doc.add_page_break()

    # ========== 第四章至第九章：各模块内容 ==========

    # 模块1
    module1 = {
        'objective': '理解品牌资产的战略意义，掌握品牌定位的基本方法，能够分析企业大学的品牌现状',
        'duration': '90分钟',
        'prerequisites': '提前了解本企业企业大学或培训体系的现状',
        'key_points': [
            '品牌资产的定义与价值（Keller模型、Aaker模型）',
            '企业大学品牌的战略定位',
            '品牌愿景、使命、价值观的转化',
            '企业大学品牌与母品牌的关系',
            '品牌现状诊断方法'
        ],
        'interactions': [
            {'name': '破冰活动', 'description': '学员自我介绍+品牌困惑一句话，每人1分钟'},
            {'name': '案例导入', 'description': '分析3个知名企业大学品牌案例（华为、西门子、GE）'},
            {'name': '小组讨论', 'description': '讨论各自企业大学品牌建设的挑战与机遇'},
        ],
        'time_distribution': [
            {'name': '开场与目标导入', 'duration': '10分钟', 'format': '讲授', 'key_points': '说明课程目标，介绍整体框架'},
            {'name': '品牌资产理论', 'duration': '25分钟', 'format': '讲授+互动', 'key_points': '讲解核心概念，学员案例思考'},
            {'name': '企业大学品牌定位', 'duration': '20分钟', 'format': '讲授', 'key_points': '定位方法论，标杆分析'},
            {'name': '茶歇', 'duration': '5分钟', 'format': '自由交流', 'key_points': '休息与社交'},
            {'name': '品牌现状诊断', 'duration': '20分钟', 'format': '工作坊', 'key_points': '使用诊断工具分析自身企业'},
            {'name': '小结与预告', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾要点，预告下节'},
        ],
        'faqs': [
            {'question': '学员认为品牌只是市场部的事', 'answer': '通过母品牌与企业大学品牌的关系分析，让学员理解企业大学品牌的独特性和全员参与的重要性'},
            {'question': '部分学员企业品牌意识薄弱', 'answer': '从业务价值角度切入，强调品牌对人才吸引、员工激励、组织文化传承的价值'},
            {'question': '讨论时有人偏离主题', 'answer': '使用"停车场"方法记录其他话题，在茶歇时简要回应'},
        ],
        'materials': [
            '品牌资产理论模型图（Keller、Aaker）',
            '企业大学品牌定位工具包',
            '品牌现状诊断问卷',
            '案例资料（华为、西门子、GE企业大学）',
            '大白纸、马克笔',
        ]
    }
    add_module_section(doc, 1, "品牌战略与定位", module1)

    # 模块2
    module2 = {
        'objective': '掌握品牌识别系统的构成要素，学会设计企业大学品牌识别系统',
        'duration': '180分钟',
        'prerequisites': '完成模块1的品牌现状诊断',
        'key_points': [
            '品牌识别系统的构成要素（MI、BI、VI）',
            '品牌理念识别（MI）的提炼方法',
            '品牌行为识别（BI）的设计要点',
            '品牌视觉识别（VI）的核心要素',
            '品牌命名与品牌故事',
            '品牌手册的编制方法'
        ],
        'interactions': [
            {'name': 'VI识别竞猜', 'description': '展示知名企业VI，学员竞猜并分析识别要素'},
            {'name': 'MI工作坊', 'description': '小组现场提炼企业大学品牌理念（使用工具卡）'},
            {'name': '品牌命名练习', 'description': '为虚构企业大学设计名称和口号'},
        ],
        'time_distribution': [
            {'name': '回顾与导入', 'duration': '10分钟', 'format': '提问', 'key_points': '回顾模块1，导入模块2'},
            {'name': '品牌识别系统理论', 'duration': '30分钟', 'format': '讲授', 'key_points': 'MI/BI/VI框架讲解'},
            {'name': 'MI设计工作坊', 'duration': '45分钟', 'format': '工作坊', 'key_points': '理念提炼实操'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': 'VI设计与应用', 'duration': '40分钟', 'format': '讲授+演示', 'key_points': '视觉识别要点'},
            {'name': '午餐', 'duration': '60分钟', 'format': '休息', 'key_points': ''},
            {'name': '品牌故事设计', 'duration': '25分钟', 'format': '工作坊', 'key_points': '故事框架与要素'},
            {'name': '品牌手册编制', 'duration': '20分钟', 'format': '讲授', 'key_points': '手册结构与要点'},
            {'name': '小组展示与点评', 'duration': '30分钟', 'format': '展示', 'key_points': '各组分享品牌理念'},
            {'name': '小结', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾要点'},
        ],
        'faqs': [
            {'question': '企业大学是否需要独立的VI系统？', 'answer': '根据企业大学独立程度决定。独立运营的企业大学需要完整VI，与母品牌强关联的可考虑子品牌策略'},
            {'question': '品牌理念与公司文化冲突怎么办？', 'answer': '以公司文化为基础进行延伸和细化，避免创造与母品牌相悖的理念'},
            {'question': '预算有限无法做完整VI？', 'answer': '聚焦核心应用（名片、PPT模板、证书），逐步完善'},
        ],
        'materials': [
            '品牌识别系统框架图',
            'MI提炼工具卡（价值观卡片、愿景卡等）',
            'VI设计案例集',
            '品牌故事模板',
            '品牌手册样例',
            '设计软件（PowerPoint/Canva）',
        ]
    }
    add_module_section(doc, 2, "品牌识别系统设计", module2)

    # 模块3
    module3 = {
        'objective': '掌握品牌传播与体验管理的核心策略，能够制定品牌传播计划',
        'duration': '150分钟',
        'prerequisites': '完成模块2的品牌识别设计',
        'key_points': [
            '品牌传播的基本原理（整合营销传播IMC）',
            '内部品牌传播策略',
            '外部品牌传播渠道',
            '品牌体验设计',
            '数字时代的品牌传播创新',
            '品牌传播效果评估'
        ],
        'interactions': [
            {'name': '传播渠道地图', 'description': '小组绘制企业大学品牌传播渠道图'},
            {'name': '角色扮演', 'description': '模拟内部品牌宣导活动（如何向员工传播品牌价值）'},
            {'name': '案例分析', 'description': '分析企业大学品牌传播成功案例'},
        ],
        'time_distribution': [
            {'name': '回顾与导入', 'duration': '10分钟', 'format': '提问', 'key_points': '回顾模块2核心产出'},
            {'name': '整合营销传播理论', 'duration': '25分钟', 'format': '讲授', 'key_points': 'IMC框架讲解'},
            {'name': '内部品牌传播', 'duration': '30分钟', 'format': '讲授+案例', 'key_points': '员工品牌认同建设'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': '外部品牌传播', 'duration': '30分钟', 'format': '讲授+讨论', 'key_points': '渠道选择与内容策略'},
            {'name': '品牌体验设计', 'duration': '25分钟', 'format': '工作坊', 'key_points': '关键时刻设计'},
            {'name': '数字传播策略', 'duration': '15分钟', 'format': '讲授', 'key_points': '新媒体、直播等'},
            {'name': '小组练习', 'duration': '25分钟', 'format': '工作坊', 'key_points': '制定传播计划'},
            {'name': '小结', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾与预告'},
        ],
        'faqs': [
            {'question': '内部品牌传播优先级不高？', 'answer': '强调员工是品牌最重要的传播者，内部认同是外部影响力的基础'},
            {'question': '如何衡量传播效果？', 'answer': '建立三层指标：曝光指标（到达率）、认知指标（理解度）、行动指标（参与度）'},
            {'question': '新媒体预算如何分配？', 'answer': '建议采用"少量付费+大量有机内容"的组合，控制获客成本'},
        ],
        'materials': [
            '整合营销传播框架图',
            '品牌传播渠道清单',
            '内部传播工具包（邮件模板、内宣材料示例）',
            '品牌体验地图模板',
            '传播效果评估表格',
            '案例视频（企业大学品牌传播案例）',
        ]
    }
    add_module_section(doc, 3, "品牌传播与体验", module3)

    # 模块4
    module4 = {
        'objective': '掌握品牌资产评估方法，能够建立品牌评估与优化体系',
        'duration': '120分钟',
        'prerequisites': '对品牌有一定认知基础',
        'key_points': [
            '品牌资产评估的意义与原则',
            '品牌资产评估的经典模型（Interbrand、BrandZ）',
            '企业大学品牌资产评估的特殊性',
            '品牌健康度调研方法',
            '品牌审计流程',
            '品牌优化策略制定'
        ],
        'interactions': [
            {'name': '品牌价值估算', 'description': '使用简化模型估算自身企业大学品牌价值'},
            {'name': '品牌审计演练', 'description': '小组对企业大学品牌进行模拟审计'},
        ],
        'time_distribution': [
            {'name': '导入', 'duration': '10分钟', 'format': '提问', 'key_points': '为什么需要评估品牌'},
            {'name': '品牌资产评估模型', 'duration': '30分钟', 'format': '讲授', 'key_points': '主流模型介绍'},
            {'name': '企业大学品牌评估', 'duration': '25分钟', 'format': '讲授+讨论', 'key_points': '特殊性分析'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': '品牌调研方法', 'duration': '20分钟', 'format': '讲授', 'key_points': '问卷、访谈、焦点小组'},
            {'name': '品牌审计工作坊', 'duration': '25分钟', 'format': '工作坊', 'key_points': '审计流程实操'},
            {'name': '小结', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾要点'},
        ],
        'faqs': [
            {'question': '品牌价值难以量化？', 'answer': '采用多种方法交叉验证：财务法（溢价能力）、市场法（市场份额）、行为法（忠诚度）'},
            {'question': '调研样本量不足怎么办？', 'answer': '采用定性+定量结合，即使小样本也能获得有价值的洞察'},
            {'question': '评估频率如何确定？', 'answer': '建议每年一次基础评估，每2-3年一次深度评估'},
        ],
        'materials': [
            '品牌资产评估模型对照表',
            '品牌健康度调研问卷模板',
            '品牌审计检查清单',
            '品牌价值计算工具',
            '案例数据（Interbrand、BrandZ报告摘要）',
        ]
    }
    add_module_section(doc, 4, "品牌评估与优化", module4)

    # 模块5
    module5 = {
        'objective': '理解品牌资产管理的整体框架，能够建立品牌治理体系',
        'duration': '150分钟',
        'prerequisites': '完成前面模块学习',
        'key_points': [
            '品牌资产管理的整体框架',
            '品牌治理结构设计',
            '品牌管理制度与流程',
            '品牌资产日常管理',
            '品牌危机管理',
            '品牌管理工具与系统'
        ],
        'interactions': [
            {'name': '品牌治理讨论', 'description': '讨论企业大学品牌治理的常见问题与解决思路'},
            {'name': '危机案例模拟', 'description': '分析品牌危机案例，演练应对策略'},
        ],
        'time_distribution': [
            {'name': '导入', 'duration': '10分钟', 'format': '讲授', 'key_points': '品牌管理的重要性'},
            {'name': '品牌管理框架', 'duration': '25分钟', 'format': '讲授', 'key_points': '整体框架讲解'},
            {'name': '品牌治理结构', 'duration': '30分钟', 'format': '讲授+讨论', 'key_points': '组织架构设计'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': '品牌管理制度', 'duration': '25分钟', 'format': '讲授', 'key_points': '制度要点'},
            {'name': '品牌危机管理', 'duration': '30分钟', 'format': '案例+模拟', 'key_points': '危机处理流程'},
            {'name': '品牌管理系统', 'duration': '15分钟', 'format': '讲授', 'key_points': '工具推荐'},
            {'name': '小结', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾'},
        ],
        'faqs': [
            {'question': '品牌管理权限如何分配？', 'answer': '建立"品牌委员会+品牌专员"的架构，明确各层级权限'},
            {'question': '小微企业如何进行品牌管理？', 'answer': '简化流程，重点把控核心触点，建立基本的品牌规范'},
            {'question': '员工违反品牌规范怎么处理？', 'answer': '事前培训明确，事中及时提醒，事后根据情节处理'},
        ],
        'materials': [
            '品牌管理框架图',
            '品牌治理组织架构模板',
            '品牌管理制度范本',
            '危机应对话术模板',
            '品牌管理工具清单',
        ]
    }
    add_module_section(doc, 5, "品牌资产管理体系", module5)

    # 模块6
    module6 = {
        'objective': '综合运用所学知识，完成企业大学品牌建设方案的策划',
        'duration': '210分钟',
        'prerequisites': '完成所有前置模块学习，准备本企业背景资料',
        'key_points': [
            '企业大学品牌战略规划流程',
            '品牌建设路线图设计',
            '资源规划与预算编制',
            '利益相关方管理',
            '品牌建设项目的推进与监控',
            '品牌建设成果的展示与汇报'
        ],
        'interactions': [
            {'name': '企业背景分析', 'description': '各小组整理本企业企业大学品牌现状'},
            {'name': '品牌战略规划', 'description': '小组制定品牌战略规划方案'},
            {'name': '路演准备', 'description': '准备品牌建设方案的汇报'},
            {'name': '最终路演', 'description': '各组展示品牌建设方案，评审点评'},
        ],
        'time_distribution': [
            {'name': '回顾与导入', 'duration': '10分钟', 'format': '讲授', 'key_points': '回顾课程框架'},
            {'name': '战略规划方法', 'duration': '20分钟', 'format': '讲授', 'key_points': '规划流程'},
            {'name': '企业背景梳理', 'duration': '30分钟', 'format': '工作坊', 'key_points': '现状分析'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': '品牌战略制定', 'duration': '45分钟', 'format': '工作坊', 'key_points': '战略规划'},
            {'name': '午餐', 'duration': '60分钟', 'format': '休息', 'key_points': ''},
            {'name': '路线图与预算', 'duration': '35分钟', 'format': '工作坊', 'key_points': '实施规划'},
            {'name': '茶歇', 'duration': '10分钟', 'format': '休息', 'key_points': ''},
            {'name': '方案完善', 'duration': '30分钟', 'format': '工作坊', 'key_points': '打磨方案'},
            {'name': '各组路演', 'duration': '60分钟', 'format': '展示', 'key_points': '每组10分钟+5分钟点评'},
            {'name': '总结', 'duration': '15分钟', 'format': '讲授', 'key_points': '课程总结'},
        ],
        'faqs': [
            {'question': '方案太理想化无法落地？', 'answer': '强调"最小可行品牌建设"理念，从低成本高价值的快速胜利开始'},
            {'question': '内部资源不足怎么办？', 'answer': '展示如何利用现有资源，以及如何争取更多资源的策略'},
            {'question': '如何获得领导支持？', 'answer': '提供领导汇报模板和说服技巧，强调品牌投资的回报'},
        ],
        'materials': [
            '企业大学品牌战略规划模板',
            '品牌建设路线图模板',
            '预算表格模板',
            '利益相关方分析工具',
            '汇报PPT模板',
            '评估标准表',
        ]
    }
    add_module_section(doc, 6, "企业大学品牌建设实战", module6)

    # ========== 第十章：评估标准与评分指引 ==========
    doc.add_heading("第十章  评估标准与评分指引", level=1)

    doc.add_heading("10.1 学员评估方式", level=2)

    assessment_methods = [
        ["评估方式", "评估内容", "权重", "评估时间"],
        ["课堂参与", "出勤、发言、讨论贡献", "10%", "全程"],
        ["案例分析", "案例理解的深度与分析框架", "15%", "模块2后"],
        ["品牌设计作业", "品牌识别系统设计的完整性与创新性", "25%", "模块2后"],
        ["品牌传播计划", "传播策略的可行性和创意", "20%", "模块3后"],
        ["毕业答辩", "品牌建设方案的完整性、可行性和呈现", "30%", "模块6后"],
    ]
    add_table_with_header(doc, assessment_methods[0], assessment_methods[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("10.2 评分标准", level=2)

    grading_rubric = [
        ["评分维度", "优秀(90-100)", "良好(80-89)", "合格(70-79)", "待改进(<70)"],
        ["品牌理解\n(20%)", "深刻理解品牌资产本质，能融会贯通", "理解较准确，能联系实际", "理解基本概念", "概念混淆"],
        ["方案设计\n(30%)", "方案完整、创新、切实可行", "方案较完整，可行性较好", "方案基本完整", "方案有明显缺陷"],
        ["创意创新\n(20%)", "有独特视角和创新点", "有一定创新", "较常规", "缺乏新意"],
        ["表达呈现\n(15%)", "表达清晰、有感染力", "表达清楚", "基本清楚", "表达混乱"],
        ["团队协作\n(15%)", "团队配合默契，贡献均衡", "配合较好", "有分工", "协作差"],
    ]
    add_table_with_header(doc, grading_rubric[0], grading_rubric[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("10.3 反馈技巧", level=2)

    feedback_tips = [
        ["反馈类型", "技巧", "示例"],
        ["正面反馈", "具体化、即时化", "你的品牌定位分析很有逻辑，特别是用PEST工具分析宏观环境这个角度很独特"],
        ["改进建议", "先肯定后建议、具体可操作", "你的传播策略覆盖面很广，建议下一步可以思考如何 prioritization（优先排序）"],
        ["质疑提问", "开放式、启发式", "我想了解一下，你为什么选择这个传播渠道？这个选择背后的逻辑是什么？"],
        ["纠错指正", "对事不对人、给出正确做法", "品牌价值观的提炼需要更具体可衡量，现在的表述有点抽象，建议用行为锚定的方式"],
    ]
    add_table_with_header(doc, feedback_tips[0], feedback_tips[1:], TABLE_HEADER_BG)

    add_info_box(doc, "反馈原则",
        "1. 及时性：在行为发生后尽快反馈\n"
        "2. 具体性：避免泛泛而谈，要指出具体表现\n"
        "3. 建设性：指出问题的同时提供解决方向\n"
        "4. 平衡性：正面与改进建议保持适当平衡", "note")

    doc.add_page_break()

    # ========== 附录 ==========
    doc.add_heading("附录  参考文献、延伸阅读、工具清单", level=1)

    doc.add_heading("附录A  参考文献", level=2)

    references = [
        "Keller, K.L. (1998). Strategic Brand Management: Building, Measuring, and Managing Brand Equity. Prentice Hall.",
        "Aaker, D.A. (1991). Managing Brand Equity: Capitalizing on the Value of a Brand Name. Free Press.",
        "Kaplan, R.S., & Norton, D.P. (1996). The Balanced Scorecard: Translating Strategy into Action. Harvard Business Review Press.",
        "Chernev, A. (2018). Strategic Marketing Management. Cerebellum Press.",
        "Holt, D. (2004). How Brands Become Icons: The Principles of Cultural Branding. Harvard Business Review Press.",
        "沈健 (2020). 企业大学白皮书. 中国企业大学联合会.",
        "ATD (2021). State of the Industry Report. Association for Talent Development.",
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(6)

    doc.add_heading("附录B  延伸阅读", level=2)

    reading_list = [
        ["书籍/文章", "作者", "推荐理由"],
        ["《品牌洗脑》", "马丁·林德斯特罗姆", "了解品牌如何影响消费者行为"],
        ["《成为抢占心智的品牌》", "杰克·特劳特", "定位理论的实战应用"],
        ["《品牌的技术与艺术》", "叶明桂", "本土品牌建设实战经验"],
        ["《企业大学宣言》", " Jay A. Conger", "企业大学建设的权威著作"],
        ["Harvard Business Review品牌专刊", "哈佛商业评论", "最新品牌管理理念和案例"],
    ]
    add_table_with_header(doc, reading_list[0], reading_list[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("附录C  工具清单", level=2)

    tools_list = [
        ["工具类别", "工具名称", "用途", "获取方式"],
        ["品牌诊断", "Brand Audit Checklist", "品牌健康度检查", "课程提供"],
        ["品牌定位", "Brand Position Canvas", "品牌定位规划", "课程提供"],
        ["视觉设计", "Canva for Enterprise", "快速视觉设计", "免费使用"],
        ["调研分析", "问卷星/腾讯问卷", "品牌调研问卷", "免费额度"],
        ["项目管理", "Trello/Notion", "品牌项目跟进", "免费使用"],
        ["数据分析", "Tableau Public", "品牌数据可视化", "免费使用"],
        ["协作沟通", "飞书/钉钉", "团队协作", "企业版"],
    ]
    add_table_with_header(doc, tools_list[0], tools_list[1:], TABLE_HEADER_BG)

    doc.add_paragraph()
    doc.add_heading("附录D  课程资源链接", level=2)

    resources = [
        "课程配套PPT素材包（包含所有讲授用的图表、案例资料）",
        "品牌工具包下载链接（包含所有工作坊使用的模板、工具卡）",
        "学员交流群二维码（课程期间用于资料分享和答疑）",
        "讲师备课指南（详细版本，包含每个模块的逐字稿）",
        "课程更新与迭代记录（持续优化课程内容）",
    ]
    for res in resources:
        p = doc.add_paragraph(res, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)

    # ========== 封底 ==========
    doc.add_page_break()
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.add_run("\n\n\n\n\n\n\n\n\n—— 完 ——")
    end_run.font.size = Pt(16)
    end_run.font.color.rgb = PRIMARY_COLOR

    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("\n\n\n企业大学品牌资产设计课程\n内部讲师使用手册\n© 版权所属 翻版必究")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    # 保存文档
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"文档已生成: {OUTPUT_PATH}")
    return OUTPUT_PATH

# ============================================================
# 执行
# ============================================================

if __name__ == "__main__":
    create_instructor_handbook()

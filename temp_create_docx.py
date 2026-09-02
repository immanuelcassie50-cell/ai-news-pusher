#!/usr/bin/env python3
"""Create resistance signal analysis report Word document"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

output_path = r'D:\新课开发\变革管理\12-抵抗信号的早期识别：变革失败之前，组织早就发出过警告\完整课程包\08-成果demo\03-抵抗信号分析报告示例.docx'

doc = Document()

# Set page margins
section = doc.sections[0]
section.page_width = Inches(11.69)  # A4 landscape
section.page_height = Inches(8.27)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# Helper functions
def set_run_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = 'Microsoft YaHei'
    r = run._r
    rPr = r.get_or_add_rPr()
    rPr.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        r = run._r
        rPr = r.get_or_add_rPr()
        rPr.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p

# === Cover Section ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.space_before = Pt(72)
title.space_after = Pt(24)
run = title.add_run('抵抗信号早期识别分析报告')
set_run_font(run, size=26, bold=True, color=(60, 60, 60))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.space_after = Pt(12)
run = subtitle.add_run('——数字化转型中的组织预警与响应策略')
set_run_font(run, size=14, color=(100, 100, 100))

doc.add_paragraph()
controlled = doc.add_paragraph()
controlled.alignment = WD_ALIGN_PARAGRAPH.LEFT
controlled.add_run('受控信息等级：内部使用').bold = True
controlled.add_run('    |    ').bold = True
controlled.add_run('编制日期：2024年11月').bold = True
controlled.add_run('    |    ').bold = True
controlled.add_run('版本：V1.0').bold = True

doc.add_page_break()

# === Section 1: Background ===
add_heading(doc, '一、背景介绍', level=1)

p = doc.add_paragraph()
p.add_run('本报告基于某制造企业数字化转型项目第二阶段（智能排产系统上线）的实际运营数据，梳理了项目启动以来观察到的各类抵抗信号。')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('变革背景：').bold = True
p.add_run('公司于2024年9月启动ERP系统升级项目，目标是通过智能化手段提升运营效率。项目分为三期，预计18个月内完成全流程数字化。')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('分析目的：').bold = True
p.add_run('通过系统识别变革过程中的抵抗信号，评估其风险等级，并提出针对性的响应建议，为管理层提供决策参考。')

# === Section 2: Signal Analysis ===
add_heading(doc, '二、信号分析（四维度）', level=1)

# 2.1 Behavioral Signals
add_heading(doc, '2.1 行为维度信号', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['观察指标', '具体表现', '风险等级']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['会议参与度', '部门例会出勤率下降15-20%，核心人员经常缺席', '高'],
    ['沟通模式', '正式渠道反馈减少，私下议论增多；会议中沉默时间延长', '中高'],
    ['工作态度', '对变更任务的主动性下降，等待观望心态明显', '中'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()

# 2.2 Financial Signals
add_heading(doc, '2.2 财务维度信号', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['观察指标', '具体表现', '风险等级']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['费用控制', '项目相关差旅和咨询费用比预算低22%，执行力度不足', '中高'],
    ['报销时效', '报销审批周期延长约40%，退单率上升', '中'],
    ['资源调配', '关键岗位人员投入项目的实际工时低于计划的60%', '高'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()

# 2.3 Organizational Signals
add_heading(doc, '2.3 组织维度信号', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['观察指标', '具体表现', '风险等级']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['人员流动', '项目核心岗位两名成员提出内部调动申请', '高'],
    ['协作意愿', '跨部门协作会议参会质量下降，配合度降低', '中高'],
    ['技能储备', '培训计划完成率仅达45%，关键用户掌握度不足', '中高'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()

# 2.4 Communication Signals
add_heading(doc, '2.4 沟通维度信号', level=2)

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['观察指标', '具体表现', '风险等级']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['反馈质量', '问题反馈变得笼统、泛化，缺少具体场景描述', '中'],
    ['信息渠道', '非正式渠道（小道消息、茶水间）信息传播活跃度上升', '中高'],
    ['变革话语', '员工在沟通中频繁使用"折腾"、"又来"等负面词汇', '中'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()

# === Section 3: Risk Assessment ===
add_heading(doc, '三、风险评估', level=1)

p = doc.add_paragraph()
p.add_run('综合以上四个维度的信号分析，当前项目的抵抗风险评估如下：')

doc.add_paragraph()

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['风险类别', '当前状态', '发展趋势', '优先级']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['项目进度风险', '中高', '上升', 'P1'],
    ['人员稳定性风险', '高', '持平', 'P1'],
    ['质量目标风险', '中', '需观察', 'P2'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()

# === Section 4: Response Recommendations ===
add_heading(doc, '四、响应建议', level=1)

add_heading(doc, '4.1 立即行动（0-2周）', level=2)
bullets = [
    '召开部门负责人专题会，明确传达项目战略价值和阶段目标',
    '启动关键人物一对一沟通计划，了解真实顾虑',
    '加强项目组与业务部门的联席沟通频率，由双周改为周例会',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b)

add_heading(doc, '4.2 短期措施（1个月内）', level=2)
bullets = [
    '重新评估培训方案，增加实操演练和个性化辅导',
    '建立员工反馈的正向激励机制，鼓励真实声音',
    '设立"过渡期缓冲"机制，允许分批切换而非一刀切',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b)

add_heading(doc, '4.3 中长期策略', level=2)
bullets = [
    '构建持续沟通机制，定期发布项目进展和成功案例',
    '识别并培养内部变革大使，形成正向影响力网络',
    '将变革管理纳入各级管理者的绩效考核维度',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(b)

doc.add_paragraph()

# === Section 5: Tracking Plan ===
add_heading(doc, '五、跟踪计划', level=1)

table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['跟踪指标', '检查频率', '责任人', '预警阈值']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['会议出勤率', '周', '项目经理', '<80%'],
    ['培训完成率', '双周', '培训负责人', '<60%'],
    ['人员流失意向', '月', 'HRBP', '>1人/月'],
    ['关键指标达成', '月', '运营总监', '<90%目标'],
    ['员工反馈满意度', '双周', '变革管理专员', '<3.5分'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, val in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.text = val

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.sections[0].footer
footer_para = footer.paragraphs[0]
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.add_run('内部文件  |  仅供授权人员使用  |  © 2024').font.size = Pt(9)

doc.save(output_path)
print(f'Created: {output_path}')
